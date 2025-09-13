#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试实际数据键名的项目信息处理
"""

import sys
import logging
from pathlib import Path
from docx import Document

# 添加项目根目录到系统路径
sys.path.append(str(Path(__file__).parent))

from ai_tender_system.modules.business_response.info_filler import InfoFiller

def setup_test_logging():
    """设置测试日志"""
    # 设置根日志器
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    
    # 创建控制台处理器
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.DEBUG)
    
    # 创建格式器
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    console_handler.setFormatter(formatter)
    
    # 添加处理器到根日志器
    root_logger.addHandler(console_handler)
    
    return root_logger

def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    # 添加包含括号格式的测试段落
    test_paragraphs = [
        "根据贵方为（项目名称）项目采购采购货物及服务的竞争性磋商公告（采购编号），签字代表经正式授权。",
        "（项目名称）",  
        "（采购编号）",
        "供应商名称 ",
    ]
    
    for text in test_paragraphs:
        doc.add_paragraph(text)
    
    return doc

def test_real_data_keys():
    """测试实际数据键名"""
    print("=" * 60)
    print("测试实际数据键名的项目信息处理")
    print("=" * 60)
    
    # 设置日志
    logger = setup_test_logging()
    
    # 创建InfoFiller实例
    filler = InfoFiller()
    
    # 创建测试文档
    doc = create_test_document()
    
    # 使用实际的数据键名（从Web应用日志中提取）
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'address': '北京市东城区王府井大街200号七层711室',
        'phone': '010-63271000',
        'email': 'lvhe@smartsteps.com',
        'fax': '010-63271000',
    }
    
    project_info = {
        '项目名称': '哈银消金2025年-2027年运营商数据采购项目',  # 中文键名!
        '招标编号': 'GXTC-C-251590031',  # 中文键名!
        'purchaserName': '哈尔滨哈银消费金融有限责任公司',
        'date': '2025年9月13日',
    }
    
    print(f"\n📋 测试数据:")
    print(f"公司信息: {company_info}")
    print(f"项目信息: {project_info}")
    
    print(f"\n📄 测试文档段落:")
    for i, paragraph in enumerate(doc.paragraphs, 1):
        print(f"  {i}. '{paragraph.text}'")
    
    print(f"\n" + "=" * 60)
    print("开始执行信息填写处理")
    print("=" * 60)
    
    # 执行信息填写
    try:
        result = filler.fill_info(doc, company_info, project_info)
        
        print(f"\n" + "=" * 60)
        print("处理完成后的文档内容")
        print("=" * 60)
        
        for i, paragraph in enumerate(doc.paragraphs, 1):
            print(f"  {i}. '{paragraph.text}'")
        
        print(f"\n📊 最终统计结果: {result}")
        
    except Exception as e:
        print(f"\n❌ 测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_real_data_keys()