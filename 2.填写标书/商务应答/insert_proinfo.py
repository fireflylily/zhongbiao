import docx
import logging
import re
from docx.enum.text import WD_UNDERLINE

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

def fill_project_info(input_file, output_file, project_name, tender_no, date_text):
    logging.info(f"加载文档: {input_file}")
    doc = docx.Document(input_file)

    # 小改动：把日期替换规则放入 replacements，避免 append 报错
    replacements = [
        # 项目名称
        (r"(项目名称[:：]\s*)$", r"\1" + project_name),
        (r"(项目名称（盖章）[:：]?\s*)$", r"\1" + project_name),
        (r"(项目名称\s+)$", r"\1" + project_name),
        (r"(项目名称)$", r"\1" + project_name),
        (r"\[项目名称\]", project_name),
        # 招标编号
        (r"(招标编号[:：]\s*)$", r"\1" + tender_no),
        (r"(招标编号\s+)$", r"\1" + tender_no),
        (r"(招标编号)$", r"\1" + tender_no),
        (r"\[招标编号\]", tender_no),
        # 日期（匹配独立“日 期”，中间允许空格）
        (r"^(日\s*期)\s*[:：]?\s*$", "日期：" + date_text),
        # 组合：项目名称、招标编号
        ("项目名称、招标编号", f"{project_name}、{tender_no}")
    ]

    # 遍历段落，先处理普通替换
    for para in doc.paragraphs:
        for run in para.runs:
            for pat, repl in replacements:
                if re.search(pat, run.text):
                    logging.info(f"匹配到: {run.text.strip()}")
                    run.text = re.sub(pat, repl, run.text)
                    break  # 一个 run 一次替换即可

    # 遍历段落，处理投标人名称相关字段
    company_name = "中国联合网络通信有限公司"
    for para in doc.paragraphs:
        for i, run in enumerate(para.runs):
            text = run.text
            # 匹配三种情况
            if "投标人名称（公章）" in text or "投标人名称（盖章）" in text or re.match(r"投标人名称\s*[:：]?", text):
                # 确定前缀
                if "（公章）" in text:
                    prefix = "投标人名称（公章）："
                elif "（盖章）" in text:
                    prefix = "投标人名称（盖章）："
                else:
                    prefix_match = re.match(r"(投标人名称\s*[:：]?)", text)
                    prefix = prefix_match.group(1) if prefix_match else "投标人名称："

                # 覆盖原 run，只保留前缀
                run.text = prefix

                # 清空后续 run（防止原有下划线空格残留）
                for j in range(i + 1, len(para.runs)):
                    para.runs[j].text = ""

                # 插入公司名称，带下划线
                new_run = para.add_run(company_name)
                new_run.underline = WD_UNDERLINE.SINGLE
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size

                logging.info(f"✅ 已填充 {prefix} {company_name}")
                break  # 一个段落只处理一次

    doc.save(output_file)
    logging.info(f"🎉 处理完成，新文件已保存到：{output_file}")


# ==============================
# 使用示例
# ==============================

input_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/项目/中信归属运营商/应答文件格式.docx"
output_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/项目/中信归属运营商/应答文件格式-1.docx"

project_name = "所属运营商数据"
tender_no = "2025-IT-0032"
date_text = "2025年 9月 4日"

fill_project_info(input_file, output_file, project_name, tender_no, date_text)
