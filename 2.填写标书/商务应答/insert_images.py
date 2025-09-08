import docx
from docx.shared import Inches
import os
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)


from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Mm  # 记得引入 Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def insert_paragraph_after(paragraph, text=None):
    """
    在指定 paragraph 后插入新段落，并返回 Paragraph 对象
    """
    # 创建一个新的 w:p 元素（段落）
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)  # 插在当前 paragraph 后面

    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para



def insert_images_after_last_matches(input_file, output_file, keyword_to_image):
    logging.info(f"加载文档: {input_file}")
    doc = docx.Document(input_file)

    # 记录每个关键字最后一次出现的位置
    last_matches = {}

    for i, para in enumerate(doc.paragraphs):
        for keyword in keyword_to_image:
            if keyword in para.text:
                last_matches[keyword] = para
                logging.info(f"匹配到关键字 [{keyword}] 段落位置：{i} - 内容：{para.text.strip()}")

    # 遍历关键字并在最后一次出现后插入图片
    for keyword, para in last_matches.items():
        img_paths = keyword_to_image[keyword]
        for img_path in img_paths:
            logging.info(f"正在处理关键字：{keyword} -> {img_path}")
            if os.path.exists(img_path):
                new_para = insert_paragraph_after(para)
                new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中
                
                run = new_para.add_run()
                
                # 根据关键字设置图片大小
                if keyword == "营业执照":
                    run.add_picture(img_path, width=Mm(150))
                elif keyword == "法定代表人身份证":
                    run.add_picture(img_path, width=Mm(65))
                elif keyword == "被授权人身份证":
                    run.add_picture(img_path, width=Mm(65))
                else:
                    run.add_picture(img_path, width=Inches(5.5))
                
                logging.info(f"✅ 已在段落后插入居中图片：{img_path}")
            else:
                logging.warning(f"⚠️ 未找到图片文件：{img_path}")

    # 保存新文档
    doc.save(output_file)
    logging.info(f"\n🎉 处理完成，新文件已保存到：{output_file}")


# ==============================
# 使用示例
# ==============================
input_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/标书模板.docx"
output_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/1-信用截图+执照+法人.docx"

keyword_to_image = {
    "失信被执行人": ["证明材料/失信被执行人-智慧足迹.png"],
    "重大税收违法案件当事人": ["证明材料/重大税收失信主体名单-智慧足迹.png"],
    "政府采购严重违法失信记录名单": ["证明材料/政府采购严重违法失信记录名单-智慧足迹.png"],
    "营业执照": ["证明材料/营业执照_智慧足迹.png"],
    "法定代表人身份证": [
        "证明材料/法人身份证_正面_智慧足迹.png",
        "证明材料/法人身份证_反面_智慧足迹.png"
    ],
        "被授权人身份证": [
        "证明材料/被授权人身份证_正面_吕贺.png",
        "证明材料/被授权人身份证_反面_吕贺.png"
    ]
}


insert_images_after_last_matches(input_file, output_file, keyword_to_image)
