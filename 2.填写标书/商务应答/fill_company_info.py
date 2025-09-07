#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fill_company_info.py
公司全信息填写处理器 - 在商务应答文件中填写公司的所有联系信息

用法:
    python3 fill_company_info.py input.docx output.docx

功能:
- 自动识别并填写公司名称、地址、电话、传真、电子邮件等信息
- 支持多种格式的标签识别
- 从配置文件自动读取公司信息
"""

import sys
import re
import json
import os
from docx import Document


class CompanyInfoProcessor:
    """公司信息填写处理器"""
    
    def __init__(self):
        # 加载公司信息
        self.company_info = self._load_company_info()
        
        # 定义各类信息的匹配模式
        self.patterns = {
            'company_name': [
                # 公司名称相关
                re.compile(r'^(?P<label>(投标人|供应商|公司)(?:名称|全称)?(?:（盖章）|（公章）|（全称、盖章）|\(盖章\)|\(公章\)|\(全称、盖章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
                re.compile(r'(?P<prefix>[\(（])\s*(?P<content>(?:请填写\s*)?(?:供应商名称|投标人名称|公司名称|单位名称))\s*(?P<suffix>[\)）])'),
            ],
            'address': [
                # 地址相关
                re.compile(r'^(?P<label>(公司|企业|投标人|供应商)(?:地址|住址|注册地址|办公地址))\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{2,}|＿+|——+)?$'),
                re.compile(r'(?P<prefix>[\(（])\s*(?P<content>地址|住址)\s*(?P<suffix>[\)）])'),
                re.compile(r'^(?P<label>地址)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{2,}|＿+|——+)?$'),
            ],
            'phone': [
                # 电话相关
                re.compile(r'^(?P<label>(联系)?电话|固定电话|办公电话)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{2,}|＿+|——+)?$'),
                re.compile(r'(?P<prefix>[\(（])\s*(?P<content>电话|联系电话)\s*(?P<suffix>[\)）])'),
            ],
            'fax': [
                # 传真相关  
                re.compile(r'^(?P<label>传真|传真号码?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{2,}|＿+|——+)?$'),
                re.compile(r'(?P<prefix>[\(（])\s*(?P<content>传真)\s*(?P<suffix>[\)）])'),
            ],
            'email': [
                # 邮箱相关
                re.compile(r'^(?P<label>电子邮件|邮箱|电子邮箱|E-?mail|Email)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{2,}|＿+|——+)?$'),
                re.compile(r'(?P<prefix>[\(（])\s*(?P<content>电子邮件|邮箱|电子邮箱)\s*(?P<suffix>[\)）])'),
            ]
        }
    
    def _load_company_info(self):
        """从配置文件加载公司信息"""
        # 查找配置文件路径
        current_dir = os.path.dirname(os.path.abspath(__file__))
        config_file = None
        
        # 从当前目录开始向上查找
        search_dirs = [
            current_dir,
            os.path.join(current_dir, '..', '配置信息'),
            os.path.join(current_dir, '..', '..', '配置信息'),
            os.path.join(current_dir, '..', '..', '..', '配置信息'),
        ]
        
        for search_dir in search_dirs:
            potential_file = os.path.join(search_dir, 'project_info.json')
            if os.path.exists(potential_file):
                config_file = potential_file
                break
        
        if not config_file:
            print("警告：未找到project_info.json配置文件，使用默认公司信息")
            return {
                'company_name': '智慧足迹数据科技有限公司',
                '注册地址': '北京市东城区王府井大街200号七层711室',
                '固定电话': '010-12345678',
                '传真': '010-87654321', 
                '电子邮件': 'info@company.com'
            }
        
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"加载配置文件失败: {e}")
            return {}
    
    def _get_company_value(self, info_type):
        """获取对应类型的公司信息"""
        mappings = {
            'company_name': self.company_info.get('company_name', ''),
            'address': self.company_info.get('注册地址', self.company_info.get('办公地址', '')),
            'phone': self.company_info.get('固定电话', ''),
            'fax': self.company_info.get('传真', ''),
            'email': self.company_info.get('电子邮件', '')
        }
        return mappings.get(info_type, '')
    
    def _process_text_with_patterns(self, text, info_type):
        """使用模式匹配处理文本"""
        if info_type not in self.patterns:
            return text, False
        
        original_text = text.strip()
        if not original_text:
            return text, False
        
        # 获取对应的公司信息
        company_value = self._get_company_value(info_type)
        if not company_value:
            return text, False
        
        # 检查是否已经包含信息，避免重复填写
        if company_value in original_text:
            return text, False
        
        # 尝试各种模式
        for pattern in self.patterns[info_type]:
            match = pattern.search(original_text)
            if match:
                groups = match.groupdict()
                
                # 括号格式处理
                if 'prefix' in groups and 'content' in groups and 'suffix' in groups:
                    prefix = groups.get('prefix', '')
                    suffix = groups.get('suffix', '')
                    new_text = f"{prefix}{company_value}{suffix}"
                    print(f"匹配{info_type}括号格式: '{original_text}' -> '{new_text}'")
                    return new_text, True
                
                # 标签格式处理
                elif 'label' in groups:
                    label = groups.get('label', '')
                    sep = groups.get('sep', ':')
                    suffix = groups.get('suffix', '')
                    
                    if not sep:
                        sep = ':'
                    if sep and not sep.endswith(' '):
                        sep += ' ' if sep in [':', '：'] else ''
                    
                    if suffix:
                        new_text = f"{label}{sep}{company_value}            {suffix}"
                    else:
                        new_text = f"{label}{sep}{company_value}"
                    
                    print(f"匹配{info_type}标签格式: '{original_text}' -> '{new_text}'")
                    return new_text, True
        
        return text, False
    
    def _process_paragraph(self, paragraph):
        """处理单个段落"""
        if not paragraph.runs:
            return False
        
        # 合并段落文本
        full_text = ''.join(run.text for run in paragraph.runs)
        if not full_text.strip():
            return False
        
        # 尝试各种信息类型的匹配
        changed = False
        for info_type in self.patterns.keys():
            new_text, is_changed = self._process_text_with_patterns(full_text, info_type)
            if is_changed:
                # 清空所有run并将新文本放入第一个run
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                changed = True
                break  # 一个段落只匹配一种信息类型
        
        return changed
    
    def _process_table(self, table):
        """处理表格"""
        changed_any = False
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if self._process_paragraph(p):
                        changed_any = True
                for tbl in cell.tables:
                    if self._process_table(tbl):
                        changed_any = True
        return changed_any
    
    def process_document(self, doc):
        """处理整个文档"""
        stats = {
            'paragraphs_changed': 0,
            'tables_changed': 0,
            'headers_changed': 0,
            'footers_changed': 0
        }
        
        print(f"开始处理文档，使用公司信息：")
        for info_type in ['company_name', 'address', 'phone', 'fax', 'email']:
            value = self._get_company_value(info_type)
            if value:
                print(f"  {info_type}: {value}")
        
        # 处理正文段落
        for para in doc.paragraphs:
            if self._process_paragraph(para):
                stats['paragraphs_changed'] += 1
        
        # 处理表格
        for table in doc.tables:
            if self._process_table(table):
                stats['tables_changed'] += 1
        
        # 处理页眉页脚
        for section in doc.sections:
            header = section.header
            footer = section.footer
            header_changed = False
            footer_changed = False
            
            for para in header.paragraphs:
                if self._process_paragraph(para):
                    header_changed = True
            for table in header.tables:
                if self._process_table(table):
                    header_changed = True
            
            for para in footer.paragraphs:
                if self._process_paragraph(para):
                    footer_changed = True
            for table in footer.tables:
                if self._process_table(table):
                    footer_changed = True
            
            if header_changed:
                stats['headers_changed'] += 1
            if footer_changed:
                stats['footers_changed'] += 1
        
        return stats


def main():
    if len(sys.argv) < 3:
        print("用法: python3 fill_company_info.py input.docx output.docx")
        return 1
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"错误: 输入文件 {input_file} 不存在")
        return 1
    
    try:
        print(f"加载文档: {input_file}")
        doc = Document(input_file)
        
        processor = CompanyInfoProcessor()
        stats = processor.process_document(doc)
        
        print(f"保存到: {output_file}")
        doc.save(output_file)
        
        print("处理完成！统计信息:")
        print(f"  修改段落数: {stats['paragraphs_changed']}")
        print(f"  修改表格数: {stats['tables_changed']}")
        print(f"  修改页眉数: {stats['headers_changed']}")
        print(f"  修改页脚数: {stats['footers_changed']}")
        
        return 0
        
    except Exception as e:
        print(f"处理文档时发生错误: {e}")
        return 1


if __name__ == '__main__':
    sys.exit(main())