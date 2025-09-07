import re
import requests
from docx import Document
from docx.oxml import OxmlElement

# ==================== 调用始皇API ====================



    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        result = data["choices"][0]["message"]["content"].strip()
        print(f"[{purpose.upper()}] {result[:100]}{'...' if len(result) > 100 else ''}")
        return result
    except Exception as e:
        print(f"调用始皇API失败（{purpose}）：", e)
        return f"（{purpose}生成失败，请手工补充）"


# ==================== 插入段落工具 ====================
def insert_paragraph_after(paragraph, text=None, style=None):
    """
    在 paragraph 后插入新段落
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


# ==================== Word处理逻辑 ====================
def insert_reply_with_llm(input_file, output_file):
    doc = Document(input_file)

    # 1️⃣ 汇总采购需求
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    summary = llm_callback(f"请总结以下采购需求，提炼关键点，简洁扼要：\n\n{full_text}", purpose="summary")

    # 2️⃣ 生成技术方案大纲
    outline = llm_callback(f"根据以下采购需求，生成技术应答方案大纲：\n\n{summary}", purpose="outline")

    # 3️⃣ 扩展完整技术方案
    plan = llm_callback(f"请根据以下大纲，扩展成完整的技术方案，语言正式、专业：\n\n{outline}", purpose="plan")

    # 4️⃣ 在文档最后追加“技术方案”部分
    doc.add_page_break()
    try:
        doc.add_paragraph("技术应答方案", style="Heading 1")
    except KeyError:
        doc.add_paragraph("技术应答方案")
    doc.add_paragraph(plan)
    doc.add_page_break()

    # 记录技术方案起始索引，避免在正文遍历时插入
    plan_start_index = len(doc.paragraphs)

    # 5️⃣ 遍历正文（跳过新加的技术方案部分）
    for para in doc.paragraphs[:plan_start_index]:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        alignment = para.paragraph_format.alignment

        if (
            text == ""
            or "Heading" in style_name
            or "标题" in style_name
            or "TOC" in style_name
            or "目录" in text
            or alignment == 1
            or re.match(r"^(\d+(\.\d+)*[、\.]?|（[一二三四五六七八九十]+）)", text)
        ):
            continue

        # 插入“应答：满足”
        new_para = insert_paragraph_after(para, "应答：满足")

        # 调用始皇API生成应答（不超过200字）
        response = llm_callback(f"针对招标文件给出的内容，写不超过200字的应答：\n\n{text}", purpose="应答")

        # 插入生成的应答内容
        insert_paragraph_after(new_para, response)

    # 保存文档
    doc.save(output_file)
    print(f"[INFO] 已生成文件：{output_file}")


if __name__ == "__main__":
    input_file = "采购需求书.docx"
    output_file = "采购需求书-点对点应答.docx"
    insert_reply_with_llm(input_file, output_file)
