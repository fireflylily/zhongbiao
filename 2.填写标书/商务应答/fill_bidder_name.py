#!/usr/bin/env python3
# coding: utf-8

"""
fill_bidder_name.py
用法:
    python fill_bidder_name.py input.docx output.docx "智慧足迹数据科技有限公司"

说明:
 - 会在段落、表格、页眉、页脚中查找常见的“投标人/供应商名称”标签并填入公司名。
 - 尝试保持原 run 的文本格式（不清空样式）。
"""

import sys
import re
from docx import Document

# 要匹配的标签变体 - 修复过度匹配问题
LABELS_RE = r'((投标人|供应商|公司)(?:名称|全称)?(?:（盖章）|（公章）|（全称、盖章）|\(盖章\)|\(公章\)|\(全称、盖章\))?)'

# 精确匹配模式，特别处理长占位符格式
BIDDER_PATTERNS = [
    # 特殊格式：处理带长空格占位符和后缀的格式 - 供应商名称：                （加盖公章）
    re.compile(r'^(?P<label>(投标人|供应商|公司)(?:名称|全称)?)\s*(?P<sep>[:：])\s*(?P<placeholder>\s{2,}|\s*[_\-\u2014\u3000]{2,}|\s*＿{2,}|\s*——{2,})\s*(?P<suffix>（[^）]*）|\([^)]*\))?\s*$'),
    
    # 基本模式：投标人名称、供应商名称、供应商全称、公司名称等
    re.compile(rf'^(?P<label>{LABELS_RE})\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{{1,}}|＿+|——+|（\s*）)?$'),
    
    # 单独标签模式  
    re.compile(rf'^(?P<label>{LABELS_RE})(?P<sep>[:：]?)\s*$'),
    
    # 特殊格式：(供应商全称) 或 （供应商全称） - 直接替换括号内容
    re.compile(r'^(?P<prefix>[\(（])(?P<content>(?:投标人|供应商|公司)(?:名称|全称|住址)?)(?P<suffix>[\)）])\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+)?$'),
    
    # 公司名称特殊格式：公司名称（全称、盖章）：、公司名称（盖章）：
    re.compile(r'^(?P<label>公司名称(?:（全称、盖章）|（盖章）|\(全称、盖章\)|\(盖章\)))\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
]

# 保持向后兼容
PLACEHOLDER_RE = BIDDER_PATTERNS[0]
LABEL_INLINE_RE = BIDDER_PATTERNS[1]

def replace_in_runs(paragraph, name):
    """段落级文本合并处理 - 解决跨run的标签拆分问题"""
    changed = False
    
    # 合并段落中所有run的文本以进行完整分析
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text.strip():
        return False
        
    # 检查是否已经包含公司名称，避免重复处理
    if name in full_text:
        print(f"跳过已包含公司名称的段落: '{full_text.strip()}'")
        return False
    
    # 在合并文本上进行模式匹配
    for pattern in BIDDER_PATTERNS:
        match = pattern.search(full_text)
        if match:
            groups = match.groupdict()
            
            # 检查是否是括号格式 (供应商名称)
            if 'prefix' in groups and 'content' in groups and 'suffix' in groups:
                prefix = groups.get('prefix', '')
                content = groups.get('content', '')
                suffix = groups.get('suffix', '')
                sep = groups.get('sep', '')
                
                # 括号格式：直接替换括号内容
                if '住址' in content:
                    # 供应商住址：简化处理，使用公司名（独立文件没有地址信息）
                    new_full_text = f"{prefix}{name}{suffix}"
                else:
                    # 供应商名称：直接替换为公司名
                    new_full_text = f"{prefix}{name}{suffix}"
                if sep:
                    new_full_text += sep
            else:
                # 普通格式处理
                label = groups.get('label', '')
                sep = groups.get('sep', ':')
                suffix = groups.get('suffix', '')
                
                # 确保分隔符格式正确
                if not sep:
                    sep = ':'
                if sep and not sep.endswith(' '):
                    sep += ' ' if sep in [':', '：'] else ''
                
                # 生成替换文本 - 特别处理带后缀的格式
                if suffix:
                    # 有后缀的情况（如："（加盖公章）"），在公司名称后添加适当空格再加后缀
                    new_full_text = f"{label}{sep}{name}            {suffix}"
                elif 'placeholder' in groups and groups.get('placeholder'):
                    # 有占位符的情况，替换整个匹配
                    new_full_text = f"{label}{sep}{name}"
                else:
                    # 没有占位符的情况，在标签后添加公司名
                    new_full_text = f"{label}{sep}{name}"
            
            # 将新文本分配回原有的run结构
            # 清空所有run的文本
            for run in paragraph.runs:
                run.text = ""
            
            # 将新文本放入第一个run（保持格式）
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
                changed = True
                print(f"段落级匹配投标人名称: '{full_text.strip()}' -> '{new_full_text}'")
            
            break  # 匹配到一个模式就停止
    
    return changed

def process_paragraph(paragraph, name):
    return replace_in_runs(paragraph, name)

def process_table(table, name):
    changed_any = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if replace_in_runs(p, name):
                    changed_any = True
            for tbl in cell.tables:
                if process_table(tbl, name):
                    changed_any = True
    return changed_any

def process_doc(doc: Document, name: str) -> dict:
    stats = {'paragraphs_changed': 0, 'tables_changed': 0, 'headers_changed': 0, 'footers_changed': 0}
    # 正文段落
    for para in doc.paragraphs:
        if process_paragraph(para, name):
            stats['paragraphs_changed'] += 1
    # 表格
    for table in doc.tables:
        if process_table(table, name):
            stats['tables_changed'] += 1
    # 页眉页脚
    for section in doc.sections:
        header = section.header
        footer = section.footer
        header_changed = False
        footer_changed = False
        for para in header.paragraphs:
            if process_paragraph(para, name):
                header_changed = True
        for table in header.tables:
            if process_table(table, name):
                header_changed = True
        for para in footer.paragraphs:
            if process_paragraph(para, name):
                footer_changed = True
        for table in footer.tables:
            if process_table(table, name):
                footer_changed = True
        if header_changed:
            stats['headers_changed'] += 1
        if footer_changed:
            stats['footers_changed'] += 1
    return stats

def main(argv):
    if len(argv) < 4:
        print("Usage: python fill_bidder_name.py input.docx output.docx \"Company Name\"")
        return 1
    input_file = argv[1]
    output_file = argv[2]
    company_name = argv[3]

    print(f"Loading {input_file} ...")
    doc = Document(input_file)
    stats = process_doc(doc, company_name)
    doc.save(output_file)
    print(f"Saved to {output_file}.")
    print("Summary of changes:")
    print(stats)
    return 0

if __name__ == '__main__':
    sys.exit(main(sys.argv))
