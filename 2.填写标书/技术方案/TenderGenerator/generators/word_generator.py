"""
Word文档生成器
将技术方案内容导出为Word文档格式，参考招标文件样式
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，Word文档生成功能不可用")

class WordGenerator:
    """Word文档生成器类"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx 不可用，无法生成Word文档")
    
    def export_proposal_to_word(self, 
                               proposal_data: Dict[str, Any],
                               outline_data: Dict[str, Any],
                               output_file: str,
                               template_style: Optional[str] = None) -> bool:
        """
        将技术方案导出为Word文档
        
        Args:
            proposal_data: 技术方案数据
            outline_data: 方案大纲数据
            output_file: 输出文件路径
            template_style: 模板样式类型
            
        Returns:
            是否成功生成
        """
        if not DOCX_AVAILABLE:
            self.logger.error("无法生成Word文档：python-docx 不可用")
            return False
        
        try:
            self.logger.info(f"开始生成Word文档: {output_file}")
            
            # 创建新文档
            doc = Document()
            
            # 设置文档样式
            self._setup_document_styles(doc)
            
            # 添加标题页
            self._add_title_page(doc, proposal_data)
            
            # 添加目录（暂时跳过，可以后续添加）
            
            # 添加技术方案主体内容
            self._add_proposal_content(doc, proposal_data, outline_data)
            
            # 保存文档
            output_path = Path(output_file).resolve()
            os.makedirs(output_path.parent, exist_ok=True)
            doc.save(str(output_path))
            
            self.logger.info(f"Word文档生成完成: {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"生成Word文档失败: {e}")
            return False
    
    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        try:
            # 设置正文样式
            styles = doc.styles
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = '宋体'
            normal_font.size = Pt(12)
            
            # 创建标题样式
            heading_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']
            heading_sizes = [Pt(18), Pt(16), Pt(14), Pt(12)]
            
            for i, (style_name, size) in enumerate(zip(heading_styles, heading_sizes)):
                if style_name in styles:
                    style = styles[style_name]
                    font = style.font
                    font.name = '黑体'
                    font.size = size
                    font.bold = True
                    
                    # 设置段落格式
                    paragraph_format = style.paragraph_format
                    paragraph_format.space_before = Pt(12) if i == 0 else Pt(6)
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15
            
            self.logger.debug("文档样式设置完成")
            
        except Exception as e:
            self.logger.warning(f"设置文档样式时出错: {e}")
    
    def _add_title_page(self, doc: Document, proposal_data: Dict[str, Any]):
        """添加标题页"""
        try:
            # 主标题
            title = proposal_data.get('title', '技术方案')
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加一些空行
            doc.add_paragraph()
            doc.add_paragraph()
            
            # 项目信息表格
            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Table Grid'
            
            # 填充项目信息
            project_info = [
                ('项目名称', proposal_data.get('project_name', '中邮保险手机号实名认证服务项目')),
                ('方案类型', '技术方案'),
                ('编制日期', datetime.now().strftime('%Y年%m月%d日')),
                ('版本号', proposal_data.get('version', 'V1.0')),
                ('编制单位', proposal_data.get('company', '')),
                ('联系方式', proposal_data.get('contact', ''))
            ]
            
            for i, (label, value) in enumerate(project_info):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = value
                
                # 设置表格样式
                for j in range(2):
                    cell = info_table.cell(i, j)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(12)
            
            # 添加分页符
            doc.add_page_break()
            
            self.logger.debug("标题页添加完成")
            
        except Exception as e:
            self.logger.warning(f"添加标题页时出错: {e}")
    
    def _add_proposal_content(self, doc: Document, proposal_data: Dict[str, Any], outline_data: Dict[str, Any]):
        """添加技术方案主体内容"""
        try:
            # 获取章节内容
            sections = proposal_data.get('sections', [])
            
            if not sections:
                # 如果没有sections，尝试从content获取
                content = proposal_data.get('content', '')
                if content:
                    self._add_text_content(doc, content)
                    return
            
            # 添加各个章节
            for i, section in enumerate(sections):
                self._add_section(doc, section, level=1)
            
            self.logger.info(f"添加了 {len(sections)} 个章节")
            
        except Exception as e:
            self.logger.error(f"添加方案内容时出错: {e}")
    
    def _add_section(self, doc: Document, section: Dict[str, Any], level: int = 1):
        """添加单个章节"""
        try:
            # 添加章节标题
            title = section.get('title', '未命名章节')
            if level <= 4:
                heading = doc.add_heading(title, level=level)
            else:
                # 超过4级的标题用加粗段落
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.bold = True
                run.font.size = Pt(12)
            
            # 添加章节内容
            content = section.get('content', '')
            if content:
                # 将内容按段落分割
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        self._add_formatted_paragraph(doc, para_text.strip())
            
            # 递归添加子章节
            subsections = section.get('subsections', [])
            for subsection in subsections:
                self._add_section(doc, subsection, level + 1)
            
        except Exception as e:
            self.logger.warning(f"添加章节时出错: {e}")
    
    def _add_formatted_paragraph(self, doc: Document, text: str):
        """添加格式化段落"""
        try:
            # 检查是否包含特殊格式
            if self._is_list_item(text):
                # 列表项
                para = doc.add_paragraph(text, style='List Paragraph')
            elif text.startswith('**') and text.endswith('**'):
                # 加粗文本
                para = doc.add_paragraph()
                run = para.add_run(text[2:-2])
                run.bold = True
            else:
                # 普通段落
                para = doc.add_paragraph(text)
            
            # 设置段落格式
            paragraph_format = para.paragraph_format
            paragraph_format.line_spacing = 1.15
            paragraph_format.space_after = Pt(6)
            
        except Exception as e:
            self.logger.warning(f"添加段落时出错: {e}")
            # fallback: 添加简单段落
            doc.add_paragraph(text)
    
    def _add_text_content(self, doc: Document, content: str):
        """添加纯文本内容（作为备用方法）"""
        try:
            # 按行分割内容
            lines = content.split('\n')
            current_heading_level = 1
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 检查是否为标题
                if self._is_heading(line):
                    level = self._get_heading_level(line)
                    clean_title = self._clean_heading_text(line)
                    doc.add_heading(clean_title, level=level)
                    current_heading_level = level
                else:
                    # 普通段落
                    if line:
                        self._add_formatted_paragraph(doc, line)
            
        except Exception as e:
            self.logger.error(f"添加文本内容时出错: {e}")
    
    def _is_heading(self, text: str) -> bool:
        """判断是否为标题"""
        # 检查各种标题格式
        patterns = [
            r'^#+\s+',  # Markdown标题
            r'^\d+\.?\s+',  # 数字标题
            r'^[一二三四五六七八九十]+[、\.]\s*',  # 中文数字标题
            r'^第[一二三四五六七八九十]+章',  # 章节标题
            r'^\d+\.\d+\.?\s+',  # 多级数字标题
        ]
        
        return any(re.match(pattern, text) for pattern in patterns)
    
    def _get_heading_level(self, text: str) -> int:
        """获取标题级别"""
        if re.match(r'^#+\s+', text):
            return min(text.count('#'), 4)
        elif re.match(r'^第[一二三四五六七八九十]+章', text):
            return 1
        elif re.match(r'^\d+\.\d+\.\d+\.?\s+', text):
            return 3
        elif re.match(r'^\d+\.\d+\.?\s+', text):
            return 2
        elif re.match(r'^\d+\.?\s+', text):
            return 1
        else:
            return 2  # 默认二级标题
    
    def _clean_heading_text(self, text: str) -> str:
        """清理标题文本"""
        # 移除标题标记
        text = re.sub(r'^#+\s*', '', text)
        text = re.sub(r'^\d+\.?\s*', '', text)
        text = re.sub(r'^[一二三四五六七八九十]+[、\.]\s*', '', text)
        text = re.sub(r'^第[一二三四五六七八九十]+章\s*', '', text)
        return text.strip()
    
    def _is_list_item(self, text: str) -> bool:
        """判断是否为列表项"""
        patterns = [
            r'^[-*+]\s+',  # markdown列表
            r'^\d+[.)\s]+',  # 数字列表
            r'^[a-zA-Z][.)\s]+',  # 字母列表
            r'^[（(]\d+[）)]\s*',  # 括号数字列表
        ]
        
        return any(re.match(pattern, text) for pattern in patterns)
    
    def create_table_from_data(self, doc: Document, data: List[List[str]], headers: Optional[List[str]] = None):
        """创建数据表格"""
        try:
            if not data:
                return None
            
            rows = len(data) + (1 if headers else 0)
            cols = len(data[0]) if data else 0
            
            if cols == 0:
                return None
            
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 添加表头
            if headers:
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    # 设置表头样式
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = '宋体'
                            run.font.size = Pt(11)
                
                # 数据从第二行开始
                start_row = 1
            else:
                start_row = 0
            
            # 添加数据
            for i, row_data in enumerate(data):
                for j, cell_data in enumerate(row_data):
                    cell = table.cell(start_row + i, j)
                    cell.text = str(cell_data)
                    # 设置单元格样式
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(10)
            
            return table
            
        except Exception as e:
            self.logger.error(f"创建表格时出错: {e}")
            return None


# 全局Word生成器实例
_word_generator = None

def get_word_generator() -> WordGenerator:
    """获取全局Word生成器实例"""
    global _word_generator
    if _word_generator is None:
        _word_generator = WordGenerator()
    return _word_generator