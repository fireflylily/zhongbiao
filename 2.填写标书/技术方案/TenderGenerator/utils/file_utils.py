"""
文件处理工具
支持PDF、Word、Excel等格式的读取和处理
"""

import os
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
import tempfile

try:
    import PyPDF2
    import docx
    from docx.document import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    import openpyxl
    PDF_AVAILABLE = True
    DOCX_AVAILABLE = True
    EXCEL_AVAILABLE = True
except ImportError as e:
    print(f"警告: 缺少依赖包 {e}，某些功能可能不可用")
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False
    EXCEL_AVAILABLE = False

try:
    from ..config import (
        MIN_PARAGRAPH_LENGTH, MAX_PARAGRAPH_LENGTH,
        SECTION_PATTERNS, MAX_DOCUMENT_PAGES
    )
except ImportError:
    import sys
    from pathlib import Path
    sys.path.append(str(Path(__file__).parent.parent))
    from config import (
        MIN_PARAGRAPH_LENGTH, MAX_PARAGRAPH_LENGTH,
        SECTION_PATTERNS, MAX_DOCUMENT_PAGES
    )

class FileUtils:
    """文件处理工具类"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def get_file_type(self, file_path: str) -> str:
        """
        获取文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件类型字符串
        """
        ext = Path(file_path).suffix.lower()
        type_mapping = {
            '.pdf': 'pdf',
            '.doc': 'word',
            '.docx': 'word',
            '.xls': 'excel',
            '.xlsx': 'excel',
            '.txt': 'text',
            '.json': 'json'
        }
        return type_mapping.get(ext, 'unknown')
    
    def read_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        读取PDF文件
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            包含页面内容的字典
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available")
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # 检查页数限制
                if len(reader.pages) > MAX_DOCUMENT_PAGES:
                    self.logger.warning(f"PDF页数({len(reader.pages)})超过限制({MAX_DOCUMENT_PAGES})")
                
                pages = []
                for i, page in enumerate(reader.pages[:MAX_DOCUMENT_PAGES]):
                    text = page.extract_text()
                    if text.strip():
                        pages.append({
                            'page_num': i + 1,
                            'content': text.strip()
                        })
                
                return {
                    'type': 'pdf',
                    'file_path': file_path,
                    'total_pages': len(reader.pages),
                    'processed_pages': len(pages),
                    'pages': pages
                }
                
        except Exception as e:
            self.logger.error(f"读取PDF文件失败: {e}")
            return {'type': 'pdf', 'error': str(e)}
    
    def read_word(self, file_path: str) -> Dict[str, Any]:
        """
        读取Word文档
        
        Args:
            file_path: Word文件路径
            
        Returns:
            包含段落和表格的字典
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        try:
            doc = docx.Document(file_path)
            
            paragraphs = []
            tables = []
            
            # 处理段落
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        'index': i,
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # 处理表格
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                tables.append({
                    'index': i,
                    'data': table_data,
                    'rows': len(table.rows),
                    'cols': len(table.rows[0].cells) if table.rows else 0
                })
            
            return {
                'type': 'word',
                'file_path': file_path,
                'paragraphs': paragraphs,
                'tables': tables,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables)
            }
            
        except Exception as e:
            self.logger.error(f"读取Word文档失败: {e}")
            return {'type': 'word', 'error': str(e)}
    
    def read_excel(self, file_path: str) -> Dict[str, Any]:
        """
        读取Excel文件
        
        Args:
            file_path: Excel文件路径
            
        Returns:
            包含工作表数据的字典
        """
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl not available")
        
        try:
            workbook = openpyxl.load_workbook(file_path)
            sheets_data = {}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                data = []
                
                # 读取所有行
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):  # 跳过空行
                        data.append(list(row))
                
                sheets_data[sheet_name] = {
                    'data': data,
                    'rows': len(data),
                    'cols': len(data[0]) if data else 0
                }
            
            return {
                'type': 'excel',
                'file_path': file_path,
                'sheets': sheets_data,
                'sheet_names': list(sheets_data.keys())
            }
            
        except Exception as e:
            self.logger.error(f"读取Excel文件失败: {e}")
            return {'type': 'excel', 'error': str(e)}
    
    def read_text(self, file_path: str, encoding: str = 'utf-8') -> Dict[str, Any]:
        """
        读取文本文件
        
        Args:
            file_path: 文本文件路径
            encoding: 文件编码
            
        Returns:
            包含文本内容的字典
        """
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            return {
                'type': 'text',
                'file_path': file_path,
                'content': content,
                'length': len(content),
                'lines': len(content.split('\n'))
            }
            
        except UnicodeDecodeError:
            # 尝试其他编码
            for enc in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=enc) as file:
                        content = file.read()
                    return {
                        'type': 'text',
                        'file_path': file_path,
                        'content': content,
                        'encoding': enc,
                        'length': len(content),
                        'lines': len(content.split('\n'))
                    }
                except:
                    continue
            
            self.logger.error(f"无法识别文本文件编码: {file_path}")
            return {'type': 'text', 'error': '编码识别失败'}
            
        except Exception as e:
            self.logger.error(f"读取文本文件失败: {e}")
            return {'type': 'text', 'error': str(e)}
    
    def read_file(self, file_path: str) -> Dict[str, Any]:
        """
        智能读取文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件内容字典
        """
        if not os.path.exists(file_path):
            return {'error': f'文件不存在: {file_path}'}
        
        file_type = self.get_file_type(file_path)
        
        try:
            if file_type == 'pdf':
                return self.read_pdf(file_path)
            elif file_type == 'word':
                return self.read_word(file_path)
            elif file_type == 'excel':
                return self.read_excel(file_path)
            elif file_type == 'text':
                return self.read_text(file_path)
            else:
                # 尝试作为文本文件读取
                return self.read_text(file_path)
                
        except Exception as e:
            self.logger.error(f"读取文件失败: {e}")
            return {'error': str(e)}
    
    def extract_text_content(self, file_data: Dict[str, Any]) -> str:
        """
        从文件数据中提取纯文本内容
        
        Args:
            file_data: 文件数据字典
            
        Returns:
            纯文本内容
        """
        file_type = file_data.get('type', '')
        text_content = ""
        
        try:
            if file_type == 'pdf':
                pages = file_data.get('pages', [])
                text_content = '\n\n'.join([page['content'] for page in pages])
                
            elif file_type == 'word':
                # 合并段落和表格内容
                paragraphs = file_data.get('paragraphs', [])
                tables = file_data.get('tables', [])
                
                # 段落文本
                para_text = '\n'.join([para['text'] for para in paragraphs])
                
                # 表格文本
                table_texts = []
                for i, table in enumerate(tables):
                    table_text = f"[表格 {i+1}]\n"
                    for row in table['data']:
                        row_text = '\t'.join([str(cell) if cell else '' for cell in row])
                        if row_text.strip():
                            table_text += row_text + '\n'
                    table_texts.append(table_text)
                
                # 合并段落和表格
                if table_texts:
                    text_content = para_text + '\n\n' + '\n\n'.join(table_texts)
                else:
                    text_content = para_text
                
            elif file_type == 'text':
                text_content = file_data.get('content', '')
                
            elif file_type == 'excel':
                # 将所有工作表的文本内容合并
                sheets = file_data.get('sheets', {})
                sheet_texts = []
                for sheet_name, sheet_data in sheets.items():
                    sheet_text = f"[工作表: {sheet_name}]\\n"
                    for row in sheet_data.get('data', []):
                        row_text = '\\t'.join([str(cell) if cell is not None else '' for cell in row])
                        if row_text.strip():
                            sheet_text += row_text + '\\n'
                    sheet_texts.append(sheet_text)
                text_content = '\\n\\n'.join(sheet_texts)
                
        except Exception as e:
            self.logger.error(f"提取文本内容失败: {e}")
        
        return text_content
    
    def split_into_sections(self, text: str) -> List[Dict[str, Any]]:
        """
        将文本分割为章节
        
        Args:
            text: 输入文本
            
        Returns:
            章节列表
        """
        sections = []
        lines = text.split('\\n')
        current_section = {
            'title': '',
            'content': '',
            'level': 0
        }
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查是否为章节标题
            section_match = None
            for i, pattern in enumerate(SECTION_PATTERNS):
                match = re.match(pattern, line)
                if match:
                    section_match = {'level': i + 1, 'title': line}
                    break
            
            if section_match:
                # 保存当前章节
                if current_section['content'].strip():
                    sections.append(current_section.copy())
                
                # 开始新章节
                current_section = {
                    'title': section_match['title'],
                    'content': '',
                    'level': section_match['level']
                }
            else:
                # 添加到当前章节内容
                if len(line) >= MIN_PARAGRAPH_LENGTH:
                    current_section['content'] += line + '\\n'
        
        # 添加最后一个章节
        if current_section['content'].strip():
            sections.append(current_section)
        
        return sections
    
    def save_json(self, data: Any, file_path: str, indent: int = 2, ensure_ascii: bool = False):
        """
        保存JSON文件
        
        Args:
            data: 要保存的数据
            file_path: 输出文件路径
            indent: 缩进空格数
            ensure_ascii: 是否确保ASCII编码
        """
        try:
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, 'w', encoding='utf-8') as file:
                json.dump(data, file, indent=indent, ensure_ascii=ensure_ascii)
            self.logger.info(f"JSON文件保存成功: {file_path}")
        except Exception as e:
            self.logger.error(f"保存JSON文件失败: {e}")
            raise
    
    def load_json(self, file_path: str) -> Any:
        """
        加载JSON文件
        
        Args:
            file_path: JSON文件路径
            
        Returns:
            解析后的数据
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except Exception as e:
            self.logger.error(f"加载JSON文件失败: {e}")
            raise
    
    def create_temp_file(self, content: str, suffix: str = '.txt') -> str:
        """
        创建临时文件
        
        Args:
            content: 文件内容
            suffix: 文件后缀
            
        Returns:
            临时文件路径
        """
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix=suffix, delete=False, encoding='utf-8')
        temp_file.write(content)
        temp_file.close()
        return temp_file.name
    
    def cleanup_temp_files(self, file_paths: List[str]):
        """
        清理临时文件
        
        Args:
            file_paths: 要清理的文件路径列表
        """
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    os.unlink(file_path)
            except Exception as e:
                self.logger.warning(f"清理临时文件失败: {file_path}, {e}")


# 全局文件工具实例
_file_utils = None

def get_file_utils() -> FileUtils:
    """获取全局文件工具实例"""
    global _file_utils
    if _file_utils is None:
        _file_utils = FileUtils()
    return _file_utils