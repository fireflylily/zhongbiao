#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import json
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def test_business_response():
    """测试商务应答处理功能"""
    
    test_input = "test_business_response_input.docx"
    test_output = "test_business_response_output.docx"
    
    # 创建测试文档，包含各种公司信息字段
    doc = Document()
    doc.add_paragraph("=== 商务应答文档测试 ===")
    doc.add_paragraph("供应商名称：                                    ")
    doc.add_paragraph("法定代表人：_________________")
    doc.add_paragraph("注册地址：________________________________")
    doc.add_paragraph("联系电话：______________")
    doc.add_paragraph("统一社会信用代码：________________________")
    doc.add_paragraph("注册资本：_________________")
    doc.add_paragraph("邮政编码：________")
    doc.add_paragraph("开户银行：________________________")
    doc.add_paragraph("银行账号：________________________")
    doc.add_paragraph("成立日期：______________")
    doc.add_paragraph("项目名称：________________________________")
    doc.add_paragraph("招标编号：__________________")
    doc.add_paragraph("日期：______________")
    doc.add_paragraph("供应商名称（加盖公章）：                     ")
    doc.save(test_input)
    
    print(f"✅ 创建测试文档: {test_input}")
    
    # 模拟公司信息（从JSON文件读取的格式）
    company_info = {
        "companyName": "智慧足迹数据科技有限公司",
        "establishDate": "2018-03-15",
        "legalRepresentative": "张三",
        "socialCreditCode": "91110000MA01234567",
        "registeredCapital": "1000万元",
        "companyType": "有限责任公司",
        "fixedPhone": "010-12345678",
        "postalCode": "100010",
        "registeredAddress": "北京市朝阳区科技园区100号",
        "officeAddress": "北京市朝阳区科技园区100号",
        "website": "www.wisepath.com",
        "employeeCount": "50-100人",
        "bankName": "中国工商银行北京科技园支行",
        "bankAccount": "620200001234567890"
    }
    
    # 项目信息
    project_name = "智慧城市数据平台建设项目"
    tender_no = "ZSCG-2025-001"
    date_text = "2025年9月6日"
    
    print(f"\n=== 测试信息 ===")
    print(f"公司名称: {company_info['companyName']}")
    print(f"项目名称: {project_name}")
    print(f"招标编号: {tender_no}")
    print(f"日期: {date_text}")
    
    try:
        # 执行处理
        processor = MCPBidderNameProcessor()
        
        print(f"\n=== 开始处理 ===")
        
        result = processor.process_business_response(
            input_file=test_input,
            output_file=test_output,
            company_info=company_info,
            project_name=project_name,
            tender_no=tender_no,
            date_text=date_text
        )
        
        print(f"\n处理结果: {result}")
        
        if result.get('success'):
            print("✅ 处理成功!")
            
            # 处理后检查
            print(f"\n=== 处理后检查 ===")
            if os.path.exists(test_output):
                doc_after = Document(test_output)
                for i, para in enumerate(doc_after.paragraphs):
                    para_text = para.text
                    if any(keyword in para_text for keyword in ["供应商名称", "法定代表人", "注册地址", "项目名称", "招标编号"]):
                        print(f"段落 #{i}: '{para_text}'")
                        
                        # 检查是否已填写信息
                        has_info = False
                        if company_info['companyName'] in para_text:
                            print(f"    ✅ 包含公司名称")
                            has_info = True
                        if company_info['legalRepresentative'] in para_text:
                            print(f"    ✅ 包含法定代表人")
                            has_info = True
                        if company_info['registeredAddress'] in para_text:
                            print(f"    ✅ 包含注册地址")
                            has_info = True
                        if project_name in para_text:
                            print(f"    ✅ 包含项目名称")
                            has_info = True
                        if tender_no in para_text:
                            print(f"    ✅ 包含招标编号")
                            has_info = True
                        if date_text in para_text:
                            print(f"    ✅ 包含日期")
                            has_info = True
                        
                        if not has_info and ("___" in para_text or "   " in para_text):
                            print(f"    ⚠️ 仍有占位符未填写")
            else:
                print(f"❌ 输出文件不存在: {test_output}")
            
            # 分析处理统计
            stats = result.get('stats', {})
            print(f"\n=== 处理统计 ===")
            print(f"总替换次数: {stats.get('total_replacements', 0)}")
            print(f"公司信息字段处理数: {stats.get('info_fields_processed', 0)}")
            
            patterns = stats.get('patterns_found', [])
            if patterns:
                print(f"处理的字段:")
                for pattern in patterns:
                    if 'field_name' in pattern:
                        # 公司信息字段
                        print(f"  - {pattern['field_name']}: '{pattern['original_text']}'")
                    else:
                        # 供应商名称字段
                        print(f"  - {pattern.get('description', 'Unknown')}")
            
        else:
            print(f"❌ 处理失败: {result.get('error', 'Unknown error')}")
            
    except Exception as e:
        print(f"❌ 处理异常: {e}")
        import traceback
        traceback.print_exc()
    
    print(f"\n=== 文件保留用于检查 ===")
    print(f"输入文件: {test_input}")
    print(f"输出文件: {test_output}")
    print("请手动检查输出文件内容")

if __name__ == "__main__":
    test_business_response()