#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
采购需求点对点应答系统
功能：从采购需求文档中提取每个需求条目，生成标准格式的技术应答
"""

import requests
import json
import logging
import configparser
import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import sys
import os
# TenderInfoExtractor dependency removed - using standalone methods
from api_config import get_api_key, API_CONFIG

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("未安装python-docx库，将无法生成Word文档")

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('requirement_response.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class RequirementProcessor:
    """采购需求处理器"""
    
    def __init__(self):
        self.config_file = 'requirement_config.ini'
        self.template_file = 'config/response_templates.json'
        self.patterns_file = 'config/requirement_patterns.json'
        
        # 加载配置
        self.load_templates()
        self.load_patterns()
    
    def load_templates(self):
        """加载应答模板"""
        try:
            if os.path.exists(self.template_file):
                with open(self.template_file, 'r', encoding='utf-8') as f:
                    self.templates = json.load(f)
                    logger.info("应答模板加载成功")
            else:
                # 使用默认模板
                self.templates = self.get_default_templates()
                logger.warning(f"模板文件不存在，使用默认模板")
        except Exception as e:
            logger.error(f"加载应答模板失败: {e}")
            self.templates = self.get_default_templates()
    
    def load_patterns(self):
        """加载需求识别模式"""
        try:
            if os.path.exists(self.patterns_file):
                with open(self.patterns_file, 'r', encoding='utf-8') as f:
                    self.patterns = json.load(f)
                    logger.info("需求识别模式加载成功")
            else:
                self.patterns = self.get_default_patterns()
                logger.warning("模式文件不存在，使用默认模式")
        except Exception as e:
            logger.error(f"加载需求识别模式失败: {e}")
            self.patterns = self.get_default_patterns()
    
    def get_default_templates(self) -> Dict:
        """获取默认应答模板"""
        return {
            "硬件配置": "应答：满足。我方提供的设备配置为：{具体配置}，完全满足采购需求。",
            "软件功能": "应答：满足。我方系统具备{功能名称}功能，采用{技术方案}实现，能够满足业务需求。",
            "性能指标": "应答：满足。系统性能指标达到{具体指标}，超过需求标准。",
            "技术规范": "应答：满足。我方产品严格遵循{标准名称}标准，符合技术规范要求。",
            "服务保障": "应答：满足。我方提供{服务内容}，确保项目顺利实施。",
            "资质证明": "应答：满足。我方具备{资质名称}，满足资质要求。",
            "通用模板": "应答：满足。针对该需求，我方的技术方案为：{具体方案}。"
        }
    
    def get_default_patterns(self) -> Dict:
        """获取默认需求识别模式"""
        return {
            "编号模式": [
                r'^(\d+)\s*[、．.]',  # 1、 1. 1．
                r'^(\d+\.\d+)\s*[、．.]',  # 1.1、 1.2.
                r'^\((\d+)\)',  # (1) (2)
                r'^([A-Z])\)',  # A) B)
                r'^([a-z])\)',  # a) b)
                r'^([IVXLC]+)\s*[、．.]',  # I、 II、 罗马数字
            ],
            "关键词": [
                "要求", "需求", "应", "必须", "应当", "需要", "具备",
                "支持", "提供", "实现", "满足", "符合", "遵循",
                "不少于", "不低于", "≥", ">=", "以上"
            ],
            "需求句式": [
                r'供应商应.*',
                r'系统需.*',
                r'设备应.*',
                r'软件需.*',
                r'.*要求.*',
                r'.*需求.*',
                r'.*应具备.*'
            ]
        }
    
    def llm_callback(self, prompt: str, purpose: str = "需求分析") -> str:
        """
        调用LLM API进行需求分析和应答生成
        """
        url = API_CONFIG["base_url"]
        headers = {
            "Authorization": f"Bearer {get_api_key()}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": API_CONFIG["model"],
            "messages": [
                {"role": "system", "content": "你是一名专业的技术方案专家，擅长为采购需求提供专业的技术应答。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": API_CONFIG["temperature"],
            "max_tokens": API_CONFIG["max_tokens"]
        }
        
        try:
            logger.info(f"[LLM调用开始] 任务: {purpose}")
            response = requests.post(url, headers=headers, json=payload, timeout=API_CONFIG["timeout"])
            
            if response.status_code != 200:
                logger.error(f"API调用失败: {response.status_code}")
                return "应答：满足。我方将根据具体需求提供相应的技术方案，完全满足采购要求。"
            
            result = response.json()
            if "choices" in result and len(result["choices"]) > 0:
                content = result["choices"][0].get("message", {}).get("content", "").strip()
                return content
            else:
                return "应答：满足。我方将根据具体需求提供相应的技术方案，完全满足采购要求。"
                
        except Exception as e:
            logger.error(f"LLM调用失败: {e}")
            return "应答：满足。我方将根据具体需求提供相应的技术方案，完全满足采购要求。"
    
    def read_document(self, file_path: str) -> str:
        """
        读取文档内容，支持多种格式
        """
        try:
            logger.info(f"正在读取文档: {file_path}")
            
            # 支持多种文档格式
            if file_path.lower().endswith('.txt'):
                content = self._read_text_file(file_path)
            elif file_path.lower().endswith(('.doc', '.docx')):
                content = self._read_word_document(file_path)
            else:
                # 默认按文本文件处理
                content = self._read_text_file(file_path)
            
            logger.info(f"文档读取成功，内容长度: {len(content)} 字符")
            return content
            
        except FileNotFoundError:
            logger.error(f"文档文件不存在: {file_path}")
            raise
        except Exception as e:
            logger.error(f"读取文档失败: {e}")
            raise
    
    def _read_text_file(self, file_path: str) -> str:
        """读取文本文件，尝试多种编码"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.info(f"使用 {encoding} 编码成功读取文本文件")
                return content
            except UnicodeDecodeError:
                continue
        
        raise UnicodeDecodeError(f"无法使用任何编码读取文件: {file_path}")
    
    def _read_word_document(self, file_path: str) -> str:
        """读取Word文档"""
        content = ""
        
        # 方法1: 使用python-docx
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            content = '\n'.join([para.text for para in doc.paragraphs])
            logger.info("使用python-docx成功读取Word文档")
            return content
        except ImportError:
            logger.warning("未安装python-docx库，尝试其他方法")
        except Exception as e:
            logger.warning(f"python-docx读取失败: {e}，尝试其他方法")
        
        # 方法2: 尝试作为zip文件解析（docx本质上是zip）
        if file_path.lower().endswith('.docx'):
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    xml_content = zip_file.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    
                    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    text_elements = root.findall('.//w:t', namespaces)
                    content = '\n'.join([elem.text or '' for elem in text_elements])
                    
                logger.info("使用ZIP解析成功读取Word文档")
                return content
            except Exception as e:
                logger.warning(f"ZIP解析失败: {e}")
        
        # 最后尝试按二进制读取并转换
        try:
            with open(file_path, 'rb') as f:
                raw_content = f.read()
            content = raw_content.decode('utf-8', errors='ignore')
            logger.warning("使用二进制读取Word文档，可能包含乱码")
            return content
        except Exception as e:
            logger.error(f"所有方法都失败了: {e}")
            raise Exception(f"无法读取Word文档 {file_path}。请尝试将文档转换为txt格式。")
    
    def extract_project_info(self, document_content: str) -> Dict[str, str]:
        """
        简单提取项目信息
        """
        project_info = {
            "project_name": "未知项目",
            "project_number": "未知编号",
            "tenderer": "",
            "agency": ""
        }
        
        # 简单的正则提取
        name_patterns = [
            r'项目名称[：:]\s*([^\n\r，,]+)',
            r'project_name[：:]\s*([^\n\r，,]+)'
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, document_content)
            if match:
                project_info["project_name"] = match.group(1).strip()
                break
        
        number_patterns = [
            r'项目编号[：:]\s*([^\n\r，,]+)',
            r'招标编号[：:]\s*([^\n\r，,]+)'
        ]
        
        for pattern in number_patterns:
            match = re.search(pattern, document_content)
            if match:
                project_info["project_number"] = match.group(1).strip()
                break
        
        return project_info
    
    def extract_requirements(self, document_content: str) -> List[Dict]:
        """
        从文档中提取需求条目
        """
        logger.info("开始提取需求条目...")
        
        # 使用LLM进行智能需求提取
        prompt = f"""
请从以下采购需求文档中提取所有的具体需求条目。

提取规则：
1. 每个需求条目应该是一个完整的要求或规格
2. 包含编号的条目优先提取
3. 包含"要求"、"需求"、"应"、"必须"等关键词的句子
4. 技术规格、功能要求、性能指标等

请按以下JSON格式返回，不要包含任何其他文字：
{{
    "requirements": [
        {{
            "id": "条目编号或序号",
            "content": "需求条目的完整内容",
            "type": "需求类型(如：硬件配置/软件功能/性能指标/服务保障等)",
            "keywords": ["关键词1", "关键词2"]
        }}
    ]
}}

文档内容：
{document_content[:3000]}  # 限制内容长度避免token超限
"""
        
        try:
            response = self.llm_callback(prompt, "需求提取")
            logger.info(f"LLM需求提取响应: {response[:200]}...")
            
            # 解析JSON响应
            if "```json" in response:
                response = response.split("```json")[1].split("```")[0].strip()
            elif "```" in response:
                response = response.split("```")[1].strip()
            
            result = json.loads(response)
            requirements = result.get("requirements", [])
            
            logger.info(f"成功提取 {len(requirements)} 个需求条目")
            return requirements
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON解析失败: {e}")
            # 降级使用正则表达式提取
            return self._fallback_extract_requirements(document_content)
        except Exception as e:
            logger.error(f"需求提取失败: {e}")
            return self._fallback_extract_requirements(document_content)
    
    def _fallback_extract_requirements(self, content: str) -> List[Dict]:
        """
        备用需求提取方法：使用正则表达式
        """
        logger.warning("使用备用方法提取需求条目")
        requirements = []
        
        # 按行分割内容
        lines = content.split('\n')
        current_id = 1
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检查是否包含需求关键词
            has_keyword = any(keyword in line for keyword in self.patterns["关键词"])
            
            # 检查编号格式
            has_number = any(re.match(pattern, line) for pattern in self.patterns["编号模式"])
            
            # 检查需求句式
            has_pattern = any(re.search(pattern, line) for pattern in self.patterns["需求句式"])
            
            if (has_keyword or has_number or has_pattern) and len(line) > 10:
                requirement = {
                    "id": str(current_id),
                    "content": line,
                    "type": "通用需求",
                    "keywords": [kw for kw in self.patterns["关键词"] if kw in line]
                }
                requirements.append(requirement)
                current_id += 1
        
        logger.info(f"备用方法提取到 {len(requirements)} 个需求条目")
        return requirements
    
    def classify_requirement(self, requirement: Dict) -> str:
        """
        对需求进行分类
        """
        content = requirement["content"].lower()
        keywords = requirement.get("keywords", [])
        
        # 硬件相关
        hardware_keywords = ["cpu", "内存", "存储", "硬盘", "服务器", "设备", "配置", "处理器"]
        if any(kw in content for kw in hardware_keywords):
            return "硬件配置"
            
        # 软件相关
        software_keywords = ["软件", "系统", "应用", "程序", "功能", "模块", "接口"]
        if any(kw in content for kw in software_keywords):
            return "软件功能"
            
        # 性能相关
        performance_keywords = ["性能", "速度", "并发", "响应", "吞吐", "处理能力", "不少于", "不低于"]
        if any(kw in content for kw in performance_keywords):
            return "性能指标"
            
        # 服务相关
        service_keywords = ["服务", "支持", "培训", "维护", "实施", "部署"]
        if any(kw in content for kw in service_keywords):
            return "服务保障"
            
        # 资质相关  
        qualification_keywords = ["资质", "证书", "认证", "经验", "案例", "业绩"]
        if any(kw in content for kw in qualification_keywords):
            return "资质证明"
            
        # 技术规范
        standard_keywords = ["标准", "规范", "协议", "符合", "遵循", "兼容"]
        if any(kw in content for kw in standard_keywords):
            return "技术规范"
            
        return "通用模板"
    
    def generate_response(self, requirement: Dict) -> str:
        """
        为单个需求生成应答
        """
        req_type = self.classify_requirement(requirement)
        template = self.templates.get(req_type, self.templates["通用模板"])
        
        # 使用LLM生成具体的技术方案
        prompt = f"""
请为以下采购需求生成一个专业的技术应答方案。

需求内容：{requirement['content']}
需求类型：{req_type}

要求：
1. 应答必须以"应答：满足。"开头
2. 后续内容要具体、专业、可信
3. 如果是技术指标，要给出具体参数
4. 如果是功能需求，要说明实现方案
5. 内容要简洁明了，1-2句话概括

请直接返回应答内容，不要包含其他解释：
"""
        
        try:
            response = self.llm_callback(prompt, "应答生成")
            
            # 确保以"应答：满足。"开头
            if not response.startswith("应答：满足。"):
                if response.startswith("应答："):
                    response = response.replace("应答：", "应答：满足。", 1)
                else:
                    response = f"应答：满足。{response}"
            
            return response.strip()
            
        except Exception as e:
            logger.error(f"生成应答失败: {e}")
            # 使用基础模板
            return template.format(具体方案="根据需求提供相应的技术方案")
    
    def process_requirements(self, file_path: str) -> Dict:
        """
        处理需求文档的主要方法
        """
        logger.info(f"开始处理需求文档: {file_path}")
        
        try:
            # 1. 读取文档内容
            document_content = self.read_document(file_path)
            logger.info(f"文档读取成功，内容长度: {len(document_content)}")
            
            # 2. 提取项目信息
            project_info = self.extract_project_info(document_content)
            
            # 3. 提取需求条目
            requirements = self.extract_requirements(document_content)
            
            if not requirements:
                raise ValueError("未能提取到任何需求条目")
            
            # 4. 为每个需求生成应答
            responses = []
            for i, req in enumerate(requirements, 1):
                logger.info(f"正在处理第 {i}/{len(requirements)} 个需求...")
                response = self.generate_response(req)
                
                responses.append({
                    "requirement": req,
                    "response": response
                })
            
            result = {
                "project_info": project_info,
                "requirements_count": len(requirements),
                "responses": responses,
                "processing_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            logger.info(f"处理完成，共生成 {len(responses)} 条应答")
            return result
            
        except Exception as e:
            logger.error(f"处理需求文档失败: {e}")
            raise
    
    def export_to_word(self, result: Dict, output_path: str = None) -> str:
        """
        导出结果到Word文档
        """
        if not DOCX_AVAILABLE:
            logger.error("未安装python-docx库，无法生成Word文档")
            return self.export_to_text(result, output_path)
        
        try:
            # 创建文档
            doc = Document()
            
            # 设置文档标题
            project_name = result["project_info"].get("project_name", "未知项目")
            title = f"采购需求应答书 - {project_name}"
            title_p = doc.add_heading(title, level=1)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加项目信息表
            doc.add_heading("项目信息", level=2)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ("项目名称", result["project_info"].get("project_name", "未知")),
                ("项目编号", result["project_info"].get("project_number", "未知")),
                ("需求条目数量", str(result["requirements_count"])),
                ("生成时间", result["processing_time"])
            ]
            
            for i, (key, value) in enumerate(info_data):
                info_table.cell(i, 0).text = key
                info_table.cell(i, 1).text = value
            
            # 添加应答汇总表
            doc.add_heading("需求应答汇总", level=2)
            
            if result["responses"]:
                # 创建表格
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                table.autofit = False
                
                # 设置表头
                header_cells = table.rows[0].cells
                header_cells[0].text = "序号"
                header_cells[1].text = "采购需求"
                header_cells[2].text = "供应商应答"
                
                # 设置列宽
                table.columns[0].width = Inches(0.6)
                table.columns[1].width = Inches(3.0)
                table.columns[2].width = Inches(3.0)
                
                # 添加数据行
                for i, resp in enumerate(result["responses"], 1):
                    row = table.add_row()
                    cells = row.cells
                    cells[0].text = str(i)
                    cells[1].text = resp["requirement"]["content"]
                    cells[2].text = resp["response"]
                    
                    # 设置行高和字体
                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
            
            # 添加详细应答部分
            doc.add_page_break()
            doc.add_heading("详细应答内容", level=2)
            
            for i, resp in enumerate(result["responses"], 1):
                # 需求标题
                req_heading = doc.add_heading(f"{i}. 需求条目", level=3)
                
                # 原始需求
                doc.add_paragraph("原文需求：", style='List Bullet')
                req_p = doc.add_paragraph(resp["requirement"]["content"])
                req_p.style = 'Quote'
                
                # 应答内容
                doc.add_paragraph("供应商应答：", style='List Bullet')
                resp_p = doc.add_paragraph(resp["response"])
                
                # 添加分隔线
                if i < len(result["responses"]):
                    separator = doc.add_paragraph("─" * 50)
                    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()  # 空行
            
            # 生成输出文件名
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_project_name = re.sub(r'[^\w\s-]', '', project_name).strip()
                output_path = f"采购需求应答书_{safe_project_name}_{timestamp}.docx"
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path) if os.path.dirname(output_path) else "output"
            if not os.path.exists(output_dir) and output_dir != "":
                os.makedirs(output_dir)
            
            # 保存文档
            if not output_path.endswith('.docx'):
                output_path += '.docx'
            
            doc.save(output_path)
            logger.info(f"Word文档已保存至: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")
            # 降级到文本格式
            return self.export_to_text(result, output_path)
    
    def export_to_text(self, result: Dict, output_path: str = None) -> str:
        """
        导出结果到文本文件（备用方案）
        """
        try:
            project_name = result["project_info"].get("project_name", "未知项目")
            
            # 生成输出文件名
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_project_name = re.sub(r'[^\w\s-]', '', project_name).strip()
                output_path = f"采购需求应答书_{safe_project_name}_{timestamp}.txt"
            
            if not output_path.endswith('.txt'):
                output_path += '.txt'
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write(f"采购需求应答书 - {project_name}\n")
                f.write("=" * 60 + "\n\n")
                
                # 项目信息
                f.write("项目信息：\n")
                f.write(f"  项目名称：{result['project_info'].get('project_name', '未知')}\n")
                f.write(f"  项目编号：{result['project_info'].get('project_number', '未知')}\n")
                f.write(f"  需求条目数量：{result['requirements_count']}\n")
                f.write(f"  生成时间：{result['processing_time']}\n\n")
                
                # 详细应答
                f.write("详细应答内容：\n")
                f.write("-" * 60 + "\n\n")
                
                for i, resp in enumerate(result["responses"], 1):
                    f.write(f"{i}. 需求条目\n")
                    f.write(f"原文需求：{resp['requirement']['content']}\n\n")
                    f.write(f"供应商应答：{resp['response']}\n")
                    f.write("\n" + "─" * 50 + "\n\n")
            
            logger.info(f"文本文件已保存至: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"导出文本文件失败: {e}")
            raise

def main():
    """主函数"""
    if len(sys.argv) != 2:
        print("使用方法: python requirement_processor.py <需求文档路径>")
        print("例如: python requirement_processor.py ./采购需求书.docx")
        return
    
    file_path = sys.argv[1]
    
    try:
        processor = RequirementProcessor()
        result = processor.process_requirements(file_path)
        
        print(f"\n处理完成!")
        print(f"项目名称: {result['project_info'].get('project_name', '未知')}")
        print(f"需求条目数量: {result['requirements_count']}")
        print(f"处理时间: {result['processing_time']}")
        
        # 导出到Word文档
        try:
            output_file = processor.export_to_word(result)
            print(f"\n应答文档已生成: {output_file}")
        except Exception as e:
            logger.error(f"导出文档失败: {e}")
            print(f"导出文档失败: {e}")
        
        # 简单输出预览
        print(f"\n前3条应答预览:")
        for i, resp in enumerate(result['responses'][:3], 1):
            print(f"{i}. 需求: {resp['requirement']['content'][:50]}...")
            print(f"   应答: {resp['response'][:80]}...")
            print()
        
        print("详细结果已保存到日志文件: requirement_response.log")
        
    except Exception as e:
        logger.error(f"程序执行失败: {e}")
        print(f"程序执行失败: {e}")

if __name__ == "__main__":
    main()