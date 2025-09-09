#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI标书点对点应答系统 - Web界面
"""

import os
import logging
import json
import hashlib
import base64
import importlib
import requests
from datetime import datetime
from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
# from enhanced_inline_reply import EnhancedInlineReplyProcessor  # 已弃用，使用MCP方法
import mcp_bidder_name_processor_enhanced
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
import tempfile
import shutil
import sys
from pathlib import Path
from web_config import get_default_api_key, get_api_config, get_web_config

# 导入新的文档处理模块
try:
    from document_processor import DocumentProcessor, ProcessingOptions
    from table_processor import TableProcessor
    from image_inserter import SmartImageInserter
    DOCUMENT_PROCESSOR_AVAILABLE = True
    TABLE_PROCESSOR_AVAILABLE = True
    IMAGE_INSERTER_AVAILABLE = True
except ImportError as e:
    print(f"警告：文档处理模块未能加载: {e}")
    DOCUMENT_PROCESSOR_AVAILABLE = False
    TABLE_PROCESSOR_AVAILABLE = False
    IMAGE_INSERTER_AVAILABLE = False
    
# 分别检查各个模块
try:
    from table_processor import TableProcessor
    TABLE_PROCESSOR_AVAILABLE = True
except ImportError:
    TABLE_PROCESSOR_AVAILABLE = False
    
try:
    from image_inserter import SmartImageInserter
    IMAGE_INSERTER_AVAILABLE = True
except ImportError:
    IMAGE_INSERTER_AVAILABLE = False

# 导入招标信息提取模块
try:
    # 添加读取信息目录到sys.path
    info_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), '1.读取信息')
    if info_path not in sys.path:
        sys.path.append(info_path)
    from read_info import TenderInfoExtractor
    TENDER_INFO_AVAILABLE = True
    print("招标信息提取模块加载成功")
except ImportError as e:
    print(f"警告：招标信息提取模块未能加载: {e}")
    TENDER_INFO_AVAILABLE = False

# 招标信息提取功能配置
print(f"招标信息提取功能状态: {'可用' if TENDER_INFO_AVAILABLE else '不可用'}")

def extract_tender_info_direct(file_path, api_key):
    """直接调用招标信息提取功能"""
    if not TENDER_INFO_AVAILABLE:
        raise Exception("招标信息提取模块不可用")
    
    try:
        extractor = TenderInfoExtractor(api_key=api_key)
        result = extractor.process_document(file_path)
        return {'success': True, 'data': result}
    except Exception as e:
        logger.error(f"招标信息提取失败: {e}")
        raise Exception(f"提取失败: {str(e)}")

def extract_tender_info_step_direct(step, file_path, result_data, api_key):
    """直接调用招标信息提取的分步功能"""
    if not TENDER_INFO_AVAILABLE:
        raise Exception("招标信息提取模块不可用")
    
    try:
        if step == '1':
            # 第一步：提取基本信息
            extractor = TenderInfoExtractor(api_key=api_key)
            tender_info = extractor.process_document(file_path)
            basic_info = {
                'tenderer': tender_info.get('tenderer', ''),
                'agency': tender_info.get('agency', ''),  
                'bidding_method': tender_info.get('bidding_method', ''),
                'bidding_location': tender_info.get('bidding_location', ''),
                'bidding_time': tender_info.get('bidding_time', ''),
                'winner_count': tender_info.get('winner_count', ''),
                'project_name': tender_info.get('project_name', ''),
                'project_number': tender_info.get('project_number', '')
            }
            return {
                'success': True,
                'basic_info': basic_info,
                'full_data': tender_info
            }
        elif step == '2':
            # 第二步：获取资质要求
            qualification_requirements = result_data.get('qualification_requirements', {})
            if not qualification_requirements:
                qualification_requirements = {
                    'business_license': {'required': False, 'description': '未检测到相关要求'},
                    'taxpayer_qualification': {'required': False, 'description': '未检测到相关要求'},
                    'performance_requirements': {'required': False, 'description': '未检测到相关要求'},
                    'authorization_requirements': {'required': False, 'description': '未检测到相关要求'},
                    'credit_china': {'required': False, 'description': '未检测到相关要求'},
                    'commitment_letter': {'required': False, 'description': '未检测到相关要求'},
                    'audit_report': {'required': False, 'description': '未检测到相关要求'},
                    'social_security': {'required': False, 'description': '未检测到相关要求'},
                    'labor_contract': {'required': False, 'description': '未检测到相关要求'},
                    'other_requirements': {'required': False, 'description': '未检测到相关要求'}
                }
            return {
                'success': True,
                'qualification_requirements': qualification_requirements
            }
        elif step == '3':
            # 第三步：获取技术评分
            technical_scoring = result_data.get('technical_scoring', {})
            if not technical_scoring:
                technical_scoring = {
                    '未检测到评分项目': {
                        'score': '未提取到分值',
                        'criteria': ['未检测到具体评分标准，请检查文档中是否包含评分表格']
                    }
                }
            return {
                'success': True,
                'technical_scoring': technical_scoring
            }
        else:
            raise Exception(f'无效的步骤: {step}')
    except Exception as e:
        logger.error(f"分步提取失败: {e}")
        raise Exception(f"提取失败: {str(e)}")

# 添加技术方案模块路径
tech_proposal_path = str(Path(__file__).parent.parent / "技术方案" / "TenderGenerator")
sys.path.insert(0, tech_proposal_path)

try:
    from main import TenderGenerator
    TECH_PROPOSAL_AVAILABLE = True
except ImportError as e:
    print(f"警告：技术方案模块未能加载: {e}")
    TECH_PROPOSAL_AVAILABLE = False

# 计算web页面目录路径，统一管理所有web页面
web_pages_dir = str(Path(__file__).parent.parent.parent / "web页面")
# 计算公司配置目录路径
company_configs_dir = str(Path(__file__).parent.parent.parent / "company_configs")
app = Flask(__name__, template_folder=web_pages_dir, static_folder=web_pages_dir)
app.secret_key = 'ai_tender_response_system_2025'

def render_template_with_layout(page_template, page_title, active_nav):
    """使用布局模板渲染页面"""
    try:
        # 读取页面内容
        page_path = os.path.join(web_pages_dir, page_template)
        with open(page_path, 'r', encoding='utf-8') as f:
            page_content = f.read()
        
        # 读取布局模板
        layout_path = os.path.join(web_pages_dir, 'layout.html')
        with open(layout_path, 'r', encoding='utf-8') as f:
            layout_content = f.read()
        
        # 获取URL参数
        from urllib.parse import urlencode
        url_params = request.args.to_dict()
        url_params_str = '?' + urlencode(url_params) if url_params else ''
        
        # 确定页面专用脚本
        page_script = page_template.replace('.html', '.js')
        page_scripts = f'<script src="js/{page_script}"></script>'
        
        # 替换模板变量
        replacements = {
            '{{page_title}}': page_title,
            '{{page_content}}': page_content,
            '{{page_scripts}}': page_scripts,
            '{{url_params}}': url_params_str,
            '{{active_home}}': 'active' if active_nav == 'active_home' else '',
            '{{active_tender_info}}': 'active' if active_nav == 'active_tender_info' else '',
            '{{active_company_selection}}': 'active' if active_nav == 'active_company_selection' else '',
            '{{active_business_response}}': 'active' if active_nav == 'active_business_response' else '',
            '{{active_point_to_point}}': 'active' if active_nav == 'active_point_to_point' else '',
            '{{active_tech_proposal}}': 'active' if active_nav == 'active_tech_proposal' else '',
        }
        
        # 执行替换
        html_content = layout_content
        for placeholder, value in replacements.items():
            html_content = html_content.replace(placeholder, value)
            
        return html_content
        
    except Exception as e:
        logging.error(f"模板渲染失败: {e}")
        return f"<h1>页面加载失败</h1><p>错误: {str(e)}</p>", 500
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# 全局进度追踪
processing_status = {}

# 配置上传文件夹
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
TECH_OUTPUT_FOLDER = 'tech_outputs'
ALLOWED_EXTENSIONS = {'docx', 'doc'}
TECH_ALLOWED_EXTENSIONS = {'docx', 'doc', 'pdf'}
TENDER_INFO_ALLOWED_EXTENSIONS = {'docx', 'doc', 'txt', 'pdf'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(TECH_OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['TECH_OUTPUT_FOLDER'] = TECH_OUTPUT_FOLDER

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_tech_file(filename):
    """检查技术方案文件类型是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in TECH_ALLOWED_EXTENSIONS

def allowed_tender_info_file(filename):
    """检查招标信息提取文件类型是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in TENDER_INFO_ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """主页"""
    return render_template_with_layout('index.html', '首页', 'active_home')

@app.route('/help.html')
def help_page():
    """帮助页面"""
    return render_template('help.html')

@app.route('/index.html')
def index_page():
    """首页"""
    return render_template_with_layout('index.html', '首页', 'active_home')

@app.route('/tender_info.html')
def tender_info():
    """招标信息提取页面"""
    return render_template_with_layout('tender_info.html', '读取信息', 'active_tender_info')

@app.route('/company_selection.html')
def company_selection():
    """公司选择页面"""
    return render_template_with_layout('company_selection.html', '选择应答人', 'active_company_selection')

@app.route('/business_response.html')
def business_response():
    """商务应答页面"""
    return render_template_with_layout('business_response.html', '商务应答', 'active_business_response')

@app.route('/point_to_point.html')
def point_to_point():
    """点对点应答页面"""
    return render_template_with_layout('point_to_point.html', '点对点应答', 'active_point_to_point')

@app.route('/tech_proposal.html')
def tech_proposal():
    """技术方案页面"""
    return render_template_with_layout('tech_proposal.html', '技术方案', 'active_tech_proposal')

@app.route('/preview-document')
def preview_document():
    """文档预览页面"""
    return render_template_with_layout('document_preview.html', '文档预览', '')

@app.route('/document_preview.html')
def document_preview():
    """文档预览页面（直接访问）"""
    return render_template_with_layout('document_preview.html', '文档预览', '')

@app.route('/debug')
def debug_upload():
    """调试上传页面"""
    return send_file('test_upload_debug.html')

@app.route('/<path:filename>')
def static_files(filename):
    """处理静态文件"""
    try:
        static_path = os.path.join(web_pages_dir, filename)
        if os.path.exists(static_path) and os.path.isfile(static_path):
            return send_file(static_path)
        else:
            return "File not found", 404
    except Exception as e:
        return f"Error serving file: {e}", 500

@app.route('/upload', methods=['POST'])
def upload_file():
    """处理文件上传"""
    upload_path = None
    try:
        logger.info("开始处理文件上传请求")
        
        if 'file' not in request.files:
            logger.warning("请求中未包含文件")
            return jsonify({'error': '没有选择文件'}), 400
        
        file = request.files['file']
        api_key = request.form.get('api_key', '').strip()
        
        if file.filename == '':
            logger.warning("文件名为空")
            return jsonify({'error': '没有选择文件'}), 400
        
        if not allowed_file(file.filename):
            logger.warning(f"不支持的文件类型: {file.filename}")
            return jsonify({'error': '不支持的文件类型，请上传.docx或.doc文件'}), 400
            
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        
        # 保存上传的文件
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        logger.info(f"保存文件到: {upload_path}")
        file.save(upload_path)
        
        # 检查文件大小
        file_size = os.path.getsize(upload_path)
        logger.info(f"文件大小: {file_size / 1024 / 1024:.2f} MB")
        
        if file_size > 50 * 1024 * 1024:  # 50MB
            os.remove(upload_path)
            return jsonify({'error': '文件过大，请上传小于50MB的文件'}), 400
        
        # 处理文档 - 统一使用MCP处理器
        logger.info("开始初始化MCP文档处理器")
        
        # 使用默认公司名称（智慧足迹数据科技有限公司）
        default_company_name = "智慧足迹数据科技有限公司"
        
        try:
            # 强制重载模块以获取最新修改
            importlib.reload(mcp_bidder_name_processor_enhanced)
            from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
            processor = MCPBidderNameProcessor()
            logger.info("初始化增强版MCP投标人名称处理器 - 已重载模块")
            
            # 生成输出文件名
            base_name = os.path.splitext(filename)[0]
            output_filename = f"{base_name}-AI应答-{timestamp}.docx"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            
            logger.info(f"开始处理文档: {filename}")
            
            # 执行MCP处理
            result = processor.process_bidder_name(
                input_file=upload_path,
                output_file=output_path,
                company_name=default_company_name
            )
            
            if not result.get('success'):
                raise Exception(result.get('error', 'MCP处理失败'))
                
            result_file = output_path
            
        except ImportError as e:
            logger.error(f"MCP处理器加载失败: {e}")
            return jsonify({'error': 'MCP处理器不可用，请检查系统配置'}), 500
        
        logger.info(f"文档处理完成: {result_file}")
        
        # 清理上传的临时文件
        if upload_path and os.path.exists(upload_path):
            os.remove(upload_path)
            logger.info("临时文件已清理")
        
        return jsonify({
            'success': True,
            'message': '文档处理完成',
            'filename': output_filename,
            'download_url': url_for('download_file', filename=output_filename)
        })
            
    except Exception as e:
        logger.error(f"文件处理失败: {e}", exc_info=True)
        
        # 清理临时文件
        if upload_path and os.path.exists(upload_path):
            try:
                os.remove(upload_path)
                logger.info("清理临时文件成功")
            except Exception as cleanup_error:
                logger.error(f"清理临时文件失败: {cleanup_error}")
        
        # 返回详细错误信息
        error_message = str(e)
        if "python-docx" in error_message:
            error_message = "文档格式错误，请确保上传的是有效的Word文档"
        elif "Memory" in error_message or "内存" in error_message:
            error_message = "文档过大导致内存不足，请尝试上传更小的文件"
        elif "API" in error_message:
            error_message = f"AI服务调用失败: {error_message}"
        
        return jsonify({'error': f'处理失败: {error_message}'}), 500

@app.route('/status/<task_id>')
def get_status(task_id):
    """获取处理状态"""
    status = processing_status.get(task_id, {
        'status': 'unknown',
        'message': '未知任务',
        'progress': 0
    })
    return jsonify(status)

@app.route('/download/<filename>')
def download_file(filename):
    """下载处理后的文件"""
    try:
        # 安全性检查：防止路径遍历攻击
        if '..' in filename or '/' in filename or '\\' in filename:
            logger.warning(f"不安全的文件名: {filename}")
            return "无效的文件名", 400
            
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        logger.info(f"尝试下载文件: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return "文件不存在", 404
            
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        logger.info(f"下载文件大小: {file_size / 1024 / 1024:.2f} MB")
        
        # 使用绝对路径和正确的mimetype
        return send_file(
            os.path.abspath(file_path), 
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"文件下载失败: {e}", exc_info=True)
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

@app.route('/outputs/images/<filename>')
def serve_editor_images(filename):
    """服务编辑器上传的图片文件"""
    try:
        # 安全性检查：防止路径遍历攻击
        if '..' in filename or '/' in filename or '\\' in filename:
            logger.warning(f"不安全的图片文件名: {filename}")
            return "无效的文件名", 400
            
        # 构建图片文件路径
        images_dir = os.path.join('outputs', 'images')
        image_path = os.path.join(images_dir, filename)
        
        if not os.path.exists(image_path):
            logger.error(f"图片文件不存在: {image_path}")
            return "文件不存在", 404
            
        # 根据文件扩展名设置正确的MIME类型
        file_ext = filename.lower().split('.')[-1]
        mime_types = {
            'png': 'image/png',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'gif': 'image/gif',
            'bmp': 'image/bmp',
            'webp': 'image/webp'
        }
        mimetype = mime_types.get(file_ext, 'application/octet-stream')
            
        return send_file(os.path.abspath(image_path), mimetype=mimetype)
        
    except Exception as e:
        logger.error(f"编辑器图片文件访问失败: {e}", exc_info=True)
        return jsonify({'error': f'访问失败: {str(e)}'}), 500

@app.route('/web页面/<path:filename>')
def serve_web_pages(filename):
    """服务web页面文件夹中的静态文件"""
    try:
        # 安全性检查：防止路径遍历攻击
        if '..' in filename:
            logger.warning(f"不安全的文件路径: {filename}")
            return "无效的文件路径", 400
            
        # 构建文件路径 - web页面文件夹位于程序目录下
        web_pages_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'web页面')
        file_path = os.path.join(web_pages_dir, filename)
        
        logger.info(f"尝试访问web页面文件: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"web页面文件不存在: {file_path}")
            return "文件不存在", 404
            
        # 根据文件扩展名设置正确的MIME类型
        if filename.endswith('.html'):
            mimetype = 'text/html'
        elif filename.endswith('.css'):
            mimetype = 'text/css'
        elif filename.endswith('.js'):
            mimetype = 'application/javascript'
        else:
            mimetype = 'application/octet-stream'
            
        return send_file(os.path.abspath(file_path), mimetype=mimetype)
        
    except Exception as e:
        logger.error(f"web页面文件访问失败: {e}", exc_info=True)
        return jsonify({'error': f'访问失败: {str(e)}'}), 500

@app.route('/api/files')
def list_files():
    """列出可下载的文件"""
    try:
        output_dir = app.config['OUTPUT_FOLDER']
        if not os.path.exists(output_dir):
            return jsonify({'files': []})
            
        files = []
        for filename in os.listdir(output_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(output_dir, filename)
                stat = os.stat(file_path)
                files.append({
                    'name': filename,
                    'size': stat.st_size,
                    'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    'download_url': url_for('download_file', filename=filename)
                })
        
        files.sort(key=lambda x: x['created'], reverse=True)
        return jsonify({'files': files})
        
    except Exception as e:
        logger.error(f"列出文件失败: {e}")
        return jsonify({'error': f'获取文件列表失败: {str(e)}'}), 500

@app.route('/api/business-files')
def list_business_files():
    """列出商务应答文件"""
    try:
        output_dir = app.config['OUTPUT_FOLDER']
        if not os.path.exists(output_dir):
            return jsonify({'files': []})
            
        files = []
        for filename in os.listdir(output_dir):
            # 只显示商务应答相关文件
            if filename.endswith('.docx') and ('商务应答' in filename or 'business' in filename.lower()):
                file_path = os.path.join(output_dir, filename)
                stat = os.stat(file_path)
                files.append({
                    'name': filename,
                    'size': stat.st_size,
                    'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    'download_url': url_for('download_file', filename=filename)
                })
        
        files.sort(key=lambda x: x['created'], reverse=True)
        return jsonify({'files': files})
        
    except Exception as e:
        logger.error(f"列出商务应答文件失败: {e}")
        return jsonify({'error': f'获取商务应答文件列表失败: {str(e)}'}), 500

@app.route('/api/test', methods=['POST'])
def test_api():
    """测试API连接"""
    try:
        api_key = request.json.get('api_key', '').strip()
        if not api_key:
            return jsonify({'error': 'API密钥不能为空'}), 400
            
        # MCP处理器不需要API连接测试，直接返回成功
        try:
            # 强制重载模块以获取最新修改
            importlib.reload(mcp_bidder_name_processor_enhanced)
            from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
            processor = MCPBidderNameProcessor()
            return jsonify({
                'success': True,
                'message': 'MCP处理器初始化成功！使用本地处理模式',
                'model': 'MCP Local Processing'
            })
        except ImportError as e:
            return jsonify({
                'success': False,
                'message': f'MCP处理器初始化失败: {str(e)}'
            })
            
    except Exception as e:
        logger.error(f"API测试失败: {e}")
        return jsonify({
            'success': False,
            'message': f'测试失败: {str(e)}'
        }), 500

@app.route('/api/save-key', methods=['POST'])
def save_api_key():
    """安全保存API密钥到服务器端"""
    try:
        api_key = request.json.get('api_key', '').strip()
        if not api_key:
            return jsonify({'error': 'API密钥不能为空'}), 400
        
        # 创建配置目录
        config_dir = 'config'
        os.makedirs(config_dir, exist_ok=True)
        
        # 使用简单加密存储（仅用于本地备份）
        key_hash = hashlib.md5(api_key.encode()).hexdigest()[:8]
        encrypted_key = base64.b64encode(api_key.encode()).decode()
        
        config_file = os.path.join(config_dir, f'api_backup_{key_hash}.json')
        
        backup_data = {
            'encrypted_key': encrypted_key,
            'created_at': datetime.now().isoformat(),
            'key_prefix': api_key[:10] + '...',
            'description': '始皇API密钥备份'
        }
        
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(backup_data, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': f'API密钥已安全备份 ({key_hash})',
            'backup_id': key_hash
        })
        
    except Exception as e:
        logger.error(f"API密钥保存失败: {e}")
        return jsonify({'error': f'保存失败: {str(e)}'}), 500

@app.route('/api/get-default-api-key', methods=['GET'])
def get_api_key():
    """获取默认API密钥"""
    try:
        default_api_key = get_default_api_key()
        if default_api_key:
            return jsonify({'api_key': default_api_key})
        else:
            return jsonify({'error': '未配置默认API密钥'}), 404
    except Exception as e:
        logger.error(f"获取默认API密钥失败: {e}")
        return jsonify({'error': f'获取API密钥失败: {str(e)}'}), 500

@app.route('/api/load-backups', methods=['GET'])
def load_api_backups():
    """加载API密钥备份列表"""
    try:
        config_dir = 'config'
        if not os.path.exists(config_dir):
            return jsonify({'backups': []})
        
        backups = []
        for filename in os.listdir(config_dir):
            if filename.startswith('api_backup_') and filename.endswith('.json'):
                try:
                    with open(os.path.join(config_dir, filename), 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        backups.append({
                            'id': filename.replace('api_backup_', '').replace('.json', ''),
                            'prefix': data.get('key_prefix', ''),
                            'created_at': data.get('created_at', ''),
                            'description': data.get('description', '')
                        })
                except Exception as e:
                    logger.warning(f"读取备份文件失败 {filename}: {e}")
        
        backups.sort(key=lambda x: x['created_at'], reverse=True)
        return jsonify({'backups': backups})
        
    except Exception as e:
        logger.error(f"加载备份列表失败: {e}")
        return jsonify({'error': f'加载失败: {str(e)}'}), 500

@app.route('/api/restore-key/<backup_id>', methods=['POST'])
def restore_api_key(backup_id):
    """恢复API密钥"""
    try:
        config_file = os.path.join('config', f'api_backup_{backup_id}.json')
        if not os.path.exists(config_file):
            return jsonify({'error': '备份文件不存在'}), 404
        
        with open(config_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        encrypted_key = data.get('encrypted_key', '')
        if not encrypted_key:
            return jsonify({'error': '备份数据无效'}), 400
        
        # 解密API密钥
        api_key = base64.b64decode(encrypted_key).decode()
        
        return jsonify({
            'success': True,
            'api_key': api_key,
            'message': 'API密钥恢复成功'
        })
        
    except Exception as e:
        logger.error(f"API密钥恢复失败: {e}")
        return jsonify({'error': f'恢复失败: {str(e)}'}), 500

@app.route('/extract-tender-info', methods=['POST'])
def extract_tender_info():
    """提取招标信息 - 通过调用独立服务"""
    try:
        logger.info("开始处理招标信息提取请求")
        logger.info(f"请求方法: {request.method}")
        logger.info(f"请求文件: {list(request.files.keys())}")
        logger.info(f"请求表单: {list(request.form.keys())}")
        
        if not TENDER_INFO_AVAILABLE:
            logger.error("招标信息提取服务不可用")
            return jsonify({'error': '招标信息提取服务不可用，请检查系统配置'}), 500
        
        # 检查文件
        if 'file' not in request.files:
            logger.error("请求中未包含文件")
            return jsonify({'error': '需要上传招标文档'}), 400
        
        file = request.files['file']
        api_key = request.form.get('api_key', '').strip()
        
        logger.info(f"上传文件名: {file.filename}")
        logger.info(f"API密钥长度: {len(api_key) if api_key else 0}")
        
        if file.filename == '':
            return jsonify({'error': '请选择招标文档'}), 400
        
        if not allowed_tender_info_file(file.filename):
            return jsonify({'error': '不支持的文件类型，请上传.docx、.doc、.txt或.pdf文件'}), 400
        
        # 保存文件并直接调用招标信息提取功能
        logger.info("直接调用招标信息提取功能")
        
        # 保存临时文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = secure_filename(file.filename)
        unique_filename = f"{timestamp}_tender_{filename}"
        upload_path = os.path.join('uploads', unique_filename)
        
        # 确保uploads目录存在
        os.makedirs('uploads', exist_ok=True)
        file.save(upload_path)
        
        try:
            result = extract_tender_info_direct(upload_path, api_key)
            
            if result.get('success'):
                logger.info("招标信息提取成功")
                return jsonify({
                    'success': True,
                    'message': '招标信息提取完成',
                    'tender_info': result.get('data', {})
                })
            else:
                logger.error(f"招标信息提取失败: {result}")
                return jsonify({'error': '提取失败'}), 500
        finally:
            # 清理临时文件
            if os.path.exists(upload_path):
                try:
                    os.remove(upload_path)
                    logger.info(f"清理临时文件: {upload_path}")
                except Exception as cleanup_error:
                    logger.error(f"清理临时文件失败: {cleanup_error}")
        
    except Exception as e:
        logger.error(f"招标信息提取失败: {e}", exc_info=True)
        
        error_message = str(e)
        if "连接" in error_message or "服务不可用" in error_message:
            error_message = "招标信息提取服务不可用，请确保读取信息服务正在运行"
        elif "API" in error_message or "OpenAI" in error_message:
            error_message = f"AI服务调用失败: {error_message}"
        elif "Memory" in error_message or "内存" in error_message:
            error_message = "文档过大导致内存不足，请尝试上传更小的文件"
        elif "No such file" in error_message:
            error_message = "文件读取失败，请确保上传的是有效的文档文件"
        
        return jsonify({'error': f'提取失败: {error_message}'}), 500

@app.route('/extract-tender-info-step', methods=['POST'])
def extract_tender_info_step():
    """分步提取招标信息 - 通过调用独立服务"""
    try:
        step = request.form.get('step', '1')
        api_key = request.form.get('api_key', '').strip()
        
        logger.info(f"分步提取招标信息 - 步骤: {step}")
        
        # 检查招标信息提取服务是否可用
        if not TENDER_INFO_AVAILABLE:
            logger.error("招标信息提取服务不可用")
            return jsonify({'error': '招标信息提取服务不可用，请检查系统配置'}), 500
        
        # 如果没有提供API密钥，使用默认密钥
        if not api_key:
            api_key = get_default_api_key()
            logger.info("使用默认API密钥")
        
        # 处理分步提取逻辑
        file_path = None
        result_data = {}
        
        if step == '1':
            # 第一步：处理文件上传
            if 'file' not in request.files:
                return jsonify({'error': '需要上传招标文档'}), 400
                
            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': '请选择招标文档'}), 400
                
            if not allowed_tender_info_file(file.filename):
                return jsonify({'error': '不支持的文件类型，请上传.docx、.doc、.txt或.pdf文件'}), 400
            
            # 保存文件
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = secure_filename(file.filename)
            unique_filename = f"{timestamp}_tender_{filename}"
            file_path = os.path.join('uploads', unique_filename)
            
            # 确保uploads目录存在
            os.makedirs('uploads', exist_ok=True)
            file.save(file_path)
            
        else:
            # 后续步骤使用之前保存的文件路径和结果数据
            file_path = request.form.get('file_path')
            result_file = request.form.get('result_file')
            
            # 尝试从临时结果文件读取
            if result_file and os.path.exists(result_file):
                try:
                    with open(result_file, 'r', encoding='utf-8') as f:
                        result_data = json.load(f)
                except Exception as e:
                    logger.error(f"读取临时结果文件失败: {e}")
        
        # 直接调用招标信息提取功能
        logger.info(f"直接调用招标信息提取功能 - 步骤 {step}")
        result = extract_tender_info_step_direct(step, file_path, result_data, api_key)
        
        # 处理步骤1的特殊逻辑，保存完整数据到临时文件
        if step == '1' and result.get('success') and 'full_data' in result:
            temp_result_file = file_path + '.result.json'
            with open(temp_result_file, 'w', encoding='utf-8') as f:
                json.dump(result['full_data'], f, ensure_ascii=False, indent=2)
            
            # 返回结果时包含文件路径信息
            result['file_path'] = file_path
            result['result_file'] = temp_result_file
            
            # 清理full_data，避免返回过大的数据
            result.pop('full_data', None)
        
        if result.get('success'):
            logger.info(f"步骤 {step} 提取成功")
            return jsonify(result)
        else:
            logger.error(f"步骤 {step} 提取失败: {result}")
            return jsonify({'error': '提取失败'}), 500
            
    except Exception as e:
        logger.error(f"分步提取失败: {e}", exc_info=True)
        
        error_message = str(e)
        if "连接" in error_message or "服务不可用" in error_message:
            error_message = "招标信息提取服务不可用，请确保读取信息服务正在运行"
        
        return jsonify({'error': f'提取失败: {error_message}'}), 500

@app.route('/health')
def health_check():
    """健康检查"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'AI标书智能生成系统',
        'tech_proposal_available': TECH_PROPOSAL_AVAILABLE,
        'tender_info_available': TENDER_INFO_AVAILABLE
    })

@app.route('/generate-proposal', methods=['POST'])
def generate_proposal():
    """生成技术方案"""
    tender_upload_path = None
    product_upload_path = None
    
    try:
        logger.info("开始处理技术方案生成请求")
        
        if not TECH_PROPOSAL_AVAILABLE:
            return jsonify({'error': '技术方案模块未加载，请检查系统配置'}), 500
        
        # 检查文件
        if 'tender_file' not in request.files or 'product_file' not in request.files:
            return jsonify({'error': '需要上传招标文件和产品文档'}), 400
        
        tender_file = request.files['tender_file']
        product_file = request.files['product_file']
        api_key = request.form.get('api_key', '').strip()
        output_prefix = request.form.get('output_prefix', '技术方案').strip()
        
        if tender_file.filename == '' or product_file.filename == '':
            return jsonify({'error': '请选择招标文件和产品文档'}), 400
        
        if not (allowed_tech_file(tender_file.filename) and allowed_tech_file(product_file.filename)):
            return jsonify({'error': '不支持的文件类型，请上传.docx、.doc或.pdf文件'}), 400
        
        # 保存上传的文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        tender_filename = secure_filename(tender_file.filename)
        tender_unique_filename = f"{timestamp}_tender_{tender_filename}"
        tender_upload_path = os.path.join(app.config['UPLOAD_FOLDER'], tender_unique_filename)
        tender_file.save(tender_upload_path)
        
        product_filename = secure_filename(product_file.filename)
        product_unique_filename = f"{timestamp}_product_{product_filename}"
        product_upload_path = os.path.join(app.config['UPLOAD_FOLDER'], product_unique_filename)
        product_file.save(product_upload_path)
        
        logger.info(f"文件上传完成: 招标文件={tender_upload_path}, 产品文档={product_upload_path}")
        
        # 初始化技术方案生成器
        logger.info("初始化技术方案生成器")
        
        # 临时修改环境变量以传递API密钥
        old_api_key = os.environ.get('OPENAI_API_KEY')
        if api_key:
            os.environ['OPENAI_API_KEY'] = api_key
        
        try:
            generator = TenderGenerator()
            
            # 生成输出文件前缀
            safe_prefix = f"{output_prefix}_{timestamp}"
            
            logger.info(f"开始生成技术方案: {safe_prefix}")
            
            # 生成技术方案
            result = generator.generate_proposal(
                tender_file=tender_upload_path,
                product_file=product_upload_path,
                output_prefix=safe_prefix
            )
            
            if result['success']:
                # 移动输出文件到技术方案输出目录
                moved_files = {}
                for file_type, file_path in result['output_files'].items():
                    if os.path.exists(file_path):
                        filename = os.path.basename(file_path)
                        new_path = os.path.join(app.config['TECH_OUTPUT_FOLDER'], filename)
                        shutil.move(file_path, new_path)
                        moved_files[file_type] = new_path
                        logger.info(f"移动文件: {file_path} -> {new_path}")
                
                result['output_files'] = moved_files
                logger.info(f"技术方案生成成功: {result}")
                
                return jsonify(result)
            else:
                logger.error(f"技术方案生成失败: {result['error']}")
                return jsonify(result), 500
                
        finally:
            # 恢复环境变量
            if old_api_key:
                os.environ['OPENAI_API_KEY'] = old_api_key
            elif 'OPENAI_API_KEY' in os.environ:
                del os.environ['OPENAI_API_KEY']
        
    except Exception as e:
        logger.error(f"技术方案生成失败: {e}", exc_info=True)
        
        error_message = str(e)
        if "API" in error_message or "OpenAI" in error_message:
            error_message = f"AI服务调用失败: {error_message}"
        elif "Memory" in error_message or "内存" in error_message:
            error_message = "文档过大导致内存不足，请尝试上传更小的文件"
        elif "No such file" in error_message:
            error_message = "文件读取失败，请确保上传的是有效的文档文件"
        
        return jsonify({'error': f'生成失败: {error_message}'}), 500
    
    finally:
        # 清理临时文件
        for path in [tender_upload_path, product_upload_path]:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                    logger.info(f"清理临时文件: {path}")
                except Exception as cleanup_error:
                    logger.error(f"清理临时文件失败: {cleanup_error}")

@app.route('/download-tech/<filename>')
def download_tech_file(filename):
    """下载技术方案文件"""
    try:
        # 安全性检查：防止路径遍历攻击
        if '..' in filename or '/' in filename or '\\' in filename:
            logger.warning(f"不安全的文件名: {filename}")
            return "无效的文件名", 400
            
        file_path = os.path.join(app.config['TECH_OUTPUT_FOLDER'], filename)
        logger.info(f"尝试下载技术方案文件: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return "文件不存在", 404
            
        # 根据文件扩展名设置正确的mimetype
        file_ext = filename.lower().split('.')[-1]
        if file_ext == 'docx':
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_ext == 'json':
            mimetype = 'application/json'
        elif file_ext == 'txt':
            mimetype = 'text/plain'
        else:
            mimetype = 'application/octet-stream'
        
        return send_file(
            os.path.abspath(file_path), 
            as_attachment=True,
            download_name=filename,
            mimetype=mimetype
        )
        
    except Exception as e:
        logger.error(f"技术方案文件下载失败: {e}", exc_info=True)
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

@app.route('/api/tech-files')
def list_tech_files():
    """列出可下载的技术方案文件"""
    try:
        output_dir = app.config['TECH_OUTPUT_FOLDER']
        if not os.path.exists(output_dir):
            return jsonify({'files': []})
            
        files = []
        for filename in os.listdir(output_dir):
            if filename.endswith(('.docx', '.json', '.txt')):
                file_path = os.path.join(output_dir, filename)
                stat = os.stat(file_path)
                files.append({
                    'name': filename,
                    'size': stat.st_size,
                    'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    'download_url': url_for('download_tech_file', filename=filename)
                })
        
        files.sort(key=lambda x: x['created'], reverse=True)
        return jsonify({'files': files})
        
    except Exception as e:
        logger.error(f"列出技术方案文件失败: {e}")
        return jsonify({'error': f'获取文件列表失败: {str(e)}'}), 500

@app.route('/api/companies')
def list_companies():
    """获取所有公司配置"""
    try:
        companies = []
        
        if os.path.exists(company_configs_dir):
            for filename in os.listdir(company_configs_dir):
                if filename.endswith('.json'):
                    file_path = os.path.join(company_configs_dir, filename)
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            company_data = json.load(f)
                            companies.append({
                                'id': company_data.get('id', filename.replace('.json', '')),
                                'companyName': company_data.get('companyName', '未命名公司'),
                                'created_at': company_data.get('created_at', ''),
                                'updated_at': company_data.get('updated_at', '')
                            })
                    except Exception as e:
                        logger.warning(f"读取公司配置文件失败 {filename}: {e}")
        
        companies.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
        return jsonify({'success': True, 'companies': companies})
        
    except Exception as e:
        logger.error(f"获取公司列表失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/companies/<company_id>')
def get_company(company_id):
    """获取指定公司的详细信息"""
    try:
        company_file = os.path.join(company_configs_dir, f'{company_id}.json')
        
        if not os.path.exists(company_file):
            return jsonify({'success': False, 'error': '公司不存在'}), 404
            
        with open(company_file, 'r', encoding='utf-8') as f:
            company_data = json.load(f)
            
        return jsonify({'success': True, 'company': company_data})
        
    except Exception as e:
        logger.error(f"获取公司信息失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/companies', methods=['POST'])
def create_company():
    """创建新公司"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': '请提供公司信息'}), 400
        
        company_name = data.get('companyName', '').strip()
        if not company_name:
            return jsonify({'success': False, 'error': '公司名称不能为空'}), 400
        
        # 生成唯一ID
        import uuid
        company_id = str(uuid.uuid4())
        
        # 准备公司数据
        company_data = {
            'id': company_id,
            'companyName': company_name,
            'establishDate': data.get('establishDate', ''),
            'legalRepresentative': data.get('legalRepresentative', ''),
            'legalRepresentativePosition': data.get('legalRepresentativePosition', ''),
            'authorizedPersonName': data.get('authorizedPersonName', ''),
            'authorizedPersonPosition': data.get('authorizedPersonPosition', ''),
            'socialCreditCode': data.get('socialCreditCode', ''),
            'registeredCapital': data.get('registeredCapital', ''),
            'companyType': data.get('companyType', ''),
            'fixedPhone': data.get('fixedPhone', ''),
            'email': data.get('email', ''),
            'postalCode': data.get('postalCode', ''),
            'registeredAddress': data.get('registeredAddress', ''),
            'officeAddress': data.get('officeAddress', ''),
            'website': data.get('website', ''),
            'employeeCount': data.get('employeeCount', ''),
            'companyDescription': data.get('companyDescription', ''),
            'businessScope': data.get('businessScope', ''),
            'bankName': data.get('bankName', ''),
            'bankAccount': data.get('bankAccount', ''),
            'fax': data.get('fax', ''),
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat()
        }
        
        # 确保目录存在
        os.makedirs(company_configs_dir, exist_ok=True)
        
        # 保存公司信息
        company_file = os.path.join(company_configs_dir, f'{company_id}.json')
        with open(company_file, 'w', encoding='utf-8') as f:
            json.dump(company_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"创建新公司: {company_name} ({company_id})")
        return jsonify({
            'success': True,
            'company_id': company_id,
            'message': '公司信息保存成功'
        })
        
    except Exception as e:
        logger.error(f"创建公司失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/companies/<company_id>', methods=['PUT'])
def update_company(company_id):
    """更新公司信息"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': '请提供公司信息'}), 400
        
        company_file = os.path.join(company_configs_dir, f'{company_id}.json')
        
        # 如果公司不存在，返回错误
        if not os.path.exists(company_file):
            return jsonify({'success': False, 'error': '公司不存在'}), 404
        
        # 读取现有数据
        with open(company_file, 'r', encoding='utf-8') as f:
            company_data = json.load(f)
        
        # 更新数据
        company_data.update({
            'companyName': data.get('companyName', company_data.get('companyName', '')),
            'establishDate': data.get('establishDate', company_data.get('establishDate', '')),
            'legalRepresentative': data.get('legalRepresentative', company_data.get('legalRepresentative', '')),
            'legalRepresentativePosition': data.get('legalRepresentativePosition', company_data.get('legalRepresentativePosition', '')),
            'authorizedPersonName': data.get('authorizedPersonName', company_data.get('authorizedPersonName', '')),
            'authorizedPersonPosition': data.get('authorizedPersonPosition', company_data.get('authorizedPersonPosition', '')),
            'socialCreditCode': data.get('socialCreditCode', company_data.get('socialCreditCode', '')),
            'registeredCapital': data.get('registeredCapital', company_data.get('registeredCapital', '')),
            'companyType': data.get('companyType', company_data.get('companyType', '')),
            'fixedPhone': data.get('fixedPhone', company_data.get('fixedPhone', '')),
            'email': data.get('email', company_data.get('email', '')),
            'postalCode': data.get('postalCode', company_data.get('postalCode', '')),
            'registeredAddress': data.get('registeredAddress', company_data.get('registeredAddress', '')),
            'officeAddress': data.get('officeAddress', company_data.get('officeAddress', '')),
            'website': data.get('website', company_data.get('website', '')),
            'employeeCount': data.get('employeeCount', company_data.get('employeeCount', '')),
            'companyDescription': data.get('companyDescription', company_data.get('companyDescription', '')),
            'businessScope': data.get('businessScope', company_data.get('businessScope', '')),
            'bankName': data.get('bankName', company_data.get('bankName', '')),
            'bankAccount': data.get('bankAccount', company_data.get('bankAccount', '')),
            'fax': data.get('fax', company_data.get('fax', '')),
            'updated_at': datetime.now().isoformat()
        })
        
        # 保存更新的数据
        with open(company_file, 'w', encoding='utf-8') as f:
            json.dump(company_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"更新公司信息: {company_data.get('companyName', '')} ({company_id})")
        return jsonify({
            'success': True,
            'message': '公司信息更新成功'
        })
        
    except Exception as e:
        logger.error(f"更新公司信息失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/companies/<company_id>/qualifications')
def get_company_qualifications(company_id):
    """获取指定公司的资质信息"""
    try:
        company_file = os.path.join(company_configs_dir, f'{company_id}.json')
        
        if not os.path.exists(company_file):
            return jsonify({'success': False, 'error': '公司不存在'}), 404
            
        with open(company_file, 'r', encoding='utf-8') as f:
            company_data = json.load(f)
        
        # 提取资质信息
        qualifications = company_data.get('qualifications', {})
        
        return jsonify({'success': True, 'qualifications': qualifications})
        
    except Exception as e:
        logger.error(f"获取公司资质信息失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/companies/<company_id>/qualifications/<key>/download')
def download_qualification_file(company_id, key):
    """下载指定公司的资质文件"""
    try:
        # 读取公司配置，获取文件信息
        company_file = os.path.join(company_configs_dir, f'{company_id}.json')
        
        if not os.path.exists(company_file):
            return jsonify({'success': False, 'error': '公司不存在'}), 404
            
        with open(company_file, 'r', encoding='utf-8') as f:
            company_data = json.load(f)
        
        # 获取资质文件信息
        qualifications = company_data.get('qualifications', {})
        if key not in qualifications:
            return jsonify({'success': False, 'error': '资质文件不存在'}), 404
        
        file_info = qualifications[key]
        safe_filename = file_info.get('safe_filename')
        
        if not safe_filename:
            return jsonify({'success': False, 'error': '文件信息不完整'}), 400
        
        # 构建文件路径
        qualifications_dir = os.path.join(os.path.dirname(__file__), 'qualifications', company_id)
        file_path = os.path.join(qualifications_dir, safe_filename)
        
        if not os.path.exists(file_path):
            logger.error(f"资质文件不存在: {file_path}")
            return jsonify({'success': False, 'error': '文件不存在'}), 404
        
        # 使用原始文件名作为下载文件名
        original_filename = file_info.get('original_filename', safe_filename)
        
        return send_file(
            file_path,
            as_attachment=False,  # 设置为False以便在浏览器中预览
            download_name=original_filename
        )
        
    except Exception as e:
        logger.error(f"下载资质文件失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/project-config')
def get_project_config():
    """获取项目配置信息"""
    try:
        import configparser
        
        # 只读取"读取信息"模块生成的配置文件
        config_file = os.path.join(tender_info_path, 'tender_config.ini')
        
        if not os.path.exists(config_file):
            return jsonify({'success': False, 'error': '项目配置文件不存在'})
            
        config = configparser.ConfigParser(interpolation=None)
        config.read(config_file, encoding='utf-8')
        
        # 提取项目信息
        project_info = {}
        if config.has_section('PROJECT_INFO'):
            project_info = {
                'projectName': config.get('PROJECT_INFO', 'project_name', fallback=''),
                'projectNumber': config.get('PROJECT_INFO', 'project_number', fallback=''),
                'tenderer': config.get('PROJECT_INFO', 'tenderer', fallback=''),
                'agency': config.get('PROJECT_INFO', 'agency', fallback=''),
                'biddingMethod': config.get('PROJECT_INFO', 'bidding_method', fallback=''),
                'biddingLocation': config.get('PROJECT_INFO', 'bidding_location', fallback=''),
                'biddingTime': config.get('PROJECT_INFO', 'bidding_time', fallback=''),
                'winnerCount': config.get('PROJECT_INFO', 'winner_count', fallback=''),
                'extractionTime': config.get('PROJECT_INFO', 'extraction_time', fallback='')
            }
        
        # 生成当前日期
        current_date = datetime.now().strftime('%Y年%m月%d日')
        project_info['currentDate'] = current_date
        
        # 添加配置文件来源信息
        project_info['configSource'] = '读取信息模块'
        
        logger.info(f"读取项目配置成功: {config_file}")
        
        return jsonify({
            'success': True, 
            'project_info': project_info
        })
        
    except Exception as e:
        logger.error(f"获取项目配置失败: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/editor/load-document', methods=['POST'])
def load_document():
    """读取Word文档转换为HTML"""
    try:
        logger.info("开始处理Word文档读取请求")
        
        if 'file' not in request.files:
            return jsonify({'error': '未上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400
        
        # 检查文件类型
        allowed_extensions = {'docx', 'doc'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_ext not in allowed_extensions:
            return jsonify({'error': '只支持 .docx 和 .doc 格式的Word文档'}), 400
        
        # 保存临时文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_filename = f"temp_editor_{timestamp}_{secure_filename(file.filename)}"
        temp_path = os.path.join('uploads', temp_filename)
        
        os.makedirs('uploads', exist_ok=True)
        file.save(temp_path)
        logger.info(f"临时文件已保存: {temp_path}")
        
        try:
            # 使用mammoth转换Word为HTML
            import mammoth
            
            with open(temp_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                
                # 获取转换警告（如果有）
                warnings = [str(warning) for warning in result.messages]
                if warnings:
                    logger.warning(f"Word转换警告: {warnings}")
            
            # 清理HTML内容，确保格式正确
            html_content = clean_html_content(html_content)
            
            logger.info(f"Word文档转换成功: {file.filename}")
            
            return jsonify({
                'success': True,
                'html_content': html_content,
                'original_filename': file.filename,
                'warnings': warnings if 'warnings' in locals() else []
            })
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
                logger.info(f"临时文件已清理: {temp_path}")
                
    except Exception as e:
        logger.error(f"Word文档读取失败: {e}", exc_info=True)
        return jsonify({'error': f'文档读取失败: {str(e)}'}), 500

@app.route('/api/editor/save-document', methods=['POST'])
def save_document():
    """将HTML内容保存为Word文档"""
    try:
        logger.info("开始处理HTML到Word转换请求")
        
        data = request.get_json()
        html_content = data.get('html_content', '')
        filename = data.get('filename', '文档')
        
        if not html_content.strip():
            return jsonify({'error': 'HTML内容为空'}), 400
        
        # 清理文件名，移除特殊字符
        filename = ''.join(c for c in filename if c.isalnum() or c in (' ', '-', '_', '（', '）', '中', '文')).strip() or '文档'
        
        # 生成输出文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{filename}_{timestamp}.docx"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
        
        # 转换HTML为Word文档
        convert_html_to_docx(html_content, output_path)
        
        logger.info(f"HTML转Word文档成功: {output_filename}")
        
        # 直接返回文件供下载
        return send_file(
            os.path.abspath(output_path),
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"HTML转Word文档失败: {e}", exc_info=True)
        return jsonify({'error': f'文档保存失败: {str(e)}'}), 500

@app.route('/api/editor/load-document-by-url', methods=['POST'])
def load_document_by_url():
    """通过文件URL加载Word文档并转换为HTML"""
    try:
        logger.info("开始处理通过URL加载Word文档请求")
        
        data = request.get_json()
        file_url = data.get('file_url', '')
        filename = data.get('filename', '文档')
        
        if not file_url:
            return jsonify({'error': '未提供文件URL'}), 400
        
        # 检查URL类型并处理
        if file_url.startswith('/download/'):
            # 处理相对下载URL，转换为实际文件路径
            filename_only = file_url.replace('/download/', '')
            file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename_only)
            logger.info(f"处理下载URL: {file_url} -> {file_path}")
        elif file_url.startswith('/download-tech/'):
            # 处理技术方案下载URL
            filename_only = file_url.replace('/download-tech/', '')
            file_path = os.path.join(app.config['TECH_OUTPUT_FOLDER'], filename_only)
            logger.info(f"处理技术方案下载URL: {file_url} -> {file_path}")
        elif file_url.startswith('/') or file_url.startswith('file://'):
            # 本地文件路径处理
            if file_url.startswith('file://'):
                file_path = file_url[7:]  # 移除 'file://' 前缀
            else:
                file_path = file_url
            logger.info(f"处理本地文件路径: {file_url} -> {file_path}")
        else:
            # 处理HTTP URL（如果需要的话）
            return jsonify({'error': '暂不支持HTTP URL下载'}), 400
        
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return jsonify({'error': '文件不存在'}), 404
            
        # 读取本地文件
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        # 检查文件类型 - 检查实际文件路径而不是传入的filename参数
        if not (file_path.endswith('.docx') or file_path.endswith('.doc')):
            return jsonify({'error': '只支持Word文档格式(.docx/.doc)'}), 400
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx' if file_path.endswith('.docx') else '.doc') as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # 使用mammoth将Word文档转换为HTML
            import mammoth
            with open(temp_file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                messages = result.messages
                
            logger.info(f"Word文档转换成功，包含 {len(messages)} 个转换消息")
            
            return jsonify({
                'success': True,
                'html_content': html_content,
                'original_filename': filename,
                'conversion_messages': [str(msg) for msg in messages]
            })
        
        finally:
            # 清理临时文件
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        
    except Exception as e:
        logger.error(f"通过URL加载Word文档失败: {e}", exc_info=True)
        return jsonify({'error': f'文档加载失败: {str(e)}'}), 500

@app.route('/api/editor/upload-image', methods=['POST'])
def upload_image():
    """上传编辑器中的图片"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': '未上传图片'}), 400
        
        file = request.files['image']
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400
        
        # 检查文件类型
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_ext not in allowed_extensions:
            return jsonify({'error': '只支持常见图片格式'}), 400
        
        # 保存图片文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        image_filename = f"editor_image_{timestamp}_{secure_filename(file.filename)}"
        
        # 创建图片存储目录
        images_dir = os.path.join('outputs', 'images')
        os.makedirs(images_dir, exist_ok=True)
        
        image_path = os.path.join(images_dir, image_filename)
        file.save(image_path)
        
        # 返回图片URL
        image_url = f'/outputs/images/{image_filename}'
        
        logger.info(f"编辑器图片上传成功: {image_filename}")
        
        return jsonify({
            'success': True,
            'location': image_url
        })
        
    except Exception as e:
        logger.error(f"编辑器图片上传失败: {e}")
        return jsonify({'error': f'图片上传失败: {str(e)}'}), 500

@app.route('/process-business-response', methods=['POST'])
def process_business_response():
    """处理商务应答请求"""
    try:
        logger.info("开始处理商务应答请求")
        
        # 检查上传的文件
        if 'template_file' not in request.files:
            return jsonify({'error': '未上传模板文件'}), 400
            
        file = request.files['template_file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
            
        # 获取表单数据
        company_id = request.form.get('company_id')
        use_mcp = request.form.get('use_mcp', 'false').lower() == 'true'
        
        if not company_id:
            return jsonify({'error': '请选择公司'}), 400
            
        # 加载公司信息
        company_file = os.path.join(company_configs_dir, f'{company_id}.json')
        if not os.path.exists(company_file):
            return jsonify({'error': '公司配置不存在'}), 404
            
        with open(company_file, 'r', encoding='utf-8') as f:
            company_info = json.load(f)
            
        # 自动从配置文件读取项目信息
        import configparser
        # 只读取"读取信息"模块生成的配置文件
        config_file = os.path.join(tender_info_path, 'tender_config.ini')
        
        project_name = ''
        tender_no = ''
        
        if os.path.exists(config_file):
            try:
                config = configparser.ConfigParser(interpolation=None)
                config.read(config_file, encoding='utf-8')
                
                if config.has_section('PROJECT_INFO'):
                    project_name = config.get('PROJECT_INFO', 'project_name', fallback='')
                    tender_no = config.get('PROJECT_INFO', 'project_number', fallback='')
                    
                logger.info(f"从配置文件读取项目信息: 项目名称='{project_name}', 项目编号='{tender_no}'")
            except Exception as e:
                logger.error(f"读取项目配置失败: {e}")
                return jsonify({'error': f'读取项目配置失败: {str(e)}'}), 500
        else:
            logger.warning("项目配置文件不存在，将使用空的项目信息")
            
        # 生成当前日期
        date_text = datetime.now().strftime('%Y年%m月%d日')
            
        # 保存上传的文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = secure_filename(file.filename)
        upload_filename = f"{timestamp}_business_template_{filename}"
        upload_path = os.path.join('uploads', upload_filename)
        
        os.makedirs('uploads', exist_ok=True)
        file.save(upload_path)
        logger.info(f"商务应答模板上传完成: {upload_path}")
        
        # 准备输出文件名
        base_name = filename.rsplit('.', 1)[0]
        output_filename = f"{base_name}-商务应答-{timestamp}.docx"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
        
        # 统一使用MCP模式处理投标人名称
        try:
            # 强制重载模块以获取最新修改
            importlib.reload(mcp_bidder_name_processor_enhanced)
            from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
            processor = MCPBidderNameProcessor()
            logger.info("初始化增强版MCP投标人名称处理器（含占位符清理修复）- 已重载模块")
            
            # 使用MCP处理器的商务应答处理方法
            result = processor.process_business_response(
                input_file=upload_path,
                output_file=output_path,
                company_info=company_info,
                project_name=project_name,
                tender_no=tender_no,
                date_text=date_text
            )
            
        except ImportError as e:
            logger.error(f"MCP处理器加载失败: {e}")
            return jsonify({'error': 'MCP处理器不可用'}), 500
        
        # 如果文本处理成功，继续进行表格处理和图片插入
        if result.get('success'):
            current_file = output_path
            processing_steps = {
                'text': {
                    'success': True,
                    'count': result.get('stats', {}).get('replacements', 0),
                    'message': '文本填写完成'
                }
            }
            
            # 第2步：表格处理
            if TABLE_PROCESSOR_AVAILABLE:
                try:
                    logger.info("开始表格处理")
                    from table_processor import TableProcessor
                    table_config_path = os.path.join(os.path.dirname(__file__), 'config', 'table_config.json')
                    table_processor = TableProcessor(table_config_path)
                    
                    # 生成表格处理后的文件名
                    table_output = current_file.replace('.docx', '_table.docx')
                    table_result = table_processor.process_document(
                        current_file,
                        company_info,
                        table_output
                    )
                    
                    # 检查处理结果
                    if table_result.get('success'):
                        processing_steps['tables'] = {
                            'success': True,
                            'count': table_result.get('tables_processed', 0),
                            'fields': table_result.get('fields_filled', 0),
                            'message': f"处理了{table_result.get('tables_processed', 0)}个表格"
                        }
                        current_file = table_output
                        logger.info(f"表格处理完成: {table_result.get('message')}")
                    else:
                        processing_steps['tables'] = {
                            'success': False,
                            'error': table_result.get('error', '未知错误'),
                            'message': '表格处理失败'
                        }
                        logger.warning(f"表格处理失败: {table_result.get('error')}")
                    
                except Exception as e:
                    logger.error(f"表格处理失败: {e}")
                    processing_steps['tables'] = {
                        'success': False,
                        'error': str(e),
                        'message': '表格处理失败'
                    }
            else:
                processing_steps['tables'] = {
                    'success': False,
                    'message': '表格处理模块未加载'
                }
            
            # 第3步：图片插入
            if IMAGE_INSERTER_AVAILABLE:
                try:
                    logger.info("开始图片插入")
                    from image_inserter import SmartImageInserter
                    image_inserter = SmartImageInserter()
                    
                    # 生成最终输出文件名
                    final_output = current_file.replace('_table.docx', '_final.docx').replace('.docx', '_final.docx') if '_table' in current_file else current_file.replace('.docx', '_final.docx')
                    
                    # 检查公司信息中是否有资质图片
                    if 'qualifications' in company_info and company_info['qualifications']:
                        # 调用图片插入器，不传递输出路径，让它自己生成
                        image_output, image_results = image_inserter.process_document(
                            current_file,
                            company_info
                        )
                        
                        success_count = len([r for r in image_results if r.success])
                        if success_count > 0:
                            processing_steps['images'] = {
                                'success': True,
                                'count': success_count,
                                'message': f"插入了{success_count}张图片"
                            }
                            current_file = image_output  # 使用返回的输出路径
                            logger.info(f"图片插入完成: 成功插入{success_count}张图片")
                        else:
                            processing_steps['images'] = {
                                'success': False,
                                'count': 0,
                                'message': '未能插入任何图片'
                            }
                            logger.warning("图片插入: 未能插入任何图片")
                    else:
                        processing_steps['images'] = {
                            'success': False,
                            'message': '未配置资质图片'
                        }
                        logger.info("跳过图片插入：未配置资质图片")
                        
                except Exception as e:
                    logger.error(f"图片插入失败: {e}")
                    processing_steps['images'] = {
                        'success': False,
                        'error': str(e),
                        'message': '图片插入失败'
                    }
            else:
                processing_steps['images'] = {
                    'success': False,
                    'message': '图片插入模块未加载'
                }
            
            # 清理临时文件
            try:
                os.remove(upload_path)
                logger.info(f"清理上传的临时文件: {upload_path}")
                # 不要清理中间文件，因为它们可能就是最终文件
                # 只记录日志，不删除文件
                logger.info(f"保留处理后的文件: {current_file}")
            except Exception as e:
                logger.warning(f"清理临时文件失败: {e}")
            
            # 重命名最终文件为原始输出文件名
            final_file = current_file  # 保存实际的最终文件路径
            logger.info(f"当前文件路径: {current_file}")
            logger.info(f"目标输出路径: {output_path}")
            
            if current_file != output_path:
                try:
                    import shutil
                    # 如果目标文件已存在，先删除
                    if os.path.exists(output_path):
                        os.remove(output_path)
                        logger.info(f"删除已存在的目标文件: {output_path}")
                    
                    shutil.move(current_file, output_path)
                    logger.info(f"移动最终文件成功: {current_file} -> {output_path}")
                    final_file = output_path
                except Exception as e:
                    logger.warning(f"移动文件失败，保持使用当前文件: {e}")
                    # 移动失败时，保持使用实际生成的文件
                    final_file = current_file
            
            # 确保使用实际存在的文件名
            actual_filename = os.path.basename(final_file)
            logger.info(f"最终文件路径: {final_file}")
            logger.info(f"最终文件名: {actual_filename}")
            
            return jsonify({
                'success': True,
                'output_file': actual_filename,
                'download_url': f'/download/{actual_filename}',
                'filename': actual_filename,  # 添加明确的filename字段供预览使用
                'processing_steps': processing_steps,
                'statistics': {
                    'text_replacements': processing_steps['text'].get('count', 0),
                    'tables_processed': processing_steps['tables'].get('count', 0),
                    'fields_filled': processing_steps['tables'].get('fields', 0),
                    'images_inserted': processing_steps['images'].get('count', 0)
                },
                'message': '商务应答文档处理完成'
            })
        else:
            # 清理临时文件
            try:
                os.remove(upload_path)
                logger.info(f"清理临时文件: {upload_path}")
            except:
                pass
            return jsonify({'error': result.get('error', '处理失败')}), 500
            
    except Exception as e:
        logger.error(f"商务应答处理失败: {e}", exc_info=True)
        return jsonify({'error': f'处理失败: {str(e)}'}), 500

def clean_html_content(html_content):
    """清理HTML内容，确保格式正确"""
    try:
        from bs4 import BeautifulSoup
        
        # 解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 清理不需要的标签和属性
        for tag in soup.find_all():
            # 保留常见的样式属性
            allowed_attrs = ['style', 'class', 'id', 'href', 'src', 'alt', 'title']
            attrs = dict(tag.attrs)
            for attr in attrs:
                if attr not in allowed_attrs:
                    del tag.attrs[attr]
        
        # 返回清理后的HTML
        return str(soup)
    except Exception as e:
        logger.warning(f"HTML内容清理失败: {e}")
        return html_content

def convert_html_to_docx(html_content, output_path):
    """HTML转Word文档"""
    try:
        from htmldocx import HtmlToDocx
        
        # 创建HTML转换器
        new_parser = HtmlToDocx()
        
        # 转换HTML到Word
        new_parser.parse_html_string(html_content)
        
        # 保存文档
        new_parser.save(output_path)
        
        logger.info(f"HTML成功转换为Word文档: {output_path}")
        
    except Exception as e:
        logger.error(f"HTML转Word失败: {e}")
        # 如果htmldocx失败，尝试使用python-docx的基础功能
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            
            doc = Document()
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 简单处理：提取文本并添加到文档
            for element in soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                text = element.get_text().strip()
                if text:
                    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        # 标题
                        doc.add_heading(text, level=int(element.name[1]))
                    else:
                        # 普通段落
                        doc.add_paragraph(text)
            
            doc.save(output_path)
            logger.info(f"使用备用方法成功转换HTML为Word文档: {output_path}")
            
        except Exception as e2:
            logger.error(f"备用HTML转Word方法也失败: {e2}")
            raise Exception(f"HTML转Word失败: {str(e)} / {str(e2)}")

# ===================== 新增：表格处理和图片插入路由 =====================

@app.route('/api/document/process', methods=['POST'])
def process_document_complete():
    """完整文档处理接口 - 包括表格填充和图片插入"""
    if not DOCUMENT_PROCESSOR_AVAILABLE:
        return jsonify({
            'success': False,
            'error': '文档处理模块未加载'
        }), 500
    
    try:
        # 获取上传的文件
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '未找到上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '未选择文件'}), 400
        
        # 获取公司ID
        company_id = request.form.get('company_id')
        if not company_id:
            return jsonify({'success': False, 'error': '未指定公司'}), 400
        
        # 获取处理选项
        process_tables = request.form.get('process_tables', 'true').lower() == 'true'
        insert_images = request.form.get('insert_images', 'true').lower() == 'true'
        process_names = request.form.get('process_names', 'false').lower() == 'true'
        
        # 保存上传的文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = secure_filename(file.filename)
        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, filename)
        file.save(input_path)
        
        # 加载公司信息
        processor = DocumentProcessor()
        company_info = processor.load_company_info(company_id)
        
        if not company_info:
            return jsonify({'success': False, 'error': '公司信息未找到'}), 404
        
        # 设置处理选项
        options = ProcessingOptions(
            process_names=process_names,
            process_tables=process_tables,
            insert_images=insert_images,
            keep_intermediate=False
        )
        
        # 处理文档
        result = processor.process_document(input_path, company_info, options)
        
        # 清理临时文件
        shutil.rmtree(temp_dir)
        
        if result.success:
            # 获取输出文件名
            output_filename = os.path.basename(result.output_path)
            
            return jsonify({
                'success': True,
                'output_file': output_filename,
                'download_url': f'/download/{output_filename}',
                'statistics': result.statistics,
                'processing_time': result.processing_time
            })
        else:
            return jsonify({
                'success': False,
                'errors': result.errors
            }), 500
            
    except Exception as e:
        logger.error(f"文档处理失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/table/analyze', methods=['POST'])
def analyze_tables():
    """分析文档中的表格结构"""
    if not DOCUMENT_PROCESSOR_AVAILABLE:
        return jsonify({
            'success': False,
            'error': '表格处理模块未加载'
        }), 500
    
    try:
        # 获取上传的文件
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '未找到上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '未选择文件'}), 400
        
        # 保存临时文件
        temp_dir = tempfile.mkdtemp()
        filename = secure_filename(file.filename)
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        
        # 分析表格
        table_config_path = os.path.join(os.path.dirname(__file__), 'config', 'table_config.json')
        processor = TableProcessor(table_config_path)
        analysis = processor.analyze_tables(file_path)
        
        # 清理临时文件
        shutil.rmtree(temp_dir)
        
        return jsonify({
            'success': True,
            'analysis': analysis
        })
        
    except Exception as e:
        logger.error(f"表格分析失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/table/process', methods=['POST'])
def process_tables_only():
    """仅处理表格填充"""
    if not DOCUMENT_PROCESSOR_AVAILABLE:
        return jsonify({
            'success': False,
            'error': '表格处理模块未加载'
        }), 500
    
    try:
        # 获取上传的文件
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '未找到上传文件'}), 400
        
        file = request.files['file']
        company_id = request.form.get('company_id')
        
        if not company_id:
            return jsonify({'success': False, 'error': '未指定公司'}), 400
        
        # 保存临时文件
        temp_dir = tempfile.mkdtemp()
        filename = secure_filename(file.filename)
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        
        # 加载公司信息
        processor = DocumentProcessor()
        company_info = processor.load_company_info(company_id)
        
        # 处理表格
        table_config_path = os.path.join(os.path.dirname(__file__), 'config', 'table_config.json')
        table_processor = TableProcessor(table_config_path)
        output_path = table_processor.process_document(file_path, company_info)
        
        # 清理临时文件
        shutil.rmtree(temp_dir)
        
        # 获取输出文件名
        output_filename = os.path.basename(output_path)
        
        return jsonify({
            'success': True,
            'output_file': output_filename,
            'download_url': f'/download/{output_filename}'
        })
        
    except Exception as e:
        logger.error(f"表格处理失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/image/insert', methods=['POST'])
def insert_images_only():
    """仅处理图片插入"""
    if not DOCUMENT_PROCESSOR_AVAILABLE:
        return jsonify({
            'success': False,
            'error': '图片插入模块未加载'
        }), 500
    
    try:
        # 获取上传的文件
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '未找到上传文件'}), 400
        
        file = request.files['file']
        company_id = request.form.get('company_id')
        
        if not company_id:
            return jsonify({'success': False, 'error': '未指定公司'}), 400
        
        # 保存临时文件
        temp_dir = tempfile.mkdtemp()
        filename = secure_filename(file.filename)
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        
        # 加载公司信息
        processor = DocumentProcessor()
        company_info = processor.load_company_info(company_id)
        
        # 插入图片
        image_inserter = SmartImageInserter()
        output_path, results = image_inserter.process_document(file_path, company_info)
        
        # 清理临时文件
        shutil.rmtree(temp_dir)
        
        # 统计结果
        success_count = sum(1 for r in results if r.success)
        
        # 获取输出文件名
        output_filename = os.path.basename(output_path)
        
        return jsonify({
            'success': True,
            'output_file': output_filename,
            'download_url': f'/download/{output_filename}',
            'images_inserted': success_count,
            'results': [
                {
                    'image_type': r.image_type,
                    'success': r.success,
                    'strategy': r.strategy_used,
                    'position': r.position
                } for r in results
            ]
        })
        
    except Exception as e:
        logger.error(f"图片插入失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ===================== 结束新增路由 =====================

def find_available_port(start_port=8080):
    """找到可用端口"""
    import socket
    for port in range(start_port, start_port + 100):
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind(('0.0.0.0', port))
                return port
        except OSError:
            continue
    return None

if __name__ == '__main__':
    port = find_available_port()
    if port:
        print("启动AI标书智能生成系统...")
        print(f"Web界面地址: http://localhost:{port}")
        print("点对点应答: 支持 .docx, .doc")
        print("技术方案: 支持 .docx, .doc, .pdf") 
        print("招标信息提取: 支持 .docx, .doc, .txt, .pdf")
        print("使用始皇API生成专业内容")
        if TECH_PROPOSAL_AVAILABLE:
            print("技术方案生成功能已加载")
        else:
            print("警告: 技术方案生成功能未加载")
        if TENDER_INFO_AVAILABLE:
            print("招标信息提取功能已加载")
        else:
            print("警告: 招标信息提取功能未加载")
        app.run(debug=True, host='0.0.0.0', port=port)
    else:
        print("错误: 无法找到可用端口，请手动指定端口运行")