#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def test_format18_no_colon():
    """测试格式18：无冒号的供应商名称处理"""
    
    test_input = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/test_format18_no_colon_input.docx"
    test_output = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/test_format18_no_colon_output.docx"
    
    # 创建测试文档，包含各种无冒号格式
    doc = Document()
    doc.add_paragraph("其他内容")
    doc.add_paragraph("供应商名称                                    ")  # 40个空格，符合格式18
    doc.add_paragraph("中间内容")
    doc.add_paragraph("供应商名称                              ")      # 30个空格，符合格式18
    doc.add_paragraph("结束内容")
    doc.add_paragraph("供应商名称            ")                       # 12个空格，不符合格式18（少于20个）
    doc.save(test_input)
    
    print(f"✅ 创建测试文档: {test_input}")
    
    # 处理前检查
    print(f"\n=== 处理前检查 ===")
    doc_before = Document(test_input)
    for i, para in enumerate(doc_before.paragraphs):
        para_text = para.text
        print(f"段落 #{i}: '{para_text}' (长度: {len(para_text)})")
        if "供应商名称" in para_text and "：" not in para_text and ":" not in para_text:
            space_count = para_text.count(' ')
            if space_count >= 20:
                print(f"  ✅ 符合格式18: {space_count}个空格")
            else:
                print(f"  ⚠️ 空格不足: {space_count}个空格（需要>=20）")
    
    try:
        # 执行处理
        processor = MCPBidderNameProcessor()
        company_name = "智慧足迹数据科技有限公司"
        
        print(f"\n=== 开始处理 ===")
        print(f"公司名称: {company_name}")
        
        result = processor.process_bidder_name(
            input_file=test_input,
            output_file=test_output,
            company_name=company_name
        )
        
        print(f"\n处理结果: {result}")
        
        if result.get('success'):
            print("✅ 处理报告成功!")
            
            # 处理后检查
            print(f"\n=== 处理后检查 ===")
            if os.path.exists(test_output):
                doc_after = Document(test_output)
                for i, para in enumerate(doc_after.paragraphs):
                    para_text = para.text
                    print(f"段落 #{i}: '{para_text}'")
                    if company_name in para_text:
                        print(f"    ✅ 找到公司名称!")
            else:
                print(f"❌ 输出文件不存在: {test_output}")
            
            # 分析处理统计
            stats = result.get('stats', {})
            patterns = stats.get('patterns_found', [])
            
            print(f"\n=== 处理统计 ===")
            print(f"总替换次数: {stats.get('total_replacements', 0)}")
            print(f"处理的模式:")
            
            format18_count = 0
            for pattern in patterns:
                rule_desc = pattern.get('description', '')
                print(f"  - 规则 #{pattern.get('rule_index')}: {rule_desc}")
                print(f"    原始文本: '{pattern.get('original_text')}'")
                print(f"    段落位置: #{pattern.get('paragraph_index')}")
                if '无冒号' in rule_desc:
                    format18_count += 1
                    print(f"    ✅ 这是格式18!")
                    
            print(f"\n✅ 格式18处理次数: {format18_count}")
            
        else:
            print(f"❌ 处理失败: {result.get('message', 'Unknown error')}")
            
    except Exception as e:
        print(f"❌ 处理异常: {e}")
        import traceback
        traceback.print_exc()
    
    print(f"\n=== 文件保留用于检查 ===")
    print(f"输入文件: {test_input}")
    print(f"输出文件: {test_output}")

if __name__ == "__main__":
    test_format18_no_colon()