#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
from docx import Document

def analyze_output_document():
    """分析输出文档中的供应商名称处理情况"""
    
    output_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/outputs/docx-商务应答-20250906_075450.docx"
    
    if not os.path.exists(output_file):
        print(f"❌ 文件不存在: {output_file}")
        return
    
    print(f"分析文件: {output_file}")
    print("=" * 80)
    
    try:
        doc = Document(output_file)
        print(f"文档共有 {len(doc.paragraphs)} 个段落\n")
        
        # 查找所有包含"供应商名称"的段落
        supplier_paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if "供应商名称" in para.text:
                supplier_paragraphs.append((i, para.text))
        
        print(f"找到 {len(supplier_paragraphs)} 个包含'供应商名称'的段落：\n")
        
        # 分析每个段落
        for para_idx, para_text in supplier_paragraphs:
            print(f"段落 #{para_idx}: '{para_text}'")
            print(f"  长度: {len(para_text)}")
            
            # 分析格式
            if "：" in para_text or ":" in para_text:
                colon_pos = para_text.find("：") if "：" in para_text else para_text.find(":")
                after_colon = para_text[colon_pos+1:]
                print(f"  有冒号，冒号后内容: '{after_colon}' (长度: {len(after_colon)})")
                
                # 检查是否已填写
                if "中国联合" in after_colon or "智慧足迹" in after_colon:
                    print(f"  ✅ 已填写公司名称")
                elif after_colon.strip() == "" or after_colon.isspace():
                    print(f"  ❌ 只有空格，未填写")
                elif "_" in after_colon:
                    print(f"  ❌ 有下划线占位符，未填写")
                else:
                    print(f"  ⚠️ 有其他内容: '{after_colon.strip()}'")
            else:
                # 无冒号的情况
                text_after_label = para_text.replace("供应商名称", "")
                space_count = text_after_label.count(" ")
                print(f"  无冒号，后面有 {space_count} 个空格")
                
                if "中国联合" in para_text or "智慧足迹" in para_text:
                    print(f"  ✅ 已填写公司名称")
                elif space_count >= 10:
                    print(f"  ❌ 有{space_count}个空格但未填写（可能需要处理）")
                else:
                    print(f"  其他格式")
            
            print()
        
        # 特别检查少量空格的情况（10-19个空格）
        print("\n=== 特别检查：少量空格的供应商名称（10-19个空格）===\n")
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            # 检查无冒号的情况
            if text.startswith("供应商名称") and "：" not in text and ":" not in text:
                space_count = text.replace("供应商名称", "").count(" ")
                if 10 <= space_count < 20:
                    print(f"段落 #{i}: '{text}'")
                    print(f"  ⚠️ 发现10-19个空格的无冒号格式（{space_count}个空格）")
                    print(f"  当前规则要求>=20个空格才处理")
                    print()
        
        # 检查带冒号但空格较少的情况
        print("\n=== 检查：带冒号但空格较少的情况 ===\n")
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if "供应商名称：" in text or "供应商名称:" in text:
                colon = "：" if "：" in text else ":"
                after_colon = text.split(colon, 1)[1] if colon in text else ""
                
                # 只看纯空格的情况
                if after_colon and all(c == " " for c in after_colon.rstrip()):
                    space_count = len(after_colon.rstrip())
                    if 5 <= space_count < 10:
                        print(f"段落 #{i}: '{text}'")
                        print(f"  有{space_count}个空格（5-9个），当前可能未被处理")
                        if not any(company in text for company in ["中国联合", "智慧足迹"]):
                            print(f"  ❌ 未填写公司名称")
                        print()
                        
    except Exception as e:
        print(f"❌ 分析出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_output_document()