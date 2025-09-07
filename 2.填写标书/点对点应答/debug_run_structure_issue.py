#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试run结构问题：为什么处理采购编号后找不到供应商名称的run
"""

import os
import re
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 配置日志
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def create_test_document_with_run_analysis():
    """创建测试文档并分析run结构"""
    doc = Document()
    doc.add_heading('Run结构问题调试', 0)
    
    # 添加一些前置段落
    for i in range(29):
        doc.add_paragraph(f'段落 #{i+1}: 前置内容...')
    
    # 段落#30 - 模拟实际的run结构
    para30 = doc.add_paragraph()
    # 根据日志，实际文档可能有复杂的run结构
    para30.add_run("供应商名称：")  # run 0
    para30.add_run("     ")       # run 1  
    para30.add_run("                  ")  # run 2
    para30.add_run("   ")         # run 3
    para30.add_run("采购编号")    # run 4
    para30.add_run("：")          # run 5  
    para30.add_run("     ")       # run 6
    para30.add_run("     ")       # run 7
    para30.add_run("     ")       # run 8
    para30.add_run("     ")       # run 9
    para30.add_run("     ")       # run 10
    
    print("创建的段落#30 run结构：")
    for i, run in enumerate(para30.runs):
        print(f"  Run {i}: '{run.text}' (长度: {len(run.text)})")
    
    full_text = ''.join(run.text for run in para30.runs)
    print(f"完整文本: '{full_text}'")
    print(f"完整文本长度: {len(full_text)}")
    
    return doc, para30

def analyze_processing_steps():
    """逐步分析处理过程中run结构的变化"""
    print("\n" + "="*60)
    print("逐步分析处理过程")
    print("="*60)
    
    doc, para30 = create_test_document_with_run_analysis()
    
    print("\n第1步：初始run结构")
    for i, run in enumerate(para30.runs):
        print(f"  Run {i}: '{run.text}'")
    
    # 模拟第一步处理：处理采购编号
    print("\n第2步：模拟处理采购编号（规则#6）")
    
    # 根据日志，处理后的结果是：
    # '供应商名称：   采购编号：GXTC-C-251590031  '
    
    # 手动修改run结构来模拟处理结果
    para30.runs[0].text = "供应商名称："
    para30.runs[1].text = "   "  
    para30.runs[2].text = ""     # 清空
    para30.runs[3].text = ""     # 清空
    para30.runs[4].text = "采购编号"
    para30.runs[5].text = "："
    para30.runs[6].text = "GXTC-C-251590031"  # 填入采购编号
    para30.runs[7].text = ""     # 清空
    para30.runs[8].text = "  "   # 保留部分空格
    if len(para30.runs) > 9:
        para30.runs[9].text = ""     # 清空
    if len(para30.runs) > 10:
        para30.runs[10].text = ""    # 清空
    
    print("处理采购编号后的run结构：")
    for i, run in enumerate(para30.runs):
        print(f"  Run {i}: '{run.text}'")
    
    full_text_after = ''.join(run.text for run in para30.runs)
    print(f"处理后完整文本: '{full_text_after}'")
    
    print("\n第3步：分析供应商名称查找问题")
    
    # 现在尝试查找包含"供应商名称"的run
    label = "供应商名称"
    label_run = None
    for i, run in enumerate(para30.runs):
        print(f"检查Run {i}: '{run.text}' -> 是否包含'{label}': {label in run.text}")
        if label in run.text:
            label_run = run
            print(f"✅ 找到包含标签的run: {i}")
            break
    
    if not label_run:
        print("❌ 问题确认：未能找到包含'供应商名称'的run")
        
        # 分析原因
        print("\n问题分析：")
        print("1. 采购编号处理过程中可能改变了run结构")
        print("2. 多个规则按顺序处理时，后续规则找不到目标文本")
        print("3. 需要改进多规则处理的协调机制")
    
    return doc, para30

def test_fix_strategy():
    """测试修复策略"""
    print("\n" + "="*60) 
    print("测试修复策略")
    print("="*60)
    
    # 策略1：在处理前收集所有匹配的规则和位置
    # 策略2：按位置排序处理，避免相互干扰
    # 策略3：改进查找逻辑，支持部分文本匹配
    
    print("推荐修复策略：")
    print("1. 批量收集：在处理前收集所有匹配的规则和位置信息")
    print("2. 智能排序：按文本位置从后往前处理，避免位置偏移")
    print("3. 容错查找：改进查找逻辑，支持run结构变化后的匹配")
    print("4. 状态保持：在处理过程中维护文本位置映射")

def main():
    """主函数"""
    print("段落#30 run结构问题调试")
    print("="*60)
    
    analyze_processing_steps()
    test_fix_strategy()

if __name__ == '__main__':
    main()