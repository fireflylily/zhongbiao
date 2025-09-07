#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试供应商名称和采购编号在同一行的处理
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_test_document():
    """创建模拟实际格式的测试文档"""
    doc = Document()
    
    # 完全模拟原始文档格式 - 供应商名称和采购编号在同一行
    para = doc.add_paragraph()
    para.add_run("供应商名称:")
    para.add_run("_______________________")
    para.add_run("采购编号:")
    para.add_run("_______________________")
    
    return doc


def analyze_paragraph_structure(doc, description):
    """分析段落结构"""
    print(f"\n{description}")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"段落{i}: '{para.text}'")
            print(f"  Run数量: {len(para.runs)}")
            for j, run in enumerate(para.runs):
                if run.text:
                    print(f"    Run{j}: '{run.text}'")


def main():
    """主测试函数"""
    print("="*60)
    print("同一行字段处理测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_same_line_input.docx')
    output_file = os.path.join(current_dir, 'test_same_line_output.docx')
    
    # 创建测试文档
    print("\n1. 创建测试文档...")
    doc = create_test_document()
    doc.save(input_file)
    
    # 分析原始结构
    analyze_paragraph_structure(doc, "原始文档结构:")
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
    }
    
    # 处理文档
    print("\n2. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 分析处理后的结构
            processed_doc = Document(output_file)
            analyze_paragraph_structure(processed_doc, "\n3. 处理后文档结构:")
            
            # 检查问题
            print("\n4. 问题分析:")
            para_count = len([p for p in processed_doc.paragraphs if p.text.strip()])
            
            if para_count == 1:
                para = processed_doc.paragraphs[0]
                text = para.text
                if '供应商名称' in text and '采购编号' in text:
                    if '中国联合网络通信有限公司' in text and 'GXTC-C-251590031' in text:
                        print("   ✅ 供应商名称和采购编号都在同一行")
                        print(f"   格式: {text}")
                    else:
                        print("   ⚠️ 字段在同一行但值未正确填写")
                else:
                    print("   ❌ 缺少某个字段")
            else:
                print(f"   ❌ 段落数量异常: {para_count} (期望1个)")
                # 检查是否被分成了多行
                for i, para in enumerate(processed_doc.paragraphs):
                    if para.text.strip():
                        print(f"   段落{i}: {para.text}")
                        if 'GXTC-C-251590031' in para.text and '采购编号' not in para.text:
                            print("      ⚠️ 采购编号值单独成行了！")
            
            print("\n5. 期望结果:")
            print("   供应商名称:中国联合网络通信有限公司采购编号:GXTC-C-251590031")
            print("   (保持在同一行)")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 保留文件供检查
        print(f"\n6. 输出文件: {output_file}")
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)


if __name__ == '__main__':
    main()