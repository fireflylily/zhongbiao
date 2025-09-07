#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试具体的盖章格式问题
"""

import re
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def test_specific_format():
    """测试具体的格式"""
    processor = MCPBidderNameProcessor()
    
    # 获取规则2
    seal_rule = processor.bidder_patterns[2]
    pattern = seal_rule['pattern']
    
    print(f"🔍 测试规则: {seal_rule['description']}")
    print(f"📝 正则表达式: {pattern.pattern}")
    print("=" * 80)
    
    # 测试具体的问题格式
    problem_text = "    公司名称（全称、盖章）：________________"
    
    print(f"❓ 问题文本: '{problem_text}'")
    print(f"   原始文本: {repr(problem_text)}")
    print(f"   长度: {len(problem_text)}")
    print(f"   前导空格数: {len(problem_text) - len(problem_text.lstrip())}")
    print(f"   下划线数量: {problem_text.count('_')}")
    
    # 测试匹配
    match = pattern.match(problem_text)
    if match:
        print("✅ 匹配成功!")
        print(f"   groups: {match.groupdict()}")
    else:
        print("❌ 匹配失败!")
        
        # 逐步调试
        print("\n🔧 调试步骤:")
        
        # 1. 检查前导空格
        stripped = problem_text.lstrip()
        print(f"1. 去除前导空格: '{stripped}'")
        match1 = pattern.match(stripped)
        if match1:
            print("   ✅ 去除前导空格后匹配成功")
        else:
            print("   ❌ 去除前导空格后仍不匹配")
        
        # 2. 检查下划线数量是否足够
        underscore_count = problem_text.count('_')
        print(f"2. 下划线数量: {underscore_count}")
        if underscore_count >= 3:
            print("   ✅ 下划线数量足够")
        else:
            print("   ❌ 下划线数量不足")
        
        # 3. 手动构建测试
        test_variants = [
            "公司名称（全称、盖章）：________________",
            "公司名称（全称、盖章）：_______________",
            "公司名称（全称、盖章）：______________",
            "公司名称（全称、盖章）：_____________",
            "公司名称（全称、盖章）：____________",
            "公司名称（全称、盖章）：___________",
            "公司名称（全称、盖章）：__________",
            "公司名称（全称、盖章）：_________",
            "公司名称（全称、盖章）：________",
            "公司名称（全称、盖章）：_______",
            "公司名称（全称、盖章）：______",
            "公司名称（全称、盖章）：_____",
            "公司名称（全称、盖章）：____",
            "公司名称（全称、盖章）：___",
        ]
        
        print("\n3. 测试不同下划线数量:")
        for test_text in test_variants:
            test_match = pattern.match(test_text)
            status = "✅" if test_match else "❌"
            print(f"   {status} {len(test_text.split('：')[1])}个下划线: '{test_text}'")

def test_with_real_document():
    """测试真实文档处理"""
    from docx import Document
    
    # 创建包含问题格式的测试文档
    doc = Document()
    problem_text = "    公司名称（全称、盖章）：________________"
    para = doc.add_paragraph()
    para.add_run(problem_text)
    
    test_file = "debug_specific_input.docx"
    doc.save(test_file)
    print(f"✅ 创建测试文档: {test_file}")
    
    # 测试处理
    processor = MCPBidderNameProcessor()
    company_name = "上海智慧足迹数据科技有限公司"
    output_file = "debug_specific_output.docx"
    
    print(f"🚀 测试处理...")
    result = processor.process_bidder_name(test_file, output_file, company_name)
    
    if result.get('success', False):
        print("✅ 处理成功")
        
        # 检查结果
        output_doc = Document(output_file)
        for para in output_doc.paragraphs:
            if para.text.strip():
                print(f"📄 结果: '{para.text}'")
                if company_name in para.text:
                    print("   ✅ 公司名称已替换")
                else:
                    print("   ❌ 公司名称未替换")
    else:
        print("❌ 处理失败")

if __name__ == "__main__":
    print("🧪 调试具体盖章格式问题")
    print("=" * 80)
    
    test_specific_format()
    
    print("\n" + "=" * 80)
    
    test_with_real_document()
    
    print("\n✅ 调试完成")