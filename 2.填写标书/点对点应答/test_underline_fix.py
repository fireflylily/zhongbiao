#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试下划线修复
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_underline_test_document():
    """创建有下划线格式的测试文档"""
    doc = Document()
    doc.add_heading('下划线修复测试', 0)
    
    # 测试1: 供应商名称和采购编号在同一行，原有下划线格式
    para1 = doc.add_paragraph()
    para1.add_run("供应商名称：")
    run1 = para1.add_run("                    ")
    run1.underline = True  # 占位符有下划线
    para1.add_run("采购编号：")
    run2 = para1.add_run("              ")
    run2.underline = True  # 占位符有下划线
    
    # 测试2: 单独的供应商名称，有下划线占位符
    para2 = doc.add_paragraph()
    para2.add_run("供应商名称：")
    run3 = para2.add_run("___________________")
    run3.underline = True  # 下划线占位符有下划线格式
    
    return doc


def check_underline_results(doc, description):
    """检查下划线结果"""
    print(f"\n{description}")
    underline_issues = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and ('中国联合' in text or '供应商名称' in text):
            print(f"  段落{i}: {text}")
            
            for j, run in enumerate(para.runs):
                if run.text.strip():
                    underline_status = "有下划线" if run.underline else "无下划线"
                    print(f"    Run{j}: '{run.text}' [{underline_status}]")
                    
                    # 检查填充内容是否有下划线
                    if ('中国联合网络通信有限公司' in run.text or 'GXTC-C-251590031' in run.text) and run.underline:
                        underline_issues.append(f"段落{i} Run{j}: 填充内容'{run.text[:20]}...'有下划线")
    
    return underline_issues


def main():
    """主测试函数"""
    print("="*60)
    print("下划线修复测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_underline_input.docx')
    output_file = os.path.join(current_dir, 'test_underline_output.docx')
    
    # 创建测试文档
    print("\n1. 创建有下划线格式的测试文档...")
    doc = create_underline_test_document()
    doc.save(input_file)
    
    # 分析原始格式
    check_underline_results(doc, "原始文档格式:")
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
    }
    
    # 处理文档
    print("\n2. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            processed_doc = Document(output_file)
            underline_issues = check_underline_results(processed_doc, "\n3. 处理后文档格式:")
            
            if underline_issues:
                print("\n❌ 下划线问题:")
                for issue in underline_issues:
                    print(f"  - {issue}")
                print("\n❌ 修复失败，填充内容仍有下划线")
            else:
                print("\n✅ 下划线修复成功！")
                print("  - 填充内容没有下划线")
                print("  - 格式正确保持")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)
        print(f"\n4. 输出文件: {output_file}")


if __name__ == '__main__':
    main()