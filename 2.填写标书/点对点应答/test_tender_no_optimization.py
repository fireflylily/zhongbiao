#!/usr/bin/env python3
"""
测试项目编号处理优化效果
验证单run优先策略是否减少字体变化问题
"""

import os
import sys
from pathlib import Path
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('tender_no_optimization_test.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    # 添加测试段落
    # 场景1：单run内的项目编号（期望单run替换，不影响格式）
    p1 = doc.add_paragraph()
    p1.add_run("项目编号：（采购编号） 测试内容")
    
    # 场景2：跨run的项目编号（需要跨run处理）
    p2 = doc.add_paragraph()
    p2.add_run("项目编号：（")
    p2.add_run("采购编号")  # 不同run
    p2.add_run("） 测试内容")
    
    # 场景3：混合格式段落（多个字段在同一段落）
    p3 = doc.add_paragraph()
    p3.add_run("供应商：（供应商名称） 编号：（采购编号） 完成。")
    
    test_file = "test_tender_no_optimization_input.docx"
    doc.save(test_file)
    logger.info(f"创建测试文档: {test_file}")
    return test_file

def test_tender_no_processing():
    """测试项目编号处理优化"""
    logger.info("=" * 60)
    logger.info("开始测试项目编号处理优化")
    logger.info("=" * 60)
    
    try:
        # 创建测试文档
        input_file = create_test_document()
        output_file = "test_tender_no_optimization_output.docx"
        
        # 创建处理器
        processor = MCPBidderNameProcessor()
        
        # 设置测试数据
        company_info = {
            'company_name': '测试科技有限公司',
            'registeredAddress': '北京市海淀区测试路123号'
        }
        
        project_name = "智能办公系统采购项目"
        tender_no = "64525343"  # 测试编号
        
        logger.info(f"测试配置:")
        logger.info(f"  公司名称: {company_info['company_name']}")
        logger.info(f"  项目名称: {project_name}")
        logger.info(f"  项目编号: {tender_no}")
        
        # 处理文档 - 修复公司名称传递问题
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name=project_name,
            tender_no=tender_no,
            date_text="2025年9月6日"
        )
        
        # 如果没有处理成功，直接测试项目编号处理
        if not result.get('stats', {}).get('total_replacements', 0):
            logger.info("使用直接方法测试项目编号处理...")
            # 设置处理器的项目编号
            processor.project_number = tender_no
            
            # 重新加载文档进行测试
            from docx import Document
            doc = Document(input_file)
            
            for i, paragraph in enumerate(doc.paragraphs):
                if '（采购编号）' in paragraph.text:
                    logger.info(f"测试段落 {i+1}: {paragraph.text}")
                    logger.info(f"  Run结构: {[run.text for run in paragraph.runs]}")
                    
                    # 模拟规则匹配
                    import re
                    pattern = re.compile(r'[\(（]\s*采购编号\s*[\)）]')
                    match = pattern.search(paragraph.text)
                    
                    if match:
                        rule = {'description': '测试-采购编号处理', 'type': 'replace_content_tender_no'}
                        success = processor._replace_content_tender_no_method(paragraph, match, rule)
                        logger.info(f"  处理结果: {success}")
                        logger.info(f"  处理后文本: {paragraph.text}")
                        logger.info(f"  处理后Run结构: {[run.text for run in paragraph.runs]}")
            
            # 保存测试结果
            doc.save(output_file)
            logger.info(f"测试处理完成，保存到: {output_file}")
        
        logger.info("=" * 60)
        logger.info("处理结果:")
        logger.info(f"  成功状态: {result.get('success', False)}")
        logger.info(f"  处理统计: {result.get('stats', {})}")
        
        # 分析输出文档
        if os.path.exists(output_file):
            analyze_output_document(output_file)
        
        logger.info("=" * 60)
        logger.info("测试完成！")
        logger.info("请查看输出文档和日志，验证:")
        logger.info("  1. 项目编号是否正确填写")
        logger.info("  2. 周围文字格式是否保持不变") 
        logger.info("  3. 日志中是否优先使用单run替换")
        logger.info("=" * 60)
        
        return result
        
    except Exception as e:
        logger.error(f"测试失败: {e}", exc_info=True)
        return None

def analyze_output_document(file_path):
    """分析输出文档内容"""
    try:
        doc = Document(file_path)
        logger.info(f"\n分析输出文档: {file_path}")
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                logger.info(f"  段落 {i+1}: {paragraph.text}")
                logger.info(f"    Run数量: {len(paragraph.runs)}")
                for j, run in enumerate(paragraph.runs):
                    if run.text:
                        logger.info(f"      Run {j+1}: '{run.text}'")
                        
    except Exception as e:
        logger.error(f"分析输出文档失败: {e}")

if __name__ == "__main__":
    test_tender_no_processing()