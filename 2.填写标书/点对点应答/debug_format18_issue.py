#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import re
from docx import Document

def debug_format18_issue():
    """调试格式18处理问题"""
    
    # 创建简单的测试文档
    test_input = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/debug_format18_input.docx"
    
    # 创建测试文档
    doc = Document()
    para = doc.add_paragraph("供应商名称                                    ")
    doc.save(test_input)
    
    try:
        # 分析原始文档
        print("=== 原始文档分析 ===")
        doc = Document(test_input)
        for para_idx, paragraph in enumerate(doc.paragraphs):
            print(f"段落 #{para_idx}: '{paragraph.text}'")
            print(f"段落长度: {len(paragraph.text)}")
            
            # 分析runs
            for run_idx, run in enumerate(paragraph.runs):
                print(f"  Run #{run_idx}: '{run.text}'")
        
        # 测试正则匹配
        print(f"\n=== 正则匹配测试 ===")
        pattern = re.compile(r'^(?P<label>供应商名称)\s*(?P<placeholder>\s{20,})\s*$')
        text = paragraph.text
        match = pattern.search(text)
        
        if match:
            print(f"✅ 正则匹配成功!")
            print(f"  完整匹配: '{match.group(0)}'")
            print(f"  标签: '{match.group('label')}'")
            print(f"  占位符: '{match.group('placeholder')}'")
            print(f"  占位符长度: {len(match.group('placeholder'))}")
        else:
            print(f"❌ 正则匹配失败")
            print(f"文本: '{text}'")
            print(f"文本长度: {len(text)}")
            print(f"是否以供应商名称开头: {text.startswith('供应商名称')}")
            
            # 检查空格
            after_label = text[4:]  # "供应商名称"之后的内容
            print(f"标签后内容: '{after_label}'")
            print(f"标签后内容长度: {len(after_label)}")
            print(f"是否全是空格: {after_label.isspace()}")
            
        # 测试替换逻辑
        if match:
            print(f"\n=== 替换逻辑测试 ===")
            label = match.group('label')
            company_name = "智慧足迹数据科技有限公司"
            new_text = f"{label} {company_name}"
            print(f"新文本: '{new_text}'")
            
            # 检查是否能在run中找到标签
            found_in_run = False
            for run_idx, run in enumerate(paragraph.runs):
                if label in run.text:
                    print(f"✅ 在Run #{run_idx}中找到标签: '{run.text}'")
                    found_in_run = True
                    
                    # 模拟替换
                    old_text = run.text
                    replacement_result = old_text.replace(match.group(0), new_text)
                    print(f"替换结果: '{old_text}' -> '{replacement_result}'")
                else:
                    print(f"❌ Run #{run_idx}中没有标签: '{run.text}'")
            
            if not found_in_run:
                print(f"❌ 在任何run中都没有找到标签 '{label}'")
                
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
    finally:
        if os.path.exists(test_input):
            os.remove(test_input)

if __name__ == "__main__":
    debug_format18_issue()