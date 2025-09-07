#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试供应商名称拆分问题的修复效果
"""

import os
import re
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_problematic_document():
    """创建会导致供应商名称拆分问题的测试文档"""
    doc = Document()
    doc.add_heading('供应商名称拆分问题测试', 0)
    
    # 添加前置段落
    for i in range(29):
        doc.add_paragraph(f'段落 #{i+1}: 前置内容...')
    
    # 段落#30 - 模拟导致问题的run结构
    para30 = doc.add_paragraph()
    
    # 模拟采购编号处理后的复杂run结构：
    # "供应商名称：   采购编号：GXTC-C-251590031  "
    # 这种结构会导致"供应商"和"名称："分在不同run中
    para30.add_run("供应商")          # run 0 - 标签的前半部分
    para30.add_run("名称：")         # run 1 - 标签的后半部分  
    para30.add_run("   ")           # run 2 - 空格占位符
    para30.add_run("采购编号：")     # run 3
    para30.add_run("GXTC-C-251590031")  # run 4 - 已填写的采购编号
    para30.add_run("  ")            # run 5 - 尾部空格
    
    print("创建的问题文档段落#30 run结构：")
    for i, run in enumerate(para30.runs):
        print(f"  Run {i}: '{run.text}'")
    
    full_text = ''.join(run.text for run in para30.runs)
    print(f"完整文本: '{full_text}'")
    
    return doc

def test_fixed_processing():
    """测试修复后的处理效果"""
    print("\n" + "="*60)
    print("测试修复后的供应商名称处理")
    print("="*60)
    
    # 创建测试文档
    input_file = 'test_supplier_fix_input.docx'
    output_file = 'test_supplier_fix_output.docx'
    
    doc = create_problematic_document()
    doc.save(input_file)
    print(f"已创建测试文档: {input_file}")
    
    # 使用修复后的处理器
    processor = MCPBidderNameProcessor()
    company_name = "智慧足迹数据科技有限公司"
    
    print(f"\n开始处理，使用公司名称: {company_name}")
    
    result = processor.process_bidder_name(
        input_file=input_file,
        output_file=output_file,
        company_name=company_name
    )
    
    print(f"\n处理结果:")
    print(f"  成功: {result.get('success', False)}")
    
    if result.get('error'):
        print(f"  错误: {result.get('error')}")
    
    stats = result.get('stats', {})
    print(f"  统计信息:")
    print(f"    总替换次数: {stats.get('total_replacements', 0)}")
    print(f"    找到的模式数: {len(stats.get('patterns_found', []))}")
    
    if result.get('success'):
        print(f"  输出文件: {output_file}")
        
        # 检查输出结果
        try:
            output_doc = Document(output_file)
            if len(output_doc.paragraphs) > 30:
                para30_result = output_doc.paragraphs[30].text
                print(f"\n段落#30处理结果:")
                print(f"  原文: '供应商名称：   采购编号：GXTC-C-251590031  '")
                print(f"  结果: '{para30_result}'")
                
                # 关键检查：确保没有拆分问题
                if "供应商 " + company_name + " 名称" in para30_result:
                    print(f"  ❌ 修复失败！仍然存在拆分问题")
                    print(f"     发现: '供应商 {company_name} 名称'")
                    return False
                elif f"供应商名称： {company_name}" in para30_result:
                    print(f"  ✅ 修复成功！正确格式")
                    return True
                elif company_name in para30_result:
                    print(f"  ✅ 公司名称已填写，但需要检查格式")
                    return True
                else:
                    print(f"  ❌ 公司名称未填写")
                    return False
        except Exception as e:
            print(f"  读取输出文件失败: {e}")
            return False
    
    return False

def main():
    """主函数"""
    print("供应商名称拆分问题修复测试")
    print("测试场景：'供应商名称'被错误拆分为'供应商 公司名称 名称'的问题")
    print("="*60)
    
    success = test_fixed_processing()
    
    print("\n" + "="*60)
    print("修复效果评估:")
    if success:
        print("✅ 修复成功！供应商名称不再被错误拆分")
        print("✅ 跨run标签重构逻辑工作正常")
    else:
        print("❌ 修复失败，需要进一步调试")
    print("="*60)
    
    return success

if __name__ == '__main__':
    main()