#!/usr/bin/env python3
"""
测试公司名称提取功能
模拟web_app传递的复杂公司信息对象
"""

import logging
from run_level_processor import RunLevelProcessor
from docx import Document

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def test_company_extraction():
    """测试公司名称提取功能"""
    logger.info("🧪 开始测试公司名称提取功能")
    
    try:
        # 创建简单的测试文档
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("代表供应商（供应商名称、地址）提交文件")
        test_input = "test_company_extract_input.docx"
        doc.save(test_input)
        
        # 创建处理器
        processor = RunLevelProcessor()
        
        # 模拟web_app传递的复杂公司信息对象（类似截图中看到的）
        complex_company_info = {
            'companyName': '中国联合网络通信有限公司',
            'establishDate': '2000-04-21',
            'legalRepresentative': '陈忠岳',
            'socialCreditCode': '91110000710939135P',
            'registeredCapital': '22539208.432769 万元',
            'companyType': '有限责任公司',
            'fixedPhone': '010-66258899',
            'postalCode': '100033',
            'registeredAddress': '北京市西城区金融大街21号',
            'officeAddress': '北京市西城区金融大街21号',
            'website': 'www.10010.com',
            'employeeCount': '1000人以上',
            'companyDescription': '中国联合网络通信有限公司（简称"中国联通"）在原中国网通和原中国联通的基础上合并组建而成...'
        }
        
        # 测试复杂对象的公司名称提取
        result = processor.process_business_response(
            input_file=test_input,
            output_file="test_company_extract_output.docx",
            company_info=complex_company_info,  # 传递复杂对象
            project_name="政府采购云平台",
            tender_no="64525343"
        )
        
        if result.get('success', False):
            logger.info("✅ 复杂公司信息处理成功！")
            
            # 验证输出文件
            output_doc = Document("test_company_extract_output.docx")
            final_text = output_doc.paragraphs[0].text
            logger.info(f"最终文本: {final_text}")
            
            # 检查是否只使用了公司名称，而不是整个复杂对象
            if "中国联合网络通信有限公司" in final_text:
                if "establishDate" not in final_text and "socialCreditCode" not in final_text:
                    logger.info("🎉 公司名称提取完美！只使用了公司名称，没有包含其他信息")
                    return True
                else:
                    logger.error("❌ 仍然包含了多余的公司信息")
                    return False
            else:
                logger.error("❌ 公司名称提取失败")
                return False
        else:
            logger.error(f"❌ 处理失败: {result}")
            return False
            
    except Exception as e:
        logger.error(f"测试异常: {e}", exc_info=True)
        return False

if __name__ == "__main__":
    success = test_company_extraction()
    if success:
        logger.info("✅ 公司名称提取功能完全正常！")
    else:
        logger.error("❌ 公司名称提取仍有问题")