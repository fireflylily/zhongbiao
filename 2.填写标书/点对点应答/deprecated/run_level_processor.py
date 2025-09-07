#!/usr/bin/env python3
"""
Run-Level Intelligent Replacement Processor
更简单、更可靠的文档处理方法
"""

import logging
import re
import os
import configparser
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class RunLevelProcessor:
    """基于Run级别的智能替换处理器"""
    
    def __init__(self):
        """初始化处理器"""
        self.company_address = "广州市黄埔区科学城中山大学达安基因股份有限公司R&D大楼A栋5层"
        
    def process_bidder_name(self, input_file: str, output_file: str, company_name: str, 
                           project_name: str = None, tender_number: str = None):
        """处理投标人名称填写"""
        try:
            logger.info(f"🚀 开始处理文档: {input_file}")
            
            # 读取文档
            doc = Document(input_file)
            
            # 加载配置数据
            config_project_name = self._load_project_name()
            config_tender_number = self._load_tender_number()
            
            # 优先使用配置文件，如果没有则使用传入参数
            final_project_name = config_project_name if config_project_name else project_name
            final_tender_number = config_tender_number if config_tender_number else tender_number
            
            logger.info(f"最终处理参数 - 公司: {company_name}, 项目: {final_project_name}, 编号: {final_tender_number}")
            
            # 处理每个段落
            processed_paragraphs = 0
            total_replacements = 0
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text and paragraph.text.strip():
                    
                    # 逐项处理每种替换类型
                    replacements_in_para = 0
                    
                    # 1. 处理供应商名称
                    company_count = self._process_company_name_in_paragraph(
                        paragraph, company_name, para_idx
                    )
                    replacements_in_para += company_count
                    
                    # 2. 处理项目名称
                    if final_project_name:
                        project_count = self._process_project_name_in_paragraph(
                            paragraph, final_project_name, para_idx
                        )
                        replacements_in_para += project_count
                    
                    # 3. 处理采购编号
                    if final_tender_number:
                        tender_count = self._process_tender_number_in_paragraph(
                            paragraph, final_tender_number, para_idx
                        )
                        replacements_in_para += tender_count
                    
                    if replacements_in_para > 0:
                        processed_paragraphs += 1
                        total_replacements += replacements_in_para
                        logger.info(f"✅ 段落#{para_idx}: 完成 {replacements_in_para} 项替换")
            
            # 保存文档
            doc.save(output_file)
            
            logger.info(f"🎉 文档处理完成!")
            logger.info(f"  处理段落数: {processed_paragraphs}")
            logger.info(f"  总替换数: {total_replacements}")
            logger.info(f"  输出文件: {output_file}")
            
            return {
                'success': True,
                'processed_paragraphs': processed_paragraphs,
                'total_replacements': total_replacements,
                'output_file': output_file
            }
            
        except Exception as e:
            logger.error(f"文档处理失败: {e}", exc_info=True)
            return {'success': False, 'error': str(e)}
    
    def process_business_response(self, input_file: str, output_file: str, company_info: str = None, 
                                project_name: str = None, tender_no: str = None, 
                                date_text: str = None, **kwargs):
        """
        兼容web服务的商务应答处理方法
        接受所有旧方法的参数，转换后调用process_bidder_name
        """
        try:
            logger.info(f"🔄 使用Run-Level处理器处理商务应答")
            logger.info(f"输入文件: {input_file}")
            logger.info(f"输出文件: {output_file}") 
            
            # 参数兼容处理 - 处理公司信息可能是字典的情况
            if company_info:
                if isinstance(company_info, dict):
                    # 如果是字典，提取公司名称
                    company_name = company_info.get('companyName', company_info.get('name', str(company_info)))
                elif isinstance(company_info, str):
                    # 如果是字符串，直接使用
                    company_name = company_info
                else:
                    # 其他情况，转换为字符串
                    company_name = str(company_info)
            else:
                company_name = kwargs.get('company_name', '未提供公司名称')
                
            logger.info(f"公司信息原始数据: {company_info}")
            logger.info(f"提取的公司名称: {company_name}")
            logger.info(f"项目名称: {project_name}")
            logger.info(f"采购编号: {tender_no}")
            
            # 调用核心处理方法，传递所有参数
            result = self.process_bidder_name(
                input_file=input_file,
                output_file=output_file,
                company_name=company_name,
                project_name=project_name,
                tender_number=tender_no
            )
            
            if result.get('success', False):
                logger.info(f"✅ Run-Level商务应答处理完成")
                # 保持与旧方法兼容的返回格式
                return {
                    'success': True,
                    'output_file': output_file,
                    'processed_paragraphs': result.get('processed_paragraphs', 0),
                    'total_replacements': result.get('total_replacements', 0),
                    'message': 'Run-Level处理器完成文档处理'
                }
            else:
                logger.error(f"❌ Run-Level商务应答处理失败")
                return result
                
        except Exception as e:
            logger.error(f"Run-Level商务应答处理异常: {e}", exc_info=True)
            return {'success': False, 'error': str(e)}
    
    def _process_company_name_in_paragraph(self, paragraph, company_name: str, para_idx: int) -> int:
        """在段落中处理供应商名称相关替换"""
        replacements = 0
        
        try:
            # 定义供应商名称替换模式
            patterns = [
                {
                    'pattern': r'（供应商名称、地址）',
                    'replacement': f'（{company_name}、{self.company_address}）',
                    'desc': '供应商名称+地址'
                },
                {
                    'pattern': r'（供应商名称）',
                    'replacement': f'（{company_name}）',
                    'desc': '供应商名称'
                },
                {
                    'pattern': r'\(供应商名称、地址\)',
                    'replacement': f'（{company_name}、{self.company_address}）',
                    'desc': '供应商名称+地址(英文括号)'
                },
                {
                    'pattern': r'\(供应商名称\)',
                    'replacement': f'（{company_name}）',
                    'desc': '供应商名称(英文括号)'
                }
            ]
            
            for pattern_info in patterns:
                count = self._replace_pattern_in_runs(
                    paragraph, 
                    pattern_info['pattern'], 
                    pattern_info['replacement'],
                    para_idx,
                    pattern_info['desc']
                )
                replacements += count
                
        except Exception as e:
            logger.error(f"段落#{para_idx} 供应商名称处理失败: {e}")
            
        return replacements
    
    def _process_project_name_in_paragraph(self, paragraph, project_name: str, para_idx: int) -> int:
        """在段落中处理项目名称相关替换"""
        replacements = 0
        
        try:
            patterns = [
                {
                    'pattern': r'（项目名称）',
                    'replacement': f'（{project_name}）',
                    'desc': '项目名称'
                },
                {
                    'pattern': r'\(项目名称\)',
                    'replacement': f'（{project_name}）',
                    'desc': '项目名称(英文括号)'
                },
                {
                    'pattern': r'为\s*[\(（][^）)]*[\)）]\s*项目',
                    'replacement': f'为（{project_name}）项目',
                    'desc': '为(xxx)项目格式'
                }
            ]
            
            for pattern_info in patterns:
                count = self._replace_pattern_in_runs(
                    paragraph, 
                    pattern_info['pattern'], 
                    pattern_info['replacement'],
                    para_idx,
                    pattern_info['desc']
                )
                replacements += count
                
        except Exception as e:
            logger.error(f"段落#{para_idx} 项目名称处理失败: {e}")
            
        return replacements
    
    def _process_tender_number_in_paragraph(self, paragraph, tender_number: str, para_idx: int) -> int:
        """在段落中处理采购编号相关替换"""
        replacements = 0
        
        try:
            patterns = [
                {
                    'pattern': r'（采购编号）',
                    'replacement': f'（{tender_number}）',
                    'desc': '采购编号'
                },
                {
                    'pattern': r'\(采购编号\)',
                    'replacement': f'（{tender_number}）',
                    'desc': '采购编号(英文括号)'
                },
                {
                    'pattern': r'（招标编号）',
                    'replacement': f'（{tender_number}）',
                    'desc': '招标编号'
                },
                {
                    'pattern': r'（项目编号）',
                    'replacement': f'（{tender_number}）',
                    'desc': '项目编号'
                }
            ]
            
            for pattern_info in patterns:
                count = self._replace_pattern_in_runs(
                    paragraph, 
                    pattern_info['pattern'], 
                    pattern_info['replacement'],
                    para_idx,
                    pattern_info['desc']
                )
                replacements += count
                
        except Exception as e:
            logger.error(f"段落#{para_idx} 采购编号处理失败: {e}")
            
        return replacements
    
    def _replace_pattern_in_runs(self, paragraph, pattern: str, replacement: str, 
                                para_idx: int, desc: str) -> int:
        """在段落的runs中查找并替换匹配的模式"""
        replacements = 0
        
        try:
            # 编译正则表达式
            regex = re.compile(pattern)
            paragraph_text = paragraph.text
            
            # 查找匹配
            matches = list(regex.finditer(paragraph_text))
            if not matches:
                return 0
                
            logger.info(f"  段落#{para_idx} 发现{len(matches)}个匹配: {desc}")
            
            # 从后往前处理，避免位置偏移问题
            for match in reversed(matches):
                success = self._replace_text_in_runs(
                    paragraph, 
                    match.group(0),  # 原文本
                    replacement,     # 替换文本
                    match.start(),   # 开始位置
                    match.end()      # 结束位置
                )
                
                if success:
                    replacements += 1
                    logger.info(f"    ✅ 替换成功: '{match.group(0)}' -> '{replacement}'")
                else:
                    logger.error(f"    ❌ 替换失败: '{match.group(0)}'")
                    
        except Exception as e:
            logger.error(f"模式替换失败 {desc}: {e}")
            
        return replacements
    
    def _replace_text_in_runs(self, paragraph, old_text: str, new_text: str, 
                             start_pos: int, end_pos: int) -> bool:
        """在runs中精确替换文本，保持格式"""
        try:
            # 方法1: 尝试单个run内替换（最简单情况）
            for run in paragraph.runs:
                if old_text in run.text:
                    # 保存原始格式
                    original_format = self._extract_run_format(run)
                    
                    # 直接替换
                    run.text = run.text.replace(old_text, new_text)
                    
                    # 应用格式（通常不需要，因为run格式会保持）
                    self._apply_format_to_run(run, original_format)
                    
                    logger.debug(f"    单run替换成功: '{old_text}' -> '{new_text}'")
                    return True
            
            # 方法2: 跨run替换（复杂情况）
            return self._replace_across_runs(paragraph, old_text, new_text, start_pos, end_pos)
                    
        except Exception as e:
            logger.error(f"run内文本替换失败: {e}")
            return False
    
    def _replace_across_runs(self, paragraph, old_text: str, new_text: str, 
                           start_pos: int, end_pos: int) -> bool:
        """处理跨run的文本替换，更好地保持格式"""
        try:
            # 找到涉及的runs
            current_pos = 0
            involved_runs = []
            
            for i, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                
                # 检查这个run是否与目标区域有重叠
                if run_end > start_pos and run_start < end_pos:
                    # 计算在这个run中的相对位置
                    relative_start = max(0, start_pos - run_start)
                    relative_end = min(len(run.text), end_pos - run_start)
                    
                    involved_runs.append({
                        'run': run,
                        'run_index': i,
                        'relative_start': relative_start,
                        'relative_end': relative_end,
                        'format': self._extract_run_format(run),
                        'text_portion': run.text[relative_start:relative_end]  # 被替换的文本部分
                    })
                
                current_pos = run_end
            
            if not involved_runs:
                logger.error("未找到涉及的runs")
                return False
            
            # 处理简单情况：只涉及一个run的部分内容
            if len(involved_runs) == 1:
                run_info = involved_runs[0]
                run = run_info['run']
                relative_start = run_info['relative_start']
                relative_end = run_info['relative_end']
                original_format = run_info['format']
                
                # 替换run中的部分文本，保持格式
                old_run_text = run.text
                new_run_text = (old_run_text[:relative_start] + 
                              new_text + 
                              old_run_text[relative_end:])
                
                run.text = new_run_text
                
                # 确保格式得到保持（特别是斜体）
                self._apply_format_to_run(run, original_format)
                
                logger.debug(f"    跨run单个替换成功，保持格式: {original_format}")
                return True
            
            # 复杂情况：涉及多个runs，采用更智能的格式保持策略
            logger.info(f"    涉及{len(involved_runs)}个runs的跨run替换")
            
            # 分析涉及runs的格式模式，找到主要格式
            dominant_format = self._analyze_dominant_format(involved_runs)
            
            # 创建优化的替换策略
            return self._perform_multi_run_replacement(paragraph, involved_runs, new_text, dominant_format)
            
        except Exception as e:
            logger.error(f"跨run替换失败: {e}")
            return False
    
    def _extract_run_format(self, run):
        """提取run的格式信息"""
        try:
            format_info = {
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': None
            }
            
            # 提取颜色信息
            if run.font.color.rgb:
                format_info['color'] = run.font.color.rgb
                
            return format_info
            
        except Exception as e:
            logger.debug(f"格式提取失败: {e}")
            return {}
    
    def _apply_format_to_run(self, run, format_info):
        """将格式应用到run"""
        try:
            if not format_info:
                return
                
            if format_info.get('font_name'):
                run.font.name = format_info['font_name']
            if format_info.get('font_size'):
                run.font.size = format_info['font_size']
            if format_info.get('bold') is not None:
                run.font.bold = format_info['bold']
            if format_info.get('italic') is not None:
                run.font.italic = format_info['italic']
            if format_info.get('underline') is not None:
                run.font.underline = format_info['underline']
            if format_info.get('color'):
                run.font.color.rgb = format_info['color']
                
        except Exception as e:
            logger.debug(f"格式应用失败: {e}")
    
    def _analyze_dominant_format(self, involved_runs):
        """分析涉及runs的主要格式特征"""
        try:
            formats = [run_info['format'] for run_info in involved_runs if run_info['format']]
            
            if not formats:
                return {}
            
            # 分析关键格式属性的出现频率
            dominant_format = {
                'font_name': self._get_most_common_value([f.get('font_name') for f in formats]),
                'font_size': self._get_most_common_value([f.get('font_size') for f in formats]),
                'bold': self._get_most_common_bool([f.get('bold') for f in formats]),
                'italic': self._get_most_common_bool([f.get('italic') for f in formats]),
                'underline': self._get_most_common_value([f.get('underline') for f in formats]),
                'color': self._get_most_common_value([f.get('color') for f in formats])
            }
            
            logger.debug(f"    主要格式: {dominant_format}")
            return dominant_format
            
        except Exception as e:
            logger.error(f"格式分析失败: {e}")
            return {}
    
    def _get_most_common_value(self, values):
        """获取最常见的值"""
        values = [v for v in values if v is not None]
        if not values:
            return None
        
        # 简单的频率统计
        from collections import Counter
        counter = Counter(values)
        most_common = counter.most_common(1)
        return most_common[0][0] if most_common else None
    
    def _get_most_common_bool(self, values):
        """获取最常见的布尔值"""
        values = [v for v in values if v is not None]
        if not values:
            return None
        
        true_count = sum(1 for v in values if v)
        false_count = len(values) - true_count
        
        return true_count > false_count
    
    def _perform_multi_run_replacement(self, paragraph, involved_runs, new_text, dominant_format):
        """执行多run替换，精确保持格式范围"""
        try:
            # 分析替换的范围和上下文
            first_run_info = involved_runs[0]
            last_run_info = involved_runs[-1]
            
            prefix_text = first_run_info['run'].text[:first_run_info['relative_start']]
            suffix_text = last_run_info['run'].text[last_run_info['relative_end']:]
            
            # 查找哪个run包含了最多的原始占位符文本
            placeholder_format = self._find_placeholder_format(involved_runs)
            
            logger.debug(f"    占位符格式: {placeholder_format}")
            logger.debug(f"    前缀: '{prefix_text}', 后缀: '{suffix_text}'")
            
            # 使用精确的三段式替换，保持精准的格式控制
            return self._create_precise_replacement(paragraph, involved_runs, 
                                                 prefix_text, new_text, suffix_text, 
                                                 placeholder_format)
            
        except Exception as e:
            logger.error(f"多run替换执行失败: {e}")
            return False
    
    def _find_placeholder_format(self, involved_runs):
        """找到占位符的主要格式特征"""
        try:
            # 分析每个run中占位符相关的文本长度
            max_placeholder_content = 0
            placeholder_format = {}
            
            for run_info in involved_runs:
                text_portion = run_info.get('text_portion', '')
                # 如果这个文本包含占位符关键字
                if any(keyword in text_portion for keyword in ['项目名称', '采购编号', '供应商名称']):
                    if len(text_portion) > max_placeholder_content:
                        max_placeholder_content = len(text_portion)
                        placeholder_format = run_info['format']
            
            # 如果没有找到特定的占位符格式，使用涉及的runs中最特殊的格式
            if not placeholder_format:
                for run_info in involved_runs:
                    fmt = run_info['format']
                    if fmt.get('italic') or fmt.get('bold') or fmt.get('underline'):
                        placeholder_format = fmt
                        break
            
            return placeholder_format if placeholder_format else involved_runs[0]['format']
            
        except Exception as e:
            logger.error(f"占位符格式查找失败: {e}")
            return {}
    
    def _create_precise_replacement(self, paragraph, involved_runs, prefix_text, new_text, suffix_text, placeholder_format):
        """创建精确的三段式替换：前缀(普通) + 新文本(特殊格式) + 后缀(普通)"""
        try:
            # 获取第一个run的普通格式（去掉特殊效果）
            first_run = involved_runs[0]['run'] 
            normal_format = self._extract_run_format(first_run)
            
            # 创建普通格式（去掉特殊效果）
            normal_format_clean = normal_format.copy()
            normal_format_clean['italic'] = None  # 清除斜体
            normal_format_clean['bold'] = None    # 清除粗体
            normal_format_clean['underline'] = None  # 清除下划线
            
            # 清空所有涉及的runs
            for run_info in involved_runs:
                run_info['run'].text = ""
            
            # 在第一个run中重建内容
            first_run = involved_runs[0]['run']
            
            # 如果有前缀文本，将其设为普通格式
            if prefix_text:
                first_run.text = prefix_text
                self._apply_format_to_run(first_run, normal_format_clean)
            
            # 为新文本创建新run（如果需要特殊格式）
            if placeholder_format and (placeholder_format.get('italic') or 
                                     placeholder_format.get('bold') or 
                                     placeholder_format.get('underline')):
                # 创建带有特殊格式的新run
                new_run = paragraph.add_run(new_text)
                self._apply_format_to_run(new_run, placeholder_format)
                logger.debug(f"    创建特殊格式run: {new_text} - {placeholder_format}")
            else:
                # 直接追加到第一个run
                first_run.text += new_text
            
            # 如果有后缀文本，创建普通格式的run
            if suffix_text:
                suffix_run = paragraph.add_run(suffix_text)
                self._apply_format_to_run(suffix_run, normal_format_clean)
            
            logger.debug(f"    精确替换完成：前缀(普通) + '{new_text}'(特殊) + 后缀(普通)")
            return True
            
        except Exception as e:
            logger.error(f"精确替换创建失败: {e}")
            return False
    
    def _load_project_name(self) -> str:
        """从配置文件加载项目名称"""
        try:
            config_file = "tender_config.ini"
            if os.path.exists(config_file):
                config = configparser.ConfigParser()
                config.read(config_file, encoding='utf-8')
                if 'PROJECT_INFO' in config and 'project_name' in config['PROJECT_INFO']:
                    project_name = config['PROJECT_INFO']['project_name']
                    if project_name and project_name != '未提供':
                        return project_name
            return None
        except Exception as e:
            logger.warning(f"无法加载项目名称: {e}")
            return None
    
    def _load_tender_number(self) -> str:
        """从配置文件加载采购编号"""
        try:
            config_file = "tender_config.ini"
            if os.path.exists(config_file):
                config = configparser.ConfigParser()
                config.read(config_file, encoding='utf-8')
                if 'PROJECT_INFO' in config and 'project_number' in config['PROJECT_INFO']:
                    tender_number = config['PROJECT_INFO']['project_number']
                    if tender_number and tender_number != '未提供':
                        return tender_number
            return None
        except Exception as e:
            logger.warning(f"无法加载采购编号: {e}")
            return None