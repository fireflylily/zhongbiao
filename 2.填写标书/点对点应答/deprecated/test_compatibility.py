#!/usr/bin/env python3
"""
快速测试Run-Level处理器与web_app兼容性
"""

import logging
from run_level_processor import RunLevelProcessor
from docx import Document

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def test_compatibility():
    """测试Run-Level处理器兼容性"""
    logger.info("🧪 开始兼容性测试")
    
    try:
        # 创建简单的测试文档
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("根据贵方为（项目名称）项目采购（采购编号），代表（供应商名称、地址）提交文件")
        test_input = "test_compat_input.docx"
        doc.save(test_input)
        
        # 创建处理器
        processor = RunLevelProcessor()
        
        # 测试process_business_response方法（模拟web_app的调用）
        result = processor.process_business_response(
            input_file=test_input,
            output_file="test_compat_output.docx",
            company_info="智慧足迹数据科技有限公司",  # 注意：使用company_info而非company_name
            project_name="政府采购云平台",
            tender_no="64525343",
            date_text="2025-09-06"
        )
        
        if result.get('success', False):
            logger.info("✅ 兼容性测试通过！")
            logger.info(f"处理统计: {result}")
            
            # 验证输出文件
            output_doc = Document("test_compat_output.docx")
            final_text = output_doc.paragraphs[0].text
            logger.info(f"最终文本: {final_text}")
            
            # 检查替换是否正确
            checks = [
                ("智慧足迹数据科技有限公司" in final_text, "公司名称替换"),
                ("政府采购云平台" in final_text, "项目名称替换"), 
                ("64525343" in final_text, "采购编号替换")
            ]
            
            success_count = sum(1 for check, desc in checks if check)
            logger.info(f"替换检查: {success_count}/{len(checks)} 项通过")
            
            if success_count == len(checks):
                logger.info("🎉 所有替换检查通过，兼容性完美！")
                return True
            else:
                logger.warning("⚠️ 部分替换检查失败")
                return False
        else:
            logger.error(f"❌ 兼容性测试失败: {result}")
            return False
            
    except Exception as e:
        logger.error(f"兼容性测试异常: {e}", exc_info=True)
        return False

if __name__ == "__main__":
    success = test_compatibility()
    if success:
        logger.info("✅ Run-Level处理器现在完全兼容web服务！")
    else:
        logger.error("❌ 仍有兼容性问题需要解决")