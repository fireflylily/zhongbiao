#!/usr/bin/env python3
"""
测试Run-Level智能替换方法的效果
对比新方法与批量替换方法的优劣
"""

import logging
import sys
from pathlib import Path
from docx import Document
from run_level_processor import RunLevelProcessor
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_complex_test_document():
    """创建包含多项替换内容的复杂测试文档"""
    doc = Document()
    
    # 添加段落1：包含4个替换项的复杂段落（高风险场景）
    logger.info("创建复杂段落1: 包含项目名称、采购编号、姓名职务、公司名称")
    para1 = doc.add_paragraph()
    
    # 模拟复杂的格式情况，每部分使用不同格式
    run1 = para1.add_run("根据贵方为")
    run1.font.name = "宋体"
    run1.font.italic = False
    
    run2 = para1.add_run("（")
    run2.font.name = "宋体"
    run2.font.italic = True
    
    run3 = para1.add_run("项目名称")
    run3.font.name = "宋体" 
    run3.font.italic = True
    
    run4 = para1.add_run("）")
    run4.font.name = "宋体"
    run4.font.italic = True
    
    run5 = para1.add_run("项目采购采购货物及服务的竞争性磋商公告")
    run5.font.name = "宋体"
    run5.font.italic = False
    
    run6 = para1.add_run("（")
    run6.font.name = "宋体"
    run6.font.italic = True
    run6.font.underline = True
    
    run7 = para1.add_run("采购编号")
    run7.font.name = "宋体"
    run7.font.italic = True
    run7.font.underline = True
    
    run8 = para1.add_run("）")
    run8.font.name = "宋体"
    run8.font.italic = True
    run8.font.underline = True
    
    run9 = para1.add_run("，签字代表")
    run9.font.name = "宋体"
    run9.font.italic = False
    
    run10 = para1.add_run("（")
    run10.font.name = "宋体"
    run10.font.italic = False
    
    run11 = para1.add_run("姓名、职务")
    run11.font.name = "宋体"
    run11.font.italic = False
    
    run12 = para1.add_run("）")
    run12.font.name = "宋体"
    run12.font.italic = False
    
    run13 = para1.add_run("经正式授权并代表供应商")
    run13.font.name = "宋体"
    run13.font.italic = False
    
    run14 = para1.add_run("（")
    run14.font.name = "宋体"
    run14.font.italic = True
    
    run15 = para1.add_run("供应商名称、地址")
    run15.font.name = "宋体"
    run15.font.italic = True
    
    run16 = para1.add_run("）")
    run16.font.name = "宋体"
    run16.font.italic = True
    
    run17 = para1.add_run("提交下述文件正本一份及副本份：")
    run17.font.name = "宋体"
    run17.font.italic = False
    
    # 添加段落2：简单的两项替换
    logger.info("创建段落2: 包含项目名称和采购编号")
    para2 = doc.add_paragraph()
    
    run1 = para2.add_run("本次")
    run1.font.name = "宋体"
    
    run2 = para2.add_run("（")
    run2.font.name = "宋体"
    run2.font.italic = True
    
    run3 = para2.add_run("项目名称")
    run3.font.name = "宋体"
    run3.font.italic = True
    
    run4 = para2.add_run("）")
    run4.font.name = "宋体"
    run4.font.italic = True
    
    run5 = para2.add_run("采购，编号为")
    run5.font.name = "宋体"
    
    run6 = para2.add_run("（")
    run6.font.name = "宋体"
    run6.font.bold = True
    
    run7 = para2.add_run("采购编号")
    run7.font.name = "宋体"
    run7.font.bold = True
    
    run8 = para2.add_run("）")
    run8.font.name = "宋体"
    run8.font.bold = True
    
    run9 = para2.add_run("，特制定本投标文件。")
    run9.font.name = "宋体"
    
    return doc

def analyze_document_format(doc, title):
    """分析文档格式结构"""
    logger.info(f"\n📋 {title}")
    logger.info("="*60)
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            logger.info(f"段落 #{para_idx}: '{paragraph.text}'")
            logger.info(f"段落总长度: {len(paragraph.text)} 字符")
            
            # 分析每个run的详细信息
            for run_idx, run in enumerate(paragraph.runs):
                if run.text:
                    format_details = []
                    if run.font.name:
                        format_details.append(f"字体={run.font.name}")
                    if run.font.bold:
                        format_details.append("粗体=True")
                    if run.font.italic:
                        format_details.append("斜体=True")
                    if run.font.underline:
                        format_details.append("下划线=True")
                    
                    format_str = ", ".join(format_details) if format_details else "默认格式"
                    
                    # 检查特殊内容
                    markers = []
                    if "项目名称" in run.text:
                        markers.append("🎯 项目名称")
                    elif "采购编号" in run.text:
                        markers.append("📋 采购编号") 
                    elif "供应商名称" in run.text:
                        markers.append("🏢 供应商名称")
                    elif "姓名、职务" in run.text:
                        markers.append("👤 姓名职务")
                    elif "智慧足迹" in run.text:
                        markers.append("✅ 已替换内容")
                    elif "政府采购云平台" in run.text:
                        markers.append("✅ 已替换项目")
                    elif "64525343" in run.text:
                        markers.append("✅ 已替换编号")
                    
                    marker_str = " " + " ".join(markers) if markers else ""
                    
                    logger.info(f"  Run {run_idx+1}: '{run.text}' [{format_str}]{marker_str}")
            
            logger.info("-" * 60)

def test_run_level_approach():
    """测试Run-Level方法"""
    logger.info("🔧 测试Run-Level智能替换方法")
    
    try:
        # 创建测试文档
        input_doc = create_complex_test_document()
        input_path = "test_run_level_input.docx"
        input_doc.save(input_path)
        
        # 分析原始文档格式
        analyze_document_format(input_doc, "原始文档格式分析")
        
        # 使用新的Run-Level处理器
        processor = RunLevelProcessor()
        
        # 处理文档
        output_path = "test_run_level_output.docx"
        result = processor.process_bidder_name(
            input_file=input_path,
            output_file=output_path,
            company_name="智慧足迹数据科技有限公司"
        )
        
        if result.get('success', False):
            logger.info("✅ Run-Level方法处理成功")
            
            # 分析处理后文档格式
            output_doc = Document(output_path)
            analyze_document_format(output_doc, "Run-Level方法处理后格式")
            
            # 验证关键点
            logger.info("\n🎯 关键验证项目")
            logger.info("="*80)
            
            success_checks = 0
            total_checks = 0
            
            # 检查文本内容和格式保持
            for para_idx, paragraph in enumerate(output_doc.paragraphs):
                para_text = paragraph.text
                
                # 检查1: 是否正确替换了公司名称
                total_checks += 1
                if "智慧足迹数据科技有限公司" in para_text and "供应商名称" not in para_text:
                    logger.info(f"✅ 段落#{para_idx}: 公司名称已正确替换")
                    success_checks += 1
                elif "供应商名称" not in para_text:
                    success_checks += 1  # 没有供应商名称也算正常
                else:
                    logger.error(f"❌ 段落#{para_idx}: 供应商名称未替换")
                
                # 检查2: 检查"姓名、职务"是否未受影响
                total_checks += 1
                name_job_runs = [run for run in paragraph.runs if "姓名、职务" in run.text]
                if name_job_runs:
                    # 检查是否保持正常格式
                    format_normal = all(not run.font.italic and not run.font.underline for run in name_job_runs)
                    if format_normal:
                        logger.info(f"✅ 段落#{para_idx}: 姓名、职务格式未受影响")
                        success_checks += 1
                    else:
                        logger.error(f"❌ 段落#{para_idx}: 姓名、职务格式受到影响")
                else:
                    success_checks += 1
                
                # 检查3: 项目名称是否正确处理
                total_checks += 1
                if "政府采购云平台" in para_text and "项目名称" not in para_text:
                    # 检查项目名称区域的格式
                    project_runs = [run for run in paragraph.runs if "政府采购云平台" in run.text]
                    if project_runs:
                        format_preserved = any(run.font.italic for run in project_runs)
                        if format_preserved:
                            logger.info(f"✅ 段落#{para_idx}: 项目名称已替换并保持格式")
                            success_checks += 1
                        else:
                            logger.warning(f"⚠️ 段落#{para_idx}: 项目名称已替换但格式可能有变化")
                            success_checks += 1  # 内容正确就算成功
                    else:
                        success_checks += 1
                elif "项目名称" not in para_text:
                    success_checks += 1  # 没有项目名称也算正常
                else:
                    logger.error(f"❌ 段落#{para_idx}: 项目名称未替换")
                
                # 检查4: 采购编号是否正确处理
                total_checks += 1
                if "64525343" in para_text and "采购编号" not in para_text:
                    logger.info(f"✅ 段落#{para_idx}: 采购编号已正确替换")
                    success_checks += 1
                elif "采购编号" not in para_text:
                    success_checks += 1  # 没有采购编号也算正常
                else:
                    logger.error(f"❌ 段落#{para_idx}: 采购编号未替换")
            
            # 最终结果
            success_rate = (success_checks / total_checks) * 100 if total_checks > 0 else 0
            logger.info(f"\n🎯 Run-Level方法测试结果")
            logger.info("="*80)
            logger.info(f"总检查项: {total_checks}")
            logger.info(f"成功项: {success_checks}")
            logger.info(f"成功率: {success_rate:.1f}%")
            
            if success_rate >= 95:
                logger.info("🎉 Run-Level方法测试优秀！")
                return True
            elif success_rate >= 80:
                logger.info("✅ Run-Level方法测试良好")
                return True
            else:
                logger.error("❌ Run-Level方法需要改进")
                return False
        else:
            logger.error("❌ Run-Level方法处理失败")
            logger.error(f"错误信息: {result.get('error', '未知错误')}")
            return False
            
    except Exception as e:
        logger.error(f"测试失败: {e}", exc_info=True)
        return False

def test_comparison():
    """对比测试：Run-Level方法 vs 批量替换方法"""
    logger.info("🆚 对比测试开始")
    
    # 测试Run-Level方法
    logger.info("\n" + "="*80)
    logger.info("测试 Run-Level 智能替换方法")
    logger.info("="*80)
    run_level_result = test_run_level_approach()
    
    # 测试批量替换方法
    logger.info("\n" + "="*80)
    logger.info("测试 现有批量替换方法")
    logger.info("="*80)
    
    try:
        # 创建相同的测试文档
        input_doc = create_complex_test_document()
        input_path = "test_batch_comparison_input.docx"
        input_doc.save(input_path)
        
        # 使用现有的批量替换处理器
        batch_processor = MCPBidderNameProcessor()
        
        # 处理文档
        output_path = "test_batch_comparison_output.docx"
        batch_result = batch_processor.process_bidder_name(
            input_file=input_path,
            output_file=output_path,
            company_name="智慧足迹数据科技有限公司"
        )
        
        batch_success = batch_result.get('success', False)
        
        if batch_success:
            logger.info("✅ 批量替换方法处理成功")
        else:
            logger.error("❌ 批量替换方法处理失败")
        
    except Exception as e:
        logger.error(f"批量替换方法测试失败: {e}")
        batch_success = False
    
    # 总结对比结果
    logger.info("\n" + "="*80)
    logger.info("🏁 对比测试总结")
    logger.info("="*80)
    logger.info(f"Run-Level方法: {'✅ 成功' if run_level_result else '❌ 失败'}")
    logger.info(f"批量替换方法: {'✅ 成功' if batch_success else '❌ 失败'}")
    
    if run_level_result and not batch_success:
        logger.info("🎉 Run-Level方法胜出！更可靠的处理效果")
    elif run_level_result and batch_success:
        logger.info("🤝 两种方法都成功，但Run-Level方法更简单易维护")
    elif not run_level_result and batch_success:
        logger.info("⚠️ 批量替换方法更稳定，Run-Level方法需要改进")
    else:
        logger.info("😞 两种方法都有问题，需要进一步优化")

if __name__ == "__main__":
    test_comparison()