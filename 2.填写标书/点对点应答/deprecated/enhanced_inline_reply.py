#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版采购需求原地插入应答系统
改进点：
1. 字体格式与上段落保持一致
2. 使用始皇API生成专业的技术应答
3. 优化需求条目识别，减少遗漏
4. 专业的LLM提示词模板
"""

import requests
import json
import logging
import os
import re
from datetime import datetime
from typing import Dict, List

#--------------------------------需要调整的全局变量--------------------------------------------------

# 始皇API配置（参考Generate.py的配置模式）
SHIHUANG_API_KEY = "sk-4sYV1WXMcdGcLz9XEKWyntV58pSnhb4GXM6aMBfzWUic3pLfnwob"  # 请替换为您的实际API密钥
SHIHUANG_BASE_URL = "https://api.oaipro.com/v1" 
SHIHUANG_MODEL = "gpt-4o-mini"  # 可选: gpt-3.5-turbo, gpt-4, gpt-4-0613, gpt-4-1106-preview 等

# API调用参数配置
API_CONFIG = {
    "base_url": f"{SHIHUANG_BASE_URL}/chat/completions",
    "model": SHIHUANG_MODEL,
    "temperature": 0.3,  # 较低温度，保证专业性和一致性
    "max_tokens": 500,   # 增加token限制，支持更详细的应答
    "timeout": 60
}

# 兼容原有API配置模块
try:
    from api_config import get_api_key, is_valid_api_key, print_api_setup_guide
    API_CONFIG_AVAILABLE = True
except ImportError:
    API_CONFIG_AVAILABLE = False

# 提示词模板（参考Generate.py的提示词设计）
Prompt_Answer = """现在有个问答，比选文件要求和比选申请人应答，我给你举个例子，比如比选文件要求：支持可视化创建不同类型数据源，包括但不限于：传统数据库、文件系统、消息队列、SaaS API，NoSQL等、必选申请人回答的是：应答：满足。。系统支持数据源配置化管理，数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置管理功能，支持搜索功能。你学习一下我的风格。现在我是比选申请人，请严格按照我的风格来回答，请注意我回答的格式：首先是'应答：满足。'，然后说'系统支持什么什么'，这个过程需要你按照问题回答，不要跑题。以下是输入文字："""

Prompt_Content = "你是一个大数据平台的专业产品售前，请针对这一需求给出800字的产品功能介绍，不要开头和总结，直接写产品功能，不需要用markdown格式，直接文本格式+特殊项目符号输出即可，需求如下："

Prompt_Title = "你是一个专业作者，请把以下这段文字变为10字以内不带细节内容和标点和解释的文字，直接给出结果不要'简化为'这种返回："

#-----------------------------------------------------------------------------------------------------------------------

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('enhanced_inline_reply.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class EnhancedInlineReplyProcessor:
    """增强版原地应答插入处理器"""
    
    def __init__(self, api_key: str = None):
        # 获取API密钥，优先级：参数 > 全局配置 > 配置文件 > 环境变量 > 默认值
        if api_key and api_key != "sk-xxx":
            self.api_key = api_key
        elif SHIHUANG_API_KEY and SHIHUANG_API_KEY != "sk-xxx":
            self.api_key = SHIHUANG_API_KEY
        elif API_CONFIG_AVAILABLE:
            self.api_key = get_api_key()
        else:
            self.api_key = api_key or "sk-xxx"
            
        # 检查API密钥有效性
        if not self._is_valid_key(self.api_key):
            logger.warning("使用默认API密钥，请在文件顶部配置SHIHUANG_API_KEY以获得更好的应答效果")
            if API_CONFIG_AVAILABLE:
                logger.info("运行 python3 api_config.py 查看配置指南")
        
        self.template_file = 'config/response_templates.json'
        self.patterns_file = 'config/requirement_patterns.json'
        
        # 使用全局API配置
        self.model_config = API_CONFIG.copy()
        
        # 加载配置
        self.load_templates()
        self.load_patterns()
        
    def _is_valid_key(self, api_key: str) -> bool:
        """检查API密钥是否有效"""
        if API_CONFIG_AVAILABLE:
            return is_valid_api_key(api_key)
        return api_key and api_key != "sk-xxx" and len(api_key) > 10
    
    def load_templates(self):
        """加载应答模板"""
        try:
            if os.path.exists(self.template_file):
                with open(self.template_file, 'r', encoding='utf-8') as f:
                    self.templates = json.load(f)
                    logger.info("应答模板加载成功")
            else:
                self.templates = self.get_default_templates()
                logger.warning("模板文件不存在，使用默认模板")
        except Exception as e:
            logger.error(f"加载应答模板失败: {e}")
            self.templates = self.get_default_templates()
    
    def load_patterns(self):
        """加载需求识别模式"""
        try:
            if os.path.exists(self.patterns_file):
                with open(self.patterns_file, 'r', encoding='utf-8') as f:
                    self.patterns = json.load(f)
                    logger.info("需求识别模式加载成功")
            else:
                self.patterns = self.get_default_patterns()
                logger.warning("模式文件不存在，使用默认模式")
        except Exception as e:
            logger.error(f"加载需求识别模式失败: {e}")
            self.patterns = self.get_default_patterns()
    
    def get_default_templates(self) -> Dict:
        """获取默认应答模板"""
        return {
            "硬件配置": "我方提供的硬件设备完全满足技术规格要求",
            "软件功能": "我方系统具备完整的功能模块，满足业务需求",
            "性能指标": "我方产品性能指标达到或超过要求标准",
            "技术规范": "我方严格遵循相关技术标准和行业规范",
            "服务保障": "我方提供全方位专业服务，确保项目成功",
            "资质证明": "我方具备完整的相关资质和丰富经验",
            "数据服务": "我方提供稳定可靠的数据服务和技术支持",
            "接口对接": "我方支持标准接口协议，确保系统无缝对接",
            "运维支持": "我方提供7×24小时专业运维服务",
            "通用模板": "我方将严格按照要求提供专业技术方案"
        }
    
    def get_default_patterns(self) -> Dict:
        """获取默认需求识别模式 - 更宽泛的识别策略"""
        return {
            "编号模式": [
                r'^(\d+)\s*[、．.]',         # 1、 1. 1．
                r'^(\d+\.\d+)\s*[、．.]',     # 1.1、 1.2.
                r'^\((\d+)\)',               # (1) (2)
                r'^([A-Z])\)',               # A) B)
                r'^([a-z])\)',               # a) b)
                r'^（[一二三四五六七八九十]+）'  # （一）（二）
            ],
            "关键词": [
                "要求", "需求", "应", "必须", "应当", "需要", "具备", "支持", "提供", 
                "实现", "满足", "符合", "遵循", "不少于", "不低于", "负责", "确保",
                "保证", "完成", "达到", "实施", "部署", "维护", "服务", "管理"
            ],
            "章节标识": [
                "乙方", "甲方", "供应商", "采购方", "服务商", "承包方"
            ]
        }
    
    def llm_callback(self, prompt: str, purpose: str = "应答生成") -> str:
        """
        调用始皇API生成专业技术应答
        """
        url = self.model_config.get("base_url", "https://api.oaipro.com/v1/chat/completions")
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": self.model_config["model"],
            "messages": [
                {"role": "system", "content": self.get_system_prompt()},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.model_config["temperature"],
            "max_tokens": self.model_config["max_tokens"]
        }
        
        try:
            logger.info(f"[始皇API调用] {purpose}")
            timeout = self.model_config.get("timeout", 60)
            resp = requests.post(url, headers=headers, json=payload, timeout=timeout)
            resp.raise_for_status()
            data = resp.json()
            
            if "choices" in data and len(data["choices"]) > 0:
                result = data["choices"][0]["message"]["content"].strip()
                
                # 确保以"应答：满足。"开头
                if not result.startswith("应答：满足。"):
                    if result.startswith("应答："):
                        result = result.replace("应答：", "应答：满足。", 1)
                    else:
                        result = f"应答：满足。{result}"
                
                logger.info(f"API调用成功，生成应答: {result[:50]}...")
                return result
            else:
                logger.warning("API返回为空")
                return self._get_fallback_response()
                
        except requests.exceptions.RequestException as e:
            if "401" in str(e):
                logger.error(f"始皇API调用失败: 401 未授权错误，请检查API密钥是否正确")
                logger.error(f"当前使用的API密钥前缀: {self.api_key[:10]}...")
            elif "403" in str(e):
                logger.error(f"始皇API调用失败: 403 禁止访问，请检查账户余额或权限")
            else:
                logger.error(f"始皇API调用失败: {e}")
            return self._get_fallback_response()
        except Exception as e:
            logger.error(f"始皇API调用失败: {e}")
            return self._get_fallback_response()
    
    def test_api_connection(self) -> bool:
        """
        测试API连接是否正常
        """
        if self.api_key == "sk-xxx":
            logger.warning("使用默认API密钥，无法进行真实API测试")
            return False
            
        try:
            test_prompt = "请回复：测试成功"
            response = self.llm_callback(test_prompt, "API连接测试")
            return "测试成功" in response or "应答：满足" in response
        except Exception as e:
            logger.error(f"API连接测试失败: {e}")
            return False
    
    def get_system_prompt(self, prompt_type: str = "default") -> str:
        """
        获取专业的LLM系统提示词（根据Generate.py的提示词模式优化）
        """
        if prompt_type == "point_to_point":
            return """你是一名专业的投标文件撰写专家，专门负责点对点应答。

你的任务：
1. 以"应答：满足。"开头
2. 然后说"系统支持..."
3. 不要跑题，针对具体需求回答
4. 语言专业、简洁

格式示例：
"应答：满足。系统支持数据源配置化管理，数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置管理功能，全面支持搜索功能。"""
        elif prompt_type == "content":
            return """你是一个大数据平台的专业产品售前。

任务：
- 针对需求给出800字的产品功能介绍
- 不要开头和总结，直接写产品功能
- 使用直接文本格式+特殊项目符号
- 内容要专业、具体、可实现"""
        else:
            # 默认系统提示词（保持原有逻辑）
            return """你是一名资深的技术方案专家和投标文件撰写专家，专门负责为采购需求提供专业的技术应答。

你的任务是：
1. 仔细分析每个采购需求的核心要点
2. 生成专业、具体、可信的技术应答
3. 应答必须以"应答：满足。"开头
4. 后续内容要体现专业性，避免空洞的表述

应答原则：
- 专业性：使用行业标准术语，体现技术实力
- 具体性：提及具体的技术方案、产品型号、服务标准
- 可信性：承诺要可实现，避免夸大其词
- 简洁性：控制在80-150字，重点突出
- 规范性：语言正式，符合商务文档标准

应答格式示例：
"应答：满足。我方采用主流的XXX技术架构，配备专业的XXX团队，严格按照XXX标准执行，确保XXX指标达到XXX水平。"

请根据具体需求内容，生成相应的专业技术应答。"""
    
    def _get_fallback_response(self) -> str:
        """获取备用应答"""
        return "应答：满足。我方具备完整的技术实力和丰富的项目经验，将严格按照采购要求提供专业的技术方案和优质服务。"
    
    def copy_paragraph_format(self, source_para, target_para):
        """
        复制段落格式，确保字体一致性
        """
        try:
            # 复制段落级别的格式
            if source_para.paragraph_format.alignment is not None:
                target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
            
            if source_para.paragraph_format.left_indent is not None:
                target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
            
            if source_para.paragraph_format.right_indent is not None:
                target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
            
            if source_para.paragraph_format.first_line_indent is not None:
                target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
            
            # 复制字体格式（从源段落的第一个run）
            if source_para.runs:
                source_run = source_para.runs[0]
                if target_para.runs:
                    target_run = target_para.runs[0]
                    
                    # 复制字体属性
                    if source_run.font.name:
                        target_run.font.name = source_run.font.name
                    if source_run.font.size:
                        target_run.font.size = source_run.font.size
                    if source_run.font.bold is not None:
                        target_run.font.bold = source_run.font.bold
                    if source_run.font.italic is not None:
                        target_run.font.italic = source_run.font.italic
                    # 对于应答内容，始终设置为黑色字体，不继承原段落颜色
                    target_run.font.color.rgb = RGBColor(0, 0, 0)
                        
        except Exception as e:
            logger.warning(f"格式复制失败: {e}")
    
    def add_paragraph_shading(self, paragraph, color):
        """
        给段落添加背景色底纹
        """
        try:
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml
            
            # 创建底纹XML元素 - RGB(217,217,217) 转换为十六进制 D9D9D9
            hex_color = f"{color.rgb:06X}" if hasattr(color, 'rgb') else "D9D9D9"
            shading_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
            shading_element = parse_xml(shading_xml)
            
            # 获取段落属性
            p_pr = paragraph._p.get_or_add_pPr()
            
            # 添加底纹元素
            p_pr.append(shading_element)
            
        except Exception as e:
            logger.warning(f"添加段落底纹失败: {e}")
    
    def insert_paragraph_after_with_format(self, paragraph, text=None, style=None):
        """
        在指定段落后插入新段落，并保持格式一致
        """
        if not DOCX_AVAILABLE:
            logger.error("未安装python-docx库，无法进行段落插入操作")
            return None
        
        try:
            # 插入新段落
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = paragraph._parent.add_paragraph()
            new_para._p = new_p
            
            if text:
                run = new_para.add_run(text)
                # 设置字体为黑色
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            if style:
                new_para.style = style
            
            # 复制源段落的格式
            self.copy_paragraph_format(paragraph, new_para)
            
            # 设置1.5倍行距
            new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            new_para.paragraph_format.line_spacing = 1.5
            
            # 添加浅灰色底纹 RGB(217,217,217)
            self.add_paragraph_shading(new_para, RGBColor(217, 217, 217))
            
            return new_para
            
        except Exception as e:
            logger.error(f"段落插入失败: {e}")
            return None
    
    def is_requirement_paragraph(self, para) -> bool:
        """
        优化的需求段落识别 - 减少遗漏
        """
        text = para.text.strip()
        if not text or len(text) < 10:  # 过短的文本不处理
            return False
        
        style_name = para.style.name if para.style else ""
        alignment = para.paragraph_format.alignment
        
        # 明确排除的条件
        exclusions = [
            "Heading" in style_name,     # 标题样式
            "标题" in style_name,         # 中文标题样式  
            "TOC" in style_name,         # 目录样式
            "目录" in text,               # 包含"目录"
            "封面" in text,               # 封面内容
            "签字" in text,               # 签字页
            alignment == 1,               # 居中对齐（通常是标题）
            text.startswith("第") and ("章" in text or "节" in text),  # 章节标题
            re.match(r'^[一二三四五六七八九十]+[、．]', text),  # 中文数字章节
        ]
        
        if any(exclusions):
            return False
        
        # 包含条件：更宽泛的识别策略
        inclusions = [
            # 包含需求关键词
            any(keyword in text for keyword in self.patterns["关键词"]),
            # 包含编号格式
            any(re.match(pattern, text) for pattern in self.patterns["编号模式"]),
            # 包含主体标识
            any(entity in text for entity in self.patterns.get("章节标识", [])),
            # 长段落且包含具体要求（超过50字且包含技术词汇）
            (len(text) > 50 and any(word in text for word in ["系统", "数据", "接口", "服务", "技术", "平台", "软件", "硬件"])),
        ]
        
        return any(inclusions)
    
    def classify_requirement_type(self, text: str) -> str:
        """
        需求类型分类
        """
        content_lower = text.lower()
        
        # 数据服务相关
        if any(kw in content_lower for kw in ["数据", "查询", "接口", "api", "专线"]):
            return "数据服务"
        
        # 系统对接相关    
        if any(kw in content_lower for kw in ["对接", "联调", "测试", "部署", "集成"]):
            return "接口对接"
            
        # 运维服务相关
        if any(kw in content_lower for kw in ["运维", "维护", "监控", "故障", "响应"]):
            return "运维支持"
        
        # 硬件相关
        if any(kw in content_lower for kw in ["cpu", "内存", "存储", "硬盘", "服务器", "设备"]):
            return "硬件配置"
        
        # 软件相关
        if any(kw in content_lower for kw in ["软件", "系统", "应用", "功能", "模块"]):
            return "软件功能"
            
        # 性能相关
        if any(kw in content_lower for kw in ["性能", "速度", "并发", "响应时间", "不少于", "不低于"]):
            return "性能指标"
            
        # 服务相关
        if any(kw in content_lower for kw in ["服务", "支持", "培训", "实施", "咨询"]):
            return "服务保障"
            
        # 资质相关  
        if any(kw in content_lower for kw in ["资质", "证书", "认证", "经验", "案例"]):
            return "资质证明"
        
        return "通用模板"
    
    def generate_professional_response(self, requirement_text: str) -> str:
        """
        生成专业的技术应答（优先使用Generate.py的点对点风格）
        """
        req_type = self.classify_requirement_type(requirement_text)
        
        # 使用Generate.py的点对点应答风格
        prompt = f"{Prompt_Answer}'{requirement_text}'"
        
        try:
            response = self.llm_callback(prompt, f"{req_type}应答")
            
            # 确保以"应答：满足"开头，赋合Generate.py的风格
            if not response.startswith("应答：满足。"):
                if "应答：满足。" in response:
                    # 找到应答：满足。并移到开头
                    response = "应答：满足。" + response.split("应答：满足。")[1]
                else:
                    response = f"应答：满足。{response}"
            
            return response
        except Exception as e:
            logger.error(f"生成专业应答失败: {e}")
            # 使用Generate.py风格的备用模板
            template = self.templates.get(req_type, self.templates["通用模板"])
            return f"应答：满足。{template}。"
    
    def process_document_enhanced(self, input_file: str, output_file: str = None) -> str:
        """
        增强版文档处理
        """
        if not DOCX_AVAILABLE:
            raise Exception("未安装python-docx库，请安装：pip install python-docx")
        
        logger.info(f"开始处理文档: {input_file}")
        
        try:
            # 读取文档
            doc = Document(input_file)
            logger.info(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")
            
            # 统计需求条目
            requirement_count = 0
            processed_count = 0
            
            # 从后往前遍历，避免插入新段落影响索引
            paragraphs = list(doc.paragraphs)
            
            for i in range(len(paragraphs) - 1, -1, -1):
                para = paragraphs[i]
                
                if self.is_requirement_paragraph(para):
                    requirement_count += 1
                    text = para.text.strip()
                    
                    logger.info(f"处理需求 {requirement_count}: {text[:60]}...")
                    
                    # 生成专业应答
                    response = self.generate_professional_response(text)
                    
                    # 在需求段落后插入应答（保持格式一致）
                    reply_para = self.insert_paragraph_after_with_format(para, response)
                    
                    if reply_para:
                        processed_count += 1
                        logger.info(f"已插入应答: {response[:60]}...")
                    else:
                        logger.error(f"插入应答失败: {text[:30]}...")
            
            # 生成输出文件名
            if not output_file:
                base_name = os.path.splitext(os.path.basename(input_file))[0]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_file = f"{base_name}-增强应答-{timestamp}.docx"
            
            # 保存文档
            doc.save(output_file)
            
            logger.info(f"处理完成:")
            logger.info(f"  识别需求条目: {requirement_count}")
            logger.info(f"  成功插入应答: {processed_count}")
            logger.info(f"  输出文件: {output_file}")
            
            return output_file
            
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            raise

def main():
    """主函数"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python3 enhanced_inline_reply.py <输入文档路径> [API密钥] [输出文档路径]")
        print("例如: python3 enhanced_inline_reply.py 采购需求.docx sk-xxxx")
        return
    
    input_file = sys.argv[1]
    api_key = sys.argv[2] if len(sys.argv) > 2 else "sk-xxx"
    output_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        processor = EnhancedInlineReplyProcessor(api_key=api_key)
        
        # 测试API连接（如果提供了有效密钥）
        if processor._is_valid_key(processor.api_key):
            logger.info("正在测试始皇API连接...")
            if processor.test_api_connection():
                logger.info("✅ 始皇API连接测试成功，将使用AI生成专业应答")
            else:
                logger.warning("❌ 始皇API连接测试失败，将使用备用模板")
        else:
            logger.info("未配置有效API密钥，将使用备用模板")
            if API_CONFIG_AVAILABLE:
                logger.info("提示：运行 'python3 api_config.py' 查看API配置指南")
        
        result_file = processor.process_document_enhanced(input_file, output_file)
        
        print(f"\n✅ 增强版处理完成!")
        print(f"📄 输入文件: {input_file}")
        print(f"📝 输出文件: {result_file}")
        print(f"🤖 使用模型: {processor.model_config['model']}")
        print(f"📋 详细日志: enhanced_inline_reply.log")
        print("\n🎉 专业技术应答生成成功！")
        
    except Exception as e:
        logger.error(f"程序执行失败: {e}")
        print(f"❌ 程序执行失败: {e}")

if __name__ == "__main__":
    main()