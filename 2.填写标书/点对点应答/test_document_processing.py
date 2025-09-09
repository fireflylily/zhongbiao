#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理功能测试脚本
Author: AI标书平台开发团队
Date: 2024-12-09
"""

import os
import json
import unittest
from docx import Document
from datetime import datetime
import tempfile
import shutil

# 导入要测试的模块
from table_processor import TableProcessor
from image_inserter import SmartImageInserter
from document_processor import DocumentProcessor, ProcessingOptions


class TestTableProcessor(unittest.TestCase):
    """表格处理器测试"""
    
    def setUp(self):
        """测试前准备"""
        self.processor = TableProcessor()
        self.test_company_info = {
            "companyName": "测试科技有限公司",
            "establishDate": "2020-01-15",
            "legalRepresentative": "测试法人",
            "registeredCapital": "5000000",
            "businessScope": "软件开发、技术服务",
            "companyAddress": "北京市海淀区测试路1号",
            "contactPhone": "010-88888888",
            "email": "test@example.com",
            "socialCreditCode": "91110108MA00TEST00"
        }
        
        # 创建测试文档
        self.test_doc_path = self._create_test_document()
    
    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)
    
    def _create_test_document(self):
        """创建测试文档"""
        doc = Document()
        doc.add_heading('测试文档', 0)
        
        # 添加测试表格
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # 填充测试数据
        table.cell(0, 0).text = '公司名称'
        table.cell(0, 1).text = ''  # 空白待填充
        
        table.cell(1, 0).text = '成立日期'
        table.cell(1, 1).text = '_____'  # 下划线占位符
        
        table.cell(2, 0).text = '法定代表人'
        table.cell(2, 1).text = '请填写'  # 文字占位符
        
        table.cell(3, 0).text = '注册资本'
        table.cell(3, 1).text = ''
        
        # 保存文档
        temp_path = tempfile.mktemp(suffix='.docx')
        doc.save(temp_path)
        return temp_path
    
    def test_field_matching(self):
        """测试字段匹配功能"""
        # 分析表格
        analysis = self.processor.analyze_tables(self.test_doc_path)
        
        # 验证结果
        self.assertEqual(analysis['total_tables'], 1)
        self.assertGreater(len(analysis['tables'][0]['matched_fields']), 0)
        
        # 检查匹配的字段
        matched_types = [f['field_type'] for f in analysis['tables'][0]['matched_fields']]
        self.assertIn('company_name', matched_types)
        self.assertIn('established_date', matched_types)
    
    def test_table_filling(self):
        """测试表格填充功能"""
        # 处理文档
        output_path = self.processor.process_document(
            self.test_doc_path,
            self.test_company_info
        )
        
        # 验证输出文件存在
        self.assertTrue(os.path.exists(output_path))
        
        # 读取输出文档验证内容
        doc = Document(output_path)
        table = doc.tables[0]
        
        # 验证填充的内容
        self.assertEqual(table.cell(0, 1).text, "测试科技有限公司")
        self.assertIn("2020", table.cell(1, 1).text)  # 日期格式可能变化
        self.assertEqual(table.cell(2, 1).text, "测试法人")
        
        # 清理输出文件
        os.remove(output_path)


class TestImageInserter(unittest.TestCase):
    """图片插入器测试"""
    
    def setUp(self):
        """测试前准备"""
        self.inserter = SmartImageInserter()
        self.test_company_info = {
            "companyName": "测试公司",
            "qualifications": {
                "business_license_path": self._create_test_image(),
                "qualification_cert_path": None,
                "legal_person_id_path": None
            }
        }
        
        # 创建测试文档
        self.test_doc_path = self._create_test_document_with_markers()
    
    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)
        
        # 清理测试图片
        license_path = self.test_company_info['qualifications']['business_license_path']
        if license_path and os.path.exists(license_path):
            os.remove(license_path)
    
    def _create_test_image(self):
        """创建测试图片"""
        from PIL import Image
        
        # 创建简单的测试图片
        img = Image.new('RGB', (100, 100), color='white')
        temp_path = tempfile.mktemp(suffix='.jpg')
        img.save(temp_path)
        return temp_path
    
    def _create_test_document_with_markers(self):
        """创建带标记的测试文档"""
        doc = Document()
        doc.add_heading('公司资质文件', 0)
        
        doc.add_paragraph('以下是公司基本资质：')
        doc.add_paragraph('[营业执照]')  # 文本标记
        doc.add_paragraph('其他内容...')
        
        # 保存文档
        temp_path = tempfile.mktemp(suffix='.docx')
        doc.save(temp_path)
        return temp_path
    
    def test_marker_detection(self):
        """测试标记检测"""
        doc = Document(self.test_doc_path)
        
        # 查找文本标记策略的插入点
        for strategy in self.inserter.strategies:
            if strategy.__class__.__name__ == 'TextMarkerStrategy':
                points = strategy.find_insertion_points(doc)
                self.assertGreater(len(points), 0)
                self.assertEqual(points[0].image_type, 'business_license')
                break
    
    def test_image_insertion(self):
        """测试图片插入"""
        # 处理文档
        output_path, results = self.inserter.process_document(
            self.test_doc_path,
            self.test_company_info
        )
        
        # 验证结果
        self.assertTrue(os.path.exists(output_path))
        self.assertGreater(len(results), 0)
        
        # 检查是否有成功的插入
        success_results = [r for r in results if r.success]
        self.assertGreater(len(success_results), 0)
        
        # 清理输出文件
        os.remove(output_path)


class TestDocumentProcessor(unittest.TestCase):
    """文档处理器集成测试"""
    
    def setUp(self):
        """测试前准备"""
        self.processor = DocumentProcessor()
        
        # 创建测试公司配置
        self.test_company_id = "test_company"
        self.test_company_info = {
            "companyName": "集成测试公司",
            "establishDate": "2019-06-01",
            "legalRepresentative": "集成测试法人",
            "registeredCapital": "10000000",
            "businessScope": "综合业务",
            "companyAddress": "测试地址",
            "contactPhone": "13800138000",
            "email": "integration@test.com",
            "socialCreditCode": "91110108MA00INTG00",
            "qualifications": {}
        }
        
        # 保存测试公司配置
        self._save_test_company_config()
        
        # 创建综合测试文档
        self.test_doc_path = self._create_comprehensive_test_document()
    
    def tearDown(self):
        """测试后清理"""
        # 清理测试文档
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)
        
        # 清理测试公司配置
        config_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            'company_configs',
            f'{self.test_company_id}.json'
        )
        if os.path.exists(config_path):
            os.remove(config_path)
    
    def _save_test_company_config(self):
        """保存测试公司配置"""
        config_dir = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            'company_configs'
        )
        os.makedirs(config_dir, exist_ok=True)
        
        config_path = os.path.join(config_dir, f'{self.test_company_id}.json')
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(self.test_company_info, f, ensure_ascii=False, indent=2)
    
    def _create_comprehensive_test_document(self):
        """创建综合测试文档"""
        doc = Document()
        doc.add_heading('综合测试文档', 0)
        
        # 添加表格
        doc.add_heading('公司信息表', 1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = '公司名称'
        table.cell(0, 1).text = ''
        table.cell(1, 0).text = '成立日期'
        table.cell(1, 1).text = ''
        table.cell(2, 0).text = '法定代表人'
        table.cell(2, 1).text = ''
        
        # 添加图片标记
        doc.add_heading('附件', 1)
        doc.add_paragraph('营业执照附件')
        
        # 保存文档
        temp_path = tempfile.mktemp(suffix='.docx')
        doc.save(temp_path)
        return temp_path
    
    def test_complete_processing(self):
        """测试完整处理流程"""
        # 设置处理选项
        options = ProcessingOptions(
            process_names=False,  # 不处理名称
            process_tables=True,   # 处理表格
            insert_images=False    # 不插入图片（因为没有实际图片）
        )
        
        # 处理文档
        result = self.processor.process_document(
            self.test_doc_path,
            self.test_company_info,
            options
        )
        
        # 验证结果
        self.assertTrue(result.success)
        self.assertTrue(os.path.exists(result.output_path))
        self.assertGreater(result.statistics['tables_processed'], 0)
        self.assertGreater(result.statistics['fields_filled'], 0)
        
        # 清理输出文件
        if os.path.exists(result.output_path):
            os.remove(result.output_path)
    
    def test_load_company_info(self):
        """测试加载公司信息"""
        loaded_info = self.processor.load_company_info(self.test_company_id)
        
        self.assertIsNotNone(loaded_info)
        self.assertEqual(loaded_info['companyName'], '集成测试公司')
        self.assertEqual(loaded_info['legalRepresentative'], '集成测试法人')


def run_tests():
    """运行所有测试"""
    # 创建测试套件
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()
    
    # 添加测试用例
    suite.addTests(loader.loadTestsFromTestCase(TestTableProcessor))
    suite.addTests(loader.loadTestsFromTestCase(TestImageInserter))
    suite.addTests(loader.loadTestsFromTestCase(TestDocumentProcessor))
    
    # 运行测试
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    # 返回测试结果
    return result.wasSuccessful()


if __name__ == '__main__':
    print("=" * 60)
    print("文档处理功能测试")
    print("=" * 60)
    
    success = run_tests()
    
    if success:
        print("\n✅ 所有测试通过！")
    else:
        print("\n❌ 有测试失败，请检查代码。")
    
    print("=" * 60)