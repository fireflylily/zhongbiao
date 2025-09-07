#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试修复后的公司名称（全称、盖章）格式处理
测试19规则版本是否能正确识别和处理特殊盖章格式
"""

import os
import sys
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def create_test_document():
    """创建包含特殊盖章格式的测试文档"""
    doc = Document()
    
    # 添加各种盖章格式进行测试
    formats_to_test = [
        "公司名称（全称、盖章）：_____________________",
        "公司名称（盖章）：_____________________", 
        "供应商名称(盖章)：_____________________",
        "供应商全称及公章：_____________________",
        "供应商名称：_____________________ （公章）",
        "投标人名称（盖章）：_____________________",
        "单位名称及公章：_____________________",
        "公司名称：_____________________",
        "供应商名称：_____________________"
    ]
    
    for i, format_text in enumerate(formats_to_test, 1):
        para = doc.add_paragraph()
        para.add_run(f"{i}. {format_text}")
        doc.add_paragraph()  # 空行分隔
    
    test_file = "test_seal_format_input.docx"
    doc.save(test_file)
    print(f"✅ 测试文档已创建: {test_file}")
    return test_file

def test_processor():
    """测试MCP处理器的盖章格式识别"""
    print("🔍 测试MCP处理器初始化...")
    
    try:
        # 创建处理器实例
        processor = MCPBidderNameProcessor()
        print("✅ MCP处理器初始化成功")
        
        # 检查属性是否正确初始化
        print(f"📋 company_name属性: {hasattr(processor, 'company_name')}")
        
        # 创建测试文档
        test_file = create_test_document()
        
        # 测试处理
        company_name = "上海智慧足迹数据科技有限公司"
        output_file = "test_seal_format_output.docx"
        
        print(f"🚀 开始处理文档...")
        print(f"   输入文件: {test_file}")
        print(f"   公司名称: {company_name}")
        print(f"   输出文件: {output_file}")
        
        result = processor.process_bidder_name(test_file, output_file, company_name)
        
        if result.get('success', False):
            print("✅ 文档处理成功!")
            
            # 读取输出文档验证结果
            output_doc = Document(output_file)
            print("\n📄 处理结果验证:")
            
            for i, para in enumerate(output_doc.paragraphs):
                if para.text.strip():
                    contains_company = company_name in para.text
                    contains_placeholder = "_____" in para.text or "___" in para.text
                    status = "✅ 已替换" if contains_company and not contains_placeholder else "❌ 未替换"
                    print(f"   {para.text[:50]}... {status}")
                    
        else:
            print("❌ 文档处理失败!")
            
    except Exception as e:
        print(f"❌ 测试过程中出现错误: {str(e)}")
        import traceback
        traceback.print_exc()

def analyze_rules():
    """分析当前规则集"""
    print("\n🔍 分析当前规则集...")
    
    try:
        processor = MCPBidderNameProcessor()
        
        # 检查是否有bidder_patterns属性
        if hasattr(processor, 'bidder_patterns'):
            print(f"📊 当前规则数量: {len(processor.bidder_patterns)}")
            
            # 查找盖章相关的规则
            seal_rules = []
            for i, pattern_info in enumerate(processor.bidder_patterns):
                pattern_str = str(pattern_info.get('pattern', ''))
                if '盖章' in pattern_str or '公章' in pattern_str:
                    seal_rules.append((i, pattern_info))
            
            print(f"🏷️  盖章相关规则数量: {len(seal_rules)}")
            
            for rule_idx, rule_info in seal_rules:
                print(f"   规则{rule_idx}: {rule_info.get('description', 'N/A')}")
                
        else:
            print("❌ 未找到bidder_patterns属性")
            
    except Exception as e:
        print(f"❌ 分析规则时出现错误: {str(e)}")

if __name__ == "__main__":
    print("🧪 测试修复后的盖章格式处理功能")
    print("=" * 50)
    
    # 分析规则
    analyze_rules()
    
    # 测试处理器
    test_processor()
    
    print("\n" + "=" * 50)
    print("✅ 测试完成")