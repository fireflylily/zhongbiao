#!/usr/bin/env python3
"""
分析采购编号段落格式变化问题
检查最新输出文件中的采购编号段落格式影响
"""

import os
import sys
from pathlib import Path
from docx import Document
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('tender_no_format_issue.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def analyze_tender_no_paragraphs(file_path):
    """分析采购编号相关段落"""
    logger.info(f"🔍 分析文件: {file_path}")
    logger.info("="*80)
    
    if not os.path.exists(file_path):
        logger.error(f"文件不存在: {file_path}")
        return
    
    try:
        doc = Document(file_path)
        logger.info(f"文档总段落数: {len(doc.paragraphs)}")
        
        # 查找包含采购编号、64525343的段落
        relevant_paragraphs = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if any(keyword in text for keyword in ['64525343', '采购编号', '（64525343）']):
                relevant_paragraphs.append((i, paragraph, text))
        
        logger.info(f"\n找到 {len(relevant_paragraphs)} 个相关段落:")
        logger.info("="*60)
        
        for para_idx, para, text in relevant_paragraphs:
            logger.info(f"\n📋 段落 #{para_idx}:")
            logger.info(f"  文本: {text[:100]}...")
            logger.info(f"  Run数量: {len(para.runs)}")
            
            # 详细分析每个run的格式
            for run_idx, run in enumerate(para.runs):
                if run.text.strip():
                    # 格式信息
                    font_name = run.font.name or "默认"
                    font_size = run.font.size.pt if run.font.size else "默认"
                    bold = run.font.bold if run.font.bold is not None else "默认"
                    italic = run.font.italic if run.font.italic is not None else "默认"
                    underline = run.font.underline if run.font.underline is not None else "默认"
                    
                    # 特别标记包含采购编号的run
                    marker = ""
                    if '64525343' in run.text:
                        marker = " ⭐ 包含采购编号"
                    elif '采购编号' in run.text:
                        marker = " ⚠️ 包含'采购编号'文本"
                    
                    logger.info(f"    Run {run_idx+1}: '{run.text}'{marker}")
                    logger.info(f"        格式: 字体={font_name}, 大小={font_size}, 粗体={bold}, 斜体={italic}, 下划线={underline}")
            
            # 检查格式一致性
            if len(para.runs) > 1:
                formats = []
                for run in para.runs:
                    if run.text.strip():
                        format_key = (
                            run.font.name,
                            run.font.size.pt if run.font.size else None,
                            bool(run.font.bold) if run.font.bold is not None else False,
                            bool(run.font.italic) if run.font.italic is not None else False,
                            bool(run.font.underline) if run.font.underline is not None else False
                        )
                        formats.append(format_key)
                
                unique_formats = set(formats)
                if len(unique_formats) > 1:
                    logger.info(f"    ❌ 段落格式不一致: {len(unique_formats)}种不同格式")
                else:
                    logger.info(f"    ✅ 段落格式一致")
        
        # 总结分析
        logger.info("\n" + "="*80)
        logger.info("📊 采购编号格式问题分析总结:")
        logger.info("="*80)
        
        if relevant_paragraphs:
            logger.info("从日志分析可以看到采购编号处理的关键信息:")
            logger.info("1. ⚠️ 单run替换失败，尝试跨run项目编号处理...")
            logger.info("2. 执行精确跨run替换: 涉及 14 个run，开始精确修改")
            logger.info("3. ✅ 精确跨run替换完成: 修改了14个run，保留了其他run的格式")
            logger.info("4. ⚠️ 项目编号跨run替换完成: (可能影响周围格式)")
            logger.info("")
            logger.info("💡 问题分析:")
            logger.info("- 采购编号'（采购编号）'跨多个run，无法使用单run替换")
            logger.info("- 需要使用跨run处理，但这会影响14个run的格式")
            logger.info("- 尽管使用了精确修改，仍可能因为run边界变化影响格式")
        else:
            logger.info("未找到相关段落，可能文件路径错误")
        
    except Exception as e:
        logger.error(f"分析失败: {e}", exc_info=True)

def main():
    logger.info("🎯 分析采购编号段落格式变化问题")
    logger.info("="*80)
    
    # 分析最新的输出文件
    output_file = "outputs/docx-商务应答-20250906_110050.docx"
    
    if os.path.exists(output_file):
        analyze_tender_no_paragraphs(output_file)
    else:
        logger.warning(f"最新文件不存在: {output_file}")
        logger.info("尝试分析其他最近的文件...")
        
        # 查找outputs目录中最新的文件
        outputs_dir = Path("outputs")
        if outputs_dir.exists():
            docx_files = list(outputs_dir.glob("docx-商务应答-*.docx"))
            if docx_files:
                # 按修改时间排序，取最新的
                latest_file = max(docx_files, key=os.path.getmtime)
                logger.info(f"分析最新文件: {latest_file}")
                analyze_tender_no_paragraphs(str(latest_file))
            else:
                logger.error("outputs目录中没有找到相关文件")
        else:
            logger.error("outputs目录不存在")

if __name__ == "__main__":
    main()