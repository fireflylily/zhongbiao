#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试格式对齐效果
验证地址、传真、电话、电子邮件等字段的格式是否正确
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_alignment_test_document():
    """创建格式对齐测试文档"""
    doc = Document()
    doc.add_heading('格式对齐测试文档', 0)
    
    # 模拟实际的文档格式
    test_paragraphs = [
        "特此证明。",
        "1. 与本磋商有关的一切正式往来信函请寄：",
        "地址________________________________  传真________________________________",
        "电话________________________________  电子邮件____________________________",
        "",
        "附：",
        "法定代表人身份证复印件(需同时提供正面及背面)"
    ]
    
    for para_text in test_paragraphs:
        doc.add_paragraph(para_text)
    
    return doc


def main():
    """主测试函数"""
    print("开始测试格式对齐效果...")
    
    # 创建测试文档
    input_file = os.path.join(current_dir, 'test_format_alignment_input.docx')
    output_file = os.path.join(current_dir, 'test_format_alignment_output.docx')
    
    print(f"创建测试文档: {input_file}")
    doc = create_alignment_test_document()
    doc.save(input_file)
    
    # 准备智慧足迹公司信息（模拟实际场景）
    company_info = {
        "companyName": "智慧足迹数据科技有限公司",
        "establishDate": "2015-12-18",
        "legalRepresentative": "李振军",
        "socialCreditCode": "91110101MA002N1D30",
        "registeredCapital": "15466.6667万元",
        "companyType": "有限责任公司",
        "fixedPhone": "010-63271000",
        "postalCode": "100010",
        "registeredAddress": "北京市东城区王府井大街200号七层711室",
        "officeAddress": "北京市西城区成方街25号长话北院",
        "website": "http://www.smartsteps.com",
        "employeeCount": "101-500人",
        # 注意：没有fax和email字段
    }
    
    # 使用处理器处理文档
    processor = MCPBidderNameProcessor()
    
    try:
        print(f"\n处理文档...")
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="哈银消金项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print(f"处理成功！")
            
            # 验证结果
            print(f"\n=== 处理后文档内容 ===")
            processed_doc = Document(output_file)
            
            for i, para in enumerate(processed_doc.paragraphs):
                if para.text.strip():
                    text = para.text
                    print(f"段落{i}: '{text}'")
                    
                    # 检查特定字段
                    if '地址' in text and '传真' in text:
                        print(f"  → 地址传真行格式检查")
                        
                        # 期望格式：地址 北京市东城区王府井大街200号七层711室    传真 未填写
                        if '北京市东城区王府井大街200号七层711室' in text:
                            print(f"    ✅ 地址内容正确")
                        else:
                            print(f"    ❌ 地址内容错误或缺失")
                        
                        if '传真 未填写' in text or '传真未填写' in text:
                            print(f"    ✅ 传真显示'未填写'")
                        elif '传真' in text and '未填写' in text:
                            print(f"    ✅ 传真字段已处理")
                        else:
                            print(f"    ❌ 传真字段未正确处理")
                        
                        # 检查格式对齐（至少有一些空格）
                        if '    ' in text:  # 至少4个空格
                            print(f"    ✅ 字段间有适当空格")
                        else:
                            print(f"    ⚠️  字段间空格可能不足")
                    
                    if '电话' in text and '电子邮件' in text:
                        print(f"  → 电话邮件行格式检查")
                        
                        if '010-63271000' in text:
                            print(f"    ✅ 电话号码正确")
                        else:
                            print(f"    ❌ 电话号码错误或缺失")
                        
                        if '电子邮件 未填写' in text or '电子邮件未填写' in text:
                            print(f"    ✅ 电子邮件显示'未填写'")
                        elif '电子邮件' in text and '未填写' in text:
                            print(f"    ✅ 电子邮件字段已处理")
                        else:
                            print(f"    ❌ 电子邮件字段未正确处理")
            
            print(f"\n=== 期望的最终格式 ===")
            print("地址 北京市东城区王府井大街200号七层711室    传真 未填写")
            print("电话 010-63271000                         电子邮件 未填写")
            
            return True
            
        else:
            print(f"处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = main()
    if success:
        print("\n🎉 格式对齐测试完成！")
    else:
        print("\n💥 格式对齐测试失败！")
    sys.exit(0 if success else 1)