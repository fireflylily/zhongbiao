#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def test_seal_no_placeholder():
    """测试公章在前无占位符的格式处理"""
    
    test_input = "test_seal_no_placeholder_input.docx"
    test_output = "test_seal_no_placeholder_output.docx"
    
    # 创建测试文档
    doc = Document()
    doc.add_paragraph("=== 测试公章在前无占位符格式 ===")
    doc.add_paragraph("供应商名称（加盖公章）：")                    # 格式9-5
    doc.add_paragraph("供应商名称(盖章)：")                         # 格式9-5 英文括号
    doc.add_paragraph("供应商名称（公章）：")                       # 格式9-5 简化版
    doc.add_paragraph("供应商名称(公司章)：")                       # 格式9-5 公司章
    doc.add_paragraph("其他内容")
    doc.add_paragraph("供应商名称（加盖公章）：                ")   # 有占位符的版本
    doc.add_paragraph("供应商名称：")                              # 无公章版本
    doc.save(test_input)
    
    print(f"✅ 创建测试文档: {test_input}")
    
    # 处理前检查
    print(f"\n=== 处理前检查 ===")
    doc_before = Document(test_input)
    for i, para in enumerate(doc_before.paragraphs):
        para_text = para.text
        if "供应商名称" in para_text:
            print(f"段落 #{i}: '{para_text}' (长度: {len(para_text)})")
            
            # 分析格式
            has_seal = "章" in para_text
            has_colon = "：" in para_text or ":" in para_text
            
            if has_seal and has_colon:
                colon_pos = para_text.find("：") if "：" in para_text else para_text.find(":")
                seal_pos = para_text.find("章")
                after_colon = para_text[colon_pos+1:] if colon_pos >= 0 else ""
                
                if seal_pos < colon_pos:
                    if not after_colon.strip():
                        print(f"  ✅ 公章在前，冒号后无内容（应匹配格式9-5）")
                    else:
                        print(f"  ✅ 公章在前，冒号后有内容：'{after_colon}'")
                else:
                    print(f"  ⚠️ 公章在冒号后")
            elif has_colon and not has_seal:
                print(f"  ⚠️ 只有冒号，无公章")
    
    try:
        # 执行处理
        processor = MCPBidderNameProcessor()
        company_name = "智慧足迹数据科技有限公司"
        
        print(f"\n=== 开始处理 ===")
        print(f"公司名称: {company_name}")
        
        result = processor.process_bidder_name(
            input_file=test_input,
            output_file=test_output,
            company_name=company_name
        )
        
        print(f"\n处理结果: {result}")
        
        if result.get('success'):
            print("✅ 处理成功!")
            
            # 处理后检查
            print(f"\n=== 处理后检查 ===")
            if os.path.exists(test_output):
                doc_after = Document(test_output)
                for i, para in enumerate(doc_after.paragraphs):
                    para_text = para.text
                    if "供应商名称" in para_text:
                        print(f"段落 #{i}: '{para_text}'")
                        if company_name in para_text:
                            print(f"    ✅ 已填写公司名称!")
                            
                            # 检查公章是否保留
                            if "章" in para_text:
                                print(f"    ✅ 公章标记已保留")
                        else:
                            print(f"    ❌ 未填写公司名称")
            else:
                print(f"❌ 输出文件不存在: {test_output}")
            
            # 分析处理统计
            stats = result.get('stats', {})
            patterns = stats.get('patterns_found', [])
            
            print(f"\n=== 处理统计 ===")
            print(f"总替换次数: {stats.get('total_replacements', 0)}")
            
            if patterns:
                print(f"处理的模式:")
                no_placeholder_count = 0
                for pattern in patterns:
                    desc = pattern.get('description', '')
                    print(f"  - 规则 #{pattern.get('rule_index')}: {desc}")
                    if '无占位符' in desc:
                        no_placeholder_count += 1
                        print(f"    ✅ 这是无占位符格式!")
                        
                print(f"\n✅ 无占位符格式处理次数: {no_placeholder_count}")
            
        else:
            print(f"❌ 处理失败: {result.get('message', 'Unknown error')}")
            
    except Exception as e:
        print(f"❌ 处理异常: {e}")
        import traceback
        traceback.print_exc()
    
    # 清理测试文件
    if os.path.exists(test_input):
        os.remove(test_input)
    if os.path.exists(test_output):
        os.remove(test_output)
    print("\n测试文件已清理")

if __name__ == "__main__":
    test_seal_no_placeholder()