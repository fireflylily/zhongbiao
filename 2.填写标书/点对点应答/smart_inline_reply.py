#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能采购需求应答系统 v2.0
改进点：
1. 基于标题层级的智能需求分段
2. 字数控制：应答不超过需求描述字数
3. 背景内容检测：避免在背景描述中添加技术指标
4. 长段落智能分割：超过1000字且有3段以上的内容进行段落级应答
"""

import requests
import json
import logging
import os
import re
from datetime import datetime
from typing import Dict, List, Tuple, Optional

# 导入API配置
try:
    from api_config import get_api_key, is_valid_api_key, API_CONFIG
    API_CONFIG_AVAILABLE = True
except ImportError:
    API_CONFIG_AVAILABLE = False

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('smart_inline_reply.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class SmartInlineReplyProcessor:
    """智能原地应答插入处理器"""
    
    def __init__(self, api_key: str = None):
        # 获取API密钥
        if api_key and self._is_valid_key(api_key):
            self.api_key = api_key
        elif API_CONFIG_AVAILABLE:
            self.api_key = get_api_key()
        else:
            self.api_key = api_key or "sk-xxx"
            
        # 检查API密钥有效性
        if not self._is_valid_key(self.api_key):
            logger.warning("使用默认API密钥，请配置有效的始皇API密钥")
            
        # API配置
        self.model_config = API_CONFIG.copy() if API_CONFIG_AVAILABLE else {
            "base_url": "https://api.oaipro.com/v1/chat/completions",
            "model": "gpt-4o-mini",
            "temperature": 0.3,
            "max_tokens": 300,
            "timeout": 60
        }
        
        # 背景内容关键词（用于识别背景描述，避免添加技术指标）
        self.background_keywords = [
            "背景", "概述", "说明", "目的", "旨在", "为了", "简介", "介绍", 
            "总体情况", "基本情况", "项目背景", "需求背景", "业务背景",
            "工作说明书", "合同", "签字生效", "有效部分"
        ]
        
        # 技术需求关键词
        self.technical_keywords = [
            "技术", "性能", "指标", "配置", "参数", "标准", "规范", "功能",
            "接口", "协议", "算法", "架构", "系统", "平台", "软件", "硬件",
            "数据库", "服务器", "网络", "安全", "加密", "备份", "监控"
        ]

    def _is_valid_key(self, api_key: str) -> bool:
        """检查API密钥是否有效"""
        if API_CONFIG_AVAILABLE:
            return is_valid_api_key(api_key)
        return api_key and api_key != "sk-xxx" and len(api_key) > 10

    def is_background_content(self, text: str) -> bool:
        """判断是否为背景描述内容"""
        if not text or len(text.strip()) < 20:
            return True
            
        # 检查背景关键词
        background_count = sum(1 for keyword in self.background_keywords if keyword in text)
        technical_count = sum(1 for keyword in self.technical_keywords if keyword in text)
        
        # 如果背景关键词多于技术关键词，且文本较短，认为是背景内容
        if background_count > technical_count and len(text) < 200:
            return True
            
        # 检查是否包含明显的背景描述模式
        background_patterns = [
            r'本.*?旨在.*?',
            r'为了.*?要求.*?',
            r'.*?说明书.*?',
            r'.*?合同.*?生效.*?',
            r'.*?工作范围.*?',
            r'.*?基本情况.*?'
        ]
        
        for pattern in background_patterns:
            if re.search(pattern, text):
                return True
                
        return False

    def analyze_document_structure(self, doc) -> List[Dict]:
        """分析文档结构，提取标题层级"""
        structure = []
        current_headings = {}  # 存储当前各级标题
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
                
            # 获取段落样式
            style_name = para.style.name if para.style else ""
            
            # 判断是否为标题
            heading_level = self._get_heading_level(para)
            
            if heading_level > 0:
                # 更新当前标题层级
                current_headings[heading_level] = {
                    'text': text,
                    'index': i
                }
                # 清除更低级别的标题
                for level in list(current_headings.keys()):
                    if level > heading_level:
                        del current_headings[level]
                        
                structure.append({
                    'type': 'heading',
                    'level': heading_level,
                    'text': text,
                    'index': i,
                    'parent_headings': current_headings.copy()
                })
            else:
                structure.append({
                    'type': 'paragraph',
                    'text': text,
                    'index': i,
                    'parent_headings': current_headings.copy(),
                    'length': len(text)
                })
                
        return structure

    def _get_heading_level(self, para) -> int:
        """获取段落的标题级别"""
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        
        # 基于样式判断
        if "Heading" in style_name:
            try:
                return int(style_name.split()[-1])
            except:
                return 1
        elif "标题" in style_name:
            if "标题 1" in style_name or "标题1" in style_name:
                return 1
            elif "标题 2" in style_name or "标题2" in style_name:
                return 2
            elif "标题 3" in style_name or "标题3" in style_name:
                return 3
            else:
                return 1
                
        # 基于格式判断（加粗、居中等）
        alignment = para.paragraph_format.alignment
        if alignment == WD_ALIGN_PARAGRAPH.CENTER and para.runs:
            first_run = para.runs[0]
            if first_run.font.bold:
                return 1
                
        # 基于编号模式判断
        if re.match(r'^[一二三四五六七八九十]+[、.]', text):
            return 1
        elif re.match(r'^（[一二三四五六七八九十]+）', text):
            return 2
        elif re.match(r'^\d+[、.]', text):
            return 2
        elif re.match(r'^\d+\.\d+', text):
            return 3
        elif re.match(r'^\(\d+\)', text):
            return 3
            
        return 0

    def group_content_by_lowest_headings(self, structure: List[Dict]) -> List[Dict]:
        """按最低级别标题分组内容"""
        groups = []
        current_group = None
        
        for item in structure:
            if item['type'] == 'heading':
                # 判断是否为最低级别标题
                is_lowest = True
                current_level = item['level']
                
                # 检查后续是否有更低级别的标题
                item_index = structure.index(item)
                for next_item in structure[item_index + 1:]:
                    if next_item['type'] == 'heading':
                        if next_item['level'] <= current_level:
                            break
                        elif next_item['level'] > current_level:
                            is_lowest = False
                            break
                            
                if is_lowest:
                    # 开始新的分组
                    if current_group:
                        groups.append(current_group)
                    current_group = {
                        'heading': item,
                        'paragraphs': [],
                        'total_length': 0
                    }
            elif item['type'] == 'paragraph' and current_group:
                current_group['paragraphs'].append(item)
                current_group['total_length'] += item['length']
                
        # 添加最后一个分组
        if current_group:
            groups.append(current_group)
            
        return groups

    def should_split_long_content(self, group: Dict) -> bool:
        """判断是否应该拆分长内容"""
        paragraphs = group['paragraphs']
        total_length = group['total_length']
        
        # 超过1000字且段落数超过3个
        return len(paragraphs) > 3 and total_length > 1000

    def generate_response(self, content: str, content_type: str = "技术需求") -> str:
        """生成应答内容"""
        content_length = len(content)
        
        # 判断是否为背景内容
        if self.is_background_content(content):
            # 背景内容使用简单应答，不添加技术指标
            return f"应答：满足。我方理解并接受上述要求，将严格按照相关规定执行。"
        
        # 构建提示词
        prompt = f"""请为以下采购需求生成专业的技术应答：

需求内容：
{content}

要求：
1. 必须以"应答：满足。"开头
2. 应答字数不超过{content_length}字
3. 内容要专业、具体、可信
4. 如果是技术需求，可以提及具体的技术方案、标准或指标
5. 如果是服务需求，重点描述服务保障措施
6. 避免空洞的表述，要体现专业性

请生成应答："""

        return self.llm_callback(prompt, f"{content_type}应答")

    def llm_callback(self, prompt: str, purpose: str = "应答生成") -> str:
        """调用始皇API生成应答"""
        if not self._is_valid_key(self.api_key):
            return self._get_fallback_response()
            
        url = self.model_config["base_url"]
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": self.model_config["model"],
            "messages": [
                {"role": "system", "content": self._get_system_prompt()},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.model_config["temperature"],
            "max_tokens": self.model_config["max_tokens"]
        }
        
        try:
            logger.info(f"[始皇API调用] {purpose}")
            response = requests.post(url, headers=headers, json=payload, 
                                   timeout=self.model_config["timeout"])
            
            if response.status_code == 200:
                result = response.json()
                if result.get("choices"):
                    content = result["choices"][0]["message"]["content"].strip()
                    logger.info(f"API调用成功，生成应答: {content[:100]}{'...' if len(content) > 100 else ''}")
                    return content
                    
        except Exception as e:
            logger.error(f"始皇API调用失败: {e}")
            
        return self._get_fallback_response()

    def _get_system_prompt(self) -> str:
        """获取系统提示词"""
        return """你是一名专业的技术方案专家，负责为采购需求提供专业的技术应答。

应答原则：
1. 专业性：使用准确的技术术语和行业标准
2. 具体性：提及具体的技术方案、产品或服务措施  
3. 简洁性：言简意赅，重点突出
4. 可信性：承诺要实际可行，避免夸大
5. 规范性：使用正式的商务语言

请根据需求内容生成相应的专业应答。"""

    def _get_fallback_response(self) -> str:
        """获取备用应答"""
        return "应答：满足。我方具备相应的技术实力和服务能力，将严格按照采购要求提供专业方案。"

    def insert_paragraph_after(self, paragraph, text: str):
        """在段落后插入新段落"""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = paragraph._parent.add_paragraph()
        new_para._p = new_p
        new_para.add_run(text)
        
        # 复制格式
        self.copy_paragraph_format(paragraph, new_para)
        
        return new_para

    def copy_paragraph_format(self, source_para, target_para):
        """复制段落格式"""
        try:
            # 复制段落格式
            if source_para.paragraph_format.alignment is not None:
                target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
            
            # 复制字体格式
            if source_para.runs and target_para.runs:
                source_run = source_para.runs[0]
                target_run = target_para.runs[0]
                
                if source_run.font.name:
                    target_run.font.name = source_run.font.name
                if source_run.font.size:
                    target_run.font.size = source_run.font.size
                if source_run.font.bold is not None:
                    target_run.font.bold = source_run.font.bold
                    
        except Exception as e:
            logger.warning(f"格式复制失败: {e}")

    def process_document(self, input_path: str, output_path: str = None):
        """处理文档"""
        logger.info(f"开始处理文档: {input_path}")
        
        # 测试API连接
        if self._is_valid_key(self.api_key):
            test_result = self.llm_callback("测试连接", "API连接测试")
            if "应答：满足" in test_result or "满足" in test_result:
                logger.info("✅ 始皇API连接测试成功")
            else:
                logger.warning("⚠️ 始皇API连接可能存在问题")
        
        # 加载文档
        doc = Document(input_path)
        logger.info(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")
        
        # 分析文档结构
        structure = self.analyze_document_structure(doc)
        logger.info(f"文档结构分析完成，识别到 {len([s for s in structure if s['type'] == 'heading'])} 个标题")
        
        # 按最低级别标题分组
        groups = self.group_content_by_lowest_headings(structure)
        logger.info(f"内容分组完成，共 {len(groups)} 个分组")
        
        # 准备应答插入点（倒序处理，避免索引混乱）
        insert_points = []
        
        for group in groups:
            if self.should_split_long_content(group):
                # 长内容按段落单独应答
                logger.info(f"分组 '{group['heading']['text'][:30]}...' 内容较长，将进行段落级应答")
                for para_info in group['paragraphs']:
                    if len(para_info['text']) > 50:  # 跳过太短的段落
                        insert_points.append({
                            'index': para_info['index'],
                            'content': para_info['text'],
                            'type': '段落应答'
                        })
            else:
                # 整个分组应答一次
                combined_content = "\n".join([p['text'] for p in group['paragraphs']])
                if combined_content.strip():
                    insert_points.append({
                        'index': group['paragraphs'][-1]['index'],  # 在最后一个段落后插入
                        'content': combined_content,
                        'type': '分组应答'
                    })
        
        # 按索引倒序排序，避免插入时索引混乱
        insert_points.sort(key=lambda x: x['index'], reverse=True)
        
        logger.info(f"准备插入 {len(insert_points)} 个应答")
        
        # 插入应答
        inserted_count = 0
        for point in insert_points:
            try:
                para = doc.paragraphs[point['index']]
                response = self.generate_response(point['content'], point['type'])
                self.insert_paragraph_after(para, response)
                inserted_count += 1
                logger.info(f"已插入应答 {inserted_count}: {response[:50]}...")
            except Exception as e:
                logger.error(f"插入应答失败: {e}")
        
        # 保存文档
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"智能应答-{timestamp}.docx"
        
        doc.save(output_path)
        logger.info(f"处理完成，共插入 {inserted_count} 个应答，保存至: {output_path}")
        
        return output_path


def main():
    """主函数"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python3 smart_inline_reply.py <输入文档> [输出文档] [API密钥]")
        return
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    api_key = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not DOCX_AVAILABLE:
        print("错误: 未安装python-docx库")
        print("请运行: pip install python-docx")
        return
    
    if not os.path.exists(input_file):
        print(f"错误: 文件不存在 {input_file}")
        return
    
    # 创建处理器
    processor = SmartInlineReplyProcessor(api_key)
    
    # 处理文档
    try:
        result_path = processor.process_document(input_file, output_file)
        print(f"\n✅ 处理完成！")
        print(f"输出文件: {result_path}")
    except Exception as e:
        logger.error(f"处理失败: {e}")
        print(f"❌ 处理失败: {e}")


if __name__ == "__main__":
    main()