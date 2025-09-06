#!/usr/bin/env python3
"""
测试格式一致性修复效果
验证修复后的精确run方法是否能保持格式一致
"""

import os
import sys
from pathlib import Path
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
import logging

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('format_consistency_test.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    # 创建一个类似实际情况的段落
    p = doc.add_paragraph()
    p.add_run("根据贵方为")
    
    # 项目名称部分 - 斜体+下划线
    project_run = p.add_run("（项目名称）")
    project_run.font.italic = True
    project_run.font.underline = True
    
    p.add_run("项目采购采购货物及服务的竞争性磋字代表")
    
    # 姓名职务部分 - 斜体+下划线
    name_run = p.add_run("（姓名、职务）")
    name_run.font.italic = True
    name_run.font.underline = True
    
    p.add_run("经正式授权并代表供应商")
    
    # 供应商名称部分 - 跨多个run，斜体+下划线
    bracket_left = p.add_run("（")
    bracket_left.font.italic = True
    bracket_left.font.underline = True
    
    supplier_name = p.add_run("供应商名称、地址")
    supplier_name.font.italic = True
    supplier_name.font.underline = True
    
    bracket_right = p.add_run("）")
    bracket_right.font.italic = True
    bracket_right.font.underline = True
    
    p.add_run("提交下述文件正本一份及副本       份：")
    
    test_file = "format_consistency_test_input.docx"
    doc.save(test_file)
    logger.info(f"创建测试文档: {test_file}")
    return test_file

def analyze_format_consistency(file_path, description):
    """分析格式一致性"""
    logger.info(f"\n📊 {description}:")
    logger.info("="*60)
    
    doc = Document(file_path)
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            logger.info(f"\n段落 #{para_idx}: {paragraph.text[:80]}...")
            logger.info(f"Run数量: {len(paragraph.runs)}")
            
            # 检查格式一致性
            format_groups = {}
            
            for run_idx, run in enumerate(paragraph.runs):
                if run.text:
                    format_key = (
                        bool(run.font.bold) if run.font.bold is not None else False,
                        bool(run.font.italic) if run.font.italic is not None else False,  
                        bool(run.font.underline) if run.font.underline is not None else False
                    )
                    
                    if format_key not in format_groups:
                        format_groups[format_key] = []
                    format_groups[format_key].append((run_idx, run.text))
            
            logger.info(f"格式组数量: {len(format_groups)}")
            
            for format_key, runs in format_groups.items():
                bold, italic, underline = format_key
                logger.info(f"  格式(粗体={bold}, 斜体={italic}, 下划线={underline}): {len(runs)}个run")
                for run_idx, text in runs:
                    logger.info(f"    Run {run_idx+1}: '{text}'")
            
            # 特别检查供应商信息区域
            supplier_runs = []
            for run_idx, run in enumerate(paragraph.runs):
                if any(keyword in run.text for keyword in ['智慧足迹', '测试科技', '供应商', '（', '）']):
                    supplier_runs.append((run_idx, run))
            
            if supplier_runs:
                logger.info(f"\n🔍 供应商信息区域 ({len(supplier_runs)}个相关run):")
                formats_consistent = True
                first_format = None
                
                for run_idx, run in supplier_runs:
                    if run.text.strip():
                        current_format = (
                            bool(run.font.bold) if run.font.bold is not None else False,
                            bool(run.font.italic) if run.font.italic is not None else False,
                            bool(run.font.underline) if run.font.underline is not None else False
                        )
                        
                        if first_format is None:
                            first_format = current_format
                        elif first_format != current_format:
                            formats_consistent = False
                        
                        logger.info(f"  Run {run_idx+1}: '{run.text}' - 粗体={current_format[0]}, 斜体={current_format[1]}, 下划线={current_format[2]}")
                
                if formats_consistent:
                    logger.info("  ✅ 供应商信息区域格式完全一致")
                else:
                    logger.info("  ❌ 供应商信息区域格式不一致")

def main():
    logger.info("🎯 测试格式一致性修复效果")
    logger.info("="*80)
    
    # 创建测试文档
    test_file = create_test_document()
    
    # 分析原始文档
    analyze_format_consistency(test_file, "原始文档格式")
    
    # 使用修复后的处理器
    processor = MCPBidderNameProcessor()
    processor.project_name = "智能办公系统采购项目"
    processor.tender_no = "64525343"
    
    result = processor.process_bidder_name(
        input_file=test_file,
        output_file="format_consistency_test_output.docx", 
        company_name="智慧足迹数据科技有限公司"
    )
    
    logger.info(f"\n处理结果: {result}")
    
    # 分析处理后文档
    if result['success']:
        analyze_format_consistency("format_consistency_test_output.docx", "处理后文档格式")
    
    logger.info("\n🎉 格式一致性修复测试完成！")

if __name__ == "__main__":
    main()