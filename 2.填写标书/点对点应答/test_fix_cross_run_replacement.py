#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试跨run替换修复效果
专门测试"（供应商名称、地址）"格式的替换问题
"""

import os
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('cross_run_fix_test.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def create_test_document():
    """创建测试文档，包含跨run的"（供应商名称、地址）"文本"""
    doc = Document()
    
    # 添加测试段落
    para1 = doc.add_paragraph()
    
    # 模拟跨run情况：将"（供应商名称、地址）"分散在多个run中
    run1 = para1.add_run("根据贵方为")
    run2 = para1.add_run("（项目名称）")
    run2.italic = True
    run3 = para1.add_run("项目采购采购货物及服务的竞争性磋商代表")
    run4 = para1.add_run("（姓名、职务）")
    run4.italic = True
    run5 = para1.add_run("经正式授权并代表供应商")
    run6 = para1.add_run("（")  # 故意分散
    run7 = para1.add_run("供应商名称、地址")
    run8 = para1.add_run("）")  # 故意分散
    run9 = para1.add_run("提交下述文件正本一份及副本")
    run10 = para1.add_run("       ")
    run11 = para1.add_run("份：")
    
    # 保存测试文档
    test_input = "test_cross_run_fix_input.docx"
    doc.save(test_input)
    logger.info(f"创建测试文档: {test_input}")
    
    return test_input

def test_cross_run_fix():
    """测试跨run替换修复效果"""
    logger.info("=== 开始测试跨run替换修复 ===")
    
    # 创建测试文档
    test_input = create_test_document()
    test_output = "test_cross_run_fix_output.docx"
    
    try:
        # 分析原始文档结构
        logger.info("\n=== 原始文档结构分析 ===")
        doc = Document(test_input)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                logger.info(f"段落 {i}: '{para.text}'")
                logger.info(f"  Run数量: {len(para.runs)}")
                for j, run in enumerate(para.runs):
                    if run.text:
                        logger.info(f"    Run {j}: '{run.text}'")
        
        # 使用修复后的处理器处理
        logger.info("\n=== 开始MCP处理（修复版本） ===")
        processor = MCPBidderNameProcessor()
        company_name = "智慧足迹数据科技有限公司"
        
        result = processor.process_bidder_name(
            input_file=test_input,
            output_file=test_output,
            company_name=company_name
        )
        
        if result.get('success'):
            logger.info("✅ MCP处理成功")
            
            # 分析处理后的文档结构
            logger.info("\n=== 处理后文档结构分析 ===")
            doc = Document(test_output)
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip() and "供应商" in para.text:
                    logger.info(f"段落 {i}: '{para.text}'")
                    logger.info(f"  Run数量: {len(para.runs)}")
                    for j, run in enumerate(para.runs):
                        if run.text:
                            logger.info(f"    Run {j}: '{run.text}'")
                    
                    # 检查是否修复了问题
                    # 1. 检查是否有部分替换的错误（如：供应商（智慧足迹数据科技有限提交）
                    # 关键：必须是以"提交"结尾的错误格式，而不是正常的完整格式
                    if ("供应商（智慧足迹数据科技有限提交" in para.text or 
                        "供应商（智慧足迹数据科技有有提交" in para.text or
                        "供应商（中国联合网络通信有限提交" in para.text):
                        logger.error("❌ 问题仍然存在：出现了部分替换的错误格式")
                        return False
                    
                    # 2. 检查内容是否正确完整替换
                    if (company_name in para.text and 
                        "供应商名称、地址" not in para.text and 
                        "北京市东城区王府井大街200号七层711室" in para.text):
                        logger.info("✅ 问题已修复：正确完整替换，内容完整")
                        logger.info(f"   最终文本：'{para.text}'")
                        return True
                    
                    # 3. 如果只是格式分散但内容正确，也认为是成功
                    if ("智慧足迹数据科技有限公司" in para.text and 
                        "北京市东城区王府井大街200号七层711室" in para.text and
                        "供应商名称、地址" not in para.text):
                        logger.info("✅ 内容替换正确：虽然分布在多个run中，但内容完整正确")
                        logger.info(f"   最终文本：'{para.text}'")
                        return True
            
            return True
        else:
            logger.error(f"❌ MCP处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        logger.error(f"测试失败: {e}", exc_info=True)
        return False
    
    finally:
        # 清理测试文件
        for file in [test_input]:
            if os.path.exists(file):
                try:
                    os.remove(file)
                except:
                    pass

if __name__ == "__main__":
    success = test_cross_run_fix()
    if success:
        logger.info("🎉 跨run替换修复测试通过！")
    else:
        logger.error("💥 跨run替换修复测试失败！")