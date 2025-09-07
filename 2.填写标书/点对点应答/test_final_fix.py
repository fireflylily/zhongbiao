#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终修复验证测试
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_test_document():
    """创建测试文档（模拟实际格式）"""
    doc = Document()
    doc.add_heading('最终修复测试', 0)
    
    # 添加测试段落
    doc.add_paragraph("特此证明。")
    doc.add_paragraph("1. 与本磋商有关的一切正式往来信函请寄：")
    doc.add_paragraph("地址________________________________  传真________________________________")
    doc.add_paragraph("电话________________________________  电子邮件____________________________")
    
    return doc


def main():
    """主测试函数"""
    print("="*60)
    print("最终修复验证测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_final_input.docx')
    output_file = os.path.join(current_dir, 'test_final_output.docx')
    
    # 创建测试文档
    print("\n1. 创建测试文档...")
    doc = create_test_document()
    doc.save(input_file)
    
    # 准备公司信息（智慧足迹，无传真和邮件）
    company_info = {
        "companyName": "智慧足迹数据科技有限公司",
        "fixedPhone": "010-63271000",
        "registeredAddress": "北京市东城区王府井大街200号七层711室",
        # 注意：没有 fax 和 email 字段
    }
    
    print("\n2. 公司信息:")
    print(f"   - 公司名称: {company_info['companyName']}")
    print(f"   - 注册地址: {company_info['registeredAddress']}")
    print(f"   - 电话: {company_info['fixedPhone']}")
    print(f"   - 传真: {company_info.get('fax', '无')}")
    print(f"   - 邮件: {company_info.get('email', '无')}")
    
    # 处理文档
    print("\n3. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-001",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            print("\n4. 处理结果:")
            processed_doc = Document(output_file)
            
            # 显示所有段落
            for i, para in enumerate(processed_doc.paragraphs):
                text = para.text.strip()
                if text:
                    print(f"   段落{i}: {text}")
            
            print("\n5. 详细分析:")
            for i, para in enumerate(processed_doc.paragraphs):
                text = para.text.strip()
                if text and ('地址' in text or '电话' in text):
                    print(f"\n   段落{i}: {text}")
                    
                    # 分析结果
                    if '地址' in text:
                        if '北京市东城区王府井大街200号七层711室' in text:
                            print("      ✅ 地址正确")
                        else:
                            print("      ❌ 地址错误")
                        
                        if '传真' in text and '未填写' in text:
                            print("      ✅ 传真显示'未填写'")
                        elif '传真' in text:
                            print("      ⚠️  传真字段存在但未显示'未填写'")
                        else:
                            print("      ❌ 传真字段丢失")
                    
                    if '电话' in text:
                        if '010-63271000' in text:
                            print("      ✅ 电话正确")
                        else:
                            print("      ❌ 电话错误")
                        
                        if '电子邮件' in text and '未填写' in text:
                            print("      ✅ 电子邮件显示'未填写'")
                        elif '电子邮件' in text:
                            print("      ⚠️  电子邮件字段存在但未显示'未填写'")
                        else:
                            print("      ❌ 电子邮件字段丢失")
            
            print("\n6. 期望结果对比:")
            print("   期望: 地址 北京市东城区王府井大街200号七层711室    传真 未填写")
            print("   期望: 电话 010-63271000                         电子邮件 未填写")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理文件
        for f in [input_file, output_file]:
            if os.path.exists(f):
                os.remove(f)
        print("\n7. 测试文件已清理")


if __name__ == '__main__':
    main()