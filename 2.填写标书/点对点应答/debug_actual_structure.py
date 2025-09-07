#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试实际文档段落#30的run结构
"""

import os
from docx import Document

def analyze_actual_document():
    """分析实际上传的商务应答文档结构"""
    # 找到最新的上传文件
    upload_files = [f for f in os.listdir('uploads') if f.startswith('20250907_072319')]
    if not upload_files:
        print("❌ 未找到上传文件")
        return
        
    latest_file = f"uploads/{upload_files[0]}"
    print(f"分析文件: {latest_file}")
    
    try:
        doc = Document(latest_file)
        
        # 找到段落#30
        if len(doc.paragraphs) <= 30:
            print(f"❌ 文档只有{len(doc.paragraphs)}个段落，无法分析段落#30")
            return
            
        para30 = doc.paragraphs[30]
        print(f"\n段落#30: '{para30.text}'")
        print(f"段落长度: {len(para30.text)}")
        print(f"Run数量: {len(para30.runs)}")
        
        print(f"\n详细run结构:")
        for i, run in enumerate(para30.runs):
            print(f"  Run {i}: '{run.text}' (长度: {len(run.text)})")
        
        # 检查是否包含"供应商名称："
        full_text = ''.join(run.text for run in para30.runs)
        print(f"\n完整文本: '{full_text}'")
        
        if "供应商名称：" in full_text:
            print("✅ 包含完整的'供应商名称：'")
            
            # 找到包含此文本的run
            for i, run in enumerate(para30.runs):
                if "供应商名称：" in run.text:
                    print(f"✅ '供应商名称：'在Run {i}中")
                    break
            else:
                print("❌ '供应商名称：'跨越了多个run")
                
                # 分析跨run情况
                accumulated = ""
                for i, run in enumerate(para30.runs):
                    accumulated += run.text
                    if "供应商名称：" in accumulated:
                        print(f"  完整标签跨越了Run 0 到 Run {i}")
                        break
        else:
            print("❌ 不包含完整的'供应商名称：'")
            
        return para30
        
    except Exception as e:
        print(f"❌ 分析失败: {e}")

if __name__ == '__main__':
    analyze_actual_document()