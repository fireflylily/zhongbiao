#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试地址字段是否保留传真标签
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_address_fax_test_document():
    """创建包含地址和传真在同一行的测试文档"""
    doc = Document()
    doc.add_heading('地址传真保留测试文档', 0)
    
    # 模拟问题情况 - 地址和传真在同一行
    test_cases = [
        "地址：                                      北京市西城区金融大街21号传真",
        "地址：北京市东城区王府井大街200号七层711室传真：010-12345678",
        "地址：                    上海市浦东新区陆家嘴环路1000号    传真号码：",
        "地址：深圳市南山区科技园传真",
    ]
    
    for case in test_cases:
        para = doc.add_paragraph(case)
    
    # 添加一些其他内容
    doc.add_paragraph("电话                                        电子邮件")
    doc.add_paragraph("service@chinaunicom.cn")
    
    return doc


def main():
    """主测试函数"""
    print("开始测试地址字段是否保留传真标签...")
    
    # 创建测试文档
    input_file = os.path.join(current_dir, 'test_address_fax_input.docx')
    output_file = os.path.join(current_dir, 'test_address_fax_output.docx')
    
    print(f"创建测试文档: {input_file}")
    doc = create_address_fax_test_document()
    doc.save(input_file)
    
    # 显示原始内容
    print("\n=== 原始文档内容 ===")
    original_doc = Document(input_file)
    for i, para in enumerate(original_doc.paragraphs):
        if para.text.strip():
            print(f"段落{i}: '{para.text}'")
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
        "establishDate": "2000-04-21",
        "legalRepresentative": "陈忠岳",
        "registeredAddress": "北京市西城区金融大街21号",
        "fixedPhone": "010-66258899",
        "fax": "010-66258866",
        "email": "service@chinaunicom.cn"
    }
    
    # 使用处理器处理文档
    processor = MCPBidderNameProcessor()
    
    try:
        print(f"\n处理文档...")
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-2025-001",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print(f"处理成功！")
            
            # 验证结果
            print(f"\n=== 处理后文档内容 ===")
            processed_doc = Document(output_file)
            
            preserved_count = 0
            lost_count = 0
            issues = []
            
            for i, para in enumerate(processed_doc.paragraphs):
                if para.text.strip():
                    text = para.text.strip()
                    print(f"段落{i}: '{text}'")
                    
                    # 检查是否是地址相关的段落
                    if '地址' in text:
                        if '传真' in text:
                            print(f"  ✅ 传真标签已保留")
                            preserved_count += 1
                        else:
                            print(f"  ❌ 传真标签丢失")
                            lost_count += 1
                            issues.append({
                                'paragraph': i,
                                'text': text,
                                'issue': '传真标签丢失'
                            })
            
            # 总结结果
            print(f"\n=== 传真标签保留测试结果 ===")
            print(f"✅ 成功保留传真标签: {preserved_count} 个段落")
            print(f"❌ 丢失传真标签: {lost_count} 个段落")
            
            if issues:
                print("\n❌ 发现的问题:")
                for issue in issues:
                    print(f"  段落{issue['paragraph']}: {issue['issue']}")
                    print(f"    内容: '{issue['text']}'")
            
            if lost_count == 0:
                print("\n🎉 所有传真标签都成功保留！")
                return True
            else:
                print(f"\n💥 有 {lost_count} 个传真标签丢失！")
                return False
            
        else:
            print(f"处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)