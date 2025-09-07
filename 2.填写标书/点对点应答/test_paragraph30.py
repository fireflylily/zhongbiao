#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专门测试段落#30的规则执行情况
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_paragraph30_test():
    """创建段落#30的测试文档"""
    doc = Document()
    doc.add_heading('段落#30测试', 0)
    
    # 模拟段落#30的确切格式
    para = doc.add_paragraph()
    para.add_run("供应商名称：                          采购编号：                   ")
    
    return doc


def main():
    """主测试函数"""
    print("="*60)
    print("段落#30规则执行测试（合并规则后）")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_para30_input.docx')
    output_file = os.path.join(current_dir, 'test_para30_output.docx')
    
    # 创建测试文档
    print("\n1. 创建段落#30测试文档...")
    doc = create_paragraph30_test()
    doc.save(input_file)
    
    # 分析原始格式
    print(f"   原始文本: {doc.paragraphs[1].text}")
    
    # 准备公司信息
    company_info = {
        "companyName": "智慧足迹数据科技有限公司",
    }
    
    # 处理文档
    print("\n2. 分析合并后的规则...")
    processor = MCPBidderNameProcessor()
    
    print(f"   规则总数: {len(processor.bidder_patterns)}")
    
    # 手动分析哪些规则会匹配段落#30
    test_text = "供应商名称：                          采购编号：                   "
    print(f"\n3. 分析文本: '{test_text}'")
    print("   会匹配的规则:")
    
    for i, rule in enumerate(processor.bidder_patterns, 1):
        pattern = rule['pattern']
        match = pattern.search(test_text)
        if match:
            print(f"   ✅ 规则{i}: {rule['description']}")
            print(f"      匹配内容: '{match.group(0)}'")
            print(f"      规则类型: {rule['type']}")
            if rule['type'] == 'fill_space_tender_no':
                print("      → 执行方法: _fill_space_tender_no_method()")
                print("      → 填写内容: GXTC-C-251590031")
            elif rule['type'] == 'fill_space':
                print("      → 执行方法: _fill_space_method()")
                print("      → 填写内容: 智慧足迹数据科技有限公司")
            print()
    
    print("\n4. 执行处理...")
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            processed_doc = Document(output_file)
            processed_text = processed_doc.paragraphs[1].text
            print(f"   处理后文本: '{processed_text}'")
            
            # 验证结果
            print("\n5. 验证结果:")
            if '智慧足迹数据科技有限公司' in processed_text:
                print("   ✅ 供应商名称填写成功")
            else:
                print("   ❌ 供应商名称填写失败")
                
            if 'GXTC-C-251590031' in processed_text:
                print("   ✅ 采购编号填写成功")
            else:
                print("   ❌ 采购编号填写失败")
            
            # 检查下划线占位符清理
            if '___' not in processed_text:
                print("   ✅ 下划线占位符清理成功")
            else:
                print(f"   ⚠️ 仍有下划线残留: {processed_text.count('_')}个")
                
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)
        print(f"\n6. 输出文件: {output_file}")


if __name__ == '__main__':
    main()