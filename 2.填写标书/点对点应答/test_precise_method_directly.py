#!/usr/bin/env python3
"""
直接测试精确run修改方法
绕过公司名称检查，直接验证新的替换逻辑
"""

import os
import sys
from pathlib import Path
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
import logging
import re

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('direct_precise_test.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    # 场景1：单run中的项目名称（应该使用单run替换）
    p1 = doc.add_paragraph()
    p1.add_run("测试项目：")
    project_run = p1.add_run("（项目名称）")
    project_run.font.italic = True
    project_run.font.underline = True
    p1.add_run(" 效果验证")
    
    # 场景2：跨run的采购编号（应该使用精确跨run替换）
    p2 = doc.add_paragraph()
    p2.add_run("竞争性磋商公告")
    
    bracket_left = p2.add_run("（")
    bracket_left.font.italic = True
    bracket_left.font.underline = True
    bracket_left.font.bold = True
    
    tender_text = p2.add_run("采购编号")
    tender_text.font.italic = True
    tender_text.font.underline = True
    tender_text.font.bold = True
    
    bracket_right = p2.add_run("）")
    bracket_right.font.italic = True
    bracket_right.font.underline = True
    bracket_right.font.bold = True
    
    p2.add_run("，签字代表授权")
    
    test_file = "direct_precise_test.docx"
    doc.save(test_file)
    logger.info(f"创建测试文档: {test_file}")
    return test_file

def test_direct_precise_replacement():
    """直接测试精确替换方法"""
    logger.info("🚀 直接测试精确run修改方法")
    logger.info("="*80)
    
    try:
        # 创建测试文档
        test_file = create_test_document()
        
        # 分析原始文档
        logger.info("\n📊 原始文档结构:")
        doc = Document(test_file)
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                logger.info(f"段落 #{i+1}: {paragraph.text}")
                logger.info(f"  Run数量: {len(paragraph.runs)}")
                for j, run in enumerate(paragraph.runs):
                    if run.text:
                        format_info = f"斜体={run.font.italic}, 粗体={run.font.bold}, 下划线={run.font.underline}"
                        logger.info(f"    Run {j+1}: '{run.text}' [{format_info}]")
        
        # 创建处理器并设置项目编号
        processor = MCPBidderNameProcessor()
        processor.project_number = "64525343"
        
        # 测试场景1：项目名称单run替换
        logger.info("\n🎯 测试场景1：项目名称单run替换")
        logger.info("="*50)
        test_paragraph1 = doc.paragraphs[0]
        pattern1 = re.compile(r'[\(（]\s*项目名称\s*[\)）]')
        match1 = pattern1.search(test_paragraph1.text)
        
        if match1:
            rule1 = {'description': '项目名称测试', 'type': 'replace_content_project'}
            success1 = processor._replace_content_project_method(test_paragraph1, match1, rule1)
            logger.info(f"项目名称替换结果: {success1}")
        
        # 测试场景2：采购编号跨run替换
        logger.info("\n🎯 测试场景2：采购编号跨run替换")
        logger.info("="*50)
        test_paragraph2 = doc.paragraphs[1]
        pattern2 = re.compile(r'[\(（]\s*采购编号\s*[\)）]')
        match2 = pattern2.search(test_paragraph2.text)
        
        if match2:
            rule2 = {'description': '采购编号测试', 'type': 'replace_content_tender_no'}
            success2 = processor._replace_content_tender_no_method(test_paragraph2, match2, rule2)
            logger.info(f"采购编号替换结果: {success2}")
        
        # 分析处理后文档
        logger.info("\n📊 处理后文档结构:")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                logger.info(f"段落 #{i+1}: {paragraph.text}")
                logger.info(f"  Run数量: {len(paragraph.runs)}")
                for j, run in enumerate(paragraph.runs):
                    if run.text:
                        format_info = f"斜体={run.font.italic}, 粗体={run.font.bold}, 下划线={run.font.underline}"
                        logger.info(f"    Run {j+1}: '{run.text}' [{format_info}]")
        
        # 保存结果
        output_file = "direct_precise_test_result.docx"
        doc.save(output_file)
        logger.info(f"\n💾 处理结果保存到: {output_file}")
        
        # 效果验证
        logger.info("\n🎯 效果验证:")
        logger.info("="*50)
        logger.info("检查要点:")
        logger.info("1. 项目名称是否从'（项目名称）'变为'（智能办公系统采购项目）'")
        logger.info("2. 采购编号是否从'（采购编号）'变为'（64525343）'")
        logger.info("3. 周围文字的格式是否保持不变")
        logger.info("4. Run结构变化是否最小化")
        
        return True
        
    except Exception as e:
        logger.error(f"直接测试失败: {e}", exc_info=True)
        return False

if __name__ == "__main__":
    test_direct_precise_replacement()