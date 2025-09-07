#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试成立日期填写优化功能
专门测试避免重复"年月日"字符的问题
"""

import os
import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_date_test_document():
    """创建专门测试日期填写的文档"""
    doc = Document()
    doc.add_heading('成立日期填写测试文档', 0)
    
    # 添加各种格式的成立日期字段
    test_fields = [
        "成立日期：            年    月    日",
        "成立时间：                    年        月        日", 
        "设立日期：_____年____月____日",
        "公司成立日期：                                年                月                日",
        "企业成立时间：          年      月      日",
    ]
    
    for field in test_fields:
        para = doc.add_paragraph()
        
        # 模拟跨run的情况 - 将文本分散到多个run中
        parts = field.split('年')
        if len(parts) > 1:
            # 添加第一部分和"年"
            run1 = para.add_run(parts[0] + '年')
            
            # 处理剩余部分
            remaining = '年'.join(parts[1:])
            month_parts = remaining.split('月')
            if len(month_parts) > 1:
                # 添加月部分
                run2 = para.add_run(month_parts[0] + '月')
                
                # 处理日部分
                day_remaining = '月'.join(month_parts[1:])
                day_parts = day_remaining.split('日')
                if len(day_parts) > 1:
                    run3 = para.add_run(day_parts[0] + '日')
                    if day_parts[1]:  # 剩余部分
                        run4 = para.add_run(day_parts[1])
                else:
                    run3 = para.add_run(day_remaining)
            else:
                run2 = para.add_run(remaining)
        else:
            # 没有年字符的情况
            para.add_run(field)
    
    # 添加表格测试
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # 表格内容
    table_data = [
        ['成立日期', '        年    月    日'],
        ['成立时间', '              年          月          日'],
        ['设立日期', '___年___月___日']
    ]
    
    for i, (label, placeholder) in enumerate(table_data):
        table.cell(i, 0).text = label
        # 在表格中也模拟跨run
        cell = table.cell(i, 1)
        parts = placeholder.split('年')
        if len(parts) > 1:
            para = cell.paragraphs[0]
            para.clear()
            para.add_run(parts[0] + '年')
            remaining = '年'.join(parts[1:])
            month_parts = remaining.split('月')
            if len(month_parts) > 1:
                para.add_run(month_parts[0] + '月')
                day_remaining = '月'.join(month_parts[1:])
                para.add_run(day_remaining)
            else:
                para.add_run(remaining)
        else:
            table.cell(i, 1).text = placeholder
    
    return doc


def main():
    """主测试函数"""
    print("开始测试成立日期填写优化功能...")
    
    # 创建测试文档
    input_file = os.path.join(current_dir, 'test_date_fix_input.docx')
    output_file = os.path.join(current_dir, 'test_date_fix_output.docx')
    
    print(f"创建测试文档: {input_file}")
    doc = create_date_test_document()
    doc.save(input_file)
    
    # 打印原始文档结构用于调试
    print("\n=== 原始文档run结构 ===")
    doc = Document(input_file)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"段落{i}: '{para.text}'")
            for j, run in enumerate(para.runs):
                if run.text:
                    print(f"  run[{j}]: '{run.text}'")
    
    # 准备公司信息
    company_info = {
        "companyName": "测试科技有限公司",
        "establishDate": "2015-12-18",  # 这会被转换为"2015年12月18日"
    }
    
    # 使用处理器处理文档
    processor = MCPBidderNameProcessor()
    
    try:
        print(f"\n处理文档...")
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-2025-001",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print(f"处理成功！")
            print(f"输出文件: {output_file}")
            
            # 验证结果
            print("\n=== 验证处理结果 ===")
            verify_doc = Document(output_file)
            
            found_issues = []
            correct_results = []
            
            for i, para in enumerate(verify_doc.paragraphs):
                if para.text.strip() and '成立' in para.text:
                    print(f"段落{i}: '{para.text}'")
                    
                    # 检查run结构
                    for j, run in enumerate(para.runs):
                        if run.text:
                            print(f"  run[{j}]: '{run.text}'")
                    
                    # 检查是否有重复的年月日
                    text = para.text
                    if '2015年12月18日' in text:
                        # 检查是否有多余的年月日字符
                        after_date = text.split('2015年12月18日', 1)
                        if len(after_date) > 1:
                            remaining = after_date[1]
                            if any(char in remaining for char in ['年', '月', '日']):
                                found_issues.append({
                                    'paragraph': i,
                                    'text': text,
                                    'issue': f"日期后还有多余字符: '{remaining}'"
                                })
                            else:
                                correct_results.append({
                                    'paragraph': i,
                                    'text': text,
                                    'status': '✓ 正确'
                                })
                    elif '成立' in text and '日期' in text:
                        found_issues.append({
                            'paragraph': i,
                            'text': text,
                            'issue': "未找到正确的日期格式"
                        })
            
            # 报告结果
            print(f"\n=== 测试结果统计 ===")
            print(f"正确处理的段落: {len(correct_results)}")
            print(f"发现问题的段落: {len(found_issues)}")
            
            if correct_results:
                print("\n✅ 正确处理的段落:")
                for result in correct_results:
                    print(f"  段落{result['paragraph']}: {result['text']}")
            
            if found_issues:
                print("\n❌ 发现的问题:")
                for issue in found_issues:
                    print(f"  段落{issue['paragraph']}: {issue['text']}")
                    print(f"    问题: {issue['issue']}")
            else:
                print("\n🎉 所有成立日期都正确填写，没有重复的年月日字符！")
            
            return len(found_issues) == 0
            
        else:
            print(f"处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)