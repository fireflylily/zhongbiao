#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
对比批量替换 vs 多次单项替换
实验验证两种方法的优缺点
"""

import os
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('batch_vs_sequential_test.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def create_test_document_multiple_items():
    """创建包含多个替换项的测试文档"""
    doc = Document()
    
    # 添加测试段落 - 包含3个替换项
    para1 = doc.add_paragraph()
    
    # 故意分散在多个run中，模拟真实Word文档的复杂结构
    run1 = para1.add_run("根据贵方为")
    run2 = para1.add_run("（项目名称）")
    run2.italic = True
    run3 = para1.add_run("项目采购货物及服务，投标人")
    run4 = para1.add_run("（")
    run5 = para1.add_run("供应商名称、地址")
    run6 = para1.add_run("）")
    run7 = para1.add_run("特此声明，编号")
    run8 = para1.add_run("（采购编号）")
    run8.italic = True
    run9 = para1.add_run("。")
    
    # 保存测试文档
    test_input = "test_multiple_items_input.docx"
    doc.save(test_input)
    logger.info(f"创建多项替换测试文档: {test_input}")
    
    return test_input

def analyze_document_structure(file_path, title):
    """分析文档结构"""
    logger.info(f"\n=== {title} ===")
    doc = Document(file_path)
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            logger.info(f"段落 {i}: '{para.text}'")
            logger.info(f"  Run数量: {len(para.runs)}")
            for j, run in enumerate(para.runs):
                if run.text:
                    format_info = f"字体={run.font.name}, 大小={run.font.size}, 粗体={run.font.bold}, 斜体={run.font.italic}"
                    logger.info(f"    Run {j}: '{run.text}' [{format_info}]")
    return doc

def test_sequential_replacement():
    """测试多次单项替换方法"""
    logger.info("\n🔄 测试方案A：多次单项替换")
    
    # 创建测试文档
    test_input = create_test_document_multiple_items()
    test_output = "test_sequential_output.docx"
    
    try:
        # 分析原始文档
        analyze_document_structure(test_input, "原始文档结构")
        
        # 临时修改系统，强制使用单项替换
        processor = MCPBidderNameProcessor()
        
        # 保存原始方法
        original_should_use_batch = processor._should_use_batch_replacement
        
        # 强制返回False，禁用批量替换
        def force_single_replacement(paragraph):
            return False
        
        processor._should_use_batch_replacement = force_single_replacement
        
        # 执行处理
        result = processor.process_bidder_name(
            input_file=test_input,
            output_file=test_output,
            company_name="智慧足迹数据科技有限公司"
        )
        
        # 恢复原始方法
        processor._should_use_batch_replacement = original_should_use_batch
        
        if result.get('success'):
            logger.info("✅ 多次单项替换完成")
            final_doc = analyze_document_structure(test_output, "多次单项替换结果")
            return final_doc.paragraphs[0].text, test_output
        else:
            logger.error(f"❌ 多次单项替换失败: {result.get('error')}")
            return None, None
            
    except Exception as e:
        logger.error(f"多次单项替换测试异常: {e}", exc_info=True)
        return None, None
    finally:
        # 清理
        if os.path.exists(test_input):
            os.remove(test_input)

def test_batch_replacement():
    """测试批量替换方法"""
    logger.info("\n🔄 测试方案B：批量替换")
    
    # 创建测试文档
    test_input = create_test_document_multiple_items()
    test_output = "test_batch_output.docx"
    
    try:
        # 执行正常的批量替换处理
        processor = MCPBidderNameProcessor()
        
        result = processor.process_bidder_name(
            input_file=test_input,
            output_file=test_output,
            company_name="智慧足迹数据科技有限公司"
        )
        
        if result.get('success'):
            logger.info("✅ 批量替换完成")
            final_doc = analyze_document_structure(test_output, "批量替换结果")
            return final_doc.paragraphs[0].text, test_output
        else:
            logger.error(f"❌ 批量替换失败: {result.get('error')}")
            return None, None
            
    except Exception as e:
        logger.error(f"批量替换测试异常: {e}", exc_info=True)
        return None, None
    finally:
        # 清理
        if os.path.exists(test_input):
            os.remove(test_input)

def compare_results():
    """对比两种方法的结果"""
    logger.info("\n" + "="*80)
    logger.info("🏁 开始对比批量替换 vs 多次单项替换")
    logger.info("="*80)
    
    # 测试多次单项替换
    sequential_text, sequential_file = test_sequential_replacement()
    
    # 测试批量替换
    batch_text, batch_file = test_batch_replacement()
    
    # 对比结果
    logger.info(f"\n📊 结果对比:")
    logger.info(f"方案A (多次单项): {sequential_text}")
    logger.info(f"方案B (批量替换): {batch_text}")
    
    if sequential_text and batch_text:
        if sequential_text == batch_text:
            logger.info("✅ 两种方法结果一致")
            return True
        else:
            logger.info("❌ 两种方法结果不一致")
            
            # 详细分析差异
            logger.info("\n🔍 差异分析:")
            
            # 检查内容完整性
            expected_items = ["智慧足迹数据科技有限公司", "北京市东城区王府井大街200号七层711室"]
            
            seq_completeness = all(item in sequential_text for item in expected_items)
            batch_completeness = all(item in batch_text for item in expected_items)
            
            logger.info(f"多次单项替换完整性: {'✅' if seq_completeness else '❌'}")
            logger.info(f"批量替换完整性: {'✅' if batch_completeness else '❌'}")
            
            return False
    else:
        logger.info("❌ 一个或多个方法执行失败")
        return False
    
    # 清理测试文件
    for file in [sequential_file, batch_file]:
        if file and os.path.exists(file):
            try:
                os.remove(file)
            except:
                pass

if __name__ == "__main__":
    success = compare_results()
    
    logger.info(f"\n🎯 结论:")
    if success:
        logger.info("两种方法都能正确工作，但可能在性能和格式保持上有差异")
    else:
        logger.info("两种方法存在明显差异，需要进一步分析原因")