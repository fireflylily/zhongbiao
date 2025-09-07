#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试格式保持效果 - 验证字体格式是否被正确保持
"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_format_test_document():
    """创建具有不同格式的测试文档"""
    doc = Document()
    doc.add_heading('格式保持测试文档', 0)
    
    # 测试案例1: 法定代表人 - 模拟不同run有不同格式
    para1 = doc.add_paragraph()
    run1 = para1.add_run("法定代表人（负")
    run1.font.size = Pt(12)
    run1.font.bold = False
    
    run2 = para1.add_run("责人）姓名：")  # 这部分可能有不同格式
    run2.font.size = Pt(12)
    run2.font.bold = True  # 设置为粗体，模拟格式差异
    
    run3 = para1.add_run("                    陈忠岳")  # 大量空格 + 内容
    run3.font.size = Pt(12)
    run3.font.bold = False
    
    # 测试案例2: 成立时间 - 不同的字体大小
    para2 = doc.add_paragraph()
    run21 = para2.add_run("成立时间：")
    run21.font.size = Pt(14)
    run21.font.bold = True
    
    run22 = para2.add_run("                    ")  # 大量空格
    run22.font.size = Pt(12)
    
    run23 = para2.add_run("年        月        日")
    run23.font.size = Pt(12)
    run23.font.italic = True
    
    # 测试案例3: 地址 - 带下划线格式
    para3 = doc.add_paragraph()
    run31 = para3.add_run("地址：")
    run31.font.size = Pt(12)
    run31.font.underline = WD_UNDERLINE.SINGLE
    
    run32 = para3.add_run("                                      ")  # 更多空格
    run32.font.size = Pt(12)
    
    run33 = para3.add_run("北京市西城区金融大街21号传真")
    run33.font.size = Pt(12)
    run33.font.bold = False
    
    return doc


def analyze_run_formats(paragraph, description):
    """分析段落中各个run的格式"""
    print(f"\n=== {description} ===")
    print(f"段落总文本: '{paragraph.text}'")
    
    for i, run in enumerate(paragraph.runs):
        if run.text:
            font = run.font
            print(f"run[{i}]: '{run.text}'")
            print(f"  字体大小: {font.size}")
            print(f"  粗体: {font.bold}")
            print(f"  斜体: {font.italic}")
            print(f"  下划线: {font.underline}")


def main():
    """主测试函数"""
    print("开始测试格式保持效果...")
    
    # 创建测试文档
    input_file = os.path.join(current_dir, 'test_format_preservation_input.docx')
    output_file = os.path.join(current_dir, 'test_format_preservation_output.docx')
    
    print(f"创建格式测试文档: {input_file}")
    doc = create_format_test_document()
    doc.save(input_file)
    
    # 分析原始文档格式
    print("\n=== 原始文档格式分析 ===")
    original_doc = Document(input_file)
    for i, para in enumerate(original_doc.paragraphs):
        if para.text.strip() and any(keyword in para.text for keyword in ['法定代表人', '成立时间', '地址']):
            analyze_run_formats(para, f"原始段落{i}")
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
        "establishDate": "2000-04-21",
        "legalRepresentative": "陈忠岳",
        "registeredAddress": "北京市西城区金融大街21号",
        "fax": "010-66258866",
        "email": "service@chinaunicom.cn"
    }
    
    # 使用处理器处理文档
    processor = MCPBidderNameProcessor()
    
    try:
        print(f"\n处理文档...")
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-2025-001",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print(f"处理成功！")
            
            # 分析处理后的文档格式
            print(f"\n=== 处理后文档格式分析 ===")
            processed_doc = Document(output_file)
            
            format_issues = []
            format_preserved = []
            
            for i, para in enumerate(processed_doc.paragraphs):
                if para.text.strip() and any(keyword in para.text for keyword in ['法定代表人', '成立时间', '地址']):
                    analyze_run_formats(para, f"处理后段落{i}")
                    
                    # 检查格式一致性
                    if len(para.runs) > 1:
                        # 检查不同run的格式是否有不合理的差异
                        first_run = para.runs[0]
                        for j, run in enumerate(para.runs[1:], 1):
                            if run.text.strip():  # 只检查有内容的run
                                if (first_run.font.bold != run.font.bold or 
                                    first_run.font.size != run.font.size):
                                    format_issues.append({
                                        'paragraph': i,
                                        'text': para.text,
                                        'issue': f'run[0]和run[{j}]格式不一致',
                                        'run0_bold': first_run.font.bold,
                                        'run_j_bold': run.font.bold,
                                        'run0_size': first_run.font.size,
                                        'run_j_size': run.font.size
                                    })
                    else:
                        format_preserved.append({
                            'paragraph': i,
                            'text': para.text,
                            'status': '单run格式保持'
                        })
            
            # 总结格式保持情况
            print(f"\n=== 格式保持总结 ===")
            print(f"✅ 格式正常保持: {len(format_preserved)} 个段落")
            print(f"⚠️  格式可能有问题: {len(format_issues)} 个段落")
            
            if format_preserved:
                print("\n✅ 格式正常的段落:")
                for item in format_preserved:
                    print(f"  段落{item['paragraph']}: {item['status']}")
            
            if format_issues:
                print("\n⚠️  格式可能有问题的段落:")
                for item in format_issues:
                    print(f"  段落{item['paragraph']}: {item['issue']}")
                    print(f"    文本: '{item['text']}'")
                    print(f"    run[0]: 粗体={item['run0_bold']}, 大小={item['run0_size']}")
                    print(f"    run[j]: 粗体={item['run_j_bold']}, 大小={item['run_j_size']}")
            
            return len(format_issues) == 0
            
        else:
            print(f"处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = main()
    if success:
        print("\n🎉 格式保持测试通过！")
    else:
        print("\n💥 格式保持测试失败，发现格式问题！")
    sys.exit(0 if success else 1)