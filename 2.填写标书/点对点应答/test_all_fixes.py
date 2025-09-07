#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
综合测试所有修复效果
验证：
1. 地址和传真之间有4个空格
2. 电话字段能正确填写
3. 电子邮件格式正确
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_comprehensive_test_document():
    """创建包含所有问题字段的测试文档"""
    doc = Document()
    doc.add_heading('综合修复测试文档', 0)
    
    # 模拟问题情况的字段
    test_fields = [
        "供应商名称：中国联合网络通信有限公司",
        "成立时间：        年    月    日",
        "法定代表人（负责人）姓名：            ",
        "",
        "与本磋商有关的一切正式往来信函请寄：",
        "地址：                                传真：                                ",
        "电话：                                电子邮件：                            ",
        "",
        "特此证明。"
    ]
    
    for field in test_fields:
        if field.strip():
            doc.add_paragraph(field)
        else:
            doc.add_paragraph("")  # 空行
    
    return doc


def analyze_results(processed_doc):
    """分析处理结果"""
    results = {
        'address_fax_spacing': {'status': 'not_found', 'details': ''},
        'phone_filled': {'status': 'not_found', 'details': ''},
        'email_format': {'status': 'not_found', 'details': ''}
    }
    
    # 收集所有文本
    all_texts = []
    for para in processed_doc.paragraphs:
        if para.text.strip():
            all_texts.append(para.text.strip())
    
    print("\n=== 处理后的文档内容 ===")
    for i, text in enumerate(all_texts):
        print(f"段落{i}: '{text}'")
    
    # 检查1: 地址和传真之间的空格
    for text in all_texts:
        if '地址' in text and '传真' in text:
            results['address_fax_spacing']['details'] = text
            # 查找地址内容后是否有4个空格
            if '北京市西城区金融大街21号    传真' in text:
                results['address_fax_spacing']['status'] = 'success'
            elif '北京市西城区金融大街21号传真' in text:
                results['address_fax_spacing']['status'] = 'failed_no_spacing'
            else:
                results['address_fax_spacing']['status'] = 'failed_other'
            break
    
    # 检查2: 电话字段是否被填写
    for text in all_texts:
        if '电话' in text and '电子邮件' in text:
            results['phone_filled']['details'] = text
            # 检查是否包含电话号码
            if '010-66258899' in text:
                results['phone_filled']['status'] = 'success'
            elif '电话                                  ' in text:
                results['phone_filled']['status'] = 'failed_not_filled'
            else:
                results['phone_filled']['status'] = 'failed_other'
            break
    
    # 检查3: 电子邮件格式
    for text in all_texts:
        if 'service@chinaunicom.cn' in text:
            results['email_format']['details'] = text
            # 检查邮件是否在正确位置（不是追加在末尾）
            if text.endswith('service@chinaunicom.cn'):
                results['email_format']['status'] = 'failed_appended'
            elif '电子邮件service@chinaunicom.cn' in text or '电子邮件：service@chinaunicom.cn' in text:
                results['email_format']['status'] = 'success'
            else:
                results['email_format']['status'] = 'failed_other'
            break
    
    return results


def main():
    """主测试函数"""
    print("开始综合测试所有修复效果...")
    
    # 创建测试文档
    input_file = os.path.join(current_dir, 'test_all_fixes_input.docx')
    output_file = os.path.join(current_dir, 'test_all_fixes_output.docx')
    
    print(f"创建测试文档: {input_file}")
    doc = create_comprehensive_test_document()
    doc.save(input_file)
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
        "establishDate": "2000-04-21",
        "legalRepresentative": "陈忠岳",
        "registeredAddress": "北京市西城区金融大街21号",
        "fixedPhone": "010-66258899",
        "fax": "010-66258866",
        "email": "service@chinaunicom.cn"
    }
    
    # 使用处理器处理文档
    processor = MCPBidderNameProcessor()
    
    try:
        print(f"\n处理文档...")
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="综合测试项目",
            tender_no="TEST-2025-ALL",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print(f"处理成功！")
            
            # 分析结果
            processed_doc = Document(output_file)
            results = analyze_results(processed_doc)
            
            print(f"\n=== 综合测试结果分析 ===")
            
            # 总结结果
            all_passed = True
            
            # 检查地址传真间距
            spacing_result = results['address_fax_spacing']
            if spacing_result['status'] == 'success':
                print("✅ 地址传真间距：成功修复，有4个空格")
            elif spacing_result['status'] == 'failed_no_spacing':
                print("❌ 地址传真间距：修复失败，缺少空格")
                all_passed = False
            else:
                print(f"❌ 地址传真间距：修复失败 ({spacing_result['status']})")
                all_passed = False
            
            if spacing_result['details']:
                print(f"    详情: '{spacing_result['details']}'")
            
            # 检查电话填写
            phone_result = results['phone_filled']
            if phone_result['status'] == 'success':
                print("✅ 电话字段填写：成功修复")
            elif phone_result['status'] == 'failed_not_filled':
                print("❌ 电话字段填写：修复失败，字段未填写")
                all_passed = False
            else:
                print(f"❌ 电话字段填写：修复失败 ({phone_result['status']})")
                all_passed = False
            
            if phone_result['details']:
                print(f"    详情: '{phone_result['details']}'")
            
            # 检查邮件格式
            email_result = results['email_format']
            if email_result['status'] == 'success':
                print("✅ 电子邮件格式：成功修复")
            elif email_result['status'] == 'failed_appended':
                print("❌ 电子邮件格式：修复失败，仍追加在末尾")
                all_passed = False
            else:
                print(f"❌ 电子邮件格式：修复失败 ({email_result['status']})")
                all_passed = False
            
            if email_result['details']:
                print(f"    详情: '{email_result['details']}'")
            
            # 总体结果
            if all_passed:
                print(f"\n🎉 所有修复都成功！三个问题都已解决。")
                return True
            else:
                print(f"\n💥 仍有问题未完全修复！")
                return False
            
        else:
            print(f"处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = main()
    if success:
        print("\n🎊 综合测试通过！所有修复都有效！")
    else:
        print("\n🚨 综合测试失败，需要继续修复！")
    sys.exit(0 if success else 1)