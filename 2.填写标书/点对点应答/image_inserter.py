#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档级图片插入模块 - 方案4智能混合策略实现
Author: AI标书平台开发团队
Date: 2024-12-09
"""

import os
import re
import json
import logging
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from abc import ABC, abstractmethod
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import datetime

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class InsertionPoint:
    """图片插入点数据结构"""
    paragraph_index: int
    image_type: str
    strategy: str
    marker: Optional[str] = None
    confidence: float = 1.0
    context: Dict = None


@dataclass
class InsertionResult:
    """插入结果"""
    success: bool
    image_type: str
    strategy_used: str
    position: str
    error_message: Optional[str] = None


class InsertionStrategy(ABC):
    """插入策略基类"""
    
    def __init__(self, config: Dict = None):
        self.config = config or {}
        self.enabled = self.config.get('enabled', True)
        self.priority = self.config.get('priority', 999)
    
    @abstractmethod
    def find_insertion_points(self, document: Document) -> List[InsertionPoint]:
        """查找插入点"""
        pass
    
    @abstractmethod
    def insert_image(self, document: Document, point: InsertionPoint, 
                    image_path: str) -> bool:
        """执行插入"""
        pass
    
    def is_enabled(self) -> bool:
        """策略是否启用"""
        return self.enabled


class TextMarkerStrategy(InsertionStrategy):
    """文本标记策略 - 查找并替换文本标记"""
    
    def __init__(self, config: Dict = None):
        super().__init__(config)
        self.markers = self.config.get('markers', {})
    
    def find_insertion_points(self, document: Document) -> List[InsertionPoint]:
        """查找文本标记位置"""
        insertion_points = []
        
        # 遍历所有段落
        for para_idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text
            
            # 检查每种图片类型的标记
            for image_type, marker_list in self.markers.items():
                for marker in marker_list:
                    if marker in text:
                        point = InsertionPoint(
                            paragraph_index=para_idx,
                            image_type=image_type,
                            strategy='text_marker',
                            marker=marker,
                            confidence=1.0
                        )
                        insertion_points.append(point)
                        logger.debug(f"找到文本标记: {marker} at paragraph {para_idx}")
        
        # 也检查表格中的文本
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text
                        
                        for image_type, marker_list in self.markers.items():
                            for marker in marker_list:
                                if marker in text:
                                    point = InsertionPoint(
                                        paragraph_index=-1,  # 表格内特殊标记
                                        image_type=image_type,
                                        strategy='text_marker',
                                        marker=marker,
                                        confidence=1.0,
                                        context={'table_cell': cell, 'paragraph': paragraph}
                                    )
                                    insertion_points.append(point)
        
        return insertion_points
    
    def insert_image(self, document: Document, point: InsertionPoint, 
                    image_path: str) -> bool:
        """在标记位置插入图片"""
        try:
            if point.paragraph_index >= 0:
                # 普通段落
                paragraph = document.paragraphs[point.paragraph_index]
                return self._replace_marker_with_image(paragraph, point.marker, image_path)
            else:
                # 表格内段落
                paragraph = point.context['paragraph']
                return self._replace_marker_with_image(paragraph, point.marker, image_path)
                
        except Exception as e:
            logger.error(f"文本标记插入失败: {str(e)}")
            return False
    
    def _replace_marker_with_image(self, paragraph, marker: str, image_path: str) -> bool:
        """替换段落中的标记为图片"""
        try:
            text = paragraph.text
            
            if marker in text:
                # 保存原始格式
                original_runs = list(paragraph.runs)
                
                # 分割文本
                parts = text.split(marker)
                
                # 清空段落
                paragraph.clear()
                
                # 重建段落内容
                for i, part in enumerate(parts):
                    if part:
                        run = paragraph.add_run(part)
                        # 尝试恢复原始格式
                        if original_runs:
                            self._copy_run_format(original_runs[0], run)
                    
                    # 在标记位置插入图片
                    if i < len(parts) - 1:
                        run = paragraph.add_run()
                        width = self._calculate_image_width()
                        run.add_picture(image_path, width=width)
                
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"替换标记失败: {str(e)}")
            return False
    
    def _copy_run_format(self, source_run, target_run):
        """复制运行格式"""
        try:
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            if source_run.font.bold is not None:
                target_run.font.bold = source_run.font.bold
        except:
            pass
    
    def _calculate_image_width(self) -> Inches:
        """计算图片宽度"""
        width_inches = self.config.get('default_width_inches', 3.5)
        return Inches(width_inches)


class ParagraphPlaceholderStrategy(InsertionStrategy):
    """段落占位符策略 - 识别独立的占位符段落"""
    
    def __init__(self, config: Dict = None):
        super().__init__(config)
        self.placeholders = self.config.get('placeholders', [])
    
    def find_insertion_points(self, document: Document) -> List[InsertionPoint]:
        """查找占位符段落"""
        insertion_points = []
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            
            # 检查是否为占位符段落
            for placeholder in self.placeholders:
                if self._is_placeholder_paragraph(text, placeholder):
                    # 推断图片类型
                    image_type = self._infer_image_type(text)
                    if image_type:
                        point = InsertionPoint(
                            paragraph_index=para_idx,
                            image_type=image_type,
                            strategy='paragraph_placeholder',
                            marker=placeholder,
                            confidence=0.9
                        )
                        insertion_points.append(point)
                        logger.debug(f"找到占位符段落: {placeholder} at paragraph {para_idx}")
        
        return insertion_points
    
    def _is_placeholder_paragraph(self, text: str, placeholder: str) -> bool:
        """判断是否为占位符段落"""
        # 完全匹配或包含关系
        return text == placeholder or placeholder in text
    
    def _infer_image_type(self, text: str) -> Optional[str]:
        """推断图片类型"""
        type_keywords = {
            'business_license': ['营业执照', '执照', 'license'],
            'qualification_cert': ['资质', '证书', '资格', 'qualification'],
            'legal_person_id': ['法人', '身份证', 'ID', '代表人']
        }
        
        text_lower = text.lower()
        for image_type, keywords in type_keywords.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    return image_type
        
        return None
    
    def insert_image(self, document: Document, point: InsertionPoint, 
                    image_path: str) -> bool:
        """替换占位符段落为图片"""
        try:
            paragraph = document.paragraphs[point.paragraph_index]
            
            # 清空段落
            paragraph.clear()
            
            # 添加图片标题（可选）
            if self.config.get('add_title', True):
                title = self._get_image_title(point.image_type)
                paragraph.add_run(f"{title}:\n")
            
            # 插入图片
            run = paragraph.add_run()
            width = Inches(self.config.get('default_width_inches', 4))
            run.add_picture(image_path, width=width)
            
            # 设置段落格式
            if self.config.get('center_align', True):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return True
            
        except Exception as e:
            logger.error(f"占位符段落插入失败: {str(e)}")
            return False
    
    def _get_image_title(self, image_type: str) -> str:
        """获取图片标题"""
        titles = {
            'business_license': '营业执照',
            'qualification_cert': '资质证书',
            'legal_person_id': '法定代表人身份证'
        }
        return titles.get(image_type, '附件')


class DocumentStructureStrategy(InsertionStrategy):
    """文档结构策略 - 基于文档结构插入图片"""
    
    def __init__(self, config: Dict = None):
        super().__init__(config)
        self.after_sections = self.config.get('after_sections', {})
    
    def find_insertion_points(self, document: Document) -> List[InsertionPoint]:
        """基于文档结构查找插入点"""
        insertion_points = []
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            # 检查是否为标题
            if self._is_section_title(paragraph):
                section_title = paragraph.text.strip()
                
                # 查找匹配的规则
                for section_keyword, image_types in self.after_sections.items():
                    if section_keyword in section_title:
                        # 在该标题后插入图片
                        for image_type in image_types:
                            point = InsertionPoint(
                                paragraph_index=para_idx + 1,  # 在标题后插入
                                image_type=image_type,
                                strategy='document_structure',
                                confidence=0.7,
                                context={'section_title': section_title}
                            )
                            insertion_points.append(point)
                            logger.debug(f"找到结构位置: after '{section_title}'")
        
        return insertion_points
    
    def _is_section_title(self, paragraph) -> bool:
        """判断是否为章节标题"""
        # 检查样式
        if paragraph.style and paragraph.style.name:
            if 'Heading' in paragraph.style.name or '标题' in paragraph.style.name:
                return True
        
        # 检查文本模式
        text = paragraph.text.strip()
        title_patterns = [
            r'^第[一二三四五六七八九十\d]+[章节部分]',
            r'^\d+[\.\、\s]+\S+',
            r'^[一二三四五六七八九十]+[\.\、\s]+\S+',
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def insert_image(self, document: Document, point: InsertionPoint, 
                    image_path: str) -> bool:
        """在文档结构位置插入图片"""
        try:
            # 在指定位置创建新段落
            if point.paragraph_index < len(document.paragraphs):
                paragraph = document.paragraphs[point.paragraph_index]
                new_paragraph = paragraph.insert_paragraph_before()
            else:
                new_paragraph = document.add_paragraph()
            
            # 添加图片
            run = new_paragraph.add_run()
            width = Inches(self.config.get('default_width_inches', 3.5))
            run.add_picture(image_path, width=width)
            
            # 设置格式
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加图片说明（可选）
            if self.config.get('add_caption', True):
                caption = self._get_image_caption(point.image_type)
                caption_para = new_paragraph.insert_paragraph_before()
                caption_para.text = caption
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return True
            
        except Exception as e:
            logger.error(f"结构位置插入失败: {str(e)}")
            return False
    
    def _get_image_caption(self, image_type: str) -> str:
        """获取图片说明"""
        captions = {
            'business_license': '企业营业执照',
            'qualification_cert': '企业资质证书',
            'legal_person_id': '法定代表人身份证明'
        }
        return captions.get(image_type, '相关证明文件')


class SmartImageInserter:
    """智能图片插入器 - 方案4主类"""
    
    def __init__(self, config_path: str = None):
        """
        初始化智能图片插入器
        
        Args:
            config_path: 配置文件路径
        """
        self.config = self._load_config(config_path)
        self.strategies = self._init_strategies()
        self.fallback_handler = SmartFallbackHandler(self.config)
        
    def _load_config(self, config_path: str = None) -> Dict:
        """加载配置文件"""
        if config_path and os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        # 默认配置
        return {
            "text_markers": {
                "enabled": True,
                "priority": 1,
                "markers": {
                    "business_license": ["[营业执照]", "【营业执照】", "{营业执照}"],
                    "qualification_cert": ["[资质证书]", "【资质证书】", "{资质证书}"],
                    "legal_person_id": ["[法人身份证]", "【法人身份证】"]
                },
                "default_width_inches": 3.5
            },
            "paragraph_placeholders": {
                "enabled": True,
                "priority": 2,
                "placeholders": [
                    "营业执照附件", "资质证书附件", "法人身份证附件",
                    "此处插入营业执照", "此处插入资质证明"
                ],
                "default_width_inches": 4,
                "add_title": True,
                "center_align": True
            },
            "document_structure": {
                "enabled": True,
                "priority": 3,
                "after_sections": {
                    "公司资质": ["business_license", "qualification_cert"],
                    "企业资质": ["business_license", "qualification_cert"],
                    "附件": ["business_license", "qualification_cert", "legal_person_id"]
                },
                "default_width_inches": 3.5,
                "add_caption": True
            },
            "fallback": {
                "enabled": True,
                "append_to_end": True,
                "create_appendix": True
            },
            "image_settings": {
                "max_width_inches": 5,
                "max_height_inches": 6,
                "quality": 85,
                "formats": ["jpg", "jpeg", "png", "gif"]
            },
            "required_images": ["business_license"],
            "optional_images": ["qualification_cert", "legal_person_id"]
        }
    
    def _init_strategies(self) -> List[InsertionStrategy]:
        """初始化策略列表"""
        strategies = []
        
        # 文本标记策略
        if self.config.get('text_markers', {}).get('enabled'):
            strategies.append(TextMarkerStrategy(self.config['text_markers']))
        
        # 段落占位符策略
        if self.config.get('paragraph_placeholders', {}).get('enabled'):
            strategies.append(ParagraphPlaceholderStrategy(self.config['paragraph_placeholders']))
        
        # 文档结构策略
        if self.config.get('document_structure', {}).get('enabled'):
            strategies.append(DocumentStructureStrategy(self.config['document_structure']))
        
        # 按优先级排序
        strategies.sort(key=lambda s: s.priority)
        
        return strategies
    
    def process_document(self, doc_path: str, company_info: Dict) -> Tuple[str, List[InsertionResult]]:
        """
        处理文档，插入所有图片
        
        Args:
            doc_path: 文档路径
            company_info: 公司信息（包含图片路径）
            
        Returns:
            (处理后文档路径, 插入结果列表)
        """
        try:
            logger.info(f"开始处理文档图片插入: {doc_path}")
            
            # 打开文档
            document = Document(doc_path)
            
            # 插入图片
            results = self._insert_all_images(document, company_info)
            
            # 保存文档
            output_path = self._save_document(doc_path, document)
            
            # 记录结果
            success_count = sum(1 for r in results if r.success)
            logger.info(f"图片插入完成: {success_count}/{len(results)} 成功")
            
            return output_path, results
            
        except Exception as e:
            logger.error(f"处理文档失败: {str(e)}")
            raise
    
    def _insert_all_images(self, document: Document, company_info: Dict) -> List[InsertionResult]:
        """执行所有图片插入"""
        results = []
        inserted_images = set()
        
        # 1. 按策略优先级尝试插入
        for strategy in self.strategies:
            if not strategy.is_enabled():
                continue
            
            # 查找插入点
            insertion_points = strategy.find_insertion_points(document)
            
            for point in insertion_points:
                # 跳过已插入的图片
                if point.image_type in inserted_images:
                    continue
                
                # 获取图片路径
                image_path = self._get_image_path(company_info, point.image_type)
                if not image_path or not os.path.exists(image_path):
                    logger.warning(f"图片文件不存在: {point.image_type}")
                    continue
                
                # 尝试插入
                success = strategy.insert_image(document, point, image_path)
                
                if success:
                    inserted_images.add(point.image_type)
                    results.append(InsertionResult(
                        success=True,
                        image_type=point.image_type,
                        strategy_used=point.strategy,
                        position=f"paragraph {point.paragraph_index}"
                    ))
                    logger.info(f"成功插入 {point.image_type} using {point.strategy}")
        
        # 2. 智能补充未插入的必要图片
        required = set(self.config.get('required_images', []))
        missing = required - inserted_images
        
        if missing and self.config.get('fallback', {}).get('enabled'):
            for image_type in missing:
                image_path = self._get_image_path(company_info, image_type)
                if image_path and os.path.exists(image_path):
                    success = self.fallback_handler.append_to_end(
                        document, image_type, image_path
                    )
                    if success:
                        results.append(InsertionResult(
                            success=True,
                            image_type=image_type,
                            strategy_used='fallback_append',
                            position='end of document'
                        ))
                        logger.info(f"回退策略: 在文档末尾添加 {image_type}")
        
        return results
    
    def _get_image_path(self, company_info: Dict, image_type: str) -> Optional[str]:
        """获取图片文件路径"""
        # 从公司信息中获取资质图片路径
        qualifications = company_info.get('qualifications', {})
        
        # 处理不同的数据格式
        # 格式1: 直接路径字符串 {'business_license': 'path/to/file.png'}
        # 格式2: 字典格式 {'business_license': {'safe_filename': 'xxx.png', ...}}
        # 格式3: 带_path后缀 {'business_license_path': 'path/to/file.png'}
        
        # 尝试直接获取
        image_data = qualifications.get(image_type)
        if image_data:
            if isinstance(image_data, dict):
                # 处理字典格式，提取safe_filename
                safe_filename = image_data.get('safe_filename')
                if safe_filename:
                    # 构建完整路径
                    return os.path.join('uploads', 'qualifications', safe_filename)
            elif isinstance(image_data, str):
                # 直接是路径字符串
                return image_data
        
        # 尝试带_path后缀的键名（向后兼容）
        path_key = f"{image_type}_path"
        if path_key in qualifications:
            return qualifications[path_key]
        
        # 特殊映射处理
        special_mappings = {
            'business_license': ['business_license', 'iso9001', 'iso20000', 'iso27001'],
            'qualification_cert': ['iso9001', 'iso20000', 'iso27001'],
            'legal_person_id': ['auth_id_front', 'auth_id_back']
        }
        
        # 尝试特殊映射
        if image_type in special_mappings:
            for key in special_mappings[image_type]:
                data = qualifications.get(key)
                if data:
                    if isinstance(data, dict) and 'safe_filename' in data:
                        return os.path.join('uploads', 'qualifications', data['safe_filename'])
                    elif isinstance(data, str):
                        return data
        
        return None
    
    def _save_document(self, original_path: str, document: Document) -> str:
        """保存处理后的文档"""
        # 生成输出文件名
        base_name = os.path.basename(original_path)
        name, ext = os.path.splitext(base_name)
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # 修复：直接保存到与原文件相同的目录，避免嵌套目录问题
        # 如果原文件已经在outputs目录，就保存在同一目录
        # 否则创建outputs目录
        original_dir = os.path.dirname(original_path)
        if 'outputs' in original_dir:
            # 已经在outputs目录中，直接使用
            output_dir = original_dir
        else:
            # 不在outputs目录，创建outputs目录
            output_dir = os.path.join(original_dir, 'outputs')
            os.makedirs(output_dir, exist_ok=True)
        
        # 保存文件
        output_path = os.path.join(output_dir, f"{name}_图片插入_{timestamp}{ext}")
        document.save(output_path)
        
        logger.info(f"文档已保存: {output_path}")
        return output_path


class SmartFallbackHandler:
    """智能回退处理器"""
    
    def __init__(self, config: Dict):
        self.config = config
    
    def append_to_end(self, document: Document, image_type: str, 
                     image_path: str) -> bool:
        """在文档末尾添加图片"""
        try:
            # 创建附件区域（如果需要）
            if self.config.get('fallback', {}).get('create_appendix'):
                # 添加分页符
                document.add_page_break()
                
                # 添加附件标题
                title_para = document.add_paragraph()
                title_run = title_para.add_run("附件：企业资质证明")
                title_run.bold = True
                title_run.font.size = Pt(14)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加图片标题
            image_title = self._get_image_title(image_type)
            document.add_paragraph(image_title)
            
            # 添加图片
            img_paragraph = document.add_paragraph()
            run = img_paragraph.add_run()
            
            # 调整图片大小
            width = Inches(self.config.get('image_settings', {}).get('max_width_inches', 4))
            run.add_picture(image_path, width=width)
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return True
            
        except Exception as e:
            logger.error(f"回退策略失败: {str(e)}")
            return False
    
    def _get_image_title(self, image_type: str) -> str:
        """获取图片标题"""
        titles = {
            'business_license': '附件1：营业执照',
            'qualification_cert': '附件2：资质证书',
            'legal_person_id': '附件3：法定代表人身份证'
        }
        return titles.get(image_type, '附件')


# 测试代码
if __name__ == "__main__":
    # 示例公司信息
    test_company_info = {
        "companyName": "智慧科技有限公司",
        "qualifications": {
            "business_license_path": "/path/to/business_license.jpg",
            "qualification_cert_path": "/path/to/qualification.jpg",
            "legal_person_id_path": "/path/to/id.jpg"
        }
    }
    
    # 测试处理
    inserter = SmartImageInserter()
    
    # 如果有测试文档
    test_doc = "test_image.docx"
    if os.path.exists(test_doc):
        output, results = inserter.process_document(test_doc, test_company_info)
        print(f"处理完成: {output}")
        
        for result in results:
            print(f"- {result.image_type}: {result.success} ({result.strategy_used})")