#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试长占位符清理功能 - 模拟实际问题场景
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_long_placeholder_test():
    """创建带长占位符的测试文档"""
    doc = Document()
    doc.add_heading('长占位符测试', 0)
    
    # 模拟实际问题情况：采购编号后有很多下划线
    para = doc.add_paragraph()
    para.add_run("供应商名称: 中国联合网络通信有限公司  __ 采购编号: GXTC-C-251590031_________")
    
    return doc


def main():
    """主测试函数"""
    print("="*60)
    print("长占位符清理测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_long_placeholder_input.docx')
    output_file = os.path.join(current_dir, 'test_long_placeholder_output.docx')
    
    # 创建测试文档
    print("\n1. 创建长占位符测试文档...")
    doc = create_long_placeholder_test()
    doc.save(input_file)
    
    # 分析原始文本
    original_text = doc.paragraphs[1].text
    print(f"   原始文本: '{original_text}'")
    print(f"   下划线数量: {original_text.count('_')}")
    
    # 准备公司信息
    company_info = {
        "companyName": "智慧足迹数据科技有限公司",
    }
    
    # 处理文档
    print("\n2. 使用增强清理功能处理...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            processed_doc = Document(output_file)
            processed_text = processed_doc.paragraphs[1].text
            print(f"   处理后文本: '{processed_text}'")
            print(f"   下划线数量: {processed_text.count('_')}")
            
            # 验证结果
            print("\n3. 验证清理效果:")
            if processed_text.count('_') == 0:
                print("   ✅ 所有下划线占位符已清理")
            else:
                print(f"   ❌ 仍有 {processed_text.count('_')} 个下划线未清理")
                
            if 'GXTC-C-251590031' in processed_text:
                print("   ✅ 采购编号填写正确")
            else:
                print("   ❌ 采购编号填写失败")
                
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)
        print(f"\n4. 输出文件: {output_file}")


if __name__ == '__main__':
    main()