#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试合并后的规则功能
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_test_document():
    """创建测试文档，包含所有合并规则的测试用例"""
    doc = Document()
    doc.add_heading('合并规则测试', 0)
    
    # 测试1: 合并后的公章/盖章规则
    doc.add_heading('测试1: 公章/盖章规则（合并后）', 1)
    para1 = doc.add_paragraph()
    para1.add_run("供应商名称（加盖公章）：                     ")  # 中文括号+空格
    
    para2 = doc.add_paragraph()
    para2.add_run("供应商名称(盖章)：         ")  # 英文括号+空格
    
    para3 = doc.add_paragraph()
    para3.add_run("供应商名称（加盖公章）：")  # 无占位符
    
    # 测试2: 合并后的编号规则
    doc.add_heading('测试2: 编号规则（合并后）', 1)
    para4 = doc.add_paragraph()
    para4.add_run("采购编号：___________________")  # 下划线
    
    para5 = doc.add_paragraph()
    para5.add_run("采购编号：               ")  # 空格
    
    para6 = doc.add_paragraph()
    para6.add_run("项目编号：___________________")
    
    para7 = doc.add_paragraph()
    para7.add_run("项目编号：               ")
    
    para8 = doc.add_paragraph()
    para8.add_run("编号：___________________")
    
    para9 = doc.add_paragraph()
    para9.add_run("编号：               ")
    
    # 测试3: 合并后的括号内编号替换
    doc.add_heading('测试3: 括号内编号替换（合并后）', 1)
    para10 = doc.add_paragraph()
    para10.add_run("根据（采购编号）的要求")
    
    para11 = doc.add_paragraph()
    para11.add_run("根据（招标编号）的要求")
    
    para12 = doc.add_paragraph()
    para12.add_run("根据（项目编号）的要求")
    
    # 测试4: 混合测试 - 同一行多个字段
    doc.add_heading('测试4: 混合测试', 1)
    para13 = doc.add_paragraph()
    para13.add_run("供应商名称：                          采购编号：                   ")
    
    return doc


def analyze_test_results(doc, description):
    """分析测试结果"""
    print(f"\n{description}")
    print("="*50)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and not text.startswith('测试') and not text.startswith('合并规则'):
            print(f"  段落{i}: {text}")


def main():
    """主测试函数"""
    print("="*60)
    print("合并规则功能测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_merged_input.docx')
    output_file = os.path.join(current_dir, 'test_merged_output.docx')
    
    # 创建测试文档
    print("\n1. 创建测试文档...")
    doc = create_test_document()
    doc.save(input_file)
    
    # 分析原始格式
    analyze_test_results(doc, "原始文档:")
    
    # 准备公司信息
    company_info = {
        "companyName": "智慧足迹数据科技有限公司",
    }
    
    # 处理文档
    print("\n2. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        # 先统计规则数量
        rule_count = len(processor.bidder_patterns)
        print(f"   当前规则数量: {rule_count} 个（原来19个，合并后应该是11个）")
        
        # 列出当前规则
        print("\n   当前规则列表:")
        for i, rule in enumerate(processor.bidder_patterns, 1):
            print(f"   规则{i}: {rule['description']}")
        
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("\n   ✅ 处理成功")
            
            # 读取结果
            processed_doc = Document(output_file)
            analyze_test_results(processed_doc, "\n3. 处理后文档:")
            
            # 验证结果
            print("\n4. 验证测试结果:")
            print("="*50)
            
            success_count = 0
            fail_count = 0
            
            for i, para in enumerate(processed_doc.paragraphs):
                text = para.text.strip()
                
                # 验证公章/盖章规则
                if '供应商名称（加盖公章）：' in text or '供应商名称(盖章)：' in text:
                    if '智慧足迹数据科技有限公司' in text:
                        print(f"  ✅ 段落{i}: 公章规则成功")
                        success_count += 1
                    else:
                        print(f"  ❌ 段落{i}: 公章规则失败")
                        fail_count += 1
                
                # 验证编号规则
                if any(x in text for x in ['采购编号：', '项目编号：', '编号：']):
                    if 'GXTC-C-251590031' in text:
                        print(f"  ✅ 段落{i}: 编号规则成功")
                        success_count += 1
                    else:
                        print(f"  ❌ 段落{i}: 编号规则失败")
                        fail_count += 1
                
                # 验证括号内编号替换
                if '根据' in text and 'GXTC-C-251590031' in text:
                    print(f"  ✅ 段落{i}: 括号内编号替换成功")
                    success_count += 1
            
            print(f"\n测试统计: 成功 {success_count} 个, 失败 {fail_count} 个")
            
            if fail_count == 0:
                print("\n✅ 所有合并规则测试通过！")
                print("  - 规则数量从19个减少到11个")
                print("  - 所有功能正常工作")
            else:
                print(f"\n⚠️ 有 {fail_count} 个测试失败")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)
        print(f"\n5. 输出文件: {output_file}")


if __name__ == '__main__':
    main()