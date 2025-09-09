#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
表格处理模块 - 自动识别和填充Word文档中的表格
Author: AI标书平台开发团队
Date: 2024-12-09
"""

import os
import re
import json
import logging
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import datetime

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class TableCell:
    """表格单元格数据结构"""
    content: str
    row_index: int
    col_index: int
    is_empty: bool
    has_underline: bool = False
    is_placeholder: bool = False
    row_span: int = 1
    col_span: int = 1


@dataclass
class FieldMatch:
    """字段匹配结果"""
    field_type: str
    label_position: Tuple[int, int, int]  # (table_idx, row_idx, col_idx)
    target_position: Tuple[int, int, int]  # (table_idx, row_idx, col_idx)
    confidence: float
    matched_text: str


class TableProcessor:
    """表格处理主类"""
    
    def __init__(self, config_path: str = None):
        """
        初始化表格处理器
        
        Args:
            config_path: 配置文件路径
        """
        self.config = self._load_config(config_path)
        self.field_rules = self._load_field_rules()
        self.processed_count = 0
        
    def _load_config(self, config_path: str = None) -> Dict:
        """加载配置文件"""
        if config_path and os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        # 默认配置
        return {
            "processing_rules": {
                "date_formats": ["YYYY-MM-DD", "YYYY年MM月DD日", "YYYY/MM/DD"],
                "currency_formats": ["￥#,###.##", "人民币#,###.##元"],
                "skip_empty_tables": True,
                "min_confidence": 0.6
            }
        }
    
    def _load_field_rules(self) -> Dict:
        """加载字段匹配规则"""
        # 如果配置文件中有字段规则，优先使用配置文件中的规则
        if hasattr(self, 'config') and 'field_rules' in self.config:
            # 合并配置文件中的所有类别的字段规则
            merged_rules = {}
            for category_rules in self.config['field_rules'].values():
                merged_rules.update(category_rules)
            return merged_rules
        
        # 否则使用默认的硬编码规则
        return {
            'company_name': {
                'keywords': ['公司名称', '企业名称', '单位名称', '投标人名称', '投标单位', 
                           '供应商名称', '承包人名称', '乙方', '卖方'],
                'patterns': [r'.*公司.*名称', r'.*企业.*名称', r'.*单位.*名称'],
                'priority': 1,
                'field_key': 'companyName'
            },
            'established_date': {
                'keywords': ['成立日期', '注册日期', '成立时间', '注册时间', '成立年月',
                           '设立日期', '创建日期'],
                'patterns': [r'.*成立.*日期', r'.*注册.*日期', r'.*成立.*时间'],
                'priority': 2,
                'field_key': 'establishDate'
            },
            'legal_representative': {
                'keywords': ['法定代表人', '法人代表', '法人', '法定代表人姓名',
                           '企业法人', '法人姓名'],
                'patterns': [r'.*法.*代表人', r'.*法人.*'],
                'priority': 1,
                'field_key': 'legalRepresentative'
            },
            'registered_capital': {
                'keywords': ['注册资本', '注册资金', '资本金', '注册资本金',
                           '实收资本', '投资总额'],
                'patterns': [r'.*注册.*资本', r'.*注册.*资金'],
                'priority': 2,
                'field_key': 'registeredCapital'
            },
            'business_scope': {
                'keywords': ['经营范围', '业务范围', '主营业务', '营业范围',
                           '经营业务', '业务内容'],
                'patterns': [r'.*经营.*范围', r'.*业务.*范围'],
                'priority': 3,
                'field_key': 'businessScope'
            },
            'company_address': {
                'keywords': ['公司地址', '企业地址', '注册地址', '地址', '住所',
                           '经营地址', '办公地址', '注册住所'],
                'patterns': [r'.*公司.*地址', r'.*注册.*地址', r'.*企业.*地址'],
                'priority': 2,
                'field_key': 'companyAddress'
            },
            'contact_phone': {
                'keywords': ['联系电话', '电话', '联系方式', '手机', '电话号码',
                           '联系人电话', '企业电话', '公司电话'],
                'patterns': [r'.*联系.*电话', r'.*电话.*', r'.*手机.*'],
                'priority': 2,
                'field_key': 'contactPhone'
            },
            'email': {
                'keywords': ['邮箱', '电子邮箱', '邮箱地址', 'E-mail', 'Email',
                           '电子邮件', '企业邮箱'],
                'patterns': [r'.*邮箱.*', r'.*[Ee]-?mail.*'],
                'priority': 3,
                'field_key': 'email'
            },
            'social_credit_code': {
                'keywords': ['统一社会信用代码', '社会信用代码', '信用代码', '组织机构代码',
                           '营业执照号', '营业执照号码', '税务登记号'],
                'patterns': [r'.*信用代码.*', r'.*营业执照.*号'],
                'priority': 1,
                'field_key': 'socialCreditCode'
            },
            'legal_address': {
                'keywords': ['法定地址', '法定住所', '注册住所地址'],
                'patterns': [r'.*法定.*地址', r'.*法定.*住所'],
                'priority': 1,
                'field_key': 'registeredAddress'
            },
            'contact_person': {
                'keywords': ['联系人', '联系人姓名', '项目联系人', '企业联系人'],
                'patterns': [r'.*联系人.*', r'.*联系.*姓名'],
                'priority': 2,
                'field_key': 'authorizedPersonName'
            },
            'fax_number': {
                'keywords': ['传真', '传真号', '传真号码', '企业传真'],
                'patterns': [r'.*传真.*', r'.*传真.*号'],
                'priority': 2,
                'field_key': 'fax'
            },
            'tax_id': {
                'keywords': ['纳税人识别号', '税务登记号', '国税登记号', '税号'],
                'patterns': [r'.*纳税人.*识别号', r'.*税.*登记号'],
                'priority': 1,
                'field_key': 'socialCreditCode'
            },
            'local_tax_id': {
                'keywords': ['地税登记号', '地方税登记号'],
                'patterns': [r'.*地税.*登记号', r'.*地方税.*登记号'],
                'priority': 1,
                'field_key': 'socialCreditCode'
            },
            'industry_certifications': {
                'keywords': ['行业相关认证情况', '资质认证', '行业认证', '相关认证'],
                'patterns': [r'.*行业.*认证.*', r'.*资质.*认证'],
                'priority': 3,
                'field_key': 'qualifications'
            },
            'employee_count': {
                'keywords': ['企业员工总人数', '员工总人数', '职工总数', '员工人数', '从业人数'],
                'patterns': [r'.*员工.*总人数', r'.*职工.*总数', r'.*从业.*人数'],
                'priority': 2,
                'field_key': 'employeeCount'
            },
            'beijing_office_address': {
                'keywords': ['北京市办公场所地址', '北京办公地址', '北京市办公地址', '办公场所地址'],
                'patterns': [r'.*北京.*办公.*地址', r'.*办公.*场所.*地址'],
                'priority': 2,
                'field_key': 'officeAddress'
            },
            'postal_code': {
                'keywords': ['邮政编码', '邮编', '邮码', 'postal code', 'zip code'],
                'patterns': [r'.*邮政编码.*', r'.*邮编.*', r'.*邮码.*'],
                'priority': 2,
                'field_key': 'postalCode'
            },
            'bank_name': {
                'keywords': ['开户银行', '开户行', '银行名称', '基本户开户行',
                           '银行', '开户行名称'],
                'patterns': [r'.*开户.*银行', r'.*开户行.*'],
                'priority': 3,
                'field_key': 'bankName'
            },
            'bank_account': {
                'keywords': ['银行账号', '账号', '银行账户', '基本账户',
                           '对公账户', '企业账户'],
                'patterns': [r'.*银行.*账号', r'.*账户.*'],
                'priority': 3,
                'field_key': 'bankAccount'
            }
        }
    
    def process_document(self, doc_path: str, company_info: Dict, output_path: str = None) -> Dict:
        """
        处理文档中的所有表格
        
        Args:
            doc_path: 文档路径
            company_info: 公司信息字典
            output_path: 输出文档路径（可选）
            
        Returns:
            包含处理结果的字典
        """
        try:
            logger.info(f"开始处理文档: {doc_path}")
            
            # 打开文档
            document = Document(doc_path)
            
            # 提取并处理所有表格
            matched_fields = []
            tables_processed = 0
            
            for table_idx, table in enumerate(document.tables):
                logger.debug(f"处理第 {table_idx + 1} 个表格")
                
                # 识别表格中的字段
                table_matches = self._process_table(table, table_idx)
                if table_matches:
                    matched_fields.extend(table_matches)
                    tables_processed += 1
            
            # 填充内容
            filled_count = self._fill_content(document, matched_fields, company_info)
            
            # 保存结果
            if output_path:
                # 使用指定的输出路径
                document.save(output_path)
                saved_path = output_path
            else:
                # 使用默认的保存方法
                saved_path = self._save_document(doc_path, document)
            
            logger.info(f"处理完成: {tables_processed}个表格，识别{len(matched_fields)}个字段，填充{filled_count}个字段")
            
            # 返回统一格式的结果
            return {
                'success': True,
                'output_path': saved_path,
                'tables_processed': tables_processed,
                'fields_filled': filled_count,
                'fields_matched': len(matched_fields),
                'message': f'成功处理{tables_processed}个表格，填充{filled_count}个字段'
            }
            
        except Exception as e:
            logger.error(f"处理文档失败: {str(e)}")
            # 返回错误结果而不是抛出异常
            return {
                'success': False,
                'output_path': doc_path,
                'tables_processed': 0,
                'fields_filled': 0,
                'fields_matched': 0,
                'error': str(e),
                'message': f'表格处理失败: {str(e)}'
            }
    
    def _process_table(self, table: Table, table_idx: int) -> List[FieldMatch]:
        """处理单个表格"""
        matches = []
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 跳过空单元格
                if not cell.text.strip():
                    continue
                
                # 尝试匹配字段
                field_type = self._match_field(cell.text)
                if field_type:
                    # 查找对应的填写位置
                    target_position = self._find_target_cell(
                        table, row_idx, col_idx
                    )
                    
                    if target_position:
                        match = FieldMatch(
                            field_type=field_type,
                            label_position=(table_idx, row_idx, col_idx),
                            target_position=(table_idx, target_position[0], target_position[1]),
                            confidence=self._calculate_confidence(cell.text, field_type),
                            matched_text=cell.text.strip()
                        )
                        matches.append(match)
                        logger.debug(f"匹配字段: {field_type} <- {cell.text.strip()}")
        
        return matches
    
    def _match_field(self, text: str) -> Optional[str]:
        """匹配字段类型"""
        text = text.strip().lower()
        
        # 移除常见的标点符号
        text_clean = re.sub(r'[：:：\s]', '', text)
        
        # 优先匹配更具体的字段
        field_priority = []
        
        for field_type, rules in self.field_rules.items():
            # 关键词精确匹配
            for keyword in rules['keywords']:
                keyword_clean = re.sub(r'[：:：\s]', '', keyword.lower())
                if keyword_clean == text_clean:
                    return field_type  # 完全匹配立即返回
                elif keyword_clean in text_clean:
                    field_priority.append((field_type, rules['priority'], len(keyword_clean)))
            
            # 正则模式匹配
            for pattern in rules['patterns']:
                if re.search(pattern, text, re.IGNORECASE):
                    field_priority.append((field_type, rules['priority'], len(pattern)))
        
        # 如果有多个匹配，选择优先级最高且关键词最长的
        if field_priority:
            # 按优先级升序、关键词长度降序排序
            field_priority.sort(key=lambda x: (x[1], -x[2]))
            return field_priority[0][0]
        
        return None
    
    def _find_target_cell(self, table: Table, row_idx: int, col_idx: int) -> Optional[Tuple[int, int]]:
        """
        查找字段对应的填写单元格
        
        Returns:
            (row_idx, col_idx) 或 None
        """
        rows = table.rows
        
        # 策略1: 检查右侧单元格（水平布局）
        if col_idx + 1 < len(rows[row_idx].cells):
            right_cell = rows[row_idx].cells[col_idx + 1]
            if self._is_fillable_cell(right_cell):
                return (row_idx, col_idx + 1)
        
        # 策略2: 检查下方单元格（垂直布局）
        if row_idx + 1 < len(rows):
            below_cell = rows[row_idx + 1].cells[col_idx]
            if self._is_fillable_cell(below_cell):
                return (row_idx + 1, col_idx)
        
        # 策略3: 检查同行的后续单元格（跳过可能的说明文字）
        if col_idx + 2 < len(rows[row_idx].cells):
            next_cell = rows[row_idx].cells[col_idx + 2]
            if self._is_fillable_cell(next_cell):
                return (row_idx, col_idx + 2)
        
        return None
    
    def _is_fillable_cell(self, cell: _Cell) -> bool:
        """判断单元格是否为可填写区域"""
        text = cell.text.strip()
        
        # 空单元格
        if not text:
            return True
        
        # 包含下划线或占位符
        fillable_patterns = [
            r'^_+$',  # 纯下划线
            r'^\.+$',  # 纯点
            r'^-+$',  # 纯横线
            r'^\s*$',  # 纯空白
            r'^（\s*）$',  # 空括号
            r'^\[\s*\]$',  # 空方括号
            r'^【\s*】$',  # 空中括号
            r'请填写',
            r'请输入',
            r'待填',
            r'此处填写'
        ]
        
        for pattern in fillable_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _calculate_confidence(self, text: str, field_type: str) -> float:
        """计算匹配置信度"""
        text_clean = re.sub(r'[：:：\s]', '', text.strip().lower())
        rules = self.field_rules[field_type]
        
        # 精确匹配关键词
        for keyword in rules['keywords']:
            keyword_clean = re.sub(r'[：:：\s]', '', keyword.lower())
            if keyword_clean == text_clean:
                return 1.0
            elif keyword_clean in text_clean:
                return 0.9
        
        # 模式匹配
        for pattern in rules['patterns']:
            if re.search(pattern, text_clean, re.IGNORECASE):
                return 0.7
        
        return 0.5
    
    def _fill_content(self, document: Document, matched_fields: List[FieldMatch], 
                     company_info: Dict) -> int:
        """填充匹配到的字段"""
        filled_count = 0
        
        for match in matched_fields:
            # 检查置信度
            if match.confidence < self.config['processing_rules']['min_confidence']:
                logger.debug(f"跳过低置信度字段: {match.field_type} ({match.confidence:.2f})")
                continue
            
            # 获取填充内容
            content = self._get_fill_content(match.field_type, company_info)
            if content is None:
                logger.debug(f"无内容可填充: {match.field_type}")
                continue
            
            # 填充到目标单元格
            table_idx, row_idx, col_idx = match.target_position
            try:
                cell = document.tables[table_idx].rows[row_idx].cells[col_idx]
                
                # 清空原内容
                cell.text = ""
                
                # 填充新内容
                if match.field_type in ['established_date']:
                    # 格式化日期
                    content = self._format_date(content)
                elif match.field_type in ['registered_capital']:
                    # 格式化货币
                    content = self._format_currency(content)
                
                # 设置内容
                cell.text = str(content)
                
                # 保持原有格式（如果有）
                if cell.paragraphs:
                    self._preserve_cell_format(cell.paragraphs[0])
                
                filled_count += 1
                logger.debug(f"已填充: {match.field_type} = {content}")
                
            except Exception as e:
                logger.error(f"填充失败 {match.field_type}: {str(e)}")
                continue
        
        return filled_count
    
    def _get_fill_content(self, field_type: str, company_info: Dict) -> Optional[str]:
        """根据字段类型获取填充内容"""
        # 获取字段对应的数据键名
        field_key = self.field_rules[field_type].get('field_key')
        if not field_key:
            return None
        
        # 特殊处理资质认证情况
        if field_key == 'qualifications':
            qualifications = company_info.get('qualifications', [])
            if isinstance(qualifications, list):
                # 过滤掉营业执照和身份证
                filtered_qualifications = [
                    qual for qual in qualifications 
                    if qual and '营业执照' not in qual and '身份证' not in qual
                ]
                return '、'.join(filtered_qualifications) if filtered_qualifications else '无'
            elif isinstance(qualifications, str):
                return qualifications if qualifications else '无'
            else:
                return '无'
        
        # 从公司信息中获取值
        value = company_info.get(field_key)
        
        # 处理嵌套字段
        if value is None and '.' in field_key:
            keys = field_key.split('.')
            value = company_info
            for key in keys:
                if isinstance(value, dict):
                    value = value.get(key)
                else:
                    value = None
                    break
        
        return value
    
    def _format_date(self, date_value: Any) -> str:
        """格式化日期"""
        if isinstance(date_value, str):
            # 尝试解析各种日期格式
            date_formats = [
                '%Y-%m-%d',
                '%Y/%m/%d',
                '%Y年%m月%d日',
                '%Y.%m.%d'
            ]
            
            for fmt in date_formats:
                try:
                    date_obj = datetime.datetime.strptime(date_value, fmt)
                    # 统一输出格式
                    return date_obj.strftime('%Y年%m月%d日')
                except:
                    continue
            
            # 如果无法解析，返回原值
            return date_value
        
        elif isinstance(date_value, (datetime.date, datetime.datetime)):
            return date_value.strftime('%Y年%m月%d日')
        
        return str(date_value)
    
    def _format_currency(self, value: Any) -> str:
        """格式化货币"""
        if isinstance(value, (int, float)):
            # 格式化为万元
            if value >= 10000:
                return f"{value/10000:.2f}万元"
            else:
                return f"{value}元"
        
        # 如果已经是字符串，直接返回
        return str(value)
    
    def _preserve_cell_format(self, paragraph: Paragraph):
        """保持单元格格式"""
        # 保持对齐方式、字体等格式
        # 这里可以根据需要扩展
        pass
    
    def _save_document(self, original_path: str, document: Document) -> str:
        """保存处理后的文档"""
        # 生成输出文件名
        base_name = os.path.basename(original_path)
        name, ext = os.path.splitext(base_name)
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # 创建输出目录
        output_dir = os.path.join(
            os.path.dirname(original_path),
            'outputs',
            'tables'
        )
        os.makedirs(output_dir, exist_ok=True)
        
        # 保存文件
        output_path = os.path.join(output_dir, f"{name}_表格填充_{timestamp}{ext}")
        document.save(output_path)
        
        logger.info(f"文档已保存: {output_path}")
        return output_path
    
    def analyze_tables(self, doc_path: str) -> Dict:
        """
        分析文档中的表格结构（用于调试和预览）
        
        Returns:
            表格分析结果
        """
        document = Document(doc_path)
        analysis = {
            'total_tables': len(document.tables),
            'tables': []
        }
        
        for table_idx, table in enumerate(document.tables):
            table_info = {
                'table_index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'matched_fields': []
            }
            
            # 识别字段
            matches = self._process_table(table, table_idx)
            for match in matches:
                table_info['matched_fields'].append({
                    'field_type': match.field_type,
                    'matched_text': match.matched_text,
                    'confidence': match.confidence,
                    'position': f"({match.label_position[1]}, {match.label_position[2]})"
                })
            
            analysis['tables'].append(table_info)
        
        return analysis


# 测试代码
if __name__ == "__main__":
    # 示例公司信息
    test_company_info = {
        "companyName": "智慧科技有限公司",
        "establishDate": "2015-12-18",
        "legalRepresentative": "张三",
        "registeredCapital": "10000000",
        "businessScope": "软件开发、技术服务、技术咨询",
        "companyAddress": "北京市海淀区中关村大街1号",
        "contactPhone": "010-12345678",
        "email": "contact@example.com",
        "socialCreditCode": "91110108MA00XXXX00",
        "bankName": "中国工商银行北京分行",
        "bankAccount": "1234567890123456789"
    }
    
    # 测试处理
    processor = TableProcessor()
    
    # 如果有测试文档，可以运行
    test_doc = "test_table.docx"
    if os.path.exists(test_doc):
        result = processor.process_document(test_doc, test_company_info)
        print(f"处理完成: {result}")
        
        # 分析表格
        analysis = processor.analyze_tables(test_doc)
        print(f"表格分析: {json.dumps(analysis, indent=2, ensure_ascii=False)}")