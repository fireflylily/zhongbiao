#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
综合测试 - 验证所有修复
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_test_document():
    """创建综合测试文档"""
    doc = Document()
    doc.add_heading('综合测试文档', 0)
    
    # 测试1: 供应商名称和采购编号在同一行
    para = doc.add_paragraph()
    para.add_run("供应商名称:")
    para.add_run("_______________________")
    para.add_run("采购编号:")
    para.add_run("_______________________")
    
    # 测试2: 跨run的供应商名称
    para2 = doc.add_paragraph()
    para2.add_run("供应商名称")
    para2.add_run("：")
    
    # 测试3: 带下划线的采购编号
    para3 = doc.add_paragraph()
    para3.add_run("采购编号：")
    run = para3.add_run("___________________")
    
    # 测试4: 括号内的采购编号
    doc.add_paragraph("根据贵方的竞争性磋商公告（采购编号），我方提交响应文件。")
    
    return doc


def analyze_formatting(paragraph, description):
    """分析段落格式"""
    print(f"\n{description}")
    print(f"  文本: '{paragraph.text}'")
    for i, run in enumerate(paragraph.runs):
        if run.text:
            format_info = []
            if run.bold:
                format_info.append("粗体")
            if run.italic:
                format_info.append("斜体")
            if run.underline:
                format_info.append("下划线")
            format_str = f" [{', '.join(format_info)}]" if format_info else ""
            print(f"    Run{i}: '{run.text}'{format_str}")


def main():
    """主测试函数"""
    print("="*60)
    print("综合修复测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_comprehensive_input.docx')
    output_file = os.path.join(current_dir, 'test_comprehensive_output.docx')
    
    # 创建测试文档
    print("\n1. 创建测试文档...")
    doc = create_test_document()
    doc.save(input_file)
    
    # 分析原始格式
    print("\n2. 原始文档格式:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            analyze_formatting(para, f"段落{i}")
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
        "fixedPhone": "010-66258899",
        "registeredAddress": "北京市西城区金融大街21号",
    }
    
    # 处理文档
    print("\n3. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            print("\n4. 处理后文档格式:")
            processed_doc = Document(output_file)
            
            for i, para in enumerate(processed_doc.paragraphs):
                if para.text.strip():
                    analyze_formatting(para, f"段落{i}")
            
            # 验证结果
            print("\n5. 验证结果:")
            issues = []
            
            # 检查第一个段落（供应商名称和采购编号）
            para1 = processed_doc.paragraphs[1]
            if '供应商名称' in para1.text and '采购编号' in para1.text:
                # 检查是否在同一行
                if '\n' not in para1.text:
                    print("   ✅ 供应商名称和采购编号在同一行")
                else:
                    issues.append("供应商名称和采购编号被分成多行")
                
                # 检查格式
                for run in para1.runs:
                    if '采购编号' in run.text or 'GXTC-C-251590031' in run.text:
                        if run.underline:
                            issues.append("采购编号或其值意外加了下划线")
                
                # 检查是否正确替换
                if '中国联合网络通信有限公司' in para1.text:
                    if '_____' not in para1.text:
                        print("   ✅ 供应商名称正确替换了占位符")
                    else:
                        issues.append("供应商名称未完全替换占位符")
                
                if 'GXTC-C-251590031' in para1.text:
                    if '_____' not in para1.text or para1.text.count('___') <= 1:
                        print("   ✅ 采购编号正确替换了占位符")
                    else:
                        issues.append("采购编号未完全替换占位符")
            
            # 检查跨run的供应商名称
            para2 = processed_doc.paragraphs[2]
            if '中国联合网络通信有限公司' in para2.text:
                print("   ✅ 跨run的供应商名称已填写")
                # 检查格式
                has_underline = any(run.underline for run in para2.runs if '中国联合网络通信有限公司' in run.text)
                if has_underline:
                    issues.append("供应商名称意外加了下划线")
            
            # 显示问题
            if issues:
                print("\n⚠️  发现的问题:")
                for issue in issues:
                    print(f"   - {issue}")
            else:
                print("\n✅ 所有测试通过！")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理文件
        for f in [input_file]:
            if os.path.exists(f):
                os.remove(f)
        print(f"\n6. 输出文件: {output_file}")


if __name__ == '__main__':
    main()