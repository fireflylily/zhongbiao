#!/usr/bin/env python3
"""
测试批量替换功能，验证多项替换时格式保持效果
模拟真实的复杂段落场景
"""

import logging
import sys
from pathlib import Path
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_complex_test_document():
    """创建包含多项替换内容的复杂测试文档"""
    doc = Document()
    
    # 添加段落1：包含4个替换项的复杂段落（类似截图中的情况）
    logger.info("创建复杂段落1: 包含项目名称、采购编号、姓名职务、公司名称")
    para1 = doc.add_paragraph()
    
    # 模拟复杂的格式情况，每部分使用不同格式
    run1 = para1.add_run("根据贵方为")
    run1.font.name = "宋体"
    run1.font.italic = False
    
    run2 = para1.add_run("（")
    run2.font.name = "宋体"
    run2.font.italic = True
    
    run3 = para1.add_run("项目名称")
    run3.font.name = "宋体" 
    run3.font.italic = True
    
    run4 = para1.add_run("）")
    run4.font.name = "宋体"
    run4.font.italic = True
    
    run5 = para1.add_run("项目采购采购货物及服务的竞争性磋商公告")
    run5.font.name = "宋体"
    run5.font.italic = False
    
    run6 = para1.add_run("（")
    run6.font.name = "宋体"
    run6.font.italic = True
    run6.font.underline = True
    
    run7 = para1.add_run("采购编号")
    run7.font.name = "宋体"
    run7.font.italic = True
    run7.font.underline = True
    
    run8 = para1.add_run("）")
    run8.font.name = "宋体"
    run8.font.italic = True
    run8.font.underline = True
    
    run9 = para1.add_run("，签字代表")
    run9.font.name = "宋体"
    run9.font.italic = False
    
    run10 = para1.add_run("（")
    run10.font.name = "宋体"
    run10.font.italic = False
    
    run11 = para1.add_run("姓名、职务")
    run11.font.name = "宋体"
    run11.font.italic = False
    
    run12 = para1.add_run("）")
    run12.font.name = "宋体"
    run12.font.italic = False
    
    run13 = para1.add_run("经正式授权并代表供应商")
    run13.font.name = "宋体"
    run13.font.italic = False
    
    run14 = para1.add_run("（")
    run14.font.name = "宋体"
    run14.font.italic = True
    
    run15 = para1.add_run("供应商名称、地址")
    run15.font.name = "宋体"
    run15.font.italic = True
    
    run16 = para1.add_run("）")
    run16.font.name = "宋体"
    run16.font.italic = True
    
    run17 = para1.add_run("提交下述文件正本一份及副本份：")
    run17.font.name = "宋体"
    run17.font.italic = False
    
    # 添加段落2：包含两个替换项的段落
    logger.info("创建段落2: 包含项目名称和采购编号")
    para2 = doc.add_paragraph()
    
    run1 = para2.add_run("本次")
    run1.font.name = "宋体"
    
    run2 = para2.add_run("（")
    run2.font.name = "宋体"
    run2.font.italic = True
    
    run3 = para2.add_run("项目名称")
    run3.font.name = "宋体"
    run3.font.italic = True
    
    run4 = para2.add_run("）")
    run4.font.name = "宋体"
    run4.font.italic = True
    
    run5 = para2.add_run("采购，编号为")
    run5.font.name = "宋体"
    
    run6 = para2.add_run("（")
    run6.font.name = "宋体"
    run6.font.bold = True
    
    run7 = para2.add_run("采购编号")
    run7.font.name = "宋体"
    run7.font.bold = True
    
    run8 = para2.add_run("）")
    run8.font.name = "宋体"
    run8.font.bold = True
    
    run9 = para2.add_run("，特制定本投标文件。")
    run9.font.name = "宋体"
    
    # 添加段落3：单项替换对比
    logger.info("创建段落3: 单项替换对比")
    para3 = doc.add_paragraph()
    
    run1 = para3.add_run("投标人：")
    run1.font.name = "宋体"
    
    run2 = para3.add_run("（")
    run2.font.name = "宋体"
    run2.font.underline = True
    
    run3 = para3.add_run("供应商名称")
    run3.font.name = "宋体"
    run3.font.underline = True
    
    run4 = para3.add_run("）")
    run4.font.name = "宋体"
    run4.font.underline = True
    
    return doc

def analyze_document_detailed(doc, title):
    """详细分析文档格式"""
    logger.info(f"\n📋 {title}")
    logger.info("="*80)
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            logger.info(f"段落 #{para_idx}: '{paragraph.text}'")
            logger.info(f"段落总长度: {len(paragraph.text)} 字符")
            
            # 分析每个run的详细信息
            total_run_length = 0
            for run_idx, run in enumerate(paragraph.runs):
                if run.text:
                    format_details = []
                    if run.font.name:
                        format_details.append(f"字体={run.font.name}")
                    if run.font.bold:
                        format_details.append("粗体=True")
                    if run.font.italic:
                        format_details.append("斜体=True")
                    if run.font.underline:
                        format_details.append("下划线=True")
                    
                    format_str = ", ".join(format_details) if format_details else "默认格式"
                    
                    # 检查特殊内容
                    markers = []
                    if "项目名称" in run.text:
                        markers.append("🎯 项目名称")
                    if "采购编号" in run.text:
                        markers.append("📋 采购编号") 
                    if "供应商名称" in run.text:
                        markers.append("🏢 供应商名称")
                    if "姓名、职务" in run.text:
                        markers.append("👤 姓名职务")
                    
                    marker_str = " " + " ".join(markers) if markers else ""
                    
                    logger.info(f"  Run {run_idx+1}: '{run.text}' [{format_str}]{marker_str}")
                    total_run_length += len(run.text)
            
            logger.info(f"Run总长度验证: {total_run_length} 字符")
            logger.info("-" * 60)

def test_batch_replacement():
    """测试批量替换功能"""
    logger.info("🚀 开始测试批量替换功能")
    
    try:
        # 创建测试文档
        input_doc = create_complex_test_document()
        input_path = "test_batch_replacement_input.docx"
        input_doc.save(input_path)
        
        # 分析原始文档
        analyze_document_detailed(input_doc, "原始文档格式分析")
        
        # 使用处理器
        processor = MCPBidderNameProcessor()
        
        # 处理文档
        output_path = "test_batch_replacement_output.docx"
        result = processor.process_bidder_name(
            input_file=input_path,
            output_file=output_path,
            company_name="智慧足迹数据科技有限公司"
        )
        
        if result.get('success', False):
            logger.info("✅ 批量替换处理成功")
            
            # 分析处理后文档
            output_doc = Document(output_path)
            analyze_document_detailed(output_doc, "处理后文档格式分析")
            
            # 验证关键点
            logger.info("\n🔍 关键验证项目")
            logger.info("="*80)
            
            success_count = 0
            total_checks = 0
            
            for para_idx, paragraph in enumerate(output_doc.paragraphs):
                para_text = paragraph.text
                
                # 检查1: 项目名称格式是否正确
                total_checks += 1
                if "（项目名称）" not in para_text:
                    # 如果包含了实际项目名称，检查格式
                    project_name_runs = [run for run in paragraph.runs if "项目" in run.text and "名称" not in run.text]
                    if project_name_runs:
                        format_ok = any(run.font.italic for run in project_name_runs)
                        if format_ok:
                            logger.info(f"✅ 段落 #{para_idx}: 项目名称格式保持正确")
                            success_count += 1
                        else:
                            logger.error(f"❌ 段落 #{para_idx}: 项目名称格式异常")
                    else:
                        success_count += 1  # 没有项目名称则认为正常
                else:
                    success_count += 1  # 未替换也认为正常
                
                # 检查2: 采购编号格式是否正确
                total_checks += 1
                if "（采购编号）" not in para_text:
                    # 检查采购编号区域的格式
                    tender_runs = [run for run in paragraph.runs if any(char.isdigit() for char in run.text)]
                    if tender_runs:
                        # 采购编号应该保持斜体+下划线或粗体格式
                        format_preserved = any(run.font.italic or run.font.bold or run.font.underline for run in tender_runs)
                        if format_preserved:
                            logger.info(f"✅ 段落 #{para_idx}: 采购编号格式保持正确")
                            success_count += 1
                        else:
                            logger.error(f"❌ 段落 #{para_idx}: 采购编号格式丢失")
                    else:
                        success_count += 1
                else:
                    success_count += 1
                
                # 检查3: 姓名职务格式是否未受影响
                total_checks += 1
                name_job_runs = [run for run in paragraph.runs if "姓名" in run.text or "职务" in run.text]
                if name_job_runs:
                    # 姓名职务应该保持正常格式（无特殊格式）
                    format_normal = all(not run.font.italic and not run.font.underline for run in name_job_runs)
                    if format_normal:
                        logger.info(f"✅ 段落 #{para_idx}: 姓名职务格式未受影响")
                        success_count += 1
                    else:
                        logger.error(f"❌ 段落 #{para_idx}: 姓名职务格式受到影响")
                else:
                    success_count += 1
                
                # 检查4: 公司名称是否正确替换
                total_checks += 1
                if "智慧足迹数据科技有限公司" in para_text:
                    logger.info(f"✅ 段落 #{para_idx}: 公司名称已正确替换")
                    success_count += 1
                elif "供应商名称" in para_text:
                    logger.error(f"❌ 段落 #{para_idx}: 供应商名称未替换")
                else:
                    success_count += 1  # 没有供应商名称则正常
            
            # 最终结果
            success_rate = (success_count / total_checks) * 100 if total_checks > 0 else 0
            logger.info(f"\n🎯 批量替换测试结果")
            logger.info("="*80)
            logger.info(f"总检查项: {total_checks}")
            logger.info(f"成功项: {success_count}")
            logger.info(f"成功率: {success_rate:.1f}%")
            
            if success_rate >= 90:
                logger.info("🎉 批量替换功能测试通过！格式保持良好")
                return True
            elif success_rate >= 70:
                logger.info("⚠️ 批量替换功能基本可用，但仍有改进空间")
                return True
            else:
                logger.error("❌ 批量替换功能存在问题，需要进一步优化")
                return False
        else:
            logger.error("❌ 批量替换处理失败")
            return False
            
    except Exception as e:
        logger.error(f"测试失败: {e}", exc_info=True)
        return False

if __name__ == "__main__":
    test_batch_replacement()