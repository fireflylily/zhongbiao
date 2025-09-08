#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
采购需求原地应答插入系统
功能：在原文档中每个需求条目后直接插入"应答：满足。"格式的技术应答
保持原文档格式和结构不变，仅做点对点插入
"""

import requests
import json
import logging
import os
import re
from datetime import datetime
from typing import Dict, List

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('inline_reply.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class InlineReplyProcessor:
    """原地应答插入处理器"""
    
    def __init__(self):
        self.template_file = 'config/response_templates.json'
        self.patterns_file = 'config/requirement_patterns.json'
        
        # 加载配置
        self.load_templates()
        self.load_patterns()
    
    def load_templates(self):
        """加载应答模板"""
        try:
            if os.path.exists(self.template_file):
                with open(self.template_file, 'r', encoding='utf-8') as f:
                    self.templates = json.load(f)
                    logger.info("应答模板加载成功")
            else:
                self.templates = self.get_default_templates()
                logger.warning("模板文件不存在，使用默认模板")
        except Exception as e:
            logger.error(f"加载应答模板失败: {e}")
            self.templates = self.get_default_templates()
    
    def load_patterns(self):
        """加载需求识别模式"""
        try:
            if os.path.exists(self.patterns_file):
                with open(self.patterns_file, 'r', encoding='utf-8') as f:
                    self.patterns = json.load(f)
                    logger.info("需求识别模式加载成功")
            else:
                self.patterns = self.get_default_patterns()
                logger.warning("模式文件不存在，使用默认模式")
        except Exception as e:
            logger.error(f"加载需求识别模式失败: {e}")
            self.patterns = self.get_default_patterns()
    
    def get_default_templates(self) -> Dict:
        """获取默认应答模板"""
        return {
            "硬件配置": "应答：满足。我方提供的硬件设备配置完全满足采购需求指标。",
            "软件功能": "应答：满足。我方系统具备相关功能，能够满足业务需求。",
            "性能指标": "应答：满足。系统性能指标完全符合并超过采购需求标准。",
            "技术规范": "应答：满足。我方产品严格遵循相关技术标准和规范。",
            "服务保障": "应答：满足。我方提供全方位服务保障，确保项目顺利实施。",
            "资质证明": "应答：满足。我方具备相关资质证书，完全满足要求。",
            "安全要求": "应答：满足。我方系统采用多层安全防护机制，符合安全要求。",
            "通用模板": "应答：满足。我方将根据具体需求提供相应的技术方案。"
        }
    
    def get_default_patterns(self) -> Dict:
        """获取默认需求识别模式"""
        return {
            "编号模式": [
                r'^(\d+)\s*[、．.]',  # 1、 1. 1．
                r'^(\d+\.\d+)\s*[、．.]',  # 1.1、 1.2.
                r'^\((\d+)\)',  # (1) (2)
                r'^([A-Z])\)',  # A) B)
                r'^([a-z])\)',  # a) b)
            ],
            "关键词": [
                "要求", "需求", "应", "必须", "应当", "需要", "具备",
                "支持", "提供", "实现", "满足", "符合", "遵循",
                "不少于", "不低于", "≥", ">=", "以上"
            ]
        }
    
    def llm_callback(self, prompt: str, purpose: str = "应答生成") -> str:
        """
        调用LLM API生成应答
        """
        url = "https://api.oaipro.com/v1/chat/completions"
        headers = {
            "Authorization": "Bearer sk-xxx",  # 需要配置有效的API密钥
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": "gpt-5",
            "messages": [
                {"role": "system", "content": "你是一名专业的技术方案专家，擅长为采购需求提供专业的技术应答。"},
                {"role": "user", "content": prompt}
            ],
            "max_completion_tokens": 200  # 限制应答长度
        }
        
        try:
            logger.info(f"[LLM调用] {purpose}")
            response = requests.post(url, headers=headers, json=payload, timeout=60)
            
            if response.status_code != 200:
                logger.error(f"API调用失败: {response.status_code}")
                return self._get_fallback_response()
            
            result = response.json()
            if "choices" in result and len(result["choices"]) > 0:
                content = result["choices"][0].get("message", {}).get("content", "").strip()
                
                # 确保以"应答：满足。"开头
                if not content.startswith("应答：满足。"):
                    if content.startswith("应答："):
                        content = content.replace("应答：", "应答：满足。", 1)
                    else:
                        content = f"应答：满足。{content}"
                
                return content
            else:
                return self._get_fallback_response()
                
        except Exception as e:
            logger.error(f"LLM调用失败: {e}")
            return self._get_fallback_response()
    
    def _get_fallback_response(self) -> str:
        """获取备用应答"""
        return "应答：满足。我方将根据具体需求提供相应的技术方案，完全满足采购要求。"
    
    def insert_paragraph_after(self, paragraph, text=None, style=None):
        """
        在指定段落后插入新段落（复用insert_reply.py的方法）
        """
        if not DOCX_AVAILABLE:
            logger.error("未安装python-docx库，无法进行段落插入操作")
            return None
        
        try:
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = paragraph._parent.add_paragraph()
            new_para._p = new_p
            if text:
                new_para.add_run(text)
            if style:
                new_para.style = style
            return new_para
        except Exception as e:
            logger.error(f"段落插入失败: {e}")
            return None
    
    def is_requirement_paragraph(self, para) -> bool:
        """
        判断段落是否为需求条目
        基于insert_reply.py的逻辑，但更精确
        """
        text = para.text.strip()
        if not text:
            return False
        
        style_name = para.style.name if para.style else ""
        alignment = para.paragraph_format.alignment
        
        # 排除条件（与insert_reply.py保持一致）
        if (
            "Heading" in style_name or  # 标题样式
            "标题" in style_name or      # 中文标题样式
            "TOC" in style_name or       # 目录样式
            "目录" in text or            # 包含"目录"的文本
            alignment == 1 or            # 居中对齐（通常是标题）
            re.match(r"^(\d+(\.\d+)*[、\.]?|（[一二三四五六七八九十]+）)", text)  # 章节编号
        ):
            return False
        
        # 包含条件：包含需求关键词或编号模式
        has_keyword = any(keyword in text for keyword in self.patterns["关键词"])
        has_number = any(re.match(pattern, text) for pattern in self.patterns["编号模式"])
        
        # 长度过滤：内容过短可能是标题或无关文本
        is_long_enough = len(text) > 15
        
        return (has_keyword or has_number) and is_long_enough
    
    def classify_requirement(self, text: str) -> str:
        """
        简单分类需求类型
        """
        content_lower = text.lower()
        
        # 硬件相关
        if any(kw in content_lower for kw in ["cpu", "内存", "存储", "硬盘", "服务器", "设备", "配置"]):
            return "硬件配置"
        
        # 软件相关
        if any(kw in content_lower for kw in ["软件", "系统", "应用", "功能", "模块"]):
            return "软件功能"
            
        # 性能相关
        if any(kw in content_lower for kw in ["性能", "速度", "并发", "响应", "不少于", "不低于"]):
            return "性能指标"
            
        # 服务相关
        if any(kw in content_lower for kw in ["服务", "支持", "培训", "维护", "实施"]):
            return "服务保障"
            
        # 资质相关  
        if any(kw in content_lower for kw in ["资质", "证书", "认证", "经验", "案例"]):
            return "资质证明"
            
        # 安全相关
        if any(kw in content_lower for kw in ["安全", "加密", "认证", "防护"]):
            return "安全要求"
        
        return "通用模板"
    
    def generate_response_for_requirement(self, text: str) -> str:
        """
        为单个需求生成应答
        """
        req_type = self.classify_requirement(text)
        
        # 使用LLM生成具体应答
        prompt = f"""
针对以下采购需求，生成专业的技术应答。

需求内容：{text}

要求：
1. 应答必须以"应答：满足。"开头
2. 后续内容要具体、专业、简洁
3. 控制在100字以内
4. 不要包含其他解释

请直接返回应答内容：
"""
        
        try:
            response = self.llm_callback(prompt, "需求应答")
            return response
        except Exception as e:
            logger.error(f"生成应答失败: {e}")
            # 使用模板应答
            template = self.templates.get(req_type, self.templates["通用模板"])
            return template
    
    def process_document_inline(self, input_file: str, output_file: str = None) -> str:
        """
        处理文档，原地插入应答
        """
        if not DOCX_AVAILABLE:
            raise Exception("未安装python-docx库，请安装：pip install python-docx")
        
        logger.info(f"开始处理文档: {input_file}")
        
        try:
            # 读取文档
            doc = Document(input_file)
            logger.info(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")
            
            # 统计需求条目
            requirement_count = 0
            processed_count = 0
            
            # 从后往前遍历，避免插入新段落影响索引
            paragraphs = list(doc.paragraphs)
            
            for i in range(len(paragraphs) - 1, -1, -1):
                para = paragraphs[i]
                
                if self.is_requirement_paragraph(para):
                    requirement_count += 1
                    text = para.text.strip()
                    
                    logger.info(f"处理需求 {requirement_count}: {text[:50]}...")
                    
                    # 生成应答
                    response = self.generate_response_for_requirement(text)
                    
                    # 在需求段落后插入应答
                    reply_para = self.insert_paragraph_after(para, response)
                    
                    if reply_para:
                        processed_count += 1
                        logger.info(f"已插入应答: {response[:50]}...")
                    else:
                        logger.error(f"插入应答失败: {text[:30]}...")
            
            # 生成输出文件名
            if not output_file:
                base_name = os.path.splitext(os.path.basename(input_file))[0]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_file = f"{base_name}-点对点应答-{timestamp}.docx"
            
            # 保存文档
            doc.save(output_file)
            
            logger.info(f"处理完成:")
            logger.info(f"  识别需求条目: {requirement_count}")
            logger.info(f"  成功插入应答: {processed_count}")
            logger.info(f"  输出文件: {output_file}")
            
            return output_file
            
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            raise

def main():
    """主函数"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python3 inline_reply_processor.py <输入文档路径> [输出文档路径]")
        print("例如: python3 inline_reply_processor.py 采购需求.docx")
        return
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        processor = InlineReplyProcessor()
        result_file = processor.process_document_inline(input_file, output_file)
        
        print(f"\n✅ 处理完成!")
        print(f"📄 输入文件: {input_file}")
        print(f"📝 输出文件: {result_file}")
        print(f"📋 详细日志: inline_reply.log")
        print("\n🎉 点对点应答插入成功！")
        
    except Exception as e:
        logger.error(f"程序执行失败: {e}")
        print(f"❌ 程序执行失败: {e}")

if __name__ == "__main__":
    main()