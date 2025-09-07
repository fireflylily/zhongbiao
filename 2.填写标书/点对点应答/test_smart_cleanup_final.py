#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试智能分隔符保留的占位符清理方法最终效果
"""

import os
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_test_document():
    """创建全面的测试文档，涵盖各种复杂情况"""
    doc = Document()
    doc.add_heading('智能占位符清理测试文档', 0)
    
    # 添加前置段落
    for i in range(29):
        doc.add_paragraph(f'段落 #{i+1}: 前置内容...')
    
    # 段落#30 - 复杂的多字段单行格式
    # 模拟采购编号已处理，供应商名称待处理的情况
    para30 = doc.add_paragraph()
    para30.add_run("                      ")     # 前导格式空格 - 应保留
    para30.add_run("供应商")                    # 分散的标签 - 第一部分
    para30.add_run("名称：")                   # 分散的标签 - 第二部分
    para30.add_run("   ")                      # 填充空格 - 应清理
    para30.add_run("采购编号：")                # 第二个标签
    para30.add_run("GXTC-C-251590031")         # 已填写的内容
    para30.add_run("          ")               # 占位符空格 - 应清理
    
    # 段落#31 - 单标签格式测试
    para31 = doc.add_paragraph()
    para31.add_run("          ")               # 前导格式空格 - 应保留
    para31.add_run("投标人名称：")              # 完整标签
    para31.add_run("________________________") # 占位符 - 应清理
    
    # 段落#32 - 公章格式测试  
    para32 = doc.add_paragraph()
    para32.add_run("供应商名称（加盖公章）：")    # 特殊格式标签
    para32.add_run("                    ")     # 占位符空格 - 应清理
    
    print("创建的测试文档结构：")
    print(f"段落#30: '{para30.text}'")
    print(f"段落#31: '{para31.text}'")  
    print(f"段落#32: '{para32.text}'")
    
    return doc

def test_comprehensive_smart_cleanup():
    """测试智能占位符清理的全面效果"""
    print("\n" + "="*60)
    print("测试智能分隔符保留的占位符清理方法")
    print("="*60)
    
    # 创建测试文档
    input_file = 'test_smart_cleanup_input.docx'
    output_file = 'test_smart_cleanup_output.docx'
    
    doc = create_test_document()
    doc.save(input_file)
    print(f"已创建测试文档: {input_file}")
    
    # 使用修复后的处理器
    processor = MCPBidderNameProcessor()
    company_name = "智慧足迹数据科技有限公司"
    
    print(f"\n开始处理，使用公司名称: {company_name}")
    
    result = processor.process_bidder_name(
        input_file=input_file,
        output_file=output_file,
        company_name=company_name
    )
    
    print(f"\n处理结果:")
    print(f"  成功: {result.get('success', False)}")
    
    if result.get('error'):
        print(f"  错误: {result.get('error')}")
        return False
    
    stats = result.get('stats', {})
    print(f"  统计信息:")
    print(f"    总替换次数: {stats.get('total_replacements', 0)}")
    print(f"    找到的模式数: {len(stats.get('patterns_found', []))}")
    
    if result.get('success'):
        print(f"  输出文件: {output_file}")
        
        # 检查输出结果
        try:
            output_doc = Document(output_file)
            success_count = 0
            total_tests = 0
            
            # 测试段落#30（多字段保留分隔符）
            if len(output_doc.paragraphs) > 30:
                total_tests += 1
                para30_result = output_doc.paragraphs[30].text
                print(f"\n段落#30处理结果:")
                print(f"  原文: '                      供应商名称：   采购编号：GXTC-C-251590031          '")
                print(f"  结果: '{para30_result}'")
                
                # 检查关键点：
                # 1. 前导空格保留
                # 2. 供应商名称正确填写
                # 3. 采购编号保留  
                # 4. 字段间分隔符保留
                # 5. 尾部占位符清理
                
                if (para30_result.startswith("                      ") and  # 前导空格保留
                    f"供应商名称：{company_name}" in para30_result and       # 正确填写
                    "   采购编号：GXTC-C-251590031" in para30_result):       # 保留已填内容且有分隔符
                    print("  ✅ 段落#30测试通过：多字段格式正确处理")
                    success_count += 1
                else:
                    print("  ❌ 段落#30测试失败：多字段格式处理有问题")
            
            # 测试段落#31（单标签前导空格保留）
            if len(output_doc.paragraphs) > 31:
                total_tests += 1
                para31_result = output_doc.paragraphs[31].text
                print(f"\n段落#31处理结果:")
                print(f"  原文: '          投标人名称：________________________'")
                print(f"  结果: '{para31_result}'")
                
                if (para31_result.startswith("          ") and                    # 前导空格保留
                    f"投标人名称：{company_name}" in para31_result and              # 正确填写（紧挨着也是正确的）
                    not "________________________" in para31_result):              # 占位符清理
                    print("  ✅ 段落#31测试通过：单标签前导空格保留")
                    success_count += 1
                else:
                    print("  ❌ 段落#31测试失败：单标签前导空格处理有问题")
                    
            # 测试段落#32（公章格式）
            if len(output_doc.paragraphs) > 32:
                total_tests += 1
                para32_result = output_doc.paragraphs[32].text  
                print(f"\n段落#32处理结果:")
                print(f"  原文: '供应商名称（加盖公章）：                    '")
                print(f"  结果: '{para32_result}'")
                
                if (f"供应商名称（加盖公章）：{company_name}" in para32_result and  # 正确填写
                    not para32_result.endswith("                    ")):          # 占位符清理
                    print("  ✅ 段落#32测试通过：公章格式正确处理")
                    success_count += 1
                else:
                    print("  ❌ 段落#32测试失败：公章格式处理有问题")
            
            print(f"\n测试汇总: {success_count}/{total_tests} 测试通过")
            return success_count == total_tests
            
        except Exception as e:
            print(f"  读取输出文件失败: {e}")
            return False
    
    return False

def main():
    """主函数"""
    print("智能分隔符保留的占位符清理方法测试")
    print("测试场景：")
    print("1. 多字段单行 - 保留字段间分隔符，清理占位符")
    print("2. 单标签行 - 保留前导格式空格，清理占位符")  
    print("3. 公章格式 - 正确处理特殊标签格式")
    print("=" * 60)
    
    success = test_comprehensive_smart_cleanup()
    
    print("\n" + "=" * 60)
    print("智能占位符清理效果评估:")
    if success:
        print("✅ 修复成功！智能分隔符保留逻辑工作正常")
        print("✅ 多字段格式保持，占位符准确清理")
        print("✅ 前导格式空格正确保留")
        print("✅ 供应商名称拆分问题已解决")
    else:
        print("❌ 修复需要进一步调试")
    print("=" * 60)
    
    return success

if __name__ == '__main__':
    main()