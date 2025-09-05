#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
商务应答处理器
集成四个现有功能模块：
1. 填写投标人/供应商名称
2. 填写项目信息（项目名称、招标编号、日期） 
3. 插入资质图片
4. 生成AI应答内容
"""

import os
import re
import logging
import json
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Mm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('business_response.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class BusinessResponseProcessor:
    """商务应答处理器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("请安装python-docx库：pip install python-docx")
        
        # 投标人名称匹配规则 - 增强版，解决跨run拆分和重复填写问题
        self.bidder_patterns = [
            # === 括号格式（最高优先级）===
            # 格式11: "（请填写供应商名称）" - 提示性括号内容替换
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>请填写\s*供应商名称)\s*(?P<suffix>[\)）])'),
            
            # 格式12: "（请填写投标人名称）" - 提示性括号内容替换
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>请填写\s*投标人名称)\s*(?P<suffix>[\)）])'),
            
            # 格式13: "（请填写公司名称）" - 提示性括号内容替换
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>请填写\s*公司名称)\s*(?P<suffix>[\)）])'),
            
            # 格式1: " (供应商全称)  " - 括号内容替换，保持字体和大小
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>供应商全称)\s*(?P<suffix>[\)）])'),
            
            # 格式10: "（供应商名称）" - 括号内容替换，保持字体和大小
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>供应商名称)\s*(?P<suffix>[\)）])'),
            
            # === 具体格式（中等优先级）===
            # 格式6: "公司名称（全称、盖章）：________________" - 横线上填写
            re.compile(r'^(?P<label>公司名称（全称、盖章）)\s*(?P<sep>[:：])\s*(?P<placeholder>_{3,})\s*$'),
            
            # 格式6-2: "公司名称（全称、盖章）：" - 冒号后填写（无占位符）
            re.compile(r'^(?P<label>公司名称（全称、盖章）)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # 格式6-3: "公司名称（全称、盖章）" - 无冒号格式（增强识别）
            re.compile(r'^(?P<label>公司名称（全称、盖章）)\s*(?P<sep>)\s*(?P<placeholder>)\s*$'),
            
            # 格式6-4: 灵活匹配"公司名称"+"全称"+"盖章"的任意组合
            re.compile(r'^(?P<label>公司名称[^:：]*全称[^:：]*盖章[^:：]*)\s*(?P<sep>[:：]?)\s*(?P<placeholder>.*?)\s*$'),
            
            # 格式7: "公司名称（盖章）：" - 冒号后填写
            re.compile(r'^(?P<label>公司名称（盖章）)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # 格式5: "供应商全称及公章：  " - 冒号后填写（最少有空格）
            re.compile(r'^(?P<label>供应商全称及公章)\s*(?P<sep>[:：])\s*(?P<placeholder>\s+)\s*$'),
            
            # 格式9: "供应商名称：                                （加盖公章）" - 中间空格处填写
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>\s{10,})\s*(?P<suffix>（[^）]*公章[^）]*）|\([^)]*公章[^)]*\))\s*$'),
            
            # 格式3: "供应商名称：                                        " - 空格处填写（长空格）
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>\s{20,})\s*$'),
            
            # 格式4: "供应商名称：               " - 中等长度空格
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>\s{10,19})\s*$'),
            
            # 格式2: "供应商名称：___________________" - 横线上填写
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>_{3,})\s*$'),
            
            # === 跨run拆分格式增强（解决供应商名称：被拆分的问题）===
            # 增强格式A: "供应商名称："精确匹配（解决跨run问题）
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # 增强格式B: "投标人名称："精确匹配
            re.compile(r'^(?P<label>投标人名称)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # 增强格式C: "公司名称："精确匹配  
            re.compile(r'^(?P<label>公司名称)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # 增强格式D: 灵活匹配含"供应商"和"名称"的任何组合
            re.compile(r'^(?P<label>[^:：]*供应商[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$'),
            
            # 增强格式E: 灵活匹配含"投标人"和"名称"的任何组合
            re.compile(r'^(?P<label>[^:：]*投标人[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$'),
            
            # === 通用格式（最低优先级）===
            # 投标人名称相关
            re.compile(r'^(?P<label>投标人名称(?:（公章）|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]*|＿*|——*)\s*$'),
        ]
        
        # 项目信息匹配规则
        self.project_replacements = [
            # 项目名称
            (r"(项目名称[:：]\s*)$", r"\1{project_name}"),
            (r"(项目名称（盖章）[:：]?\s*)$", r"\1{project_name}"),
            (r"(项目名称\s+)$", r"\1{project_name}"),
            (r"(项目名称)$", r"\1{project_name}"),
            (r"\[项目名称\]", "{project_name}"),
            # 招标编号
            (r"(招标编号[:：]\s*)$", r"\1{tender_no}"),
            (r"(招标编号\s+)$", r"\1{tender_no}"),
            (r"(招标编号)$", r"\1{tender_no}"),
            (r"\[招标编号\]", "{tender_no}"),
            # 日期
            (r"^(日\s*期)\s*[:：]?\s*$", "日期：{date_text}"),
            # 组合
            ("项目名称、招标编号", "{project_name}、{tender_no}")
        ]
        
        # 公司详细信息匹配规则 - 增强多行格式支持
        self.company_info_patterns = [
            # 地址相关 - 支持多种占位符格式
            re.compile(r'(?P<label>(注册地址|办公地址|公司地址|联系地址|通讯地址|地\s*址)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{3,}|＿{3,}|——{1,}|（\s*）)?$', re.MULTILINE),
            
            # 邮编 - 支持多种占位符格式  
            re.compile(r'(?P<label>(邮\s*编|邮政编码|邮码)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{3,}|＿{3,}|——{1,}|（\s*）)?$', re.MULTILINE),
            
            # 电话 - 支持多种占位符格式
            re.compile(r'(?P<label>(联系电话|固定电话|电\s*话|办公电话|公司电话)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{3,}|＿{3,}|——{1,}|（\s*）)?$', re.MULTILINE),
            
            # 电子邮箱 - 新增支持
            re.compile(r'(?P<label>(电子邮箱|邮箱|电子邮件|E-mail|Email|email)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{3,}|＿{3,}|——{1,}|（\s*）)?$', re.MULTILINE),
            
            # 成立时间
            re.compile(r'(?P<label>(成立时间|成立日期|注册时间|注册日期|成\s*立)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
            
            # 经营范围
            re.compile(r'(?P<label>(经营范围|业务范围|主营业务|经营内容)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
            
            # 法定代表人
            re.compile(r'(?P<label>(法定代表人|法人代表|法\s*人)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
            
            # 统一社会信用代码
            re.compile(r'(?P<label>(统一社会信用代码|社会信用代码|信用代码|统一代码)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
            
            # 注册资本
            re.compile(r'(?P<label>(注册资本|注册资金|资本金)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
            
            # 新增供应商相关字段（移除供应商全称及公章，避免与fill_bidder_name重复）
            re.compile(r'(?P<label>(供应商代表签字或印章)(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'),
        ]
        
    def insert_paragraph_after(self, paragraph, text=None, style=None):
        """在指定段落后插入新段落"""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = paragraph._parent.add_paragraph()
        new_para._p = new_p
        if text:
            new_para.add_run(text)
        if style:
            new_para.style = style
        return new_para
    
    def _handle_multiline_format(self, paragraphs, current_index, company_name, company_info, processed_indices):
        """处理分行格式的供应商名称（如：第一行是标签，第二行是印章）"""
        
        current_para = paragraphs[current_index]
        current_text = ''.join(run.text for run in current_para.runs).strip()
        
        # 检查当前段落是否匹配供应商名称模式（无印章后缀）
        current_matched = False
        matched_pattern = None
        match_groups = None
        
        for pattern in self.bidder_patterns:
            match = pattern.search(current_text)
            if match:
                groups = match.groupdict()
                # 只处理标准格式，不处理括号格式
                if 'prefix' in groups and 'content' in groups:
                    continue
                    
                label = groups.get('label', '')
                if label and any(keyword in label for keyword in ['投标人', '供应商', '公司']):
                    # 检查是否已经包含印章标识
                    if not any(seal in label for seal in ['盖章', '公章']):
                        current_matched = True
                        matched_pattern = pattern
                        match_groups = groups
                        break
        
        if not current_matched:
            return False
        
        # 检查下一个段落是否包含印章
        next_index = current_index + 1
        if next_index >= len(paragraphs):
            return False
            
        next_para = paragraphs[next_index]
        next_text = ''.join(run.text for run in next_para.runs).strip()
        
        # 检查下一行是否是印章格式
        has_seal = any(seal in next_text for seal in ['（公章）', '（盖章）', '(公章)', '(盖章)'])
        if not has_seal:
            return False
        
        logger.info(f"检测到分行格式 - 第{current_index+1}行: '{current_text}', 第{next_index+1}行: '{next_text}'")
        
        # 处理分行格式 - 保持分行显示
        label = match_groups.get('label', '')
        sep = match_groups.get('sep', ':')
        
        # 确定使用哪种印章类型
        if '盖章' in next_text:
            seal_text = '（盖章）' if '（盖章）' in next_text else '(盖章)'
        else:
            seal_text = '（公章）' if '（公章）' in next_text else '(公章)'
        
        # 确保分隔符正确
        if not sep:
            sep = ':'
        if sep and not sep.endswith(' '):
            sep += ' ' if sep in [':', '：'] else ''
        
        # 修改当前段落 - 保持原有标签格式，添加公司名称
        # 找到最佳源run用于复制格式
        source_run = self._find_best_source_run(current_para)
        
        # 清空当前段落内容
        for run in current_para.runs:
            run.text = ""
        
        # 重建第一行内容：保持原标签格式，添加公司名称
        if current_para.runs:
            label_run = current_para.runs[0]
        else:
            label_run = current_para.add_run("")
            
        label_run.text = f"{label}{sep}"
        
        # 复制原标签的字体格式到标签run
        if source_run:
            self._copy_font_format(source_run, label_run)
        
        # 添加公司名称，保持相同字体样式但加下划线
        company_run = current_para.add_run(company_name)
        company_run.underline = True
        
        # 复制源run的字体格式到公司名称run
        if source_run:
            self._copy_font_format(source_run, company_run)
        
        # 修改下一个段落 - 保持印章在第二行
        # 找到下一个段落的最佳源run用于复制格式
        seal_source_run = self._find_best_source_run(next_para)
        
        for run in next_para.runs:
            run.text = ""
        
        # 重建印章行
        if next_para.runs:
            seal_run = next_para.runs[0]
        else:
            seal_run = next_para.add_run("")
        seal_run.text = seal_text
        
        # 复制原印章行的字体格式
        if seal_source_run:
            self._copy_font_format(seal_source_run, seal_run)
        
        logger.info(f"分行格式处理完成: 第一行: '{label}{sep}{company_name}', 第二行: '{seal_text}'")
        
        # 标记两个段落都已处理
        processed_indices.add(current_index)
        processed_indices.add(next_index)
        
        return True
    
    def _might_be_multiline_format(self, current_text, paragraphs, current_index):
        """检查当前段落是否可能是分行格式的第一部分"""
        if not paragraphs or current_index is None:
            return False
            
        next_index = current_index + 1
        if next_index >= len(paragraphs):
            return False
            
        next_para = paragraphs[next_index]
        next_text = ''.join(run.text for run in next_para.runs).strip()
        
        # 检查下一行是否只包含印章（没有其他文字）
        has_seal = any(seal in next_text for seal in ['（公章）', '（盖章）', '(公章)', '(盖章)'])
        if not has_seal:
            return False
            
        # 检查下一行是否主要是空格和印章
        next_clean = next_text.strip()
        is_mainly_seal = any(next_clean == seal for seal in ['（公章）', '（盖章）', '(公章)', '(盖章)'])
        
        return is_mainly_seal
    
    def _copy_font_format(self, source_run, target_run):
        """完整复制字体格式，确保一致性"""
        try:
            if not source_run or not target_run or not source_run.font:
                return
            
            source_font = source_run.font
            target_font = target_run.font
            
            # 复制字体名称
            if source_font.name:
                target_font.name = source_font.name
            
            # 复制字体大小
            if source_font.size:
                target_font.size = source_font.size
            
            # 复制粗体
            if source_font.bold is not None:
                target_font.bold = source_font.bold
            
            # 复制斜体
            if source_font.italic is not None:
                target_font.italic = source_font.italic
            
            # 复制字体颜色
            if source_font.color and hasattr(source_font.color, 'rgb') and source_font.color.rgb:
                target_font.color.rgb = source_font.color.rgb
            
            # 复制上标/下标
            if hasattr(source_font, 'superscript') and source_font.superscript is not None:
                target_font.superscript = source_font.superscript
            if hasattr(source_font, 'subscript') and source_font.subscript is not None:
                target_font.subscript = source_font.subscript
                
        except Exception as e:
            logger.warning(f"复制字体格式时出错: {e}")
    
    def _find_best_source_run(self, paragraph):
        """找到段落中最合适的源run用于复制格式"""
        if not paragraph.runs:
            return None
            
        # 优先选择有文字内容的run
        for run in paragraph.runs:
            if run.text.strip():
                return run
        
        # 如果都没有文字，返回第一个run
        return paragraph.runs[0]
    
    def _update_paragraph_with_format(self, paragraph, new_full_text, company_name, original_text):
        """保持原有布局结构的智能文本替换"""
        try:
            if not paragraph.runs:
                return
            
            # 使用精确替换策略，保持原有布局
            self._precise_text_replacement(paragraph, original_text, new_full_text, company_name)
                
        except Exception as e:
            logger.warning(f"更新段落格式时出错: {e}")
            # 降级处理：使用简单替换
            self._fallback_text_replacement(paragraph, new_full_text, company_name)
    
    def _precise_text_replacement(self, paragraph, original_text, new_text, company_name):
        """精确的文本替换，保持原有run结构和格式"""
        try:
            # 收集所有run的信息
            run_info = []
            current_pos = 0
            
            for i, run in enumerate(paragraph.runs):
                run_length = len(run.text)
                run_info.append({
                    'index': i,
                    'original_text': run.text,
                    'start_pos': current_pos,
                    'end_pos': current_pos + run_length,
                    'font': run.font
                })
                current_pos += run_length
            
            # 找到需要替换公司名称的位置
            if company_name in new_text:
                company_start = new_text.find(company_name)
                company_end = company_start + len(company_name)
                
                # 逐个处理run，保持原有结构
                current_new_pos = 0
                for run_data in run_info:
                    run = paragraph.runs[run_data['index']]
                    original_start = run_data['start_pos']
                    original_end = run_data['end_pos']
                    run_length = original_end - original_start
                    
                    # 计算这个run在新文本中对应的位置
                    new_start = current_new_pos
                    new_end = current_new_pos + run_length
                    
                    # 提取对应的新文本
                    if new_end <= len(new_text):
                        new_run_text = new_text[new_start:new_end]
                    else:
                        new_run_text = new_text[new_start:]
                    
                    # 检查这个run是否包含公司名称
                    if (new_start <= company_start < new_end or
                        new_start < company_end <= new_end or
                        (company_start <= new_start and company_end >= new_end)):
                        
                        # 这个run包含公司名称，需要特殊处理
                        self._handle_run_with_company_name(run, new_run_text, company_name, run_data)
                    else:
                        # 普通run，直接替换文本
                        run.text = new_run_text
                    
                    current_new_pos = new_end
                    
        except Exception as e:
            logger.warning(f"精确替换失败: {e}")
            self._fallback_text_replacement(paragraph, new_text, company_name)
    
    def _handle_run_with_company_name(self, run, new_run_text, company_name, run_data):
        """处理包含公司名称的run"""
        try:
            if company_name in new_run_text:
                # 分割文本
                parts = new_run_text.split(company_name, 1)
                prefix = parts[0]
                suffix = parts[1] if len(parts) > 1 else ""
                
                # 设置前缀
                run.text = prefix
                
                # 在这个run后面插入公司名称run
                para = run._element.getparent().getparent()  # 获取段落元素
                run_index = list(para).index(run._element)
                
                # 创建公司名称run
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                
                company_r = OxmlElement('w:r')
                company_t = OxmlElement('w:t')
                company_t.text = company_name
                
                # 添加下划线格式
                rPr = OxmlElement('w:rPr')
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                company_r.append(rPr)
                company_r.append(company_t)
                
                # 插入到正确位置
                para.insert(run_index + 1, company_r)
                
                # 如果有后缀，创建后缀run
                if suffix:
                    suffix_r = OxmlElement('w:r')
                    suffix_t = OxmlElement('w:t')
                    suffix_t.text = suffix
                    suffix_r.append(suffix_t)
                    para.insert(run_index + 2, suffix_r)
            else:
                # 不包含公司名称，直接设置
                run.text = new_run_text
                
        except Exception as e:
            logger.warning(f"处理包含公司名称的run时出错: {e}")
            run.text = new_run_text
    
    def _fallback_text_replacement(self, paragraph, new_text, company_name):
        """降级处理：简单的文本替换"""
        try:
            # 清空所有run
            for run in paragraph.runs:
                run.text = ""
            
            if company_name in new_text:
                # 找到公司名称位置
                company_start = new_text.find(company_name)
                prefix = new_text[:company_start] 
                suffix = new_text[company_start + len(company_name):]
                
                # 设置前缀
                if paragraph.runs:
                    paragraph.runs[0].text = prefix
                
                # 添加公司名称（带下划线）
                company_run = paragraph.add_run(company_name)
                company_run.underline = True
                
                # 添加后缀
                if suffix:
                    paragraph.add_run(suffix)
            else:
                # 直接设置文本
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    
        except Exception as e:
            logger.error(f"降级处理失败: {e}")
            if paragraph.runs:
                paragraph.runs[0].text = new_text
    
    def _simple_replace_paragraph_text(self, paragraph, new_text):
        """精确替换段落文本内容，保持原有格式"""
        try:
            # 获取原始文本
            original_text = ''.join(run.text for run in paragraph.runs)
            
            # 如果文本没有变化，直接返回
            if original_text == new_text:
                return False
            
            # 尝试精确替换，保持run结构不变
            return self._preserve_format_replace(paragraph, original_text, new_text)
                
        except Exception as e:
            logger.error(f"替换段落文本失败: {e}")
            return False
    
    def _preserve_format_replace(self, paragraph, original_text, new_text):
        """保持格式的精确替换"""
        try:
            # 如果只有一个run，直接替换
            if len(paragraph.runs) == 1:
                paragraph.runs[0].text = new_text
                return True
            
            # 多个run的情况，尝试智能替换
            # 找出需要替换的部分在哪些run中
            runs = paragraph.runs
            total_length = 0
            run_positions = []
            
            for i, run in enumerate(runs):
                start_pos = total_length
                end_pos = total_length + len(run.text)
                run_positions.append((i, start_pos, end_pos, run.text))
                total_length += len(run.text)
            
            # 找出差异部分
            diff_start, diff_end = self._find_text_difference(original_text, new_text)
            
            if diff_start == -1:  # 没有找到差异，降级为简单替换
                runs[0].text = new_text
                for i in range(1, len(runs)):
                    runs[i].text = ""
                return True
            
            # 找出需要修改的run
            replacement_text = new_text[diff_start:diff_start + (len(new_text) - len(original_text) + (diff_end - diff_start))]
            
            # 找到包含差异的run
            target_run_idx = -1
            for i, (run_idx, start_pos, end_pos, run_text) in enumerate(run_positions):
                if start_pos <= diff_start < end_pos:
                    target_run_idx = run_idx
                    break
            
            if target_run_idx >= 0:
                # 计算在run内的位置
                run_start_pos = run_positions[target_run_idx][1]
                local_start = diff_start - run_start_pos
                local_end = min(len(runs[target_run_idx].text), local_start + (diff_end - diff_start))
                
                # 替换run中的文本
                old_run_text = runs[target_run_idx].text
                new_run_text = old_run_text[:local_start] + replacement_text + old_run_text[local_end:]
                runs[target_run_idx].text = new_run_text
                
                return True
            else:
                # 降级为简单替换
                runs[0].text = new_text
                for i in range(1, len(runs)):
                    runs[i].text = ""
                return True
                
        except Exception as e:
            logger.error(f"保持格式替换失败: {e}")
            # 最后的降级方案
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for i in range(1, len(paragraph.runs)):
                    paragraph.runs[i].text = ""
            return True
    
    def _find_text_difference(self, original, new):
        """找出两个文本之间的差异位置"""
        try:
            # 找到第一个不同的字符位置
            min_len = min(len(original), len(new))
            start_diff = 0
            
            while start_diff < min_len and original[start_diff] == new[start_diff]:
                start_diff += 1
            
            if start_diff == min_len:
                # 一个是另一个的前缀
                return start_diff, len(original)
            
            # 从后往前找最后一个不同的位置
            end_diff = 0
            orig_idx = len(original) - 1
            new_idx = len(new) - 1
            
            while (orig_idx >= start_diff and new_idx >= start_diff and 
                   original[orig_idx] == new[new_idx]):
                end_diff += 1
                orig_idx -= 1
                new_idx -= 1
            
            return start_diff, len(original) - end_diff
            
        except Exception as e:
            logger.error(f"查找文本差异失败: {e}")
            return -1, -1
    
    def _precise_replace_in_runs(self, paragraph, original_text, new_text, company_name):
        """精确替换run中的文本，保持原有格式和布局"""
        try:
            # 遍历所有run，寻找需要替换的内容
            for run_idx, run in enumerate(paragraph.runs):
                if not run.text:
                    continue
                
                # 检查这个run是否包含需要替换的占位符或标签
                run_text = run.text
                
                # 查找各种可能的占位符模式
                placeholder_patterns = [
                    r'[_\-\u2014\u3000]{2,}',  # 下划线、短横线等
                    r'＿{2,}',                 # 全角下划线
                    r'——{2,}',                 # 长横线
                    r'\([^)]*\)',             # 空括号或占位符括号
                ]
                
                # 寻找占位符并替换为公司名称
                for pattern in placeholder_patterns:
                    import re
                    matches = list(re.finditer(pattern, run_text))
                    if matches:
                        # 从后往前替换，避免位置偏移
                        new_run_text = run_text
                        for match in reversed(matches):
                            # 只替换最长的占位符（通常是公司名称的位置）
                            if len(match.group()) >= 4:  # 至少4个字符的占位符才替换
                                start, end = match.span()
                                # 替换为公司名称，但保持前后的空格和格式
                                new_run_text = new_run_text[:start] + company_name + new_run_text[end:]
                                break
                        
                        if new_run_text != run_text:
                            # 需要特殊处理：分割成多个run以添加下划线
                            company_pos = new_run_text.find(company_name)
                            if company_pos >= 0:
                                # 分割文本
                                prefix = new_run_text[:company_pos]
                                suffix = new_run_text[company_pos + len(company_name):]
                                
                                # 保存原始字体格式
                                original_font = run.font
                                
                                # 设置前缀
                                run.text = prefix
                                
                                # 插入公司名称run（带下划线）
                                company_run = paragraph.add_run(company_name)
                                company_run.underline = True
                                self._copy_font_format(run, company_run)
                                
                                # 插入后缀（如果有）
                                if suffix:
                                    suffix_run = paragraph.add_run(suffix)
                                    self._copy_font_format(run, suffix_run)
                                
                                return True
                
            return False
            
        except Exception as e:
            logger.warning(f"精确替换失败: {e}")
            return False
    
    def fill_bidder_name(self, doc: Document, company_name: str, company_info: Dict = {}) -> Dict:
        """填写投标人/供应商名称"""
        stats = {'paragraphs_changed': 0, 'tables_changed': 0, 'headers_changed': 0, 'footers_changed': 0}
        
        def replace_in_runs(paragraph, name, company_info={}, check_multiline=False, paragraphs=None, current_index=None):
            """段落级文本合并处理 - 解决跨run的标签拆分问题"""
            changed = False
            
            # 合并段落中所有run的文本以进行完整分析
            full_text = ''.join(run.text for run in paragraph.runs)
            if not full_text.strip():
                return False
            
            # 记录原始run结构用于调试
            logger.debug(f"段落 #{current_index+1 if current_index is not None else '?'}: 完整文本='{full_text}', run数量={len(paragraph.runs)}")
            for i, run in enumerate(paragraph.runs):
                logger.debug(f"  Run {i+1}: '{run.text}' (长度: {len(run.text)})")
            
            # 特殊处理：检测跨run拆分的标签
            if len(paragraph.runs) > 1:
                # 检查是否存在标签被拆分的情况
                potential_labels = ['供应商名称', '投标人名称', '公司名称', '公司名称（全称、盖章）', '公司名称（盖章）']
                for label in potential_labels:
                    if label in full_text and ':' in full_text or '：' in full_text:
                        # 找到可能被拆分的标签，记录调试信息
                        label_pos = full_text.find(label)
                        colon_pos = max(full_text.find(':'), full_text.find('：'))
                        if label_pos >= 0 and colon_pos > label_pos:
                            logger.info(f"检测到可能的跨run标签拆分: '{label}' 在位置 {label_pos}, 冒号在位置 {colon_pos}")
                            break
                
            # 增强的重复检测 - 解决问题1：重复填写
            # 检查是否已经包含公司名称（完整匹配或部分匹配）
            if name in full_text:
                logger.info(f"跳过已包含公司名称的段落: '{full_text.strip()}'")
                return False
            
            # 检查是否已经被处理过（通过查找下划线格式的公司名称片段）
            for run in paragraph.runs:
                if run.underline and name in run.text:
                    logger.info(f"跳过已有下划线格式公司名称的段落: '{full_text.strip()}'")
                    return False
            
            # 检查是否包含公司名称的关键词（避免部分匹配时的重复处理）
            company_keywords = name.split()[:2] if len(name) > 8 else [name]  # 取公司名前两个词或全名
            if len(company_keywords) > 1:
                keyword_matches = sum(1 for keyword in company_keywords if keyword in full_text)
                if keyword_matches >= len(company_keywords):
                    logger.info(f"跳过可能已包含公司关键词的段落: '{full_text.strip()}'")
                    return False
            
            # 在合并文本上进行模式匹配
            new_full_text = full_text
            matched = False
            
            for pattern_index, pattern in enumerate(self.bidder_patterns):
                # 检查是否是括号格式模式
                test_match = pattern.search(full_text)
                if test_match:
                    logger.debug(f"模式 #{pattern_index+1} 匹配成功: {test_match.groupdict()}")
                
                if test_match and 'prefix' in test_match.groupdict() and 'content' in test_match.groupdict():
                    # 括号格式：处理所有匹配项
                    matches = list(pattern.finditer(full_text))
                    if matches:
                        # 从后往前替换，避免位置偏移
                        current_text = full_text
                        for match in reversed(matches):
                            groups = match.groupdict()
                            prefix = groups.get('prefix', '')
                            content = groups.get('content', '')
                            suffix = groups.get('suffix', '')
                            
                            # 生成替换内容
                            if '住址' in content:
                                # 供应商住址应该保持原样，不进行替换
                                logger.warning(f"检测到住址字段 '{content}'，跳过替换以保护原文结构")
                                continue
                            else:
                                # 供应商名称：直接替换为公司名
                                replacement = f"{prefix}{name}{suffix}"
                            
                            # 替换匹配的部分
                            start, end = match.span()
                            current_text = current_text[:start] + replacement + current_text[end:]
                        
                        new_full_text = current_text
                        matched = True
                        break
                else:
                    # 普通格式：只处理第一个匹配
                    match = pattern.search(full_text)
                    if match:
                        groups = match.groupdict()
                        # 普通格式处理
                        label = groups.get('label', '')
                        sep = groups.get('sep', ':')
                        suffix = groups.get('suffix', '')
                        
                        # 确保分隔符格式正确
                        if not sep:
                            sep = ':'
                        if sep and not sep.endswith(' '):
                            sep += ' ' if sep in [':', '：'] else ''
                        
                        # 生成替换文本 - 根据不同格式进行处理
                        placeholder = groups.get('placeholder', '')
                        
                        logger.debug(f"处理标签: '{label}', 分隔符: '{sep}', 占位符: '{placeholder}' (长度: {len(placeholder)})")
                        
                        # 特殊处理格式9：带后缀的格式（如："供应商名称：                                （加盖公章）"）
                        if suffix:
                            # 格式9: 中间空格处填写，保持后缀
                            if placeholder and len(placeholder.strip()) == 0 and len(placeholder) > 10:
                                # 保持原有空格长度，在中间填入公司名称
                                spaces_before_suffix = "                    "  # 适量空格
                                new_full_text = f"{label}{sep}{name}{spaces_before_suffix}{suffix}"
                            else:
                                # 其他有后缀的情况，直接在标签后填写
                                new_full_text = f"{label}{sep}{name} {suffix}"
                        elif placeholder is not None:
                            # 检查占位符类型
                            if '_' in placeholder and len(placeholder) >= 3:
                                # 格式2,6: 横线格式，直接替换横线
                                new_full_text = f"{label}{sep}{name}"
                                logger.debug(f"横线格式替换: {new_full_text}")
                            elif placeholder.strip() == '' and len(placeholder) >= 20:
                                # 格式3: 长空格格式，检查是否可能是分行格式
                                if check_multiline and self._might_be_multiline_format(full_text, paragraphs, current_index):
                                    # 可能是分行格式，跳过当前处理，留给分行格式处理
                                    logger.debug("检测到分行格式，跳过处理")
                                    return False
                                # 格式3: 空格处填写
                                new_full_text = f"{label}{sep}{name}"
                                logger.debug(f"长空格格式替换: {new_full_text}")
                            elif placeholder.strip() == '' and len(placeholder) >= 10:
                                # 格式4: 中等长度空格，可能是分行格式，检查下一行
                                if check_multiline and self._might_be_multiline_format(full_text, paragraphs, current_index):
                                    # 可能是分行格式，跳过当前处理，留给分行格式处理
                                    logger.debug("检测到分行格式，跳过处理")
                                    return False
                                # 格式4: 空格处填写
                                new_full_text = f"{label}{sep}{name}"
                                logger.debug(f"中等空格格式替换: {new_full_text}")
                            elif placeholder.strip() == '' and len(placeholder) > 0:
                                # 格式5: 短空格格式
                                new_full_text = f"{label}{sep} {name}"
                                logger.debug(f"短空格格式替换: {new_full_text}")
                            elif placeholder == '':
                                # 格式7,8: 空占位符（冒号后直接填写）
                                new_full_text = f"{label}{sep} {name}"
                                logger.debug(f"空占位符格式替换: {new_full_text}")
                            else:
                                # 其他占位符情况
                                new_full_text = f"{label}{sep}{name}"
                                logger.debug(f"其他格式替换: {new_full_text}")
                        else:
                            # 没有占位符的情况，在标签后添加公司名
                            new_full_text = f"{label}{sep} {name}"
                            logger.debug(f"无占位符格式替换: {new_full_text}")
                        
                        matched = True
                        break
            
            # 如果有匹配，直接替换整个段落文本
            if matched:
                # 简化替换逻辑：直接替换所有run的文本内容
                changed = self._simple_replace_paragraph_text(paragraph, new_full_text)
                if changed:
                    logger.info(f"段落级匹配投标人名称: '{full_text.strip()}' -> '{new_full_text.strip()}'")
            
            return changed
        
        def process_table(table, name, company_info={}):
            changed_any = False
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if replace_in_runs(p, name, company_info):
                            changed_any = True
                    for tbl in cell.tables:
                        if process_table(tbl, name, company_info):
                            changed_any = True
            return changed_any
        
        # 处理正文段落
        paragraphs = doc.paragraphs
        processed_indices = set()  # 记录已处理的段落索引
        
        for i, para in enumerate(paragraphs):
            if i in processed_indices:
                continue
                
            # 尝试标准段落处理（启用分行格式检查）
            if replace_in_runs(para, company_name, company_info, check_multiline=True, paragraphs=paragraphs, current_index=i):
                stats['paragraphs_changed'] += 1
                processed_indices.add(i)
            else:
                # 检查是否是分行格式（供应商名称在当前行，印章在下一行）
                if self._handle_multiline_format(paragraphs, i, company_name, company_info, processed_indices):
                    stats['paragraphs_changed'] += 1
        
        # 处理表格
        for table in doc.tables:
            if process_table(table, company_name, company_info):
                stats['tables_changed'] += 1
        
        # 处理页眉页脚
        for section in doc.sections:
            header = section.header
            footer = section.footer
            header_changed = False
            footer_changed = False
            
            for para in header.paragraphs:
                if replace_in_runs(para, company_name, company_info):
                    header_changed = True
            for table in header.tables:
                if process_table(table, company_name, company_info):
                    header_changed = True
            
            for para in footer.paragraphs:
                if replace_in_runs(para, company_name, company_info):
                    footer_changed = True
            for table in footer.tables:
                if process_table(table, company_name, company_info):
                    footer_changed = True
            
            if header_changed:
                stats['headers_changed'] += 1
            if footer_changed:
                stats['footers_changed'] += 1
        
        logger.info(f"填写投标人名称完成: {stats}")
        return stats
    
    def fill_project_info(self, doc: Document, project_name: str, tender_no: str, date_text: str, company_name: str) -> Dict:
        """填写项目信息"""
        stats = {'replacements_made': 0, 'bidder_fields_filled': 0}
        
        # 准备替换模板
        replacements = []
        for pat, repl in self.project_replacements:
            replacements.append((
                pat,
                repl.format(project_name=project_name, tender_no=tender_no, date_text=date_text)
            ))
        
        # 处理普通替换
        for para in doc.paragraphs:
            for run in para.runs:
                for pat, repl in replacements:
                    if re.search(pat, run.text):
                        logger.info(f"匹配项目信息: {run.text.strip()}")
                        run.text = re.sub(pat, repl, run.text)
                        stats['replacements_made'] += 1
                        break
        
        # 处理投标人名称相关字段（带下划线格式）- 使用增强的模式匹配
        processed_paragraphs = set()  # 记录已处理的段落，避免重复
        
        for para_idx, para in enumerate(doc.paragraphs):
            if para_idx in processed_paragraphs:
                continue
                
            for i, run in enumerate(para.runs):
                text = run.text.strip()
                if not text:
                    continue
                
                # 检查是否已经包含公司名称，避免重复处理
                if company_name in text:
                    logger.info(f"跳过已包含公司名称的段落: '{text}'")
                    continue
                
                # 使用多个模式尝试匹配投标人名称字段
                matched_pattern = None
                for pattern in self.bidder_patterns:
                    if pattern.search(text):
                        matched_pattern = pattern
                        break
                
                if matched_pattern:
                    match = matched_pattern.search(text)
                    if match:
                        groups = match.groupdict()
                        
                        # 检查是否是括号格式 (供应商名称)
                        if 'prefix' in groups and 'content' in groups and 'suffix' in groups:
                            # 括号格式：直接替换括号内容，不处理项目信息
                            continue
                        
                        # 普通格式处理
                        label = groups.get('label', '')
                        if not label:
                            continue
                            
                        # 确定前缀格式
                        if "（公章）" in label or "公章" in label:
                            prefix = f"{label}："
                        elif "（盖章）" in label or "盖章" in label:
                            prefix = f"{label}："
                        else:
                            # 确保有冒号分隔符
                            if not label.endswith((':', '：')):
                                prefix = f"{label}："
                            else:
                                prefix = label
                        
                        # 覆盖原run
                        run.text = prefix
                        
                        # 清空后续run
                        for j in range(i + 1, len(para.runs)):
                            para.runs[j].text = ""
                        
                        # 插入带下划线的公司名称
                        new_run = para.add_run(company_name)
                        new_run.underline = WD_UNDERLINE.SINGLE
                        
                        # 使用完整的字体格式复制
                        self._copy_font_format(run, new_run)
                        
                        logger.info(f"填充投标人名称: {prefix} {company_name}")
                        stats['bidder_fields_filled'] += 1
                        processed_paragraphs.add(para_idx)  # 标记已处理
                        break
        
        logger.info(f"填写项目信息完成: {stats}")
        return stats
    
    def fill_company_details(self, doc: Document, company_info: Dict) -> Dict:
        """填写公司详细信息（地址、邮编、电话、成立时间、经营范围等）"""
        stats = {'details_filled': 0, 'fields_processed': 0}
        
        # 公司信息字段映射
        field_mapping = {
            # 地址类字段
            '注册地址': company_info.get('registeredAddress', ''),
            '办公地址': company_info.get('officeAddress', ''),
            '公司地址': company_info.get('officeAddress', '') or company_info.get('registeredAddress', ''),
            '联系地址': company_info.get('officeAddress', '') or company_info.get('registeredAddress', ''),
            '通讯地址': company_info.get('officeAddress', '') or company_info.get('registeredAddress', ''),
            '地址': company_info.get('officeAddress', '') or company_info.get('registeredAddress', ''),
            
            # 邮编
            '邮编': company_info.get('postalCode', ''),
            '邮政编码': company_info.get('postalCode', ''),
            '邮码': company_info.get('postalCode', ''),
            
            # 电话
            '联系电话': company_info.get('fixedPhone', ''),
            '固定电话': company_info.get('fixedPhone', ''),
            '电话': company_info.get('fixedPhone', ''),
            '办公电话': company_info.get('fixedPhone', ''),
            '公司电话': company_info.get('fixedPhone', ''),
            
            # 电子邮箱
            '电子邮箱': company_info.get('email', ''),
            '邮箱': company_info.get('email', ''),
            '电子邮件': company_info.get('email', ''),
            'E-mail': company_info.get('email', ''),
            'Email': company_info.get('email', ''),
            'email': company_info.get('email', ''),
            
            # 成立时间
            '成立时间': company_info.get('establishDate', ''),
            '成立日期': company_info.get('establishDate', ''),
            '注册时间': company_info.get('establishDate', ''),
            '注册日期': company_info.get('establishDate', ''),
            '成立': company_info.get('establishDate', ''),
            
            # 经营范围
            '经营范围': company_info.get('businessScope', ''),
            '业务范围': company_info.get('businessScope', ''),
            '主营业务': company_info.get('businessScope', ''),
            '经营内容': company_info.get('businessScope', ''),
            
            # 法定代表人
            '法定代表人': company_info.get('legalRepresentative', ''),
            '法人代表': company_info.get('legalRepresentative', ''),
            '法人': company_info.get('legalRepresentative', ''),
            
            # 统一社会信用代码
            '统一社会信用代码': company_info.get('socialCreditCode', ''),
            '社会信用代码': company_info.get('socialCreditCode', ''),
            '信用代码': company_info.get('socialCreditCode', ''),
            '统一代码': company_info.get('socialCreditCode', ''),
            
            # 注册资本
            '注册资本': company_info.get('registeredCapital', ''),
            '注册资金': company_info.get('registeredCapital', ''),
            '资本金': company_info.get('registeredCapital', ''),
            
            # 新增字段 - 供应商相关（供应商全称及公章已在fill_bidder_name中处理）
            '供应商代表签字或印章': company_info.get('legalRepresentative', ''),
        }
        
        def replace_in_runs(paragraph, patterns, field_mapping):
            changed = False
            for run in paragraph.runs:
                original_text = run.text
                text = original_text.strip()
                if not text:
                    continue
                
                # 尝试匹配公司信息模式
                for pattern in patterns:
                    match = pattern.search(text)
                    if match:
                        label = match.group('label')
                        sep = match.group('sep') if 'sep' in match.groupdict() else ':'
                        
                        # 查找对应的字段值
                        field_value = None
                        for field_key, value in field_mapping.items():
                            if field_key in label and value:
                                field_value = value
                                break
                        
                        if field_value:
                            # 检查是否已经包含该值，避免重复处理
                            if field_value in text:
                                logger.info(f"跳过已包含信息的文本: '{text}'")
                                continue
                            
                            # 确保分隔符格式正确
                            if not sep:
                                sep = ':'
                            if sep and not sep.endswith(' '):
                                sep += ' ' if sep in [':', '：'] else ''
                            
                            # 生成替换文本
                            if 'placeholder' in match.groupdict() and match.group('placeholder'):
                                # 有占位符的情况，替换整个匹配
                                new_text = f"{label}{sep}{field_value}"
                            else:
                                # 没有占位符的情况，在标签后添加信息
                                new_text = f"{label}{sep}{field_value}"
                            
                            # 用新文本替换原文本
                            run.text = original_text.replace(match.group(0), new_text)
                            changed = True
                            logger.info(f"填写公司信息: '{text}' -> '{run.text}'")
                            stats['details_filled'] += 1
                            break  # 匹配到一个模式就停止
                
                stats['fields_processed'] += 1
            
            return changed
        
        def process_table(table, patterns, field_mapping):
            changed_any = False
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if replace_in_runs(p, patterns, field_mapping):
                            changed_any = True
                    for tbl in cell.tables:
                        if process_table(tbl, patterns, field_mapping):
                            changed_any = True
            return changed_any
        
        # 处理正文段落
        for para in doc.paragraphs:
            replace_in_runs(para, self.company_info_patterns, field_mapping)
        
        # 处理表格
        for table in doc.tables:
            process_table(table, self.company_info_patterns, field_mapping)
        
        # 处理页眉页脚
        for section in doc.sections:
            header = section.header
            footer = section.footer
            
            for para in header.paragraphs:
                replace_in_runs(para, self.company_info_patterns, field_mapping)
            for table in header.tables:
                process_table(table, self.company_info_patterns, field_mapping)
            
            for para in footer.paragraphs:
                replace_in_runs(para, self.company_info_patterns, field_mapping)
            for table in footer.tables:
                process_table(table, self.company_info_patterns, field_mapping)
        
        logger.info(f"填写公司详细信息完成: {stats}")
        return stats
    
    def insert_qualification_images(self, doc: Document, keyword_to_image: Dict[str, List[str]]) -> Dict:
        """插入资质证明图片"""
        stats = {'images_inserted': 0, 'keywords_found': 0}
        
        # 记录每个关键字最后一次出现的位置
        last_matches = {}
        
        for i, para in enumerate(doc.paragraphs):
            for keyword in keyword_to_image:
                if keyword in para.text:
                    last_matches[keyword] = para
                    logger.info(f"匹配到关键字 [{keyword}] 段落位置：{i}")
        
        stats['keywords_found'] = len(last_matches)
        
        # 在最后一次出现后插入图片
        for keyword, para in last_matches.items():
            img_paths = keyword_to_image[keyword]
            for img_path in img_paths:
                if os.path.exists(img_path):
                    new_para = self.insert_paragraph_after(para)
                    new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    run = new_para.add_run()
                    
                    # 根据关键字设置图片大小
                    if keyword == "营业执照":
                        run.add_picture(img_path, width=Mm(150))
                    elif keyword in ["法定代表人身份证", "被授权人身份证"]:
                        run.add_picture(img_path, width=Mm(65))
                    else:
                        run.add_picture(img_path, width=Inches(5.5))
                    
                    logger.info(f"插入图片: {keyword} -> {img_path}")
                    stats['images_inserted'] += 1
                else:
                    logger.warning(f"未找到图片文件: {img_path}")
        
        logger.info(f"插入资质图片完成: {stats}")
        return stats
    
    def process_business_response(self, 
                                input_file: str, 
                                output_file: str,
                                company_info: Dict,
                                project_info: Dict,
                                qualification_images: Dict = None) -> Dict:
        """
        处理商务应答文档
        
        Args:
            input_file: 输入文档路径
            output_file: 输出文档路径
            company_info: 公司信息 {"companyName": "xxx"}
            project_info: 项目信息 {"projectName": "xxx", "tenderNo": "xxx", "date": "xxx"}
            qualification_images: 资质图片路径映射 {"关键字": ["图片路径1", "图片路径2"]}
        """
        logger.info(f"开始处理商务应答文档: {input_file}")
        
        try:
            # 加载文档
            doc = Document(input_file)
            logger.info(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")
            
            # 统计信息
            total_stats = {
                'bidder_name': {},
                'project_info': {},
                'company_details': {},
                'qualification_images': {},
                'processing_time': 0
            }
            
            start_time = datetime.now()
            
            # 1. 填写投标人/供应商名称
            company_name = company_info.get('companyName', '')
            if company_name:
                total_stats['bidder_name'] = self.fill_bidder_name(doc, company_name, company_info)
            
            # 2. 填写项目信息
            project_name = project_info.get('projectName', '')
            tender_no = project_info.get('tenderNo', '')
            date_text = project_info.get('date', '')
            
            if any([project_name, tender_no, date_text]):
                total_stats['project_info'] = self.fill_project_info(
                    doc, project_name, tender_no, date_text, company_name
                )
            
            # 3. 填写公司详细信息（地址、邮编、电话、成立时间、经营范围等）
            if company_info:
                total_stats['company_details'] = self.fill_company_details(doc, company_info)
            
            # 4. 插入资质证明图片
            if qualification_images:
                total_stats['qualification_images'] = self.insert_qualification_images(
                    doc, qualification_images
                )
            
            # 保存文档
            doc.save(output_file)
            
            end_time = datetime.now()
            total_stats['processing_time'] = (end_time - start_time).total_seconds()
            
            logger.info(f"商务应答处理完成: {output_file}")
            logger.info(f"处理统计: {total_stats}")
            
            return {
                'success': True,
                'output_file': output_file,
                'stats': total_stats,
                'message': '商务应答文档处理完成'
            }
            
        except Exception as e:
            logger.error(f"处理商务应答失败: {e}", exc_info=True)
            return {
                'success': False,
                'error': str(e),
                'message': f'处理失败: {str(e)}'
            }


def main():
    """测试主函数"""
    processor = BusinessResponseProcessor()
    
    # 测试数据
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'establishDate': '2015-12-18',
        'legalRepresentative': '李振军',
        'socialCreditCode': '91110101MA002N1D30',
        'registeredCapital': '15466.6667万元',
        'fixedPhone': '010-63271000',
        'postalCode': '100010',
        'registeredAddress': '北京市东城区王府井大街200号七层711室',
        'officeAddress': '北京市西城区成方街25号长话北院',
        'email': 'service@smartsteps.com',
        'businessScope': '计算机信息技术；技术推广、技术开发、技术转让、技术服务、技术咨询；计算机软件开发、设计；计算机软件服务；计算机系统服务；计算机及系统集成；数据处理等。'
    }
    
    project_info = {
        'projectName': '所属运营商数据服务项目',
        'tenderNo': '2025-IT-0032',
        'date': '2025年 9月 3日'
    }
    
    qualification_images = {
        "营业执照": ["证明材料/营业执照_智慧足迹.png"],
        "法定代表人身份证": [
            "证明材料/法人身份证_正面_智慧足迹.png",
            "证明材料/法人身份证_反面_智慧足迹.png"
        ]
    }
    
    result = processor.process_business_response(
        "标书模板.docx",
        "商务应答-输出.docx",
        company_info,
        project_info,
        qualification_images
    )
    
    print(f"处理结果: {result}")


if __name__ == "__main__":
    main()