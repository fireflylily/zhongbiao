#!/usr/bin/env python3
"""
分析跨run处理方法
探究为什么需要重构整个段落，以及是否有更精细的替换方法
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
import logging
import re

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('cross_run_methods_analysis.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def create_test_document_with_cross_run():
    """创建跨run的测试文档"""
    doc = Document()
    
    # 创建跨run的"（采购编号）"
    p = doc.add_paragraph()
    p.add_run("根据贵方购采购货物及服务的竞争性磋商公告")
    
    # 分别创建三个run来模拟跨run情况
    run1 = p.add_run("（")
    run1.font.italic = True
    run1.font.underline = True
    
    run2 = p.add_run("采购编号")
    run2.font.italic = True 
    run2.font.underline = True
    
    run3 = p.add_run("）")
    run3.font.italic = True
    run3.font.underline = True
    
    p.add_run("，签字代表经正式授权")
    
    test_file = "cross_run_test.docx"
    doc.save(test_file)
    logger.info(f"创建跨run测试文档: {test_file}")
    return test_file, p

def method1_precise_run_modification(paragraph, old_text, new_text):
    """方法1：精确run修改法 - 只修改涉及的run，保留其他run不变"""
    logger.info("\n🔧 方法1：精确run修改法")
    logger.info("=" * 40)
    
    try:
        # 查找包含目标文本的run范围
        full_text = paragraph.text
        logger.info(f"段落完整文本: {full_text}")
        logger.info(f"要替换: '{old_text}' -> '{new_text}'")
        
        # 找到目标文本的位置
        start_pos = full_text.find(old_text)
        if start_pos == -1:
            logger.warning("未找到目标文本")
            return False
        
        end_pos = start_pos + len(old_text)
        logger.info(f"目标文本位置: {start_pos}-{end_pos}")
        
        # 找出涉及的run范围
        current_pos = 0
        affected_runs = []
        
        for i, run in enumerate(paragraph.runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            # 检查这个run是否与目标文本有重叠
            if (run_start < end_pos and run_end > start_pos):
                overlap_start = max(start_pos, run_start)
                overlap_end = min(end_pos, run_end)
                affected_runs.append({
                    'index': i,
                    'run': run,
                    'run_start': run_start,
                    'run_end': run_end,
                    'overlap_start': overlap_start,
                    'overlap_end': overlap_end,
                    'overlap_text': full_text[overlap_start:overlap_end]
                })
            
            current_pos = run_end
        
        logger.info(f"涉及的run数量: {len(affected_runs)}")
        for info in affected_runs:
            logger.info(f"  Run #{info['index']+1}: '{info['run'].text}' (重叠: '{info['overlap_text']}')")
        
        # 执行精确替换
        if len(affected_runs) == 1:
            # 只涉及一个run，直接替换
            run_info = affected_runs[0]
            run = run_info['run']
            original_text = run.text
            new_run_text = original_text.replace(old_text, new_text)
            run.text = new_run_text
            logger.info(f"✅ 单run精确替换: '{original_text}' -> '{new_run_text}'")
            return True
            
        elif len(affected_runs) > 1:
            # 涉及多个run，需要精细处理
            logger.info("⚠️ 涉及多个run，需要精细重构...")
            
            # 保存第一个run的格式作为模板
            first_run = affected_runs[0]['run']
            template_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'color': first_run.font.color.rgb if first_run.font.color.rgb else None
            }
            
            # 构建新的文本片段
            new_segment = ""
            current_pos = affected_runs[0]['run_start']
            
            # 处理第一个run的前缀部分
            first_run_info = affected_runs[0]
            if first_run_info['overlap_start'] > first_run_info['run_start']:
                prefix_len = first_run_info['overlap_start'] - first_run_info['run_start']
                new_segment += first_run_info['run'].text[:prefix_len]
            
            # 添加替换文本
            new_segment += new_text
            
            # 处理最后一个run的后缀部分
            last_run_info = affected_runs[-1]
            if last_run_info['overlap_end'] < last_run_info['run_end']:
                suffix_start = last_run_info['overlap_end'] - last_run_info['run_start']
                new_segment += last_run_info['run'].text[suffix_start:]
            
            # 清空所有受影响的run
            for info in affected_runs:
                info['run'].text = ""
            
            # 在第一个run中设置新文本
            first_run.text = new_segment
            
            # 恢复格式
            if template_format['font_name']:
                first_run.font.name = template_format['font_name']
            if template_format['font_size']:
                first_run.font.size = template_format['font_size']
            if template_format['bold'] is not None:
                first_run.font.bold = template_format['bold']
            if template_format['italic'] is not None:
                first_run.font.italic = template_format['italic']
            if template_format['underline'] is not None:
                first_run.font.underline = template_format['underline']
            if template_format['color']:
                first_run.font.color.rgb = template_format['color']
            
            logger.info(f"✅ 多run精确重构完成: '{new_segment}'")
            return True
            
    except Exception as e:
        logger.error(f"精确run修改失败: {e}", exc_info=True)
        return False

def method2_current_safe_replace(paragraph, old_text, new_text):
    """方法2：当前的_safe_replace_paragraph_text方法"""
    logger.info("\n🔧 方法2：当前的安全替换方法")
    logger.info("=" * 40)
    
    try:
        # 方法1：尝试在现有run中替换
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = new_text
                logger.info(f"✅ 在run中直接替换")
                return True
        
        # 方法2：重构整个段落
        if paragraph.text == old_text:
            # 保存第一个run的格式
            original_format = None
            if paragraph.runs:
                first_run = paragraph.runs[0]
                original_format = {
                    'font_name': first_run.font.name,
                    'font_size': first_run.font.size,
                    'bold': first_run.font.bold,
                    'italic': first_run.font.italic,
                    'underline': first_run.font.underline
                }
            
            # 清空段落并重新创建
            for run in paragraph.runs:
                run.text = ""
            
            # 添加新文本
            new_run = paragraph.add_run(new_text)
            
            # 恢复格式
            if original_format:
                if original_format['font_name']:
                    new_run.font.name = original_format['font_name']
                if original_format['font_size']:
                    new_run.font.size = original_format['font_size']
                if original_format['bold'] is not None:
                    new_run.font.bold = original_format['bold']
                if original_format['italic'] is not None:
                    new_run.font.italic = original_format['italic']
                if original_format['underline'] is not None:
                    new_run.font.underline = original_format['underline']
            
            logger.info(f"✅ 重构整个段落")
            return True
        else:
            # 部分文本替换 - 这里是关键问题所在！
            original_text = paragraph.text
            new_paragraph_text = original_text.replace(old_text, new_text)
            
            if new_paragraph_text != original_text:
                logger.info("⚠️ 需要部分文本替换，但当前方法会重构整个段落")
                
                # 保存第一个run的格式（这里是问题！只保存第一个run格式）
                original_format = None
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    original_format = {
                        'font_name': first_run.font.name,
                        'font_size': first_run.font.size,
                        'bold': first_run.font.bold,
                        'italic': first_run.font.italic,
                        'underline': first_run.font.underline
                    }
                
                # 清空所有run并重新创建（问题在这里！）
                for run in paragraph.runs:
                    run.text = ""
                
                # 创建新的单一run包含所有文本
                new_run = paragraph.add_run(new_paragraph_text)
                
                # 只恢复第一个run的格式（问题！丢失了其他run的格式）
                if original_format:
                    if original_format['font_name']:
                        new_run.font.name = original_format['font_name']
                    if original_format['font_size']:
                        new_run.font.size = original_format['font_size']
                    if original_format['bold'] is not None:
                        new_run.font.bold = original_format['bold']
                    if original_format['italic'] is not None:
                        new_run.font.italic = original_format['italic']
                    if original_format['underline'] is not None:
                        new_run.font.underline = original_format['underline']
                
                logger.info(f"⚠️ 整段重构完成，但可能丢失格式信息")
                return True
                
        return False
        
    except Exception as e:
        logger.error(f"当前安全替换方法失败: {e}")
        return False

def compare_methods():
    """对比不同方法的效果"""
    logger.info("\n" + "=" * 80)
    logger.info("对比不同跨run处理方法")
    logger.info("=" * 80)
    
    # 创建测试文档
    test_file, original_paragraph = create_test_document_with_cross_run()
    
    # 分析原始状态
    logger.info("\n📊 原始段落状态:")
    logger.info(f"文本: {original_paragraph.text}")
    logger.info(f"Run数量: {len(original_paragraph.runs)}")
    for i, run in enumerate(original_paragraph.runs):
        if run.text:
            logger.info(f"  Run {i+1}: '{run.text}' [斜体={run.font.italic}, 下划线={run.font.underline}]")
    
    # 方法1测试
    doc1 = Document(test_file)
    test_para1 = doc1.paragraphs[0]
    success1 = method1_precise_run_modification(test_para1, "（采购编号）", "（64525343）")
    
    if success1:
        logger.info("\n📊 方法1处理后:")
        logger.info(f"文本: {test_para1.text}")
        logger.info(f"Run数量: {len(test_para1.runs)}")
        for i, run in enumerate(test_para1.runs):
            if run.text:
                logger.info(f"  Run {i+1}: '{run.text}' [斜体={run.font.italic}, 下划线={run.font.underline}]")
        doc1.save("method1_result.docx")
    
    # 方法2测试
    doc2 = Document(test_file)
    test_para2 = doc2.paragraphs[0]
    success2 = method2_current_safe_replace(test_para2, "（采购编号）", "（64525343）")
    
    if success2:
        logger.info("\n📊 方法2处理后:")
        logger.info(f"文本: {test_para2.text}")
        logger.info(f"Run数量: {len(test_para2.runs)}")
        for i, run in enumerate(test_para2.runs):
            if run.text:
                logger.info(f"  Run {i+1}: '{run.text}' [斜体={run.font.italic}, 下划线={run.font.underline}]")
        doc2.save("method2_result.docx")
    
    # 总结对比
    logger.info("\n🎯 方法对比总结:")
    logger.info("=" * 50)
    logger.info("方法1 (精确run修改):")
    logger.info("  ✅ 只修改涉及的run，保留无关run的格式")
    logger.info("  ✅ 最小化格式影响")
    logger.info("  ⚠️ 实现复杂度较高")
    
    logger.info("\n方法2 (当前整段重构):")
    logger.info("  ⚠️ 清空所有run，重新创建单一run")
    logger.info("  ⚠️ 只保留第一个run的格式，丢失其他格式")
    logger.info("  ✅ 实现简单，稳定性高")
    
    logger.info("\n💡 为什么当前方法要重构整个段落？")
    logger.info("  1. 简单可靠：避免复杂的run边界计算")
    logger.info("  2. 一致性：确保替换后的文本格式统一")
    logger.info("  3. 兼容性：处理各种复杂的跨run情况")
    logger.info("  4. 但代价是：可能影响周围文字的格式")

if __name__ == "__main__":
    compare_methods()