#!/usr/bin/env python3
"""
测试采购编号格式修复效果
验证"姓名、职务"部分不再受到采购编号替换的格式影响
"""

import logging
import sys
from pathlib import Path
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_test_document():
    """创建包含采购编号和姓名职务的测试文档"""
    doc = Document()
    
    # 添加测试段落，模拟真实情况下的复杂格式
    paragraph = doc.add_paragraph()
    
    # 添加多个run模拟复杂的格式情况
    run1 = paragraph.add_run("根据贵方购采购货物及服务的竞争性磋商公告")
    run1.font.name = "宋体"
    run1.font.bold = False
    run1.font.italic = False
    
    run2 = paragraph.add_run("（")
    run2.font.name = "宋体"
    run2.font.bold = False
    run2.font.italic = True
    run2.font.underline = True
    
    run3 = paragraph.add_run("采购编号")
    run3.font.name = "宋体"
    run3.font.bold = False
    run3.font.italic = True
    run3.font.underline = True
    
    run4 = paragraph.add_run("）")
    run4.font.name = "宋体"
    run4.font.bold = False
    run4.font.italic = True
    run4.font.underline = True
    
    run5 = paragraph.add_run("，签字代表")
    run5.font.name = "宋体"
    run5.font.bold = False
    run5.font.italic = False
    
    run6 = paragraph.add_run("（")
    run6.font.name = "宋体"
    run6.font.bold = False
    run6.font.italic = False
    
    run7 = paragraph.add_run("姓名、职务")
    run7.font.name = "宋体"
    run7.font.bold = False
    run7.font.italic = False
    
    run8 = paragraph.add_run("）")
    run8.font.name = "宋体"
    run8.font.bold = False
    run8.font.italic = False
    
    run9 = paragraph.add_run("经正式授权并代表供应交下述文件正本一份及副本份：")
    run9.font.name = "宋体"
    run9.font.bold = False
    run9.font.italic = False
    
    return doc

def analyze_document_format(doc, title):
    """分析文档格式"""
    logger.info(f"\n📋 {title}")
    logger.info("="*60)
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            logger.info(f"段落 #{para_idx}: '{paragraph.text}'")
            
            for run_idx, run in enumerate(paragraph.runs):
                if run.text:
                    format_info = f"字体={run.font.name}, 粗体={run.font.bold}, 斜体={run.font.italic}, 下划线={run.font.underline}"
                    
                    # 特殊标记关键文本
                    marker = ""
                    if "64525343" in run.text:
                        marker = " ⭐ 新采购编号"
                    elif "采购编号" in run.text:
                        marker = " ⚠️ 原采购编号"
                    elif "姓名、职务" in run.text:
                        marker = " 🎯 姓名职务（重点关注）"
                    
                    logger.info(f"  Run {run_idx+1}: '{run.text}' [{format_info}]{marker}")

def test_format_fix():
    """测试采购编号格式修复效果"""
    logger.info("🔧 测试采购编号格式修复效果")
    
    try:
        # 创建测试文档
        input_doc = create_test_document()
        input_path = "test_format_fix_input.docx"
        input_doc.save(input_path)
        
        # 分析原始文档格式
        analyze_document_format(input_doc, "原始文档格式")
        
        # 使用修复后的处理器
        processor = MCPBidderNameProcessor()
        
        # 处理文档，传递必要参数
        output_path = "test_format_fix_output.docx"
        result = processor.process_bidder_name(
            input_file=input_path,
            output_file=output_path,
            company_name="智慧足迹数据科技有限公司"
        )
        success = result.get('success', False)
        
        if success:
            logger.info("✅ 文档处理成功")
            
            # 分析处理后文档格式
            output_doc = Document(output_path)
            analyze_document_format(output_doc, "处理后文档格式")
            
            # 🎯 验证关键点：检查"姓名、职务"的格式是否保持正常
            logger.info("\n🎯 关键验证：检查姓名、职务格式是否正常")
            logger.info("="*60)
            
            name_job_format_correct = True
            for para_idx, paragraph in enumerate(output_doc.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    if "姓名、职务" in run.text:
                        # 检查格式是否正常（不应该有斜体或下划线）
                        if run.font.italic or run.font.underline:
                            logger.error(f"❌ 发现格式异常！Run {run_idx+1} '姓名、职务' 意外包含格式: 斜体={run.font.italic}, 下划线={run.font.underline}")
                            name_job_format_correct = False
                        else:
                            logger.info(f"✅ 格式正常：Run {run_idx+1} '姓名、职务' 格式正确")
            
            # 验证采购编号是否正确替换
            logger.info("\n📝 验证采购编号替换结果")
            logger.info("="*60)
            
            full_text = "\n".join([p.text for p in output_doc.paragraphs])
            if "64525343" in full_text:
                logger.info("✅ 采购编号已成功替换为 64525343")
            else:
                logger.error("❌ 采购编号替换失败")
                name_job_format_correct = False
            
            if "采购编号" in full_text:
                logger.error("❌ 仍存在未替换的'采购编号'文本")
                name_job_format_correct = False
            else:
                logger.info("✅ 所有'采购编号'占位符已正确替换")
            
            # 最终结果
            logger.info(f"\n{'='*60}")
            if name_job_format_correct:
                logger.info("🎉 修复成功！姓名、职务格式未受采购编号替换影响")
                return True
            else:
                logger.error("❌ 修复失败！姓名、职务格式仍然受到影响")
                return False
            
        else:
            logger.error("❌ 文档处理失败")
            return False
            
    except Exception as e:
        logger.error(f"测试失败: {e}", exc_info=True)
        return False

if __name__ == "__main__":
    test_format_fix()