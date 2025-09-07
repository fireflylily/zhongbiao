#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试采购编号格式问题
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_test_document():
    """创建包含采购编号的测试文档"""
    doc = Document()
    doc.add_heading('采购编号格式测试', 0)
    
    # 测试不同的采购编号格式
    doc.add_paragraph("供应商名称：__________________ 采购编号：__________________")
    doc.add_paragraph("")
    doc.add_paragraph("根据贵方的竞争性磋商公告（采购编号），我方提交响应文件。")
    doc.add_paragraph("")
    
    # 创建一个表格测试
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "供应商名称："
    table.cell(0, 1).text = "_____________________"
    table.cell(1, 0).text = "采购编号："
    table.cell(1, 1).text = "_____________________"
    
    return doc


def main():
    """主测试函数"""
    print("="*60)
    print("采购编号格式测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_purchase_number_input.docx')
    output_file = os.path.join(current_dir, 'test_purchase_number_output.docx')
    
    # 创建测试文档
    print("\n1. 创建测试文档...")
    doc = create_test_document()
    doc.save(input_file)
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
        "fixedPhone": "010-66258899",
        "registeredAddress": "北京市西城区金融大街21号",
    }
    
    # 处理文档
    print("\n2. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            print("\n3. 处理结果:")
            processed_doc = Document(output_file)
            
            for i, para in enumerate(processed_doc.paragraphs):
                text = para.text.strip()
                if text and ('采购编号' in text or 'GXTC-C-251590031' in text):
                    print(f"\n   段落{i}: {text}")
                    
                    # 分析格式
                    if 'GXTC-C-251590031' in text:
                        if '采购编号：GXTC-C-251590031' in text:
                            print("      ✅ 格式正确：采购编号：GXTC-C-251590031")
                        elif '（GXTC-C-251590031）' in text:
                            print("      ✅ 格式正确：括号内替换")
                        else:
                            print(f"      ⚠️  格式可能有问题")
                            # 显示run结构
                            print("      Run结构:")
                            for j, run in enumerate(para.runs):
                                if run.text:
                                    print(f"        Run{j}: '{run.text}'")
            
            # 检查表格
            print("\n4. 表格内容:")
            for table in processed_doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if '采购编号' in text or 'GXTC-C-251590031' in text:
                            print(f"   表格[{row_idx},{col_idx}]: {text}")
            
            print("\n5. 期望格式:")
            print("   - 段落中: 供应商名称：中国联合网络通信有限公司 采购编号：GXTC-C-251590031")
            print("   - 括号中: （GXTC-C-251590031）")
            print("   - 表格中: 采购编号：GXTC-C-251590031")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 保留文件供检查
        print(f"\n6. 输出文件保存在: {output_file}")


if __name__ == '__main__':
    main()