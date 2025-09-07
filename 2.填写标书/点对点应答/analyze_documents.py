#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析处理前后的文档
"""

from docx import Document
import os

def analyze_document(file_path, title):
    """分析文档内容"""
    print(f"\n{'='*60}")
    print(f"{title}")
    print(f"文件: {file_path}")
    print(f"{'='*60}")
    
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return None
    
    try:
        doc = Document(file_path)
        
        # 查找包含供应商名称和采购编号的段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if ('供应商名称' in text and '采购编号' in text) or ('供应商名称' in text and len(text) > 4):
                print(f"\n段落{i}: {repr(text)}")
                print(f"  Run数量: {len(para.runs)}")
                for j, run in enumerate(para.runs):
                    format_info = []
                    if run.bold:
                        format_info.append("粗体")
                    if run.italic:
                        format_info.append("斜体")
                    if run.underline:
                        format_info.append("下划线")
                    format_str = f" [{', '.join(format_info)}]" if format_info else ""
                    print(f"    Run{j}: {repr(run.text)}{format_str}")
        
        return doc
        
    except Exception as e:
        print(f"❌ 读取文件失败: {e}")
        return None


def main():
    """主函数"""
    # 查找处理前的文件
    possible_input_files = [
        "/Users/lvhe/Downloads/中邮保险商务应答格式_测试.docx",
        "/Users/lvhe/Desktop/中邮保险商务应答格式_测试.docx",
        # 在当前目录的父目录中查找
    ]
    
    input_file = None
    for path in possible_input_files:
        if os.path.exists(path):
            input_file = path
            break
    
    # 也可以查找最近上传的文件
    uploads_dir = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/uploads"
    if os.path.exists(uploads_dir):
        for f in os.listdir(uploads_dir):
            if "20250906_212952" in f or "business_template" in f:
                potential_input = os.path.join(uploads_dir, f)
                if os.path.exists(potential_input):
                    input_file = potential_input
                    break
    
    output_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/outputs/docx-商务应答-20250906_212952.docx"
    
    print("分析文档内容...")
    
    if input_file:
        print(f"处理前文件: {input_file}")
        analyze_document(input_file, "处理前文档")
    else:
        print("⚠️ 未找到处理前文件，请确认文件位置")
    
    print(f"处理后文件: {output_file}")
    analyze_document(output_file, "处理后文档")
    
    # 对比分析
    print(f"\n{'='*60}")
    print("问题分析")
    print(f"{'='*60}")
    
    if input_file and os.path.exists(output_file):
        try:
            input_doc = Document(input_file)
            output_doc = Document(output_file)
            
            # 查找相关段落
            input_para = None
            output_para = None
            
            for para in input_doc.paragraphs:
                if '供应商名称' in para.text:
                    input_para = para
                    break
            
            for para in output_doc.paragraphs:
                if '供应商名称' in para.text and '中国联合' in para.text:
                    output_para = para
                    break
            
            if input_para and output_para:
                print(f"\n原文: {repr(input_para.text)}")
                print(f"处理后: {repr(output_para.text)}")
                
                # 检查问题
                issues = []
                
                # 检查是否有下划线
                for run in output_para.runs:
                    if '中国联合' in run.text and run.underline:
                        issues.append("✗ 公司名称被加了下划线")
                
                # 检查是否是插入而非替换
                if '___' in input_para.text and '___' in output_para.text:
                    remaining_underscores = output_para.text.count('_')
                    original_underscores = input_para.text.count('_')
                    if remaining_underscores >= original_underscores:
                        issues.append("✗ 占位符未被替换，公司名称是插入的")
                
                # 检查格式
                if '供应商名称' in output_para.text and '采购编号' in output_para.text:
                    if '\n' in output_para.text:
                        issues.append("✗ 供应商名称和采购编号被分到不同行")
                    else:
                        print("✓ 供应商名称和采购编号在同一行")
                
                if issues:
                    print("\n发现的问题:")
                    for issue in issues:
                        print(f"  {issue}")
                else:
                    print("\n✓ 未发现问题")
                
        except Exception as e:
            print(f"对比分析失败: {e}")


if __name__ == '__main__':
    main()