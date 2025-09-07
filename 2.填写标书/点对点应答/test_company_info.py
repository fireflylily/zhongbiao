#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试公司信息填写功能
"""

import os
import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_test_document():
    """创建测试文档"""
    doc = Document()
    doc.add_heading('公司信息填写测试文档', 0)
    
    # 添加各种格式的公司信息字段
    test_fields = [
        "投标人名称：                    ",
        "供应商名称（盖章）：            ",
        "公司名称：                      ",
        "地址：                          ",
        "注册地址：                      ",
        "电话：                          ",
        "联系电话：                      ",
        "传真：                          ",
        "传真号码：                      ",
        "电子邮件：                      ",
        "邮箱：                          ",
        "成立日期：                      ",
        "法定代表人：                    ",
        "统一社会信用代码：              ",
        "注册资本：                      ",
        "开户银行：                      ",
        "银行账号：                      ",
    ]
    
    for field in test_fields:
        para = doc.add_paragraph(field)
    
    # 添加表格测试
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 表格内容
    table_data = [
        ['投标人名称', '                    '],
        ['地址', '                        '],
        ['电话', '                        '],
        ['传真', '                        '],
        ['电子邮件', '                    ']
    ]
    
    for i, (label, placeholder) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    return doc


def main():
    """主测试函数"""
    print("开始测试公司信息填写功能...")
    
    # 创建测试文档
    input_file = os.path.join(current_dir, 'test_company_info_input.docx')
    output_file = os.path.join(current_dir, 'test_company_info_output.docx')
    
    print(f"创建测试文档: {input_file}")
    doc = create_test_document()
    doc.save(input_file)
    
    # 准备公司信息 - 使用英文字段名格式
    company_info = {
        "companyName": "中国联合网络通信有限公司",
        "establishDate": "2000-04-21",
        "legalRepresentative": "陈忠岳",
        "socialCreditCode": "91110000710939135P",
        "registeredCapital": "22539208.432769万元",
        "companyType": "有限责任公司（台港澳法人独资）",
        "fixedPhone": "010-66258899",
        "registeredAddress": "北京市西城区金融大街21号",
        "officeAddress": "北京市西城区金融大街21号",
        "website": "www.10010.com",
        "bankName": "中国工商银行北京市长安支行",
        "bankAccount": "200003309221111116",
        "fax": "010-66258866",
        "email": "service@chinaunicom.cn"
    }
    
    # 使用处理器处理文档
    processor = MCPBidderNameProcessor()
    
    try:
        print(f"处理文档...")
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-2025-001",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print(f"处理成功！")
            print(f"输出文件: {output_file}")
            print(f"处理统计: {result.get('stats', {})}")
            
            # 验证结果
            print("\n验证处理结果...")
            verify_doc = Document(output_file)
            all_text = []
            
            # 收集所有文本
            for para in verify_doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text.strip())
            
            for table in verify_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text.append(cell.text.strip())
            
            # 检查是否填入了公司信息
            company_name = company_info['companyName']
            company_address = company_info['registeredAddress']
            company_phone = company_info['fixedPhone']
            company_fax = company_info['fax']
            company_email = company_info['email']
            
            found_info = {
                '公司名称': False,
                '地址': False,
                '电话': False,
                '传真': False,
                '邮箱': False
            }
            
            for text in all_text:
                if company_name in text:
                    found_info['公司名称'] = True
                if company_address in text:
                    found_info['地址'] = True
                if company_phone in text:
                    found_info['电话'] = True
                if company_fax in text:
                    found_info['传真'] = True
                if company_email in text:
                    found_info['邮箱'] = True
            
            print("填写验证结果：")
            for field, found in found_info.items():
                status = "✓ 已填写" if found else "✗ 未填写"
                print(f"  {field}: {status}")
            
            return True
            
        else:
            print(f"处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)