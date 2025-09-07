#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试run结构修复效果
"""

import os
import re
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_complex_run_structure_document():
    """创建复杂run结构的测试文档，模拟采购编号处理后的情况"""
    doc = Document()
    doc.add_heading('Run结构修复测试', 0)
    
    # 添加一些前置段落
    for i in range(29):
        doc.add_paragraph(f'段落 #{i+1}: 前置内容...')
    
    # 段落#30 - 模拟采购编号处理后的复杂run结构
    para30 = doc.add_paragraph()
    
    # 模拟处理后的run结构（根据日志中的实际情况）
    # 处理前: "供应商名称：                          采购编号：                   "
    # 处理后: "供应商名称：   采购编号：GXTC-C-251590031  "
    
    para30.add_run("供应商名")     # run 0 - 标签被分割了！
    para30.add_run("称：")        # run 1 - 标签的后半部分
    para30.add_run("   ")         # run 2 - 空格
    para30.add_run("采购编号")     # run 3
    para30.add_run("：")          # run 4
    para30.add_run("GXTC-C-251590031")  # run 5 - 已填写的采购编号
    para30.add_run("  ")          # run 6 - 剩余空格
    
    print("创建的复杂run结构段落#30:")
    for i, run in enumerate(para30.runs):
        print(f"  Run {i}: '{run.text}' (长度: {len(run.text)})")
    
    full_text = ''.join(run.text for run in para30.runs)
    print(f"完整文本: '{full_text}'")
    
    return doc

def test_improved_run_handling():
    """测试改进后的run处理逻辑"""
    print("\n" + "="*60)
    print("测试改进后的run处理逻辑")
    print("="*60)
    
    # 创建测试文档
    input_file = 'test_run_structure_fix_input.docx'
    output_file = 'test_run_structure_fix_output.docx'
    
    doc = create_complex_run_structure_document()
    doc.save(input_file)
    print(f"已创建测试文档: {input_file}")
    
    # 使用改进后的MCP处理器处理
    processor = MCPBidderNameProcessor()
    company_name = "智慧足迹数据科技有限公司"
    
    print(f"\n开始处理，使用公司名称: {company_name}")
    
    result = processor.process_bidder_name(
        input_file=input_file,
        output_file=output_file,
        company_name=company_name
    )
    
    print(f"\n处理结果:")
    print(f"  成功: {result.get('success', False)}")
    print(f"  错误: {result.get('error', 'N/A')}")
    
    stats = result.get('stats', {})
    print(f"  统计信息:")
    print(f"    总替换次数: {stats.get('total_replacements', 0)}")
    print(f"    替换内容次数: {stats.get('replace_content_count', 0)}")  
    print(f"    空格填写次数: {stats.get('fill_space_count', 0)}")
    
    patterns_found = stats.get('patterns_found', [])
    print(f"    找到的模式:")
    for pattern in patterns_found:
        print(f"      规则#{pattern.get('rule_index', 'N/A')}: {pattern.get('description', 'N/A')}")
        print(f"      原文: '{pattern.get('original_text', 'N/A')}'")
    
    if result.get('success'):
        print(f"  输出文件: {output_file}")
        
        # 读取输出文件查看段落#30的结果
        try:
            output_doc = Document(output_file)
            if len(output_doc.paragraphs) > 30:
                para30_result = output_doc.paragraphs[30].text
                print(f"\n段落#30处理结果:")
                print(f"  原文: '供应商名称：   采购编号：GXTC-C-251590031  '")
                print(f"  结果: '{para30_result}'")
                
                # 检查是否有填写供应商名称
                if company_name in para30_result:
                    print(f"  ✅ 成功填写了供应商名称: {company_name}")
                else:
                    print(f"  ❌ 没有填写供应商名称")
                    
                # 检查run结构
                para30_runs = output_doc.paragraphs[30].runs
                print(f"  输出文档run结构:")
                for i, run in enumerate(para30_runs):
                    print(f"    Run {i}: '{run.text}'")
                    
        except Exception as e:
            print(f"  读取输出文件失败: {e}")
    
    return result

def main():
    """主函数"""
    print("Run结构修复测试")
    print("="*60)
    
    result = test_improved_run_handling()
    
    # 结论
    print("\n" + "="*60)
    print("修复效果评估:")
    if result.get('success'):
        patterns_found = result.get('stats', {}).get('patterns_found', [])
        supplier_patterns = [p for p in patterns_found if '供应商名称' in p.get('description', '')]
        
        if supplier_patterns:
            print("✅ 修复成功！改进后的run查找逻辑能够处理复杂run结构")
            print("✅ 供应商名称字段被正确识别和填写")
        else:
            print("❌ 修复不完整，供应商名称仍未被处理")
    else:
        print("❌ 修复失败，处理过程出现错误")
        
    print("="*60)

if __name__ == '__main__':
    main()