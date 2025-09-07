#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试具体的文字格式：
"供应商名称：                          采购编号：                   "
"""

import os
import re
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_specific_test_doc():
    """创建用户提供的具体格式测试文档"""
    doc = Document()
    doc.add_heading('用户提供的具体格式测试', 0)
    
    # 添加一些前置段落
    for i in range(29):
        doc.add_paragraph(f'段落 #{i+1}: 前置内容...')
    
    # 段落#30 - 用户提供的确切格式
    para30 = doc.add_paragraph()
    # 使用用户提供的确切文字，包含长空格
    para30.add_run("供应商名称：                          采购编号：                   ")
    
    print(f"创建的段落#30:")
    print(f"  文本: '{para30.text}'")
    print(f"  文本长度: {len(para30.text)}")
    print(f"  Run数量: {len(para30.runs)}")
    
    return doc

def analyze_specific_format_matching():
    """分析用户提供格式的规则匹配情况"""
    processor = MCPBidderNameProcessor()
    
    # 用户提供的确切文本
    test_text = "供应商名称：                          采购编号：                   "
    
    print(f"\n分析文本: '{test_text}'")
    print(f"文本长度: {len(test_text)}")
    print("="*60)
    
    # 检查每个规则是否匹配
    matching_rules = []
    for i, rule in enumerate(processor.bidder_patterns):
        pattern = rule['pattern']
        match = pattern.search(test_text)
        
        print(f"\n规则#{i+1}: {rule['description']}")
        print(f"  模式: {pattern.pattern}")
        
        if match:
            print(f"  ✅ 匹配成功")
            print(f"  匹配文本: '{match.group(0)}'")
            if hasattr(match, 'groups') and match.groups():
                print(f"  匹配组:")
                try:
                    if hasattr(match, 'groupdict') and match.groupdict():
                        for key, value in match.groupdict().items():
                            print(f"    {key}: '{value}'")
                    else:
                        for j, group in enumerate(match.groups()):
                            print(f"    组{j+1}: '{group}'")
                except:
                    print(f"    {match.groups()}")
            matching_rules.append((i+1, rule, match))
        else:
            print(f"  ❌ 不匹配")
    
    print(f"\n匹配规则总数: {len(matching_rules)}")
    return matching_rules

def test_specific_format_processing():
    """测试用户提供格式的实际处理"""
    print("\n" + "="*60)
    print("实际处理测试")
    print("="*60)
    
    # 创建测试文档
    input_file = 'test_specific_format_input.docx'
    output_file = 'test_specific_format_output.docx'
    
    doc = create_specific_test_doc()
    doc.save(input_file)
    print(f"已创建测试文档: {input_file}")
    
    # 使用MCP处理器处理
    processor = MCPBidderNameProcessor()
    company_name = "智慧足迹数据科技有限公司"
    
    print(f"\n开始处理，使用公司名称: {company_name}")
    
    result = processor.process_bidder_name(
        input_file=input_file,
        output_file=output_file,
        company_name=company_name
    )
    
    print(f"\n处理结果:")
    print(f"  成功: {result.get('success', False)}")
    print(f"  错误: {result.get('error', 'N/A')}")
    
    stats = result.get('stats', {})
    print(f"  统计信息:")
    print(f"    总替换次数: {stats.get('total_replacements', 0)}")
    print(f"    替换内容次数: {stats.get('replace_content_count', 0)}")  
    print(f"    空格填写次数: {stats.get('fill_space_count', 0)}")
    
    patterns_found = stats.get('patterns_found', [])
    print(f"    找到的模式:")
    for pattern in patterns_found:
        print(f"      规则#{pattern.get('rule_index', 'N/A')}: {pattern.get('description', 'N/A')}")
        print(f"      原文: '{pattern.get('original_text', 'N/A')}'")
    
    if result.get('success'):
        print(f"  输出文件: {output_file}")
        
        # 读取输出文件查看段落#30的结果
        try:
            output_doc = Document(output_file)
            if len(output_doc.paragraphs) > 30:
                para30_result = output_doc.paragraphs[30].text
                print(f"\n段落#30处理结果:")
                print(f"  原文: '供应商名称：                          采购编号：                   '")
                print(f"  结果: '{para30_result}'")
                print(f"  长度变化: {len('供应商名称：                          采购编号：                   ')} -> {len(para30_result)}")
                
                # 检查是否有填写
                if company_name in para30_result:
                    print(f"  ✅ 成功填写了公司名称: {company_name}")
                else:
                    print(f"  ❌ 没有填写公司名称")
        except Exception as e:
            print(f"  读取输出文件失败: {e}")
    
    return result

def main():
    """主函数"""
    print("测试用户提供的具体格式")
    print("格式：'供应商名称：                          采购编号：                   '")
    print("="*60)
    
    # 1. 规则匹配分析
    print("\n1. 规则匹配分析:")
    matching_rules = analyze_specific_format_matching()
    
    # 2. 实际处理测试
    print("\n2. 实际处理测试:")
    result = test_specific_format_processing()
    
    # 3. 结论
    print("\n" + "="*60)
    print("结论:")
    if not matching_rules:
        print("❌ 该格式没有匹配到任何规则，不会被处理")
        print("原因：长空格可能不被现有规则识别为有效占位符")
    elif len(matching_rules) > 0:
        print(f"✅ 该格式匹配到 {len(matching_rules)} 个规则:")
        for rule_idx, rule, match in matching_rules:
            print(f"   规则#{rule_idx}: {rule['description']}")
        
        if result.get('success'):
            print("✅ 处理成功")
        else:
            print("❌ 处理失败，需要检查处理逻辑")
    
    print("="*60)

if __name__ == '__main__':
    main()