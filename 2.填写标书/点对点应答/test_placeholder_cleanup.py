#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试下划线占位符清理功能
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_placeholder_test_document():
    """创建有下划线占位符的测试文档"""
    doc = Document()
    doc.add_heading('下划线占位符清理测试', 0)
    
    # 测试1: 段落#30风格 - 供应商名称和采购编号之间有下划线占位符
    para1 = doc.add_paragraph()
    para1.add_run("供应商名称：")
    para1.add_run("__________________")  # 长下划线占位符
    para1.add_run("采购编号：")  
    para1.add_run("______________")
    
    # 测试2: 另一种格式 - 在不同run中
    para2 = doc.add_paragraph()
    para2.add_run("供应商名称：")
    para2.add_run("                    ")  # 空格占位符
    para2.add_run("____")  # 短下划线
    para2.add_run("采购编号：")
    para2.add_run("              ")
    
    return doc


def analyze_placeholder_results(doc, description):
    """分析占位符清理结果"""
    print(f"\n{description}")
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and ('供应商名称' in text or '中国联合' in text):
            print(f"  段落{i}: {text}")
            
            # 检查是否有残留下划线
            underscore_count = text.count('_')
            if underscore_count > 0:
                issues.append(f"段落{i}: 还有{underscore_count}个下划线残留")
                print(f"    ❌ 还有{underscore_count}个下划线")
            else:
                print(f"    ✅ 无下划线残留")
            
            # 检查公司名称和采购编号之间的间距
            if '中国联合网络通信有限公司' in text and '采购编号' in text:
                # 查找公司名称和采购编号之间的内容
                company_end = text.find('中国联合网络通信有限公司') + len('中国联合网络通信有限公司')
                purchase_start = text.find('采购编号')
                if purchase_start > company_end:
                    middle_content = text[company_end:purchase_start]
                    print(f"    间距内容: '{middle_content}' (长度: {len(middle_content)})")
                    
                    if len(middle_content) > 10:
                        issues.append(f"段落{i}: 间距过长({len(middle_content)}个字符)")
                    elif '__' in middle_content:
                        issues.append(f"段落{i}: 间距中有下划线残留")
                    else:
                        print(f"    ✅ 间距合理")
    
    return issues


def main():
    """主测试函数"""
    print("="*60)
    print("下划线占位符清理测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_placeholder_input.docx')
    output_file = os.path.join(current_dir, 'test_placeholder_output.docx')
    
    # 创建测试文档
    print("\n1. 创建有下划线占位符的测试文档...")
    doc = create_placeholder_test_document()
    doc.save(input_file)
    
    # 分析原始格式
    analyze_placeholder_results(doc, "原始文档:")
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
    }
    
    # 处理文档
    print("\n2. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            processed_doc = Document(output_file)
            issues = analyze_placeholder_results(processed_doc, "\n3. 处理后文档:")
            
            if issues:
                print("\n❌ 占位符清理问题:")
                for issue in issues:
                    print(f"  - {issue}")
                print("\n❌ 占位符清理不完全")
            else:
                print("\n✅ 占位符清理成功！")
                print("  - 所有下划线占位符已清理")
                print("  - 字段间距合理")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)
        print(f"\n4. 输出文件: {output_file}")


if __name__ == '__main__':
    main()