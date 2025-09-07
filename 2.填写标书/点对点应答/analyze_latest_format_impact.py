#!/usr/bin/env python3
"""
分析最新输出文件的格式影响问题
对比处理前后的run结构，找出精确run修改法仍然影响格式的原因
"""

import os
import sys
from pathlib import Path
from docx import Document
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('format_impact_analysis.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def detailed_run_analysis(paragraph, paragraph_name):
    """详细分析段落的run结构"""
    logger.info(f"\n📊 {paragraph_name}:")
    logger.info(f"  完整文本: {paragraph.text}")
    logger.info(f"  Run数量: {len(paragraph.runs)}")
    
    for i, run in enumerate(paragraph.runs):
        if run.text:
            # 格式信息
            font_name = run.font.name or "默认"
            font_size = run.font.size.pt if run.font.size else "默认"
            bold = run.font.bold if run.font.bold is not None else "默认"
            italic = run.font.italic if run.font.italic is not None else "默认" 
            underline = run.font.underline if run.font.underline is not None else "默认"
            
            # 颜色信息
            color = "默认"
            if run.font.color and run.font.color.rgb:
                color = f"RGB({run.font.color.rgb.r},{run.font.color.rgb.g},{run.font.color.rgb.b})"
            
            logger.info(f"    Run {i+1}: '{run.text}'")
            logger.info(f"      格式: 字体={font_name}, 大小={font_size}, 粗体={bold}, 斜体={italic}, 下划线={underline}, 颜色={color}")

def analyze_specific_problematic_areas(doc):
    """分析具体的问题区域"""
    logger.info("\n🎯 重点分析可能影响格式的区域:")
    logger.info("="*80)
    
    problem_areas = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        # 查找包含项目名称和采购编号的段落
        if '（所属运营商数据）' in text and '（64525343）' in text:
            problem_areas.append((i, paragraph, "包含项目名称和采购编号的段落"))
        elif '（所属运营商数据）' in text:
            problem_areas.append((i, paragraph, "包含项目名称的段落"))
        elif '（64525343）' in text:
            problem_areas.append((i, paragraph, "包含采购编号的段落"))
        elif '智慧足迹数据科技有限公司' in text and len(paragraph.runs) > 5:
            problem_areas.append((i, paragraph, "包含公司名称的多run段落"))
    
    for para_idx, para, description in problem_areas:
        logger.info(f"\n🔍 段落 #{para_idx}: {description}")
        detailed_run_analysis(para, f"段落{para_idx}")
        
        # 检查run格式一致性
        if len(para.runs) > 1:
            first_run = para.runs[0]
            inconsistent_runs = []
            
            for i, run in enumerate(para.runs[1:], 1):
                if run.text and (
                    run.font.name != first_run.font.name or
                    run.font.size != first_run.font.size or
                    run.font.bold != first_run.font.bold or
                    run.font.italic != first_run.font.italic or
                    run.font.underline != first_run.font.underline
                ):
                    inconsistent_runs.append(i)
            
            if inconsistent_runs:
                logger.info(f"    ⚠️ 格式不一致的run: {inconsistent_runs}")
            else:
                logger.info(f"    ✅ 所有run格式一致")

def main():
    logger.info("🔍 开始分析最新输出文件的格式影响")
    logger.info("="*80)
    
    # 分析最新输出文件
    output_file = "outputs/docx-商务应答-20250906_105452.docx"
    
    if not os.path.exists(output_file):
        logger.error(f"文件不存在: {output_file}")
        return
    
    try:
        doc = Document(output_file)
        logger.info(f"成功加载文档: {output_file}")
        logger.info(f"文档总段落数: {len(doc.paragraphs)}")
        
        # 分析具体问题区域
        analyze_specific_problematic_areas(doc)
        
        # 总结分析结果
        logger.info("\n" + "="*80)
        logger.info("📋 格式影响分析总结:")
        logger.info("="*80)
        logger.info("基于日志分析，精确run修改法的实际执行情况:")
        logger.info("1. ✅ 项目名称 - 单run替换成功")
        logger.info("2. ⚠️ 采购编号 - 跨run替换（涉及14个run）")
        logger.info("3. ⚠️ 供应商信息 - 跨run替换（涉及16个run）")
        logger.info("")
        logger.info("💡 可能的格式影响原因:")
        logger.info("1. 跨run替换时，虽然保留了其他run，但可能改变了run的边界")
        logger.info("2. 智能格式保持替换可能在重组run时影响了细微的格式属性")
        logger.info("3. Word文档的某些隐藏格式属性可能没有被完全保留")
        
    except Exception as e:
        logger.error(f"分析失败: {e}", exc_info=True)

if __name__ == "__main__":
    main()