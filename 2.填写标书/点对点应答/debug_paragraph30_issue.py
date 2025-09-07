#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专门调试段落#30供应商名称填写问题
"""

import os
import re
import logging
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_paragraph30_test_doc():
    """创建段落#30的测试文档，模拟用户上传的格式"""
    doc = Document()
    doc.add_heading('段落#30供应商名称填写测试', 0)
    
    # 添加一些前置段落来模拟真实文档结构
    for i in range(29):
        doc.add_paragraph(f'段落 #{i+1}: 前置内容...')
    
    # 段落#30 - 关键测试段落，这是用户上传文档中的实际格式
    para30 = doc.add_paragraph()
    run1 = para30.add_run("供应商名称：")
    run2 = para30.add_run("__ ")
    run3 = para30.add_run("采购编号：")  
    run4 = para30.add_run("GXTC-C-251590031")
    
    print(f"段落#30构成:")
    print(f"  Run1: '{run1.text}'")
    print(f"  Run2: '{run2.text}'")  
    print(f"  Run3: '{run3.text}'")
    print(f"  Run4: '{run4.text}'")
    print(f"  完整文本: '{''.join([run1.text, run2.text, run3.text, run4.text])}'")
    
    return doc

def analyze_paragraph30_matching():
    """分析段落#30的规则匹配情况"""
    processor = MCPBidderNameProcessor()
    
    # 模拟段落#30的文本
    test_text = "供应商名称：__ 采购编号：GXTC-C-251590031"
    
    print(f"\n分析文本: '{test_text}'")
    print("="*60)
    
    # 检查每个规则是否匹配
    matching_rules = []
    for i, rule in enumerate(processor.bidder_patterns):
        pattern = rule['pattern']
        match = pattern.search(test_text)
        
        if match:
            print(f"✅ 规则#{i+1} 匹配: {rule['description']}")
            print(f"   模式: {pattern.pattern}")
            print(f"   匹配文本: '{match.group(0)}'")
            print(f"   匹配组: {match.groups() if hasattr(match, 'groups') else 'N/A'}")
            matching_rules.append((i+1, rule, match))
        else:
            print(f"❌ 规则#{i+1} 不匹配: {rule['description']}")
    
    print(f"\n总计匹配的规则数: {len(matching_rules)}")
    
    return matching_rules

def test_actual_processing():
    """测试实际的处理过程"""
    print("\n" + "="*60)
    print("实际处理测试")
    print("="*60)
    
    # 创建测试文档
    input_file = 'debug_para30_input.docx'
    output_file = 'debug_para30_output.docx'
    
    doc = create_paragraph30_test_doc()
    doc.save(input_file)
    print(f"已创建测试文档: {input_file}")
    
    # 使用MCP处理器处理
    processor = MCPBidderNameProcessor()
    company_name = "智慧足迹数据科技有限公司"
    
    result = processor.process_bidder_name(
        input_file=input_file,
        output_file=output_file,
        company_name=company_name
    )
    
    print(f"\n处理结果:")
    print(f"  成功: {result.get('success', False)}")
    print(f"  统计: {result.get('stats', {})}")
    
    if result.get('success'):
        print(f"  输出文件: {output_file}")
        
        # 读取输出文件查看段落#30的结果
        output_doc = Document(output_file)
        if len(output_doc.paragraphs) > 30:
            para30_result = output_doc.paragraphs[30].text
            print(f"  段落#30结果: '{para30_result}'")
        
    return result

def main():
    """主函数"""
    print("段落#30供应商名称填写问题调试")
    print("="*60)
    
    # 1. 分析规则匹配
    print("\n1. 规则匹配分析:")
    matching_rules = analyze_paragraph30_matching()
    
    # 2. 实际处理测试
    print("\n2. 实际处理测试:")
    result = test_actual_processing()
    
    # 3. 问题诊断
    print("\n3. 问题诊断:")
    if not matching_rules:
        print("❌ 问题：段落#30的文本格式没有匹配到任何规则")
        print("建议：需要添加或修改规则来支持这种格式")
    elif len(matching_rules) == 1 and matching_rules[0][1]['type'] == 'fill_space_tender_no':
        print("❌ 问题：只匹配到了采购编号规则，没有匹配到供应商名称规则") 
        print("建议：需要优化供应商名称匹配规则")
    else:
        print("✅ 规则匹配正常，需要检查处理逻辑")
        
    # 4. 推荐解决方案
    print("\n4. 推荐解决方案:")
    print("   - 检查供应商名称规则是否覆盖了'供应商名称：__ '格式")
    print("   - 确认多字段同行处理逻辑是否正确")
    print("   - 验证占位符识别逻辑")

if __name__ == '__main__':
    main()