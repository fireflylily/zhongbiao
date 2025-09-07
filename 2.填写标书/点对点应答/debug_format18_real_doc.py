#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import re
from docx import Document

def debug_format18_real_doc():
    """调试实际文档中的格式18问题"""
    
    # 使用最近处理过的文件
    recent_files = [
        "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/uploads/20250905_195135_business_template_docx",
        "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/uploads/20250903_095446_tender_document.docx"
    ]
    
    input_file = None
    for file in recent_files:
        if os.path.exists(file):
            input_file = file
            break
    
    if not input_file:
        print("❌ 没有找到可用的输入文件")
        return
    
    print(f"使用文件: {input_file}")
    
    try:
        doc = Document(input_file)
        print(f"文档加载成功，共有 {len(doc.paragraphs)} 个段落")
        
        # 格式18的正则模式
        pattern = re.compile(r'^(?P<label>供应商名称)\s*(?P<placeholder>\s{20,})\s*$')
        
        print(f"\n=== 搜索格式18模式 ===")
        found_count = 0
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text
            match = pattern.search(para_text)
            
            if match:
                found_count += 1
                print(f"\n🎯 在段落 #{para_idx} 找到格式18:")
                print(f"  段落文本: '{para_text}'")
                print(f"  文本长度: {len(para_text)}")
                print(f"  标签: '{match.group('label')}'")
                print(f"  占位符: '{match.group('placeholder')}'")
                print(f"  占位符长度: {len(match.group('placeholder'))}")
                
                # 分析runs结构
                print(f"  总run数: {len(paragraph.runs)}")
                for run_idx, run in enumerate(paragraph.runs):
                    print(f"    Run #{run_idx}: '{run.text}'")
                    if match.group('label') in run.text:
                        print(f"      ✅ 包含标签")
                    if len(run.text) > 20 and run.text.isspace():
                        print(f"      ✅ 可能是占位符run")
                
                # 模拟处理逻辑
                print(f"  \n=== 模拟处理逻辑 ===")
                company_name = "智慧足迹数据科技有限公司"
                new_text = f"{match.group('label')} {company_name}"
                print(f"  新文本: '{new_text}'")
                
                # 检查是否能在run中找到标签
                label_found_in_run = False
                for run_idx, run in enumerate(paragraph.runs):
                    if match.group('label') in run.text:
                        print(f"  ✅ 在Run #{run_idx}中找到标签")
                        old_text = run.text
                        new_run_text = old_text.replace(match.group(0), new_text)
                        print(f"  替换: '{old_text}' -> '{new_run_text}'")
                        label_found_in_run = True
                        break
                
                if not label_found_in_run:
                    print(f"  ❌ 在任何run中都找不到标签，这是问题所在！")
                    print(f"  可能原因: 标签被分割到多个run中")
        
        if found_count == 0:
            print("❌ 文档中没有找到格式18模式")
            
            # 搜索可能的供应商名称行
            print(f"\n=== 搜索可能的供应商名称行 ===")
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                if "供应商名称" in para_text and len(para_text) > 15:
                    print(f"段落 #{para_idx}: '{para_text}' (长度: {len(para_text)})")
                    
                    # 检查是否有长空格
                    space_count = para_text.count(' ')
                    if space_count > 10:
                        print(f"  可能的候选: 包含 {space_count} 个空格")
        else:
            print(f"\n✅ 总共找到 {found_count} 个格式18模式")
            
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_format18_real_doc()