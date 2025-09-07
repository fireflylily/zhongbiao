#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def test_format18_on_para25():
    """测试格式18对段落25类型的处理"""
    
    # 创建测试文档，模拟段落25的内容
    test_input = "test_para25_input.docx"
    test_output = "test_para25_output.docx"
    
    doc = Document()
    # 添加一个完全模拟段落25的内容
    doc.add_paragraph("供应商名称                                    ")  # 36个空格，与段落25相同
    doc.save(test_input)
    
    print("创建测试文档，内容与段落25相同")
    
    # 读取并验证
    doc_check = Document(test_input)
    para_text = doc_check.paragraphs[0].text
    print(f"段落内容: '{para_text}'")
    print(f"长度: {len(para_text)}")
    print(f"空格数: {para_text.count(' ')}")
    
    # 检查是否符合格式18
    pattern18 = re.compile(r'^(?P<label>供应商名称)\s*(?P<placeholder>\s{20,})\s*$')
    match = pattern18.search(para_text)
    
    if match:
        print(f"✅ 符合格式18模式")
        print(f"  标签: '{match.group('label')}'")
        print(f"  占位符长度: {len(match.group('placeholder'))}")
    else:
        print(f"❌ 不符合格式18模式")
        
    # 使用处理器处理
    print("\n=== 开始MCP处理 ===")
    processor = MCPBidderNameProcessor()
    company_name = "中国联合网络通信有限公司"
    
    result = processor.process_bidder_name(
        input_file=test_input,
        output_file=test_output,
        company_name=company_name
    )
    
    print(f"处理结果: {result}")
    
    # 检查输出
    if result.get('success'):
        doc_after = Document(test_output)
        para_after = doc_after.paragraphs[0].text
        print(f"\n处理后内容: '{para_after}'")
        
        if company_name in para_after:
            print("✅ 成功填写公司名称！")
        else:
            print("❌ 未填写公司名称")
            
        # 查看处理统计
        stats = result.get('stats', {})
        patterns = stats.get('patterns_found', [])
        if patterns:
            print("\n匹配的规则：")
            for p in patterns:
                print(f"  - {p['description']}")
        else:
            print("\n❌ 没有匹配任何规则！")
            
    # 清理测试文件
    import os
    if os.path.exists(test_input):
        os.remove(test_input)
    if os.path.exists(test_output):
        os.remove(test_output)
    print("\n测试文件已清理")

if __name__ == "__main__":
    test_format18_on_para25()