#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
深度调试占位符清理问题
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def detailed_run_analysis(paragraph):
    """详细分析段落中的每个run"""
    print("  详细run分析:")
    for i, run in enumerate(paragraph.runs):
        if run.text:
            print(f"    run #{i}: '{run.text}' (长度: {len(run.text)})")
            # 分析空格数量
            space_count = run.text.count(' ')
            tab_count = run.text.count('\t')
            print(f"             空格数: {space_count}, Tab数: {tab_count}")
    print(f"  总文本: '{paragraph.text}'")
    print(f"  总长度: {len(paragraph.text)}")


def main():
    """主调试函数"""
    print("=" * 60)
    print("深度调试占位符清理问题")
    print("=" * 60)
    
    input_file = os.path.join(current_dir, 'debug_placeholder_input.docx')
    output_file = os.path.join(current_dir, 'debug_placeholder_output.docx')
    
    # 创建测试文档
    print("\n1. 创建测试文档...")
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("供应商名称：                          采购编号：                   ")
    doc.save(input_file)
    
    print("原始状态:")
    detailed_run_analysis(para)
    
    # 准备公司信息
    company_info = {
        "companyName": "智慧足迹数据科技有限公司",
    }
    
    # 处理文档
    print("\n2. 执行处理...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果文档
            processed_doc = Document(output_file)
            processed_para = None
            
            # 找到对应的段落
            for para in processed_doc.paragraphs:
                if '智慧足迹数据科技有限公司' in para.text or 'GXTC-C-251590031' in para.text:
                    processed_para = para
                    break
            
            if processed_para:
                print("\n3. 处理后状态:")
                detailed_run_analysis(processed_para)
                
                # 计算各种空格占位符
                text = processed_para.text
                consecutive_spaces = []
                current_spaces = 0
                for char in text:
                    if char == ' ' or char == '\t':
                        current_spaces += 1
                    else:
                        if current_spaces >= 3:
                            consecutive_spaces.append(current_spaces)
                        current_spaces = 0
                if current_spaces >= 3:
                    consecutive_spaces.append(current_spaces)
                
                print(f"\n4. 发现的连续空格组: {consecutive_spaces}")
                if consecutive_spaces:
                    print(f"   ⚠️ 仍有 {len(consecutive_spaces)} 个连续空格组需要清理")
                    print(f"   最长连续空格: {max(consecutive_spaces)} 个")
                else:
                    print("   ✅ 没有发现需要清理的连续空格组")
                    
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)
        print(f"\n5. 输出文件: {output_file}")


if __name__ == '__main__':
    main()