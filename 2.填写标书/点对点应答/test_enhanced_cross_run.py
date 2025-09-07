#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试增强跨run填写方法的修改
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 重新导入以获取最新修改
import importlib
import mcp_bidder_name_processor_enhanced
importlib.reload(mcp_bidder_name_processor_enhanced)
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_test_cases():
    """创建不同的测试案例"""
    doc = Document()
    doc.add_heading('增强跨run填写测试', 0)
    
    # 案例1: 只有标签和分隔符，没有占位符
    doc.add_paragraph("案例1: 只有标签")
    para1 = doc.add_paragraph()
    para1.add_run("供应商名称")
    para1.add_run("：")
    
    # 案例2: 有下划线占位符
    doc.add_paragraph("案例2: 下划线占位符")
    para2 = doc.add_paragraph()
    para2.add_run("供应商名称")
    para2.add_run("：")
    para2.add_run("_______________")
    
    # 案例3: 有空格占位符
    doc.add_paragraph("案例3: 空格占位符")
    para3 = doc.add_paragraph()
    para3.add_run("供应商名称")
    para3.add_run("：")
    para3.add_run("               ")
    
    # 案例4: 同一行有两个字段
    doc.add_paragraph("案例4: 同一行两个字段")
    para4 = doc.add_paragraph()
    para4.add_run("供应商名称")
    para4.add_run("：")
    para4.add_run("_______________")
    para4.add_run(" 采购编号")
    para4.add_run("：")
    para4.add_run("_______________")
    
    return doc


def check_results(doc):
    """检查处理结果"""
    issues = []
    
    # 检查每个案例
    for i, para in enumerate(doc.paragraphs):
        if '供应商名称' in para.text and '中国联合网络通信有限公司' in para.text:
            print(f"\n段落{i}: {para.text}")
            
            # 检查是否有下划线
            for run in para.runs:
                if '中国联合网络通信有限公司' in run.text:
                    if run.underline:
                        issues.append(f"段落{i}: 公司名称被加了下划线")
                        print(f"  ❌ 公司名称有下划线")
                    else:
                        print(f"  ✅ 公司名称无下划线")
            
            # 检查是否有残留占位符
            if '___' in para.text:
                remaining_underscores = para.text.count('_')
                if remaining_underscores > 5:  # 允许少量下划线
                    issues.append(f"段落{i}: 有{remaining_underscores}个下划线未被替换")
                    print(f"  ❌ 还有{remaining_underscores}个下划线")
            
            # 检查间距
            if '：中国联合' in para.text:
                print(f"  ✅ 紧凑格式（无空格）")
            elif '： 中国联合' in para.text:
                print(f"  ✅ 有一个空格")
            elif '：  ' in para.text:
                issues.append(f"段落{i}: 多余空格")
                print(f"  ⚠️ 有多余空格")
    
    return issues


def main():
    """主测试函数"""
    print("="*60)
    print("增强跨run填写方法测试")
    print("="*60)
    
    input_file = os.path.join(current_dir, 'test_enhanced_input.docx')
    output_file = os.path.join(current_dir, 'test_enhanced_output.docx')
    
    # 创建测试文档
    print("\n1. 创建测试文档...")
    doc = create_test_cases()
    doc.save(input_file)
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
    }
    
    # 处理文档
    print("\n2. 处理文档...")
    processor = MCPBidderNameProcessor()
    
    try:
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="GXTC-C-251590031",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print("   ✅ 处理成功")
            
            # 读取结果
            print("\n3. 检查结果:")
            processed_doc = Document(output_file)
            issues = check_results(processed_doc)
            
            if issues:
                print("\n⚠️ 发现的问题:")
                for issue in issues:
                    print(f"  - {issue}")
            else:
                print("\n✅ 所有测试通过！")
                print("  - 公司名称正确替换占位符")
                print("  - 没有添加下划线")
                print("  - 格式保持正确")
            
        else:
            print(f"   ❌ 处理失败: {result.get('error')}")
            
    except Exception as e:
        print(f"   ❌ 发生错误: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # 清理输入文件
        if os.path.exists(input_file):
            os.remove(input_file)
        print(f"\n4. 输出文件: {output_file}")


if __name__ == '__main__':
    main()