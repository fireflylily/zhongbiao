#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import re
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def test_format18_no_separator():
    """测试格式18 - 无冒号的供应商名称填写"""
    
    # 创建测试文档
    test_input = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/test_format18_input.docx"
    test_output = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/test_format18_output.docx"
    
    # 创建测试文档
    doc = Document()
    
    # 添加测试格式
    para1 = doc.add_paragraph()
    run1 = para1.add_run("供应商名称")
    run2 = para1.add_run("                                    ")  # 长空格
    
    para2 = doc.add_paragraph("其他文本内容")
    
    # 保存测试文档
    doc.save(test_input)
    print(f"创建测试文档: {test_input}")
    
    try:
        # 测试处理
        processor = MCPBidderNameProcessor()
        company_name = "智慧足迹数据科技有限公司"
        
        print(f"开始处理，公司名称: {company_name}")
        
        result = processor.process_bidder_name(
            input_file=test_input,
            output_file=test_output,
            company_name=company_name
        )
        
        print(f"处理结果: {result}")
        
        if result.get('success'):
            print("✅ 处理成功!")
            
            # 检查结果
            stats = result.get('stats', {})
            patterns_found = stats.get('patterns_found', [])
            
            print(f"找到的模式:")
            for pattern in patterns_found:
                print(f"  - {pattern.get('description')}: {pattern.get('original_text')}")
                if '无冒号' in pattern.get('description', ''):
                    print(f"    ✅ 格式18已处理!")
            
            # 验证输出文档
            if os.path.exists(test_output):
                output_doc = Document(test_output)
                for para_idx, paragraph in enumerate(output_doc.paragraphs):
                    if company_name in paragraph.text:
                        print(f"✅ 在段落#{para_idx}找到填写结果: {paragraph.text}")
                        break
                else:
                    print("❌ 未在输出文档中找到填写结果")
            
            return True
        else:
            print("❌ 处理失败")
            return False
            
    except Exception as e:
        print(f"❌ 测试过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理测试文件
        for file_path in [test_input, test_output]:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    print(f"清理测试文件: {file_path}")
                except:
                    pass

if __name__ == "__main__":
    test_format18_no_separator()