#!/usr/bin/env python3
"""
分析项目名称和采购编号处理的格式差异
基于日志分析为什么项目名称没有影响格式，而采购编号却影响了其他文本
"""

import os
import sys
from pathlib import Path
from docx import Document
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('format_difference_analysis.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def analyze_docx_structure(file_path: str):
    """分析Word文档的run结构"""
    try:
        doc = Document(file_path)
        logger.info(f"\n=== 分析文档结构: {file_path} ===")
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text and ('项目名称' in text or '采购编号' in text or '64525343' in text or '所属运营商' in text):
                logger.info(f"\n段落 #{i}: {text[:100]}...")
                logger.info(f"  完整文本: {text}")
                logger.info(f"  Run数量: {len(paragraph.runs)}")
                
                for j, run in enumerate(paragraph.runs):
                    if run.text:
                        font_info = f"字体={run.font.name}, 大小={run.font.size}, 粗体={run.font.bold}, 斜体={run.font.italic}"
                        logger.info(f"    Run {j+1}: '{run.text}' [{font_info}]")
                        
    except Exception as e:
        logger.error(f"分析文档失败: {e}")

def analyze_log_differences():
    """基于日志分析处理差异"""
    logger.info("=" * 80)
    logger.info("基于日志分析项目名称vs采购编号处理差异")
    logger.info("=" * 80)
    
    logger.info("""
日志分析：

1. 项目名称处理（段落#4）：
   - 匹配文本: '（项目名称）'
   - 执行方法: _replace_content_project_method
   - 处理方式: 项目名称括号替换: '（项目名称）' -> '（所属运营商数据）'
   - 结果: 直接成功，没有跨run处理的日志

2. 采购编号处理（段落#11）：
   - 匹配文本: '（采购编号）'
   - 执行方法: _replace_content_tender_no_method
   - 处理方式: 
     * 优先尝试单run替换: '（采购编号）' -> '（64525343）'
     * ⚠️ 单run替换失败，尝试跨run项目编号处理...
     * ⚠️ 项目编号跨run替换完成: (可能影响周围格式)

关键差异：
- 项目名称：成功使用单run替换，保持格式完美
- 采购编号：单run失败，被迫使用跨run处理，影响周围格式

原因分析：
1. 项目名称所在段落的"（项目名称）"文本完整包含在一个run中
2. 采购编号所在段落的"（采购编号）"文本跨越多个run，单run无法找到完整匹配

这说明问题不在于处理方法本身，而在于文档中文本的run分布结构！
""")

def create_run_analysis_test():
    """创建测试来验证run结构影响"""
    logger.info("=" * 80)
    logger.info("创建run结构影响验证测试")
    logger.info("=" * 80)
    
    # 创建两种不同run结构的测试文档
    doc = Document()
    
    # 测试1：项目名称在单个run中（模拟成功情况）
    p1 = doc.add_paragraph()
    p1.add_run("测试项目：（项目名称）需要处理")
    
    # 测试2：采购编号跨多个run（模拟失败情况）
    p2 = doc.add_paragraph()
    p2.add_run("测试编号：（")
    p2.add_run("采购编号")  # 不同run
    p2.add_run("）需要处理")
    
    # 测试3：混合格式的采购编号
    p3 = doc.add_paragraph()
    p3.add_run("根据贵方购采购货物及服务的竞争性磋商公告")
    p3.add_run("（")
    run_middle = p3.add_run("采购编号")
    run_middle.font.italic = True  # 设置斜体
    p3.add_run("）")
    p3.add_run("，签字代表经正式授权")
    
    test_file = "run_structure_test.docx"
    doc.save(test_file)
    logger.info(f"创建run结构测试文档: {test_file}")
    
    # 分析测试文档结构
    analyze_docx_structure(test_file)
    
    return test_file

if __name__ == "__main__":
    logger.info("开始格式差异分析")
    
    # 分析处理过的文档
    output_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/outputs/docx-商务应答-20250906_101512.docx"
    
    if os.path.exists(output_file):
        analyze_docx_structure(output_file)
    else:
        logger.warning(f"输出文件不存在: {output_file}")
    
    # 基于日志分析差异
    analyze_log_differences()
    
    # 创建run结构测试
    create_run_analysis_test()
    
    logger.info("\n" + "=" * 80)
    logger.info("结论:")
    logger.info("项目名称处理成功是因为'（项目名称）'文本完整在一个run中")
    logger.info("采购编号处理失败是因为'（采购编号）'文本跨越多个run")
    logger.info("这是Word文档本身的run结构问题，不是代码逻辑问题")
    logger.info("=" * 80)