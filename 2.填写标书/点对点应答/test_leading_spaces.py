#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试前导空格的盖章格式处理
"""

from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor

def test_leading_spaces():
    """测试各种前导空格情况"""
    
    # 创建测试文档
    doc = Document()
    
    # 添加各种前导空格的格式
    test_formats = [
        "    公司名称（全称、盖章）：________________",  # 4个空格
        "        供应商名称（盖章）：________________",  # 8个空格
        "\t\t投标人名称（盖章）：________________",    # 2个制表符
        "  \t  单位名称及公章：________________",       # 混合空格和制表符
        "公司名称（全称、盖章）：________________",     # 无前导空格
        "1.  公司名称（盖章）：________________",       # 数字编号+空格
        "   2. 供应商全称及公章：________________",     # 空格+数字编号
    ]
    
    for i, format_text in enumerate(test_formats):
        para = doc.add_paragraph()
        para.add_run(format_text)
        if i < len(test_formats) - 1:
            doc.add_paragraph()  # 空行分隔
    
    input_file = "test_leading_spaces_input.docx"
    doc.save(input_file)
    print(f"✅ 创建测试文档: {input_file}")
    
    # 处理文档
    processor = MCPBidderNameProcessor()
    company_name = "上海智慧足迹数据科技有限公司"
    output_file = "test_leading_spaces_output.docx"
    
    print(f"🚀 开始处理文档...")
    result = processor.process_bidder_name(input_file, output_file, company_name)
    
    if result.get('success', False):
        print("✅ 文档处理成功!")
        
        # 验证结果
        output_doc = Document(output_file)
        print("\n📄 处理结果验证:")
        
        for i, para in enumerate(output_doc.paragraphs):
            if para.text.strip():
                contains_company = company_name in para.text
                contains_placeholder = "___" in para.text
                status = "✅ 已替换" if contains_company and not contains_placeholder else "❌ 未替换"
                
                # 显示原始格式和结果
                original = test_formats[i//2] if i//2 < len(test_formats) else "未知"
                print(f"   原始: {repr(original[:30])}...")
                print(f"   结果: {repr(para.text[:50])}... {status}")
                print()
    else:
        print("❌ 文档处理失败!")

if __name__ == "__main__":
    print("🧪 测试前导空格的盖章格式处理")
    print("=" * 60)
    
    test_leading_spaces()
    
    print("=" * 60)
    print("✅ 测试完成")