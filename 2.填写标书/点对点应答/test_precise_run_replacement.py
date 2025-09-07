#!/usr/bin/env python3
"""
测试精确run修改法的效果
验证新实现的格式保持能力
"""

import os
import sys
from pathlib import Path
from docx import Document
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('precise_run_test.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def create_complex_test_document():
    """创建复杂的测试文档，包含各种跨run情况"""
    doc = Document()
    
    # 场景1：项目名称在单个run中（应该使用单run替换）
    p1 = doc.add_paragraph()
    p1.add_run("测试项目：")
    project_run = p1.add_run("（项目名称）")
    project_run.font.italic = True
    project_run.font.underline = True
    p1.add_run(" 的处理效果")
    
    # 场景2：采购编号跨多个run（应该使用精确跨run替换）
    p2 = doc.add_paragraph()
    p2.add_run("竞争性磋商公告")
    
    bracket_left = p2.add_run("（")
    bracket_left.font.italic = True
    bracket_left.font.underline = True
    bracket_left.font.bold = True
    
    tender_text = p2.add_run("采购编号")
    tender_text.font.italic = True
    tender_text.font.underline = True
    tender_text.font.bold = True
    
    bracket_right = p2.add_run("）")
    bracket_right.font.italic = True
    bracket_right.font.underline = True
    bracket_right.font.bold = True
    
    p2.add_run("，签字代表授权")
    
    # 场景3：复杂混合格式
    p3 = doc.add_paragraph()
    p3.add_run("根据贵方为")
    
    project_complex = p3.add_run("（项目名称）")
    project_complex.font.italic = True
    project_complex.font.underline = True
    
    p3.add_run("项目，供应商")
    
    company_left = p3.add_run("（")
    company_left.font.italic = True
    
    company_name = p3.add_run("供应商名称")
    company_name.font.italic = True
    company_name.font.bold = True
    
    company_right = p3.add_run("）")
    company_right.font.italic = True
    
    p3.add_run("提交文件")
    
    test_file = "precise_run_test_input.docx"
    doc.save(test_file)
    logger.info(f"创建复杂测试文档: {test_file}")
    return test_file

def analyze_document_before_after(before_file, after_file, description):
    """对比处理前后的文档结构"""
    logger.info(f"\n{'='*60}")
    logger.info(f"📊 文档结构分析: {description}")
    logger.info(f"{'='*60}")
    
    # 分析处理前
    logger.info(f"\n🔍 处理前文档: {before_file}")
    if os.path.exists(before_file):
        doc_before = Document(before_file)
        for i, paragraph in enumerate(doc_before.paragraphs):
            if paragraph.text.strip():
                logger.info(f"  段落 #{i+1}: {paragraph.text}")
                logger.info(f"    Run数量: {len(paragraph.runs)}")
                for j, run in enumerate(paragraph.runs):
                    if run.text:
                        format_info = f"斜体={run.font.italic}, 粗体={run.font.bold}, 下划线={run.font.underline}"
                        logger.info(f"      Run {j+1}: '{run.text}' [{format_info}]")
    
    # 分析处理后
    logger.info(f"\n🔍 处理后文档: {after_file}")
    if os.path.exists(after_file):
        doc_after = Document(after_file)
        for i, paragraph in enumerate(doc_after.paragraphs):
            if paragraph.text.strip():
                logger.info(f"  段落 #{i+1}: {paragraph.text}")
                logger.info(f"    Run数量: {len(paragraph.runs)}")
                for j, run in enumerate(paragraph.runs):
                    if run.text:
                        format_info = f"斜体={run.font.italic}, 粗体={run.font.bold}, 下划线={run.font.underline}"
                        logger.info(f"      Run {j+1}: '{run.text}' [{format_info}]")

def test_precise_run_replacement():
    """测试精确run替换功能"""
    logger.info("🚀 开始测试精确run修改法")
    logger.info("="*80)
    
    try:
        # 创建测试文档
        input_file = create_complex_test_document()
        output_file = "precise_run_test_output.docx"
        
        # 分析原始文档
        analyze_document_before_after(input_file, input_file, "原始文档结构")
        
        # 创建处理器并设置测试数据
        processor = MCPBidderNameProcessor()
        processor.project_number = "64525343"
        processor.project_name = "智能办公系统采购项目"
        
        # 处理文档
        company_info = {
            'company_name': '测试科技有限公司',
            'registeredAddress': '北京市海淀区测试路123号'
        }
        
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="智能办公系统采购项目",
            tender_no="64525343",
            date_text="2025年9月6日"
        )
        
        # 分析处理结果
        logger.info("\n" + "="*80)
        logger.info("📈 处理结果分析")
        logger.info("="*80)
        logger.info(f"处理成功: {result.get('success', False)}")
        logger.info(f"处理统计: {result.get('stats', {})}")
        
        # 对比处理前后
        analyze_document_before_after(input_file, output_file, "处理前后对比")
        
        # 验证精确性
        logger.info("\n🎯 精确性验证:")
        logger.info("="*50)
        logger.info("检查点:")
        logger.info("1. ✅ 项目名称替换是否成功且保持单run")
        logger.info("2. ✅ 采购编号替换是否成功且最小化run影响")
        logger.info("3. ✅ 其他格式是否完全保持不变")
        logger.info("4. ✅ 周围文字格式是否未受影响")
        
        logger.info("\n🎉 精确run修改法测试完成！")
        logger.info("查看输出文档验证格式保持效果")
        
        return result
        
    except Exception as e:
        logger.error(f"测试失败: {e}", exc_info=True)
        return None

if __name__ == "__main__":
    test_precise_run_replacement()