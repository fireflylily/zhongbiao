#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试空格问题修复效果
验证方案A是否成功解决了填写值前面有大量空格的问题
"""

import os
import sys
from datetime import datetime
from docx import Document

# 添加当前目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor


def create_spacing_problem_document():
    """创建有空格问题的测试文档"""
    doc = Document()
    doc.add_heading('空格问题修复测试文档', 0)
    
    # 模拟图片中看到的有问题的格式
    problem_fields = [
        "供应商名称：中国联合网络通信有限公司",
        "单位性质：",
        "成立时间：                    年    月    日",
        "经营期限：",
        "法定代表人（负责人）姓名：                    陈忠岳",
        "性别：                        ",
        "职务：                        ",
        "",
        "特此证明。",
        "1. 与本磋商有关的一切正式往来信函请寄：",
        "地址：                                      北京市西城区金融大街21号传真",
        "电话                                        电子邮件",
        "service@chinaunicom.cn"
    ]
    
    for field in problem_fields:
        if field.strip():
            para = doc.add_paragraph(field)
        else:
            doc.add_paragraph("")  # 空行
    
    # 添加表格测试
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # 表格内容 - 模拟有空格问题的情况
    table_data = [
        ['成立时间', '                    年    月    日'],
        ['法定代表人', '                    '],
        ['地址', '                                      '],
        ['电话', '                        ']
    ]
    
    for i, (label, content) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = content
    
    return doc


def main():
    """主测试函数"""
    print("开始测试空格问题修复效果...")
    
    # 创建测试文档
    input_file = os.path.join(current_dir, 'test_spacing_fix_input.docx')
    output_file = os.path.join(current_dir, 'test_spacing_fix_output.docx')
    
    print(f"创建测试文档: {input_file}")
    doc = create_spacing_problem_document()
    doc.save(input_file)
    
    # 准备公司信息
    company_info = {
        "companyName": "中国联合网络通信有限公司",
        "establishDate": "2000-04-21",
        "legalRepresentative": "陈忠岳",
        "socialCreditCode": "91110000710939135P",
        "registeredCapital": "22539208.432769万元",
        "fixedPhone": "010-66258899",
        "registeredAddress": "北京市西城区金融大街21号",
        "fax": "010-66258866",
        "email": "service@chinaunicom.cn"
    }
    
    # 使用处理器处理文档
    processor = MCPBidderNameProcessor()
    
    try:
        print(f"\n处理文档...")
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-2025-001",
            date_text=datetime.now().strftime("%Y年%m月%d日")
        )
        
        if result.get('success'):
            print(f"处理成功！")
            print(f"输出文件: {output_file}")
            
            # 详细验证结果
            print("\n=== 详细验证修复效果 ===")
            verify_doc = Document(output_file)
            
            # 检查的问题字段
            problem_checks = [
                {
                    'pattern': '成立时间',
                    'expected_compact': '成立时间：2000年4月21日',
                    'problem_format': '成立时间：                    2000年4月21日'
                },
                {
                    'pattern': '法定代表人',
                    'expected_compact': '法定代表人（负责人）姓名：陈忠岳',
                    'problem_format': '法定代表人（负责人）姓名：                    陈忠岳'
                },
                {
                    'pattern': '地址',
                    'expected_compact': '地址：北京市西城区金融大街21号',
                    'problem_format': '地址：                                      北京市西城区金融大街21号'
                }
            ]
            
            results = {
                'fixed': [],
                'still_problems': [],
                'not_found': []
            }
            
            # 收集所有文本
            all_paragraphs = []
            for para in verify_doc.paragraphs:
                if para.text.strip():
                    all_paragraphs.append(para.text.strip())
            
            for table in verify_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_paragraphs.append(cell.text.strip())
            
            # 检查每个问题字段
            for check in problem_checks:
                pattern = check['pattern']
                expected = check['expected_compact']
                problem_format = check['problem_format']
                
                found_texts = [text for text in all_paragraphs if pattern in text]
                
                if not found_texts:
                    results['not_found'].append(pattern)
                    continue
                
                for text in found_texts:
                    print(f"\n字段: {pattern}")
                    print(f"找到文本: '{text}'")
                    print(f"文本长度: {len(text)}")
                    
                    # 检查是否修复
                    if text == expected or (expected in text and len(text) <= len(expected) + 5):
                        print(f"✅ 已修复 - 格式紧凑")
                        results['fixed'].append({
                            'field': pattern,
                            'text': text,
                            'status': '修复成功'
                        })
                    elif problem_format in text or len(text) > len(expected) + 20:
                        print(f"❌ 仍有问题 - 存在大量空格")
                        results['still_problems'].append({
                            'field': pattern,
                            'text': text,
                            'status': '仍有空格问题'
                        })
                    else:
                        print(f"⚠️  部分修复 - 格式有改善但不完美")
                        results['fixed'].append({
                            'field': pattern,
                            'text': text,
                            'status': '部分修复'
                        })
            
            # 总结结果
            print(f"\n=== 修复效果总结 ===")
            print(f"✅ 成功修复: {len(results['fixed'])} 个字段")
            print(f"❌ 仍有问题: {len(results['still_problems'])} 个字段")
            print(f"⚠️  未找到: {len(results['not_found'])} 个字段")
            
            if results['fixed']:
                print("\n✅ 成功修复的字段:")
                for item in results['fixed']:
                    print(f"  {item['field']}: {item['status']}")
            
            if results['still_problems']:
                print("\n❌ 仍有问题的字段:")
                for item in results['still_problems']:
                    print(f"  {item['field']}: {item['status']}")
                    print(f"    文本: '{item['text']}'")
            
            if results['not_found']:
                print("\n⚠️  未找到的字段:")
                for field in results['not_found']:
                    print(f"  {field}")
            
            # 返回结果
            return len(results['still_problems']) == 0 and len(results['fixed']) > 0
            
        else:
            print(f"处理失败: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = main()
    if success:
        print("\n🎉 空格问题修复测试通过！")
    else:
        print("\n💥 空格问题修复测试失败！")
    sys.exit(0 if success else 1)