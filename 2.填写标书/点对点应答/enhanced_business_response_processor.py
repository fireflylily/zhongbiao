#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版商务应答处理器
解决第5行"供应商名称："跨run拆分识别问题
增加强化调试日志和更健壮的模式匹配
"""

import os
import re
import logging
import json
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Mm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 配置增强日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - [%(lineno)d] %(message)s',
    handlers=[
        logging.FileHandler('enhanced_business_response.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class EnhancedBusinessResponseProcessor:
    """增强版商务应答处理器 - 解决跨run拆分问题"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("请安装python-docx库：pip install python-docx")
        
        # 原有的投标人名称匹配规则
        self.original_bidder_patterns = [
            # === 括号格式（最高优先级）===
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>请填写\s*供应商名称)\s*(?P<suffix>[\)）])'),
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>请填写\s*投标人名称)\s*(?P<suffix>[\)）])'),
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>请填写\s*公司名称)\s*(?P<suffix>[\)）])'),
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>供应商全称)\s*(?P<suffix>[\)）])'),
            re.compile(r'(?P<prefix>[\(（])\s*(?P<content>供应商名称)\s*(?P<suffix>[\)）])'),
            
            # === 具体格式（中等优先级）===
            re.compile(r'^(?P<label>公司名称（全称、盖章）)\s*(?P<sep>[:：])\s*(?P<placeholder>_{3,})\s*$'),
            re.compile(r'^(?P<label>公司名称（全称、盖章）)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            re.compile(r'^(?P<label>公司名称（盖章）)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            re.compile(r'^(?P<label>供应商全称及公章)\s*(?P<sep>[:：])\s*(?P<placeholder>\s+)\s*$'),
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>\s{10,})\s*(?P<suffix>（[^）]*公章[^）]*）|\([^)]*公章[^)]*\))\s*$'),
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>\s{20,})\s*$'),
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>\s{10,19})\s*$'),
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>_{3,})\s*$'),
            
            # === 通用格式（最低优先级）===
            re.compile(r'^(?P<label>投标人名称(?:（公章）|\(公章\))?)\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]*|＿*|——*)\s*$'),
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
        ]
        
        # 增强的跨run拆分模式 - 特别针对"供应商名称："跨run问题
        self.enhanced_bidder_patterns = [
            # === 直接全匹配模式（处理跨run拆分后的完整文本）===
            # 模式1: "供应商名称：" - 完全精确匹配（处理拆分后合并的文本）
            re.compile(r'^(?P<label>供应商名称)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # 模式2: "投标人名称：" - 完全精确匹配
            re.compile(r'^(?P<label>投标人名称)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # 模式3: "公司名称：" - 完全精确匹配
            re.compile(r'^(?P<label>公司名称)\s*(?P<sep>[:：])\s*(?P<placeholder>)\s*$'),
            
            # === 灵活匹配模式（处理各种变形）===
            # 模式4: 包含"供应商"和"名称"的任何组合
            re.compile(r'^(?P<label>[^:：]*供应商[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$'),
            
            # 模式5: 包含"投标人"和"名称"的任何组合
            re.compile(r'^(?P<label>[^:：]*投标人[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$'),
            
            # 模式6: 包含"公司"和"名称"的任何组合
            re.compile(r'^(?P<label>[^:：]*公司[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$'),
            
            # === 宽松匹配模式（最后的备选方案）===
            # 模式7: 任何带"名称"的标签
            re.compile(r'^(?P<label>[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$'),
        ]
        
        # 合并所有模式
        self.bidder_patterns = self.enhanced_bidder_patterns + self.original_bidder_patterns
    
    def analyze_paragraph_structure(self, paragraph, index=None):
        """分析段落结构，返回详细信息"""
        full_text = ''.join(run.text for run in paragraph.runs)
        
        analysis = {
            'index': index,
            'full_text': full_text,
            'text_length': len(full_text),
            'is_empty': not full_text.strip(),
            'run_count': len(paragraph.runs),
            'run_details': []
        }
        
        for j, run in enumerate(paragraph.runs):
            run_analysis = {
                'run_index': j,
                'text': run.text,
                'length': len(run.text),
                'char_details': []
            }
            
            # 字符编码详情（限制显示前15个字符）
            for k, char in enumerate(run.text):
                run_analysis['char_details'].append(f"'{char}'(U+{ord(char):04X})")
                if k >= 14:  # 限制显示前15个字符
                    run_analysis['char_details'].append("...")
                    break
            
            analysis['run_details'].append(run_analysis)
        
        return analysis
    
    def test_patterns_with_debug(self, text, patterns, pattern_names=None):
        """测试模式匹配并返回详细调试信息"""
        results = {
            'text': text,
            'successful_matches': [],
            'failed_patterns': [],
            'total_patterns': len(patterns)
        }
        
        for i, pattern in enumerate(patterns):
            pattern_name = pattern_names[i] if pattern_names and i < len(pattern_names) else f"模式{i+1}"
            
            try:
                match = pattern.search(text)
                if match:
                    match_info = {
                        'pattern_id': i + 1,
                        'pattern_name': pattern_name,
                        'pattern_regex': pattern.pattern,
                        'matched_text': match.group(0),
                        'groups': match.groupdict(),
                        'span': match.span()
                    }
                    results['successful_matches'].append(match_info)
                    logger.debug(f"✓ {pattern_name} 匹配成功: {match.groupdict()}")
                else:
                    results['failed_patterns'].append({
                        'pattern_id': i + 1,
                        'pattern_name': pattern_name,
                        'pattern_regex': pattern.pattern
                    })
                    logger.debug(f"✗ {pattern_name} 匹配失败")
                    
            except Exception as e:
                logger.error(f"! {pattern_name} 测试异常: {e}")
                results['failed_patterns'].append({
                    'pattern_id': i + 1,
                    'pattern_name': pattern_name,
                    'pattern_regex': pattern.pattern,
                    'error': str(e)
                })
        
        return results
    
    def suggest_pattern_for_text(self, text):
        """为给定文本建议新的匹配模式"""
        suggestions = []
        
        # 分析文本特征
        has_colon = ':' in text or '：' in text
        has_supplier = '供应商' in text
        has_bidder = '投标人' in text  
        has_company = '公司' in text
        has_name = '名称' in text
        
        logger.info(f"文本特征分析: 冒号={has_colon}, 供应商={has_supplier}, 投标人={has_bidder}, 公司={has_company}, 名称={has_name}")
        
        if has_colon and has_name:
            # 基于冒号分割
            colon_pos = text.find(':') if ':' in text else text.find('：')
            if colon_pos > 0:
                label_part = text[:colon_pos].strip()
                placeholder_part = text[colon_pos+1:].strip()
                
                # 转义特殊字符用于正则表达式
                escaped_label = re.escape(label_part)
                
                if placeholder_part == '':
                    pattern = f'^(?P<label>{escaped_label})\\s*(?P<sep>[:：])\\s*(?P<placeholder>)\\s*$'
                elif '_' in placeholder_part:
                    pattern = f'^(?P<label>{escaped_label})\\s*(?P<sep>[:：])\\s*(?P<placeholder>_{{3,}})\\s*$'
                elif placeholder_part.strip() == '' and len(placeholder_part) > 3:
                    space_count = len(placeholder_part)
                    pattern = f'^(?P<label>{escaped_label})\\s*(?P<sep>[:：])\\s*(?P<placeholder>\\s{{{space_count//2},}})\\s*$'
                else:
                    pattern = f'^(?P<label>{escaped_label})\\s*(?P<sep>[:：])\\s*(?P<placeholder>.*?)\\s*$'
                
                suggestions.append({
                    'type': '基于冒号分割',
                    'pattern': pattern,
                    'description': f'为标签"{label_part}"设计的模式'
                })
        
        # 生成通用建议
        if has_supplier and has_name:
            suggestions.append({
                'type': '供应商名称通用',
                'pattern': r'^(?P<label>[^:：]*供应商[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$',
                'description': '匹配包含"供应商"和"名称"的任何格式'
            })
        
        if has_bidder and has_name:
            suggestions.append({
                'type': '投标人名称通用',
                'pattern': r'^(?P<label>[^:：]*投标人[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$',
                'description': '匹配包含"投标人"和"名称"的任何格式'
            })
        
        if has_company and has_name:
            suggestions.append({
                'type': '公司名称通用',
                'pattern': r'^(?P<label>[^:：]*公司[^:：]*名称[^:：]*)\s*(?P<sep>[:：])\s*(?P<placeholder>.*?)\s*$',
                'description': '匹配包含"公司"和"名称"的任何格式'
            })
        
        return suggestions
    
    def enhanced_fill_bidder_name(self, doc: Document, company_name: str, company_info: Dict = {}) -> Dict:
        """增强版填写投标人/供应商名称 - 增加调试日志"""
        logger.info(f"开始增强版填写投标人名称: {company_name}")
        stats = {'paragraphs_changed': 0, 'tables_changed': 0, 'headers_changed': 0, 'footers_changed': 0, 'debug_info': []}
        
        def enhanced_replace_in_runs(paragraph, name, paragraph_index=None):
            """增强版段落文本替换，包含详细调试"""
            logger.debug(f"\n=== 处理段落 {paragraph_index+1 if paragraph_index is not None else '?'} ===")
            
            # 分析段落结构
            analysis = self.analyze_paragraph_structure(paragraph, paragraph_index)
            logger.debug(f"段落分析: {analysis['full_text']!r} (长度: {analysis['text_length']}, runs: {analysis['run_count']})")
            
            # 记录run详情
            for run_detail in analysis['run_details']:
                logger.debug(f"  Run {run_detail['run_index']+1}: {run_detail['text']!r} (长度: {run_detail['length']})")
            
            # 检查是否为空
            if analysis['is_empty']:
                logger.debug("段落为空，跳过处理")
                return False
            
            # 检查是否已包含公司名称
            if name in analysis['full_text']:
                logger.info(f"段落已包含公司名称，跳过: {analysis['full_text']!r}")
                return False
            
            # 测试模式匹配
            pattern_names = [
                "精确供应商名称", "精确投标人名称", "精确公司名称",
                "灵活供应商名称", "灵活投标人名称", "灵活公司名称", 
                "通用名称格式", "原始模式1", "原始模式2", "原始模式3",
                "原始模式4", "原始模式5", "原始模式6", "原始模式7",
                "原始模式8", "原始模式9", "原始模式10", "原始模式11",
                "原始模式12", "原始模式13", "原始模式14", "原始模式15"
            ]
            
            match_results = self.test_patterns_with_debug(
                analysis['full_text'], 
                self.bidder_patterns,
                pattern_names[:len(self.bidder_patterns)]
            )
            
            # 保存调试信息
            debug_entry = {
                'paragraph_index': paragraph_index,
                'text': analysis['full_text'],
                'successful_matches': match_results['successful_matches'],
                'failed_count': len(match_results['failed_patterns'])
            }
            stats['debug_info'].append(debug_entry)
            
            # 处理匹配结果
            if match_results['successful_matches']:
                # 使用第一个成功的匹配
                best_match = match_results['successful_matches'][0]
                logger.info(f"使用最佳匹配: {best_match['pattern_name']}")
                
                groups = best_match['groups']
                
                # 检查是否是括号格式
                if 'prefix' in groups and 'content' in groups:
                    # 括号格式处理
                    prefix = groups.get('prefix', '')
                    content = groups.get('content', '')
                    suffix = groups.get('suffix', '')
                    
                    if '住址' in content:
                        logger.warning(f"检测到住址字段，跳过替换: {content}")
                        return False
                    
                    replacement = f"{prefix}{name}{suffix}"
                    logger.info(f"括号格式替换: {analysis['full_text']!r} -> {replacement!r}")
                    
                    # 简单替换整个段落文本
                    self._simple_replace_paragraph_text(paragraph, replacement)
                    return True
                
                else:
                    # 普通格式处理
                    label = groups.get('label', '')
                    sep = groups.get('sep', ':')
                    suffix = groups.get('suffix', '')
                    
                    # 确保分隔符格式正确
                    if not sep:
                        sep = ':'
                    if sep and not sep.endswith(' '):
                        sep += ' ' if sep in [':', '：'] else ''
                    
                    # 生成新文本
                    if suffix:
                        # 有后缀的格式（如格式9）
                        spaces = "                    "  # 适量空格
                        new_text = f"{label}{sep}{name}{spaces}{suffix}"
                    else:
                        # 简单格式
                        new_text = f"{label}{sep} {name}"
                    
                    logger.info(f"普通格式替换: {analysis['full_text']!r} -> {new_text!r}")
                    
                    # 替换段落文本
                    success = self._simple_replace_paragraph_text(paragraph, new_text)
                    
                    if success:
                        # 给公司名称添加下划线
                        self._add_underline_to_company_name(paragraph, name)
                        return True
                    
                    return False
            else:
                logger.warning(f"没有匹配的模式，尝试建议新模式")
                suggestions = self.suggest_pattern_for_text(analysis['full_text'])
                for suggestion in suggestions:
                    logger.info(f"建议模式: {suggestion['type']} - {suggestion['pattern']}")
                    
                    # 尝试建议的模式
                    try:
                        test_pattern = re.compile(suggestion['pattern'])
                        match = test_pattern.search(analysis['full_text'])
                        if match:
                            logger.info(f"建议模式匹配成功！将应用替换")
                            # 应用建议的模式
                            groups = match.groupdict()
                            label = groups.get('label', '')
                            sep = groups.get('sep', ':')
                            
                            if not sep.endswith(' '):
                                sep += ' ' if sep in [':', '：'] else ''
                            
                            new_text = f"{label}{sep} {name}"
                            success = self._simple_replace_paragraph_text(paragraph, new_text)
                            if success:
                                self._add_underline_to_company_name(paragraph, name)
                                return True
                    except Exception as e:
                        logger.error(f"建议模式测试失败: {e}")
            
            return False
        
        # 处理正文段落
        logger.info(f"处理正文段落，共 {len(doc.paragraphs)} 个")
        for i, para in enumerate(doc.paragraphs):
            if enhanced_replace_in_runs(para, company_name, i):
                stats['paragraphs_changed'] += 1
        
        logger.info(f"增强版填写投标人名称完成: {stats}")
        return stats
    
    def _simple_replace_paragraph_text(self, paragraph, new_text):
        """简化的段落文本替换"""
        try:
            # 清空所有run
            for run in paragraph.runs:
                run.text = ""
            
            # 设置新文本到第一个run
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
            
            return True
        except Exception as e:
            logger.error(f"替换段落文本失败: {e}")
            return False
    
    def _add_underline_to_company_name(self, paragraph, company_name):
        """为段落中的公司名称添加下划线"""
        try:
            # 查找包含公司名称的run
            for run in paragraph.runs:
                if company_name in run.text:
                    # 如果run只包含公司名称，直接添加下划线
                    if run.text.strip() == company_name:
                        run.underline = True
                        logger.debug(f"为公司名称添加下划线: {company_name}")
                        return
                    
                    # 如果run包含更多内容，需要分割
                    text = run.text
                    company_start = text.find(company_name)
                    if company_start >= 0:
                        prefix = text[:company_start]
                        suffix = text[company_start + len(company_name):]
                        
                        # 重构run
                        run.text = prefix
                        
                        # 添加公司名称run（带下划线）
                        company_run = paragraph.add_run(company_name)
                        company_run.underline = True
                        
                        # 添加后缀run
                        if suffix:
                            paragraph.add_run(suffix)
                        
                        logger.debug(f"分割run并为公司名称添加下划线: {company_name}")
                        return
        except Exception as e:
            logger.warning(f"添加下划线失败: {e}")
    
    def test_with_document(self, doc_path, company_name="测试公司"):
        """测试文档处理"""
        logger.info(f"开始测试文档处理: {doc_path}")
        
        if not os.path.exists(doc_path):
            logger.error(f"文档不存在: {doc_path}")
            return None
        
        try:
            doc = Document(doc_path)
            logger.info(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")
            
            # 执行增强填写
            result = self.enhanced_fill_bidder_name(doc, company_name)
            
            # 保存结果文档
            output_path = doc_path.replace('.docx', '_enhanced_output.docx')
            doc.save(output_path)
            logger.info(f"结果已保存到: {output_path}")
            
            return {
                'success': True,
                'input_file': doc_path,
                'output_file': output_path,
                'stats': result
            }
            
        except Exception as e:
            logger.error(f"测试处理失败: {e}", exc_info=True)
            return {
                'success': False,
                'error': str(e)
            }


def main():
    """主函数 - 测试增强版处理器"""
    processor = EnhancedBusinessResponseProcessor()
    
    # 测试文档路径
    doc_path = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/项目/2-数研所二次卡/商务应答文件格式_测试.docx"
    company_name = "智慧足迹数据科技有限公司"
    
    result = processor.test_with_document(doc_path, company_name)
    
    print(f"\n处理结果: {result}")
    print("\n详细日志已保存到 enhanced_business_response.log")


if __name__ == "__main__":
    main()