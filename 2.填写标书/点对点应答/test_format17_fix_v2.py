#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
from mcp_bidder_name_processor_enhanced import MCPBidderNameProcessor
from docx import Document

def test_format17_fix_v2():
    """测试格式17修复版本2 - 格式保持"""
    
    # 使用包含模式的文件
    input_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/uploads/20250903_095446_tender_document.docx"
    output_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/2.填写标书/点对点应答/test_format17_fix_v2_result.docx"
    company_name = "智慧足迹数据科技有限公司"
    
    if not os.path.exists(input_file):
        print(f"输入文件不存在: {input_file}")
        return False
        
    print("=== 处理前格式分析 ===")
    
    # 先分析原始文件的格式
    try:
        doc = Document(input_file)
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if "供应商名称、地址" in paragraph.text:
                print(f"找到目标段落 #{para_idx}: {paragraph.text}")
                print(f"总run数: {len(paragraph.runs)}")
                
                # 显示格式分布
                italic_runs = []
                for run_idx, run in enumerate(paragraph.runs):
                    if run.font.italic:
                        italic_runs.append(f"#{run_idx}('{run.text}')")
                
                print(f"斜体run: {', '.join(italic_runs) if italic_runs else '无'}")
                break
        
    except Exception as e:
        print(f"分析原始文件时出错: {e}")
    
    print("\n=== 开始处理 ===")
    
    try:
        processor = MCPBidderNameProcessor()
        
        result = processor.process_bidder_name(
            input_file=input_file,
            output_file=output_file,
            company_name=company_name
        )
        
        print(f"处理结果: {result}")
        
        if result.get('success'):
            print("✅ 处理成功!")
            
            stats = result.get('stats', {})
            patterns_found = stats.get('patterns_found', [])
            
            for pattern in patterns_found:
                if '供应商名称、地址' in pattern.get('description', ''):
                    print(f"✅ 格式17已处理: {pattern.get('description')}")
            
            # 分析处理后的格式
            print("\n=== 处理后格式分析 ===")
            try:
                doc_after = Document(output_file)
                for para_idx, paragraph in enumerate(doc_after.paragraphs):
                    if "智慧足迹数据科技有限公司" in paragraph.text and "北京市东城区王府井大街200号" in paragraph.text:
                        print(f"找到处理后段落 #{para_idx}")
                        print(f"文本: {paragraph.text[:150]}...")
                        print(f"总run数: {len(paragraph.runs)}")
                        
                        # 检查格式分布
                        format_info = []
                        for run_idx, run in enumerate(paragraph.runs):
                            if run.text:  # 只显示有文本的run
                                italic = "斜体" if run.font.italic else "正常"
                                format_info.append(f"#{run_idx}({italic}): '{run.text[:20]}{'...' if len(run.text) > 20 else ''}'")
                        
                        print("Run格式分布:")
                        for info in format_info:
                            print(f"  {info}")
                        
                        break
                        
            except Exception as e:
                print(f"分析处理后文件时出错: {e}")
            
            return True
        else:
            print("❌ 处理失败")
            return False
            
    except Exception as e:
        print(f"❌ 处理过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_format17_fix_v2()