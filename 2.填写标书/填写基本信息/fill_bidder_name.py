#!/usr/bin/env python3
# coding: utf-8

"""
fill_bidder_name.py
用法:
    python fill_bidder_name.py input.docx output.docx "智慧足迹数据科技有限公司"

说明:
 - 会在段落、表格、页眉、页脚中查找常见的“投标人/供应商名称”标签并填入公司名。
 - 尝试保持原 run 的文本格式（不清空样式）。
"""

import sys
import re
from docx import Document

# 要匹配的标签变体
LABELS_RE = r'((投标人|供应商)(?:名称|全称)?(?:（盖章）|（公章）|\(盖章\)|\(公章\))?)'

# 标签 + 占位符（行尾）
PLACEHOLDER_RE = re.compile(
    rf'(?P<label>{LABELS_RE})\s*(?P<sep>[:：]?)\s*(?P<placeholder>[_\-\u2014\s\u3000]{1,}|＿+|——+|（\s*）)?$'
)

# 标签单独行尾
LABEL_INLINE_RE = re.compile(rf'(?P<label>{LABELS_RE})(?P<sep>[:：]?)\s*$')

def replace_in_runs(paragraph, name):
    """只修改 runs 的 text，不清空格式"""
    changed = False
    for run in paragraph.runs:
        text = run.text
        # 标签 + 占位符
        m = PLACEHOLDER_RE.search(text)
        if m:
            run.text = re.sub(PLACEHOLDER_RE, f"{m.group('label')}{m.group('sep')}{name}", text)
            changed = True
            continue
        # 标签单独行尾
        m2 = LABEL_INLINE_RE.search(text)
        if m2:
            run.text = re.sub(LABEL_INLINE_RE, f"{m2.group('label')}{m2.group('sep')} {name}", text)
            changed = True
    return changed

def process_paragraph(paragraph, name):
    return replace_in_runs(paragraph, name)

def process_table(table, name):
    changed_any = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if replace_in_runs(p, name):
                    changed_any = True
            for tbl in cell.tables:
                if process_table(tbl, name):
                    changed_any = True
    return changed_any

def process_doc(doc: Document, name: str) -> dict:
    stats = {'paragraphs_changed': 0, 'tables_changed': 0, 'headers_changed': 0, 'footers_changed': 0}
    # 正文段落
    for para in doc.paragraphs:
        if process_paragraph(para, name):
            stats['paragraphs_changed'] += 1
    # 表格
    for table in doc.tables:
        if process_table(table, name):
            stats['tables_changed'] += 1
    # 页眉页脚
    for section in doc.sections:
        header = section.header
        footer = section.footer
        header_changed = False
        footer_changed = False
        for para in header.paragraphs:
            if process_paragraph(para, name):
                header_changed = True
        for table in header.tables:
            if process_table(table, name):
                header_changed = True
        for para in footer.paragraphs:
            if process_paragraph(para, name):
                footer_changed = True
        for table in footer.tables:
            if process_table(table, name):
                footer_changed = True
        if header_changed:
            stats['headers_changed'] += 1
        if footer_changed:
            stats['footers_changed'] += 1
    return stats

def main(argv):
    if len(argv) < 4:
        print("Usage: python fill_bidder_name.py input.docx output.docx \"Company Name\"")
        return 1
    input_file = argv[1]
    output_file = argv[2]
    company_name = argv[3]

    print(f"Loading {input_file} ...")
    doc = Document(input_file)
    stats = process_doc(doc, company_name)
    doc.save(output_file)
    print(f"Saved to {output_file}.")
    print("Summary of changes:")
    print(stats)
    return 0

if __name__ == '__main__':
    sys.exit(main(sys.argv))
