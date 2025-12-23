#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模拟搜索"第一章：供应商须知前附表"的过程
"""

import sys
import re
from pathlib import Path

# 添加项目路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / 'ai_tender_system'))

from docx import Document

def extract_core_keywords(text: str) -> str:
    """提取核心关键词：去除编号和常见前缀"""
    # 移除编号
    text = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章|\d+\.|\d+\.\d+|[一二三四五六七八九十]+、)\s*', '', text)
    # 移除"附件"前缀
    text = re.sub(r'^附件[-:：]?', '', text)
    # 移除分隔符
    text = re.sub(r'[-_\t]+', '', text)
    # 移除空格
    text = re.sub(r'\s+', '', text)
    return text

def aggressive_normalize(text: str) -> str:
    """激进规范化"""
    # 移除所有空白字符
    text = re.sub(r'\s+', '', text)
    # 统一冒号
    text = re.sub(r'[:：]', ':', text)
    return text

def main():
    doc_path = "ai_tender_system/data/uploads/tender_processing/2025/12/20251221_180214_中邮保险手机号实名认证服务采购项目竞争性磋商采购文件_fcfdcec9.docx"
    doc = Document(doc_path)

    print("=" * 100)
    print("模拟搜索'第一章：供应商须知前附表'")
    print("=" * 100)

    title = "第一章：供应商须知前附表"
    start_idx = 29  # 从段落29开始（假设之前找到了"竞争性磋商公告"在段落28）

    # 清理标题
    clean_title = re.sub(r'\s+', '', title)
    aggressive_title = aggressive_normalize(title)
    core_keywords = extract_core_keywords(aggressive_title)

    print(f"搜索标题: '{title}'")
    print(f"  clean_title: '{clean_title}'")
    print(f"  aggressive_title: '{aggressive_title}'")
    print(f"  core_keywords: '{core_keywords}'")
    print(f"  从段落 {start_idx} 开始搜索\n")

    # 遍历段落
    print("扫描段落...")
    print("-" * 100)

    matches_found = []

    for i in range(start_idx, min(150, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        para_text = para.text.strip()

        if not para_text:
            continue

        # 清理段落文本
        clean_para = re.sub(r'\s+', '', para_text)
        aggressive_para = aggressive_normalize(para_text)
        para_keywords = extract_core_keywords(aggressive_para)

        # 检查样式
        style_name = para.style.name if para.style else "None"

        # Level 1: 完全匹配或包含匹配
        level1_match = clean_title == clean_para or clean_title in clean_para

        # Level 2: 激进规范化后的完全匹配
        level2_match = aggressive_title == aggressive_para or aggressive_title in aggressive_para

        # Level 3: 去除编号后的匹配
        title_without_number = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章|\d+\.|\d+\.\d+|[一二三四五六七八九十]+、)\s*', '', clean_title)
        para_without_number = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章|\d+\.|\d+\.\d+|[一二三四五六七八九十]+、)\s*', '', clean_para)
        level3_match = title_without_number and para_without_number and title_without_number == para_without_number

        # 如果有任何级别匹配，记录
        if level1_match or level2_match or level3_match:
            matches_found.append({
                'idx': i,
                'text': para_text,
                'style': style_name,
                'level1': level1_match,
                'level2': level2_match,
                'level3': level3_match,
                'clean_para': clean_para,
                'para_keywords': para_keywords
            })

            match_type = []
            if level1_match:
                match_type.append("Level 1")
            if level2_match:
                match_type.append("Level 2")
            if level3_match:
                match_type.append("Level 3")

            print(f"{i:4d}. [{style_name:20s}] {', '.join(match_type):20s} {para_text[:60]}")

    print("\n" + "=" * 100)
    print("分析")
    print("=" * 100)

    if matches_found:
        print(f"\n找到 {len(matches_found)} 个匹配:")
        for m in matches_found:
            print(f"\n段落 {m['idx']} ({m['style']}):")
            print(f"  文本: {m['text']}")
            print(f"  clean_para: {m['clean_para']}")
            print(f"  para_keywords: {m['para_keywords']}")
            print(f"  匹配级别: ", end="")
            levels = []
            if m['level1']:
                levels.append("Level 1 (完全匹配)")
            if m['level2']:
                levels.append("Level 2 (规范化匹配)")
            if m['level3']:
                levels.append("Level 3 (去编号匹配)")
            print(", ".join(levels))

        print(f"\n⭐ 第一个匹配是段落 {matches_found[0]['idx']}，应该返回这个位置")
        if matches_found[0]['idx'] == 63:
            print("✅ 正确！段落63是真正的第一章标题（Heading 1）")
        elif matches_found[0]['idx'] == 92:
            print("❌ 错误！段落92是文件构成说明列表中的，不应该匹配")
    else:
        print("\n❌ 没有找到任何匹配！")

    print()

if __name__ == '__main__':
    main()
