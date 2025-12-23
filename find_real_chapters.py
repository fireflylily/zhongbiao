#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
查找真正的章节标题位置
"""

from docx import Document


def find_real_chapters():
    """查找所有包含'第X部分'的段落"""

    doc_path = "/Users/lvhe/Downloads/zhongbiao/zhongbiao/ai_tender_system/data/uploads/tender_processing/2025/12/20251221_180155_招标文件-哈银消金_b965f323.docx"

    doc = Document(doc_path)

    print("=" * 100)
    print("📄 查找所有'第X部分'的段落")
    print("=" * 100)

    chapter_list = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 查找包含"第X部分"的段落
        if '第' in text and '部分' in text:
            chapter_list.append({
                'para_idx': i,
                'text': text,
                'word_count': len(text.replace(' ', '').replace('\n', ''))
            })

    print(f"\n共找到 {len(chapter_list)} 个包含'第X部分'的段落:\n")

    for item in chapter_list:
        marker = ""
        if item['para_idx'] in [27, 28, 29, 30, 31, 32]:  # 目录页
            marker = "📋 [目录页]"
        elif item['para_idx'] in [119, 120, 121, 122, 123, 124]:  # 文件构成清单
            marker = "📝 [文件构成清单]"
        elif item['para_idx'] in [38, 104]:  # 已知的真实章节
            marker = "✅ [已识别的真实章节]"

        print(f"段落 {item['para_idx']:3d} ({item['word_count']:3d}字) {marker}")
        print(f"  内容: {item['text']}")
        print()

    # 分析可能的真实章节
    print("=" * 100)
    print("推断真实章节位置")
    print("=" * 100)

    real_chapters = []
    for item in chapter_list:
        # 跳过目录页(25-32)和文件构成清单(118-125)
        if 25 <= item['para_idx'] <= 32:
            continue
        if 118 <= item['para_idx'] <= 125:
            continue

        real_chapters.append(item)

    print(f"\n真实章节候选（排除目录页和文件构成清单）: {len(real_chapters)} 个\n")

    for item in real_chapters:
        print(f"段落 {item['para_idx']:3d}: {item['text']}")

    # 计算"第二部分"的真实范围
    if len(real_chapters) >= 2:
        second_part = real_chapters[1]  # 第二部分
        third_part = real_chapters[2] if len(real_chapters) > 2 else None

        print("\n" + "=" * 100)
        print("'第二部分'的真实范围")
        print("=" * 100)

        print(f"\n起始段落: {second_part['para_idx']}")
        if third_part:
            print(f"结束段落: {third_part['para_idx'] - 1}")
            para_count = third_part['para_idx'] - second_part['para_idx']
            print(f"段落数量: {para_count}")

            # 计算字数
            total_chars = 0
            for i in range(second_part['para_idx'] + 1, third_part['para_idx']):
                if i < len(doc.paragraphs):
                    text = doc.paragraphs[i].text
                    total_chars += len(text.replace(' ', '').replace('\n', ''))

            print(f"内容字数: {total_chars:,}字")
            print(f"\n对比:")
            print(f"  当前识别: 段落[104, 120], 字数=409")
            print(f"  应该是:   段落[{second_part['para_idx']}, {third_part['para_idx']-1}], 字数={total_chars:,}")


if __name__ == '__main__':
    find_real_chapters()
