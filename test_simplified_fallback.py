#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试简化后的fallback机制
"""

import sys
import logging
from pathlib import Path
from docx import Document

# 添加项目根目录到系统路径
sys.path.append(str(Path(__file__).parent))

from ai_tender_system.modules.business_response.info_filler import InfoFiller

def setup_test_logging():
    """设置测试日志"""
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    
    # 清除现有处理器
    for handler in root_logger.handlers[:]:
        root_logger.removeHandler(handler)
    
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    console_handler.setFormatter(formatter)
    
    root_logger.addHandler(console_handler)
    
    return root_logger

def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    test_paragraphs = [
        "（项目名称）",  
        "（采购编号）",
        "供应商名称：____",
        "电话：____",
        "地址：____",
        "传真：____",
        "致：采购人",
        "根据贵方为（项目名称）项目采购采购货物及服务的竞争性磋商公告（采购编号），签字代表经正式授权。",
    ]
    
    for text in test_paragraphs:
        doc.add_paragraph(text)
    
    return doc

def test_simplified_fallback():
    """测试简化后的fallback机制"""
    print("=" * 80)
    print("测试简化后的fallback机制")
    print("=" * 80)
    
    # 设置日志
    logger = setup_test_logging()
    
    # 创建InfoFiller实例
    filler = InfoFiller()
    
    # 创建测试文档
    doc = create_test_document()
    
    # 模拟实际的数据结构（来自BusinessResponseProcessor和Web应用）
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'address': '北京市东城区王府井大街200号七层711室',
        'fixedPhone': '010-63271000',  # 注意：Web传入的是fixedPhone，不是phone
        'email': 'lvhe@smartsteps.com',
        'fax': '010-63271000',
    }
    
    project_info = {
        'projectName': '哈银消金2025年-2027年运营商数据采购项目',
        'projectNumber': 'GXTC-C-251590031',
        'date': '2025年9月13日',
        'purchaserName': '哈尔滨哈银消费金融有限责任公司',
        'projectOwner': '哈尔滨哈银消费金融有限责任公司'  # 作为purchaserName的备份
    }
    
    print(f"\n📋 测试数据（简化后的键名）:")
    print(f"公司信息键: {list(company_info.keys())}")
    print(f"项目信息键: {list(project_info.keys())}")
    
    print(f"\n📄 原始文档段落:")
    for i, paragraph in enumerate(doc.paragraphs, 1):
        print(f"  {i}. '{paragraph.text}'")
    
    print(f"\n" + "=" * 80)
    print("执行信息填写处理")
    print("=" * 80)
    
    # 执行信息填写
    try:
        result = filler.fill_info(doc, company_info, project_info)
        
        print(f"\n" + "=" * 80)
        print("处理完成后的文档内容")
        print("=" * 80)
        
        for i, paragraph in enumerate(doc.paragraphs, 1):
            print(f"  {i}. '{paragraph.text}'")
        
        print(f"\n📊 统计结果: {result}")
        
        # 验证关键字段
        print(f"\n🔍 关键字段验证:")
        
        # 检查项目信息（固定键名）
        if "哈银消金" in doc.paragraphs[0].text:
            print("✅ 项目名称处理成功（使用固定键名projectName）")
        else:
            print("❌ 项目名称处理失败")
            
        if "GXTC-C" in doc.paragraphs[1].text:
            print("✅ 项目编号处理成功（使用固定键名projectNumber）")
        else:
            print("❌ 项目编号处理失败")
            
        # 检查电话（修复后的键名匹配）
        phone_filled = False
        for para in doc.paragraphs:
            if "电话" in para.text and "010-63271000" in para.text:
                phone_filled = True
                break
        if phone_filled:
            print("✅ 电话处理成功（fixedPhone → phone映射）")
        else:
            print("❌ 电话处理失败")
            
        # 检查地址（保留必要的fallback）
        address_filled = False
        for para in doc.paragraphs:
            if "地址" in para.text and "王府井" in para.text:
                address_filled = True
                break
        if address_filled:
            print("✅ 地址处理成功（address fallback正常）")
        else:
            print("❌ 地址处理失败")
            
        # 检查采购人（保留必要的fallback）
        purchaser_filled = False
        for para in doc.paragraphs:
            if "哈尔滨哈银" in para.text:
                purchaser_filled = True
                break
        if purchaser_filled:
            print("✅ 采购人处理成功（purchaserName → projectOwner fallback）")
        else:
            print("❌ 采购人处理失败")
        
    except Exception as e:
        print(f"\n❌ 测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_simplified_fallback()