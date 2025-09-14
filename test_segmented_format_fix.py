#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试分段格式修复 - 验证真实场景下的格式隔离

基于用户反馈的实际问题：
- 上传文件："授权"字段没有下划线
- 输出文件："授权"字段被错误地加上了下划线
"""

import os
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent))

from docx import Document
from ai_tender_system.modules.business_response.info_filler import InfoFiller


def create_real_problem_scenario():
    """创建真实问题场景的测试文档"""
    doc = Document()

    # 根据用户实际文档创建相同的结构
    p1 = doc.add_paragraph()

    # Run 0: 前缀空格 + 供应商全称字段（有下划线）
    run1 = p1.add_run("   (供应商全称)       ")
    run1.font.name = "华文细黑"
    run1.font.underline = True  # 下划线

    # Run 1: "授权 "（没有下划线）
    run2 = p1.add_run("授权 ")
    run2.font.name = "华文细黑"
    run2.font.underline = False  # 没有下划线

    # Run 2: 供应商代表姓名字段（有下划线）
    run3 = p1.add_run("(供应商代表姓名)      ")
    run3.font.name = "华文细黑"
    run3.font.underline = True  # 下划线

    # Run 3: "，"（没有下划线）
    run4 = p1.add_run(" ，")
    run4.font.name = "华文细黑"
    run4.font.underline = False

    # Run 4: 职务字段（有下划线）
    run5 = p1.add_run(" (职务、职称)        ")
    run5.font.name = "华文细黑"
    run5.font.underline = True

    # Run 5: 普通文本（没有下划线）
    run6 = p1.add_run("为我方代表，参加贵方组织的")
    run6.font.name = "华文细黑"
    run6.font.underline = False

    # Run 6: 项目字段（有下划线）
    run7 = p1.add_run("  (项目名称、项目编号)                 ")
    run7.font.name = "华文细黑"
    run7.font.underline = True

    # Run 7: 普通文本（没有下划线）
    run8 = p1.add_run("谈判的有关活动，并对此项目进行应答。")
    run8.font.name = "华文细黑"
    run8.font.underline = False

    return doc


def test_segmented_format_fix():
    """测试分段格式修复功能"""
    print("🔧 测试分段格式修复功能")
    print("=" * 60)

    # 创建真实场景测试文档
    doc = create_real_problem_scenario()
    test_file = "test_segmented_fix_input.docx"
    doc.save(test_file)

    print("📋 原始文档分析（模拟用户上传文件）:")
    for i, para in enumerate(doc.paragraphs):
        if para.runs:
            print(f"段落 #{i+1}: '{para.text}'")
            for j, run in enumerate(para.runs):
                underline = bool(run.font.underline)
                font_name = run.font.name or "默认"
                print(f"  Run {j}: '{run.text}' | 字体: {font_name} | 下划线: {underline}")
    print()

    # 测试数据
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'authorizedPersonName': '吕贺',
    }

    project_info = {
        'projectName': '哈银消金2025年-2027年运营商数据采购项目',
        'projectNumber': 'GXTC-C-251590031'
    }

    # 执行处理
    info_filler = InfoFiller()
    doc_loaded = Document(test_file)

    print("🔧 执行分段格式修复处理...")
    stats = info_filler.fill_info(doc_loaded, company_info, project_info)

    output_file = "test_segmented_fix_output.docx"
    doc_loaded.save(output_file)

    print(f"📊 处理统计: {stats}")
    print()

    # 分析处理结果
    print("📋 处理后文档分析:")
    doc_result = Document(output_file)

    authorization_has_underline = False
    company_has_underline = False

    for i, para in enumerate(doc_result.paragraphs):
        if para.runs and '授权' in para.text:
            print(f"段落 #{i+1}: '{para.text}'")
            for j, run in enumerate(para.runs):
                if run.text.strip():
                    underline = bool(run.font.underline)
                    font_name = run.font.name or "默认"
                    print(f"  Run {j}: '{run.text}' | 字体: {font_name} | 下划线: {underline}")

                    # 检查关键内容
                    if '授权' in run.text and underline:
                        authorization_has_underline = True
                    if '智慧足迹数据科技有限公司' in run.text and underline:
                        company_has_underline = True
    print()

    # 测试结果评估
    print("📋 格式修复测试结果:")
    print("-" * 40)

    tests = [
        ("公司名称格式隔离", not company_has_underline, "公司名称不应继承下划线"),
        ("授权字段格式隔离", not authorization_has_underline, "授权文字不应继承下划线"),
    ]

    passed_tests = 0
    for test_name, passed, description in tests:
        status = "✅ PASS" if passed else "❌ FAIL"
        print(f"{status} {test_name}: {description}")
        if passed:
            passed_tests += 1

    print()
    print(f"📊 测试总结: {passed_tests}/{len(tests)} 通过")

    if passed_tests == len(tests):
        print("🎉 分段格式修复成功！")
        print("✅ '授权' 字段没有错误继承下划线")
        print("✅ 公司名称等业务内容格式隔离正常")
        return True
    else:
        print("❌ 格式修复仍有问题，需要进一步调试")
        return False


if __name__ == '__main__':
    print("🚀 启动分段格式修复验证测试")
    print("=" * 60)

    success = test_segmented_format_fix()

    if success:
        print("\n🎉 分段格式修复验证成功!")
        print("✅ 已解决用户反馈的格式继承问题")
        sys.exit(0)
    else:
        print("\n⚠️  分段格式修复需要进一步完善")
        sys.exit(1)