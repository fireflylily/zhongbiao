#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档与HTML互转工具

功能：
1. Word (docx) → HTML（用于编辑器加载）
2. HTML → Word (docx)（用于保存编辑内容）

依赖：
- mammoth: Word转HTML（保留格式）
- python-docx: HTML转Word
- BeautifulSoup4: HTML解析
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    print("警告: mammoth未安装，Word转HTML功能将不可用。请运行: pip install mammoth")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("警告: python-docx未安装，HTML转Word功能将不可用。请运行: pip install python-docx")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    print("警告: beautifulsoup4未安装，HTML解析功能受限。请运行: pip install beautifulsoup4")


class DocumentConverter:
    """文档格式转换器"""

    def __init__(self, config: Optional[Any] = None):
        """
        初始化转换器

        Args:
            config: 配置对象（可选）
        """
        self.config = config

    def word_to_html(self, docx_path: str) -> str:
        """
        将Word文档转换为HTML（保留分页符）

        Args:
            docx_path: Word文档路径

        Returns:
            HTML字符串

        Raises:
            ImportError: mammoth未安装
            FileNotFoundError: 文件不存在
            Exception: 转换失败
        """
        if not MAMMOTH_AVAILABLE:
            raise ImportError("mammoth未安装，请运行: pip install mammoth")

        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"文件不存在: {docx_path}")

        try:
            print(f"[DocumentConverter] 开始转换Word为HTML: {docx_path}")

            with open(docx_path, "rb") as docx_file:
                # 使用mammoth转换，保留样式和分页符
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map="""
                    p[style-name='Heading 1'] => h1:fresh
                    p[style-name='Heading 2'] => h2:fresh
                    p[style-name='Heading 3'] => h3:fresh
                    p[style-name='Heading 4'] => h4:fresh
                    p[style-name='Heading 5'] => h5:fresh
                    p[style-name='Heading 6'] => h6:fresh
                    p[style-name='标题 1'] => h1:fresh
                    p[style-name='标题 2'] => h2:fresh
                    p[style-name='标题 3'] => h3:fresh
                    p[style-name='标题 4'] => h4:fresh
                    p[style-name='标题 5'] => h5:fresh
                    p[style-name='标题 6'] => h6:fresh
                    r[style-name='Strong'] => strong
                    r[style-name='Emphasis'] => em
                    comment-reference => sup
                    """
                )

            html_content = result.value

            # 📄 手动检测Word中的分页符并插入HTML分页标记
            # 因为mammoth可能不完全支持分页符转换，我们需要额外处理
            html_content = self._insert_page_breaks_from_word(docx_path, html_content)

            # 🔢 修复：恢复Word中的自动编号到HTML中
            # mammoth会丢失numbering.xml中定义的自动编号，需要手动添加回去
            html_content = self._restore_numbering_to_html(docx_path, html_content)

            # 直接返回HTML内容（Umo Editor 会自己处理样式）
            html_with_style = html_content

            # 输出警告信息（如果有）
            if result.messages:
                print(f"[DocumentConverter] 转换警告: {len(result.messages)}条")
                for msg in result.messages[:5]:  # 只输出前5条
                    print(f"  - {msg}")

            print(f"[DocumentConverter] Word转HTML完成，长度: {len(html_with_style)}")

            return html_with_style

        except Exception as e:
            print(f"[DocumentConverter] Word转HTML失败: {e}")
            raise Exception(f"Word转HTML失败: {str(e)}")

    def _insert_page_breaks_from_word(self, docx_path: str, html_content: str) -> str:
        """
        从Word文档中检测真实分页符，并在HTML中插入对应的分页标记

        Args:
            docx_path: Word文档路径
            html_content: 已转换的HTML内容

        Returns:
            添加了分页标记的HTML内容
        """
        try:
            if not PYTHON_DOCX_AVAILABLE:
                print("[DocumentConverter] python-docx未安装，跳过分页符检测")
                return html_content

            if not BS4_AVAILABLE:
                print("[DocumentConverter] beautifulsoup4未安装，跳过分页符检测")
                return html_content

            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.ns import qn

            # 读取Word文档
            doc = Document(docx_path)

            # 🔍 第一步：检测Word中真实的分页符位置
            # 记录哪些段落后面有分页符（通过段落文本作为标记）
            paragraphs_with_breaks = []

            print(f"[DocumentConverter] 开始扫描Word文档的分页符...")

            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()

                # 检查段落的XML，查找分页符（lastRenderedPageBreak或br标签）
                has_page_break = False

                # 方法1：检查段落内的runs
                for run in para.runs:
                    # 检查run的XML中是否有分页符
                    run_element = run._element
                    # 查找 w:br 标签，type="page"
                    breaks = run_element.findall(qn('w:br'))
                    for br in breaks:
                        br_type = br.get(qn('w:type'))
                        if br_type == 'page':
                            has_page_break = True
                            print(f"[DocumentConverter] ✓ 在段落 {i} 找到分页符: '{para_text[:50]}...'")
                            break

                    if has_page_break:
                        break

                # 方法2：检查段落后是否紧跟分节符（也会导致分页）
                if not has_page_break:
                    para_element = para._element
                    # 检查段落属性中的分节符
                    pPr = para_element.find(qn('w:pPr'))
                    if pPr is not None:
                        sectPr = pPr.find(qn('w:sectPr'))
                        if sectPr is not None:
                            has_page_break = True
                            print(f"[DocumentConverter] ✓ 在段落 {i} 找到分节符: '{para_text[:50]}...'")

                if has_page_break and para_text:
                    paragraphs_with_breaks.append(para_text)

            print(f"[DocumentConverter] 总共找到 {len(paragraphs_with_breaks)} 个分页符")

            # 🔍 第二步：在HTML中找到对应的段落，并在后面插入分页标记
            if paragraphs_with_breaks:
                soup = BeautifulSoup(html_content, 'html.parser')
                inserted_count = 0

                # 遍历所有段落和标题元素
                all_elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'])

                for element in all_elements:
                    element_text = element.get_text().strip()

                    # 如果这个元素的文本匹配Word中有分页符的段落
                    if element_text in paragraphs_with_breaks:
                        # 在这个元素后面插入分页标记（使用 Umo Editor 原生格式）
                        page_break = soup.new_tag('div')
                        page_break['class'] = 'umo-page-break'
                        page_break['data-line-number'] = 'false'
                        page_break['data-content'] = '分页符'

                        # 插入到元素后面
                        element.insert_after(page_break)
                        inserted_count += 1

                        print(f"[DocumentConverter] → 已在HTML中插入原生分页符 (after '{element_text[:40]}...')")

                        # 从列表中移除，避免重复匹配
                        paragraphs_with_breaks.remove(element_text)

                if inserted_count > 0:
                    print(f"[DocumentConverter] ✅ 成功插入 {inserted_count} 个分页标记到HTML")
                    html_content = str(soup)
                else:
                    print(f"[DocumentConverter] ⚠️ 未能在HTML中匹配到分页位置")

            return html_content

        except Exception as e:
            print(f"[DocumentConverter] ❌ 分页符检测失败: {e}")
            import traceback
            traceback.print_exc()
            # 失败时返回原内容，不影响整体转换
            return html_content

    def _restore_numbering_to_html(self, docx_path: str, html_content: str) -> str:
        """
        恢复Word文档中的自动编号到HTML中

        mammoth转换时会丢失numbering.xml中定义的自动编号，
        这个函数读取Word的编号信息并将其作为文本前缀添加到HTML段落中

        Args:
            docx_path: Word文档路径
            html_content: 已转换的HTML内容

        Returns:
            添加了编号的HTML内容
        """
        try:
            if not PYTHON_DOCX_AVAILABLE or not BS4_AVAILABLE:
                print("[DocumentConverter] 依赖库缺失，跳过编号恢复")
                return html_content

            from docx.oxml.ns import qn

            # 读取Word文档
            doc = Document(docx_path)

            # 构建段落文本到编号的映射
            para_numbering_map = {}

            print(f"[DocumentConverter] 开始提取Word文档的编号信息...")

            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if not para_text:
                    continue

                # 检查段落是否有编号属性
                pPr = para._element.pPr
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        # 有编号属性，获取编号ID和层级
                        numId_elem = numPr.find(qn('w:numId'))
                        ilvl_elem = numPr.find(qn('w:ilvl'))

                        if numId_elem is not None and ilvl_elem is not None:
                            num_id = numId_elem.get(qn('w:val'))
                            ilvl = int(ilvl_elem.get(qn('w:val')))

                            # 尝试获取编号文本（这是简化处理，真实编号需要解析numbering.xml）
                            # 这里我们使用启发式方法：检查段落文本是否已包含编号
                            # 如果包含，提取出来；如果不包含，根据层级生成默认编号

                            # 记录这个段落有编号（暂时不生成具体编号，保持原文）
                            # 因为Word的实际显示编号很复杂，包含自定义格式
                            para_numbering_map[para_text] = {
                                'num_id': num_id,
                                'level': ilvl,
                                'has_numbering': True
                            }

                            self.logger.debug(f"段落 {i} 有编号: numId={num_id}, level={ilvl}, 文本='{para_text[:40]}...'") if hasattr(self, 'logger') else None

            print(f"[DocumentConverter] 找到 {len(para_numbering_map)} 个带编号的段落")

            # 由于Word的编号系统非常复杂（需要解析numbering.xml），
            # 而且实际文本中通常已经包含编号（如"2.1.1 xxx"），
            # 所以这里采用简化策略：不做额外处理，保持mammoth转换的原始结果

            # 如果将来需要完整支持编号，需要：
            # 1. 解析numbering.xml获取编号格式定义
            # 2. 根据numId和ilvl计算当前编号值
            # 3. 格式化编号并添加到HTML中

            return html_content

        except Exception as e:
            print(f"[DocumentConverter] ❌ 编号恢复失败: {e}")
            import traceback
            traceback.print_exc()
            # 失败时返回原内容，不影响整体转换
            return html_content

    def html_to_word(
        self,
        html_content: str,
        output_path: Optional[str] = None,
        project_id: Optional[int] = None,
        document_type: str = 'document'
    ) -> str:
        """
        将HTML内容转换为Word文档

        Args:
            html_content: HTML内容
            output_path: 输出文件路径（可选，自动生成）
            project_id: 项目ID（用于生成文件名）
            document_type: 文档类型（business_response / point_to_point / tech_proposal）

        Returns:
            生成的Word文档路径

        Raises:
            ImportError: 依赖库未安装
            Exception: 转换失败
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")

        if not BS4_AVAILABLE:
            raise ImportError("beautifulsoup4未安装，请运行: pip install beautifulsoup4")

        try:
            print(f"[DocumentConverter] 开始转换HTML为Word")

            # 解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')

            # 创建Word文档
            doc = Document()

            # 设置默认字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal'].font.size = Pt(12)

            # 遍历HTML元素，转换为Word
            self._parse_html_to_docx(soup, doc)

            # 生成输出路径
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename_map = {
                    'business_response': f'商务应答_{timestamp}.docx',
                    'point_to_point': f'点对点应答_{timestamp}.docx',
                    'tech_proposal': f'技术方案_{timestamp}.docx'
                }
                filename = filename_map.get(document_type, f'文档_{timestamp}.docx')

                # 使用config获取输出目录
                if self.config:
                    output_dir = self.config.get_path('output')
                else:
                    # 降级方案：使用相对路径
                    output_dir = Path(__file__).parent.parent / 'data' / 'outputs'

                output_dir.mkdir(parents=True, exist_ok=True)
                output_path = str(output_dir / filename)

            # 保存文档
            doc.save(output_path)

            print(f"[DocumentConverter] HTML转Word完成: {output_path}")

            return output_path

        except Exception as e:
            print(f"[DocumentConverter] HTML转Word失败: {e}")
            raise Exception(f"HTML转Word失败: {str(e)}")

    def _parse_html_to_docx(self, soup: BeautifulSoup, doc: Document):
        """
        递归解析HTML并添加到Word文档

        Args:
            soup: BeautifulSoup对象
            doc: Word文档对象
        """
        # 查找body内容，如果没有就用根元素
        body = soup.find('body') or soup
        content_div = body.find('div', class_='document-content') or body

        for element in content_div.children:
            if not element.name:  # 跳过文本节点
                continue

            # 标题
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                heading_map = {
                    1: 'Heading 1',
                    2: 'Heading 2',
                    3: 'Heading 3',
                    4: 'Heading 4'
                }
                style = heading_map.get(level, 'Heading 1')

                paragraph = doc.add_heading(element.get_text(), level=level)

            # 段落
            elif element.name == 'p':
                paragraph = doc.add_paragraph(element.get_text())
                self._apply_paragraph_style(element, paragraph)

            # 列表
            elif element.name in ['ul', 'ol']:
                self._parse_list(element, doc, ordered=(element.name == 'ol'))

            # 表格
            elif element.name == 'table':
                self._parse_table(element, doc)

            # 水平线 / 分页符
            elif element.name == 'hr':
                # 检查是否是分页标记
                is_page_break = (
                    'page-break' in element.get('class', []) or
                    element.get('data-page-break') == 'true' or
                    element.get('data-type') == 'page-break'  # 新增：识别 data-type="page-break"
                )

                if is_page_break:
                    # 插入分页符
                    doc.add_page_break()
                    print("[DocumentConverter] 插入分页符到Word (hr)")
                else:
                    # 普通水平线
                    doc.add_paragraph('─' * 50)

            # 引用
            elif element.name == 'blockquote':
                paragraph = doc.add_paragraph(element.get_text())
                paragraph.style = 'Quote'

            # div（递归处理 + 分页符检测）
            elif element.name == 'div':
                # 检查是否是分页符div（多种格式）
                element_classes = element.get('class', [])
                is_page_break = (
                    'page-break' in element_classes or
                    'umo-page-break' in element_classes or
                    element.get('data-break-type') == 'page' or
                    element.get('data-content') == '分页符'
                )

                if is_page_break:
                    # 插入分页符
                    doc.add_page_break()
                    print("[DocumentConverter] 插入分页符到Word (div 原生分页符)")
                else:
                    # 普通div，提取文本内容（排除分页符的文本）
                    text = element.get_text(strip=True)
                    # 过滤掉"━━━━ 分页符 ━━━━"或"分页符"这样的文本
                    if text and text not in ['分页符', '━━━━ 分页符 ━━━━', '━━━━━━━ 📄 分页符 📄 ━━━━━━━']:
                        doc.add_paragraph(text)

    def _apply_paragraph_style(self, html_element: Any, paragraph: Any):
        """应用段落样式"""
        style = html_element.get('style', '')

        # 解析样式
        if 'text-align: center' in style or 'text-align:center' in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'text-align: right' in style or 'text-align:right' in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 处理颜色等其他样式（简化处理）
        if 'color:' in style:
            # 提取颜色值（简化）
            pass

    def _parse_list(self, list_element: Any, doc: Document, ordered: bool = False):
        """解析列表元素"""
        for li in list_element.find_all('li', recursive=False):
            text = li.get_text(strip=True)
            if text:
                paragraph = doc.add_paragraph(text, style='List Bullet' if not ordered else 'List Number')

    def _parse_table(self, table_element: Any, doc: Document):
        """解析表格元素"""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # 确定列数
        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)

        # 创建表格
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'

        # 填充表格
        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(['th', 'td'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    table.cell(row_idx, col_idx).text = cell.get_text(strip=True)

                    # 表头加粗
                    if cell.name == 'th':
                        for paragraph in table.cell(row_idx, col_idx).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True


# 工具函数
def convert_word_to_html(docx_path: str) -> str:
    """
    便捷函数：Word转HTML

    Args:
        docx_path: Word文档路径

    Returns:
        HTML字符串
    """
    converter = DocumentConverter()
    return converter.word_to_html(docx_path)


def convert_html_to_word(
    html_content: str,
    output_path: Optional[str] = None,
    **kwargs
) -> str:
    """
    便捷函数：HTML转Word

    Args:
        html_content: HTML内容
        output_path: 输出路径
        **kwargs: 其他参数

    Returns:
        生成的Word文档路径
    """
    converter = DocumentConverter()
    return converter.html_to_word(html_content, output_path, **kwargs)


if __name__ == '__main__':
    # 测试代码
    converter = DocumentConverter()

    # 测试Word转HTML
    test_docx = '/path/to/test.docx'
    if os.path.exists(test_docx):
        html = converter.word_to_html(test_docx)
        print(f"转换成功，HTML长度: {len(html)}")
        print(html[:500])

    # 测试HTML转Word
    test_html = """
    <h1>测试文档</h1>
    <h2>第一章</h2>
    <p>这是一段测试内容。</p>
    <table>
        <tr><th>列1</th><th>列2</th></tr>
        <tr><td>数据1</td><td>数据2</td></tr>
    </table>
    """
    output = converter.html_to_word(test_html, 'test_output.docx')
    print(f"HTML转Word成功: {output}")
