# -*- coding: utf-8 -*-
"""
统一文档处理器
整合所有文档读取逻辑，支持多种格式
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, Optional, List, Union
from abc import ABC, abstractmethod

from .config import get_config
from .exceptions import DocumentError, DocumentFormatError, DocumentParsingError


class BaseDocumentProcessor(ABC):
    """文档处理器基类"""
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.config = get_config()
    
    @abstractmethod
    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理该文件"""
        pass
    
    @abstractmethod
    def extract_content(self, file_path: str) -> str:
        """提取文档内容"""
        pass
    
    def validate_file(self, file_path: str) -> bool:
        """验证文件是否有效"""
        if not os.path.exists(file_path):
            raise DocumentError(f"文件不存在: {file_path}")
        
        if not os.path.isfile(file_path):
            raise DocumentError(f"不是文件: {file_path}")
        
        # 检查文件大小
        max_size = self._parse_size(self.config.app.max_file_size)
        file_size = os.path.getsize(file_path)
        if file_size > max_size:
            raise DocumentError(f"文件过大: {file_size} bytes, 最大允许: {max_size} bytes")
        
        return True
    
    def _parse_size(self, size_str: str) -> int:
        """解析文件大小字符串，如 '50MB' -> 字节数"""
        size_str = size_str.upper().strip()
        
        units = {
            'B': 1,
            'KB': 1024,
            'MB': 1024 * 1024,
            'GB': 1024 * 1024 * 1024
        }
        
        for unit, multiplier in units.items():
            if size_str.endswith(unit):
                number = size_str[:-len(unit)].strip()
                try:
                    return int(float(number) * multiplier)
                except ValueError:
                    break
        
        # 默认按字节处理
        try:
            return int(size_str)
        except ValueError:
            return 50 * 1024 * 1024  # 默认50MB


class TextDocumentProcessor(BaseDocumentProcessor):
    """文本文档处理器"""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.txt')
    
    def extract_content(self, file_path: str) -> str:
        """读取文本文件，尝试多种编码"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                self.logger.info(f"使用 {encoding} 编码成功读取文本文件")
                return content
            except UnicodeDecodeError:
                continue
        
        raise DocumentParsingError(f"无法使用任何编码读取文件: {file_path}")


class WordDocumentProcessor(BaseDocumentProcessor):
    """Word文档处理器"""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.doc', '.docx'))
    
    def extract_content(self, file_path: str) -> str:
        """读取Word文档，尝试多种方法"""
        content = ""
        
        # 方法1: python-docx（优先使用，功能最完善）
        content = self._try_python_docx(file_path)
        if content:
            return content
        
        # 方法2: mammoth（备用方案1）
        content = self._try_mammoth(file_path)
        if content:
            return content
        
        # 方法3: win32com（Windows环境）
        content = self._try_win32com(file_path)
        if content:
            return content
        
        # 方法4: antiword（老格式.doc文件）
        if file_path.lower().endswith('.doc'):
            content = self._try_antiword(file_path)
            if content:
                return content
        
        # 方法5: ZIP解析（docx文件）
        if file_path.lower().endswith('.docx'):
            content = self._try_zip_parsing(file_path)
            if content:
                return content
        
        # 方法6: 二进制读取（最后手段）
        content = self._try_binary_reading(file_path)
        if content:
            return content
        
        raise DocumentParsingError(f"无法读取Word文档: {file_path}")
    
    def _try_python_docx(self, file_path: str) -> Optional[str]:
        """使用python-docx读取"""
        try:
            import docx
            doc = docx.Document(file_path)
            
            # 提取段落内容
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # 提取表格内容
            table_content = []
            for table_idx, table in enumerate(doc.tables):
                table_content.append(f"\n=== 表格 {table_idx + 1} ===")
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        table_content.append(f"第{row_idx + 1}行: " + " | ".join(row_cells))
            
            # 合并内容
            all_content = paragraphs + table_content
            content = '\n'.join(all_content)
            
            self.logger.info(f"python-docx读取成功（{len(doc.paragraphs)}段落，{len(doc.tables)}表格）")
            return content
            
        except ImportError:
            self.logger.debug("python-docx未安装")
            return None
        except Exception as e:
            self.logger.debug(f"python-docx读取失败: {e}")
            return None
    
    def _try_mammoth(self, file_path: str) -> Optional[str]:
        """使用mammoth读取"""
        try:
            import mammoth
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                content = result.value
            self.logger.info("mammoth读取成功")
            return content
        except ImportError:
            self.logger.debug("mammoth未安装")
            return None
        except Exception as e:
            self.logger.debug(f"mammoth读取失败: {e}")
            return None
    
    def _try_win32com(self, file_path: str) -> Optional[str]:
        """使用win32com读取（Windows环境）"""
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
            self.logger.info("win32com读取成功")
            return content
        except ImportError:
            self.logger.debug("pywin32未安装或非Windows环境")
            return None
        except Exception as e:
            self.logger.debug(f"win32com读取失败: {e}")
            return None
    
    def _try_antiword(self, file_path: str) -> Optional[str]:
        """使用antiword读取老格式.doc文件"""
        try:
            import antiword
            content = antiword.extract_text(file_path)
            self.logger.info("antiword读取成功")
            return content
        except ImportError:
            self.logger.debug("antiword未安装")
            return None
        except Exception as e:
            self.logger.debug(f"antiword读取失败: {e}")
            return None
    
    def _try_zip_parsing(self, file_path: str) -> Optional[str]:
        """ZIP解析docx文件"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                xml_content = zip_file.read('word/document.xml')
                root = ET.fromstring(xml_content)
                
                # 提取文本内容
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                text_elements = root.findall('.//w:t', namespaces)
                content = '\n'.join([elem.text or '' for elem in text_elements])
                
            self.logger.info("ZIP解析读取成功")
            return content
        except Exception as e:
            self.logger.debug(f"ZIP解析失败: {e}")
            return None
    
    def _try_binary_reading(self, file_path: str) -> Optional[str]:
        """二进制读取并清理"""
        try:
            with open(file_path, 'rb') as f:
                raw_content = f.read()
            
            # 尝试多种解码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'cp1252', 'latin1']:
                try:
                    content = raw_content.decode(encoding, errors='ignore')
                    # 清理二进制垃圾字符
                    cleaned_content = re.sub(
                        r'[^\u4e00-\u9fff\u3400-\u4dbfa-zA-Z0-9\s\-_.,;:!?()（）【】《》""''@#￥%…&*+=<>/\\|`~\n\r，。；？！、]', 
                        '', content
                    )
                    
                    if len(cleaned_content) > 100:
                        self.logger.warning(f"二进制读取成功 ({encoding} 编码)")
                        return cleaned_content
                        
                except UnicodeDecodeError:
                    continue
            
            return None
            
        except Exception as e:
            self.logger.debug(f"二进制读取失败: {e}")
            return None


class PDFDocumentProcessor(BaseDocumentProcessor):
    """PDF文档处理器"""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def extract_content(self, file_path: str) -> str:
        """读取PDF文档"""
        # 方法1: PyPDF2
        content = self._try_pypdf2(file_path)
        if content:
            return content
        
        # 方法2: pdfplumber
        content = self._try_pdfplumber(file_path)
        if content:
            return content
        
        # 方法3: fitz (PyMuPDF)
        content = self._try_fitz(file_path)
        if content:
            return content
        
        raise DocumentParsingError(f"无法读取PDF文档: {file_path}")
    
    def _try_pypdf2(self, file_path: str) -> Optional[str]:
        """使用PyPDF2读取"""
        try:
            import PyPDF2
            content_list = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content_list.append(page.extract_text())
            
            content = '\n'.join(content_list)
            self.logger.info("PyPDF2读取成功")
            return content
        except ImportError:
            self.logger.debug("PyPDF2未安装")
            return None
        except Exception as e:
            self.logger.debug(f"PyPDF2读取失败: {e}")
            return None
    
    def _try_pdfplumber(self, file_path: str) -> Optional[str]:
        """使用pdfplumber读取"""
        try:
            import pdfplumber
            content_list = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_list.append(text)
            
            content = '\n'.join(content_list)
            self.logger.info("pdfplumber读取成功")
            return content
        except ImportError:
            self.logger.debug("pdfplumber未安装")
            return None
        except Exception as e:
            self.logger.debug(f"pdfplumber读取失败: {e}")
            return None
    
    def _try_fitz(self, file_path: str) -> Optional[str]:
        """使用fitz (PyMuPDF)读取"""
        try:
            import fitz
            content_list = []
            
            pdf_document = fitz.open(file_path)
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                content_list.append(page.get_text())
            
            pdf_document.close()
            content = '\n'.join(content_list)
            self.logger.info("fitz读取成功")
            return content
        except ImportError:
            self.logger.debug("fitz (PyMuPDF)未安装")
            return None
        except Exception as e:
            self.logger.debug(f"fitz读取失败: {e}")
            return None


class DocumentProcessor:
    """统一文档处理器"""
    
    def __init__(self):
        self.processors = [
            TextDocumentProcessor(),
            WordDocumentProcessor(),
            PDFDocumentProcessor()
        ]
        self.logger = logging.getLogger(__name__)
        self.config = get_config()
    
    def process_document(self, file_path: str) -> str:
        """处理文档，返回清理后的内容"""
        file_path = str(Path(file_path).resolve())
        
        # 找到合适的处理器
        processor = self._find_processor(file_path)
        if not processor:
            supported_formats = ['.txt', '.doc', '.docx', '.pdf']
            raise DocumentFormatError(f"不支持的文件格式。支持的格式: {supported_formats}")
        
        # 验证文件
        processor.validate_file(file_path)
        
        # 提取内容
        self.logger.info(f"开始处理文档: {file_path}")
        content = processor.extract_content(file_path)
        
        # 清理和截断内容
        cleaned_content = self._clean_content(content)
        
        self.logger.info(f"文档处理完成，内容长度: {len(cleaned_content)}")
        return cleaned_content
    
    def _find_processor(self, file_path: str) -> Optional[BaseDocumentProcessor]:
        """找到合适的文档处理器"""
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None
    
    def _clean_content(self, content: str) -> str:
        """清理文档内容"""
        if not content:
            raise DocumentParsingError("文档内容为空")
        
        # 移除特殊字符和控制字符
        # 保留中文、英文、数字、常用标点符号
        cleaned = re.sub(
            r'[^\u4e00-\u9fff\u3400-\u4dbf\w\s\-_.,;:!?()（）【】《》""''@#￥%…&*+=<>/\\|`~\n\r]', 
            '', content
        )
        
        # 移除多余的空白字符
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        
        # 如果内容太长，智能截断
        max_length = self.config.get('modules.tender_info.max_content_length', 6000)
        if len(cleaned) > max_length:
            cleaned = self._smart_truncate(cleaned, max_length)
        
        return cleaned
    
    def _smart_truncate(self, content: str, max_length: int) -> str:
        """智能截断内容"""
        if len(content) <= max_length:
            return content
        
        # 尝试找到合适的截断点（优先在段落或句子结束）
        truncated = content[:max_length]
        
        # 寻找最后一个句号、问号或感叹号
        for delimiter in ['。', '！', '？', '.', '!', '?', '\n']:
            last_pos = truncated.rfind(delimiter)
            if last_pos > max_length * 0.8:  # 至少保留80%的内容
                truncated = truncated[:last_pos + 1]
                break
        
        self.logger.warning(f"内容过长，已截断至 {len(truncated)} 字符")
        return truncated
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        formats = []
        for processor in self.processors:
            if hasattr(processor, 'supported_formats'):
                formats.extend(processor.supported_formats)
        
        if not formats:
            formats = ['.txt', '.doc', '.docx', '.pdf']
        
        return formats


# 全局文档处理器实例
_document_processor = None


def get_document_processor() -> DocumentProcessor:
    """获取文档处理器实例"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor


def process_document(file_path: str) -> str:
    """便捷函数：处理文档"""
    processor = get_document_processor()
    return processor.process_document(file_path)


if __name__ == "__main__":
    # 测试文档处理器
    import sys
    
    if len(sys.argv) != 2:
        print("使用方法: python document_processor.py <文档路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        content = process_document(file_path)
        print(f"文档内容长度: {len(content)}")
        print(f"前500字符: {content[:500]}...")
        print("✅ 文档处理测试成功")
    except Exception as e:
        print(f"❌ 文档处理测试失败: {e}")
        sys.exit(1)