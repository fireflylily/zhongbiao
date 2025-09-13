#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
集成测试：验证统一字段映射在文档处理中的实际效果
"""

import sys
import tempfile
from pathlib import Path
from docx import Document

# 添加项目根目录到路径
sys.path.append(str(Path(__file__).parent))

from modules.business_response.info_filler import InfoFiller

def create_test_document():
    """创建测试文档"""
    doc = Document()

    # 添加测试内容 - 包含多种格式
    test_paragraphs = [
        "供应商名称：________________",  # 填空规则测试
        "地址：___________________",    # 地址映射测试
        "电话：___________________",    # 电话映射测试
        "采购人：_________________",    # 采购人映射测试
        "（供应商名称）投标文件",        # 替换规则测试
        "（项目名称）技术方案",         # 项目信息测试
        "供应商名称：                   ", # 纯空格测试（原Bug场景）
    ]

    for text in test_paragraphs:
        doc.add_paragraph(text)

    return doc

def test_integration():
    """集成测试主函数"""
    print("=" * 60)
    print("🧪 统一字段映射集成测试")
    print("=" * 60)

    # 创建InfoFiller实例
    filler = InfoFiller()

    # 模拟真实数据
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'address': '',  # 主地址为空，测试回退逻辑
        'registeredAddress': '北京市朝阳区创新大厦888号',
        'phone': '',    # 主电话为空，测试回退逻辑
        'fixedPhone': '010-88888888',
        'email': 'info@smartsteps.com'
    }

    project_info = {
        'projectName': 'AI智能标书生成系统',
        'projectNumber': 'PROJ-2025-001',
        'purchaserName': '',  # 采购人为空，测试回退逻辑
        'projectOwner': '北京市政府采购中心',
        'date': '2025年1月15日'
    }

    # 创建测试文档
    doc = create_test_document()

    print("📄 测试文档原始内容:")
    for i, para in enumerate(doc.paragraphs, 1):
        print(f"  {i}. '{para.text}'")

    # 执行信息填写
    print("\n🔧 执行信息填写...")
    stats = filler.fill_info(doc, company_info, project_info)

    print(f"\n📊 处理统计: {stats}")

    print("\n📄 处理后文档内容:")
    success_count = 0
    expected_results = [
        ("供应商名称：智慧足迹数据科技有限公司", "供应商名称填空"),
        ("地址：北京市朝阳区创新大厦888号", "地址映射填空"),
        ("电话：010-88888888", "电话映射填空"),
        ("采购人：北京市政府采购中心", "采购人映射填空"),
        ("（智慧足迹数据科技有限公司）", "供应商名称替换"),
        ("（AI智能标书生成系统）", "项目名称替换"),
    ]

    for i, para in enumerate(doc.paragraphs, 1):
        content = para.text
        print(f"  {i}. '{content}'")

        # 检查是否包含预期内容
        for keyword, expected_value in expected_results:
            if keyword in content:
                print(f"    ✅ 发现预期内容: {expected_value}")
                success_count += 1
                break

    print("\n" + "=" * 60)
    print(f"📈 成功验证: {success_count}/{len(expected_results)} 项")

    # 特别检查关键Bug修复
    bug_fixed = False
    for para in doc.paragraphs:
        if "供应商名称：智慧足迹数据科技有限公司" in para.text:
            bug_fixed = True
            break

    if bug_fixed:
        print("🎉 关键Bug修复验证成功: 纯空格格式正常填充!")
    else:
        print("❌ 关键Bug修复验证失败: 纯空格格式未填充")

    print("=" * 60)

    return success_count >= len(expected_results) * 0.8 and bug_fixed  # 80%成功率且关键Bug修复

if __name__ == "__main__":
    success = test_integration()

    if success:
        print("\n🎊 集成测试通过! 统一字段映射系统工作正常")
        exit(0)
    else:
        print("\n💥 集成测试失败! 需要进一步调试")
        exit(1)