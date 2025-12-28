#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标书智能处理 HITL API路由
支持三步人工确认流程
"""

import json
import re
import uuid
import os
import shutil
from datetime import datetime
from flask import request, jsonify, send_file, render_template
from pathlib import Path

from common import get_module_logger, resolve_file_path
from common.database import get_knowledge_base_db

# 导入结构解析器
import sys
sys.path.append(str(Path(__file__).parent.parent))
from modules.tender_processing.structure_parser import DocumentStructureParser

logger = get_module_logger("api_hitl")


def register_hitl_routes(app):
    """注册 HITL API 路由"""

    # 获取数据库连接
    db = get_knowledge_base_db()

    # ============================================
    # 页面路由
    # ============================================

    @app.route('/tender-processing-hitl')
    def tender_processing_hitl_page():
        """渲染HITL流程页面"""
        return render_template('tender_processing_hitl.html')

    # ============================================
    # 步骤1：章节选择相关API
    # ============================================

    @app.route('/api/tender-processing/parse-structure', methods=['POST'])
    def parse_document_structure():
        """
        解析文档结构（步骤1的第一步）

        请求参数（FormData）：
        - file: Word文档文件（与 file_path 二选一）
        - file_path: 历史文件路径（与 file 二选一）
        - company_id: 公司ID（必填）
        - project_id: 项目ID（可选，为空时自动创建新项目）
        - methods: 解析方法列表（可选，JSON数组字符串）
                  例如: '["toc_exact", "style"]'
                  可选值: toc_exact, semantic_anchors, style, hybrid, azure, outline_level, gemini
        - fallback: 是否启用回退机制（可选，true/false字符串，默认true）

        返回：
        {
            "success": True/False,
            "project_id": xxx,  # 新建或已有项目ID
            "chapters": [...],  # 章节树
            "statistics": {...},
            "method": "使用的解析方法"
        }
        """
        try:
            # 获取参数
            company_id = request.form.get('company_id')
            project_id = request.form.get('project_id')  # 可选
            file_path_param = request.form.get('file_path')  # 历史文件路径（可选）

            # 新增：解析方法参数
            methods_param = request.form.get('methods')  # 可选
            fallback_param = request.form.get('fallback', 'true')  # 可选，默认true

            if not company_id:
                return jsonify({'success': False, 'error': '缺少company_id参数'}), 400

            # 解析methods参数（如果提供）
            methods = None
            if methods_param:
                try:
                    methods = json.loads(methods_param)
                    if not isinstance(methods, list):
                        return jsonify({'success': False, 'error': 'methods参数必须是JSON数组'}), 400
                    logger.info(f"使用指定解析方法: {methods}")
                except json.JSONDecodeError:
                    return jsonify({'success': False, 'error': 'methods参数格式错误'}), 400

            # 解析fallback参数
            fallback = fallback_param.lower() in ('true', '1', 'yes')

            # 检查文件：支持两种模式
            file_path = None
            original_filename = None

            if file_path_param:
                # 模式1：使用历史文件路径
                import os
                if not os.path.exists(file_path_param):
                    return jsonify({'success': False, 'error': f'文件不存在: {file_path_param}'}), 400

                file_path = file_path_param
                original_filename = os.path.basename(file_path_param)
                logger.info(f"使用历史文件: {file_path}")

            elif 'file' in request.files:
                # 模式2：上传新文件
                file = request.files['file']
                if not file.filename:
                    return jsonify({'success': False, 'error': '文件名为空'}), 400

                # 保存文件到临时目录
                from core.storage_service import storage_service
                file_metadata = storage_service.store_file(
                    file_obj=file,
                    original_name=file.filename,
                    category='tender_processing',
                    business_type='tender_hitl_document'
                )

                file_path = file_metadata.file_path
                # 修正路径：storage_service返回的是相对路径，需要补充ai_tender_system前缀
                if not file_path.startswith('ai_tender_system/'):
                    file_path = f"ai_tender_system/{file_path}"
                original_filename = file.filename
                logger.info(f"新文件已保存: {file_path}")

            else:
                return jsonify({'success': False, 'error': '未提供文件或文件路径'}), 400

            # 如果没有提供project_id，创建新项目
            if not project_id:
                logger.info(f"未提供project_id，自动创建新项目 (company_id: {company_id})")

                # 创建新项目记录（让数据库自动生成project_id）
                db.execute_query("""
                    INSERT INTO tender_projects (
                        company_id, project_name,
                        tender_document_path, original_filename,
                        status, created_at
                    ) VALUES (?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
                """, (
                    company_id,
                    f"标书项目_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                    file_path,
                    original_filename,
                    'draft'
                ))

                # 获取刚刚插入的project_id
                result = db.execute_query("""
                    SELECT project_id FROM tender_projects
                    WHERE company_id = ?
                    ORDER BY created_at DESC
                    LIMIT 1
                """, (company_id,), fetch_one=True)

                if result:
                    project_id = result['project_id']
                    logger.info(f"✅ 新项目已创建: {project_id}")
                else:
                    raise Exception("创建项目后无法获取project_id")

            # 解析文档结构
            parser = DocumentStructureParser()
            result = parser.parse_document_structure(
                file_path,
                methods=methods,
                fallback=fallback
            )

            if not result["success"]:
                return jsonify(result), 500

            # 删除旧的章节数据（如果存在）
            db.execute_query("""
                DELETE FROM tender_document_chapters
                WHERE project_id = ?
            """, (project_id,))
            logger.info(f"已清除项目 {project_id} 的旧章节数据")

            # 保存章节到数据库
            chapter_ids = _save_chapters_to_db(
                db, result["chapters"], project_id
            )

            # 检查是否已存在HITL任务
            existing_task = db.execute_query("""
                SELECT project_id FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if existing_task:
                # 更新现有任务
                # 🔧 修复：同时更新 tender_document_path 和 original_filename 字段
                # 确保悬浮按钮可以检测到文件路径
                db.execute_query("""
                    UPDATE tender_projects
                    SET step1_status = 'in_progress',
                        step1_data = ?,
                        tender_document_path = ?,
                        original_filename = ?,
                        hitl_estimated_words = ?,
                        hitl_estimated_cost = ?,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE project_id = ?
                """, (
                    json.dumps({
                        'file_path': file_path,
                        'file_name': original_filename,
                        'chapters': result["chapters"]  # 🆕 保存章节数据（包含 chapter_type）
                    }),
                    file_path,
                    original_filename,
                    result["statistics"].get("total_words", 0),
                    result["statistics"].get("estimated_processing_cost", 0.0),
                    project_id
                ))
                logger.info(f"HITL任务已更新，project_id: {project_id}, 文件路径: {file_path}")
            else:
                # 创建新任务
                db.execute_query("""
                    INSERT INTO tender_projects (
                        project_id,
                        step1_status, step1_data,
                        hitl_estimated_words, hitl_estimated_cost
                    ) VALUES (?, ?, ?, ?, ?)
                """, (
                    project_id,
                    'in_progress',
                    json.dumps({
                        'file_path': file_path,
                        'file_name': original_filename,
                        'chapters': result["chapters"]  # 🆕 保存章节数据（包含 chapter_type）
                    }),
                    result["statistics"].get("total_words", 0),
                    result["statistics"].get("estimated_processing_cost", 0.0)
                ))
                logger.info(f"HITL任务已创建，project_id: {project_id}")

            return jsonify({
                'success': True,
                'project_id': project_id,  # 返回项目ID（新建或已有）
                'chapters': result["chapters"],
                'statistics': result["statistics"]
            })

        except Exception as e:
            logger.error(f"解析文档结构失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/select-chapters', methods=['POST'])
    def submit_chapter_selection():
        """
        提交章节选择结果（步骤1的第二步）

        请求参数（JSON）：
        {
            "project_id": xxx,
            "selected_chapter_ids": ["ch_0", "ch_1", ...]
        }

        返回：
        {
            "success": True/False,
            "selected_count": 5,
            "selected_words": 8000,
            "estimated_cost": 0.016
        }
        """
        try:
            data = request.get_json()
            project_id = data.get('project_id')
            selected_ids = data.get('selected_chapter_ids', [])

            if not project_id:
                return jsonify({'success': False, 'error': '缺少project_id参数'}), 400

            db = get_knowledge_base_db()

            # 更新章节选择状态
            for chapter_id in selected_ids:
                db.execute_query("""
                    UPDATE tender_document_chapters
                    SET is_selected = 1, updated_at = CURRENT_TIMESTAMP
                    WHERE chapter_node_id = ? AND project_id = ?
                """, (chapter_id, project_id))

            # 记录用户操作
            db.execute_query("""
                INSERT INTO tender_user_actions (
                    project_id, action_type, action_step, action_data
                ) VALUES (?, 'chapter_selected', 1, ?)
            """, (project_id, json.dumps({'selected_ids': selected_ids})))

            # 统计选中章节
            stats = db.execute_query("""
                SELECT
                    COUNT(*) as selected_count,
                    SUM(word_count) as selected_words
                FROM tender_document_chapters
                WHERE project_id = ? AND is_selected = 1
            """, (project_id,), fetch_one=True)

            selected_words = stats['selected_words'] or 0
            estimated_cost = (selected_words / 1000) * 0.002  # 假设成本

            # 读取原有的step1_data,保留file_path
            existing_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            # 合并step1_data,保留file_path
            step1_data = {}
            if existing_data and existing_data.get('step1_data'):
                step1_data = json.loads(existing_data['step1_data'])

            # 更新选择信息
            step1_data.update({
                'selected_ids': selected_ids,
                'selected_count': stats['selected_count']
            })

            # 更新 HITL 任务状态
            db.execute_query("""
                UPDATE tender_projects
                SET step1_status = 'completed',
                    step1_completed_at = CURRENT_TIMESTAMP,
                    step1_data = ?,
                    hitl_estimated_words = ?,
                    hitl_estimated_cost = ?
                WHERE project_id = ?
            """, (
                json.dumps(step1_data),
                selected_words,
                estimated_cost,
                project_id
            ))

            logger.info(f"步骤1完成: 选中 {stats['selected_count']} 个章节, {selected_words} 字")

            # 更新任务状态到步骤2（移除自动提取逻辑）
            db.execute_query("""
                UPDATE tender_projects
                SET current_step = 2,
                    step2_status = 'in_progress'
                WHERE project_id = ?
            """, (project_id,))

            logger.info(f"任务状态已更新到步骤2，等待用户在Tab 3触发AI提取")

            return jsonify({
                'success': True,
                'project_id': project_id,
                'selected_count': stats['selected_count'],
                'selected_words': selected_words,
                'estimated_cost': estimated_cost
            })

        except Exception as e:
            logger.error(f"提交章节选择失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/extract-eligibility-requirements/<int:project_id>', methods=['POST'])
    def extract_eligibility_requirements(project_id):
        """
        提取13条供应商资格要求（关键词匹配）
        使用固定的13项清单模板，使用关键词匹配替代AI调用
        """
        try:
            logger.info(f"开始提取19条供应商资格要求（关键词匹配）: project_id={project_id}")

            # 查询任务信息
            task_info = db.execute_query("""
                SELECT project_id, step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_info:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            project_id = task_info['project_id']
            step1_data = json.loads(task_info['step1_data']) if task_info.get('step1_data') else {}

            doc_path = step1_data.get('file_path')
            selected_ids = step1_data.get('selected_ids', [])

            if not doc_path or not selected_ids:
                return jsonify({'success': False, 'error': '缺少文档路径或选中章节'}), 400

            # 检查是否已有资格要求数据，清空旧数据
            existing_count = db.execute_query("""
                SELECT COUNT(*) as count FROM tender_requirements
                WHERE project_id = ? AND category = 'qualification'
            """, (project_id,), fetch_one=True)

            if existing_count and existing_count['count'] > 0:
                db.execute_query("""
                    DELETE FROM tender_requirements
                    WHERE project_id = ? AND category = 'qualification'
                """, (project_id,))
                logger.info(f"清除了 {existing_count['count']} 个旧资格要求记录")

            # 重新解析文档以获取章节内容
            parser = DocumentStructureParser()
            result = parser.parse_document_structure(doc_path)

            if not result["success"]:
                return jsonify({'success': False, 'error': '文档解析失败'}), 500

            # 从解析结果中提取选中章节的内容
            def find_chapter_by_id(chapters, target_id):
                """递归查找章节"""
                for ch in chapters:
                    if ch.get('id') == target_id:
                        return ch
                    if ch.get('children'):
                        found = find_chapter_by_id(ch['children'], target_id)
                        if found:
                            return found
                return None

            # 读取文档内容
            from docx import Document
            doc = Document(doc_path)
            all_paragraphs = [p.text for p in doc.paragraphs]

            # 获取选中章节的内容
            selected_content = []
            for chapter_id in selected_ids:
                chapter_info = find_chapter_by_id(result["chapters"], chapter_id)
                if chapter_info:
                    start_idx = chapter_info.get('para_start_idx', 0)
                    end_idx = chapter_info.get('para_end_idx', len(all_paragraphs))
                    chapter_text = '\n'.join(all_paragraphs[start_idx:end_idx])

                    selected_content.append({
                        'chapter_id': chapter_id,
                        'chapter_title': chapter_info.get('title', ''),
                        'content': chapter_text
                    })

            logger.info(f"提取了 {len(selected_content)} 个章节的内容")

            # 合并所有选中章节的内容为一个文本
            full_text = '\n'.join([c['content'] for c in selected_content])
            logger.info(f"合并文本总长度: {len(full_text)} 字")

            # 使用关键词匹配提取13条资格要求
            from modules.tender_info.extractor import TenderInfoExtractor
            extractor = TenderInfoExtractor()

            checklist_results = extractor.extract_supplier_eligibility_checklist(full_text)

            # 保存到数据库
            total_saved = 0
            for checklist_item in checklist_results:
                if checklist_item['found']:
                    for req in checklist_item['requirements']:
                        db.execute_query("""
                            INSERT INTO tender_requirements (
                                project_id, constraint_type, category,
                                subcategory, detail, summary, source_location, priority,
                                extraction_confidence
                            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """, (
                            project_id,
                            req.get('constraint_type', 'mandatory'),
                            'qualification',
                            checklist_item['checklist_name'],
                            req.get('detail', ''),
                            req.get('summary', ''),
                            req.get('source_location', ''),
                            'high',
                            req.get('extraction_confidence', 0.8)
                        ))
                        total_saved += 1

            # 统计找到的项数
            found_count = sum(1 for item in checklist_results if item['found'])

            logger.info(f"19条资格要求提取完成（关键词匹配）: 找到 {found_count} 项，未找到 {19 - found_count} 项，保存 {total_saved} 条要求")

            # 【新增】同步到项目表
            try:
                sync_requirements_to_project(project_id, 'qualification')
            except Exception as sync_error:
                logger.warning(f"同步资格要求到项目表失败，但不影响主流程: {sync_error}")

            return jsonify({
                'success': True,
                'checklist': checklist_results,
                'method': 'keyword_matching',  # 标识使用关键词匹配
                'requirements_saved': total_saved,
                'found_count': found_count,
                'not_found_count': 19 - found_count
            })

        except Exception as e:
            logger.error(f"提取13条资格要求失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/export-chapter/<int:project_id>/<chapter_id>', methods=['GET'])
    def export_chapter_as_template(project_id, chapter_id):
        """
        导出单个章节为Word文档模板（步骤1功能扩展）

        Args:
            project_id: 项目ID (如 "hitl_abc123")
            chapter_id: 章节ID (如 "ch_4")

        Returns:
            Word文档文件流 (application/vnd.openxmlformats-officedocument.wordprocessingml.document)

        使用示例:
            GET /api/tender-processing/export-chapter/hitl_abc123/ch_4
            => 下载: 第五部分_响应文件格式_模板_20251006.docx
        """
        try:
            from datetime import datetime
            from flask import send_file
            import re

            db = get_knowledge_base_db()

            # 1. 查询HITL任务，获取原始文档路径
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            step1_data = json.loads(task_data['step1_data'])
            doc_path = step1_data.get('file_path')

            if not doc_path or not Path(doc_path).exists():
                return jsonify({'success': False, 'error': '原始文档不存在'}), 404

            # 2. 调用结构解析器导出章节
            from modules.tender_processing.structure_parser import DocumentStructureParser
            parser = DocumentStructureParser()
            result = parser.export_chapter_to_docx(doc_path, chapter_id)

            if not result['success']:
                return jsonify(result), 500

            # 3. 返回文件流
            output_path = result['file_path']
            chapter_title = result['chapter_title']

            # 生成下载文件名（移除特殊字符）
            safe_title = re.sub(r'[^\w\s-]', '', chapter_title).strip()
            safe_title = re.sub(r'[\s]+', '_', safe_title)  # 空格转下划线
            download_name = f"{safe_title}_应答模板_{datetime.now().strftime('%Y%m%d')}.docx"

            logger.info(f"导出章节模板: {chapter_title} -> {download_name}")

            return send_file(
                output_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=download_name
            )

        except Exception as e:
            logger.error(f"导出章节失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/export-chapters/<int:project_id>', methods=['POST'])
    def export_multiple_chapters(project_id):
        """批量导出多个章节为单个Word文档"""
        try:
            data = request.get_json()
            chapter_ids = data.get('chapter_ids', [])

            if not chapter_ids:
                return jsonify({"error": "未提供章节ID"}), 400

            # 查询任务信息获取文档路径
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            doc_path = step1_data.get('file_path')

            # 调用parser批量导出
            parser = DocumentStructureParser()
            result = parser.export_multiple_chapters_to_docx(doc_path, chapter_ids)

            if not result['success']:
                return jsonify({"error": result.get('error', '导出失败')}), 500

            output_path = result['file_path']
            chapter_titles = result.get('chapter_titles', [])

            # 生成文件名
            safe_titles = '_'.join(chapter_titles[:3])  # 最多3个标题
            if len(chapter_titles) > 3:
                safe_titles += f'_等{len(chapter_titles)}章节'
            safe_titles = re.sub(r'[^\w\s-]', '', safe_titles).strip()
            download_name = f"{safe_titles}_应答模板_{datetime.now().strftime('%Y%m%d')}.docx"

            logger.info(f"批量导出章节模板: {len(chapter_ids)}个章节 -> {download_name}")

            return send_file(
                output_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=download_name
            )

        except Exception as e:
            logger.error(f"批量导出章节失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/tender-processing/save-response-file/<int:project_id>', methods=['POST'])
    def save_response_file(project_id):
        """保存应答文件到服务器"""
        try:
            data = request.get_json()
            chapter_ids = data.get('chapter_ids', [])
            custom_filename = data.get('filename')

            if not chapter_ids:
                return jsonify({"error": "未提供章节ID"}), 400

            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data, project_id FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            doc_path = step1_data.get('file_path')
            project_id = task_data['project_id']

            # 调用parser导出文件
            parser = DocumentStructureParser()
            result = parser.export_multiple_chapters_to_docx(doc_path, chapter_ids)

            if not result['success']:
                return jsonify({"error": result.get('error', '导出失败')}), 500

            # 创建存储目录（使用绝对路径）
            now = datetime.now()
            # 获取项目根目录（ai_tender_system）
            project_root = Path(__file__).parent.parent
            save_dir = os.path.join(
                project_root,
                'data/uploads/response_files',
                str(now.year),
                f"{now.month:02d}",
                str(project_id)
            )
            os.makedirs(save_dir, exist_ok=True)

            # 生成文件名
            chapter_titles = result.get('chapter_titles', [])
            safe_titles = '_'.join(chapter_titles[:2]) if chapter_titles else '应答文件'
            safe_titles = re.sub(r'[^\w\s-]', '', safe_titles).strip()

            if custom_filename:
                filename = custom_filename if custom_filename.endswith('.docx') else f"{custom_filename}.docx"
            else:
                filename = f"{safe_titles}_应答模板_{now.strftime('%Y%m%d_%H%M%S')}.docx"

            # 复制文件到目标位置（使用copy而非move，避免python-docx文件引用问题）
            target_path = os.path.join(save_dir, filename)
            shutil.copy(result['file_path'], target_path)

            # 计算文件大小
            file_size = os.path.getsize(target_path)

            # 更新任务的step1_data - 直接保存路径字符串，符合前端设计
            step1_data['response_file_path'] = target_path

            # ⭐ 同时更新独立的路径字段（新设计）
            db.execute_query("""
                UPDATE tender_projects
                SET step1_data = ?,
                    response_template_path = ?
                WHERE project_id = ?
            """, (json.dumps(step1_data), target_path, project_id))

            logger.info(f"保存应答文件: {filename} ({file_size} bytes)")

            return jsonify({
                "success": True,
                "file_path": target_path,
                "file_url": f"/api/tender-processing/download-response-file/{project_id}",
                "filename": filename,
                "file_size": file_size,
                "saved_at": now.isoformat()
            })

        except Exception as e:
            logger.error(f"保存应答文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/tender-processing/download-response-file/<int:project_id>', methods=['GET'])
    def download_response_file(project_id):
        """下载已保存的应答文件"""
        try:
            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            response_file = step1_data.get('response_file')

            if not response_file:
                return jsonify({"error": "应答文件不存在"}), 404

            file_path = response_file['file_path']
            if not os.path.exists(file_path):
                return jsonify({"error": "文件已被删除"}), 404

            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=response_file['filename']
            )

        except Exception as e:
            logger.error(f"下载应答文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/tender-processing/preview-response-file/<int:project_id>', methods=['GET'])
    def preview_response_file(project_id):
        """预览已保存的应答文件"""
        try:
            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            response_file = step1_data.get('response_file')

            if not response_file:
                return jsonify({"error": "应答文件不存在"}), 404

            file_path = response_file['file_path']
            if not os.path.exists(file_path):
                return jsonify({"error": "文件已被删除"}), 404

            # 预览模式：as_attachment=False，浏览器会尝试在线打开
            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=False,
                download_name=response_file['filename']
            )

        except Exception as e:
            logger.error(f"预览应答文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/tender-processing/response-file-info/<int:project_id>', methods=['GET'])
    def get_response_file_info(project_id):
        """获取应答文件信息"""
        try:
            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"success": False, "error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            response_file = step1_data.get('response_file')

            if not response_file:
                return jsonify({
                    "success": True,
                    "has_file": False
                })

            # 检查文件是否存在
            file_exists = os.path.exists(response_file['file_path'])

            return jsonify({
                "success": True,
                "has_file": file_exists,
                "filename": response_file.get('filename'),
                "file_path": response_file.get('file_path'),  # ✅ 添加 file_path 字段
                "file_size": response_file.get('file_size'),
                "saved_at": response_file.get('saved_at'),
                "download_url": f"/api/tender-processing/download-response-file/{project_id}"
            })

        except Exception as e:
            logger.error(f"获取应答文件信息失败: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500

    # ============================================
    # 技术需求章节相关API
    # ============================================

    @app.route('/api/tender-processing/save-technical-chapters/<int:project_id>', methods=['POST'])
    def save_technical_chapters(project_id):
        """保存技术需求章节"""
        try:
            data = request.get_json()
            chapter_ids = data.get('chapter_ids', [])

            if not chapter_ids:
                return jsonify({"success": False, "error": "未选择章节"}), 400

            # 获取任务信息
            task_info = db.execute_query("""
                SELECT project_id, step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_info:
                return jsonify({"success": False, "error": "任务不存在"}), 404

            project_id = task_info['project_id']

            # 获取原始文档路径
            step1_data = json.loads(task_info['step1_data'])
            doc_path = step1_data.get('file_path')  # 与应答文件API保持一致

            if not doc_path or not os.path.exists(doc_path):
                return jsonify({"success": False, "error": "原始文档不存在"}), 404

            # 调用DocumentStructureParser导出章节
            parser = DocumentStructureParser()
            result = parser.export_multiple_chapters_to_docx(doc_path, chapter_ids)

            if not result.get('success'):
                return jsonify({"success": False, "error": result.get('error', '导出失败')}), 500

            # 创建存储目录（使用绝对路径）
            now = datetime.now()
            # 获取项目根目录（ai_tender_system）
            project_root = Path(__file__).parent.parent
            save_dir = os.path.join(
                project_root,
                'data/uploads/technical_files',
                str(now.year),
                f"{now.month:02d}",
                str(project_id)
            )
            os.makedirs(save_dir, exist_ok=True)

            # 生成文件名
            chapter_titles = result.get('chapter_titles', [])
            safe_titles = '_'.join(chapter_titles[:2]) if chapter_titles else '技术需求'
            safe_titles = re.sub(r'[^\w\s-]', '', safe_titles).strip()
            filename = f"{safe_titles}_技术需求_{now.strftime('%Y%m%d_%H%M%S')}.docx"

            # 复制文件到目标位置（使用copy而非move，避免python-docx文件引用问题）
            target_path = os.path.join(save_dir, filename)
            shutil.copy(result['file_path'], target_path)

            # 计算文件大小
            file_size = os.path.getsize(target_path)

            # 更新step1_data - 直接保存路径字符串，符合前端设计
            step1_data['technical_file_path'] = target_path

            # ⭐ 更新数据库（同时更新独立字段）
            db.execute_query("""
                UPDATE tender_projects
                SET step1_data = ?,
                    technical_requirement_path = ?
                WHERE project_id = ?
            """, (json.dumps(step1_data, ensure_ascii=False), target_path, project_id))

            logger.info(f"✅ 技术需求章节已保存: {filename} ({file_size} bytes)")

            return jsonify({
                "success": True,
                "file_path": target_path,
                "file_url": f"/api/tender-processing/download-technical-file/{project_id}",
                "filename": filename,
                "file_size": file_size,
                "saved_at": now.isoformat(),
                "message": "技术需求章节已成功保存"
            })

        except Exception as e:
            logger.error(f"保存技术需求章节失败: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route('/api/tender-processing/technical-file-info/<int:project_id>', methods=['GET'])
    def get_technical_file_info(project_id):
        """获取技术需求文件信息(支持文件系统扫描fallback)"""
        try:
            file_info = None

            # 尝试从数据库获取文件信息
            try:
                task_data = db.execute_query("""
                    SELECT step1_data FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if task_data:
                    step1_data = json.loads(task_data['step1_data'])
                    file_info = step1_data.get('technical_file')
            except Exception as db_error:
                logger.warning(f"数据库查询失败,将尝试文件系统扫描: {str(db_error)}")

            # 如果数据库中没有文件信息,尝试从文件系统扫描
            if not file_info:
                logger.info(f"数据库中无技术需求文件信息,尝试从文件系统扫描")
                file_path = _find_file_in_filesystem(project_id, 'technical_file')

                if file_path:
                    # 构建文件信息
                    file_stat = file_path.stat()
                    file_info = {
                        'file_path': str(file_path),
                        'filename': file_path.name,
                        'file_size': file_stat.st_size,
                        'saved_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                    }
                    logger.info(f"从文件系统找到技术需求文件: {file_path.name}")

            if not file_info:
                return jsonify({
                    "success": True,
                    "has_file": False
                })

            # 检查文件是否存在
            file_exists = os.path.exists(file_info.get('file_path', ''))

            # 如果数据库中的文件路径不存在,尝试从文件系统扫描
            if not file_exists and 'file_path' in file_info:
                logger.warning(f"数据库中的文件路径不存在: {file_info['file_path']}, 尝试文件系统扫描")
                file_path = _find_file_in_filesystem(project_id, 'technical_file')

                if file_path:
                    # 更新文件信息
                    file_stat = file_path.stat()
                    file_info = {
                        'file_path': str(file_path),
                        'filename': file_path.name,
                        'file_size': file_stat.st_size,
                        'saved_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                    }
                    file_exists = True
                    logger.info(f"从文件系统找到替代文件: {file_path.name}")

            return jsonify({
                "success": True,
                "has_file": file_exists,
                "filename": file_info.get('filename'),
                "file_size": file_info.get('file_size'),
                "saved_at": file_info.get('saved_at'),
                "download_url": f"/api/tender-processing/download-technical-file/{project_id}"
            })

        except Exception as e:
            logger.error(f"获取技术需求文件信息失败: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route('/api/tender-processing/download-technical-file/<int:project_id>', methods=['GET'])
    def download_technical_file(project_id):
        """下载技术需求文件(支持文件系统扫描fallback)"""
        try:
            file_info = None
            file_path = None

            # 尝试从数据库获取文件信息
            try:
                task_data = db.execute_query("""
                    SELECT step1_data FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if task_data:
                    step1_data = json.loads(task_data['step1_data'])
                    file_info = step1_data.get('technical_file')
                    if file_info:
                        file_path = file_info.get('file_path')
            except Exception as db_error:
                logger.warning(f"数据库查询失败,将尝试文件系统扫描: {str(db_error)}")

            # 如果数据库中没有或文件不存在,尝试从文件系统扫描
            if not file_path or not os.path.exists(file_path):
                logger.info(f"数据库中的文件路径无效,尝试从文件系统扫描")
                found_path = _find_file_in_filesystem(project_id, 'technical_file')
                if found_path:
                    file_path = str(found_path)
                    logger.info(f"从文件系统找到技术需求文件: {found_path.name}")

            if not file_path or not os.path.exists(file_path):
                return jsonify({"error": "技术需求文件不存在或已被删除"}), 404

            # 获取文件名
            filename = file_info.get('filename') if file_info else Path(file_path).name

            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )

        except Exception as e:
            logger.error(f"下载技术需求文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/tender-processing/preview-technical-file/<int:project_id>', methods=['GET'])
    def preview_technical_file(project_id):
        """预览技术需求文件(支持文件系统扫描fallback)"""
        try:
            file_info = None
            file_path = None

            # 尝试从数据库获取文件信息
            try:
                task_data = db.execute_query("""
                    SELECT step1_data FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if task_data:
                    step1_data = json.loads(task_data['step1_data'])
                    file_info = step1_data.get('technical_file')
                    if file_info:
                        file_path = file_info.get('file_path')
            except Exception as db_error:
                logger.warning(f"数据库查询失败,将尝试文件系统扫描: {str(db_error)}")

            # 如果数据库中没有或文件不存在,尝试从文件系统扫描
            if not file_path or not os.path.exists(file_path):
                logger.info(f"数据库中的文件路径无效,尝试从文件系统扫描")
                found_path = _find_file_in_filesystem(project_id, 'technical_file')
                if found_path:
                    file_path = str(found_path)
                    logger.info(f"从文件系统找到技术需求文件: {found_path.name}")

            if not file_path or not os.path.exists(file_path):
                return jsonify({"error": "技术需求文件不存在或已被删除"}), 404

            # 获取文件名
            filename = file_info.get('filename') if file_info else Path(file_path).name

            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=False,
                download_name=filename
            )

        except Exception as e:
            logger.error(f"预览技术需求文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    # ============================================
    # 获取章节列表API（用于步骤3章节选择）
    # ============================================

    @app.route('/api/tender-processing/chapters/<int:project_id>', methods=['GET'])
    def get_chapters_list(project_id):
        """
        获取任务的章节列表（从数据库读取）

        用于步骤3的章节选择功能，当步骤1没有执行时也能获取章节数据

        返回：
        {
            "success": True,
            "chapters": [
                {
                    "id": "ch_1",
                    "level": 1,
                    "title": "第一章 项目概述",
                    "word_count": 1500,
                    "para_start_idx": 0,
                    "para_end_idx": 10,
                    "preview_text": "...",
                    "auto_selected": True,
                    "skip_recommended": False,
                    "content_tags": ["技术需求"]
                },
                ...
            ]
        }
        """
        try:
            db = get_knowledge_base_db()

            # 验证HITL任务是否存在
            hitl_task = db.execute_query("""
                SELECT project_id FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not hitl_task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 从数据库获取章节列表
            chapters_raw = db.execute_query("""
                SELECT
                    chapter_node_id,
                    level,
                    title,
                    para_start_idx,
                    para_end_idx,
                    word_count,
                    preview_text,
                    is_selected,
                    auto_selected,
                    skip_recommended,
                    parent_chapter_id
                FROM tender_document_chapters
                WHERE project_id = ?
                ORDER BY para_start_idx ASC
            """, (project_id,))

            if not chapters_raw:
                return jsonify({
                    'success': False,
                    'error': '该任务还没有章节数据，请先在步骤1解析文档'
                }), 404

            # 转换为前端需要的格式
            chapters = []
            for ch in chapters_raw:
                # 从chapter_node_id提取出简单的ID（如ch_1, ch_1_2等）
                chapter_data = {
                    'id': ch['chapter_node_id'],
                    'level': ch['level'],
                    'title': ch['title'],
                    'word_count': ch['word_count'] or 0,
                    'para_start_idx': ch['para_start_idx'],
                    'para_end_idx': ch['para_end_idx'],
                    'preview_text': ch['preview_text'] or '',
                    'auto_selected': bool(ch['auto_selected']),
                    'skip_recommended': bool(ch['skip_recommended']),
                    'content_tags': []  # TODO: 如果有标签数据，从其他表获取
                }
                chapters.append(chapter_data)

            logger.info(f"✅ 获取章节列表成功: {len(chapters)}个章节")

            return jsonify({
                'success': True,
                'chapters': chapters,
                'total': len(chapters)
            })

        except Exception as e:
            logger.error(f"获取章节列表失败: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route('/api/tender-processing/chapter-content/<int:project_id>/<chapter_id>', methods=['GET'])
    def get_chapter_full_content(project_id, chapter_id):
        """
        实时从原始文档提取章节的完整内容

        Args:
            project_id: 项目ID
            chapter_id: 章节ID (如 ch_1, ch_1_2)

        Returns:
            {
                "success": True,
                "content": "章节完整文本内容...",
                "word_count": 1500,
                "title": "章节标题"
            }
        """
        try:
            from docx import Document

            db = get_knowledge_base_db()

            # 1. 获取章节信息
            hitl_task = db.execute_query("""
                SELECT project_id, step1_data
                FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not hitl_task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 2. 查询章节的段落范围
            chapter = db.execute_query("""
                SELECT
                    title,
                    para_start_idx,
                    para_end_idx,
                    word_count
                FROM tender_document_chapters
                WHERE project_id = ? AND chapter_node_id = ?
            """, (project_id, chapter_id), fetch_one=True)

            if not chapter:
                return jsonify({'success': False, 'error': '章节不存在'}), 404

            # 3. 获取原始文档路径
            step1_data = hitl_task.get('step1_data')
            if isinstance(step1_data, str):
                import json
                step1_data = json.loads(step1_data)

            if not step1_data or 'file_path' not in step1_data:
                return jsonify({
                    'success': False,
                    'error': '原始文档路径未找到，无法提取完整内容'
                }), 404

            file_path = Path(step1_data['file_path'])

            if not file_path.exists():
                return jsonify({
                    'success': False,
                    'error': '原始文档文件不存在，可能已被删除'
                }), 404

            # 4. 从文档中提取指定段落范围的文本
            doc = Document(str(file_path))
            para_start = chapter['para_start_idx']
            para_end = chapter['para_end_idx']

            # 如果para_end为None，表示到文档末尾
            if para_end is None:
                para_end = len(doc.paragraphs) - 1

            # 提取段落内容
            content_lines = []
            for i in range(para_start, min(para_end + 1, len(doc.paragraphs))):
                text = doc.paragraphs[i].text.strip()
                if text:  # 只添加非空段落
                    content_lines.append(text)

            full_content = '\n\n'.join(content_lines)

            logger.info(f"✅ 成功提取章节完整内容: {chapter_id}, 共{len(full_content)}字符")

            return jsonify({
                'success': True,
                'content': full_content,
                'word_count': len(full_content.replace(' ', '').replace('\n', '')),
                'title': chapter['title'],
                'para_range': f"{para_start}-{para_end}"
            })

        except Exception as e:
            logger.error(f"获取章节完整内容失败: {str(e)}", exc_info=True)
            return jsonify({
                "success": False,
                "error": f"提取章节内容失败: {str(e)}"
            }), 500

    # ============================================
    # 步骤2：章节要求预览相关API
    # ============================================

    @app.route('/api/tender-processing/chapter-requirements/<int:project_id>', methods=['GET'])
    def get_chapter_requirements_summary(project_id):
        """
        获取各章节的要求聚合信息（步骤2）

        按章节分组统计AI提取的要求数量

        返回：
        {
            "success": True,
            "chapters": [
                {
                    "chapter_id": 123,
                    "title": "第一章 项目概述",
                    "word_count": 1500,
                    "is_selected": True,
                    "requirement_stats": {
                        "total": 15,
                        "mandatory": 10,
                        "scoring": 3,
                        "optional": 2
                    }
                },
                ...
            ],
            "summary": {
                "total_chapters": 10,
                "selected_chapters": 8,
                "total_requirements": 50
            }
        }
        """
        try:
            db = get_knowledge_base_db()

            # 获取任务关联的project_id
            hitl_task = db.execute_query("""
                SELECT project_id FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not hitl_task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 获取所有章节（步骤1选中的章节）
            chapters = db.execute_query("""
                SELECT
                    chapter_id,
                    title,
                    word_count,
                    is_selected,
                    level
                FROM tender_document_chapters
                WHERE project_id = ?
                ORDER BY chapter_node_id
            """, (project_id,))

            # 为每个章节统计要求数量
            chapter_list = []
            total_requirements = 0

            for chapter in chapters:
                # 统计该章节的要求数量（通过source_location匹配章节标题）
                req_stats = db.execute_query("""
                    SELECT
                        constraint_type,
                        COUNT(*) as count
                    FROM tender_requirements
                    WHERE project_id = ?
                      AND source_location LIKE ?
                    GROUP BY constraint_type
                """, (project_id, f"%{chapter['title']}%"))

                # 构建统计数据
                stats = {'total': 0, 'mandatory': 0, 'scoring': 0, 'optional': 0}
                for stat in req_stats:
                    count = stat['count']
                    stats['total'] += count
                    stats[stat['constraint_type']] = count

                total_requirements += stats['total']

                chapter_list.append({
                    'chapter_id': chapter['chapter_id'],
                    'title': chapter['title'],
                    'word_count': chapter['word_count'] or 0,
                    'is_selected': bool(chapter['is_selected']),
                    'level': chapter['level'],
                    'requirement_stats': stats
                })

            # 汇总统计
            selected_count = sum(1 for c in chapter_list if c['is_selected'])

            return jsonify({
                'success': True,
                'chapters': chapter_list,
                'summary': {
                    'total_chapters': len(chapter_list),
                    'selected_chapters': selected_count,
                    'total_requirements': total_requirements
                }
            })

        except Exception as e:
            logger.error(f"获取章节要求汇总失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/filtered-blocks/<int:project_id>', methods=['GET'])
    def get_filtered_blocks(project_id):
        """
        获取被AI筛选掉的文本块（步骤2）

        返回：
        {
            "success": True,
            "filtered_blocks": [
                {
                    "chunk_id": 123,
                    "content": "...",
                    "ai_decision": "NON-REQUIREMENT",
                    "ai_confidence": 0.85,
                    "ai_reasoning": "这是描述性文本"
                },
                ...
            ],
            "statistics": {
                "total_filtered": 50,
                "high_confidence": 40,
                "low_confidence": 10
            }
        }
        """
        try:
            db = get_knowledge_base_db()

            # 获取任务关联的project_id
            hitl_task = db.execute_query("""
                SELECT project_id FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not hitl_task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 获取被过滤的块（只查询当前HITL任务的数据）
            filtered_blocks = db.execute_query("""
                SELECT
                    c.chunk_id,
                    c.content,
                    c.chunk_type,
                    c.is_valuable,
                    c.filter_confidence,
                    r.ai_decision,
                    r.ai_confidence,
                    r.ai_reasoning,
                    r.user_decision,
                    r.reviewed_at
                FROM tender_document_chunks c
                LEFT JOIN tender_filter_review r ON c.chunk_id = r.chunk_id
                WHERE c.project_id = ? AND c.is_valuable = 0
                ORDER BY c.chunk_index
            """, (project_id,))

            # 统计（处理None值）
            high_conf = sum(1 for b in filtered_blocks if (b.get('ai_confidence') or 0) >= 0.7)
            low_conf = len(filtered_blocks) - high_conf

            return jsonify({
                'success': True,
                'filtered_blocks': filtered_blocks,
                'statistics': {
                    'total_filtered': len(filtered_blocks),
                    'high_confidence': high_conf,
                    'low_confidence': low_conf
                }
            })

        except Exception as e:
            logger.error(f"获取筛选块失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/update-chapter-selection', methods=['POST'])
    def update_chapter_selection_step2():
        """
        更新章节筛选状态（步骤2）

        允许用户在步骤2取消某些章节，该章节的要求不会出现在步骤3

        请求参数（JSON）：
        {
            "task_id": "hitl_xxx",
            "chapter_ids": [123, 456, ...],  // 选中的章节ID
            "deselected_chapter_ids": [789, ...]  // 取消选中的章节ID
        }

        返回：
        {
            "success": True,
            "selected_count": 8,
            "deselected_count": 2
        }
        """
        try:
            data = request.get_json()
            project_id = data.get('project_id')
            selected_ids = data.get('chapter_ids', [])
            deselected_ids = data.get('deselected_chapter_ids', [])

            if not project_id:
                return jsonify({'success': False, 'error': '缺少project_id'}), 400

            db = get_knowledge_base_db()

            # 更新选中状态
            if selected_ids:
                placeholders = ','.join(['?' for _ in selected_ids])
                db.execute_query(f"""
                    UPDATE tender_document_chapters
                    SET is_selected = 1, updated_at = CURRENT_TIMESTAMP
                    WHERE chapter_id IN ({placeholders})
                """, selected_ids)

            # 更新取消选中状态
            if deselected_ids:
                placeholders = ','.join(['?' for _ in deselected_ids])
                db.execute_query(f"""
                    UPDATE tender_document_chapters
                    SET is_selected = 0, updated_at = CURRENT_TIMESTAMP
                    WHERE chapter_id IN ({placeholders})
                """, deselected_ids)

            # 记录用户操作
            db.execute_query("""
                INSERT INTO tender_user_actions (
                    project_id, action_type, action_step, action_data
                ) SELECT project_id, 'chapter_reselection', 2, ?
                FROM tender_projects WHERE project_id = ?
            """, (json.dumps({
                'selected_count': len(selected_ids),
                'deselected_count': len(deselected_ids)
            }), project_id))

            logger.info(f"步骤2更新章节选择: 选中{len(selected_ids)}个, 取消{len(deselected_ids)}个")

            return jsonify({
                'success': True,
                'selected_count': len(selected_ids),
                'deselected_count': len(deselected_ids)
            })

        except Exception as e:
            logger.error(f"更新章节选择失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/restore-blocks', methods=['POST'])
    def restore_filtered_blocks():
        """
        恢复被误判的文本块（步骤2）

        请求参数（JSON）：
        {
            "task_id": "hitl_xxx",
            "chunk_ids": [123, 456, ...]
        }

        返回：
        {
            "success": True,
            "restored_count": 3
        }
        """
        try:
            data = request.get_json()
            project_id = data.get('project_id')
            chunk_ids = data.get('chunk_ids', [])

            if not project_id or not chunk_ids:
                return jsonify({'success': False, 'error': '参数不完整'}), 400

            db = get_knowledge_base_db()

            # 恢复块（标记为高价值）
            for chunk_id in chunk_ids:
                db.execute_query("""
                    UPDATE tender_document_chunks
                    SET is_valuable = 1, updated_at = CURRENT_TIMESTAMP
                    WHERE chunk_id = ?
                """, (chunk_id,))

                # 记录复核操作
                db.execute_query("""
                    INSERT OR REPLACE INTO tender_filter_review (
                        chunk_id, project_id,
                        ai_decision, user_decision, reviewed_by, reviewed_at
                    ) SELECT
                        ?, project_id, 'NON-REQUIREMENT', 'restore', 'user', CURRENT_TIMESTAMP
                    FROM tender_projects WHERE project_id = ?
                """, (chunk_id, project_id))

            # 记录用户操作
            db.execute_query("""
                INSERT INTO tender_user_actions (
                    project_id, action_type, action_step, action_data
                ) SELECT project_id, 'chunk_restored', 2, ?
                FROM tender_projects WHERE project_id = ?
            """, (json.dumps({'chunk_ids': chunk_ids}), project_id))

            logger.info(f"恢复了 {len(chunk_ids)} 个被过滤的块")

            return jsonify({
                'success': True,
                'restored_count': len(chunk_ids)
            })

        except Exception as e:
            logger.error(f"恢复块失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    # ============================================
    # 步骤3：可编辑表格相关API
    # ============================================

    @app.route('/api/tender-processing/requirements/<int:project_id>', methods=['GET'])
    def get_task_requirements(project_id):
        """
        获取HITL任务的所有要求及统计信息（步骤3）

        注意：现在按 project_id 过滤
        这确保只显示当前项目中选中章节提取的需求

        返回：
        {
            "success": True,
            "requirements": [{...}, ...],
            "summary": [
                {"constraint_type": "mandatory", "count": 10},
                {"constraint_type": "optional", "count": 5},
                {"constraint_type": "scoring", "count": 3}
            ]
        }
        """
        try:
            db = get_knowledge_base_db()

            # 获取当前项目的所有要求
            requirements = db.execute_query("""
                SELECT * FROM tender_requirements
                WHERE project_id = ?
                ORDER BY requirement_id
            """, (project_id,))

            # 统计各类型数量
            summary = db.execute_query("""
                SELECT
                    constraint_type,
                    COUNT(*) as count
                FROM tender_requirements
                WHERE project_id = ?
                GROUP BY constraint_type
            """, (project_id,))

            logger.info(f"加载HITL任务 {project_id} 的 {len(requirements)} 个要求")

            return jsonify({
                'success': True,
                'requirements': requirements,
                'summary': summary
            })

        except Exception as e:
            logger.error(f"获取任务要求失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/requirements/<int:req_id>', methods=['PATCH'])
    def update_requirement(req_id):
        """
        编辑单个要求（步骤3）

        请求参数（JSON）：
        {
            "constraint_type": "mandatory",
            "category": "technical",
            "detail": "更新后的内容",
            ...
        }

        返回：
        {
            "success": True,
            "requirement_id": 123
        }
        """
        try:
            data = request.get_json()
            project_id = data.get('project_id')  # 可选

            if not data:
                return jsonify({'success': False, 'error': '缺少更新数据'}), 400

            db = get_knowledge_base_db()

            # 构建更新SQL
            update_fields = []
            params = []

            allowed_fields = [
                'constraint_type', 'category', 'subcategory',
                'detail', 'source_location', 'priority'
            ]

            for field in allowed_fields:
                if field in data:
                    update_fields.append(f"{field} = ?")
                    params.append(data[field])

            if not update_fields:
                return jsonify({'success': False, 'error': '没有可更新的字段'}), 400

            params.append(req_id)

            # 执行更新
            db.execute_query(f"""
                UPDATE tender_requirements
                SET {', '.join(update_fields)}, updated_at = CURRENT_TIMESTAMP
                WHERE requirement_id = ?
            """, tuple(params))

            # 如果提供了project_id，记录到草稿表
            if project_id:
                db.execute_query("""
                    INSERT INTO tender_requirements_draft (
                        requirement_id, project_id,
                        constraint_type, category, subcategory, detail, source_location, priority,
                        operation, edited_by, edited_at
                    ) SELECT
                        ?, project_id, constraint_type, category, subcategory,
                        detail, source_location, priority, 'edit', 'user', CURRENT_TIMESTAMP
                    FROM tender_requirements
                    WHERE requirement_id = ?
                """, (req_id, req_id))

            logger.info(f"要求 {req_id} 已更新")

            return jsonify({
                'success': True,
                'requirement_id': req_id
            })

        except Exception as e:
            logger.error(f"更新要求失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/requirements/batch', methods=['POST'])
    def batch_requirement_operations():
        """
        批量操作要求（步骤3）

        请求参数（JSON）：
        {
            "task_id": "hitl_xxx",
            "operations": [
                {"action": "add", "data": {...}},
                {"action": "delete", "requirement_id": 123},
                ...
            ]
        }

        返回：
        {
            "success": True,
            "results": [...]
        }
        """
        try:
            data = request.get_json()
            project_id = data.get('project_id')
            operations = data.get('operations', [])

            if not project_id or not operations:
                return jsonify({'success': False, 'error': '参数不完整'}), 400

            db = get_knowledge_base_db()
            results = []

            for op in operations:
                action = op.get('action')

                if action == 'add':
                    # 添加新要求
                    req_data = op.get('data', {})
                    db.execute_query("""
                        INSERT INTO tender_requirements (
                            project_id, constraint_type, category, subcategory,
                            detail, source_location, priority
                        ) VALUES (?, ?, ?, ?, ?, ?, ?)
                    """, (
                        project_id,
                        req_data.get('constraint_type'),
                        req_data.get('category'),
                        req_data.get('subcategory'),
                        req_data.get('detail'),
                        req_data.get('source_location'),
                        req_data.get('priority', 'medium')
                    ))
                    results.append({'action': 'add', 'success': True})

                elif action == 'delete':
                    # 删除要求
                    req_id = op.get('requirement_id')
                    db.execute_query("""
                        DELETE FROM tender_requirements
                        WHERE requirement_id = ?
                    """, (req_id,))
                    results.append({'action': 'delete', 'requirement_id': req_id, 'success': True})

            return jsonify({
                'success': True,
                'results': results
            })

        except Exception as e:
            logger.error(f"批量操作失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/export-draft/<int:project_id>', methods=['GET'])
    def export_draft_requirements(project_id):
        """
        导出草稿要求（步骤3）

        返回：Excel 文件
        """
        try:
            import pandas as pd
            from io import BytesIO

            db = get_knowledge_base_db()

            # 获取草稿要求
            drafts = db.execute_query("""
                SELECT * FROM tender_requirements_draft
                WHERE project_id = ?
                ORDER BY created_at DESC
            """, (project_id,))

            if not drafts:
                return jsonify({'success': False, 'error': '没有草稿数据'}), 404

            # 转换为 DataFrame
            df = pd.DataFrame(drafts)

            # 生成 Excel
            output = BytesIO()
            df.to_excel(output, index=False, sheet_name='草稿要求')
            output.seek(0)

            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'requirements_draft_{project_id}.xlsx'
            )

        except Exception as e:
            logger.error(f"导出草稿失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    # ============================================
    # 步骤3增强：基本信息和资质信息提取API
    # ============================================

    @app.route('/api/tender-processing/extract-basic-info/<int:project_id>', methods=['POST'])
    def extract_basic_info_step3(project_id):
        """
        提取基本信息（步骤3）

        使用现有的 TenderInfoExtractor 提取项目基本信息

        请求参数（JSON，可选）：
        {
            "model_name": "yuanjing-deepseek-v3"  # 可选，默认使用联通元景模型
        }

        返回：
        {
            "success": True,
            "data": {
                "project_name": "...",
                "project_number": "...",
                "tender_party": "...",
                ...
            }
        }
        """
        try:
            from modules.tender_info.extractor import TenderInfoExtractor
            from common import get_config

            db = get_knowledge_base_db()

            # 获取请求参数（支持前端传递模型选择）
            request_data = request.get_json() if request.is_json else {}
            model_name = request_data.get('model_name', 'yuanjing-deepseek-v3')  # 默认使用联通元景模型

            logger.info(f"基本信息提取 - 任务ID: {project_id}, 使用模型: {model_name}")

            # 获取任务信息
            hitl_task = db.execute_query("""
                SELECT project_id, step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not hitl_task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 获取文档路径
            step1_data = json.loads(hitl_task['step1_data'])
            doc_path = step1_data.get('file_path')

            # 使用智能路径解析（兼容阿里云/本地/Docker等多种环境）
            logger.info(f"=== 路径解析调试信息(extract-basic-info) ===")
            logger.info(f"原始路径: {doc_path}")
            resolved_path = resolve_file_path(doc_path, logger=logger)
            if not resolved_path:
                logger.error(f"文档路径解析失败: {doc_path}")
                return jsonify({'success': False, 'error': f'文档路径解析失败: {doc_path}'}), 404

            doc_path = str(resolved_path)  # 转为字符串供extractor使用
            logger.info(f"解析后路径: {doc_path}")
            logger.info(f"=== 路径解析调试信息结束 ===")

            # 创建提取器 - 使用指定的模型
            config = get_config()
            extractor = TenderInfoExtractor(model_name=model_name)

            # 读取文档内容（用于向后兼容，新方法使用章节识别）
            # text = extractor.read_document(doc_path)

            # 提取基本信息 - 使用章节识别方法
            basic_info = extractor.extract_basic_info(project_id=project_id)

            # 🔧 字段名映射：将数据库列名映射为前端期望的字段名
            # 保持向后兼容，同时提供两套字段名
            response_data = {
                **basic_info,  # 保留原始字段名
                # 添加前端期望的字段名别名
                'tender_party': basic_info.get('tenderer'),
                'tender_agent': basic_info.get('agency'),
                'tender_method': basic_info.get('bidding_method'),
                'tender_location': basic_info.get('bidding_location'),
                'tender_deadline': basic_info.get('bidding_time'),
                # budget_amount 已经在basic_info中，无需额外映射
            }

            logger.info(f"基本信息提取成功: {project_id}, tenderer={basic_info.get('tenderer')}, product_category={basic_info.get('product_category')}")

            return jsonify({
                'success': True,
                'data': response_data
            })

        except Exception as e:
            logger.error(f"提取基本信息失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/save-basic-info/<int:project_id>', methods=['POST'])
    def save_basic_info_step3(project_id):
        """
        保存基本信息（步骤3）

        将提取的基本信息保存到 tender_projects 表

        请求参数（JSON）：
        {
            "task_id": "hitl_xxx",
            "project_name": "...",
            "project_number": "...",
            ...
        }

        返回：
        {
            "success": True,
            "project_id": 123
        }
        """
        try:
            data = request.get_json()
            project_id = data.get('project_id')

            if not data:
                return jsonify({'success': False, 'error': '缺少数据'}), 400

            db = get_knowledge_base_db()

            # 字段名映射：支持前端字段名和数据库列名两种方式
            # LLM返回的是数据库列名，前端可能使用不同的字段名
            field_mapping = {
                'tender_party': 'tenderer',
                'tender_agent': 'agency',
                'tender_method': 'bidding_method',
                'tender_location': 'bidding_location',
                'tender_deadline': 'bidding_time',
            }

            # 获取字段值，优先使用数据库列名，其次使用前端字段名
            def get_field_value(field_name):
                db_column = field_mapping.get(field_name, field_name)
                # 先尝试数据库列名，再尝试前端字段名
                return data.get(db_column, data.get(field_name, ''))

            # 检查项目是否已存在
            existing_project = db.execute_query("""
                SELECT project_id FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if existing_project:
                # 处理产品分类信息
                product_category = data.get('product_category')
                product_items = data.get('product_items', [])
                product_category_id = None
                product_category_name = None

                # 根据分类名称查找分类ID
                if product_category:
                    category_result = db.execute_query(
                        "SELECT category_id, category_name FROM product_categories WHERE category_name = ? AND is_active = 1",
                        [product_category],
                        fetch_one=True
                    )
                    if category_result:
                        product_category_id = category_result['category_id']
                        product_category_name = category_result['category_name']
                        logger.info(f"找到产品分类: {product_category_name} (ID: {product_category_id})")

                # 序列化 product_items
                product_items_json = json.dumps(product_items, ensure_ascii=False) if product_items else None

                # 更新现有项目
                db.execute_query("""
                    UPDATE tender_projects
                    SET project_name = ?,
                        project_number = ?,
                        tenderer = ?,
                        agency = ?,
                        bidding_method = ?,
                        bidding_location = ?,
                        bidding_time = ?,
                        winner_count = ?,
                        authorized_person_name = ?,
                        authorized_person_id = ?,
                        authorized_person_position = ?,
                        product_category_id = ?,
                        product_category_name = ?,
                        product_items = ?,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE project_id = ?
                """, (
                    data.get('project_name', ''),
                    data.get('project_number', ''),
                    get_field_value('tender_party'),
                    get_field_value('tender_agent'),
                    get_field_value('tender_method'),
                    get_field_value('tender_location'),
                    get_field_value('tender_deadline'),
                    data.get('winner_count', ''),
                    data.get('authorized_person_name', ''),
                    data.get('authorized_person_id', ''),
                    data.get('authorized_person_position', ''),
                    product_category_id,
                    product_category_name,
                    product_items_json,
                    project_id
                ))

                logger.info(f"项目基本信息已更新: {project_id}, 产品分类: {product_category_name}")
            else:
                # 创建新项目（理论上不应该走这个分支，因为项目应该已经创建）
                logger.warning(f"项目 {project_id} 不存在，跳过保存")
                return jsonify({'success': False, 'error': '项目不存在'}), 404

            return jsonify({
                'success': True,
                'project_id': project_id
            })

        except Exception as e:
            logger.error(f"保存基本信息失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/extract-qualifications/<int:project_id>', methods=['POST'])
    def extract_qualifications_step3(project_id):
        """
        提取资质要求并对比公司状态（步骤3）

        使用关键词匹配提取资质要求，并对比公司已上传的资质

        请求参数（JSON）：
        {
            "project_id": 123,
            "model_name": "yuanjing-deepseek-v3"  # 可选，默认使用联通元景模型
        }

        返回：
        {
            "success": True,
            "data": {
                "qualifications": {...},
                "summary": {
                    "required_count": 10,
                    "uploaded_count": 8,
                    "missing_count": 2
                }
            }
        }
        """
        try:
            from modules.tender_info.extractor import TenderInfoExtractor
            from common import get_config

            data = request.get_json() or {}
            # project_id already comes from URL path parameter, don't overwrite it
            model_name = data.get('model_name', 'yuanjing-deepseek-v3')  # 默认使用联通元景模型

            logger.info(f"资质提取 - 任务ID: {project_id}, 使用模型: {model_name}")

            db = get_knowledge_base_db()

            # 获取任务信息
            hitl_task = db.execute_query("""
                SELECT project_id, step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not hitl_task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 获取文档路径
            step1_data = json.loads(hitl_task['step1_data'])
            doc_path = step1_data.get('file_path')

            # 使用智能路径解析（兼容阿里云/本地/Docker等多种环境）
            logger.info(f"=== 路径解析调试信息(extract-qualifications) ===")
            logger.info(f"原始路径: {doc_path}")
            resolved_path = resolve_file_path(doc_path, logger=logger)
            if not resolved_path:
                logger.error(f"文档路径解析失败: {doc_path}")
                return jsonify({'success': False, 'error': f'文档路径解析失败: {doc_path}'}), 404

            doc_path = str(resolved_path)  # 转为字符串供extractor使用
            logger.info(f"解析后路径: {doc_path}")
            logger.info(f"=== 路径解析调试信息结束 ===")

            # 创建提取器 - 使用指定的模型
            config = get_config()
            extractor = TenderInfoExtractor(model_name=model_name)

            # 读取文档内容
            text = extractor.read_document(doc_path)

            # 提取资质要求（使用关键词匹配）
            qualifications = extractor.extract_qualification_requirements(text)

            # 获取公司ID（如果存在）
            company_id = None
            if project_id:
                project_info = db.execute_query("""
                    SELECT company_id FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if project_info:
                    company_id = project_info.get('company_id')

            # 整合公司资质状态
            if company_id:
                # 导入资质对比函数
                from web.blueprints.api_tender_bp import enrich_qualification_with_company_status
                enriched_data = enrich_qualification_with_company_status(qualifications, company_id)
            else:
                # 没有公司ID，只返回提取的资质要求
                enriched_data = qualifications
                # 添加空的公司状态
                for key in enriched_data.get('qualifications', {}):
                    enriched_data['qualifications'][key]['company_status'] = {
                        'uploaded': False,
                        'original_filename': None,
                        'upload_time': None
                    }

                # 添加汇总信息
                required_count = len([q for q in enriched_data.get('qualifications', {}).values()
                                     if q.get('required')])
                enriched_data['summary'] = {
                    'required_count': required_count,
                    'uploaded_count': 0,
                    'missing_count': required_count
                }

            # 保存资质数据到数据库
            if project_id:
                qualifications_json = json.dumps(
                    enriched_data.get('qualifications', {}),
                    ensure_ascii=False  # 保留中文字符
                )

                db.execute_query("""
                    UPDATE tender_projects
                    SET qualifications_data = ?,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE project_id = ?
                """, (qualifications_json, project_id))

                logger.info(f"资质数据已保存到项目 {project_id}: {len(enriched_data.get('qualifications', {}))}项")

            logger.info(f"资质要求提取成功: {project_id}, 检测到{len(enriched_data.get('qualifications', {}))}项资质")

            return jsonify({
                'success': True,
                'data': enriched_data
            })

        except Exception as e:
            logger.error(f"提取资质要求失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/trigger-extraction/<int:project_id>', methods=['POST'])
    def trigger_requirement_extraction(project_id):
        """
        手动触发要求提取（用于已有任务补充提取）

        对指定的HITL任务，从已有的有价值文本块中提取详细要求
        """
        try:
            db = get_knowledge_base_db()

            # 检查任务是否存在
            task_info = db.execute_query("""
                SELECT project_id FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_info:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 先清除该任务的旧要求
            db.execute_query("""
                DELETE FROM tender_requirements
                WHERE project_id = ?
            """, (project_id,))

            # 提取要求
            from modules.tender_processing.requirement_extractor import RequirementExtractor
            extractor = RequirementExtractor()

            # 查询有价值的文本块
            valuable_chunks = db.execute_query("""
                SELECT chunk_id, content, metadata
                FROM tender_document_chunks
                WHERE project_id = ? AND is_valuable = 1
                ORDER BY chunk_index
            """, (project_id,))

            if not valuable_chunks:
                return jsonify({
                    'success': False,
                    'error': '没有找到有价值的文本块，请先完成步骤1的章节选择'
                }), 400

            total_extracted = 0
            for chunk in valuable_chunks:
                # 提取要求（使用extract_chunk方法，需要传入字典格式）
                chunk_dict = {
                    'content': chunk['content'],
                    'chunk_type': chunk.get('chunk_type', 'paragraph'),
                    'metadata': json.loads(chunk['metadata']) if chunk.get('metadata') else {}
                }
                requirements = extractor.extract_chunk(chunk_dict)

                # 解析metadata获取章节标题
                metadata = json.loads(chunk['metadata']) if chunk.get('metadata') else {}
                chapter_title = metadata.get('chapter_title', '未知章节')

                # 保存到数据库
                for req in requirements:
                    db.execute_query("""
                        INSERT INTO tender_requirements (
                            project_id, constraint_type, category,
                            subcategory, detail, summary, source_location, priority,
                            extraction_confidence
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        project_id,
                        req.constraint_type,
                        req.category,
                        req.subcategory,
                        req.detail,
                        req.summary,
                        f"{chapter_title} - {req.source_location}" if req.source_location else chapter_title,
                        req.priority,
                        req.extraction_confidence
                    ))
                    total_extracted += 1

            logger.info(f"手动触发提取完成: 从 {len(valuable_chunks)} 个块中提取了 {total_extracted} 个要求")

            return jsonify({
                'success': True,
                'chunks_processed': len(valuable_chunks),
                'requirements_extracted': total_extracted
            })

        except Exception as e:
            logger.error(f"手动触发提取失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    # ============================================
    # 步骤3：保存和完成相关API
    # ============================================

    @app.route('/api/tender-processing/save-basic-info/<int:project_id>', methods=['POST'])
    def save_basic_info(project_id):
        """保存基本信息到step3_data"""
        try:
            data = request.get_json()
            basic_info = data.get('basic_info', {})

            # 获取当前step3_data和step1_data
            task_data = db.execute_query("""
                SELECT step3_data, step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"success": False, "error": "任务不存在"}), 404

            # 解析现有step3_data
            step3_data = json.loads(task_data['step3_data']) if task_data['step3_data'] else {}

            # 解析step1_data，获取应答文件信息
            step1_data = json.loads(task_data['step1_data']) if task_data['step1_data'] else {}
            response_file = step1_data.get('response_file')

            # 更新基本信息
            step3_data['basic_info'] = basic_info
            step3_data['basic_info_saved_at'] = datetime.now().isoformat()

            # 如果有应答文件信息，也保存到step3_data
            if response_file:
                step3_data['response_file'] = response_file

            # 保存回数据库
            db.execute_query("""
                UPDATE tender_projects
                SET step3_data = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE project_id = ?
            """, (json.dumps(step3_data, ensure_ascii=False), project_id))

            logger.info(f"✅ 保存基本信息成功: {project_id}")
            return jsonify({
                "success": True,
                "message": "基本信息保存成功"
            })

        except Exception as e:
            logger.error(f"保存基本信息失败: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route('/api/tender-processing/complete-hitl/<int:project_id>', methods=['POST'])
    def complete_hitl(project_id):
        """完成HITL流程"""
        try:
            # 更新任务状态
            db.execute_query("""
                UPDATE tender_projects
                SET step3_status = 'completed',
                    step3_completed_at = CURRENT_TIMESTAMP,
                    overall_status = 'completed',
                    updated_at = CURRENT_TIMESTAMP
                WHERE project_id = ?
            """, (project_id,))

            logger.info(f"✅ HITL流程完成: {project_id}")
            return jsonify({
                "success": True,
                "message": "HITL流程完成"
            })

        except Exception as e:
            logger.error(f"完成HITL流程失败: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({"success": False, "error": str(e)}), 500


# ============================================
# 辅助函数
# ============================================

    def _save_chapters_to_db(db, chapters, project_id, parent_id=None):
        """递归保存章节树到数据库"""
        chapter_ids = []

        for chapter in chapters:
            # 插入章节
            chapter_db_id = db.execute_query("""
                INSERT INTO tender_document_chapters (
                    project_id, chapter_node_id, level, title,
                    para_start_idx, para_end_idx, word_count, preview_text,
                    is_selected, auto_selected, skip_recommended, parent_chapter_id
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                project_id, chapter['id'], chapter['level'], chapter['title'],
                chapter['para_start_idx'], chapter.get('para_end_idx'),
                chapter.get('word_count', 0), chapter.get('preview_text', ''),
                False, chapter.get('auto_selected', False),
                chapter.get('skip_recommended', False), parent_id
            ))

            # execute_query 已经返回 lastrowid
            chapter_ids.append(chapter_db_id)

            # 递归保存子章节
            if chapter.get('children'):
                child_ids = _save_chapters_to_db(
                    db, chapter['children'], project_id, chapter_db_id
                )
                chapter_ids.extend(child_ids)

        return chapter_ids


        # ============================================
        # 数据同步相关API（新增）
        # ============================================

    @app.route('/api/tender-processing/requirements/<int:project_id>', methods=['GET'])
    def get_project_requirements(project_id):
        """
        获取项目的需求数据（支持按类别筛选）

        Query参数：
        - category: 可选，筛选类别 (qualification/technical/commercial/service)
        """
        try:
            category = request.args.get('category')

            query = """
                SELECT * FROM tender_requirements
                WHERE project_id = ?
            """
            params = [project_id]

            if category:
                query += " AND category = ?"
                params.append(category)

            query += " ORDER BY created_at DESC"

            requirements = db.execute_query(query, params)

            logger.info(f"查询项目 {project_id} 的需求数据，类别: {category or '全部'}，找到 {len(requirements)} 条")

            return jsonify({
                'success': True,
                'requirements': requirements,
                'has_extracted': len(requirements) > 0,
                'count': len(requirements)
            })

        except Exception as e:
            logger.error(f"获取项目需求失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500


    @app.route('/api/tender-processing/hitl-tasks', methods=['GET'])
    def get_hitl_tasks():
        """
        查询HITL任务（支持按项目ID筛选，返回最新任务）

        Query参数：
        - project_id: 必选，项目ID
        - latest: 可选，是否只返回最新任务 (true/false)
        """
        try:
            project_id = request.args.get('project_id', type=int)
            latest = request.args.get('latest', 'false').lower() == 'true'

            if not project_id:
                return jsonify({'success': False, 'error': '缺少project_id参数'}), 400

            query = """
                SELECT * FROM tender_projects
                WHERE project_id = ?
                ORDER BY created_at DESC
            """

            if latest:
                query += " LIMIT 1"

            tasks = db.execute_query(query, [project_id])

            logger.info(f"查询项目 {project_id} 的HITL任务，找到 {len(tasks)} 个")

            if latest and tasks:
                return jsonify({
                    'success': True,
                    'task': tasks[0]
                })

            return jsonify({
                'success': True,
                'tasks': tasks,
                'count': len(tasks)
            })

        except Exception as e:
            logger.error(f"查询HITL任务失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500


    @app.route('/api/tender-processing/hitl-tasks/<int:project_id>', methods=['GET'])
    def get_hitl_task_by_id(project_id):
        """
        根据 project_id 获取单个HITL任务详情

        URL参数：
        - project_id: 项目ID

        返回：
        - success: 是否成功
        - task: 任务详情对象（包含step1_data, step2_data, step3_data等）
        """
        try:
            logger.info(f"获取HITL任务详情: {project_id}")

            query = """
                SELECT * FROM tender_projects
                WHERE project_id = ?
            """

            tasks = db.execute_query(query, [project_id])

            if not tasks or len(tasks) == 0:
                logger.warning(f"未找到HITL任务: {project_id}")
                return jsonify({
                    'success': False,
                    'error': f'未找到ID为 {project_id} 的HITL任务'
                }), 404

            task = tasks[0]
            logger.info(f"成功获取HITL任务: {project_id}")

            return jsonify({
                'success': True,
                'task': task
            })

        except Exception as e:
            logger.error(f"获取HITL任务详情失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500


    def sync_requirements_to_project(project_id, category='qualification'):
        """
        将需求明细同步到项目汇总字段

        Args:
            project_id: 项目ID
            category: 需求类别 (qualification/technical/commercial/service)
        """
        try:
            logger.info(f"开始同步项目 {project_id} 的 {category} 数据到项目表")

            # 1. 查询该类别的所有需求
            requirements = db.execute_query("""
                SELECT * FROM tender_requirements
                WHERE project_id = ? AND category = ?
                ORDER BY created_at
            """, (project_id, category))

            if not requirements:
                logger.info(f"项目 {project_id} 没有 {category} 类别的需求数据，跳过同步")
                return

            # 2. 转换为JSON格式
            aggregated_data = {}
            for req in requirements:
                # 使用 subcategory 作为key，如果没有则用 summary 或 requirement_id
                key = req.get('subcategory') or req.get('summary') or f"requirement_{req['requirement_id']}"

                aggregated_data[key] = {
                    'requirement_id': req['requirement_id'],
                    'constraint_type': req.get('constraint_type'),
                    'detail': req.get('detail'),
                    'summary': req.get('summary'),
                    'source_location': req.get('source_location'),
                    'priority': req.get('priority'),
                    'extraction_confidence': req.get('extraction_confidence'),
                    'is_verified': req.get('is_verified', False),
                    'created_at': req.get('created_at')
                }

            # 3. 确定更新的字段名
            if category == 'qualification':
                field_name = 'qualifications_data'
            elif category == 'technical':
                field_name = 'scoring_data'  # 暂时复用 scoring_data 字段
            else:
                logger.warning(f"未知的类别 {category}，使用 scoring_data 字段")
                field_name = 'scoring_data'

            # 4. 更新项目表
            db.execute_query(f"""
                UPDATE tender_projects
                SET {field_name} = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE project_id = ?
            """, (json.dumps(aggregated_data, ensure_ascii=False), project_id))

            logger.info(f"✅ 已同步 {len(requirements)} 条 {category} 数据到项目表的 {field_name} 字段")

        except Exception as e:
            logger.error(f"同步需求数据到项目表失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise

    # ============================================
    # 填充应答文件API
    # ============================================
    # ============================================
    # 统一文件同步API (支持多种文件类型)
    # ============================================

    # 文件同步配置
    HITL_FILE_SYNC_CONFIG = {
        'business_response': {
            'dir_name': 'completed_response_files',
            'field_name': 'business_response_file',
            'suffix': '_应答完成',
            'display_name': '商务应答文件'
        },
        'point_to_point': {
            'dir_name': 'point_to_point_files',
            'field_name': 'technical_point_to_point_file',
            'suffix': '_点对点应答',
            'display_name': '点对点应答文件'
        },
        'tech_proposal': {
            'dir_name': 'tech_proposal_files',
            'field_name': 'technical_proposal_file',
            'suffix': '_技术方案',
            'display_name': '技术方案文件'
        }
    }

    def _find_file_in_filesystem(project_id, field_name):
        """
        从文件系统中查找文件(fallback机制,当数据库中没有数据时使用)

        Args:
            project_id: 项目ID
            field_name: 字段名(如 'business_response_file')

        Returns:
            文件路径的Path对象,如果找不到返回None
        """
        try:
            # 根据field_name确定目录(支持多目录搜索)
            field_to_dirs = {
                'business_response_file': ['completed_response_files'],
                'technical_point_to_point_file': ['point_to_point_files'],
                'technical_proposal_file': ['tech_proposal_files'],
                # response_file 需要搜索两个目录: 原始模板目录 和 完成文件目录
                'response_file': ['response_files', 'completed_response_files'],
                # technical_file 在技术需求文件目录中(三级目录结构: 年/月/任务ID)
                'technical_file': ['technical_files'],
                'technical_file_path': ['technical_files']  # 兼容字段名
            }

            dir_names = field_to_dirs.get(field_name)
            if not dir_names:
                logger.warning(f"未知的字段名: {field_name}")
                return None

            # 构建基础目录
            project_root = Path(__file__).parent.parent

            # 遍历所有可能的目录
            for dir_name in dir_names:
                base_dir = project_root / 'data' / 'uploads' / dir_name

                if not base_dir.exists():
                    logger.debug(f"目录不存在,跳过: {base_dir}")
                    continue

                # tender_processing 目录使用扁平结构(年/月/文件),其他使用三级结构(年/月/任务ID/)
                if dir_name == 'tender_processing':
                    # 扁平结构: 遍历 year/month/ 查找包含task_id的文件
                    for year_dir in base_dir.iterdir():
                        if not year_dir.is_dir():
                            continue
                        for month_dir in year_dir.iterdir():
                            if not month_dir.is_dir():
                                continue
                            # 查找文件名中包含task_id或者匹配招标文件的docx文件
                            for file_path in month_dir.iterdir():
                                if file_path.suffix.lower() in ['.docx', '.doc']:
                                    logger.info(f"从文件系统找到技术文件: {file_path}")
                                    return file_path
                else:
                    # 三级结构: year/month/task_id/文件
                    for year_dir in base_dir.iterdir():
                        if not year_dir.is_dir():
                            continue
                        for month_dir in year_dir.iterdir():
                            if not month_dir.is_dir():
                                continue
                            task_dir = month_dir / str(project_id)
                            if task_dir.exists() and task_dir.is_dir():
                                # 查找docx文件
                                for file_path in task_dir.iterdir():
                                    if file_path.suffix.lower() == '.docx':
                                        logger.info(f"从文件系统找到文件: {file_path} (目录: {dir_name})")
                                        return file_path

            logger.warning(f"文件系统中未找到文件: task_id={project_id}, field_name={field_name}, 搜索目录={dir_names}")
            return None

        except Exception as e:
            logger.error(f"文件系统扫描失败: {str(e)}")
            return None

    @app.route('/api/tender-processing/sync-file/<int:project_id>', methods=['POST'])
    def sync_file_to_hitl(project_id):
        """
        统一的文件同步API - 支持多种文件类型

        请求体:
        {
            "file_path": "源文件路径",
            "file_type": "business_response" | "point_to_point" | "tech_proposal"
        }
        """
        try:
            data = request.get_json()
            source_file_path = data.get('file_path')
            file_type = data.get('file_type')

            # 验证参数
            if not source_file_path:
                return jsonify({"success": False, "error": "未提供文件路径"}), 400

            if not file_type:
                return jsonify({"success": False, "error": "未提供文件类型"}), 400

            # 获取文件类型配置
            config = HITL_FILE_SYNC_CONFIG.get(file_type)
            if not config:
                return jsonify({
                    "success": False,
                    "error": f"不支持的文件类型: {file_type}。支持的类型: {', '.join(HITL_FILE_SYNC_CONFIG.keys())}"
                }), 400

            # 检查源文件是否存在
            if not os.path.exists(source_file_path):
                return jsonify({"success": False, "error": "源文件不存在"}), 404

            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"success": False, "error": "任务不存在"}), 404

            # 解析step1_data，如果为空则初始化为空字典
            step1_data_raw = task_data.get('step1_data') or task_data.get('step1_data', '{}')
            if not step1_data_raw or step1_data_raw == '':
                step1_data_raw = '{}'
            step1_data = json.loads(step1_data_raw)

            # 创建存储目录
            now = datetime.now()
            project_root = Path(__file__).parent.parent
            save_dir = os.path.join(
                project_root,
                f'data/uploads/{config["dir_name"]}',
                str(now.year),
                f"{now.month:02d}",
                str(project_id)
            )
            os.makedirs(save_dir, exist_ok=True)

            # 生成文件名
            source_filename = os.path.basename(source_file_path)
            if '_' in source_filename:
                base_name = source_filename.rsplit('.', 1)[0]
                filename = f"{base_name}{config['suffix']}.docx"
            else:
                filename = f"{config['suffix']}_{now.strftime('%Y%m%d_%H%M%S')}.docx"

            # 复制文件到目标位置
            target_path = os.path.join(save_dir, filename)
            shutil.copy2(source_file_path, target_path)

            # 计算文件大小
            file_size = os.path.getsize(target_path)

            # 更新任务的step1_data
            file_info = {
                "file_path": target_path,
                "filename": filename,
                "file_size": file_size,
                "saved_at": now.isoformat(),
                "source_file": source_file_path
            }
            step1_data[config['field_name']] = file_info

            db.execute_query("""
                UPDATE tender_projects
                SET step1_data = ?
                WHERE project_id = ?
            """, (json.dumps(step1_data), project_id))

            logger.info(f"同步{config['display_name']}到HITL任务: {project_id}, 文件: {filename} ({file_size} bytes)")

            return jsonify({
                "success": True,
                "message": f"{config['display_name']}已成功同步到投标项目",
                "file_path": target_path,
                "filename": filename,
                "file_size": file_size,
                "file_type": file_type,
                "saved_at": file_info['saved_at']
            })

        except Exception as e:
            logger.error(f"同步文件失败: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({"success": False, "error": str(e)}), 500

    # ============================================
    # 应答完成文件相关API (从商务应答同步)
    # ============================================

    @app.route('/api/tender-processing/sync-business-response/<int:project_id>', methods=['POST'])
    def sync_business_response_to_hitl(project_id):
        """
        将商务应答生成的文件同步到HITL投标项目（向后兼容API）
        内部调用统一的sync_file_to_hitl API
        """
        # 构造统一API所需的请求数据
        data = request.get_json()
        unified_data = {
            'file_path': data.get('file_path'),
            'file_type': 'business_response'
        }

        # 临时替换request数据
        from flask import g
        original_json = request.get_json
        request.get_json = lambda: unified_data

        # 调用统一API
        result = sync_file_to_hitl(project_id)

        # 恢复原始request
        request.get_json = original_json

        return result

    @app.route('/api/tender-processing/completed-response-info/<int:project_id>', methods=['GET'])
    def get_completed_response_info(project_id):
        """获取应答完成文件信息"""
        try:
            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"success": False, "error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            completed_response = step1_data.get('completed_response_file')

            if not completed_response:
                return jsonify({
                    "success": True,
                    "has_file": False
                })

            # 检查文件是否存在
            file_exists = os.path.exists(completed_response['file_path'])

            return jsonify({
                "success": True,
                "has_file": file_exists,
                "filename": completed_response.get('filename'),
                "file_size": completed_response.get('file_size'),
                "saved_at": completed_response.get('saved_at'),
                "source": completed_response.get('source', 'unknown'),
                "download_url": f"/api/tender-processing/download-completed-response/{project_id}"
            })

        except Exception as e:
            logger.error(f"获取应答完成文件信息失败: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route('/api/tender-processing/download-completed-response/<int:project_id>', methods=['GET'])
    def download_completed_response(project_id):
        """下载应答完成文件"""
        try:
            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            completed_response = step1_data.get('completed_response_file')

            if not completed_response:
                return jsonify({"error": "应答完成文件不存在"}), 404

            file_path = completed_response['file_path']
            if not os.path.exists(file_path):
                return jsonify({"error": "文件已被删除"}), 404

            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=completed_response['filename']
            )

        except Exception as e:
            logger.error(f"下载应答完成文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/tender-processing/preview-completed-response/<int:project_id>', methods=['GET'])
    def preview_completed_response(project_id):
        """预览应答完成文件"""
        try:
            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_projects
                WHERE project_id = ?
            """, (project_id,), fetch_one=True)

            if not task_data:
                return jsonify({"error": "任务不存在"}), 404

            step1_data = json.loads(task_data['step1_data'])
            completed_response = step1_data.get('completed_response_file')

            if not completed_response:
                return jsonify({"error": "应答完成文件不存在"}), 404

            file_path = completed_response['file_path']
            if not os.path.exists(file_path):
                return jsonify({"error": "文件已被删除"}), 404

            # 预览模式：as_attachment=False，浏览器会尝试在线打开
            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=False,
                download_name=completed_response['filename']
            )

        except Exception as e:
            logger.error(f"预览应答完成文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    # ============================================
    # 新的三种文件类型的获取API
    # ============================================

    def _get_file_info(project_id, field_name, file_display_name):
        """
        通用的文件信息获取函数

        Args:
            project_id: 项目ID
            field_name: step1_data中的字段名（如 'technical_point_to_point_file'）
            file_display_name: 文件显示名称（用于日志和错误提示）

        Returns:
            JSON响应
        """
        try:
            file_info = None

            # 尝试从数据库获取文件信息
            try:
                task_data = db.execute_query("""
                    SELECT step1_data FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if task_data:
                    step1_data = json.loads(task_data['step1_data'])
                    file_info = step1_data.get(field_name)
            except Exception as db_error:
                logger.warning(f"数据库查询失败,将尝试文件系统扫描: {str(db_error)}")

            # 如果数据库中没有文件信息,尝试从文件系统扫描
            if not file_info:
                logger.info(f"数据库中无{file_display_name}信息,尝试从文件系统扫描")
                file_path = _find_file_in_filesystem(project_id, field_name)

                if file_path:
                    # 构建文件信息
                    file_stat = file_path.stat()
                    file_info = {
                        'file_path': str(file_path),
                        'filename': file_path.name,
                        'file_size': file_stat.st_size,
                        'saved_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                    }
                    logger.info(f"从文件系统找到{file_display_name}: {file_path.name}")

            if not file_info:
                return jsonify({
                    "success": True,
                    "has_file": False
                })

            # 检查文件是否存在
            file_exists = os.path.exists(file_info.get('file_path', ''))

            # 如果数据库中的文件路径不存在,尝试从文件系统扫描
            if not file_exists and 'file_path' in file_info:
                logger.warning(f"数据库中的文件路径不存在: {file_info['file_path']}, 尝试文件系统扫描")
                file_path = _find_file_in_filesystem(project_id, field_name)

                if file_path:
                    # 更新文件信息
                    file_stat = file_path.stat()
                    file_info = {
                        'file_path': str(file_path),
                        'filename': file_path.name,
                        'file_size': file_stat.st_size,
                        'saved_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                    }
                    file_exists = True
                    logger.info(f"从文件系统找到替代文件: {file_path.name}")

            return jsonify({
                "success": True,
                "has_file": file_exists,
                "filename": file_info.get('filename'),
                "file_size": file_info.get('file_size'),
                "saved_at": file_info.get('saved_at'),
                "download_url": f"/api/tender-processing/download-file/{project_id}/{field_name}"
            })

        except Exception as e:
            logger.error(f"获取{file_display_name}信息失败: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route('/api/tender-processing/technical-point-to-point-info/<int:project_id>', methods=['GET'])
    def get_technical_point_to_point_info(project_id):
        """获取技术需求点对点应答文件信息"""
        return _get_file_info(project_id, 'technical_point_to_point_file', '技术需求点对点应答文件')

    @app.route('/api/tender-processing/technical-proposal-info/<int:project_id>', methods=['GET'])
    def get_technical_proposal_info(project_id):
        """获取技术方案文件信息"""
        return _get_file_info(project_id, 'technical_proposal_file', '技术方案文件')

    @app.route('/api/tender-processing/business-response-info/<int:project_id>', methods=['GET'])
    def get_business_response_info(project_id):
        """获取商务应答完成文件信息（仅返回从商务应答同步的完成文件）"""
        try:
            file_info = None
            field_name = 'business_response_file'  # 固定使用商务应答完成文件字段

            # 尝试从数据库获取文件信息
            try:
                task_data = db.execute_query("""
                    SELECT step1_data FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if task_data:
                    step1_data = json.loads(task_data['step1_data'])

                    # 仅使用 business_response_file（从商务应答同步的完成文件）
                    file_info = step1_data.get('business_response_file')
                    field_name = 'business_response_file'
            except Exception as db_error:
                logger.warning(f"数据库查询失败,将尝试文件系统扫描: {str(db_error)}")

            # 如果数据库中没有文件信息,尝试从文件系统扫描
            if not file_info:
                logger.info(f"数据库中无商务应答完成文件,尝试从文件系统扫描")
                file_path = _find_file_in_filesystem(project_id, field_name)

                if file_path:
                    # 构建文件信息
                    file_stat = file_path.stat()
                    file_info = {
                        'file_path': str(file_path),
                        'filename': file_path.name,
                        'file_size': file_stat.st_size,
                        'saved_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                    }
                    logger.info(f"从文件系统找到商务应答文件: {file_path.name}")

            if not file_info:
                return jsonify({
                    "success": True,
                    "has_file": False
                })

            # 检查文件是否存在
            file_exists = os.path.exists(file_info.get('file_path', ''))

            # 如果数据库中的文件路径不存在,尝试从文件系统扫描
            if not file_exists and 'file_path' in file_info:
                logger.warning(f"数据库中的文件路径不存在: {file_info['file_path']}, 尝试文件系统扫描")
                file_path = _find_file_in_filesystem(project_id, field_name)

                if file_path:
                    # 更新文件信息
                    file_stat = file_path.stat()
                    file_info = {
                        'file_path': str(file_path),
                        'filename': file_path.name,
                        'file_size': file_stat.st_size,
                        'saved_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                    }
                    file_exists = True
                    logger.info(f"从文件系统找到替代文件: {file_path.name}")

            return jsonify({
                "success": True,
                "has_file": file_exists,
                "filename": file_info.get('filename'),
                "file_size": file_info.get('file_size'),
                "saved_at": file_info.get('saved_at'),
                "download_url": f"/api/tender-processing/download-file/{project_id}/{field_name}"
            })

        except Exception as e:
            logger.error(f"获取商务应答文件信息失败: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500

    # ============================================
    # 统一的文件下载和预览API
    # ============================================

    @app.route('/api/tender-processing/download-file/<int:project_id>/<field_name>', methods=['GET'])
    def download_unified_file(project_id, field_name):
        """统一的文件下载接口(支持数据库和文件系统fallback)"""
        try:
            # 验证field_name合法性（防止SQL注入）
            valid_fields = ['technical_point_to_point_file', 'technical_proposal_file', 'response_file', 'business_response_file']
            if field_name not in valid_fields:
                return jsonify({"error": "无效的文件类型"}), 400

            file_info = None

            # 尝试从数据库获取文件信息
            try:
                task_data = db.execute_query("""
                    SELECT step1_data FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if task_data:
                    step1_data = json.loads(task_data['step1_data'])
                    file_info = step1_data.get(field_name)
            except Exception as db_error:
                logger.warning(f"数据库查询失败,将尝试文件系统扫描: {str(db_error)}")

            # 如果数据库中没有文件信息,尝试从文件系统扫描
            if not file_info:
                logger.info(f"数据库中无文件信息,尝试从文件系统扫描: {field_name}")
                file_path_obj = _find_file_in_filesystem(project_id, field_name)

                if file_path_obj:
                    file_info = {
                        'file_path': str(file_path_obj),
                        'filename': file_path_obj.name
                    }

            if not file_info:
                return jsonify({"error": "文件不存在"}), 404

            file_path = file_info['file_path']
            if not os.path.exists(file_path):
                return jsonify({"error": "文件已被删除"}), 404

            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=file_info.get('filename', os.path.basename(file_path))
            )

        except Exception as e:
            logger.error(f"下载文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/tender-processing/preview-file/<int:project_id>/<field_name>', methods=['GET'])
    def preview_unified_file(project_id, field_name):
        """统一的文件预览接口(支持数据库和文件系统fallback)"""
        try:
            # 验证field_name合法性
            valid_fields = ['technical_point_to_point_file', 'technical_proposal_file', 'response_file', 'business_response_file']
            if field_name not in valid_fields:
                return jsonify({"error": "无效的文件类型"}), 400

            file_info = None

            # 尝试从数据库获取文件信息
            try:
                task_data = db.execute_query("""
                    SELECT step1_data FROM tender_projects
                    WHERE project_id = ?
                """, (project_id,), fetch_one=True)

                if task_data:
                    step1_data = json.loads(task_data['step1_data'])
                    file_info = step1_data.get(field_name)
            except Exception as db_error:
                logger.warning(f"数据库查询失败,将尝试文件系统扫描: {str(db_error)}")

            # 如果数据库中没有文件信息,尝试从文件系统扫描
            if not file_info:
                logger.info(f"数据库中无文件信息,尝试从文件系统扫描: {field_name}")
                file_path_obj = _find_file_in_filesystem(project_id, field_name)

                if file_path_obj:
                    file_info = {
                        'file_path': str(file_path_obj),
                        'filename': file_path_obj.name
                    }

            if not file_info:
                return jsonify({"error": "文件不存在"}), 404

            file_path = file_info['file_path']
            if not os.path.exists(file_path):
                return jsonify({"error": "文件已被删除"}), 404

            # 预览模式：as_attachment=False
            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=False,
                download_name=file_info.get('filename', os.path.basename(file_path))
            )

        except Exception as e:
            logger.error(f"预览文件失败: {str(e)}")
            return jsonify({"error": str(e)}), 500




# ============================================
# 辅助函数（在 register_hitl_routes 之外）
# ============================================

def convert_image_urls_to_paths(image_config_urls, company_id, db):
    """
    将图片URL转换为实际文件路径

    Args:
        image_config_urls: 图片配置字典 {资质名称: 图片URL}
        company_id: 公司ID
        db: 数据库连接

    Returns:
        转换后的图片配置字典 {资质名称: 文件路径}
    """
    image_config = {}

    for qual_name, image_url in image_config_urls.items():
        try:
            # URL格式通常是: /uploads/<company_id>/<qual_key>/<filename>
            # 从URL中提取qualification_id
            # 由于前端传递的是image_url，我们需要通过URL找到对应的资质记录

            # 简单方法：通过image_url直接查询数据库
            qualification = db.execute_query("""
                SELECT file_path, image_path
                FROM company_qualifications
                WHERE company_id = ?
                AND (image_url = ? OR image_path = ?)
            """, (company_id, image_url, image_url), fetch_one=True)

            if qualification:
                # 优先使用file_path，如果没有则使用image_path
                file_path = qualification.get('file_path') or qualification.get('image_path')
                if file_path and os.path.exists(file_path):
                    image_config[qual_name] = file_path
                    logger.info(f"  转换资质 '{qual_name}': {image_url} -> {file_path}")
                else:
                    logger.warning(f"  资质 '{qual_name}' 文件不存在: {file_path}")
            else:
                logger.warning(f"  未找到资质 '{qual_name}' 对应的数据库记录: {image_url}")

        except Exception as e:
            logger.error(f"转换图片路径失败 ({qual_name}): {e}")
            continue

    return image_config


