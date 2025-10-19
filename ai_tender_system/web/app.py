#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI标书系统统一Web应用
整合招标信息提取、点对点应答、技术方案生成等所有功能
"""

import os
import sys
import tempfile
import hashlib
import time
import re
import urllib.parse
from datetime import datetime
from pathlib import Path
from functools import wraps
from flask import Flask, request, jsonify, render_template, send_file, send_from_directory, session, redirect, url_for
from flask_cors import CORS
from flask_wtf.csrf import CSRFProtect, generate_csrf
from werkzeug.utils import secure_filename

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

# 导入公共组件
from common import (
    get_config, setup_logging, get_module_logger,
    AITenderSystemError, format_error_response, handle_api_error,
    safe_filename, allowed_file, ensure_dir
)

# 导入业务模块
try:
    from modules.tender_info.extractor import TenderInfoExtractor
    TENDER_INFO_AVAILABLE = True
except ImportError as e:
    print(f"招标信息提取模块加载失败: {e}")
    TENDER_INFO_AVAILABLE = False

# 商务应答模块（原点对点应答）
try:
    from modules.business_response.processor import BusinessResponseProcessor, PointToPointProcessor
    BUSINESS_RESPONSE_AVAILABLE = True
    POINT_TO_POINT_AVAILABLE = True  # 保持向后兼容
except ImportError as e:
    print(f"商务应答模块加载失败: {e}")
    BUSINESS_RESPONSE_AVAILABLE = False
    POINT_TO_POINT_AVAILABLE = False

# 技术需求回复模块
try:
    from modules.point_to_point.tech_responder import TechResponder
    TECH_RESPONDER_AVAILABLE = True
except ImportError as e:
    print(f"技术需求回复模块加载失败: {e}")
    TECH_RESPONDER_AVAILABLE = False

# 全局pipeline实例存储（用于分步处理）
PIPELINE_INSTANCES = {}

def create_app() -> Flask:
    """创建Flask应用"""
    # 初始化配置和日志
    config = get_config()
    setup_logging()
    logger = get_module_logger("web_app")
    
    # 创建Flask应用
    app = Flask(__name__, 
                template_folder=str(config.get_path('templates')),
                static_folder=str(config.get_path('static')))
    
    # 配置应用
    web_config = config.get_web_config()
    app.config.update({
        'SECRET_KEY': web_config['secret_key'],
        'MAX_CONTENT_LENGTH': config.get_upload_config()['max_file_size']
    })

    # 开发模式下禁用静态文件缓存
    if app.debug:
        app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
        logger.info("开发模式：已禁用静态文件缓存")

    # 启用CORS
    CORS(app, supports_credentials=True)

    # 启用CSRF保护
    csrf = CSRFProtect(app)
    logger.info("CSRF保护已启用")

    # 提供CSRF token的API端点
    @app.route('/api/csrf-token', methods=['GET'])
    def get_csrf_token():
        """获取CSRF token（用于AJAX请求）"""
        token = generate_csrf()
        return jsonify({'csrf_token': token})

    # 注册知识库API蓝图
    try:
        from modules.knowledge_base.api import knowledge_base_api
        app.register_blueprint(knowledge_base_api.get_blueprint())
        logger.info("知识库API模块注册成功")
    except ImportError as e:
        logger.warning(f"知识库API模块加载失败: {e}")

    # 注册向量搜索API蓝图
    try:
        from modules.vector_search_api import vector_search_api
        app.register_blueprint(vector_search_api.get_blueprint())
        logger.info("向量搜索API模块注册成功")
    except ImportError as e:
        logger.warning(f"向量搜索API模块加载失败: {e}")

    # 注册RAG知识库API蓝图
    try:
        from modules.knowledge_base.rag_api import rag_api
        app.register_blueprint(rag_api, url_prefix='/api')
        logger.info("RAG知识库API模块注册成功")
    except ImportError as e:
        logger.warning(f"RAG知识库API模块加载失败: {e}")

    # 注册技术方案大纲生成API蓝图
    try:
        from web.api_outline_generator import api_outline_bp
        app.register_blueprint(api_outline_bp)
        logger.info("技术方案大纲生成API模块注册成功")
    except ImportError as e:
        logger.warning(f"技术方案大纲生成API模块加载失败: {e}")

    # 注册案例库API蓝图
    try:
        from modules.case_library.api import case_library_api
        app.register_blueprint(case_library_api.get_blueprint())
        logger.info("案例库API模块注册成功")
    except ImportError as e:
        logger.warning(f"案例库API模块加载失败: {e}")

    # 注册内部蓝图（新架构）
    from web.blueprints import register_all_blueprints
    register_all_blueprints(app, config, logger)

    # 注册路由（旧架构 - 将逐步迁移）
    register_routes(app, config, logger)
    
    logger.info("AI标书系统Web应用初始化完成")
    return app

def register_routes(app: Flask, config, logger):
    """注册所有路由"""

    # 初始化知识库管理器
    from modules.knowledge_base.manager import KnowledgeBaseManager
    kb_manager = KnowledgeBaseManager()

    # ===================
    # 辅助函数 - 已迁移的函数
    # login_required -> middleware/auth.py
    # enrich_qualification_with_company_status -> blueprints/api_tender_bp.py
    # ===================

    # ===================
    # 已迁移到蓝图的路由 (Phase 1 + Phase 2)
    # ===================
    # Phase 1: 认证和静态页面 -> blueprints/auth_bp.py, pages_bp.py, static_files_bp.py
    # Phase 2: 核心API -> blueprints/api_core_bp.py (/api/health, /api/config)
    # Phase 2: 文件管理 -> blueprints/api_files_bp.py (/upload, /download/<filename>)
    # Phase 2: 招标信息 -> blueprints/api_tender_bp.py (/extract-tender-info, /extract-tender-info-step)
    # ===================
    
    # ===================
    # 商务应答相关路由（暂时占位）
    # ===================

    def build_image_config_from_db(company_id: int) -> dict:
        """
        从数据库加载公司资质信息并构建图片配置

        Args:
            company_id: 公司ID

        Returns:
            图片配置字典，包含：
            - seal_path: 公章图片路径
            - license_path: 营业执照图片路径
            - qualification_paths: 资质证书图片路径列表
        """
        try:
            # 从数据库获取公司的所有资质
            qualifications = kb_manager.db.get_company_qualifications(company_id)

            if not qualifications:
                logger.warning(f"公司 {company_id} 没有上传任何资质文件")
                return {}

            logger.info(f"从数据库加载公司 {company_id} 的资质信息，共 {len(qualifications)} 个资质")

            image_config = {}
            qualification_paths = []

            # 遍历所有资质，按类型分类
            for qual in qualifications:
                qual_key = qual.get('qualification_key')
                file_path = qual.get('file_path')

                if not file_path:
                    continue

                # 营业执照
                if qual_key == 'business_license':
                    image_config['license_path'] = file_path
                    logger.info(f"  - 营业执照: {file_path}")

                # 公章
                elif qual_key == 'company_seal':
                    image_config['seal_path'] = file_path
                    logger.info(f"  - 公章: {file_path}")

                # 资质证书 - 包括各类ISO认证、CMMI等
                elif qual_key in ['iso9001', 'iso14001', 'iso20000', 'iso27001',
                                 'cmmi', 'itss', 'safety_production',
                                 'software_copyright', 'patent_certificate']:
                    qualification_paths.append(file_path)
                    logger.info(f"  - 资质证书 ({qual_key}): {file_path}")

            # 添加资质证书列表
            if qualification_paths:
                image_config['qualification_paths'] = qualification_paths

            logger.info(f"构建的图片配置: {len(image_config)} 个类型，{len(qualification_paths)} 个资质证书")
            return image_config

        except Exception as e:
            logger.error(f"从数据库构建图片配置失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return {}

    def generate_output_filename(project_name: str, file_type: str, timestamp: str = None) -> str:
        """
        生成统一格式的输出文件名: {项目名称}_{类型}_{时间戳}.docx

        Args:
            project_name: 项目名称
            file_type: 文件类型（如：商务应答、点对点应答、技术方案）
            timestamp: 时间戳，如果未提供则自动生成

        Returns:
            格式化的文件名
        """
        if not timestamp:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        # 使用safe_filename确保项目名称安全，但不添加时间戳（避免重复）
        safe_project = safe_filename(project_name, timestamp=False) if project_name else "未命名项目"

        return f"{safe_project}_{file_type}_{timestamp}.docx"

    @app.route('/process-business-response', methods=['POST'])
    def process_business_response():
        """处理商务应答"""
        if not BUSINESS_RESPONSE_AVAILABLE:
            return jsonify({
                'success': False,
                'message': '商务应答模块不可用'
            })
        
        try:
            # 获取上传的文件
            if 'template_file' not in request.files:
                raise ValueError("没有选择模板文件")
            
            file = request.files['template_file']
            if file.filename == '':
                raise ValueError("文件名为空")
            
            # 获取表单数据
            data = request.form.to_dict()
            company_id = data.get('company_id', '')
            project_name = data.get('project_name', '')
            tender_no = data.get('tender_no', '')
            date_text = data.get('date_text', '')
            use_mcp = data.get('use_mcp', 'false').lower() == 'true'

            # 验证必填字段
            if not company_id:
                raise ValueError("请选择应答公司")

            # 转换公司ID为整数
            company_id_int = int(company_id)

            # 从数据库获取项目相关信息（如果有项目名称）
            purchaser_name = ''
            db_project_number = ''
            authorized_rep_name = ''
            authorized_rep_position = ''
            if project_name:
                try:
                    query = """SELECT tenderer, project_number,
                               authorized_representative_name,
                               authorized_representative_position
                               FROM tender_projects WHERE project_name = ? LIMIT 1"""
                    result = kb_manager.db.execute_query(query, [project_name])
                    if result and len(result) > 0:
                        purchaser_name = result[0].get('tenderer', '')
                        db_project_number = result[0].get('project_number', '')
                        authorized_rep_name = result[0].get('authorized_representative_name', '')
                        authorized_rep_position = result[0].get('authorized_representative_position', '')
                        if purchaser_name:
                            logger.info(f"从数据库获取采购人信息: {purchaser_name}")
                        if db_project_number:
                            logger.info(f"从数据库获取项目编号: {db_project_number}")
                        if authorized_rep_name:
                            logger.info(f"从数据库获取授权人信息: {authorized_rep_name} ({authorized_rep_position})")
                except Exception as e:
                    logger.warning(f"查询项目信息失败: {e}")

            # 如果表单没有提供项目编号，使用数据库中的项目编号
            if not tender_no and db_project_number:
                tender_no = db_project_number
                logger.info(f"使用数据库项目编号: {tender_no}")

            # 从数据库直接加载图片配置（新方案：消除前端时序问题）
            image_config = build_image_config_from_db(company_id_int)

            if image_config:
                logger.info(f"成功从数据库加载图片配置，包含 {len(image_config)} 个类型")
            else:
                logger.warning(f"公司 {company_id} 没有可用的资质图片")

            # 从数据库获取公司信息
            company_db_data = kb_manager.get_company_detail(company_id_int)
            if not company_db_data:
                raise ValueError(f"未找到公司信息: {company_id}")

            # 使用现有字段映射反向转换为业务处理器期望的格式
            field_mapping = {
                'companyName': 'company_name',
                'establishDate': 'establish_date',
                'legalRepresentative': 'legal_representative',
                'legalRepresentativePosition': 'legal_representative_position',
                'legalRepresentativeGender': 'legal_representative_gender',
                'legalRepresentativeAge': 'legal_representative_age',
                'socialCreditCode': 'social_credit_code',
                'registeredCapital': 'registered_capital',
                'companyType': 'company_type',
                'registeredAddress': 'registered_address',
                'businessScope': 'business_scope',
                'companyDescription': 'description',
                'fixedPhone': 'fixed_phone',
                'fax': 'fax',
                'postalCode': 'postal_code',
                'email': 'email',
                'officeAddress': 'office_address',
                'employeeCount': 'employee_count',
                'bankName': 'bank_name',
                'bankAccount': 'bank_account'
            }
            reverse_mapping = {v: k for k, v in field_mapping.items()}
            company_data = {reverse_mapping.get(k, k): v for k, v in company_db_data.items()}

            # 添加采购人信息到company_data（采购人是项目信息，但为了方便传递，加到这里）
            if purchaser_name:
                company_data['purchaserName'] = purchaser_name

            # 添加授权人信息到company_data
            if authorized_rep_name:
                company_data['representativeName'] = authorized_rep_name
            if authorized_rep_position:
                company_data['representativeTitle'] = authorized_rep_position

            # 保存模板文件 - 使用统一服务
            from core.storage_service import storage_service
            file_metadata = storage_service.store_file(
                file_obj=file,
                original_name=file.filename,
                category='business_templates',
                business_type='business_response',
                company_id=company_id
            )
            template_path = Path(file_metadata.file_path)
            filename = file_metadata.safe_name

            logger.info(f"开始处理商务应答: {file_metadata.original_name}")
            
            # 公共的输出文件路径设置（移到外面，两个分支都需要）
            output_dir = ensure_dir(config.get_path('output'))
            # 使用新的文件命名规则：{项目名称}_商务应答_{时间戳}.docx
            output_filename = generate_output_filename(project_name, "商务应答")
            output_path = output_dir / output_filename

            # 添加调试日志
            logger.info(f"[文件命名调试] project_name参数: {project_name}")
            logger.info(f"[文件命名调试] 生成的文件名: {output_filename}")
            logger.info(f"[文件命名调试] 输出目录: {output_dir}")
            logger.info(f"[文件命名调试] 完整输出路径: {output_path}")

            logger.info(f"公司数据验证:")
            logger.info(f"  - 公司名称: {company_data.get('companyName', 'N/A')}")
            logger.info(f"  - 联系电话: {company_data.get('fixedPhone', 'N/A')}")
            logger.info(f"  - 电子邮件: {company_data.get('email', 'N/A')}")
            logger.info(f"  - 公司地址: {company_data.get('address', 'N/A')}")
            logger.info(f"  - 传真号码: {company_data.get('fax', 'N/A')}")
            logger.info(f"  - 项目名称: {project_name}")
            logger.info(f"  - 招标编号: {tender_no}")
            logger.info(f"  - 日期文本: {date_text}")
            
            # 使用MCP处理器处理商务应答
            if use_mcp:
                # 使用新架构的商务应答处理器
                processor = BusinessResponseProcessor()

                # 使用MCP处理器的完整商务应答处理方法，包含日期字段处理和图片插入
                result_stats = processor.process_business_response(
                    str(template_path),
                    str(output_path),
                    company_data,
                    project_name,
                    tender_no,
                    date_text,
                    image_config  # 传递图片配置
                )
                
                output_path = str(output_path)
                
                # 检查处理结果并构建响应
                if result_stats.get('success'):
                    logger.info(f"新架构处理器执行成功: {result_stats.get('message', '无消息')}")
                    logger.info(f"处理统计: {result_stats.get('stats', {})}")
                    
                    # 构建成功结果
                    result = {
                        'success': True,
                        'message': result_stats.get('message', '商务应答处理完成'),
                        'output_file': output_path,
                        'download_url': f'/download/{os.path.basename(output_path)}',
                        'stats': result_stats.get('stats', {})
                    }
                else:
                    logger.error(f"新架构处理器执行失败: {result_stats.get('error', '未知错误')}")
                    result = {
                        'success': False,
                        'error': result_stats.get('error', '处理失败'),
                        'message': result_stats.get('message', '商务应答处理失败')
                    }
            else:
                # 使用向后兼容的处理器（实际上还是新的BusinessResponseProcessor）
                processor = PointToPointProcessor()  # 这是BusinessResponseProcessor的别名
                result_stats = processor.process_business_response(
                    str(template_path),
                    str(output_path),
                    company_data,
                    project_name,
                    tender_no,
                    date_text,
                    image_config  # 传递图片配置
                )
                
                # 统一返回格式处理
                if result_stats.get('success'):
                    result = {
                        'success': True,
                        'message': result_stats.get('message', '商务应答处理完成'),
                        'output_file': str(output_path),
                        'download_url': f'/download/{os.path.basename(output_path)}',
                        'stats': result_stats.get('summary', {})
                    }
                else:
                    result = {
                        'success': False,
                        'error': result_stats.get('error', '处理失败'),
                        'message': result_stats.get('message', '商务应答处理失败')
                    }
            
            logger.info("商务应答处理完成")
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"商务应答处理失败: {e}")
            return jsonify(format_error_response(e))
    
    @app.route('/api/document/process', methods=['POST'])
    def process_document():
        """处理文档 - 通用接口"""
        if not BUSINESS_RESPONSE_AVAILABLE:
            return jsonify({
                'success': False,
                'message': '文档处理模块不可用'
            })
        
        try:
            data = request.get_json()
            file_path = data.get('file_path', '')
            options = data.get('options', {})
            
            if not file_path:
                raise ValueError("文件路径不能为空")
            
            # 这是一个通用接口，根据选项决定使用哪个处理器
            doc_type = options.get('type', 'business_response')
            
            if doc_type == 'tech_requirements' and TECH_RESPONDER_AVAILABLE:
                result = {
                    'success': True,
                    'message': '技术需求处理功能可用，请使用 /process-tech-requirements 接口',
                    'redirect': '/process-tech-requirements'
                }
            else:
                result = {
                    'success': True,
                    'message': '商务应答处理功能可用，请使用 /process-business-response 接口',
                    'redirect': '/process-business-response'
                }
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            return jsonify(format_error_response(e))
    
    @app.route('/process-point-to-point', methods=['POST'])
    def process_point_to_point():
        """处理点对点应答 - 使用内联回复功能（原地插入应答）"""
        if not BUSINESS_RESPONSE_AVAILABLE:
            return jsonify({
                'success': False,
                'message': '点对点应答模块不可用'
            })

        try:
            # 检查是否使用HITL技术需求文件
            use_hitl_file = request.form.get('use_hitl_technical_file') == 'true'
            hitl_task_id = request.form.get('hitl_task_id')

            if use_hitl_file and hitl_task_id:
                # 从HITL任务获取技术需求文件
                logger.info(f"使用HITL任务的技术需求文件: {hitl_task_id}")

                # 查询HITL任务的技术需求文件路径
                # 技术需求文件保存在 technical_files/{year}/{month}/{task_id}/ 目录下
                from datetime import datetime
                now = datetime.now()
                year = now.strftime('%Y')
                month = now.strftime('%m')

                technical_dir = Path(config.get_path('upload')) / 'technical_files' / year / month / hitl_task_id

                if not technical_dir.exists():
                    raise ValueError(f"HITL任务技术需求文件目录不存在: {technical_dir}")

                # 查找目录中的docx文件
                docx_files = list(technical_dir.glob('*.docx'))
                if not docx_files:
                    raise ValueError(f"HITL任务目录中没有找到技术需求文件: {technical_dir}")

                # 使用第一个找到的docx文件
                file_path = docx_files[0]
                filename = file_path.name
                logger.info(f"使用HITL技术需求文件: {filename}, 路径: {file_path}")
            else:
                # 获取上传的文件
                if 'file' not in request.files:
                    raise ValueError("没有选择文件")

                file = request.files['file']
                if file.filename == '':
                    raise ValueError("文件名为空")

                # 保存文件 - 使用统一服务
                from core.storage_service import storage_service
                file_metadata = storage_service.store_file(
                    file_obj=file,
                    original_name=file.filename,
                    category='point_to_point',
                    business_type='point_to_point_response'
                )
                file_path = Path(file_metadata.file_path)
                filename = file_metadata.original_name

            logger.info(f"开始处理点对点应答: {filename}")

            # 获取公司ID参数
            company_id = request.form.get('companyId')
            if not company_id:
                return jsonify({
                    'success': False,
                    'error': '缺少公司ID参数'
                })

            # 从数据库获取公司信息
            company_id_int = int(company_id)
            company_db_data = kb_manager.get_company_detail(company_id_int)
            if not company_db_data:
                return jsonify({
                    'success': False,
                    'error': f'未找到公司数据: {company_id}'
                })

            # 使用现有字段映射反向转换为业务处理器期望的格式
            field_mapping = {
                'companyName': 'company_name',
                'establishDate': 'establish_date',
                'legalRepresentative': 'legal_representative',
                'legalRepresentativePosition': 'legal_representative_position',
                'legalRepresentativeGender': 'legal_representative_gender',
                'legalRepresentativeAge': 'legal_representative_age',
                'socialCreditCode': 'social_credit_code',
                'registeredCapital': 'registered_capital',
                'companyType': 'company_type',
                'registeredAddress': 'registered_address',
                'businessScope': 'business_scope',
                'companyDescription': 'description',
                'fixedPhone': 'fixed_phone',
                'fax': 'fax',
                'postalCode': 'postal_code',
                'email': 'email',
                'officeAddress': 'office_address',
                'employeeCount': 'employee_count',
                'bankName': 'bank_name',
                'bankAccount': 'bank_account'
            }
            reverse_mapping = {v: k for k, v in field_mapping.items()}
            company_data = {reverse_mapping.get(k, k): v for k, v in company_db_data.items()}

            logger.info(f"使用公司信息: {company_data.get('companyName', 'N/A')}")

            # 获取配置参数
            response_frequency = request.form.get('responseFrequency', 'every_paragraph')
            response_mode = request.form.get('responseMode', 'simple')
            ai_model = request.form.get('aiModel', 'shihuang-gpt4o-mini')

            # 根据模型选择映射到正确的模型名称
            model_mapping = {
                'gpt-4o-mini': 'shihuang-gpt4o-mini',
                'gpt-4': 'shihuang-gpt4',
                'deepseek-v3': 'yuanjing-deepseek-v3',
                'qwen-235b': 'yuanjing-qwen-235b'
            }
            actual_model = model_mapping.get(ai_model, ai_model)

            logger.info(f"配置参数 - 应答频次: {response_frequency}, 应答方式: {response_mode}, AI模型: {actual_model}")

            # 创建商务应答处理器（使用内联回复功能）
            processor = BusinessResponseProcessor(model_name=actual_model)

            # 获取项目名称参数
            project_name = request.form.get('projectName')

            # 生成输出文件路径
            output_dir = ensure_dir(config.get_path('output'))
            # 使用新的文件命名规则：{项目名称}_点对点应答_{时间戳}.docx
            # 如果有项目名称，使用项目名称；否则使用原文件名
            base_name = project_name if project_name else Path(filename).stem
            output_filename = generate_output_filename(base_name, "点对点应答")
            output_path = output_dir / output_filename

            # 判断是否使用AI（根据应答方式）
            use_ai = response_mode == 'ai'
            logger.info(f"处理模式: {'AI智能应答' if use_ai else '简单模板应答'}")

            # 使用新的内联回复处理方法
            result_stats = processor.process_inline_reply(
                str(file_path),
                str(output_path),
                use_ai=use_ai
            )

            if result_stats.get('success'):
                logger.info(f"内联回复处理成功: {result_stats.get('message')}")

                # 生成下载URL
                download_url = f'/download/{output_filename}'

                return jsonify({
                    'success': True,
                    'message': '内联回复处理完成，应答已插入到原文档中（灰色底纹标记）',
                    'download_url': download_url,
                    'filename': output_filename,
                    'output_file': str(output_path),
                    'model_used': actual_model,
                    'requirements_count': result_stats.get('requirements_count', 0),
                    'responses_count': result_stats.get('responses_count', 0),
                    'response_mode': response_mode,
                    'model_name': actual_model if use_ai else None,
                    'features': result_stats.get('features', {}),
                    'stats': {
                        'inline_reply': True,
                        'gray_shading': True,
                        'format_preserved': True,
                        'requirements_count': result_stats.get('requirements_count', 0),
                        'responses_count': result_stats.get('responses_count', 0),
                        'response_mode': response_mode,
                        'model_name': actual_model if use_ai else None
                    }
                })
            else:
                logger.error(f"内联回复处理失败: {result_stats.get('error')}")
                return jsonify({
                    'success': False,
                    'error': result_stats.get('error', '处理失败'),
                    'message': result_stats.get('message', '内联回复处理失败')
                })

        except Exception as e:
            logger.error(f"点对点应答处理失败: {e}")
            return jsonify(format_error_response(e))

    # ===================
    # 技术需求回复路由
    # ===================
    
    @app.route('/process-tech-requirements', methods=['POST'])
    def process_tech_requirements():
        """处理技术需求回复"""
        if not TECH_RESPONDER_AVAILABLE:
            return jsonify({
                'success': False,
                'message': '技术需求回复模块不可用'
            })
        
        try:
            # 获取上传的文件
            if 'requirements_file' not in request.files:
                raise ValueError("没有选择需求文件")
            
            file = request.files['requirements_file']
            if file.filename == '':
                raise ValueError("文件名为空")
            
            # 获取表单数据
            data = request.form.to_dict()
            company_id = data.get('company_id', '')
            response_strategy = data.get('response_strategy', 'comprehensive')
            
            # 验证必填字段
            if not company_id:
                raise ValueError("请选择应答公司")
            
            # 保存上传的文件 - 使用统一服务
            from core.storage_service import storage_service
            file_metadata = storage_service.store_file(
                file_obj=file,
                original_name=file.filename,
                category='tech_proposals',
                business_type='tech_requirements',
                company_id=int(company_id)
            )
            requirements_path = Path(file_metadata.file_path)
            
            # 从数据库获取公司信息
            company_id_int = int(company_id)
            company_db_data = kb_manager.get_company_detail(company_id_int)
            if not company_db_data:
                raise ValueError(f"未找到公司数据: {company_id}")

            # 使用现有字段映射反向转换为业务处理器期望的格式
            field_mapping = {
                'companyName': 'company_name',
                'establishDate': 'establish_date',
                'legalRepresentative': 'legal_representative',
                'legalRepresentativePosition': 'legal_representative_position',
                'legalRepresentativeGender': 'legal_representative_gender',
                'legalRepresentativeAge': 'legal_representative_age',
                'socialCreditCode': 'social_credit_code',
                'registeredCapital': 'registered_capital',
                'companyType': 'company_type',
                'registeredAddress': 'registered_address',
                'businessScope': 'business_scope',
                'companyDescription': 'description',
                'fixedPhone': 'fixed_phone',
                'fax': 'fax',
                'postalCode': 'postal_code',
                'email': 'email',
                'officeAddress': 'office_address',
                'employeeCount': 'employee_count',
                'bankName': 'bank_name',
                'bankAccount': 'bank_account'
            }
            reverse_mapping = {v: k for k, v in field_mapping.items()}
            company_data = {reverse_mapping.get(k, k): v for k, v in company_db_data.items()}
            
            logger.info(f"开始处理技术需求: {filename}")
            
            # 创建技术需求回复处理器
            responder = TechResponder()
            
            # 生成输出文件路径
            output_dir = ensure_dir(config.get_path('output'))
            # 使用新的文件命名规则：{项目名称}_技术方案_{时间戳}.docx
            # 对于技术需求回复，使用原文件名作为项目名称（因为没有单独的project_name字段）
            base_name = Path(file_metadata.original_name).stem
            output_filename = generate_output_filename(base_name, "技术方案")
            output_path = output_dir / output_filename
            
            # 处理技术需求
            result_stats = responder.process_tech_requirements(
                str(requirements_path),
                str(output_path),
                company_data,
                response_strategy
            )
            
            if result_stats.get('success'):
                logger.info(f"技术需求处理成功: {result_stats.get('message')}")
                result = {
                    'success': True,
                    'message': result_stats.get('message', '技术需求处理完成'),
                    'output_file': str(output_path),
                    'download_url': f'/download/{os.path.basename(output_path)}',
                    'stats': {
                        'requirements_count': result_stats.get('requirements_count', 0),
                        'responses_count': result_stats.get('responses_count', 0)
                    }
                }
            else:
                logger.error(f"技术需求处理失败: {result_stats.get('error')}")
                result = {
                    'success': False,
                    'error': result_stats.get('error', '处理失败'),
                    'message': result_stats.get('message', '技术需求处理失败')
                }
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"技术需求处理异常: {e}")
            return jsonify(format_error_response(e))
    
    # 文档预览和编辑API
    @app.route('/api/document/preview/<filename>', methods=['GET'])
    def preview_document(filename):
        """预览文档内容 - 直接返回.docx文件供前端mammoth.js转换"""
        try:
            # URL解码文件名（处理中文字符等）
            filename = urllib.parse.unquote(filename)

            # 只进行基本的安全检查，避免路径遍历攻击
            if '..' in filename or '/' in filename or '\\' in filename:
                return jsonify({'success': False, 'error': '非法文件名'}), 400

            # 先尝试从output目录查找，如果不存在则从upload目录查找
            file_path = config.get_path('output') / filename

            if not file_path.exists():
                file_path = config.get_path('upload') / filename

            if not file_path.exists():
                return jsonify({'success': False, 'error': f'文档不存在: {filename}'}), 404

            file_ext = file_path.suffix.lower()

            # 只处理Word文档
            if file_ext not in ['.doc', '.docx']:
                return jsonify({'success': False, 'error': f'不支持的文件格式: {file_ext}'}), 400

            # 如果是.doc格式，提示需要转换
            if file_ext == '.doc':
                return jsonify({
                    'success': False,
                    'error': '旧版.doc格式预览失败。建议：\n1. 将文件另存为.docx格式\n2. 或直接进行信息提取（系统会自动处理）'
                }), 400

            # 直接返回.docx文件，让前端mammoth.js处理
            return send_file(
                file_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=False,
                download_name=filename
            )

        except Exception as e:
            logger.error(f"文档预览失败: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500
    
    @app.route('/api/editor/load-document', methods=['POST'])
    def load_document_for_edit():
        """加载文档用于编辑"""
        try:
            from docx import Document
            from markupsafe import Markup
            
            if 'file' not in request.files:
                raise ValueError("没有选择文件")
            
            file = request.files['file']
            if file.filename == '':
                raise ValueError("文件名为空")
            
            # 读取Word文档
            doc = Document(file)
            
            # 转换为HTML格式用于编辑器
            html_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    style_name = paragraph.style.name if paragraph.style else ''
                    
                    if 'Heading 1' in style_name:
                        html_content.append(f'<h1>{paragraph.text}</h1>')
                    elif 'Heading 2' in style_name:
                        html_content.append(f'<h2>{paragraph.text}</h2>')
                    elif 'Heading 3' in style_name:
                        html_content.append(f'<h3>{paragraph.text}</h3>')
                    else:
                        html_content.append(f'<p>{paragraph.text}</p>')
            
            # 处理表格
            for table in doc.tables:
                html_content.append('<table>')
                for row in table.rows:
                    html_content.append('<tr>')
                    for cell in row.cells:
                        html_content.append(f'<td>{cell.text}</td>')
                    html_content.append('</tr>')
                html_content.append('</table>')
            
            return jsonify({
                'success': True,
                'html_content': ''.join(html_content),
                'original_filename': file.filename
            })
            
        except Exception as e:
            logger.error(f"文档加载失败: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            })
    
    @app.route('/api/editor/save-document', methods=['POST'])
    def save_edited_document():
        """保存编辑后的文档"""
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            import re
            
            data = request.get_json()
            html_content = data.get('html_content', '')
            filename = data.get('filename', 'document')
            
            if not html_content:
                raise ValueError("文档内容为空")
            
            # 创建新文档
            doc = Document()
            
            # 解析HTML内容
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 处理各种HTML元素
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table']):
                if element.name == 'h1':
                    doc.add_heading(element.get_text(), level=1)
                elif element.name == 'h2':
                    doc.add_heading(element.get_text(), level=2)
                elif element.name == 'h3':
                    doc.add_heading(element.get_text(), level=3)
                elif element.name == 'p':
                    text = element.get_text()
                    if text.strip():
                        doc.add_paragraph(text)
                elif element.name == 'table':
                    # 计算表格行列数
                    rows = element.find_all('tr')
                    if rows:
                        cols = len(rows[0].find_all(['td', 'th']))
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = 'Table Grid'
                        
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols:
                                    table.cell(i, j).text = cell.get_text()
            
            # 保存文档
            output_dir = ensure_dir(config.get_path('output'))
            output_path = output_dir / f"{filename}.docx"
            doc.save(str(output_path))
            
            # 返回文件供下载
            return send_file(
                str(output_path),
                as_attachment=True,
                download_name=f"{filename}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        except Exception as e:
            logger.error(f"文档保存失败: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            })
    
    @app.route('/api/editor/upload-image', methods=['POST'])
    def upload_editor_image():
        """上传编辑器图片"""
        try:
            if 'image' not in request.files:
                raise ValueError("没有选择图片")
            
            file = request.files['image']
            if file.filename == '':
                raise ValueError("文件名为空")
            
            # 保存图片 - 使用统一服务
            from core.storage_service import storage_service
            file_metadata = storage_service.store_file(
                file_obj=file,
                original_name=file.filename,
                category='processed_results',  # 编辑器图片属于处理结果
                business_type='editor_image'
            )
            file_path = Path(file_metadata.file_path)
            filename = file_metadata.safe_name
            
            # 返回图片URL
            image_url = f'/static/uploads/images/{filename}'
            
            return jsonify({
                'success': True,
                'location': image_url
            })
            
        except Exception as e:
            logger.error(f"图片上传失败: {e}")
            return jsonify({
                'success': False,
                'error': str(e)
            })
    
    @app.route('/api/table/analyze', methods=['POST'])
    def analyze_table():
        """分析表格"""
        if not POINT_TO_POINT_AVAILABLE:
            return jsonify({
                'success': False,
                'message': '表格处理模块不可用'
            })
        
        try:
            data = request.get_json()
            table_data = data.get('table_data', {})
            
            processor = TableProcessor()
            result = processor.analyze_table(table_data)
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"表格分析失败: {e}")
            return jsonify(format_error_response(e))
    
    @app.route('/api/table/process', methods=['POST'])
    def process_table():
        """处理表格"""
        if not POINT_TO_POINT_AVAILABLE:
            return jsonify({
                'success': False,
                'message': '表格处理模块不可用'
            })
        
        try:
            data = request.get_json()
            table_data = data.get('table_data', {})
            options = data.get('options', {})
            
            processor = TableProcessor()
            result = processor.process_table(table_data, options)
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"表格处理失败: {e}")
            return jsonify(format_error_response(e))
    
    # ===================
    # ===================
    # 公司管理API
    # ===================
    
    @app.route('/api/companies')
    def list_companies():
        """获取所有公司配置"""
        try:
            companies = kb_manager.get_companies()

            # 转换字段格式以保持前端兼容性，过滤无效公司ID
            result_companies = []
            for company in companies:
                company_id = company.get('company_id')
                # 跳过没有有效 company_id 的记录
                if company_id is None:
                    logger.warning(f"跳过无效的公司记录，company_id为None: {company.get('company_name', '未知')}")
                    continue

                result_companies.append({
                    'company_id': company_id,
                    'company_name': company.get('company_name', '未命名公司'),
                    'created_at': company.get('created_at', ''),
                    'updated_at': company.get('updated_at', ''),
                    'product_count': company.get('product_count', 0),
                    'document_count': company.get('document_count', 0)
                })

            # 安全排序，处理可能的 None 值
            result_companies.sort(key=lambda x: x.get('updated_at') or '', reverse=True)
            return jsonify({'success': True, 'data': result_companies})

        except Exception as e:
            logger.error(f"获取公司列表失败: {e}")
            return jsonify({'error': str(e)}), 500

    @app.route('/api/companies/<company_id>')
    def get_company(company_id):
        """获取指定公司的详细信息"""
        try:
            # 转换字符串ID为整数ID
            company_id_int = int(company_id)

            company_data = kb_manager.get_company_detail(company_id_int)

            # DEBUG: 记录从数据库获取的原始数据
            logger.info(f"[DEBUG GET] 公司 {company_id} - 数据库返回的原始数据: {company_data}")
            if company_data and 'registered_capital' in company_data:
                logger.info(f"[DEBUG GET] registered_capital 字段存在: {company_data['registered_capital']!r}")
            elif company_data:
                logger.info(f"[DEBUG GET] registered_capital 字段不在返回数据中，可用字段: {list(company_data.keys())}")

            if not company_data:
                return jsonify({'success': False, 'error': '公司不存在'}), 404

            # 转换字段格式以保持前端兼容性 - 保持原有格式
            result_company = company_data

            return jsonify({'success': True, 'data': result_company})

        except ValueError:
            return jsonify({'success': False, 'error': '无效的公司ID'}), 400
        except Exception as e:
            logger.error(f"获取公司信息失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies', methods=['POST'])
    def create_company():
        """创建新公司"""
        try:
            data = request.get_json()
            if not data:
                return jsonify({'success': False, 'error': '请提供公司信息'}), 400

            company_name = data.get('companyName', '').strip()
            if not company_name:
                return jsonify({'success': False, 'error': '公司名称不能为空'}), 400

            # 使用知识库管理器创建公司
            result = kb_manager.create_company(
                company_name=company_name,
                company_code=data.get('companyCode', None),
                industry_type=data.get('industryType', None),
                description=data.get('companyDescription', None)
            )

            if result['success']:
                # 导入datetime
                from datetime import datetime

                # 返回格式与前端兼容
                company_data = {
                    'id': str(result['company_id']),
                    'companyName': company_name,
                    'created_at': datetime.now().isoformat(),
                    'updated_at': datetime.now().isoformat()
                }

                logger.info(f"创建公司成功: {company_name} (ID: {result['company_id']})")
                return jsonify({'success': True, 'company': company_data})
            else:
                return jsonify({'success': False, 'error': result['error']}), 400

        except Exception as e:
            logger.error(f"创建公司失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies/<company_id>', methods=['PUT'])
    def update_company(company_id):
        """更新公司信息"""
        try:
            # 转换字符串ID为整数ID
            company_id_int = int(company_id)

            data = request.get_json()
            if not data:
                return jsonify({'success': False, 'error': '请提供公司信息'}), 400

            # 使用知识库管理器更新公司信息
            result = kb_manager.update_company(company_id_int, data)

            if result['success']:
                # 获取更新后的公司详情
                updated_company = kb_manager.get_company_detail(company_id_int)

                if updated_company:
                    # 转换格式与前端兼容
                    result_company = {
                        'id': str(updated_company.get('company_id', '')),
                        'companyName': updated_company.get('company_name', ''),
                        'establishDate': updated_company.get('establish_date', ''),
                        'legalRepresentative': updated_company.get('legal_representative', ''),
                        'legalRepresentativePosition': updated_company.get('legal_representative_position', ''),
                        'socialCreditCode': updated_company.get('social_credit_code', ''),
                        'registeredCapital': updated_company.get('registered_capital', ''),
                        'companyType': updated_company.get('company_type', ''),
                        'registeredAddress': updated_company.get('registered_address', ''),
                        'businessScope': updated_company.get('business_scope', ''),
                        'companyDescription': updated_company.get('description', ''),
                        'fixedPhone': updated_company.get('fixed_phone', ''),
                        'fax': updated_company.get('fax', ''),
                        'postalCode': updated_company.get('postal_code', ''),
                        'email': updated_company.get('email', ''),
                        'officeAddress': updated_company.get('office_address', ''),
                        'employeeCount': updated_company.get('employee_count', ''),
                        'created_at': updated_company.get('created_at', ''),
                        'updated_at': updated_company.get('updated_at', '')
                    }

                    logger.info(f"更新公司成功: {updated_company.get('company_name', '')} (ID: {company_id})")
                    return jsonify({'success': True, 'company': result_company, 'message': '公司信息更新成功'})
                else:
                    return jsonify({'success': False, 'error': '获取更新后的公司信息失败'}), 500
            else:
                return jsonify({'success': False, 'error': result['error']}), 400

        except ValueError:
            return jsonify({'success': False, 'error': '无效的公司ID'}), 400
        except Exception as e:
            logger.error(f"更新公司失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies/<company_id>', methods=['DELETE'])
    def delete_company(company_id):
        """删除公司"""
        try:
            # 转换字符串ID为整数ID
            company_id_int = int(company_id)

            # 使用知识库管理器删除公司
            result = kb_manager.delete_company(company_id_int)

            if result['success']:
                logger.info(f"删除公司成功: {company_id}")
                return jsonify({'success': True, 'message': '公司删除成功'})
            else:
                if '不存在' in result['error']:
                    return jsonify({'success': False, 'error': result['error']}), 404
                else:
                    return jsonify({'success': False, 'error': result['error']}), 500

        except ValueError:
            return jsonify({'success': False, 'error': '无效的公司ID'}), 400
        except Exception as e:
            logger.error(f"删除公司失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies/<company_id>/qualifications')
    def get_company_qualifications(company_id):
        """获取公司资质文件列表"""
        try:
            # 验证公司ID并获取资质列表
            try:
                company_id_int = int(company_id)
                # 使用数据库方法获取资质列表
                qualifications = kb_manager.get_company_qualifications(company_id_int)

                # 转换为前端期望的格式
                qualifications_dict = {}
                for qual in qualifications:
                    qual_key = qual['qualification_key']
                    qualifications_dict[qual_key] = {
                        'original_filename': qual['original_filename'],
                        'safe_filename': qual['safe_filename'],
                        'file_size': qual['file_size'],
                        'upload_time': qual['upload_time'],
                        'custom_name': qual.get('custom_name'),
                        'expire_date': qual.get('expire_date'),
                        'verify_status': qual.get('verify_status', 'pending')
                    }

                logger.info(f"获取公司 {company_id} 的资质文件列表，共 {len(qualifications_dict)} 个")
                return jsonify({
                    'success': True,
                    'qualifications': qualifications_dict
                })

            except ValueError:
                return jsonify({'success': False, 'error': '无效的公司ID'}), 400

        except Exception as e:
            logger.error(f"获取公司资质文件失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies/<company_id>/qualifications/upload', methods=['POST'])
    def upload_company_qualifications(company_id):
        """上传公司资质文件"""
        try:
            import json
            import shutil
            from werkzeug.utils import secure_filename

            # 首先检查数据库中是否存在该公司
            try:
                company_id_int = int(company_id)
                company_data = kb_manager.get_company_detail(company_id_int)
                if not company_data:
                    return jsonify({'success': False, 'error': '公司不存在'}), 404
            except ValueError:
                return jsonify({'success': False, 'error': '无效的公司ID'}), 400
            
            # 处理上传的文件
            uploaded_files = {}
            qualification_names = request.form.get('qualification_names', '{}')
            qualification_names = json.loads(qualification_names) if qualification_names else {}

            for key, file in request.files.items():
                if key.startswith('qualifications[') and file.filename:
                    # 提取资质键名
                    qual_key = key.replace('qualifications[', '').replace(']', '')

                    # 使用数据库方法上传资质文件
                    result = kb_manager.upload_qualification(
                        company_id=company_id_int,
                        qualification_key=qual_key,
                        file_obj=file,
                        original_filename=file.filename,
                        qualification_name=qualification_names.get(qual_key, qual_key),
                        custom_name=qualification_names.get(qual_key) if qual_key.startswith('custom') else None
                    )

                    if result['success']:
                        uploaded_files[qual_key] = {
                            'filename': file.filename,
                            'qualification_id': result['qualification_id'],
                            'message': result['message']
                        }
            
            logger.info(f"公司 {company_id} 上传了 {len(uploaded_files)} 个资质文件")
            return jsonify({
                'success': True, 
                'message': f'成功上传 {len(uploaded_files)} 个资质文件',
                'uploaded_files': uploaded_files
            })
            
        except Exception as e:
            logger.error(f"上传公司资质文件失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies/<company_id>/qualifications/<qualification_key>/download')
    def download_qualification_file(company_id, qualification_key):
        """下载公司资质文件"""
        try:
            # 验证公司ID
            try:
                company_id_int = int(company_id)
                company_data = kb_manager.get_company_detail(company_id_int)
                if not company_data:
                    return jsonify({'success': False, 'error': '公司不存在'}), 404
            except ValueError:
                return jsonify({'success': False, 'error': '无效的公司ID'}), 400

            # 从数据库获取资质文件信息
            qualification = kb_manager.db.get_qualification_by_key(company_id_int, qualification_key)
            if not qualification:
                return jsonify({'success': False, 'error': '资质文件不存在'}), 404

            # 检查文件是否存在
            file_path = Path(qualification['file_path'])
            if not file_path.exists():
                return jsonify({'success': False, 'error': '文件不存在'}), 404

            # 返回文件
            original_filename = qualification['original_filename']
            logger.info(f"下载资质文件: {original_filename}")
            return send_file(str(file_path), as_attachment=True, download_name=original_filename)

        except Exception as e:
            logger.error(f"下载资质文件失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies/<company_id>/qualifications/<qualification_key>/preview')
    def preview_qualification_file(company_id, qualification_key):
        """预览资质文件 - 返回JSON格式（符合全站架构）"""
        try:
            import base64

            # 验证公司ID
            try:
                company_id_int = int(company_id)
                company_data = kb_manager.get_company_detail(company_id_int)
                if not company_data:
                    return jsonify({'success': False, 'error': '公司不存在'}), 404
            except ValueError:
                return jsonify({'success': False, 'error': '无效的公司ID'}), 400

            # 从数据库获取资质文件信息
            qualification = kb_manager.db.get_qualification_by_key(company_id_int, qualification_key)
            if not qualification:
                return jsonify({'success': False, 'error': '资质文件不存在'}), 404

            # 检查文件是否存在
            file_path = Path(qualification['file_path'])
            if not file_path.exists():
                return jsonify({'success': False, 'error': '文件不存在'}), 404

            file_type = qualification['file_type'].lower() if qualification['file_type'] else ''
            filename = qualification['original_filename']

            # 根据文件类型生成HTML内容
            if file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                # 图片：base64编码嵌入
                with open(file_path, 'rb') as f:
                    img_data = base64.b64encode(f.read()).decode()
                html_content = f'''
                    <div class="text-center p-4">
                        <img src="data:image/{file_type};base64,{img_data}"
                             class="img-fluid"
                             style="max-width: 100%; height: auto; box-shadow: 0 2px 8px rgba(0,0,0,0.1);" />
                    </div>
                '''
            elif file_type == 'pdf':
                # PDF：提示下载
                html_content = f'''
                    <div class="alert alert-info m-4">
                        <h5><i class="bi bi-file-pdf"></i> PDF文档预览</h5>
                        <p class="mb-0">文件名: {filename}</p>
                        <p class="text-muted">PDF预览功能正在开发中，请使用下载功能查看完整内容。</p>
                    </div>
                '''
            else:
                return jsonify({'success': False, 'error': f'不支持的文件格式: {file_type}'}), 400

            logger.info(f"预览资质文件: {filename}")
            return jsonify({
                'success': True,
                'content': html_content,
                'filename': filename
            })

        except Exception as e:
            logger.error(f"预览资质文件失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/companies/<company_id>/qualifications/<qualification_key>', methods=['DELETE'])
    def delete_qualification_file(company_id, qualification_key):
        """删除公司资质文件"""
        try:
            # 验证公司ID
            try:
                company_id_int = int(company_id)
                company_data = kb_manager.get_company_detail(company_id_int)
                if not company_data:
                    return jsonify({'success': False, 'error': '公司不存在'}), 404
            except ValueError:
                return jsonify({'success': False, 'error': '无效的公司ID'}), 400

            # 使用新的数据库方法删除资质文件
            result = kb_manager.delete_qualification_by_key(company_id_int, qualification_key)

            if result['success']:
                return jsonify(result)
            else:
                return jsonify(result), 400

        except Exception as e:
            logger.error(f"删除资质文件失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    # ===================
    # 商务文件管理API
    # ===================
    
    @app.route('/api/business-files')
    def list_business_files():
        """获取商务应答文件列表"""
        try:
            import os
            from datetime import datetime

            def format_size(size_bytes):
                """格式化文件大小"""
                for unit in ['B', 'KB', 'MB', 'GB']:
                    if size_bytes < 1024.0:
                        return f"{size_bytes:.1f} {unit}"
                    size_bytes /= 1024.0
                return f"{size_bytes:.1f} TB"

            files = []
            output_dir = config.get_path('output')

            if output_dir.exists():
                for filename in os.listdir(output_dir):
                    # 过滤商务应答文件（只匹配包含"商务应答"的文件）
                    if filename.endswith(('.docx', '.doc', '.pdf')) and '商务应答' in filename:
                        file_path = output_dir / filename
                        try:
                            stat = file_path.stat()
                            modified_time = datetime.fromtimestamp(stat.st_mtime)
                            files.append({
                                'name': filename,
                                'size': format_size(stat.st_size),
                                'date': modified_time.strftime('%Y-%m-%d %H:%M:%S'),
                                'download_url': f'/download/{filename}',
                                'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                                'modified': modified_time.isoformat()
                            })
                        except Exception as e:
                            logger.warning(f"读取文件信息失败 {filename}: {e}")

            files.sort(key=lambda x: x.get('modified', ''), reverse=True)
            return jsonify({'success': True, 'files': files})

        except Exception as e:
            logger.error(f"获取商务文件列表失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/point-to-point/files')
    def list_point_to_point_files():
        """获取点对点应答文件列表"""
        try:
            import os
            from datetime import datetime

            files = []
            output_dir = config.get_path('output')

            if output_dir.exists():
                for filename in os.listdir(output_dir):
                    # 过滤点对点应答文件（只匹配包含"点对点应答"或"点对点"的文件）
                    if filename.endswith(('.docx', '.doc', '.pdf')) and ('点对点应答' in filename or '点对点' in filename or 'point-to-point' in filename.lower()):
                        file_path = output_dir / filename
                        try:
                            stat = file_path.stat()
                            files.append({
                                'id': hashlib.md5(str(file_path).encode()).hexdigest()[:8],
                                'filename': filename,
                                'original_filename': filename,
                                'file_path': str(file_path),
                                'output_path': str(file_path),
                                'size': stat.st_size,
                                'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                                'process_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                                'status': 'completed',
                                'company_name': '未知公司'  # 暂时使用默认值，后续会从数据库获取
                            })
                        except Exception as e:
                            logger.warning(f"读取文件信息失败 {filename}: {e}")

            files.sort(key=lambda x: x.get('process_time', ''), reverse=True)
            return jsonify({'success': True, 'data': files})

        except Exception as e:
            logger.error(f"获取点对点应答文件列表失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/point-to-point/preview')
    def preview_point_to_point_document():
        """预览点对点应答文档 - 直接返回.docx文件供前端mammoth.js转换"""
        try:
            # 获取参数
            file_id = request.args.get('file_id')
            file_path = request.args.get('file_path')

            if not file_id and not file_path:
                return jsonify({'success': False, 'error': '缺少文件ID或文件路径参数'}), 400

            # 根据参数确定文件路径
            if file_path:
                target_file = Path(file_path)
            else:
                # 如果只有file_id，需要从输出目录查找文件
                output_dir = config.get_path('output')
                target_file = None
                if output_dir.exists():
                    for filename in os.listdir(output_dir):
                        full_path = output_dir / filename
                        if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                            target_file = full_path
                            break

                if not target_file:
                    return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

            if not target_file.exists():
                return jsonify({'success': False, 'error': '文件不存在'}), 404

            # 根据文件类型进行预览
            file_extension = target_file.suffix.lower()

            if file_extension == '.docx':
                # 直接返回.docx文件，让前端mammoth.js处理
                return send_file(
                    target_file,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=False,
                    download_name=target_file.name
                )

            elif file_extension == '.doc':
                return jsonify({
                    'success': False,
                    'error': '旧版.doc格式预览失败。建议：\n1. 将文件另存为.docx格式\n2. 或直接进行信息提取（系统会自动处理）'
                }), 400

            elif file_extension == '.pdf':
                # PDF预览（简单实现，返回提示信息）
                return jsonify({
                    'success': True,
                    'content': f'<div class="alert alert-info"><h4>PDF文档预览</h4><p>文件名: {target_file.name}</p><p>PDF文件预览功能正在开发中，请使用下载功能查看完整内容。</p></div>',
                    'filename': target_file.name
                })

            else:
                return jsonify({'success': False, 'error': '不支持的文件格式'}), 400

        except Exception as e:
            logger.error(f"文档预览失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/point-to-point/edit', methods=['GET', 'POST'])
    def edit_point_to_point_document():
        """编辑点对点应答文档"""

        if request.method == 'GET':
            # 获取文档内容用于编辑
            try:
                from docx import Document
                import html

                # 获取参数
                file_id = request.args.get('file_id')
                file_path = request.args.get('file_path')

                if not file_id and not file_path:
                    return jsonify({'success': False, 'error': '缺少文件ID或文件路径参数'}), 400

                # 根据参数确定文件路径
                if file_path:
                    target_file = Path(file_path)
                else:
                    # 如果只有file_id，需要从输出目录查找文件
                    output_dir = config.get_path('output')
                    target_file = None
                    if output_dir.exists():
                        for filename in os.listdir(output_dir):
                            full_path = output_dir / filename
                            if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                                target_file = full_path
                                break

                    if not target_file:
                        return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

                if not target_file.exists():
                    return jsonify({'success': False, 'error': '文件不存在'}), 404

                # 只支持Word文档编辑
                file_extension = target_file.suffix.lower()
                if file_extension not in ['.docx', '.doc']:
                    return jsonify({'success': False, 'error': '只支持Word文档编辑'}), 400

                try:
                    # 读取Word文档并转换为可编辑的HTML
                    doc = Document(target_file)
                    html_content = []

                    # 处理段落
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            style_name = paragraph.style.name if paragraph.style else ''
                            text = html.escape(paragraph.text)

                            if 'Heading 1' in style_name or 'heading 1' in style_name.lower():
                                html_content.append(f'<h1>{text}</h1>')
                            elif 'Heading 2' in style_name or 'heading 2' in style_name.lower():
                                html_content.append(f'<h2>{text}</h2>')
                            elif 'Heading 3' in style_name or 'heading 3' in style_name.lower():
                                html_content.append(f'<h3>{text}</h3>')
                            else:
                                html_content.append(f'<p>{text}</p>')

                    # 处理表格（简化为文本形式）
                    for table in doc.tables:
                        html_content.append('<table border="1">')
                        for row in table.rows:
                            html_content.append('<tr>')
                            for cell in row.cells:
                                cell_text = html.escape(cell.text)
                                html_content.append(f'<td>{cell_text}</td>')
                            html_content.append('</tr>')
                        html_content.append('</table>')

                    return jsonify({
                        'success': True,
                        'content': '\n'.join(html_content),
                        'filename': target_file.name
                    })

                except Exception as e:
                    logger.error(f"读取文档内容失败: {e}")
                    return jsonify({'success': False, 'error': f'读取文档内容失败: {str(e)}'}), 500

            except Exception as e:
                logger.error(f"获取编辑内容失败: {e}")
                return jsonify({'success': False, 'error': str(e)}), 500

        elif request.method == 'POST':
            # 保存编辑后的文档
            try:
                from docx import Document
                from bs4 import BeautifulSoup
                import re

                # 获取参数
                file_id = request.args.get('file_id')
                file_path = request.args.get('file_path')

                # 获取POST数据
                data = request.get_json()
                if not data or 'content' not in data:
                    return jsonify({'success': False, 'error': '缺少文档内容'}), 400

                new_content = data['content']

                # 根据参数确定文件路径
                if file_path:
                    target_file = Path(file_path)
                else:
                    # 如果只有file_id，需要从输出目录查找文件
                    output_dir = config.get_path('output')
                    target_file = None
                    if output_dir.exists():
                        for filename in os.listdir(output_dir):
                            full_path = output_dir / filename
                            if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                                target_file = full_path
                                break

                    if not target_file:
                        return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

                if not target_file.exists():
                    return jsonify({'success': False, 'error': '文件不存在'}), 404

                try:
                    # 解析HTML内容
                    soup = BeautifulSoup(new_content, 'html.parser')

                    # 创建新的Word文档
                    doc = Document()

                    # 遍历解析的HTML元素
                    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table']):
                        if element.name in ['h1', 'h2', 'h3']:
                            # 添加标题
                            heading_level = int(element.name[1])
                            paragraph = doc.add_heading(element.get_text().strip(), level=heading_level)
                        elif element.name == 'p':
                            # 添加段落
                            doc.add_paragraph(element.get_text().strip())
                        elif element.name == 'table':
                            # 添加表格
                            rows = element.find_all('tr')
                            if rows:
                                cols = len(rows[0].find_all(['td', 'th']))
                                table = doc.add_table(rows=len(rows), cols=cols)
                                table.style = 'Table Grid'

                                for i, row in enumerate(rows):
                                    cells = row.find_all(['td', 'th'])
                                    for j, cell in enumerate(cells):
                                        if i < len(table.rows) and j < len(table.rows[i].cells):
                                            table.rows[i].cells[j].text = cell.get_text().strip()

                    # 保存文档
                    doc.save(str(target_file))

                    logger.info(f"文档保存成功: {target_file}")

                    return jsonify({
                        'success': True,
                        'message': '文档保存成功',
                        'filename': target_file.name
                    })

                except Exception as e:
                    logger.error(f"保存文档失败: {e}")
                    return jsonify({'success': False, 'error': f'保存文档失败: {str(e)}'}), 500

            except Exception as e:
                logger.error(f"编辑文档失败: {e}")
                return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/point-to-point/download')
    def download_point_to_point_document():
        """下载点对点应答文档"""
        try:
            # 获取参数
            file_id = request.args.get('file_id')
            file_path = request.args.get('file_path')

            if not file_id and not file_path:
                return jsonify({'success': False, 'error': '缺少文件ID或文件路径参数'}), 400

            # 根据参数确定文件路径
            if file_path:
                target_file = Path(file_path)
            else:
                # 如果只有file_id，需要从输出目录查找文件
                output_dir = config.get_path('output')
                target_file = None
                if output_dir.exists():
                    for filename in os.listdir(output_dir):
                        full_path = output_dir / filename
                        if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                            target_file = full_path
                            break

                if not target_file:
                    return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

            if not target_file.exists():
                return jsonify({'success': False, 'error': '文件不存在'}), 404

            # 确定MIME类型
            file_extension = target_file.suffix.lower()
            if file_extension == '.docx':
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif file_extension == '.doc':
                mimetype = 'application/msword'
            elif file_extension == '.pdf':
                mimetype = 'application/pdf'
            else:
                mimetype = 'application/octet-stream'

            # 生成下载文件名
            download_filename = target_file.name

            logger.info(f"开始下载文件: {target_file}")

            return send_file(
                str(target_file),
                as_attachment=True,
                download_name=download_filename,
                mimetype=mimetype
            )

        except Exception as e:
            logger.error(f"文档下载失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    # ===================
    # 项目配置API
    # ===================
    # 招标项目管理API
    # ===================

    @app.route('/api/tender-projects', methods=['GET'])
    def get_tender_projects():
        """获取招标项目列表"""
        try:
            company_id = request.args.get('company_id')
            status = request.args.get('status')

            query = "SELECT * FROM tender_projects WHERE 1=1"
            params = []

            if company_id:
                query += " AND company_id = ?"
                params.append(company_id)

            if status:
                query += " AND status = ?"
                params.append(status)

            query += " ORDER BY created_at DESC LIMIT 100"

            projects = kb_manager.db.execute_query(query, params)

            return jsonify({
                'success': True,
                'data': projects or []
            })
        except Exception as e:
            logger.error(f"获取项目列表失败: {e}")
            return jsonify({
                'success': False,
                'message': str(e),
                'data': []
            })

    @app.route('/api/tender-projects', methods=['POST'])
    def create_tender_project():
        """创建新招标项目"""
        try:
            import json
            data = request.get_json()

            # 【新增】检查是否存在相同项目（防止重复创建）
            company_id = data.get('company_id')
            project_name = data.get('project_name')
            project_number = data.get('project_number')

            if company_id and project_name:
                check_query = """
                    SELECT project_id FROM tender_projects
                    WHERE company_id = ? AND project_name = ?
                """
                check_params = [company_id, project_name]

                if project_number:
                    check_query += " AND project_number = ?"
                    check_params.append(project_number)

                existing = kb_manager.db.execute_query(check_query, check_params, fetch_one=True)

                if existing:
                    logger.warning(f"项目已存在，返回已有项目ID: {existing['project_id']}")
                    return jsonify({
                        'success': True,
                        'project_id': existing['project_id'],
                        'message': '项目已存在',
                        'is_existing': True
                    })

            # 序列化资质和评分数据为JSON
            qualifications_json = None
            scoring_json = None

            if data.get('qualifications_data'):
                qualifications_json = json.dumps(data.get('qualifications_data'), ensure_ascii=False)
            if data.get('scoring_data'):
                scoring_json = json.dumps(data.get('scoring_data'), ensure_ascii=False)

            query = """
                INSERT INTO tender_projects (
                    project_name, project_number, tenderer, agency,
                    bidding_method, bidding_location, bidding_time,
                    tender_document_path, original_filename,
                    company_id, qualifications_data, scoring_data,
                    status, created_by
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """

            params = [
                data.get('project_name'),
                data.get('project_number'),
                data.get('tenderer'),
                data.get('agency'),
                data.get('bidding_method'),
                data.get('bidding_location'),
                data.get('bidding_time'),
                data.get('tender_document_path'),
                data.get('original_filename'),
                data.get('company_id'),
                qualifications_json,
                scoring_json,
                'draft',
                'system'
            ]

            project_id = kb_manager.db.execute_query(query, params)

            logger.info(f"创建项目成功，ID: {project_id}")

            return jsonify({
                'success': True,
                'project_id': project_id,
                'message': '项目创建成功'
            })
        except Exception as e:
            logger.error(f"创建项目失败: {e}")
            return jsonify({
                'success': False,
                'message': str(e)
            })

    @app.route('/api/tender-projects/<int:project_id>', methods=['GET'])
    def get_tender_project(project_id):
        """获取单个项目详情"""
        try:
            query = "SELECT * FROM tender_projects WHERE project_id = ?"
            projects = kb_manager.db.execute_query(query, [project_id])

            if projects and len(projects) > 0:
                return jsonify({
                    'success': True,
                    'data': projects[0]
                })
            else:
                return jsonify({
                    'success': False,
                    'message': '项目不存在'
                })
        except Exception as e:
            logger.error(f"获取项目详情失败: {e}")
            return jsonify({
                'success': False,
                'message': str(e)
            })

    @app.route('/api/tender-projects/<int:project_id>', methods=['PUT'])
    def update_tender_project(project_id):
        """更新招标项目"""
        try:
            import json
            data = request.get_json()

            # 序列化资质和评分数据为JSON
            qualifications_json = None
            scoring_json = None

            if data.get('qualifications_data'):
                qualifications_json = json.dumps(data.get('qualifications_data'), ensure_ascii=False)
            if data.get('scoring_data'):
                scoring_json = json.dumps(data.get('scoring_data'), ensure_ascii=False)

            query = """
                UPDATE tender_projects SET
                    project_name = ?,
                    project_number = ?,
                    tenderer = ?,
                    agency = ?,
                    bidding_method = ?,
                    bidding_location = ?,
                    bidding_time = ?,
                    tender_document_path = ?,
                    original_filename = ?,
                    company_id = ?,
                    qualifications_data = ?,
                    scoring_data = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE project_id = ?
            """

            params = [
                data.get('project_name'),
                data.get('project_number'),
                data.get('tenderer'),
                data.get('agency'),
                data.get('bidding_method'),
                data.get('bidding_location'),
                data.get('bidding_time'),
                data.get('tender_document_path'),
                data.get('original_filename'),
                data.get('company_id'),
                qualifications_json,
                scoring_json,
                project_id
            ]

            kb_manager.db.execute_query(query, params)

            logger.info(f"更新项目成功，ID: {project_id}")

            return jsonify({
                'success': True,
                'project_id': project_id,
                'message': '项目更新成功'
            })
        except Exception as e:
            logger.error(f"更新项目失败: {e}")
            return jsonify({
                'success': False,
                'message': str(e)
            })

    # ===================

    @app.route('/api/project-config')
    def get_project_config():
        """获取项目配置信息"""
        try:
            import configparser
            
            # 读取招标信息提取模块生成的配置文件
            config_file = config.get_path('config') / 'tender_config.ini'
            
            if not config_file.exists():
                return jsonify({'success': False, 'error': '项目配置文件不存在'})
                
            ini_config = configparser.ConfigParser(interpolation=None)
            ini_config.read(config_file, encoding='utf-8')
            
            # 提取项目信息
            project_info = {}
            if ini_config.has_section('PROJECT_INFO'):
                project_info = {
                    'projectName': ini_config.get('PROJECT_INFO', 'project_name', fallback=''),
                    'projectNumber': ini_config.get('PROJECT_INFO', 'project_number', fallback=''),
                    'tenderer': ini_config.get('PROJECT_INFO', 'tenderer', fallback=''),
                    'agency': ini_config.get('PROJECT_INFO', 'agency', fallback=''),
                    'biddingMethod': ini_config.get('PROJECT_INFO', 'bidding_method', fallback=''),
                    'biddingLocation': ini_config.get('PROJECT_INFO', 'bidding_location', fallback=''),
                    'biddingTime': ini_config.get('PROJECT_INFO', 'bidding_time', fallback=''),
                }
            
            return jsonify({'success': True, 'projectInfo': project_info})
            
        except Exception as e:
            logger.error(f"获取项目配置失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500
    
    # ===================
    # 模型管理API
    # ===================

    @app.route('/api/models', methods=['GET'])
    def get_available_models():
        """获取可用的AI模型列表"""
        try:
            from common.llm_client import get_available_models

            models = get_available_models()

            # 添加模型状态检查
            for model in models:
                try:
                    # 这里可以添加模型可用性检查的逻辑
                    model['status'] = 'available' if model['has_api_key'] else 'no_api_key'
                    model['status_message'] = '已配置' if model['has_api_key'] else '未配置API密钥'
                except (KeyError, TypeError, AttributeError) as e:
                    logger.warning(f"处理模型状态时出错: {e}")
                    model['status'] = 'unknown'
                    model['status_message'] = '状态未知'

            logger.info(f"获取模型列表成功，共 {len(models)} 个模型")
            return jsonify({
                'success': True,
                'models': models,
                'count': len(models)
            })

        except Exception as e:
            logger.error(f"获取模型列表失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/models/<model_name>/validate', methods=['POST'])
    def validate_model_config(model_name):
        """验证指定模型的配置"""
        try:
            from common.llm_client import create_llm_client

            # 创建模型客户端并验证配置
            client = create_llm_client(model_name)
            validation_result = client.validate_config()

            logger.info(f"模型 {model_name} 配置验证结果: {validation_result['valid']}")
            return jsonify({
                'success': True,
                'validation': validation_result
            })

        except Exception as e:
            logger.error(f"验证模型 {model_name} 配置失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    # ===================
    # 标书智能处理API
    # ===================

    @app.route('/api/tender-processing/start', methods=['POST'])
    def start_tender_processing():
        """启动标书智能处理流程"""
        try:
            # 获取表单数据
            project_id = request.form.get('project_id')
            filter_model = request.form.get('filter_model', 'gpt-4o-mini')
            extract_model = request.form.get('extract_model', 'yuanjing-deepseek-v3')
            step = int(request.form.get('step', 1))  # 默认只执行第1步（分块）

            if not project_id:
                return jsonify({'success': False, 'error': '缺少project_id参数'}), 400

            # 检查文件上传
            if 'file' not in request.files:
                return jsonify({'success': False, 'error': '未上传文件'}), 400

            file = request.files['file']
            if not file.filename:
                return jsonify({'success': False, 'error': '文件名为空'}), 400

            # 保存文件到临时目录
            import os
            import tempfile
            from pathlib import Path

            temp_dir = Path(tempfile.gettempdir()) / 'tender_processing'
            temp_dir.mkdir(exist_ok=True)

            file_ext = Path(file.filename).suffix
            temp_file = temp_dir / f"tender_{project_id}{file_ext}"
            file.save(str(temp_file))

            logger.info(f"文件已保存: {temp_file}")

            # 使用ParserManager解析文档
            import asyncio
            from modules.document_parser.parser_manager import ParserManager

            async def parse_document():
                parser = ParserManager()
                result = await parser.parse_document(doc_id=int(project_id), file_path=str(temp_file))
                return result

            # 运行异步解析
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            parse_result = loop.run_until_complete(parse_document())
            loop.close()

            if parse_result.status.value != 'completed':
                return jsonify({
                    'success': False,
                    'error': f'文档解析失败: {parse_result.error_message}'
                }), 500

            document_text = parse_result.content
            logger.info(f"启动标书智能处理 - 项目ID: {project_id}, 文档长度: {len(document_text)}")

            # 导入处理流程
            from modules.tender_processing.processing_pipeline import TenderProcessingPipeline

            # 创建流程实例（异步处理需要在后台线程中运行）
            import threading

            result_holder = {'task_id': None, 'error': None}

            def run_pipeline():
                try:
                    pipeline = TenderProcessingPipeline(
                        project_id=project_id,
                        document_text=document_text,
                        filter_model=filter_model,
                        extract_model=extract_model
                    )
                    result_holder['task_id'] = pipeline.task_id

                    # 保存pipeline实例到全局字典供后续步骤使用
                    PIPELINE_INSTANCES[pipeline.task_id] = pipeline

                    # 运行指定步骤
                    result = pipeline.run_step(step)
                    result_holder['result'] = result

                    logger.info(f"步骤 {step} 处理完成 - 任务ID: {pipeline.task_id}, 成功: {result['success']}")
                except Exception as e:
                    logger.error(f"处理流程执行失败: {e}")
                    result_holder['error'] = str(e)

            # 启动后台线程
            thread = threading.Thread(target=run_pipeline, daemon=True)
            thread.start()

            # 等待task_id生成（最多等待2秒）
            import time
            for _ in range(20):
                if result_holder['task_id'] or result_holder['error']:
                    break
                time.sleep(0.1)

            if result_holder['error']:
                return jsonify({'success': False, 'error': result_holder['error']}), 500

            if not result_holder['task_id']:
                return jsonify({'success': False, 'error': '任务启动超时'}), 500

            return jsonify({
                'success': True,
                'task_id': result_holder['task_id'],
                'message': '处理任务已启动，请使用task_id查询进度'
            })

        except Exception as e:
            logger.error(f"启动标书处理失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/continue/<task_id>', methods=['POST'])
    def continue_tender_processing(task_id):
        """继续执行下一步骤"""
        try:
            # 获取参数
            data = request.get_json()
            step = data.get('step', 2)  # 默认执行第2步

            # 从全局字典中获取pipeline实例
            if task_id not in PIPELINE_INSTANCES:
                return jsonify({'success': False, 'error': f'找不到任务 {task_id} 的pipeline实例'}), 404

            pipeline = PIPELINE_INSTANCES[task_id]

            # 在后台线程中执行步骤
            import threading

            result_holder = {'result': None, 'error': None}

            def run_step():
                try:
                    result = pipeline.run_step(step)
                    result_holder['result'] = result
                    logger.info(f"步骤 {step} 处理完成 - 任务ID: {task_id}, 成功: {result['success']}")
                except Exception as e:
                    logger.error(f"步骤 {step} 执行失败: {e}")
                    result_holder['error'] = str(e)

            # 启动后台线程
            thread = threading.Thread(target=run_step, daemon=True)
            thread.start()

            # 等待步骤完成（最多等待5秒以返回响应）
            import time
            for _ in range(50):
                if result_holder['result'] or result_holder['error']:
                    break
                time.sleep(0.1)

            if result_holder['error']:
                return jsonify({'success': False, 'error': result_holder['error']}), 500

            if not result_holder['result']:
                # 步骤仍在执行中，返回处理中状态
                return jsonify({
                    'success': True,
                    'task_id': task_id,
                    'message': f'步骤 {step} 正在处理中，请查询状态'
                })

            # 如果是最后一步（第3步），清理pipeline实例
            if step == 3:
                PIPELINE_INSTANCES.pop(task_id, None)
                logger.info(f"任务 {task_id} 已完成，清理pipeline实例")

            return jsonify({
                'success': True,
                'task_id': task_id,
                'result': result_holder['result'],
                'message': f'步骤 {step} 处理完成'
            })

        except Exception as e:
            logger.error(f"继续处理失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/status/<task_id>', methods=['GET'])
    def get_processing_status(task_id):
        """查询处理进度"""
        try:
            from common.database import get_knowledge_base_db

            db = get_knowledge_base_db()

            # 获取任务信息
            task = db.get_processing_task(task_id)

            if not task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 获取处理日志
            logs = db.get_processing_logs(task_id=task_id)

            # 获取统计信息
            stats = db.get_processing_statistics(task_id)

            return jsonify({
                'success': True,
                'task': dict(task),
                'logs': [dict(log) for log in logs],
                'statistics': dict(stats) if stats else None
            })

        except Exception as e:
            logger.error(f"查询处理状态失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/chunks/<int:project_id>', methods=['GET'])
    def get_tender_chunks(project_id):
        """获取文档分块列表"""
        try:
            from common.database import get_knowledge_base_db

            valuable_only = request.args.get('valuable_only', 'false').lower() == 'true'

            db = get_knowledge_base_db()
            chunks = db.get_tender_chunks(project_id, valuable_only=valuable_only)

            # 解析metadata JSON
            for chunk in chunks:
                if chunk.get('metadata'):
                    try:
                        chunk['metadata'] = json.loads(chunk['metadata'])
                    except (json.JSONDecodeError, TypeError) as e:
                        logger.warning(f"解析chunk metadata失败: {e}")
                        chunk['metadata'] = {}

            return jsonify({
                'success': True,
                'chunks': chunks,
                'total': len(chunks)
            })

        except Exception as e:
            logger.error(f"获取分块列表失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/requirements/<int:project_id>', methods=['GET'])
    def get_tender_requirements(project_id):
        """获取提取的要求列表"""
        try:
            from common.database import get_knowledge_base_db

            constraint_type = request.args.get('constraint_type')
            category = request.args.get('category')

            db = get_knowledge_base_db()
            requirements = db.get_tender_requirements(
                project_id=project_id,
                constraint_type=constraint_type,
                category=category
            )

            # 获取汇总统计
            summary = db.get_requirements_summary(project_id)

            return jsonify({
                'success': True,
                'requirements': requirements,
                'total': len(requirements),
                'summary': summary,
                'has_extracted': len(requirements) > 0
            })

        except Exception as e:
            logger.error(f"获取要求列表失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/analytics/<int:project_id>', methods=['GET'])
    def get_processing_analytics(project_id):
        """获取处理统计分析"""
        try:
            from common.database import get_knowledge_base_db

            db = get_knowledge_base_db()

            # 获取分块统计
            all_chunks = db.get_tender_chunks(project_id)
            valuable_chunks = db.get_tender_chunks(project_id, valuable_only=True)

            # 获取要求统计
            requirements = db.get_tender_requirements(project_id)
            summary = db.get_requirements_summary(project_id)

            # 计算分块类型分布
            chunk_type_dist = {}
            for chunk in all_chunks:
                chunk_type = chunk.get('chunk_type', 'unknown')
                chunk_type_dist[chunk_type] = chunk_type_dist.get(chunk_type, 0) + 1

            # 计算筛选效果
            filter_rate = (len(all_chunks) - len(valuable_chunks)) / len(all_chunks) * 100 if all_chunks else 0

            analytics = {
                'chunks': {
                    'total': len(all_chunks),
                    'valuable': len(valuable_chunks),
                    'filtered': len(all_chunks) - len(valuable_chunks),
                    'filter_rate': round(filter_rate, 2),
                    'type_distribution': chunk_type_dist
                },
                'requirements': {
                    'total': len(requirements),
                    'by_type': summary.get('by_type', {}),
                    'by_category': summary.get('by_category', {})
                }
            }

            return jsonify({
                'success': True,
                'analytics': analytics
            })

        except Exception as e:
            logger.error(f"获取处理统计失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/tender-processing/export/<int:project_id>', methods=['GET'])
    def export_requirements(project_id):
        """导出提取的要求为Excel"""
        try:
            from common.database import get_knowledge_base_db
            import pandas as pd
            from io import BytesIO

            db = get_knowledge_base_db()
            requirements = db.get_tender_requirements(project_id)

            if not requirements:
                return jsonify({'success': False, 'error': '没有可导出的数据'}), 404

            # 转换为DataFrame
            df = pd.DataFrame(requirements)

            # 选择需要的列
            columns = [
                'requirement_id', 'constraint_type', 'category', 'subcategory',
                'detail', 'source_location', 'priority', 'extraction_confidence',
                'is_verified', 'extracted_at'
            ]
            df = df[[col for col in columns if col in df.columns]]

            # 重命名列（中文）
            df.columns = [
                '要求ID', '类型', '分类', '子分类', '详情', '来源',
                '优先级', '置信度', '已验证', '提取时间'
            ][:len(df.columns)]

            # 生成Excel
            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='投标要求', index=False)
            output.seek(0)

            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'项目{project_id}_投标要求.xlsx'
            )

        except ImportError:
            return jsonify({
                'success': False,
                'error': '缺少pandas或openpyxl库，请安装：pip install pandas openpyxl'
            }), 500
        except Exception as e:
            logger.error(f"导出要求失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    # ===================
    # HITL（Human-in-the-Loop）API - 三步人工确认流程
    # ===================

    # 注册 HITL API 路由
    from web.api_tender_processing_hitl import register_hitl_routes
    register_hitl_routes(app)
    logger.info("HITL API 路由已注册")

    # ===================
    # 错误处理
    # ===================
    
    @app.errorhandler(404)
    def not_found_error(error):
        return jsonify({'error': 'Not Found'}), 404
    
    @app.errorhandler(500)
    def internal_error(error):
        logger.error(f"内部服务器错误: {error}")
        return jsonify({'error': 'Internal Server Error'}), 500
    
    @app.errorhandler(413)
    def file_too_large(error):
        return jsonify({'error': '文件太大'}), 413

    @app.route('/api/tender-processing/sync-point-to-point/<task_id>', methods=['POST'])
    def sync_point_to_point_to_hitl(task_id):
        """
        同步点对点应答文件到HITL投标项目
        接收点对点应答生成的文件路径,复制到HITL任务目录,保存为"应答完成文件"
        """
        try:
            import json
            import shutil
            from common.database import get_knowledge_base_db

            data = request.get_json()
            source_file_path = data.get('file_path')

            if not source_file_path:
                return jsonify({
                    'success': False,
                    'error': '未提供文件路径'
                }), 400

            # 如果传入的是下载URL(以/api/downloads/或/downloads/开头),转换为实际文件路径
            if source_file_path.startswith('/api/downloads/') or source_file_path.startswith('/downloads/'):
                filename = source_file_path.replace('/api/downloads/', '').replace('/downloads/', '')
                # 使用URL解码处理中文文件名
                from urllib.parse import unquote
                filename = unquote(filename)
                project_root = Path(__file__).parent.parent
                source_file_path = os.path.join(project_root, 'data/outputs', filename)
                logger.info(f"从下载URL转换为文件路径: {source_file_path}")

            # 检查源文件是否存在
            if not os.path.exists(source_file_path):
                return jsonify({
                    'success': False,
                    'error': '源文件不存在'
                }), 404

            logger.info(f"同步点对点应答文件到HITL项目: task_id={task_id}, file_path={source_file_path}")

            # 获取数据库实例
            db = get_knowledge_base_db()

            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_hitl_tasks
                WHERE hitl_task_id = ?
            """, (task_id,), fetch_one=True)

            if not task_data:
                return jsonify({
                    'success': False,
                    'error': '任务不存在'
                }), 404

            step1_data = json.loads(task_data['step1_data'])

            # 创建存储目录
            now = datetime.now()
            project_root = Path(__file__).parent.parent
            save_dir = os.path.join(
                project_root,
                'data/uploads/completed_response_files',
                str(now.year),
                f"{now.month:02d}",
                task_id
            )
            os.makedirs(save_dir, exist_ok=True)

            # 生成文件名
            source_filename = os.path.basename(source_file_path)
            # 从源文件名提取,如果包含时间戳则保留,否则添加时间戳
            if '_' in source_filename:
                base_name = source_filename.rsplit('.', 1)[0]
                filename = f"{base_name}_应答完成.docx"
            else:
                filename = f"点对点应答_{now.strftime('%Y%m%d_%H%M%S')}_应答完成.docx"

            # 复制文件到目标位置
            target_path = os.path.join(save_dir, filename)
            shutil.copy2(source_file_path, target_path)

            # 计算文件大小
            file_size = os.path.getsize(target_path)

            # 更新任务的step1_data - 使用独立字段存储点对点应答文件
            point_to_point_file_info = {
                "file_path": target_path,
                "filename": filename,
                "file_size": file_size,
                "saved_at": now.isoformat(),
                "source_file": source_file_path
            }
            step1_data['technical_point_to_point_file'] = point_to_point_file_info

            db.execute_query("""
                UPDATE tender_hitl_tasks
                SET step1_data = ?
                WHERE hitl_task_id = ?
            """, (json.dumps(step1_data), task_id))

            logger.info(f"同步点对点应答文件到HITL任务: {task_id}, 文件: {filename} ({file_size} bytes)")

            return jsonify({
                'success': True,
                'message': '点对点应答文件已成功同步到投标项目',
                'file_path': target_path,
                'filename': filename,
                'file_size': file_size,
                'saved_at': point_to_point_file_info['saved_at']
            })

        except Exception as e:
            logger.error(f"同步点对点应答文件失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'error': f'同步失败: {str(e)}'
            }), 500

    @app.route('/api/tender-processing/sync-tech-proposal/<task_id>', methods=['POST'])
    def sync_tech_proposal_to_hitl(task_id):
        """
        同步技术方案文件到HITL投标项目
        接收技术方案生成的文件路径,复制到HITL任务目录,保存为"应答完成文件"
        """
        try:
            import json
            import shutil
            from common.database import get_knowledge_base_db

            data = request.get_json()
            source_file_path = data.get('file_path')
            output_files = data.get('output_files', {})  # 可能包含多个输出文件

            if not source_file_path:
                return jsonify({
                    'success': False,
                    'error': '未提供文件路径'
                }), 400

            # 如果传入的是下载URL(以/api/downloads/开头),转换为实际文件路径
            if source_file_path.startswith('/api/downloads/'):
                filename = source_file_path.replace('/api/downloads/', '')
                # 使用URL解码处理中文文件名
                from urllib.parse import unquote
                filename = unquote(filename)
                project_root = Path(__file__).parent.parent
                source_file_path = os.path.join(project_root, 'data/outputs', filename)
                logger.info(f"从下载URL转换为文件路径: {source_file_path}")

            # 检查源文件是否存在
            if not os.path.exists(source_file_path):
                return jsonify({
                    'success': False,
                    'error': '源文件不存在'
                }), 404

            logger.info(f"同步技术方案文件到HITL项目: task_id={task_id}, file_path={source_file_path}")

            # 获取数据库实例
            db = get_knowledge_base_db()

            # 查询任务信息
            task_data = db.execute_query("""
                SELECT step1_data FROM tender_hitl_tasks
                WHERE hitl_task_id = ?
            """, (task_id,), fetch_one=True)

            if not task_data:
                return jsonify({
                    'success': False,
                    'error': '任务不存在'
                }), 404

            step1_data = json.loads(task_data['step1_data'])

            # 创建存储目录
            now = datetime.now()
            project_root = Path(__file__).parent.parent
            save_dir = os.path.join(
                project_root,
                'data/uploads/completed_response_files',
                str(now.year),
                f"{now.month:02d}",
                task_id
            )
            os.makedirs(save_dir, exist_ok=True)

            # 生成文件名
            source_filename = os.path.basename(source_file_path)
            # 从源文件名提取,如果包含时间戳则保留,否则添加时间戳
            if '_' in source_filename:
                base_name = source_filename.rsplit('.', 1)[0]
                filename = f"{base_name}_应答完成.docx"
            else:
                filename = f"技术方案_{now.strftime('%Y%m%d_%H%M%S')}_应答完成.docx"

            # 复制文件到目标位置
            target_path = os.path.join(save_dir, filename)
            shutil.copy2(source_file_path, target_path)

            # 计算文件大小
            file_size = os.path.getsize(target_path)

            # 更新任务的step1_data - 使用独立字段存储技术方案文件
            tech_proposal_file_info = {
                "file_path": target_path,
                "filename": filename,
                "file_size": file_size,
                "saved_at": now.isoformat(),
                "source_file": source_file_path,
                "output_files": output_files  # 保存所有输出文件信息
            }
            step1_data['technical_proposal_file'] = tech_proposal_file_info

            db.execute_query("""
                UPDATE tender_hitl_tasks
                SET step1_data = ?
                WHERE hitl_task_id = ?
            """, (json.dumps(step1_data), task_id))

            logger.info(f"同步技术方案文件到HITL任务: {task_id}, 文件: {filename} ({file_size} bytes)")

            return jsonify({
                'success': True,
                'message': '技术方案文件已成功同步到投标项目',
                'file_path': target_path,
                'filename': filename,
                'file_size': file_size,
                'saved_at': tech_proposal_file_info['saved_at']
            })

        except Exception as e:
            logger.error(f"同步技术方案文件失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'error': f'同步失败: {str(e)}'
            }), 500

def main():
    """主函数"""
    app = create_app()
    config = get_config()
    web_config = config.get_web_config()
    
    print(f"启动AI标书系统Web应用...")
    print(f"访问地址: http://{web_config['host']}:{web_config['port']}")
    
    app.run(
        host=web_config['host'],
        port=web_config['port'],
        debug=web_config['debug']
    )

if __name__ == '__main__':
    main()