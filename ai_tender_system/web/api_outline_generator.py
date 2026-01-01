#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
技术方案大纲生成API
提供技术方案生成的Web接口
"""

import os
import json
import traceback
from datetime import datetime
from pathlib import Path
from flask import Blueprint, request, jsonify, send_file, Response, stream_with_context
from werkzeug.utils import secure_filename

# 导入核心模块
import sys
sys.path.append(str(Path(__file__).parent.parent))
from common import get_module_logger, get_config, get_prompt_manager
from modules.outline_generator import (
    RequirementAnalyzer,
    OutlineGenerator,
    ProductMatcher,
    ProposalAssembler,
    WordExporter
)

# 创建蓝图
api_outline_bp = Blueprint('api_outline', __name__, url_prefix='/api')

# 全局变量
logger = get_module_logger("api_outline_generator")
config = get_config()

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {'doc', 'docx', 'pdf', 'xlsx', 'xls'}


def allowed_file(filename: str) -> bool:
    """检查文件扩展名是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_hitl_technical_file_path(project_id: str) -> Path:
    """
    从数据库查询HITL项目的技术需求文件路径

    Args:
        project_id: 项目ID

    Returns:
        技术需求文件路径，如果未找到则返回 None
    """
    if not project_id or project_id == 'default':
        return None

    try:
        from ai_tender_system.common.database import get_knowledge_base_db
        db = get_knowledge_base_db()

        # 查询项目的技术需求文件路径
        db_row = db.execute_query(
            "SELECT step1_data, technical_requirement_path FROM tender_projects WHERE project_id = ?",
            (project_id,),
            fetch_one=True
        )

        if db_row:
            tender_path = None
            # 优先使用 technical_requirement_path
            if db_row[1] if isinstance(db_row, (list, tuple)) else db_row.get('technical_requirement_path'):
                path_value = db_row[1] if isinstance(db_row, (list, tuple)) else db_row.get('technical_requirement_path')
                tender_path = Path(path_value)
                logger.info(f"从 technical_requirement_path 获取文件路径: {tender_path}")
            # 备选：从 step1_data 中获取
            else:
                step1_data_str = db_row[0] if isinstance(db_row, (list, tuple)) else db_row.get('step1_data')
                if step1_data_str:
                    try:
                        step1_data_json = json.loads(step1_data_str)
                        if step1_data_json.get('technical_file_path'):
                            tender_path = Path(step1_data_json['technical_file_path'])
                            logger.info(f"从 step1_data.technical_file_path 获取文件路径: {tender_path}")
                    except Exception as e:
                        logger.warning(f"解析 step1_data 失败: {e}")

            if tender_path and tender_path.exists():
                logger.info(f"找到HITL技术需求文件: {tender_path}")
                return tender_path
            else:
                logger.warning(f"技术需求文件不存在: {tender_path}")
        else:
            logger.warning(f"未找到项目记录: project_id={project_id}")
    except Exception as e:
        logger.warning(f"从数据库查询技术需求文件失败: {e}", exc_info=True)

    return None


def get_tech_proposal_output_dir(project_id=None) -> Path:
    """
    获取技术方案输出目录

    如果有 project_id，保存到 tech_proposal_files/{年}/{月}/{项目ID}/
    否则保存到默认的 outputs/ 目录

    Args:
        project_id: 项目ID（可选）

    Returns:
        输出目录路径
    """
    if project_id and project_id != 'default':
        now = datetime.now()
        output_dir = config.get_path('upload') / 'tech_proposal_files' / str(now.year) / str(now.month).zfill(2) / str(project_id)
    else:
        output_dir = config.get_path('output')

    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


@api_outline_bp.route('/generate-proposal', methods=['POST'])
def generate_proposal():
    """
    生成技术方案API

    请求参数（multipart/form-data）:
    - tender_file: 技术需求文档文件
    - product_file: 产品文档文件（可选）
    - company_id: 公司ID（可选）
    - output_prefix: 输出文件名前缀（可选，默认"技术方案"）
    - include_analysis: 是否包含需求分析报告（可选，默认false）
    - include_mapping: 是否生成需求匹配表（可选，默认false）
    - include_summary: 是否生成生成报告（可选，默认false）

    返回:
    {
        "success": true,
        "requirements_count": 50,
        "sections_count": 5,
        "matches_count": 45,
        "output_files": {
            "proposal": "/downloads/xxx.docx",
            "analysis": "/downloads/xxx.docx",
            "mapping": "/downloads/xxx.xlsx",
            "summary": "/downloads/xxx.txt"
        }
    }
    """
    try:
        logger.info("收到技术方案生成请求")

        # 1. 验证请求 - 支持两种文件来源：直接上传或从HITL传递
        use_hitl_file = request.form.get('use_hitl_technical_file', 'false').lower() == 'true'
        project_id = request.form.get('hitl_task_id') or request.form.get('project_id')

        if use_hitl_file and project_id:
            # 使用HITL传递的技术需求文件 - 从数据库查询路径
            logger.info(f"使用HITL项目的技术需求文件: project_id={project_id}")
            tender_path = get_hitl_technical_file_path(project_id)

            if not tender_path:
                return jsonify({
                    'success': False,
                    'error': f'未找到项目的技术需求文件: project_id={project_id}'
                }), 400
        else:
            # 使用上传的文件
            if 'tender_file' not in request.files:
                return jsonify({
                    'success': False,
                    'error': '缺少技术需求文档文件'
                }), 400

            tender_file = request.files['tender_file']

            if tender_file.filename == '':
                return jsonify({
                    'success': False,
                    'error': '未选择技术需求文档文件'
                }), 400

            if not allowed_file(tender_file.filename):
                return jsonify({
                    'success': False,
                    'error': '不支持的文件格式，请上传 .doc, .docx, .pdf, .xlsx 或 .xls 文件'
                }), 400

            # 保存上传的文件
            upload_dir = config.get_path('uploads')
            upload_dir.mkdir(parents=True, exist_ok=True)

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            tender_filename = secure_filename(f"{timestamp}_{tender_file.filename}")
            tender_path = upload_dir / tender_filename

            tender_file.save(str(tender_path))
            logger.info(f"需求文档已保存: {tender_path}")

        # 2. 获取参数
        company_id = request.form.get('companyId')
        project_name = request.form.get('projectName', '')  # 获取项目名称
        output_prefix = request.form.get('output_prefix', '技术方案')

        options = {
            'include_analysis': request.form.get('includeAnalysis', 'false').lower() == 'true',
            'include_mapping': request.form.get('includeMapping', 'false').lower() == 'true',
            'include_summary': request.form.get('includeSummary', 'false').lower() == 'true'
        }

        logger.info(f"生成选项: {options}")

        # 4. 阶段1：需求分析
        logger.info("开始阶段1：需求分析...")
        analyzer = RequirementAnalyzer()
        analysis_result = analyzer.analyze_document(str(tender_path))

        logger.info("需求分析完成")

        # 5. 阶段2：大纲生成
        logger.info("开始阶段2：大纲生成...")
        outline_gen = OutlineGenerator()
        outline_data = outline_gen.generate_outline(
            analysis_result,
            project_name=output_prefix
        )

        logger.info("大纲生成完成")

        # 6. 阶段3：产品文档匹配
        logger.info("开始阶段3：产品文档匹配...")
        matcher = ProductMatcher()
        matched_docs = matcher.match_documents(
            analysis_result.get('requirement_categories', []),
            company_id=int(company_id) if company_id else None
        )

        logger.info(f"产品文档匹配完成，共匹配 {sum(len(v) for v in matched_docs.values())} 份文档")

        # 7. 阶段4：方案组装
        logger.info("开始阶段4：方案组装...")
        assembler = ProposalAssembler()
        proposal = assembler.assemble_proposal(
            outline_data,
            analysis_result,
            matched_docs,
            options
        )

        logger.info("方案组装完成")

        # 8. 导出文件
        logger.info("开始导出文件...")
        exporter = WordExporter()

        output_dir = config.get_path('output')
        output_dir.mkdir(parents=True, exist_ok=True)

        output_files = {}

        # 智能文件命名：有项目名称时使用"项目名_类型_时间"，无项目名称时使用"前缀_时间"
        if project_name:
            proposal_filename = f"{project_name}_技术方案_{timestamp}.docx"
            analysis_filename = f"{project_name}_需求分析_{timestamp}.docx"
            mapping_filename = f"{project_name}_需求匹配表_{timestamp}.xlsx"
            summary_filename = f"{project_name}_生成报告_{timestamp}.txt"
        else:
            proposal_filename = f"{output_prefix}_{timestamp}.docx"
            analysis_filename = f"{output_prefix}_需求分析_{timestamp}.docx"
            mapping_filename = f"{output_prefix}_需求匹配表_{timestamp}.xlsx"
            summary_filename = f"{output_prefix}_生成报告_{timestamp}.txt"

        # 导出主方案（简洁模式，不显示大纲指导信息）
        proposal_path = output_dir / proposal_filename
        exporter.export_proposal(proposal, str(proposal_path), show_guidance=False)
        output_files['proposal'] = f"/api/downloads/{proposal_filename}"

        logger.info(f"主方案已导出: {proposal_path}")

        # 导出附件
        if options['include_analysis']:
            analysis_path = output_dir / analysis_filename
            exporter.export_analysis_report(analysis_result, str(analysis_path))
            output_files['analysis'] = f"/api/downloads/{analysis_filename}"

            logger.info(f"需求分析报告已导出: {analysis_path}")

        if options['include_mapping']:
            mapping_path = output_dir / mapping_filename

            mapping_data = []
            for attachment in proposal.get('attachments', []):
                if attachment['type'] == 'mapping':
                    mapping_data = attachment['data']
                    break

            exporter.export_mapping_table(mapping_data, str(mapping_path))
            output_files['mapping'] = f"/api/downloads/{mapping_filename}"

            logger.info(f"需求匹配表已导出: {mapping_path}")

        if options['include_summary']:
            summary_path = output_dir / summary_filename

            summary_data = {}
            for attachment in proposal.get('attachments', []):
                if attachment['type'] == 'summary':
                    summary_data = attachment['data']
                    break

            exporter.export_summary_report(summary_data, str(summary_path))
            output_files['summary'] = f"/api/downloads/{summary_filename}"

            logger.info(f"生成报告已导出: {summary_path}")

        # 9. 统计信息
        requirements_count = analysis_result.get('document_summary', {}).get('total_requirements', 0)
        sections_count = len(outline_data.get('chapters', []))
        matches_count = sum(len(docs) for docs in matched_docs.values())

        # 10. 返回结果
        response = {
            'success': True,
            'requirements_count': requirements_count,
            'features_count': 0,  # 暂时为0
            'sections_count': sections_count,
            'matches_count': matches_count,
            'output_files': output_files
        }

        logger.info("技术方案生成成功")
        return jsonify(response)

    except Exception as e:
        logger.error(f"技术方案生成失败: {e}", exc_info=True)
        error_trace = traceback.format_exc()

        return jsonify({
            'success': False,
            'error': str(e),
            'trace': error_trace if config.get('debug', False) else None
        }), 500


@api_outline_bp.route('/generate-proposal-stream', methods=['POST'])
def generate_proposal_stream():
    """
    生成技术方案API（流式SSE版本）
    实时推送生成进度

    请求参数（multipart/form-data）:
    - tender_file: 技术需求文档文件
    - product_file: 产品文档文件（可选）
    - outputPrefix: 输出文件名前缀
    - companyId: 公司ID
    - projectName: 项目名称（可选，从HITL传递）
    - technicalFileTaskId: HITL任务ID（可选）
    - includeAnalysis: 是否包含需求分析
    - includeMapping: 是否生成匹配表
    - includeSummary: 是否生成总结报告

    返回: text/event-stream
    """

    def generate_events():
        """生成SSE事件流"""
        try:
            # 发送初始进度
            yield f"data: {json.dumps({'stage': 'init', 'progress': 0, 'message': '准备生成技术方案...'}, ensure_ascii=False)}\n\n"

            # 1. 参数解析（复用原有逻辑）
            yield f"data: {json.dumps({'stage': 'init', 'progress': 5, 'message': '解析请求参数...'}, ensure_ascii=False)}\n\n"

            tender_file = request.files.get('tender_file')
            product_file = request.files.get('product_file')
            output_prefix = request.form.get('outputPrefix', '技术方案')
            company_id = request.form.get('companyId')
            project_name = request.form.get('projectName', '')
            project_id = request.form.get('technicalFileTaskId', '') or request.form.get('projectId', '')

            # 生成选项
            options = {
                'include_analysis': request.form.get('includeAnalysis', 'false').lower() == 'true',
                'include_mapping': request.form.get('includeMapping', 'false').lower() == 'true',
                'include_summary': request.form.get('includeSummary', 'false').lower() == 'true'
            }

            # 2. 获取技术需求文件路径
            if project_id:
                # 从HITL项目加载 - 使用数据库查询
                yield f"data: {json.dumps({'stage': 'init', 'progress': 10, 'message': f'从投标项目加载技术需求文件 (project_id={project_id})...'}, ensure_ascii=False)}\n\n"

                tender_path = get_hitl_technical_file_path(project_id)

                if not tender_path:
                    logger.error(f'未找到项目的技术需求文件: project_id={project_id}')
                    raise ValueError(f'未找到项目的技术需求文件: project_id={project_id}')
            elif tender_file:
                # 上传文件
                yield f"data: {json.dumps({'stage': 'init', 'progress': 10, 'message': '保存上传的技术需求文件...'}, ensure_ascii=False)}\n\n"
                if not allowed_file(tender_file.filename):
                    raise ValueError('文件类型不支持')

                upload_dir = config.get_path('uploads') / 'tender_processing' / datetime.now().strftime('%Y/%m')
                upload_dir.mkdir(parents=True, exist_ok=True)

                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = secure_filename(f"{timestamp}_{tender_file.filename}")
                tender_path = upload_dir / filename
                tender_file.save(str(tender_path))
            else:
                raise ValueError('未提供技术需求文档文件')

            # 3. 阶段1：需求分析
            yield f"data: {json.dumps({'stage': 'analysis', 'progress': 15, 'message': '🔍 正在分析技术需求文档...'}, ensure_ascii=False)}\n\n"

            analyzer = RequirementAnalyzer()
            analysis_result = analyzer.analyze_document(str(tender_path))

            yield f"data: {json.dumps({'stage': 'analysis', 'progress': 30, 'message': '✓ 需求分析完成'}, ensure_ascii=False)}\n\n"

            # 发送完整的需求分析结果供前端展示
            try:
                # 确保analysis_result可以被JSON序列化
                analysis_result_serializable = json.loads(json.dumps(analysis_result, ensure_ascii=False, default=str))
                yield f"data: {json.dumps({'stage': 'analysis_completed', 'analysis_result': analysis_result_serializable}, ensure_ascii=False)}\n\n"
            except Exception as e:
                logger.warning(f"无法序列化需求分析结果: {e}, 跳过前端展示")
                # 继续执行,不影响后续流程

            # 4. 阶段2：大纲生成
            yield f"data: {json.dumps({'stage': 'outline', 'progress': 35, 'message': '📝 正在生成技术方案大纲...'}, ensure_ascii=False)}\n\n"

            outline_gen = OutlineGenerator()
            outline_data = outline_gen.generate_outline(
                analysis_result,
                project_name=output_prefix
            )

            yield f"data: {json.dumps({'stage': 'outline', 'progress': 55, 'message': '✓ 大纲生成完成'}, ensure_ascii=False)}\n\n"

            # 发送完整的大纲数据供前端展示
            try:
                # 确保outline_data可以被JSON序列化
                outline_data_serializable = json.loads(json.dumps(outline_data, ensure_ascii=False, default=str))
                yield f"data: {json.dumps({'stage': 'outline_completed', 'outline_data': outline_data_serializable}, ensure_ascii=False)}\n\n"
            except Exception as e:
                logger.warning(f"无法序列化大纲数据: {e}, 跳过前端展示")
                # 继续执行,不影响后续流程

            # 5. 阶段3：产品文档匹配
            yield f"data: {json.dumps({'stage': 'matching', 'progress': 60, 'message': '🔗 正在匹配产品文档...'}, ensure_ascii=False)}\n\n"

            matcher = ProductMatcher()
            matched_docs = matcher.match_documents(
                analysis_result.get('requirement_categories', []),
                company_id=int(company_id) if company_id else None
            )

            matches_count = sum(len(v) for v in matched_docs.values())
            yield f"data: {json.dumps({'stage': 'matching', 'progress': 70, 'message': f'✓ 文档匹配完成（匹配到 {matches_count} 份文档）'}, ensure_ascii=False)}\n\n"

            # 6. 阶段4：方案组装
            yield f"data: {json.dumps({'stage': 'assembly', 'progress': 75, 'message': '⚙️ 正在组装技术方案...'}, ensure_ascii=False)}\n\n"

            assembler = ProposalAssembler()
            proposal = assembler.assemble_proposal(
                outline_data,
                analysis_result,
                matched_docs,
                options
            )

            yield f"data: {json.dumps({'stage': 'assembly', 'progress': 85, 'message': '✓ 方案组装完成'}, ensure_ascii=False)}\n\n"

            # 7. 导出文件
            yield f"data: {json.dumps({'stage': 'export', 'progress': 90, 'message': '💾 正在导出文件...'}, ensure_ascii=False)}\n\n"

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            # 使用统一的输出目录函数，直接保存到 tech_proposal_files/{年}/{月}/{项目ID}/
            output_dir = get_tech_proposal_output_dir(project_id)

            exporter = WordExporter()
            output_files = {}

            # 文件命名：项目ID_项目名称_类型_时间戳（项目ID确保唯一性）
            project_id_str = f"P{project_id}" if project_id and project_id != 'default' else ''
            name_part = project_name if project_name else output_prefix
            if project_id_str:
                proposal_filename = f"{project_id_str}_{name_part}_技术方案_{timestamp}.docx"
                analysis_filename = f"{project_id_str}_{name_part}_需求分析_{timestamp}.docx"
                mapping_filename = f"{project_id_str}_{name_part}_需求匹配表_{timestamp}.xlsx"
                summary_filename = f"{project_id_str}_{name_part}_生成报告_{timestamp}.txt"
            else:
                proposal_filename = f"{name_part}_技术方案_{timestamp}.docx"
                analysis_filename = f"{name_part}_需求分析_{timestamp}.docx"
                mapping_filename = f"{name_part}_需求匹配表_{timestamp}.xlsx"
                summary_filename = f"{name_part}_生成报告_{timestamp}.txt"

            # 导出主方案（简洁模式，不显示大纲指导信息）
            proposal_path = output_dir / proposal_filename
            exporter.export_proposal(proposal, str(proposal_path), show_guidance=False)
            output_files['proposal'] = f"/api/downloads/{proposal_filename}"

            # 导出附件
            if options['include_analysis']:
                analysis_path = output_dir / analysis_filename
                exporter.export_analysis_report(analysis_result, str(analysis_path))
                output_files['analysis'] = f"/api/downloads/{analysis_filename}"

            if options['include_mapping']:
                mapping_path = output_dir / mapping_filename
                mapping_data = []
                for attachment in proposal.get('attachments', []):
                    if attachment['type'] == 'mapping':
                        mapping_data = attachment['data']
                        break
                exporter.export_mapping_table(mapping_data, str(mapping_path))
                output_files['mapping'] = f"/api/downloads/{mapping_filename}"

            if options['include_summary']:
                summary_path = output_dir / summary_filename
                summary_data = {}
                for attachment in proposal.get('attachments', []):
                    if attachment['type'] == 'summary':
                        summary_data = attachment['data']
                        break
                exporter.export_summary_report(summary_data, str(summary_path))
                output_files['summary'] = f"/api/downloads/{summary_filename}"

            # 统计信息
            requirements_count = analysis_result.get('document_summary', {}).get('total_requirements', 0)
            sections_count = len(outline_data.get('chapters', []))

            # 8. 完成
            result = {
                'stage': 'completed',
                'progress': 100,
                'message': '✅ 技术方案生成成功！',
                'success': True,
                'requirements_count': requirements_count,
                'features_count': 0,
                'sections_count': sections_count,
                'matches_count': matches_count,
                'output_file': str(proposal_path),  # ✅ 添加文件系统完整路径，与商务应答/点对点应答保持一致
                'output_files': output_files  # 保留下载URL字典
            }

            yield f"data: {json.dumps(result, ensure_ascii=False)}\n\n"
            logger.info("技术方案生成成功（SSE流式）")

        except Exception as e:
            logger.error(f"技术方案生成失败（SSE流式）: {e}", exc_info=True)
            error_data = {
                'stage': 'error',
                'progress': 0,
                'message': f'生成失败: {str(e)}',
                'success': False,
                'error': str(e)
            }
            yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate_events()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no'
        }
    )


@api_outline_bp.route('/downloads/<filename>', methods=['GET'])
def download_file(filename):
    """
    下载生成的文件

    Args:
        filename: 文件名

    Returns:
        文件内容
    """
    try:
        output_dir = config.get_path('output')
        file_path = output_dir / filename

        if not file_path.exists():
            return jsonify({
                'success': False,
                'error': '文件不存在'
            }), 404

        # 根据文件扩展名设置MIME类型
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

        mime_types = {
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'txt': 'text/plain',
            'pdf': 'application/pdf'
        }

        mimetype = mime_types.get(ext, 'application/octet-stream')

        return send_file(
            str(file_path),
            as_attachment=True,
            download_name=filename,
            mimetype=mimetype
        )

    except Exception as e:
        logger.error(f"文件下载失败: {e}", exc_info=True)

        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@api_outline_bp.route('/generate-proposal-stream-v2', methods=['POST'])
def generate_proposal_stream_v2():
    """
    生成技术方案API（流式SSE版本 V2 - 支持实时内容推送）
    实时推送生成进度和章节内容

    请求参数（multipart/form-data）:
    - tender_file: 技术需求文档文件
    - outputPrefix: 输出文件名前缀
    - companyId: 公司ID
    - projectName: 项目名称
    - projectId: 项目ID（从HITL传递）
    - includeAnalysis: 是否包含需求分析
    - includeMapping: 是否生成匹配表
    - includeSummary: 是否生成总结报告
    - useStreamingContent: 是否使用流式内容生成（默认true）

    返回: text/event-stream
    """

    def generate_events():
        """生成SSE事件流"""
        try:
            # 初始化
            yield f"data: {json.dumps({'stage': 'init', 'progress': 0, 'message': '准备生成技术方案...'}, ensure_ascii=False)}\n\n"

            # 参数解析
            yield f"data: {json.dumps({'stage': 'init', 'progress': 5, 'message': '解析请求参数...'}, ensure_ascii=False)}\n\n"

            tender_file = request.files.get('tender_file')
            output_prefix = request.form.get('outputPrefix', '技术方案')
            company_id = request.form.get('companyId')
            project_name = request.form.get('projectName', '')
            project_id = request.form.get('projectId', '')
            ai_model = request.form.get('aiModel', 'shihuang-gpt4o-mini')  # ✅ 获取AI模型参数，默认gpt4o-mini
            use_streaming_content = request.form.get('useStreamingContent', 'true').lower() == 'true'
            proposal_mode = request.form.get('proposalMode', 'basic')  # ✅ 获取方案模式参数，默认basic

            logger.info(f"使用AI模型: {ai_model}, 方案模式: {proposal_mode}")

            # 生成选项
            options = {
                'include_analysis': request.form.get('includeAnalysis', 'false').lower() == 'true',
                'include_mapping': request.form.get('includeMapping', 'false').lower() == 'true',
                'include_summary': request.form.get('includeSummary', 'false').lower() == 'true'
            }

            # 获取技术需求文件路径
            if project_id:
                yield f"data: {json.dumps({'stage': 'init', 'progress': 10, 'message': f'从投标项目加载技术需求文件...'}, ensure_ascii=False)}\n\n"

                tender_path = get_hitl_technical_file_path(project_id)

                if not tender_path:
                    raise ValueError(f'未找到项目的技术需求文件: project_id={project_id}')
            elif tender_file:
                yield f"data: {json.dumps({'stage': 'init', 'progress': 10, 'message': '保存上传的技术需求文件...'}, ensure_ascii=False)}\n\n"

                if not allowed_file(tender_file.filename):
                    raise ValueError('文件类型不支持')

                upload_dir = config.get_path('uploads') / 'tender_processing' / datetime.now().strftime('%Y/%m')
                upload_dir.mkdir(parents=True, exist_ok=True)

                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = secure_filename(f"{timestamp}_{tender_file.filename}")
                tender_path = upload_dir / filename
                tender_file.save(str(tender_path))
            else:
                raise ValueError('未提供技术需求文档文件')

            # 阶段1：需求分析
            yield f"data: {json.dumps({'stage': 'analysis', 'progress': 15, 'message': '🔍 正在分析技术需求文档...'}, ensure_ascii=False)}\n\n"

            analyzer = RequirementAnalyzer(model_name=ai_model)  # ✅ 传递AI模型参数
            analysis_result = analyzer.analyze_document(str(tender_path))

            yield f"data: {json.dumps({'stage': 'analysis', 'progress': 30, 'message': '✓ 需求分析完成'}, ensure_ascii=False)}\n\n"

            # 阶段2：大纲生成
            yield f"data: {json.dumps({'stage': 'outline', 'progress': 35, 'message': '📝 正在生成技术方案大纲...'}, ensure_ascii=False)}\n\n"

            outline_gen = OutlineGenerator(model_name=ai_model)  # ✅ 传递AI模型参数
            outline_data = outline_gen.generate_outline(analysis_result, project_name=output_prefix)

            yield f"data: {json.dumps({'stage': 'outline', 'progress': 55, 'message': '✓ 大纲生成完成'}, ensure_ascii=False)}\n\n"

            # 发送完整的大纲数据供前端展示（与V1接口保持一致）
            try:
                outline_data_serializable = json.loads(json.dumps(outline_data, ensure_ascii=False, default=str))
                yield f"data: {json.dumps({'stage': 'outline_completed', 'outline_data': outline_data_serializable}, ensure_ascii=False)}\n\n"
            except Exception as e:
                logger.warning(f"无法序列化大纲数据: {e}, 跳过前端展示")
                # 继续执行，不影响后续流程

            # 阶段3：产品文档匹配
            yield f"data: {json.dumps({'stage': 'matching', 'progress': 60, 'message': '🔗 正在匹配产品文档...'}, ensure_ascii=False)}\n\n"

            matcher = ProductMatcher()
            matched_docs = matcher.match_documents(
                analysis_result.get('requirement_categories', []),
                company_id=int(company_id) if company_id else None
            )

            matches_count = sum(len(v) for v in matched_docs.values())
            yield f"data: {json.dumps({'stage': 'matching', 'progress': 70, 'message': f'✓ 文档匹配完成（匹配到 {matches_count} 份文档）'}, ensure_ascii=False)}\n\n"

            # 阶段4：方案组装（流式）
            yield f"data: {json.dumps({'stage': 'assembly', 'progress': 75, 'message': '⚙️ 正在组装技术方案...'}, ensure_ascii=False)}\n\n"

            assembler = ProposalAssembler(model_name=ai_model)  # ✅ 传递AI模型参数

            # 选择流式或非流式组装
            if use_streaming_content:
                logger.info("使用流式内容生成模式")
                proposal = None

                for event in assembler.assemble_proposal_stream(outline_data, analysis_result, matched_docs, options, proposal_mode):
                    event_type = event.get('type')

                    if event_type == 'chapter_start':
                        # 推送章节开始
                        chapter_start_data = {
                            'stage': 'content_generation',
                            'event': 'chapter_start',
                            'chapter_number': event.get('chapter_number', ''),
                            'chapter_title': event.get('chapter_title', ''),
                            'message': f"📄 开始生成 {event.get('chapter_number', '')} {event.get('chapter_title', '')}..."
                        }
                        yield f"data: {json.dumps(chapter_start_data, ensure_ascii=False)}\n\n"

                    elif event_type == 'content_chunk':
                        # 推送内容片段
                        content_chunk_data = {
                            'stage': 'content_generation',
                            'event': 'content_chunk',
                            'chapter_number': event.get('chapter_number', ''),
                            'content': event.get('chunk', '')
                        }
                        yield f"data: {json.dumps(content_chunk_data, ensure_ascii=False)}\n\n"

                    elif event_type == 'chapter_end':
                        # 推送章节完成
                        chapter_end_data = {
                            'stage': 'content_generation',
                            'event': 'chapter_end',
                            'chapter_number': event.get('chapter_number', ''),
                            'message': f"✓ {event.get('chapter_title', '')} 生成完成"
                        }
                        yield f"data: {json.dumps(chapter_end_data, ensure_ascii=False)}\n\n"

                    elif event_type == 'completed':
                        # 保存方案数据
                        proposal = event['proposal']

                    elif event_type == 'error':
                        raise Exception(event['error'])

                if not proposal:
                    raise ValueError("流式组装未返回完整方案")

                yield f"data: {json.dumps({'stage': 'assembly', 'progress': 85, 'message': '✓ 方案组装完成（流式）'}, ensure_ascii=False)}\n\n"
            else:
                # 使用非流式组装（原有逻辑）
                proposal = assembler.assemble_proposal(outline_data, analysis_result, matched_docs, options, proposal_mode)
                yield f"data: {json.dumps({'stage': 'assembly', 'progress': 85, 'message': '✓ 方案组装完成'}, ensure_ascii=False)}\n\n"

            # 导出文件
            yield f"data: {json.dumps({'stage': 'export', 'progress': 90, 'message': '💾 正在导出文件...'}, ensure_ascii=False)}\n\n"

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            # 使用统一的输出目录函数，直接保存到 tech_proposal_files/{年}/{月}/{项目ID}/
            output_dir = get_tech_proposal_output_dir(project_id)

            exporter = WordExporter()
            output_files = {}

            # 文件命名：项目ID_项目名称_类型_时间戳（项目ID确保唯一性）
            project_id_str = f"P{project_id}" if project_id and project_id != 'default' else ''
            name_part = project_name if project_name else output_prefix
            if project_id_str:
                proposal_filename = f"{project_id_str}_{name_part}_技术方案_{timestamp}.docx"
                analysis_filename = f"{project_id_str}_{name_part}_需求分析_{timestamp}.docx"
                mapping_filename = f"{project_id_str}_{name_part}_需求匹配表_{timestamp}.xlsx"
                summary_filename = f"{project_id_str}_{name_part}_生成报告_{timestamp}.txt"
            else:
                proposal_filename = f"{name_part}_技术方案_{timestamp}.docx"
                analysis_filename = f"{name_part}_需求分析_{timestamp}.docx"
                mapping_filename = f"{name_part}_需求匹配表_{timestamp}.xlsx"
                summary_filename = f"{name_part}_生成报告_{timestamp}.txt"

            # 导出主方案（简洁模式，不显示大纲指导信息）
            proposal_path = output_dir / proposal_filename
            exporter.export_proposal(proposal, str(proposal_path), show_guidance=False)
            output_files['proposal'] = f"/api/downloads/{proposal_filename}"

            # 导出附件
            if options['include_analysis']:
                analysis_path = output_dir / analysis_filename
                exporter.export_analysis_report(analysis_result, str(analysis_path))
                output_files['analysis'] = f"/api/downloads/{analysis_filename}"

            if options['include_mapping']:
                mapping_path = output_dir / mapping_filename
                mapping_data = []
                for attachment in proposal.get('attachments', []):
                    if attachment['type'] == 'mapping':
                        mapping_data = attachment['data']
                        break
                exporter.export_mapping_table(mapping_data, str(mapping_path))
                output_files['mapping'] = f"/api/downloads/{mapping_filename}"

            if options['include_summary']:
                summary_path = output_dir / summary_filename
                summary_data = {}
                for attachment in proposal.get('attachments', []):
                    if attachment['type'] == 'summary':
                        summary_data = attachment['data']
                        break
                exporter.export_summary_report(summary_data, str(summary_path))
                output_files['summary'] = f"/api/downloads/{summary_filename}"

            # 统计信息
            requirements_count = analysis_result.get('document_summary', {}).get('total_requirements', 0)
            sections_count = len(outline_data.get('chapters', []))

            # 完成
            result = {
                'stage': 'completed',
                'progress': 100,
                'message': '✅ 技术方案生成成功！',
                'success': True,
                'requirements_count': requirements_count,
                'features_count': 0,
                'sections_count': sections_count,
                'matches_count': matches_count,
                'output_file': str(proposal_path),
                'output_files': output_files,
                'streaming_mode': use_streaming_content
            }

            yield f"data: {json.dumps(result, ensure_ascii=False)}\n\n"
            logger.info("技术方案生成成功（SSE流式V2）")

        except Exception as e:
            logger.error(f"技术方案生成失败（SSE流式V2）: {e}", exc_info=True)
            error_data = {
                'stage': 'error',
                'progress': 0,
                'message': f'生成失败: {str(e)}',
                'success': False,
                'error': str(e)
            }
            yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate_events()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no'
        }
    )


@api_outline_bp.route('/prompts/outline-generation', methods=['GET'])
def get_outline_generation_prompts():
    """
    获取大纲生成提示词配置

    返回:
    {
        "success": true,
        "prompts": {
            "analyze_requirements": "...",
            "generate_outline": "...",
            "generate_response_suggestions": "...",
            "recommend_product_docs": "..."
        }
    }
    """
    try:
        logger.info("获取大纲生成提示词配置")

        # 获取 prompt manager
        prompt_manager = get_prompt_manager()

        # 获取所有大纲生成相关的提示词
        prompts = {
            'analyze_requirements': prompt_manager.get_prompt('outline_generation', 'analyze_requirements'),
            'generate_outline': prompt_manager.get_prompt('outline_generation', 'generate_outline'),
            'generate_response_suggestions': prompt_manager.get_prompt('outline_generation', 'generate_response_suggestions'),
            'recommend_product_docs': prompt_manager.get_prompt('outline_generation', 'recommend_product_docs')
        }

        return jsonify({
            'success': True,
            'prompts': prompts
        })

    except Exception as e:
        logger.error(f"获取提示词配置失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ==================== 智能体API ====================

@api_outline_bp.route('/agent/generate', methods=['POST'])
def generate_with_agent():
    """
    使用智能体生成技术方案API（SSE流式响应）

    请求参数（multipart/form-data 或 JSON）:
    - generation_mode: 生成模式 ("按评分点写" | "按招标书目录写" | "编写专项章节")
    - tender_file: 招标文档文件 (FormData) 或 tender_doc: 文本 (JSON)
    - page_count: 目标页数
    - content_style: 内容风格配置(JSON字符串或对象)
    - template_name: 模板名称 (编写专项章节模式必填)
    - projectId: 项目ID (可选，从HITL加载)

    返回: SSE流式事件
    - stage: init/analysis/analysis_completed/outline_completed/export/completed/error
    - progress: 0-100
    - message: 进度消息
    - 各阶段特定数据
    """
    # 先解析请求参数（在generator外部，避免上下文问题）
    try:
        from ai_tender_system.modules.outline_generator.agents import AgentRouter

        logger.info("【智能体API】收到请求")

        # 支持两种格式: FormData 和 JSON
        is_json_request = request.is_json
        if is_json_request:
            # JSON格式
            req_data = request.json
            generation_mode = req_data.get('generation_mode')
            tender_doc = req_data.get('tender_doc', '')
            page_count = int(req_data.get('page_count', 200))
            content_style = req_data.get('content_style', {})
            template_name = req_data.get('template_name', '政府采购标准')
            scoring_points = req_data.get('scoring_points')
            project_id = req_data.get('projectId', 'default')
            project_name = req_data.get('projectName', '')
            output_prefix = req_data.get('outputPrefix', '技术方案')
        else:
            # FormData格式
            generation_mode = request.form.get('generation_mode')
            tender_doc = ''
            page_count = int(request.form.get('page_count', 200))
            template_name = request.form.get('template_name', '政府采购标准')
            scoring_points = None
            project_id = request.form.get('projectId') or request.form.get('project_id') or 'default'
            project_name = request.form.get('projectName', '')
            output_prefix = request.form.get('outputPrefix', '技术方案')

            # content_style可能是JSON字符串
            content_style_str = request.form.get('content_style', '')
            if content_style_str:
                try:
                    content_style = json.loads(content_style_str)
                except:
                    content_style = {'tables': '适量', 'flowcharts': '流程图', 'images': '少量'}
            else:
                content_style = {'tables': '适量', 'flowcharts': '流程图', 'images': '少量'}

            # 从HITL或上传文件获取招标文档
            use_hitl_file = request.form.get('use_hitl_technical_file', 'false').lower() == 'true'

            if use_hitl_file and project_id and project_id != 'default':
                # 使用辅助函数从数据库查询技术需求文件路径
                logger.info(f"[智能体API] 从HITL项目加载技术需求文件: project_id={project_id}")
                tender_path = get_hitl_technical_file_path(project_id)

                if tender_path:
                    logger.info(f"[智能体API] 找到HITL技术需求文件: {tender_path}")
                    from docx import Document
                    doc = Document(str(tender_path))
                    content_parts = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_parts.append(text)
                    for table in doc.tables:
                        content_parts.append('\n[表格内容]')
                        for tbl_row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in tbl_row.cells if cell.text.strip()])
                            if row_text:
                                content_parts.append(row_text)
                        content_parts.append('[表格结束]\n')
                    tender_doc = '\n'.join(content_parts)
                else:
                    logger.warning(f"[智能体API] 未找到HITL技术需求文件: project_id={project_id}")

            elif 'tender_file' in request.files:
                tender_file = request.files['tender_file']
                if allowed_file(tender_file.filename):
                    upload_dir = config.get_path('uploads')
                    upload_dir.mkdir(parents=True, exist_ok=True)
                    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = secure_filename(f"{ts}_{tender_file.filename}")
                    tender_path = upload_dir / filename
                    tender_file.save(str(tender_path))

                    from docx import Document
                    doc = Document(str(tender_path))
                    content_parts = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_parts.append(text)
                    for table in doc.tables:
                        content_parts.append('\n[表格内容]')
                        for row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                content_parts.append(row_text)
                        content_parts.append('[表格结束]\n')
                    tender_doc = '\n'.join(content_parts)
                    logger.info(f"已解析上传文件: {tender_path}, 文本长度: {len(tender_doc)}")

        # 参数验证
        param_error = None
        if not generation_mode:
            param_error = '缺少必填参数: generation_mode'
        elif not tender_doc:
            param_error = '缺少必填参数: tender_doc 或 tender_file'

    except Exception as e:
        param_error = f'参数解析失败: {str(e)}'
        generation_mode = None
        tender_doc = None
        page_count = 200
        content_style = {}
        template_name = '政府采购标准'
        scoring_points = None
        project_id = 'default'
        project_name = ''
        output_prefix = '技术方案'

    def generate_events():
        """生成SSE事件流"""
        nonlocal param_error, generation_mode, tender_doc, page_count, content_style
        nonlocal template_name, scoring_points, project_id, project_name, output_prefix

        try:
            # 检查参数错误
            if param_error:
                yield f"data: {json.dumps({'stage': 'error', 'error': param_error, 'message': f'❌ {param_error}'}, ensure_ascii=False)}\n\n"
                return

            # 1. 初始化阶段
            yield f"data: {json.dumps({'stage': 'init', 'progress': 5, 'message': '🚀 开始生成技术方案...'}, ensure_ascii=False)}\n\n"

            # 2. 解析参数完成
            yield f"data: {json.dumps({'stage': 'init', 'progress': 10, 'message': '📄 参数解析完成'}, ensure_ascii=False)}\n\n"

            # 3. 需求分析阶段
            yield f"data: {json.dumps({'stage': 'analysis', 'progress': 20, 'message': '🔍 正在分析技术需求...'}, ensure_ascii=False)}\n\n"

            # 创建路由器并生成
            from ai_tender_system.modules.outline_generator.agents import AgentRouter
            router = AgentRouter()

            logger.info(f"调用智能体路由: mode={generation_mode}, pages={page_count}, template={template_name}")

            result = router.route(
                generation_mode=generation_mode,
                tender_doc=tender_doc,
                page_count=page_count,
                content_style=content_style,
                scoring_points=scoring_points,
                template_name=template_name
            )

            logger.info(f"【智能体API】生成成功，模式: {generation_mode}")

            # 4. 分析完成
            yield f"data: {json.dumps({'stage': 'analysis', 'progress': 35, 'message': '✓ 需求分析完成'}, ensure_ascii=False)}\n\n"

            # 发送分析结果（如果有）
            analysis_result = result.get('analysis', {})
            if analysis_result:
                yield f"data: {json.dumps({'stage': 'analysis_completed', 'progress': 40, 'analysis_result': analysis_result}, ensure_ascii=False)}\n\n"

            # 5. 大纲生成完成
            outline_data = result.get('outline', {})
            chapters = result.get('chapters', [])
            yield f"data: {json.dumps({'stage': 'outline', 'progress': 50, 'message': '📝 正在生成大纲...'}, ensure_ascii=False)}\n\n"

            # 构建大纲数据供前端显示
            outline_for_frontend = {
                'chapters': outline_data.get('chapters', []),
                'total_chapters': len(chapters),
                'estimated_pages': result.get('metadata', {}).get('total_pages', page_count)
            }
            yield f"data: {json.dumps({'stage': 'outline_completed', 'progress': 60, 'outline_data': outline_for_frontend}, ensure_ascii=False)}\n\n"

            # 6. 导出Word文档
            yield f"data: {json.dumps({'stage': 'export', 'progress': 75, 'message': '💾 正在导出Word文档...'}, ensure_ascii=False)}\n\n"

            from ai_tender_system.modules.outline_generator import WordExporter

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            # 使用统一的输出目录函数，直接保存到 tech_proposal_files/{年}/{月}/{项目ID}/
            output_dir = get_tech_proposal_output_dir(project_id)

            # 文件命名：项目ID_项目名称_类型_时间戳（项目ID确保唯一性）
            project_id_str = f"P{project_id}" if project_id and project_id != 'default' else ''
            name_part = project_name if project_name else output_prefix
            if project_id_str:
                proposal_filename = f"{project_id_str}_{name_part}_技术方案_{timestamp}.docx"
            else:
                proposal_filename = f"{name_part}_技术方案_{timestamp}.docx"

            proposal_path = output_dir / proposal_filename

            logger.info(f"【智能体API】文件将保存到: {proposal_path}")

            # 构建proposal数据结构
            proposal_data = {
                'metadata': {
                    'title': outline_data.get('title', '技术方案'),
                    'generation_time': timestamp,
                    'total_chapters': len(chapters),
                    'estimated_pages': result.get('metadata', {}).get('estimated_pages', 0)
                },
                'chapters': []
            }

            for chapter in chapters:
                # 类型安全检查
                if not isinstance(chapter, dict):
                    logger.warning(f"【智能体API】章节不是字典类型: {type(chapter)}")
                    if isinstance(chapter, list) and len(chapter) > 0:
                        chapter = chapter[0] if isinstance(chapter[0], dict) else {}
                    else:
                        chapter = {}

                # 智能体返回的字段是 'title' 和 'content'
                content = chapter.get('content', '')
                if isinstance(content, list):
                    content = '\n'.join(str(item) for item in content if item)
                elif not isinstance(content, str):
                    content = str(content) if content else ''

                chapter_data = {
                    'level': chapter.get('level', 1),
                    'chapter_number': chapter.get('chapter_number', ''),
                    'title': chapter.get('title', ''),  # 智能体返回 'title' 而非 'chapter_title'
                    'ai_generated_content': content,
                    'subsections': chapter.get('subsections', []) if isinstance(chapter.get('subsections'), list) else []
                }
                proposal_data['chapters'].append(chapter_data)

            logger.info(f"【智能体API】构建的章节数量: {len(proposal_data['chapters'])}")
            if proposal_data['chapters']:
                first_chapter = proposal_data['chapters'][0]
                logger.info(f"【智能体API】第一章标题: {first_chapter.get('title', 'N/A')}")
                logger.info(f"【智能体API】第一章内容长度: {len(first_chapter.get('ai_generated_content', ''))}")

            exporter = WordExporter()
            exporter.export_proposal(proposal_data, str(proposal_path), show_guidance=False)

            logger.info(f"【智能体API】Word文件已导出: {proposal_path}")

            yield f"data: {json.dumps({'stage': 'export', 'progress': 90, 'message': '✓ Word文档导出完成'}, ensure_ascii=False)}\n\n"

            # 7. 完成
            completed_data = {
                'stage': 'completed',
                'progress': 100,
                'success': True,
                'message': '✅ 技术方案生成成功！',
                'output_file': str(proposal_path),
                'output_files': {
                    'proposal': f"/api/downloads/{proposal_filename}"
                },
                'sections_count': len(chapters),
                'requirements_count': result.get('metadata', {}).get('requirement_categories_count', 0),
                'coverage_rate': result.get('metadata', {}).get('coverage_rate', 0)
            }
            yield f"data: {json.dumps(completed_data, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"【智能体API】生成失败: {e}", exc_info=True)
            error_data = {
                'stage': 'error',
                'error': str(e),
                'message': f'❌ 生成失败: {str(e)}'
            }
            yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate_events()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no'
        }
    )


@api_outline_bp.route('/agent/templates', methods=['GET'])
def list_templates():
    """
    获取可用模板列表

    返回:
    {
        "success": true,
        "data": [
            {
                "template_id": "政府采购标准",
                "name": "政府采购标准模板",
                "description": "...",
                "chapters_count": 5,
                "chapters": [...]
            },
            ...
        ]
    }
    """
    try:
        from ai_tender_system.modules.outline_generator.agents import TemplateAgent

        templates = TemplateAgent.list_templates()

        return jsonify({
            'success': True,
            'data': templates
        })

    except Exception as e:
        logger.error(f"获取模板列表失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@api_outline_bp.route('/agent/modes', methods=['GET'])
def get_generation_modes():
    """
    获取可用的生成模式

    返回:
    {
        "success": true,
        "data": {
            "modes": [...],
            "templates": [...]
        }
    }
    """
    try:
        from ai_tender_system.modules.outline_generator.agents import AgentRouter

        router = AgentRouter()
        modes_info = router.get_available_modes()

        return jsonify({
            'success': True,
            'data': modes_info
        })

    except Exception as e:
        logger.error(f"获取生成模式失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@api_outline_bp.route('/agent/scoring/extract', methods=['POST'])
def extract_scoring_points():
    """
    从招标文档中提取评分点（场景1）

    请求参数（JSON）:
    {
        "tender_doc": "招标文档文本"
    }

    返回:
    {
        "success": true,
        "data": {
            "scoring_points": [...]
        }
    }
    """
    try:
        from ai_tender_system.modules.outline_generator.agents import ScoringPointAgent

        data = request.json
        tender_doc = data.get('tender_doc', '')

        if not tender_doc:
            return jsonify({
                'success': False,
                'error': '缺少必填参数: tender_doc'
            }), 400

        agent = ScoringPointAgent()
        scoring_points = agent.extract_scoring_points(tender_doc)

        return jsonify({
            'success': True,
            'data': {
                'scoring_points': scoring_points
            }
        })

    except Exception as e:
        logger.error(f"评分点提取失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@api_outline_bp.route('/agent/scoring/evaluate', methods=['POST'])
def evaluate_tender():
    """
    AI评分API（场景2）

    请求参数（JSON）:
    {
        "tender_doc": "标书文档文本",
        "scoring_points": [...]
    }

    返回:
    {
        "success": true,
        "data": {
            "overall_score": 85.5,
            "dimension_scores": [...],
            "risk_analysis": {...},
            "improvement_suggestions": [...]
        }
    }
    """
    try:
        from ai_tender_system.modules.outline_generator.agents import ScoringPointAgent

        data = request.json
        tender_doc = data.get('tender_doc', '')
        scoring_points = data.get('scoring_points', [])

        if not tender_doc or not scoring_points:
            return jsonify({
                'success': False,
                'error': '缺少必填参数: tender_doc 或 scoring_points'
            }), 400

        agent = ScoringPointAgent()
        result = agent.evaluate_tender_document(tender_doc, scoring_points)

        return jsonify({
            'success': True,
            'data': result
        })

    except Exception as e:
        logger.error(f"AI评分失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@api_outline_bp.route('/agent/generate-crew', methods=['POST'])
def generate_with_crew():
    """
    使用 ProposalCrew 生成技术方案（Quality-First 模式）

    请求参数（multipart/form-data 或 JSON）:
    - tender_file / tender_doc: 招标文档
    - page_count: 目标页数 (默认200)
    - projectId: 项目ID
    - projectName: 项目名称
    - companyId: 公司ID
    - aiModel: AI模型
    - crew_config: CrewConfig配置（JSON字符串）
      - skip_product_matching: bool
      - skip_material_retrieval: bool
      - enable_expert_review: bool
      - max_iterations: int
      - min_review_score: float
    - content_style: 内容风格配置

    返回: SSE流式事件
    """
    # 先解析请求参数（在generator外部，避免上下文问题）
    try:
        logger.info("【Quality-First API】收到请求")

        # 支持两种格式: FormData 和 JSON
        is_json_request = request.is_json
        if is_json_request:
            req_data = request.json
            tender_doc = req_data.get('tender_doc', '')
            page_count = int(req_data.get('page_count', 200))
            company_id = req_data.get('companyId', 1)
            project_id = req_data.get('projectId', 'default')
            project_name = req_data.get('projectName', '')
            customer_name = ''  # 客户名称（招标方）
            ai_model = req_data.get('aiModel', 'shibing624-gpt4o-mini')
            crew_config_raw = req_data.get('crew_config', {})
            content_style = req_data.get('content_style', {})
        else:
            # FormData格式
            tender_doc = ''
            page_count = int(request.form.get('page_count', 200))
            company_id = request.form.get('companyId', 1)
            project_id = request.form.get('projectId') or request.form.get('project_id') or 'default'
            project_name = request.form.get('projectName', '')
            customer_name = ''  # 客户名称（招标方）
            ai_model = request.form.get('aiModel', 'shibing624-gpt4o-mini')

            # crew_config 是 JSON 字符串
            crew_config_str = request.form.get('crew_config', '{}')
            try:
                crew_config_raw = json.loads(crew_config_str)
            except:
                crew_config_raw = {}

            # content_style 是 JSON 字符串
            content_style_str = request.form.get('content_style', '{}')
            try:
                content_style = json.loads(content_style_str)
            except:
                content_style = {'tables': '适量', 'flowcharts': '流程图', 'images': '少量'}

            # 从HITL或上传文件获取招标文档
            use_hitl_file = request.form.get('use_hitl_technical_file', 'false').lower() == 'true'

            if use_hitl_file and project_id and project_id != 'default':
                # 使用辅助函数从数据库查询技术需求文件路径
                logger.info(f"[Quality-First API] 从HITL项目加载技术需求文件: project_id={project_id}")
                tender_path = get_hitl_technical_file_path(project_id)

                if tender_path:
                    logger.info(f"[Quality-First API] 找到HITL技术需求文件: {tender_path}")
                    from docx import Document
                    doc = Document(str(tender_path))
                    content_parts = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_parts.append(text)
                    for table in doc.tables:
                        content_parts.append('\n[表格内容]')
                        for tbl_row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in tbl_row.cells if cell.text.strip()])
                            if row_text:
                                content_parts.append(row_text)
                        content_parts.append('[表格结束]\n')
                    tender_doc = '\n'.join(content_parts)
                else:
                    logger.warning(f"[Quality-First API] 未找到HITL技术需求文件: project_id={project_id}")

            elif 'tender_file' in request.files:
                tender_file = request.files['tender_file']
                if allowed_file(tender_file.filename):
                    upload_dir = config.get_path('uploads')
                    upload_dir.mkdir(parents=True, exist_ok=True)
                    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = secure_filename(f"{ts}_{tender_file.filename}")
                    tender_path = upload_dir / filename
                    tender_file.save(str(tender_path))

                    from docx import Document
                    doc = Document(str(tender_path))
                    content_parts = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_parts.append(text)
                    for table in doc.tables:
                        content_parts.append('\n[表格内容]')
                        for row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                content_parts.append(row_text)
                        content_parts.append('[表格结束]\n')
                    tender_doc = '\n'.join(content_parts)
                    logger.info(f"已解析上传文件: {tender_path}, 文本长度: {len(tender_doc)}")

        # 从数据库查询项目信息（项目名称和客户名称）
        if project_id and project_id != 'default':
            try:
                from ai_tender_system.common.database import get_db_connection
                with get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT project_name, tenderer FROM tender_projects WHERE project_id = ?",
                        (project_id,)
                    )
                    row = cursor.fetchone()
                    if row:
                        # 使用数据库中的值作为后备（如果前端没传）
                        if not project_name and row[0]:
                            project_name = row[0]
                        if row[1]:
                            customer_name = row[1]
                        logger.info(f"从数据库获取项目信息: project_name={project_name}, customer_name={customer_name}")
            except Exception as e:
                logger.warning(f"查询项目信息失败: {e}")

        # 参数验证
        param_error = None
        if not tender_doc:
            param_error = '缺少必填参数: tender_doc 或 tender_file'

    except Exception as e:
        param_error = f'参数解析失败: {str(e)}'
        tender_doc = ''
        page_count = 200
        company_id = 1
        project_id = 'default'
        project_name = ''
        customer_name = ''
        ai_model = 'shibing624-gpt4o-mini'
        crew_config_raw = {}
        content_style = {}

    def generate_events():
        """生成SSE事件流"""
        nonlocal param_error, tender_doc, page_count, company_id, project_id
        nonlocal project_name, customer_name, ai_model, crew_config_raw, content_style

        try:
            # 检查参数错误
            if param_error:
                yield f"data: {json.dumps({'stage': 'error', 'status': 'error', 'error': param_error, 'message': f'❌ {param_error}'}, ensure_ascii=False)}\n\n"
                return

            # 初始化事件
            yield f"data: {json.dumps({'stage': 'init', 'status': 'running', 'progress': 0, 'message': '🚀 初始化 Quality-First 模式...'}, ensure_ascii=False)}\n\n"

            # 创建 ProposalCrew
            from ai_tender_system.modules.outline_generator.agents import ProposalCrew, CrewConfig

            crew_config = CrewConfig(
                model_name=ai_model,
                company_id=int(company_id) if company_id else 1,
                skip_product_matching=crew_config_raw.get('skip_product_matching', False),
                skip_material_retrieval=crew_config_raw.get('skip_material_retrieval', False),
                enable_expert_review=crew_config_raw.get('enable_expert_review', True),
                max_iterations=crew_config_raw.get('max_iterations', 2),
                min_review_score=crew_config_raw.get('min_review_score', 85.0),
                target_word_count=page_count * 700,
                content_style=content_style,
                project_name=project_name,      # 项目名称
                customer_name=customer_name     # 客户名称（招标方）
            )

            crew = ProposalCrew(crew_config)

            logger.info(f"【Quality-First】开始生成: pages={page_count}, company_id={company_id}")

            # 阶段进度映射
            phase_progress = {
                'scoring_extraction': 12,
                'product_matching': 22,
                'strategy_planning': 32,
                'material_retrieval': 42,
                'outline_generation': 55,
                'content_writing': 75,
                'expert_review': 90,
                'iteration': 95,
                'complete': 100
            }

            phase_messages = {
                'scoring_extraction': '🎯 正在提取评分点...',
                'product_matching': '🔗 正在匹配产品能力...',
                'strategy_planning': '📊 正在制定评分策略...',
                'material_retrieval': '📚 正在检索历史素材...',
                'outline_generation': '📝 正在生成技术方案大纲...',
                'content_writing': '✍️ 正在撰写方案内容...',
                'expert_review': '👨‍🔬 专家评审中...',
                'iteration': '🔄 正在迭代优化...',
                'complete': '✅ 技术方案生成完成！'
            }

            # 流式运行 ProposalCrew
            for event in crew.run_stream(tender_doc, page_count):
                phase = event.get('phase', 'unknown')
                status = event.get('status', 'running')

                # 处理错误事件
                if phase == 'error' or status == 'error':
                    error_data = {
                        'stage': 'error',
                        'status': 'error',
                        'progress': 0,
                        'error': event.get('error', '未知错误'),
                        'message': event.get('message', f"❌ {event.get('error', '未知错误')}"),
                        'traceback': event.get('traceback', '')
                    }
                    yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"
                    return  # 遇到错误就停止

                # 转换为前端期望的格式
                sse_data = {
                    'stage': phase,
                    'phase': phase,  # 同时传递 phase 字段供前端使用
                    'status': status,
                    'progress': phase_progress.get(phase, 0),
                    'message': event.get('message') or phase_messages.get(phase, '')
                }

                # ========================================
                # 传递细粒度进度详情 (新增)
                # ========================================
                if 'detail' in event:
                    sse_data['detail'] = event['detail']

                # ========================================
                # 传递章节流式内容预览 (新增)
                # ========================================
                if 'content' in event:
                    sse_data['content'] = event['content']

                # 附加阶段特定数据
                if status == 'complete' and 'result' in event:
                    result = event['result']
                    sse_data['result'] = result  # 同时传递原始 result
                    if phase == 'scoring_extraction':
                        sse_data['scoring_points'] = result
                    elif phase == 'product_matching':
                        sse_data['product_match'] = result
                    elif phase == 'strategy_planning':
                        sse_data['scoring_strategy'] = result
                    elif phase == 'material_retrieval':
                        sse_data['materials'] = result
                    elif phase == 'outline_generation':
                        sse_data['outline_data'] = result
                        # 兼容现有前端
                        sse_data['stage'] = 'outline_completed'
                    elif phase == 'expert_review':
                        sse_data['review_result'] = result

                # 内容撰写的进度事件
                if phase == 'content_writing' and event.get('event') == 'chapter_progress':
                    chapter_idx = event.get('chapter_index', 0)
                    total_chapters = event.get('total_chapters', 1)
                    sse_data['progress'] = 55 + int(chapter_idx / total_chapters * 20)
                    sse_data['chapter_title'] = event.get('chapter_title', '')
                    sse_data['event'] = 'chapter_progress'

                # 章节完成事件
                if phase == 'content_writing' and event.get('event') == 'chapter_complete':
                    sse_data['chapter'] = event.get('chapter')
                    sse_data['event'] = 'chapter_complete'

                yield f"data: {json.dumps(sse_data, ensure_ascii=False)}\n\n"

            # 获取最终结果
            final_result = crew._build_final_result()

            if final_result.get('success'):
                # 导出 Word 文档
                from ai_tender_system.modules.outline_generator import WordExporter

                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                # 使用统一的输出目录函数，直接保存到 tech_proposal_files/{年}/{月}/{项目ID}/
                output_dir = get_tech_proposal_output_dir(project_id)

                # 文件命名：项目ID_项目名称_技术方案_时间戳.docx（项目ID确保唯一性）
                project_id_str = f"P{project_id}" if project_id and project_id != 'default' else ''
                name_part = project_name if project_name else '技术方案'
                proposal_filename = f"{project_id_str}_{name_part}_技术方案_{timestamp}.docx" if project_id_str else f"{name_part}_{timestamp}.docx"
                proposal_path = output_dir / proposal_filename

                exporter = WordExporter()
                proposal_content = final_result.get('proposal_content', {})

                # 转换章节格式以匹配 WordExporter 的期望
                def convert_chapter(chapter, prefix: str = '') -> dict:
                    """将 ContentWriterAgent 输出转换为 WordExporter 期望的格式"""
                    # 类型安全检查：确保 chapter 是字典
                    if not isinstance(chapter, dict):
                        logger.warning(f"convert_chapter: 收到非字典类型的chapter: {type(chapter)}")
                        # 如果是列表，尝试取第一个元素
                        if isinstance(chapter, list) and len(chapter) > 0:
                            chapter = chapter[0] if isinstance(chapter[0], dict) else {}
                        else:
                            chapter = {}

                    # 生成章节编号
                    chapter_num = prefix or chapter.get('id', '') or ''

                    # 安全提取 content，确保是字符串
                    content = chapter.get('content', '')
                    if isinstance(content, dict):
                        content = content.get('content', '') or content.get('text', '') or str(content)
                    elif isinstance(content, list):
                        # 如果 content 是列表，合并为字符串
                        content = '\n'.join(str(item) for item in content if item)
                    if not isinstance(content, str):
                        content = str(content) if content else ''

                    converted = {
                        'chapter_number': chapter_num,
                        'title': chapter.get('title', ''),
                        'level': chapter.get('level', 1),
                        'ai_generated_content': content,
                        'subsections': []
                    }

                    # 递归转换子章节
                    children = chapter.get('children', [])
                    if isinstance(children, list):
                        for i, child in enumerate(children, 1):
                            child_prefix = f"{chapter_num}.{i}" if chapter_num else str(i)
                            converted['subsections'].append(convert_chapter(child, child_prefix))

                    return converted

                # 转换所有章节
                converted_chapters = []
                chapters_data = proposal_content.get('chapters', [])
                if isinstance(chapters_data, list):
                    for i, chapter in enumerate(chapters_data, 1):
                        converted_chapters.append(convert_chapter(chapter, str(i)))
                else:
                    logger.warning(f"proposal_content['chapters'] 不是列表: {type(chapters_data)}")

                # 构建 WordExporter 期望的数据结构（添加 metadata）
                proposal_for_export = {
                    'metadata': {
                        'title': project_name or 'Quality-First技术方案',
                        'generation_time': timestamp,
                        'total_chapters': len(converted_chapters),
                        'estimated_pages': final_result.get('metadata', {}).get('total_words', 0) // 700
                    },
                    'chapters': converted_chapters
                }
                exporter.export_proposal(proposal_for_export, str(proposal_path), show_guidance=False)

                logger.info(f"【Quality-First】导出完成: {proposal_path}")

                # 发送完成事件
                completed_data = {
                    'stage': 'completed',
                    'status': 'complete',
                    'progress': 100,
                    'success': True,
                    'message': '✅ Quality-First 技术方案生成完成！',
                    'output_file': str(proposal_path),  # 实际文件系统路径，用于同步到HITL
                    'output_files': {
                        'proposal': f'/api/downloads/{proposal_filename}'  # 下载URL
                    },
                    'final_result': {
                        'total_words': final_result.get('metadata', {}).get('total_words', 0),
                        'chapter_count': final_result.get('metadata', {}).get('chapter_count', 0),
                        'final_score': final_result.get('metadata', {}).get('final_score', 0),
                        'pass_recommendation': final_result.get('metadata', {}).get('pass_recommendation', False),
                        'elapsed_time': final_result.get('elapsed_time', 0)
                    }
                }
                yield f"data: {json.dumps(completed_data, ensure_ascii=False)}\n\n"
            else:
                yield f"data: {json.dumps({'stage': 'error', 'status': 'error', 'error': final_result.get('error', '未知错误')}, ensure_ascii=False)}\n\n"

        except Exception as e:
            import traceback
            error_msg = str(e) if str(e) else '未知错误'
            error_trace = traceback.format_exc()
            logger.error(f"【Quality-First API】生成失败: {error_msg}\n{error_trace}")
            error_data = {
                'stage': 'error',
                'status': 'error',
                'progress': 0,
                'error': error_msg,
                'message': f'❌ 生成失败: {error_msg}',
                'traceback': error_trace[:500] if error_trace else ''  # 截断避免过长
            }
            yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate_events()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no'
        }
    )


@api_outline_bp.route('/tech-proposal-files', methods=['GET'])
def list_tech_proposal_files():
    """
    获取技术方案文件列表

    查询参数:
    - project_id: 项目ID（可选，如果提供则只返回该项目的文件）

    返回:
    {
        "success": true,
        "files": [
            {
                "filename": "xxx_技术方案_20251222_110612_技术方案.docx",
                "size": "46K",
                "date": "2025-12-22 11:06:12",
                "download_url": "/api/files/download/xxx.docx",
                "project_id": "61",
                "created": "2025-12-22T11:06:12",
                "modified": "2025-12-22T11:06:12"
            },
            ...
        ]
    }
    """
    try:
        import os
        from datetime import datetime

        def format_size(size_bytes):
            """格式化文件大小"""
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024.0:
                    return f"{size_bytes:.1f} {unit}"
                size_bytes /= 1024.0
            return f"{size_bytes:.1f} TB"

        project_id = request.args.get('project_id')
        files = []

        # 技术方案文件保存在 tech_proposal_files/{年份}/{月份}/{项目ID}/ 目录下
        tech_proposal_base = config.get_path('upload') / 'tech_proposal_files'

        if tech_proposal_base.exists():
            # 遍历年份目录
            for year_dir in tech_proposal_base.glob('*'):
                if not year_dir.is_dir():
                    continue

                # 遍历月份目录
                for month_dir in year_dir.glob('*'):
                    if not month_dir.is_dir():
                        continue

                    # 遍历项目目录
                    for proj_dir in month_dir.glob('*'):
                        if not proj_dir.is_dir():
                            continue

                        # 如果指定了project_id,只处理该项目
                        if project_id and proj_dir.name != str(project_id):
                            continue

                        # 遍历项目目录下的文件
                        for file_path in proj_dir.glob('*.docx'):
                            try:
                                stat = file_path.stat()
                                modified_time = datetime.fromtimestamp(stat.st_mtime)

                                # 构建相对于uploads目录的路径用于下载
                                relative_path = file_path.relative_to(config.get_path('upload'))

                                files.append({
                                    'filename': file_path.name,
                                    'size': format_size(stat.st_size),
                                    'date': modified_time.strftime('%Y-%m-%d %H:%M:%S'),
                                    'download_url': f'/api/files/download/{file_path.name}',
                                    'file_path': str(file_path),  # 完整路径用于预览
                                    'project_id': proj_dir.name,
                                    'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                                    'modified': modified_time.isoformat()
                                })
                            except Exception as e:
                                logger.warning(f"读取文件信息失败 {file_path}: {e}")

        # 按修改时间倒序排列
        files.sort(key=lambda x: x.get('modified', ''), reverse=True)

        logger.info(f"找到{len(files)}个技术方案文件" + (f" (project_id={project_id})" if project_id else ""))

        return jsonify({'success': True, 'files': files})

    except Exception as e:
        logger.error(f"获取技术方案文件列表失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


# =========================
# 任务管理API（断点续传支持）
# =========================

@api_outline_bp.route('/tasks', methods=['GET'])
def list_tasks():
    """
    获取任务列表

    查询参数:
    - project_id: 项目ID（可选）
    - status: 任务状态过滤（可选: pending/running/completed/failed）
    - limit: 返回数量限制（默认20）

    返回:
    {
        "success": true,
        "tasks": [...],
        "total": 10
    }
    """
    try:
        from modules.outline_generator.task_manager import get_task_manager

        task_manager = get_task_manager()
        project_id = request.args.get('project_id', type=int)
        limit = request.args.get('limit', 20, type=int)

        if project_id:
            tasks = task_manager.get_tasks_by_project(project_id, limit=limit)
        else:
            # 获取所有可恢复的任务
            tasks = task_manager.get_resumable_tasks()

        return jsonify({
            'success': True,
            'tasks': tasks,
            'total': len(tasks)
        })

    except Exception as e:
        logger.error(f"获取任务列表失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@api_outline_bp.route('/tasks/<task_id>', methods=['GET'])
def get_task_detail(task_id):
    """
    获取任务详情

    返回:
    {
        "success": true,
        "task": {...},
        "logs": [...],
        "stats": {...}
    }
    """
    try:
        from modules.outline_generator.task_manager import get_task_manager

        task_manager = get_task_manager()

        task = task_manager.get_task(task_id)
        if not task:
            return jsonify({'success': False, 'error': '任务不存在'}), 404

        logs = task_manager.get_execution_logs(task_id)
        stats = task_manager.get_task_stats(task_id)

        return jsonify({
            'success': True,
            'task': task,
            'logs': logs,
            'stats': stats
        })

    except Exception as e:
        logger.error(f"获取任务详情失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@api_outline_bp.route('/tasks/<task_id>/resume', methods=['POST'])
def resume_task(task_id):
    """
    恢复任务执行

    返回:
    {
        "success": true,
        "task_id": "xxx",
        "message": "任务已恢复，请连接流式接口获取进度"
    }
    """
    try:
        from modules.outline_generator.task_manager import get_task_manager

        task_manager = get_task_manager()

        # 检查任务是否可恢复
        if not task_manager.can_resume(task_id):
            task = task_manager.get_task(task_id)
            if not task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404
            return jsonify({
                'success': False,
                'error': f'任务状态为 {task.get("overall_status")}，无法恢复'
            }), 400

        return jsonify({
            'success': True,
            'task_id': task_id,
            'message': '任务已准备恢复，请调用 /api/agent/generate-with-recovery 接口'
        })

    except Exception as e:
        logger.error(f"恢复任务失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@api_outline_bp.route('/tasks/<task_id>/cancel', methods=['POST'])
def cancel_task(task_id):
    """取消任务"""
    try:
        from modules.outline_generator.task_manager import get_task_manager

        task_manager = get_task_manager()

        if task_manager.cancel_task(task_id):
            return jsonify({'success': True, 'message': '任务已取消'})
        else:
            return jsonify({'success': False, 'error': '取消失败'}), 400

    except Exception as e:
        logger.error(f"取消任务失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@api_outline_bp.route('/agent/generate-with-recovery', methods=['POST'])
def generate_with_recovery():
    """
    支持断点恢复的技术方案生成API

    请求参数:
    - task_id: 任务ID（可选，用于恢复已有任务）
    - tender_file: 招标文件（新任务必需）
    - project_id: 项目ID
    - company_id: 公司ID
    - page_count: 目标页数
    - ai_model: AI模型
    - crew_config: Crew配置（JSON字符串）

    返回:
    SSE流式响应
    """
    try:
        from modules.outline_generator.task_manager import get_task_manager
        from modules.outline_generator.agents.proposal_crew import ProposalCrew, CrewConfig

        task_manager = get_task_manager()

        # 获取参数
        task_id = request.form.get('task_id')
        project_id = request.form.get('projectId', type=int)
        company_id = request.form.get('companyId', 1, type=int)
        page_count = request.form.get('page_count', 200, type=int)
        ai_model = request.form.get('aiModel', 'deepseek-v3')

        # Crew配置
        crew_config_str = request.form.get('crew_config', '{}')
        try:
            crew_config_dict = json.loads(crew_config_str)
        except:
            crew_config_dict = {}

        # 处理招标文件
        tender_doc = None
        tender_file_path = None

        if task_id:
            # 恢复模式：从任务中获取文件路径
            task = task_manager.get_task(task_id)
            if not task:
                return jsonify({'success': False, 'error': '任务不存在'}), 404

            # 并发锁：使用原子操作抢占任务，防止重复恢复
            # 只有非 running 状态才能抢占
            lock_acquired = task_manager.try_acquire_task_lock(task_id)
            if not lock_acquired:
                return jsonify({
                    'success': False,
                    'error': '任务正在执行中，请勿重复操作'
                }), 409  # 409 Conflict

            tender_file_path = task.get('tender_file_path')
            if tender_file_path and Path(tender_file_path).exists():
                # 读取文件内容
                tender_doc = read_document_content(tender_file_path)
            else:
                # 释放锁
                task_manager.fail_task(task_id, '招标文件不存在')
                return jsonify({'success': False, 'error': '招标文件不存在'}), 400

            project_id = task.get('project_id')
            company_id = task.get('company_id', 1)

        else:
            # 新任务模式
            if 'tender_file' not in request.files:
                return jsonify({'success': False, 'error': '缺少招标文件'}), 400

            tender_file = request.files['tender_file']

            # 保存文件
            upload_dir = config.get_path('upload') / 'tech_proposal_uploads'
            upload_dir.mkdir(parents=True, exist_ok=True)

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_name = secure_filename(tender_file.filename)
            saved_filename = f"{timestamp}_{safe_name}"
            tender_file_path = str(upload_dir / saved_filename)
            tender_file.save(tender_file_path)

            # 读取文件内容
            tender_doc = read_document_content(tender_file_path)

            # 创建新任务
            task_id = task_manager.create_task(
                project_id=project_id or 0,
                company_id=company_id,
                generation_mode="quality_first",
                ai_model=ai_model,
                crew_config=crew_config_dict,
                tender_file_path=tender_file_path,
                page_count=page_count
            )

        if not tender_doc:
            return jsonify({'success': False, 'error': '无法读取招标文件内容'}), 400

        # 创建Crew配置
        crew_config = CrewConfig(
            model_name=ai_model,
            company_id=company_id,
            skip_product_matching=crew_config_dict.get('skip_product_matching', False),
            skip_material_retrieval=crew_config_dict.get('skip_material_retrieval', False),
            enable_expert_review=crew_config_dict.get('enable_expert_review', True),
            max_iterations=crew_config_dict.get('max_iterations', 2),
            min_review_score=crew_config_dict.get('min_review_score', 70.0)
        )

        # 创建Crew
        crew = ProposalCrew(config=crew_config, task_manager=task_manager)

        def generate_events():
            try:
                # 发送任务ID
                yield f"data: {json.dumps({'stage': 'init', 'task_id': task_id, 'message': '任务已创建'}, ensure_ascii=False)}\n\n"

                # 使用支持恢复的流式运行
                for event in crew.run_stream_with_recovery(task_id, tender_doc, page_count):
                    phase = event.get('phase', 'unknown')
                    status = event.get('status', 'running')

                    # 转换为前端期望的格式
                    sse_data = {
                        'stage': phase,
                        'phase': phase,  # 同时传递 phase 字段供前端使用
                        'status': status,
                        'task_id': task_id,
                        'message': event.get('message', ''),
                        'can_resume': event.get('can_resume', False)
                    }

                    # 传递细粒度进度详情 (新增)
                    if 'detail' in event:
                        sse_data['detail'] = event['detail']

                    # 传递章节流式内容预览 (新增)
                    if 'content' in event:
                        sse_data['content'] = event['content']

                    # 添加结果数据
                    if 'result' in event:
                        sse_data['result'] = event['result']
                    if 'error' in event:
                        sse_data['error'] = event['error']

                    yield f"data: {json.dumps(sse_data, ensure_ascii=False)}\n\n"

            except Exception as e:
                import traceback
                error_msg = str(e) if str(e) else '未知错误'
                error_trace = traceback.format_exc()
                logger.error(f"生成失败: {error_msg}\n{error_trace}")

                yield f"data: {json.dumps({'stage': 'error', 'status': 'error', 'error': error_msg, 'task_id': task_id, 'can_resume': True}, ensure_ascii=False)}\n\n"

        return Response(
            stream_with_context(generate_events()),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no'
            }
        )

    except Exception as e:
        logger.error(f"生成任务失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


def read_document_content(file_path: str) -> str:
    """读取文档内容"""
    try:
        from modules.outline_generator import RequirementAnalyzer

        # 使用RequirementAnalyzer的文档读取功能
        analyzer = RequirementAnalyzer()

        # 解析文档
        file_ext = Path(file_path).suffix.lower()
        if file_ext in ['.docx', '.doc']:
            return analyzer._read_docx(file_path)
        elif file_ext == '.pdf':
            return analyzer._read_pdf(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return analyzer._read_excel(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

    except Exception as e:
        logger.error(f"读取文档失败: {e}")
        return None


# 导出蓝图
__all__ = ['api_outline_bp']
