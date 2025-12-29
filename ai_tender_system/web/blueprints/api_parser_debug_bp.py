#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
目录解析方法对比调试 API
功能：
- 上传标书文档并运行所有解析方法
- 对比不同解析方法的准确率
- 支持人工标注正确答案
- 计算准确率指标（P/R/F1）
"""

import os
import json
import uuid
import time
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from flask import Blueprint, request, jsonify, send_file, Response
from werkzeug.utils import secure_filename
from docx import Document

from common import get_module_logger, get_config
from common.database import get_knowledge_base_db
from modules.tender_processing.structure_parser import DocumentStructureParser, ChapterNode
from modules.tender_processing.level_analyzer import LevelAnalyzer

# 尝试导入 Azure 解析器
try:
    from modules.tender_processing.azure_parser import AzureDocumentParser, is_azure_available
    AZURE_PARSER_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Azure 解析器不可用: {e}")
    AZURE_PARSER_AVAILABLE = False

    def is_azure_available():
        return False

# 尝试导入 Gemini 解析器
try:
    from modules.tender_processing.parsers.gemini_parser import GeminiParser
    GEMINI_PARSER_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Gemini 解析器不可用: {e}")
    GEMINI_PARSER_AVAILABLE = False

logger = get_module_logger("api_parser_debug")

api_parser_debug_bp = Blueprint('api_parser_debug', __name__, url_prefix='/api/parser-debug')


class ParserDebugger:
    """解析方法对比调试器"""

    def __init__(self, doc_path: str):
        """
        初始化调试器

        Args:
            doc_path: Word文档路径
        """
        self.doc_path = doc_path
        self.parser = DocumentStructureParser()
        self.doc = Document(doc_path)

        # 文档信息
        self.total_paragraphs = len(self.doc.paragraphs)
        self.has_toc = False
        self.toc_items_count = 0
        self.toc_start_idx = None
        self.toc_end_idx = None

        # 计算文档总字数 (去除空格)
        self.total_chars = sum(len(p.text.replace(' ', '').replace('\t', '')) for p in self.doc.paragraphs)

        # 预先检测目录
        self._detect_toc_info()

    def _detect_toc_info(self):
        """检测目录信息"""
        try:
            toc_idx = self.parser._find_toc_section(self.doc)
            if toc_idx is not None:
                self.has_toc = True
                self.toc_start_idx = toc_idx
                toc_items, toc_end_idx = self.parser._parse_toc_items(self.doc, toc_idx)
                self.toc_items_count = len(toc_items)
                self.toc_end_idx = toc_end_idx
                logger.info(f"检测到目录: {self.toc_items_count} 项，位于段落 {toc_idx}-{toc_end_idx}")
        except Exception as e:
            logger.warning(f"目录检测失败: {e}")

    def get_document_info(self) -> Dict:
        """获取文档基本信息"""
        return {
            'filename': Path(self.doc_path).name,
            'total_paragraphs': self.total_paragraphs,
            'total_chars': self.total_chars,  # 文档总字数
            'has_toc': self.has_toc,
            'toc_items_count': self.toc_items_count,
            'toc_start_idx': self.toc_start_idx,
            'toc_end_idx': self.toc_end_idx
        }

    def run_all_methods(self) -> Dict:
        """
        运行所有解析方法

        Returns:
            {
                'semantic': {...},
                'style': {...},
                'azure': {...},  # 可选
                'docx_native': {...},
                'gemini': {...}  # 可选
            }
        """
        results = {}

        # 方法4: Azure Form Recognizer（如果可用）
        if is_azure_available() and AZURE_PARSER_AVAILABLE:
            results['azure'] = self._run_with_timing(
                self._run_azure_parser,
                "Azure Form Recognizer"
            )
        else:
            results['azure'] = {
                'success': False,
                'error': 'Azure Form Recognizer 未配置或SDK未安装',
                'chapters': [],
                'method_name': 'Azure Form Recognizer',
                'performance': {'elapsed': 0}
            }

        # 方法5: Word大纲级别识别
        results['docx_native'] = self._run_with_timing(
            self._run_docx_native,
            "Word大纲级别识别"
        )

        # 方法6: Gemini AI解析器（如果可用）
        if GEMINI_PARSER_AVAILABLE:
            try:
                gemini_parser = GeminiParser()
                if gemini_parser.is_available():
                    results['gemini'] = self._run_with_timing(
                        lambda: self._run_gemini_parser(gemini_parser),
                        "Gemini AI解析器"
                    )
                else:
                    results['gemini'] = {
                        'success': False,
                        'error': 'Gemini API密钥未配置',
                        'chapters': [],
                        'method_name': 'Gemini AI解析器',
                        'performance': {'elapsed': 0}
                    }
            except Exception as e:
                logger.error(f"初始化Gemini解析器失败: {e}")
                results['gemini'] = {
                    'success': False,
                    'error': f'Gemini解析器初始化失败: {str(e)}',
                    'chapters': [],
                    'method_name': 'Gemini AI解析器',
                    'performance': {'elapsed': 0}
                }
        else:
            results['gemini'] = {
                'success': False,
                'error': 'Gemini SDK未安装 (pip install google-generativeai)',
                'chapters': [],
                'method_name': 'Gemini AI解析器',
                'performance': {'elapsed': 0}
            }

        return results

    def _run_with_timing(self, method_func, method_name: str) -> Dict:
        """
        运行方法并计时

        Args:
            method_func: 要运行的方法
            method_name: 方法名称（用于日志）

        Returns:
            包含结果和性能指标的字典
        """
        logger.info(f"开始运行: {method_name}")
        start_time = time.time()

        try:
            result = method_func()
            elapsed = time.time() - start_time

            result['performance'] = {
                'elapsed': round(elapsed, 3),
                'elapsed_formatted': f"{elapsed:.3f}s"
            }

            logger.info(f"{method_name} 完成，耗时 {elapsed:.3f}s")
            return result

        except Exception as e:
            elapsed = time.time() - start_time
            logger.error(f"{method_name} 失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'chapters': [],
                'performance': {'elapsed': round(elapsed, 3)}
            }

    def _run_toc_exact_match(self) -> Dict:
        """方法0: 精确匹配(基于目录) - 直接使用目录定位章节"""
        return self.parser.parse_by_toc_exact(self.doc_path)

    def _run_azure_parser(self) -> Dict:
        """方法4: Azure Form Recognizer 解析"""
        try:
            azure_parser = AzureDocumentParser()
            result = azure_parser.parse_document_structure(self.doc_path)
            return result
        except Exception as e:
            logger.error(f"Azure 解析失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise

    def _run_docx_native(self) -> Dict:
        """方法5: Word大纲级别识别（微软官方API）"""
        return self.parser.parse_by_outline_level(self.doc_path)

    def _run_gemini_parser(self, gemini_parser: 'GeminiParser') -> Dict:
        """方法6: Gemini AI解析器"""
        try:
            # 调用Gemini解析器的parse_structure方法
            result = gemini_parser.parse_structure(self.doc_path)

            # Gemini返回的结果已经是标准格式,包含success, chapters, statistics等
            # 但需要确保chapters是dict格式而不是ChapterNode对象
            if result.get('success') and result.get('chapters'):
                # 如果chapters是ChapterNode对象列表,需要转换为dict
                chapters = result['chapters']
                if chapters and hasattr(chapters[0], 'to_dict'):
                    result['chapters'] = [ch.to_dict() if hasattr(ch, 'to_dict') else ch for ch in chapters]

            return result

        except Exception as e:
            logger.error(f"Gemini解析失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise

    def _run_llm_level_analyzer(self, model_name: str = "deepseek-v3") -> Dict:
        """
        方法7: LLM智能层级分析

        使用LLM理解目录语义，智能判断每个目录项的层级
        """
        try:
            # 首先需要获取目录项列表（使用现有的TOC解析）
            toc_idx = self.parser._find_toc_section(self.doc)
            if toc_idx is None:
                return {
                    'success': False,
                    'error': '未找到目录节',
                    'chapters': [],
                    'method_name': 'LLM智能层级分析'
                }

            # 解析目录项
            toc_items, toc_end_idx = self.parser._parse_toc_items(self.doc, toc_idx)
            if not toc_items:
                return {
                    'success': False,
                    'error': '目录解析失败',
                    'chapters': [],
                    'method_name': 'LLM智能层级分析'
                }

            logger.info(f"LLM层级分析: 解析到 {len(toc_items)} 个目录项")

            # 使用LLM分析层级
            level_analyzer = LevelAnalyzer()
            levels = level_analyzer.analyze_toc_hierarchy_with_llm(toc_items, model_name)

            # 将层级结果应用到目录项
            for i, item in enumerate(toc_items):
                if i < len(levels):
                    item['level'] = levels[i]
                else:
                    item['level'] = 2  # 默认2级

            # 🔧 使用 _locate_chapters_by_toc 定位章节并计算字数
            # 这样可以获取每个章节的 word_count
            # 注意: _locate_chapters_by_toc 内部已经调用了 _build_chapter_tree，返回的就是树形结构
            chapter_tree = self.parser._locate_chapters_by_toc(self.doc, toc_items, toc_end_idx)

            # 转换为字典格式
            chapters = [ch.to_dict() for ch in chapter_tree]

            # 🆕 添加章节类型分类
            classified_chapters, key_sections = self.parser._classify_chapters(chapters)

            # 计算统计信息
            stats = self.parser._calculate_statistics(chapter_tree)

            return {
                'success': True,
                'chapters': classified_chapters,
                'method_name': 'LLM智能层级分析',
                'model_used': model_name,
                'key_sections': key_sections,
                'statistics': {
                    'total_items': len(toc_items),
                    'total_words': stats.get('total_words', self.total_chars),
                    'level_1_count': sum(1 for l in levels if l == 1),
                    'level_2_count': sum(1 for l in levels if l == 2),
                    'level_3_count': sum(1 for l in levels if l == 3),
                }
            }

        except Exception as e:
            logger.error(f"LLM层级分析失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise

    def _build_chapter_tree_from_toc(self, toc_items: List[Dict]) -> List[Dict]:
        """
        从带层级的目录项列表构建章节树

        Args:
            toc_items: 带level的目录项列表

        Returns:
            层级结构的章节列表
        """
        if not toc_items:
            return []

        chapters = []
        stack = []  # 用于跟踪父节点

        for item in toc_items:
            chapter = {
                'title': item.get('title', ''),
                'level': item.get('level', 1),
                'page_num': item.get('page_num'),
                'children': []
            }

            level = chapter['level']

            # 清空比当前层级高或相等的节点
            while stack and stack[-1]['level'] >= level:
                stack.pop()

            if not stack:
                # 没有父节点，添加到根
                chapters.append(chapter)
            else:
                # 添加到最近的父节点
                stack[-1]['children'].append(chapter)

            stack.append(chapter)

        return chapters

    @staticmethod
    def calculate_accuracy(detected_chapters: List[Dict], ground_truth_chapters: List[Dict]) -> Dict:
        """
        计算准确率指标

        Args:
            detected_chapters: 检测到的章节列表
            ground_truth_chapters: 正确答案章节列表

        Returns:
            {
                'precision': 0.0-1.0,
                'recall': 0.0-1.0,
                'f1_score': 0.0-1.0,
                'matched_count': int,
                'detected_count': int,
                'ground_truth_count': int,
                'details': [...]
            }
        """
        if not ground_truth_chapters:
            return {
                'precision': 0.0,
                'recall': 0.0,
                'f1_score': 0.0,
                'matched_count': 0,
                'detected_count': len(detected_chapters),
                'ground_truth_count': 0
            }

        # 扁平化章节列表（包含子章节）
        def flatten_chapters(chapters_list):
            flat = []
            for ch in chapters_list:
                flat.append(ch)
                if 'children' in ch and ch['children']:
                    flat.extend(flatten_chapters(ch['children']))
            return flat

        detected_flat = flatten_chapters(detected_chapters)
        truth_flat = flatten_chapters(ground_truth_chapters)

        # 规范化标题（用于匹配）
        def normalize_title(title: str) -> str:
            import re
            # 移除所有空格、编号
            cleaned = re.sub(r'^\d+\.\s*', '', title)
            cleaned = re.sub(r'^\d+\.\d+\s*', '', cleaned)
            cleaned = re.sub(r'^第[一二三四五六七八九十\d]+[章节部分]\s*', '', cleaned)
            cleaned = re.sub(r'\s+', '', cleaned)
            return cleaned.lower()

        # 构建真实答案的标题集合
        truth_titles = {normalize_title(ch['title']): ch for ch in truth_flat}
        detected_titles = {normalize_title(ch['title']): ch for ch in detected_flat}

        # 计算匹配
        matched_titles = set(truth_titles.keys()) & set(detected_titles.keys())
        matched_count = len(matched_titles)

        # 计算指标
        precision = matched_count / len(detected_flat) if detected_flat else 0.0
        recall = matched_count / len(truth_flat) if truth_flat else 0.0
        f1_score = (2 * precision * recall / (precision + recall)) if (precision + recall) > 0 else 0.0

        # 详细匹配信息
        details = []
        for title in truth_titles.keys():
            if title in matched_titles:
                details.append({
                    'title': truth_titles[title]['title'],
                    'status': 'matched',
                    'detected': True
                })
            else:
                details.append({
                    'title': truth_titles[title]['title'],
                    'status': 'missed',
                    'detected': False
                })

        # 检测多余的（误检）
        for title in detected_titles.keys():
            if title not in matched_titles:
                details.append({
                    'title': detected_titles[title]['title'],
                    'status': 'false_positive',
                    'detected': True
                })

        return {
            'precision': round(precision, 4),
            'recall': round(recall, 4),
            'f1_score': round(f1_score, 4),
            'matched_count': matched_count,
            'detected_count': len(detected_flat),
            'ground_truth_count': len(truth_flat),
            'details': details
        }


@api_parser_debug_bp.route('/parse-smart/<document_id>', methods=['POST'])
def parse_smart(document_id):
    """
    智能解析：结构识别 + 类型分类

    流程：
    1. 检查是否有目录
       - 有目录 → 精确匹配(toc_exact)
       - 无目录 → Word大纲识别(docx_native)
    2. 检查结果是否异常
       - 正常 → 对章节进行类型分类
       - 异常 → 回退到LLM层级分析

    请求参数:
        - classify: 是否进行章节类型分类 (默认true)

    响应:
        {
            "success": true,
            "result": {
                "chapters": [...],
                "method_used": "toc_exact",
                "fallback_from": null,
                "fallback_reason": null,
                "key_sections": {...}
            }
        }
    """
    try:
        # 获取请求参数
        data = request.get_json() or {}
        classify_chapters = data.get('classify', True)

        # 获取文件路径
        db = get_knowledge_base_db()
        row = db.execute_query(
            "SELECT file_path FROM parser_debug_tests WHERE document_id = ?",
            (document_id,),
            fetch_one=True
        )

        if not row:
            return jsonify({'success': False, 'error': '文档不存在'}), 404

        file_path = row['file_path']

        # 初始化解析器
        parser = DocumentStructureParser()

        # 调用智能解析方法
        start_time = time.time()
        result = parser.parse_smart(file_path, classify_chapters=classify_chapters)
        elapsed = time.time() - start_time

        # 添加性能信息
        result['performance'] = {
            'elapsed': round(elapsed, 3),
            'elapsed_formatted': f"{elapsed:.3f}s"
        }

        logger.info(f"智能解析完成: method={result.get('method_used')}, "
                   f"fallback={result.get('fallback_from')}, "
                   f"chapters={len(result.get('chapters', []))}, "
                   f"elapsed={elapsed:.3f}s")

        return jsonify({
            'success': True,
            'result': result
        })

    except Exception as e:
        logger.error(f"智能解析失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/upload', methods=['POST'])
def upload_document():
    """
    上传文档并返回document_id（流式解析请使用 /upload-stream）

    请求:
        - file: .docx文件

    响应:
        {
            "success": true,
            "document_id": "uuid",
            "document_info": {...}
        }
    """
    try:
        # 检查文件
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '没有上传文件'}), 400

        file = request.files['file']
        if not file.filename:
            return jsonify({'success': False, 'error': '文件名为空'}), 400

        if not file.filename.endswith('.docx'):
            return jsonify({'success': False, 'error': '仅支持 .docx 格式文件'}), 400

        # 保存文件
        document_id = str(uuid.uuid4())
        original_filename = file.filename

        config = get_config()
        upload_dir = config.get_path('data') / 'parser_debug'
        upload_dir.mkdir(parents=True, exist_ok=True)

        file_path = upload_dir / f"{document_id}.docx"
        file.save(str(file_path))

        logger.info(f"文件已保存: {file_path}")

        # 获取文档基本信息
        debugger = ParserDebugger(str(file_path))
        document_info = debugger.get_document_info()
        document_info['filename'] = original_filename

        # 初始化数据库记录（结果为空，稍后由流式API更新）
        db = get_knowledge_base_db()
        db.execute_query("""
            INSERT INTO parser_debug_tests (
                document_id, filename, file_path,
                total_paragraphs, has_toc, toc_items_count, toc_start_idx, toc_end_idx
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            document_id,
            original_filename,
            str(file_path),
            document_info['total_paragraphs'],
            document_info['has_toc'],
            document_info['toc_items_count'],
            document_info['toc_start_idx'],
            document_info['toc_end_idx']
        ))

        return jsonify({
            'success': True,
            'document_id': document_id,
            'document_info': document_info
        })

    except Exception as e:
        logger.error(f"上传处理失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/parse-single/<document_id>/<method>', methods=['POST'])
def parse_single_method(document_id, method):
    """
    解析单个方法

    Args:
        document_id: 文档ID
        method: 解析方法 (toc_exact/azure/docx_native/gemini)

    返回:
        {
            "success": true,
            "result": {...}
        }
    """
    try:
        # 获取文件路径
        db = get_knowledge_base_db()
        row = db.execute_query(
            "SELECT file_path FROM parser_debug_tests WHERE document_id = ?",
            (document_id,),
            fetch_one=True
        )

        if not row:
            return jsonify({'success': False, 'error': '文档不存在'}), 404

        file_path = row['file_path']
        debugger = ParserDebugger(file_path)

        # 根据method选择对应的解析器
        method_map = {
            'toc_exact': (debugger._run_toc_exact_match, '精确匹配(基于目录)'),
            'docx_native': (debugger._run_docx_native, 'Word大纲级别识别'),
            'azure': (debugger._run_azure_parser, 'Azure Form Recognizer'),
            'gemini': (lambda: debugger._run_gemini_parser(GeminiParser()) if GEMINI_PARSER_AVAILABLE else None, 'Gemini AI解析器'),
            'llm_level': (debugger._run_llm_level_analyzer, 'LLM智能层级分析')
        }

        if method not in method_map:
            return jsonify({'success': False, 'error': f'不支持的解析方法: {method}'}), 400

        method_func, method_name = method_map[method]

        # 特殊处理Azure和Gemini
        if method == 'azure' and not (is_azure_available() and AZURE_PARSER_AVAILABLE):
            result = {
                'success': False,
                'error': 'Azure Form Recognizer 未配置或SDK未安装',
                'chapters': [],
                'method_name': method_name,
                'performance': {'elapsed': 0}
            }
        elif method == 'gemini' and not GEMINI_PARSER_AVAILABLE:
            result = {
                'success': False,
                'error': 'Gemini SDK未安装 (pip install google-generativeai)',
                'chapters': [],
                'method_name': method_name,
                'performance': {'elapsed': 0}
            }
        elif method == 'gemini':
            try:
                gemini_parser = GeminiParser()
                if not gemini_parser.is_available():
                    result = {
                        'success': False,
                        'error': 'Gemini API密钥未配置',
                        'chapters': [],
                        'method_name': method_name,
                        'performance': {'elapsed': 0}
                    }
                else:
                    result = debugger._run_with_timing(lambda: debugger._run_gemini_parser(gemini_parser), method_name)
            except Exception as e:
                logger.error(f"Gemini解析失败: {e}")
                result = {
                    'success': False,
                    'error': str(e),
                    'chapters': [],
                    'method_name': method_name,
                    'performance': {'elapsed': 0}
                }
        else:
            # 执行解析
            result = debugger._run_with_timing(method_func, method_name)

        # 更新数据库
        db.execute_query(f"""
            UPDATE parser_debug_tests
            SET {method}_result = ?,
                {method}_elapsed = ?,
                {method}_chapters_count = ?
            WHERE document_id = ?
        """, (
            json.dumps(result, ensure_ascii=False),
            result['performance']['elapsed'],
            len(result.get('chapters', [])),
            document_id
        ))

        logger.info(f"单方法解析完成: {method_name}, 耗时 {result['performance']['elapsed']}s")

        return jsonify({
            'success': True,
            'result': result
        })

    except Exception as e:
        logger.error(f"单方法解析失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/parse-stream/<document_id>', methods=['GET'])
def parse_document_stream(document_id):
    """
    流式解析文档 - 每完成一个解析器就立即返回结果

    使用Server-Sent Events (SSE)格式

    事件格式:
        data: {"method": "style", "result": {...}, "progress": "2/5"}
        data: {"method": "complete", "document_id": "xxx"}
    """
    def generate():
        try:
            # 获取文件路径
            db = get_knowledge_base_db()
            row = db.execute_query(
                "SELECT file_path FROM parser_debug_tests WHERE document_id = ?",
                (document_id,),
                fetch_one=True
            )

            if not row:
                yield f"data: {json.dumps({'error': '文档不存在'}, ensure_ascii=False)}\n\n"
                return

            file_path = row['file_path']
            debugger = ParserDebugger(file_path)

            # 定义要运行的解析器列表（按新顺序：Gemini → Word大纲 → 精确匹配 → Azure → 其他）
            parsers = []

            # 1. Gemini AI解析器（如果可用）
            if GEMINI_PARSER_AVAILABLE:
                try:
                    gemini_parser = GeminiParser()
                    if gemini_parser.is_available():
                        parsers.append(('gemini', lambda: debugger._run_gemini_parser(gemini_parser), 'Gemini AI解析器'))
                except:
                    pass

            # 2. Word大纲级别识别
            parsers.append(('docx_native', debugger._run_docx_native, 'Word大纲级别识别'))

            # 3. 精确匹配(基于目录)
            parsers.append(('toc_exact', debugger._run_toc_exact_match, '精确匹配(基于目录)'))

            # 4. Azure Form Recognizer（如果可用）
            parsers.append(('azure', lambda: debugger._run_azure_parser() if is_azure_available() and AZURE_PARSER_AVAILABLE else {
                'success': False, 'error': 'Azure未配置', 'chapters': [], 'method_name': 'Azure Form Recognizer', 'performance': {'elapsed': 0}
            }, 'Azure Form Recognizer'))

            # 5. LLM智能层级分析
            parsers.append(('llm_level', debugger._run_llm_level_analyzer, 'LLM智能层级分析'))

            total = len(parsers)
            results_dict = {
                'semantic': {
                    'success': False,
                    'error': '语义锚点解析已禁用（性能优化中）',
                    'chapters': [],
                    'method_name': '语义锚点解析',
                    'performance': {'elapsed': 0}
                }
            }

            # 逐个运行解析器并流式返回结果
            for idx, (method_key, method_func, method_name) in enumerate(parsers, 1):
                try:
                    logger.info(f"[流式解析 {idx}/{total}] 开始运行: {method_name}")
                    result = debugger._run_with_timing(method_func, method_name)
                    results_dict[method_key] = result

                    # 立即发送结果给前端
                    event_data = {
                        'method': method_key,
                        'method_name': method_name,
                        'result': result,
                        'progress': f"{idx}/{total}",
                        'progress_percent': int((idx / total) * 100)
                    }
                    yield f"data: {json.dumps(event_data, ensure_ascii=False)}\n\n"

                    # 立即更新数据库
                    db.execute_query(f"""
                        UPDATE parser_debug_tests
                        SET {method_key}_result = ?,
                            {method_key}_elapsed = ?,
                            {method_key}_chapters_count = ?
                        WHERE document_id = ?
                    """, (
                        json.dumps(result, ensure_ascii=False),
                        result['performance']['elapsed'],
                        len(result.get('chapters', [])),
                        document_id
                    ))

                except Exception as e:
                    logger.error(f"解析器 {method_name} 失败: {e}")
                    error_result = {
                        'success': False,
                        'error': str(e),
                        'chapters': [],
                        'method_name': method_name,
                        'performance': {'elapsed': 0}
                    }
                    results_dict[method_key] = error_result

                    event_data = {
                        'method': method_key,
                        'method_name': method_name,
                        'result': error_result,
                        'progress': f"{idx}/{total}",
                        'progress_percent': int((idx / total) * 100)
                    }
                    yield f"data: {json.dumps(event_data, ensure_ascii=False)}\n\n"

            # 同步语义锚点解析结果到数据库
            db.execute_query("""
                UPDATE parser_debug_tests
                SET semantic_result = ?, semantic_elapsed = ?, semantic_chapters_count = ?
                WHERE document_id = ?
            """, (
                json.dumps(results_dict['semantic'], ensure_ascii=False),
                0,
                0,
                document_id
            ))

            # 完成信号
            yield f"data: {json.dumps({'method': 'complete', 'document_id': document_id}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"流式解析失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    return Response(generate(), mimetype='text/event-stream')


@api_parser_debug_bp.route('/<document_id>', methods=['GET'])
def get_test_result(document_id):
    """
    获取测试结果

    响应:
        {
            "success": true,
            "document_info": {...},
            "results": {...},
            "ground_truth": {...},
            "accuracy": {...}
        }
    """
    try:
        db = get_knowledge_base_db()
        row = db.execute_query(
            "SELECT * FROM parser_debug_tests WHERE document_id = ?",
            (document_id,),
            fetch_one=True
        )

        if not row:
            return jsonify({'success': False, 'error': '测试记录不存在'}), 404

        # 解析结果
        results = {
            'toc_exact': json.loads(row['toc_exact_result']) if row.get('toc_exact_result') else None,
            'semantic': json.loads(row['semantic_result']) if row['semantic_result'] else None,
            'style': json.loads(row['style_result']) if row['style_result'] else None,
            'hybrid': json.loads(row['hybrid_result']) if row.get('hybrid_result') else None,
            'azure': json.loads(row['azure_result']) if row.get('azure_result') else None,
            'docx_native': json.loads(row['docx_native_result']) if row.get('docx_native_result') else None,
            'gemini': json.loads(row['gemini_result']) if row.get('gemini_result') else None,
            'llm_level': json.loads(row['llm_level_result']) if row.get('llm_level_result') else None,
        }

        document_info = {
            'filename': row['filename'],
            'total_paragraphs': row['total_paragraphs'],
            'has_toc': bool(row['has_toc']),
            'toc_items_count': row['toc_items_count'],
            'upload_time': row['upload_time']
        }

        ground_truth = json.loads(row['ground_truth']) if row['ground_truth'] else None

        # 如果有ground_truth，返回准确率数据
        accuracy = None
        if ground_truth:
            accuracy = {
                'toc_exact': {
                    'precision': row.get('toc_exact_precision'),
                    'recall': row.get('toc_exact_recall'),
                    'f1_score': row.get('toc_exact_f1')
                } if row.get('toc_exact_precision') else None,
                'semantic': {
                    'precision': row['semantic_precision'],
                    'recall': row['semantic_recall'],
                    'f1_score': row['semantic_f1']
                },
                'style': {
                    'precision': row['style_precision'],
                    'recall': row['style_recall'],
                    'f1_score': row['style_f1']
                },
                'hybrid': {
                    'precision': row.get('hybrid_precision'),
                    'recall': row.get('hybrid_recall'),
                    'f1_score': row.get('hybrid_f1')
                } if row.get('hybrid_precision') else None,
                'azure': {
                    'precision': row.get('azure_precision'),
                    'recall': row.get('azure_recall'),
                    'f1_score': row.get('azure_f1')
                } if row.get('azure_precision') else None,
                'docx_native': {
                    'precision': row.get('docx_native_precision'),
                    'recall': row.get('docx_native_recall'),
                    'f1_score': row.get('docx_native_f1')
                } if row.get('docx_native_precision') else None,
                'gemini': {
                    'precision': row.get('gemini_precision'),
                    'recall': row.get('gemini_recall'),
                    'f1_score': row.get('gemini_f1')
                } if row.get('gemini_precision') else None,
                'llm_level': {
                    'precision': row.get('llm_level_precision'),
                    'recall': row.get('llm_level_recall'),
                    'f1_score': row.get('llm_level_f1')
                } if row.get('llm_level_precision') else None,
                'best_method': row['best_method'],
                'best_f1_score': row['best_f1_score']
            }

        return jsonify({
            'success': True,
            'document_id': document_id,
            'document_info': document_info,
            'results': results,
            'ground_truth': ground_truth,
            'accuracy': accuracy
        })

    except Exception as e:
        logger.error(f"获取测试结果失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/<document_id>/ground-truth', methods=['POST'])
def save_ground_truth(document_id):
    """
    保存人工标注的正确答案

    请求:
        {
            "chapters": [...],  # 正确的章节列表
            "annotator": "用户名"
        }

    响应:
        {
            "success": true,
            "accuracy": {...}  # 自动计算的准确率
        }
    """
    try:
        data = request.get_json()
        if not data or 'chapters' not in data:
            return jsonify({'success': False, 'error': '缺少章节数据'}), 400

        chapters = data['chapters']
        annotator = data.get('annotator', 'unknown')

        # 获取现有测试结果
        db = get_knowledge_base_db()
        row = db.execute_query(
            "SELECT toc_exact_result, semantic_result, style_result, hybrid_result, azure_result, docx_native_result, gemini_result, llm_level_result FROM parser_debug_tests WHERE document_id = ?",
            (document_id,),
            fetch_one=True
        )

        if not row:
            return jsonify({'success': False, 'error': '测试记录不存在'}), 404

        # 解析各方法的结果
        toc_exact_chapters = json.loads(row['toc_exact_result'])['chapters'] if row.get('toc_exact_result') else []
        semantic_chapters = json.loads(row['semantic_result'])['chapters'] if row['semantic_result'] else []
        style_chapters = json.loads(row['style_result'])['chapters'] if row['style_result'] else []
        hybrid_chapters = json.loads(row['hybrid_result'])['chapters'] if row.get('hybrid_result') else []
        azure_chapters = json.loads(row['azure_result'])['chapters'] if row.get('azure_result') else []
        docx_native_chapters = json.loads(row['docx_native_result'])['chapters'] if row.get('docx_native_result') else []
        gemini_chapters = json.loads(row['gemini_result'])['chapters'] if row.get('gemini_result') else []
        llm_level_chapters = json.loads(row['llm_level_result'])['chapters'] if row.get('llm_level_result') else []

        # 计算各方法的准确率
        toc_exact_acc = ParserDebugger.calculate_accuracy(toc_exact_chapters, chapters) if toc_exact_chapters else None
        semantic_acc = ParserDebugger.calculate_accuracy(semantic_chapters, chapters)
        style_acc = ParserDebugger.calculate_accuracy(style_chapters, chapters)
        hybrid_acc = ParserDebugger.calculate_accuracy(hybrid_chapters, chapters) if hybrid_chapters else None
        azure_acc = ParserDebugger.calculate_accuracy(azure_chapters, chapters) if azure_chapters else None
        docx_native_acc = ParserDebugger.calculate_accuracy(docx_native_chapters, chapters) if docx_native_chapters else None
        gemini_acc = ParserDebugger.calculate_accuracy(gemini_chapters, chapters) if gemini_chapters else None
        llm_level_acc = ParserDebugger.calculate_accuracy(llm_level_chapters, chapters) if llm_level_chapters else None

        # 找出最佳方法
        all_f1 = {
            'semantic': semantic_acc['f1_score'],
            'style': style_acc['f1_score'],
        }
        if toc_exact_acc:
            all_f1['toc_exact'] = toc_exact_acc['f1_score']
        if hybrid_acc:
            all_f1['hybrid'] = hybrid_acc['f1_score']
        if azure_acc:
            all_f1['azure'] = azure_acc['f1_score']
        if docx_native_acc:
            all_f1['docx_native'] = docx_native_acc['f1_score']
        if gemini_acc:
            all_f1['gemini'] = gemini_acc['f1_score']
        if llm_level_acc:
            all_f1['llm_level'] = llm_level_acc['f1_score']
        best_method = max(all_f1, key=all_f1.get)
        best_f1_score = all_f1[best_method]

        # 更新数据库
        update_params = [
            json.dumps(chapters, ensure_ascii=False),
            annotator,
            datetime.now().isoformat(),
            len(chapters),
            semantic_acc['precision'], semantic_acc['recall'], semantic_acc['f1_score'],
            style_acc['precision'], style_acc['recall'], style_acc['f1_score'],
        ]

        # 如果有 toc_exact 结果，添加其准确率
        if toc_exact_acc:
            update_params.extend([toc_exact_acc['precision'], toc_exact_acc['recall'], toc_exact_acc['f1_score']])
        else:
            update_params.extend([None, None, None])

        # 如果有 hybrid 结果，添加其准确率
        if hybrid_acc:
            update_params.extend([hybrid_acc['precision'], hybrid_acc['recall'], hybrid_acc['f1_score']])
        else:
            update_params.extend([None, None, None])

        # 如果有 Azure 结果，添加其准确率
        if azure_acc:
            update_params.extend([azure_acc['precision'], azure_acc['recall'], azure_acc['f1_score']])
        else:
            update_params.extend([None, None, None])

        # 如果有 docx_native 结果，添加其准确率
        if docx_native_acc:
            update_params.extend([docx_native_acc['precision'], docx_native_acc['recall'], docx_native_acc['f1_score']])
        else:
            update_params.extend([None, None, None])

        # 如果有 Gemini 结果，添加其准确率
        if gemini_acc:
            update_params.extend([gemini_acc['precision'], gemini_acc['recall'], gemini_acc['f1_score']])
        else:
            update_params.extend([None, None, None])

        # 如果有 LLM层级分析 结果，添加其准确率
        if llm_level_acc:
            update_params.extend([llm_level_acc['precision'], llm_level_acc['recall'], llm_level_acc['f1_score']])
        else:
            update_params.extend([None, None, None])

        update_params.extend([best_method, best_f1_score, document_id])

        db.execute_query("""
            UPDATE parser_debug_tests SET
                ground_truth = ?, annotator = ?, annotation_time = ?, ground_truth_count = ?,
                semantic_precision = ?, semantic_recall = ?, semantic_f1 = ?,
                style_precision = ?, style_recall = ?, style_f1 = ?,
                toc_exact_precision = ?, toc_exact_recall = ?, toc_exact_f1 = ?,
                hybrid_precision = ?, hybrid_recall = ?, hybrid_f1 = ?,
                azure_precision = ?, azure_recall = ?, azure_f1 = ?,
                docx_native_precision = ?, docx_native_recall = ?, docx_native_f1 = ?,
                gemini_precision = ?, gemini_recall = ?, gemini_f1 = ?,
                llm_level_precision = ?, llm_level_recall = ?, llm_level_f1 = ?,
                best_method = ?, best_f1_score = ?
            WHERE document_id = ?
        """, tuple(update_params))

        accuracy_result = {
            'semantic': semantic_acc,
            'style': style_acc,
            'best_method': best_method,
            'best_f1_score': best_f1_score
        }

        if toc_exact_acc:
            accuracy_result['toc_exact'] = toc_exact_acc
        if hybrid_acc:
            accuracy_result['hybrid'] = hybrid_acc
        if azure_acc:
            accuracy_result['azure'] = azure_acc
        if docx_native_acc:
            accuracy_result['docx_native'] = docx_native_acc
        if gemini_acc:
            accuracy_result['gemini'] = gemini_acc
        if llm_level_acc:
            accuracy_result['llm_level'] = llm_level_acc

        return jsonify({
            'success': True,
            'accuracy': accuracy_result
        })

    except Exception as e:
        logger.error(f"保存ground truth失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/history', methods=['GET'])
def get_history():
    """
    获取历史测试列表

    查询参数:
        - limit: 返回数量限制（默认20）
        - has_ground_truth: 是否只返回已标注的（可选）

    响应:
        {
            "success": true,
            "tests": [...]
        }
    """
    try:
        limit = request.args.get('limit', 20, type=int)
        has_ground_truth = request.args.get('has_ground_truth', type=bool)

        db = get_knowledge_base_db()

        sql = "SELECT * FROM v_parser_debug_summary"
        params = []

        if has_ground_truth is not None:
            sql += " WHERE has_ground_truth = ?"
            params.append(1 if has_ground_truth else 0)

        sql += " LIMIT ?"
        params.append(limit)

        rows = db.execute_query(sql, tuple(params))

        tests = []
        for row in rows:
            tests.append(dict(row))

        return jsonify({
            'success': True,
            'tests': tests,
            'total': len(tests)
        })

    except Exception as e:
        logger.error(f"获取历史记录失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/<document_id>/delete', methods=['DELETE'])
def delete_test(document_id):
    """删除测试记录"""
    try:
        db = get_knowledge_base_db()

        # 获取文件路径并删除文件
        row = db.execute_query(
            "SELECT file_path FROM parser_debug_tests WHERE document_id = ?",
            (document_id,),
            fetch_one=True
        )

        if row and row['file_path']:
            file_path = Path(row['file_path'])
            if file_path.exists():
                file_path.unlink()
                logger.info(f"已删除文件: {file_path}")

        # 删除数据库记录
        db.execute_query(
            "DELETE FROM parser_debug_tests WHERE document_id = ?",
            (document_id,)
        )

        return jsonify({'success': True})

    except Exception as e:
        logger.error(f"删除测试记录失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/export/<document_id>', methods=['GET'])
def export_comparison_report(document_id):
    """
    导出对比报告（JSON格式）

    响应:
        完整的JSON报告文件
    """
    try:
        db = get_knowledge_base_db()
        row = db.execute_query(
            "SELECT * FROM parser_debug_tests WHERE document_id = ?",
            (document_id,),
            fetch_one=True
        )

        if not row:
            return jsonify({'success': False, 'error': '测试记录不存在'}), 404

        # 构建完整报告
        report = {
            'document_id': document_id,
            'filename': row['filename'],
            'upload_time': row['upload_time'],
            'document_info': {
                'total_paragraphs': row['total_paragraphs'],
                'has_toc': bool(row['has_toc']),
                'toc_items_count': row['toc_items_count']
            },
            'results': {
                'toc_exact': json.loads(row['toc_exact_result']) if row.get('toc_exact_result') else None,
                'semantic': json.loads(row['semantic_result']) if row['semantic_result'] else None,
                'style': json.loads(row['style_result']) if row['style_result'] else None,
                'hybrid': json.loads(row['hybrid_result']) if row.get('hybrid_result') else None,
                'azure': json.loads(row['azure_result']) if row.get('azure_result') else None,
                'docx_native': json.loads(row['docx_native_result']) if row.get('docx_native_result') else None,
                'gemini': json.loads(row['gemini_result']) if row.get('gemini_result') else None,
                'llm_level': json.loads(row['llm_level_result']) if row.get('llm_level_result') else None,
            },
            'ground_truth': json.loads(row['ground_truth']) if row['ground_truth'] else None,
            'accuracy': None
        }

        # 如果有标注，添加准确率数据
        if row['ground_truth']:
            report['accuracy'] = {
                'semantic': {
                    'precision': row['semantic_precision'],
                    'recall': row['semantic_recall'],
                    'f1_score': row['semantic_f1']
                },
                'style': {
                    'precision': row['style_precision'],
                    'recall': row['style_recall'],
                    'f1_score': row['style_f1']
                },
                'best_method': row['best_method'],
                'best_f1_score': row['best_f1_score']
            }

            # 添加toc_exact结果(如果存在)
            if row.get('toc_exact_precision'):
                report['accuracy']['toc_exact'] = {
                    'precision': row['toc_exact_precision'],
                    'recall': row['toc_exact_recall'],
                    'f1_score': row['toc_exact_f1']
                }

            # 添加hybrid结果(如果存在)
            if row.get('hybrid_precision'):
                report['accuracy']['hybrid'] = {
                    'precision': row['hybrid_precision'],
                    'recall': row['hybrid_recall'],
                    'f1_score': row['hybrid_f1']
                }

            # 添加azure结果(如果存在)
            if row.get('azure_precision'):
                report['accuracy']['azure'] = {
                    'precision': row['azure_precision'],
                    'recall': row['azure_recall'],
                    'f1_score': row['azure_f1']
                }

            # 添加docx_native结果(如果存在)
            if row.get('docx_native_precision'):
                report['accuracy']['docx_native'] = {
                    'precision': row['docx_native_precision'],
                    'recall': row['docx_native_recall'],
                    'f1_score': row['docx_native_f1']
                }

            # 添加gemini结果(如果存在)
            if row.get('gemini_precision'):
                report['accuracy']['gemini'] = {
                    'precision': row['gemini_precision'],
                    'recall': row['gemini_recall'],
                    'f1_score': row['gemini_f1']
                }

            # 添加llm_level结果(如果存在)
            if row.get('llm_level_precision'):
                report['accuracy']['llm_level'] = {
                    'precision': row['llm_level_precision'],
                    'recall': row['llm_level_recall'],
                    'f1_score': row['llm_level_f1']
                }

        # 保存为临时JSON文件
        config = get_config()
        temp_dir = config.get_path('data') / 'temp'
        temp_dir.mkdir(parents=True, exist_ok=True)

        report_file = temp_dir / f"parser_comparison_{document_id}.json"
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        return send_file(
            report_file,
            as_attachment=True,
            download_name=f"parser_comparison_{row['filename']}.json",
            mimetype='application/json'
        )

    except Exception as e:
        logger.error(f"导出报告失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_parser_debug_bp.route('/preview/<document_id>', methods=['GET'])
def preview_document(document_id):
    """
    获取文档预览信息（文件路径）

    Args:
        document_id: 文档ID

    Returns:
        文档信息（包含文件路径，用于 DocumentPreview 组件）
    """
    try:
        db = get_knowledge_base_db()

        # 获取文件路径
        row = db.execute_query(
            "SELECT file_path, filename FROM parser_debug_tests WHERE document_id = ?",
            (document_id,),
            fetch_one=True
        )

        if not row:
            return jsonify({'success': False, 'error': '文档不存在'}), 404

        file_path = Path(row['file_path'])
        if not file_path.exists():
            return jsonify({'success': False, 'error': '文件不存在'}), 404

        # 返回文件路径信息（供 DocumentPreview 组件使用）
        return jsonify({
            'success': True,
            'file_path': str(file_path),
            'filename': row['filename']
        })

    except Exception as e:
        logger.error(f"获取文档预览信息失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# 注册蓝图到应用（需要在app.py中调用）
def register_parser_debug_bp(app):
    """注册解析调试蓝图"""
    app.register_blueprint(api_parser_debug_bp)
    logger.info("解析调试API已注册")
