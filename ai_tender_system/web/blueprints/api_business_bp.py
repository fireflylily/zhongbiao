#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
商务应答和点对点应答API蓝图
处理商务应答文档生成和点对点应答处理
"""

import os
import sys
import hashlib
import urllib.parse
import html
from datetime import datetime
from pathlib import Path
from flask import Blueprint, request, jsonify, send_file

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from common import (
    get_module_logger, get_config, format_error_response,
    safe_filename, ensure_dir
)
from web.shared.instances import get_kb_manager

# 创建蓝图
api_business_bp = Blueprint('api_business', __name__, url_prefix='/api')

# 日志记录器
logger = get_module_logger("web.api_business")

# 获取配置和知识库管理器
config = get_config()
kb_manager = get_kb_manager()

# 检查商务应答模块可用性
BUSINESS_RESPONSE_AVAILABLE = False
POINT_TO_POINT_AVAILABLE = False
try:
    from modules.business_response.processor import BusinessResponseProcessor, PointToPointProcessor
    # 【重构】导入统一的图片配置构建器
    from modules.business_response.image_config_builder import build_image_config_from_db
    BUSINESS_RESPONSE_AVAILABLE = True
    POINT_TO_POINT_AVAILABLE = True  # 保持向后兼容
except ImportError:
    pass


# ===================
# 辅助函数
# ===================


def generate_output_filename(project_name: str, file_type: str, timestamp: str = None) -> str:
    """
    生成统一格式的输出文件名: {项目名称}_{类型}_{时间戳}.docx

    Args:
        project_name: 项目名称
        file_type: 文件类型（如：商务应答、点对点应答、技术方案）
        timestamp: 时间戳，如果未提供则自动生成

    Returns:
        格式化的文件名
    """
    if not timestamp:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # 使用safe_filename确保项目名称安全，但不添加时间戳（避免重复）
    safe_project = safe_filename(project_name, timestamp=False) if project_name else "未命名项目"

    return f"{safe_project}_{file_type}_{timestamp}.docx"


# ===================
# 商务应答路由
# ===================

@api_business_bp.route('/process-business-response', methods=['POST'])
def process_business_response():
    """处理商务应答"""
    if not BUSINESS_RESPONSE_AVAILABLE:
        return jsonify({
            'success': False,
            'message': '商务应答模块不可用'
        })

    try:
        # 检查是否从HITL传递了文件路径
        hitl_file_path = request.form.get('hitl_file_path')

        if hitl_file_path:
            # 使用HITL文件路径，跳过文件上传和storage_service
            template_path = Path(hitl_file_path)

            if not template_path.exists():
                raise ValueError(f"HITL文件不存在: {hitl_file_path}")

            filename = template_path.name
            logger.info(f"使用HITL文件路径: {hitl_file_path}")
        else:
            # 原有逻辑：处理上传文件
            if 'template_file' not in request.files:
                raise ValueError("没有选择模板文件")

            file = request.files['template_file']
            if file.filename == '':
                raise ValueError("文件名为空")

        # 获取表单数据
        data = request.form.to_dict()
        company_id = data.get('company_id', '')
        project_name = data.get('project_name', '')
        tender_no = data.get('tender_no', '')
        date_text = data.get('date_text', '')
        use_mcp = data.get('use_mcp', 'false').lower() == 'true'

        # 验证必填字段
        if not company_id:
            raise ValueError("请选择应答公司")

        # 转换公司ID为整数
        company_id_int = int(company_id)

        # 从数据库获取项目相关信息（如果有项目名称）
        purchaser_name = ''
        db_project_number = ''
        db_deadline = ''  # 投标截止时间
        if project_name:
            try:
                query = """SELECT tenderer, project_number, bidding_time
                           FROM tender_projects WHERE project_name = ? LIMIT 1"""
                result = kb_manager.db.execute_query(query, [project_name])
                if result and len(result) > 0:
                    purchaser_name = result[0].get('tenderer', '')
                    db_project_number = result[0].get('project_number', '')
                    db_deadline = result[0].get('bidding_time', '')
                    if purchaser_name:
                        logger.info(f"从数据库获取采购人信息: {purchaser_name}")
                    if db_project_number:
                        logger.info(f"从数据库获取项目编号: {db_project_number}")
                    if db_deadline:
                        logger.info(f"从数据库获取投标截止时间: {db_deadline}")
            except Exception as e:
                logger.warning(f"查询项目信息失败: {e}")

        # 如果表单没有提供项目编号，使用数据库中的项目编号
        if not tender_no and db_project_number:
            tender_no = db_project_number
            logger.info(f"使用数据库项目编号: {tender_no}")

        # 统一日期格式处理：移除时分，只保留到日
        import re

        if date_text and date_text.strip():
            # 用户传入了日期，格式化：移除"XX时XX分"或实际的时分部分
            # 支持格式：2025年11月24日XX时XX分、2025年11月24日 12:30、2025-11-24 12:30
            date_text = re.sub(r'XX时XX分.*$', '', date_text).strip()  # 移除"XX时XX分"占位符
            date_text = re.sub(r'\s+\d{1,2}:\d{2}.*$', '', date_text).strip()  # 移除实际时间（空格+时:分）
            logger.info(f"格式化日期（移除时分）: {date_text}")
        elif db_deadline:
            # 使用项目截止日期（格式化为只包含日期）
            if isinstance(db_deadline, str):
                # 移除时分部分
                date_text = re.sub(r'XX时XX分.*$', '', db_deadline).strip()
                date_text = re.sub(r'\s+\d{1,2}:\d{2}.*$', '', date_text).strip()
            else:
                date_text = str(db_deadline).split()[0]
            logger.info(f"使用数据库截止日期（已格式化）: {date_text}")
        else:
            logger.info("未提供日期且无项目截止日期，date字段将不填充")
            # 注意：这里不设置当前日期，而是保持为空，让后端填充器跳过

        # 从数据库加载所有资质（模板驱动）
        image_config, required_quals = build_image_config_from_db(company_id_int, project_name, kb_manager)

        if image_config:
            logger.info(f"成功从数据库加载图片配置，包含 {len(image_config)} 个类型")
        else:
            logger.warning(f"公司 {company_id} 没有可用的资质图片")

        # 输出项目资格要求信息
        if required_quals:
            logger.info(f"📋 项目资格要求: {len(required_quals)} 个资质")

        # 从数据库获取公司信息
        company_db_data = kb_manager.get_company_detail(company_id_int)
        if not company_db_data:
            raise ValueError(f"未找到公司信息: {company_id}")

        # 使用现有字段映射反向转换为业务处理器期望的格式
        field_mapping = {
            'companyName': 'company_name',
            'establishDate': 'establish_date',
            'legalRepresentative': 'legal_representative',
            'legalRepresentativePosition': 'legal_representative_position',
            'legalRepresentativeGender': 'legal_representative_gender',
            'legalRepresentativeAge': 'legal_representative_age',
            'socialCreditCode': 'social_credit_code',
            'registeredCapital': 'registered_capital',
            'companyType': 'company_type',
            'registeredAddress': 'registered_address',
            'businessScope': 'business_scope',
            'companyDescription': 'description',
            'fixedPhone': 'fixed_phone',
            'fax': 'fax',
            'postalCode': 'postal_code',
            'email': 'email',
            'officeAddress': 'office_address',
            'employeeCount': 'employee_count',
            'bankName': 'bank_name',
            'bankAccount': 'bank_account',
            # 被授权人信息（2025-11-07添加）
            'representativeName': 'authorized_person_name',
            'representativeTitle': 'authorized_person_position',
            'authorizedPersonId': 'authorized_person_id',
            # 股权结构字段（2025-10-27添加）
            'actual_controller': 'actual_controller',
            'controlling_shareholder': 'controlling_shareholder',
            'shareholders_info': 'shareholders_info',
            # 管理关系字段（2025-10-28添加）
            'managing_unit_name': 'managing_unit_name',
            'managed_unit_name': 'managed_unit_name'
        }
        reverse_mapping = {v: k for k, v in field_mapping.items()}
        company_data = {reverse_mapping.get(k, k): v for k, v in company_db_data.items()}

        # 确保company_id被正确传递(用于案例库和简历库查询)
        company_data['company_id'] = company_id_int
        logger.info(f"✅ 已添加company_id到company_data: company_id={company_id_int}")

        # 添加采购人信息到company_data（采购人是项目信息，但为了方便传递，加到这里）
        if purchaser_name:
            company_data['purchaserName'] = purchaser_name
            logger.info(f"✅ 已添加采购人到company_data: purchaserName={purchaser_name}")
        else:
            logger.warning("⚠️  项目无采购人信息（tenderer字段为空）")

        # 🆕 加载公司资质信息（用于表格填充）
        try:
            qualifications_query = """
                SELECT qualification_key, qualification_name, file_path
                FROM company_qualifications
                WHERE company_id = ? AND is_valid = 1
                ORDER BY upload_time DESC
            """
            qualifications = kb_manager.db.execute_query(qualifications_query, (company_id_int,))

            # 将资质信息添加到company_data中
            for qual in qualifications:
                qual_key = qual.get('qualification_key')
                file_path = qual.get('file_path')
                qual_name = qual.get('qualification_name', '')

                # 特殊处理：审计报告（填充友好提示文本而不是文件路径）
                if qual_key == 'audit_report' and file_path:
                    # 填充提示文本："见附件-财务审计报告"
                    company_data['audit_report'] = f"见附件-{qual_name}" if qual_name else "见附件"
                    company_data['audit_organization'] = "见审计报告附件"  # 审计机构也填充提示
                    logger.info(f"✅ 已添加审计报告到company_data: {company_data['audit_report']}")

                # 其他资质也可以类似处理（后续可扩展）

            logger.info(f"✅ 成功加载{len(qualifications)}个资质记录")
        except Exception as e:
            logger.warning(f"⚠️  加载资质信息失败: {e}")

        # 🆕 加载人员信息摘要（用于表格字段填充）
        try:
            from modules.resume_library.manager import ResumeLibraryManager
            resume_manager = ResumeLibraryManager()

            # 查询该公司的简历（或所有简历，如果company_id为None）
            resume_result = resume_manager.get_resumes(company_id=company_id_int, page_size=100)
            resumes = resume_result.get('resumes', [])

            if resumes:
                # 生成人员摘要：列出主要人员姓名和职位
                staff_summary = []
                for idx, resume in enumerate(resumes[:5]):  # 最多显示5个
                    name = resume.get('name', '')
                    position = resume.get('current_position', '')
                    if name and position:
                        staff_summary.append(f"{name}({position})")
                    elif name:
                        staff_summary.append(name)

                if staff_summary:
                    company_data['staff_info'] = '、'.join(staff_summary) + (f'等{len(resumes)}人' if len(resumes) > 5 else '')
                    company_data['project_manager_name'] = resumes[0].get('name', '') if resumes else ''
                    logger.info(f"✅ 已添加人员信息到company_data: {len(resumes)}份简历")
                else:
                    company_data['staff_info'] = "见简历附件"
            else:
                # 没有简历数据，填充提示
                company_data['staff_info'] = "见简历附件"
                logger.info("ℹ️  公司暂无简历数据")

        except Exception as e:
            logger.warning(f"⚠️  加载人员信息失败: {e}")
            company_data['staff_info'] = "见简历附件"

        # 如果没有使用HITL文件路径,才需要保存上传的文件
        if not hitl_file_path:
            # 保存模板文件 - 使用统一服务
            from core.storage_service import storage_service
            file_metadata = storage_service.store_file(
                file_obj=file,
                original_name=file.filename,
                category='business_templates',
                business_type='business_response',
                company_id=company_id
            )
            template_path = Path(file_metadata.file_path)
            filename = file_metadata.safe_name
            logger.info(f"开始处理商务应答: {file_metadata.original_name}")
        else:
            logger.info(f"开始处理商务应答(使用HITL文件): {filename}")

        # 公共的输出文件路径设置（移到外面，两个分支都需要）
        output_dir = ensure_dir(config.get_path('output'))
        # 使用新的文件命名规则：{项目名称}_商务应答_{时间戳}.docx
        output_filename = generate_output_filename(project_name, "商务应答")
        output_path = output_dir / output_filename

        # 添加调试日志
        logger.info(f"[文件命名调试] project_name参数: {project_name}")
        logger.info(f"[文件命名调试] 生成的文件名: {output_filename}")
        logger.info(f"[文件命名调试] 输出目录: {output_dir}")
        logger.info(f"[文件命名调试] 完整输出路径: {output_path}")

        logger.info(f"公司数据验证:")
        logger.info(f"  - 公司名称: {company_data.get('companyName', 'N/A')}")
        logger.info(f"  - 联系电话: {company_data.get('fixedPhone', 'N/A')}")
        logger.info(f"  - 电子邮件: {company_data.get('email', 'N/A')}")
        logger.info(f"  - 公司地址: {company_data.get('address', 'N/A')}")
        logger.info(f"  - 传真号码: {company_data.get('fax', 'N/A')}")
        logger.info(f"  - 项目名称: {project_name}")
        logger.info(f"  - 招标编号: {tender_no}")
        logger.info(f"  - 日期文本: {date_text}")

        # 使用MCP处理器处理商务应答
        if use_mcp:
            # 使用新架构的商务应答处理器
            processor = BusinessResponseProcessor()

            # 使用MCP处理器的完整商务应答处理方法，包含日期字段处理和图片插入
            result_stats = processor.process_business_response(
                str(template_path),
                str(output_path),
                company_data,
                project_name,
                tender_no,
                date_text,
                image_config,     # 传递图片配置
                required_quals    # 传递项目资格要求列表（用于追加和统计）
            )

            output_path = str(output_path)

            # 检查处理结果并构建响应
            if result_stats.get('success'):
                logger.info(f"新架构处理器执行成功: {result_stats.get('message', '无消息')}")
                logger.info(f"处理统计: {result_stats.get('stats', {})}")

                # 构建成功结果
                result = {
                    'success': True,
                    'message': result_stats.get('message', '商务应答处理完成'),
                    'output_file': output_path,
                    'download_url': f'/api/files/download/{os.path.basename(output_path)}',
                    'stats': result_stats.get('summary', {})  # 修复：使用'summary'字段（与processor返回的字段名一致）
                }
            else:
                logger.error(f"新架构处理器执行失败: {result_stats.get('error', '未知错误')}")
                result = {
                    'success': False,
                    'error': result_stats.get('error', '处理失败'),
                    'message': result_stats.get('message', '商务应答处理失败')
                }
        else:
            # 使用向后兼容的处理器（实际上还是新的BusinessResponseProcessor）
            processor = PointToPointProcessor()  # 这是BusinessResponseProcessor的别名
            result_stats = processor.process_business_response(
                str(template_path),
                str(output_path),
                company_data,
                project_name,
                tender_no,
                date_text,
                image_config,     # 传递图片配置
                required_quals    # 传递项目资格要求列表（用于追加和统计）
            )

            # 统一返回格式处理
            if result_stats.get('success'):
                result = {
                    'success': True,
                    'message': result_stats.get('message', '商务应答处理完成'),
                    'output_file': str(output_path),
                    'download_url': f'/api/files/download/{os.path.basename(output_path)}',
                    'stats': result_stats.get('summary', {})
                }
            else:
                result = {
                    'success': False,
                    'error': result_stats.get('error', '处理失败'),
                    'message': result_stats.get('message', '商务应答处理失败')
                }

        # 【优化】如果处理成功，直接同步文件信息到数据库（不使用HTTP调用）
        if result.get('success') and project_name:
            try:
                import json

                # 1. 查询项目信息（包括project_id和step1_data）
                query = """
                    SELECT project_id, step1_data
                    FROM tender_projects
                    WHERE project_name = ? AND company_id = ?
                    LIMIT 1
                """
                project_result = kb_manager.db.execute_query(
                    query, [project_name, company_id_int], fetch_one=True
                )

                if project_result:
                    project_id = project_result['project_id']

                    # 2. 解析现有的step1_data
                    step1_data_raw = project_result.get('step1_data', '{}')
                    step1_data = json.loads(step1_data_raw) if step1_data_raw else {}

                    # 3. 构建文件信息
                    now = datetime.now()
                    file_info = {
                        "file_path": str(output_path),
                        "filename": os.path.basename(output_path),
                        "file_name": os.path.basename(output_path),  # 兼容性字段
                        "file_size": os.path.getsize(output_path),
                        "file_url": f"/api/files/download/{os.path.basename(output_path)}",  # 可访问的下载URL
                        "download_url": f"/api/files/download/{os.path.basename(output_path)}",  # 与API响应保持一致
                        "saved_at": now.isoformat(),
                        "source": "business_response_api"
                    }

                    # 4. 更新step1_data
                    step1_data['business_response_file'] = file_info

                    # 5. 保存到数据库
                    update_query = """
                        UPDATE tender_projects
                        SET step1_data = ?, updated_at = CURRENT_TIMESTAMP
                        WHERE project_id = ?
                    """
                    kb_manager.db.execute_query(
                        update_query,
                        [json.dumps(step1_data, ensure_ascii=False), project_id]
                    )

                    logger.info(
                        f"✅ 商务应答文件已同步到数据库: "
                        f"project_id={project_id}, "
                        f"file={os.path.basename(output_path)}, "
                        f"size={file_info['file_size']} bytes"
                    )
                else:
                    logger.warning(
                        f"⚠️  未找到匹配的项目记录: "
                        f"project_name='{project_name}', "
                        f"company_id={company_id_int}"
                    )
            except Exception as e:
                logger.error(
                    f"❌ 同步商务应答文件到数据库失败: {e}",
                    exc_info=True  # 打印完整堆栈
                )
                # 注意：不影响主流程，文件已生成成功

        logger.info("商务应答处理完成")
        return jsonify(result)

    except Exception as e:
        logger.error(f"商务应答处理失败: {e}")
        return jsonify(format_error_response(e))


@api_business_bp.route('/api/document/process', methods=['POST'])
def process_document():
    """处理文档 - 通用接口"""
    if not BUSINESS_RESPONSE_AVAILABLE:
        return jsonify({
            'success': False,
            'message': '文档处理模块不可用'
        })

    try:
        data = request.get_json()
        file_path = data.get('file_path', '')
        options = data.get('options', {})

        if not file_path:
            raise ValueError("文件路径不能为空")

        # 这是一个通用接口，根据选项决定使用哪个处理器
        doc_type = options.get('type', 'business_response')

        if doc_type == 'tech_requirements':
            result = {
                'success': True,
                'message': '技术需求处理功能可用，请使用 /process-tech-requirements 接口',
                'redirect': '/process-tech-requirements'
            }
        else:
            result = {
                'success': True,
                'message': '商务应答处理功能可用，请使用 /process-business-response 接口',
                'redirect': '/process-business-response'
            }

        return jsonify(result)

    except Exception as e:
        logger.error(f"文档处理失败: {e}")
        return jsonify(format_error_response(e))


@api_business_bp.route('/api/document/preview/<filename>', methods=['GET'])
def preview_document(filename):
    """预览文档内容 - 直接返回.docx文件供前端mammoth.js转换"""
    try:
        # URL解码文件名（处理中文字符等）
        filename = urllib.parse.unquote(filename)

        # 只进行基本的安全检查，避免路径遍历攻击
        if '..' in filename or '/' in filename or '\\' in filename:
            return jsonify({'success': False, 'error': '非法文件名'}), 400

        # 先尝试从output目录查找，如果不存在则从upload目录查找
        file_path = config.get_path('output') / filename

        if not file_path.exists():
            file_path = config.get_path('upload') / filename

        if not file_path.exists():
            return jsonify({'success': False, 'error': f'文档不存在: {filename}'}), 404

        file_ext = file_path.suffix.lower()

        # 只处理Word文档
        if file_ext not in ['.doc', '.docx']:
            return jsonify({'success': False, 'error': f'不支持的文件格式: {file_ext}'}), 400

        # 如果是.doc格式，提示需要转换
        if file_ext == '.doc':
            return jsonify({
                'success': False,
                'error': '旧版.doc格式预览失败。建议：\n1. 将文件另存为.docx格式\n2. 或直接进行信息提取（系统会自动处理）'
            }), 400

        # 直接返回.docx文件，让前端mammoth.js处理
        return send_file(
            file_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=False,
            download_name=filename
        )

    except Exception as e:
        logger.error(f"文档预览失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@api_business_bp.route('/business-files')
def list_business_files():
    """获取商务应答文件列表"""
    try:
        import os
        from datetime import datetime

        def format_size(size_bytes):
            """格式化文件大小"""
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024.0:
                    return f"{size_bytes:.1f} {unit}"
                size_bytes /= 1024.0
            return f"{size_bytes:.1f} TB"

        files = []
        output_dir = config.get_path('output')

        if output_dir.exists():
            for filename in os.listdir(output_dir):
                # 过滤商务应答文件（只匹配包含"商务应答"的文件）
                if filename.endswith(('.docx', '.doc', '.pdf')) and '商务应答' in filename:
                    file_path = output_dir / filename
                    try:
                        stat = file_path.stat()
                        modified_time = datetime.fromtimestamp(stat.st_mtime)
                        files.append({
                            'name': filename,
                            'size': format_size(stat.st_size),
                            'date': modified_time.strftime('%Y-%m-%d %H:%M:%S'),
                            'download_url': f'/api/files/download/{filename}',
                            'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                            'modified': modified_time.isoformat()
                        })
                    except Exception as e:
                        logger.warning(f"读取文件信息失败 {filename}: {e}")

        files.sort(key=lambda x: x.get('modified', ''), reverse=True)
        return jsonify({'success': True, 'files': files})

    except Exception as e:
        logger.error(f"获取商务文件列表失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ===================
# 点对点应答路由
# ===================

@api_business_bp.route('/process-point-to-point', methods=['POST'])
def process_point_to_point():
    """处理点对点应答 - 使用内联回复功能（原地插入应答）"""
    if not BUSINESS_RESPONSE_AVAILABLE:
        return jsonify({
            'success': False,
            'message': '点对点应答模块不可用'
        })

    try:
        # 检查是否使用HITL技术需求文件
        use_hitl_file = request.form.get('use_hitl_technical_file') == 'true'
        project_id = request.form.get('project_id')

        if use_hitl_file and project_id:
            # 从HITL任务获取技术需求文件
            logger.info(f"使用HITL项目的技术需求文件: project_id={project_id}")

            # 查询HITL任务的技术需求文件路径
            # 技术需求文件保存在 technical_files/{year}/{month}/{project_id}/ 目录下
            from datetime import datetime
            now = datetime.now()
            year = now.strftime('%Y')
            month = now.strftime('%m')

            technical_dir = Path(config.get_path('upload')) / 'technical_files' / year / month / str(project_id)

            if not technical_dir.exists():
                raise ValueError(f"HITL任务技术需求文件目录不存在: {technical_dir}")

            # 查找目录中的docx文件
            docx_files = list(technical_dir.glob('*.docx'))
            if not docx_files:
                raise ValueError(f"HITL任务目录中没有找到技术需求文件: {technical_dir}")

            # 使用第一个找到的docx文件
            file_path = docx_files[0]
            filename = file_path.name
            logger.info(f"使用HITL技术需求文件: {filename}, 路径: {file_path}")
        else:
            # 获取上传的文件
            if 'file' not in request.files:
                raise ValueError("没有选择文件")

            file = request.files['file']
            if file.filename == '':
                raise ValueError("文件名为空")

            # 保存文件 - 使用统一服务
            from core.storage_service import storage_service
            file_metadata = storage_service.store_file(
                file_obj=file,
                original_name=file.filename,
                category='point_to_point',
                business_type='point_to_point_response'
            )
            file_path = Path(file_metadata.file_path)
            filename = file_metadata.original_name

        logger.info(f"开始处理点对点应答: {filename}")

        # 获取公司ID参数
        company_id = request.form.get('companyId')
        if not company_id:
            return jsonify({
                'success': False,
                'error': '缺少公司ID参数'
            })

        # 从数据库获取公司信息
        company_id_int = int(company_id)
        company_db_data = kb_manager.get_company_detail(company_id_int)
        if not company_db_data:
            return jsonify({
                'success': False,
                'error': f'未找到公司数据: {company_id}'
            })

        # 使用现有字段映射反向转换为业务处理器期望的格式
        field_mapping = {
            'companyName': 'company_name',
            'establishDate': 'establish_date',
            'legalRepresentative': 'legal_representative',
            'legalRepresentativePosition': 'legal_representative_position',
            'legalRepresentativeGender': 'legal_representative_gender',
            'legalRepresentativeAge': 'legal_representative_age',
            'socialCreditCode': 'social_credit_code',
            'registeredCapital': 'registered_capital',
            'companyType': 'company_type',
            'registeredAddress': 'registered_address',
            'businessScope': 'business_scope',
            'companyDescription': 'description',
            'fixedPhone': 'fixed_phone',
            'fax': 'fax',
            'postalCode': 'postal_code',
            'email': 'email',
            'officeAddress': 'office_address',
            'employeeCount': 'employee_count',
            'bankName': 'bank_name',
            'bankAccount': 'bank_account',
            # 被授权人信息（2025-11-07添加）
            'representativeName': 'authorized_person_name',
            'representativeTitle': 'authorized_person_position',
            'authorizedPersonId': 'authorized_person_id',
            # 股权结构字段（2025-10-27添加）
            'actual_controller': 'actual_controller',
            'controlling_shareholder': 'controlling_shareholder',
            'shareholders_info': 'shareholders_info',
            # 管理关系字段（2025-10-28添加）
            'managing_unit_name': 'managing_unit_name',
            'managed_unit_name': 'managed_unit_name'
        }
        reverse_mapping = {v: k for k, v in field_mapping.items()}
        company_data = {reverse_mapping.get(k, k): v for k, v in company_db_data.items()}

        logger.info(f"使用公司信息: {company_data.get('companyName', 'N/A')}")

        # 获取配置参数
        response_frequency = request.form.get('responseFrequency', 'every_paragraph')
        response_mode = request.form.get('responseMode', 'simple')
        ai_model = request.form.get('aiModel', 'shihuang-gpt4o-mini')

        # 根据模型选择映射到正确的模型名称
        model_mapping = {
            'gpt-4o-mini': 'shihuang-gpt4o-mini',
            'gpt-4': 'shihuang-gpt4',
            'deepseek-v3': 'yuanjing-deepseek-v3',
            'qwen-235b': 'yuanjing-qwen-235b'
        }
        actual_model = model_mapping.get(ai_model, ai_model)

        logger.info(f"配置参数 - 应答频次: {response_frequency}, 应答方式: {response_mode}, AI模型: {actual_model}")

        # 创建商务应答处理器（使用内联回复功能）
        processor = BusinessResponseProcessor(model_name=actual_model)

        # 获取项目名称参数
        project_name = request.form.get('projectName')

        # 生成输出文件路径
        output_dir = ensure_dir(config.get_path('output'))
        # 使用新的文件命名规则：{项目名称}_点对点应答_{时间戳}.docx
        # 如果有项目名称，使用项目名称；否则使用原文件名
        base_name = project_name if project_name else Path(filename).stem
        output_filename = generate_output_filename(base_name, "点对点应答")
        output_path = output_dir / output_filename

        # 判断是否使用AI（根据应答方式）
        use_ai = response_mode == 'ai'
        logger.info(f"处理模式: {'AI智能应答' if use_ai else '简单模板应答'}")

        # 使用新的内联回复处理方法
        result_stats = processor.process_inline_reply(
            str(file_path),
            str(output_path),
            use_ai=use_ai
        )

        if result_stats.get('success'):
            logger.info(f"内联回复处理成功: {result_stats.get('message')}")

            # 【优化】如果处理成功，直接同步文件信息到数据库（不使用HTTP调用）
            if project_name or project_id:
                try:
                    import json

                    # 1. 查询项目信息（包括project_id和step1_data）
                    # 优先使用project_id，如果没有则使用project_name和company_id
                    if project_id:
                        query = """
                            SELECT project_id, step1_data
                            FROM tender_projects
                            WHERE project_id = ?
                            LIMIT 1
                        """
                        query_params = [project_id]
                    else:
                        query = """
                            SELECT project_id, step1_data
                            FROM tender_projects
                            WHERE project_name = ? AND company_id = ?
                            LIMIT 1
                        """
                        query_params = [project_name, company_id_int]

                    project_result = kb_manager.db.execute_query(
                        query, query_params, fetch_one=True
                    )

                    if project_result:
                        # 2. 解析现有的step1_data
                        step1_data_raw = project_result.get('step1_data', '{}')
                        step1_data = json.loads(step1_data_raw) if step1_data_raw else {}

                        # 3. 构建文件信息
                        from datetime import datetime
                        now = datetime.now()
                        file_info = {
                            "file_path": str(output_path),
                            "filename": output_filename,
                            "file_name": output_filename,  # 兼容性字段
                            "file_size": os.path.getsize(output_path),
                            "file_url": f"/api/files/download/{output_filename}",
                            "download_url": f"/api/files/download/{output_filename}",
                            "saved_at": now.isoformat(),
                            "source": "point_to_point_api"
                        }

                        # 4. 更新step1_data
                        step1_data['technical_point_to_point_file'] = file_info

                        # 5. 保存到数据库
                        update_query = """
                            UPDATE tender_projects
                            SET step1_data = ?, updated_at = CURRENT_TIMESTAMP
                            WHERE project_id = ?
                        """
                        kb_manager.db.execute_query(
                            update_query,
                            [json.dumps(step1_data, ensure_ascii=False), project_id]
                        )

                        logger.info(
                            f"✅ 点对点应答文件已同步到数据库: "
                            f"project_id={project_id}, "
                            f"file={output_filename}, "
                            f"size={file_info['file_size']} bytes"
                        )
                    else:
                        logger.warning(
                            f"⚠️  未找到匹配的项目记录: "
                            f"project_id={project_id}"
                        )
                except Exception as e:
                    logger.error(
                        f"❌ 同步点对点应答文件到数据库失败: {e}",
                        exc_info=True  # 打印完整堆栈
                    )
                    # 注意：不影响主流程，文件已生成成功

            # 生成下载URL（使用 /api/downloads/ 路由）
            download_url = f'/api/downloads/{output_filename}'

            return jsonify({
                'success': True,
                'message': '内联回复处理完成，应答已插入到原文档中（灰色底纹标记）',
                'download_url': download_url,
                'filename': output_filename,
                'output_file': str(output_path),
                'model_used': actual_model,
                'requirements_count': result_stats.get('requirements_count', 0),
                'responses_count': result_stats.get('responses_count', 0),
                'response_mode': response_mode,
                'model_name': actual_model if use_ai else None,
                'features': result_stats.get('features', {}),
                'stats': {
                    'inline_reply': True,
                    'gray_shading': True,
                    'format_preserved': True,
                    'requirements_count': result_stats.get('requirements_count', 0),
                    'responses_count': result_stats.get('responses_count', 0),
                    'response_mode': response_mode,
                    'model_name': actual_model if use_ai else None
                }
            })
        else:
            logger.error(f"内联回复处理失败: {result_stats.get('error')}")
            return jsonify({
                'success': False,
                'error': result_stats.get('error', '处理失败'),
                'message': result_stats.get('message', '内联回复处理失败')
            })

    except Exception as e:
        logger.error(f"点对点应答处理失败: {e}")
        return jsonify(format_error_response(e))


@api_business_bp.route('/point-to-point/files')
def list_point_to_point_files():
    """获取点对点应答文件列表（自动去重，每个项目只保留最新的文件）"""
    try:
        import os
        from datetime import datetime
        import re

        files = []
        output_dir = config.get_path('output')

        if output_dir.exists():
            for filename in os.listdir(output_dir):
                # 过滤点对点应答文件（只匹配包含"点对点应答"或"点对点"的文件）
                if filename.endswith(('.docx', '.doc', '.pdf')) and ('点对点应答' in filename or '点对点' in filename or 'point-to-point' in filename.lower()):
                    file_path = output_dir / filename
                    try:
                        stat = file_path.stat()

                        # 提取项目名称(文件名格式: {项目名称}_点对点应答_{时间戳}.docx)
                        # 使用正则提取项目名称部分(去掉后面的时间戳)
                        project_name_match = re.match(r'^(.+?)_点对点应答_\d{8}_\d{6}\.', filename)
                        project_name = project_name_match.group(1) if project_name_match else filename

                        files.append({
                            'id': hashlib.md5(str(file_path).encode()).hexdigest()[:8],
                            'filename': filename,
                            'original_filename': filename,
                            'file_path': str(file_path),
                            'output_path': str(file_path),
                            'size': stat.st_size,
                            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                            'process_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                            'status': 'completed',
                            'company_name': '未知公司',  # 暂时使用默认值，后续会从数据库获取
                            'project_name': project_name  # 添加项目名称用于去重
                        })
                    except Exception as e:
                        logger.warning(f"读取文件信息失败 {filename}: {e}")

        # 按时间排序(最新的在前)
        files.sort(key=lambda x: x.get('process_time', ''), reverse=True)

        # 去重:每个项目只保留最新的文件
        unique_files = []
        seen_projects = set()
        for file in files:
            project_name = file.get('project_name', file['filename'])
            if project_name not in seen_projects:
                unique_files.append(file)
                seen_projects.add(project_name)
                logger.debug(f"保留文件: {file['filename']} (项目: {project_name})")
            else:
                logger.debug(f"跳过重复项目文件: {file['filename']} (项目: {project_name})")

        logger.info(f"文件去重: 原始{len(files)}个 -> 去重后{len(unique_files)}个")

        return jsonify({'success': True, 'data': unique_files})

    except Exception as e:
        logger.error(f"获取点对点应答文件列表失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_business_bp.route('/point-to-point/preview')
def preview_point_to_point_document():
    """预览点对点应答文档 - 直接返回.docx文件供前端mammoth.js转换"""
    try:
        # 获取参数
        file_id = request.args.get('file_id')
        file_path = request.args.get('file_path')

        if not file_id and not file_path:
            return jsonify({'success': False, 'error': '缺少文件ID或文件路径参数'}), 400

        # 根据参数确定文件路径
        if file_path:
            target_file = Path(file_path)
        else:
            # 如果只有file_id，需要从输出目录查找文件
            output_dir = config.get_path('output')
            target_file = None
            if output_dir.exists():
                for filename in os.listdir(output_dir):
                    full_path = output_dir / filename
                    if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                        target_file = full_path
                        break

            if not target_file:
                return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

        if not target_file.exists():
            return jsonify({'success': False, 'error': '文件不存在'}), 404

        # 根据文件类型进行预览
        file_extension = target_file.suffix.lower()

        if file_extension == '.docx':
            # 直接返回.docx文件，让前端mammoth.js处理
            return send_file(
                target_file,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=False,
                download_name=target_file.name
            )

        elif file_extension == '.doc':
            return jsonify({
                'success': False,
                'error': '旧版.doc格式预览失败。建议：\n1. 将文件另存为.docx格式\n2. 或直接进行信息提取（系统会自动处理）'
            }), 400

        elif file_extension == '.pdf':
            # PDF预览（简单实现，返回提示信息）
            return jsonify({
                'success': True,
                'content': f'<div class="alert alert-info"><h4>PDF文档预览</h4><p>文件名: {target_file.name}</p><p>PDF文件预览功能正在开发中，请使用下载功能查看完整内容。</p></div>',
                'filename': target_file.name
            })

        else:
            return jsonify({'success': False, 'error': '不支持的文件格式'}), 400

    except Exception as e:
        logger.error(f"文档预览失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_business_bp.route('/point-to-point/edit', methods=['GET', 'POST'])
def edit_point_to_point_document():
    """编辑点对点应答文档"""

    if request.method == 'GET':
        # 获取文档内容用于编辑
        try:
            from docx import Document

            # 获取参数
            file_id = request.args.get('file_id')
            file_path = request.args.get('file_path')

            if not file_id and not file_path:
                return jsonify({'success': False, 'error': '缺少文件ID或文件路径参数'}), 400

            # 根据参数确定文件路径
            if file_path:
                target_file = Path(file_path)
            else:
                # 如果只有file_id，需要从输出目录查找文件
                output_dir = config.get_path('output')
                target_file = None
                if output_dir.exists():
                    for filename in os.listdir(output_dir):
                        full_path = output_dir / filename
                        if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                            target_file = full_path
                            break

                if not target_file:
                    return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

            if not target_file.exists():
                return jsonify({'success': False, 'error': '文件不存在'}), 404

            # 只支持Word文档编辑
            file_extension = target_file.suffix.lower()
            if file_extension not in ['.docx', '.doc']:
                return jsonify({'success': False, 'error': '只支持Word文档编辑'}), 400

            try:
                # 读取Word文档并转换为可编辑的HTML
                doc = Document(target_file)
                html_content = []

                # 处理段落
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        style_name = paragraph.style.name if paragraph.style else ''
                        text = html.escape(paragraph.text)

                        if 'Heading 1' in style_name or 'heading 1' in style_name.lower():
                            html_content.append(f'<h1>{text}</h1>')
                        elif 'Heading 2' in style_name or 'heading 2' in style_name.lower():
                            html_content.append(f'<h2>{text}</h2>')
                        elif 'Heading 3' in style_name or 'heading 3' in style_name.lower():
                            html_content.append(f'<h3>{text}</h3>')
                        else:
                            html_content.append(f'<p>{text}</p>')

                # 处理表格（简化为文本形式）
                for table in doc.tables:
                    html_content.append('<table border="1">')
                    for row in table.rows:
                        html_content.append('<tr>')
                        for cell in row.cells:
                            cell_text = html.escape(cell.text)
                            html_content.append(f'<td>{cell_text}</td>')
                        html_content.append('</tr>')
                    html_content.append('</table>')

                return jsonify({
                    'success': True,
                    'content': '\n'.join(html_content),
                    'filename': target_file.name
                })

            except Exception as e:
                logger.error(f"读取文档内容失败: {e}")
                return jsonify({'success': False, 'error': f'读取文档内容失败: {str(e)}'}), 500

        except Exception as e:
            logger.error(f"获取编辑内容失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    elif request.method == 'POST':
        # 保存编辑后的文档
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            import re

            # 获取参数
            file_id = request.args.get('file_id')
            file_path = request.args.get('file_path')

            # 获取POST数据
            data = request.get_json()
            if not data or 'content' not in data:
                return jsonify({'success': False, 'error': '缺少文档内容'}), 400

            new_content = data['content']

            # 根据参数确定文件路径
            if file_path:
                target_file = Path(file_path)
            else:
                # 如果只有file_id，需要从输出目录查找文件
                output_dir = config.get_path('output')
                target_file = None
                if output_dir.exists():
                    for filename in os.listdir(output_dir):
                        full_path = output_dir / filename
                        if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                            target_file = full_path
                            break

                if not target_file:
                    return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

            if not target_file.exists():
                return jsonify({'success': False, 'error': '文件不存在'}), 404

            try:
                # 解析HTML内容
                soup = BeautifulSoup(new_content, 'html.parser')

                # 创建新的Word文档
                doc = Document()

                # 遍历解析的HTML元素
                for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table']):
                    if element.name in ['h1', 'h2', 'h3']:
                        # 添加标题
                        heading_level = int(element.name[1])
                        paragraph = doc.add_heading(element.get_text().strip(), level=heading_level)
                    elif element.name == 'p':
                        # 添加段落
                        doc.add_paragraph(element.get_text().strip())
                    elif element.name == 'table':
                        # 添加表格
                        rows = element.find_all('tr')
                        if rows:
                            cols = len(rows[0].find_all(['td', 'th']))
                            table = doc.add_table(rows=len(rows), cols=cols)
                            table.style = 'Table Grid'

                            for i, row in enumerate(rows):
                                cells = row.find_all(['td', 'th'])
                                for j, cell in enumerate(cells):
                                    if i < len(table.rows) and j < len(table.rows[i].cells):
                                        table.rows[i].cells[j].text = cell.get_text().strip()

                # 保存文档
                doc.save(str(target_file))

                logger.info(f"文档保存成功: {target_file}")

                return jsonify({
                    'success': True,
                    'message': '文档保存成功',
                    'filename': target_file.name
                })

            except Exception as e:
                logger.error(f"保存文档失败: {e}")
                return jsonify({'success': False, 'error': f'保存文档失败: {str(e)}'}), 500

        except Exception as e:
            logger.error(f"编辑文档失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500


@api_business_bp.route('/point-to-point/download')
def download_point_to_point_document():
    """下载点对点应答文档"""
    try:
        # 获取参数
        file_id = request.args.get('file_id')
        file_path = request.args.get('file_path')

        if not file_id and not file_path:
            return jsonify({'success': False, 'error': '缺少文件ID或文件路径参数'}), 400

        # 根据参数确定文件路径
        if file_path:
            target_file = Path(file_path)
        else:
            # 如果只有file_id，需要从输出目录查找文件
            output_dir = config.get_path('output')
            target_file = None
            if output_dir.exists():
                for filename in os.listdir(output_dir):
                    full_path = output_dir / filename
                    if hashlib.md5(str(full_path).encode()).hexdigest()[:8] == file_id:
                        target_file = full_path
                        break

            if not target_file:
                return jsonify({'success': False, 'error': '找不到指定的文件'}), 404

        if not target_file.exists():
            return jsonify({'success': False, 'error': '文件不存在'}), 404

        # 确定MIME类型
        file_extension = target_file.suffix.lower()
        if file_extension == '.docx':
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_extension == '.doc':
            mimetype = 'application/msword'
        elif file_extension == '.pdf':
            mimetype = 'application/pdf'
        else:
            mimetype = 'application/octet-stream'

        # 生成下载文件名
        download_filename = target_file.name

        logger.info(f"开始下载文件: {target_file}")

        return send_file(
            str(target_file),
            as_attachment=True,
            download_name=download_filename,
            mimetype=mimetype
        )

    except Exception as e:
        logger.error(f"文档下载失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


__all__ = ['api_business_bp']
