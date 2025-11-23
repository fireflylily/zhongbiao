#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档编辑器和表格处理API蓝图
提供文档加载、保存、图片上传、表格处理和AI助手功能

路由列表:
- POST /api/editor/load-document - 加载文档
- POST /api/editor/save-document - 保存文档
- POST /api/editor/upload-image - 上传图片
- POST /api/editor/ai-assistant - AI助手（改写、扩写、总结、翻译）
- POST /api/table/analyze - 分析表格
- POST /api/table/process - 处理表格
"""

import sys
from pathlib import Path
from flask import Blueprint, request, jsonify, send_file

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

# 导入公共组件
from common import (
    get_config, get_module_logger,
    format_error_response, ensure_dir
)

# 创建蓝图
api_editor_bp = Blueprint('api_editor', __name__, url_prefix='/api')

# 获取配置和日志器
config = get_config()
logger = get_module_logger("api_editor_bp")

# 检查表格处理模块可用性
try:
    from modules.business_response.table_processor import TableProcessor
    POINT_TO_POINT_AVAILABLE = True
except ImportError as e:
    logger.warning(f"表格处理模块加载失败: {e}")
    POINT_TO_POINT_AVAILABLE = False


# ===================
# 文档编辑器API
# ===================

@api_editor_bp.route('/editor/load-document', methods=['POST'])
def load_document_for_edit():
    """加载文档用于编辑"""
    try:
        from docx import Document
        from markupsafe import Markup

        if 'file' not in request.files:
            raise ValueError("没有选择文件")

        file = request.files['file']
        if file.filename == '':
            raise ValueError("文件名为空")

        # 读取Word文档
        doc = Document(file)

        # 转换为HTML格式用于编辑器
        html_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                style_name = paragraph.style.name if paragraph.style else ''

                if 'Heading 1' in style_name:
                    html_content.append(f'<h1>{paragraph.text}</h1>')
                elif 'Heading 2' in style_name:
                    html_content.append(f'<h2>{paragraph.text}</h2>')
                elif 'Heading 3' in style_name:
                    html_content.append(f'<h3>{paragraph.text}</h3>')
                else:
                    html_content.append(f'<p>{paragraph.text}</p>')

        # 处理表格
        for table in doc.tables:
            html_content.append('<table>')
            for row in table.rows:
                html_content.append('<tr>')
                for cell in row.cells:
                    html_content.append(f'<td>{cell.text}</td>')
                html_content.append('</tr>')
            html_content.append('</table>')

        return jsonify({
            'success': True,
            'html_content': ''.join(html_content),
            'original_filename': file.filename
        })

    except Exception as e:
        logger.error(f"文档加载失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        })


@api_editor_bp.route('/editor/save-document', methods=['POST'])
def save_edited_document():
    """保存编辑后的文档"""
    try:
        from docx import Document
        from bs4 import BeautifulSoup
        import re

        data = request.get_json()
        html_content = data.get('html_content', '')
        filename = data.get('filename', 'document')

        if not html_content:
            raise ValueError("文档内容为空")

        # 创建新文档
        doc = Document()

        # 解析HTML内容
        soup = BeautifulSoup(html_content, 'html.parser')

        # 处理各种HTML元素
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table']):
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p':
                text = element.get_text()
                if text.strip():
                    doc.add_paragraph(text)
            elif element.name == 'table':
                # 计算表格行列数
                rows = element.find_all('tr')
                if rows:
                    cols = len(rows[0].find_all(['td', 'th']))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Table Grid'

                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < cols:
                                table.cell(i, j).text = cell.get_text()

        # 保存文档
        output_dir = ensure_dir(config.get_path('output'))
        output_path = output_dir / f"{filename}.docx"
        doc.save(str(output_path))

        # 返回文件供下载
        return send_file(
            str(output_path),
            as_attachment=True,
            download_name=f"{filename}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"文档保存失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        })


@api_editor_bp.route('/editor/upload-image', methods=['POST'])
def upload_editor_image():
    """上传编辑器图片"""
    try:
        if 'image' not in request.files:
            raise ValueError("没有选择图片")

        file = request.files['image']
        if file.filename == '':
            raise ValueError("文件名为空")

        # 保存图片 - 使用统一服务
        from core.storage_service import storage_service
        file_metadata = storage_service.store_file(
            file_obj=file,
            original_name=file.filename,
            category='processed_results',  # 编辑器图片属于处理结果
            business_type='editor_image'
        )
        file_path = Path(file_metadata.file_path)
        filename = file_metadata.safe_name

        # 返回图片URL
        image_url = f'/static/uploads/images/{filename}'

        return jsonify({
            'success': True,
            'location': image_url
        })

    except Exception as e:
        logger.error(f"图片上传失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        })


# ===================
# 表格处理API
# ===================

@api_editor_bp.route('/table/analyze', methods=['POST'])
def analyze_table():
    """分析表格"""
    if not POINT_TO_POINT_AVAILABLE:
        return jsonify({
            'success': False,
            'message': '表格处理模块不可用'
        })

    try:
        data = request.get_json()
        table_data = data.get('table_data', {})

        processor = TableProcessor()
        result = processor.analyze_table(table_data)

        return jsonify(result)

    except Exception as e:
        logger.error(f"表格分析失败: {e}")
        return jsonify(format_error_response(e))


@api_editor_bp.route('/table/process', methods=['POST'])
def process_table():
    """处理表格"""
    if not POINT_TO_POINT_AVAILABLE:
        return jsonify({
            'success': False,
            'message': '表格处理模块不可用'
        })

    try:
        data = request.get_json()
        table_data = data.get('table_data', {})
        options = data.get('options', {})

        processor = TableProcessor()
        result = processor.process_table(table_data, options)

        return jsonify(result)

    except Exception as e:
        logger.error(f"表格处理失败: {e}")
        return jsonify(format_error_response(e))


# ===================
# AI 助手API
# ===================

@api_editor_bp.route('/editor/ai-assistant', methods=['POST'])
def ai_assistant():
    """
    UmoEditor AI 助手接口

    支持的 command 类型：
    - rewrite: 改写内容
    - expand: 扩写内容
    - summarize: 总结内容
    - translate: 翻译内容
    - custom: 自定义生成
    """
    try:
        data = request.get_json()

        command = data.get('command')        # AI 操作类型
        input_text = data.get('input')       # 用户选中的文本
        output_format = data.get('output', 'html')  # 输出格式
        lang = data.get('lang', 'zh-CN')     # 语言
        company_id = data.get('company_id')  # 公司ID
        project_name = data.get('project_name')  # 项目名称

        if not input_text:
            return jsonify({
                'success': False,
                'error': '缺少输入内容'
            }), 400

        # 构建 AI 提示词（根据 command 类型）
        system_prompt, user_prompt = build_ai_prompts(
            command, input_text, lang, company_id, project_name
        )

        # 调用 LLM 生成内容
        from common.llm_client import create_llm_client

        # 使用配置的默认模型（或指定模型）
        client = create_llm_client("gpt-4o-mini")

        # 调用 AI（非流式）
        generated_content = client.call(
            prompt=user_prompt,
            system_prompt=system_prompt,
            temperature=0.7,
            max_tokens=2000,
            purpose=f"AI助手-{command}"
        )

        logger.info(f"AI助手生成成功: command={command}, 输入长度={len(input_text)}, 输出长度={len(generated_content)}")

        return jsonify({
            'success': True,
            'content': generated_content
        })

    except Exception as e:
        logger.error(f"AI助手调用失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def build_ai_prompts(command: str, input_text: str, lang: str,
                     company_id, project_name):
    """
    根据 command 类型构建 AI 提示词

    Returns:
        (system_prompt, user_prompt) 元组
    """
    # 系统提示词（定义 AI 的角色和行为）
    system_prompt = """你是一个专业的标书撰写助手，擅长编写规范、专业、有说服力的标书内容。

你的职责：
1. 根据用户的要求生成或改写标书内容
2. 确保内容专业、规范、符合招标要求
3. 输出格式使用 HTML 标签（支持 h1, h2, h3, p, ul, ol, li, table, strong, em 等）
4. 保持内容简洁明了，重点突出

输出要求：
- 使用正确的 HTML 标签格式
- 段落使用 <p></p>
- 列表使用 <ul><li></li></ul> 或 <ol><li></li></ol>
- 表格使用 <table><tr><td></td></tr></table>
- 标题使用 <h2></h2> 或 <h3></h3>
- 强调使用 <strong></strong>
"""

    # 根据 command 类型构建用户提示词

    # 【新增】标书专用指令
    if command == 'auto_fill' or '智能填写' in str(command):
        # 智能填写：识别占位符类型并生成内容
        user_prompt = f"""请识别以下占位符或待填写内容，并生成专业的标书内容：

选中内容：
{input_text}

识别规则：
- 如果包含"格式自拟"：生成该部分的完整专业内容（300-500字）
- 如果包含"待填写"：根据上下文生成合适内容
- 如果是描述性需求：直接生成对应内容

要求：
1. 内容专业、规范、有说服力
2. 符合标书编写规范
3. 结构清晰、重点突出
4. 输出格式为 HTML（使用 p, ul, h3 等标签）"""

    elif command == 'generate_table' or '生成表格' in str(command):
        # 生成表格：根据描述生成标准表格
        user_prompt = f"""请根据以下描述生成一个标准的HTML表格：

用户要求：
{input_text}

要求：
1. 使用标准的HTML table标签格式
2. 表格包含表头（thead）和表体（tbody）
3. 内容专业、规范、符合标书要求
4. 如果未指定列数和行数，请合理设计
5. 输出纯HTML代码（不要markdown）

示例格式：
<table border="1">
  <thead>
    <tr><th>列1</th><th>列2</th></tr>
  </thead>
  <tbody>
    <tr><td>数据1</td><td>数据2</td></tr>
  </tbody>
</table>"""

    elif command == 'rewrite':
        user_prompt = f"""请改写以下内容，使其更加专业、规范、有说服力：

原始内容：
{input_text}

要求：
1. 保持原意，但提升专业性和可读性
2. 使用规范的标书用语
3. 输出格式为 HTML"""

    elif command == 'expand':
        user_prompt = f"""请扩写以下内容，补充更多细节和说明：

原始内容：
{input_text}

要求：
1. 扩展为 300-500 字的详细描述
2. 补充技术细节、实施方案或优势说明
3. 保持逻辑清晰，层次分明
4. 输出格式为 HTML"""

    elif command == 'summarize':
        user_prompt = f"""请总结以下内容的核心要点：

原始内容：
{input_text}

要求：
1. 提取 3-5 个核心要点
2. 使用列表形式展示
3. 保持简洁明了
4. 输出格式为 HTML 列表"""

    elif command == 'translate':
        target_lang = '英文' if lang == 'zh-CN' else '中文'
        user_prompt = f"""请将以下内容翻译为{target_lang}：

原始内容：
{input_text}

要求：
1. 保持专业术语的准确性
2. 符合标书用语规范
3. 输出格式为 HTML"""

    elif command == 'generate' or command == 'custom':
        # 自定义生成（用户输入的是生成需求描述）
        user_prompt = f"""请根据以下要求生成标书内容：

用户要求：
{input_text}

要求：
1. 生成 300-800 字的专业内容
2. 内容符合标书规范
3. 结构清晰，逻辑严谨
4. 输出格式为 HTML"""

    else:
        # 默认：改写
        user_prompt = f"""请优化以下标书内容：

{input_text}

要求：输出格式为 HTML"""

    return (system_prompt, user_prompt)


__all__ = ['api_editor_bp']
