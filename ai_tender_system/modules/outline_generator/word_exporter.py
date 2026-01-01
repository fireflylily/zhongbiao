#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档导出器
将方案数据导出为Word文档格式
"""

import os
import re
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import openpyxl

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger


class WordExporter:
    """Word文档导出器"""

    def __init__(self):
        """初始化Word导出器"""
        self.logger = get_module_logger("word_exporter")
        # 延迟导入 MermaidRenderer，避免循环依赖
        self._mermaid_renderer = None
        self.logger.info("Word导出器初始化完成")

    @property
    def mermaid_renderer(self):
        """延迟初始化 Mermaid 渲染器"""
        if self._mermaid_renderer is None:
            from .mermaid_renderer import MermaidRenderer
            self._mermaid_renderer = MermaidRenderer()
        return self._mermaid_renderer

    def export_proposal(
        self,
        proposal: Dict[str, Any],
        output_path: str,
        show_guidance: bool = False
    ) -> str:
        """
        导出技术方案为Word文档

        Args:
            proposal: 方案数据
            output_path: 输出文件路径
            show_guidance: 是否显示大纲指导信息（默认False，简洁模式）

        Returns:
            生成的文件路径
        """
        try:
            self.logger.info(f"开始导出技术方案到: {output_path}，指导信息: {'显示' if show_guidance else '隐藏'}")

            # 创建Word文档
            doc = Document()

            # 设置文档样式
            self._setup_document_styles(doc)

            # 添加标题
            self._add_title(doc, proposal['metadata']['title'])

            # 添加元信息
            self._add_metadata(doc, proposal['metadata'])

            # ✅ 添加目录
            self._add_table_of_contents(doc)

            # 分页符（目录独立一页）
            doc.add_page_break()

            # 添加章节内容
            for chapter in proposal['chapters']:
                self._add_chapter(doc, chapter, show_guidance=show_guidance)

            # 保存文档
            doc.save(output_path)

            self.logger.info(f"技术方案导出完成: {output_path}")
            return output_path

        except Exception as e:
            self.logger.error(f"导出技术方案失败: {e}", exc_info=True)
            raise

    def export_analysis_report(
        self,
        analysis: Dict[str, Any],
        output_path: str
    ) -> str:
        """
        导出需求分析报告

        Args:
            analysis: 需求分析数据
            output_path: 输出文件路径

        Returns:
            生成的文件路径
        """
        try:
            self.logger.info(f"开始导出需求分析报告到: {output_path}")

            doc = Document()
            self._setup_document_styles(doc)

            # 标题
            self._add_title(doc, "技术需求分析报告")

            # 文档摘要
            summary = analysis.get('document_summary', {})
            doc.add_heading('一、文档摘要', level=1)

            doc.add_paragraph(f"总需求数量: {summary.get('total_requirements', 0)}")
            doc.add_paragraph(f"强制需求: {summary.get('mandatory_count', 0)}")
            doc.add_paragraph(f"可选需求: {summary.get('optional_count', 0)}")
            doc.add_paragraph(f"评分项目: {summary.get('scoring_items', 0)}")
            doc.add_paragraph(f"复杂度: {summary.get('complexity_level', 'medium')}")

            # 需求分类
            doc.add_heading('二、需求分类', level=1)

            for i, category in enumerate(analysis.get('requirement_categories', []), start=1):
                doc.add_heading(f"2.{i} {category.get('category', '未分类')}", level=2)

                doc.add_paragraph(f"需求数量: {category.get('requirements_count', 0)}")
                doc.add_paragraph(f"优先级: {category.get('priority', 'medium')}")
                doc.add_paragraph(f"摘要: {category.get('summary', '')}")

                # 关键点
                if category.get('key_points'):
                    doc.add_paragraph("关键需求:")
                    for point in category['key_points']:
                        doc.add_paragraph(point, style='List Bullet')

            # 保存
            doc.save(output_path)

            self.logger.info(f"需求分析报告导出完成: {output_path}")
            return output_path

        except Exception as e:
            self.logger.error(f"导出需求分析报告失败: {e}", exc_info=True)
            raise

    def export_mapping_table(
        self,
        mapping_data: List[Dict],
        output_path: str
    ) -> str:
        """
        导出需求匹配表为Excel

        Args:
            mapping_data: 匹配表数据
            output_path: 输出文件路径

        Returns:
            生成的文件路径
        """
        try:
            self.logger.info(f"开始导出需求匹配表到: {output_path}")

            # 创建Excel工作簿
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "需求匹配表"

            # 添加表头
            headers = ['序号', '需求类别', '具体需求', '匹配产品文档', '匹配状态']
            ws.append(headers)

            # 设置表头样式
            for cell in ws[1]:
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(
                    start_color="CCE5FF",
                    end_color="CCE5FF",
                    fill_type="solid"
                )

            # 添加数据行
            for i, row in enumerate(mapping_data, start=1):
                ws.append([
                    i,
                    row.get('category', ''),
                    row.get('requirement', ''),
                    ', '.join(row.get('matched_docs', [])),
                    row.get('match_status', '')
                ])

            # 调整列宽
            ws.column_dimensions['A'].width = 8
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 40
            ws.column_dimensions['D'].width = 30
            ws.column_dimensions['E'].width = 12

            # 保存
            wb.save(output_path)

            self.logger.info(f"需求匹配表导出完成: {output_path}")
            return output_path

        except Exception as e:
            self.logger.error(f"导出需求匹配表失败: {e}", exc_info=True)
            raise

    def export_summary_report(
        self,
        report_data: Dict[str, Any],
        output_path: str
    ) -> str:
        """
        导出生成报告为文本文件

        Args:
            report_data: 报告数据
            output_path: 输出文件路径

        Returns:
            生成的文件路径
        """
        try:
            self.logger.info(f"开始导出生成报告到: {output_path}")

            lines = [
                "=" * 60,
                "技术方案生成报告",
                "=" * 60,
                "",
                "一、生成信息",
                f"  方案标题: {report_data['generation_info']['title']}",
                f"  生成时间: {report_data['generation_info']['generation_time']}",
                f"  总章节数: {report_data['generation_info']['total_chapters']}",
                "",
                "二、需求统计",
                f"  总需求数: {report_data['requirements_summary']['total_requirements']}",
                f"  强制需求: {report_data['requirements_summary']['mandatory_count']}",
                f"  可选需求: {report_data['requirements_summary']['optional_count']}",
                f"  复杂度: {report_data['requirements_summary']['complexity_level']}",
                "",
                "三、匹配统计",
                f"  匹配文档数: {report_data['matching_summary']['total_documents_matched']}",
                f"  匹配类别数: {report_data['matching_summary']['categories_with_matches']}",
                f"  匹配成功率: {report_data['matching_summary']['match_rate']}%",
                "",
                "=" * 60
            ]

            # 写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))

            self.logger.info(f"生成报告导出完成: {output_path}")
            return output_path

        except Exception as e:
            self.logger.error(f"导出生成报告失败: {e}", exc_info=True)
            raise

    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_title(self, doc: Document, title: str):
        """添加文档标题"""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 设置标题字体
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(18)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def _add_metadata(self, doc: Document, metadata: Dict):
        """添加元信息"""
        doc.add_paragraph(f"生成时间: {metadata.get('generation_time', '')}")
        doc.add_paragraph(f"总章节数: {metadata.get('total_chapters', 0)}")
        doc.add_paragraph(f"预计页数: {metadata.get('estimated_pages', 0)}")
        doc.add_paragraph()  # 空行

    def _add_table_of_contents(self, doc: Document):
        """添加目录（Table of Contents）"""
        # 添加"目录"标题
        heading = doc.add_heading('目录', level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 使用python-docx添加TOC域
        # 注意：TOC域需要在Word中手动更新（右键->更新域）或使用宏自动更新
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # 添加TOC域代码
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # 显示1-3级标题

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        doc.add_paragraph()  # 空行

    def _add_chapter(self, doc: Document, chapter: Dict, show_guidance: bool = False):
        """
        添加章节

        Args:
            doc: Word文档对象
            chapter: 章节数据
            show_guidance: 是否显示大纲指导信息（默认False）
        """
        level = chapter.get('level', 1)
        chapter_num = chapter.get('chapter_number', '')
        title = chapter.get('title', '')

        # 清理 title 中已有的编号前缀，避免重复（如 "第四章 系统架构" -> "系统架构"）
        import re
        title = re.sub(r'^第[一二三四五六七八九十百千]+章\s*', '', title)
        title = re.sub(r'^\d+[\.\s]+', '', title)
        title = title.strip()

        # 添加章节标题
        doc.add_heading(f"{chapter_num} {title}", level=level)

        # ✅ 可选：显示大纲指导信息（默认不显示，交付文档更简洁）
        if show_guidance:
            # 添加章节描述
            if chapter.get('description'):
                doc.add_paragraph(f"【本章说明】{chapter['description']}")

            # 添加应答策略
            if chapter.get('response_strategy'):
                p = doc.add_paragraph()
                p.add_run("【应答策略】").bold = True
                p.add_run(chapter['response_strategy'])

            # 添加内容提示
            if chapter.get('content_hints'):
                p = doc.add_paragraph()
                p.add_run("【内容提示】").bold = True
                for hint in chapter['content_hints']:
                    doc.add_paragraph(hint, style='List Bullet')

            # 添加应答建议
            if chapter.get('response_tips'):
                p = doc.add_paragraph()
                p.add_run("【应答建议】").bold = True
                p.add_run().font.color.rgb = RGBColor(0, 112, 192)  # 蓝色
                for tip in chapter['response_tips']:
                    doc.add_paragraph(tip, style='List Bullet')

            # 添加建议引用文档
            if chapter.get('suggested_references'):
                p = doc.add_paragraph()
                p.add_run("【建议引用文档】").bold = True
                for ref in chapter['suggested_references']:
                    doc_name = ref.get('doc_name', '')
                    reason = ref.get('reason', '')
                    doc.add_paragraph(f"• {doc_name} - {reason}")

            # 添加证明材料清单
            if chapter.get('evidence_needed'):
                p = doc.add_paragraph()
                p.add_run("【需提供证明材料】").bold = True
                p.add_run().font.color.rgb = RGBColor(255, 0, 0)  # 红色
                for evidence in chapter['evidence_needed']:
                    doc.add_paragraph(f"• {evidence}")

            # 分隔线
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("【AI生成内容】").bold = True
            p.add_run().font.color.rgb = RGBColor(0, 176, 80)  # 绿色标注

        # 添加AI生成的章节内容（总是显示）
        if chapter.get('ai_generated_content'):
            # ✅ 解析AI生成的Markdown格式内容并转换为Word格式
            ai_content = chapter['ai_generated_content']
            self._add_markdown_content(doc, ai_content)

        # 添加表格（新增）
        for table_data in chapter.get('tables', []):
            self._add_markdown_table(doc, table_data)

        # 添加流程图（新增）
        for flowchart_data in chapter.get('flowcharts', []):
            self._add_mermaid_flowchart(doc, flowchart_data)

        # 添加子章节
        for subsection in chapter.get('subsections', []):
            self._add_chapter(doc, subsection, show_guidance=show_guidance)

        # 章节结束后添加空行
        doc.add_paragraph()

    def _add_markdown_content(self, doc: Document, markdown_text):
        """
        将Markdown格式文本转换为Word格式并添加到文档

        Args:
            doc: Word文档对象
            markdown_text: Markdown格式的文本，或包含 'content' 键的字典
        """
        import re

        def extract_string_content(data, max_depth=5):
            """递归提取字符串内容，防止无限循环"""
            if max_depth <= 0:
                return str(data) if data else ''

            if isinstance(data, str):
                return data
            elif isinstance(data, dict):
                # 优先尝试 'content' 键
                if 'content' in data:
                    return extract_string_content(data['content'], max_depth - 1)
                # 其次尝试 'text' 键
                elif 'text' in data:
                    return extract_string_content(data['text'], max_depth - 1)
                # 尝试将所有字符串值连接起来
                else:
                    parts = []
                    for v in data.values():
                        if isinstance(v, str):
                            parts.append(v)
                    return '\n'.join(parts) if parts else str(data)
            elif isinstance(data, list):
                # 处理列表：连接所有字符串元素
                parts = [extract_string_content(item, max_depth - 1) for item in data]
                return '\n'.join(filter(None, parts))
            else:
                return str(data) if data else ''

        # 递归提取字符串内容
        markdown_text = extract_string_content(markdown_text)

        # 最终确保是字符串
        if not isinstance(markdown_text, str):
            markdown_text = str(markdown_text) if markdown_text else ''

        lines = markdown_text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 处理标题：### 三级标题 → Heading 3
            if line.startswith('### '):
                title_text = line[4:].strip()
                doc.add_heading(title_text, level=3)
            # 处理标题：## 二级标题 → Heading 2
            elif line.startswith('## '):
                title_text = line[3:].strip()
                doc.add_heading(title_text, level=2)
            # 处理有序列表：1. 内容
            elif re.match(r'^\d+\.\s', line):
                # 移除数字和点号
                text = re.sub(r'^\d+\.\s+', '', line)
                # 处理加粗标记
                text = self._process_inline_markdown(text)
                # 添加为列表项
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, text)
            # 处理无序列表：- 内容 或 * 内容
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                text = self._process_inline_markdown(text)
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, text)
            # 普通段落
            else:
                text = self._process_inline_markdown(line)
                p = doc.add_paragraph()
                self._add_formatted_text(p, text)

    def _process_inline_markdown(self, text: str) -> str:
        """处理行内Markdown标记（暂时保留标记，由_add_formatted_text处理）"""
        return text

    def _add_formatted_text(self, paragraph, text: str):
        """
        添加格式化文本到段落（处理加粗等格式）

        Args:
            paragraph: Word段落对象
            text: 包含Markdown标记的文本
        """
        import re

        # 正则匹配 **加粗文本**
        pattern = r'\*\*(.+?)\*\*'

        last_end = 0
        for match in re.finditer(pattern, text):
            # 添加普通文本
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            # 添加加粗文本
            bold_text = match.group(1)
            run = paragraph.add_run(bold_text)
            run.bold = True

            last_end = match.end()

        # 添加剩余文本
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_markdown_table(self, doc: Document, table_data: Dict):
        """
        将 Markdown 表格转换为 Word 表格

        Args:
            doc: Word 文档对象
            table_data: 表格数据
                - markdown: Markdown 格式的表格
                - caption: 表格标题
                - type: 表格类型
        """
        markdown = table_data.get('markdown', '')
        caption = table_data.get('caption', '')

        if not markdown or '|' not in markdown:
            self.logger.warning("无效的 Markdown 表格数据")
            return

        try:
            # 解析 Markdown 表格
            headers, rows = self._parse_markdown_table(markdown)

            if not headers or not rows:
                self.logger.warning("表格解析失败: 无表头或数据行")
                return

            # 添加表格标题（在表格上方）
            if caption:
                caption_para = doc.add_paragraph()
                run = caption_para.add_run(f"表：{caption}")
                run.bold = True
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 创建 Word 表格
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = 'Table Grid'
            table.autofit = True

            # 填充表头
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header.strip()
                # 设置表头样式（加粗 + 浅灰色背景）
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                # 添加背景色 (浅灰色 D9D9D9)
                self._set_cell_shading(cell, "D9D9D9")

            # 填充数据行
            for row_idx, row_data in enumerate(rows):
                table_row = table.rows[row_idx + 1]
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(table_row.cells):
                        table_row.cells[col_idx].text = cell_text.strip()

            # 添加空行
            doc.add_paragraph()

            self.logger.info(f"添加表格成功: {caption or '无标题'}")

        except Exception as e:
            self.logger.error(f"添加表格失败: {e}")

    def _parse_markdown_table(self, markdown: str) -> Tuple[List[str], List[List[str]]]:
        """
        解析 Markdown 表格为表头和数据行

        Args:
            markdown: Markdown 格式的表格

        Returns:
            (headers, rows): 表头列表和数据行列表
        """
        lines = [l.strip() for l in markdown.strip().split('\n') if l.strip()]

        if len(lines) < 2:
            return [], []

        # 第一行是表头
        headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]

        # 跳过分隔行（第二行，通常是 |---|---|---|）
        # 检测分隔行
        separator_idx = 1
        if len(lines) > 1 and all(c in '-| ' for c in lines[1]):
            separator_idx = 2

        # 其余是数据行
        rows = []
        for line in lines[separator_idx:]:
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if row:
                rows.append(row)

        return headers, rows

    def _set_cell_shading(self, cell, color_hex: str):
        """
        设置单元格背景色

        Args:
            cell: Word 表格单元格
            color_hex: 16进制颜色值（不含#）
        """
        try:
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            self.logger.warning(f"设置单元格背景色失败: {e}")

    def _add_mermaid_flowchart(self, doc: Document, flowchart_data: Dict):
        """
        渲染 Mermaid 流程图并嵌入 Word

        Args:
            doc: Word 文档对象
            flowchart_data: 流程图数据
                - mermaid_code: Mermaid 语法代码
                - caption: 图片标题
                - type: 流程图类型
        """
        mermaid_code = flowchart_data.get('mermaid_code', '')
        caption = flowchart_data.get('caption', '')

        if not mermaid_code:
            self.logger.warning("无 Mermaid 代码")
            return

        try:
            # 渲染为 PNG
            png_path = self.mermaid_renderer.render_to_png(mermaid_code)

            if png_path and os.path.exists(png_path):
                # 添加图片
                doc.add_picture(png_path, width=Inches(5.5))

                # 设置图片居中
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # 添加图片标题
                if caption:
                    caption_para = doc.add_paragraph()
                    run = caption_para.add_run(f"图：{caption}")
                    run.italic = True
                    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # 添加空行
                doc.add_paragraph()

                self.logger.info(f"添加流程图成功: {caption or '无标题'}")

            else:
                # 渲染失败时，添加占位文本
                p = doc.add_paragraph()
                run = p.add_run(f"[流程图: {caption or '未命名'}]")
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                self.logger.warning(f"流程图渲染失败，添加占位符: {caption}")

        except Exception as e:
            self.logger.error(f"添加流程图失败: {e}")
            # 添加错误占位
            p = doc.add_paragraph()
            run = p.add_run(f"[流程图加载失败: {caption or '未命名'}]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
