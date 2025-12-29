#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档结构解析器 - 用于 HITL 1 (人工章节选择)
功能：
- 解析 Word 文档的目录结构
- 识别章节层级（H1/H2/H3）
- 基于白名单自动推荐章节
- 提供章节预览文本
"""

import re
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, asdict
from docx import Document
from docx.oxml import CT_Tbl, CT_P
from difflib import SequenceMatcher

from common import get_module_logger
from common.utils import resolve_file_path
from .level_analyzer import LevelAnalyzer

logger = get_module_logger("structure_parser")


@dataclass
class ChapterNode:
    """章节节点数据类"""
    id: str                      # 唯一ID，如 "ch_1_2_3"
    level: int                   # 层级：1=H1, 2=H2, 3=H3
    title: str                   # 章节标题
    para_start_idx: int          # 起始段落索引
    para_end_idx: int            # 结束段落索引（可能为None）
    word_count: int              # 字数统计
    preview_text: str            # 预览文本（前5行）
    has_table: bool = False      # 章节是否包含表格
    content_tags: List[str] = None  # 内容标签（基于内容关键词检测）
    content_sample: str = None   # 内容样本（用于合同识别，1500-2000字）
    children: List['ChapterNode'] = None  # 子章节列表

    def __post_init__(self):
        if self.children is None:
            self.children = []
        if self.content_tags is None:
            self.content_tags = []

    def to_dict(self) -> Dict:
        """转换为字典（用于JSON序列化）"""
        data = asdict(self)
        data['children'] = [child.to_dict() for child in self.children]
        return data


class DocumentStructureParser:
    """文档结构解析器"""

    def __init__(self):
        """初始化解析器"""
        self.logger = get_module_logger("structure_parser")

        # ========================================
        # 新增：编号模式（用于识别章节锚点）（改进4：扩展编号模式）
        # ========================================
        self.NUMBERING_PATTERNS = [
            # 中文部分/章节编号
            r'^第[一二三四五六七八九十百\d]+部分\s*',
            r'^第[一二三四五六七八九十百\d]+章\s*',
            r'^第[一二三四五六七八九十百\d]+节\s*',
            # 数字编号
            r'^\d+\.\s*',           # 1.
            r'^\d+\.\d+\s*',        # 1.1
            r'^\d+\.\d+\.\d+\s*',   # 1.1.1
            r'^\d+\.\d+\.\d+\.\d+\s*',  # 1.1.1.1（四级编号，改进4新增）
            # 中文序号
            r'^[一二三四五六七八九十]+、\s*',
            r'^（[一二三四五六七八九十]+）\s*',
            r'^\([一二三四五六七八九十]+\)\s*',
            # 字母编号
            r'^[A-Z]\.\s*',
            r'^[a-z]\.\s*',
            r'^\([A-Za-z]\)\s*',
            # 罗马数字
            r'^[IVX]+\.\s*',
            r'^[ivx]+\.\s*',
            # 附件/附表编号（改进4新增）
            r'^附件[一二三四五六七八九十\d]+[:：]\s*',  # 附件1: 或 附件一：
            r'^附表[一二三四五六七八九十\d]+[:：]\s*',  # 附表1: 或 附表一：
            r'^附录[一二三四五六七八九十\d]+[:：]\s*',  # 附录1: 或 附录一：
        ]


        # 标题样式名称映射（中英文）
        self.HEADING_STYLES = {
            1: ['Heading 1', '标题 1', 'heading 1', '1级标题'],
            2: ['Heading 2', '标题 2', 'heading 2', '2级标题'],
            3: ['Heading 3', '标题 3', 'heading 3', '3级标题'],
        }

    def parse_document_structure(
        self,
        doc_path: str,
        methods: Optional[List[str]] = None,
        fallback: bool = True
    ) -> Dict:
        """
        解析文档结构 - 总调用器

        支持指定解析方法或使用默认智能策略

        Args:
            doc_path: Word文档路径
            methods: 可选，指定要使用的解析方法列表，按优先级顺序尝试
                    可选值: ['toc_exact', 'semantic_anchors', 'style', 'hybrid',
                            'azure', 'outline_level', 'gemini']
                    默认None表示使用智能策略（根据文档特征自动选择）
            fallback: 可选，是否启用回退机制（当前方法失败时尝试下一个）
                     默认True

        Returns:
            {
                "success": True/False,
                "chapters": [ChapterNode.to_dict(), ...],
                "statistics": {
                    "total_chapters": 10,
                    "total_words": 15000
                },
                "method": "使用的解析方法名称",
                "error": "错误信息（如果失败）"
            }
        """
        # 方法映射表
        method_map = {
            'toc_exact': self.parse_by_toc_exact,
            'azure': self.parse_by_azure,
            'outline_level': self.parse_by_outline_level,
            'gemini': self.parse_by_gemini
        }

        try:
            self.logger.info(f"开始解析文档结构: {doc_path}")

            # ⭐️ 文件格式检测：不支持旧版.doc格式
            if doc_path.lower().endswith('.doc'):
                error_message = (
                    "暂不支持旧版 .doc 格式文档的章节解析。\n\n"
                    "请按以下步骤操作：\n"
                    "1. 使用 Microsoft Word 或 WPS Office 打开该文件\n"
                    '2. 点击"文件" → "另存为"\n'
                    '3. 在"保存类型"中选择 "Word 文档 (*.docx)"\n'
                    "4. 保存后重新上传 .docx 文件\n\n"
                    "提示：.docx 格式兼容性更好，推荐使用。"
                )
                self.logger.error(f".doc 文件不支持: {doc_path}")
                raise ValueError(error_message)

            # 场景1: 用户指定了具体方法
            if methods is not None:
                self.logger.info(f"使用指定方法: {methods}, fallback={fallback}")

                for method_name in methods:
                    if method_name not in method_map:
                        self.logger.warning(f"未知方法: {method_name}，跳过")
                        continue

                    self.logger.info(f"尝试方法: {method_name}")
                    result = method_map[method_name](doc_path)

                    # 判断成功标准
                    if result.get('success'):
                        # 检查是否识别到足够的章节
                        chapters = result.get('chapters', [])
                        if len(chapters) >= 1:  # 至少识别到1个章节
                            self.logger.info(f"方法 {method_name} 成功，识别到 {len(chapters)} 个章节")
                            return result
                        else:
                            self.logger.warning(f"方法 {method_name} 未识别到章节")

                    # 如果不启用fallback，直接返回结果（即使失败）
                    if not fallback:
                        self.logger.info(f"fallback=False，直接返回 {method_name} 的结果")
                        return result

                    # 否则继续尝试下一个方法
                    self.logger.warning(f"方法 {method_name} 失败或效果不佳，尝试下一个方法")

                # 所有方法都失败
                return {
                    "success": False,
                    "chapters": [],
                    "statistics": {},
                    "error": f"所有指定方法({methods})都失败",
                    "method": "none"
                }

            # 场景2: 默认智能策略（使用 parse_smart 智能识别 + 章节分类）
            self.logger.info("使用智能策略：parse_smart（精确/大纲 → 异常检测 → LLM回退 → 章节分类）")

            # 调用智能解析方法
            result = self.parse_smart(doc_path, classify_chapters=True)

            # parse_smart 返回的结果需要转换为 parse_document_structure 的标准格式
            return {
                "success": result.get('success', False),
                "chapters": result.get('chapters', []),
                "statistics": result.get('statistics', {}),
                "method": result.get('method_used', 'smart'),
                "fallback_from": result.get('fallback_from'),
                "fallback_reason": result.get('fallback_reason'),
                "key_sections": result.get('key_sections', {})
            }

        except Exception as e:
            self.logger.error(f"文档结构解析失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return {
                "success": False,
                "chapters": [],
                "statistics": {},
                "error": str(e),
                "method": "error"
            }

    # ============================================
    # 独立解析方法 - 可单独调用或组合使用
    # ============================================

    def parse_by_toc_exact(self, doc_path: str) -> Dict:
        """
        方法0: 精确匹配(基于目录)

        直接使用目录项精确匹配正文位置,速度快、准确率高

        Args:
            doc_path: Word文档路径

        Returns:
            {
                "success": True/False,
                "chapters": [...],
                "statistics": {...},
                "method": "toc_exact"
            }
        """
        try:
            doc_path_abs = resolve_file_path(doc_path)
            if not doc_path_abs:
                raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

            doc = Document(str(doc_path_abs))

            # 检测目录
            toc_idx = self._find_toc_section(doc)
            if toc_idx is None:
                return {
                    "success": False,
                    "error": "未检测到目录,无法使用精确匹配方法",
                    "chapters": [],
                    "statistics": {},
                    "method": "toc_exact"
                }

            toc_items, toc_end_idx = self._parse_toc_items(doc, toc_idx)
            if not toc_items:
                return {
                    "success": False,
                    "error": "目录解析失败",
                    "chapters": [],
                    "statistics": {},
                    "method": "toc_exact"
                }

            # 使用精确匹配
            chapters = self._locate_chapters_by_toc(doc, toc_items, toc_end_idx)
            chapter_tree = self._build_chapter_tree(chapters)
            stats = self._calculate_statistics(chapter_tree)

            return {
                "success": True,
                "chapters": [ch.to_dict() for ch in chapter_tree],
                "statistics": stats,
                "method": "toc_exact"
            }

        except Exception as e:
            self.logger.error(f"精确匹配解析失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "chapters": [],
                "statistics": {},
                "method": "toc_exact"
            }

    def parse_structure_quick(self, doc_path: str) -> Dict:
        """
        快速解析文档结构（阶段1）- 只提取目录项并使用LLM分析层级

        不做正文定位和字数统计，这些在 enrich_chapters 阶段完成。

        Args:
            doc_path: Word文档路径

        Returns:
            {
                "success": True/False,
                "chapters": [...],  # 章节树（无字数信息）
                "toc_end_idx": int,
                "method": "llm_quick"
            }
        """
        try:
            doc_path_abs = resolve_file_path(doc_path)
            if not doc_path_abs:
                raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

            doc = Document(str(doc_path_abs))

            # 1. 检测目录位置
            toc_idx = self._find_toc_section(doc)
            if toc_idx is None:
                # 无目录，尝试从文档开头识别章节标题
                self.logger.info("未检测到目录，尝试从文档中识别章节标题")
                toc_items = self._extract_chapter_titles_from_doc(doc)
                toc_end_idx = 0
            else:
                # 有目录，提取目录项
                toc_items, toc_end_idx = self._parse_toc_items(doc, toc_idx)

            if not toc_items:
                return {
                    "success": False,
                    "error": "未能提取到目录项",
                    "chapters": [],
                    "toc_end_idx": 0,
                    "method": "llm_quick"
                }

            self.logger.info(f"提取到 {len(toc_items)} 个目录项，开始LLM层级分析")

            # 2. 使用LLM分析层级
            from modules.tender_processing.level_analyzer import LevelAnalyzer
            analyzer = LevelAnalyzer()
            levels = analyzer.analyze_toc_hierarchy_with_llm(toc_items)

            # 3. 构建章节节点（不含正文定位信息）
            chapters = []
            for i, item in enumerate(toc_items):
                level = levels[i] if i < len(levels) else 1
                chapter = ChapterNode(
                    id=f"ch_{i}",
                    level=level,
                    title=item.get('title', ''),
                    para_start_idx=-1,  # 未定位
                    para_end_idx=-1,    # 未定位
                    word_count=0,       # 未统计
                    preview_text=""     # 未提取
                )
                chapters.append(chapter)

            # 4. 构建树形结构
            chapter_tree = self._build_chapter_tree(chapters)

            # 5. 章节分类
            classified_chapters, _ = self._classify_chapters(
                [ch.to_dict() for ch in chapter_tree]
            )

            return {
                "success": True,
                "chapters": classified_chapters,
                "toc_end_idx": toc_end_idx,
                "method": "llm_quick"
            }

        except Exception as e:
            self.logger.error(f"快速解析失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return {
                "success": False,
                "error": str(e),
                "chapters": [],
                "toc_end_idx": 0,
                "method": "llm_quick"
            }

    def enrich_chapters(self, doc_path: str, chapters: List[Dict], toc_end_idx: int = 0) -> Dict:
        """
        补充章节信息（阶段2）- 执行正文定位和字数统计

        根据阶段1识别的章节标题，在文档正文中定位每个章节的位置，
        并统计字数、提取预览文本。

        Args:
            doc_path: Word文档路径
            chapters: 阶段1返回的章节树（嵌套结构）
            toc_end_idx: 目录结束的段落索引

        Returns:
            {
                "success": True/False,
                "chapters": [...],  # 补充完整信息的章节树
                "statistics": {
                    "total_words": int,
                    "chapter_count": int,
                    "estimated_processing_cost": float
                }
            }
        """
        try:
            doc_path_abs = resolve_file_path(doc_path)
            if not doc_path_abs:
                raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

            doc = Document(str(doc_path_abs))
            self.logger.info(f"[enrich_chapters] 开始补充章节信息，共 {len(chapters)} 个根章节")

            # 1. 将嵌套的章节树展平为列表，方便处理
            flat_chapters = self._flatten_chapters(chapters)
            self.logger.info(f"[enrich_chapters] 展平后共 {len(flat_chapters)} 个章节")

            # 2. 定位每个章节在文档中的位置
            last_found_idx = toc_end_idx + 1
            located_chapters = []

            for order_idx, chapter in enumerate(flat_chapters):
                title = chapter.get('title', '')
                level = chapter.get('level', 1)
                chapter_id = chapter.get('id', '')

                # 在文档中查找标题位置
                para_idx = self._find_paragraph_by_title_simple(doc, title, last_found_idx)

                if para_idx is None:
                    self.logger.warning(f"⚠️  未找到章节: [{level}级] {title}")
                    # 仍然添加到列表，但标记为未定位
                    located_chapters.append({
                        'id': chapter_id,
                        'title': title,
                        'level': level,
                        'para_idx': -1,
                        'para_end_idx': -1,
                        'located': False,
                        'order_index': order_idx,  # 保持原始目录顺序
                        'original': chapter
                    })
                    continue

                located_chapters.append({
                    'id': chapter_id,
                    'title': title,
                    'level': level,
                    'para_idx': para_idx,
                    'para_end_idx': None,  # 稍后计算
                    'located': True,
                    'order_index': order_idx,  # 保持原始目录顺序
                    'original': chapter
                })

                last_found_idx = para_idx + 1
                self.logger.debug(f"  ✓ 找到 [{level}级] {title} (段落 {para_idx})")

            # 3. 计算每个章节的结束位置
            for i, chapter_info in enumerate(located_chapters):
                if not chapter_info['located']:
                    continue

                current_level = chapter_info['level']
                next_start = len(doc.paragraphs)

                # 找下一个同级或更高级的章节
                for j in range(i + 1, len(located_chapters)):
                    if located_chapters[j]['located'] and located_chapters[j]['level'] <= current_level:
                        next_start = located_chapters[j]['para_idx']
                        break

                chapter_info['para_end_idx'] = next_start - 1

            # 4. 计算字数和提取预览文本
            total_words = 0
            for chapter_info in located_chapters:
                if not chapter_info['located']:
                    chapter_info['word_count'] = 0
                    chapter_info['preview_text'] = "(未定位)"
                    chapter_info['has_table'] = False
                    continue

                para_idx = chapter_info['para_idx']
                para_end_idx = chapter_info['para_end_idx']

                # 提取章节内容
                content_text, preview_text, has_table = self._extract_chapter_content_with_tables(
                    doc, para_idx, para_end_idx
                )

                word_count = self._calculate_word_count(content_text)

                chapter_info['word_count'] = word_count
                chapter_info['preview_text'] = preview_text if preview_text else "(无内容)"
                chapter_info['has_table'] = has_table

                total_words += word_count

            # 5. 更新原始章节树的信息
            enriched_chapters = self._update_chapter_tree(chapters, located_chapters)

            # 6. 统计信息
            statistics = {
                "total_words": total_words,
                "chapter_count": len(flat_chapters),
                "located_count": sum(1 for c in located_chapters if c['located']),
                "estimated_processing_cost": round(total_words * 0.000002, 4)  # 假设每字0.000002元
            }

            self.logger.info(
                f"✅ [enrich_chapters] 完成: 定位 {statistics['located_count']}/{statistics['chapter_count']} 个章节, "
                f"总字数 {total_words}"
            )

            return {
                "success": True,
                "chapters": enriched_chapters,
                "statistics": statistics
            }

        except Exception as e:
            self.logger.error(f"补充章节信息失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return {
                "success": False,
                "error": str(e),
                "chapters": chapters,
                "statistics": {}
            }
    def _find_paragraph_by_title_simple(self, doc: Document, title: str, start_idx: int) -> int:
        """
        简化版标题定位：在文档中查找与标题匹配的段落

        Args:
            doc: Word文档对象
            title: 要查找的标题
            start_idx: 开始搜索的段落索引

        Returns:
            段落索引，未找到返回None
        """
        import re
        from difflib import SequenceMatcher

        # 标准化标题：去除多余空格、制表符
        def normalize(text: str) -> str:
            text = text.strip()
            text = re.sub(r'\s+', ' ', text)  # 多个空白字符合并为一个空格
            text = re.sub(r'\t+', ' ', text)  # 制表符转空格
            return text

        title_clean = normalize(title)
        if not title_clean:
            return None

        for i in range(start_idx, len(doc.paragraphs)):
            para_text = normalize(doc.paragraphs[i].text)
            if not para_text:
                continue

            # 完全匹配
            if para_text == title_clean:
                return i

            # 模糊匹配（相似度 > 0.85）
            ratio = SequenceMatcher(None, para_text, title_clean).ratio()
            if ratio > 0.85:
                return i

            # 包含匹配（段落以标题开头或标题以段落开头）
            if para_text.startswith(title_clean) or title_clean.startswith(para_text):
                if len(para_text) < 150:  # 避免匹配到长段落
                    return i

            # 去除编号后匹配（处理"第三章  用户需求书" vs "第三章 用户需求书"）
            title_no_space = title_clean.replace(' ', '')
            para_no_space = para_text.replace(' ', '')
            if para_no_space == title_no_space:
                return i

        return None

    def _update_chapter_tree(self, chapters: List[Dict], located_chapters: List[Dict]) -> List[Dict]:
        """
        更新章节树，将定位信息合并回嵌套结构

        Args:
            chapters: 原始嵌套章节树
            located_chapters: 展平后带定位信息的章节列表

        Returns:
            更新后的嵌套章节树
        """
        # 建立ID到定位信息的映射
        location_map = {c['id']: c for c in located_chapters}

        def update_recursive(chapter_list: List[Dict]) -> List[Dict]:
            result = []
            for chapter in chapter_list:
                chapter_id = chapter.get('id', '')
                updated = dict(chapter)

                if chapter_id in location_map:
                    loc_info = location_map[chapter_id]
                    updated['para_start_idx'] = loc_info.get('para_idx', -1)
                    updated['para_end_idx'] = loc_info.get('para_end_idx', -1)
                    updated['word_count'] = loc_info.get('word_count', 0)
                    updated['preview_text'] = loc_info.get('preview_text', '')
                    updated['has_table'] = loc_info.get('has_table', False)
                    updated['order_index'] = loc_info.get('order_index', 0)  # 保持原始目录顺序

                # 递归处理子章节
                if 'children' in updated and updated['children']:
                    updated['children'] = update_recursive(updated['children'])

                result.append(updated)
            return result

        return update_recursive(chapters)

    def _extract_chapter_titles_from_doc(self, doc: Document) -> List[Dict]:
        """
        从文档中提取可能的章节标题（用于无目录的情况）

        Args:
            doc: Word文档对象

        Returns:
            目录项列表格式：[{'title': '...', 'page_num': 0}, ...]
        """
        from modules.tender_processing.level_analyzer import LevelAnalyzer
        analyzer = LevelAnalyzer()

        titles = []
        for i, para in enumerate(doc.paragraphs[:500]):  # 只扫描前500段
            text = para.text.strip()
            if not text or len(text) > 100:
                continue

            # 检测是否有编号模式
            pattern_info = analyzer.extract_numbering_pattern(text)
            if pattern_info.get('type') != 'none':
                titles.append({
                    'title': text,
                    'page_num': 0,
                    'index': i
                })

        self.logger.info(f"从文档中提取到 {len(titles)} 个潜在章节标题")
        return titles

    def parse_by_outline_level(self, doc_path: str) -> Dict:
        """
        方法5: Word大纲级别识别

        使用Word文档大纲级别(Outline Level)识别章节

        Args:
            doc_path: Word文档路径

        Returns:
            {
                "success": True/False,
                "chapters": [...],
                "statistics": {...},
                "method": "outline_level"
            }
        """
        try:
            doc_path_abs = resolve_file_path(doc_path)
            if not doc_path_abs:
                raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

            doc = Document(str(doc_path_abs))

            # 1. 基于Word大纲级别识别章节
            chapters = self._parse_chapters_by_outline_level(doc)

            # 2. 🆕 使用智能层级分析器修正层级(与精确识别保持一致)
            if chapters:
                from modules.tender_processing.level_analyzer import LevelAnalyzer

                # 转换为类似目录项的格式
                toc_like_items = [
                    {'title': ch.title, 'level': ch.level}
                    for ch in chapters
                ]

                # 使用contextual方法智能分析层级
                analyzer = LevelAnalyzer()
                corrected_levels = analyzer.analyze_toc_hierarchy_contextual(toc_like_items)

                # 更新章节层级
                for i, level in enumerate(corrected_levels):
                    chapters[i].level = level

                # 统计日志
                from collections import Counter
                level_dist = Counter(corrected_levels)
                self.logger.info(f"✅ 智能层级分析完成,修正 {len(chapters)} 个章节: {dict(level_dist)}")

            # 3. 后续处理(定位内容、构建树、统计)
            chapters = self._locate_chapter_content(doc, chapters)
            chapter_tree = self._build_chapter_tree(chapters)
            stats = self._calculate_statistics(chapter_tree)

            return {
                "success": True,
                "chapters": [ch.to_dict() for ch in chapter_tree],
                "statistics": stats,
                "method": "outline_level"
            }

        except Exception as e:
            self.logger.error(f"大纲级别识别失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "chapters": [],
                "statistics": {},
                "method": "outline_level"
            }

    def parse_by_azure(self, doc_path: str) -> Dict:
        """
        方法4: Azure Form Recognizer解析

        使用Azure AI服务识别文档结构

        Args:
            doc_path: Word文档路径

        Returns:
            {
                "success": True/False,
                "chapters": [...],
                "statistics": {...},
                "method": "azure"
            }
        """
        try:
            # 检查Azure解析器是否可用
            try:
                from modules.tender_processing.azure_parser import AzureDocumentParser, is_azure_available

                if not is_azure_available():
                    return {
                        "success": False,
                        "error": "Azure Form Recognizer未配置或SDK未安装",
                        "chapters": [],
                        "statistics": {},
                        "method": "azure"
                    }

                doc_path_abs = resolve_file_path(doc_path)
                if not doc_path_abs:
                    raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

                # 使用Azure解析器
                azure_parser = AzureDocumentParser()
                result = azure_parser.parse_document_structure(str(doc_path_abs))
                return result

            except ImportError as e:
                return {
                    "success": False,
                    "error": f"Azure解析器不可用: {str(e)}",
                    "chapters": [],
                    "statistics": {},
                    "method": "azure"
                }

        except Exception as e:
            self.logger.error(f"Azure解析失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "chapters": [],
                "statistics": {},
                "method": "azure"
            }

    def parse_by_gemini(self, doc_path: str) -> Dict:
        """
        方法6: Gemini AI解析

        使用Google Gemini AI模型识别文档结构

        Args:
            doc_path: Word文档路径

        Returns:
            {
                "success": True/False,
                "chapters": [...],
                "statistics": {...},
                "method": "gemini"
            }
        """
        try:
            # 检查Gemini解析器是否可用
            try:
                from modules.tender_processing.parsers.gemini_parser import GeminiParser

                gemini_parser = GeminiParser()

                if not gemini_parser.is_available():
                    return {
                        "success": False,
                        "error": "Gemini API密钥未配置",
                        "chapters": [],
                        "statistics": {},
                        "method": "gemini"
                    }

                doc_path_abs = resolve_file_path(doc_path)
                if not doc_path_abs:
                    raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

                # 使用Gemini解析器
                result = gemini_parser.parse_structure(str(doc_path_abs))

                # 确保chapters是dict格式
                if result.get('success') and result.get('chapters'):
                    chapters = result['chapters']
                    if chapters and hasattr(chapters[0], 'to_dict'):
                        result['chapters'] = [ch.to_dict() if hasattr(ch, 'to_dict') else ch for ch in chapters]

                return result

            except ImportError as e:
                return {
                    "success": False,
                    "error": f"Gemini解析器不可用: {str(e)} (pip install google-generativeai)",
                    "chapters": [],
                    "statistics": {},
                    "method": "gemini"
                }

        except Exception as e:
            self.logger.error(f"Gemini解析失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "chapters": [],
                "statistics": {},
                "method": "gemini"
            }

    def parse_smart(self, doc_path: str, classify_chapters: bool = True) -> Dict:
        """
        智能解析：结构识别 + 类型分类

        流程：
        1. 检查文档是否有目录
           - 有目录 → 精确匹配(toc_exact)
           - 无目录 → Word大纲识别(docx_native)
        2. 检查结果是否异常（章节太少、编号跳跃等）
           - 正常 → 进行章节分类
           - 异常 → 回退到LLM层级分析
        3. 对章节进行类型分类（可选）

        Args:
            doc_path: Word文档路径
            classify_chapters: 是否对章节进行类型分类

        Returns:
            {
                "success": True/False,
                "chapters": [...],
                "statistics": {...},
                "method": "smart",
                "primary_method": "toc_exact|docx_native|llm_level",
                "fallback_from": None|"toc_exact"|"docx_native",
                "fallback_reason": None|str,
                "key_sections": {
                    "business_response": [...],
                    "technical_spec": [...],
                    "contract_content": [...]
                }
            }
        """
        import time
        start_time = time.time()

        try:
            doc_path_abs = resolve_file_path(doc_path)
            if not doc_path_abs:
                raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

            doc = Document(str(doc_path_abs))
            doc_paragraph_count = len(doc.paragraphs)

            # ========================================
            # 阶段1: 章节结构识别（智能回退）
            # ========================================
            fallback_from = None
            fallback_reason = None
            primary_method = None

            # 检测是否有目录
            toc_idx = self._find_toc_section(doc)
            has_toc = toc_idx is not None

            if has_toc:
                # 有目录：使用精确匹配
                self.logger.info("📌 智能解析: 检测到目录，使用精确匹配")
                result = self.parse_by_toc_exact(doc_path)
                primary_method = "toc_exact"
            else:
                # 无目录：使用大纲识别
                self.logger.info("📌 智能解析: 未检测到目录，使用大纲识别")
                result = self.parse_by_outline_level(doc_path)
                primary_method = "docx_native"

            # 检查结果是否异常
            if result.get('success') and result.get('chapters'):
                chapters = result['chapters']
                is_suspicious, reason = self._is_result_suspicious(chapters, doc_paragraph_count)

                if is_suspicious:
                    self.logger.warning(f"⚠️ {primary_method} 结果异常: {reason}，回退到LLM层级分析")
                    fallback_from = primary_method
                    fallback_reason = reason

                    # 回退到LLM层级分析
                    llm_result = self._parse_by_llm_level(doc_path)
                    if llm_result.get('success') and llm_result.get('chapters'):
                        result = llm_result
                        primary_method = "llm_level"
                    else:
                        self.logger.warning("LLM层级分析也失败，保留原结果")
            elif not result.get('success'):
                # 主方法失败，尝试备选方法
                self.logger.warning(f"⚠️ {primary_method} 失败，尝试备选方法")
                fallback_from = primary_method

                if has_toc:
                    # 精确匹配失败，尝试大纲识别
                    result = self.parse_by_outline_level(doc_path)
                    primary_method = "docx_native"
                else:
                    # 大纲识别失败，尝试LLM
                    result = self._parse_by_llm_level(doc_path)
                    primary_method = "llm_level"

                if not result.get('success'):
                    fallback_reason = "所有方法都失败"

            # ========================================
            # 阶段2: 章节类型分类
            # ========================================
            key_sections = {}
            if classify_chapters and result.get('success') and result.get('chapters'):
                try:
                    classified_chapters, key_sections = self._classify_chapters(result['chapters'])
                    result['chapters'] = classified_chapters
                except Exception as e:
                    self.logger.warning(f"章节分类失败: {e}")

            # 构建返回结果
            elapsed = time.time() - start_time
            return {
                "success": result.get('success', False),
                "chapters": result.get('chapters', []),
                "statistics": result.get('statistics', {}),
                "method": "smart",
                "primary_method": primary_method,
                "fallback_from": fallback_from,
                "fallback_reason": fallback_reason,
                "key_sections": key_sections,
                "performance": {
                    "elapsed": elapsed,
                    "elapsed_formatted": f"{elapsed:.2f}s"
                }
            }

        except Exception as e:
            self.logger.error(f"智能解析失败: {e}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e),
                "chapters": [],
                "statistics": {},
                "method": "smart"
            }

    def _is_result_suspicious(self, chapters: List[Dict], doc_paragraph_count: int = 0) -> Tuple[bool, Optional[str]]:
        """
        检查识别结果是否异常

        Args:
            chapters: 章节列表
            doc_paragraph_count: 文档段落总数

        Returns:
            (是否异常, 异常原因)
        """
        # 扁平化章节列表（包括子章节）
        def flatten(chs):
            result = []
            for ch in chs:
                result.append(ch)
                if ch.get('children'):
                    result.extend(flatten(ch['children']))
            return result

        all_chapters = flatten(chapters) if chapters else []
        level1_chapters = [c for c in all_chapters if c.get('level') == 1]
        level1_count = len(level1_chapters)

        # 规则1: 一级章节太少（招标文件通常有5-10个一级章节）
        if level1_count < 3:
            return True, f"一级章节太少（仅{level1_count}个）"

        # 规则2: 章节编号有跳跃（如只有第一章和第五章）
        has_gap, gap_info = self._has_chapter_number_gap(level1_chapters)
        if has_gap:
            return True, f"章节编号有跳跃: {gap_info}"

        # 规则3: 文档很长但章节很少（每章平均超过100个段落不太合理）
        if doc_paragraph_count > 0 and level1_count > 0:
            avg_paras = doc_paragraph_count / level1_count
            if avg_paras > 150:
                return True, f"每章平均段落数过多（{avg_paras:.0f}段/章）"

        return False, None

    def _has_chapter_number_gap(self, level1_chapters: List[Dict]) -> Tuple[bool, Optional[str]]:
        """
        检查章节编号是否有跳跃

        Args:
            level1_chapters: 一级章节列表

        Returns:
            (是否有跳跃, 跳跃信息)
        """
        # 中文数字转换
        chinese_nums = {
            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
            '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
            '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15
        }

        def chinese_to_number(s):
            if s.isdigit():
                return int(s)
            return chinese_nums.get(s, None)

        # 提取章节编号
        numbers = []
        for ch in level1_chapters:
            title = ch.get('title', '')
            # 匹配 "第一章"、"第1章"、"第一部分" 等格式
            match = re.search(r'第([一二三四五六七八九十\d]+)[章部分篇]', title)
            if match:
                num = chinese_to_number(match.group(1))
                if num:
                    numbers.append(num)

        # 检查连续性
        if len(numbers) >= 2:
            numbers.sort()
            for i in range(1, len(numbers)):
                if numbers[i] - numbers[i-1] > 1:
                    return True, f"第{numbers[i-1]}→第{numbers[i]}"

        return False, None

    def _parse_by_llm_level(self, doc_path: str) -> Dict:
        """
        使用LLM层级分析解析文档结构

        Args:
            doc_path: Word文档路径

        Returns:
            解析结果字典
        """
        try:
            from modules.tender_processing.level_analyzer import LevelAnalyzer

            doc_path_abs = resolve_file_path(doc_path)
            if not doc_path_abs:
                raise FileNotFoundError(f"无法解析文件路径: {doc_path}")

            doc = Document(str(doc_path_abs))
            analyzer = LevelAnalyzer()

            # 提取所有可能的章节标题
            potential_titles = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text or len(text) > 100:
                    continue

                # 检测是否有编号模式
                pattern_info = analyzer.extract_numbering_pattern(text)
                if pattern_info.get('type') != 'none':
                    potential_titles.append({
                        'index': i,
                        'title': text,
                        'pattern': pattern_info
                    })

            if not potential_titles:
                return {
                    "success": False,
                    "error": "未检测到章节标题",
                    "chapters": [],
                    "method": "llm_level"
                }

            # 使用LLM分析层级
            toc_items = [{'title': t['title'], 'index': t['index']} for t in potential_titles]
            analyzed = analyzer.analyze_toc_hierarchy(toc_items)

            # 构建章节节点
            chapters = []
            for item in analyzed:
                chapter = ChapterNode(
                    id=f"ch_{item.get('index', len(chapters))}",
                    level=item.get('level', 1),
                    title=item.get('title', ''),
                    para_start_idx=item.get('index', 0),
                    para_end_idx=item.get('index', 0),
                    word_count=0,
                    preview_text=""
                )
                chapters.append(chapter)

            # 构建树形结构
            chapter_tree = self._build_chapter_tree(chapters)

            return {
                "success": True,
                "chapters": [ch.to_dict() for ch in chapter_tree],
                "statistics": self._calculate_statistics(chapter_tree),
                "method": "llm_level"
            }

        except Exception as e:
            self.logger.error(f"LLM层级分析失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "chapters": [],
                "method": "llm_level"
            }

    def _classify_chapters(self, chapters: List[Dict]) -> Tuple[List[Dict], Dict]:
        """
        对章节进行类型分类

        章节类型：
        - invitation: 投标邀请书
        - bidder_notice: 投标人须知
        - evaluation: 评标办法/评分标准
        - contract_terms: 合同条款
        - contract_content: 合同正文（实际需填写的合同）
        - business_response: 商务应答模板
        - technical_spec: 技术规范书/技术需求
        - appendix: 附件/附录
        - other: 其他

        Args:
            chapters: 章节列表

        Returns:
            (带类型标注的章节列表, 重点区域汇总)
        """
        # 分类规则（基于关键词匹配）
        classification_rules = {
            'invitation': [
                r'投标邀请', r'招标公告', r'邀请书', r'询价公告', r'比选公告'
            ],
            'bidder_notice': [
                r'投标人须知', r'应答人须知', r'供应商须知', r'磋商须知', r'响应人须知'
            ],
            'evaluation': [
                r'评标办法', r'评分标准', r'评审办法', r'评审标准', r'评分细则', r'评标标准'
            ],
            'contract_terms': [
                r'合同条款', r'合同条件', r'协议条款', r'合同格式'
            ],
            'contract_content': [
                r'服务合同', r'采购合同', r'合同正文', r'合同书', r'合同范本',
                r'货物供应合同', r'技术服务合同', r'框架合同'
            ],
            'business_response': [
                r'投标文件格式', r'响应文件格式', r'商务部分', r'资格证明', r'报价要求',
                r'投标函', r'法定代表人', r'授权委托书', r'投标保证金',
                r'业绩证明', r'财务报表', r'资质证明'
            ],
            'technical_spec': [
                r'技术规范', r'技术要求', r'技术需求', r'服务要求', r'功能需求', r'技术参数',
                r'技术方案', r'服务规范', r'技术服务', r'技术（服务）',
                r'需求说明', r'产品规格'
            ],
            'appendix': [
                r'^附件', r'^附录', r'^附表', r'^附图'
            ]
        }

        def classify_single(title: str) -> str:
            """对单个标题进行分类"""
            title_clean = title.strip()
            for category, patterns in classification_rules.items():
                for pattern in patterns:
                    if re.search(pattern, title_clean, re.IGNORECASE):
                        return category
            return 'other'

        def classify_recursive(chs: List[Dict]) -> List[Dict]:
            """递归分类章节及其子章节"""
            result = []
            for ch in chs:
                ch_copy = ch.copy()
                ch_copy['chapter_type'] = classify_single(ch.get('title', ''))

                if ch.get('children'):
                    ch_copy['children'] = classify_recursive(ch['children'])

                result.append(ch_copy)
            return result

        # 执行分类
        classified = classify_recursive(chapters)

        # 汇总重点区域
        key_sections = {
            'business_response': [],
            'technical_spec': [],
            'contract_content': []
        }

        def collect_key_sections(chs: List[Dict]):
            for ch in chs:
                ch_type = ch.get('chapter_type')
                if ch_type in key_sections:
                    key_sections[ch_type].append(ch.get('title', ''))
                if ch.get('children'):
                    collect_key_sections(ch['children'])

        collect_key_sections(classified)

        return classified, key_sections
    def _is_valid_chapter_title(self, text: str) -> bool:
        """
        判断文本是否是合法的章节标题 (过滤列表项、说明性内容等)

        Args:
            text: 段落文本

        Returns:
            True 表示可能是章节标题，False 表示应该被过滤
        """
        if not text or len(text.strip()) == 0:
            return False

        # 1. 过滤纯编号（如 "1"、"1.1"）
        if re.match(r'^[\d\.]+$', text):
            return False

        # 2. 过滤纯列表项标记（如 "1."、"2."、"(1)"）
        if re.match(r'^[\d]+\.$', text):  # "1." "2."
            return False
        if re.match(r'^\([\d]+\)$', text):  # "(1)" "(2)"
            return False
        if re.match(r'^[一二三四五六七八九十]+、$', text):  # "一、" "二、"
            return False

        # 3. 过滤说明性内容的特征关键词
        note_keywords = ['注：', '注:', '说明：', '说明:', '备注：', '备注:']
        if any(text.startswith(kw) for kw in note_keywords):
            # "注："本身可能不是章节，但如果后面有内容可能是
            # 如果只有"注："两个字，肯定不是章节标题
            if len(text) <= 4:
                return False

        # 4. 过滤超长内容（章节标题一般不超过50个字符）
        # 但允许包含长条款编号的章节（如 "7.3 双方均应当..."）
        if len(text) > 100:  # 超过100字符肯定不是标题
            return False

        # 5. 过滤纯列表项格式（开头是编号+简短内容）
        # 例如: "1.以上响应文件的构成为必须包含的内容..."
        if re.match(r'^[\d]+\.[\u4e00-\u9fa5]{2,}', text):
            # 如果匹配 "数字.中文"，且没有明确的章节特征词，则可能是列表项
            # 检查是否包含章节特征词
            chapter_markers = ['章', '节', '部分', '第', '条款', '要求', '标准', '方法', '原则']
            if not any(marker in text for marker in chapter_markers):
                # 进一步检查：如果内容很长（>20字符），可能是说明性列表项
                if len(text) > 30:
                    return False

        # 6. 特殊过滤：排除纯条款编号（如 "1.1.1"、"2.3.4.5"）
        if re.match(r'^[\d]+\.[\d]+\.[\d]+(\.[\d]+)*$', text):
            return False

        # 7. 特殊过滤：排除以条款编号开头且过长的内容
        # 例如: "7.3 双方均应当严格按照相关法律法规..."（这不是章节标题）
        if re.match(r'^[\d]+\.[\d]+\s+', text) and len(text) > 50:
            return False

        return True

    def _parse_chapters_by_outline_level(self, doc: Document) -> List[ChapterNode]:
        """
        方法五：基于Word大纲级别（outlineLevel）识别章节

        使用微软官方的大纲级别API，支持0-8级标题识别，
        包含3层精细过滤规则和Heading样式备用检测。

        Args:
            doc: python-docx Document 对象

        Returns:
            章节列表（扁平结构，未构建树）
        """
        import re

        chapters = []
        chapter_counter = 0

        self.logger.info("使用方法五：大纲级别识别（增强版）")

        # 直接从Word文档提取标题 - 使用微软官方的大纲级别API
        for para_idx, paragraph in enumerate(doc.paragraphs):
            is_heading = False
            level = 0
            detection_method = ""

            # ⭐ 优先级1: 检查大纲级别 (Outline Level) - 微软官方语义标记
            # 这是Word导航窗格和大纲视图使用的结构，准确度最高
            try:
                pPr = paragraph._element.pPr
                if pPr is not None:
                    outlineLvl = pPr.outlineLvl
                    if outlineLvl is not None:
                        outline_level_val = int(outlineLvl.val)
                        # Word大纲级别: 0-8表示标题(0=一级), 9表示正文
                        if outline_level_val <= 8:
                            # 🔧 添加过滤规则，排除噪音内容
                            text = paragraph.text.strip()
                            should_skip = False

                            # 过滤1: 跳过文档前30段的封面/元数据（Level 0）
                            if para_idx < 30 and outline_level_val == 0:
                                metadata_keywords = ['项目编号', '招标人', '代理机构', '联系人', '联系方式',
                                                    '地址', '电话', '传真', '邮编', '网址', 'http']
                                if any(kw in text for kw in metadata_keywords):
                                    should_skip = True
                                    self.logger.debug(f"过滤封面: 段落{para_idx} '{text[:30]}'")

                            # 过滤2: 跳过Level 3-4的长条款内容
                            if not should_skip and outline_level_val >= 3:
                                # 形如 "1.1 这是一个很长的说明文字..." 的是条款，不是标题
                                if re.match(r'^\d+\.\d+\s+.{15,}', text):
                                    should_skip = True
                                    self.logger.debug(f"过滤条款: 段落{para_idx} '{text[:30]}'")

                            # 过滤3: 标题长度限制（超过50字的通常不是标题）
                            if not should_skip and len(text) > 50:
                                # 除非有明确的章节编号
                                if not re.match(r'^第[一二三四五六七八九十\d]+[章部分]', text):
                                    should_skip = True
                                    self.logger.debug(f"过滤长文本: 段落{para_idx} '{text[:30]}'")

                            if not should_skip:
                                is_heading = True
                                level = outline_level_val + 1  # 转换: 0→1级, 1→2级, ...
                                detection_method = f"大纲级别{outline_level_val}"
            except (AttributeError, TypeError, ValueError):
                pass  # 没有大纲级别，继续其他方法

            # 优先级2: 检查标准Heading样式 (备用方案)
            if not is_heading:
                style_name = paragraph.style.name if paragraph.style else ""

                # 只接受标准的Heading样式（精确匹配，避免误识别）
                if style_name.startswith('Heading '):  # 'Heading 1', 'Heading 2'
                    match = re.search(r'Heading (\d+)', style_name)
                    if match:
                        is_heading = True
                        level = int(match.group(1))
                        detection_method = f"样式{style_name}"
                elif style_name.startswith('标题 '):  # '标题 1', '标题 2'
                    match = re.search(r'标题 (\d+)', style_name)
                    if match:
                        is_heading = True
                        level = int(match.group(1))
                        detection_method = f"样式{style_name}"

            if is_heading and paragraph.text.strip():
                title = paragraph.text.strip()

                chapter = ChapterNode(
                    id=f"docx_{chapter_counter}",
                    level=level if level > 0 else 1,
                    title=title,
                    para_start_idx=para_idx,
                    para_end_idx=None,  # 稍后计算
                    word_count=0,       # 稍后计算
                    preview_text="",    # 稍后提取
                    content_tags=['docx_native', detection_method]
                )

                chapters.append(chapter)
                chapter_counter += 1

                self.logger.debug(
                    f"识别标题: 段落{para_idx} [{detection_method}] '{title[:50]}'"
                )

        self.logger.info(f"✅ 方法五识别完成：找到 {len(chapters)} 个标题")
        return chapters

    def _get_heading_level(self, paragraph) -> int:
        """
        获取段落的标题层级 (优化6: 添加前置过滤和大纲级别检查)

        Args:
            paragraph: python-docx Paragraph 对象

        Returns:
            0: 不是标题
            1-3: 标题层级
        """
        text = paragraph.text.strip()

        # ⚠️ 方法0：前置过滤 - 排除明显不是章节的内容
        if not self._is_valid_chapter_title(text):
            return 0

        # ⭐ 方法1：大纲级别判断（最可靠，优先级最高）
        try:
            pPr = paragraph._element.pPr
            if pPr is not None:
                outlineLvl = pPr.outlineLvl
                if outlineLvl is not None:
                    level = int(outlineLvl.val)
                    if level <= 2:  # 0=一级, 1=二级, 2=三级
                        self.logger.debug(f"✓ 大纲级别识别: Level {level} → '{text[:30]}'")
                        return level + 1  # 转换为1-3
        except (AttributeError, TypeError):
            pass  # 没有大纲级别，继续其他方法

        # 方法2：通过样式名判断 (优先，最可靠)
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name
            for level, style_names in self.HEADING_STYLES.items():
                if any(sn.lower() in style_name.lower() for sn in style_names):
                    return level

        # 方法3：通过 XML 样式属性判断（更准确）
        try:
            pPr = paragraph._element.pPr
            if pPr is not None:
                pStyle = pPr.pStyle
                if pStyle is not None:
                    style_val = pStyle.val
                    if 'heading1' in style_val.lower() or style_val.lower() == '1':
                        return 1
                    elif 'heading2' in style_val.lower() or style_val.lower() == '2':
                        return 2
                    elif 'heading3' in style_val.lower() or style_val.lower() == '3':
                        return 3
        except (AttributeError, TypeError):
            pass  # XML属性不可用时使用其他方法

        # 方法4：通过文本格式启发式判断 (优化: 收集所有run的信息)
        if paragraph.runs:
            # 收集所有run的字体信息
            sizes = []
            bold_count = 0

            for run in paragraph.runs:
                if run.font.size:
                    sizes.append(run.font.size.pt)
                if run.bold:
                    bold_count += 1

            # 至少一半的run是加粗，才认为是标题
            if sizes and bold_count >= len(paragraph.runs) / 2:
                avg_size = sum(sizes) / len(sizes)

                # 调整阈值: 更灵活的判断 (降低阈值以提高识别率)
                if avg_size >= 16:  # 16pt+ → Level 1 (原18pt)
                    return 1
                elif avg_size >= 13:  # 13-15pt → Level 2 (原15pt)
                    return 2
                elif avg_size >= 10:  # 10-12pt → Level 3 (原12pt)
                    return 3

        # 方法5：通过编号模式判断 (增强fallback机制)
        # 一级标题模式
        if re.match(r'^第[一二三四五六七八九十\d]+[章部分]', text):
            return 1
        if re.match(r'^\d+\.\s+\S', text) and not re.match(r'^\d+\.\d+', text):  # 1. xxx (不包含1.1)
            return 1

        # 二级标题模式 (如果文本较短且有编号)
        if re.match(r'^\d+\.\d+\s+\S', text) and not re.match(r'^\d+\.\d+\.\d+', text):
            if len(text.strip()) <= 100:  # 长度限制
                return 2

        # 三级标题模式
        if re.match(r'^\d+\.\d+\.\d+\s+\S', text):
            if len(text.strip()) <= 100:
                return 3

        return 0

    def _clean_title_v2(self, title: str, aggressive=False) -> str:
        """
        分阶段清理标题文本 (优化3: 温和/激进两种模式)

        Args:
            title: 原始标题
            aggressive: 是否使用激进清理模式

        Returns:
            清理后的标题
        """
        if not aggressive:
            # 温和清理: 只删除明显的编号和空格
            cleaned = re.sub(r'^\d+\.\s*', '', title)  # 删除 "1. "
            cleaned = re.sub(r'^\d+\.\d+\s*', '', cleaned)  # 删除 "1.1 "
            cleaned = re.sub(r'^\d+\.\d+\.\d+\s*', '', cleaned)  # 删除 "1.1.1 "
            cleaned = re.sub(r'^第[一二三四五六七八九十\d]+[章节部分]\s*', '', cleaned)  # 删除 "第一章 "
            cleaned = re.sub(r'\s+', '', cleaned)  # 删除空格
            return cleaned
        else:
            # 激进清理: 删除所有编号、符号和空格
            cleaned = title
            cleaned = re.sub(r'^[一二三四五六七八九十]+、\s*', '', cleaned)
            cleaned = re.sub(r'^\([一二三四五六七八九十\d]+\)\s*', '', cleaned)
            cleaned = re.sub(r'^\d+[-\.]\d*\s*', '', cleaned)
            cleaned = re.sub(r'^[A-Za-z]+\.\s*', '', cleaned)
            cleaned = re.sub(r'^第[一二三四五六七八九十\d]+[章节部分]\s*', '', cleaned)
            cleaned = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', cleaned)  # 只保留中英文数字
            return cleaned

    def fuzzy_match_title_v2(self, title1: str, title2: str, threshold=0.7) -> float:
        """
        模糊匹配标题，支持多级清理尝试 (优化3: 分阶段匹配)

        Args:
            title1: 标题1
            title2: 标题2
            threshold: 相似度阈值

        Returns:
            相似度得分 (0.0-1.0)
        """
        # Level 1: 原始比较
        if title1 == title2:
            return 1.0

        # Level 2: 温和清理后比较
        clean1 = self._clean_title_v2(title1, aggressive=False)
        clean2 = self._clean_title_v2(title2, aggressive=False)

        if clean1 == clean2:
            return 0.95

        if clean1 in clean2 or clean2 in clean1:
            return 0.90

        # Level 3: 激进清理后比较
        aggr1 = self._clean_title_v2(title1, aggressive=True)
        aggr2 = self._clean_title_v2(title2, aggressive=True)

        if aggr1 == aggr2:
            return 0.85

        if aggr1 in aggr2 or aggr2 in aggr1:
            shorter = aggr1 if len(aggr1) <= len(aggr2) else aggr2
            longer = aggr2 if len(aggr1) <= len(aggr2) else aggr1
            return len(shorter) / len(longer) * 0.80  # 包含匹配，根据长度比例打分

        # Level 4: SequenceMatcher相似度
        similarity = SequenceMatcher(None, aggr1, aggr2).ratio()

        return similarity


    def _calculate_contract_density(self, text: str) -> float:
        """
        计算文本的合同密度（合同特征词出现频率）

        Args:
            text: 待分析的文本内容

        Returns:
            合同密度（0-1之间的浮点数）
        """
        if not text or len(text) < 100:  # 文本太短，无法判断
            return 0.0

        # 定义合同特征词及其权重
        contract_keywords = {
            # 强合同特征（权重 x3）
            '甲方': 3, '乙方': 3, '丙方': 3,
            # 中等合同特征（权重 x2）
            '违约金': 2, '违约责任': 2, '履约保证金': 2, '本合同': 2,
            '合同生效': 2, '合同终止': 2, '合同解除': 2,
            # 弱合同特征（权重 x1）
            '付款': 1, '验收': 1, '保密': 1, '争议': 1,
            '仲裁': 1, '管辖': 1, '双方': 1, '签订': 1
        }

        # 计算加权出现次数
        total_weight = 0
        for keyword, weight in contract_keywords.items():
            count = text.count(keyword)
            total_weight += count * weight

        # 计算密度（每1000字符的加权关键词出现次数）
        text_length = len(text)
        density = (total_weight / text_length) * 1000

        # 归一化到0-1之间（假设密度>50就是100%的合同）
        normalized_density = min(density / 50.0, 1.0)

        return normalized_density

    def _is_contract_chapter(self, title: str, content_sample: str = None) -> tuple:
        """
        判断章节是否为合同章节（增强版：结合标题、内容、排除规则）

        Args:
            title: 章节标题
            content_sample: 章节内容样本（可选）

        Returns:
            (is_contract, density, reason): 是否为合同章节、合同密度、判断理由
        """
        # 🆕 0. 排除规则优先（避免误判）
        exclude_keywords = [
            "投标人须知", "供应商须知", "投标邀请", "投标须知", "招标须知",
            "附件", "投标文件格式", "响应文件格式", "投标文件组成", "谈判响应文件格式",
            "用户需求", "需求书", "技术要求", "技术规格", "技术需求书", "技术需求",
            "评分", "评审", "评标", "开标", "报价", "投标报价",
            "投标邀请函", "采购邀请", "谈判邀请"
        ]

        for exclude_kw in exclude_keywords:
            if exclude_kw in title:
                self.logger.debug(f"  ⛔ 排除章节: '{title}' - 匹配排除关键词'{exclude_kw}'")
                return (False, 0.0, f"排除规则: '{exclude_kw}'")

        # 1. 基于标题的强合同标识（保留原有逻辑）
        strong_contract_keywords = [
            "合同条款", "合同文本", "合同范本", "合同格式", "合同协议",
            "通用条款", "专用条款", "合同主要条款", "合同草稿", "拟签合同",
            "服务合同", "采购合同", "买卖合同", "销售合同", "施工合同",
            "分包合同", "劳务合同", "租赁合同", "委托合同", "代理合同",
        ]

        for keyword in strong_contract_keywords:
            if keyword in title:
                return (True, 1.0, f"标题强匹配: '{keyword}'")

        # 2. 基于内容的合同密度检测（🆕 提高阈值，降低误判）
        if content_sample:
            density = self._calculate_contract_density(content_sample)

            # 🆕 高密度阈值：从5%提高到10%
            high_density_threshold = 0.10

            if density > high_density_threshold:
                # 高密度需要标题验证，避免误判"引用合同内容"的章节
                title_has_contract_kw = any(kw in title for kw in ["合同", "协议", "条款", "甲方", "乙方"])

                if title_has_contract_kw:
                    return (True, density, f"高密度+标题验证: {density:.1%}")
                else:
                    # 密度高但标题不像合同，可能是引用合同内容
                    self.logger.debug(f"  ⚠️  高密度({density:.1%})但标题不匹配: '{title}'")
                    return (False, density, f"高密度但标题不匹配: {density:.1%}")

        # 3. 弱合同标识（标题模糊但可能是合同）- 保留原有逻辑
        weak_contract_patterns = [
            r'(第[一二三四五六七八九十\d]+部分|附件\d*)[^\u4e00-\u9fa5]*(协议|条款|权利|义务)',
            r'双方.*权利.*义务',
            r'甲.*乙.*方',
        ]

        import re
        for pattern in weak_contract_patterns:
            if re.search(pattern, title):
                # 如果有内容样本，进一步验证
                if content_sample:
                    density = self._calculate_contract_density(content_sample)
                    # 🆕 提高弱匹配的密度阈值：从3%提高到5%
                    if density > 0.05:
                        return (True, density, f"标题弱匹配+内容验证: {density:.1%}")

        return (False, 0.0, "非合同章节")

    def _extract_content_sample(self, doc: Document, para_start_idx: int, para_end_idx: int, sample_size: int = 2000) -> str:
        """
        提取章节内容样本（用于合同识别）

        Args:
            doc: Word文档对象
            para_start_idx: 起始段落索引
            para_end_idx: 结束段落索引
            sample_size: 样本大小（字符数）

        Returns:
            内容样本字符串
        """
        if para_end_idx is None or para_end_idx <= para_start_idx:
            return ""

        # 提取章节内容（跳过标题本身）
        content_paras = doc.paragraphs[para_start_idx + 1 : para_end_idx + 1]

        # 合并文本
        sample_text = ""
        for para in content_paras:
            text = para.text.strip()
            if text:
                sample_text += text + "\n"
                if len(sample_text) >= sample_size:
                    break

        # 截取指定长度
        return sample_text[:sample_size]
    def _detect_contract_cluster_in_chapter(self, doc: Document, start_idx: int, end_idx: int) -> Optional[Dict]:
        """
        检测章节内的合同段落聚集区（精确定位起始位置）

        策略：
        1. 用50段滑动窗口扫描整个章节
        2. 找到第一个合同密度>20%的区域
        3. 从该区域向前精确定位聚集区起点（找到第一个包含合同特征的段落）

        Args:
            doc: Word文档对象
            start_idx: 章节起始段落索引
            end_idx: 章节结束段落索引

        Returns:
            如果发现聚集区：{'start': int, 'end': int, 'density': float}
            否则返回 None
        """
        # 🆕 参数验证：防止 end_idx 为 None 或无效值
        if end_idx is None or start_idx is None:
            self.logger.debug(f"  ⚠️  参数无效: start_idx={start_idx}, end_idx={end_idx}，跳过合同聚集区检测")
            return None

        if end_idx <= start_idx:
            self.logger.debug(f"  ⚠️  章节范围无效: start_idx={start_idx}, end_idx={end_idx}，跳过合同聚集区检测")
            return None

        if end_idx - start_idx < 50:
            # 章节太短，不需要检测
            return None

        window_size = 50  # 窗口大小：50个段落
        density_threshold = 0.2  # 密度阈值：20%
        step_size = 10  # 滑动步长：每次移动10段

        # 滑动窗口扫描
        for i in range(start_idx, end_idx - window_size, step_size):
            window_end = min(i + window_size, end_idx)

            # 提取窗口内文本
            window_text = ""
            for j in range(i, window_end):
                if j < len(doc.paragraphs):
                    para_text = doc.paragraphs[j].text.strip()
                    if para_text:
                        window_text += para_text + "\n"

            # 计算合同密度
            density = self._calculate_contract_density(window_text)

            if density > density_threshold:
                # 找到高密度区域，向前精确定位起点
                cluster_start = i

                # 向前查找第一个包含强合同特征的段落
                strong_contract_keywords = ['甲方', '乙方', '本合同', '合同的组成', '合同组成']

                for j in range(i, start_idx - 1, -1):  # 向前查找
                    if j < len(doc.paragraphs):
                        para_text = doc.paragraphs[j].text.strip()
                        if any(kw in para_text for kw in strong_contract_keywords):
                            cluster_start = j
                        else:
                            # 找到不含合同特征的段落，停止
                            break

                # 确定聚集区结束位置（向后扫描，找到密度降低的位置）
                cluster_end = end_idx
                for j in range(window_end, end_idx, 10):
                    check_end = min(j + 50, end_idx)
                    check_text = "\n".join(
                        doc.paragraphs[k].text.strip()
                        for k in range(j, check_end)
                        if k < len(doc.paragraphs) and doc.paragraphs[k].text.strip()
                    )
                    check_density = self._calculate_contract_density(check_text)

                    if check_density < density_threshold:
                        cluster_end = j - 1
                        break

                # 🆕 计算聚集区占章节的比例
                chapter_length = end_idx - start_idx
                cluster_length = cluster_end - cluster_start
                cluster_ratio = cluster_length / chapter_length if chapter_length > 0 else 0

                # 🆕 占比过高（>80%）：可能整章都是合同，不拆分
                if cluster_ratio > 0.8:
                    self.logger.info(
                        f"  ⚠️  合同聚集区占比过高({cluster_ratio:.1%})，不拆分章节 "
                        f"(段落{cluster_start}-{cluster_end}, 可能整章都是合同或误判)"
                    )
                    return None

                # 🆕 占比过低（<20%）：可能是误判（只是引用了部分合同内容）
                if cluster_ratio < 0.2:
                    self.logger.info(
                        f"  ⚠️  合同聚集区占比过低({cluster_ratio:.1%})，可能是误判 "
                        f"(段落{cluster_start}-{cluster_end}, 只是引用了部分合同内容)"
                    )
                    return None

                # 占比适中（20%-80%），执行拆分
                self.logger.info(
                    f"检测到合同聚集区: 段落{cluster_start}-{cluster_end} "
                    f"(章节范围:{start_idx}-{end_idx}, 合同密度:{density:.1%}, 占比:{cluster_ratio:.1%})"
                )

                return {
                    'start': cluster_start,
                    'end': cluster_end,
                    'density': density,
                    'ratio': cluster_ratio
                }

        return None
    def _find_toc_section(self, doc: Document) -> Optional[int]:
        """
        查找文档中的目录部分 (优化版: 扩展关键词 + SDT支持)

        Args:
            doc: Word文档对象

        Returns:
            目录起始段落索引，如果未找到则返回None
        """
        # 优化2: 扩展目录关键词列表
        TOC_KEYWORDS = [
            # 中文
            "目录", "目  录", "索引", "章节目录", "内容目录",
            # 英文
            "contents", "table of contents", "catalogue", "index"
        ]

        # 第零轮: 检测SDT容器中的TOC域（Word自动目录）
        try:
            body = doc.element.body
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # 查找所有SDT元素
            sdt_elements = body.findall('.//w:sdt', namespaces=ns)

            for sdt_idx, sdt in enumerate(sdt_elements):
                # 方法1: 检查是否是TOC类型的SDT（标准Word自动目录）
                docpart = sdt.find('.//w:docPartObj/w:docPartGallery', namespaces=ns)
                is_toc_sdt = False

                if docpart is not None:
                    gallery_val = docpart.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if gallery_val == 'Table of Contents':
                        is_toc_sdt = True
                        self.logger.info(f"检测到标准TOC类型的SDT (docPartGallery)")

                # 方法2: 检查SDT中是否包含"目录"关键词（手工创建的目录SDT）
                if not is_toc_sdt:
                    sdt_paras = sdt.findall('.//w:p', namespaces=ns)
                    if sdt_paras:
                        # 获取第一个段落的文本
                        first_para_texts = sdt_paras[0].findall('.//w:t', namespaces=ns)
                        first_para_text = ''.join([t.text for t in first_para_texts if t.text])

                        # 检查是否包含目录关键词
                        for keyword in TOC_KEYWORDS:
                            if keyword in first_para_text:
                                is_toc_sdt = True
                                self.logger.info(f"检测到包含'{keyword}'的SDT，可能是手工目录")
                                break

                # 如果确认是目录SDT，返回对应的段落索引
                if is_toc_sdt:
                    # 方法1: 尝试在doc.paragraphs中查找"目录"标题
                    # （适用于目录标题同时出现在doc.paragraphs中的情况）
                    toc_keywords_normalized = ["目录", "contents", "tableofcontents", "索引", "章节目录", "内容目录"]
                    for idx, para in enumerate(doc.paragraphs[:50]):
                        text = para.text.strip().replace(" ", "").replace("\u3000", "").lower()
                        if text in toc_keywords_normalized:
                            self.logger.info(f"检测到目录SDT，目录标题在paragraphs[{idx}]")
                            return idx

                    # 方法2: 如果目录标题完全在SDT内部（不在doc.paragraphs中），
                    # 通过计算SDT在body中的位置来推算段落索引
                    sdt_paras = sdt.findall('.//w:p', namespaces=ns)
                    if sdt_paras:
                        para_count = 0
                        for child in body:
                            child_tag = child.tag.split('}')[-1]
                            if child == sdt:
                                # 找到当前SDT，此时para_count就是它前面的段落数
                                self.logger.info(f"检测到目录SDT（目录在SDT内），位于body的第{para_count}个段落位置")
                                self.logger.info(f"目录位于SDT中，包含{len(sdt_paras)}个段落")
                                return para_count
                            elif child_tag == 'p':
                                para_count += 1
                            elif child_tag == 'sdt':
                                # 之前的SDT中的段落也要计数
                                prev_sdt_paras = child.findall('.//w:p', namespaces=ns)
                                para_count += len(prev_sdt_paras)

                    self.logger.debug("SDT目录检测失败，回退到关键词检测")
        except Exception as e:
            self.logger.debug(f"SDT检测失败（正常情况）: {e}")
            pass

        # 第一轮: 检测显式目录标题
        for i, para in enumerate(doc.paragraphs[:50]):  # 只检查前50段
            text = para.text.strip()

            # 跳过空段落
            if not text:
                continue

            # 检测目录标题 (使用扩展关键词列表)
            text_lower = text.lower()
            for keyword in TOC_KEYWORDS:
                if text_lower == keyword.lower() or text.replace(" ", "") == keyword.replace(" ", ""):
                    self.logger.info(f"检测到目录标题 (关键词: '{keyword}')，位于段落 {i}: {text}")
                    return i

            # 检测TOC域（Word自动生成的目录）
            # 通过检查段落的XML来识别
            try:
                xml_str = para._element.xml.decode() if isinstance(para._element.xml, bytes) else str(para._element.xml)
                if 'TOC' in xml_str and 'fldChar' in xml_str:
                    self.logger.info(f"检测到Word TOC域，位于段落 {i}")
                    return i
            except (AttributeError, UnicodeDecodeError, TypeError):
                pass  # 无法获取XML或解析失败时跳过该段落

        self.logger.info("未检测到目录，将使用标题样式识别方案")
        return None

    def _parse_toc_items(self, doc: Document, toc_start_idx: int) -> Tuple[List[Dict], int]:
        """
        解析目录项（改进3：确保 toc_end_idx > toc_start_idx + SDT支持）

        Args:
            doc: Word文档对象
            toc_start_idx: 目录起始段落索引

        Returns:
            (目录项列表, 目录结束索引)
            目录项列表格式：[{'title': '...', 'page_num': 1, 'level': 1}, ...]
        """
        toc_items = []
        consecutive_non_toc = 0  # 连续非目录项计数
        toc_end_idx = toc_start_idx  # 目录结束位置

        # ⭐️ 目录项数量限制（避免扫描过远）
        MAX_TOC_ITEMS = 200

        # 🆕 特殊处理：检查目录是否在SDT中
        try:
            body = doc.element.body
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # 计算toc_start_idx对应的body子元素
            para_count = 0
            target_sdt = None

            for child in body:
                child_tag = child.tag.split('}')[-1]
                if child_tag == 'p':
                    if para_count == toc_start_idx:
                        # 找到了,但这个位置不是SDT
                        break
                    para_count += 1
                elif child_tag == 'sdt':
                    sdt_para_count = len(child.findall('.//w:p', namespaces=ns))
                    if para_count == toc_start_idx:
                        # toc_start_idx正好指向这个SDT
                        target_sdt = child
                        self.logger.info(f"目录位于SDT中，包含{sdt_para_count}个段落")
                        break
                    para_count += sdt_para_count

            # 如果找到SDT目录，直接从SDT中提取目录项
            if target_sdt is not None:
                sdt_paras = target_sdt.findall('.//w:p', namespaces=ns)
                self.logger.info(f"从SDT中提取目录项，共{len(sdt_paras)}个段落")

                for para_elem in sdt_paras[1:]:  # 跳过第一个"目录"标题段落
                    # 提取段落文本
                    texts = para_elem.findall('.//w:t', namespaces=ns)
                    text = ''.join([t.text for t in texts if t.text])
                    text = text.strip()

                    if not text:
                        continue

                    # 解析目录项（使用相同的正则模式）
                    # 匹配目录项格式（带页码）
                    toc_pattern = re.compile(
                        r'^(.+?)'  # 标题（非贪婪）
                        r'[\s\.\u2026]*'  # 可能的点、空格或省略号
                        r'(\d+)$'  # 页码
                    )

                    match = toc_pattern.match(text)
                    if match:
                        title = match.group(1).strip()
                        page_num = int(match.group(2))

                        if len(toc_items) >= MAX_TOC_ITEMS:
                            self.logger.info(f"SDT目录项已达上限({MAX_TOC_ITEMS})，停止解析")
                            break

                        # 从XML元素推断层级（因为无法直接访问Paragraph对象）
                        level = 1
                        if re.match(r'^\d+\.\d+\.\d+', title):  # 先检查3级
                            level = 3
                        elif re.match(r'^\d+\.\d+[^\d]', title):  # 再检查2级
                            level = 2

                        toc_items.append({
                            'title': title,
                            'page_num': page_num,
                            'level': level
                        })

                        self.logger.debug(f"SDT目录项 [{level}级]: {title} (页码:{page_num})")

                # SDT目录解析完成
                if toc_items:
                    toc_end_idx = toc_start_idx - 1  # SDT不占用paragraphs索引,SDT后第一个段落就是toc_start_idx
                    self.logger.info(f"SDT目录解析完成，共{len(toc_items)}项")
                    return (toc_items, toc_end_idx)

        except Exception as e:
            self.logger.debug(f"SDT目录解析失败，回退到常规方法: {e}")

        # 常规方法：从python-docx的paragraphs中解析
        # 🔧 扩展范围从100到150，以覆盖更长的目录（如包含100+个目录项的招标文件）
        # 🆕 从 toc_start_idx 开始（而不是 +1），因为 toc_start_idx 可能本身就是第一个目录项
        #    如果 toc_start_idx 段落只包含"目录"标题，会被后续逻辑跳过
        for i in range(toc_start_idx, min(toc_start_idx + 150, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()

            # 跳过空行
            if not text:
                continue

            # 尝试匹配目录项格式
            # 格式1: "标题文本    页码" (多个空格)
            match = re.match(r'^(.+?)\s{2,}(\d+)$', text)
            if not match:
                # 格式2: "标题文本....页码" (点号填充)
                match = re.match(r'^(.+?)\.{2,}(\d+)$', text)
            if not match:
                # 格式3: "标题文本\t页码" (制表符)
                match = re.match(r'^(.+?)\t+(\d+)$', text)

            # 🆕 格式4: "第X章 标题 页码" (单个空格，专门匹配章节标题)
            # 使用非贪婪匹配，确保页码前的空格被正确识别
            if not match:
                match = re.match(r'^(第[一二三四五六七八九十百千\d]+[章部分节篇].+?)\s+(\d+)$', text)

            # 🆕 格式5: "标题文本 - 页码 -" (带短横线的页码格式)
            # 某些招标文件使用此格式，如 "第一章	投标邀请书	- 3 -"
            if not match:
                match = re.match(r'^(.+?)\s*-\s*(\d+)\s*-\s*$', text)

            if match:
                title = match.group(1).strip()
                page_num = int(match.group(2))

                # ⭐️ 重复检测：如果与第一项重复，说明扫描到正文了
                # 使用模糊匹配代替严格相等，以应对文本略有差异的情况
                if toc_items:
                    first_title = toc_items[0]['title']
                    similarity = self.fuzzy_match_title_v2(title, first_title)
                    if similarity >= 0.70:  # 70%以上相似度认为是重复
                        self.logger.info(f"检测到重复目录项（与第1项相似度{similarity:.0%}），目录解析结束")
                        self.logger.debug(f"  第1项: '{first_title}' vs 当前: '{title}'")
                        break

                # ⭐️ 数量限制检查
                if len(toc_items) >= MAX_TOC_ITEMS:
                    self.logger.info(f"目录项已达上限({MAX_TOC_ITEMS})，停止解析")
                    break

                # 层级初始化为1，后续由 LevelAnalyzer 修正
                level = 1

                toc_items.append({
                    'title': title,
                    'page_num': page_num,
                    'level': level
                })

                self.logger.debug(f"目录项 [{level}级]: {title} -> 第{page_num}页")

                # 更新目录结束位置
                toc_end_idx = i
                # 重置计数
                consecutive_non_toc = 0
            else:
                # 新增：尝试识别无页码的简单目录项
                # 🆕 首先过滤目录标题本身，避免将"目录"当作目录项
                TOC_TITLE_KEYWORDS = ['目录', '目  录', '目 录', 'contents', 'index', 'catalogue']
                text_normalized = text.replace(' ', '').replace('\u3000', '').lower()
                if text_normalized in [k.replace(' ', '').lower() for k in TOC_TITLE_KEYWORDS]:
                    self.logger.debug(f"跳过目录标题本身: '{text}'")
                    continue

                # 特征1：有统一缩进（目录项通常缩进对齐）
                has_indent = (para.paragraph_format.left_indent and
                              para.paragraph_format.left_indent > 200000)

                # 特征2：匹配章节模式
                is_chapter_pattern = (
                    re.match(r'^第[一二三四五六七八九十\d]+章', text) or
                    re.match(r'^第[一二三四五六七八九十\d]+部分', text) or
                    text in ['竞争性磋商公告', '招标公告', '采购公告', '谈判公告', '询价公告']
                )

                # 🆕 排除 Heading 样式的段落 - Heading 样式说明这是正文标题，不是目录项
                # 目录项通常使用 TOC 样式或普通样式，不会使用 Heading 样式
                is_heading_style = para.style and para.style.name.startswith('Heading')
                if is_heading_style and is_chapter_pattern:
                    self.logger.info(f"检测到Heading样式的章节标题 '{text}'，判定为正文开始，目录解析结束")
                    break

                # 如果满足条件，作为无页码目录项
                if (has_indent or is_chapter_pattern) and len(text) < 50 and not text.startswith('项目'):
                    title = text.strip()

                    # ⭐️ 重复检测：如果与第一项重复，说明扫描到正文了
                    # 使用模糊匹配代替严格相等，以应对文本略有差异的情况
                    if toc_items:
                        first_title = toc_items[0]['title']
                        similarity = self.fuzzy_match_title_v2(title, first_title)
                        if similarity >= 0.70:  # 70%以上相似度认为是重复
                            self.logger.info(f"检测到重复目录项（与第1项相似度{similarity:.0%}），目录解析结束")
                            self.logger.debug(f"  第1项: '{first_title}' vs 当前: '{title}'")
                            break

                    # ⭐️ 数量限制检查
                    if len(toc_items) >= MAX_TOC_ITEMS:
                        self.logger.info(f"目录项已达上限({MAX_TOC_ITEMS})，停止解析")
                        break

                    page_num = -1  # 标记为无页码
                    level = 1  # 初始值，后续由 LevelAnalyzer 修正

                    toc_items.append({
                        'title': title,
                        'page_num': page_num,
                        'level': level
                    })

                    self.logger.debug(f"目录项(无页码) [{level}级]: {title}")
                    toc_end_idx = i
                    consecutive_non_toc = 0
                    continue  # 成功识别，继续下一个

                # 非目录项
                consecutive_non_toc += 1
                # 连续5行不匹配，认为目录结束
                if consecutive_non_toc >= 5 and len(toc_items) > 0:
                    self.logger.info(f"目录解析完成，共 {len(toc_items)} 项，结束于段落 {toc_end_idx}")
                    break
                # 新增：如果找到了目录项但遇到空行后的Heading样式，认为目录结束
                if (len(toc_items) > 0 and consecutive_non_toc >= 2 and
                    para.style and para.style.name.startswith('Heading')):
                    self.logger.info(f"检测到Heading样式，目录结束于段落 {toc_end_idx}")
                    break

        # 改进3：确保 toc_end_idx 严格大于 toc_start_idx
        if toc_end_idx == toc_start_idx:
            # 如果没有找到任何目录项，至少向后移动1个段落
            toc_end_idx = toc_start_idx + 1
            self.logger.warning(f"未解析到目录项，将目录结束位置设为 {toc_end_idx}（避免逻辑错误）")

        # 使用层级分析器修正层级（基于统计频率的上下文分析）
        if toc_items:
            self.logger.info(f"⚡️ 开始使用contextual方法分析 {len(toc_items)} 个TOC项的层级")
            analyzer = LevelAnalyzer()
            corrected_levels = analyzer.analyze_toc_hierarchy_contextual(toc_items)
            for i, level in enumerate(corrected_levels):
                toc_items[i]['level'] = level

            from collections import Counter
            level_dist = Counter(corrected_levels)
            self.logger.info(f"使用contextual方法修正层级，分布: {dict(level_dist)}")

        # 🆕 轻量级合同潜在标记（基于标题关键词）
        if toc_items:
            # 排除规则：这些章节不应被标记为合同
            exclude_keywords = [
                "投标人须知", "供应商须知", "投标邀请", "投标须知", "招标须知",
                "附件", "投标文件格式", "响应文件格式", "投标文件组成", "谈判响应文件格式",
                "用户需求", "需求书", "技术要求", "技术规格", "技术需求书", "技术需求",
                "评分", "评审", "评标", "开标", "报价", "投标报价",
                "投标邀请函", "采购邀请", "谈判邀请"
            ]

            # 合同关键词：只有包含这些关键词才标记为潜在合同
            contract_keywords = [
                "合同条款", "合同文本", "合同范本", "合同协议", "合同格式",
                "通用条款", "专用条款", "拟签合同", "合同草稿", "合同主要条款"
            ]

            contract_potential_count = 0
            for item in toc_items:
                title = item['title']

                # 排除规则优先
                is_excluded = any(kw in title for kw in exclude_keywords)
                is_contract_potential = any(kw in title for kw in contract_keywords)

                # 只有在不被排除且包含合同关键词时才标记
                item['is_contract_potential'] = is_contract_potential and not is_excluded

                if item['is_contract_potential']:
                    contract_potential_count += 1
                    self.logger.debug(f"  🏷️  潜在合同章节: '{title}'")
                elif is_excluded:
                    self.logger.debug(f"  ⛔ 排除章节: '{title}' (匹配排除规则)")

            if contract_potential_count > 0:
                self.logger.info(f"🏷️  标记了 {contract_potential_count} 个潜在合同章节")

        return toc_items, toc_end_idx

    def _find_paragraph_by_title(self, doc: Document, title: str, start_idx: int = 0, toc_items: Optional[List[Dict]] = None, toc_end_idx: Optional[int] = None) -> Optional[int]:
        """
        在文档中搜索与标题匹配的段落

        Args:
            doc: Word文档对象
            title: 要搜索的标题文本
            start_idx: 开始搜索的段落索引
            toc_items: 目录项列表（可选），用于智能检测元数据列表
            toc_end_idx: 目录结束的段落索引（可选），用于限制元数据检测不扫描目录区域

        Returns:
            段落索引，如果未找到则返回None
        """
        def aggressive_normalize(text: str) -> str:
            """激进文本规范化：移除所有分隔符、前缀、空格"""
            # 移除"附件-"、"附件:"等前缀
            text = re.sub(r'^附件[-:：]?', '', text)
            # 移除连字符、下划线、制表符
            text = re.sub(r'[-_\t]+', '', text)
            # 移除所有空格和冒号（将"第一章："和"第一章 "视为等价）
            text = re.sub(r'[:：\s]+', '', text)
            return text

        def extract_core_keywords(text: str) -> str:
            """提取核心关键词：去除编号和常见前缀"""
            # 移除编号（支持冒号和空格：第X章：、第X章 、第X章、第X部分：等）
            text = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章)[:：\s]*', '', text)
            text = re.sub(r'^(\d+\.|\d+\.\d+|[一二三四五六七八九十]+、)\s*', '', text)
            # 移除"附件"前缀
            text = re.sub(r'^附件[-:：]?', '', text)
            # 移除分隔符
            text = re.sub(r'[-_\t]+', '', text)
            # 移除空格和冒号
            text = re.sub(r'[:：\s]+', '', text)
            return text

        def calculate_similarity(str1: str, str2: str) -> float:
            """计算两个字符串的相似度（使用 SequenceMatcher）"""
            if not str1 or not str2:
                return 0.0

            # 使用 SequenceMatcher 计算真正的字符串相似度
            # 这比基于子串的方法更能处理同义词替换（如"响应"vs"应答"）
            return SequenceMatcher(None, str1, str2).ratio()

        # 清理标题（移除多余空格）
        clean_title = re.sub(r'\s+', '', title)

        # 激进规范化的标题
        aggressive_title = aggressive_normalize(title)

        # 提取核心关键词
        core_keywords = extract_core_keywords(aggressive_title)

        self.logger.info(f"搜索标题: '{title}' (清理后: '{clean_title}', 核心: '{core_keywords}'), 从段落 {start_idx} 开始")

        # 候选匹配列表（用于宽松匹配级别）
        # 格式: [(段落索引, 匹配级别, 得分, 段落文本, 匹配原因)]
        loose_match_candidates = []

        # 记录已跳过的元数据列表区域（避免重复检测和日志）
        skipped_ranges = []

        # 🔑 第一轮：严格匹配 (Level 1-3)，找到立即返回
        for i in range(start_idx, len(doc.paragraphs)):
            # 检查是否在已跳过的区域中
            if any(start <= i <= end for start, end in skipped_ranges):
                continue
            para = doc.paragraphs[i]
            para_text = para.text.strip()

            if not para_text:
                continue

            # 清理段落文本
            clean_para = re.sub(r'\s+', '', para_text)

            # 激进规范化的段落
            aggressive_para = aggressive_normalize(para_text)

            # 段落核心关键词
            para_keywords = extract_core_keywords(aggressive_para)

            # Level 1: 完全匹配或包含匹配
            level1_exact_match = (clean_title == clean_para)
            level1_contain_match = (clean_title in clean_para)

            if level1_exact_match or level1_contain_match:
                # ⭐️ 检查是否在连续章节标题列表中（如"文件构成说明"）
                list_range = self._detect_chapter_title_list_range(doc, i, toc_items, toc_end_idx)

                if list_range:
                    start, end, titles = list_range

                    # 🔑 关键：对于"第X章"、"第X部分"这种明确的一级章节标题
                    # 只有当它具有 Heading 样式时才认为是真正的章节标题
                    # 否则可能只是文档中的一个章节列表/索引区域
                    is_primary_chapter = re.match(r'^第[一二三四五六七八九十百千\d]+[章部分]', para_text.strip())
                    is_heading_style = para.style and ('heading' in para.style.name.lower() or '标题' in para.style.name.lower())

                    if is_primary_chapter and is_heading_style:
                        self.logger.info(f"  ✓ 找到一级章节标题（Heading样式，不跳过）: 段落 {i}: '{para_text}'")
                        return i

                    # 如果是"第X章"但不是Heading样式，很可能是文档中的章节列表区域
                    if is_primary_chapter and not is_heading_style:
                        self.logger.info(f"  ⏭️  跳过段落 {i}（第X章格式但非Heading样式，可能是章节列表）: '{para_text}'")
                        # 记录跳过区域
                        skipped_ranges.append((start, end))
                        continue

                    self.logger.info(
                        f"  ⚠️  检测到'文件构成说明'列表 (段落{start}-{end}，"
                        f"包含{len(titles)}个章节标题)，跳过该区域"
                    )
                    # 显示列表中的章节（最多5个）
                    for idx, title_text in titles[:5]:
                        self.logger.info(f"      - 段落{idx}: {title_text}")
                    if len(titles) > 5:
                        self.logger.info(f"      ... 还有{len(titles)-5}个")

                    # 记录跳过区域
                    skipped_ranges.append((start, end))

                    # 继续搜索（不返回此匹配）
                    continue

                # 🔑 对于包含匹配（非完全匹配），需要额外验证
                if level1_contain_match and not level1_exact_match:
                    # 检查是否是Heading样式（更可能是真正的章节标题）
                    is_heading = para.style and ('heading' in para.style.name.lower() or '标题' in para.style.name.lower())
                    # 或者文本很短（≤20字，更可能是标题而不是正文）
                    is_short = len(para_text) <= 20
                    # 🆕 检查标题是否在段落开头位置（章节标题应该在开头，不是中间）
                    title_at_start = clean_para.startswith(clean_title)

                    if not is_heading:
                        # 非 Heading 样式时，需要满足：短文本 或 标题在开头
                        if not (is_short or title_at_start):
                            # 标题在段落中间 → 可能是正文中提到标题 → 跳过
                            self.logger.debug(f"  ⏭️  跳过段落 {i}（包含匹配但标题不在开头且非短文本）: '{para_text[:60]}'")
                            continue

                self.logger.info(f"  ✓ 找到匹配 (Level 1-完全): 段落 {i}: '{para_text}'")
                return i

            # Level 2: 激进规范化后的完全匹配或包含匹配
            level2_exact_match = (aggressive_title == aggressive_para)
            level2_contain_match = (aggressive_title in aggressive_para)

            if level2_exact_match or level2_contain_match:
                # 🔑 对于包含匹配（非完全匹配），需要额外验证
                if level2_contain_match and not level2_exact_match:
                    is_heading = para.style and ('heading' in para.style.name.lower() or '标题' in para.style.name.lower())
                    is_short = len(para_text) <= 20
                    # 🆕 检查标题是否在段落开头位置（章节标题应该在开头，不是中间）
                    title_at_start = aggressive_para.startswith(aggressive_title)

                    if not is_heading:
                        # 非 Heading 样式时，需要满足：短文本 或 标题在开头
                        if not (is_short or title_at_start):
                            # 标题在段落中间 → 可能是正文中提到标题 → 跳过
                            self.logger.debug(f"  ⏭️  跳过段落 {i}（包含匹配但标题不在开头且非短文本）: '{para_text[:60]}'")
                            continue

                self.logger.info(f"  ✓ 找到匹配 (Level 2-规范化): 段落 {i}: '{para_text}'")
                return i

            # 检查标题和段落是否包含"第X部分"（用于Level 3和Level 4约束）
            title_has_part_number = bool(re.search(r'第[一二三四五六七八九十\d]+部分', title))
            para_has_part_number = bool(re.search(r'第[一二三四五六七八九十\d]+部分', para_text))

            # Level 3: 去除编号后的匹配
            # 支持多种编号格式：第X部分、第X章、1.、1.1、一、等
            title_without_number = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章|\d+\.|\d+\.\d+|[一二三四五六七八九十]+、)\s*', '', clean_title)
            para_without_number = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章|\d+\.|\d+\.\d+|[一二三四五六七八九十]+、)\s*', '', clean_para)

            if title_without_number and para_without_number and title_without_number == para_without_number:
                # 如果TOC标题有"第X部分"，则段落也必须有"第X部分"（避免匹配到TOC内的编号内容）
                # 🆕 例外：如果正文段落使用 Heading 样式，说明是真正的章节标题，即使没有"第X部分"前缀也应匹配
                is_heading = para.style and ('heading' in para.style.name.lower() or '标题' in para.style.name.lower())
                if title_has_part_number and not para_has_part_number:
                    if is_heading:
                        # 正文使用 Heading 样式，视为有效的章节标题，可以匹配
                        self.logger.info(f"  ✓ 找到匹配 (Level 3-去编号+Heading样式): 段落 {i}: '{para_text}'")
                        return i
                    else:
                        pass  # 跳过，不匹配
                else:
                    self.logger.info(f"  ✓ 找到匹配 (Level 3-去编号): 段落 {i}: '{para_text}'")
                    return i

        # 🔑 第二轮：宽松匹配 (Level 4-7)，收集所有候选
        self.logger.info(f"  严格匹配 (Level 1-3) 未找到，开始宽松匹配 (Level 4-7)")

        for i in range(start_idx, len(doc.paragraphs)):
            # 检查是否在已跳过的区域中
            if any(start <= i <= end for start, end in skipped_ranges):
                continue
            para = doc.paragraphs[i]
            para_text = para.text.strip()

            if not para_text:
                continue

            # 清理段落文本
            clean_para = re.sub(r'\s+', '', para_text)

            # 激进规范化的段落
            aggressive_para = aggressive_normalize(para_text)

            # 段落核心关键词
            para_keywords = extract_core_keywords(aggressive_para)

            # 检查标题和段落是否包含"第X部分"
            title_has_part_number = bool(re.search(r'第[一二三四五六七八九十\d]+部分', title))
            para_has_part_number = bool(re.search(r'第[一二三四五六七八九十\d]+部分', para_text))

            # 去除编号
            title_without_number = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章|\d+\.|\d+\.\d+|[一二三四五六七八九十]+、)\s*', '', clean_title)
            para_without_number = re.sub(r'^(第[一二三四五六七八九十\d]+部分|第[一二三四五六七八九十\d]+章|\d+\.|\d+\.\d+|[一二三四五六七八九十\d]+、)\s*', '', clean_para)

            # Level 4: 核心关键词匹配（长度≥4字）
            if len(core_keywords) >= 4 and len(para_keywords) >= 4:
                # 完全相等匹配（得分最高）
                if core_keywords == para_keywords:
                    loose_match_candidates.append((i, 4, 100, para_text, f"关键词完全相等: '{core_keywords}'"))
                # 标题关键词包含段落关键词（得分中等）
                elif core_keywords in para_keywords:
                    loose_match_candidates.append((i, 4, 70, para_text, f"关键词包含: '{core_keywords}' in '{para_keywords}'"))
                # 段落关键词包含标题关键词（得分较低，容易误匹配）
                elif para_keywords in core_keywords:
                    loose_match_candidates.append((i, 4, 50, para_text, f"被包含: '{para_keywords}' in '{core_keywords}'"))

            # Level 4.5: 部分子串匹配
            if len(core_keywords) >= 6 and title_has_part_number:
                for substr_len in range(len(core_keywords), 5, -1):
                    substr = core_keywords[:substr_len]
                    if substr in para_keywords and len(substr) >= 6:
                        if para_has_part_number and len(para_text) <= 50:
                            match_ratio = len(substr) / len(core_keywords)
                            score = 65 + match_ratio * 10  # 65-75分
                            loose_match_candidates.append((i, 4.5, score, para_text, f"部分子串{match_ratio:.0%}: '{substr}'"))
                        break

            # Level 5: 相似度匹配（相似度≥60%）
            if len(core_keywords) >= 4:
                similarity = calculate_similarity(core_keywords, para_keywords)
                if similarity >= 0.6:
                    # 🆕 检查章节编号是否匹配（如"第三章"）
                    title_chapter_match = re.match(r'^第([一二三四五六七八九十\d]+)章', title)
                    para_chapter_match = re.match(r'^第([一二三四五六七八九十\d]+)章', para_text)

                    if title_chapter_match and para_chapter_match:
                        # 两者都有章节编号
                        if title_chapter_match.group(1) == para_chapter_match.group(1):
                            # 章节编号相同（如都是"第三章"），大幅提高分数
                            # 这种情况下，即使内容略有不同（如"响应"vs"应答"），也应该优先匹配
                            score = 80 + similarity * 20  # 80-100分
                            loose_match_candidates.append((i, 5, score, para_text, f"章节编号匹配+相似度{similarity:.0%}"))
                        else:
                            # 章节编号不同，降低分数
                            score = similarity * 30  # 18-30分
                            loose_match_candidates.append((i, 5, score, para_text, f"相似度{similarity:.0%}(编号不同)"))
                    else:
                        # 普通相似度匹配
                        score = similarity * 60  # 36-60分
                        loose_match_candidates.append((i, 5, score, para_text, f"相似度{similarity:.0%}"))

            # Level 6: 宽松关键词匹配（至少6字标题）
            if len(title_without_number) >= 6:
                if title_without_number in clean_para:
                    loose_match_candidates.append((i, 6, 40, para_text, f"包含去编号标题: '{title_without_number}'"))

            # Level 7: 转换编号后匹配
            def convert_chinese_to_number(text):
                """将第一/第二/第三等转换为1/2/3"""
                mapping = {'一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
                          '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'}
                match = re.match(r'^第([一二三四五六七八九十]+)部分(.*)$', text)
                if match:
                    num = mapping.get(match.group(1), match.group(1))
                    return f"{num}.{match.group(2)}"
                return text

            converted_title = convert_chinese_to_number(clean_title)
            if converted_title != clean_title and clean_para.startswith(converted_title[:3]):
                converted_para_without_num = re.sub(r'^\d+\.', '', clean_para)
                converted_title_without_num = re.sub(r'^\d+\.', '', converted_title)
                if converted_title_without_num == converted_para_without_num:
                    loose_match_candidates.append((i, 7, 30, para_text, f"转换编号后匹配"))

        # 🔑 从候选中选择得分最高的
        if loose_match_candidates:
            # 按得分排序
            loose_match_candidates.sort(key=lambda x: x[2], reverse=True)
            best = loose_match_candidates[0]
            para_idx, level, score, para_text, reason = best

            self.logger.info(f"  ✓ 从 {len(loose_match_candidates)} 个宽松候选中选择最佳匹配:")
            self.logger.info(f"     段落 {para_idx} (Level {level}, 得分{score:.0f}): '{para_text[:60]}'")
            self.logger.info(f"     匹配原因: {reason}")

            # 显示其他候选（前3个）
            if len(loose_match_candidates) > 1:
                self.logger.info(f"  其他候选:")
                for candidate in loose_match_candidates[1:4]:
                    c_idx, c_level, c_score, c_text, c_reason = candidate
                    self.logger.info(f"     段落 {c_idx} (Level {c_level}, 得分{c_score:.0f}): '{c_text[:60]}'")

            return para_idx

        # 未找到任何匹配
        self.logger.warning(f"未找到标题匹配: '{title}'")
        return None
    def _locate_chapters_by_toc(self, doc: Document, toc_items: List[Dict], toc_end_idx: int) -> List[ChapterNode]:
        """
        根据目录项定位章节在文档中的位置，并构建树形结构

        核心改进: 完全使用TOC中已识别的层级信息，不再重新检测子章节
        这样可以保留 analyze_toc_hierarchy_contextual() 的精确分析结果

        Args:
            doc: Word文档对象
            toc_items: 目录项列表（已包含精确的层级信息）
            toc_end_idx: 目录结束的段落索引

        Returns:
            章节树（只包含根节点，子节点在children属性中）
        """
        # 步骤1: 定位所有TOC项在文档中的位置
        all_chapters = []
        last_found_idx = toc_end_idx + 1
        self.logger.info(f"⭐ 开始定位 {len(toc_items)} 个TOC项的位置（目录结束于段落 {toc_end_idx}）")

        for i, item in enumerate(toc_items):
            title = item['title']
            level = item['level']

            # 在文档中查找标题位置（传入toc_items和toc_end_idx用于智能检测元数据列表）
            para_idx = self._find_paragraph_by_title(doc, title, last_found_idx, toc_items, toc_end_idx)

            if para_idx is None:
                self.logger.warning(f"⚠️  未找到目录项: [{level}级] {title}")
                continue

            # 记录找到的章节（暂不确定结束位置）
            all_chapters.append({
                'toc_index': i,
                'title': title,
                'level': level,
                'para_idx': para_idx,
                'para_end_idx': None  # 稍后计算
            })

            last_found_idx = para_idx + 1
            self.logger.debug(f"  ✓ 找到 [{level}级] {title} (段落 {para_idx})")

        # 步骤2: 计算每个章节的结束位置
        # ⭐ 关键修复：采用"同级或更高级"逻辑，确保父章节包含所有子章节的内容
        for i, chapter_info in enumerate(all_chapters):
            current_level = chapter_info['level']
            next_start = len(doc.paragraphs)  # 默认到文档末尾

            # 找下一个同级或更高级的章节（与方法2逻辑保持一致）
            for j in range(i + 1, len(all_chapters)):
                if all_chapters[j]['level'] <= current_level:
                    next_start = all_chapters[j]['para_idx']
                    break

            chapter_info['para_end_idx'] = next_start - 1

        # 🆕 步骤2.5: 重型合同检测（仅处理标记为潜在合同的章节）
        contract_chapters_to_insert = []  # 存储需要插入的合同章节
        contract_check_count = 0
        contract_confirmed_count = 0

        for i, chapter_info in enumerate(all_chapters):
            toc_item = toc_items[chapter_info['toc_index']]

            # 🎯 只对标记为潜在合同的章节执行重型检测
            if toc_item.get('is_contract_potential', False):
                contract_check_count += 1
                self.logger.info(f"  🔍 对潜在合同章节执行重型检测: '{chapter_info['title']}'")

                para_idx = chapter_info['para_idx']
                para_end_idx = chapter_info['para_end_idx']

                # 重型作业1: 提取完整内容样本（2000字）
                content_sample = self._extract_content_sample(
                    doc, para_idx, para_end_idx, sample_size=2000
                )

                # 重型作业2: 计算精确的合同密度
                is_contract, density, reason = self._is_contract_chapter(
                    chapter_info['title'],
                    content_sample
                )

                if is_contract:
                    contract_confirmed_count += 1
                    chapter_info['is_contract_confirmed'] = True
                    self.logger.info(f"  ✓ 合同章节确认: '{chapter_info['title']}' - {reason}")

                    # 重型作业3: 检测合同聚集区（用于章节拆分）
                    contract_cluster = self._detect_contract_cluster_in_chapter(
                        doc, para_idx, para_end_idx
                    )

                    if contract_cluster:
                        # 占比已在检测方法中限制（20%-80%），这里直接处理
                        cluster_start = contract_cluster['start']
                        cluster_end = contract_cluster['end']
                        cluster_ratio = contract_cluster['ratio']

                        # 检查是否满足拆分条件
                        min_content_length = 1000
                        min_paragraph_gap = 5

                        if cluster_start > para_idx + min_paragraph_gap:
                            # 计算前半部分字数
                            front_content = "\n".join(
                                doc.paragraphs[j].text.strip()
                                for j in range(para_idx + 1, cluster_start)
                                if j < len(doc.paragraphs)
                            )
                            front_word_count = self._calculate_word_count(front_content)

                            if front_word_count >= min_content_length:
                                self.logger.warning(
                                    f"✂️  准备拆分章节: '{chapter_info['title']}' "
                                    f"→ 正常部分({para_idx}-{cluster_start-1}, {front_word_count}字) "
                                    f"+ 合同部分({cluster_start}-{para_end_idx}, 占比{cluster_ratio:.1%})"
                                )

                                # 截断原章节（只保留合同之前的部分）
                                chapter_info['para_end_idx'] = cluster_start - 1

                                # 创建合同章节（记录待插入）
                                contract_chapter = {
                                    'toc_index': chapter_info['toc_index'] + 0.5,  # 标记为插入的章节
                                    'title': '[检测到的合同条款-需人工确认]',
                                    'level': chapter_info['level'],  # 与原章节同级
                                    'para_idx': cluster_start,
                                    'para_end_idx': para_end_idx,
                                    'is_contract_confirmed': True
                                }

                                # 记录待插入位置（插入到原章节后面）
                                contract_chapters_to_insert.append((i + 1, contract_chapter))
                            else:
                                self.logger.info(
                                    f"  ⏭️  跳过拆分: '{chapter_info['title']}' - "
                                    f"前半部分内容不足({front_word_count}字 < {min_content_length}字)"
                                )
                        else:
                            self.logger.info(
                                f"  ⏭️  跳过拆分: '{chapter_info['title']}' - "
                                f"合同聚集区起点太靠前(段落{cluster_start})"
                            )
                else:
                    self.logger.info(f"  ✗ 非合同章节: '{chapter_info['title']}' - {reason}")
            else:
                # 非潜在合同章节，跳过重型检测
                self.logger.debug(f"  ⏭️  跳过重型检测: '{chapter_info['title']}'")

        # 插入拆分出的合同章节（倒序插入，避免索引偏移）
        for insert_pos, contract_chapter in reversed(contract_chapters_to_insert):
            all_chapters.insert(insert_pos, contract_chapter)
            self.logger.info(
                f"  ➕ 插入合同章节: '{contract_chapter['title']}' "
                f"(段落{contract_chapter['para_idx']}-{contract_chapter['para_end_idx']})"
            )

        if contract_check_count > 0:
            self.logger.info(
                f"🔍 重型检测完成: 检测了 {contract_check_count} 个潜在章节，确认 {contract_confirmed_count} 个合同章节，"
                f"拆分出 {len(contract_chapters_to_insert)} 个合同章节"
            )

        # 步骤3: 创建所有ChapterNode对象
        chapter_nodes = []
        for chapter_info in all_chapters:
            para_idx = chapter_info['para_idx']
            para_end_idx = chapter_info['para_end_idx']

            # 提取章节内容（包括段落和表格）
            content_text, preview_text, has_table = self._extract_chapter_content_with_tables(
                doc, para_idx, para_end_idx
            )

            # 计算字数
            word_count = self._calculate_word_count(content_text)

            # 如果没有预览文本，设置默认值
            if not preview_text:
                preview_text = "(无内容)"

            node = ChapterNode(
                id=f"ch_{chapter_info['toc_index']}",
                level=chapter_info['level'],  # ⭐ 使用TOC中的层级，不重新检测
                title=chapter_info['title'],
                para_start_idx=para_idx,
                para_end_idx=para_end_idx,
                word_count=word_count,
                preview_text=preview_text,
                has_table=has_table
            )

            chapter_nodes.append(node)

        # 步骤4: 构建树形结构（基于level属性）
        root_chapters = self._build_chapter_tree(chapter_nodes)

        self.logger.info(f"✅ 成功构建章节树: {len(root_chapters)} 个根章节")
        return root_chapters
    def _locate_chapter_content(self, doc: Document, chapters: List[ChapterNode]) -> List[ChapterNode]:
        """
        定位每个章节的内容范围

        Args:
            doc: Word 文档对象
            chapters: 章节列表

        Returns:
            更新后的章节列表（包含 para_end_idx、word_count、preview_text）
        """
        # ⭐️ 关键修复：按段落索引排序，确保章节顺序与文档物理顺序一致
        # 防止索引倒置问题（如 para_start_idx=542 > para_end_idx=62）
        chapters_sorted = sorted(chapters, key=lambda ch: ch.para_start_idx)
        self.logger.info(f"章节已按段落索引排序，共 {len(chapters_sorted)} 个章节")

        total_paras = len(doc.paragraphs)

        # 🆕 用于收集需要插入的合同章节
        contract_chapters_to_insert = []

        for i, chapter in enumerate(chapters_sorted):
            # 确定章节结束位置（下一个同级或更高级标题的前一个段落）
            next_start = total_paras  # 默认到文档末尾

            for j in range(i + 1, len(chapters_sorted)):
                if chapters_sorted[j].level <= chapter.level:
                    next_start = chapters_sorted[j].para_start_idx
                    break

            chapter.para_end_idx = next_start - 1

            # 提取章节内容（包括段落和表格）
            content_text, preview_text, has_table = self._extract_chapter_content_with_tables(
                doc, chapter.para_start_idx, chapter.para_end_idx
            )

            # 计算字数
            chapter.word_count = self._calculate_word_count(content_text)
            chapter.preview_text = preview_text if preview_text else "(无内容)"
            chapter.has_table = has_table

            # 【新增】对于level 1-2的章节，提取内容样本并进行合同识别
            if chapter.level <= 2:
                # 提取内容样本用于合同识别
                chapter.content_sample = self._extract_content_sample(
                    doc, chapter.para_start_idx, chapter.para_end_idx, sample_size=2000
                )

                # 基于内容进行合同识别
                is_contract, density, reason = self._is_contract_chapter(
                    chapter.title, chapter.content_sample
                )

                if is_contract:
                    self.logger.info(
                        f"  ✓ 合同章节识别: '{chapter.title}' - {reason}"
                    )

            # 🆕 新增：检测章节内是否有合同聚集区（用于拆分章节）
            contract_cluster = self._detect_contract_cluster_in_chapter(
                doc, chapter.para_start_idx, chapter.para_end_idx
            )

            if contract_cluster:
                cluster_start = contract_cluster['start']
                cluster_end = contract_cluster['end']
                density = contract_cluster['density']

                # 确保聚集区起点在章节内且有足够的前置内容
                min_content_length = 1000  # 前半部分至少1000字

                if cluster_start > chapter.para_start_idx + 5:  # 至少跳过5个段落
                    # 计算前半部分的字数
                    front_content = "\n".join(
                        doc.paragraphs[j].text.strip()
                        for j in range(chapter.para_start_idx + 1, cluster_start)
                        if j < len(doc.paragraphs)
                    )
                    front_word_count = self._calculate_word_count(front_content)

                    if front_word_count >= min_content_length:
                        self.logger.warning(
                            f"⚠️ 章节将被拆分: '{chapter.title}' "
                            f"→ 正常部分({chapter.para_start_idx}-{cluster_start-1}, {front_word_count}字) "
                            f"+ 合同部分({cluster_start}-{cluster_end}, 密度{density:.1%})"
                        )

                        # 截断当前章节（只保留合同之前的部分）
                        original_end = chapter.para_end_idx
                        chapter.para_end_idx = cluster_start - 1

                        # 重新计算缩短后的章节内容
                        content_text, preview_text, has_table = self._extract_chapter_content_with_tables(
                            doc, chapter.para_start_idx, chapter.para_end_idx
                        )
                        chapter.word_count = self._calculate_word_count(content_text)
                        chapter.preview_text = preview_text
                        chapter.has_table = has_table

                        # 🆕 创建合同章节（标记为待插入）
                        contract_chapter = ChapterNode(
                            id=f"ch_{i}_contract",  # 临时ID，后续会重新分配
                            level=chapter.level,  # 与原章节同级
                            title="[检测到的合同条款-需人工确认]",
                            para_start_idx=cluster_start,
                            para_end_idx=original_end,
                            word_count=0,
                            preview_text=""
                        )

                        # 计算合同章节内容
                        contract_content, contract_preview, contract_has_table = self._extract_chapter_content_with_tables(
                            doc, contract_chapter.para_start_idx, contract_chapter.para_end_idx
                        )
                        contract_chapter.word_count = self._calculate_word_count(contract_content)
                        contract_chapter.preview_text = contract_preview
                        contract_chapter.has_table = contract_has_table

                        # 添加到待插入列表（记录插入位置）
                        contract_chapters_to_insert.append((i + 1, contract_chapter))

                        self.logger.info(
                            f"✂️ 章节拆分完成: "
                            f"正常部分({chapter.para_start_idx}-{chapter.para_end_idx}, {chapter.word_count}字) "
                            f"+ 合同部分({contract_chapter.para_start_idx}-{contract_chapter.para_end_idx}, {contract_chapter.word_count}字)"
                        )
                    else:
                        self.logger.info(
                            f"跳过拆分: '{chapter.title}' 前半部分内容不足({front_word_count}字 < {min_content_length}字)"
                        )
                else:
                    self.logger.info(
                        f"跳过拆分: '{chapter.title}' 合同聚集区起点太靠前(段落{cluster_start})"
                    )

        # 🆕 插入所有检测到的合同章节
        if contract_chapters_to_insert:
            # 按插入位置倒序插入（避免索引偏移）
            for insert_pos, contract_chapter in reversed(contract_chapters_to_insert):
                chapters_sorted.insert(insert_pos, contract_chapter)

            self.logger.info(f"已插入 {len(contract_chapters_to_insert)} 个合同章节")

            # 🆕 重新分配章节ID，确保ID连续
            for idx, ch in enumerate(chapters_sorted):
                ch.id = f"ch_{idx}"

            self.logger.info(f"章节ID已重新分配，当前共 {len(chapters_sorted)} 个章节")

        return chapters_sorted

    def _extract_chapter_content_with_tables(self, doc: Document, para_start_idx: int, para_end_idx: int) -> tuple:
        """
        提取章节内容,包括段落和表格

        Args:
            doc: Word文档对象
            para_start_idx: 起始段落索引
            para_end_idx: 结束段落索引

        Returns:
            (完整内容文本, 预览文本, 是否包含表格)
        """
        # 构建段落索引到body元素索引的映射
        para_count = 0
        start_body_idx = None
        end_body_idx = None

        for body_idx, element in enumerate(doc.element.body):
            if isinstance(element, CT_P):
                if para_count == para_start_idx and start_body_idx is None:
                    start_body_idx = body_idx
                if para_count == para_end_idx:
                    end_body_idx = body_idx
                    break
                para_count += 1

        if start_body_idx is None:
            return "", "", False

        if end_body_idx is None:
            end_body_idx = len(doc.element.body) - 1

        # 提取内容(跳过章节标题,从start+1开始)
        content_parts = []
        preview_lines = []
        preview_limit = 5
        has_table = False  # 标记是否包含表格

        for body_idx in range(start_body_idx + 1, end_body_idx + 1):
            element = doc.element.body[body_idx]

            if isinstance(element, CT_P):
                # 段落
                from docx.text.paragraph import Paragraph
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    content_parts.append(text)
                    # 添加到预览
                    if len(preview_lines) < preview_limit:
                        preview_lines.append(text[:100] + ('...' if len(text) > 100 else ''))

            elif isinstance(element, CT_Tbl):
                # 表格
                from docx.table import Table
                table = Table(element, doc)

                # 提取表格文本
                table_text_parts = []
                table_preview_parts = []

                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                        row_data.append(cell_text)

                    if any(cell.strip() for cell in row_data):  # 非空行
                        row_text = ' | '.join(row_data)
                        table_text_parts.append(row_text)

                        # 添加到预览(表格的前几行)
                        if len(preview_lines) < preview_limit and row_idx < 3:
                            table_preview_parts.append(row_text[:100] + ('...' if len(row_text) > 100 else ''))

                if table_text_parts:
                    has_table = True  # 发现表格
                    # 添加表格标识
                    table_content = f"[表格]\n" + '\n'.join(table_text_parts)
                    content_parts.append(table_content)

                    # 添加表格预览
                    if len(preview_lines) < preview_limit:
                        preview_lines.append("[表格]")
                        preview_lines.extend(table_preview_parts[:preview_limit - len(preview_lines)])

        full_content = '\n'.join(content_parts)
        preview_text = '\n'.join(preview_lines)

        return full_content, preview_text, has_table

    def _calculate_word_count(self, text: str) -> int:
        """
        计算文本的字数（中文字符数）

        统计规则：去除空格和换行后的字符总数

        Args:
            text: 待统计的文本

        Returns:
            字数（去除空格和换行后的字符数）

        Examples:
            >>> parser._calculate_word_count("Hello World\\n测试")
            13  # "HelloWorld测试" = 13个字符

            >>> parser._calculate_word_count("")
            0

        Note:
            此方法提供统一的字数统计逻辑，用于：
            - 章节内容字数统计
            - 判断内容是否足够（如合同拆分前的字数检查）
            - 预览显示的字数信息
        """
        if not text:
            return 0
        return len(text.replace(' ', '').replace('\n', ''))

    def _build_chapter_tree(self, chapters: List[ChapterNode]) -> List[ChapterNode]:
        """
        构建章节层级树

        Args:
            chapters: 扁平章节列表

        Returns:
            根级章节列表（包含子章节）
        """
        if not chapters:
            return []

        # 使用栈来构建树
        root_chapters = []
        stack = []  # [(level, chapter), ...]

        for chapter in chapters:
            # 弹出所有层级 >= 当前章节层级的节点
            while stack and stack[-1][0] >= chapter.level:
                stack.pop()

            if not stack:
                # 当前是根级章节
                root_chapters.append(chapter)
            else:
                # 当前是子章节，添加到父章节的 children
                parent = stack[-1][1]
                parent.children.append(chapter)
                # 更新子章节ID（包含父级ID）
                chapter.id = f"{parent.id}_{len(parent.children)}"

            # 将当前章节入栈
            stack.append((chapter.level, chapter))

        return root_chapters

    def _calculate_statistics(self, chapter_tree: List[ChapterNode]) -> Dict:
        """
        计算统计信息

        Args:
            chapter_tree: 章节树

        Returns:
            统计字典
        """
        stats = {
            "total_chapters": 0,
            "total_words": 0,
            "estimated_processing_cost": 0.0
        }

        def traverse(chapters, is_root_level=False):
            for ch in chapters:
                stats["total_chapters"] += 1

                # 🔑 只统计根节点（1级章节）的字数，避免重复统计
                # 因为父节点的字数已经包含了所有子节点的内容
                if is_root_level:
                    stats["total_words"] += ch.word_count

                # 递归遍历子章节
                if ch.children:
                    traverse(ch.children, is_root_level=False)

        traverse(chapter_tree, is_root_level=True)

        # 估算处理成本（基于字数）
        # 假设：1000字 ≈ 1500 tokens ≈ $0.002（GPT-4o-mini）
        stats["estimated_processing_cost"] = (stats["total_words"] / 1000) * 0.002

        return stats

    def get_selected_chapter_content(self, doc_path: str, selected_chapter_ids: List[str]) -> Dict:
        """
        根据用户选择的章节ID，提取对应的文本内容

        Args:
            doc_path: Word 文档路径
            selected_chapter_ids: 选中的章节ID列表，如 ["ch_0", "ch_1_2"]

        Returns:
            {
                "success": True/False,
                "chapters": [
                    {
                        "id": "ch_0",
                        "title": "第一章 项目概述",
                        "content": "完整章节文本内容...",
                        "word_count": 1500
                    },
                    ...
                ],
                "total_words": 8000
            }
        """
        try:
            # 重新解析文档（获取完整章节信息）
            result = self.parse_document_structure(doc_path)
            if not result["success"]:
                return result

            chapters = self._flatten_chapters(result["chapters"])

            # 打开文档
            doc = Document(doc_path)

            # 提取选中章节的内容
            selected_chapters = []
            total_words = 0

            for chapter_dict in chapters:
                if chapter_dict["id"] in selected_chapter_ids:
                    # 提取内容
                    start_idx = chapter_dict["para_start_idx"]
                    end_idx = chapter_dict["para_end_idx"]

                    content_paras = doc.paragraphs[start_idx : end_idx + 1]
                    content = '\n'.join(p.text for p in content_paras)

                    selected_chapters.append({
                        "id": chapter_dict["id"],
                        "title": chapter_dict["title"],
                        "content": content,
                        "word_count": self._calculate_word_count(content)
                    })

                    total_words += selected_chapters[-1]["word_count"]

            self.logger.info(f"提取了 {len(selected_chapters)} 个章节，共 {total_words} 字")

            return {
                "success": True,
                "chapters": selected_chapters,
                "total_words": total_words
            }

        except Exception as e:
            self.logger.error(f"提取章节内容失败: {e}")
            return {
                "success": False,
                "chapters": [],
                "total_words": 0,
                "error": str(e)
            }

    def export_chapter_to_docx(self, doc_path: str, chapter_id: str,
                              output_path: str = None) -> Dict:
        """
        将指定章节导出为独立的Word文档（保留原始格式）

        Args:
            doc_path: 原始Word文档路径
            chapter_id: 章节ID (如 "ch_4")
            output_path: 输出文件路径（可选，默认临时目录）

        Returns:
            {
                "success": True,
                "file_path": "/path/to/exported.docx",
                "chapter_title": "第五部分 响应文件格式",
                "word_count": 1500
            }
        """
        try:
            from docx import Document
            from tempfile import NamedTemporaryFile
            from copy import deepcopy

            # 1. 解析文档结构，定位目标章节
            result = self.parse_document_structure(doc_path)
            if not result["success"]:
                return result

            chapters = self._flatten_chapters(result["chapters"])
            target_chapter = None

            for ch in chapters:
                if ch["id"] == chapter_id:
                    target_chapter = ch
                    break

            if not target_chapter:
                return {
                    "success": False,
                    "error": f"未找到章节ID: {chapter_id}"
                }

            # 2. 打开原始文档
            source_doc = Document(doc_path)

            # 3. 创建新文档
            new_doc = Document()

            # 4. 复制章节内容（保留格式）
            para_start = target_chapter["para_start_idx"]
            para_end = target_chapter.get("para_end_idx")

            if para_end is None:
                para_end = len(source_doc.paragraphs) - 1

            self.logger.info(f"导出章节: {target_chapter['title']}")
            self.logger.info(f"段落范围: {para_start} - {para_end}")

            # 复制段落（使用深拷贝保留格式）
            for i in range(para_start, min(para_end + 1, len(source_doc.paragraphs))):
                source_para = source_doc.paragraphs[i]

                # 使用XML深拷贝（最佳格式保留）
                # 导入段落的完整XML节点
                new_para_element = deepcopy(source_para._element)
                new_doc.element.body.append(new_para_element)

            # 5. 保存到临时文件或指定路径
            if output_path is None:
                # 使用临时文件
                temp_file = NamedTemporaryFile(
                    delete=False,
                    suffix='.docx',
                    prefix='chapter_template_'
                )
                output_path = temp_file.name
                temp_file.close()

            new_doc.save(output_path)

            self.logger.info(f"章节已导出: {output_path}")

            return {
                "success": True,
                "file_path": output_path,
                "chapter_title": target_chapter["title"],
                "word_count": target_chapter.get("word_count", 0),
                "para_count": para_end - para_start + 1
            }

        except Exception as e:
            self.logger.error(f"导出章节失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return {
                "success": False,
                "error": str(e)
            }

    def export_multiple_chapters_to_docx(self, doc_path: str, chapter_ids: List[str], output_path: str = None) -> Dict:
        """
        将多个章节导出为单个Word文档

        Args:
            doc_path: 源文档路径
            chapter_ids: 章节ID列表
            output_path: 输出路径（可选）

        Returns:
            {
                "success": bool,
                "file_path": str,
                "chapter_titles": List[str],
                "chapter_count": int
            }
        """
        from docx import Document
        from tempfile import NamedTemporaryFile
        from copy import deepcopy

        try:
            # 使用智能路径解析（兼容多种环境）
            resolved_path = resolve_file_path(doc_path)
            if not resolved_path:
                return {"success": False, "error": f"文档路径解析失败: {doc_path}"}

            doc_path = str(resolved_path)  # 使用解析后的绝对路径

            # 解析文档结构
            result = self.parse_document_structure(doc_path)
            chapters = self._flatten_chapters(result["chapters"])

            # ⭐ 关键修复：过滤掉父章节已被选中的子章节，避免重复导出
            filtered_chapter_ids = self._filter_redundant_chapters(chapter_ids)
            if len(filtered_chapter_ids) < len(chapter_ids):
                removed_count = len(chapter_ids) - len(filtered_chapter_ids)
                self.logger.info(f"去重完成：移除了 {removed_count} 个冗余子章节")
                self.logger.info(f"  原始章节列表: {chapter_ids}")
                self.logger.info(f"  去重后章节列表: {filtered_chapter_ids}")

            # 按ID查找目标章节并按原文档顺序排序
            target_chapters = []
            for ch in chapters:
                if ch["id"] in filtered_chapter_ids:
                    target_chapters.append(ch)

            if not target_chapters:
                return {"success": False, "error": "未找到指定章节"}

            # 打开源文档
            source_doc = Document(doc_path)
            # 使用源文档作为模板，保留所有样式和页面设置
            new_doc = Document(doc_path)

            # 清空模板文档的所有body内容（段落+表格），保留样式定义、页面设置、页眉页脚等
            for element in list(new_doc.element.body):
                element.getparent().remove(element)

            chapter_titles = []

            # 依次复制每个章节
            for i, chapter in enumerate(target_chapters):
                chapter_titles.append(chapter["title"])

                # 添加章节标题（除第一个章节外，其他章节前加分页符）
                if i > 0:
                    new_doc.add_page_break()

                # 复制章节内容（包括段落和表格）
                para_start = chapter["para_start_idx"]
                para_end = chapter.get("para_end_idx", len(source_doc.paragraphs) - 1)

                # 构建段落索引到body索引的映射
                para_count = 0
                start_body_idx = None
                end_body_idx = None

                for body_idx, element in enumerate(source_doc.element.body):
                    if isinstance(element, CT_P):
                        if para_count == para_start and start_body_idx is None:
                            start_body_idx = body_idx
                        if para_count == para_end:
                            end_body_idx = body_idx
                            break
                        para_count += 1

                # 复制范围内的所有元素（段落+表格）
                if start_body_idx is not None and end_body_idx is not None:
                    for body_idx in range(start_body_idx, end_body_idx + 1):
                        element = source_doc.element.body[body_idx]
                        new_element = deepcopy(element)
                        new_doc.element.body.append(new_element)

            # 保存到临时文件
            if output_path is None:
                temp_file = NamedTemporaryFile(delete=False, suffix='.docx', prefix='chapters_template_')
                output_path = temp_file.name
                temp_file.close()

            new_doc.save(output_path)

            logger.info(f"批量导出成功: {len(target_chapters)}个章节 -> {output_path}")

            return {
                "success": True,
                "file_path": output_path,
                "chapter_titles": chapter_titles,
                "chapter_count": len(target_chapters)
            }

        except Exception as e:
            logger.error(f"批量导出失败: {str(e)}")
            return {"success": False, "error": str(e)}

    def _flatten_chapters(self, chapter_dicts: List[Dict]) -> List[Dict]:
        """将章节树扁平化为列表"""
        flat = []

        def traverse(chapters):
            for ch in chapters:
                flat.append(ch)
                if ch.get("children"):
                    traverse(ch["children"])

        traverse(chapter_dicts)
        return flat

    def _filter_redundant_chapters(self, chapter_ids: List[str]) -> List[str]:
        """
        过滤掉父章节已被选中的子章节，避免重复导出

        例如：如果同时选择了 ch_3 和 ch_3_2，则只保留 ch_3

        Args:
            chapter_ids: 原始章节ID列表

        Returns:
            去重后的章节ID列表
        """
        filtered_ids = []

        for chapter_id in chapter_ids:
            # 检查是否有父章节已在列表中
            has_parent_selected = False

            # 分析章节ID的层级结构（ch_3_2_1 -> ["ch_3", "ch_3_2", "ch_3_2_1"]）
            parts = chapter_id.split('_')
            for i in range(1, len(parts)):
                # 构建可能的父章节ID
                parent_id = '_'.join(parts[:i+1])
                if parent_id != chapter_id and parent_id in chapter_ids:
                    has_parent_selected = True
                    self.logger.debug(f"跳过子章节 {chapter_id}，因为父章节 {parent_id} 已被选中")
                    break

            # 如果没有父章节被选中，保留该章节
            if not has_parent_selected:
                filtered_ids.append(chapter_id)

        return filtered_ids

    # ========================================
        """
        检测连续章节标题列表的完整范围（用于跳过"文件构成说明"等元数据列表）

        策略（基于TOC对比，更智能）：
        1. 从当前段落向前向后扫描短文本（≤50字）
        2. 检查这些文本是否与TOC项匹配
        3. 如果连续3个以上TOC匹配，且标题间无实质内容，判定为元数据列表

        Args:
            doc: Word文档对象
            para_idx: 当前段落索引
            toc_items: 目录项列表（可选），如果提供则基于TOC对比
            toc_end_idx: 目录结束的段落索引（可选），用于限制不扫描目录区域

        Returns:
            (start, end, titles) 或 None
            - start: 列表起始段落索引
            - end: 列表结束段落索引
            - titles: [(idx, text), ...] 列表中的所有章节标题
        """
        # 如果没有TOC，无法进行智能检测
        if not toc_items:
            return None

        # 🔑 扫描范围：当前段落前后，但不扫描目录区域（目录和正文分离）
        if toc_end_idx is not None:
            # 确保扫描起点在目录结束之后（正文区域）
            scan_start = max(toc_end_idx + 1, para_idx - 10)
        else:
            scan_start = max(0, para_idx - 10)
        scan_end = min(len(doc.paragraphs), para_idx + 20)

        # 收集与TOC匹配的段落
        toc_matched_paras = []

        for i in range(scan_start, scan_end):
            para = doc.paragraphs[i]
            text = para.text.strip()

            # 只检查短文本（可能是标题）
            if not text or len(text) > 50:
                continue

            # 清理文本（去除空格、制表符）
            clean_text = re.sub(r'[\s\u3000]+', '', text)

            # 检查是否与任何TOC项匹配
            for toc_item in toc_items:
                toc_title = toc_item['title']
                clean_toc = re.sub(r'[\s\u3000]+', '', toc_title)

                # 🔑 只使用完全匹配（避免误判子标题为元数据）
                # 例如："投标人须知前附表"不应该匹配"投标人须知前附表及投标人须知"
                if clean_text == clean_toc:
                    toc_matched_paras.append((i, text, toc_item))
                    break

        # 至少需要3个连续的TOC匹配
        if len(toc_matched_paras) < 3:
            return None

        # 检查是否连续（允许间隔1-3个段落）
        consecutive_groups = []
        current_group = [toc_matched_paras[0]]

        for j in range(1, len(toc_matched_paras)):
            gap = toc_matched_paras[j][0] - toc_matched_paras[j-1][0]

            if gap <= 4:  # 允许间隔最多3个空段
                current_group.append(toc_matched_paras[j])
            else:
                if len(current_group) >= 3:
                    consecutive_groups.append(current_group)
                current_group = [toc_matched_paras[j]]

        # 检查最后一组
        if len(current_group) >= 3:
            consecutive_groups.append(current_group)

        # 如果没有找到连续组，返回None
        if not consecutive_groups:
            return None

        # 选择包含当前段落的组（如果有多个，选择最大的）
        target_group = None
        for group in consecutive_groups:
            group_start = group[0][0]
            group_end = group[-1][0]
            if group_start <= para_idx <= group_end:
                if target_group is None or len(group) > len(target_group):
                    target_group = group

        # 🔑 关键修改：如果当前段落不在任何组中，返回 None（不是元数据）
        if target_group is None:
            return None

        # 检查该组的标题之间是否无实质内容
        has_substantial_content = False
        for j in range(len(target_group) - 1):
            start_idx = target_group[j][0]
            end_idx = target_group[j + 1][0]

            # 计算两个标题之间的内容字数
            content_chars = sum(
                len(doc.paragraphs[k].text.strip())
                for k in range(start_idx + 1, end_idx)
            )

            # 如果标题间有超过100字，说明不是元数据列表
            if content_chars > 100:
                has_substantial_content = True
                break

        if not has_substantial_content:
            # ⭐ 关键逻辑：即使标题间无实质内容，仍需验证这不是真正章节的开始
            #
            # 检查：如果列表中所有段落都有Heading样式 → 这是真正的章节序列，不是元数据
            # 示例：连续的Heading 1段落 = 真正章节
            #       连续的Normal样式 = 元数据列表（文件构成说明）
            all_headings = all(
                doc.paragraphs[idx].style and
                ('heading' in doc.paragraphs[idx].style.name.lower() or '标题' in doc.paragraphs[idx].style.name.lower())
                for idx, _, _ in target_group
            )

            if all_headings:
                self.logger.info(f"  检测到连续Heading样式章节（段落{target_group[0][0]}-{target_group[-1][0]}），"
                                f"判定为真正章节序列，不是元数据")
                return None  # 不是元数据列表，是真正的章节

            # 如果不是全部Heading，则判定为元数据列表
            # 返回列表范围
            return (
                target_group[0][0],  # start
                target_group[-1][0],  # end
                [(idx, text) for idx, text, _ in target_group]  # titles
            )

        return None
if __name__ == '__main__':
    # 测试代码
    import sys

    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        print("用法: python structure_parser.py <word文档路径>")
        sys.exit(1)

    parser = DocumentStructureParser()
    result = parser.parse_document_structure(doc_path)

    if result["success"]:
        print(f"\n✅ 解析成功！")
        print(f"统计信息: {result['statistics']}")
        print(f"\n章节结构:")

        def print_tree(chapters, indent=0):
            for ch in chapters:
                prefix = "  " * indent
                print(f"{prefix}[{ch['level']}级] {ch['title']} ({ch['word_count']}字)")
                if ch.get("children"):
                    print_tree(ch["children"], indent + 1)

        print_tree(result["chapters"])
    else:
        print(f"\n❌ 解析失败: {result.get('error')}")
