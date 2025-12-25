#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公司信息提取器
从历史标书中智能提取公司基本信息、资质文件等
"""

import re
import io
import base64
import json
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    print("需要安装python-docx: pip install python-docx")
    raise

import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_config, get_module_logger, get_prompt_manager
from common.llm_client import create_llm_client

logger = get_module_logger("company_info.extractor")


class CompanyInfoExtractor:
    """从标书中提取公司信息"""

    # 资质图片识别关键词映射
    QUALIFICATION_IMAGE_KEYWORDS = {
        'business_license': ['营业执照', '工商执照', '统一社会信用代码', '企业法人营业执照'],
        'legal_id_front': ['法定代表人身份证', '法人身份证', '身份证正面'],
        'legal_id_back': ['身份证反面', '身份证背面', '国徽面'],
        'auth_id_front': ['被授权人身份证', '授权代表身份证', '委托代理人身份证'],
        'auth_id_back': ['被授权人身份证反面', '授权代表身份证反面'],
        'iso9001': ['ISO9001', 'ISO 9001', '质量管理体系认证', 'GB/T19001'],
        'iso14001': ['ISO14001', 'ISO 14001', '环境管理体系认证', 'GB/T24001'],
        'iso20000': ['ISO20000', 'ISO 20000', 'IT服务管理体系', 'ISO/IEC 20000'],
        'iso27001': ['ISO27001', 'ISO 27001', '信息安全管理体系', 'ISO/IEC 27001'],
        'cmmi': ['CMMI', '能力成熟度模型', 'CMMI-DEV', 'CMMI3', 'CMMI5'],
        'itss': ['ITSS', '信息技术服务标准', '运维能力成熟度'],
        'value_added_telecom_permit': ['增值电信业务经营许可证', 'ICP许可证', 'ISP许可证', 'B1', 'B2'],
        'basic_telecom_permit': ['基础电信业务经营许可证'],
        'level_protection': ['等级保护', '等保测评', '信息系统安全等级保护', '等保三级'],
        'software_copyright': ['软件著作权', '计算机软件著作权登记证书', '软著'],
        'patent_certificate': ['专利证书', '发明专利', '实用新型专利', '外观设计专利'],
        'high_tech_enterprise': ['高新技术企业证书', '高企证书', '国家高新技术企业'],
        'audit_report': ['审计报告', '财务审计', '年度审计', '会计师事务所'],
        'bank_account_permit': ['开户许可证', '银行开户许可', '基本存款账户'],
        'tax_registration': ['税务登记证', '纳税人资格', '一般纳税人'],
    }

    def __init__(self, model_name: str = "gpt-4o-mini"):
        self.config = get_config()
        self.prompt_manager = get_prompt_manager()
        self.llm_client = create_llm_client(model_name)
        self.model_name = model_name
        logger.info(f"公司信息提取器初始化完成，使用模型: {model_name}")

    def extract_from_tender(self, file_path: str) -> Dict[str, Any]:
        """
        从标书文档中提取公司信息

        Args:
            file_path: Word文档路径

        Returns:
            提取结果字典，包含：
            - success: 是否成功
            - company_info: 公司基本信息
            - authorized_person: 被授权人信息
            - financial_info: 财务信息
            - qualification_images: 识别到的资质图片列表
        """
        try:
            logger.info(f"开始从标书提取公司信息: {file_path}")

            file_path = Path(file_path)
            if not file_path.exists():
                return {'success': False, 'error': f'文件不存在: {file_path}'}

            # 1. 解析文档
            doc = Document(str(file_path))

            # 2. 提取文档全文
            full_text = self._extract_full_text(doc)
            logger.info(f"文档全文提取完成，共 {len(full_text)} 字符")

            # 3. 提取嵌入图片
            images = self._extract_embedded_images(doc)
            logger.info(f"共提取到 {len(images)} 张嵌入图片")

            # 4. 使用LLM提取公司信息
            extracted_info = self._extract_company_info_by_llm(full_text)

            # 5. 智能分类资质图片
            classified_images = self._classify_qualification_images(images, full_text)

            result = {
                'success': True,
                'company_info': extracted_info.get('basic_info', {}),
                'authorized_person': extracted_info.get('authorized_person', {}),
                'financial_info': extracted_info.get('financial_info', {}),
                'qualification_images': classified_images,
                'extraction_time': datetime.now().isoformat(),
                'model_used': self.model_name
            }

            # 统计提取结果
            info_count = sum(1 for v in result['company_info'].values() if v)
            person_count = sum(1 for v in result['authorized_person'].values() if v)
            logger.info(f"提取完成: 基本信息字段={info_count}, 被授权人字段={person_count}, 资质图片={len(classified_images)}")

            return result

        except Exception as e:
            logger.error(f"提取公司信息失败: {e}", exc_info=True)
            return {
                'success': False,
                'error': str(e)
            }

    def _extract_full_text(self, doc: Document) -> str:
        """提取文档完整文本"""
        paragraphs = []

        # 提取段落文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(' | '.join(row_texts))

        return '\n'.join(paragraphs)

    def _extract_embedded_images(self, doc: Document) -> List[Dict]:
        """
        提取Word文档中的嵌入图片

        Returns:
            图片列表，每个图片包含：
            - image_data: base64编码的图片数据
            - content_type: MIME类型
            - filename: 原始文件名
            - size: 文件大小
        """
        images = []

        try:
            # 遍历文档关系获取图片
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref.lower():
                    try:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        content_type = image_part.content_type

                        # 过滤太小的图片（可能是装饰图标）
                        if len(image_data) < 5000:  # 小于5KB的跳过
                            continue

                        # 转为base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')

                        # 提取文件名
                        filename = Path(rel.target_ref).name if rel.target_ref else f"image_{rel_id}.png"

                        images.append({
                            'rel_id': rel_id,
                            'image_data': image_base64,
                            'content_type': content_type,
                            'filename': filename,
                            'size': len(image_data)
                        })

                    except Exception as e:
                        logger.warning(f"提取图片失败 {rel_id}: {e}")

        except Exception as e:
            logger.error(f"提取嵌入图片失败: {e}")

        return images

    def _extract_company_info_by_llm(self, text: str) -> Dict[str, Any]:
        """使用LLM提取公司信息"""

        # 获取提示词
        prompt_template = self.prompt_manager.get_prompt(
            'company_extraction',
            'extract_company_info'
        )

        # 如果提示词文件不存在，使用默认提示词
        if not prompt_template:
            prompt_template = self._get_default_extraction_prompt()

        # 限制文本长度，避免超出token限制
        max_text_length = 15000
        if len(text) > max_text_length:
            # 优先保留前面的内容（通常包含公司基本信息）
            text = text[:max_text_length] + "\n...[文本已截断]..."

        prompt = prompt_template.format(text=text)

        try:
            logger.info("调用LLM提取公司信息...")
            response = self.llm_client.call(
                prompt=prompt,
                temperature=0.3,  # 使用较低温度保证一致性
                max_tokens=2000,
                purpose="提取公司信息"
            )

            # 解析JSON响应
            result = self._parse_llm_response(response)
            logger.info(f"LLM提取完成，解析到 {len(result)} 个顶级字段")
            return result

        except Exception as e:
            logger.error(f"LLM提取失败: {e}")
            return {}

    def _get_default_extraction_prompt(self) -> str:
        """获取默认的提取提示词"""
        return '''请从以下历史标书文档中提取投标公司的信息，以JSON格式返回。

文档内容：
{text}

请提取以下信息并返回JSON格式：

{{
  "basic_info": {{
    "company_name": "公司全称",
    "social_credit_code": "统一社会信用代码（18位）",
    "legal_representative": "法定代表人姓名",
    "legal_representative_position": "法定代表人职务",
    "legal_representative_gender": "法定代表人性别（男/女）",
    "registered_capital": "注册资本",
    "establish_date": "成立日期（YYYY-MM-DD格式）",
    "company_type": "公司类型",
    "registered_address": "注册地址",
    "office_address": "办公地址",
    "business_scope": "经营范围",
    "fixed_phone": "固定电话",
    "fax": "传真号码",
    "postal_code": "邮政编码",
    "email": "电子邮箱",
    "employee_count": "员工人数"
  }},
  "authorized_person": {{
    "name": "被授权人/投标代表姓名",
    "id_number": "身份证号码",
    "gender": "性别",
    "position": "职位",
    "title": "职称",
    "phone": "联系电话"
  }},
  "financial_info": {{
    "bank_name": "开户银行全称",
    "bank_account": "银行账号"
  }}
}}

提取规则：
1. 公司名称通常出现在投标函、法人授权书或公司介绍部分
2. 统一社会信用代码为18位字母数字组合
3. 被授权人信息通常在法人授权委托书中
4. 财务信息可能出现在公司基本情况表或银行资信证明中
5. 如果某字段在文档中未找到，请设为null
6. 只返回JSON，不要添加任何说明文字'''

    def _parse_llm_response(self, response: str) -> Dict:
        """解析LLM响应的JSON"""
        if not response or not response.strip():
            return {}

        try:
            # 清理markdown代码块
            response = re.sub(r'^\s*```json\s*', '', response.strip())
            response = re.sub(r'\s*```\s*$', '', response.strip())

            # 查找JSON对象
            json_start = response.find('{')
            json_end = response.rfind('}') + 1

            if json_start != -1 and json_end > json_start:
                json_str = response[json_start:json_end]
                return json.loads(json_str)

            return {}

        except json.JSONDecodeError as e:
            logger.error(f"JSON解析失败: {e}")
            logger.debug(f"原始响应: {response[:500]}...")
            return {}

    def _classify_qualification_images(self, images: List[Dict],
                                       doc_text: str) -> List[Dict]:
        """
        智能分类资质图片
        基于文档上下文和关键词匹配
        """
        classified = []

        # 将文档文本按段落分割，用于上下文匹配
        text_lower = doc_text.lower()

        for idx, img in enumerate(images):
            # 基于文件名推断类型
            filename = img.get('filename', '').lower()
            guessed_type, confidence = self._guess_image_type(filename, text_lower, idx, len(images))

            classified.append({
                'index': idx,
                'image_data': img['image_data'],
                'content_type': img['content_type'],
                'original_filename': img.get('filename', f'image_{idx}.png'),
                'guessed_type': guessed_type,
                'guessed_type_name': self._get_qualification_type_name(guessed_type),
                'confidence': confidence,
                'size': img.get('size', 0),
                'confirmed': confidence >= 0.6  # 高置信度自动确认
            })

        return classified

    def _guess_image_type(self, filename: str, doc_text: str,
                          img_index: int, total_images: int) -> Tuple[str, float]:
        """
        根据文件名和上下文猜测图片类型

        Returns:
            (资质类型, 置信度)
        """
        # 基于文件名匹配
        for qual_type, keywords in self.QUALIFICATION_IMAGE_KEYWORDS.items():
            for keyword in keywords:
                if keyword.lower() in filename:
                    return qual_type, 0.9

        # 基于文档中关键词出现频率推断
        # 统计各类资质关键词在文档中的出现次数
        keyword_counts = {}
        for qual_type, keywords in self.QUALIFICATION_IMAGE_KEYWORDS.items():
            count = 0
            for keyword in keywords:
                count += doc_text.count(keyword.lower())
            if count > 0:
                keyword_counts[qual_type] = count

        # 如果文档中有明显的资质关键词，根据图片位置推断
        if keyword_counts:
            # 按出现次数排序
            sorted_types = sorted(keyword_counts.items(), key=lambda x: x[1], reverse=True)

            # 根据图片在文档中的位置，匹配对应的资质类型
            if img_index < len(sorted_types):
                return sorted_types[img_index][0], 0.5

        # 基于图片位置的启发式规则
        if total_images > 0:
            position_ratio = img_index / total_images

            # 第一张图片通常是营业执照
            if img_index == 0:
                return 'business_license', 0.6
            # 前几张可能是身份证
            elif img_index <= 2:
                return 'legal_id_front', 0.4

        return 'unknown', 0.2

    def _get_qualification_type_name(self, type_key: str) -> str:
        """获取资质类型的中文名称"""
        type_names = {
            'business_license': '营业执照',
            'legal_id_front': '法人身份证正面',
            'legal_id_back': '法人身份证反面',
            'auth_id_front': '被授权人身份证正面',
            'auth_id_back': '被授权人身份证反面',
            'iso9001': 'ISO9001质量管理体系认证',
            'iso14001': 'ISO14001环境管理体系认证',
            'iso20000': 'ISO20000信息技术服务管理认证',
            'iso27001': 'ISO27001信息安全管理体系认证',
            'cmmi': 'CMMI认证',
            'itss': 'ITSS信息技术服务标准认证',
            'value_added_telecom_permit': '增值电信业务经营许可证',
            'basic_telecom_permit': '基础电信业务经营许可证',
            'level_protection': '等级保护认证',
            'software_copyright': '软件著作权',
            'patent_certificate': '专利证书',
            'high_tech_enterprise': '高新技术企业证书',
            'audit_report': '审计报告',
            'bank_account_permit': '开户许可证',
            'tax_registration': '税务登记证/纳税人资格',
            'unknown': '未识别',
        }
        return type_names.get(type_key, type_key)


# 便捷函数
def extract_company_from_tender(file_path: str, model_name: str = "gpt-4o-mini") -> Dict[str, Any]:
    """
    从标书中提取公司信息的便捷函数

    Args:
        file_path: Word文档路径
        model_name: 使用的LLM模型名称

    Returns:
        提取结果字典
    """
    extractor = CompanyInfoExtractor(model_name=model_name)
    return extractor.extract_from_tender(file_path)
