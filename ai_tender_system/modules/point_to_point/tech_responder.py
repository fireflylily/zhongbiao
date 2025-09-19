#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
技术需求回复模块 - 处理招标文件中的技术需求点对点回复
根据技术需求自动生成技术响应文档
"""

import json
import re
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
import requests

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import (
    get_config, get_module_logger,
    APIError, FileProcessingError,
    ensure_dir
)
from common.llm_client import create_llm_client

# 专业提示词模板 - 与enhanced_inline_reply.py保持完全一致
PROMPT_ANSWER = """现在有个问答，比选文件要求和比选申请人应答，我给你举个例子，比如比选文件要求：支持可视化创建不同类型数据源，包括但不限于：传统数据库、文件系统、消息队列、SaaS API，NoSQL等、必选申请人回答的是：应答：满足。。系统支持数据源配置化管理，数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置管理功能，支持搜索功能。你学习一下我的风格。现在我是比选申请人，请严格按照我的风格来回答，请注意我回答的格式：首先是'应答：满足。'，然后说'系统支持什么什么'，这个过程需要你按照问题回答，不要跑题。以下是输入文字："""

PROMPT_CONTENT = """你是一个大数据平台的专业产品售前，请针对这一需求给出800字的产品功能介绍，不要开头和总结，直接写产品功能，不需要用markdown格式，直接文本格式+特殊项目符号输出即可，需求如下："""

PROMPT_TITLE = """你是一个专业作者，请把以下这段文字变为10字以内不带细节内容和标点和解释的文字，直接给出结果不要'简化为'这种返回："""

class TechResponder:
    """技术需求回复处理器"""
    
    def __init__(self, api_key: Optional[str] = None, model_name: str = "gpt-4o-mini"):
        self.config = get_config()
        self.logger = get_module_logger("tech_responder")

        # 创建LLM客户端
        self.llm_client = create_llm_client(model_name, api_key)

        # 保持向后兼容性的配置
        api_config = self.config.get_api_config()
        self.api_key = api_key or api_config['api_key']
        self.api_endpoint = api_config.get('api_endpoint', 'https://api.openai.com/v1/chat/completions')
        self.model_name = model_name

        self.logger.info(f"技术需求回复处理器初始化完成，使用模型: {model_name}")

    def copy_paragraph_format(self, source_para, target_para):
        """
        复制段落格式，确保字体一致性 - 与enhanced_inline_reply.py保持一致
        """
        try:
            # 复制段落级别的格式
            if source_para.paragraph_format.alignment is not None:
                target_para.paragraph_format.alignment = source_para.paragraph_format.alignment

            if source_para.paragraph_format.left_indent is not None:
                target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent

            if source_para.paragraph_format.right_indent is not None:
                target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent

            if source_para.paragraph_format.first_line_indent is not None:
                target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent

            # 复制字体格式（从源段落的第一个run）
            if source_para.runs:
                source_run = source_para.runs[0]
                if target_para.runs:
                    target_run = target_para.runs[0]

                    # 复制字体属性
                    if source_run.font.name:
                        target_run.font.name = source_run.font.name
                    if source_run.font.size:
                        target_run.font.size = source_run.font.size
                    if source_run.font.bold is not None:
                        target_run.font.bold = source_run.font.bold
                    if source_run.font.italic is not None:
                        target_run.font.italic = source_run.font.italic
                    # 对于应答内容，始终设置为黑色字体，不继承原段落颜色
                    target_run.font.color.rgb = RGBColor(0, 0, 0)

        except Exception as e:
            self.logger.warning(f"格式复制失败: {e}")

    def add_paragraph_shading(self, paragraph, color):
        """
        给段落添加背景色底纹 - 与enhanced_inline_reply.py保持一致
        """
        try:
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml

            # 创建底纹XML元素 - RGB(217,217,217) 转换为十六进制 D9D9D9
            hex_color = f"{color.rgb:06X}" if hasattr(color, 'rgb') else "D9D9D9"
            shading_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
            shading_element = parse_xml(shading_xml)

            # 获取段落属性
            p_pr = paragraph._p.get_or_add_pPr()

            # 添加底纹元素
            p_pr.append(shading_element)

        except Exception as e:
            self.logger.warning(f"添加段落底纹失败: {e}")

    def insert_paragraph_after_with_format(self, paragraph, text=None, style=None):
        """
        在指定段落后插入新段落，并保持格式一致 - 与enhanced_inline_reply.py保持一致
        """
        try:
            # 插入新段落
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = paragraph._parent.add_paragraph()
            new_para._p = new_p

            if text:
                run = new_para.add_run(text)
                # 设置字体为黑色
                run.font.color.rgb = RGBColor(0, 0, 0)

            if style:
                new_para.style = style

            # 复制源段落的格式
            self.copy_paragraph_format(paragraph, new_para)

            # 设置1.5倍行距
            new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            new_para.paragraph_format.line_spacing = 1.5

            # 添加浅灰色底纹 RGB(217,217,217)
            self.add_paragraph_shading(new_para, RGBColor(217, 217, 217))

            return new_para

        except Exception as e:
            self.logger.error(f"段落插入失败: {e}")
            return None

    def process_tech_requirements(self,
                                 requirements_file: str,
                                 output_file: str,
                                 company_info: Dict[str, Any],
                                 response_strategy: str = "comprehensive",
                                 response_frequency: str = "every_paragraph",
                                 response_mode: str = "simple",
                                 ai_model: str = "gpt-4o-mini",
                                 output_mode: str = "document") -> Dict[str, Any]:
        """
        处理技术需求并生成回复文档

        Args:
            requirements_file: 技术需求文档路径
            output_file: 输出文档路径
            company_info: 公司信息
            response_strategy: 回复策略 (comprehensive/concise/detailed)
            response_frequency: 应答频次 (every_paragraph/major_headings)
            response_mode: 应答方式 (simple/ai)
            ai_model: AI模型 (gpt-4o-mini/unicom-yuanjing)
            output_mode: 输出模式 (document/inline) - 新增参数

        Returns:
            处理结果统计
        """
        try:
            self.logger.info(f"开始处理技术需求文档: {requirements_file}")
            self.logger.info(f"配置参数 - 应答频次: {response_frequency}, 应答方式: {response_mode}, AI模型: {ai_model}, 输出模式: {output_mode}")

            # 根据输出模式选择处理方式
            if output_mode == "inline":
                # 使用原地插入模式
                return self.process_inline_requirements(
                    requirements_file=requirements_file,
                    output_file=output_file,
                    company_info=company_info,
                    ai_model=ai_model
                )
            else:
                # 使用独立文档生成模式（原有逻辑）
                # 第1步：提取技术需求
                requirements = self._extract_requirements(requirements_file, response_frequency)
                self.logger.info(f"提取到{len(requirements)}个技术需求")

                # 第2步：生成技术响应
                responses = self._generate_responses(requirements, company_info, response_strategy, response_mode, ai_model)

                # 第3步：创建响应文档
                self._create_response_document(responses, output_file, company_info)

                stats = {
                    'success': True,
                    'requirements_count': len(requirements),
                    'responses_count': len(responses),
                    'output_file': output_file,
                    'message': f'成功处理{len(requirements)}个技术需求'
                }

                self.logger.info(f"技术需求处理完成: {stats['message']}")
                return stats
            
        except Exception as e:
            self.logger.error(f"技术需求处理失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': '处理失败'
            }
    
    def _extract_requirements(self, requirements_file: str, response_frequency: str = "every_paragraph") -> List[Dict[str, Any]]:
        """从文档中提取技术需求"""
        requirements = []
        
        try:
            doc = Document(requirements_file)
            current_section = ""
            requirement_id = 1
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 识别章节标题
                if self._is_section_title(text):
                    current_section = text
                    continue
                
                # 识别技术需求点
                if self._is_requirement(text):
                    # 根据应答频次过滤
                    if response_frequency == "major_headings":
                        # 只对大标题条目应答，跳过段落级别的需求
                        if not self._is_major_heading(text):
                            continue

                    requirement = {
                        'id': requirement_id,
                        'section': current_section,
                        'content': text,
                        'type': self._classify_requirement(text)
                    }
                    requirements.append(requirement)
                    requirement_id += 1
                    self.logger.debug(f"提取需求#{requirement_id}: {text[:50]}...")
            
            # 也从表格中提取需求
            for table in doc.tables:
                table_requirements = self._extract_table_requirements(table)
                for req in table_requirements:
                    req['id'] = requirement_id
                    requirements.append(req)
                    requirement_id += 1
            
        except Exception as e:
            self.logger.error(f"提取需求失败: {e}")
            raise
        
        return requirements

    def is_requirement_paragraph_enhanced(self, para) -> bool:
        """
        优化的需求段落识别 - 与enhanced_inline_reply.py保持一致，减少遗漏
        """
        text = para.text.strip()
        if not text or len(text) < 10:  # 过短的文本不处理
            return False

        style_name = para.style.name if para.style else ""
        alignment = para.paragraph_format.alignment

        # 明确排除的条件
        exclusions = [
            "Heading" in style_name,     # 标题样式
            "标题" in style_name,         # 中文标题样式
            "TOC" in style_name,         # 目录样式
            "目录" in text,               # 包含"目录"
            "封面" in text,               # 封面内容
            "签字" in text,               # 签字页
            alignment == 1,               # 居中对齐（通常是标题）
            text.startswith("第") and ("章" in text or "节" in text),  # 章节标题
            re.match(r'^[一二三四五六七八九十]+[、．]', text),  # 中文数字章节
        ]

        if any(exclusions):
            return False

        # 增强的关键词识别 - 更宽泛的识别策略
        enhanced_keywords = [
            "要求", "需求", "应", "必须", "应当", "需要", "具备", "支持", "提供",
            "实现", "满足", "符合", "遵循", "不少于", "不低于", "负责", "确保",
            "保证", "完成", "达到", "实施", "部署", "维护", "服务", "管理"
        ]

        # 编号模式识别 - 更宽泛的编号格式
        enhanced_patterns = [
            r'^(\d+)\s*[、．.]',         # 1、 1. 1．
            r'^(\d+\.\d+)\s*[、．.]',     # 1.1、 1.2.
            r'^\((\d+)\)',               # (1) (2)
            r'^([A-Z])\)',               # A) B)
            r'^([a-z])\)',               # a) b)
            r'^（[一二三四五六七八九十]+）'  # （一）（二）
        ]

        # 章节标识
        section_identifiers = [
            "乙方", "甲方", "供应商", "采购方", "服务商", "承包方"
        ]

        # 包含条件：更宽泛的识别策略
        inclusions = [
            # 包含需求关键词
            any(keyword in text for keyword in enhanced_keywords),
            # 包含编号格式
            any(re.match(pattern, text) for pattern in enhanced_patterns),
            # 包含主体标识
            any(entity in text for entity in section_identifiers),
            # 长段落且包含具体要求（超过50字且包含技术词汇）
            (len(text) > 50 and any(word in text for word in ["系统", "数据", "接口", "服务", "技术", "平台", "软件", "硬件"])),
        ]

        return any(inclusions)

    def _is_section_title(self, text: str) -> bool:
        """判断是否为章节标题"""
        patterns = [
            r'^\d+\.\s+',  # 1. 标题
            r'^第[一二三四五六七八九十]+[章节部分]',  # 第一章
            r'^[一二三四五六七八九十]+、',  # 一、标题
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        # 检查是否包含关键词
        title_keywords = ['技术要求', '功能要求', '性能要求', '安全要求', '接口要求']
        return any(keyword in text for keyword in title_keywords)
    
    def _is_requirement(self, text: str) -> bool:
        """判断是否为技术需求"""
        # 需求通常包含的关键词
        requirement_keywords = [
            '应', '必须', '需要', '要求', '支持', '提供', '实现',
            '具备', '满足', '确保', '保证', '能够', '可以'
        ]
        
        return any(keyword in text for keyword in requirement_keywords)
    
    def _classify_requirement(self, text: str) -> str:
        """对需求进行分类"""
        if any(word in text for word in ['功能', '模块', '组件']):
            return 'functional'
        elif any(word in text for word in ['性能', '响应', '并发', '吞吐']):
            return 'performance'
        elif any(word in text for word in ['安全', '权限', '加密', '认证']):
            return 'security'
        elif any(word in text for word in ['接口', 'API', '集成', '对接']):
            return 'interface'
        elif any(word in text for word in ['数据', '存储', '备份', '恢复']):
            return 'data'
        else:
            return 'general'
    
    def _extract_table_requirements(self, table) -> List[Dict[str, Any]]:
        """从表格中提取需求"""
        requirements = []
        
        # 分析表格结构
        if len(table.rows) < 2:
            return requirements
        
        # 假设第一行是表头
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # 查找需求相关的列
        requirement_col_idx = -1
        for idx, header in enumerate(headers):
            if any(keyword in header for keyword in ['需求', '要求', '功能', '描述']):
                requirement_col_idx = idx
                break
        
        if requirement_col_idx >= 0:
            for row in table.rows[1:]:
                if requirement_col_idx < len(row.cells):
                    req_text = row.cells[requirement_col_idx].text.strip()
                    if req_text and self._is_requirement(req_text):
                        requirements.append({
                            'section': '表格需求',
                            'content': req_text,
                            'type': self._classify_requirement(req_text)
                        })
        
        return requirements

    def get_inline_system_prompt(self, prompt_type: str = "point_to_point") -> str:
        """
        获取专业的内联应答系统提示词 - 与enhanced_inline_reply.py保持一致
        """
        if prompt_type == "point_to_point":
            return """你是一名专业的投标文件撰写专家，专门负责点对点应答。

你的任务：
1. 以"应答：满足。"开头
2. 然后说"系统支持..."
3. 不要跑题，针对具体需求回答
4. 语言专业、简洁

格式示例：
"应答：满足。系统支持数据源配置化管理，数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置管理功能，全面支持搜索功能。"""
        else:
            # 默认系统提示词
            return """你是一名资深的技术方案专家和投标文件撰写专家，专门负责为采购需求提供专业的技术应答。

你的任务是：
1. 仔细分析每个采购需求的核心要点
2. 生成专业、具体、可信的技术应答
3. 应答必须以"应答：满足。"开头
4. 后续内容要体现专业性，避免空洞的表述

应答原则：
- 专业性：使用行业标准术语，体现技术实力
- 具体性：提及具体的技术方案、产品型号、服务标准
- 可信性：承诺要可实现，避免夸大其词
- 简洁性：控制在80-150字，重点突出
- 规范性：语言正式，符合商务文档标准

应答格式示例：
"应答：满足。我方采用主流的XXX技术架构，配备专业的XXX团队，严格按照XXX标准执行，确保XXX指标达到XXX水平。"

请根据具体需求内容，生成相应的专业技术应答。"""

    def generate_inline_response(self, requirement_text: str, company_info: Dict[str, Any]) -> str:
        """
        生成专业的内联应答 - 使用元景大模型和专业提示词
        """
        try:
            # 使用与enhanced_inline_reply.py相同的提示词格式
            prompt = f"{PROMPT_ANSWER}'{requirement_text}'"

            # 使用统一的LLM客户端调用元景大模型
            response = self.llm_client.call(
                prompt=prompt,
                system_prompt=self.get_inline_system_prompt("point_to_point"),
                temperature=0.3,  # 保持与enhanced_inline_reply.py一致的温度设置
                purpose="点对点应答生成"
            )

            # 确保以"应答：满足。"开头，与enhanced_inline_reply.py保持一致
            if not response.startswith("应答：满足。"):
                if "应答：满足。" in response:
                    # 找到应答：满足。并移到开头
                    response = "应答：满足。" + response.split("应答：满足。")[1]
                else:
                    response = f"应答：满足。{response}"

            self.logger.info(f"生成内联应答: {response[:50]}...")
            return response

        except Exception as e:
            self.logger.error(f"生成内联应答失败: {e}")
            # 使用备用应答模板，与enhanced_inline_reply.py风格一致
            return "应答：满足。我方具备完整的技术实力和丰富的项目经验，将严格按照采购要求提供专业的技术方案和优质服务。"

    def classify_requirement_type_enhanced(self, text: str) -> str:
        """
        增强的需求类型分类 - 与enhanced_inline_reply.py保持一致
        """
        content_lower = text.lower()

        # 数据服务相关
        if any(kw in content_lower for kw in ["数据", "查询", "接口", "api", "专线"]):
            return "数据服务"

        # 系统对接相关
        if any(kw in content_lower for kw in ["对接", "联调", "测试", "部署", "集成"]):
            return "接口对接"

        # 运维服务相关
        if any(kw in content_lower for kw in ["运维", "维护", "监控", "故障", "响应"]):
            return "运维支持"

        # 硬件相关
        if any(kw in content_lower for kw in ["cpu", "内存", "存储", "硬盘", "服务器", "设备"]):
            return "硬件配置"

        # 软件相关
        if any(kw in content_lower for kw in ["软件", "系统", "应用", "功能", "模块"]):
            return "软件功能"

        # 性能相关
        if any(kw in content_lower for kw in ["性能", "速度", "并发", "响应时间", "不少于", "不低于"]):
            return "性能指标"

        # 服务相关
        if any(kw in content_lower for kw in ["服务", "支持", "培训", "实施", "咨询"]):
            return "服务保障"

        # 资质相关
        if any(kw in content_lower for kw in ["资质", "证书", "认证", "经验", "案例"]):
            return "资质证明"

        return "通用模板"

    def _generate_responses(self, requirements: List[Dict[str, Any]],
                          company_info: Dict[str, Any],
                          strategy: str,
                          response_mode: str = "simple",
                          ai_model: str = "gpt-4o-mini") -> List[Dict[str, Any]]:
        """生成技术响应"""
        responses = []
        
        for requirement in requirements:
            self.logger.info(f"生成响应 #{requirement['id']}: {requirement['content'][:50]}...")
            
            # 根据需求类型选择响应模板
            response_template = self._get_response_template(requirement['type'])
            
            # 生成具体响应
            if response_mode == "ai" and self.api_key:
                # 使用AI生成响应
                response_text = self._generate_ai_response(requirement, company_info, strategy, ai_model)
            else:
                # 使用模板生成响应
                response_text = self._generate_template_response(requirement, response_template, company_info)
            
            response = {
                'requirement_id': requirement['id'],
                'requirement': requirement['content'],
                'response': response_text,
                'type': requirement['type'],
                'section': requirement['section']
            }
            
            responses.append(response)
        
        return responses
    
    def _get_response_template(self, requirement_type: str) -> str:
        """获取响应模板"""
        templates = {
            'functional': """
我公司完全理解并响应该技术需求。我们的解决方案将：
1. 全面实现所要求的功能
2. 采用成熟稳定的技术架构
3. 确保系统的可扩展性和可维护性
具体实现方案：{details}
            """,
            'performance': """
我公司充分理解性能要求的重要性，承诺：
1. 系统响应时间满足要求
2. 支持并发用户数达到指定要求
3. 提供性能监控和优化方案
技术保障措施：{details}
            """,
            'security': """
我公司高度重视系统安全，将采取以下措施：
1. 实施多层次安全防护体系
2. 采用行业标准的加密算法
3. 建立完善的权限管理机制
安全保障方案：{details}
            """,
            'interface': """
针对接口集成需求，我公司方案包括：
1. 提供标准化的API接口
2. 支持多种数据交换格式
3. 确保接口的稳定性和兼容性
接口设计方案：{details}
            """,
            'data': """
关于数据管理需求，我公司将：
1. 建立完善的数据管理体系
2. 实施数据备份和恢复策略
3. 确保数据的安全性和完整性
数据管理方案：{details}
            """,
            'general': """
我公司完全响应该项技术需求：
1. 严格按照要求实施
2. 提供专业的技术支持
3. 确保达到预期效果
实施方案：{details}
            """
        }
        
        return templates.get(requirement_type, templates['general'])

    def _is_major_heading(self, text: str) -> bool:
        """判断文本是否为大标题条目"""
        # 大标题的特征：
        # 1. 以数字开头（如 1. 2. 3.）
        # 2. 以中文数字开头（如 一、二、三、）
        # 3. 长度相对较短
        # 4. 不包含详细描述性语言

        # 检查是否以数字或中文数字开头
        major_heading_patterns = [
            r'^\d+[\.\uff0e]',  # 1. 2. 3.
            r'^[一二三四五六七八九十]+[\u3001\uff0c]',  # 一、二、三、
            r'^\(\d+\)',  # (1) (2) (3)
            r'^\d+\)',  # 1) 2) 3)
            r'^第[一二三四五六七八九十\d]+[章节条款项]',  # 第一章 第二节
        ]

        for pattern in major_heading_patterns:
            if re.match(pattern, text):
                return True

        # 检查长度和内容特征
        if len(text) <= 50 and not any(keyword in text for keyword in ['应', '须', '必须', '要求', '不得', '应当']):
            return True

        return False

    def _get_model_config(self, ai_model: str) -> Dict[str, str]:
        """获取指定AI模型的配置"""
        # 使用配置系统获取模型配置
        model_config = self.config.get_model_config(ai_model)

        return {
            'endpoint': model_config.get('api_endpoint', self.api_endpoint),
            'api_key': model_config.get('api_key', self.api_key),
            'model_name': model_config.get('model_name', ai_model),
            'max_tokens': model_config.get('max_tokens', 500),
            'timeout': model_config.get('timeout', 30)
        }

    def _generate_ai_response(self, requirement: Dict[str, Any],
                            company_info: Dict[str, Any],
                            strategy: str,
                            ai_model: str = "gpt-4o-mini") -> str:
        """使用AI生成响应 - 使用统一的LLM客户端"""
        try:
            # 创建指定模型的LLM客户端
            llm_client = create_llm_client(ai_model)

            system_prompt = "你是一个专业的技术方案撰写专家，擅长为企业撰写技术响应文档。"

            prompt = f"""
作为{company_info.get('companyName', '投标方')}的技术专家，请针对以下技术需求生成专业的响应：

技术需求：{requirement['content']}
需求类型：{requirement['type']}
响应策略：{strategy}

要求：
1. 响应要专业、具体、可行
2. 突出公司的技术优势
3. 明确承诺满足需求
4. 提供具体的实现方案
5. 字数控制在200-300字

请生成响应：
            """

            response_content = llm_client.call(
                prompt=prompt,
                system_prompt=system_prompt,
                temperature=0.7,
                purpose=f"生成技术响应 - {requirement['type']}"
            )

            return response_content

        except Exception as e:
            self.logger.error(f"AI生成响应失败: {e}")
            # 回退到模板响应
            return self._generate_template_response(requirement,
                                                   self._get_response_template(requirement['type']),
                                                   company_info)
    
    def _generate_template_response(self, requirement: Dict[str, Any],
                                   template: str,
                                   company_info: Dict[str, Any]) -> str:
        """使用模板生成响应"""
        # 根据需求内容生成具体细节
        details = f"针对'{requirement['content'][:50]}...'的要求，我们将提供完整的解决方案，确保满足所有技术指标。"
        
        response = template.format(details=details)
        
        # 添加公司名称
        response = response.replace('我公司', company_info.get('companyName', '我公司'))
        
        return response.strip()
    
    def _create_response_document(self, responses: List[Dict[str, Any]],
                                 output_file: str,
                                 company_info: Dict[str, Any]):
        """创建响应文档"""
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('技术需求响应文档', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加公司信息
        doc.add_paragraph(f"投标单位：{company_info.get('companyName', '')}")
        doc.add_paragraph(f"联系电话：{company_info.get('phone', '')}")
        doc.add_paragraph(f"电子邮箱：{company_info.get('email', '')}")
        doc.add_paragraph()
        
        # 按章节组织响应
        current_section = ""
        for response in responses:
            # 如果进入新章节，添加章节标题
            if response['section'] != current_section:
                current_section = response['section']
                if current_section:
                    doc.add_heading(current_section, 1)
            
            # 添加需求
            req_para = doc.add_paragraph()
            req_para.add_run(f"需求{response['requirement_id']}：").bold = True
            req_para.add_run(response['requirement'])
            
            # 添加响应
            resp_para = doc.add_paragraph()
            resp_para.add_run("响应：").bold = True
            resp_para.add_run(response['response'])
            
            # 添加分隔
            doc.add_paragraph()
        
        # 保存文档
        doc.save(output_file)
        self.logger.info(f"响应文档已保存: {output_file}")

    def process_inline_requirements(self,
                                   requirements_file: str,
                                   output_file: str = None,
                                   company_info: Dict[str, Any] = None,
                                   ai_model: str = "unicom-yuanjing") -> Dict[str, Any]:
        """
        原地插入应答处理 - 与enhanced_inline_reply.py功能完全一致
        """
        from datetime import datetime

        try:
            self.logger.info(f"开始原地插入处理: {requirements_file}")
            self.logger.info(f"使用AI模型: {ai_model}")

            # 读取文档
            doc = Document(requirements_file)
            self.logger.info(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")

            # 统计需求条目
            requirement_count = 0
            processed_count = 0

            # 从后往前遍历，避免插入新段落影响索引
            paragraphs = list(doc.paragraphs)

            for i in range(len(paragraphs) - 1, -1, -1):
                para = paragraphs[i]

                if self.is_requirement_paragraph_enhanced(para):
                    requirement_count += 1
                    text = para.text.strip()

                    self.logger.info(f"处理需求 {requirement_count}: {text[:60]}...")

                    # 生成专业应答
                    response = self.generate_inline_response(text, company_info or {})

                    # 在需求段落后插入应答（保持格式一致）
                    reply_para = self.insert_paragraph_after_with_format(para, response)

                    if reply_para:
                        processed_count += 1
                        self.logger.info(f"已插入应答: {response[:60]}...")
                    else:
                        self.logger.error(f"插入应答失败: {text[:30]}...")

            # 生成输出文件名
            if not output_file:
                import os
                base_name = os.path.splitext(os.path.basename(requirements_file))[0]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_file = f"{base_name}-点对点应答-{timestamp}.docx"

            # 保存文档
            doc.save(output_file)

            stats = {
                'success': True,
                'requirements_count': requirement_count,
                'responses_count': processed_count,
                'output_file': output_file,
                'message': f'成功处理{requirement_count}个需求条目，插入{processed_count}个应答'
            }

            self.logger.info(f"原地插入处理完成:")
            self.logger.info(f"  识别需求条目: {requirement_count}")
            self.logger.info(f"  成功插入应答: {processed_count}")
            self.logger.info(f"  输出文件: {output_file}")

            return stats

        except Exception as e:
            self.logger.error(f"原地插入处理失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': '原地插入处理失败'
            }