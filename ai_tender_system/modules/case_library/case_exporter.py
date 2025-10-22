#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
案例库导出器
功能：将多个案例导出为Word文档
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from common.logger import get_module_logger
from common.config import get_config

logger = get_module_logger("case_library.case_exporter")


class CaseExporter:
    """案例库导出器"""

    def __init__(self):
        """初始化案例导出器"""
        self.config = get_config()
        self.logger = logger
        self.logger.info("案例导出器初始化完成")

    def export_cases_to_word(
        self,
        cases: List[Dict[str, Any]],
        output_path: str,
        options: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        导出案例列表为Word文档

        Args:
            cases: 案例数据列表
            output_path: 输出文件路径
            options: 导出选项
                - title: 文档标题（默认：案例汇总文档）
                - include_attachments: 是否包含附件信息（默认：True）
                - include_statistics: 是否包含统计信息（默认：True）

        Returns:
            生成的文件路径
        """
        try:
            self.logger.info(f"开始导出 {len(cases)} 个案例到: {output_path}")

            if not cases:
                raise ValueError("案例列表不能为空")

            # 处理选项
            options = options or {}
            doc_title = options.get('title', '案例汇总文档')
            include_attachments = options.get('include_attachments', True)
            include_statistics = options.get('include_statistics', True)

            # 创建Word文档
            doc = Document()

            # 设置文档样式
            self._setup_document_styles(doc)

            # 添加封面
            self._add_cover_page(doc, doc_title, len(cases))

            # 添加目录提示（Word需手动更新域代码）
            self._add_toc_placeholder(doc)

            # 添加每个案例
            for i, case in enumerate(cases, start=1):
                self._add_case_section(doc, case, i, include_attachments)

            # 添加统计汇总
            if include_statistics:
                self._add_statistics_section(doc, cases)

            # 保存文档
            doc.save(output_path)

            self.logger.info(f"案例导出完成: {output_path}")
            return output_path

        except Exception as e:
            self.logger.error(f"导出案例失败: {e}", exc_info=True)
            raise

    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置段落间距
        style.paragraph_format.line_spacing = 1.5

    def _add_cover_page(self, doc: Document, title: str, case_count: int):
        """添加封面"""
        # 标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 设置标题字体
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(24)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 添加空行
        doc.add_paragraph()
        doc.add_paragraph()

        # 文档信息
        info_para = doc.add_paragraph()
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info_para.add_run(f"案例数量：{case_count}\n").font.size = Pt(14)
        info_para.add_run(f"导出时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}\n").font.size = Pt(14)

        # 分页
        doc.add_page_break()

    def _add_toc_placeholder(self, doc: Document):
        """添加目录占位符"""
        doc.add_heading('目录', level=1)
        toc_para = doc.add_paragraph()
        toc_para.add_run('【在Word中右键点击此处，选择"更新域"以生成完整目录】')
        toc_para.runs[0].font.italic = True
        toc_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # 添加分页
        doc.add_page_break()

    def _add_case_section(self, doc: Document, case: Dict, index: int, include_attachments: bool):
        """添加单个案例章节"""
        # 案例标题
        case_title = case.get('case_title', f'案例 {index}')
        doc.add_heading(f"{index}. {case_title}", level=1)

        # 基本信息表格
        self._add_basic_info_table(doc, case)

        # 合同信息表格
        self._add_contract_info_table(doc, case)

        # 甲乙方信息
        self._add_party_info_tables(doc, case)

        # 附件信息（如果有）
        if include_attachments and case.get('attachments'):
            self._add_attachments_section(doc, case['attachments'])

        # 章节结束后添加分页
        doc.add_page_break()

    def _add_basic_info_table(self, doc: Document, case: Dict):
        """添加基本信息表格"""
        doc.add_heading('基本信息', level=2)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # 填充数据
        cells = table.rows[0].cells
        cells[0].text = '案例标题'
        cells[1].text = case.get('case_title', '')

        cells = table.rows[1].cells
        cells[0].text = '案例编号'
        cells[1].text = case.get('case_number', '')

        cells = table.rows[2].cells
        cells[0].text = '客户名称'
        cells[1].text = case.get('customer_name', '')

        cells = table.rows[3].cells
        cells[0].text = '所属行业'
        cells[1].text = case.get('industry', '')

        cells = table.rows[4].cells
        cells[0].text = '案例状态'
        cells[1].text = self._get_status_text(case.get('case_status', ''))

        # 设置列宽
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)

        doc.add_paragraph()  # 空行

    def _add_contract_info_table(self, doc: Document, case: Dict):
        """添加合同信息表格"""
        doc.add_heading('合同信息', level=2)

        # 计算行数
        row_count = 5
        if case.get('contract_type') == '订单' and case.get('final_customer_name'):
            row_count = 6

        table = doc.add_table(rows=row_count, cols=2)
        table.style = 'Light Grid Accent 1'

        row_idx = 0

        # 合同类型
        cells = table.rows[row_idx].cells
        cells[0].text = '合同类型'
        cells[1].text = case.get('contract_type', '')
        row_idx += 1

        # 最终客户（仅订单类型）
        if case.get('contract_type') == '订单' and case.get('final_customer_name'):
            cells = table.rows[row_idx].cells
            cells[0].text = '最终客户'
            cells[1].text = case.get('final_customer_name', '')
            row_idx += 1

        # 合同金额
        cells = table.rows[row_idx].cells
        cells[0].text = '合同金额'
        amount = case.get('contract_amount', '')
        if amount:
            # 如果是纯数字，添加"万元"单位
            try:
                float_amount = float(amount)
                cells[1].text = f"{amount} 万元"
            except (ValueError, TypeError):
                cells[1].text = str(amount)
        else:
            cells[1].text = ''
        row_idx += 1

        # 开始日期
        cells = table.rows[row_idx].cells
        cells[0].text = '合同开始日期'
        cells[1].text = case.get('contract_start_date', '')
        row_idx += 1

        # 结束日期
        cells = table.rows[row_idx].cells
        cells[0].text = '合同结束日期'
        cells[1].text = case.get('contract_end_date', '')
        row_idx += 1

        # 设置列宽
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)

        doc.add_paragraph()  # 空行

    def _add_party_info_tables(self, doc: Document, case: Dict):
        """添加甲乙方信息表格"""
        # 甲方信息
        if any([case.get('party_a_contact_name'), case.get('party_a_contact_phone')]):
            doc.add_heading('甲方联系信息', level=2)

            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'

            cells = table.rows[0].cells
            cells[0].text = '联系人'
            cells[1].text = case.get('party_a_contact_name', '')

            cells = table.rows[1].cells
            cells[0].text = '联系电话'
            cells[1].text = case.get('party_a_contact_phone', '')

            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(4.5)

            doc.add_paragraph()  # 空行

        # 乙方信息
        if any([case.get('party_b_contact_name'), case.get('party_b_contact_phone')]):
            doc.add_heading('乙方联系信息', level=2)

            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'

            cells = table.rows[0].cells
            cells[0].text = '联系人'
            cells[1].text = case.get('party_b_contact_name', '')

            cells = table.rows[1].cells
            cells[0].text = '联系电话'
            cells[1].text = case.get('party_b_contact_phone', '')

            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(4.5)

            doc.add_paragraph()  # 空行

    def _add_attachments_section(self, doc: Document, attachments: List[Dict]):
        """添加附件信息"""
        doc.add_heading('附件清单', level=2)

        if not attachments:
            doc.add_paragraph('无附件')
        else:
            table = doc.add_table(rows=len(attachments) + 1, cols=3)
            table.style = 'Light Grid Accent 1'

            # 表头
            header_cells = table.rows[0].cells
            header_cells[0].text = '序号'
            header_cells[1].text = '文件名'
            header_cells[2].text = '类型'

            # 数据行
            for i, att in enumerate(attachments, start=1):
                row_cells = table.rows[i].cells
                row_cells[0].text = str(i)
                row_cells[1].text = att.get('original_filename', '')
                row_cells[2].text = self._get_attachment_type_label(att.get('attachment_type', ''))

            # 设置列宽
            table.columns[0].width = Inches(0.8)
            table.columns[1].width = Inches(4.0)
            table.columns[2].width = Inches(1.7)

        doc.add_paragraph()  # 空行

    def _add_statistics_section(self, doc: Document, cases: List[Dict]):
        """添加统计汇总章节"""
        doc.add_heading('统计汇总', level=1)

        # 基本统计
        total_cases = len(cases)
        doc.add_paragraph(f'总案例数：{total_cases}')

        # 按行业统计
        industry_stats = {}
        for case in cases:
            industry = case.get('industry', '未分类')
            industry_stats[industry] = industry_stats.get(industry, 0) + 1

        if industry_stats:
            doc.add_heading('按行业分布', level=2)
            table = doc.add_table(rows=len(industry_stats) + 1, cols=2)
            table.style = 'Light Grid Accent 1'

            header_cells = table.rows[0].cells
            header_cells[0].text = '行业'
            header_cells[1].text = '案例数'

            for i, (industry, count) in enumerate(sorted(industry_stats.items()), start=1):
                row_cells = table.rows[i].cells
                row_cells[0].text = industry
                row_cells[1].text = str(count)

            table.columns[0].width = Inches(3.5)
            table.columns[1].width = Inches(2.0)

        # 按合同类型统计
        contract_type_stats = {}
        for case in cases:
            contract_type = case.get('contract_type', '未分类')
            contract_type_stats[contract_type] = contract_type_stats.get(contract_type, 0) + 1

        if contract_type_stats:
            doc.add_paragraph()
            doc.add_heading('按合同类型分布', level=2)
            table = doc.add_table(rows=len(contract_type_stats) + 1, cols=2)
            table.style = 'Light Grid Accent 1'

            header_cells = table.rows[0].cells
            header_cells[0].text = '合同类型'
            header_cells[1].text = '案例数'

            for i, (contract_type, count) in enumerate(sorted(contract_type_stats.items()), start=1):
                row_cells = table.rows[i].cells
                row_cells[0].text = contract_type
                row_cells[1].text = str(count)

            table.columns[0].width = Inches(3.5)
            table.columns[1].width = Inches(2.0)

    def _get_status_text(self, status: str) -> str:
        """获取案例状态文本"""
        status_map = {
            'success': '成功',
            '进行中': '进行中',
            '待验收': '待验收'
        }
        return status_map.get(status, status)

    def _get_attachment_type_label(self, type_code: str) -> str:
        """获取附件类型标签"""
        type_map = {
            'contract_order': '合同/订单',
            'invoice': '发票',
            'statement': '对账单',
            'contract': '合同文件',
            'acceptance': '验收证明',
            'testimony': '客户证明',
            'photo': '项目照片',
            'other': '其他'
        }
        return type_map.get(type_code, '其他')


# 测试代码
if __name__ == "__main__":
    # 测试数据
    test_cases = [
        {
            'case_id': 1,
            'case_title': 'XX市政府云平台建设项目',
            'case_number': 'XM-2024-001',
            'customer_name': 'XX市政府',
            'industry': '政府',
            'contract_type': '合同',
            'contract_amount': '500',
            'contract_start_date': '2024-01-01',
            'contract_end_date': '2024-12-31',
            'party_a_contact_name': '张三',
            'party_a_contact_phone': '010-12345678',
            'case_status': 'success',
            'attachments': [
                {'original_filename': '合同文件.pdf', 'attachment_type': 'contract_order'},
                {'original_filename': '验收报告.docx', 'attachment_type': 'acceptance'}
            ]
        },
        {
            'case_id': 2,
            'case_title': 'XX银行数据中心项目',
            'case_number': 'XM-2024-002',
            'customer_name': 'XX银行',
            'industry': '金融',
            'contract_type': '合同',
            'contract_amount': '800',
            'contract_start_date': '2024-03-01',
            'contract_end_date': '2024-12-31',
            'party_a_contact_name': '李四',
            'party_a_contact_phone': '010-87654321',
            'case_status': '进行中',
            'attachments': []
        }
    ]

    exporter = CaseExporter()
    output_file = '/tmp/test_cases_export.docx'

    try:
        result = exporter.export_cases_to_word(
            cases=test_cases,
            output_path=output_file,
            options={'title': '测试案例汇总', 'include_attachments': True}
        )
        print(f"导出成功: {result}")
    except Exception as e:
        print(f"导出失败: {e}")
