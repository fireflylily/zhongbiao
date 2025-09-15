#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
格式后处理模块 - 专业化的文档格式清理和美化

专门负责商务应答文档的格式后处理：
- 年月日格式处理 (Date Format Processing)
- 装饰性格式优化 (Decorative Format Optimization)
- 空白字符规范化 (Whitespace Normalization)  
- 后处理美化机制 (Final Beautification)

设计理念：将格式处理从info_filler.py中分离，实现专业化的格式清理
"""

import re
from typing import Dict, List, Optional, Tuple, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging

logger = logging.getLogger(__name__)

# =====================================
# 1. 年月日格式处理器 (Date Format Processor)
# =====================================

class DateFormatProcessor:
    """专门处理日期格式的清理和标准化"""
    
    # 年月日格式模式
    DATE_PATTERNS = {
        'year_month_day': r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
        'spaced_date': r'(\d{4})\s+年\s+(\d{1,2})\s+月\s+(\d{1,2})\s+日',
        'mixed_date': r'(\d{4})\s*[年]\s*(\d{1,2})\s*[月]\s*(\d{1,2})\s*[日]',
        'template_date': r'(\s+年\s+月\s+日\s*)',
        'empty_date': r'____年____月____日',
        'placeholder_date': r'___年___月___日'
    }
    
    # 年月日格式填充匹配模式
    DATE_FILL_PATTERNS = [
        r'^\s{2,}年\s{2,}月\s{2,}日$',      # 空格分隔的年月日格式（独立行）
        r'(\s+)年(\s+)月(\s+)日(\s*)$',      # 末尾格式：空格+年+空格+月+空格+日
        r'(\n\s*)年(\s+)月(\s+)日(\s*)$',    # 换行+空格+年月日格式
        r'(\s+)年(\s+)月(\s+)日',           # 通用格式：空格+年+空格+月+空格+日
    ]
    
    # 日期清理规则
    CLEANING_RULES = [
        # 去除多余空格的年月日
        (r'(\d{4})\s{2,}年\s{2,}(\d{1,2})\s{2,}月\s{2,}(\d{1,2})\s{2,}日', r'\1年\2月\3日'),
        # 标准化年月日间距
        (r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', r'\1年\2月\3日'),
        # 清理多余的占位符
        (r'____年____月____日', r'    年    月    日'),
        (r'___年___月___日', r'   年   月   日'),
        # 处理单独的年月日标记
        (r'年\s+月\s+日(?=\s|$)', r'年  月  日')
    ]
    
    @classmethod
    def clean_date_format(cls, text: str) -> str:
        """
        清理日期格式，去除多余的年月日标识和空格
        
        Args:
            text: 输入文本
            
        Returns:
            清理后的文本
        """
        result = text
        
        # 应用清理规则
        for pattern, replacement in cls.CLEANING_RULES:
            result = re.sub(pattern, replacement, result)
            
        return result
    
    @classmethod
    def extract_date_components(cls, text: str) -> Optional[Dict[str, str]]:
        """
        从文本中提取日期组件
        
        Args:
            text: 包含日期的文本
            
        Returns:
            日期组件字典或None
        """
        for pattern_name, pattern in cls.DATE_PATTERNS.items():
            match = re.search(pattern, text)
            if match and len(match.groups()) >= 3:
                return {
                    'year': match.group(1),
                    'month': match.group(2), 
                    'day': match.group(3),
                    'pattern': pattern_name,
                    'original': match.group(0)
                }
        return None
    
    @classmethod
    def format_date(cls, year: str, month: str, day: str, 
                   format_type: str = 'standard') -> str:
        """
        格式化日期字符串
        
        Args:
            year: 年份
            month: 月份
            day: 日期
            format_type: 格式类型 ('standard', 'spaced', 'formal')
            
        Returns:
            格式化后的日期字符串
        """
        formats = {
            'standard': f'{year}年{month}月{day}日',
            'spaced': f'{year} 年 {month} 月 {day} 日',
            'formal': f'{year}年{month.zfill(2)}月{day.zfill(2)}日'
        }
        
        return formats.get(format_type, formats['standard'])
    
    @classmethod
    def format_date_from_string(cls, date_str: str) -> str:
        """
        从字符串格式化日期
        
        Args:
            date_str: 日期字符串
            
        Returns:
            格式化后的日期字符串
        """
        if not date_str:
            return ''
            
        # 提取年月日数字
        date_match = re.search(r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})', date_str)
        if date_match:
            year, month, day = date_match.groups()
            return f'{year}年{month}月{day}日'
            
        return date_str
    
    @classmethod
    def clean_date_redundancy_and_placeholders(cls, text: str, logger=None) -> str:
        """
        清理日期相关的重复内容和占位符残留
        处理各种"年月日"重复模式和占位符
        
        Args:
            text: 输入文本
            logger: 日志对象
            
        Returns:
            清理后的文本
        """
        original_text = text
        
        # 第一组：处理直接重复的年月日字符
        redundant_patterns = [
            r'(\d+年\d+月\d+日)年',              # 2015年12月18日年
            r'(\d+年\d+月\d+日)月',              # 2015年12月18日月  
            r'(\d+年\d+月\d+日)日',              # 2015年12月18日日
            r'(\d+年\d+月\d+日)年\s*月',         # 2015年12月18日年月
            r'(\d+年\d+月\d+日)年\s*月\s*日',    # 2015年12月18日年月日
            r'(\d+年\d+月\d+日)月\s*日',         # 2015年12月18日月日
        ]
        
        for pattern in redundant_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                if logger:
                    logger.debug(f"清理重复字符: '{old_text}' -> '{text}'")
        
        # 第二组：处理带空格的重复模式
        spaced_redundant_patterns = [
            r'(\d+年\d+月\d+日)\s+年',           # 2015年12月18日 年
            r'(\d+年\d+月\d+日)\s+月',           # 2015年12月18日 月
            r'(\d+年\d+月\d+日)\s+日',           # 2015年12月18日 日
            r'(\d+年\d+月\d+日)\s+年\s*月',      # 2015年12月18日 年月
            r'(\d+年\d+月\d+日)\s+年\s*月\s*日', # 2015年12月18日 年月日
            r'(\d+年\d+月\d+日)\s+月\s*日',      # 2015年12月18日 月日
            r'(\d+年\d+月\d+日)\s*年\s*月\s*日', # 2015年12月18日年月日（任意空格）
            r'(\d+年\d+月\d+日)\s+月\s+日',      # 2015年12月18日  月  日
            r'(\d+年\d+月\d+日)\s+年\s+月',      # 2015年12月18日  年  月
        ]
        
        for pattern in spaced_redundant_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                if logger:
                    logger.debug(f"清理空格重复: '{old_text}' -> '{text}'")
        
        # 第三组：处理占位符残留
        placeholder_cleanup_patterns = [
            # 清理日期后的下划线占位符
            r'(\d+年\d+月\d+日)_+月_+日',        # 2025年09月07日_____月_____日
            r'(\d+年\d+月\d+日)_+月',            # 2025年09月07日_____月
            r'(\d+年\d+月\d+日)_+日',            # 2025年09月07日_____日
            r'(\d+年\d+月\d+日)_+年_+月_+日',    # 2025年09月07日_____年_____月_____日
            r'(\d+年\d+月\d+日)_+年_+月',        # 2025年09月07日_____年_____月
            r'(\d+年\d+月\d+日)_+年',            # 2025年09月07日_____年
            
            # 清理空格和下划线混合的情况
            r'(\d+年\d+月\d+日)[\s_]+月[\s_]*日', # 2025年09月07日 ___月___日
            r'(\d+年\d+月\d+日)[\s_]+年[\s_]*月[\s_]*日', # 带空格的混合情况
            
            # 清理横线形式的占位符
            r'(\d+年\d+月\d+日)-+',              # 2025年09月07日--------
            r'(\d+年\d+月\d+日)[\s-]+$',         # 2025年09月07日 ---- (行末)
            r'(\d+年\d+月\d+日)_+$',             # 2025年09月07日_____ (行末)
            
            # 清理日期后的任意组合占位符（更通用的模式）
            r'(\d+年\d+月\d+日)[\s_-]+.*?$',     # 日期后任意占位符到行末
        ]
        
        for pattern in placeholder_cleanup_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                if logger:
                    logger.debug(f"清理占位符残留: '{old_text}' -> '{text}'")
        
        # 第四组：处理重复日期
        duplicate_date_patterns = [
            r'(\d{4}年\d{1,2}月\d{1,2}日)\s+(\d{4}年\d{1,2}月\d{1,2}日)',  # 2025年9月12日 2025年9月12日
            r'(\d{4}年\d{1,2}月\d{1,2}日)年[_\s]*月[_\s]*日',              # 2025年9月12日年_____月_____日
        ]
        
        for pattern in duplicate_date_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                if logger:
                    logger.debug(f"清理重复日期: '{old_text}' -> '{text}'")
        
        # 如果有清理操作，记录日志
        if text != original_text and logger:
            logger.info(f"📅 日期清理完成: '{original_text}' -> '{text}'")
        
        return text
    
    @classmethod
    def process_date_format_filling(cls, doc, date_value: str, utils_module, logger=None) -> int:
        """
        处理年月日格式填充
        识别和填充"    年    月    日"等各种空格分隔的年月日格式
        
        Args:
            doc: Word文档对象
            date_value: 日期值
            utils_module: 工具模块（用于精确替换）
            logger: 日志对象
            
        Returns:
            int: 处理的段落数量
        """
        if logger:
            logger.debug("📅 开始处理年月日格式填充")
        
        if not date_value:
            if logger:
                logger.debug("⚠️ 日期值为空，跳过年月日格式填充")
            return 0
        
        # 格式化日期
        formatted_date = cls.format_date_from_string(date_value)
        if not formatted_date:
            if logger:
                logger.warning("⚠️ 日期格式化失败，跳过年月日格式填充")
            return 0
            
        if logger:
            logger.debug(f"📅 准备填充的日期值: '{formatted_date}'")
        
        processed_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if not text or not text.strip():
                continue
                
            # 检查是否匹配年月日格式
            for i, pattern in enumerate(cls.DATE_FILL_PATTERNS, 1):
                match = re.search(pattern, text)
                if match:
                    if logger:
                        logger.info(f"✅ 年月日模式{i}匹配成功: '{match.group()}' 在段落: '{text[:50]}...'")
                    
                    # 构建新文本
                    if i == 1:  # 独立的空格年月日格式
                        # 直接替换整个匹配内容
                        new_text = re.sub(pattern, formatted_date, text)
                    elif i == 3:  # 换行+空格+年月日格式
                        # 保留换行符，只替换年月日部分
                        new_text = re.sub(pattern, rf'\1{formatted_date}', text)
                    else:
                        # 标准替换：整个匹配的年月日模式为完整日期
                        new_text = re.sub(pattern, formatted_date, text)
                    
                    if new_text != text:
                        # 使用精确格式处理进行替换
                        escaped_original = re.escape(text.strip())
                        if utils_module and hasattr(utils_module, 'precise_replace'):
                            if utils_module.precise_replace(paragraph, escaped_original, new_text.strip(), logger):
                                if logger:
                                    logger.info(f"🔄 年月日格式填充成功: '{text}' -> '{new_text}'")
                                processed_count += 1
                            else:
                                # 后备方案：直接文本替换
                                paragraph.text = new_text.strip()
                                if logger:
                                    logger.info(f"🔄 年月日格式填充成功(后备): '{text}' -> '{new_text}'")
                                processed_count += 1
                        
                        # 找到一个匹配后，跳出模式循环
                        break
        
        if processed_count > 0 and logger:
            logger.info(f"📊 年月日格式填充完成，共处理 {processed_count} 个段落")
        elif logger:
            logger.debug("📊 未发现需要填充的年月日格式")
        
        return processed_count

# =====================================
# 2. 装饰性格式优化器 (Decorative Format Optimizer)
# =====================================

class DecorativeFormatOptimizer:
    """处理文档中的装饰性格式元素"""
    
    # 装饰性元素模式
    DECORATIVE_PATTERNS = {
        'seal_marks': r'[（(]\s*盖章\s*[）)]',
        'signature_marks': r'[（(]\s*签字\s*[）)]',
        'official_marks': r'[（(]\s*公章\s*[）)]',
        'bracket_spacing': r'[（(]\s+([^）)]+)\s+[）)]',
        'colon_spacing': r'([^：:\s]+)\s*[:：]\s*',
        'punctuation_excess': r'([。，！？；]){2,}',
        'whitespace_excess': r'\s{3,}'
    }
    
    # 优化规则
    OPTIMIZATION_RULES = [
        # 括号内容空格标准化
        (r'[（(]\s+([^）)]+)\s+[）)]', r'（\1）'),
        # 冒号空格标准化
        (r'([^：:\s]+)\s*[:：]\s+', r'\1：'),
        # 多余标点符号清理
        (r'([。，！？；]){2,}', r'\1'),
        # 多余空白字符清理
        (r'\s{3,}', r'  '),
        # 清理行末空格
        (r'\s+$', r'', re.MULTILINE),
        # 清理行首空格（保留缩进）
        (r'^\s{5,}', r'    ', re.MULTILINE)
    ]
    
    @classmethod
    def optimize_decorative_format(cls, text: str) -> str:
        """
        优化装饰性格式元素
        
        Args:
            text: 输入文本
            
        Returns:
            优化后的文本
        """
        result = text
        
        # 应用优化规则
        for pattern, replacement, *flags in cls.OPTIMIZATION_RULES:
            flag = flags[0] if flags else 0
            result = re.sub(pattern, replacement, result, flags=flag)
            
        return result
    
    @classmethod
    def clean_signature_areas(cls, text: str) -> str:
        """
        清理签字区域格式
        
        Args:
            text: 输入文本
            
        Returns:
            清理后的文本
        """
        # 保留签字相关的格式不做修改
        signature_protected = [
            r'法定代表人\s*[（(]?\s*签字\s*[）)]?',
            r'授权代表人\s*[（(]?\s*签字\s*[）)]?',
            r'项目负责人\s*[（(]?\s*签字\s*[）)]?'
        ]
        
        result = text
        for pattern in signature_protected:
            # 这些区域保持原格式，不做清理
            matches = re.finditer(pattern, result, re.IGNORECASE)
            for match in matches:
                logger.debug(f"保护签字区域格式: {match.group(0)}")
                
        return result

# =====================================
# 3. 空白字符规范化器 (Whitespace Normalizer)
# =====================================

class WhitespaceNormalizer:
    """专门处理文档中的空白字符规范化"""
    
    # 空白字符类型
    WHITESPACE_TYPES = {
        'space': ' ',           # 普通空格
        'tab': '\t',           # 制表符
        'newline': '\n',       # 换行符
        'carriage': '\r',      # 回车符
        'nbsp': '\u00a0',      # 不间断空格
        'ideographic': '\u3000' # 全角空格
    }
    
    # 规范化规则
    NORMALIZATION_RULES = [
        # 全角空格转换为半角空格
        ('\u3000', ' '),
        # 不间断空格转换为普通空格
        ('\u00a0', ' '),
        # 制表符转换为4个空格
        ('\t', '    '),
        # Windows换行符标准化
        ('\r\n', '\n'),
        # 多个连续空格合并
        (r' {2,}', '  '),
        # 清理行末空格
        (r' +$', '', re.MULTILINE),
    ]
    
    @classmethod
    def normalize_spacing(cls, text: str) -> str:
        """
        规范化文本中的空白字符
        
        Args:
            text: 输入文本
            
        Returns:
            规范化后的文本
        """
        result = text
        
        # 应用规范化规则
        for pattern, replacement, *flags in cls.NORMALIZATION_RULES:
            if flags:
                result = re.sub(pattern, replacement, result, flags=flags[0])
            else:
                result = result.replace(pattern, replacement)
                
        return result
    
    @classmethod
    def preserve_table_spacing(cls, text: str) -> str:
        """
        保持表格区域的空格格式
        
        Args:
            text: 输入文本
            
        Returns:
            处理后的文本
        """
        # 识别表格区域并保护其空格格式
        table_patterns = [
            r'^\s*\|.*\|\s*$',      # Markdown表格
            r'^\s*[─┌┐└┘├┤┬┴┼│]+\s*$',  # 文本表格边框
            r'^\s*[^\s]+\s+[^\s]+\s+[^\s]+\s*$'  # 多列对齐数据
        ]
        
        lines = text.split('\n')
        result_lines = []
        
        for line in lines:
            is_table_line = any(re.match(pattern, line) for pattern in table_patterns)
            
            if is_table_line:
                # 表格行保持原格式
                result_lines.append(line)
                logger.debug(f"保护表格行格式: {line[:30]}...")
            else:
                # 非表格行正常处理
                result_lines.append(cls.normalize_spacing(line))
                
        return '\n'.join(result_lines)

# =====================================
# 4. 文档美化处理器 (Document Beautifier)
# =====================================

class DocumentBeautifier:
    """文档的最终美化处理"""
    
    @classmethod
    def apply_final_beautification(cls, doc: Document) -> Document:
        """
        应用最终的文档美化处理
        
        Args:
            doc: Word文档对象
            
        Returns:
            美化后的文档对象
        """
        # 处理所有段落
        for paragraph in doc.paragraphs:
            cls._beautify_paragraph(paragraph)
            
        # 处理表格
        for table in doc.tables:
            cls._beautify_table(table)
            
        return doc
    
    @classmethod
    def _beautify_paragraph(cls, paragraph) -> None:
        """美化段落格式"""
        if not paragraph.text.strip():
            return
            
        # 设置行距
        paragraph_format = paragraph.paragraph_format
        if paragraph_format.line_spacing is None:
            paragraph_format.line_spacing = 1.2
            
        # 处理段落中的运行
        for run in paragraph.runs:
            cls._beautify_run(run)
    
    @classmethod
    def _beautify_run(cls, run) -> None:
        """美化文本运行格式"""
        if not run.text.strip():
            return
            
        # 设置字体大小（如果未设置）
        if run.font.size is None:
            run.font.size = Pt(12)
            
        # 清理文本内容
        original_text = run.text
        cleaned_text = cls._clean_run_text(original_text)
        
        if cleaned_text != original_text:
            run.text = cleaned_text
            logger.debug(f"美化运行文本: '{original_text}' -> '{cleaned_text}'")
    
    @classmethod
    def _clean_run_text(cls, text: str) -> str:
        """清理运行中的文本"""
        # 应用所有清理规则
        result = text
        
        # 日期格式清理
        result = DateFormatProcessor.clean_date_format(result)
        
        # 装饰性格式优化
        result = DecorativeFormatOptimizer.optimize_decorative_format(result)
        
        # 空白字符规范化
        result = WhitespaceNormalizer.normalize_spacing(result)
        
        return result
    
    @classmethod
    def _beautify_table(cls, table) -> None:
        """美化表格格式"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cls._beautify_paragraph(paragraph)

# =====================================
# 5. 格式清理协调器 (Format Cleaner Coordinator)
# =====================================

class FormatCleaner:
    """格式清理的主协调器，统一管理所有格式处理流程"""
    
    def __init__(self):
        self.date_processor = DateFormatProcessor()
        self.decorative_optimizer = DecorativeFormatOptimizer()
        self.whitespace_normalizer = WhitespaceNormalizer()
        self.document_beautifier = DocumentBeautifier()
        
    def clean_text(self, text: str, options: Dict[str, bool] = None) -> str:
        """
        清理文本格式
        
        Args:
            text: 输入文本
            options: 清理选项
            
        Returns:
            清理后的文本
        """
        default_options = {
            'clean_dates': True,
            'optimize_decorative': True,
            'normalize_whitespace': True,
            'preserve_table_format': True
        }
        
        options = {**default_options, **(options or {})}
        result = text
        
        try:
            # 1. 日期格式清理
            if options['clean_dates']:
                result = self.date_processor.clean_date_format(result)
                logger.debug("应用日期格式清理")
            
            # 2. 装饰性格式优化
            if options['optimize_decorative']:
                result = self.decorative_optimizer.optimize_decorative_format(result)
                result = self.decorative_optimizer.clean_signature_areas(result)
                logger.debug("应用装饰性格式优化")
            
            # 3. 空白字符规范化
            if options['normalize_whitespace']:
                if options['preserve_table_format']:
                    result = self.whitespace_normalizer.preserve_table_spacing(result)
                else:
                    result = self.whitespace_normalizer.normalize_spacing(result)
                logger.debug("应用空白字符规范化")
                
        except Exception as e:
            logger.error(f"格式清理过程中出现错误: {e}")
            return text  # 返回原文本
            
        return result
    
    def clean_document(self, doc: Document, options: Dict[str, bool] = None) -> Document:
        """
        清理整个Word文档的格式
        
        Args:
            doc: Word文档对象
            options: 清理选项
            
        Returns:
            清理后的文档对象
        """
        default_options = {
            'apply_beautification': True,
            'clean_paragraph_text': True
        }
        
        options = {**default_options, **(options or {})}
        
        try:
            # 清理段落文本
            if options['clean_paragraph_text']:
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        cleaned_text = self.clean_text(original_text)
                        
                        if cleaned_text != original_text:
                            # 更新段落文本，保持格式
                            for run in paragraph.runs:
                                if run.text:
                                    run.text = self.clean_text(run.text)
            
            # 应用文档美化
            if options['apply_beautification']:
                doc = self.document_beautifier.apply_final_beautification(doc)
                
        except Exception as e:
            logger.error(f"文档格式清理过程中出现错误: {e}")
            
        return doc
    
    def get_cleaning_statistics(self, original_text: str, cleaned_text: str) -> Dict[str, Any]:
        """
        获取清理统计信息
        
        Args:
            original_text: 原始文本
            cleaned_text: 清理后文本
            
        Returns:
            统计信息字典
        """
        return {
            'original_length': len(original_text),
            'cleaned_length': len(cleaned_text),
            'reduction_percentage': (len(original_text) - len(cleaned_text)) / len(original_text) * 100,
            'lines_original': len(original_text.split('\n')),
            'lines_cleaned': len(cleaned_text.split('\n')),
            'whitespace_reduction': original_text.count(' ') - cleaned_text.count(' ')
        }

# 导出主要类
__all__ = [
    'DateFormatProcessor',
    'DecorativeFormatOptimizer', 
    'WhitespaceNormalizer',
    'DocumentBeautifier',
    'FormatCleaner'
]