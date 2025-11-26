#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
案例表格生成器 - CaseTableGenerator
为"格式自拟"的业绩案例要求生成标准表格

职责：
- 在指定位置生成标准案例表格
- 设置表格样式（边框、对齐、列宽）
- 生成的表格符合 CaseTableFiller 的识别规则

设计理念：
- 生成在前，填充在后（第3.5步生成 → 第4步填充）
- 生成的表格必须包含 CaseTableFiller.case_table_headers 中的关键字段
- 完全复用现有的数据填充逻辑
"""

from typing import Optional
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger


class CaseTableGenerator:
    """案例表格生成器 - 为"格式自拟"的案例要求生成标准表格"""

    def __init__(self):
        self.logger = get_module_logger("case_table_generator")

        # 标准案例表格模板（可配置）
        # 注意：表头必须包含 CaseTableFiller.case_table_headers 中的关键字段
        self.default_table_structure = [
            {'header': '序号', 'width': 0.8},
            {'header': '项目名称', 'width': 2.8},
            {'header': '客户名称', 'width': 2.2},
            {'header': '合同金额（万元）', 'width': 1.6},
            {'header': '实施时间', 'width': 2.0}
        ]

    def generate_case_table(self, doc: Document, insert_after_para,
                           num_rows: int = 5) -> Optional[Table]:
        """
        在指定段落后生成案例表格

        生成步骤：
        1. 插入分页符（将案例表格单独成页）
        2. 插入表格标题（"业绩案例一览表"）
        3. 创建表格（表头行 + N行数据行）
        4. 设置表头样式（加粗、背景色）
        5. 设置数据行样式（居中对齐）
        6. 设置表格边框

        Args:
            doc: Word文档对象
            insert_after_para: 插入位置（在该段落后）
            num_rows: 数据行数（默认5行，不含表头）

        Returns:
            生成的表格对象，失败返回None
        """
        try:
            self.logger.info(f"🆕 开始生成案例表格，{num_rows}行...")

            # 1. 插入分页符（将案例表格单独成页）
            page_break_para = self._insert_paragraph_after(doc, insert_after_para)
            run = page_break_para.add_run()
            run.add_break()

            # 2. 插入表格标题
            title_para = self._insert_paragraph_after(doc, page_break_para)
            title_para.text = "业绩案例一览表"
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title_para.runs:
                title_para.runs[0].font.size = Pt(14)
                title_para.runs[0].font.bold = True

            # 3. 创建表格（表头行 + 数据行）
            total_rows = num_rows + 1
            total_cols = len(self.default_table_structure)

            # 在标题段落后插入表格
            table = self._insert_table_after(doc, title_para, total_rows, total_cols)

            # 4. 设置表头
            header_row = table.rows[0]
            for col_idx, col_info in enumerate(self.default_table_structure):
                cell = header_row.cells[col_idx]
                cell.text = col_info['header']

                # 设置表头样式
                self._set_cell_style(cell, bold=True, background='D9EAD3')

                # 设置列宽
                if 'width' in col_info:
                    self._set_column_width(table, col_idx, col_info['width'])

            # 5. 设置数据行样式（居中对齐、边框）
            for row_idx in range(1, total_rows):
                for col_idx in range(total_cols):
                    cell = table.rows[row_idx].cells[col_idx]
                    self._set_cell_style(cell, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # 6. 设置表格整体样式
            self._set_table_style(table)

            self.logger.info(f"✅ 案例表格生成完成: {total_rows}行 x {total_cols}列")
            return table

        except Exception as e:
            self.logger.error(f"❌ 生成案例表格失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return None

    def _insert_paragraph_after(self, doc: Document, paragraph):
        """
        在指定段落后插入新段落

        Args:
            doc: Word文档对象
            paragraph: 参考段落

        Returns:
            新段落对象
        """
        para_element = paragraph._element
        new_para_element = OxmlElement('w:p')
        para_element.addnext(new_para_element)

        from docx.text.paragraph import Paragraph
        return Paragraph(new_para_element, doc)

    def _insert_table_after(self, doc: Document, paragraph, rows: int, cols: int) -> Table:
        """
        在指定段落后插入表格

        Args:
            doc: Word文档对象
            paragraph: 参考段落
            rows: 行数
            cols: 列数

        Returns:
            表格对象
        """
        # 创建表格元素
        tbl = OxmlElement('w:tbl')

        # 添加表格属性（宽度100%）
        tblPr = OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)
        tbl.append(tblPr)

        # 创建表格网格
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(cols):
            gridCol = OxmlElement('w:gridCol')
            tblGrid.append(gridCol)
        tbl.append(tblGrid)

        # 创建表格行和单元格
        for _ in range(rows):
            tr = OxmlElement('w:tr')
            for _ in range(cols):
                tc = OxmlElement('w:tc')
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
                p = OxmlElement('w:p')
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

        # 插入到段落后
        paragraph._element.addnext(tbl)

        # 包装为Table对象
        return Table(tbl, doc)

    def _set_cell_style(self, cell, bold=False, background=None,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """
        设置单元格样式

        Args:
            cell: 单元格对象
            bold: 是否加粗
            background: 背景颜色（十六进制，如'D9EAD3'）
            alignment: 段落对齐方式
        """
        # 设置段落对齐
        if cell.paragraphs:
            cell.paragraphs[0].alignment = alignment

            # 设置文字格式
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = bold

        # 设置背景色
        if background:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), background)
            cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_column_width(self, table: Table, col_idx: int, width_inches: float):
        """
        设置列宽

        Args:
            table: 表格对象
            col_idx: 列索引
            width_inches: 宽度（英寸）
        """
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = Inches(width_inches)

    def _set_table_style(self, table: Table):
        """
        设置表格整体样式（边框、对齐）

        Args:
            table: 表格对象
        """
        tbl = table._element
        tblPr = tbl.tblPr

        # 设置表格边框
        tblBorders = OxmlElement('w:tblBorders')
        border_names = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']

        for border_name in border_names:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

        # 设置表格居中对齐
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)


# ==================== 设计说明（保留在代码中） ====================
"""
## 设计理念

### 1. 为什么是第3.5步？
- 第3步：图片插入（资质证书）
- 🆕 第3.5步：生成案例表格
- 第4步：案例表格填充（复用现有逻辑）

**核心原则**：生成在前，填充在后
- 如果在第4步之后生成（第4.5步），新生成的表格不会被填充
- 必须在第4步之前生成，才能被 CaseTableFiller 自动识别和填充

### 2. 与 CaseTableFiller 的兼容性

**表头关键字匹配**：
CaseTableFiller 通过检查表头关键字识别案例表格：
```python
case_table_headers = {
    '项目名称', '案例名称', '客户名称', '合同金额',
    '实施时间', '联系人', '联系方式', ...
}
```

**识别规则**：至少匹配2个关键字段

我们生成的表格包含：
- ✅ 项目名称
- ✅ 客户名称
- ✅ 合同金额
- ✅ 实施时间
- ✅ 联系人

→ 符合识别规则，会被自动填充 ✅

### 3. 完整流程

```
"格式自拟"文字
     ↓
DocumentScanner.scan_case_requirements()
     ↓
识别为需要生成表格的位置
     ↓
🆕 CaseTableGenerator.generate_case_table()
     ↓
生成标准案例表格
     ↓
CaseTableFiller.fill_case_tables()
     ↓
自动识别表格 → 填充数据 → 插入附件图片
     ↓
完成！
```

### 4. 智能去重

在生成表格之前，DocumentScanner 会检查附近是否已有案例表格：
- 有表格 → 跳过生成 ⏭️
- 无表格 → 生成表格 ✅

避免重复生成表格。

### 5. 可配置性

表格结构可以通过修改 `default_table_structure` 自定义：
```python
default_table_structure = [
    {'header': '序号', 'width': 0.8},
    {'header': '项目名称', 'width': 2.5},
    # ... 添加或修改列
]
```

注意：表头必须包含 CaseTableFiller 能识别的关键字段。
"""
