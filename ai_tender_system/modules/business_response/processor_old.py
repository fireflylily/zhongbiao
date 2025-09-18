#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
点对点应答处理器 - 统一字段处理框架版本
整合所有优化功能，包含智能双字段表格处理和采购人信息识别
1、供应商名称、项目编号、项目名称、法定代
   表人的填写，存在的规则比较多。
   1.1 替换格式，比如（供应商名称），填写完成
   是（智慧足迹数据科技有限公司）
   1.2 填空规则，比如 供应商名称：
   。填写完成是
   供应商名称：智慧足迹数据科技有限公司
   1.3 多字段替换规则，比如（供应商名称、地址
   ）填写完成是（智慧足迹数据科技有限公司、北
   京市东城区王府井大街200号七层）
   1.4 供应商名称变体较多，可能是公司名称、应
   答人名称、供应商名称（盖章）、供应商名称：
   。         （公章）等格式。

   2、公司其他信息填写。仅做填空规则即可。不需
   要做组合规则，按单标签循环即可。
   \
   3、例外格式：\
   3.1 法定代表人，授权代表人，后面带有"签字"
   字样的，不需要做填空规则或替换格式。
   3.2 日期，日期两个字中间可能会出现不定数量
   的空格。日期后面带有 年  月
   日标识，需要在后处理机制中去掉。
   3.3 有些标签，可能是采购人或招标人用来公示
   信息的，比如地址：中信大厦1204室。这种情况
   就不需要填写内容。
   \
   \
   4. 规则细化，
   4.1标签名称变体，比如邮箱同电子邮件等情况。
   4.2格式变化，比如带冒号，空格，占位签

   5。格式要求。
   完全不改变原有格式，替换规则直接使用原有文
   本的第一个字的格式即可。填写规则的新增内容
   也使用第一个规则即可。（建议使用重构前原代
   码）

   6。后处理机制。清除多余的占位符等。（建议使
   用重构前原代码）

   所有项目信息和公司信息都来源于系统配置文件
   。

   商务应答模块重构实施方案                │
     │                                         │
     │ 第一步：目录结构调整                    │
     │                                         │
     │ ai_tender_system/modules/               │
     │ ├── business_response/        #         │
     │ 商务应答（原point_to_point改名）        │
     │ │   ├── __init__.py                     │
     │ │   ├── processor.py          #         │
     │ 主处理器，协调三个子模块                │
     │ │   ├── info_filler.py        #         │
     │ 信息填写模块                            │
     │ │   ├── table_processor.py    #         │
     │ 表格处理模块                            │
     │ │   ├── image_handler.py      #         │
     │ 图片插入模块                            │
     │ │   └── utils.py              #         │
     │ 共享工具函数                            │
     │ └── point_to_point/           #         │
     │ 新建：技术需求点对点回复                │
     │     ├── __init__.py                     │
     │     └── tech_responder.py     #         │
     │ 技术需求回复处理器                      │
     │                                         │
     │ 第二步：信息填写模块(info_filler.py)实  │
     │ 现                                      │
     │                                         │
     │ 2.1 核心字段处理规则                    │
     │                                         │
     │ 供应商名称类（支持多种规则）：          │
     │ - 替换规则：（供应商名称） →            │
     │ （智慧足迹数据科技有限公司）            │
     │ - 填空规则：供应商名称：____ →          │
     │ 供应商名称：智慧足迹数据科技有限公司    │
     │ - 组合规则：（供应商名称、地址） → （智 │
     │ 慧足迹数据科技有限公司、北京市东城区）  │
     │ - 变体处理：公司名称、应答人名称、供应  │
     │ 商名称（盖章）等                        │
     │                                         │
     │ 其他信息字段（仅填空）：                │
     │ - 电话、邮箱、地址、邮编、传真等        │
     │ - 支持标签变体（如：邮箱/电子邮件）     │
     │ - 支持格式变化（冒号、空格、占位符）    │
     │                                         │
     │ 2.2 例外处理                            │
     │                                         │
     │ - 跳过"签字"相关字段（法定代表人签字、  │
     │ 授权代表人签字）                        │
     │ - 智能日期处理（处理空格、去除多余的年  │
     │ 月日）                                  │
     │ - 识别并跳过采购人/招标人信息           │
     │                                         │
     │ 第三步：表格处理模块(table_processor.py │
     │ )                                       │
     │                                         │
     │ - 识别表格中的待填字段                  │
     │ - 保持表格格式不变                      │
     │ - 处理合并单元格                        │
     │ - 支持表格内的字段组合                  │
     │                                         │
     │ 第四步：图片插入模块(image_handler.py)  │
     │                                         │
     │ - 公司公章图片插入                      │
     │ - 资质证明图片插入                      │
     │ - 保持文档布局                          │
     │ - 图片尺寸自适应                        │
     │                                         │
     │ 第五步：技术需求回复模块(tech_responder │
     │ .py)                                    │
     │                                         │
     │ - 恢复原有的技术需求点对点回复功能      │
     │ - 基于需求自动生成技术响应              │
     │ - 支持技术参数匹配                      │
     │ - 技术方案模板填充                      │
     │                                         │
     │ 实施步骤：                              │
     │                                         │
     │ 1. 备份现有代码（5分钟）                │
     │   - 备份当前processor.py                │
     │   - 保存测试用例                        │
     │ 2. 创建新目录结构（10分钟）             │
     │   -                                     │
     │ 重命名point_to_point为business_response │
     │   - 创建新的point_to_point目录          │
     │ 3. 拆分info_filler.py（2小时）          │
     │   - 提取信息填写相关代码                │
     │   - 实现六大规则类型                    │
     │   - 添加例外处理逻辑                    │
     │ 4. 实现table_processor.py（1小时）      │
     │   - 提取表格处理逻辑                    │
     │   - 优化表格识别算法                    │
     │ 5. 实现image_handler.py（1小时）        │
     │   - 实现图片插入功能                    │
     │   - 处理图片定位和缩放                  │
     │ 6. 恢复tech_responder.py（2小时）       │
     │   - 查找原有技术回复代码                │
     │   - 重新实现技术需求响应                │
     │ 7. 集成测试（1小时）                    │
     │   - 测试商务应答三大功能                │
     │   - 验证技术需求回复                    │
     │   - 确保格式保持完整                    │
     │                                         │
     │ 预期效果：                              │
     │                                         │
     │ - 代码结构清晰，每个模块200行以内       │
     │ - 功能独立，便于维护和测试              │
     │ - 恢复丢失的技术需求回复功能            │
     │ - 保持原有的格式处理能力                │
     │ - 提高字段识别准确率到95%+  
"""

import json
import re
from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import (
    get_config, get_module_logger, 
    BusinessResponseError, APIError, FileProcessingError,
    safe_filename, ensure_dir
)

class PointToPointProcessor:
    """点对点应答处理器 - 统一字段处理框架版本"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.config = get_config()
        self.logger = get_module_logger("point_to_point")
        
        # API配置
        api_config = self.config.get_api_config()
        self.api_key = api_key or api_config['api_key']
        
        # 实例变量初始化
        self.company_info = {}
        self.project_name = ""
        self.tender_no = ""
        self.date_text = ""
        
        self.logger.info("点对点应答处理器初始化完成（统一字段处理框架版本）")
    
    def process_business_response(self, 
                                 input_file: str, 
                                 output_file: str,
                                 company_info: Dict[str, Any],
                                 project_name: str = "",
                                 tender_no: str = "", 
                                 date_text: str = "") -> Dict[str, Any]:
        """
        处理商务应答文档 - 使用统一字段处理框架
        
        Args:
            input_file: 输入文档路径
            output_file: 输出文档路径
            company_info: 完整的公司信息字典
            project_name: 项目名称
            tender_no: 招标编号
            date_text: 日期文本
            
        Returns:
            dict: 处理结果
        """
        try:
            self.logger.info(f"开始处理商务应答文档（统一框架版本）")
            self.logger.info(f"输入文件: {input_file}")
            self.logger.info(f"输出文件: {output_file}")
            self.logger.info(f"公司名称: {company_info.get('companyName', 'N/A')}")
            self.logger.info(f"项目名称: {project_name}")
            self.logger.info(f"招标编号: {tender_no}")
            self.logger.info(f"日期文本: {date_text}")
            
            # 保存信息到实例变量
            self.company_info = company_info
            self.project_name = project_name
            self.tender_no = tender_no
            self.date_text = date_text
            
            # 第1步：处理投标人名称
            self.logger.info("第1步：处理投标人名称")
            name_result = self._process_bidder_name(input_file, output_file, company_info.get('companyName', ''))
            
            if not name_result.get('success'):
                return name_result
            
            # 第2步：处理其他公司信息字段（统一框架）
            self.logger.info("第2步：处理其他公司信息字段（统一框架）")
            info_result = self._process_company_info_fields(output_file, company_info, project_name, tender_no, date_text)
            
            # 合并结果
            combined_stats = name_result.get('stats', {})
            info_stats = info_result.get('stats', {})
            
            # 更新总计数
            combined_stats['total_replacements'] = combined_stats.get('total_replacements', 0) + info_stats.get('total_replacements', 0)
            combined_stats['info_fields_processed'] = info_stats.get('info_fields_processed', 0)
            
            # 合并处理的模式列表
            if 'patterns_found' in combined_stats and 'patterns_found' in info_stats:
                combined_stats['patterns_found'].extend(info_stats['patterns_found'])
            
            self.logger.info(f"商务应答文档处理完成，总共处理了{combined_stats.get('total_replacements', 0)}个字段")
            
            return {
                'success': True,
                'stats': combined_stats,
                'message': f'商务应答文档处理完成，处理了{combined_stats.get("total_replacements", 0)}个字段',
                'output_file': output_file
            }
            
        except Exception as e:
            self.logger.error(f"商务应答文档处理失败: {e}")
            return {
                'success': False,
                'error': f'处理失败: {str(e)}',
                'message': '处理失败'
            }
    
    def _process_bidder_name(self, input_file: str, output_file: str, company_name: str) -> Dict[str, Any]:
        """处理投标人名称 - 简化版本"""
        try:
            self.logger.info(f"处理投标人名称: {company_name}")
            
            # 复制输入文件到输出文件
            import shutil
            shutil.copy2(input_file, output_file)
            
            doc = Document(output_file)
            total_replacements = 0
            
            # 简单的投标人名称替换
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if '供应商名称' in para_text and company_name:
                    # 查找供应商名称字段并替换
                    for run in paragraph.runs:
                        if '供应商名称' in run.text:
                            # 简单替换逻辑
                            if ':' in run.text or '：' in run.text:
                                new_text = re.sub(r'(供应商名称[:：]\s*)([_\s]*)', f'\\1{company_name}', run.text)
                                if new_text != run.text:
                                    run.text = new_text
                                    total_replacements += 1
                                    self.logger.info(f"替换供应商名称: {company_name}")
            
            doc.save(output_file)
            
            return {
                'success': True,
                'stats': {
                    'total_replacements': total_replacements,
                    'patterns_found': ['供应商名称'] if total_replacements > 0 else []
                }
            }
            
        except Exception as e:
            self.logger.error(f"处理投标人名称失败: {e}")
            return {
                'success': False,
                'error': f'投标人名称处理失败: {str(e)}'
            }
    
    def _process_company_info_fields(self, file_path: str, company_info: dict, 
                                   project_name: str, tender_no: str, date_text: str):
        """
        使用统一框架处理公司信息字段 - 完整迁移版本
        
        Args:
            file_path: 文档路径
            company_info: 公司信息
            project_name: 项目名称
            tender_no: 招标编号  
            date_text: 日期文本
            
        Returns:
            dict: 处理结果
        """
        try:
            doc = Document(file_path)
            self.logger.info(f"开始处理公司信息字段（统一框架），文档共有 {len(doc.paragraphs)} 个段落")
            self.logger.info(f"公司数据: fixedPhone={company_info.get('fixedPhone')}, email={company_info.get('email')}, address={company_info.get('address')}")
            
            total_replacements = 0
            patterns_found = []
            processed_paragraphs = set()  # 记录已处理的段落，防止重复处理
            
            # 获取统一的字段配置
            field_configs = self._create_unified_field_config(company_info, project_name, tender_no, date_text)
            self.logger.info(f"创建了 {len(field_configs)} 个字段配置")
            
            # 处理所有段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if para_idx in processed_paragraphs:
                    continue
                    
                para_text = paragraph.text.strip()
                if not para_text:
                    continue
                
                self.logger.info(f"检查段落 #{para_idx}: '{para_text[:80]}{'...' if len(para_text) > 80 else ''}'")
                
                # 使用统一框架处理字段
                for field_config in field_configs:
                    result = self._process_unified_field(paragraph, para_text, field_config, para_idx)
                    if result['modified']:
                        total_replacements += result['replacements']
                        patterns_found.extend(result['patterns'])
                        processed_paragraphs.add(para_idx)
                        self.logger.info(f"段落 #{para_idx} 字段处理成功: {result['patterns']}")
                        break  # 一个段落只处理一个字段，避免冲突
            
            # 后处理美化机制 - 统一清理和优化格式
            beautified_paragraphs = self._post_process_beautification(doc, total_replacements)
            
            # 保存处理后的文档
            doc.save(file_path)
            
            self.logger.info(f"公司信息字段处理完成，处理了 {total_replacements} 个字段，美化了 {beautified_paragraphs} 个段落")
            
            return {
                'success': True,
                'stats': {
                    'total_replacements': total_replacements,
                    'info_fields_processed': total_replacements,
                    'patterns_found': patterns_found,
                    'beautified_paragraphs': beautified_paragraphs
                },
                'message': f'处理了{total_replacements}个公司信息字段'
            }
            
        except Exception as e:
            self.logger.error(f"统一字段处理失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return {
                'success': False,
                'stats': {'total_replacements': 0, 'info_fields_processed': 0, 'patterns_found': []},
                'error': f'统一字段处理失败: {str(e)}'
            }
    
    def _create_unified_field_config(self, company_info: dict, project_name: str, tender_no: str, date_text: str):
        """
        创建统一的字段处理配置 - 完整迁移版本
        """
        return [
            # 联系电话字段
            {
                'field_names': ['电话', '联系电话', '固定电话', '电话号码', '联系方式'],
                'value': company_info.get('fixedPhone', ''),
                'display_name': '联系电话',
                'field_type': 'contact',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理', '联系人', '项目联系人', '业主'],
                'formats': {
                    'with_colon': True,
                    'without_colon': True,
                    'placeholder_types': ['underline', 'space', 'mixed'],
                    'table_layout': True
                },
                'table_combinations': [
                    {
                        'partner_field': '电子邮件',
                        'pattern': r'(电话|联系电话)(\s{8,})(电子邮件|电子邮箱|邮箱)',
                        'description': '电话+邮件表格组合'
                    }
                ]
            },
            # 电子邮件字段
            {
                'field_names': ['电子邮件', '电子邮箱', '邮箱', 'email', 'Email'],
                'value': company_info.get('email', ''),
                'display_name': '电子邮件',
                'field_type': 'contact',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理', '联系人', '项目联系人', '业主'],
                'formats': {
                    'with_colon': True,
                    'without_colon': True,
                    'placeholder_types': ['underline', 'space', 'mixed'],
                    'table_layout': True
                },
                'table_combinations': [
                    {
                        'partner_field': '电话',
                        'pattern': r'(电话|联系电话)(\s{8,})(电子邮件|电子邮箱|邮箱)',
                        'description': '电话+邮件表格组合（作为第二字段）'
                    }
                ]
            },
            # 地址字段
            {
                'field_names': ['地址', '注册地址', '办公地址', '联系地址', '通讯地址'],
                'value': (company_info.get('address', '') or 
                         company_info.get('registeredAddress', '') or 
                         company_info.get('officeAddress', '') or
                         '北京市东城区王府井大街200号七层711室'),  # 默认地址
                'display_name': '地址',
                'field_type': 'basic_info',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': True,
                    'placeholder_types': ['underline', 'space', 'mixed'],
                    'table_layout': True
                },
                'table_combinations': [
                    {
                        'partner_field': '传真',
                        'pattern': r'(地址|注册地址|办公地址|联系地址)(\s{8,})(传真)',
                        'description': '地址+传真表格组合'
                    }
                ]
            },
            # 邮政编码字段
            {
                'field_names': ['邮政编码', '邮编', '邮码'],
                'value': company_info.get('postalCode', '') or '100006',
                'display_name': '邮政编码',
                'field_type': 'basic_info',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': True,
                    'placeholder_types': ['underline', 'space', 'mixed'],
                    'table_layout': False
                }
            },
            # 传真字段
            {
                'field_names': ['传真', '传真号码', '传真号', 'fax', 'Fax'],
                'value': company_info.get('fax', ''),
                'display_name': '传真',
                'field_type': 'contact',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理', '联系人', '项目联系人', '业主'],
                'formats': {
                    'with_colon': True,
                    'without_colon': True,
                    'placeholder_types': ['underline', 'space', 'mixed'],
                    'table_layout': True
                }
            },
            # 日期字段 - 特殊处理
            {
                'field_names': ['日期'],
                'value': date_text or '2025年9月12日',
                'display_name': '日期',
                'field_type': 'date',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'mixed'],
                    'table_layout': False
                }
            }
        ]
    
    def _process_unified_field(self, paragraph: Paragraph, para_text: str, field_config: dict, para_idx: int):
        """统一字段处理方法 - 简化版本"""
        result = {'modified': False, 'replacements': 0, 'patterns': []}
        
        try:
            field_value = field_config.get('value', '')
            field_names = field_config.get('field_names', [])
            field_type = field_config.get('field_type', '')
            
            if not field_value:
                return result
            
            # 检查是否是采购人信息（需要排除）
            if self._is_purchaser_info(para_text, field_config.get('exclude_contexts', [])):
                self.logger.info(f"段落 #{para_idx} 识别为采购人信息，跳过处理")
                return result
            
            # 首先尝试表格式双字段处理
            if field_config.get('formats', {}).get('table_layout'):
                self.logger.info(f"尝试表格式处理: 字段={field_names[0]}, 段落='{para_text}'")
                table_result = self._handle_dual_field_table_layout(para_text, field_names[0], field_value)
                if table_result != para_text:
                    # 替换整个段落文本
                    paragraph.clear()
                    paragraph.add_run(table_result)
                    result['modified'] = True
                    result['replacements'] = 1
                    result['patterns'] = [f'{field_names[0]}(表格式)']
                    self.logger.info(f"双字段表格处理成功: '{para_text}' -> '{table_result}'")
                    return result
                else:
                    self.logger.info(f"双字段表格处理无匹配: '{para_text}'")
            
            # 标准字段处理
            for field_name in field_names:
                # 日期字段特殊处理
                if field_type == 'date' and field_name == '日期':
                    # 匹配 "____年____月____日" 格式
                    date_pattern = r'_{2,}年_{2,}月_{2,}日'
                    if re.search(date_pattern, para_text):
                        for run in paragraph.runs:
                            if re.search(date_pattern, run.text):
                                new_text = re.sub(date_pattern, field_value, run.text)
                                if new_text != run.text:
                                    run.text = new_text
                                    result['modified'] = True
                                    result['replacements'] = 1
                                    result['patterns'] = [f'{field_name}(日期格式)']
                                    self.logger.info(f"日期字段替换成功: {field_value}")
                                    return result
                    
                    # 匹配 "日期：____年____月____日" 格式
                    date_pattern2 = rf'({field_name}[:：]\s*)_{2,}年_{2,}月_{2,}日'
                    if re.search(date_pattern2, para_text):
                        for run in paragraph.runs:
                            if re.search(date_pattern2, run.text):
                                new_text = re.sub(date_pattern2, lambda m: f'{m.group(1)}{field_value}', run.text)
                                if new_text != run.text:
                                    run.text = new_text
                                    result['modified'] = True
                                    result['replacements'] = 1
                                    result['patterns'] = [f'{field_name}(带冒号日期)']
                                    self.logger.info(f"带冒号日期字段替换成功: {field_value}")
                                    return result
                
                # 检查带冒号的格式
                pattern = rf'({field_name}[:：]\s*)([_\s]*)'
                self.logger.info(f"尝试标准字段处理: {field_name}, 模式='{pattern}', 段落='{para_text}'")
                if re.search(pattern, para_text):
                    self.logger.info(f"模式匹配成功: {field_name}")
                    for run in paragraph.runs:
                        self.logger.info(f"检查Run: '{run.text}'")
                        if field_name in run.text and (':' in run.text or '：' in run.text):
                            old_text = run.text
                            new_text = re.sub(pattern, lambda m: f'{m.group(1)}{field_value}', run.text)
                            if new_text != run.text:
                                run.text = new_text
                                result['modified'] = True
                                result['replacements'] = 1
                                result['patterns'] = [field_name]
                                self.logger.info(f"字段替换成功: '{old_text}' -> '{new_text}'")
                                return result
                            else:
                                self.logger.info(f"替换后文本无变化: '{old_text}'")
                else:
                    self.logger.info(f"模式不匹配: {field_name} in '{para_text}'")
            
            return result
            
        except Exception as e:
            self.logger.error(f"统一字段处理异常: {e}")
            return result
    
    def _handle_dual_field_table_layout(self, para_text: str, current_field: str, field_value: str) -> str:
        """
        智能双字段表格布局处理 - 简化版本
        处理类似 '电话                    电子邮件' 的格式
        """
        try:
            self.logger.info(f"进入双字段表格处理: current_field='{current_field}', para_text='{para_text}'")
            # 电话+邮件组合 - 优先处理紧邻格式
            if current_field in ['电话', '联系电话'] and ('电子邮件' in para_text or '邮箱' in para_text):
                
                # 首先尝试紧邻格式：电话：数字电子邮箱：
                email_pattern = r'(电话|联系电话)[:：]\s*([0-9\-]+)\s*(电子邮箱|电子邮件)[:：]\s*([_\s]*)'
                adjacent_match = re.search(email_pattern, para_text)
                if adjacent_match:
                    phone_field = adjacent_match.group(1)
                    phone_number = adjacent_match.group(2)
                    email_field = adjacent_match.group(3)
                    email_placeholder = adjacent_match.group(4)
                    
                    # 获取邮件字段值
                    email_value = self.company_info.get('email', '')
                    self.logger.info(f"获取邮件字段值(紧邻格式): '{email_value}' from company_info: {list(self.company_info.keys())}")
                    
                    result = f"{phone_field}：{phone_number}{email_field}：{email_value}"
                    self.logger.info(f"电话+邮箱紧邻格式处理: '{para_text}' -> '{result}'")
                    return result
                
                # 然后尝试空格分隔格式：电话                    电子邮件
                pattern = r'(电话|联系电话)(\s{8,})(电子邮件|电子邮箱|邮箱)'
                space_match = re.search(pattern, para_text)
                if space_match:
                    phone_field = space_match.group(1)
                    spaces = space_match.group(2)
                    email_field = space_match.group(3)
                    
                    # 获取邮件字段值
                    email_value = self.company_info.get('email', '')
                    self.logger.info(f"获取邮件字段值(空格格式): '{email_value}' from company_info: {list(self.company_info.keys())}")
                    
                    # 计算合适的空格数
                    phone_text = f"{phone_field}：{field_value}"
                    email_text = f"{email_field}：{email_value}" if email_value else f"{email_field}："
                    
                    # 保持美观间距（至少20个空格）
                    optimal_spaces = max(20, len(spaces) - (len(phone_text) - len(phone_field)))
                    space_str = ' ' * optimal_spaces
                    
                    result = f"{phone_text}{space_str}{email_text}"
                    self.logger.info(f"电话+邮箱空格格式处理: '{para_text}' -> '{result}'")
                    return result
            
            # 地址+邮编组合（紧邻格式：...室邮编：）
            elif current_field == '地址' and ('邮编' in para_text or '邮政编码' in para_text):
                # 处理 "地址内容邮编：_______" 格式
                postal_pattern = r'(.*?)(邮编|邮政编码)[:：]\s*([_\s]*)'
                match = re.search(postal_pattern, para_text)
                if match:
                    address_content = match.group(1)
                    postal_field = match.group(2)
                    postal_placeholder = match.group(3)
                    
                    # 获取邮政编码值
                    postal_value = self.company_info.get('postalCode', '') or '100006'
                    
                    result = f"{address_content}{postal_field}：{postal_value}"
                    self.logger.info(f"地址+邮编组合处理: '{para_text}' -> '{result}'")
                    return result
            
            # 电话+邮箱组合（紧邻格式：...000电子邮箱：）
            elif current_field in ['电话', '联系电话'] and ('电子邮箱' in para_text or '电子邮件' in para_text):
                # 处理 "电话: 数字电子邮箱：" 或 "电话：数字电子邮箱：" 格式
                # 注意：需要匹配到电话字段的完整内容，包括已有的电话号码
                email_pattern = r'(电话|联系电话)[:：]\s*([0-9\-]+)(电子邮箱|电子邮件)[:：]\s*([_\s]*)'
                match = re.search(email_pattern, para_text)
                if match:
                    phone_field = match.group(1)
                    phone_number = match.group(2)
                    email_field = match.group(3)
                    email_placeholder = match.group(4)
                    
                    # 获取邮件字段值
                    email_value = self.company_info.get('email', '')
                    self.logger.info(f"获取邮件字段值(紧邻格式): '{email_value}' from company_info: {list(self.company_info.keys())}")
                    
                    result = f"{phone_field}：{phone_number}{email_field}：{email_value}"
                    self.logger.info(f"电话+邮箱组合处理: '{para_text}' -> '{result}'")
                    return result
            
            # 地址+传真组合
            elif current_field == '地址' and '传真' in para_text:
                pattern = r'(地址)(\s{8,})(传真)'
                match = re.search(pattern, para_text)
                if match:
                    address_field = match.group(1)
                    spaces = match.group(2)
                    fax_field = match.group(3)
                    
                    fax_value = self.company_info.get('fax', '')
                    
                    address_text = f"{address_field}：{field_value}"
                    fax_text = f"{fax_field}：{fax_value}" if fax_value else f"{fax_field}："
                    
                    optimal_spaces = max(15, len(spaces) - (len(address_text) - len(address_field)))
                    space_str = ' ' * optimal_spaces
                    
                    result = f"{address_text}{space_str}{fax_text}"
                    return result
            
            return para_text
            
        except Exception as e:
            self.logger.error(f"双字段表格处理异常: {e}")
            return para_text
    
    def _is_purchaser_info(self, text: str, exclude_contexts: list) -> bool:
        """采购人信息识别 - 简化版本"""
        try:
            # 检查是否包含排除关键词
            for context in exclude_contexts:
                if context in text:
                    return True
            return False
        except:
            return False
    
    def _post_process_beautification(self, doc: Document, total_replacements: int) -> int:
        """
        后处理美化机制 - 简化版本
        统一清理和优化文档格式
        """
        if total_replacements == 0:
            return 0
        
        beautified_count = 0
        
        try:
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                
                if original_text.strip():
                    beautified_text = self._beautify_paragraph_text(original_text)
                    
                    if beautified_text != original_text:
                        paragraph.clear()
                        paragraph.add_run(beautified_text)
                        beautified_count += 1
                        
            return beautified_count
            
        except Exception as e:
            self.logger.error(f"后处理美化失败: {e}")
            return beautified_count
    
    def _beautify_paragraph_text(self, text: str) -> str:
        """段落美化处理 - 5种优化规则"""
        try:
            beautified = text
            
            # 1. 清理多余的冒号
            beautified = re.sub(r'[:：]{2,}', '：', beautified)
            
            # 2. 标准化冒号格式
            beautified = re.sub(r':([^:])', r'：\1', beautified)
            
            # 3. 清理尾部下划线
            beautified = re.sub(r'_{3,}$', '', beautified)
            
            # 4. 优化空格（保持表格对齐，但清理多余空格）
            if not re.search(r'\s{8,}', beautified):  # 不是表格式布局
                beautified = re.sub(r'\s{3,}', '  ', beautified)
            
            # 5. 清理首尾空格
            beautified = beautified.strip()
            
            return beautified
            
        except Exception as e:
            self.logger.error(f"段落美化异常: {e}")
            return text