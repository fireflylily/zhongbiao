#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
身份证插入器 - 处理身份证正反面表格插入
"""

import os
from pathlib import Path
from typing import Optional, Dict
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger
from .document_utils import DocumentUtils


class IdCardInserter:
    """身份证插入器 - 处理身份证正反面表格插入"""

    def __init__(self, utils: DocumentUtils):
        self.logger = get_module_logger("id_card_inserter")
        self.utils = utils

    def insert_id_card(self, doc: Document, front_path: str, back_path: str,
                       insert_point: Optional[Dict], id_type: str) -> bool:
        """
        插入身份证图片（正面和反面并排显示）

        支持两种模式：
        1. 如果段落后有现有表格，插入到表格单元格中
        2. 如果没有表格，创建新表格

        Args:
            doc: Word文档对象
            front_path: 身份证正面图片路径
            back_path: 身份证反面图片路径
            insert_point: 插入点信息
            id_type: 身份证类型（如 '法定代表人' 或 '被授权人'）

        Returns:
            bool: 插入是否成功
        """
        try:
            # 解析并验证图片路径（支持相对路径）
            if not front_path:
                self.logger.error(f"{id_type}身份证正面图片路径为空")
                return False

            front_path_resolved = self.utils.resolve_file_path(front_path)
            if not os.path.exists(front_path_resolved):
                self.logger.error(f"{id_type}身份证正面图片不存在: {front_path} (resolved: {front_path_resolved})")
                return False

            if not back_path:
                self.logger.error(f"{id_type}身份证反面图片路径为空")
                return False

            back_path_resolved = self.utils.resolve_file_path(back_path)
            if not os.path.exists(back_path_resolved):
                self.logger.error(f"{id_type}身份证反面图片不存在: {back_path} (resolved: {back_path_resolved})")
                return False

            # 使用解析后的路径
            front_path = front_path_resolved
            back_path = back_path_resolved

            # 使用7厘米宽度
            id_width_cm = 7

            if insert_point and insert_point['type'] == 'paragraph':
                # 在找到的段落位置插入
                target_para = insert_point['paragraph']

                # 检查段落后是否有现有表格
                existing_table = self.utils.find_next_table_after_paragraph(target_para)

                if existing_table:
                    # 模式1：使用现有表格
                    self.logger.info(f"检测到段落后有现有表格，将插入到表格中")
                    return self._insert_id_into_existing_table(
                        existing_table, front_path, back_path, id_width_cm, id_type
                    )
                else:
                    # 模式2：创建新表格
                    self.logger.info(f"段落后没有表格，将创建新表格")

                    # 【修复】先验证图片文件，避免后续失败
                    try:
                        # 验证正面图片
                        img_front = Image.open(front_path)
                        front_size = img_front.size
                        self.logger.info(f"  验证正面图片: {Path(front_path).name}, 尺寸={front_size}")
                        img_front.close()

                        # 验证反面图片
                        img_back = Image.open(back_path)
                        back_size = img_back.size
                        self.logger.info(f"  验证反面图片: {Path(back_path).name}, 尺寸={back_size}")
                        img_back.close()
                    except Exception as e:
                        self.logger.error(f"❌ 图片验证失败: {e}")
                        self.logger.error(f"  正面图片: {front_path}, 存在={os.path.exists(front_path)}")
                        self.logger.error(f"  反面图片: {back_path}, 存在={os.path.exists(back_path)}")
                        return False

                    # 【修复】使用简化的表格创建逻辑（避免复杂DOM操作）
                    try:
                        # 插入分页符
                        page_break_para = self.utils.insert_paragraph_after(target_para)
                        page_break_para.add_run().add_break()
                        self.logger.info(f"  ✓ 已插入分页符")

                        # 插入标题
                        title = self.utils.insert_paragraph_after(page_break_para)
                        title.text = f"{id_type}身份证"
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if title.runs:
                            title.runs[0].font.bold = True
                        self.logger.info(f"  ✓ 已插入标题: {id_type}身份证")

                        # 【关键修复】使用最简单可靠的方法：在文档末尾创建表格，然后移动到正确位置
                        # 这种方法避免了复杂的DOM操作，更加稳定
                        from docx.table import Table

                        # 【修复】确保文档有section（节），python-docx创建表格需要section信息
                        if len(doc.sections) == 0:
                            self.logger.warning(f"  ⚠️ 文档缺少section定义，正在添加默认section")
                            doc.add_section()
                            self.logger.info(f"  ✓ 已添加默认section")

                        # 方法1：直接在title后添加表格（最简单）
                        # 先创建一个临时段落
                        temp_para = self.utils.insert_paragraph_after(title)

                        # 在文档末尾创建表格
                        table = doc.add_table(rows=2, cols=2)
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # 将表格移动到临时段落的位置
                        table._element.getparent().remove(table._element)
                        temp_para._element.addprevious(table._element)

                        # 删除临时段落
                        temp_para._element.getparent().remove(temp_para._element)

                        self.logger.info(f"  ✓ 已创建表格 (2行x2列)")

                        # 第一行：标签
                        table.rows[0].cells[0].text = "正面"
                        table.rows[0].cells[1].text = "反面"
                        for cell in table.rows[0].cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if cell.paragraphs[0].runs:
                                cell.paragraphs[0].runs[0].font.bold = True
                        self.logger.info(f"  ✓ 已设置表格标题行")

                        # 第二行：图片
                        self.logger.info(f"  开始插入图片...")

                        # 插入正面图片
                        front_cell = table.rows[1].cells[0]
                        front_cell.text = ""
                        front_para = front_cell.paragraphs[0]
                        front_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        front_run = front_para.add_run()
                        front_run.add_picture(front_path, width=Cm(id_width_cm))
                        self.logger.info(f"  ✓ 正面图片已插入: {Path(front_path).name}")

                        # 插入反面图片
                        back_cell = table.rows[1].cells[1]
                        back_cell.text = ""
                        back_para = back_cell.paragraphs[0]
                        back_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        back_run = back_para.add_run()
                        back_run.add_picture(back_path, width=Cm(id_width_cm))
                        self.logger.info(f"  ✓ 反面图片已插入: {Path(back_path).name}")

                        self.logger.info(f"✅ 成功在指定位置插入{id_type}身份证（新建表格）")
                        return True

                    except Exception as table_error:
                        self.logger.error(f"❌ 创建表格或插入图片失败: {table_error}")
                        self.logger.error(f"  错误类型: {type(table_error).__name__}")
                        import traceback
                        self.logger.error(f"  完整堆栈:\n{traceback.format_exc()}")

                        # 【TODO】理想情况下应该回滚已插入的标题和分页符，但由于复杂性暂时保留
                        # 至少在日志中清晰标记失败
                        return False

            elif insert_point and insert_point['type'] == 'table_cell':
                # 【修复】处理表格单元格类型的插入点
                # 通过 table_index 从 doc.tables 获取表格对象
                table_idx = insert_point['table_index']
                self.logger.info(f"检测到table_cell类型插入点，表格索引={table_idx}")

                # 从文档中获取表格对象
                table = doc.tables[table_idx]

                # 直接使用现有的表格插入方法
                self.logger.info(f"将使用现有表格插入身份证图片")
                return self._insert_id_into_existing_table(
                    table, front_path, back_path, id_width_cm, id_type
                )

            else:
                # 降级：添加到文档末尾
                self.logger.info(f"未找到插入点，将在文档末尾创建{id_type}身份证")

                # 【修复】先验证图片文件
                try:
                    # 验证正面图片
                    img_front = Image.open(front_path)
                    front_size = img_front.size
                    self.logger.info(f"  验证正面图片: {Path(front_path).name}, 尺寸={front_size}")
                    img_front.close()

                    # 验证反面图片
                    img_back = Image.open(back_path)
                    back_size = img_back.size
                    self.logger.info(f"  验证反面图片: {Path(back_path).name}, 尺寸={back_size}")
                    img_back.close()
                except Exception as e:
                    self.logger.error(f"❌ 图片验证失败: {e}")
                    self.logger.error(f"  正面图片: {front_path}, 存在={os.path.exists(front_path)}")
                    self.logger.error(f"  反面图片: {back_path}, 存在={os.path.exists(back_path)}")
                    return False

                # 【修复】添加详细的步骤日志
                try:
                    doc.add_page_break()
                    self.logger.info(f"  ✓ 已添加分页符")

                    title = doc.add_paragraph(f"{id_type}身份证")
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if title.runs:
                        title.runs[0].font.bold = True
                    self.logger.info(f"  ✓ 已添加标题: {id_type}身份证")

                    # 【修复】确保文档有section（节），python-docx创建表格需要section信息
                    if len(doc.sections) == 0:
                        self.logger.warning(f"  ⚠️ 文档缺少section定义，正在添加默认section")
                        doc.add_section()
                        self.logger.info(f"  ✓ 已添加默认section")

                    # 创建表格（2行2列）
                    table = doc.add_table(rows=2, cols=2)
                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.logger.info(f"  ✓ 已创建表格 (2行x2列)")

                    # 第一行：标签
                    table.rows[0].cells[0].text = "正面"
                    table.rows[0].cells[1].text = "反面"
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].font.bold = True
                    self.logger.info(f"  ✓ 已设置表格标题行")

                    # 第二行：图片
                    self.logger.info(f"  开始插入图片...")

                    # 插入正面图片
                    front_cell = table.rows[1].cells[0]
                    front_cell.text = ""
                    front_para = front_cell.paragraphs[0]
                    front_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    front_run = front_para.add_run()
                    front_run.add_picture(front_path, width=Cm(id_width_cm))
                    self.logger.info(f"  ✓ 正面图片已插入: {Path(front_path).name}")

                    # 插入反面图片
                    back_cell = table.rows[1].cells[1]
                    back_cell.text = ""
                    back_para = back_cell.paragraphs[0]
                    back_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    back_run = back_para.add_run()
                    back_run.add_picture(back_path, width=Cm(id_width_cm))
                    self.logger.info(f"  ✓ 反面图片已插入: {Path(back_path).name}")

                    self.logger.info(f"✅ 在文档末尾插入{id_type}身份证成功")
                    return True

                except Exception as fallback_error:
                    self.logger.error(f"❌ 在文档末尾插入身份证失败: {fallback_error}")
                    self.logger.error(f"  错误类型: {type(fallback_error).__name__}")
                    import traceback
                    self.logger.error(f"  完整堆栈:\n{traceback.format_exc()}")
                    return False

        except Exception as e:
            self.logger.error(f"❌ 插入{id_type}身份证失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return False

    def _insert_id_into_existing_table(self, table, front_path: str, back_path: str,
                                       id_width_cm: float, id_type: str) -> bool:
        """
        将身份证图片插入到现有表格中

        Args:
            table: 现有表格对象
            front_path: 身份证正面图片路径
            back_path: 身份证反面图片路径
            id_width_cm: 图片宽度（厘米）
            id_type: 身份证类型

        Returns:
            bool: 插入是否成功
        """
        try:
            # 【修复】增强边界检查：验证表格结构
            if not table or not hasattr(table, 'columns') or not hasattr(table, 'rows'):
                self.logger.error(f"❌ 无效的表格对象")
                return False

            num_cols = len(table.columns)
            num_rows = len(table.rows)

            # 【修复】检查表格是否为空
            if num_cols == 0 or num_rows == 0:
                self.logger.error(f"❌ 表格为空: {num_rows}行 x {num_cols}列")
                return False

            self.logger.info(f"现有表格结构: {num_rows}行 x {num_cols}列")

            # 输出表格第一行的内容（标题行）
            if num_rows > 0:
                try:
                    header_texts = [cell.text.strip() for cell in table.rows[0].cells]
                    self.logger.info(f"表格标题行: {header_texts}")
                except Exception as e:
                    self.logger.warning(f"⚠️ 无法读取表格标题行: {e}")

            if num_cols >= 2:
                # 情况1: 表格有2列或更多列
                # 智能识别"头像面"和"国徽面"列
                front_col_idx = None
                back_col_idx = None

                # 扫描第一行，识别列标题
                if num_rows > 0:
                    for col_idx, cell in enumerate(table.rows[0].cells):
                        cell_text = cell.text.strip()

                        # 识别正面列（头像面）
                        if any(keyword in cell_text for keyword in ['头像面', '正面', '人像面']):
                            front_col_idx = col_idx
                            self.logger.info(f"✅ 识别到正面列: 第{col_idx}列 ('{cell_text}')")

                        # 识别反面列（国徽面）
                        if any(keyword in cell_text for keyword in ['国徽面', '反面', '国徽']):
                            back_col_idx = col_idx
                            self.logger.info(f"✅ 识别到反面列: 第{col_idx}列 ('{cell_text}')")

                # 降级策略：如果无法识别列标题，使用默认索引
                if front_col_idx is None or back_col_idx is None:
                    if num_cols == 2:
                        # 2列表格：假设 [正面, 反面]
                        front_col_idx = 0
                        back_col_idx = 1
                        self.logger.warning(f"⚠️ 无法识别列标题，使用默认2列模式: 正面=列0, 反面=列1")
                    else:
                        # 3+列表格：假设 [序号, 正面, 反面]（跳过第一列）
                        front_col_idx = 1
                        back_col_idx = 2
                        self.logger.warning(f"⚠️ 无法识别列标题，使用默认3+列模式: 正面=列1, 反面=列2")

                # 确定插入的行（优先第二行，即索引1）
                target_row_idx = 1 if num_rows >= 2 else 0

                # 【修复】边界检查：确保目标行存在
                if target_row_idx >= num_rows:
                    self.logger.error(f"❌ 目标行索引{target_row_idx}超出范围(总行数={num_rows})")
                    return False

                target_row = table.rows[target_row_idx]

                # 【修复】边界检查：确保列索引有效
                if front_col_idx >= num_cols or back_col_idx >= num_cols:
                    self.logger.error(
                        f"❌ 列索引超出范围: 正面列{front_col_idx}, 反面列{back_col_idx}, "
                        f"总列数={num_cols}"
                    )
                    return False

                self.logger.info(f"📍 将插入到: 行{target_row_idx}, 正面列{front_col_idx}, 反面列{back_col_idx}")

                # 【修复】增强错误处理：插入正面图片
                try:
                    front_cell = target_row.cells[front_col_idx]
                    front_cell.text = ""  # 清空现有文本
                    front_para = front_cell.paragraphs[0] if front_cell.paragraphs else front_cell.add_paragraph()
                    front_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    front_run = front_para.add_run()
                    front_run.add_picture(front_path, width=Cm(id_width_cm))
                    self.logger.info(f"  ✅ 正面图片已插入到列{front_col_idx}")
                except IndexError as e:
                    self.logger.error(
                        f"❌ 访问单元格失败: 行{target_row_idx}, 列{front_col_idx}, "
                        f"表格结构={num_rows}x{num_cols}, 错误: {e}"
                    )
                    return False
                except Exception as e:
                    self.logger.error(f"❌ 插入正面图片失败: {e}")
                    import traceback
                    self.logger.error(traceback.format_exc())
                    return False

                # 【修复】增强错误处理：插入反面图片
                try:
                    back_cell = target_row.cells[back_col_idx]
                    back_cell.text = ""  # 清空现有文本
                    back_para = back_cell.paragraphs[0] if back_cell.paragraphs else back_cell.add_paragraph()
                    back_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    back_run = back_para.add_run()
                    back_run.add_picture(back_path, width=Cm(id_width_cm))
                    self.logger.info(f"  ✅ 反面图片已插入到列{back_col_idx}")
                except IndexError as e:
                    self.logger.error(
                        f"❌ 访问单元格失败: 行{target_row_idx}, 列{back_col_idx}, "
                        f"表格结构={num_rows}x{num_cols}, 错误: {e}"
                    )
                    return False
                except Exception as e:
                    self.logger.error(f"❌ 插入反面图片失败: {e}")
                    import traceback
                    self.logger.error(traceback.format_exc())
                    return False

                self.logger.info(f"✅ 已将{id_type}身份证插入到现有表格（行{target_row_idx}，正面=列{front_col_idx}，反面=列{back_col_idx}）")
                return True

            elif num_cols == 1:
                # 情况2: 表格只有1列（垂直布局）
                # 需要找到"人像面"和"国徽面"标题行，分别在它们下方插入图片
                front_row_idx = None
                back_row_idx = None

                # 扫描表格，查找"人像面"和"国徽面"标题行
                for row_idx, row in enumerate(table.rows):
                    cell_text = row.cells[0].text.strip()

                    # 识别"人像面"标题行
                    if any(keyword in cell_text for keyword in ['人像面', '头像面', '正面']):
                        front_row_idx = row_idx
                        self.logger.info(f"✅ 识别到正面标题行: 第{row_idx}行 ('{cell_text}')")

                    # 识别"国徽面"标题行
                    if any(keyword in cell_text for keyword in ['国徽面', '反面', '国徽']):
                        back_row_idx = row_idx
                        self.logger.info(f"✅ 识别到反面标题行: 第{row_idx}行 ('{cell_text}')")

                # 【修复】增强错误处理：插入正面图片（在"人像面"标题的下一行）
                if front_row_idx is not None and front_row_idx + 1 < num_rows:
                    try:
                        front_cell = table.rows[front_row_idx + 1].cells[0]
                        front_cell.text = ""  # 清空现有文本
                        front_para = front_cell.paragraphs[0] if front_cell.paragraphs else front_cell.add_paragraph()
                        front_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        front_run = front_para.add_run()
                        front_run.add_picture(front_path, width=Cm(id_width_cm))
                        self.logger.info(f"✅ 已插入正面图片到第{front_row_idx + 1}行")
                    except IndexError as e:
                        self.logger.error(
                            f"❌ 访问单元格失败: 行{front_row_idx + 1}, 列0, "
                            f"表格结构={num_rows}x{num_cols}, 错误: {e}"
                        )
                    except Exception as e:
                        self.logger.error(f"❌ 插入正面图片失败: {e}")
                        import traceback
                        self.logger.error(traceback.format_exc())
                else:
                    self.logger.warning(f"⚠️ 未找到正面插入位置 (front_row_idx={front_row_idx}, num_rows={num_rows})")

                # 【修复】增强错误处理：插入反面图片（在"国徽面"标题的下一行）
                if back_row_idx is not None and back_row_idx + 1 < num_rows:
                    try:
                        back_cell = table.rows[back_row_idx + 1].cells[0]
                        back_cell.text = ""  # 清空现有文本
                        back_para = back_cell.paragraphs[0] if back_cell.paragraphs else back_cell.add_paragraph()
                        back_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        back_run = back_para.add_run()
                        back_run.add_picture(back_path, width=Cm(id_width_cm))
                        self.logger.info(f"✅ 已插入反面图片到第{back_row_idx + 1}行")
                    except IndexError as e:
                        self.logger.error(
                            f"❌ 访问单元格失败: 行{back_row_idx + 1}, 列0, "
                            f"表格结构={num_rows}x{num_cols}, 错误: {e}"
                        )
                    except Exception as e:
                        self.logger.error(f"❌ 插入反面图片失败: {e}")
                        import traceback
                        self.logger.error(traceback.format_exc())
                else:
                    self.logger.warning(f"⚠️ 未找到反面插入位置 (back_row_idx={back_row_idx}, num_rows={num_rows})")

                self.logger.info(f"✅ 已将{id_type}身份证插入到现有表格（1列垂直模式）")
                return True

            else:
                self.logger.error(f"表格列数异常: {num_cols}")
                return False

        except Exception as e:
            self.logger.error(f"插入到现有表格失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return False
