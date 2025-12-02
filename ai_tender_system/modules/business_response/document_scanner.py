#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档扫描器 - 扫描Word文档查找图片插入位置
"""

import sys
from pathlib import Path
from typing import Dict, Any
from docx import Document

sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger


class DocumentScanner:
    """文档扫描器 - 负责扫描Word文档查找图片插入位置"""

    def __init__(self):
        self.logger = get_module_logger("document_scanner")

        # 图片类型关键词映射
        self.image_keywords = {
            'license': ['营业执照', '营业执照副本', '执照'],
            'qualification': [],  # 清空通用关键词，只使用具体资质类型匹配（避免误匹配"相关资质证书"等泛指文字）
            'authorization': ['授权书', '授权委托书', '法人授权'],
            'certificate': ['证书', '认证', '资格证'],
            'legal_id': [
                '法定代表人身份证复印件', '法定代表人身份证', '法人身份证', '法定代表人身份证明',
                '法人代表身份证', '法人代表身份证复印件',
                '法定代表人居民身份证',  # 新增：正式表述
                '法人居民身份证'  # 新增：简化表述
            ],
            'auth_id': [
                '授权代表身份证', '授权人身份证', '被授权人身份证',
                '授权代表身份证复印件', '被授权人身份证复印件',
                '委托代理人身份证', '代理人身份证复印件',
                '被授权代表身份证', '被授权代表身份证复印件',  # 新增：被授权代表变体
                '授权代表人身份证', '授权代表人身份证复印件',  # 新增：授权代表人变体
                '授权代表居民身份证', '被授权人居民身份证',  # 新增：正式表述
                '身份证复印件',  # 新增：通用表述（优先级低，会在法人身份证后匹配）
                '身份证',  # 新增：最简化表述（优先级最低）
                '居民身份证'  # 新增：正式简化表述
            ],
            'dishonest_executor': ['失信被执行人', '失信被执行人名单'],
            'tax_violation_check': ['重大税收违法', '税收违法案件当事人名单'],
            # 注意：gov_procurement_creditchina 和 gov_procurement_ccgp 使用特殊AND逻辑处理，见下方扫描代码
            'audit_report': ['审计报告', '财务审计报告', '年度审计报告', '会计师事务所出具']
        }

    def scan_insert_points(self, doc: Document, image_config: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        扫描文档，查找图片插入点（两阶段识别法：核心词+上下文分类）

        Args:
            doc: Word文档对象
            image_config: 图片配置（可选），包含qualification_details用于精确匹配

        Returns:
            插入点字典，键可以是通用类型(license/qualification)或具体资质(iso9001/cmmi等)
        """
        import re

        # 候选位置字典：{img_type: [candidate_dict, ...]}
        candidates = {}

        # 从qualification_matcher导入映射表
        from .qualification_matcher import QUALIFICATION_MAPPING

        # 获取文档总段落数
        total_paragraphs = len(doc.paragraphs)

        # ===== 阶段1：扫描段落，基于核心词识别 =====
        self.logger.info(f"📄 开始扫描文档（共{total_paragraphs}个段落）")

        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # 获取段落样式名
            style_name = paragraph.style.name if paragraph.style else ''

            # ===== 1. 营业执照识别 =====
            if "营业执照" in text:
                category, bonus = self._classify_paragraph(text, para_idx, total_paragraphs, style_name)
                if category != 'exclude':
                    candidates.setdefault('license', []).append({
                        'type': 'paragraph',
                        'index': para_idx,
                        'paragraph': paragraph,
                        'category': category,
                        'bonus_score': bonus,
                        'text': text[:60]
                    })
                    self.logger.info(f"🔍 营业执照候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

            # ===== 2. 身份证识别（支持组合判断）=====
            if "身份证" in text:
                category, bonus = self._classify_paragraph(text, para_idx, total_paragraphs, style_name)
                if category != 'exclude':
                    # 判断是哪种身份证
                    has_legal = any(kw in text for kw in ["法定代表人", "法人", "法人代表"])
                    has_auth = any(kw in text for kw in ["授权", "被授权", "代理人", "委托"])

                    # 法人身份证
                    if has_legal:
                        candidates.setdefault('legal_id', []).append({
                            'type': 'paragraph',
                            'index': para_idx,
                            'paragraph': paragraph,
                            'category': category,
                            'bonus_score': bonus,
                            'text': text[:60]
                        })
                        self.logger.info(f"🔍 法人身份证候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

                    # 被授权人身份证
                    if has_auth:
                        candidates.setdefault('auth_id', []).append({
                            'type': 'paragraph',
                            'index': para_idx,
                            'paragraph': paragraph,
                            'category': category,
                            'bonus_score': bonus,
                            'text': text[:60]
                        })
                        self.logger.info(f"🔍 被授权人身份证候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

                    # 如果两者都没有，可能是通用身份证要求（两者都需要）
                    if not has_legal and not has_auth:
                        # 同时为两种身份证添加候选
                        for id_type in ['legal_id', 'auth_id']:
                            candidates.setdefault(id_type, []).append({
                                'type': 'paragraph',
                                'index': para_idx,
                                'paragraph': paragraph,
                                'category': category,
                                'bonus_score': bonus,
                                'text': text[:60]
                            })
                        self.logger.info(f"🔍 通用身份证候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

            # ===== 4. 授权书识别 =====
            if "授权" in text and ("授权书" in text or "授权委托书" in text):
                category, bonus = self._classify_paragraph(text, para_idx, total_paragraphs, style_name)
                if category != 'exclude':
                    candidates.setdefault('authorization', []).append({
                        'type': 'paragraph',
                        'index': para_idx,
                        'paragraph': paragraph,
                        'category': category,
                        'bonus_score': bonus,
                        'text': text[:60]
                    })
                    self.logger.info(f"🔍 授权书候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

            # ===== 5. 特殊处理：政府采购资质（使用优先级判断）=====
            #
            # 【重要】两种不同的资质类型（必须正确区分）：
            # 1. gov_procurement_creditchina: 在"信用中国"网站查询政府采购违法记录
            # 2. gov_procurement_ccgp: 在"中国政府采购网"查询政府采购违法记录
            #
            # 【测试案例A - 同时提到两个网站】：
            # 文本："响应方在'信用中国'网站...的信用报告及'中国政府采购网'...的政府采购严重违法失信行为信息记录"
            # 预期结果：
            #   - creditchina_credit_report（信用报告） ← 对应"信用中国"
            #   - gov_procurement_ccgp（政府采购网查询） ← 对应"政府采购网"
            #   - 不应该识别为 gov_procurement_creditchina ❌
            #
            # 【测试案例B - 只提信用中国】：
            # 文本："未被列入'信用中国'网站失信被执行人、重大税收违法、政府采购严重违法失信记录名单"
            # 预期结果：
            #   - dishonest_executor（失信被执行人）
            #   - tax_violation_check（税收违法）
            #   - gov_procurement_creditchina（信用中国的政府采购查询） ✅
            #
            # 【识别逻辑】：
            # 优先级1：明确提到"政府采购网"或"ccgp" → gov_procurement_ccgp
            # 优先级2：只提到"信用中国"（且没有政府采购网） → gov_procurement_creditchina
            #
            if "政府采购严重违法失信" in text:
                category, bonus = self._classify_paragraph(text, para_idx, total_paragraphs, style_name)
                if category != 'exclude':
                    # 优先级1：明确提到"政府采购网" → 识别为政府采购网查询
                    if "政府采购网" in text or "ccgp" in text.lower():
                        candidates.setdefault('gov_procurement_ccgp', []).append({
                            'type': 'paragraph',
                            'index': para_idx,
                            'paragraph': paragraph,
                            'category': category,
                            'bonus_score': bonus,
                            'text': text[:60]
                        })
                        self.logger.info(f"🔍 政府采购-政采网候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

                    # 优先级2：只提到"信用中国"（没有政府采购网）→ 识别为信用中国查询
                    elif "信用中国" in text or "creditchina" in text.lower():
                        candidates.setdefault('gov_procurement_creditchina', []).append({
                            'type': 'paragraph',
                            'index': para_idx,
                            'paragraph': paragraph,
                            'category': category,
                            'bonus_score': bonus,
                            'text': text[:60]
                        })
                        self.logger.info(f"🔍 政府采购-信用中国候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

            # ===== 6. 查找具体资质类型（ISO9001, CMMI等）=====
            for qual_key, qual_info in QUALIFICATION_MAPPING.items():
                # 跳过已特殊处理的政府采购资质
                if qual_key in ['gov_procurement_creditchina', 'gov_procurement_ccgp']:
                    continue

                keywords = qual_info.get('keywords', [])
                if any(keyword in text for keyword in keywords):
                    category, bonus = self._classify_paragraph(text, para_idx, total_paragraphs, style_name)
                    if category != 'exclude':
                        candidates.setdefault(qual_key, []).append({
                            'type': 'paragraph',
                            'index': para_idx,
                            'paragraph': paragraph,
                            'category': category,
                            'bonus_score': bonus,
                            'text': text[:60]
                        })
                        matched_kw = next((kw for kw in keywords if kw in text), keywords[0])
                        self.logger.info(f"🔍 {qual_key}候选: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 关键词='{matched_kw}'")
                    # 修复：删除break，允许一个标题匹配多个资质（如"基础电信业务经营许可证及增值电信业务经营许可证"）

        # ===== 扫描文本框中的插入点（特殊处理）=====
        self.logger.info(f"📦 开始扫描文本框...")
        textbox_candidates = self._scan_textboxes(doc)

        # 将文本框候选合并到candidates字典
        for img_type, textbox_list in textbox_candidates.items():
            for textbox_candidate in textbox_list:
                candidates.setdefault(img_type, []).append(textbox_candidate)

        # ===== 扫描表格中的身份证插入点（特殊处理）=====
        self.logger.info(f"📋 开始扫描表格（共{len(doc.tables)}个表格）")

        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue

                    # 身份证表格特殊处理（检测表格特征）
                    if "身份证" in cell_text:
                        # 检测是否为身份证表格（包含"正反面"、"头像面"等特征）
                        id_table_features = ['正、反面', '正反面', '头像面', '国徽面', '人像面']
                        is_id_table = any(feature in cell_text for feature in id_table_features)

                        if is_id_table:
                            # 判断是哪种身份证
                            has_legal = any(kw in cell_text for kw in ["法定代表人", "法人"])
                            has_auth = any(kw in cell_text for kw in ["授权", "被授权", "代理"])

                            # 法人身份证表格（优先级高）
                            if has_legal:
                                candidates.setdefault('legal_id', []).append({
                                    'type': 'table_cell',
                                    'table_index': table_idx,
                                    'cell': cell,
                                    'category': 'strong_attach',  # 表格特征明确，设为强附件
                                    'bonus_score': 0,  # 表格候选不额外加分
                                    'text': cell_text[:60]
                                })
                                self.logger.info(f"🔍 法人身份证表格: 表格#{table_idx}, 文本='{cell_text[:60]}'")

                            # 被授权人身份证表格
                            if has_auth:
                                candidates.setdefault('auth_id', []).append({
                                    'type': 'table_cell',
                                    'table_index': table_idx,
                                    'cell': cell,
                                    'category': 'strong_attach',
                                    'bonus_score': 0,  # 表格候选不额外加分
                                    'text': cell_text[:60]
                                })
                                self.logger.info(f"🔍 被授权人身份证表格: 表格#{table_idx}, 文本='{cell_text[:60]}'")

                            # 通用身份证表格
                            if not has_legal and not has_auth:
                                for id_type in ['legal_id', 'auth_id']:
                                    candidates.setdefault(id_type, []).append({
                                        'type': 'table_cell',
                                        'table_index': table_idx,
                                        'cell': cell,
                                        'category': 'strong_attach',
                                        'bonus_score': 0,  # 表格候选不额外加分
                                        'text': cell_text[:60]
                                    })
                                self.logger.info(f"🔍 通用身份证表格: 表格#{table_idx}, 文本='{cell_text[:60]}'")

        # ===== 阶段2：选择最佳位置（基于分类优先级）=====
        self.logger.info(f"📊 开始选择最佳插入位置...")

        # 定义分类优先级（数字越大越优先）
        category_priority = {
            'strong_attach': 100,        # 强附件标记（最优）
            'weak_attach': 80,           # 弱附件标记
            'neutral': 50,               # 中性位置
            'chapter': 30,               # 章节标题
            'toc': 10,                   # 目录
            'reference': 5,              # 正文引用
            'requirement_clause': -10,   # 招标要求条款（可用但不推荐）
            'header_noise': -50,         # 文档标题噪音（基本不用）
            'exclude': -999,             # 绝对排除（页眉页脚、附件清单标题）
        }

        insert_points = {}

        for img_type, candidate_list in candidates.items():
            if not candidate_list:
                self.logger.warning(f"⚠️ {img_type}未找到任何候选位置，将使用降级策略（文档末尾）")
                continue

            # 按优先级选择最佳候选
            # 排序规则：1. 总分（类别+奖励） 2. 文本简短（简短优先） 3. 位置靠后（靠后优先）
            best_candidate = max(candidate_list, key=lambda x: (
                category_priority.get(x['category'], 0) + x.get('bonus_score', 0),  # 总分 = 基础分 + 奖励分
                -len(x['text']),                          # 文本越短越好（负号实现）
                x.get('index', 0)                         # 位置越靠后越好
            ))

            best_category = best_candidate['category']
            best_priority = category_priority.get(best_category, 0)
            best_bonus = best_candidate.get('bonus_score', 0)
            total_score = best_priority + best_bonus

            # 构建插入点信息
            insert_point = {
                'type': best_candidate['type'],
                'category': best_category,
                'matched_keyword': best_candidate.get('text', '')[:30]
            }

            if best_candidate['type'] == 'paragraph':
                insert_point['index'] = best_candidate['index']
                insert_point['paragraph'] = best_candidate['paragraph']
            elif best_candidate['type'] == 'table_cell':
                insert_point['table_index'] = best_candidate['table_index']
                insert_point['cell'] = best_candidate['cell']

            insert_points[img_type] = insert_point

            # 友好的日志输出（根据总分级别 - 包含奖励分）
            # 显示格式: [category+bonus] 总分=xxx
            score_display = f"[{best_category}+{best_bonus}] 总分={total_score}" if best_bonus > 0 else f"[{best_category}] 总分={total_score}"

            if total_score >= 80:
                self.logger.info(
                    f"✅ {img_type}: 找到优质位置 {score_display} "
                    f"'{best_candidate['text']}' (共{len(candidate_list)}个候选)"
                )
            elif total_score >= 30:
                self.logger.info(
                    f"☑️ {img_type}: 找到可用位置 {score_display} "
                    f"'{best_candidate['text']}' (共{len(candidate_list)}个候选)"
                )
            elif total_score >= 0:
                self.logger.warning(
                    f"⚠️ {img_type}: 找到次优位置 {score_display} "
                    f"'{best_candidate['text']}' (共{len(candidate_list)}个候选)"
                )
            else:
                self.logger.warning(
                    f"🔻 {img_type}: 仅找到低质量位置 {score_display} "
                    f"'{best_candidate['text']}' (共{len(candidate_list)}个候选)"
                )

        # 输出扫描总结
        self.logger.info(f"📊 扫描完成: 找到 {len(insert_points)} 个插入点 - {list(insert_points.keys())}")
        return insert_points

    def _classify_paragraph(self, text: str, para_idx: int, total_paras: int,
                           style_name: str = ''):
        """
        段落分类（符合人的判断逻辑）+ 质量评分

        分类优先级（从高到低）：
        1. strong_attach      - 强附件标记（编号附件、附件标题） 100分
        2. weak_attach        - 弱附件标记（说明性文字、"后附"） 80分
        3. neutral            - 中性位置（普通段落） 50分
        4. chapter            - 章节标题（不理想但可接受） 30分
        5. toc                - 目录（很不理想） 10分
        6. reference          - 正文引用（最不理想） 5分
        7. requirement_clause - 招标要求条款（可用但不推荐） -10分
        8. header_noise       - 文档标题噪音（基本不用） -50分
        9. exclude            - 绝对排除（页眉页脚、附件清单标题） -999分

        质量评分（bonus_score，0-50分）：
        - 身份证专业术语（国徽面、人像面）：+20分
        - 插入指示词（扫描件、复印件）：+10分
        - 操作指引（需同时提供）：+5分
        - 格式说明句式（如提供...）：+5分

        Args:
            text: 段落文本
            para_idx: 段落索引
            total_paras: 文档总段落数
            style_name: Word样式名（可选）

        Returns:
            (category: str, bonus_score: int) 元组
            - category: 分类字符串
            - bonus_score: 质量奖励分（0-50分）
        """
        import re

        # ========== 1. exclude（绝对排除 - 仅保留技术性禁区）==========

        # 页眉页脚（技术禁区）
        if style_name and ('Header' in style_name or 'Footer' in style_name):
            return ('exclude', 0)

        # 附件清单标题（语义禁区）
        if "附件清单" in text or "附件目录" in text:
            return ('exclude', 0)

        # ========== 2. strong_attach（强附件标记）==========

        # 编号附件（最强信号）- "5-1 营业执照"
        if re.match(r'^\d+[-.]?\d*\s+', text):
            return ('strong_attach', 0)

        # 附件标题 - "附件：营业执照"、"附：营业执照"
        if text.startswith("附件") or text.startswith("附："):
            # 短文本（<50字符）直接识别为强附件
            if len(text) < 50:
                return ('strong_attach', 0)

            # 长文本 - 需要智能判断 + 质量评分
            elif any(kw in text for kw in ['身份证', '营业执照', '资质', '证书', '许可证', '授权书', '合同', '报告']):
                bonus_score = 0

                # 🌟 质量加分1：身份证专业术语（+20分）
                id_card_terms = ['国徽面', '人像面', '正、反面', '正反面', '头像面']
                if any(term in text for term in id_card_terms):
                    bonus_score += 20

                # 🌟 质量加分2：明确的插入指示词（+10分）
                insert_indicators = ['扫描件', '复印件', '影印件', '照片', '彩色扫描件']
                if any(indicator in text for indicator in insert_indicators):
                    bonus_score += 10

                # 🌟 质量加分3：操作指引性语言（+5分）
                instruction_patterns = ['需同时提供', '需提供', '应提供', '须提供', '请提供']
                if any(pattern in text for pattern in instruction_patterns):
                    bonus_score += 5

                # 🌟 质量加分4：格式说明句式（+5分）
                if '(如提供' in text or '（如提供' in text:
                    bonus_score += 5

                return ('strong_attach', bonus_score)

        # ========== 3. weak_attach（弱附件标记）==========

        # 说明性指示
        if any(pattern in text for pattern in [
            "后附", "如下", "见下", "以下为", "如下所示", "见后"
        ]) and len(text) < 50:
            return ('weak_attach', 0)

        # 包含"附件"但较长（可能是附件说明）
        if "附件" in text and 20 < len(text) < 80:
            return ('weak_attach', 0)

        # ========== 4. chapter（章节标题）==========

        # 检测章节标题
        is_chapter = any([
            text.startswith("第") and ("章" in text or "节" in text or "部分" in text),
            re.match(r'^[一二三四五六七八九十]+[、．.]', text),
            'Heading' in style_name,  # Word样式为标题
        ])

        if is_chapter:
            # 特殊情况：小节标题且简短，可能是插入点
            # 如 "5.1 营业执照副本"
            if re.match(r'^\d+\.\d+', text) and len(text) < 30:
                return ('weak_attach', 0)  # 升级为弱附件
            return ('chapter', 0)

        # ========== 5. toc（目录）==========

        if any([
            "目录" in text,
            "......" in text or "…………" in text,  # 目录特征
            para_idx < total_paras * 0.05,  # 文档前5%
            "TOC" in style_name,  # Word目录样式
        ]):
            return ('toc', 0)

        # ========== 6. reference（正文引用）==========

        # 正文中的引用/描述
        if any(keyword in text for keyword in [
            "根据", "依据", "按照", "参照",
            "记载", "所示", "显示", "颁发的",
        ]) and len(text) > 30:  # 较长的句子
            return ('reference', 0)

        # ========== 7. requirement_clause（招标要求条款 -10分）==========
        # 📌 从exclude降级为负分评分项

        # 招标文件的要求条款
        if any(pattern in text for pattern in [
            "须在响应文件中提供",
            "应在投标文件中提供",
            "投标人须提供", "响应方须提供",
            "投标人需提供", "响应方需提供",
        ]):
            return ('requirement_clause', 0)

        if ("如响应方" in text or "如投标人" in text) and "须" in text:
            return ('requirement_clause', 0)

        # ========== 8. header_noise（文档标题噪音 -50分）==========

        # 文档开头的极短文本
        if len(text) < 10 and para_idx < 3:
            # 如果是关键词标题，正常评分
            if any(kw in text for kw in ['营业执照', '身份证', '授权书', '资质', '证书']):
                return ('neutral', 0)
            else:
                return ('header_noise', 0)

        # ========== 9. 通用质量检测（检测强插入信号）==========
        # 🌟 新增：不仅检测"附："开头，还要检测其他强信号

        # 检测"粘贴处"等最强插入信号
        strong_insert_markers = ['粘贴处', '粘贴页', '贴在此处', '贴于此处', '张贴处']
        has_strong_marker = any(marker in text for marker in strong_insert_markers)

        # 检测插入指示词
        insert_indicators = ['扫描件', '复印件', '影印件', '照片', '彩色扫描件']
        has_insert_indicator = any(indicator in text for indicator in insert_indicators)

        # 检测资质关键词
        qual_keywords = ['身份证', '营业执照', '资质', '证书', '许可证', '授权书', '合同', '报告']
        has_qual_keyword = any(kw in text for kw in qual_keywords)

        # 如果包含强插入信号，提升为strong_attach并计算奖励分
        if has_strong_marker and has_qual_keyword:
            bonus_score = 0

            # 🎯 "粘贴处"是最强、最明确的位置信号 +30分（提高权重！）
            # 理由：这是直接的空间标记，明确告诉你"在这里贴"，无需任何推断
            if has_strong_marker:
                bonus_score += 30

            # 插入指示词 +10分
            if has_insert_indicator:
                bonus_score += 10

            # 身份证专业术语 +20分
            id_card_terms = ['国徽面', '人像面', '正、反面', '正反面', '头像面']
            if any(term in text for term in id_card_terms):
                bonus_score += 20

            return ('strong_attach', bonus_score)

        # ========== 10. neutral（中性位置 - 默认）==========

        return ('neutral', 0)

    # ==================== 🆕 案例表格处理相关方法 ====================

    def scan_case_requirements(self, doc: Document) -> list:
        """
        扫描文档中"格式自拟"的业绩案例要求

        功能：
        1. 识别包含"案例"+"格式自拟"关键词的段落
        2. 智能去重：检查附近是否已有案例表格
        3. 返回需要生成表格的位置列表

        识别规则：
        - 案例关键词：业绩案例、资格案例、项目案例、项目经验等
        - 格式自拟关键词：格式自拟、自拟格式、自行编制、格式不限等

        智能去重：
        - 检查段落后10个段落范围内是否已有案例表格
        - 如果有表格，跳过（避免重复生成）
        - 如果无表格，标记为需要生成

        Args:
            doc: Word文档对象

        Returns:
            案例要求位置列表
            [
                {
                    'type': 'paragraph',
                    'paragraph': paragraph对象,
                    'index': 段落索引,
                    'text': '八、资格案例（加盖公章）',
                    'requirement_text': '格式自拟',
                    'insert_position': 'after'
                },
                ...
            ]
        """
        case_requirements = []

        # 案例相关关键词
        case_keywords = [
            '业绩案例', '资格案例', '项目案例', '类似案例',
            '业绩', '项目经验', '同类项目', '以往项目',
            '项目实施经验', '完成的项目', '项目业绩'
        ]

        # "格式自拟"指示词
        format_free_keywords = [
            '格式自拟', '自拟格式', '自行编制',
            '自行设计', '格式不限', '自定义格式',
            '格式由投标人自定', '按投标人格式', '自行提供格式'
        ]

        self.logger.info(f"📋 开始扫描格式自拟的案例要求...")

        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # 检查是否包含案例关键词
            has_case_keyword = any(kw in text for kw in case_keywords)

            # 检查是否包含"格式自拟"指示
            has_format_free = any(kw in text for kw in format_free_keywords)

            if has_case_keyword:
                # 情况1：当前段落同时包含案例+格式自拟
                if has_format_free:
                    # 智能去重：检查附近是否已有案例表格
                    nearby_has_case_table = self._check_nearby_case_table(
                        doc, para_idx, search_range=10
                    )

                    if nearby_has_case_table:
                        self.logger.info(
                            f"⏭️  段落#{para_idx}附近已有案例表格，跳过生成: '{text[:60]}'"
                        )
                        continue

                    # 需要生成表格
                    case_requirements.append({
                        'type': 'paragraph',
                        'paragraph': paragraph,
                        'index': para_idx,
                        'text': text,
                        'requirement_text': '格式自拟',
                        'insert_position': 'after',
                        'reason': '当前段落包含案例+格式自拟'
                    })
                    self.logger.info(
                        f"🔍 识别到需要生成案例表格: 段落#{para_idx}, '{text[:60]}'"
                    )

                # 情况2：下一段落包含"格式自拟"（标题和内容分段的情况）
                elif para_idx + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[para_idx + 1]
                    next_text = next_para.text.strip()
                    if any(kw in next_text for kw in format_free_keywords):
                        # 智能去重
                        nearby_has_case_table = self._check_nearby_case_table(
                            doc, para_idx + 1, search_range=10
                        )

                        if nearby_has_case_table:
                            self.logger.info(
                                f"⏭️  段落#{para_idx}附近已有案例表格，跳过生成: '{text[:60]}'"
                            )
                            continue

                        # 使用下一段落作为插入点
                        case_requirements.append({
                            'type': 'paragraph',
                            'paragraph': next_para,
                            'index': para_idx + 1,
                            'text': text,
                            'requirement_text': next_text,
                            'insert_position': 'after',
                            'reason': '案例标题在当前段，格式说明在下一段'
                        })
                        self.logger.info(
                            f"🔍 识别到需要生成案例表格: 段落#{para_idx}, '{text[:60]}' "
                            f"(格式说明在下一段)"
                        )

        self.logger.info(f"📊 扫描完成: 识别到 {len(case_requirements)} 处需要生成案例表格的位置")
        return case_requirements

    def _check_nearby_case_table(self, doc: Document, para_idx: int,
                                 search_range: int = 10) -> bool:
        """
        检查指定段落附近是否已有案例表格

        检测范围：
        - 向后搜索N个段落（默认10个）
        - 如果遇到新的章节标题，停止搜索

        检测方法：
        - 遍历文档中的所有表格
        - 检查表格位置是否在搜索范围内
        - 使用CaseTableFiller的识别逻辑判断是否为案例表格

        Args:
            doc: Word文档对象
            para_idx: 起始段落索引
            search_range: 向后搜索的段落数量

        Returns:
            True: 附近有案例表格
            False: 附近无案例表格
        """
        # 计算搜索范围的结束段落索引
        end_idx = min(para_idx + search_range, len(doc.paragraphs))

        # 获取搜索范围内的段落元素
        search_paragraphs = []
        for i in range(para_idx, end_idx):
            if i >= len(doc.paragraphs):
                break

            para = doc.paragraphs[i]

            # 如果遇到新的章节标题，停止搜索
            if self._is_chapter_title(para.text):
                self.logger.debug(f"  遇到新章节，停止搜索: {para.text[:30]}")
                break

            search_paragraphs.append(para._element)

        if not search_paragraphs:
            return False

        # 遍历文档中的所有表格
        for table in doc.tables:
            table_element = table._element

            # 检查表格是否在搜索范围内
            if self._is_table_in_range(table_element, search_paragraphs):
                # 使用CaseTableFiller的识别逻辑
                # 临时导入，避免循环依赖
                try:
                    from .case_table_filler import CaseTableFiller
                    temp_filler = CaseTableFiller(None)  # 临时创建，只用于识别

                    if temp_filler._is_case_table(table):
                        self.logger.debug(f"  ✅ 在搜索范围内找到案例表格")
                        return True
                except ImportError:
                    self.logger.warning("  ⚠️ 无法导入CaseTableFiller，跳过表格检测")
                    return False

        self.logger.debug(f"  ❌ 搜索范围内未找到案例表格")
        return False

    def _is_table_in_range(self, table_element, paragraph_elements: list) -> bool:
        """
        检查表格是否在段落范围内

        通过遍历父元素的子元素，检查表格是否出现在段落之后

        Args:
            table_element: 表格元素
            paragraph_elements: 段落元素列表

        Returns:
            True: 表格在范围内
            False: 表格不在范围内
        """
        if not paragraph_elements:
            return False

        # 获取第一个段落的父元素（通常是body或section）
        parent = paragraph_elements[0].getparent()
        if parent is None:
            return False

        # 标记是否进入搜索范围
        in_range = False

        for child in parent:
            # 如果遇到起始段落，开始搜索
            if child == paragraph_elements[0]:
                in_range = True
                continue

            # 如果在范围内且遇到表格，检查是否是目标表格
            if in_range and child.tag.endswith('}tbl'):
                if child == table_element:
                    return True

            # 如果遇到最后一个搜索段落之后的段落，停止
            if in_range and child in paragraph_elements:
                # 还在范围内，继续
                continue
            elif in_range and child not in paragraph_elements:
                # 检查是否已经超出范围
                # （超出最后一个段落）
                found_last = False
                for para_elem in paragraph_elements:
                    if child.getprevious() == para_elem:
                        found_last = True
                        break
                if found_last:
                    # 已经超出范围
                    break

        return False

    def _is_chapter_title(self, text: str) -> bool:
        """
        判断是否为章节标题

        常见章节标题特征：
        - 一、二、三、...
        - 第一章、第二章、...
        - 1.、2.、3.、...

        Args:
            text: 段落文本

        Returns:
            True: 是章节标题
            False: 不是章节标题
        """
        import re

        text = text.strip()
        if not text:
            return False

        patterns = [
            r'^[一二三四五六七八九十]+[、．.]',  # 一、 二、
            r'^第[一二三四五六七八九十]+章',      # 第一章
            r'^第[一二三四五六七八九十]+节',      # 第一节
            r'^\d+[、．.]',                        # 1、 2、
            r'^\d+\.\d+',                          # 1.1 2.3
        ]

        for pattern in patterns:
            if re.match(pattern, text):
                return True

        return False

    def _scan_textboxes(self, doc: Document) -> Dict[str, list]:
        """
        扫描文档中的文本框，查找插入点

        文本框通常用于占位符，如"法定代表人身份证扫描件粘贴处"
        这些文本框是非常明确的插入位置指示

        Args:
            doc: Word文档对象

        Returns:
            候选位置字典：{img_type: [candidate_dict, ...]}
        """
        from docx.oxml.ns import qn

        candidates = {}
        textbox_count = 0

        # 遍历所有段落，查找包含文本框的段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_elem = paragraph._element

            # 查找文本框 (w:txbxContent)
            textboxes = para_elem.findall('.//w:txbxContent', namespaces={
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            })

            if not textboxes:
                continue

            # 提取文本框内容
            for tb_idx, textbox in enumerate(textboxes):
                # 提取文本框中的所有文本
                tb_text_elements = textbox.findall('.//w:t', namespaces={
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                })
                text = ''.join([t.text for t in tb_text_elements if t.text])
                text = text.strip()

                if not text:
                    continue

                textbox_count += 1

                # 使用质量评分系统评估文本框
                category, bonus = self._classify_paragraph(text, para_idx, len(doc.paragraphs), '')

                # 文本框通常是明确的插入位置，如果分类不是exclude，应该考虑
                if category == 'exclude':
                    continue

                # 识别文本框类型
                # ===== 1. 身份证识别 =====
                if "身份证" in text:
                    # 判断是哪种身份证
                    has_legal = any(kw in text for kw in ["法定代表人", "法人", "法人代表"])
                    has_auth = any(kw in text for kw in ["授权", "被授权", "代理人", "委托"])

                    # 法人身份证
                    if has_legal:
                        candidates.setdefault('legal_id', []).append({
                            'type': 'textbox',
                            'index': para_idx,
                            'paragraph': paragraph,
                            'category': category,
                            'bonus_score': bonus,
                            'text': text[:60]
                        })
                        self.logger.info(f"🔍 法人身份证文本框: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

                    # 被授权人身份证
                    if has_auth:
                        candidates.setdefault('auth_id', []).append({
                            'type': 'textbox',
                            'index': para_idx,
                            'paragraph': paragraph,
                            'category': category,
                            'bonus_score': bonus,
                            'text': text[:60]
                        })
                        self.logger.info(f"🔍 被授权人身份证文本框: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

                    # 通用身份证
                    if not has_legal and not has_auth:
                        for id_type in ['legal_id', 'auth_id']:
                            candidates.setdefault(id_type, []).append({
                                'type': 'textbox',
                                'index': para_idx,
                                'paragraph': paragraph,
                                'category': category,
                                'bonus_score': bonus,
                                'text': text[:60]
                            })
                        self.logger.info(f"🔍 通用身份证文本框: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

                # ===== 2. 营业执照识别 =====
                elif "营业执照" in text:
                    candidates.setdefault('license', []).append({
                        'type': 'textbox',
                        'index': para_idx,
                        'paragraph': paragraph,
                        'category': category,
                        'bonus_score': bonus,
                        'text': text[:60]
                    })
                    self.logger.info(f"🔍 营业执照文本框: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

                # ===== 3. 授权书识别 =====
                elif "授权" in text and ("授权书" in text or "授权委托书" in text):
                    candidates.setdefault('authorization', []).append({
                        'type': 'textbox',
                        'index': para_idx,
                        'paragraph': paragraph,
                        'category': category,
                        'bonus_score': bonus,
                        'text': text[:60]
                    })
                    self.logger.info(f"🔍 授权书文本框: 段落#{para_idx}, 类别={category}, 奖励分={bonus}, 文本='{text[:60]}'")

        if textbox_count > 0:
            self.logger.info(f"📦 文本框扫描完成: 找到 {textbox_count} 个文本框")
        else:
            self.logger.info(f"📦 文本框扫描完成: 未找到文本框")

        return candidates
