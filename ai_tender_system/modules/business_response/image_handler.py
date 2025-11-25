#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
图片插入模块 - 处理商务应答模板中的图片插入
包括公司公章、资质证明等图片的插入
"""

import os
from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger

# 导入拆分后的模块
from .document_utils import DocumentUtils
from .document_scanner import DocumentScanner
from .id_card_inserter import IdCardInserter
from .qualification_validator import QualificationValidator


class ImageHandler:
    """图片处理器"""

    def __init__(self):
        self.logger = get_module_logger("image_handler")

        # 默认图片尺寸（英寸）
        self.default_sizes = {
            'license': (6, 0),    # 营业执照：宽6英寸（约15.24厘米）
            'qualification': (6, 0),  # 资质证书：宽6英寸（约15.24厘米）
            'authorization': (6, 0),   # 授权书：宽6英寸（约15.24厘米）
            'certificate': (6, 0),      # 其他证书：宽6英寸（约15.24厘米）
            'legal_id': (2.165, 0),  # 法人身份证：宽2.165英寸（5.5厘米）
            'auth_id': (2.165, 0),    # 被授权人身份证：宽2.165英寸（5.5厘米）
            'dishonest_executor': (6, 0),              # 失信被执行人查询截图：宽6英寸
            'tax_violation_check': (6, 0),             # 税收违法查询截图：宽6英寸
            'gov_procurement_creditchina': (6, 0),     # 信用中国政采查询截图：宽6英寸
            'gov_procurement_ccgp': (6, 0),            # 政府采购网查询截图：宽6英寸
            'creditchina_report': (6, 0),              # 信用报告（信用中国）：宽6英寸
            'enterprise_credit_report': (6, 0),        # 国家企业信用信息公示系统信息报告：宽6英寸
            'audit_report': (6, 0)                     # 审计报告：宽6英寸
        }

        # 组合各模块
        self.utils = DocumentUtils()
        self.scanner = DocumentScanner()
        self.id_card_inserter = IdCardInserter(self.utils, self.default_sizes)
        self.validator = QualificationValidator(self.utils, self.default_sizes)

    def insert_images(self, doc: Document, image_config: Dict[str, Any],
                     required_quals: List[Dict] = None) -> Dict[str, Any]:
        """
        插入图片主方法（模板驱动 + 统计追踪）

        核心逻辑：
        1. 扫描模板占位符
        2. 填充所有有文件的占位符（成功填充）
        3. 记录有占位符但无文件的资质（缺失资质）
        4. 追加项目要求但模板没有占位符的资质（追加资质）

        Args:
            doc: Word文档对象
            image_config: 图片配置信息，包含所有资质
                {
                    'license_path': '营业执照路径',
                    'qualification_paths': ['资质证书路径列表'],
                    'qualification_details': [  # 资质详细信息
                        {
                            'qual_key': 'iso9001',
                            'file_path': '/path/to/iso9001.jpg',
                            'insert_hint': 'ISO9001质量管理体系'
                        }
                    ]
                }
            required_quals: 项目资格要求列表（可选，用于追加和统计）

        Returns:
            详细统计信息：
            {
                'images_inserted': 10,
                'images_types': ['营业执照', 'iso9001', ...],
                'errors': [],
                'filled_qualifications': [{'qual_key': 'iso9001', 'qual_name': '...'}],
                'missing_qualifications': [{'qual_key': 'cmmi', 'qual_name': '...'}],
                'appended_qualifications': [{'qual_key': 'level_protection', ...}]
            }
        """
        # 初始化统计数据
        stats = {
            'images_inserted': 0,
            'images_types': [],
            'errors': [],
            'filled_qualifications': [],      # 成功填充的资质
            'missing_qualifications': [],     # 缺失的资质（有占位符无文件）
            'appended_qualifications': []     # 追加的资质（项目要求但无占位符）
        }

        # 扫描文档，查找图片插入位置
        insert_points = self.scanner.scan_insert_points(doc, image_config)

        # 从qualification_matcher导入映射表（用于获取资质名称）
        from .qualification_matcher import QUALIFICATION_MAPPING

        # 【重构】构建统一的资质列表
        all_resources = self._build_resource_list(image_config)
        self.logger.info(f"📋 构建资质列表完成，共 {len(all_resources)} 项资质待插入")

        # 【重构】统一循环插入所有资质
        for idx, resource in enumerate(all_resources):
            resource_key = resource.get('key')
            # 查找插入点（优先使用具体key，降级使用通用key）
            insert_point = insert_points.get(resource_key)
            if not insert_point and resource.get('type') == 'single_image' and resource_key != 'license':
                # 资质证书可以降级使用通用 'qualification' 插入点
                insert_point = insert_points.get('qualification')

            # 调用统一分发方法
            self._insert_resource(doc, resource, insert_point, stats, idx)

        # 步骤：检测缺失的资质（模板有占位符但公司无文件）
        self.validator.detect_missing_qualifications(insert_points, image_config, stats, QUALIFICATION_MAPPING)

        # 步骤：追加项目要求但模板没有占位符的资质
        if required_quals:
            self.validator.append_required_qualifications(
                doc, required_quals, insert_points, image_config, stats, QUALIFICATION_MAPPING
            )

        # 输出统计摘要
        self.logger.info(f"📊 图片插入完成:")
        self.logger.info(f"  - 插入图片: {stats['images_inserted']}张")
        self.logger.info(f"  - 成功填充资质: {len(stats['filled_qualifications'])}个")
        self.logger.info(f"  - 缺失资质: {len(stats['missing_qualifications'])}个")
        self.logger.info(f"  - 追加资质: {len(stats['appended_qualifications'])}个")

        return stats

    def validate_images(self, image_paths: List[str]) -> Dict[str, Any]:
        """验证图片文件（委托给validator）"""
        return self.validator.validate_images(image_paths)

    def _build_resource_list(self, image_config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        构建统一的资质列表（用于循环插入）

        将image_config转换为统一格式的资质列表，每个资质包含：
        - type: 资质类型 ('single_image' 或 'id_card')
        - key: 资质键（用于查找插入点）
        - title: 显示标题
        - metadata: 其他信息（路径、页码等）

        Args:
            image_config: 图片配置字典

        Returns:
            统一格式的资质列表
        """
        resources = []

        # 1. 营业执照
        if image_config.get('license_path'):
            resources.append({
                'type': 'single_image',
                'key': 'license',
                'path': image_config['license_path'],
                'title': '营业执照副本',
                'qual_key': 'license',
                'insert_hint': None,
                'is_first_page': True,
                'page_num': 1
            })

        # 2. 被授权人身份证
        auth_id = image_config.get('auth_id')
        if auth_id and isinstance(auth_id, dict):
            front = auth_id.get('front')
            back = auth_id.get('back')
            if front and back:
                resources.append({
                    'type': 'id_card',
                    'key': 'auth_id',
                    'front': front,
                    'back': back,
                    'title': '被授权人身份证',
                    'id_type': '被授权人'
                })

        # 3. 法人身份证
        legal_id = image_config.get('legal_id')
        if legal_id and isinstance(legal_id, dict):
            front = legal_id.get('front')
            back = legal_id.get('back')
            if front and back:
                resources.append({
                    'type': 'id_card',
                    'key': 'legal_id',
                    'front': front,
                    'back': back,
                    'title': '法定代表人身份证',
                    'id_type': '法定代表人'
                })

        # 4. 资质证书（分组处理多页PDF）
        qualification_details = image_config.get('qualification_details', [])
        if qualification_details:
            # 按qual_key分组
            grouped_quals = {}
            for qual_detail in qualification_details:
                qual_key = qual_detail.get('qual_key')
                if qual_key:
                    if qual_key not in grouped_quals:
                        grouped_quals[qual_key] = []
                    grouped_quals[qual_key].append(qual_detail)

            # 对每组内的页面按page_num排序
            for qual_key, details in grouped_quals.items():
                details.sort(key=lambda x: x.get('page_num', 0))

            # 为每一页创建resource
            for qual_key, details_group in grouped_quals.items():
                is_multi_page = len(details_group) > 1

                for page_idx, qual_detail in enumerate(details_group):
                    resources.append({
                        'type': 'single_image',
                        'key': qual_key,
                        'path': qual_detail.get('file_path'),
                        'title': None,  # 由插入方法生成
                        'qual_key': qual_key,
                        'insert_hint': qual_detail.get('insert_hint', ''),
                        'is_first_page': (page_idx == 0),
                        'is_multi_page': is_multi_page,
                        'page_num': qual_detail.get('page_num', page_idx + 1),
                        'total_pages': len(details_group)
                    })

        return resources

    def _insert_single_image(self, doc: Document, resource: Dict[str, Any],
                            insert_point: Optional[Dict], index: int = 0) -> bool:
        """
        插入单张图片（营业执照、资质证书通用方法）

        合并了 _insert_license 和 _insert_qualification 的公共逻辑

        Args:
            doc: Word文档对象
            resource: 资质信息字典
            insert_point: 插入点信息
            index: 索引（用于生成默认标题）

        Returns:
            bool: 插入是否成功
        """
        try:
            # 提取资质信息
            image_path = resource.get('path')
            qual_key = resource.get('qual_key', resource.get('key'))
            insert_hint = resource.get('insert_hint')
            is_first_page = resource.get('is_first_page', True)
            page_num = resource.get('page_num', 1)

            # 解析路径（支持相对路径）
            resolved_path = self.utils.resolve_file_path(image_path)
            if not os.path.exists(resolved_path):
                self.logger.error(f"{qual_key}图片不存在: {image_path} (resolved: {resolved_path})")
                return False
            image_path = resolved_path

            # 生成标题
            if resource.get('title'):
                # 使用预设标题（如"营业执照副本"）
                title_text = resource['title']
            else:
                # 智能生成标题（资质证书）
                from .qualification_matcher import QUALIFICATION_MAPPING
                if qual_key and qual_key in QUALIFICATION_MAPPING:
                    qual_info = QUALIFICATION_MAPPING[qual_key]
                    if 'display_title' in qual_info:
                        title_text = qual_info['display_title']
                    elif insert_hint:
                        title_text = insert_hint[:50]
                    else:
                        title_text = f"{qual_info['category']}认证证书"
                elif insert_hint:
                    title_text = insert_hint[:50]
                else:
                    title_text = f"资质证书 {index + 1}"

            # 获取图片宽度
            width_inches = self.default_sizes.get(qual_key, (6, 0))[0]

            # 插入逻辑
            if insert_point and insert_point['type'] == 'paragraph':
                # 在找到的段落位置插入
                target_para = insert_point['paragraph']

                # 多页优化：只在第一页插入分页符和标题
                if is_first_page:
                    # 插入分页符
                    page_break_para = self.utils.insert_paragraph_after(target_para)
                    page_break_para.add_run().add_break()

                    # 插入标题
                    title = self.utils.insert_paragraph_after(page_break_para)
                    title.text = title_text
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if title.runs:
                        title.runs[0].font.bold = True

                    # 记录位置供后续页使用
                    self._last_insert_para = title
                    log_msg = f"✅ 在指定位置插入 {qual_key} 标题: {title_text}"
                else:
                    # 后续页：从上次插入位置继续
                    if hasattr(self, '_last_insert_para'):
                        title = self._last_insert_para
                    else:
                        title = target_para
                    log_msg = f"✅ 继续插入 {qual_key} 第{page_num}页"

                # 插入图片
                img_para = self.utils.insert_paragraph_after(title)
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(image_path, width=Inches(width_inches))

                # 更新插入位置
                self._last_insert_para = img_para

                self.logger.info(log_msg)
                return True

            else:
                # 没找到占位符 → 跳过不填充（由步骤5处理项目要求的资质）
                if is_first_page:
                    self.logger.info(f"ℹ️  {qual_key} 未找到模板占位符，跳过填充（如项目要求会在步骤5追加）")
                return False

        except Exception as e:
            self.logger.error(f"❌ 插入{resource.get('key')}失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return False

    def _insert_resource(self, doc: Document, resource: Dict[str, Any],
                        insert_point: Optional[Dict], stats: Dict, index: int = 0) -> None:
        """
        统一的资质插入分发方法

        根据资质类型分发到具体的插入方法，并统一更新统计信息

        Args:
            doc: Word文档对象
            resource: 资质信息字典
            insert_point: 插入点信息
            stats: 统计信息字典（会被修改）
            index: 索引（用于生成默认标题）
        """
        resource_type = resource.get('type')
        resource_key = resource.get('key')

        try:
            # 根据类型分发
            if resource_type == 'single_image':
                # 单张图片（营业执照、资质证书）
                success = self._insert_single_image(doc, resource, insert_point, index)

                if success:
                    stats['images_inserted'] += 1
                    page_num = resource.get('page_num', 1)
                    stats['images_types'].append(f"{resource_key}_p{page_num}" if page_num > 1 else resource_key)

                    # 只在第一页记录到 filled_qualifications
                    if resource.get('is_first_page', True) and resource_key != 'license':
                        from .qualification_matcher import QUALIFICATION_MAPPING
                        qual_name = QUALIFICATION_MAPPING.get(resource_key, {}).get('category', resource_key)
                        stats['filled_qualifications'].append({
                            'qual_key': resource_key,
                            'qual_name': qual_name,
                            'file_path': resource.get('path'),
                            'total_pages': resource.get('total_pages', 1)
                        })
                        total_pages = resource.get('total_pages', 1)
                        self.logger.info(f"✅ 填充资质: {resource_key} ({qual_name}), {total_pages}页")
                else:
                    page_num = resource.get('page_num', 1)
                    stats['errors'].append(f"{resource_key}_p{page_num}插入失败" if page_num > 1 else f"{resource_key}插入失败")

            elif resource_type == 'id_card':
                # 身份证（正反面表格）
                front_path = resource.get('front')
                back_path = resource.get('back')
                id_type = resource.get('id_type', '身份证')

                success = self.id_card_inserter.insert_id_card(doc, front_path, back_path, insert_point, id_type)

                if success:
                    stats['images_inserted'] += 2  # 正反两面
                    stats['images_types'].append(f"{id_type}身份证")
                else:
                    stats['errors'].append(f"{id_type}身份证插入失败")

        except Exception as e:
            self.logger.error(f"❌ 插入资质 {resource_key} 异常: {e}")
            stats['errors'].append(f"{resource_key}插入异常")
