#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
案例表格填充器 - CaseTableFiller
从案例库查询案例数据并填充到Word文档的案例表格中
"""

import re
from typing import Dict, Any, List, Optional
from docx import Document
from docx.table import Table
from pathlib import Path

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger


class CaseTableFiller:
    """案例表格填充器"""

    def __init__(self, case_manager, image_handler=None):
        """
        初始化案例表格填充器

        Args:
            case_manager: 案例库管理器实例
            image_handler: 图片处理器实例(可选,用于插入附件图片)
        """
        self.case_manager = case_manager
        self.image_handler = image_handler
        self.logger = get_module_logger("case_table_filler")

        # 案例表格的标识字段(表头关键字)
        self.case_table_headers = {
            '项目名称', '案例名称', '案例标题', '合同名称',
            '客户名称', '购买方', '甲方', '用户单位', '用户名称',
            '合同金额', '项目金额', '合同价格', '金额',
            '合同类型', '项目类型', '产品名称',
            '实施时间', '项目周期', '合同期限', '合同签约时间', '签约时间',
            '行业', '所属行业',
            '项目规模', '合同编号', '案例编号', '序号',
            '联系人', '联系方式', '用户联系人'
        }

        # 字段映射表:表头文字 -> 数据库字段名
        self.field_mapping = {
            # 项目/案例基本信息
            '项目名称': 'case_title',
            '案例名称': 'case_title',
            '案例标题': 'case_title',
            '合同名称': 'contract_name',
            '案例编号': 'case_number',
            '项目编号': 'case_number',
            '合同编号': 'case_number',
            '序号': 'case_number',  # 新增：序号字段

            # 客户信息
            '客户名称': 'customer_name',
            '购买方': 'customer_name',
            '甲方': 'customer_name',           # 甲方即客户名称
            '甲方单位名称': 'customer_name',  # 甲方单位名称即客户名称
            '甲方单位': 'customer_name',      # 甲方单位即客户名称
            '用户名称': 'customer_name',      # 用户名称即客户名称
            '用户单位': 'customer_name',      # 用户单位即客户名称
            '使用单位': 'customer_name',      # 使用单位即客户名称
            '最终用户': 'final_customer_name',

            # 合同信息
            '合同内容': 'contract_name',  # 新增:用合同名称代替合同内容
            '项目内容': 'contract_name',  # 新增
            '合同类型': 'contract_type',
            '项目类型': 'contract_type',
            '产品名称': 'contract_name',  # 新增：产品名称映射到合同名称
            '产品名称、数量': 'contract_name',  # 新增：组合列（产品名称+数量）
            '合同金额': 'contract_amount',
            '项目金额': 'contract_amount',
            '合同价格': 'contract_amount',
            '金额': 'contract_amount',    # 新增：金额字段
            '金额（元）': 'contract_amount',  # 新增：带单位的金额
            '数量': 'contract_amount',  # 新增:用合同金额代替数量
            '合同数量': 'contract_amount',  # 新增

            # 时间信息
            '合同签订时间': 'contract_start_date',  # 新增:用开始时间代替签订时间
            '签订时间': 'contract_start_date',      # 新增
            '合同签约时间': 'contract_start_date',  # 新增：签约时间
            '签约时间': 'contract_start_date',      # 新增
            '合同开始时间': 'contract_start_date',
            '合同开始日期': 'contract_start_date',
            '项目开始时间': 'contract_start_date',
            '合同结束时间': 'contract_end_date',
            '合同结束日期': 'contract_end_date',
            '项目结束时间': 'contract_end_date',
            '实施时间': 'contract_period',  # 组合字段
            '项目周期': 'contract_period',   # 组合字段
            '合同期限': 'contract_period',   # 组合字段
            '年份': 'contract_year',        # 虚拟字段，从contract_start_date提取年份
            '年度': 'contract_year',        # 年度即年份
            '时间': 'contract_year',        # 可能表示年份

            # 其他信息
            '行业': 'industry',
            '所属行业': 'industry',
            '项目规模': 'contract_amount',  # 可以用金额表示规模

            # 产品信息
            '产品类别': 'product_category',
            '产品分类': 'product_category',

            # 联系信息
            '联系人': 'party_a_contact_name',
            '用户联系人': 'party_a_contact_name',  # 新增：用户联系人
            '用户联系人及联系方式': 'party_a_contact_combined',  # 新增：组合字段（联系人+电话）
            '联系电话': 'party_a_contact_phone',
            '联系方式': 'party_a_contact_phone',  # 新增：联系方式
            '联系邮箱': 'party_a_contact_email'
        }

    def fill_case_tables(self, doc: Document, company_id: int, max_cases: int = 10) -> Dict[str, Any]:
        """
        识别并填充文档中的所有案例表格

        Args:
            doc: Word文档对象
            company_id: 公司ID
            max_cases: 最多填充的案例数量(默认10个)

        Returns:
            填充统计信息
        """
        stats = {
            'tables_filled': 0,
            'rows_filled': 0,
            'cases_used': 0,
            'images_inserted': 0,  # 新增:插入的图片数量
            'skipped_tables': 0
        }

        self.logger.info(f"开始处理案例表格,公司ID: {company_id}")

        # 从案例库查询所有案例
        cases = self._query_cases(company_id, max_cases)

        if not cases:
            self.logger.warning(f"公司 {company_id} 没有可用的案例数据")
            return stats

        self.logger.info(f"从案例库查询到 {len(cases)} 个案例")

        # 遍历文档中的所有表格
        for table_idx, table in enumerate(doc.tables):
            self.logger.debug(f"检查表格 #{table_idx + 1}")

            # 识别是否为案例表格
            if self._is_case_table(table):
                self.logger.info(f"识别到案例表格 #{table_idx + 1}")

                # 填充案例数据
                filled_rows = self._fill_table(table, cases)

                if filled_rows > 0:
                    stats['tables_filled'] += 1
                    stats['rows_filled'] += filled_rows
                    stats['cases_used'] = min(filled_rows, len(cases))
                    self.logger.info(f"  ✅ 填充了 {filled_rows} 行案例数据")

                    # 新增:在表格后插入案例附件图片
                    if self.image_handler:
                        images_count = self._insert_case_images_after_table(
                            doc, table, cases[:filled_rows]
                        )
                        stats['images_inserted'] += images_count
                    else:
                        self.logger.debug("  未提供image_handler,跳过图片插入")
                else:
                    stats['skipped_tables'] += 1
                    self.logger.warning(f"  ⚠️  表格识别为案例表格,但填充失败")

        self.logger.info(f"案例表格填充完成: 填充了 {stats['tables_filled']} 个表格, "
                        f"{stats['rows_filled']} 行数据, 使用了 {stats['cases_used']} 个案例, "
                        f"插入了 {stats['images_inserted']} 张图片")

        return stats

    def _is_case_table(self, table: Table) -> bool:
        """
        识别是否为案例表格(通过表头关键字)

        Args:
            table: Word表格对象

        Returns:
            是否为案例表格
        """
        if not table.rows or len(table.rows) < 2:
            return False

        # 检查第一行的表头
        header_row = table.rows[0]
        header_texts = [cell.text.strip() for cell in header_row.cells]

        # 匹配案例表格特征字段(至少包含2个关键字段)
        matched = 0
        for header_text in header_texts:
            # 移除空格和括号等干扰字符
            clean_header = re.sub(r'[\s()（）]', '', header_text)

            for keyword in self.case_table_headers:
                if keyword in clean_header:
                    matched += 1
                    self.logger.debug(f"    匹配到案例表格关键字: {keyword} (表头: {header_text})")
                    break

        # 至少匹配2个关键字段才认为是案例表格
        is_case_table = matched >= 2

        if is_case_table:
            self.logger.debug(f"    ✅ 识别为案例表格(匹配 {matched} 个关键字段)")
        else:
            self.logger.debug(f"    ❌ 非案例表格(仅匹配 {matched} 个关键字段)")

        return is_case_table

    def _query_cases(self, company_id: int, limit: int = 10) -> List[Dict]:
        """
        从案例库查询案例数据

        Args:
            company_id: 公司ID
            limit: 最多返回的案例数量

        Returns:
            案例列表(包含image_attachments字段)
        """
        try:
            # 调用案例库管理器查询案例
            cases = self.case_manager.get_cases(company_id=company_id)

            # 限制返回数量
            if len(cases) > limit:
                cases = cases[:limit]

            # 为每个案例加载附件,并处理PDF转换后的图片
            for case in cases:
                case_id = case['case_id']
                attachments = self.case_manager.get_attachments(case_id)

                # 处理每个附件,提取可用的图片
                case['image_attachments'] = []

                for att in attachments:
                    # 检查是否有PDF转换后的图片
                    converted_images = att.get('converted_images')
                    if converted_images:
                        import json
                        try:
                            images = json.loads(converted_images)
                            # 多页PDF: 添加所有页
                            for img_data in images:
                                case['image_attachments'].append({
                                    'attachment_id': att['attachment_id'],
                                    'file_path': img_data['file_path'],
                                    'page_num': img_data.get('page_num', 1),
                                    'is_multi_page': len(images) > 1,
                                    'description': att.get('attachment_description'),
                                    'type': att.get('attachment_type')
                                })
                        except Exception as e:
                            self.logger.warning(f"解析converted_images失败: {e}")

                    # 检查是否为图片文件
                    elif self._is_image_file(att.get('file_path')):
                        case['image_attachments'].append({
                            'attachment_id': att['attachment_id'],
                            'file_path': att['file_path'],
                            'page_num': 1,
                            'is_multi_page': False,
                            'description': att.get('attachment_description'),
                            'type': att.get('attachment_type')
                        })

            total_images = sum(len(c.get('image_attachments', [])) for c in cases)
            self.logger.info(f"查询到 {len(cases)} 个案例, 共 {total_images} 张可插入图片")

            return cases

        except Exception as e:
            self.logger.error(f"查询案例失败: {e}")
            return []

    def _fill_table(self, table: Table, cases: List[Dict]) -> int:
        """
        将案例数据填充到表格中

        Args:
            table: Word表格对象
            cases: 案例数据列表

        Returns:
            填充的行数
        """
        if not cases:
            return 0

        # 分析表头,建立列映射(表头文字 -> 列索引)
        column_mapping = self._build_column_mapping(table)

        if not column_mapping:
            self.logger.warning("  ⚠️  未识别到有效的列映射")
            return 0

        self.logger.debug(f"  列映射: {column_mapping}")

        # 从第2行开始填充(第1行是表头)
        filled_count = 0
        for idx, case in enumerate(cases):
            row_idx = idx + 1  # 跳过表头

            # 检查表格行数是否足够
            if row_idx >= len(table.rows):
                self.logger.warning(f"  ⚠️  表格行数不足,已填充 {filled_count} 行,剩余 {len(cases) - idx} 个案例未填充")
                break

            row = table.rows[row_idx]

            # 填充当前行
            for col_idx, field_key in column_mapping.items():
                if col_idx >= len(row.cells):
                    continue

                # 获取字段值
                value = self._get_field_value(case, field_key)

                if value:
                    # 填充单元格
                    self._fill_cell(row.cells[col_idx], str(value))
                    self.logger.debug(f"    填充单元格[{row_idx},{col_idx}]: {field_key} = {value}")

            filled_count += 1

        return filled_count

    def _build_column_mapping(self, table: Table) -> Dict[int, str]:
        """
        分析表头,建立列映射

        Args:
            table: Word表格对象

        Returns:
            列映射字典 {列索引: 数据库字段名}
        """
        if not table.rows:
            return {}

        header_row = table.rows[0]
        column_mapping = {}

        for col_idx, cell in enumerate(header_row.cells):
            header_text = cell.text.strip()

            # 移除空格和括号
            clean_header = re.sub(r'[\s()（）]', '', header_text)

            # 查找匹配的字段映射
            for field_name, field_key in self.field_mapping.items():
                clean_field_name = re.sub(r'[\s()（）]', '', field_name)

                if clean_field_name == clean_header or clean_field_name in clean_header:
                    column_mapping[col_idx] = field_key
                    self.logger.debug(f"    列{col_idx}: {header_text} -> {field_key}")
                    break

        return column_mapping

    def _get_field_value(self, case: Dict, field_key: str) -> Optional[str]:
        """
        获取案例字段值(支持组合字段和虚拟字段)

        Args:
            case: 案例数据字典
            field_key: 字段键名

        Returns:
            字段值
        """
        # 处理组合字段：合同期限
        if field_key == 'contract_period':
            # 合同期限 = 开始日期 ~ 结束日期
            start_date = case.get('contract_start_date', '')
            end_date = case.get('contract_end_date', '')

            if start_date and end_date:
                return f"{start_date} ~ {end_date}"
            elif start_date:
                return f"{start_date} 起"
            elif end_date:
                return f"至 {end_date}"
            else:
                return None

        # 处理组合字段：用户联系人及联系方式
        if field_key == 'party_a_contact_combined':
            # 用户联系人及联系方式 = 联系人 + 电话
            contact_name = case.get('party_a_contact_name', '')
            contact_phone = case.get('party_a_contact_phone', '')

            if contact_name and contact_phone:
                return f"{contact_name} {contact_phone}"
            elif contact_name:
                return contact_name
            elif contact_phone:
                return contact_phone
            else:
                return None

        # 处理虚拟字段：年份（从合同签订日期提取）
        if field_key == 'contract_year':
            start_date = case.get('contract_start_date', '')
            if start_date:
                # 支持多种日期格式
                # "2024-05-20" -> "2024"
                # "2024年05月20日" -> "2024"
                # "2024/05/20" -> "2024"
                import re
                year_match = re.match(r'^(\d{4})', str(start_date))
                if year_match:
                    return year_match.group(1)
            return None

        # 🆕 特殊处理：contract_name 为空时回退到 case_title
        # 原因：数据库中大部分案例只填写了case_title，contract_name为空
        # 根据schema设计，case_title等同于contract_name
        if field_key == 'contract_name':
            value = case.get('contract_name')
            if not value or (isinstance(value, str) and value.strip() == ''):
                # 回退到case_title
                value = case.get('case_title')
                if value:
                    self.logger.debug(f"contract_name为空，回退到case_title: {value}")
            return value if value else None

        # 普通字段
        value = case.get(field_key)

        # 过滤空值
        if value is None or (isinstance(value, str) and value.strip() == ''):
            return None

        return value

    def _fill_cell(self, cell, value: str):
        """
        填充单元格

        Args:
            cell: 单元格对象
            value: 要填充的值
        """
        if not cell.paragraphs:
            cell.add_paragraph(value)
            return

        # 只处理第一个段落
        paragraph = cell.paragraphs[0]

        # 如果段落有Run,保存第一个Run的格式
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_properties = {
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'size': first_run.font.size,
                'name': first_run.font.name
            }

            # 清空并重新设置文本
            paragraph.clear()
            new_run = paragraph.add_run(value)

            # 恢复格式
            for prop, val in font_properties.items():
                if val is not None:
                    setattr(new_run.font, prop, val)
        else:
            # 没有Run,直接设置文本
            if paragraph.text:
                paragraph.clear()
            paragraph.add_run(value)

    def _insert_case_images_after_table(self, doc: Document, table: Table, cases: List[Dict]) -> int:
        """
        在案例表格后插入案例附件图片(复用ImageHandler的插入逻辑)

        处理逻辑:
        1. PDF附件: 使用转换后的图片(支持多页)
        2. 图片附件: 直接插入
        3. Word附件: 跳过(无法直接插入)

        Args:
            doc: Word文档对象
            table: 案例表格对象
            cases: 已填充到表格中的案例列表

        Returns:
            插入的图片数量
        """
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from common import resolve_file_path
        import os

        images_inserted = 0

        # 找到表格后的插入点
        last_insert_para = self._find_para_after_table(doc, table)

        if not last_insert_para:
            self.logger.warning("  ⚠️  无法找到表格后的插入点,无法插入图片")
            return 0

        self.logger.info("  开始插入案例附件图片...")

        # 为每个案例插入附件图片
        for case_idx, case in enumerate(cases):
            image_attachments = case.get('image_attachments', [])

            if not image_attachments:
                self.logger.debug(f"    案例{case_idx+1}无图片附件")
                continue

            case_title = case.get('case_title', f'案例{case_idx+1}')
            self.logger.info(f"    案例 '{case_title}' 有 {len(image_attachments)} 张图片")

            # 按附件顺序插入图片
            for img_att in image_attachments:
                file_path = resolve_file_path(img_att['file_path'])

                if not file_path or not os.path.exists(file_path):
                    self.logger.warning(f"      图片文件不存在: {img_att['file_path']}")
                    continue

                # 生成标题
                att_desc = img_att.get('description') or \
                          self._get_attachment_type_name(img_att.get('type'))

                # 多页PDF特殊处理:第一页插入标题,后续页不插入
                page_num = img_att.get('page_num', 1)
                is_multi_page = img_att.get('is_multi_page', False)

                if is_multi_page:
                    if page_num == 1:
                        title_text = f"{case_title} - {att_desc}"
                    else:
                        title_text = None  # 后续页不插入标题
                else:
                    title_text = f"{case_title} - {att_desc}"

                # 插入标题(第一页或单页)
                if title_text:
                    title_para = self.image_handler._insert_paragraph_after(last_insert_para)
                    title_para.text = title_text
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if title_para.runs:
                        title_para.runs[0].font.bold = True
                    last_insert_para = title_para
                    self.logger.debug(f"      插入标题: {title_text}")

                # 插入图片
                img_para = self.image_handler._insert_paragraph_after(last_insert_para)
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(file_path, width=Inches(6))  # 6英寸(与资质证书一致)

                last_insert_para = img_para
                images_inserted += 1

                if title_text:
                    self.logger.info(f"      ✅ 已插入: {title_text}")
                else:
                    self.logger.info(f"      ✅ 已插入第{page_num}页")

        self.logger.info(f"  案例附件图片插入完成,共插入 {images_inserted} 张图片")
        return images_inserted

    def _find_para_after_table(self, doc: Document, table: Table):
        """
        找到表格后的段落作为插入点

        Args:
            doc: Word文档对象
            table: 表格对象

        Returns:
            表格后的段落对象,如果找不到则返回None
        """
        try:
            # 获取表格元素
            table_element = table._element

            # 尝试找到表格后的第一个段落
            for para in doc.paragraphs:
                # 检查段落是否在表格之后
                if para._element.getprevious() == table_element:
                    self.logger.debug(f"    找到表格后的段落: '{para.text[:50]}'")
                    return para

            # 降级:遍历表格后的所有元素,找到第一个段落
            next_element = table_element.getnext()
            while next_element is not None:
                if next_element.tag.endswith('}p'):  # 找到段落元素
                    # 包装成Paragraph对象
                    from docx.text.paragraph import Paragraph
                    para = Paragraph(next_element, doc)
                    self.logger.debug(f"    通过遍历找到表格后的段落")
                    return para
                next_element = next_element.getnext()

            # 再降级:返回文档末尾段落
            if doc.paragraphs:
                self.logger.debug(f"    使用文档末尾段落作为插入点")
                return doc.paragraphs[-1]
            else:
                # 创建新段落
                self.logger.debug(f"    文档无段落,创建新段落")
                return doc.add_paragraph()

        except Exception as e:
            self.logger.error(f"查找表格后段落失败: {e}")
            # 降级:返回文档末尾
            if doc.paragraphs:
                return doc.paragraphs[-1]
            return doc.add_paragraph()

    def _is_image_file(self, file_path: str) -> bool:
        """
        判断是否为图片文件

        Args:
            file_path: 文件路径

        Returns:
            是否为图片文件
        """
        if not file_path:
            return False
        valid_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'}
        ext = Path(file_path).suffix.lower()
        return ext in valid_extensions

    def _get_attachment_type_name(self, attachment_type: str) -> str:
        """
        获取附件类型的中文名称

        Args:
            attachment_type: 附件类型代码

        Returns:
            附件类型中文名称
        """
        type_map = {
            'contract': '合同',
            'acceptance': '验收证明',
            'testimony': '客户证明',
            'photo': '项目照片',
            'other': '其他附件'
        }
        return type_map.get(attachment_type, '附件')
