#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版内联回复处理器
功能特性：
1. 原地插入应答 - 在原文档每个需求后直接插入
2. 灰色底纹标记 - 应答内容添加浅灰色背景(RGB 217,217,217)
3. 完美格式保持 - 复制源段落的字体、大小、样式
4. 智能需求识别 - 改进的模式匹配，减少遗漏
5. 专业提示词模板 - 点对点应答、内容生成、标题生成
"""

import re
import logging
from typing import Dict, Optional, List, Any
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_LINE_SPACING
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import (
    get_config, get_module_logger,
    LLMClient, BusinessResponseError,
    get_prompt_manager
)


class InlineReplyProcessor:
    """增强版内联回复处理器"""

    def __init__(self, model_name: str = "shihuang-gpt4o-mini"):
        """
        初始化处理器

        Args:
            model_name: 使用的模型名称
        """
        self.config = get_config()
        self.logger = get_module_logger("inline_reply")

        # 初始化提示词管理器
        self.prompt_manager = get_prompt_manager()

        # 初始化LLM客户端
        self.llm_client = LLMClient(model_name=model_name)
        self.model_name = model_name

        # 加载配置
        self.patterns = self._get_requirement_patterns()
        self.templates = self._get_response_templates()

        self.logger.info(f"内联回复处理器初始化完成，使用模型: {model_name}")

    def _get_requirement_patterns(self) -> Dict:
        """获取需求识别模式"""
        return {
            "编号模式": [
                r'^(\d+)\s*[、．.]',         # 1、 1. 1．
                r'^(\d+\.\d+)\s*[、．.]',    # 1.1、 1.2.
                r'^\((\d+)\)',              # (1) (2)
                r'^([A-Z])\)',              # A) B)
                r'^([a-z])\)',              # a) b)
                r'^（[一二三四五六七八九十]+）'  # （一）（二）
            ],
            "关键词": [
                "要求", "需求", "应", "必须", "应当", "需要", "具备", "支持", "提供",
                "实现", "满足", "符合", "遵循", "不少于", "不低于", "负责", "确保",
                "保证", "完成", "达到", "实施", "部署", "维护", "服务", "管理"
            ],
            "章节标识": [
                "乙方", "甲方", "供应商", "采购方", "服务商", "承包方"
            ]
        }

    def _get_response_templates(self) -> Dict:
        """获取应答模板"""
        return {
            "硬件配置": "我方提供的硬件设备完全满足技术规格要求",
            "软件功能": "我方系统具备完整的功能模块，满足业务需求",
            "性能指标": "我方产品性能指标达到或超过要求标准",
            "技术规范": "我方严格遵循相关技术标准和行业规范",
            "服务保障": "我方提供全方位专业服务，确保项目成功",
            "资质证明": "我方具备完整的相关资质和丰富经验",
            "数据服务": "我方提供稳定可靠的数据服务和技术支持",
            "接口对接": "我方支持标准接口协议，确保系统无缝对接",
            "运维支持": "我方提供7×24小时专业运维服务",
            "通用模板": "我方将严格按照要求提供专业技术方案"
        }

    def is_requirement_paragraph(self, para) -> bool:
        """
        判断段落是否为需求条目（优化的识别策略）

        Args:
            para: 段落对象

        Returns:
            是否为需求段落
        """
        text = para.text.strip()
        if not text or len(text) < 10:  # 过短的文本不处理
            return False

        style_name = para.style.name if para.style else ""
        alignment = para.paragraph_format.alignment

        # 明确排除的条件
        exclusions = [
            "Heading" in style_name,     # 标题样式
            "标题" in style_name,         # 中文标题样式
            "TOC" in style_name,         # 目录样式
            "目录" in text,              # 包含"目录"
            "封面" in text,              # 封面内容
            "签字" in text,              # 签字页
            alignment == 1,              # 居中对齐（通常是标题）
            text.startswith("第") and ("章" in text or "节" in text),  # 章节标题
            re.match(r'^[一二三四五六七八九十]+[、．]', text),  # 中文数字章节
        ]

        if any(exclusions):
            return False

        # 包含条件：更宽泛的识别策略
        inclusions = [
            # 包含需求关键词
            any(keyword in text for keyword in self.patterns["关键词"]),
            # 包含编号格式
            any(re.match(pattern, text) for pattern in self.patterns["编号模式"]),
            # 包含主体标识
            any(entity in text for entity in self.patterns.get("章节标识", [])),
            # 长段落且包含具体要求（超过50字且包含技术词汇）
            (len(text) > 50 and any(word in text for word in
             ["系统", "数据", "接口", "服务", "技术", "平台", "软件", "硬件"])),
        ]

        return any(inclusions)

    def classify_requirement_type(self, text: str) -> str:
        """
        分类需求类型

        Args:
            text: 需求文本

        Returns:
            需求类型
        """
        content_lower = text.lower()

        # 数据服务相关
        if any(kw in content_lower for kw in ["数据", "查询", "接口", "api", "专线"]):
            return "数据服务"

        # 系统对接相关
        if any(kw in content_lower for kw in ["对接", "联调", "测试", "部署", "集成"]):
            return "接口对接"

        # 运维服务相关
        if any(kw in content_lower for kw in ["运维", "维护", "监控", "故障", "响应"]):
            return "运维支持"

        # 硬件相关
        if any(kw in content_lower for kw in ["cpu", "内存", "存储", "硬盘", "服务器", "设备"]):
            return "硬件配置"

        # 软件相关
        if any(kw in content_lower for kw in ["软件", "系统", "应用", "功能", "模块"]):
            return "软件功能"

        # 性能相关
        if any(kw in content_lower for kw in ["性能", "速度", "并发", "响应时间", "不少于", "不低于"]):
            return "性能指标"

        # 服务相关
        if any(kw in content_lower for kw in ["服务", "支持", "培训", "实施", "咨询"]):
            return "服务保障"

        # 资质相关
        if any(kw in content_lower for kw in ["资质", "证书", "认证", "经验", "案例"]):
            return "资质证明"

        return "通用模板"

    def generate_professional_response(self, requirement_text: str, use_ai: bool = True) -> str:
        """
        生成专业的技术应答

        Args:
            requirement_text: 需求文本
            use_ai: 是否使用AI生成（False则使用简单模板）

        Returns:
            专业应答文本
        """
        req_type = self.classify_requirement_type(requirement_text)

        # 如果不使用AI，直接返回模板应答
        if not use_ai:
            template = self.templates.get(req_type, self.templates["通用模板"])
            return f"应答：满足。{template}。"

        # 使用点对点应答风格 - 从提示词管理器获取
        prompt_template = self.prompt_manager.get_prompt('business_response', 'point_to_point',
            default="以下是输入文字：")
        prompt = f"{prompt_template}'{requirement_text}'"

        try:
            # 调用LLM生成应答 - 使用提示词管理器获取系统提示词
            system_prompt = self.prompt_manager.get_prompt('business_response', 'system_default',
                default="你是一名资深的技术方案专家。")
            response = self.llm_client.call(
                prompt=prompt,
                system_prompt=system_prompt,
                temperature=0.3,
                purpose=f"{req_type}应答生成"
            )

            # 确保以"应答：满足。"开头
            if not response.startswith("应答：满足。"):
                if "应答：满足。" in response:
                    # 找到应答：满足。并移到开头
                    response = "应答：满足。" + response.split("应答：满足。")[1]
                else:
                    response = f"应答：满足。{response}"

            return response
        except Exception as e:
            self.logger.error(f"生成专业应答失败: {e}")
            # 使用备用模板
            template = self.templates.get(req_type, self.templates["通用模板"])
            return f"应答：满足。{template}。"

    def copy_paragraph_format(self, source_para, target_para):
        """
        复制段落格式，确保字体一致性

        Args:
            source_para: 源段落
            target_para: 目标段落
        """
        try:
            # 复制段落级别的格式
            if source_para.paragraph_format.alignment is not None:
                target_para.paragraph_format.alignment = source_para.paragraph_format.alignment

            if source_para.paragraph_format.left_indent is not None:
                target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent

            if source_para.paragraph_format.right_indent is not None:
                target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent

            if source_para.paragraph_format.first_line_indent is not None:
                target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent

            # 复制行距设置
            if source_para.paragraph_format.line_spacing_rule is not None:
                target_para.paragraph_format.line_spacing_rule = source_para.paragraph_format.line_spacing_rule
            if source_para.paragraph_format.line_spacing is not None:
                target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
            if source_para.paragraph_format.space_before is not None:
                target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
            if source_para.paragraph_format.space_after is not None:
                target_para.paragraph_format.space_after = source_para.paragraph_format.space_after

            # 复制字体格式（从源段落的第一个run）
            if source_para.runs:
                source_run = source_para.runs[0]
                if target_para.runs:
                    target_run = target_para.runs[0]

                    # 复制字体属性
                    if source_run.font.name:
                        target_run.font.name = source_run.font.name
                    if source_run.font.size:
                        target_run.font.size = source_run.font.size
                    if source_run.font.bold is not None:
                        target_run.font.bold = source_run.font.bold
                    if source_run.font.italic is not None:
                        target_run.font.italic = source_run.font.italic
                    # 对于应答内容，始终设置为黑色字体，不继承原段落颜色
                    target_run.font.color.rgb = RGBColor(0, 0, 0)

        except Exception as e:
            self.logger.warning(f"格式复制失败: {e}")

    def copy_paragraph_format_except_line_spacing(self, source_para, target_para):
        """
        复制段落格式，但排除行距设置（专门为内联回复使用）

        Args:
            source_para: 源段落
            target_para: 目标段落
        """
        try:
            # 复制段落级别的格式（排除行距相关）
            if source_para.paragraph_format.alignment is not None:
                target_para.paragraph_format.alignment = source_para.paragraph_format.alignment

            if source_para.paragraph_format.left_indent is not None:
                target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent

            if source_para.paragraph_format.right_indent is not None:
                target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent

            # 统一设置应答段落的首行缩进为4个空格（约24磅）
            target_para.paragraph_format.first_line_indent = Pt(24)

            # 复制段前段后间距，但不复制行距
            if source_para.paragraph_format.space_before is not None:
                target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
            if source_para.paragraph_format.space_after is not None:
                target_para.paragraph_format.space_after = source_para.paragraph_format.space_after

            # 复制字体格式（从源段落的第一个run）
            if source_para.runs and target_para.runs:
                source_run = source_para.runs[0]
                target_run = target_para.runs[0]

                # 复制字体属性
                if source_run.font.name:
                    target_run.font.name = source_run.font.name
                if source_run.font.size:
                    target_run.font.size = source_run.font.size
                if source_run.font.bold is not None:
                    target_run.font.bold = source_run.font.bold
                if source_run.font.italic is not None:
                    target_run.font.italic = source_run.font.italic
                # 对于应答内容，始终设置为黑色字体
                target_run.font.color.rgb = RGBColor(0, 0, 0)

        except Exception as e:
            self.logger.warning(f"格式复制（排除行距）失败: {e}")

    def add_paragraph_shading(self, paragraph, color):
        """
        给段落添加背景色底纹

        Args:
            paragraph: 段落对象
            color: 颜色对象
        """
        try:
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml

            # 创建底纹XML元素 - RGB(217,217,217) 转换为十六进制 D9D9D9
            hex_color = f"{color.rgb:06X}" if hasattr(color, 'rgb') else "D9D9D9"
            shading_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
            shading_element = parse_xml(shading_xml)

            # 获取段落属性
            p_pr = paragraph._p.get_or_add_pPr()

            # 添加底纹元素
            p_pr.append(shading_element)

        except Exception as e:
            self.logger.warning(f"添加段落底纹失败: {e}")

    def insert_paragraph_after_with_format(self, paragraph, text=None, style=None):
        """
        在指定段落后插入新段落，并保持格式一致

        Args:
            paragraph: 源段落
            text: 要插入的文本
            style: 段落样式

        Returns:
            新插入的段落
        """
        if not DOCX_AVAILABLE:
            self.logger.error("未安装python-docx库，无法进行段落插入操作")
            return None

        try:
            # 使用正确的方式插入新段落
            # 获取段落的父元素（通常是body或table cell）
            parent = paragraph._parent

            # 找到当前段落在父元素中的位置
            para_index = None
            for i, p in enumerate(parent.paragraphs):
                if p._p == paragraph._p:
                    para_index = i
                    break

            if para_index is None:
                self.logger.error("无法找到段落位置")
                return None

            # 尝试使用insert_paragraph_after方法（如果可用）
            new_para = None
            try:
                if hasattr(parent, 'insert_paragraph_after'):
                    new_para = parent.insert_paragraph_after(text or "", paragraph)
            except Exception as e:
                self.logger.debug(f"insert_paragraph_after方法不可用: {e}")

            # 如果insert_paragraph_after不可用，使用备用方法
            if new_para is None:
                # 备用方法：直接在父容器中添加段落
                new_para = parent.add_paragraph()
                if text:
                    new_para.text = text

                # 手动调整段落位置（将新段落移动到目标段落后面）
                try:
                    # 使用正确的段落元素操作
                    target_p = paragraph._element
                    new_p = new_para._element
                    target_p.addnext(new_p)
                except Exception as move_error:
                    self.logger.warning(f"段落位置调整失败: {move_error}")

            # 设置样式
            if style:
                new_para.style = style

            # 首先明确设置1.5倍行距（在添加文本和格式之前）
            new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            new_para.paragraph_format.line_spacing = 1.5
            self.logger.debug(f"初始行距设置: rule={new_para.paragraph_format.line_spacing_rule}, spacing={new_para.paragraph_format.line_spacing}")

            # 如果需要添加文本且之前没有添加
            if text and not new_para.text:
                run = new_para.add_run(text)
                # 设置字体为黑色
                run.font.color.rgb = RGBColor(0, 0, 0)

            # 复制源段落的部分格式（除了行距，因为我们要强制使用1.5倍）
            self.copy_paragraph_format_except_line_spacing(paragraph, new_para)

            # 再次确保行距设置（在格式复制之后）
            new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            new_para.paragraph_format.line_spacing = 1.5
            self.logger.debug(f"格式复制后重新设置行距: rule={new_para.paragraph_format.line_spacing_rule}, spacing={new_para.paragraph_format.line_spacing}")

            # 添加浅灰色底纹 RGB(217,217,217)
            self.add_paragraph_shading(new_para, RGBColor(217, 217, 217))

            # 最终验证行距设置
            final_rule = new_para.paragraph_format.line_spacing_rule
            final_spacing = new_para.paragraph_format.line_spacing
            self.logger.info(f"最终行距设置: rule={final_rule}({final_rule.name if hasattr(final_rule, 'name') else final_rule}), spacing={final_spacing}")

            return new_para

        except Exception as e:
            self.logger.error(f"段落插入失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return None

    def process_document(self, input_file: str, output_file: Optional[str] = None, use_ai: bool = True) -> Dict[str, Any]:
        """
        处理文档，插入内联回复

        Args:
            input_file: 输入文档路径
            output_file: 输出文档路径（可选）
            use_ai: 是否使用AI生成应答（False则使用简单模板）

        Returns:
            dict: 包含输出文件路径和统计信息
            {
                'output_file': str,
                'requirements_count': int,
                'responses_count': int
            }

        Raises:
            BusinessResponseError: 处理失败
        """
        if not DOCX_AVAILABLE:
            error_msg = "未安装python-docx库，请安装：pip install python-docx"
            self.logger.error(error_msg)
            raise BusinessResponseError(error_msg)

        self.logger.info(f"开始处理文档: {input_file}")

        try:
            # 读取文档
            doc = Document(input_file)
            self.logger.info(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")

            # 统计需求条目
            requirement_count = 0
            processed_count = 0

            # 从后往前遍历，避免插入新段落影响索引
            paragraphs = list(doc.paragraphs)

            for i in range(len(paragraphs) - 1, -1, -1):
                para = paragraphs[i]

                if self.is_requirement_paragraph(para):
                    requirement_count += 1
                    text = para.text.strip()

                    self.logger.info(f"处理需求 {requirement_count}: {text[:60]}...")

                    # 生成专业应答
                    response = self.generate_professional_response(text, use_ai)

                    # 在需求段落后插入应答（保持格式一致）
                    reply_para = self.insert_paragraph_after_with_format(para, response)

                    if reply_para:
                        processed_count += 1
                        self.logger.info(f"已插入应答: {response[:60]}...")
                    else:
                        self.logger.error(f"插入应答失败: {text[:30]}...")

            # 生成输出文件名
            if not output_file:
                base_name = Path(input_file).stem
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_file = str(Path(input_file).parent / f"{base_name}-内联应答-{timestamp}.docx")

            # 保存文档
            doc.save(output_file)

            self.logger.info(f"处理完成:")
            self.logger.info(f"  识别需求条目: {requirement_count}")
            self.logger.info(f"  成功插入应答: {processed_count}")
            self.logger.info(f"  输出文件: {output_file}")

            return {
                'output_file': output_file,
                'requirements_count': requirement_count,
                'responses_count': processed_count
            }

        except Exception as e:
            self.logger.error(f"文档处理失败: {e}")
            raise BusinessResponseError(f"文档处理失败: {e}")


# 保持向后兼容
class PointToPointProcessor(InlineReplyProcessor):
    """点对点处理器 - 向后兼容别名"""
    pass