#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能文档填写器 - SmartDocumentFiller
基于规则引擎和模式识别的通用文档填写解决方案

特性：
1. 自动识别5种填空模式
2. 智能字段映射和变体识别
3. 保持原文档格式
4. 详细的匹配和填充日志
5. 模块化、可扩展的设计

作者：AI Tender System
版本：2.0
日期：2025-10-19
"""

from typing import Dict, Any, List
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger

# 导入Word文档工具
from .utils import WordDocumentUtils, normalize_data_keys

# 导入子模块
from .pattern_matcher import PatternMatcher
from .field_recognizer import FieldRecognizer
from .content_filler import ContentFiller


class SmartDocumentFiller:
    """智能文档填写器主类"""

    def __init__(self):
        self.logger = get_module_logger("smart_filler")

        # 初始化子模块
        self.pattern_matcher = PatternMatcher()
        self.field_recognizer = FieldRecognizer()
        self.content_filler = ContentFiller(self.logger)

        self.logger.info("智能文档填写器初始化完成")

    def fill_document(self,
                     doc: Document,
                     data: Dict[str, Any]) -> Dict[str, Any]:
        """
        填写文档主方法

        Args:
            doc: Word文档对象
            data: 数据字典，包含所有需要填充的信息

        Returns:
            填充统计信息
        """
        # 标准化数据键名（兼容旧数据格式）
        data = self._normalize_data_keys(data)

        stats = {
            'total_filled': 0,
            'pattern_counts': {},
            'unfilled_fields': [],
            'errors': []
        }

        self.logger.info("="*60)
        self.logger.info("开始智能文档填充")
        self.logger.info(f"可用数据字段: {list(data.keys())}")

        # ✅ 专门检测 purchaserName 字段
        if 'purchaserName' in data:
            self.logger.info(f"✅ 检测到purchaserName字段: {data.get('purchaserName')}")
        else:
            self.logger.warning("⚠️  未检测到purchaserName字段")

        # ✅ 专门检测 companyName 字段（响应人名称映射到此字段）
        if 'companyName' in data:
            self.logger.info(f"✅ 检测到companyName字段: {data.get('companyName')}")
            self.logger.info(f"   → 此字段应填充到: 响应人名称/应答人名称/供应商名称等")
        else:
            self.logger.warning("⚠️  未检测到companyName字段")

        self.logger.info("="*60)

        # 处理所有段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue

            # 匹配并填充
            result = self._process_paragraph(paragraph, data, para_idx)

            # 更新统计
            if result['filled']:
                stats['total_filled'] += 1
                pattern = result['pattern']
                stats['pattern_counts'][pattern] = stats['pattern_counts'].get(pattern, 0) + 1
            elif result['unfilled_fields']:
                stats['unfilled_fields'].extend(result['unfilled_fields'])

            if result['errors']:
                stats['errors'].extend(result['errors'])

        # 处理所有表格
        self.logger.info("开始处理表格...")
        table_para_idx = len(doc.paragraphs)  # 表格段落从这个索引开始编号
        for table_idx, table in enumerate(doc.tables):
            self.logger.debug(f"处理表格#{table_idx}")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for cell_para_idx, paragraph in enumerate(cell.paragraphs):
                        if not paragraph.text.strip():
                            continue

                        # 匹配并填充（使用统一的段落处理逻辑）
                        para_id = f"table{table_idx}_row{row_idx}_cell{cell_idx}_para{cell_para_idx}"
                        result = self._process_paragraph(paragraph, data, table_para_idx)
                        table_para_idx += 1

                        # 更新统计
                        if result['filled']:
                            stats['total_filled'] += 1
                            pattern = result['pattern']
                            stats['pattern_counts'][pattern] = stats['pattern_counts'].get(pattern, 0) + 1
                            self.logger.debug(f"  表格填充成功: {para_id}")
                        elif result['unfilled_fields']:
                            stats['unfilled_fields'].extend(result['unfilled_fields'])

                        if result['errors']:
                            stats['errors'].extend(result['errors'])

        # 过滤未填充字段，排除误识别内容
        if stats['unfilled_fields']:
            filtered_unfilled = self._filter_invalid_fields(stats['unfilled_fields'])
            stats['filtered_unfilled_fields'] = filtered_unfilled  # 添加过滤后的字段
            stats['original_unfilled_count'] = len(stats['unfilled_fields'])  # 保留原始数量用于调试
        else:
            stats['filtered_unfilled_fields'] = []
            stats['original_unfilled_count'] = 0

        # 输出统计报告
        self._print_stats(stats)

        # 后处理：清理多余占位符
        # self._post_process(doc)

        return stats

    def _process_paragraph(self,
                          paragraph: Paragraph,
                          data: Dict[str, Any],
                          para_idx: int) -> Dict[str, Any]:
        """
        处理单个段落

        Returns:
            {
                'filled': bool,
                'pattern': str,
                'unfilled_fields': list,
                'errors': list
            }
        """
        text = paragraph.text.strip()
        result = {
            'filled': False,
            'pattern': None,
            'unfilled_fields': [],
            'errors': []
        }

        self.logger.debug(f"段落#{para_idx}: {text[:80]}")

        # 2. 按优先级尝试填充
        # 注意：每次填充后重新检测模式，因为文本已改变，位置信息会失效
        filled_patterns = []

        for pattern_type in ['combo', 'bracket', 'date', 'colon', 'space_fill']:
            try:
                # 每次都重新检测当前文本的模式
                # 注意：不要strip，因为位置信息必须与paragraph.text一致
                text = paragraph.text
                patterns = self.pattern_matcher.detect_patterns(text)

                if pattern_type not in patterns or not patterns[pattern_type]:
                    continue

                self.logger.debug(f"  尝试 {pattern_type} 模式，检测到 {len(patterns[pattern_type])} 个匹配")

                # date模式特殊处理：避免与colon冲突
                if pattern_type == 'colon' and 'date' in filled_patterns:
                    # 如果date已经填充，跳过colon中的日期字段
                    colon_has_date = any('日期' in m['field'] for m in patterns['colon'])
                    if colon_has_date:
                        self.logger.debug("  跳过colon模式中的日期字段（已由date模式处理）")
                        patterns['colon'] = [m for m in patterns['colon'] if '日期' not in m['field']]
                        if not patterns['colon']:
                            continue

                filled = self._fill_by_pattern(
                    paragraph,
                    text,
                    pattern_type,
                    patterns[pattern_type],
                    data
                )

                if filled:
                    filled_patterns.append(pattern_type)
                    self.logger.info(f"  ✅ 使用 {pattern_type} 模式成功填充")

            except Exception as e:
                error_msg = f"段落#{para_idx} 填充失败({pattern_type}): {e}"
                self.logger.error(error_msg)
                result['errors'].append(error_msg)

        # 如果有任何模式填充成功
        if filled_patterns:
            result['filled'] = True
            result['pattern'] = '+'.join(filled_patterns)  # 记录使用了哪些模式
            return result

        # 3. 记录未填充的字段
        if patterns:
            for pattern_type, matches in patterns.items():
                for match in matches:
                    result['unfilled_fields'].append({
                        'para_idx': para_idx,
                        'text': text[:100],
                        'pattern': pattern_type,
                        'field': match
                    })

        return result

    def _fill_by_pattern(self,
                        paragraph: Paragraph,
                        text: str,
                        pattern_type: str,
                        matches: List,
                        data: Dict[str, Any]) -> bool:
        """根据模式类型填充"""

        if pattern_type == 'combo':
            return self.content_filler.fill_combo_field(paragraph, text, matches, data)

        elif pattern_type == 'bracket':
            return self.content_filler.fill_bracket_field(paragraph, text, matches, data)

        elif pattern_type == 'colon':
            return self.content_filler.fill_colon_field(paragraph, text, matches, data)

        elif pattern_type == 'space_fill':
            return self.content_filler.fill_space_field(paragraph, text, matches, data)

        elif pattern_type == 'date':
            return self.content_filler.fill_date_field(paragraph, text, matches, data)

        return False

    def _normalize_data_keys(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        标准化数据键名，兼容旧格式（调用公共工具）

        将旧的字段名映射到新的标准字段名:
        - fixedPhone -> phone (驼峰)
        - fixed_phone -> phone (下划线)
        - registeredAddress -> address
        - postal_code -> postalCode
        等等
        """
        # 调用公共工具函数
        return normalize_data_keys(data, logger=self.logger)

    def _print_stats(self, stats: Dict[str, Any]):
        """打印统计报告"""
        self.logger.info("="*60)
        self.logger.info("填充统计报告")
        self.logger.info("="*60)
        self.logger.info(f"总填充数: {stats['total_filled']}")

        if stats['pattern_counts']:
            self.logger.info("按模式统计:")
            for pattern, count in stats['pattern_counts'].items():
                self.logger.info(f"  - {pattern}: {count}")

        # 过滤未填充字段：排除明显的误识别内容
        if stats['unfilled_fields']:
            filtered_unfilled = self._filter_invalid_fields(stats['unfilled_fields'])

            if filtered_unfilled:
                self.logger.warning(f"真正未填充字段数: {len(filtered_unfilled)}")
                for uf in filtered_unfilled[:10]:  # 只显示前10个
                    self.logger.warning(f"  - 段落#{uf['para_idx']}: {uf['field']}")

            # 显示过滤掉的误识别数量
            filtered_count = len(stats['unfilled_fields']) - len(filtered_unfilled)
            if filtered_count > 0:
                self.logger.debug(f"过滤掉的误识别字段数: {filtered_count}")

        if stats['errors']:
            self.logger.error(f"错误数: {len(stats['errors'])}")
            for err in stats['errors'][:5]:  # 只显示前5个
                self.logger.error(f"  - {err}")

    def _filter_invalid_fields(self, unfilled_fields: List[Dict]) -> List[Dict]:
        """
        过滤掉明显不是数据字段的误识别内容

        Args:
            unfilled_fields: 原始未填充字段列表

        Returns:
            过滤后的列表（仅包含真正的数据字段）
        """
        # 排除关键词：说明性文字、标题、提示
        exclude_keywords = [
            # 说明性文字
            '行贿犯罪记录', '时间从', '以相关', '以中国', '本条', '即增值税',
            '我公司', '我方', '本投标', '有关的一切', '请寄', '通讯',
            # 提示性文字
            '盖章', '签字', '公章', '签名',
            # 格式和标题
            '格式', '附件', '正本', '副本', '电子版',
            # 条款标题
            '情况', '承诺', '声明', '说明', '注意', '备注',
            # 网址和特殊字符
            'http', 'www.', '://','满足招标','见附件',
            # 文档材料类关键词（文档清单项）- 精确化，避免过滤"身份证号"
            '身份证复印件', '身份证扫描件', '提供身份证',  # 精确化
            '复印件', '证明', '原件', '扫描件', '材料', '文件'
        ]

        # 过滤逻辑
        filtered = []
        for field_info in unfilled_fields:
            field = field_info.get('field', '')

            # 如果是字典（bracket、colon模式），取field字段
            if isinstance(field, dict):
                field_text = field.get('field', '')
            else:
                field_text = str(field)

            # 检查是否包含排除关键词
            should_exclude = any(keyword in field_text for keyword in exclude_keywords)

            # 检查长度（字段名通常不超过15个字符）
            if not should_exclude and len(field_text) > 15:
                should_exclude = True

            if not should_exclude:
                filtered.append(field_info)

        return filtered


# 向后兼容：提供与旧API相同的接口
class InfoFillerV2(SmartDocumentFiller):
    """InfoFiller V2 - 基于 SmartDocumentFiller 的兼容层"""

    def fill_info(self, doc: Document, company_info: Dict[str, Any],
                  project_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        兼容旧版本的 fill_info 接口
        """
        # 合并数据
        data = {**company_info, **project_info}

        # 调用新方法
        stats = self.fill_document(doc, data)

        # 转换为旧格式
        return {
            'total_replacements': stats['total_filled'],
            'replacement_rules': stats['pattern_counts'].get('bracket', 0),
            'fill_rules': stats['pattern_counts'].get('colon', 0),
            'combination_rules': stats['pattern_counts'].get('combo', 0),
            'skipped_fields': 0,
            'none': len(stats['unfilled_fields'])
        }
