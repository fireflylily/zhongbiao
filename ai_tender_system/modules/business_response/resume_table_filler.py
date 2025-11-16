#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历表格填充器 - ResumeTableFiller
从简历库查询简历数据并填充到Word文档的简历表格中
"""

import re
import json
from typing import Dict, Any, List, Optional
from docx import Document
from docx.table import Table
from pathlib import Path

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger


class ResumeTableFiller:
    """简历表格填充器"""

    def __init__(self, resume_manager, image_handler=None):
        """
        初始化简历表格填充器

        Args:
            resume_manager: 简历库管理器实例
            image_handler: 图片处理器实例(可选,用于插入附件图片)
        """
        self.resume_manager = resume_manager
        self.image_handler = image_handler
        self.logger = get_module_logger("resume_table_filler")

        # 简历表格的标识字段(表头关键字)
        self.resume_table_headers = {
            '姓名', '人员姓名', '项目经理', '技术负责人',
            '职位', '岗位', '职务',
            '学历', '最高学历',
            '专业', '所学专业',
            '工作年限', '从业年限', '工作经验',
            '项目经验', '主要业绩', '工作经历',
            '资格证书', '职称', '专业技术职称'
        }

        # 字段映射表:表头文字 -> 数据库字段名
        self.field_mapping = {
            # 基本信息
            '姓名': 'name',
            '人员姓名': 'name',
            '项目经理': 'name',  # 如果表格标题是"项目经理",填充姓名
            '技术负责人': 'name',
            '性别': 'gender',
            '年龄': 'age',  # 组合字段,需要从出生日期计算
            '出生年月': 'birth_date',
            '出生日期': 'birth_date',
            '民族': 'nationality',
            '籍贯': 'native_place',
            '政治面貌': 'political_status',
            '身份证号': 'id_number',

            # 教育背景
            '学历': 'education_level',
            '最高学历': 'education_level',
            '学位': 'degree',
            '毕业院校': 'university',
            '毕业学校': 'university',
            '专业': 'major',
            '所学专业': 'major',
            '毕业时间': 'graduation_date',
            '毕业日期': 'graduation_date',

            # 工作信息
            '职位': 'current_position',
            '岗位': 'current_position',
            '职务': 'current_position',
            '职称': 'professional_title',
            '专业技术职称': 'professional_title',
            '工作年限': 'work_years',
            '从业年限': 'work_years',
            '工作经验': 'work_years',
            '工作单位': 'current_company',
            '现任单位': 'current_company',
            '部门': 'department',

            # 技能和证书
            '技能': 'skills',
            '专业技能': 'skills',
            '资格证书': 'certificates',
            '职业资格': 'certificates',
            '证书': 'certificates',
            '语言能力': 'languages',
            '外语水平': 'languages',

            # 项目经验
            '项目经验': 'project_experience',
            '主要业绩': 'project_experience',
            '工作经历': 'work_experience',
            '项目案例': 'project_experience',

            # 联系方式
            '联系电话': 'phone',
            '手机': 'phone',
            '电话': 'phone',
            '邮箱': 'email',
            '电子邮箱': 'email',
            '联系地址': 'address',
            '地址': 'address',

            # 其他
            '个人简介': 'introduction',
            '自我评价': 'introduction',
            '获奖情况': 'awards',
            '荣誉奖项': 'awards'
        }

    def fill_resume_tables(self, doc: Document, company_id: int, max_resumes: int = 10) -> Dict[str, Any]:
        """
        识别并填充文档中的所有简历表格

        Args:
            doc: Word文档对象
            company_id: 公司ID
            max_resumes: 最多填充的简历数量(默认10个)

        Returns:
            填充统计信息
        """
        stats = {
            'tables_filled': 0,
            'rows_filled': 0,
            'resumes_used': 0,
            'images_inserted': 0,  # 新增:插入的图片数量
            'skipped_tables': 0
        }

        self.logger.info(f"开始处理简历表格,公司ID: {company_id}")

        # 从简历库查询所有简历
        resumes = self._query_resumes(company_id, max_resumes)

        if not resumes:
            self.logger.warning(f"公司 {company_id} 没有可用的简历数据")
            return stats

        self.logger.info(f"从简历库查询到 {len(resumes)} 份简历")

        # 遍历文档中的所有表格
        for table_idx, table in enumerate(doc.tables):
            self.logger.debug(f"检查表格 #{table_idx + 1}")

            # 识别是否为简历表格
            if self._is_resume_table(table):
                self.logger.info(f"识别到简历表格 #{table_idx + 1}")

                # 填充简历数据
                filled_rows = self._fill_table(table, resumes)

                if filled_rows > 0:
                    stats['tables_filled'] += 1
                    stats['rows_filled'] += filled_rows
                    stats['resumes_used'] = min(filled_rows, len(resumes))
                    self.logger.info(f"  ✅ 填充了 {filled_rows} 行简历数据")

                    # 新增:在表格后插入简历附件图片
                    if self.image_handler:
                        images_count = self._insert_resume_images_after_table(
                            doc, table, resumes[:filled_rows]
                        )
                        stats['images_inserted'] += images_count
                    else:
                        self.logger.debug("  未提供image_handler,跳过图片插入")
                else:
                    stats['skipped_tables'] += 1
                    self.logger.warning(f"  ⚠️  表格识别为简历表格,但填充失败")

        self.logger.info(f"简历表格填充完成: 填充了 {stats['tables_filled']} 个表格, "
                        f"{stats['rows_filled']} 行数据, 使用了 {stats['resumes_used']} 份简历, "
                        f"插入了 {stats['images_inserted']} 张图片")

        return stats

    def _is_resume_table(self, table: Table) -> bool:
        """
        识别是否为简历表格(通过表头关键字)

        Args:
            table: Word表格对象

        Returns:
            是否为简历表格
        """
        if not table.rows or len(table.rows) < 2:
            return False

        # 检查第一行的表头
        header_row = table.rows[0]
        header_texts = [cell.text.strip() for cell in header_row.cells]

        # 匹配简历表格特征字段(至少包含2个关键字段)
        matched = 0
        for header_text in header_texts:
            # 移除空格和括号等干扰字符
            clean_header = re.sub(r'[\s()（）]', '', header_text)

            for keyword in self.resume_table_headers:
                if keyword in clean_header:
                    matched += 1
                    self.logger.debug(f"    匹配到简历表格关键字: {keyword} (表头: {header_text})")
                    break

        # 至少匹配2个关键字段才认为是简历表格
        is_resume_table = matched >= 2

        if is_resume_table:
            self.logger.debug(f"    ✅ 识别为简历表格(匹配 {matched} 个关键字段)")
        else:
            self.logger.debug(f"    ❌ 非简历表格(仅匹配 {matched} 个关键字段)")

        return is_resume_table

    def _query_resumes(self, company_id: Optional[int], limit: int = 10) -> List[Dict]:
        """
        从简历库查询简历数据

        Args:
            company_id: 公司ID(可选,如果为None则查询所有简历)
            limit: 最多返回的简历数量

        Returns:
            简历列表(包含image_attachments字段)
        """
        try:
            # 调用简历库管理器查询简历
            result = self.resume_manager.get_resumes(
                company_id=company_id,  # 可以为None
                page=1,
                page_size=limit,
                status='active'  # 只查询激活状态的简历
            )

            resumes = result.get('resumes', [])

            # 为每个简历加载附件,并处理PDF转换后的图片
            for resume in resumes:
                resume_id = resume['resume_id']
                attachments = self.resume_manager.get_attachments(resume_id)

                # 处理每个附件,提取可用的图片
                resume['image_attachments'] = []

                for att in attachments:
                    # 检查是否有PDF转换后的图片
                    converted_images = att.get('converted_images')
                    if converted_images:
                        import json
                        try:
                            images = json.loads(converted_images)
                            # 多页PDF: 添加所有页
                            for img_data in images:
                                resume['image_attachments'].append({
                                    'attachment_id': att['attachment_id'],
                                    'file_path': img_data['file_path'],
                                    'page_num': img_data.get('page_num', 1),
                                    'is_multi_page': len(images) > 1,
                                    'description': att.get('attachment_description'),
                                    'category': att.get('attachment_category')
                                })
                        except Exception as e:
                            self.logger.warning(f"解析converted_images失败: {e}")

                    # 检查是否为图片文件
                    elif self._is_image_file(att.get('file_path')):
                        resume['image_attachments'].append({
                            'attachment_id': att['attachment_id'],
                            'file_path': att['file_path'],
                            'page_num': 1,
                            'is_multi_page': False,
                            'description': att.get('attachment_description'),
                            'category': att.get('attachment_category')
                        })

            total_images = sum(len(r.get('image_attachments', [])) for r in resumes)
            self.logger.info(f"查询到 {len(resumes)} 份简历, 共 {total_images} 张可插入图片")

            return resumes

        except Exception as e:
            self.logger.error(f"查询简历失败: {e}")
            return []

    def _fill_table(self, table: Table, resumes: List[Dict]) -> int:
        """
        将简历数据填充到表格中

        Args:
            table: Word表格对象
            resumes: 简历数据列表

        Returns:
            填充的行数
        """
        if not resumes:
            return 0

        # 分析表头,建立列映射(表头文字 -> 列索引)
        column_mapping = self._build_column_mapping(table)

        if not column_mapping:
            self.logger.warning("  ⚠️  未识别到有效的列映射")
            return 0

        self.logger.debug(f"  列映射: {column_mapping}")

        # 从第2行开始填充(第1行是表头)
        filled_count = 0
        for idx, resume in enumerate(resumes):
            row_idx = idx + 1  # 跳过表头

            # 检查表格行数是否足够
            if row_idx >= len(table.rows):
                self.logger.warning(f"  ⚠️  表格行数不足,已填充 {filled_count} 行,剩余 {len(resumes) - idx} 份简历未填充")
                break

            row = table.rows[row_idx]

            # 填充当前行
            for col_idx, field_key in column_mapping.items():
                if col_idx >= len(row.cells):
                    continue

                # 获取字段值
                value = self._get_field_value(resume, field_key)

                if value:
                    # 填充单元格
                    self._fill_cell(row.cells[col_idx], str(value))
                    self.logger.debug(f"    填充单元格[{row_idx},{col_idx}]: {field_key} = {value[:50] if len(str(value)) > 50 else value}")

            filled_count += 1

        return filled_count

    def _build_column_mapping(self, table: Table) -> Dict[int, str]:
        """
        分析表头,建立列映射

        Args:
            table: Word表格对象

        Returns:
            列映射字典 {列索引: 数据库字段名}
        """
        if not table.rows:
            return {}

        header_row = table.rows[0]
        column_mapping = {}

        for col_idx, cell in enumerate(header_row.cells):
            header_text = cell.text.strip()

            # 移除空格和括号
            clean_header = re.sub(r'[\s()（）]', '', header_text)

            # 查找匹配的字段映射
            for field_name, field_key in self.field_mapping.items():
                clean_field_name = re.sub(r'[\s()（）]', '', field_name)

                if clean_field_name == clean_header or clean_field_name in clean_header:
                    column_mapping[col_idx] = field_key
                    self.logger.debug(f"    列{col_idx}: {header_text} -> {field_key}")
                    break

        return column_mapping

    def _get_field_value(self, resume: Dict, field_key: str) -> Optional[str]:
        """
        获取简历字段值(支持组合字段和JSON字段)

        Args:
            resume: 简历数据字典
            field_key: 字段键名

        Returns:
            字段值
        """
        # 处理组合字段:年龄(从出生日期计算)
        if field_key == 'age':
            birth_date = resume.get('birth_date')
            if birth_date:
                try:
                    from datetime import datetime
                    birth_year = datetime.fromisoformat(str(birth_date)).year
                    current_year = datetime.now().year
                    age = current_year - birth_year
                    return str(age)
                except Exception as e:
                    self.logger.debug(f"计算年龄失败: {e}")
                    return None
            return None

        # 获取字段值
        value = resume.get(field_key)

        # 过滤空值
        if value is None:
            return None

        # 处理JSON字段(skills, certificates, languages, project_experience等)
        if field_key in ['skills', 'certificates', 'languages', 'project_experience', 'work_experience']:
            return self._format_json_field(value, field_key)

        # 处理字符串值
        if isinstance(value, str) and value.strip() == '':
            return None

        return str(value)

    def _format_json_field(self, json_value: Any, field_key: str) -> Optional[str]:
        """
        格式化JSON字段为可读文本

        Args:
            json_value: JSON字段值(可能是字符串或已解析的对象)
            field_key: 字段键名

        Returns:
            格式化后的文本
        """
        try:
            # 如果是字符串,尝试解析JSON
            if isinstance(json_value, str):
                try:
                    data = json.loads(json_value)
                except json.JSONDecodeError:
                    # 如果解析失败,直接返回字符串
                    return json_value
            else:
                data = json_value

            # 处理列表类型
            if isinstance(data, list):
                if field_key == 'skills':
                    # 技能列表:用顿号连接
                    return '、'.join([str(item) for item in data if item])

                elif field_key == 'certificates':
                    # 证书列表:用顿号连接
                    return '、'.join([str(item) for item in data if item])

                elif field_key == 'languages':
                    # 语言能力列表
                    return '、'.join([str(item) for item in data if item])

                elif field_key in ['project_experience', 'work_experience']:
                    # 项目/工作经验:简化显示(只显示前3个,每个最多100字)
                    items = []
                    for idx, item in enumerate(data[:3]):  # 最多显示3个
                        if isinstance(item, dict):
                            # 如果是字典,提取关键信息
                            proj_name = item.get('project_name', item.get('name', ''))
                            role = item.get('role', item.get('position', ''))
                            desc = item.get('description', '')

                            if proj_name:
                                item_text = f"{idx+1}. {proj_name}"
                                if role:
                                    item_text += f" ({role})"
                                if desc:
                                    # 截取描述的前50个字符
                                    desc_short = desc[:50] + '...' if len(desc) > 50 else desc
                                    item_text += f": {desc_short}"
                                items.append(item_text)
                        elif isinstance(item, str):
                            # 如果是字符串,直接添加
                            items.append(f"{idx+1}. {item[:100]}")

                    return '\n'.join(items) if items else None

            # 处理字典类型
            elif isinstance(data, dict):
                # 简化显示:只显示前3个键值对
                items = []
                for idx, (key, val) in enumerate(list(data.items())[:3]):
                    items.append(f"{key}: {val}")
                return '\n'.join(items) if items else None

            # 其他类型,转为字符串
            return str(data)

        except Exception as e:
            self.logger.error(f"格式化JSON字段失败({field_key}): {e}")
            return str(json_value) if json_value else None

    def _fill_cell(self, cell, value: str):
        """
        填充单元格

        Args:
            cell: 单元格对象
            value: 要填充的值
        """
        if not cell.paragraphs:
            cell.add_paragraph(value)
            return

        # 只处理第一个段落
        paragraph = cell.paragraphs[0]

        # 如果段落有Run,保存第一个Run的格式
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_properties = {
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'size': first_run.font.size,
                'name': first_run.font.name
            }

            # 清空并重新设置文本
            paragraph.clear()
            new_run = paragraph.add_run(value)

            # 恢复格式
            for prop, val in font_properties.items():
                if val is not None:
                    setattr(new_run.font, prop, val)
        else:
            # 没有Run,直接设置文本
            if paragraph.text:
                paragraph.clear()
            paragraph.add_run(value)

    def _insert_resume_images_after_table(self, doc: Document, table: Table, resumes: List[Dict]) -> int:
        """
        在简历表格后插入简历附件图片(复用ImageHandler的插入逻辑)

        处理逻辑:
        1. PDF附件: 使用转换后的图片(支持多页)
        2. 图片附件: 直接插入
        3. Word附件: 跳过(无法直接插入)

        Args:
            doc: Word文档对象
            table: 简历表格对象
            resumes: 已填充到表格中的简历列表

        Returns:
            插入的图片数量
        """
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from common import resolve_file_path
        import os

        images_inserted = 0

        # 找到表格后的插入点
        last_insert_para = self._find_para_after_table(doc, table)

        if not last_insert_para:
            self.logger.warning("  ⚠️  无法找到表格后的插入点,无法插入图片")
            return 0

        self.logger.info("  开始插入简历附件图片...")

        # 为每个简历插入附件图片
        for resume_idx, resume in enumerate(resumes):
            image_attachments = resume.get('image_attachments', [])

            if not image_attachments:
                self.logger.debug(f"    简历{resume_idx+1}无图片附件")
                continue

            resume_name = resume.get('name', f'简历{resume_idx+1}')
            self.logger.info(f"    简历 '{resume_name}' 有 {len(image_attachments)} 张图片")

            # 按附件顺序插入图片
            for img_att in image_attachments:
                file_path = resolve_file_path(img_att['file_path'])

                if not file_path or not os.path.exists(file_path):
                    self.logger.warning(f"      图片文件不存在: {img_att['file_path']}")
                    continue

                # 生成标题
                att_desc = img_att.get('description') or \
                          self._get_attachment_category_name(img_att.get('category'))

                # 多页PDF特殊处理:第一页插入标题,后续页不插入
                page_num = img_att.get('page_num', 1)
                is_multi_page = img_att.get('is_multi_page', False)

                if is_multi_page:
                    if page_num == 1:
                        title_text = f"{resume_name} - {att_desc}"
                    else:
                        title_text = None  # 后续页不插入标题
                else:
                    title_text = f"{resume_name} - {att_desc}"

                # 插入标题(第一页或单页)
                if title_text:
                    title_para = self.image_handler._insert_paragraph_after(last_insert_para)
                    title_para.text = title_text
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if title_para.runs:
                        title_para.runs[0].font.bold = True
                    last_insert_para = title_para
                    self.logger.debug(f"      插入标题: {title_text}")

                # 插入图片
                img_para = self.image_handler._insert_paragraph_after(last_insert_para)
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(file_path, width=Inches(6))  # 6英寸(与资质证书一致)

                last_insert_para = img_para
                images_inserted += 1

                if title_text:
                    self.logger.info(f"      ✅ 已插入: {title_text}")
                else:
                    self.logger.info(f"      ✅ 已插入第{page_num}页")

        self.logger.info(f"  简历附件图片插入完成,共插入 {images_inserted} 张图片")
        return images_inserted

    def _find_para_after_table(self, doc: Document, table: Table):
        """
        找到表格后的段落作为插入点

        Args:
            doc: Word文档对象
            table: 表格对象

        Returns:
            表格后的段落对象,如果找不到则返回None
        """
        try:
            # 获取表格元素
            table_element = table._element

            # 尝试找到表格后的第一个段落
            for para in doc.paragraphs:
                # 检查段落是否在表格之后
                if para._element.getprevious() == table_element:
                    self.logger.debug(f"    找到表格后的段落: '{para.text[:50]}'")
                    return para

            # 降级:遍历表格后的所有元素,找到第一个段落
            next_element = table_element.getnext()
            while next_element is not None:
                if next_element.tag.endswith('}p'):  # 找到段落元素
                    # 包装成Paragraph对象
                    from docx.text.paragraph import Paragraph
                    para = Paragraph(next_element, doc)
                    self.logger.debug(f"    通过遍历找到表格后的段落")
                    return para
                next_element = next_element.getnext()

            # 再降级:返回文档末尾段落
            if doc.paragraphs:
                self.logger.debug(f"    使用文档末尾段落作为插入点")
                return doc.paragraphs[-1]
            else:
                # 创建新段落
                self.logger.debug(f"    文档无段落,创建新段落")
                return doc.add_paragraph()

        except Exception as e:
            self.logger.error(f"查找表格后段落失败: {e}")
            # 降级:返回文档末尾
            if doc.paragraphs:
                return doc.paragraphs[-1]
            return doc.add_paragraph()

    def _is_image_file(self, file_path: str) -> bool:
        """
        判断是否为图片文件

        Args:
            file_path: 文件路径

        Returns:
            是否为图片文件
        """
        if not file_path:
            return False
        valid_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'}
        ext = Path(file_path).suffix.lower()
        return ext in valid_extensions

    def _get_attachment_category_name(self, category: str) -> str:
        """
        获取附件分类的中文名称

        Args:
            category: 附件分类代码

        Returns:
            附件分类中文名称
        """
        category_map = {
            'resume': '简历',
            'id_card': '身份证',
            'education': '学历证书',
            'degree': '学位证书',
            'qualification': '资质证书',
            'award': '获奖证书',
            'other': '其他附件'
        }
        return category_map.get(category, '附件')
