#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
资质验证器 - 验证图片和检测缺失资质
"""

import os
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger
from .document_utils import DocumentUtils


class QualificationValidator:
    """资质验证器 - 验证图片和检测缺失资质"""

    def __init__(self, utils: DocumentUtils, default_sizes: Dict):
        self.logger = get_module_logger("qualification_validator")
        self.utils = utils
        self.default_sizes = default_sizes

    def validate_images(self, image_paths: List[str]) -> Dict[str, Any]:
        """验证图片文件

        Args:
            image_paths: 图片路径列表

        Returns:
            验证结果字典，包含valid、invalid、missing三个列表
        """
        validation_result = {
            'valid': [],
            'invalid': [],
            'missing': []
        }

        for path in image_paths:
            if not path:
                continue

            if not os.path.exists(path):
                validation_result['missing'].append(path)
            elif not self._is_valid_image(path):
                validation_result['invalid'].append(path)
            else:
                validation_result['valid'].append(path)

        return validation_result

    def _is_valid_image(self, path: str) -> bool:
        """检查是否为有效的图片文件

        Args:
            path: 图片路径

        Returns:
            是否为有效图片格式
        """
        valid_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'}
        ext = Path(path).suffix.lower()
        return ext in valid_extensions

    def detect_missing_qualifications(self, insert_points: Dict, image_config: Dict,
                                     stats: Dict, qual_mapping: Dict) -> None:
        """
        检测缺失的资质（模板有占位符但公司无对应文件）

        Args:
            insert_points: 扫描到的插入点字典
            image_config: 图片配置（包含公司已上传的资质）
            stats: 统计信息字典（会被修改）
            qual_mapping: 资质映射表（QUALIFICATION_MAPPING）
        """
        # 获取公司已上传的资质keys
        uploaded_qual_keys = set()
        qualification_details = image_config.get('qualification_details', [])
        for qual_detail in qualification_details:
            qual_key = qual_detail.get('qual_key')
            if qual_key:
                uploaded_qual_keys.add(qual_key)

        # 遍历所有发现的占位符
        for placeholder_key in insert_points.keys():
            # 跳过基础类型（license, legal_id等，这些不是资质证书）
            if placeholder_key in ['license', 'qualification', 'legal_id', 'auth_id',
                                   'authorization', 'certificate']:
                continue

            # 检查该占位符是否有对应的公司资质文件
            if placeholder_key not in uploaded_qual_keys:
                # 有占位符但无文件 → 缺失资质
                qual_name = qual_mapping.get(placeholder_key, {}).get('category', placeholder_key)
                stats['missing_qualifications'].append({
                    'qual_key': placeholder_key,
                    'qual_name': qual_name,
                    'placeholder': insert_points[placeholder_key].get('matched_keyword', '')
                })
                self.logger.warning(f"⚠️  缺失资质: {placeholder_key} ({qual_name}) - 模板有占位符但公司未上传")

    def append_required_qualifications(self, doc: Document, required_quals: List[Dict],
                                      insert_points: Dict, image_config: Dict,
                                      stats: Dict, qual_mapping: Dict) -> None:
        """
        追加项目要求但模板没有占位符的资质

        Args:
            doc: Word文档对象
            required_quals: 项目资格要求列表
            insert_points: 已扫描的插入点
            image_config: 图片配置
            stats: 统计信息字典（会被修改）
            qual_mapping: 资质映射表
        """
        # 获取公司已上传的资质（key -> file_path映射）
        uploaded_quals_map = {}
        qualification_details = image_config.get('qualification_details', [])
        for qual_detail in qualification_details:
            qual_key = qual_detail.get('qual_key')
            file_path = qual_detail.get('file_path')
            if qual_key and file_path:
                uploaded_quals_map[qual_key] = qual_detail

        # 遍历项目要求的资质
        for req_qual in required_quals:
            qual_key = req_qual.get('qual_key')
            if not qual_key:
                continue

            # 判断条件：项目要求 + 公司有文件 + 模板无占位符
            has_file = qual_key in uploaded_quals_map
            has_placeholder = (qual_key in insert_points or 'qualification' in insert_points)

            if has_file and not has_placeholder:
                # 需要追加：项目要求且公司有文件，但模板没有对应占位符
                qual_detail = uploaded_quals_map[qual_key]
                file_path = qual_detail['file_path']
                insert_hint = req_qual.get('source_detail', '')
                qual_name = qual_mapping.get(qual_key, {}).get('category', qual_key)

                # 在文档末尾追加该资质
                try:
                    if self._append_qualification_to_end(doc, file_path, qual_key, insert_hint):
                        stats['images_inserted'] += 1
                        stats['images_types'].append(f'{qual_key}_appended')
                        stats['appended_qualifications'].append({
                            'qual_key': qual_key,
                            'qual_name': qual_name,
                            'file_path': file_path,
                            'reason': '项目要求但模板无占位符'
                        })
                        self.logger.info(f"✅ 追加资质: {qual_key} ({qual_name}) - 项目要求但模板无占位符")
                    else:
                        self.logger.error(f"❌ 追加资质失败: {qual_key}")
                except Exception as e:
                    self.logger.error(f"❌ 追加资质异常: {qual_key}, 错误: {e}")

    def _append_qualification_to_end(self, doc: Document, image_path: str,
                                    qual_key: str, insert_hint: str = None) -> bool:
        """
        在文档末尾追加资质证书

        Args:
            doc: Word文档对象
            image_path: 图片路径
            qual_key: 资质键
            insert_hint: 插入提示（用于生成标题）

        Returns:
            bool: 是否成功
        """
        try:
            # 解析路径（支持相对路径）
            resolved_path = self.utils.resolve_file_path(image_path)
            if not os.path.exists(resolved_path):
                self.logger.error(f"资质图片不存在: {image_path} (resolved: {resolved_path})")
                return False
            image_path = resolved_path  # 使用解析后的路径

            # 生成标题（优先级: display_title > insert_hint > category + "认证证书"）
            from .qualification_matcher import QUALIFICATION_MAPPING
            if qual_key in QUALIFICATION_MAPPING:
                qual_info = QUALIFICATION_MAPPING[qual_key]
                # 优先使用 display_title（如果存在）
                if 'display_title' in qual_info:
                    title_text = qual_info['display_title']
                elif insert_hint:
                    title_text = insert_hint[:50]
                else:
                    title_text = f"{qual_info['category']}认证证书"
            elif insert_hint:
                title_text = insert_hint[:50]
            else:
                title_text = f"资质证书 ({qual_key})"

            # 添加分页符
            doc.add_page_break()

            # 添加标题
            title = doc.add_paragraph(title_text)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title.runs:
                title.runs[0].font.bold = True

            # 添加图片
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(self.default_sizes.get(qual_key, (6, 0))[0]))

            self.logger.info(f"✅ 已在文档末尾追加资质: {title_text}")
            return True

        except Exception as e:
            self.logger.error(f"❌ 追加资质到文档末尾失败: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return False
