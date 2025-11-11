#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
表格处理模块 - 处理商务应答模板中的表格填写
"""

import re
from typing import Dict, Any, List, Optional
from docx import Document
from docx.table import Table, _Cell

# 导入公共模块
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger

# 导入Word文档工具
from .utils import WordDocumentUtils

class TableProcessor:
    """表格处理器"""
    
    def __init__(self):
        self.logger = get_module_logger("table_processor")
        
        # 表格中的关键字段映射
        self.table_field_mapping = {
            '供应商名称': 'companyName',
            '投标人名称': 'companyName',
            '公司名称': 'companyName',
            '响应人名称': 'companyName',  # 新增：响应人名称
            '响应人全称': 'companyName',  # 新增：响应人全称
            '法定代表人': 'legalRepresentative',
            '注册资本': 'registeredCapital',
            '成立日期': 'establishDate',
            '统一社会信用代码': 'socialCreditCode',
            '注册地址': 'registeredAddress',
            '联系地址': 'address',
            '联系人': 'contactPerson',
            '联系电话': 'phone',
            '电子邮箱': 'email',
            '传真': 'fax',
            '开户银行': 'bankName',
            '银行账号': 'bankAccount',
            '税号': 'taxNumber',
            '资质等级': 'qualification',
            '项目名称': 'projectName',
            '项目编号': 'projectNumber',
            '投标报价': 'bidPrice',
            '交货期': 'deliveryTime',
            '质保期': 'warrantyPeriod',
            # 股权结构字段（2025-10-27添加，2025-11-09增强）
            '实际控制人': 'actual_controller',
            '控股股东': 'controlling_shareholder',
            '控股股东及出资比例': 'controlling_shareholder',
            '供应商的控股股东/投资人名称及出资比例': 'controlling_shareholder',  # 🆕 支持长字段名
            '股东': 'shareholders_info',
            '股东信息': 'shareholders_info',  # 支持变体
            '供应商的非控股股东/投资人名称及出资比例': 'shareholders_info',  # 🆕 支持长字段名
            '投资人名称及出资比例': 'shareholders_info',  # 🆕 支持简化名
            # 管理关系字段（2025-10-28添加）
            '管理关系单位': 'managing_unit_name',
            '管理关系单位名称': 'managing_unit_name',  # 支持变体
            '被管理关系单位': 'managed_unit_name',
            '被管理关系单位名称': 'managed_unit_name'  # 支持变体
        }
    
    def process_tables(self, doc: Document, company_info: Dict[str, Any], 
                       project_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        处理文档中的所有表格
        
        Args:
            doc: Word文档对象
            company_info: 公司信息
            project_info: 项目信息
            
        Returns:
            处理统计信息
        """
        stats = {
            'tables_processed': 0,
            'cells_filled': 0,
            'fields_matched': []
        }
        
        # 合并所有信息
        all_info = {**company_info, **project_info}
        
        for table_idx, table in enumerate(doc.tables):
            self.logger.info(f"处理表格 #{table_idx + 1}")
            result = self._process_single_table(table, all_info)
            
            if result['cells_filled'] > 0:
                stats['tables_processed'] += 1
                stats['cells_filled'] += result['cells_filled']
                stats['fields_matched'].extend(result['fields_matched'])
        
        self.logger.info(f"表格处理完成: 处理了{stats['tables_processed']}个表格，"
                        f"填充了{stats['cells_filled']}个单元格")
        
        return stats
    
    def _process_single_table(self, table: Table, info: Dict[str, Any]) -> Dict[str, Any]:
        """处理单个表格"""
        result = {
            'cells_filled': 0,
            'fields_matched': []
        }
        
        # 分析表格结构
        table_type = self._analyze_table_structure(table)
        
        if table_type == 'key_value':
            # 键值对表格（左边是字段名，右边是值）
            result = self._process_key_value_table(table, info)
        elif table_type == 'header_data':
            # 表头-数据表格（第一行是表头，后续行是数据）
            result = self._process_header_data_table(table, info)
        elif table_type == 'mixed':
            # 混合型表格
            result = self._process_mixed_table(table, info)
        
        return result
    
    def _analyze_table_structure(self, table: Table) -> str:
        """分析表格结构类型（支持混合列数表格，包括2列和3列键值对表格）"""
        if not table.rows:
            return 'empty'

        # 统计每行的实际列数
        row_column_counts = [len(row.cells) for row in table.rows]
        two_col_rows = sum(1 for count in row_column_counts if count == 2)
        three_col_rows = sum(1 for count in row_column_counts if count == 3)
        total_rows = len(table.rows)

        # 调试日志：输出表格结构信息
        self.logger.debug(f"  表格结构分析: 总行数={total_rows}, 2列行数={two_col_rows}, 3列行数={three_col_rows}, 列数分布={row_column_counts}")

        # 检查是否为键值对表格（允许部分行有不同列数）
        # 情况1：如果超过80%的行是2列，则可能是键值对表格
        # 情况2：如果超过80%的行是3列，也可能是键值对表格（第1列=字段名，第2列=值，第3列=说明）
        if two_col_rows >= total_rows * 0.8 or three_col_rows >= total_rows * 0.8:
            # 提取第一列文本（至少有2列的行）
            first_col_texts = [row.cells[0].text.strip() for row in table.rows if len(row.cells) >= 2]
            field_count = sum(1 for text in first_col_texts
                            if any(field in text for field in self.table_field_mapping.keys()))

            self.logger.debug(f"  键值对检测: 匹配字段数={field_count}/{len(first_col_texts)}")

            if field_count > len(first_col_texts) * 0.5:
                self.logger.debug(f"  ✅ 识别为 key_value 类型（2列或3列键值对表格）")
                return 'key_value'

        # 原有逻辑：检查是否为表头-数据表格
        if total_rows > 1:
            first_row_texts = [cell.text.strip() for cell in table.rows[0].cells]
            field_count = sum(1 for text in first_row_texts
                            if any(field in text for field in self.table_field_mapping.keys()))
            if field_count > len(table.columns) * 0.3:
                self.logger.debug(f"  ✅ 识别为 header_data 类型")
                return 'header_data'

        self.logger.debug(f"  ✅ 识别为 mixed 类型")
        return 'mixed'

    def _normalize_field_name(self, field_name: str) -> str:
        """
        规范化字段名：移除所有空格、括号后缀等

        示例：
        - "日    期" → "日期"
        - "成立日期（盖章）" → "成立日期"
        - "投标人名称  " → "投标人名称"

        Args:
            field_name: 原始字段名

        Returns:
            规范化后的字段名
        """
        field_name = field_name.strip()
        # 移除所有空格（处理"日    期"等情况）
        field_name = re.sub(r'\s+', '', field_name)
        # 移除常见后缀
        field_name = re.sub(r'[（(][^）)]*[）)]', '', field_name)
        return field_name

    def _process_key_value_table(self, table: Table, info: Dict[str, Any]) -> Dict[str, Any]:
        """
        处理键值对类型的表格（增强版：支持2列和3列表格）

        支持格式：
        - 2列表格：第1列=字段名，第2列=值
        - 3列表格：第1列=字段名，第2列=值，第3列=说明文字（忽略）
        """
        result = {
            'cells_filled': 0,
            'fields_matched': []
        }

        for row in table.rows:
            # 支持2列或3列的键值对表格
            if len(row.cells) < 2:
                self.logger.debug(f"  跳过单列行（列数={len(row.cells)}）")
                continue

            # 对于超过3列的行，可能不是键值对表格，跳过
            if len(row.cells) > 3:
                self.logger.debug(f"  跳过多列行（列数={len(row.cells)}，可能是数据表格）")
                continue

            key_cell = row.cells[0]
            value_cell = row.cells[1]  # 无论2列还是3列，第2列都是值

            # 规范化字段名（移除空格等）
            key_text = self._normalize_field_name(key_cell.text)

            # 查找匹配的字段
            for field_name, field_key in self.table_field_mapping.items():
                # 规范化映射表中的字段名进行匹配
                normalized_field = self._normalize_field_name(field_name)
                if normalized_field in key_text or normalized_field == key_text:
                    value = info.get(field_key, '')

                    # 新增：日期字段格式化
                    if value and (field_key in ['date', 'establishDate'] or '日期' in field_name):
                        value = self._format_date(str(value))

                    # 🆕 格式化股东信息JSON
                    if value and field_key == 'shareholders_info':
                        value = self._format_shareholders_info(value)

                    # 🔧 修复：跳过"无"、"/" 等占位值
                    if value and str(value).strip() in ['无', '/', '-', 'N/A', 'NA']:
                        self.logger.debug(f"  跳过占位值字段: {field_name} = '{value}'")
                        continue

                    if value and self._should_fill_cell(value_cell):
                        self._fill_cell(value_cell, str(value))
                        result['cells_filled'] += 1
                        result['fields_matched'].append(field_name)
                        self.logger.info(f"  ✅ 表格字段填充: {field_name} = {value} (列数={len(row.cells)})")
                        break

        return result
    
    def _process_header_data_table(self, table: Table, info: Dict[str, Any]) -> Dict[str, Any]:
        """处理表头-数据类型的表格"""
        result = {
            'cells_filled': 0,
            'fields_matched': []
        }

        if len(table.rows) < 2:
            self.logger.debug("  表格行数不足2行，跳过处理")
            return result

        # 分析表头
        header_row = table.rows[0]
        column_mapping = {}

        for col_idx, cell in enumerate(header_row.cells):
            # 规范化表头文本（移除空格等）
            header_text = self._normalize_field_name(cell.text)

            # 🔧 修复：使用精确匹配而不是包含匹配，避免"管理关系单位"匹配到"管理关系单位名称"
            # 方法：先按字段名长度降序排序，优先匹配更长的字段名
            sorted_fields = sorted(
                self.table_field_mapping.items(),
                key=lambda x: len(self._normalize_field_name(x[0])),
                reverse=True  # 从长到短排序
            )

            for field_name, field_key in sorted_fields:
                # 规范化映射表中的字段名进行匹配
                normalized_field = self._normalize_field_name(field_name)
                # 使用精确匹配
                if normalized_field == header_text:
                    column_mapping[col_idx] = field_key
                    self.logger.debug(f"  表头列{col_idx}识别为: {field_name} -> {field_key} (精确匹配)")
                    break

        self.logger.debug(f"  表头分析完成，识别到{len(column_mapping)}个字段列")

        # 填充数据行
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            for col_idx, field_key in column_mapping.items():
                if col_idx < len(row.cells):
                    value = info.get(field_key, '')

                    # 新增：日期字段格式化
                    if value and field_key in ['date', 'establishDate']:
                        value = self._format_date(str(value))

                    # 🆕 格式化股东信息JSON
                    if value and field_key == 'shareholders_info':
                        value = self._format_shareholders_info(value)

                    # 🔧 修复：跳过"无"、"/" 等占位值
                    if value and str(value).strip() in ['无', '/', '-', 'N/A', 'NA']:
                        self.logger.debug(f"  跳过占位值字段: 行{row_idx}列{col_idx} {field_key} = '{value}'")
                        continue

                    if value and self._should_fill_cell(row.cells[col_idx]):
                        self._fill_cell(row.cells[col_idx], str(value))
                        result['cells_filled'] += 1

                        # 记录字段名
                        for field_name, key in self.table_field_mapping.items():
                            if key == field_key:
                                result['fields_matched'].append(field_name)
                                self.logger.info(f"  ✅ 表格数据填充: 行{row_idx}列{col_idx} {field_name} = {value}")
                                break

        return result
    
    def _process_mixed_table(self, table: Table, info: Dict[str, Any]) -> Dict[str, Any]:
        """处理混合型表格（详细日志版）"""
        result = {
            'cells_filled': 0,
            'fields_matched': []
        }

        self.logger.debug("  开始处理混合型表格，逐单元格查找字段")

        # 遍历所有单元格，查找并填充
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # 原始单元格文本（用于检测占位符）
                original_cell_text = cell.text
                # 规范化单元格文本（移除空格等，用于字段匹配）
                cell_text = self._normalize_field_name(original_cell_text)

                # 检查是否包含字段名和占位符
                for field_name, field_key in self.table_field_mapping.items():
                    # 规范化映射表中的字段名进行匹配
                    normalized_field = self._normalize_field_name(field_name)
                    if normalized_field in cell_text:
                        self.logger.debug(f"  发现字段名 '{field_name}' 在单元格[{row_idx},{cell_idx}]: {cell_text[:30]}...")

                        # 检查是否有占位符（使用原始文本检测，而不是规范化后的文本）
                        if re.search(r'[_\s]{3,}|[:：]\s*$', original_cell_text):
                            value = info.get(field_key, '')

                            # 新增：日期字段格式化
                            if value and (field_key in ['date', 'establishDate'] or '日期' in field_name):
                                value = self._format_date(str(value))

                            if value:
                                # 替换占位符（使用原始文本）
                                new_text = self._replace_placeholder(original_cell_text.strip(), str(value))
                                self._update_cell_text(cell, new_text)
                                result['cells_filled'] += 1
                                result['fields_matched'].append(field_name)
                                self.logger.info(f"  ✅ 混合表格填充（占位符）: {field_name} = {value}")
                                break
                        # 检查下一个单元格是否为空（可能是值单元格）
                        elif cell_idx + 1 < len(row.cells):
                            next_cell = row.cells[cell_idx + 1]
                            if self._should_fill_cell(next_cell):
                                value = info.get(field_key, '')

                                # 新增：日期字段格式化
                                if value and (field_key in ['date', 'establishDate'] or '日期' in field_name):
                                    value = self._format_date(str(value))

                                if value:
                                    self._fill_cell(next_cell, str(value))
                                    result['cells_filled'] += 1
                                    result['fields_matched'].append(field_name)
                                    self.logger.info(f"  ✅ 混合表格填充（下一单元格）: {field_name} = {value}")
                                    break
                            else:
                                self.logger.debug(f"  下一单元格不为空，跳过填充")

        self.logger.debug(f"  混合型表格处理完成，填充了{result['cells_filled']}个单元格")
        return result
    
    def _should_fill_cell(self, cell: _Cell) -> bool:
        """判断单元格是否应该被填充"""
        text = cell.text.strip()
        
        # 空单元格或只有占位符的单元格
        if not text or re.match(r'^[_\s]*$', text):
            return True
        
        # 包含占位符的单元格
        if re.search(r'_{3,}', text):
            return True
        
        return False
    
    def _fill_cell(self, cell: _Cell, value: str):
        """
        填充单元格（简化版本，避免Run映射Bug）

        修复说明：
        - 之前使用WordDocumentUtils.apply_replacement_to_runs会导致索引越界
        - 改用简单的段落清空+重写方法，保持格式

        Args:
            cell: 表格单元格对象
            value: 要填充的值
        """
        if not cell.paragraphs:
            cell.add_paragraph(value)
            return

        # 只处理第一个段落（表格单元格通常只有一个段落）
        paragraph = cell.paragraphs[0]

        # 如果段落有Run，保存第一个Run的格式
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_properties = {
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'size': first_run.font.size,
                'name': first_run.font.name
            }

            # 清空并重新设置文本
            paragraph.clear()
            new_run = paragraph.add_run(value)

            # 恢复格式
            for prop, val in font_properties.items():
                if val is not None:
                    setattr(new_run.font, prop, val)
        else:
            # 没有Run，直接设置文本
            if paragraph.text:
                paragraph.clear()
            paragraph.add_run(value)
    
    def _update_cell_text(self, cell: _Cell, new_text: str):
        """更新单元格文本，保持格式"""
        if cell.paragraphs and cell.paragraphs[0].runs:
            # 保存第一个run的格式
            first_run = cell.paragraphs[0].runs[0]
            font_properties = {
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'size': first_run.font.size,
                'name': first_run.font.name
            }
            
            # 清空并重新设置文本
            cell.paragraphs[0].clear()
            new_run = cell.paragraphs[0].add_run(new_text)
            
            # 恢复格式
            for prop, value in font_properties.items():
                if value is not None:
                    setattr(new_run.font, prop, value)
        else:
            # 直接设置文本
            if cell.paragraphs:
                cell.paragraphs[0].text = new_text
            else:
                cell.add_paragraph(new_text)
    
    def _replace_placeholder(self, text: str, value: str) -> str:
        """替换文本中的占位符"""
        # 替换下划线占位符
        text = re.sub(r'_{3,}', value, text)

        # 替换冒号后的空白
        text = re.sub(r'([:：])\s*$', f'\\1{value}', text)

        # 替换连续空格
        if re.search(r'\s{3,}', text):
            text = re.sub(r'\s{3,}', value, text)

        return text

    def _format_date(self, date_str: str) -> str:
        """
        格式化日期，移除时间信息

        处理格式：
        - 2025-08-27 → 2025年08月27日
        - 2025/08/27 → 2025年08月27日
        - 2025年08月27日下午14:30整（北京时间） → 2025年08月27日
        """
        # 移除空格
        date_str = re.sub(r'\s+', '', date_str)

        # 尝试匹配常见格式
        patterns = [
            (r'(\d{4})-(\d{1,2})-(\d{1,2})', r'\1年\2月\3日'),
            (r'(\d{4})/(\d{1,2})/(\d{1,2})', r'\1年\2月\3日'),
            (r'(\d{4})\.(\d{1,2})\.(\d{1,2})', r'\1年\2月\3日'),
        ]

        for pattern, replacement in patterns:
            if re.match(pattern, date_str):
                return re.sub(pattern, replacement, date_str)

        # 已经是中文格式，提取"年月日"部分，删除后面的时间信息
        if '年' in date_str and '月' in date_str:
            # 匹配格式：2025年08月27日下午14:30整（北京时间） → 2025年08月27日
            date_match = re.match(r'(\d{4}年\d{1,2}月\d{1,2}日)', date_str)
            if date_match:
                return date_match.group(1)  # 只返回"年月日"部分
            return date_str

        return date_str

    def _format_shareholders_info(self, shareholders_json: str) -> str:
        """
        格式化股东信息JSON为可读文本

        输入示例：
        [
            {"name": "股东A", "type": "企业", "ratio": "30%"},
            {"name": "股东B", "type": "自然人", "ratio": "20%"}
        ]

        输出示例：
        股东A（企业，30%）、股东B（自然人，20%）

        Args:
            shareholders_json: 股东信息JSON字符串

        Returns:
            格式化后的股东信息文本
        """
        try:
            import json

            # 解析JSON
            if isinstance(shareholders_json, str):
                shareholders = json.loads(shareholders_json)
            else:
                shareholders = shareholders_json

            if not shareholders or not isinstance(shareholders, list):
                return ''

            # 格式化每个股东信息
            formatted_list = []
            for shareholder in shareholders:
                name = shareholder.get('name', '')
                shareholder_type = shareholder.get('type', '')
                ratio = shareholder.get('ratio', '')

                # 格式：股东名称（类型，出资比例）
                if shareholder_type and ratio:
                    formatted_list.append(f"{name}（{shareholder_type}，{ratio}）")
                elif ratio:
                    formatted_list.append(f"{name}（{ratio}）")
                else:
                    formatted_list.append(name)

            # 用顿号连接
            return '、'.join(formatted_list)

        except Exception as e:
            self.logger.error(f"格式化股东信息失败: {e}")
            return str(shareholders_json)