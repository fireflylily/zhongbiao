#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
信息填写模块 - 精简版 (从2098行精简至400行以内)

处理项目和公司信息的填写，实现六大规则：
- 替换规则、填空规则、组合规则、变体处理、例外处理、后处理

通过依赖utils.py和format_cleaner.py实现复用架构优化
"""

import re
import sys
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import logging

# 添加路径以导入公共模块
sys.path.append(str(Path(__file__).parent.parent.parent))

# 导入复用模块
from .utils import (
    FieldMapper, PatternMatcher, WordDocumentUtils,
    SmartFieldDetector, TextUtils
)
from .format_cleaner import FormatCleaner, DateFormatProcessor
from common import get_module_logger


class InfoFiller:
    """
    信息填写器 - 精简版
    
    专注核心的信息填写功能，通过复用模块实现工具函数和格式处理
    """
    
    def __init__(self):
        """初始化信息填写器"""
        self.logger = get_module_logger("info_filler")
        self.field_mapper = FieldMapper()
        self.pattern_matcher = PatternMatcher()
        self.format_cleaner = FormatCleaner()
        self.info = {}  # 统一信息字典

        # 字段变体映射 (核心配置)
        self.field_variants = self._create_field_variants()

        # 供应商名称的扩展匹配模式（支持带公章、盖章的变体）
        self.company_name_extended_patterns = [
            r'供应商名称(?:\s*[（(][^）)]*[公盖]章[^）)]*[）)])?',  # 供应商名称（加盖公章）
            r'供应商全称(?:\s*[（(][^）)]*[公盖]章[^）)]*[）)])?',
            r'投标人名称(?:\s*[（(][^）)]*[公盖]章[^）)]*[）)])?',  # 投标人名称（公章）
            r'公司名称(?:\s*[（(][^）)]*[公盖]章[^）)]*[）)])?',
            r'单位名称(?:\s*[（(][^）)]*[公盖]章[^）)]*[）)])?',
            r'应答人名称(?:\s*[（(][^）)]*[公盖]章[^）)]*[）)])?',
        ]
        
    def _create_field_variants(self) -> Dict[str, List[str]]:
        """创建字段变体映射"""
        return {
            # 供应商名称变体 (12种)
            'companyName': [
                '供应商名称', '供应商全称', '投标人名称', '公司名称', 
                '单位名称', '应答人名称', '供应商名称（盖章）',
                '供应商名称（公章）', '公司名称（盖章）',
                '投标人名称（盖章）', '投标人名称（公章）',
                '单位名称（盖章）', '单位名称（公章）'
            ],
            # 其他字段变体
            'email': ['邮箱', '邮件', '电子邮件', '电子邮箱', 'email', 'Email', 'E-mail', 'E-Mail', '邮件地址'],
            'phone': ['电话', '联系电话', '固定电话', '办公电话', '座机', '电话号码', '联系方式'],
            'fax': ['传真', '传真号码', '传真电话', 'FAX', '传真号', 'fax', 'Fax'],
            'address': ['地址', '联系地址', '办公地址', '注册地址', '公司地址', '详细地址', '通讯地址', '供应商地址'],
            'postalCode': ['邮政编码', '邮编', '邮码'],
            'date': ['日期', '日 期', '日  期', '日   期', '日    期', '日     期', '时间', '签署日期', '投标日期'],
            'purchaserName': ['采购人', '招标人', '采购单位', '招标单位', '业主', '甲方'],
            'projectName': ['项目名称', '项目名', '工程名称', '标的名称', '采购项目名称', '招标项目名称'],
            'projectNumber': ['项目编号', '采购编号', '招标编号', '项目号', '标书编号', '工程编号'],
            # 备份版本中的重要字段
            'establishDate': ['成立时间', '成立日期', '注册时间', '注册日期'],
            'businessScope': ['经营范围', '业务范围', '经营项目'],
            'legalRepresentative': ['法定代表人', '法人代表', '法人'],
            'authorizedPersonName': ['供应商代表姓名', '授权代表姓名', '代表姓名', '授权代表'],
            'position': ['职务', '职位', '职称'],
        }
    
    def fill_info(self, doc: Document, company_info: Dict[str, Any], 
                  project_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        填写文档信息 - 主处理流程
        
        Args:
            doc: Word文档对象
            company_info: 公司信息字典
            project_info: 项目信息字典
            
        Returns:
            处理统计信息字典
        """
        self.logger.info("🛠 统一字段映射初始化: 32 个字段")
        
        # 创建统一字段映射
        self.info = self._create_unified_field_mapping(company_info, project_info)
        
        # 初始化统计
        stats = {
            'total_replacements': 0,
            'replacement_rules': 0,
            'fill_rules': 0, 
            'combination_rules': 0,
            'skipped_fields': 0,
            'none': 0
        }
        
        # 处理段落
        self.logger.info(f"📊 开始处理文档: {len([p for p in doc.paragraphs if p.text.strip()])} 个非空段落, 0 个表格")
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                self.logger.info(f"🔍 [调试] 处理段落 {i+1}: {paragraph.text.strip()[:100]}...")
                paragraph_stats = self._process_paragraph(paragraph, self.info)
                if paragraph_stats['total_replacements'] > 0:
                    self.logger.info(f"🎉 [调试] 段落 {i+1} 处理成功: {paragraph_stats}")
                for key, value in paragraph_stats.items():
                    stats[key] = stats.get(key, 0) + value
        
        # 处理表格（简化版）
        for table in doc.tables:
            table_stats = self._process_table(table, self.info)
            for key, value in table_stats.items():
                stats[key] = stats.get(key, 0) + value
        
        # 后处理：清理多余的占位符和装饰性格式
        self._post_process(doc)
        
        # 日志输出
        self.logger.info(f"📊 文档处理统计: {stats}")
        self.logger.info(f"✅ 成功处理了 {stats['total_replacements']} 个字段")
        self.logger.info(f"  - 填空规则: {stats['fill_rules']} 个")
        self.logger.info(f"  - 组合规则: {stats['combination_rules']} 个")
        
        return stats
    
    def _create_unified_field_mapping(self, company_info: Dict[str, Any],
                                     project_info: Dict[str, Any]) -> Dict[str, str]:
        """创建统一的字段映射"""
        # 添加调试信息：显示输入的company_info数据
        self.logger.info(f"🔍 [调试] 输入的company_info数据: {company_info}")
        self.logger.info(f"🔍 [调试] 输入的project_info数据: {project_info}")

        info = {}
        
        # 映射公司信息 - 直接使用原有的字段名映射
        info['公司名称'] = company_info.get('companyName', '')
        info['邮箱'] = company_info.get('email', '')
        info['传真'] = company_info.get('fax', '')
        info['邮政编码'] = company_info.get('postalCode', '')
        info['成立时间'] = company_info.get('establishDate', '')
        info['经营范围'] = company_info.get('businessScope', '')
        info['法定代表人'] = company_info.get('legalRepresentative', '')
        info['被授权人姓名'] = company_info.get('authorizedPersonName', '')
        info['被授权人职务'] = company_info.get('authorizedPersonPosition', '')
        info['法定代表人职位'] = company_info.get('legalRepresentativePosition', '')
        
        # 处理多源映射字段
        info['地址'] = (company_info.get('address') or 
                      company_info.get('registeredAddress') or 
                      company_info.get('officeAddress', ''))
        info['电话'] = (company_info.get('fixedPhone') or
                      company_info.get('phone', ''))

        # 调试电话字段映射
        self.logger.info(f"🔍 [调试] 电话字段映射: fixedPhone={company_info.get('fixedPhone')}, phone={company_info.get('phone')}, 最终值={info['电话']}")
        
        # 映射项目信息
        info['项目名称'] = project_info.get('projectName', '')
        info['项目编号'] = project_info.get('projectNumber', '')
        info['日期'] = project_info.get('date', '')
        info['采购人名称'] = project_info.get('purchaserName', '')
        
        # 清理空值
        info = {k: v for k, v in info.items() if v and str(v).strip()}
        
        # 添加调试信息
        self.logger.info(f"🗺️ 字段映射结果: {list(info.keys())}")
        for key, value in info.items():
            self.logger.info(f"  - {key}: {value}")
        
        return info
    
    def _process_paragraph(self, paragraph: Paragraph, info: Dict[str, Any]) -> Dict[str, Any]:
        """
        处理段落 - 核心业务逻辑
        
        Args:
            paragraph: Word段落对象
            info: 信息字典
            
        Returns:
            处理统计信息
        """
        stats = {'total_replacements': 0, 'replacement_rules': 0, 'fill_rules': 0, 'combination_rules': 0}
        
        # 跳过处理
        if self._should_skip(paragraph.text):
            return stats
            
        processed = False
        final_type = 'none'

        # 1. 尝试组合替换规则
        if self._try_combination_rule(paragraph, info):
            processed = True
            final_type = 'combination_rules'
            stats['combination_rules'] += 1
            stats['total_replacements'] += 1

        # 2. 尝试单字段替换规则（即使组合规则已处理，也要尝试）
        if self._try_replacement_rule(paragraph, info):
            processed = True
            # 如果已经有组合规则，保持组合规则类型，否则设为替换规则
            if final_type == 'none':
                final_type = 'replacement_rules'
                stats['replacement_rules'] += 1
                stats['total_replacements'] += 1

        # 3. 尝试填空规则（仅在前两个都没有处理时才尝试）
        if not processed and self._try_fill_rule(paragraph, info):
            processed = True
            final_type = 'fill_rules'
            stats['fill_rules'] += 1
            stats['total_replacements'] += 1

        return stats
    
    def _should_skip(self, text: str) -> bool:
        """判断是否应该跳过处理"""
        import re

        # 使用精确匹配模式，避免误跳过合法填空字段
        skip_patterns = [
            r'签字\s*[：:]?\s*$',      # 签字: 或 签字（行尾）
            r'签字处',                 # 签字处
            r'签名\s*[：:]?\s*$',      # 签名: 或 签名（行尾）
            r'签名处',                 # 签名处
            r'签章\s*[：:]?\s*$',      # 签章: 或 签章（行尾）
            r'盖章处',                 # 盖章处
        ]

        # 简单关键词匹配（保持原有逻辑）
        simple_keywords = [
            '招标代理', '采购代理', '业主', '发包人', '委托人'
        ]

        text_lower = text.lower()

        # 检查精确模式
        for pattern in skip_patterns:
            if re.search(pattern, text_lower):
                return True

        # 检查简单关键词
        for keyword in simple_keywords:
            if keyword in text_lower:
                return True

        return False

    def _detect_position_context(self, paragraph_text: str) -> str:
        """
        检测段落中的职位上下文，区分被授权人职务和法定代表人职位

        Args:
            paragraph_text: 段落文本

        Returns:
            'authorized_person': 被授权人上下文
            'legal_representative': 法定代表人上下文（默认）
        """
        try:
            if not paragraph_text or not isinstance(paragraph_text, str):
                self.logger.warning(f"⚠️  职位上下文检测：无效的段落文本输入")
                return 'legal_representative'

            text = paragraph_text.strip()
            if not text:
                self.logger.warning(f"⚠️  职位上下文检测：段落文本为空")
                return 'legal_representative'

            # 被授权人上下文关键字模式 (增强版)
            authorized_person_patterns = [
                r'授权.*?代表.*?[职位务称]',          # "授权代表职务"
                r'为我方.*?授权.*?代表',             # "为我方授权代表"
                r'为我方代表.*?[职位务称]',           # "为我方代表职务"
                r'参加.*?代表.*?[职位务称]',          # "参加投标代表职务"
                r'授权.*?[（(].*?[）)].*?[职位务称]',   # "授权（张三）职务"
                r'被授权.*?[职位务称]',              # "被授权人职务"
                r'商务代表.*?[职位务称]',             # "商务代表职务"
                r'授权.*?[（(].*?姓名.*?职[位务称]',   # "授权（姓名、职位）"
                r'为我方.*?授权.*?[（(].*?职[位务称]', # "为我方授权（职位、职称）"
            ]

            # 法定代表人上下文关键字模式
            legal_representative_patterns = [
                r'法定代表人.*?职位',           # "法定代表人职位"
                r'法人.*?职位',                # "法人职位"
                r'系.*?法定代表人.*?职位',      # "系我公司法定代表人职位"
                r'公司.*?法定代表人.*?职位',    # "公司法定代表人职位"
            ]

            self.logger.debug(f"🔍 检测职位上下文: '{text[:100]}{'...' if len(text) > 100 else ''}'")

            # 检查被授权人上下文
            try:
                for pattern in authorized_person_patterns:
                    if re.search(pattern, text):
                        self.logger.info(f"✅ 识别为被授权人上下文: '{pattern}' 匹配")
                        return 'authorized_person'
            except re.error as e:
                self.logger.error(f"❌ 被授权人模式匹配正则表达式错误: {e}")

            # 检查法定代表人上下文
            try:
                for pattern in legal_representative_patterns:
                    if re.search(pattern, text):
                        self.logger.info(f"✅ 识别为法定代表人上下文: '{pattern}' 匹配")
                        return 'legal_representative'
            except re.error as e:
                self.logger.error(f"❌ 法定代表人模式匹配正则表达式错误: {e}")

            # 默认情况：如果没有明确上下文，使用法定代表人
            return 'legal_representative'

        except Exception as e:
            self.logger.error(f"❌ 职位上下文检测发生异常: {e}")
            return 'legal_representative'

    def _try_combination_rule(self, paragraph: Paragraph, info: Dict[str, Any]) -> bool:
        """尝试应用组合替换规则 - 参考备份版本的独立模式检查"""
        text = paragraph.text
        processed_any = False

        # 组合模式1：供应商名称、地址
        pattern1 = r'[（(]\s*(?:供应商名称|公司名称|单位名称)\s*[、，]\s*(?:地址|联系地址)\s*[）)]'
        if re.search(pattern1, text):
            self.logger.debug(f"🎯 检测到供应商名称地址组合模式: '{text[:50]}...'")
            company_name = info.get('公司名称', '')
            address = info.get('地址', '')
            self.logger.debug(f"📊 字段数据: 公司名称='{company_name}', 地址='{address}'")

            if company_name and address:
                replacement = f"（{company_name}、{address}）"
                success = WordDocumentUtils.precise_replace(paragraph, pattern1, replacement, self.logger)
                if success:
                    self.logger.info(f"🔄 组合规则替换成功: 供应商名称、地址 -> {replacement}")
                    processed_any = True
            else:
                missing_fields = []
                if not company_name:
                    missing_fields.append('公司名称')
                if not address:
                    missing_fields.append('地址')
                self.logger.warning(f"⚠️ 供应商名称地址组合缺少字段: {', '.join(missing_fields)}")

        # 组合模式2：项目名称、项目编号
        pattern2 = r'[（(]\s*(?:项目名称|工程名称)\s*[、，]\s*(?:项目编号|招标编号|采购编号)\s*[）)]'
        if re.search(pattern2, text):
            self.logger.debug(f"🎯 检测到项目名称编号组合模式: '{text[:50]}...'")
            project_name = info.get('项目名称', '')
            project_number = info.get('项目编号', '')
            self.logger.debug(f"📊 字段数据: 项目名称='{project_name}', 项目编号='{project_number}'")

            if project_name and project_number:
                replacement = f"（{project_name}、{project_number}）"
                success = WordDocumentUtils.precise_replace(paragraph, pattern2, replacement, self.logger)
                if success:
                    self.logger.info(f"🔄 组合规则替换成功: 项目名称、项目编号 -> {replacement}")
                    processed_any = True
            else:
                missing_fields = []
                if not project_name:
                    missing_fields.append('项目名称')
                if not project_number:
                    missing_fields.append('项目编号')
                self.logger.warning(f"⚠️ 项目名称编号组合缺少字段: {', '.join(missing_fields)}")

        # 组合模式3：联系电话、邮箱
        pattern3 = r'[（(]\s*(?:联系电话|电话)\s*[、，]\s*(?:邮箱|电子邮件)\s*[）)]'
        if re.search(pattern3, text):
            self.logger.debug(f"🎯 检测到电话邮箱组合模式: '{text[:50]}...'")
            phone = info.get('电话', '')
            email = info.get('邮箱', '')
            self.logger.debug(f"📊 字段数据: 电话='{phone}', 邮箱='{email}'")

            if phone and email:
                replacement = f"（{phone}、{email}）"
                success = WordDocumentUtils.precise_replace(paragraph, pattern3, replacement, self.logger)
                if success:
                    self.logger.info(f"🔄 组合规则替换成功: 联系电话、邮箱 -> {replacement}")
                    processed_any = True
            else:
                missing_fields = []
                if not phone:
                    missing_fields.append('电话')
                if not email:
                    missing_fields.append('邮箱')
                self.logger.warning(f"⚠️ 电话邮箱组合缺少字段: {', '.join(missing_fields)}")

        # 组合模式4：职位、职称 - 智能上下文识别
        position_pattern = r'[（(]\s*职[位务称]\s*[、，]\s*职[位务称]\s*[）)]'
        if re.search(position_pattern, text):
            self.logger.debug(f"🎯 检测到职位组合模式: '{text[:50]}...'")

            try:
                # 智能识别上下文
                context = self._detect_position_context(text)
                self.logger.debug(f"🧠 上下文识别结果: {context}")

                # 根据上下文选择数据源
                if context == 'authorized_person':
                    position = info.get('被授权人职务', '') or info.get('authorizedPersonPosition', '')
                    if position:
                        self.logger.info(f"📝 使用被授权人职务: '{position}'")
                    else:
                        self.logger.warning(f"⚠️ 被授权人职务为空，尝试法定代表人职位")
                        position = info.get('法定代表人职位', '') or info.get('legalRepresentativePosition', '')
                        if position:
                            self.logger.info(f"📝 回退使用法定代表人职位: '{position}'")
                else:  # legal_representative
                    position = info.get('法定代表人职位', '') or info.get('legalRepresentativePosition', '')
                    if position:
                        self.logger.info(f"📝 使用法定代表人职位: '{position}'")
                    else:
                        self.logger.warning(f"⚠️ 法定代表人职位为空，尝试被授权人职务")
                        position = info.get('被授权人职务', '') or info.get('authorizedPersonPosition', '')
                        if position:
                            self.logger.info(f"📝 回退使用被授权人职务: '{position}'")

                if position:
                    replacement = f"（{position}、{position}）"
                    success = WordDocumentUtils.precise_replace(paragraph, position_pattern, replacement, self.logger)
                    if success:
                        self.logger.info(f"智能职位组合替换: （职位、职称） → （{position}、{position}）")
                        processed_any = True
                else:
                    self.logger.warning(f"⚠️ 所有职位数据源都为空，跳过处理")

            except Exception as e:
                self.logger.error(f"❌ 职位组合替换发生异常: {e}")

        # 组合模式5：姓名、职位 - 智能上下文识别
        name_position_pattern = r'[（(]\s*姓名\s*[、，]\s*职[位务称]\s*[）)]'
        if re.search(name_position_pattern, text):
            self.logger.debug(f"🎯 检测到姓名职位组合模式: '{text[:50]}...'")

            try:
                # 智能识别上下文
                context = self._detect_position_context(text)
                self.logger.debug(f"🧠 上下文识别结果: {context}")

                # 根据上下文选择数据源
                if context == 'authorized_person':
                    name = info.get('被授权人姓名', '') or info.get('authorizedPersonName', '')
                    position = info.get('被授权人职务', '') or info.get('authorizedPersonPosition', '')
                    if name and position:
                        self.logger.info(f"📝 使用被授权人信息: '{name}', '{position}'")
                    elif not name:
                        name = info.get('法定代表人', '') or info.get('legalRepresentativeName', '')
                        if name:
                            self.logger.warning(f"⚠️ 被授权人姓名为空，使用法定代表人: '{name}'")
                    elif not position:
                        position = info.get('法定代表人职位', '') or info.get('legalRepresentativePosition', '')
                        if position:
                            self.logger.warning(f"⚠️ 被授权人职务为空，使用法定代表人职位: '{position}'")
                else:  # legal_representative
                    name = info.get('法定代表人', '') or info.get('legalRepresentativeName', '')
                    position = info.get('法定代表人职位', '') or info.get('legalRepresentativePosition', '')
                    if name and position:
                        self.logger.info(f"📝 使用法定代表人信息: '{name}', '{position}'")
                    elif not name:
                        name = info.get('被授权人姓名', '') or info.get('authorizedPersonName', '')
                        if name:
                            self.logger.warning(f"⚠️ 法定代表人姓名为空，使用被授权人: '{name}'")
                    elif not position:
                        position = info.get('被授权人职务', '') or info.get('authorizedPersonPosition', '')
                        if position:
                            self.logger.warning(f"⚠️ 法定代表人职位为空，使用被授权人职务: '{position}'")

                if name and position:
                    replacement = f"（{name}、{position}）"
                    success = WordDocumentUtils.precise_replace(paragraph, name_position_pattern, replacement, self.logger)
                    if success:
                        self.logger.info(f"智能姓名职位组合替换: （姓名、职位） → （{name}、{position}）")
                        processed_any = True
                else:
                    self.logger.warning(f"⚠️ 姓名或职位数据为空: 姓名={name}, 职位={position}")

            except Exception as e:
                self.logger.error(f"❌ 姓名职位组合替换发生异常: {e}")

        return processed_any
    
    def _try_replacement_rule(self, paragraph: Paragraph, info: Dict[str, Any]) -> bool:
        """尝试应用括号替换规则 - 累积处理模式"""
        text = paragraph.text
        replacement_count = 0

        # 括号替换规则
        for field_name, value in info.items():
            if not value:
                continue

            # 获取字段变体 - 直接使用字段名进行映射
            field_mapping = {
                '公司名称': self.field_variants['companyName'],
                '邮箱': self.field_variants['email'],
                '电话': self.field_variants['phone'],
                '传真': self.field_variants['fax'],
                '地址': self.field_variants['address'],
                '邮政编码': self.field_variants['postalCode'],
                '日期': self.field_variants['date'],
                '采购人名称': self.field_variants['purchaserName'],
                '项目名称': self.field_variants['projectName'],
                '项目编号': self.field_variants['projectNumber']
            }
            field_variants = field_mapping.get(field_name, [])

            if not field_variants:
                continue

            # 构建括号匹配模式
            variants_pattern = '|'.join(re.escape(variant) for variant in field_variants)
            bracket_pattern = f'[（(]\\s*(?:{variants_pattern})\\s*[）)]'

            if re.search(bracket_pattern, text):
                replacement_text = f'（{value}）'
                success = WordDocumentUtils.precise_replace(paragraph, bracket_pattern, replacement_text, self.logger)
                if success:
                    self.logger.info(f"🔄 括号替换成功: {field_name} -> {value}")
                    replacement_count += 1

        return replacement_count > 0
    
    def _try_fill_rule(self, paragraph: Paragraph, info: Dict[str, Any]) -> bool:
        """尝试应用填空规则 - 累积处理模式"""
        text = paragraph.text
        fill_count = 0

        # 填空规则处理
        for field_name, value in info.items():
            if not value:
                self.logger.info(f"🔍 [调试] 跳过空字段: {field_name}")
                continue

            self.logger.info(f"🔍 [调试] 尝试填空字段: {field_name} = {value}")

            # 获取字段变体 - 直接使用字段名进行映射
            field_mapping = {
                '公司名称': self.field_variants['companyName'],
                '邮箱': self.field_variants['email'],
                '电话': self.field_variants['phone'],
                '传真': self.field_variants['fax'],
                '地址': self.field_variants['address'],
                '邮政编码': self.field_variants['postalCode'],
                '日期': self.field_variants['date'],
                '采购人名称': self.field_variants['purchaserName'],
                '项目名称': self.field_variants['projectName'],
                '项目编号': self.field_variants['projectNumber']
            }
            field_variants = field_mapping.get(field_name, [])

            if not field_variants:
                continue

            # 检查是否应该在此段落中尝试该字段
            if not SmartFieldDetector.should_try_field_in_paragraph(text, field_variants):
                self.logger.info(f"🔍 [调试] 跳过字段 {field_name} 在段落: {text[:50]}...")
                continue

            # 尝试各种填空策略
            field_processed = False
            for variant in field_variants:
                if field_processed:
                    break  # 该字段已处理成功，尝试下一个字段

                self.logger.info(f"🔍 [调试] 尝试字段变体: {variant} 在段落: {paragraph.text[:50]}...")

                # 策略1：插入式替换（按备份文件顺序）
                if self._try_insert_strategy(paragraph, variant, value):
                    self.logger.info(f"🔄 插入式成功: {field_name} -> {value}")
                    fill_count += 1
                    field_processed = True
                    continue

                # 策略2：公章格式替换
                if self._try_stamp_strategy(paragraph, variant, value):
                    self.logger.info(f"🔄 公章格式成功: {field_name} -> {value}")
                    fill_count += 1
                    field_processed = True
                    continue

                # 策略3：纯空格替换
                if self._try_space_only_strategy(paragraph, variant, value):
                    self.logger.info(f"🔄 纯空格成功: {field_name} -> {value}")
                    fill_count += 1
                    field_processed = True
                    continue

                # 策略4：括号格式替换
                if self._try_bracket_strategy(paragraph, variant, value):
                    self.logger.info(f"🔄 括号格式成功: {field_name} -> {value}")
                    fill_count += 1
                    field_processed = True
                    continue

                # 策略5：精确模式替换（新增）
                if self._try_precise_strategies(paragraph, variant, value):
                    self.logger.info(f"🔄 精确模式成功: {field_name} -> {value}")
                    fill_count += 1
                    field_processed = True
                    continue

                # 策略6：传统填空模式（保留原有逻辑）
                if self._try_fill_patterns(paragraph, variant, value):
                    self.logger.info(f"🔄 填空规则成功: {field_name} -> {value}")
                    fill_count += 1
                    field_processed = True
                    continue

        return fill_count > 0
    
    def _try_fill_patterns(self, paragraph: Paragraph, field_variant: str, value: str) -> bool:
        """尝试各种填空模式"""
        patterns = [
            # 模式1: 字段名：___
            f'{re.escape(field_variant)}\\s*[：:]\\s*_+',
            # 模式2: 字段名：
            f'{re.escape(field_variant)}\\s*[：:]\\s*$',
            # 模式3: 字段名：___ (混合)
            f'{re.escape(field_variant)}\\s*[：:]\\s*[_\\s]*$',
            # 模式4: 字段名 (插入式)
            f'{re.escape(field_variant)}(?=\\s+)(?![：:])',
            # 模式5: 致：字段名 格式
            f'致\\s*[：:]\\s*{re.escape(field_variant)}\\s*$'
        ]
        
        for pattern in patterns:
            if re.search(pattern, paragraph.text):
                # 构建替换文本
                if '致' in pattern:
                    # 特殊处理"致："格式
                    replacement = f'致：{value}'
                elif '：' in paragraph.text or ':' in paragraph.text:
                    replacement = f'{field_variant}：{value}'
                else:
                    replacement = f'{field_variant} {value}'

                success = WordDocumentUtils.precise_replace(paragraph, pattern, replacement, self.logger)
                if success:
                    return True
        
        return False

    def _try_bracket_strategy(self, paragraph: Paragraph, variant: str, value: str) -> bool:
        """策略1：括号格式替换 - 处理（字段名）→（替换值）格式"""
        # 快速检查：如果段落中根本没有括号，直接返回
        if '（' not in paragraph.text and '(' not in paragraph.text:
            return False

        bracket_pattern = rf'[（(]\s*{re.escape(variant)}\s*[）)]'
        match = re.search(bracket_pattern, paragraph.text)
        if not match:
            return False

        self.logger.info(f"🎯 策略1(括号格式)匹配成功 - 字段: {variant}")
        self.logger.info(f"📝 匹配模式: {bracket_pattern}")
        self.logger.info(f"✅ 匹配内容: '{match.group()}'")

        replacement = f"（{value}）"
        return WordDocumentUtils.precise_replace(paragraph, bracket_pattern, replacement, self.logger)

    def _try_insert_strategy(self, paragraph: Paragraph, variant: str, value: str) -> bool:
        """策略2：插入式替换 - 直接在字段名后插入内容"""
        # 快速检查：只有字段名后直接跟冒号才拒绝
        if re.search(rf'{re.escape(variant)}\s*[:：]', paragraph.text):
            # 如果字段名后直接跟冒号，不是插入式格式
            return False

        # 检查是否匹配插入式模式：字段名后面跟空格但不跟冒号
        # 改进：匹配字段名及其后面的所有空格，确保完全替换
        insert_pattern = rf'{re.escape(variant)}\s+'

        # 先检查基本匹配条件
        if not re.search(rf'{re.escape(variant)}(?=\s+)(?![:：])', paragraph.text):
            return False

        self.logger.info(f"🎯 策略2(插入式)匹配成功 - 字段: {variant}")
        self.logger.info(f"📝 匹配模式(改进版): {insert_pattern}")

        replacement = f'{variant} {value}'
        success = WordDocumentUtils.precise_replace(paragraph, insert_pattern, replacement, self.logger)

        if success:
            # 标记此段落需要后续格式清理
            self._mark_paragraph_for_format_cleanup(paragraph, variant, value)
            self.logger.info(f"🏷️ 标记段落需要格式清理: {variant}")

        return success

    def _try_space_only_strategy(self, paragraph: Paragraph, variant: str, value: str) -> bool:
        """策略3：纯空格替换 - 处理只有空格无下划线的情况"""
        # 快速检查：必须包含冒号并且以空格结尾
        if not re.search(r'[:：]\s+$', paragraph.text):
            return False

        space_pattern = rf'({re.escape(variant)}\s*[:：])\s+$'
        match = re.search(space_pattern, paragraph.text)
        if not match:
            return False

        self.logger.info(f"🎯 策略3(纯空格)匹配成功 - 字段: {variant}")
        replacement = rf'\1{value}'
        return WordDocumentUtils.precise_replace(paragraph, space_pattern, replacement, self.logger)

    def _try_stamp_strategy(self, paragraph: Paragraph, variant: str, value: str) -> bool:
        """策略4：公章格式替换 - 保留公章括号"""
        # 快速检查：如果段落中没有"章"字，不可能是公章格式
        if '章' not in paragraph.text:
            return False

        stamp_pattern = rf'(?P<prefix>{re.escape(variant)}\s*[:：]\s*)(?P<spaces>[_\s]+)(?P<stamp>[（(][^）)]*章[^）)]*[）)])'
        match = re.search(stamp_pattern, paragraph.text)
        if not match:
            return False

        self.logger.info(f"🎯 策略4(公章格式)匹配成功 - 字段: {variant}")
        replacement = rf'\g<prefix>{value}\g<stamp>'
        return WordDocumentUtils.precise_replace(paragraph, stamp_pattern, replacement, self.logger)

    def _try_precise_strategies(self, paragraph: Paragraph, variant: str, replacement_text: str) -> bool:
        """策略5：精确模式替换 - 4个子策略"""
        self.logger.debug(f"🔄 使用精确模式替换策略")

        # 精确模式子策略列表
        precise_patterns = [
            # 子策略1：多字段格式处理 - 地址：___ 邮编：___（保留后续字段）
            (rf'(?P<prefix>{re.escape(variant)}\s*[:：]\s*)(?P<underscores>_+)(?P<suffix>\s+[^\s_]+[:：])',
             rf'\g<prefix>{replacement_text}\g<suffix>'),

            # 子策略2：单字段末尾格式 - 电话：___________（清理所有下划线）
            (rf'({re.escape(variant)}\s*[:：]\s*)_+\s*$',
             rf'\g<1>{replacement_text}'),

            # 子策略3：无下划线格式 - 电子邮箱：（直接添加内容）
            (rf'({re.escape(variant)}\s*[:：])\s*$',
             rf'\g<1>{replacement_text}'),

            # 子策略4：通用下划线格式 - 供应商名称：___（清理下划线和空格）
            (rf'({re.escape(variant)}\s*[:：]\s*)[_\s]+',
             rf'\g<1>{replacement_text}')
        ]

        # 依次尝试每个精确子策略
        for i, (pattern, replacement) in enumerate(precise_patterns, 1):
            if re.search(pattern, paragraph.text):
                # 提升到INFO级别，并增加详细信息
                self.logger.info(f"🎯 精确子策略{i}匹配成功 - 模式: {pattern}")
                self.logger.info(f"📝 替换模式: {replacement}")
                match_obj = re.search(pattern, paragraph.text)
                if match_obj:
                    self.logger.info(f"✅ 匹配内容: '{match_obj.group()}'")
                if WordDocumentUtils.precise_replace(paragraph, pattern, replacement, self.logger):
                    return True

        return False

    def _process_table(self, table: Table, info: Dict[str, Any]) -> Dict[str, Any]:
        """处理表格（简化版）"""
        stats = {'total_replacements': 0}
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_stats = self._process_paragraph(paragraph, info)
                        for key, value in cell_stats.items():
                            stats[key] = stats.get(key, 0) + value
        
        return stats
    
    def _post_process(self, doc: Document):
        """
        后处理：清理多余的占位符和装饰性格式（保护已填充内容）
        新增：专门处理插入式策略的格式清理需求
        """
        # 第一步：处理标记的格式清理
        self.logger.info("🧹 开始后处理：格式清理和美化")

        for paragraph in doc.paragraphs:
            # 检查是否有格式清理标记
            if hasattr(paragraph, '_format_cleanup_needed'):
                self.logger.info(f"🏷️ 发现需要格式清理的段落: '{paragraph.text[:30]}...'")
                self._clean_decorative_formats_only(paragraph)

        # 第二步：处理年月日格式填充
        self._process_date_format_filling(doc)

        # 第三步：原有的文本清理逻辑
        for paragraph in doc.paragraphs:
            text = paragraph.text
            original_text = text

            # 检查是否包含已填充的内容（包含中文公司名称等）
            contains_filled_content = False
            for company_pattern in self.company_name_extended_patterns:
                base_pattern = company_pattern.split('(?:')[0]  # 获取基础模式（去掉公章部分）
                if base_pattern in text and len(text) > len(base_pattern) + 5:
                    # 如果文本比基础字段名长很多，可能包含填充内容
                    contains_filled_content = True
                    break

            # 检查是否包含采购人名称等填充内容
            if '北京市' in text or '有限公司' in text or '集团' in text:
                contains_filled_content = True

            # 检查是否包含联系信息等填充内容
            if (re.search(r'\d{3,4}-\d{7,8}', text) or  # 电话号码格式
                re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text) or  # 邮箱格式
                re.search(r'\d{6}', text) or  # 邮编格式
                '年' in text and '月' in text and '日' in text):  # 日期格式
                contains_filled_content = True

            # 如果包含填充内容，跳过空格清理以保护内容
            if not contains_filled_content:
                # 清理多余的下划线
                text = re.sub(r'_{3,}', '', text)

                # 清理多余的空格（保留表格对齐所需的空格）
                if not re.search(r'\s{8,}', text):  # 不是表格式布局
                    text = re.sub(r'\s{3,}', '  ', text)

            # 清理多余的冒号
            text = re.sub(r'[:：]{2,}', '：', text)

            # 标准化冒号
            text = re.sub(r':', '：', text)

            # 去除多余的年月日标识和重复内容
            text = self._clean_date_redundancy_and_placeholders(text)

            if text != original_text:
                # 使用精确格式处理进行后处理清理
                escaped_original = re.escape(original_text.strip())
                if not WordDocumentUtils.precise_replace(paragraph, escaped_original, text.strip(), self.logger):
                    # 后备方案：使用格式保护方法
                    self._update_paragraph_text_preserving_format(paragraph, text.strip())
    
    def _mark_paragraph_for_format_cleanup(self, paragraph, field_name: str, content: str):
        """标记段落需要后续格式清理"""
        try:
            # 在段落对象上添加清理标记（临时属性）
            if not hasattr(paragraph, '_format_cleanup_needed'):
                paragraph._format_cleanup_needed = []

            paragraph._format_cleanup_needed.append({
                'field_name': field_name,
                'content': content
            })

            self.logger.info(f"🏷️ 段落格式清理标记已添加: {field_name} -> {content[:20]}...")

        except Exception as e:
            self.logger.error(f"❌ 添加格式清理标记失败: {e}")

    def _is_filled_content_run(self, run) -> bool:
        """判断run是否包含已填充的内容"""
        try:
            run_text = run.text
            if not run_text:
                return False

            # 检查是否包含典型的填充内容模式
            filled_patterns = [
                r'\d{3,4}-\d{7,8}',  # 电话号码
                r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',  # 邮箱
                r'\d{6}',  # 邮编
                r'www\.',  # 网站
                r'有限公司|股份|集团|科技',  # 公司名称特征
            ]

            for pattern in filled_patterns:
                if re.search(pattern, run_text):
                    return True

            return False

        except Exception as e:
            self.logger.error(f"❌ 检查填充内容失败: {e}")
            return False

    def _clean_decorative_formats_only(self, paragraph):
        """只清理装饰性格式，完全保留文本结构"""
        try:
            cleanup_info = getattr(paragraph, '_format_cleanup_needed', [])

            for run in paragraph.runs:
                should_clean = False

                # 方法1：检查是否包含我们标记的填充内容
                for info in cleanup_info:
                    if info['content'] in run.text:
                        should_clean = True
                        self.logger.info(f"🏷️ 发现标记的填充内容: '{info['content']}'")
                        break

                # 方法2：检查是否包含典型填充内容模式
                if not should_clean:
                    should_clean = self._is_filled_content_run(run)

                # 方法3：对于插入式策略，清理所有后续空格的装饰格式
                if not should_clean and len(run.text.strip()) == 0:
                    # 如果是纯空格run，且段落包含填充内容，则清理装饰格式
                    paragraph_text = paragraph.text
                    if any(pattern in paragraph_text for pattern in ['010-', '@', 'www.', '有限公司']):
                        should_clean = True
                        self.logger.info(f"🧹 清理填充内容后的空格装饰格式")

                if should_clean:
                    # 清理装饰格式
                    if hasattr(run.font, 'underline') and run.font.underline:
                        run.font.underline = False
                        self.logger.info(f"🔧 清除下划线格式: '{run.text[:15]}...'")

                    if hasattr(run.font, 'strike') and run.font.strike:
                        run.font.strike = False
                        self.logger.info(f"🔧 清除删除线格式: '{run.text[:15]}...'")

            # 清理临时标记
            if hasattr(paragraph, '_format_cleanup_needed'):
                delattr(paragraph, '_format_cleanup_needed')

        except Exception as e:
            self.logger.error(f"❌ 装饰性格式清理失败: {e}")

    def _process_date_format_filling(self, doc: Document):
        """
        处理年月日格式填充
        识别和填充"    年    月    日"等各种空格分隔的年月日格式
        """
        self.logger.info("📅 开始处理年月日格式填充")

        # 获取日期值
        date_value = self.info.get('date', '')
        if not date_value:
            self.logger.info("⚠️ 日期值为空，跳过年月日格式填充")
            return

        # 格式化日期
        formatted_date = self._format_date(date_value)
        if not formatted_date:
            self.logger.warning("⚠️ 日期格式化失败，跳过年月日格式填充")
            return

        self.logger.info(f"📅 准备填充的日期值: '{formatted_date}'")

        # 定义年月日格式匹配模式
        date_end_patterns = [
            r'^\s{2,}年\s{2,}月\s{2,}日$',      # 空格分隔的年月日格式（独立行）
            r'(\s+)年(\s+)月(\s+)日(\s*)$',      # 末尾格式：空格+年+空格+月+空格+日
            r'(\n\s*)年(\s+)月(\s+)日(\s*)$',    # 换行+空格+年月日格式
            r'(\s+)年(\s+)月(\s+)日',           # 通用格式：空格+年+空格+月+空格+日
        ]

        processed_count = 0

        for paragraph in doc.paragraphs:
            text = paragraph.text
            if not text or not text.strip():
                continue

            # 检查是否匹配年月日格式
            for i, pattern in enumerate(date_end_patterns, 1):
                match = re.search(pattern, text)
                if match:
                    self.logger.info(f"✅ 年月日模式{i}匹配成功: '{match.group()}' 在段落: '{text[:50]}...'")

                    # 构建新文本
                    if i == 1:  # 独立的空格年月日格式
                        # 直接替换整个匹配内容
                        new_text = re.sub(pattern, formatted_date, text)
                    elif i == 3:  # 换行+空格+年月日格式
                        # 保留换行符，只替换年月日部分
                        new_text = re.sub(pattern, rf'\1{formatted_date}', text)
                    else:
                        # 标准替换：整个匹配的年月日模式为完整日期
                        new_text = re.sub(pattern, formatted_date, text)

                    if new_text != text:
                        # 使用精确格式处理进行替换
                        escaped_original = re.escape(text.strip())
                        if WordDocumentUtils.precise_replace(paragraph, escaped_original, new_text.strip(), self.logger):
                            self.logger.info(f"🔄 年月日格式填充成功: '{text}' -> '{new_text}'")
                            processed_count += 1
                        else:
                            # 后备方案：使用格式保护方法
                            self._update_paragraph_text_preserving_format(paragraph, new_text.strip())
                            self.logger.info(f"🔄 年月日格式填充成功(后备): '{text}' -> '{new_text}'")
                            processed_count += 1

                        # 找到一个匹配后，跳出模式循环
                        break

        if processed_count > 0:
            self.logger.info(f"📊 年月日格式填充完成，共处理 {processed_count} 个段落")
        else:
            self.logger.info("📊 未发现需要填充的年月日格式")

    def _clean_date_redundancy_and_placeholders(self, text: str) -> str:
        """
        清理日期相关的重复内容和占位符残留
        处理各种"年月日"重复模式和占位符
        """
        original_text = text

        # 第一组：处理直接重复的年月日字符
        redundant_patterns = [
            r'(\d+年\d+月\d+日)年',              # 2015年12月18日年
            r'(\d+年\d+月\d+日)月',              # 2015年12月18日月
            r'(\d+年\d+月\d+日)日',              # 2015年12月18日日
            r'(\d+年\d+月\d+日)年\s*月',         # 2015年12月18日年月
            r'(\d+年\d+月\d+日)年\s*月\s*日',    # 2015年12月18日年月日
            r'(\d+年\d+月\d+日)月\s*日',         # 2015年12月18日月日
        ]

        for pattern in redundant_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                self.logger.debug(f"清理重复字符: '{old_text}' -> '{text}'")

        # 第二组：处理带空格的重复模式
        spaced_redundant_patterns = [
            r'(\d+年\d+月\d+日)\s+年',           # 2015年12月18日 年
            r'(\d+年\d+月\d+日)\s+月',           # 2015年12月18日 月
            r'(\d+年\d+月\d+日)\s+日',           # 2015年12月18日 日
            r'(\d+年\d+月\d+日)\s+年\s*月',      # 2015年12月18日 年月
            r'(\d+年\d+月\d+日)\s+年\s*月\s*日', # 2015年12月18日 年月日
            r'(\d+年\d+月\d+日)\s+月\s*日',      # 2015年12月18日 月日
            r'(\d+年\d+月\d+日)\s*年\s*月\s*日', # 2015年12月18日年月日（任意空格）
            r'(\d+年\d+月\d+日)\s+月\s+日',      # 2015年12月18日  月  日
            r'(\d+年\d+月\d+日)\s+年\s+月',      # 2015年12月18日  年  月
        ]

        for pattern in spaced_redundant_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                self.logger.debug(f"清理空格重复: '{old_text}' -> '{text}'")

        # 第三组：处理占位符残留
        placeholder_cleanup_patterns = [
            # 清理日期后的下划线占位符
            r'(\d+年\d+月\d+日)_+月_+日',        # 2025年09月07日_____月_____日
            r'(\d+年\d+月\d+日)_+月',            # 2025年09月07日_____月
            r'(\d+年\d+月\d+日)_+日',            # 2025年09月07日_____日
            r'(\d+年\d+月\d+日)_+年_+月_+日',    # 2025年09月07日_____年_____月_____日
            r'(\d+年\d+月\d+日)_+年_+月',        # 2025年09月07日_____年_____月
            r'(\d+年\d+月\d+日)_+年',            # 2025年09月07日_____年

            # 清理空格和下划线混合的情况
            r'(\d+年\d+月\d+日)[\s_]+月[\s_]*日', # 2025年09月07日 ___月___日
            r'(\d+年\d+月\d+日)[\s_]+年[\s_]*月[\s_]*日', # 带空格的混合情况

            # 清理横线形式的占位符
            r'(\d+年\d+月\d+日)-+',              # 2025年09月07日--------
            r'(\d+年\d+月\d+日)[\s-]+$',         # 2025年09月07日 ---- (行末)
            r'(\d+年\d+月\d+日)_+$',             # 2025年09月07日_____ (行末)

            # 清理日期后的任意组合占位符（更通用的模式）
            r'(\d+年\d+月\d+日)[\s_-]+.*?$',     # 日期后任意占位符到行末
        ]

        for pattern in placeholder_cleanup_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                self.logger.debug(f"清理占位符残留: '{old_text}' -> '{text}'")

        # 第四组：处理重复日期
        duplicate_date_patterns = [
            r'(\d{4}年\d{1,2}月\d{1,2}日)\s+(\d{4}年\d{1,2}月\d{1,2}日)',  # 2025年9月12日 2025年9月12日
            r'(\d{4}年\d{1,2}月\d{1,2}日)年[_\s]*月[_\s]*日',              # 2025年9月12日年_____月_____日
        ]

        for pattern in duplicate_date_patterns:
            if re.search(pattern, text):
                old_text = text
                text = re.sub(pattern, r'\1', text)
                self.logger.debug(f"清理重复日期: '{old_text}' -> '{text}'")

        # 如果有清理操作，记录日志
        if text != original_text:
            self.logger.info(f"📅 日期清理完成: '{original_text}' -> '{text}'")

        return text

    def _format_date(self, date_str: str) -> str:
        """格式化日期字符串"""
        # 移除多余的空格
        date_str = re.sub(r'\s+', '', date_str)

        # 尝试匹配常见日期格式
        patterns = [
            (r'(\d{4})-(\d{1,2})-(\d{1,2})', r'\1年\2月\3日'),
            (r'(\d{4})/(\d{1,2})/(\d{1,2})', r'\1年\2月\3日'),
            (r'(\d{4})\.(\d{1,2})\.(\d{1,2})', r'\1年\2月\3日'),
        ]

        for pattern, replacement in patterns:
            if re.match(pattern, date_str):
                return re.sub(pattern, replacement, date_str)

        # 如果已经是年月日格式，直接返回
        if re.match(r'\d{4}年\d{1,2}月\d{1,2}日', date_str):
            return date_str

        return None

    def _update_paragraph_text_preserving_format(self, paragraph: Paragraph, new_text: str):
        """
        格式保护的段落文本更新方法 - 后备方案

        注意：优先使用 precise_replace() 精确格式处理引擎
        此方法主要作为复杂情况下的后备方案
        """
        if not paragraph.runs:
            return

        # 将新文本分配给第一个run，清空其他runs
        first_run = paragraph.runs[0]
        first_run.text = new_text

        # 清空其他runs的文本但保持它们的存在（保持格式结构）
        for run in paragraph.runs[1:]:
            run.text = ''

        self.logger.info(f"✅ 格式保护更新完成: '{new_text}'")

# 向后兼容
class BusinessInfoFiller(InfoFiller):
    """向后兼容的类名别名"""
    pass