#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ä¿¡æ¯å¡«å†™æ¨¡å— - å¤„ç†é¡¹ç›®å’Œå…¬å¸ä¿¡æ¯çš„å¡«å†™
å®ç°å…­å¤§è§„åˆ™ï¼šæ›¿æ¢è§„åˆ™ã€å¡«ç©ºè§„åˆ™ã€ç»„åˆè§„åˆ™ã€å˜ä½“å¤„ç†ã€ä¾‹å¤–å¤„ç†ã€åå¤„ç†

 1. ä¾›åº”å•†åç§°å¤„ç†ï¼ˆæ”¯æŒ4ç§è§„åˆ™ï¼‰
      - æ›¿æ¢è§„åˆ™ï¼šï¼ˆä¾›åº”å•†åç§°ï¼‰â†’ï¼ˆå…¬å¸åï¼‰
      - å¡«ç©ºè§„åˆ™ï¼šä¾›åº”å•†åç§°ï¼š___ â†’ 
    ä¾›åº”å•†åç§°ï¼šå…¬å¸å
      - ç»„åˆè§„åˆ™ï¼šï¼ˆä¾›åº”å•†åç§°ã€åœ°å€ï¼‰â†’ï¼ˆå…¬å¸
    åã€åœ°å€ï¼‰
      - å˜ä½“å¤„ç†ï¼šå…¬å¸åç§°ã€åº”ç­”äººåç§°ã€ä¾›åº”
    å•†åç§°ï¼ˆç›–ç« ï¼‰ç­‰
    2. å…¶ä»–ä¿¡æ¯å­—æ®µï¼ˆä»…å¡«ç©ºè§„åˆ™ï¼‰
      - ç”µè¯ã€é‚®ç®±ã€åœ°å€ã€é‚®ç¼–ã€ä¼ çœŸ
      - æ”¯æŒæ ‡ç­¾å˜ä½“ï¼ˆé‚®ç®±/ç”µå­é‚®ä»¶ï¼‰
      - æ”¯æŒæ ¼å¼å˜åŒ–ï¼ˆå†’å·ã€ç©ºæ ¼ã€å ä½ç¬¦ï¼‰
    3. ä¾‹å¤–å¤„ç†
      - è·³è¿‡"ç­¾å­—"ç›¸å…³å­—æ®µ
      - æ™ºèƒ½æ—¥æœŸå¤„ç†
      - è¯†åˆ«å¹¶è·³è¿‡é‡‡è´­äºº/æ‹›æ ‡äººä¿¡æ¯
    4. æ ¼å¼ä¿æŒ
      - ç»§æ‰¿ç¬¬ä¸€ä¸ªå­—ç¬¦çš„æ ¼å¼
      - ä¿æŒåŸæœ‰æ–‡æ¡£æ ·å¼

"""

import re
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

# å¯¼å…¥å…¬å…±æ¨¡å—
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import get_module_logger

class InfoFiller:
    """ä¿¡æ¯å¡«å†™å¤„ç†å™¨"""
    
    def __init__(self):
        self.logger = get_module_logger("info_filler")
        
        # ä¾›åº”å•†åç§°çš„å˜ä½“
        self.company_name_variants = [
            'ä¾›åº”å•†åç§°', 'ä¾›åº”å•†å…¨ç§°', 'æŠ•æ ‡äººåç§°', 'å…¬å¸åç§°',
            'å•ä½åç§°', 'åº”ç­”äººåç§°', 'ä¾›åº”å•†åç§°ï¼ˆç›–ç« ï¼‰',
            'ä¾›åº”å•†åç§°ï¼ˆå…¬ç« ï¼‰', 'å…¬å¸åç§°ï¼ˆç›–ç« ï¼‰', 'æŠ•æ ‡äººï¼ˆç›–ç« ï¼‰',
            'å“åº”äººåç§°', 'å“åº”äººåç§°ï¼ˆç›–ç« ï¼‰', 'å“åº”äººåç§°ï¼ˆå…¬ç« ï¼‰',  # æ–°å¢ï¼šå“åº”äººåç§°å˜ä½“
            'å“åº”äººå…¨ç§°'  # æ–°å¢ï¼šå“åº”äººå…¨ç§°
        ]
        
        # å…¶ä»–å­—æ®µçš„å˜ä½“æ˜ å°„
        self.field_variants = {
            'email': ['ç”µå­é‚®ä»¶', 'ç”µå­é‚®ç®±', 'é‚®ç®±', 'email', 'Email', 'E-mail'],
            'phone': ['ç”µè¯', 'è”ç³»ç”µè¯', 'å›ºå®šç”µè¯', 'ç”µè¯å·ç ', 'è”ç³»æ–¹å¼'],
            'fax': ['ä¼ çœŸ', 'ä¼ çœŸå·ç ', 'ä¼ çœŸå·', 'fax', 'Fax'],
            'address': ['åœ°å€', 'æ³¨å†Œåœ°å€', 'åŠå…¬åœ°å€', 'è”ç³»åœ°å€', 'é€šè®¯åœ°å€', 'ä¾›åº”å•†åœ°å€', 'å…¬å¸åœ°å€'],
            'postalCode': ['é‚®æ”¿ç¼–ç ', 'é‚®ç¼–', 'é‚®ç '],
            'legalRepresentative': ['æ³•å®šä»£è¡¨äºº', 'æ³•äººä»£è¡¨', 'æ³•äºº'],
            'projectName': ['é¡¹ç›®åç§°', 'é‡‡è´­é¡¹ç›®åç§°', 'æ‹›æ ‡é¡¹ç›®åç§°'],
            'projectNumber': ['é¡¹ç›®ç¼–å·', 'é‡‡è´­ç¼–å·', 'æ‹›æ ‡ç¼–å·', 'é¡¹ç›®å·']
        }
        
        # éœ€è¦è·³è¿‡çš„å…³é”®è¯ï¼ˆé‡‡è´­äºº/æ‹›æ ‡äººä¿¡æ¯ï¼‰
        self.skip_keywords = [
            'é‡‡è´­äºº', 'æ‹›æ ‡äºº', 'ç”²æ–¹', 'ä»£ç†', 'æ‹›æ ‡ä»£ç†',
            'é‡‡è´­ä»£ç†', 'ä¸šä¸»', 'å‘åŒ…äºº', 'å§”æ‰˜äºº'
        ]
        
        # éœ€è¦è·³è¿‡çš„ç­¾å­—ç›¸å…³è¯
        self.signature_keywords = ['ç­¾å­—', 'ç­¾å', 'ç­¾ç« ', 'ç›–ç« å¤„']
        
    def fill_info(self, doc: Document, company_info: Dict[str, Any], 
                  project_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        å¡«å†™ä¿¡æ¯ä¸»æ–¹æ³•
        
        Args:
            doc: Wordæ–‡æ¡£å¯¹è±¡
            company_info: å…¬å¸ä¿¡æ¯å­—å…¸
            project_info: é¡¹ç›®ä¿¡æ¯å­—å…¸ï¼ˆåŒ…å«é¡¹ç›®åç§°ã€ç¼–å·ã€æ—¥æœŸç­‰ï¼‰
            
        Returns:
            å¤„ç†ç»Ÿè®¡ä¿¡æ¯
        """
        stats = {
            'total_replacements': 0,
            'replacement_rules': 0,
            'fill_rules': 0,
            'combination_rules': 0,
            'skipped_fields': 0,
            'none': 0  # æ·»åŠ å¯¹æœªå¤„ç†æ®µè½çš„ç»Ÿè®¡
        }
        
        # åˆå¹¶æ‰€æœ‰ä¿¡æ¯
        all_info = {**company_info, **project_info}
        
        # å¤„ç†æ‰€æœ‰æ®µè½
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                result = self._process_paragraph(paragraph, all_info)
                stats['total_replacements'] += result['count']
                
                # å®‰å…¨åœ°æ›´æ–°ç»Ÿè®¡ä¿¡æ¯
                result_type = result['type']
                if result_type in stats:
                    if result_type == 'skipped':
                        stats[result_type] += 1
                    else:
                        stats[result_type] += result['count']
                else:
                    # å¦‚æœç±»å‹ä¸å­˜åœ¨ï¼Œè®°å½•ä¸º'none'
                    stats['none'] += result['count']
        
        # å¤„ç†è¡¨æ ¼
        for table in doc.tables:
            result = self._process_table(table, all_info)
            stats['total_replacements'] += result['count']
            stats['fill_rules'] += result['count']
        
        # åå¤„ç†ï¼šæ¸…ç†å¤šä½™çš„å ä½ç¬¦
        self._post_process(doc)
        
        self.logger.info(f"ä¿¡æ¯å¡«å†™å®Œæˆ: {stats}")
        return stats
    
    def _process_paragraph(self, paragraph: Paragraph, info: Dict[str, Any]) -> Dict[str, Any]:
        """å¤„ç†å•ä¸ªæ®µè½"""
        result = {'count': 0, 'type': 'none'}
        para_text = paragraph.text.strip()
        
        # æ£€æŸ¥æ˜¯å¦éœ€è¦è·³è¿‡
        if self._should_skip(para_text):
            self.logger.debug(f"è·³è¿‡æ®µè½: {para_text[:50]}")
            return {'count': 0, 'type': 'skipped'}
        
        # 1. å°è¯•ç»„åˆæ›¿æ¢è§„åˆ™
        if self._try_combination_rule(paragraph, info):
            return {'count': 1, 'type': 'combination_rules'}
        
        # 2. å°è¯•å•å­—æ®µæ›¿æ¢è§„åˆ™
        if self._try_replacement_rule(paragraph, info):
            return {'count': 1, 'type': 'replacement_rules'}
        
        # 3. å°è¯•å¡«ç©ºè§„åˆ™
        if self._try_fill_rule(paragraph, info):
            return {'count': 1, 'type': 'fill_rules'}
        
        return result
    
    def _should_skip(self, text: str) -> bool:
        """æ£€æŸ¥æ˜¯å¦åº”è¯¥è·³è¿‡è¯¥æ–‡æœ¬"""
        # æ£€æŸ¥æ˜¯å¦åŒ…å«é‡‡è´­äºº/æ‹›æ ‡äººç­‰å…³é”®è¯
        for keyword in self.skip_keywords:
            if keyword in text:
                return True
        
        # æ£€æŸ¥æ˜¯å¦åŒ…å«ç­¾å­—ç›¸å…³è¯
        for keyword in self.signature_keywords:
            if keyword in text:
                return True
        
        return False
    
    def _try_combination_rule(self, paragraph: Paragraph, info: Dict[str, Any]) -> bool:
        """
        å°è¯•ç»„åˆæ›¿æ¢è§„åˆ™
        å¦‚ï¼šï¼ˆä¾›åº”å•†åç§°ã€åœ°å€ï¼‰â†’ï¼ˆå…¬å¸åã€åœ°å€ï¼‰
        """
        text = paragraph.text
        
        # ç»„åˆæ¨¡å¼1ï¼šä¾›åº”å•†åç§°ã€åœ°å€
        pattern1 = r'[ï¼ˆ(]\s*ä¾›åº”å•†åç§°\s*[ã€ï¼Œ]\s*åœ°å€\s*[ï¼‰)]'
        if re.search(pattern1, text):
            company_name = info.get('companyName', '')
            address = info.get('address', '') or info.get('registeredAddress', '')
            if company_name and address:
                replacement = f"ï¼ˆ{company_name}ã€{address}ï¼‰"
                new_text = re.sub(pattern1, replacement, text)
                self._update_paragraph_text(paragraph, new_text)
                self.logger.info(f"ç»„åˆæ›¿æ¢: ä¾›åº”å•†åç§°ã€åœ°å€")
                return True
        
        # ç»„åˆæ¨¡å¼2ï¼šé¡¹ç›®åç§°ã€é¡¹ç›®ç¼–å·
        pattern2 = r'[ï¼ˆ(]\s*é¡¹ç›®åç§°\s*[ã€ï¼Œ]\s*é¡¹ç›®ç¼–å·\s*[ï¼‰)]'
        if re.search(pattern2, text):
            project_name = info.get('projectName', '')
            project_number = info.get('projectNumber', '')
            if project_name and project_number:
                replacement = f"ï¼ˆ{project_name}ã€{project_number}ï¼‰"
                new_text = re.sub(pattern2, replacement, text)
                self._update_paragraph_text(paragraph, new_text)
                self.logger.info(f"ç»„åˆæ›¿æ¢: é¡¹ç›®åç§°ã€é¡¹ç›®ç¼–å·")
                return True
        
        return False
    
    def _try_replacement_rule(self, paragraph: Paragraph, info: Dict[str, Any]) -> bool:
        """
        å°è¯•å•å­—æ®µæ›¿æ¢è§„åˆ™
        å¦‚ï¼šï¼ˆä¾›åº”å•†åç§°ï¼‰â†’ï¼ˆå…¬å¸åï¼‰
        """
        text = paragraph.text
        
        # å¤„ç†ä¾›åº”å•†åç§°ç±»
        for variant in self.company_name_variants:
            pattern = rf'[ï¼ˆ(]\s*{re.escape(variant)}\s*[ï¼‰)]'
            if re.search(pattern, text):
                company_name = info.get('companyName', '')
                if company_name:
                    replacement = f"ï¼ˆ{company_name}ï¼‰"
                    new_text = re.sub(pattern, replacement, text)
                    self._update_paragraph_text(paragraph, new_text)
                    self.logger.info(f"æ›¿æ¢è§„åˆ™: {variant} â†’ {company_name}")
                    return True
        
        # å¤„ç†å…¶ä»–å­—æ®µ
        for field_key, variants in self.field_variants.items():
            for variant in variants:
                pattern = rf'[ï¼ˆ(]\s*{re.escape(variant)}\s*[ï¼‰)]'
                if re.search(pattern, text):
                    value = info.get(field_key, '')
                    if value:
                        replacement = f"ï¼ˆ{value}ï¼‰"
                        new_text = re.sub(pattern, replacement, text)
                        self._update_paragraph_text(paragraph, new_text)
                        self.logger.info(f"æ›¿æ¢è§„åˆ™: {variant} â†’ {value}")
                        return True
        
        return False
    
    def _try_fill_rule(self, paragraph: Paragraph, info: Dict[str, Any]) -> bool:
        """
        å°è¯•å¡«ç©ºè§„åˆ™
        å¦‚ï¼šä¾›åº”å•†åç§°ï¼š______ â†’ ä¾›åº”å•†åç§°ï¼šå…¬å¸å
        """
        text = paragraph.text
        
        # å¤„ç†ä¾›åº”å•†åç§°ç±»çš„å¡«ç©º
        for variant in self.company_name_variants:
            # å¤šç§å¡«ç©ºæ ¼å¼
            patterns = [
                rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]*$',  # å†’å·åè·Ÿä¸‹åˆ’çº¿æˆ–ç©ºæ ¼
                rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]+[ã€‚\.]',  # å†’å·åè·Ÿä¸‹åˆ’çº¿ï¼Œä»¥å¥å·ç»“æŸ
                rf'{re.escape(variant)}\s+[_\s]+$',  # ç©ºæ ¼åè·Ÿä¸‹åˆ’çº¿
                rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]{{10,}}',  # ğŸ†• 10+ä¸ªç©ºæ ¼/ä¸‹åˆ’çº¿
                rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]+ï¼ˆ[^ï¼‰]*ï¼‰',  # ğŸ†• ç©ºæ ¼+ï¼ˆç›–ç« ï¼‰
            ]

            for pattern in patterns:
                if re.search(pattern, text):
                    company_name = info.get('companyName', '')
                    if company_name:
                        # ä¿ç•™åç¼€ï¼ˆå¦‚"ï¼ˆåŠ ç›–å…¬ç« ï¼‰"ï¼‰
                        suffix_pattern = r'ï¼ˆ[^ï¼‰]*[ç« å­—å°]ï¼‰'
                        suffix_match = re.search(suffix_pattern, text)
                        suffix = suffix_match.group(0) if suffix_match else ''

                        # æå–è¦æ›¿æ¢çš„éƒ¨åˆ†
                        text_without_suffix = text.replace(suffix, '') if suffix else text

                        # æ›¿æ¢ç©ºç™½ç¬¦
                        new_text = re.sub(r'[_\s]+', company_name, text_without_suffix).rstrip()

                        # å¦‚æœæœ‰åç¼€ï¼Œæ·»åŠ ç©ºæ ¼åˆ†éš”
                        if suffix:
                            new_text = new_text + '  ' + suffix

                        self._update_paragraph_text(paragraph, new_text)
                        self.logger.info(f"å¡«ç©ºè§„åˆ™: {variant} å¡«å…¥ {company_name}")
                        return True
        
        # å¤„ç†å…¶ä»–å­—æ®µçš„å¡«ç©º
        for field_key, variants in self.field_variants.items():
            for variant in variants:
                patterns = [
                    rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]*$',
                    rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]+[ã€‚\.]',
                    rf'{re.escape(variant)}\s+[_\s]+$',
                    rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]{{10,}}',  # ğŸ†• 10+ä¸ªç©ºæ ¼/ä¸‹åˆ’çº¿
                    rf'{re.escape(variant)}\s*[:ï¼š]\s*[_\s]+ï¼ˆ[^ï¼‰]*ï¼‰',  # ğŸ†• ç©ºæ ¼+ï¼ˆå¤‡æ³¨ï¼‰
                ]

                for pattern in patterns:
                    if re.search(pattern, text):
                        value = info.get(field_key, '')
                        if value:
                            # ä¿ç•™åç¼€ï¼ˆå¦‚"ï¼ˆç›–ç« ï¼‰"ï¼‰
                            suffix_pattern = r'ï¼ˆ[^ï¼‰]*[ç« å­—å°æ³¨]ï¼‰'
                            suffix_match = re.search(suffix_pattern, text)
                            suffix = suffix_match.group(0) if suffix_match else ''

                            # æå–è¦æ›¿æ¢çš„éƒ¨åˆ†
                            text_without_suffix = text.replace(suffix, '') if suffix else text

                            # æ›¿æ¢ç©ºç™½ç¬¦
                            new_text = re.sub(r'[_\s]+', value, text_without_suffix).rstrip()

                            # å¦‚æœæœ‰åç¼€ï¼Œæ·»åŠ ç©ºæ ¼åˆ†éš”
                            if suffix:
                                new_text = new_text + '  ' + suffix

                            self._update_paragraph_text(paragraph, new_text)
                            self.logger.info(f"å¡«ç©ºè§„åˆ™: {variant} å¡«å…¥ {value}")
                            return True
        
        # ç‰¹æ®Šå¤„ç†æ—¥æœŸ
        # æ”¯æŒå¤šç§æ—¥æœŸå ä½ç¬¦æ ¼å¼ï¼š____å¹´____æœˆ____æ—¥ æˆ– XXXXå¹´XæœˆXæ—¥ æˆ– XXå¹´ XXæœˆXXæ—¥
        date_patterns = [
            r'æ—¥\s*æœŸ\s*[:ï¼š]?\s*[_\s]*å¹´[_\s]*æœˆ[_\s]*æ—¥',  # æ—¥æœŸï¼š____å¹´____æœˆ____æ—¥
            r'æ—¥\s*æœŸ\s*[:ï¼š]?\s*[Xx]{1,4}\s*å¹´\s*[Xx]{1,2}\s*æœˆ\s*[Xx]{1,2}\s*æ—¥',  # æ—¥æœŸï¼šXXXXå¹´XæœˆXæ—¥ æˆ– XXå¹´ XXæœˆXXæ—¥ï¼ˆæ”¯æŒç©ºæ ¼å’Œå°å†™ï¼‰
            r'[Xx]{1,4}\s*å¹´\s*[Xx]{1,2}\s*æœˆ\s*[Xx]{1,2}\s*æ—¥',  # XXXXå¹´XæœˆXæ—¥ æˆ– XXå¹´ XXæœˆXXæ—¥ï¼ˆä¸å¸¦æ—¥æœŸå‰ç¼€ï¼Œæ”¯æŒç©ºæ ¼å’Œå°å†™ï¼‰
        ]

        for date_pattern in date_patterns:
            if re.search(date_pattern, text):
                date_text = info.get('date', '')
                if date_text:
                    # æ ¼å¼åŒ–æ—¥æœŸ
                    formatted_date = self._format_date(date_text)
                    # å¦‚æœæ¨¡å¼åŒ…å«"æ—¥æœŸ"ï¼Œä¿ç•™æ ¼å¼ï¼›å¦åˆ™ç›´æ¥æ›¿æ¢
                    if 'æ—¥' in date_pattern and 'æœŸ' in date_pattern:
                        new_text = re.sub(date_pattern, f'æ—¥æœŸï¼š{formatted_date}', text)
                    else:
                        new_text = re.sub(date_pattern, formatted_date, text)
                    self._update_paragraph_text(paragraph, new_text)
                    self.logger.info(f"æ—¥æœŸå¡«ç©º: {formatted_date}")
                    return True
        
        return False
    
    def _process_table(self, table: Table, info: Dict[str, Any]) -> Dict[str, Any]:
        """å¤„ç†è¡¨æ ¼ä¸­çš„ä¿¡æ¯å¡«å†™"""
        count = 0
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if self._try_fill_rule(paragraph, info):
                        count += 1
                    elif self._try_replacement_rule(paragraph, info):
                        count += 1
        
        return {'count': count}
    
    def _format_date(self, date_str: str) -> str:
        """æ ¼å¼åŒ–æ—¥æœŸå­—ç¬¦ä¸²"""
        # ç§»é™¤å¤šä½™çš„ç©ºæ ¼
        date_str = re.sub(r'\s+', '', date_str)
        
        # å°è¯•åŒ¹é…å¸¸è§æ—¥æœŸæ ¼å¼
        patterns = [
            (r'(\d{4})-(\d{1,2})-(\d{1,2})', r'\1å¹´\2æœˆ\3æ—¥'),
            (r'(\d{4})/(\d{1,2})/(\d{1,2})', r'\1å¹´\2æœˆ\3æ—¥'),
            (r'(\d{4})\.(\d{1,2})\.(\d{1,2})', r'\1å¹´\2æœˆ\3æ—¥'),
        ]
        
        for pattern, replacement in patterns:
            if re.match(pattern, date_str):
                return re.sub(pattern, replacement, date_str)
        
        # å¦‚æœå·²ç»æ˜¯ä¸­æ–‡æ ¼å¼ï¼Œç›´æ¥è¿”å›
        if 'å¹´' in date_str and 'æœˆ' in date_str and 'æ—¥' in date_str:
            return date_str
        
        # é»˜è®¤è¿”å›åŸå­—ç¬¦ä¸²
        return date_str
    
    def _update_paragraph_text(self, paragraph: Paragraph, new_text: str):
        """æ›´æ–°æ®µè½æ–‡æœ¬ï¼Œä¿æŒæ ¼å¼"""
        # ä¿å­˜ç¬¬ä¸€ä¸ªrunçš„æ ¼å¼
        if paragraph.runs:
            first_run = paragraph.runs[0]
            # ä¿å­˜æ ¼å¼å±æ€§
            font = first_run.font
            bold = font.bold
            italic = font.italic
            underline = font.underline
            font_size = font.size
            font_name = font.name
            
            # æ¸…ç©ºæ®µè½
            paragraph.clear()
            
            # æ·»åŠ æ–°æ–‡æœ¬å¹¶æ¢å¤æ ¼å¼
            new_run = paragraph.add_run(new_text)
            if bold is not None:
                new_run.font.bold = bold
            if italic is not None:
                new_run.font.italic = italic
            if underline is not None:
                new_run.font.underline = underline
            if font_size:
                new_run.font.size = font_size
            if font_name:
                new_run.font.name = font_name
        else:
            # å¦‚æœæ²¡æœ‰runsï¼Œç›´æ¥è®¾ç½®æ–‡æœ¬
            paragraph.text = new_text
    
    def _post_process(self, doc: Document):
        """åå¤„ç†ï¼šæ¸…ç†å¤šä½™çš„å ä½ç¬¦å’Œæ ¼å¼"""
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # æ¸…ç†å¤šä½™çš„ä¸‹åˆ’çº¿
            text = re.sub(r'_{3,}', '', text)
            
            # æ¸…ç†å¤šä½™çš„ç©ºæ ¼ï¼ˆä¿ç•™è¡¨æ ¼å¯¹é½æ‰€éœ€çš„ç©ºæ ¼ï¼‰
            if not re.search(r'\s{8,}', text):  # ä¸æ˜¯è¡¨æ ¼å¼å¸ƒå±€
                text = re.sub(r'\s{3,}', '  ', text)
            
            # æ¸…ç†å¤šä½™çš„å†’å·
            text = re.sub(r'[:ï¼š]{2,}', 'ï¼š', text)
            
            # æ ‡å‡†åŒ–å†’å·
            text = re.sub(r':', 'ï¼š', text)
            
            # å»é™¤å¤šä½™çš„å¹´æœˆæ—¥æ ‡è¯†
            text = re.sub(r'(\d{4}å¹´\d{1,2}æœˆ\d{1,2}æ—¥)\s*å¹´\s*æœˆ\s*æ—¥', r'\1', text)
            
            if text != paragraph.text:
                self._update_paragraph_text(paragraph, text.strip())