# ai_tender_system/modules/document_merger/merger_service.py

import os
import re
import json
from typing import Optional, Callable
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn


class DocumentMergerService:
    """投标文档智能整合服务 - V2版本"""

    def __init__(self):
        pass

    def merge_documents_v2(
        self,
        project_id: int,
        file_paths: dict,
        config: dict,
        output_dir: str,
        progress_callback: Optional[Callable] = None
    ) -> dict:
        """
        智能整合投标文档（V2增强版）

        Args:
            project_id: 项目ID
            file_paths: {
                "business": str,  # 商务应答文件路径
                "p2p": str,       # 点对点应答文件路径（可选）
                "tech": str       # 技术方案文件路径
            }
            config: {
                "include_p2p": bool,           # 是否包含点对点
                "doc_order": list,             # 文档顺序 ["business", "p2p", "tech"]
                "generate_toc": bool,          # 生成目录
                "remove_blanks": bool,         # 删除空白
                "unify_styles": bool,          # 统一样式
                "add_section_breaks": bool,    # 添加分节符
                "index_config": dict,          # 索引配置
                "output_filename": str         # 输出文件名
            }
            output_dir: 输出目录
            progress_callback: 进度回调函数 fn(percent, message)

        Returns:
            {
                "docx_path": str,
                "stats": {
                    "total_pages": int,
                    "business_pages": int,
                    "p2p_pages": int,
                    "tech_pages": int,
                    "toc_pages": int,
                    "index_pages": int,
                    "removed_blanks": int
                }
            }
        """
        self._report_progress(progress_callback, 5, "开始文档整合...")

        # 1. 加载文档
        self._report_progress(progress_callback, 10, "正在读取商务应答文件...")
        business_doc = Document(file_paths["business"])

        docs_to_merge = []
        doc_labels = []

        # 根据配置确定合并顺序
        doc_order = config.get("doc_order", ["business", "p2p", "tech"])

        for doc_type in doc_order:
            if doc_type == "business":
                docs_to_merge.append((business_doc, "商务应答"))
                doc_labels.append("business")
            elif doc_type == "p2p" and config.get("include_p2p", True):
                if file_paths.get("p2p"):
                    self._report_progress(progress_callback, 15, "正在读取点对点应答文件...")
                    p2p_doc = Document(file_paths["p2p"])
                    docs_to_merge.append((p2p_doc, "技术点对点应答"))
                    doc_labels.append("p2p")
            elif doc_type == "tech":
                self._report_progress(progress_callback, 20, "正在读取技术方案文件...")
                tech_doc = Document(file_paths["tech"])
                docs_to_merge.append((tech_doc, "技术方案"))
                doc_labels.append("tech")

        # 2. 创建主文档（以商务应答为基准）
        self._report_progress(progress_callback, 25, "正在创建主文档...")
        main_doc = Document()

        # 复制商务应答的样式
        if config.get("unify_styles", True):
            self._copy_styles(main_doc, business_doc)

        # 3. 合并文档内容
        progress_step = 30
        for i, (doc, label) in enumerate(docs_to_merge):
            self._report_progress(
                progress_callback,
                progress_step + i * 15,
                f"正在合并{label}..."
            )

            # 如果不是第一个文档，添加分节符
            if i > 0 and config.get("add_section_breaks", True):
                main_doc.add_page_break()

            # 复制内容
            self._copy_document_content(main_doc, doc)

        # 4. 删除空白段落
        removed_blanks = 0
        if config.get("remove_blanks", True):
            self._report_progress(progress_callback, 65, "正在检查并删除空白段落...")
            removed_blanks = self._detect_and_remove_blanks(main_doc)

        # 5. 生成目录
        toc_pages = 0
        if config.get("generate_toc", True):
            self._report_progress(progress_callback, 75, "正在生成目录...")
            toc_pages = self._generate_toc_advanced(main_doc)

        # 6. 生成索引
        index_pages = 0
        index_config = config.get("index_config", {})
        if index_config.get("required", False):
            self._report_progress(progress_callback, 85, "正在生成索引...")
            index_pages = self._generate_index(main_doc, index_config)

        # 7. 保存文档
        self._report_progress(progress_callback, 90, "正在保存文档...")
        output_filename = config.get("output_filename", f"项目{project_id}_最终标书")
        if not output_filename.endswith('.docx'):
            output_filename += '.docx'

        output_path = os.path.join(output_dir, output_filename)
        main_doc.save(output_path)

        # 8. 计算统计信息
        self._report_progress(progress_callback, 95, "正在计算文档统计...")
        stats = {
            "total_pages": self._estimate_page_count(main_doc),
            "toc_pages": toc_pages,
            "index_pages": index_pages,
            "removed_blanks": removed_blanks,
            "business_pages": 0,  # TODO: 精确计算
            "p2p_pages": 0,
            "tech_pages": 0
        }

        self._report_progress(progress_callback, 100, "文档整合完成！")

        return {
            "docx_path": output_path,
            "stats": stats
        }

    def _copy_styles(self, target_doc: Document, source_doc: Document):
        """复制源文档的样式到目标文档"""
        try:
            # 复制样式定义
            for style in source_doc.styles:
                try:
                    if style.name not in [s.name for s in target_doc.styles]:
                        # 注意：完整的样式复制比较复杂，这里做简化处理
                        pass
                except:
                    pass
        except Exception as e:
            print(f"样式复制警告: {e}")

    def _copy_document_content(self, target_doc: Document, source_doc: Document):
        """复制源文档的所有内容到目标文档"""
        # 复制段落
        for para in source_doc.paragraphs:
            # 尝试使用原样式，如果不存在则使用None（默认样式）
            style_name = None
            if para.style:
                try:
                    # 检查样式是否存在于目标文档
                    target_doc.styles[para.style.name]
                    style_name = para.style.name
                except KeyError:
                    # 样式不存在，使用默认样式
                    print(f"警告：样式 '{para.style.name}' 不存在，使用默认样式")
                    style_name = None

            new_para = target_doc.add_paragraph(para.text, style=style_name)
            new_para.alignment = para.alignment

            # 复制段落格式
            if para.paragraph_format:
                try:
                    new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                    new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                    new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                    new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                except Exception as e:
                    # 格式复制失败，忽略
                    pass

            # 复制字体格式
            if para.runs:
                for i, run in enumerate(para.runs):
                    try:
                        if i < len(new_para.runs):
                            new_run = new_para.runs[i]
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.name:
                                new_run.font.name = run.font.name
                    except Exception as e:
                        # 字体格式复制失败，忽略
                        pass

        # 复制表格
        for table in source_doc.tables:
            new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))

            # 尝试应用表格样式
            try:
                if table.style:
                    new_table.style = table.style
            except Exception as e:
                # 样式不存在，使用默认表格样式
                print(f"警告：表格样式不存在，使用默认样式")

            # 复制单元格内容
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text

    def _detect_and_remove_blanks(self, doc: Document) -> int:
        """检测并删除空白段落"""
        removed = 0

        # 倒序遍历，避免索引问题
        for i in range(len(doc.paragraphs) - 1, -1, -1):
            para = doc.paragraphs[i]
            if self._is_blank_paragraph(para):
                # 删除段落
                p_element = para._element
                p_element.getparent().remove(p_element)
                removed += 1

        return removed

    def _is_blank_paragraph(self, para) -> bool:
        """判断段落是否为空白"""
        # 1. 检查文本内容
        if para.text.strip():
            return False

        # 2. 检查是否包含图片
        if para._element.xpath('.//w:drawing') or \
           para._element.xpath('.//w:pict'):
            return False

        # 3. 检查是否包含表格
        if para._element.xpath('.//w:tbl'):
            return False

        return True

    def _generate_toc_advanced(self, doc: Document) -> int:
        """生成高级目录（在文档开头插入）"""
        # 在文档开头插入目录
        if len(doc.paragraphs) == 0:
            doc.add_paragraph()

        # 添加目录标题
        toc_heading = doc.paragraphs[0]._element.addprevious(
            doc._element.body.add_p()
        )
        toc_para = doc.paragraphs[0]
        toc_para.text = '目录'
        toc_para.style = 'Heading 1'
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 插入目录域
        toc_field_para = toc_para._element.addnext(
            doc._element.body.add_p()
        )

        # 获取新插入的段落对象
        toc_field = None
        for para in doc.paragraphs:
            if para._element == toc_field_para:
                toc_field = para
                break

        if toc_field:
            run = toc_field.add_run()

            # 添加TOC域代码
            fldChar1 = run._element.add_fldChar()
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = run._element.add_instrText()
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

            fldChar2 = run._element.add_fldChar()
            fldChar2.set(qn('w:fldCharType'), 'end')

        # 添加分页符
        doc.add_page_break()

        return 2  # 目录通常占2页

    def _generate_index(self, doc: Document, index_config: dict) -> int:
        """根据配置生成索引"""
        index_type = index_config.get("type", "none")

        if index_type == "fixed_format":
            return self._generate_fixed_index(doc, index_config.get("template", ""))
        elif index_type == "score_based":
            return self._generate_score_index(doc, index_config.get("score_items", []))

        return 0

    def _generate_fixed_index(self, doc: Document, template: str) -> int:
        """生成固定格式索引"""
        doc.add_page_break()

        heading = doc.add_heading('索引', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        lines = template.split('\n')
        for line in lines:
            if not line.strip():
                continue

            # TODO: 智能查找页码
            # 现在先保留原样
            doc.add_paragraph(line)

        return 1

    def _generate_score_index(self, doc: Document, score_items: list) -> int:
        """生成评分标准索引"""
        doc.add_page_break()

        heading = doc.add_heading('评分标准对照索引', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '评分项'
        header_cells[1].text = '对应章节'
        header_cells[2].text = '页码'

        # 填充评分项
        for item in score_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = '待定位'  # TODO: 智能匹配
            row_cells[2].text = '第X页'

        return 1

    def _estimate_page_count(self, doc: Document) -> int:
        """估算文档页数（简化版）"""
        # 简单估算：段落数 / 25 + 表格数 * 0.5
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        estimated_pages = int(para_count / 25 + table_count * 0.5)
        return max(estimated_pages, 1)

    def _report_progress(self, callback, percent: int, message: str):
        """报告进度"""
        if callback:
            callback(percent, message)

    # ========== 保留旧版本方法以兼容现有代码 ==========

    def merge_documents(
        self,
        business_doc_path: str,
        p2p_doc_path: str,
        tech_doc_path: str,
        output_dir: str,
        project_name: str,
        style_option: str = "business_style",
        progress_callback=None
    ) -> dict:
        """
        旧版本合并方法（保留兼容性）
        """
        if progress_callback:
            progress_callback("开始文件融合...")

        # 使用新版本方法
        file_paths = {
            "business": business_doc_path,
            "p2p": p2p_doc_path,
            "tech": tech_doc_path
        }

        config = {
            "include_p2p": True,
            "doc_order": ["business", "p2p", "tech"],
            "generate_toc": True,
            "remove_blanks": True,
            "unify_styles": style_option == "business_style",
            "add_section_breaks": True,
            "index_config": {"required": False},
            "output_filename": f"{project_name}_最终标书.docx"
        }

        result = self.merge_documents_v2(
            project_id=0,
            file_paths=file_paths,
            config=config,
            output_dir=output_dir,
            progress_callback=lambda p, m: progress_callback(m) if progress_callback else None
        )

        return {
            "merged_document_path": result["docx_path"],
            "index_file_path": None
        }
