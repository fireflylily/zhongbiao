#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标信息提取器 - 重构版本
从招标文档中提取项目信息、资质要求和技术评分标准
"""

import requests
import json
import re
import threading
import configparser
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, List, Any
import tempfile
import os

# 导入公共模块
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from common import (
    get_config, get_module_logger,
    TenderInfoExtractionError, APIError, FileProcessingError,
    get_prompt_manager
)
from common.llm_client import create_llm_client
from common.database import get_knowledge_base_db

class TenderInfoExtractor:
    """招标信息提取器"""
    
    def __init__(self, api_key: Optional[str] = None, model_name: str = "gpt-4o-mini"):
        self.config = get_config()
        self.logger = get_module_logger("tender_info")

        # 初始化提示词管理器
        self.prompt_manager = get_prompt_manager()

        # 创建LLM客户端
        self.llm_client = create_llm_client(model_name, api_key)

        # 初始化数据库连接
        self.db = get_knowledge_base_db()

        # 保持向后兼容性的配置
        api_config = self.config.get_api_config()
        self.api_key = api_key or api_config['api_key']
        self.api_endpoint = api_config['api_endpoint']
        self.model_name = model_name
        self.max_tokens = api_config['max_tokens']
        self.timeout = api_config['timeout']

        self.logger.info(f"招标信息提取器初始化完成，使用模型: {model_name}")

    def _safe_json_parse(self, response: str, task_name: str) -> Optional[Dict]:
        """
        安全的JSON解析，增强容错性
        处理各种格式的LLM响应
        """
        if not response or not response.strip():
            self.logger.error(f"{task_name} - 响应为空")
            return None

        # 移除可能的markdown代码块标记
        response = re.sub(r'^```json\s*', '', response.strip())
        response = re.sub(r'\s*```$', '', response.strip())

        # 尝试多种JSON提取策略
        json_candidates = []

        # 策略1: 查找最外层的完整JSON对象
        json_start = response.find('{')
        json_end = response.rfind('}') + 1
        if json_start != -1 and json_end > json_start:
            json_candidates.append(response[json_start:json_end])

        # 策略2: 使用正则表达式查找JSON模式
        json_pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
        matches = re.findall(json_pattern, response, re.DOTALL)
        json_candidates.extend(matches)

        # 策略3: 如果响应看起来已经是JSON，直接使用
        if response.strip().startswith('{') and response.strip().endswith('}'):
            json_candidates.append(response.strip())

        # 尝试解析每个候选
        for candidate in json_candidates:
            if not candidate.strip():
                continue

            try:
                # 清理常见的JSON格式问题
                cleaned = self._clean_json_string(candidate)
                result = json.loads(cleaned)
                self.logger.info(f"{task_name} - JSON解析成功")
                return result
            except json.JSONDecodeError as e:
                self.logger.debug(f"{task_name} - JSON候选解析失败: {e}")
                continue

        # 所有策略都失败
        self.logger.error(f"{task_name} - 所有JSON解析策略都失败")
        self.logger.error(f"{task_name} - 原始响应: {response[:500]}...")  # 改为error级别，确保能看到
        return None

    def _clean_json_string(self, json_str: str) -> str:
        """
        清理JSON字符串中的常见问题
        """
        # 移除前后空白
        json_str = json_str.strip()

        # 修复单引号为双引号（但要小心字符串内容中的单引号）
        # 简单的启发式规则：仅在看起来像属性名的地方替换
        json_str = re.sub(r"'([^']*)':", r'"\1":', json_str)

        # 修复结尾缺少逗号的问题
        json_str = re.sub(r'"\s*\n\s*"', '",\n"', json_str)

        # 移除JSON中的注释（// 和 /* */ 风格）
        json_str = re.sub(r'//.*?$', '', json_str, flags=re.MULTILINE)
        json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)

        return json_str

    def _validate_project_name(self, project_name: str) -> bool:
        """
        验证项目名称是否合理

        Args:
            project_name: 待验证的项目名称

        Returns:
            True if valid, False otherwise
        """
        if not project_name or not project_name.strip():
            return False

        project_name = project_name.strip()

        # === 基础检查 ===

        # 1. 长度检查（通常不超过100字符）
        if len(project_name) > 100:
            self.logger.warning(f"项目名称过长（{len(project_name)}字符）: {project_name[:50]}...")
            return False

        # 2. 最短有效长度检查（去除空格和标点后至少10个有效字符）
        effective_chars = re.sub(r'[\s\.\d\-、，。]', '', project_name)
        if len(effective_chars) < 10:
            self.logger.warning(f"项目名称有效字符过少（{len(effective_chars)}字符）: {project_name}")
            return False

        # === 条款编号模式检测 ===

        # 3. 检测条款编号（如"5.4"、"3.2.1"、"第三条"）
        clause_number_patterns = [
            r'^\d+\.\d+',              # "5.4"、"3.2"
            r'^\d+\.\d+\.\d+',         # "3.2.1"
            r'^第[一二三四五六七八九十\d]+条',  # "第三条"
            r'^第[一二三四五六七八九十\d]+款',  # "第五款"
            r'^第[一二三四五六七八九十\d]+项',  # "第二项"
            r'^\(\d+\)',               # "(1)"、"(2)"
            r'^（[一二三四五六七八九十\d]+）',  # "（一）"
        ]

        for pattern in clause_number_patterns:
            if re.match(pattern, project_name):
                self.logger.warning(f"项目名称疑似条款编号（模式: {pattern}）: {project_name}")
                return False

        # === 单字/双字无意义词检测 ===

        # 4. 检测单字或双字无意义前缀/后缀（如"本"、"该"、"此"）
        meaningless_words = ['本', '该', '此', '其', '上述', '前述', '下述', '如下', '所述']
        for word in meaningless_words:
            # 检查是否以无意义词开头或结尾
            if project_name.startswith(word) or project_name.endswith(word):
                self.logger.warning(f"项目名称包含无意义词（'{word}'）: {project_name}")
                return False
            # 检查整个名称是否只有无意义词+标点
            if re.match(f'^{word}[\\s\\.、，。]*$', project_name):
                self.logger.warning(f"项目名称只是无意义词: {project_name}")
                return False

        # === 描述性词汇检测 ===

        # 5. 检测描述性词汇（这些词通常出现在描述性语句中，而非真实项目名称）
        descriptive_patterns = [
            r'技术.*规范.*中.*所述',  # "技术规范中所述"
            r'本项目.*的.*相关',      # "本项目的相关"
            r'上述.*项目',            # "上述项目"
            r'前述.*内容',            # "前述内容"
            r'如下.*所述',            # "如下所述"
            r'详见.*附件',            # "详见附件"
            r'参见.*文件',            # "参见文件"
            r'按照.*要求',            # "按照要求"
            r'^该项目',               # "该项目XX"
            r'^本项目',               # "本项目XX"
            r'^此项目',               # "此项目XX"
        ]

        for pattern in descriptive_patterns:
            if re.search(pattern, project_name):
                self.logger.warning(f"项目名称包含描述性词汇（模式: {pattern}）: {project_name}")
                return False

        # === 文件类型描述检测 ===

        # 6. 检测是否只是文件类型描述
        file_type_only_patterns = [
            r'^(招标|磋商|询价|采购)文件$',
            r'^技术(服务)?规范$',
            r'^投标须知$',
            r'^供应商须知$',
        ]

        for pattern in file_type_only_patterns:
            if re.match(pattern, project_name):
                self.logger.warning(f"项目名称只是文件类型描述: {project_name}")
                return False

        return True

    def _extract_project_name_by_regex(self, text: str) -> Optional[str]:
        """
        使用正则表达式提取项目名称（辅助方法）

        Args:
            text: 文档文本

        Returns:
            提取的项目名称，如果未找到则返回None
        """
        # 正则表达式模式（按优先级排序）
        patterns = [
            # 明确的"项目名称："字段
            r"项目名称[：:]\s*([^\n\r]{5,100})",
            # "关于XX项目的公告"格式（贪婪匹配，保留"项目"）
            r"关于[""\"']?(.{5,80}?项目)[""\"']?(?:的|之)",
            # "XX项目招标/磋商/采购文件"格式
            r"([^\n\r]{5,80}?)项目.*?(?:招标|磋商|采购)文件",
            # 标题中的"XX项目"
            r"^([^\n\r]{5,80}?)项目",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                project_name = match.group(1).strip()
                # 清理项目名称中的特殊字符
                project_name = re.sub(r'\s+', ' ', project_name)  # 多个空格合并为一个

                # 验证提取的项目名称
                if self._validate_project_name(project_name):
                    self.logger.info(f"正则提取到项目名称: {project_name} (模式: {pattern})")
                    return project_name

        return None

    def _get_qualification_keywords(self) -> Dict[str, List[str]]:
        """
        获取资质关键字匹配规则
        基于数据库中实际的18种资质类型
        """
        return {
            # 基础资质类（考虑三证合一）
            'business_license': [
                # 营业执照相关
                '营业执照', '三证合一', '企业法人营业执照',
                '工商营业执照', '统一社会信用代码', '一照一码',
                '企业法人', '工商登记',

                # 注册资本相关
                '注册资金', '注册资本', '注册资本金', '实缴资本',
                '认缴资本', '注册资本要求',

                # 法人资格相关
                '独立法人', '法人资格', '独立法人资格',
                '独立承担民事责任', '民事责任能力',
                '独立承担民事责任的能力',

                # 成立时间相关
                '成立时间', '成立年限', '注册时间'
            ],
            'legal_id_front': [
                '法人身份证正面', '法定代表人身份证', '法人身份证',
                '法人代表身份证', '企业法人身份证'
            ],
            'legal_id_back': [
                '法人身份证反面', '法定代表人身份证反面',
                '法人身份证背面'
            ],
            'auth_id_front': [
                '被授权人身份证正面', '被授权人身份证', '委托代理人身份证',
                '代理人身份证', '投标代表身份证', '受托人身份证'
            ],
            'auth_id_back': [
                '被授权人身份证反面', '被授权人身份证背面',
                '代理人身份证反面', '投标代表身份证反面', '受托人身份证反面'
            ],
            'authorization_letter': [
                '法人授权委托书', '授权书', '法人授权', '委托书',
                '投标授权书', '授权委托书', '法定代表人授权书'
            ],

            # 认证证书类
            'iso9001': [
                'ISO9001', 'ISO 9001', '质量管理体系', '质量认证',
                'GB/T19001', 'iso9001', '质量体系认证'
            ],
            'iso20000': [
                'ISO20000', 'ISO 20000', '信息技术服务管理', 'IT服务管理',
                'ISO/IEC 20000', 'iso20000', '信息技术服务'
            ],
            'iso27001': [
                'ISO27001', 'ISO 27001', '信息安全管理', '信息安全认证',
                'ISO/IEC 27001', 'iso27001', '信息安全体系'
            ],
            'cmmi': [
                'CMMI', '能力成熟度', '软件能力成熟度', 'CMMI认证',
                'cmmi', '软件过程改进', '能力成熟度集成'
            ],
            'itss': [
                'ITSS', '信息技术服务标准', 'ITSS认证', '运维服务能力',
                'itss', 'IT服务标准', '信息技术服务'
            ],

            # 行业资质类
            'telecom_license': [
                '基础电信业务许可证', '基础电信业务经营许可证', '基础电信许可证',
                '增值电信业务许可证', '增值电信业务经营许可证', '增值电信许可证',
                'ICP许可证', 'ICP经营许可证', 'ISP许可证', 'IDC许可证',
                '电信业务许可证', '电信经营许可证', '电信业务经营许可'
            ],
            'value_added_telecom_license': [
                '增值电信业务许可证', '增值电信业务经营许可证', '增值电信许可证',
                'ICP许可证', 'ICP经营许可证', 'ISP许可证', 'IDC许可证', 'CDN许可证',
                '增值电信', 'ICP', 'IDC', 'ISP', 'CDN'
            ],
            'basic_telecom_license': [
                '基础电信业务许可证', '基础电信业务经营许可证', '基础电信许可证',
                '电信业务经营许可证', '基础电信'
            ],
            'level_protection': [
                '等保三级', '等级保护三级', '信息安全等级保护', '等保',
                '三级等保', '等级保护备案', '等级保护', '网络安全等级保护',
                '等保测评', '等保认证'
            ],
            'software_copyright': [
                '软件著作权', '软著', '计算机软件著作权', '软件版权',
                '软件著作权登记证书', '著作权登记'
            ],
            'patent_certificate': [
                '专利证书', '专利', '发明专利', '实用新型', '外观设计',
                '知识产权', '专利权'
            ],
            'audit_report': [
                # 审计报告类
                '审计报告', '财务审计', '年度审计报告',
                '近三年审计报告', '审计', '会计师事务所',
                '财务审计报告', '审计证明', '审计意见',
                '审计师', '注册会计师',

                # 财务报表类（新增）
                '资产负债表', '利润表', '损益表',
                '现金流量表', '收入费用表', '财务报表',
                '资产负债', '利润状况', '三大报表',
                '财务会计报表', '会计报表',

                # 财务证明类（新增）
                '财务状况证明', '财务证明', '财务情况说明',
                '银行资信证明', '资信证明',

                # 财务制度类（新增）
                '财务会计制度', '财务管理制度', '财务体系',
                '会计制度', '会计核算体系', '财务管理体系',

                # 财务能力描述（新增）
                '财务状况良好', '良好的财务状况', '财务能力',
                '财务健康', '具有履约能力的财务状况'
            ],
            'project_performance': [
                '项目案例', '项目业绩', '业绩证明', '类似项目',
                '成功案例', '项目经验', '业绩要求', '项目经历',
                '类似业绩', '同类项目', '项目实施案例', '项目实绩',
                '承建项目', '项目实施经验', '合同业绩'
            ],

            # 信用资质类
            'social_security': [
                '社保', '社会保险', '社保缴纳', '社保证明',
                '养老保险', '医疗保险', '失业保险', '工伤保险', '生育保险',
                '社会保障', '社保缴费', '参保证明', '社会保险证明',
                '社保缴纳证明', '五险', '社会保险费'
            ],
            # 信用资质类（4个独立资质，与后端 qualification_matcher.py 保持一致）
            'dishonest_executor': [
                '失信被执行人', '失信名单', '失信被执行人名单',
                '不得被列入失信', '未被列入失信', '失信黑名单'
            ],
            'tax_violation_check': [
                '重大税收违法', '重大税收违法案件当事人名单',
                '重大税收违法失信主体', '税收违法', '税收黑名单',
                '不得被列入.*重大税收违法', '未被列入.*重大税收违法'
            ],
            'gov_procurement_creditchina': [
                '信用中国', 'www.creditchina.gov.cn', 'creditchina.gov.cn',
                '政府采购严重违法失信', '信用中国.*政府采购',
                '不得被列入.*政府采购.*信用中国'
            ],
            'gov_procurement_ccgp': [
                '中国政府采购网', 'www.ccgp.gov.cn', 'ccgp.gov.cn',
                '政府采购严重违法失信行为信息记录', '政府采购网.*违法',
                '不得被列入.*政府采购网'
            ],

            # 补充13条供应商资格要求中缺失的关键词
            'tax_compliance': [
                '依法纳税', '增值税', '纳税证明', '完税证明',
                '税收缴纳', '纳税凭证', '税务证明', '缴税证明',
                '增值税缴纳证明', '税收缴纳记录'
            ],
            'commitment_letter': [
                '承诺函', '投标承诺', '承诺书', '投标人承诺',
                '资格承诺', '诚信承诺', '投标承诺函', '承诺声明'
            ],
            'property_certificate': [
                '房产证明', '办公场所', '营业场所', '场地证明',
                '房屋产权', '租赁合同', '场所证明', '经营场所',
                '办公场地', '房产证', '产权证明'
            ],
            'deposit_requirement': [
                '保证金', '投标保证金', '履约保证金', '质量保证金',
                '担保金', '押金', '保证金缴纳', '保证金金额'
            ],
            'purchaser_blacklist': [
                '采购人黑名单', '黑名单', '禁入名单',
                '不良行为记录', '政府采购黑名单', '采购黑名单'
            ]
        }

    def extract_qualification_requirements_by_keywords(self, text: str) -> Dict[str, Any]:
        """
        使用关键字匹配提取资质要求
        替代LLM调用，提高性能和准确性

        增强功能：
        1. 否定词过滤 - 识别"本项目不适用"等排除性表述
        2. 上下文验证 - 确保是真实的资格要求
        """
        try:
            self.logger.info("开始使用关键字匹配提取资质要求")

            # 定义否定标记词
            negation_markers = [
                '本项目不适用', '不适用', '本次不适用', '本次采购不适用',
                '免除', '无需提供', '不需要', '不要求', '不涉及',
                '除外', '本项目除外', '已删除', '取消', '不作要求',
                '删除', '暂不要求', '可不提供', '非必须',
                '（不适用）', '【不适用】', '(不适用)', '[不适用]'
            ]

            # 获取关键字匹配规则
            keywords_mapping = self._get_qualification_keywords()

            # 存储检测结果
            qualification_results = {}

            # 转换文本为小写以便匹配
            text_lower = text.lower()

            for qual_key, keywords in keywords_mapping.items():
                found_keywords = []
                qualification_required = False
                description_parts = []

                # 检查每个关键字
                for keyword in keywords:
                    if keyword.lower() in text_lower:
                        # 使用句子提取方法，获取包含关键词的完整句子（更精准）
                        context = self._extract_sentence_with_keyword(text, keyword)

                        # 检查上下文中是否有否定标记
                        is_negated = False
                        matched_negation = None
                        for neg_marker in negation_markers:
                            if neg_marker in context:
                                is_negated = True
                                matched_negation = neg_marker
                                break

                        if is_negated:
                            # 跳过这个被否定的匹配
                            self.logger.debug(
                                f"跳过否定匹配 [{qual_key}]: "
                                f"关键词='{keyword}', 否定词='{matched_negation}', "
                                f"上下文='{context[:80]}...'"
                            )
                            continue

                        # 正常记录
                        found_keywords.append(keyword)
                        qualification_required = True

                        # 提取关键字周围的上下文作为描述
                        if context:
                            description_parts.append(context)

                # 如果找到关键字，记录结果
                if found_keywords:
                    # 去重并保留所有匹配的描述（支持多条匹配展示）
                    unique_descriptions = []
                    seen_descriptions = set()

                    for desc in description_parts:
                        # 截断到500字符
                        desc_truncated = desc[:500] if len(desc) > 500 else desc
                        # 去重（避免同一句子被重复添加）
                        if desc_truncated not in seen_descriptions:
                            unique_descriptions.append(desc_truncated)
                            seen_descriptions.add(desc_truncated)

                    # 如果有多条匹配，用换行符或特殊分隔符连接
                    # 如果只有一条，直接使用
                    if unique_descriptions:
                        description = unique_descriptions[0] if len(unique_descriptions) == 1 else '\n'.join(unique_descriptions)
                    else:
                        description = f"需要提供{found_keywords[0]}"

                    qualification_results[qual_key] = {
                        'required': True,
                        'keywords_found': found_keywords,
                        'description': description,
                        'match_count': len(unique_descriptions)  # 新增：匹配数量
                    }

            # 构建返回结果
            result = {
                'qualifications': qualification_results,
                'total_required': len(qualification_results),
                'keywords_method': True
            }

            self.logger.info(f"关键字匹配完成，检测到{len(qualification_results)}项资质要求")
            return result

        except Exception as e:
            self.logger.error(f"关键字匹配提取资质要求失败: {e}")
            return {}

    def _extract_context_around_keyword(self, text: str, keyword: str, context_length: int = 50) -> str:
        """
        提取关键字周围的上下文文本作为描述
        """
        try:
            # 找到关键字位置
            keyword_pos = text.lower().find(keyword.lower())
            if keyword_pos == -1:
                return ""

            # 计算上下文范围
            start = max(0, keyword_pos - context_length)
            end = min(len(text), keyword_pos + len(keyword) + context_length)

            context = text[start:end].strip()

            # 清理文本，移除换行符和多余空格
            context = ' '.join(context.split())

            return context

        except Exception as e:
            self.logger.debug(f"提取上下文失败: {e}")
            return ""

    def _extract_sentence_with_keyword(self, text: str, keyword: str) -> str:
        """
        提取包含关键词的完整句子（精准提取，避免过多上下文）

        Args:
            text: 完整文本
            keyword: 关键词

        Returns:
            包含关键词的完整句子
        """
        try:
            # 按句子分隔符分割文本（保留分隔符）
            import re

            # 使用正则表达式分割，但保留分隔符
            sentences = re.split(r'([。；！？])', text)

            # 重组句子（将分隔符附加回去）
            complete_sentences = []
            for i in range(0, len(sentences) - 1, 2):
                if i + 1 < len(sentences):
                    complete_sentences.append(sentences[i] + sentences[i + 1])

            # 查找包含关键词的句子（优先查找最短的）
            matching_sentences = []
            for sentence in complete_sentences:
                if keyword.lower() in sentence.lower():
                    matching_sentences.append(sentence.strip())

            # 如果找到匹配的句子，返回最短的那个（通常是最精准的）
            if matching_sentences:
                shortest = min(matching_sentences, key=len)
                # 如果句子太长（>500字符），尝试进一步分割
                if len(shortest) > 500:
                    # 尝试在逗号处分割，找到包含关键词的最小片段
                    parts = re.split(r'([，,])', shortest)
                    for i in range(0, len(parts) - 1, 2):
                        if i + 1 < len(parts):
                            part = parts[i] + parts[i + 1]
                            if keyword.lower() in part.lower():
                                return part.strip()
                return shortest

            # 如果没找到完整句子，尝试使用正则表达式提取
            # 针对"供应商...关键词..."这种常见格式
            pattern = f'供应商[^。；]*{re.escape(keyword)}[^。；]*[。；]'
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0).strip()

            # 最后兜底：返回关键词周围100字符
            self.logger.debug(f"未找到包含'{keyword}'的完整句子，使用上下文提取")
            return self._extract_context_around_keyword(text, keyword, 100)

        except Exception as e:
            self.logger.debug(f"句子提取失败: {e}，回退到上下文提取")
            return self._extract_context_around_keyword(text, keyword, 100)

    def _timeout_regex_search(self, pattern: str, text: str, timeout: int = 5):
        """带超时的正则表达式搜索，防止灾难性回溯"""
        result = None
        exception = None
        
        def search_target():
            nonlocal result, exception
            try:
                result = re.search(pattern, text)
            except Exception as e:
                exception = e
        
        thread = threading.Thread(target=search_target)
        thread.daemon = True
        thread.start()
        thread.join(timeout)
        
        if thread.is_alive():
            self.logger.warning(f"正则表达式搜索超时，跳过模式: {pattern[:50]}...")
            return None
        
        if exception:
            self.logger.warning(f"正则表达式搜索出错: {str(exception)}")
            return None
            
        return result
    
    def _timeout_regex_search_ignore_case(self, pattern: str, text: str, timeout: int = 5):
        """带超时的正则表达式搜索（忽略大小写）"""
        result = None
        exception = None
        
        def search_target():
            nonlocal result, exception
            try:
                result = re.search(pattern, text, re.IGNORECASE)
            except Exception as e:
                exception = e
        
        thread = threading.Thread(target=search_target)
        thread.daemon = True
        thread.start()
        thread.join(timeout)
        
        if thread.is_alive():
            self.logger.warning(f"正则表达式搜索超时（忽略大小写）: {pattern[:50]}...")
            return None
        
        if exception:
            self.logger.warning(f"正则表达式搜索出错（忽略大小写）: {str(exception)}")
            return None
            
        return result
    
    def llm_callback(self, prompt: str, purpose: str = "应答", max_retries: int = 3) -> str:
        """调用LLM API - 使用统一的LLM客户端"""
        try:
            return self.llm_client.call(
                prompt=prompt,
                temperature=0.7,  # 使用稍低的温度以获得更一致的结果
                max_retries=max_retries,
                purpose=purpose
            )
        except Exception as e:
            self.logger.error(f"LLM调用失败 - {purpose}: {str(e)}")
            raise APIError(f"LLM调用失败 - {purpose}: {str(e)}")
    
    def read_document(self, file_path: str) -> str:
        """读取文档内容"""
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                raise FileProcessingError(f"文件不存在: {file_path}")

            if file_path.suffix.lower() == '.pdf':
                return self._read_pdf(file_path)
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                try:
                    return self._read_word(file_path)
                except FileProcessingError as e:
                    # 如果Word读取失败，尝试使用备用方法
                    self.logger.warning(f"Word文档读取失败，尝试备用方法: {str(e)}")
                    return self._read_word_alternative(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._read_text(file_path)
            else:
                raise FileProcessingError(f"不支持的文件格式: {file_path.suffix}")

        except Exception as e:
            if isinstance(e, FileProcessingError):
                raise
            else:
                raise FileProcessingError(f"读取文档失败: {str(e)}")
    
    def _read_pdf(self, file_path: Path) -> str:
        """读取PDF文件"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                if not text.strip():
                    raise FileProcessingError("PDF文件内容为空或无法提取文本")
                
                self.logger.info(f"PDF文件读取成功，共{len(pdf_reader.pages)}页")
                return text
                
        except ImportError:
            raise FileProcessingError("缺少PyPDF2库，无法处理PDF文件")
        except Exception as e:
            raise FileProcessingError(f"PDF文件读取失败: {str(e)}")
    
    def _read_doc_with_antiword(self, file_path: Path) -> str:
        """使用antiword读取.doc文件"""
        try:
            # 尝试使用antiword命令行工具
            result = subprocess.run(
                ['/opt/homebrew/bin/antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )

            if result.returncode == 0:
                text = result.stdout.strip()
                if text:
                    self.logger.info(f"使用antiword成功读取.doc文件，内容长度: {len(text)}")
                    return text
                else:
                    raise FileProcessingError(".doc文件内容为空")
            else:
                raise FileProcessingError(f"antiword读取失败: {result.stderr}")

        except subprocess.TimeoutExpired:
            raise FileProcessingError("antiword执行超时")
        except FileNotFoundError:
            raise FileProcessingError("antiword工具未安装，请运行: brew install antiword")
        except Exception as e:
            raise FileProcessingError(f"使用antiword读取.doc文件失败: {str(e)}")

    def _read_doc_with_libreoffice(self, file_path: Path) -> str:
        """使用LibreOffice转换.doc文件为.docx后读取"""
        temp_dir = None
        temp_docx = None

        try:
            # 检查LibreOffice是否可用
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                '/usr/bin/libreoffice',
                '/usr/local/bin/libreoffice',
                'libreoffice'
            ]

            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run([path, '--version'],
                                          capture_output=True, timeout=5)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        break
                except (subprocess.TimeoutExpired, subprocess.SubprocessError, OSError, FileNotFoundError) as e:
                    self.logger.debug(f"检测LibreOffice路径 {path} 失败: {e}")
                    continue

            if not libreoffice_cmd:
                raise FileProcessingError("LibreOffice未安装或不可用")

            # 创建临时目录
            temp_dir = tempfile.mkdtemp()

            # 使用LibreOffice转换为docx
            self.logger.info(f"使用LibreOffice转换.doc文件: {file_path}")
            result = subprocess.run([
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', temp_dir,
                str(file_path)
            ], capture_output=True, text=True, timeout=60)

            if result.returncode != 0:
                raise FileProcessingError(f"LibreOffice转换失败: {result.stderr}")

            # 找到转换后的docx文件
            docx_filename = file_path.stem + '.docx'
            temp_docx = Path(temp_dir) / docx_filename

            if not temp_docx.exists():
                raise FileProcessingError("LibreOffice转换后的docx文件未找到")

            # 使用python-docx读取转换后的文件
            from docx import Document
            doc = Document(str(temp_docx))

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            text = '\n'.join(text_parts)

            if text.strip():
                self.logger.info(f"LibreOffice转换成功，内容长度: {len(text)}")
                return text
            else:
                raise FileProcessingError("转换后的文档内容为空")

        except subprocess.TimeoutExpired:
            raise FileProcessingError("LibreOffice转换超时")
        except ImportError:
            raise FileProcessingError("缺少python-docx库")
        except Exception as e:
            raise FileProcessingError(f"LibreOffice转换.doc文件失败: {str(e)}")
        finally:
            # 清理临时文件
            if temp_docx and temp_docx.exists():
                try:
                    os.remove(temp_docx)
                except (OSError, PermissionError) as e:
                    self.logger.warning(f"清理临时文件失败 {temp_docx}: {e}")
            if temp_dir and os.path.exists(temp_dir):
                try:
                    os.rmdir(temp_dir)
                except (OSError, PermissionError) as e:
                    self.logger.warning(f"清理临时目录失败 {temp_dir}: {e}")

    def _detect_office_suite(self, file_path: Path) -> str:
        """检测文档的创建软件"""
        try:
            # 使用 errors='ignore' 来处理编码问题
            result = subprocess.run(['file', str(file_path)],
                                  capture_output=True, timeout=10)

            if result.returncode == 0:
                # 尝试不同的编码方式
                try:
                    file_info = result.stdout.decode('utf-8').lower()
                except UnicodeDecodeError:
                    try:
                        file_info = result.stdout.decode('utf-8', errors='ignore').lower()
                    except (UnicodeDecodeError, AttributeError):
                        file_info = result.stdout.decode('latin-1', errors='ignore').lower()

                if 'wps' in file_info:
                    return 'WPS Office'
                elif 'microsoft' in file_info or 'office' in file_info:
                    return 'Microsoft Office'
                elif 'libreoffice' in file_info or 'openoffice' in file_info:
                    return 'LibreOffice/OpenOffice'

            # 作为备选方案，尝试读取文件的二进制开头来检测格式
            with open(file_path, 'rb') as f:
                header = f.read(512)
                header_str = header.decode('latin-1', errors='ignore').lower()

                if 'wps' in header_str:
                    return 'WPS Office'

            return 'Unknown'
        except Exception as e:
            self.logger.debug(f"Office suite detection failed: {e}")
            return 'Unknown'

    def _convert_doc_to_docx(self, file_path: Path) -> str:
        """将.doc文件转换为.docx"""
        try:
            import doc2docx

            # 创建临时文件
            temp_dir = tempfile.gettempdir()
            temp_docx = os.path.join(temp_dir, f"temp_extractor_{os.getpid()}_{file_path.stem}.docx")

            self.logger.info(f"开始转换.doc文件: {file_path} -> {temp_docx}")

            # 转换文件
            doc2docx.convert(str(file_path), temp_docx)

            # 检查转换后的文件是否存在
            if os.path.exists(temp_docx) and os.path.getsize(temp_docx) > 0:
                self.logger.info(f".doc文件转换成功: {temp_docx}")
                return temp_docx
            else:
                raise FileProcessingError("doc2docx转换失败，转换后的文件不存在或为空")

        except ImportError:
            raise FileProcessingError("缺少doc2docx库，无法转换.doc文件。请运行: pip install doc2docx")
        except Exception as e:
            self.logger.error(f".doc文件转换失败: {str(e)}")
            raise FileProcessingError(f"无法转换.doc文件: {e}")

    def _read_word(self, file_path: Path) -> str:
        """读取Word文件"""
        temp_file = None

        try:
            from docx import Document
            import zipfile

            # 检查是否是.doc格式，如果是则先尝试直接读取
            if file_path.suffix.lower() == '.doc':
                try:
                    # 先尝试使用antiword直接读取
                    text = self._read_doc_with_antiword(file_path)
                    return text
                except FileProcessingError as e:
                    self.logger.warning(f"antiword读取失败，尝试LibreOffice: {str(e)}")
                    # 如果antiword失败，尝试LibreOffice
                    try:
                        text = self._read_doc_with_libreoffice(file_path)
                        return text
                    except FileProcessingError as libreoffice_error:
                        self.logger.warning(f"LibreOffice转换失败，尝试doc2docx: {str(libreoffice_error)}")
                        # 如果LibreOffice也失败，尝试doc2docx转换
                        temp_file = self._convert_doc_to_docx(file_path)
                        file_path = Path(temp_file)
                        self.logger.info(f"使用转换后的临时文件: {file_path}")

            # 首先检查文件是否为有效的ZIP文件（docx本质上是ZIP文件）
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    # 检查是否包含Word文档的必要文件
                    required_files = ['word/document.xml', '[Content_Types].xml']
                    file_list = zip_ref.namelist()

                    for required_file in required_files:
                        if required_file not in file_list:
                            raise FileProcessingError(f"文件不是有效的Word文档格式：缺少 {required_file}")

                    self.logger.info("Word文档格式验证通过")

            except zipfile.BadZipFile:
                raise FileProcessingError("文件不是有效的Word文档格式：不是ZIP文件")
            except Exception as e:
                if "不是有效的Word文档格式" in str(e):
                    raise e
                else:
                    self.logger.warning(f"Word文档格式检查警告: {str(e)}")

            # 尝试读取Word文档
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # 读取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            if not text.strip():
                raise FileProcessingError("Word文档内容为空")

            self.logger.info(f"Word文档读取成功，共{len(doc.paragraphs)}段落，{len(doc.tables)}表格")

            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                    self.logger.info(f"已清理临时文件: {temp_file}")
                except Exception as cleanup_error:
                    self.logger.warning(f"清理临时文件失败: {cleanup_error}")

            return text

        except ImportError:
            raise FileProcessingError("缺少python-docx库，无法处理Word文件")
        except FileProcessingError:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                    self.logger.info(f"已清理临时文件: {temp_file}")
                except (OSError, PermissionError) as e:
                    self.logger.warning(f"清理临时文件失败 {temp_file}: {e}")
            # 重新抛出已知的文件处理错误
            raise
        except Exception as e:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                    self.logger.info(f"已清理临时文件: {temp_file}")
                except (OSError, PermissionError) as e:
                    self.logger.warning(f"清理临时文件失败 {temp_file}: {e}")
            # 针对特定的docx错误提供更详细的信息
            error_msg = str(e)
            if "no relationship of type" in error_msg:
                raise FileProcessingError(f"Word文档格式错误：文件可能损坏或不是有效的Word文档格式。详细错误：{error_msg}")
            elif "BadZipFile" in error_msg:
                raise FileProcessingError("文件不是有效的Word文档格式：文件可能损坏")
            else:
                raise FileProcessingError(f"Word文档读取失败: {error_msg}")

    def _read_word_alternative(self, file_path: Path) -> str:
        """备用Word文档读取方法"""
        temp_file = None
        original_file_path = file_path

        try:
            # 检查是否是.doc格式
            if file_path.suffix.lower() == '.doc':
                # 先尝试直接使用antiword读取原始.doc文件
                try:
                    text = self._read_doc_with_antiword(file_path)
                    return text
                except FileProcessingError as e:
                    self.logger.warning(f"备用方法antiword读取失败: {str(e)}")

                # 如果antiword失败，尝试LibreOffice转换
                try:
                    text = self._read_doc_with_libreoffice(file_path)
                    return text
                except FileProcessingError as e:
                    self.logger.warning(f"LibreOffice转换失败: {str(e)}")

                # 如果LibreOffice也失败，尝试doc2docx转换
                try:
                    temp_file = self._convert_doc_to_docx(file_path)
                    file_path = Path(temp_file)
                    self.logger.info(f"备用方法使用转换后的临时文件: {file_path}")
                except FileProcessingError as e:
                    self.logger.warning(f"备用方法转换失败: {str(e)}")

                    # 检测文档创建软件，提供针对性建议
                    office_suite = self._detect_office_suite(original_file_path)

                    if office_suite == 'WPS Office':
                        error_msg = "检测到WPS Office格式的.doc文件，建议：\n1. 使用WPS Office将文件另存为.docx格式\n2. 或使用Microsoft Word打开后另存为.docx格式\n3. 确保文件未损坏"
                    elif office_suite == 'Microsoft Office':
                        error_msg = "Microsoft Word .doc文件处理失败，建议：\n1. 使用Word将文件另存为.docx格式\n2. 检查文件是否损坏\n3. 尝试在其他电脑上打开文件"
                    else:
                        error_msg = f"无法处理.doc文件（检测到：{office_suite}），建议：\n1. 将文件转换为.docx格式\n2. 检查文件是否损坏\n3. 确认文件为有效的Word文档"

                    raise FileProcessingError(error_msg)

            # 方法1：尝试使用mammoth库
            try:
                import mammoth
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
                    if text.strip():
                        self.logger.info("使用mammoth库成功读取Word文档")
                        if temp_file and os.path.exists(temp_file):
                            os.remove(temp_file)
                        return text
            except ImportError:
                self.logger.info("mammoth库未安装，跳过该方法")
            except Exception as e:
                self.logger.warning(f"mammoth库读取失败: {str(e)}")

            # 方法2：尝试使用docx2txt库
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                if text.strip():
                    self.logger.info("使用docx2txt库成功读取Word文档")
                    if temp_file and os.path.exists(temp_file):
                        os.remove(temp_file)
                    return text
            except ImportError:
                self.logger.info("docx2txt库未安装，跳过该方法")
            except Exception as e:
                self.logger.warning(f"docx2txt库读取失败: {str(e)}")

            # 方法3：尝试将文件当作ZIP读取XML
            try:
                import zipfile
                import xml.etree.ElementTree as ET

                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    # 读取主文档内容
                    try:
                        xml_content = zip_ref.read('word/document.xml')
                        root = ET.fromstring(xml_content)

                        # 提取文本内容
                        text_parts = []
                        for elem in root.iter():
                            if elem.text:
                                text_parts.append(elem.text)

                        text = '\n'.join(text_parts)
                        if text.strip():
                            self.logger.info("使用XML解析成功读取Word文档内容")
                            if temp_file and os.path.exists(temp_file):
                                os.remove(temp_file)
                            return text
                    except Exception as e:
                        self.logger.warning(f"XML解析失败: {str(e)}")

            except Exception as e:
                self.logger.warning(f"ZIP/XML方法失败: {str(e)}")

            # 如果所有方法都失败了
            raise FileProcessingError("无法使用任何方法读取Word文档，文件可能损坏或格式不支持")

        except FileProcessingError:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except (OSError, PermissionError) as e:
                    self.logger.warning(f"清理临时文件失败 {temp_file}: {e}")
            raise
        except Exception as e:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except (OSError, PermissionError) as e:
                    self.logger.warning(f"清理临时文件失败 {temp_file}: {e}")
            raise FileProcessingError(f"备用Word文档读取方法失败: {str(e)}")

    def _read_text(self, file_path: Path) -> str:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text.strip():
                raise FileProcessingError("文本文件内容为空")
            
            self.logger.info(f"文本文件读取成功，长度: {len(text)}字符")
            return text
            
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    text = file.read()
                self.logger.info(f"使用GBK编码读取文本文件成功")
                return text
            except (UnicodeDecodeError, LookupError):
                raise FileProcessingError("文本文件编码不支持（尝试了UTF-8、GBK编码）")
        except Exception as e:
            raise FileProcessingError(f"文本文件读取失败: {str(e)}")
    
    def extract_chapters_for_basic_info(self, project_id: int) -> str:
        """
        根据章节识别结果，智能提取包含基础信息的章节文本

        Args:
            project_id: 项目ID

        Returns:
            合并的相关章节文本
        """
        try:
            self.logger.info(f"开始基于章节识别提取项目基础信息文本: project_id={project_id}")

            # 定义包含项目基础信息的章节关键词
            BASIC_INFO_KEYWORDS = [
                "项目概况", "项目基本信息", "招标公告", "采购公告",
                "投标邀请", "竞争性谈判公告", "竞争性磋商公告",
                "投标须知", "供应商须知", "投标人须知", "响应人须知",
                "项目简介", "采购需求概况", "项目背景",
                "单一来源", "询价公告", "谈判邀请", "磋商邀请"
            ]

            # 1. 查询相关章节
            query = """
                SELECT chapter_node_id, title, para_start_idx, para_end_idx, word_count
                FROM tender_document_chapters
                WHERE project_id = ?
                AND level <= 2
                ORDER BY para_start_idx
            """
            chapters = self.db.execute_query(query, [project_id])

            if not chapters:
                self.logger.warning(f"项目 {project_id} 未找到章节数据，将回退到传统方法")
                return None

            # ⭐️ 容错处理：二次排序 + 过滤异常章节
            # 确保章节按段落索引顺序排列（数据库排序可能不可靠）
            chapters = sorted(chapters, key=lambda ch: ch.get('para_start_idx', 0))

            # 过滤索引异常的章节（para_start_idx >= para_end_idx）
            valid_chapters = []
            invalid_count = 0
            skipped_critical_chapters = []  # 记录被跳过的关键章节

            # 定义关键章节（包含项目基本信息）
            CRITICAL_CHAPTER_KEYWORDS = ["公告", "邀请", "项目概况", "项目基本信息", "项目简介"]

            for chapter in chapters:
                start_idx = chapter.get('para_start_idx')
                end_idx = chapter.get('para_end_idx')
                title = chapter.get('title', '')

                # 检查索引有效性
                if end_idx is None or start_idx >= end_idx:
                    # 检查是否为关键章节
                    is_critical = any(kw in title for kw in CRITICAL_CHAPTER_KEYWORDS)
                    if is_critical:
                        skipped_critical_chapters.append(title)

                    self.logger.warning(
                        f"⚠️ 跳过异常章节: '{title}' "
                        f"(索引: {start_idx} >= {end_idx}) "
                        f"{'【关键章节！】' if is_critical else ''}"
                    )
                    invalid_count += 1
                    continue

                valid_chapters.append(chapter)

            if invalid_count > 0:
                self.logger.warning(f"共发现 {invalid_count} 个异常章节，已过滤")

            # ⭐️ 关键改进：如果关键章节被跳过，直接返回None触发智能回退
            if skipped_critical_chapters:
                self.logger.error(
                    f"❌ 关键章节被跳过: {skipped_critical_chapters}，"
                    f"建议使用智能回退读取文档原文"
                )
                return None

            if not valid_chapters:
                self.logger.warning("所有章节数据均异常，回退到传统方法")
                return None

            # 使用过滤后的有效章节列表
            chapters = valid_chapters
            self.logger.info(f"有效章节数量: {len(chapters)}")

            # 2. 根据关键词匹配章节
            selected_chapters = []
            for chapter in chapters:
                title = chapter['title']
                for keyword in BASIC_INFO_KEYWORDS:
                    if keyword in title:
                        selected_chapters.append(chapter)
                        self.logger.info(f"匹配到相关章节: {title} (关键词: {keyword})")
                        break

            # 3. 如果没有匹配到，选择前几个章节作为兜底
            if not selected_chapters:
                selected_chapters = chapters[:3]  # 默认前3个章节
                self.logger.info(f"未匹配到特定关键词，使用前 {len(selected_chapters)} 个章节")
            else:
                # 最多选择前5个匹配的章节
                selected_chapters = selected_chapters[:5]
                self.logger.info(f"共匹配到 {len(selected_chapters)} 个相关章节")

            # 4. 获取文档路径
            project_data = self.db.execute_query(
                "SELECT step1_data FROM tender_projects WHERE project_id = ?",
                [project_id],
                fetch_one=True
            )

            if not project_data or not project_data.get('step1_data'):
                self.logger.warning("无法获取文档路径")
                return None

            step1_data = json.loads(project_data['step1_data'])
            doc_path = step1_data.get('file_path')

            if not doc_path or not Path(doc_path).exists():
                self.logger.warning(f"文档路径无效: {doc_path}")
                return None

            # 5. 提取章节文本
            from docx import Document
            doc = Document(doc_path)
            paragraphs = [p.text for p in doc.paragraphs]

            chapter_texts = []
            total_chars = 0
            MAX_CHARS = 20000  # 最大字符数限制

            for chapter in selected_chapters:
                start_idx = chapter['para_start_idx']
                end_idx = chapter['para_end_idx'] or len(paragraphs)

                # 提取章节文本
                chapter_content = "\n".join(paragraphs[start_idx:end_idx])

                # 检查是否超过限制
                if total_chars + len(chapter_content) > MAX_CHARS:
                    self.logger.info(f"已达到字符数限制 ({MAX_CHARS})，停止添加更多章节")
                    break

                chapter_texts.append(f"【{chapter['title']}】\n{chapter_content}")
                total_chars += len(chapter_content)
                self.logger.info(f"提取章节: {chapter['title']} ({len(chapter_content)} 字符)")

            # 6. 合并文本
            merged_text = "\n\n".join(chapter_texts)

            self.logger.info(f"章节文本提取完成，共 {len(chapter_texts)} 个章节，总计 {len(merged_text)} 字符")

            return merged_text

        except Exception as e:
            self.logger.warning(f"基于章节提取文本失败: {e}，将回退到传统方法")
            import traceback
            self.logger.debug(traceback.format_exc())
            return None

    def extract_basic_info(self, text: str = None, project_id: int = None) -> Dict[str, str]:
        """
        提取基本项目信息 - 使用提示词管理器

        Args:
            text: 文档全文（向后兼容）
            project_id: 项目ID（新方法，优先使用章节识别）

        Returns:
            提取的基础信息字典
        """
        try:
            self.logger.info("开始提取基本项目信息")

            # 优先使用章节识别方法
            text_for_extraction = None
            if project_id:
                try:
                    text_for_extraction = self.extract_chapters_for_basic_info(project_id)
                    if text_for_extraction:
                        self.logger.info(f"✅ 使用章节识别方法，提取文本长度: {len(text_for_extraction)} 字符")
                    else:
                        self.logger.info("章节识别方法未返回文本，回退到传统方法")
                except Exception as e:
                    self.logger.warning(f"章节识别方法失败，回退到传统方法: {e}")

            # 回退方案：使用原有的前3000字符方法
            if not text_for_extraction:
                if text:
                    text_for_extraction = text[:3000] + "..." if len(text) > 3000 else text
                    self.logger.info(f"使用传统方法（前3000字符），实际文本长度: {len(text_for_extraction)} 字符")
                elif project_id:
                    # ⭐️ 新增智能回退：当章节提取失败时，直接读取文档前N段
                    self.logger.warning("章节识别失败，尝试直接读取文档前50段")
                    try:
                        # 获取文档路径
                        project_data = self.db.execute_query(
                            "SELECT step1_data FROM tender_projects WHERE project_id = ?",
                            [project_id],
                            fetch_one=True
                        )

                        if project_data and project_data.get('step1_data'):
                            import json
                            step1_data = json.loads(project_data['step1_data'])
                            doc_path = step1_data.get('file_path')

                            if doc_path and Path(doc_path).exists():
                                from docx import Document
                                doc = Document(doc_path)

                                # 提取前50段（通常包含公告和项目基本信息）
                                first_paras = [p.text for p in doc.paragraphs[:50]]
                                text_for_extraction = '\n'.join(first_paras)

                                self.logger.info(f"✅ 直接读取文档前50段，文本长度: {len(text_for_extraction)} 字符")
                            else:
                                raise TenderInfoExtractionError("文档路径无效")
                        else:
                            raise TenderInfoExtractionError("无法获取文档路径")

                    except Exception as fallback_error:
                        self.logger.error(f"智能回退失败: {fallback_error}")
                        raise TenderInfoExtractionError("未提供文本且所有提取方法均失败")
                else:
                    raise TenderInfoExtractionError("未提供文本且无法使用章节识别方法")

            # 从提示词管理器获取提示词模板
            prompt_template = self.prompt_manager.get_prompt(
                'tender_info',
                'extract_basic_info',
                default="请从招标文档中提取基本信息并以JSON格式返回。"
            )

            # 使用提取的文本生成提示词
            prompt = prompt_template.format(text=text_for_extraction)

            response = self.llm_callback(prompt, "基本信息提取")

            # 添加详细日志记录LLM原始响应
            self.logger.info(f"LLM原始响应（前500字符）: {response[:500]}")

            # 解析JSON响应
            basic_info = self._safe_json_parse(response, "基本信息提取")

            # 记录解析结果
            self.logger.info(f"JSON解析结果: {basic_info}")

            if basic_info:
                # ===== 项目名称验证和修正逻辑 =====
                project_name = basic_info.get('project_name')

                if project_name:
                    # 验证LLM提取的项目名称
                    if self._validate_project_name(project_name):
                        self.logger.info(f"✅ LLM提取的项目名称验证通过: {project_name}")
                    else:
                        self.logger.warning(f"❌ LLM提取的项目名称验证失败，尝试正则提取: {project_name}")
                        # 验证失败，尝试使用正则提取
                        if text_for_extraction:
                            regex_project_name = self._extract_project_name_by_regex(text_for_extraction)
                            if regex_project_name:
                                self.logger.info(f"✅ 正则提取成功，替换项目名称: {regex_project_name}")
                                basic_info['project_name'] = regex_project_name
                            else:
                                self.logger.warning("⚠️ 正则提取也失败，保留原始值但标记为可疑")
                                # 保留原值但添加警告标记
                                basic_info['project_name'] = project_name
                else:
                    # LLM未提取到项目名称，尝试正则提取
                    self.logger.warning("LLM未提取到项目名称，尝试正则提取")
                    if text_for_extraction:
                        regex_project_name = self._extract_project_name_by_regex(text_for_extraction)
                        if regex_project_name:
                            self.logger.info(f"✅ 正则提取成功: {regex_project_name}")
                            basic_info['project_name'] = regex_project_name
                        else:
                            self.logger.warning("⚠️ 正则提取也失败，项目名称为None")
                            basic_info['project_name'] = None

                # 确保返回的字典包含所有期望的字段
                return basic_info
            else:
                self.logger.warning("JSON解析失败或返回空字典，返回默认结构")
                # 返回包含所有期望字段的默认结构，值为None
                return {
                    "project_name": None,
                    "project_number": None,
                    "tender_party": None,
                    "tender_agent": None,
                    "tender_method": None,
                    "tender_location": None,
                    "tender_deadline": None,
                    "winner_count": None,
                    "tenderer_contact_person": None,
                    "tenderer_contact_method": None,
                    "agency_contact_person": None,
                    "agency_contact_method": None
                }

        except Exception as e:
            self.logger.error(f"提取基本信息失败: {e}")
            raise TenderInfoExtractionError(f"基本信息提取失败: {str(e)}")
    
    def extract_qualification_requirements(self, text: str) -> Dict[str, Any]:
        """提取资质要求 - 使用关键字匹配方法"""
        try:
            # 使用新的关键字匹配方法
            return self.extract_qualification_requirements_by_keywords(text)

        except Exception as e:
            self.logger.error(f"提取资质要求失败: {e}")
            return {}

    def extract_supplier_eligibility_checklist(self, text: str) -> List[Dict[str, Any]]:
        """
        提取18条供应商资格要求清单（关键词匹配）
        返回固定18项清单，每项标记found=True/False

        返回格式：
        [
            {
                "checklist_id": 1,
                "checklist_name": "营业执照信息",
                "found": True,
                "requirements": [{
                    "summary": "需要提供营业执照",
                    "detail": "供应商须提供有效的营业执照...",
                    "constraint_type": "mandatory",
                    "source_location": "第二章 投标人资格要求",
                    "extraction_confidence": 0.85
                }]
            },
            ...
        ]
        """
        try:
            self.logger.info("开始提取18条供应商资格要求清单")

            # 调用关键词匹配方法
            qual_results = self.extract_qualification_requirements_by_keywords(text)
            qualifications = qual_results.get('qualifications', {})

            # 定义18项清单与关键词的映射关系
            checklist_mapping = [
                {
                    "checklist_id": 1,
                    "checklist_name": "营业执照信息",
                    "qual_keys": ["business_license"]  # 对应关键词类型
                },
                {
                    "checklist_id": 2,
                    "checklist_name": "财务要求",
                    "qual_keys": ["audit_report"]
                },
                {
                    "checklist_id": 3,
                    "checklist_name": "依法纳税",
                    "qual_keys": ["tax_compliance"]
                },
                {
                    "checklist_id": 4,
                    "checklist_name": "缴纳社保",
                    "qual_keys": ["social_security"]
                },
                {
                    "checklist_id": 5,
                    "checklist_name": "失信被执行人",
                    "qual_keys": ["dishonest_executor"]
                },
                {
                    "checklist_id": 6,
                    "checklist_name": "政府采购严重违法失信记录（信用中国）",
                    "qual_keys": ["gov_procurement_creditchina"]
                },
                {
                    "checklist_id": 6.5,
                    "checklist_name": "政府采购严重违法失信行为信息记录（政府采购网）",
                    "qual_keys": ["gov_procurement_ccgp"]
                },
                {
                    "checklist_id": 7,
                    "checklist_name": "信用中国重大税收违法",
                    "qual_keys": ["tax_violation_check"]
                },
                {
                    "checklist_id": 8,
                    "checklist_name": "采购人黑名单",
                    "qual_keys": ["purchaser_blacklist"]
                },
                {
                    "checklist_id": 9,
                    "checklist_name": "承诺函",
                    "qual_keys": ["commitment_letter"]
                },
                {
                    "checklist_id": 10,
                    "checklist_name": "营业办公场所房产证明",
                    "qual_keys": ["property_certificate"]
                },
                {
                    "checklist_id": 11,
                    "checklist_name": "业绩案例要求",
                    "qual_keys": ["project_performance"]
                },
                {
                    "checklist_id": 12,
                    "checklist_name": "保证金要求",
                    "qual_keys": ["deposit_requirement"]
                },
                {
                    "checklist_id": 13,
                    "checklist_name": "增值电信业务许可证",
                    "qual_keys": ["value_added_telecom_license"]
                },
                {
                    "checklist_id": 14,
                    "checklist_name": "基础电信业务许可证",
                    "qual_keys": ["basic_telecom_license"]
                },
                {
                    "checklist_id": 15,
                    "checklist_name": "ISO9001质量管理体系认证",
                    "qual_keys": ["iso9001"]
                },
                {
                    "checklist_id": 16,
                    "checklist_name": "ISO20000信息技术服务管理体系认证",
                    "qual_keys": ["iso20000"]
                },
                {
                    "checklist_id": 17,
                    "checklist_name": "ISO27001信息安全管理体系认证",
                    "qual_keys": ["iso27001"]
                },
                {
                    "checklist_id": 18,
                    "checklist_name": "等保三级认证",
                    "qual_keys": ["level_protection"]
                }
            ]

            # 构建18项清单结果
            checklist_results = []

            for item in checklist_mapping:
                checklist_item = {
                    "checklist_id": item["checklist_id"],
                    "checklist_name": item["checklist_name"],
                    "found": False,
                    "requirements": []
                }

                # 检查是否找到相关资质
                for qual_key in item["qual_keys"]:
                    if qual_key in qualifications:
                        qual_data = qualifications[qual_key]
                        checklist_item["found"] = True

                        # 构建requirement对象
                        requirement = {
                            "summary": f"需要提供{item['checklist_name']}",
                            "detail": qual_data.get('description', ''),
                            "constraint_type": "mandatory",
                            "source_location": ""  # 关键词匹配无法确定具体位置
                        }

                        checklist_item["requirements"].append(requirement)

                checklist_results.append(checklist_item)

            # 统计
            found_count = sum(1 for item in checklist_results if item["found"])
            self.logger.info(f"18条清单提取完成：找到 {found_count} 项，未找到 {18 - found_count} 项")

            return checklist_results

        except Exception as e:
            self.logger.error(f"提取18条供应商资格要求清单失败: {e}")
            # 返回空清单
            return [
                {
                    "checklist_id": i,
                    "checklist_name": ["营业执照信息", "财务要求", "依法纳税", "缴纳社保",
                                      "失信被执行人", "政府采购严重违法失信记录",
                                      "信用中国重大税收违法", "采购人黑名单", "承诺函",
                                      "营业办公场所房产证明", "业绩案例要求", "保证金要求",
                                      "增值电信业务许可证", "基础电信业务许可证",
                                      "ISO9001质量管理体系认证", "ISO20000信息技术服务管理体系认证",
                                      "ISO27001信息安全管理体系认证", "等保三级认证"][i-1],
                    "found": False,
                    "requirements": []
                }
                for i in range(1, 19)
            ]

    def extract_technical_scoring(self, text: str) -> Dict[str, Any]:
        """提取技术评分标准 - 使用提示词管理器"""
        try:
            self.logger.info("开始提取技术评分标准")

            # 从提示词管理器获取提示词模板
            prompt_template = self.prompt_manager.get_prompt(
                'tender_info',
                'extract_technical_scoring',
                default="请从招标文档中提取技术评分标准并以JSON格式返回。"
            )

            # 使用模板生成提示词
            prompt = prompt_template.format(text=text[:4000] + "...")

            response = self.llm_callback(prompt, "技术评分提取")

            scoring_info = self._safe_json_parse(response, "技术评分提取")
            if scoring_info:
                return scoring_info
            else:
                return {}

        except Exception as e:
            self.logger.error(f"提取技术评分标准失败: {e}")
            return {}
    
    def process_document(self, file_path: str, file_info: Dict[str, Any] = None) -> Dict[str, Any]:
        """处理完整文档提取"""
        try:
            self.logger.info(f"开始处理文档: {file_path}")

            # 准备文件信息
            file_path_obj = Path(file_path)
            if not file_info:
                file_info = {}

            # 补充文件信息
            if not file_info.get('file_path'):
                file_info['file_path'] = str(file_path)
            if not file_info.get('original_filename'):
                file_info['original_filename'] = file_path_obj.name
            if not file_info.get('file_size') and file_path_obj.exists():
                file_info['file_size'] = file_path_obj.stat().st_size
            if not file_info.get('file_hash') and file_path_obj.exists():
                # 计算文件哈希值
                import hashlib
                with open(file_path, 'rb') as f:
                    file_info['file_hash'] = hashlib.md5(f.read()).hexdigest()

            # 读取文档
            text = self.read_document(file_path)

            # 提取各项信息
            basic_info = self.extract_basic_info(text)
            qualification_info = self.extract_qualification_requirements(text)
            scoring_info = self.extract_technical_scoring(text)

            # 合并结果
            result = {
                **basic_info,
                **qualification_info,
                **scoring_info,
                'extraction_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'file_path': str(file_path)
            }

            # 保存到数据库
            project_id = self.save_to_database(result, file_info)

            # 添加项目ID到返回结果
            result['project_id'] = project_id

            # 保存到配置文件（向后兼容）
            self.save_to_config(result)

            self.logger.info(f"文档处理完成，项目ID: {project_id}")
            return result

        except Exception as e:
            self.logger.error(f"文档处理失败: {e}")
            raise TenderInfoExtractionError(f"文档处理失败: {str(e)}")
    
    def save_to_database(self, data: Dict[str, Any], file_info: Dict[str, Any] = None) -> int:
        """保存数据到数据库"""
        try:
            # 检查是否已存在相同项目（防止重复创建）
            project_name = data.get('project_name', '')
            project_number = data.get('project_number', '')

            if project_name and project_number:
                check_sql = """
                    SELECT project_id FROM tender_projects
                    WHERE project_name = ? AND project_number = ?
                    LIMIT 1
                """
                with self.db.get_connection() as conn:
                    cursor = conn.execute(check_sql, (project_name, project_number))
                    existing = cursor.fetchone()

                    if existing:
                        existing_id = existing[0]
                        self.logger.info(f"项目已存在（名称: {project_name}, 编号: {project_number}），使用现有项目ID: {existing_id}")
                        return existing_id

            # 准备文件信息
            file_info = file_info or {}
            file_path = file_info.get('file_path', data.get('file_path', ''))
            original_filename = file_info.get('original_filename', Path(file_path).name if file_path else '')
            file_size = file_info.get('file_size', 0)
            file_hash = file_info.get('file_hash', '')

            with self.db.get_connection() as conn:
                # 创建投标项目记录
                project_sql = """
                    INSERT INTO tender_projects (
                        project_name, project_number, tenderer, agency, bidding_method,
                        bidding_location, bidding_time, winner_count, tender_document_path,
                        original_filename, file_size, file_hash, extraction_time,
                        extraction_summary, total_score, items_count,
                        tenderer_contact_person, tenderer_contact_method,
                        agency_contact_person, agency_contact_method,
                        status, created_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'analyzed', ?)
                """

                project_params = (
                    data.get('project_name', ''),
                    data.get('project_number', ''),
                    data.get('tenderer', ''),
                    data.get('agency', ''),
                    data.get('bidding_method', ''),
                    data.get('bidding_location', ''),
                    data.get('bidding_time', ''),
                    data.get('winner_count', 0),
                    file_path,
                    original_filename,
                    file_size,
                    file_hash,
                    data.get('extraction_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                    data.get('extraction_summary', ''),
                    data.get('total_score', ''),
                    data.get('items_count', 0),
                    data.get('tenderer_contact_person', ''),
                    data.get('tenderer_contact_method', ''),
                    data.get('agency_contact_person', ''),
                    data.get('agency_contact_method', ''),
                    datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                )

                cursor = conn.execute(project_sql, project_params)
                project_id = cursor.lastrowid

                # 保存资质要求
                self._save_qualification_requirements(conn, project_id, data)

                # 保存技术评分标准
                self._save_scoring_criteria(conn, project_id, data)

                conn.commit()
                self.logger.info(f"数据已保存到数据库，项目ID: {project_id}")

                return project_id

        except Exception as e:
            self.logger.error(f"保存到数据库失败: {e}")
            raise TenderInfoExtractionError(f"保存到数据库失败: {str(e)}")

    def _save_qualification_requirements(self, conn, project_id: int, data: Dict[str, Any]) -> None:
        """保存资质要求到数据库"""
        # 查找所有资质要求相关的字段
        qual_types = set()
        for key in data.keys():
            if key.endswith('_required') or key.endswith('_description'):
                qual_type = key.replace('_required', '').replace('_description', '')
                qual_types.add(qual_type)

        # 为每种资质类型创建记录
        for qual_type in qual_types:
            required_key = f"{qual_type}_required"
            desc_key = f"{qual_type}_description"

            is_required = data.get(required_key, False)
            description = data.get(desc_key, '')

            # 只保存有实际内容的资质要求
            if is_required or description:
                sql = """
                    INSERT INTO tender_qualification_requirements (
                        project_id, requirement_type, is_required, requirement_description
                    ) VALUES (?, ?, ?, ?)
                """
                params = (project_id, qual_type, is_required, description)
                conn.execute(sql, params)

    def _save_scoring_criteria(self, conn, project_id: int, data: Dict[str, Any]) -> None:
        """保存技术评分标准到数据库"""
        items_count = data.get('items_count', 0)
        if not items_count:
            return

        try:
            items_count = int(items_count)
        except (ValueError, TypeError):
            items_count = 0

        # 保存每个评分项目
        for i in range(1, items_count + 1):
            item_name = data.get(f'item_{i}_name', '')
            item_weight = data.get(f'item_{i}_weight', '')
            item_criteria = data.get(f'item_{i}_criteria', '')
            item_source = data.get(f'item_{i}_source', '')

            # 只保存有实际内容的评分项目
            if item_name or item_criteria:
                sql = """
                    INSERT INTO tender_scoring_criteria (
                        project_id, item_name, item_weight, criteria_description,
                        source_location, item_order
                    ) VALUES (?, ?, ?, ?, ?, ?)
                """
                params = (project_id, item_name, item_weight, item_criteria, item_source, i)
                conn.execute(sql, params)

    def save_to_config(self, data: Dict[str, Any]) -> None:
        """保存数据到配置文件（保持向后兼容）"""
        try:
            config_file = self.config.get_path('config') / 'tender_config.ini'

            config = configparser.ConfigParser()

            # 项目基本信息
            config['PROJECT_INFO'] = {}
            for key in ['project_name', 'project_number', 'extraction_time',
                       'tenderer', 'agency', 'bidding_method', 'bidding_location',
                       'bidding_time', 'winner_count',
                       'tenderer_contact_person', 'tenderer_contact_method',
                       'agency_contact_person', 'agency_contact_method']:
                if key in data:
                    config['PROJECT_INFO'][key] = str(data[key])

            # 资质要求
            config['QUALIFICATION_REQUIREMENTS'] = {}
            qual_keys = [k for k in data.keys() if k.endswith('_required') or k.endswith('_description')]
            for key in qual_keys:
                config['QUALIFICATION_REQUIREMENTS'][key] = str(data[key])

            # 技术评分
            config['TECHNICAL_SCORING'] = {}
            scoring_keys = [k for k in data.keys() if k.startswith('total_score') or
                           k.startswith('extraction_summary') or k.startswith('items_count') or
                           k.startswith('item_')]
            for key in scoring_keys:
                config['TECHNICAL_SCORING'][key] = str(data[key])

            # 保存文件
            with open(config_file, 'w', encoding='utf-8') as f:
                config.write(f)

            self.logger.info(f"配置已保存到: {config_file}")

        except Exception as e:
            self.logger.error(f"保存配置失败: {e}")
            # 配置文件保存失败不影响主流程
            pass

if __name__ == "__main__":
    # 测试代码
    extractor = TenderInfoExtractor()
    print("招标信息提取器初始化完成")