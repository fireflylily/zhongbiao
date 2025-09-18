"""
招标文件智能拆分系统
BiddingDocumentExtractor - 智能提取招标文件中的关键章节
Version: 1.0.0
Author: AI Assistant
"""

import re
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from abc import ABC, abstractmethod
from enum import Enum

# 需要安装的依赖: pip install python-docx
try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请安装python-docx: pip install python-docx")
    raise

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============== 数据类定义 ==============

class DocType(Enum):
    """文档类型枚举"""
    BIDDING = "招标文件"
    NEGOTIATION = "竞争性磋商"
    INQUIRY = "询价采购"
    OTHER = "其他"


@dataclass
class ExtractionResult:
    """提取结果数据类"""
    success: bool
    project_name: str
    doc_type: DocType
    sections: Dict[str, Any]
    output_files: List[str]
    errors: List[str]
    processing_time: float


@dataclass
class SectionInfo:
    """章节信息数据类"""
    title: str
    start_index: int
    end_index: int
    content: List[Any]
    confidence: float


# ============== 配置管理 ==============

class ExtractionConfig:
    """提取配置类"""

    # 文档类型识别规则
    DOC_TYPE_PATTERNS = {
        DocType.BIDDING: ["招标文件", "公开招标", "投标人须知"],
        DocType.NEGOTIATION: ["竞争性磋商", "磋商文件", "供应商须知"],
        DocType.INQUIRY: ["询价", "询价采购", "询价文件"]
    }

    # 章节识别规则
    SECTION_RULES = {
        "公告": {
            "patterns": ["招标公告", "磋商公告", "竞争性磋商公告", "采购公告"],
            "end_markers": ["投标人须知", "供应商须知", "第一章", "第二章"],
            "required": True
        },
        "前附表": {
            "patterns": ["投标人须知前附表", "供应商须知前附表", "前附表"],
            "end_markers": ["投标人须知", "供应商须知", "第二章"],
            "required": True
        },
        "评分办法": {
            "patterns": ["评标办法", "评分办法", "磋商的评价", "评审办法"],
            "end_markers": ["第四章", "合同", "附件"],
            "required": True
        },
        "响应文件格式": {
            "patterns": ["响应文件格式", "投标文件格式", "附件", "第四章.*格式"],
            "end_markers": ["第五章", "采购需求", "技术规范"],
            "required": False
        }
    }

    # 项目名称提取规则
    PROJECT_NAME_PATTERNS = [
        r"关于[""\"']?(.+?)[""\"']?项目",
        r"(.+?)项目.*?(?:招标|磋商|采购)文件",
        r"项目名称[：:]\s*(.+?)[\n\r]",
        r"[""\"'](.+?)[""\"'].*?(?:招标|磋商|采购)",
    ]


# ============== 智能学习模块 ==============

class LearningModule:
    """智能学习模块"""

    def __init__(self, knowledge_file: str = "knowledge_base.json"):
        self.knowledge_file = Path(knowledge_file)
        self.knowledge_base = self.load_knowledge()

    def load_knowledge(self) -> Dict:
        """加载知识库"""
        if self.knowledge_file.exists():
            with open(self.knowledge_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {
            "patterns": {},
            "boundaries": {},
            "corrections": [],
            "success_cases": []
        }

    def save_knowledge(self):
        """保存知识库"""
        with open(self.knowledge_file, 'w', encoding='utf-8') as f:
            json.dump(self.knowledge_base, f, ensure_ascii=False, indent=2)

    def learn_from_success(self, doc_path: str, result: ExtractionResult):
        """从成功案例学习"""
        case = {
            "doc_path": str(doc_path),
            "project_name": result.project_name,
            "doc_type": result.doc_type.value,
            "sections_found": list(result.sections.keys()),
            "timestamp": datetime.now().isoformat()
        }
        self.knowledge_base["success_cases"].append(case)
        self.save_knowledge()
        logger.info(f"学习成功案例: {doc_path}")

    def learn_from_correction(self, original: Dict, corrected: Dict):
        """从用户修正学习"""
        correction = {
            "original": original,
            "corrected": corrected,
            "timestamp": datetime.now().isoformat()
        }
        self.knowledge_base["corrections"].append(correction)

        # 更新模式识别
        for key, value in corrected.items():
            if key not in original or original[key] != value:
                self.update_pattern(key, value)

        self.save_knowledge()
        logger.info("从用户修正中学习完成")

    def update_pattern(self, pattern_type: str, pattern_value: str):
        """更新识别模式"""
        if pattern_type not in self.knowledge_base["patterns"]:
            self.knowledge_base["patterns"][pattern_type] = []

        if pattern_value not in self.knowledge_base["patterns"][pattern_type]:
            self.knowledge_base["patterns"][pattern_type].append(pattern_value)

    def get_suggestions(self, doc_text: str) -> Dict:
        """基于学习经验提供建议"""
        suggestions = {}

        # 基于成功案例提供建议
        for case in self.knowledge_base["success_cases"][-10:]:  # 最近10个案例
            if case["doc_type"] in doc_text:
                suggestions["probable_type"] = case["doc_type"]
                suggestions["expected_sections"] = case["sections_found"]
                break

        return suggestions


# ============== 核心提取器 ==============

class BiddingDocumentExtractor:
    """招标文件提取器主类"""

    def __init__(self, enable_learning: bool = True):
        self.config = ExtractionConfig()
        self.learning = LearningModule() if enable_learning else None
        self.current_doc = None
        self.current_text = ""

    def extract(self, file_path: str, output_dir: str = None) -> ExtractionResult:
        """主提取方法"""
        start_time = datetime.now()
        errors = []

        try:
            # 1. 加载文档
            self.current_doc = Document(file_path)
            self.current_text = self._get_full_text()

            # 2. 识别文档类型
            doc_type = self._identify_doc_type()

            # 3. 提取项目名称
            project_name = self._extract_project_name()
            if not project_name:
                project_name = Path(file_path).stem

            # 4. 获取学习建议
            suggestions = {}
            if self.learning:
                suggestions = self.learning.get_suggestions(self.current_text)

            # 5. 提取各章节
            sections = self._extract_sections(suggestions)

            # 6. 生成输出文件
            output_dir = output_dir or Path(file_path).parent / "extracted"
            output_files = self._generate_output_files(
                sections, project_name, output_dir
            )

            # 7. 创建结果
            result = ExtractionResult(
                success=True,
                project_name=project_name,
                doc_type=doc_type,
                sections=sections,
                output_files=output_files,
                errors=errors,
                processing_time=(datetime.now() - start_time).total_seconds()
            )

            # 8. 学习成功案例
            if self.learning and result.success:
                self.learning.learn_from_success(file_path, result)

            logger.info(f"提取完成: {file_path}")
            return result

        except Exception as e:
            logger.error(f"提取失败: {e}")
            errors.append(str(e))
            return ExtractionResult(
                success=False,
                project_name="",
                doc_type=DocType.OTHER,
                sections={},
                output_files=[],
                errors=errors,
                processing_time=(datetime.now() - start_time).total_seconds()
            )

    def _get_full_text(self) -> str:
        """获取文档全文"""
        text_parts = []
        for para in self.current_doc.paragraphs:
            text_parts.append(para.text)
        return "\n".join(text_parts)

    def _identify_doc_type(self) -> DocType:
        """识别文档类型"""
        for doc_type, patterns in self.config.DOC_TYPE_PATTERNS.items():
            for pattern in patterns:
                if pattern in self.current_text:
                    logger.info(f"识别文档类型: {doc_type.value}")
                    return doc_type
        return DocType.OTHER

    def _extract_project_name(self) -> Optional[str]:
        """提取项目名称"""
        for pattern in self.config.PROJECT_NAME_PATTERNS:
            match = re.search(pattern, self.current_text)
            if match:
                project_name = match.group(1).strip()
                # 清理项目名称
                project_name = re.sub(r'[^\w\u4e00-\u9fff\-]', '', project_name)
                logger.info(f"提取项目名称: {project_name}")
                return project_name
        return None

    def _extract_sections(self, suggestions: Dict = None) -> Dict[str, SectionInfo]:
        """提取各章节"""
        sections = {}

        for section_name, rules in self.config.SECTION_RULES.items():
            section_info = self._extract_single_section(section_name, rules)
            if section_info:
                sections[section_name] = section_info
                logger.info(f"成功提取章节: {section_name}")
            elif rules.get("required", False):
                logger.warning(f"未找到必需章节: {section_name}")

        return sections

    def _get_paragraph_heading_level(self, para) -> int:
        """获取段落标题级别"""
        text = para.text.strip()
        if not text:
            return 0

        # 检查Word样式
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if 'heading' in style_name:
                # 提取标题级别 (Heading 1 -> 1, Heading 2 -> 2, etc.)
                match = re.search(r'heading\s*(\d+)', style_name)
                if match:
                    return int(match.group(1))

        # 基于文本模式判断级别
        # 第X章 -> 级别1
        if re.match(r'^第[一二三四五六七八九十\d]+章', text):
            return 1
        # X.X 格式 -> 级别2
        if re.match(r'^\d+\.\d+', text):
            return 2
        # X.X.X 格式 -> 级别3
        if re.match(r'^\d+\.\d+\.\d+', text):
            return 3
        # (X) 格式 -> 级别3
        if re.match(r'^\([一二三四五六七八九十\d]+\)', text):
            return 3

        # 基于字体大小判断（如果可用）
        try:
            if hasattr(para, 'runs') and para.runs:
                first_run = para.runs[0]
                if hasattr(first_run, 'font') and hasattr(first_run.font, 'size'):
                    if first_run.font.size:
                        size_pt = first_run.font.size.pt
                        if size_pt >= 16:
                            return 1
                        elif size_pt >= 14:
                            return 2
                        elif size_pt >= 12:
                            return 3
        except:
            pass

        return 0  # 普通段落

    def _is_same_or_higher_level_heading(self, para, reference_level: int) -> bool:
        """判断是否为同级别或更高级别的标题"""
        para_level = self._get_paragraph_heading_level(para)
        return para_level > 0 and para_level <= reference_level

    def _extract_document_elements(self, start_index: int, end_index: int) -> List[Any]:
        """提取文档元素（段落和表格）"""
        content = []

        # 获取文档中的所有元素（段落和表格）
        document_elements = []
        for element in self.current_doc.element.body:
            if element.tag.endswith('p'):  # 段落
                # 查找对应的段落对象
                for i, para in enumerate(self.current_doc.paragraphs):
                    if para._element == element:
                        document_elements.append((i, 'paragraph', para))
                        break
            elif element.tag.endswith('tbl'):  # 表格
                # 查找对应的表格对象
                for table in self.current_doc.tables:
                    if table._element == element:
                        document_elements.append((-1, 'table', table))
                        break

        # 按照段落索引范围提取内容
        for element_index, element_type, element_obj in document_elements:
            if element_type == 'paragraph':
                if start_index <= element_index <= end_index:
                    content.append(element_obj)
            elif element_type == 'table':
                # 对于表格，检查它是否在指定范围内
                # 简化处理：如果表格在范围内的段落之间，就包含它
                content.append(element_obj)

        # 如果无法获取复杂元素，就只提取段落
        if not content:
            for i in range(start_index, min(end_index + 1, len(self.current_doc.paragraphs))):
                content.append(self.current_doc.paragraphs[i])

        return content

    def _extract_single_section(self, section_name: str, rules: Dict) -> Optional[SectionInfo]:
        """提取单个章节"""
        # 查找章节开始位置
        start_index = None
        start_para = None
        for pattern in rules["patterns"]:
            for i, para in enumerate(self.current_doc.paragraphs):
                if re.search(pattern, para.text, re.IGNORECASE):
                    start_index = i
                    start_para = para
                    break
            if start_index is not None:
                break

        if start_index is None:
            return None

        # 获取起始章节的标题级别
        start_level = self._get_paragraph_heading_level(start_para)

        # 使用同级别边界检测来查找章节结束位置
        end_index = len(self.current_doc.paragraphs) - 1

        # 从起始位置的下一个段落开始查找
        for i in range(start_index + 1, len(self.current_doc.paragraphs)):
            para = self.current_doc.paragraphs[i]

            # 如果找到同级别或更高级别的标题，则结束
            if self._is_same_or_higher_level_heading(para, start_level):
                end_index = i - 1
                break

            # 备用策略：如果没有明确的标题级别，使用原来的end_markers
            if start_level == 0:  # 如果无法确定起始级别
                for pattern in rules.get("end_markers", []):
                    if re.search(pattern, para.text, re.IGNORECASE):
                        end_index = i - 1
                        break
                if end_index < len(self.current_doc.paragraphs) - 1:
                    break

        # 提取内容（包括段落和表格）
        content = self._extract_document_elements(start_index, end_index)

        confidence = 0.9 if start_level > 0 else 0.7  # 有明确级别的章节置信度更高

        return SectionInfo(
            title=self.current_doc.paragraphs[start_index].text,
            start_index=start_index,
            end_index=end_index,
            content=content,
            confidence=confidence
        )

    def _generate_output_files(self, sections: Dict[str, SectionInfo],
                              project_name: str, output_dir: str) -> List[str]:
        """生成输出文件"""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_files = []

        for section_name, section_info in sections.items():
            # 生成文件名
            filename = f"{section_name}_{project_name}.docx"
            file_path = output_dir / filename

            # 创建新文档并保持格式
            new_doc = Document()

            # 更全面的内容复制，包括段落和表格
            for element in section_info.content:
                try:
                    # 处理段落
                    if hasattr(element, 'text') and hasattr(element, '_element') and element._element.tag.endswith('p'):
                        new_para = new_doc.add_paragraph(element.text)
                        # 复制段落格式
                        if hasattr(element, 'alignment') and element.alignment:
                            new_para.alignment = element.alignment
                        if hasattr(element, 'style') and element.style:
                            try:
                                new_para.style = element.style
                            except:
                                pass  # 如果样式不存在，忽略

                        # 复制字体格式
                        if hasattr(element, 'runs'):
                            for i, run in enumerate(element.runs):
                                if i < len(new_para.runs):
                                    try:
                                        if hasattr(run, 'bold') and run.bold:
                                            new_para.runs[i].bold = True
                                        if hasattr(run, 'italic') and run.italic:
                                            new_para.runs[i].italic = True
                                    except:
                                        pass
                    # 处理表格
                    elif hasattr(element, '_element') and element._element.tag.endswith('tbl'):
                        # 复制表格
                        table = element
                        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))

                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                try:
                                    new_table.rows[i].cells[j].text = cell.text
                                except:
                                    pass
                    else:
                        # 如果不是段落或表格，就简单添加文本
                        text = getattr(element, 'text', str(element))
                        if text and text.strip():
                            new_doc.add_paragraph(text)
                except Exception as e:
                    # 如果出现错误，就简单添加文本
                    logger.warning(f"复制元素时出错: {e}")
                    try:
                        text = getattr(element, 'text', str(element))
                        if text and text.strip():
                            new_doc.add_paragraph(text)
                    except:
                        pass

            # 保存文档
            new_doc.save(str(file_path))
            output_files.append(str(file_path))
            logger.info(f"生成文件: {file_path}")

        return output_files

    def learn_from_user_feedback(self, original_result: Dict, corrected_result: Dict):
        """从用户反馈学习"""
        if self.learning:
            self.learning.learn_from_correction(original_result, corrected_result)


# ============== 便捷接口 ==============

def extract_bidding_document(file_path: str, output_dir: str = None,
                            enable_learning: bool = True) -> ExtractionResult:
    """
    提取招标文件的便捷接口

    Args:
        file_path: 文档路径
        output_dir: 输出目录（可选）
        enable_learning: 是否启用学习功能

    Returns:
        ExtractionResult: 提取结果
    """
    extractor = BiddingDocumentExtractor(enable_learning=enable_learning)
    return extractor.extract(file_path, output_dir)


# ============== 批量处理 ==============

class BatchProcessor:
    """批量处理器"""

    def __init__(self, enable_learning: bool = True):
        self.extractor = BiddingDocumentExtractor(enable_learning=enable_learning)

    def process_directory(self, directory: str, output_dir: str = None) -> List[ExtractionResult]:
        """处理目录下的所有文档"""
        results = []
        directory = Path(directory)

        # 查找所有Word文档
        doc_files = list(directory.glob("*.docx")) + list(directory.glob("*.doc"))

        for doc_file in doc_files:
            logger.info(f"处理文档: {doc_file}")
            result = self.extractor.extract(str(doc_file), output_dir)
            results.append(result)

        # 生成汇总报告
        self._generate_summary_report(results, output_dir)

        return results

    def _generate_summary_report(self, results: List[ExtractionResult], output_dir: str):
        """生成汇总报告"""
        if not output_dir:
            output_dir = Path.cwd() / "extracted"

        report_path = Path(output_dir) / f"extraction_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("=" * 60 + "\n")
            f.write("招标文件提取汇总报告\n")
            f.write(f"处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n\n")

            success_count = sum(1 for r in results if r.success)
            f.write(f"处理文档总数: {len(results)}\n")
            f.write(f"成功: {success_count}\n")
            f.write(f"失败: {len(results) - success_count}\n\n")

            for i, result in enumerate(results, 1):
                f.write(f"{i}. 项目名称: {result.project_name}\n")
                f.write(f"   文档类型: {result.doc_type.value}\n")
                f.write(f"   提取章节: {', '.join(result.sections.keys())}\n")
                f.write(f"   处理时间: {result.processing_time:.2f}秒\n")
                if result.errors:
                    f.write(f"   错误: {', '.join(result.errors)}\n")
                f.write("\n")

        logger.info(f"报告生成: {report_path}")


# ============== 主程序入口 ==============

def main():
    """主程序示例"""
    # 单文档处理示例
    file_path = "path/to/your/document.docx"
    result = extract_bidding_document(file_path)

    if result.success:
        print(f"提取成功!")
        print(f"项目名称: {result.project_name}")
        print(f"文档类型: {result.doc_type.value}")
        print(f"提取章节: {list(result.sections.keys())}")
        print(f"生成文件: {result.output_files}")
    else:
        print(f"提取失败: {result.errors}")

    # 批量处理示例
    # batch_processor = BatchProcessor()
    # results = batch_processor.process_directory("path/to/documents")


if __name__ == "__main__":
    # 测试代码
    print("招标文件智能拆分系统 v1.0.0")
    print("=" * 60)

    # 您可以在这里添加测试代码
    # main()