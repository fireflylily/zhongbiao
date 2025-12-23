#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查段落63的实际内容
"""

import sys
from pathlib import Path

# 添加项目路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / 'ai_tender_system'))

from docx import Document

def main():
    doc_path = "ai_tender_system/data/uploads/tender_processing/2025/12/20251221_180214_中邮保险手机号实名认证服务采购项目竞争性磋商采购文件_fcfdcec9.docx"
    doc = Document(doc_path)

    print("=" * 100)
    print("检查段落63的详细信息")
    print("=" * 100)

    para = doc.paragraphs[63]
    print(f"\n段落63:")
    print(f"  文本: '{para.text}'")
    print(f"  文本长度: {len(para.text)}")
    print(f"  样式: {para.style.name if para.style else 'None'}")

    # 检查文本的字符编码
    print(f"\n  字符分析:")
    for i, char in enumerate(para.text):
        print(f"    {i}: '{char}' (U+{ord(char):04X})")

    # 检查大纲级别
    if hasattr(para, '_element') and hasattr(para._element, 'pPr') and para._element.pPr is not None:
        outline_lvl = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
        if outline_lvl is not None:
            outline_level = int(outline_lvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
            print(f"\n  大纲级别: {outline_level}")

    print("\n" + "=" * 100)
    print("段落92（对比）:")
    print("=" * 100)

    para92 = doc.paragraphs[92]
    print(f"\n段落92:")
    print(f"  文本: '{para92.text}'")
    print(f"  文本长度: {len(para92.text)}")
    print(f"  样式: {para92.style.name if para92.style else 'None'}")

    print(f"\n  字符分析:")
    for i, char in enumerate(para92.text):
        print(f"    {i}: '{char}' (U+{ord(char):04X})")

    print("\n" + "=" * 100)
    print("对比分析")
    print("=" * 100)

    print(f"\n段落63 vs 段落92:")
    print(f"  段落63: '{para.text}'")
    print(f"  段落92: '{para92.text}'")
    print(f"  相同: {para.text == para92.text}")

    # 去除空格后对比
    clean63 = para.text.replace(' ', '').replace('\t', '')
    clean92 = para92.text.replace(' ', '').replace('\t', '')
    print(f"\n去除空格后:")
    print(f"  段落63: '{clean63}'")
    print(f"  段落92: '{clean92}'")
    print(f"  相同: {clean63 == clean92}")

    print()

if __name__ == '__main__':
    main()
