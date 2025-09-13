#!/usr/bin/env python3
"""
调试脚本：检查商务应答文档中的电子邮件字段格式
"""
import sys
import os
from pathlib import Path

# 添加项目路径到sys.path
script_dir = Path(__file__).parent
project_root = script_dir / 'ai_tender_system'
sys.path.insert(0, str(project_root))

try:
    from docx import Document
    
    # 找到最新的输出文档
    output_dir = project_root / 'data' / 'outputs'
    latest_file = None
    latest_time = 0
    
    for file in output_dir.glob('*.docx'):
        if file.stat().st_mtime > latest_time:
            latest_time = file.stat().st_mtime
            latest_file = file
    
    if not latest_file:
        print("❌ 没有找到输出文档")
        sys.exit(1)
    
    print(f"🔍 检查文档: {latest_file.name}")
    
    # 读取文档
    doc = Document(str(latest_file))
    
    # 搜索包含邮件相关关键词的段落
    email_keywords = ['邮件', '邮箱', 'email', 'Email', 'EMAIL']
    
    print(f"\n📧 搜索包含邮件关键词的段落:")
    found_email_paragraphs = False
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if any(keyword in text for keyword in email_keywords):
            print(f"段落 #{i}: '{text}'")
            # 显示段落的详细结构
            if paragraph.runs:
                print(f"  Run结构:")
                for j, run in enumerate(paragraph.runs):
                    print(f"    Run {j}: '{run.text}' (字体: {run.font.name})")
            found_email_paragraphs = True
    
    if not found_email_paragraphs:
        print("❌ 文档中没有找到包含邮件关键词的段落")
        
        # 显示所有段落用于调试
        print(f"\n📄 文档中的所有段落 (共{len(doc.paragraphs)}个):")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # 只显示非空段落
                print(f"段落 #{i}: '{text[:80]}{'...' if len(text) > 80 else ''}'")
    
    # 检查表格中是否有邮件字段
    print(f"\n📋 检查表格中的邮件字段:")
    found_email_in_tables = False
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if any(keyword in text for keyword in email_keywords):
                    print(f"表格 #{table_idx}, 行 #{row_idx}, 列 #{cell_idx}: '{text}'")
                    found_email_in_tables = True
    
    if not found_email_in_tables:
        print("❌ 表格中也没有找到包含邮件关键词的单元格")
    
    print(f"\n✅ 文档分析完成")
    
except Exception as e:
    print(f"❌ 执行错误: {e}")
    import traceback
    traceback.print_exc()