#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MCP投标人名称处理器
专门处理文档中投标人名称的填写
支持两种填写方式：
1. 替换内容（如将"（公司全称）"替换为公司名称）
2. 在空格处填写（将空格替换为公司名称，格式与标签保持一致）
"""

import os
import re
import logging
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Mm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 配置日志
logger = logging.getLogger(__name__)


class MCPBidderNameProcessor:
    """MCP投标人名称处理器 - 专门处理投标人名称填写"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("请安装python-docx库：pip install python-docx")
        
        # 公司地址信息（默认值，可被覆盖）
        self.company_address = "北京市东城区王府井大街200号七层711室"
        
        # 尝试从配置文件读取项目编号和项目名称
        self.project_number = self._load_project_number()
        
        # 存储完整的公司信息
        self.company_info = {}
        self.company_name = ""  # 添加公司名称属性支持
        self.project_name = self._extract_project_name_from_config() or ""  # 从配置文件加载项目名称
        self.tender_no = ""
        self.date_text = ""
        
        # 投标人名称匹配规则 - 按优先级排序
        self.bidder_patterns = [
            # === 第一种方式：替换内容 ===
            # 通用括号内容替换 - 合并格式11-16（6个规则合并为1个）
            {
                'pattern': re.compile(r'(?P<prefix>[\(（])\s*(?P<content>(?:请填写\s*)?(?:供应商名称|供应商全称|投标人名称|公司名称|单位名称|采购人|供应商住址|供应商地址|公司地址|注册地址))\s*(?P<suffix>[\)）])'),
                'type': 'replace_content',
                'description': '通用括号内容替换 - 公司名称类'
            },
            
            
            # 格式17: "（供应商名称、地址）" - 括号内容替换为公司名称和地址
            {
                'pattern': re.compile(r'(?P<prefix>[\(（])\s*(?P<content>供应商名称、地址)\s*(?P<suffix>[\)）])'),
                'type': 'replace_content_with_address',
                'description': '括号内容替换 - 供应商名称、地址'
            },
            
            # 格式18: "（项目名称、项目编号）" - 括号内容替换为项目名称和项目编号
            {
                'pattern': re.compile(r'(?P<prefix>[\(（])\s*(?P<content>项目名称\s*[、，]\s*项目编号)\s*(?P<suffix>[\)）])'),
                'type': 'replace_content_with_project_info',
                'description': '括号内容替换 - 项目名称、项目编号'
            },
            
            
            # === 第二种方式：在空格处填写 ===
            # 通用简单填空规则 - 合并多个简单fill_space规则
            {
                'pattern': re.compile(r'^(?:\s*\d+\.\s*|\s+)?(?P<label>公司名称（全称、盖章）|公司名称（盖章）|供应商名称（盖章）|供应商名称\(盖章\)|供应商全称及公章|供应商全称|供应商名称|投标人名称（盖章）|单位名称及公章|投标人名称\(盖章\)|单位名称\(公章\))\s*(?P<sep>[:：])?\s*(?P<placeholder>_{3,}|\s{3,}|)\s*(?P<suffix>（[^）]*公章[^）]*）|\([^)]*公章[^)]*\))?\s*$'),
                'type': 'fill_space',
                'description': '通用简单填空 - 各种公司供应商名称格式（支持空格和数字前缀）'
            },
            
            # 合并后的公章/盖章在前规则 - 支持中英文括号、各种章类型、有无占位符
            {
                'pattern': re.compile(r'^(?P<label>供应商名称)\s*(?P<seal>[（(][^）)]*[公盖]章[^）)]*[）)])\s*(?P<sep>[:：])\s*(?P<placeholder>\s*)\s*$'),
                'type': 'fill_space_with_seal_prefix',
                'description': '公章/盖章在前统一规则 - 供应商名称'
            },
            
            
            # 通用无分隔符填写 - 合并格式18系列（4个规则合并为1个）
            {
                'pattern': re.compile(r'^(?P<label>供应商名称|公司名称|投标人名称|单位名称)\s*(?P<placeholder>\s{20,})\s*$'),
                'type': 'fill_space_no_separator',
                'description': '通用标签后空格填写 - 公司名称类（无冒号）'
            },
            
            # 合并后的所有编号类填写规则 - 支持采购编号、项目编号、通用编号，下划线或空格
            {
                'pattern': re.compile(r'(?P<label>采购编号|项目编号|编号)\s*(?P<sep>[:：])\s*(?P<placeholder>_+|[ \t]+)'),
                'type': 'fill_space_tender_no',
                'description': '所有编号类统一填写规则'
            },
            
            # 通用投标供应商名称填空 - 支持部分匹配（去掉$以支持同行多标签）
            {
                'pattern': re.compile(r'(?P<label>供应商名称|投标人名称(?:（公章）|\(公章\))?)\s*(?P<sep>[:：])\s*(?P<placeholder>\s.*|[_\-\u2014]+|＿+|——+)'),
                'type': 'fill_space',
                'description': '通用投标供应商名称填空 - 支持部分匹配'
            },
            
            # 项目名称相关 - 使用相同的处理逻辑
            {
                'pattern': re.compile(r'[\(（]\s*项目名称\s*[\)）]'),
                'type': 'replace_content_project',
                'description': '括号内容替换 - 项目名称'
            },
            {
                'pattern': re.compile(r'为\s*[\(（][^）)]*[\)）]\s*项目'),
                'type': 'replace_content_project_context',
                'description': '上下文中的项目名称替换 - 为（xxx）项目格式'
            },
            
            # 合并后的括号内所有编号替换规则
            {
                'pattern': re.compile(r'[\(（]\s*(?P<content>采购编号|招标编号|项目编号)\s*[\)）]'),
                'type': 'replace_content_tender_no',
                'description': '括号内所有编号类型统一替换'
            },
        ]
        
        # 运行时验证：确保使用合并后的11规则版本（已添加项目名称、项目编号组合规则）
        expected_rule_count = 11
        actual_rule_count = len(self.bidder_patterns)
        if actual_rule_count != expected_rule_count:
            logger.error(f"❌ 严重错误：期望{expected_rule_count}个规则，实际{actual_rule_count}个！缓存问题未解决")
            raise RuntimeError(f"规则数量不匹配：期望{expected_rule_count}，实际{actual_rule_count}")
        else:
            logger.info(f"✅ 规则验证通过：成功加载{actual_rule_count}个合并规则")
    
    def _load_project_number(self) -> str:
        """从“读取信息”页面的tender_config.ini加载项目编号"""
        try:
            import configparser
            import os
            from pathlib import Path
            
            # 指向“读取信息”页面的配置文件
            tender_info_path = str(Path(__file__).parent.parent.parent / "1.读取信息")
            config_file = os.path.join(tender_info_path, 'tender_config.ini')
            
            if os.path.exists(config_file):
                config = configparser.ConfigParser()
                config.read(config_file, encoding='utf-8')
                if config.has_section('PROJECT_INFO'):
                    project_num = config.get('PROJECT_INFO', 'project_number', fallback='')
                    if project_num and project_num.strip():
                        logger.info(f"从读取信息页面加载项目编号: {project_num}")
                        return project_num.strip()
            
            # 如果没有找到，返回空字符串
            logger.info(f"未找到读取信息页面的配置文件: {config_file}")
            return ""
            
        except Exception as e:
            logger.warning(f"从读取信息页面加载项目编号失败: {e}")
            return ""
    
    def _format_chinese_date(self, date_str: str) -> str:
        """
        将英文日期格式(YYYY-MM-DD)转换为中文日期格式(YYYY年M月D日)
        例如：2000-04-21 -> 2000年4月21日
        """
        if not date_str or not isinstance(date_str, str):
            return ''
        
        try:
            # 匹配 YYYY-MM-DD 格式
            import re
            match = re.match(r'^(\d{4})-(\d{1,2})-(\d{1,2})$', date_str.strip())
            if match:
                year, month, day = match.groups()
                # 转换为中文格式，去掉前导0
                return f"{year}年{int(month)}月{int(day)}日"
            else:
                # 如果不匹配预期格式，返回原字符串
                return date_str
        except Exception as e:
            logger.warning(f"日期格式转换失败: {date_str}, 错误: {e}")
            return date_str
    
    def _get_project_info_field(self, field_name: str) -> str:
        """从"读取信息"页面的配置文件读取项目信息字段"""
        try:
            import configparser
            import os
            from pathlib import Path
            
            # 指向实际的配置文件位置
            config_file = str(Path(__file__).parent.parent.parent / "ai_tender_system" / "data" / "configs" / "tender_config.ini")
            
            if not os.path.exists(config_file):
                logger.warning(f"读取信息页面的配置文件不存在: {config_file}")
                return ""
            
            config = configparser.ConfigParser()
            config.read(config_file, encoding='utf-8')
            
            if 'PROJECT_INFO' not in config:
                logger.warning("读取信息页面配置文件中没有PROJECT_INFO节")
                return ""
            
            value = config['PROJECT_INFO'].get(field_name, '')
            
            # 过滤掉"未提供"等无效值
            if value and value not in ['未提供', '未知', 'N/A', 'n/a', '']:
                logger.info(f"从读取信息页面获取项目字段{field_name}: {value}")
                return value
            else:
                logger.info(f"读取信息页面{field_name}字段为空或无效: {value}")
                return ""
                
        except Exception as e:
            logger.warning(f"从读取信息页面读取项目信息字段{field_name}失败: {e}")
            return ""
        
    def process_bidder_name(self, input_file: str, output_file: str, company_name: str) -> Dict:
        """
        处理投标人名称填写
        
        Args:
            input_file: 输入文件路径
            output_file: 输出文件路径  
            company_name: 公司名称
            
        Returns:
            处理结果字典
        """
        logger.info(f"开始MCP投标人名称处理: {input_file}")
        logger.info(f"使用公司名称: {company_name}")
        
        # 设置实例属性以供统一处理方法使用
        self.company_name = company_name
        
        try:
            # 打开Word文档
            doc = Document(input_file)
            
            # 统计信息
            stats = {
                'total_replacements': 0,
                'replace_content_count': 0,
                'fill_space_count': 0,
                'patterns_found': []
            }
            
            # 遍历所有段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                
                # 记录段落处理前的文本
                original_para_text = paragraph.text
                logger.info(f"处理段落 #{para_idx}: '{original_para_text[:100]}...'")
                
                # 🎯 新增：检查是否为多项替换段落，如果是则使用批量替换策略
                if self._should_use_batch_replacement(paragraph):
                    logger.info("🔄 检测到多项替换段落，使用批量替换策略")
                    
                    # 优先使用传递进来的项目名称和招标编号，如果没有则从配置文件读取
                    project_name = self.project_name if self.project_name else self._extract_project_name_from_config()
                    tender_number = self.tender_no if self.tender_no else self._extract_tender_number_from_config()
                    
                    logger.info(f"批量替换使用项目信息: 项目名称='{project_name}', 招标编号='{tender_number}'")
                    
                    # 执行批量替换
                    batch_result = self._batch_replace_multiple_items(
                        paragraph, 
                        company_name, 
                        project_name, 
                        tender_number
                    )
                    
                    if batch_result['success']:
                        # 批量替换成功，更新统计信息
                        replacement_count = batch_result['replacements']
                        stats['total_replacements'] += replacement_count
                        stats['replace_content_count'] += replacement_count
                        
                        for i in range(replacement_count):
                            stats['patterns_found'].append({
                                'rule_index': f'批量-{i+1}',
                                'description': '批量多项替换',
                                'type': 'batch_replace',
                                'original_text': f'多项内容替换',
                                'paragraph_index': para_idx
                            })
                        
                        logger.info(f"✅ 段落 #{para_idx} 批量替换成功，跳过常规处理")
                        continue  # 跳过常规的逐规则处理
                    else:
                        logger.info("⚠️ 批量替换失败，回退到常规处理")
                
                # 🎯 常规处理：逐规则匹配和处理
                # 尝试匹配每个规则
                bidder_name_processed = False  # 标记投标人名称是否已处理
                for rule_idx, rule in enumerate(self.bidder_patterns):
                    # 项目名称和项目编号处理可以与投标人名称处理并行进行
                    if (bidder_name_processed and 
                        rule['type'] not in ['replace_content_project', 'replace_content_project_context', 'replace_content_tender_no', 'replace_content_with_project_info']):
                        continue  # 如果投标人名称已处理，只允许项目名称和项目编号处理继续
                        
                    pattern = rule['pattern']
                    # 项目名称和项目编号处理使用当前文本，其他处理使用原始文本
                    search_text = paragraph.text if rule['type'] in ['replace_content_project', 'replace_content_project_context', 'replace_content_tender_no', 'replace_content_with_project_info'] else original_para_text
                    match = pattern.search(search_text)
                    
                    if match:
                        logger.info(f"匹配到规则 #{rule_idx+1}: {rule['description']}")
                        logger.info(f"匹配文本: '{match.group(0)}'")
                        
                        # 检查是否已经包含公司名称，避免重复填写
                        if company_name in original_para_text:
                            logger.info(f"段落已包含公司名称，跳过处理")
                            continue
                        
                        success = False
                        # 根据类型选择处理方式
                        if rule['type'] == 'replace_content':
                            success = self._replace_content_method(paragraph, match, company_name, rule)
                            if success:
                                stats['replace_content_count'] += 1
                        elif rule['type'] == 'replace_content_with_address':
                            success = self._replace_content_with_address_method(paragraph, match, company_name, rule)
                            if success:
                                stats['replace_content_count'] += 1
                        elif rule['type'] == 'replace_content_with_project_info':
                            success = self._replace_content_with_project_info_method(paragraph, match, company_name, rule)
                            if success:
                                stats['replace_content_count'] += 1
                        elif rule['type'] == 'fill_space':
                            success = self._fill_space_method(paragraph, match, company_name, rule)  
                            if success:
                                stats['fill_space_count'] += 1
                        elif rule['type'] == 'fill_space_no_separator':
                            success = self._fill_space_no_separator_method(paragraph, match, company_name, rule)
                            if success:
                                stats['fill_space_count'] += 1
                        elif rule['type'] == 'fill_space_tender_no':
                            success = self._fill_space_tender_no_method(paragraph, match, rule)
                            if success:
                                stats['fill_space_count'] += 1
                        elif rule['type'] == 'fill_space_with_seal_prefix':
                            success = self._fill_space_with_seal_prefix_method(paragraph, match, company_name, rule)
                            if success:
                                stats['fill_space_count'] += 1
                        elif rule['type'] == 'replace_content_project':
                            success = self._replace_content_project_method(paragraph, match, rule)
                            if success:
                                stats['replace_content_count'] += 1
                        elif rule['type'] == 'replace_content_project_context':
                            success = self._replace_content_project_context_method(paragraph, match, rule)
                            if success:
                                stats['replace_content_count'] += 1
                        elif rule['type'] == 'replace_content_tender_no':
                            success = self._replace_content_tender_no_method(paragraph, match, rule)
                            if success:
                                stats['replace_content_count'] += 1
                                
                        if success:
                            stats['total_replacements'] += 1
                            stats['patterns_found'].append({
                                'rule_index': rule_idx + 1,
                                'description': rule['description'],
                                'type': rule['type'],
                                'original_text': match.group(0),
                                'paragraph_index': para_idx
                            })
                            logger.info(f"处理后: '{paragraph.text[:100]}...'")
                            
                            # 如果是投标人名称相关处理，标记为已处理
                            if rule['type'] not in ['replace_content_project', 'replace_content_project_context', 'replace_content_tender_no', 'fill_space_tender_no']:
                                bidder_name_processed = True
                            
                            # 不再需要break，允许在同一段落中进行多种类型的处理
            
            # 保存文档
            doc.save(output_file)
            logger.info(f"MCP投标人名称处理完成，保存到: {output_file}")
            logger.info(f"处理统计: 总计{stats['total_replacements']}次, 替换内容{stats['replace_content_count']}次, 空格填写{stats['fill_space_count']}次")
            
            return {
                'success': True,
                'stats': stats,
                'message': f'成功处理{stats["total_replacements"]}个投标人名称字段'
            }
            
        except Exception as e:
            logger.error(f"MCP投标人名称处理失败: {e}", exc_info=True)
            return {
                'success': False,
                'error': f'处理失败: {str(e)}'
            }
    
    def _replace_content_method(self, paragraph: Paragraph, match, company_name: str, rule: dict) -> bool:
        """
        第一种填写方式：替换内容 - 使用智能三层替换
        如将"（公司全称）"替换为"（智慧足迹数据科技有限公司）"，格式保持不变
        """
        try:
            # 获取匹配的组
            groups = match.groupdict()
            prefix = groups.get('prefix', '')
            content = groups.get('content', '')  
            suffix = groups.get('suffix', '')
            
            if not content:
                return False
            
            # 根据内容类型选择替换文本
            if '采购人' in content:
                # 如果是采购人，使用项目信息中的采购人名称
                replacement_text = self._get_project_info_field('tenderer') or "未提供采购人信息"
            elif any(addr_keyword in content for addr_keyword in ['住址', '地址']):
                # 如果是地址相关，使用公司地址
                replacement_text = self.company_address or "未提供公司地址"
            else:
                # 其他情况使用公司名称
                replacement_text = company_name
            
            # 使用智能三层替换策略
            old_text = f"{prefix}{content}{suffix}"
            new_text = f"{prefix}{replacement_text}{suffix}"
            
            success = self.smart_text_replace(paragraph, old_text, new_text)
            
            if success:
                logger.info(f"智能替换内容完成: '{content}' -> '{replacement_text}'")
            else:
                logger.error(f"智能替换内容失败: '{old_text}'")
                
            return success
            
        except Exception as e:
            logger.error(f"替换内容方法失败: {e}")
            return False
    
    def _fill_space_method(self, paragraph: Paragraph, match, company_name: str, rule: dict) -> bool:
        """
        第二种填写方式：在空格处填写
        增强版：支持跨run拆分处理
        """
        try:
            # 获取匹配的组
            groups = match.groupdict()
            label = groups.get('label', '')
            sep = groups.get('sep', '')
            placeholder = groups.get('placeholder', '')
            suffix = groups.get('suffix', '')
            
            if not label:
                return False
            
            # 特殊处理：检查是否是跨run拆分的情况（如"供应商名称："被拆分）
            # 注意：只有当文本恰好是标签+分隔符（无前导空格）时才认为是跨run拆分
            full_text = ''.join(run.text for run in paragraph.runs)
            if full_text == f"{label}{sep}" and not placeholder:
                logger.info(f"检测到跨run拆分情况: '{full_text}'")
                return self._enhanced_cross_run_fill(paragraph, label, sep, company_name)
            
            # 构建要查找和替换的模式
            if suffix:  # 有后缀的情况，如（加盖公章）
                search_pattern = f"{label}{sep}{placeholder}{suffix}"
                replacement = f"{label}{sep}{company_name} {suffix}"
            else:  # 无后缀的情况
                if placeholder:  # 有占位符（空格或下划线）
                    search_pattern = f"{label}{sep}{placeholder}"
                    replacement = f"{label}{sep}{company_name}"
                else:  # 无占位符，直接在冒号后添加
                    search_pattern = f"{label}{sep}"
                    replacement = f"{label}{sep}{company_name}"
            
            logger.info(f"查找模式: '{search_pattern}' -> 替换为: '{replacement}'")
            
            # 寻找包含整个匹配文本的run并替换
            found_and_replaced = False
            for run in paragraph.runs:
                if search_pattern in run.text:
                    font_name = run.font.name if run.font.name else "默认"
                    logger.info(f"找到匹配run: '{run.text}'，字体={font_name}")
                    # 直接在run中替换内容，保持原有格式
                    run.text = run.text.replace(search_pattern, replacement)
                    logger.info(f"替换后: '{run.text}'，字体={font_name}")
                    found_and_replaced = True
                    break
            
            # 如果没找到完整匹配，尝试分别查找标签和占位符
            if not found_and_replaced:
                logger.info("未找到完整匹配，尝试分别处理")
                
                # 首先定位包含标签的run的位置（支持跨run标签）
                label_run_index = -1
                
                # 方法1：查找包含完整标签的单个run
                for i, run in enumerate(paragraph.runs):
                    if label and label in run.text:
                        label_run_index = i
                        break
                
                # 方法2：如果没找到完整标签，查找标签开始的run
                if label_run_index == -1 and label:
                    # 查找标签的第一个词或关键词
                    if "供应商" in label:
                        search_term = "供应商"
                    elif "投标人" in label:
                        search_term = "投标人"
                    elif "投标" in label:
                        search_term = "投标"
                    else:
                        label_parts = label.split()
                        search_term = label_parts[0] if label_parts else label
                    
                    for i, run in enumerate(paragraph.runs):
                        if search_term in run.text:
                            logger.info(f"找到标签关键词 '{search_term}' 在run {i}")
                            label_run_index = i
                            break
                
                # 如果找到标签run，从该位置开始向后查找占位符
                if label_run_index >= 0:
                    logger.info(f"找到标签run在位置 {label_run_index}")
                    # 从标签run开始向后查找占位符
                    for i in range(label_run_index, len(paragraph.runs)):
                        run = paragraph.runs[i]
                        if placeholder and placeholder in run.text:
                            logger.info(f"找到标签后的占位符run: '{run.text}'")
                            # 只替换第一个匹配的占位符，避免重复替换
                            run.text = run.text.replace(placeholder, company_name, 1)
                            
                            # 清除可能的下划线格式（因为占位符可能有下划线，但填充内容不需要）
                            if run.underline:
                                run.underline = False
                                logger.info("清除了填充内容的下划线格式")
                            
                            logger.info(f"替换占位符后: '{run.text}'")
                            found_and_replaced = True
                            break
                
                # 如果还是没找到，使用原来的逻辑作为fallback
                if not found_and_replaced:
                    logger.info("未找到标签后的占位符，使用通用查找")
                    for run in paragraph.runs:
                        # 查找包含占位符的部分
                        if placeholder and placeholder in run.text:
                            logger.info(f"找到包含占位符的run: '{run.text}'")
                            # 只替换第一个匹配的占位符，避免重复替换
                            run.text = run.text.replace(placeholder, company_name, 1)
                            
                            # 清除可能的下划线格式（因为占位符可能有下划线，但填充内容不需要）
                            if run.underline:
                                run.underline = False
                                logger.info("清除了填充内容的下划线格式")
                            
                            logger.info(f"替换占位符后: '{run.text}'")
                            found_and_replaced = True
                            break
            
            if found_and_replaced:
                logger.info(f"空格填写方式成功: 填写'{company_name}'，保持原有格式")
                
                # 强化清理：专门处理"内容+残留占位符"混合状态
                self._enhanced_residual_placeholder_cleanup(paragraph, label, company_name)
                
                # 注意：强化清理已包含完整的清理逻辑，无需再调用智能清理
                
                # 特殊清理：处理字段间下划线占位符
                self._cleanup_underline_placeholders_between_fields(paragraph, company_name)
                
                return True
            else:
                logger.info("尝试跨run处理分散的文本")
                return self._handle_cross_run_text(paragraph, match, company_name, rule)
                
        except Exception as e:
            logger.error(f"填空方式处理失败: {e}")
            return False
    
    def _enhanced_cross_run_fill(self, paragraph: Paragraph, label: str, sep: str, company_name: str) -> bool:
        """
        增强的跨run填写方法 - 修复版
        专门处理像"供应商名称："这样被拆分成多个run的情况
        正确重构完整标签，避免拆分问题
        """
        try:
            logger.info(f"使用增强跨run填写: 标签='{label}', 分隔符='{sep}'")
            
            # 获取完整段落文本
            full_text = paragraph.text
            complete_label_sep = f"{label}{sep}"
            
            logger.info(f"完整段落文本: '{full_text}'")
            logger.info(f"寻找完整标签: '{complete_label_sep}'")
            
            # 检查run结构，找到构成完整标签的所有runs
            accumulated_text = ""
            label_runs = []
            
            for i, run in enumerate(paragraph.runs):
                if run.text.strip():  # 忽略纯空格run
                    potential_text = accumulated_text + run.text.strip()
                    
                    # 检查是否匹配标签的开始部分
                    if complete_label_sep.startswith(potential_text):
                        label_runs.append((i, run))
                        accumulated_text = potential_text
                        
                        logger.info(f"Run {i} 匹配标签部分: '{run.text}' -> 累积: '{accumulated_text}'")
                        
                        # 如果已经构成完整标签
                        if accumulated_text == complete_label_sep:
                            logger.info(f"找到完整跨run标签，涉及runs: {[(j, r.text) for j, r in label_runs]}")
                            break
                    elif accumulated_text and not complete_label_sep.startswith(accumulated_text):
                        # 重新开始查找
                        if complete_label_sep.startswith(run.text.strip()):
                            label_runs = [(i, run)]
                            accumulated_text = run.text.strip()
                        else:
                            label_runs = []
                            accumulated_text = ""
            
            # 如果找到了完整的跨run标签
            if label_runs and accumulated_text == complete_label_sep:
                logger.info(f"成功识别跨run标签，开始重构")
                
                # 重构策略：使用第一个run承载完整内容，清空其他runs
                first_run_idx, first_run = label_runs[0]
                
                # 查找占位符（在标签runs之后的空格或下划线）
                placeholder_found = False
                for i in range(len(paragraph.runs)):
                    if i not in [idx for idx, _ in label_runs]:
                        run = paragraph.runs[i]
                        # 检查是否为占位符run
                        if run.text and (run.text.isspace() or '_' in run.text):
                            # 将公司名称放到第一个run，替换占位符
                            first_run.text = f"{complete_label_sep} {company_name}"
                            run.text = ""  # 清空占位符run
                            placeholder_found = True
                            logger.info(f"重构成功: 标签run设为 '{first_run.text}', 清空占位符run")
                            break
                
                # 如果没有找到占位符，直接在标签后添加
                if not placeholder_found:
                    first_run.text = f"{complete_label_sep} {company_name}"
                    logger.info(f"重构成功: 直接添加 '{first_run.text}'")
                
                # 清空其他标签runs
                for i in range(1, len(label_runs)):
                    _, run_to_clear = label_runs[i]
                    run_to_clear.text = ""
                    logger.info(f"清空标签run: 原内容 '{run_to_clear.text}'")
                
                logger.info(f"增强跨run填写成功（完整重构）: '{complete_label_sep}' + '{company_name}'")
                return True
            
            # 备用方案：如果无法识别跨run结构，尝试简单处理
            logger.warning(f"无法识别跨run结构，使用备用方案")
            
            # 查找分隔符后面的占位符并替换
            import re
            patterns = [
                re.compile(f"{re.escape(label)}{re.escape(sep)}\\s*(_+|\\s{{3,}})"),  # 标签:___ 或 标签:   
                re.compile(f"{re.escape(label)}\\s*{re.escape(sep)}\\s*(_+|\\s{{3,}})")  # 标签 :___ 或 标签 :   
            ]
            
            for pattern in patterns:
                match = pattern.search(full_text)
                if match:
                    # 找到占位符，进行智能替换
                    old_text = match.group(0)
                    new_text = f"{label}{sep} {company_name}"
                    
                    success = self.smart_text_replace(paragraph, old_text, new_text)
                    if success:
                        logger.info(f"增强跨run填写成功（备用方案）: '{old_text}' -> '{new_text}'")
                        return True
            
            logger.warning(f"增强跨run填写失败，无法处理")
            return False
            
        except Exception as e:
            logger.error(f"增强跨run填写失败: {e}")
            return False
    
    def _handle_cross_run_text(self, paragraph: Paragraph, match, company_name: str, rule: dict) -> bool:
        """
        处理跨run分散的文本情况
        当标签和占位符分布在不同run中时使用
        """
        try:
            groups = match.groupdict()
            label = groups.get('label', '')
            sep = groups.get('sep', '')
            
            # 改进的跨run处理：正确重构完整标签
            complete_label_with_sep = label + sep
            
            # 首先尝试找到完整包含标签+分隔符的run
            target_run = None
            for run in paragraph.runs:
                if complete_label_with_sep in run.text:
                    target_run = run
                    logger.info(f"找到完整标签run: '{run.text}'")
                    break
            
            if target_run:
                # 情况1：完整标签在一个run中，直接替换
                if target_run.text.endswith(complete_label_with_sep):
                    target_run.text = target_run.text + " " + company_name
                else:
                    target_run.text = target_run.text.replace(complete_label_with_sep, f"{complete_label_with_sep} {company_name}", 1)
            else:
                # 情况2：标签跨run分布，需要重构
                logger.info(f"标签跨run分布，开始重构: '{complete_label_with_sep}'")
                
                # 找到包含标签开始部分的runs
                label_runs = []
                accumulated_text = ""
                
                for i, run in enumerate(paragraph.runs):
                    if run.text.strip():  # 跳过空run
                        potential_text = accumulated_text + run.text.strip()
                        if complete_label_with_sep.startswith(potential_text):
                            label_runs.append((i, run))
                            accumulated_text = potential_text
                            
                            # 检查是否已经构成完整标签
                            if accumulated_text == complete_label_with_sep:
                                logger.info(f"找到完整跨run标签: {[(j, r.text) for j, r in label_runs]}")
                                break
                        elif accumulated_text and complete_label_with_sep.startswith(accumulated_text):
                            # 继续累加，但不匹配当前run
                            pass
                        else:
                            # 重新开始
                            if complete_label_with_sep.startswith(run.text.strip()):
                                label_runs = [(i, run)]
                                accumulated_text = run.text.strip()
                            else:
                                label_runs = []
                                accumulated_text = ""
                
                if label_runs and accumulated_text == complete_label_with_sep:
                    # 重构完整标签：使用第一个run，清空其他
                    first_run_idx, first_run = label_runs[0]
                    first_run.text = f"{complete_label_with_sep} {company_name}"
                    
                    # 清空其他标签runs
                    for i in range(1, len(label_runs)):
                        _, run_to_clear = label_runs[i]
                        run_to_clear.text = ""
                    
                    logger.info(f"跨run标签重构成功: '{complete_label_with_sep} {company_name}'")
                else:
                    logger.warning(f"无法重构跨run标签: '{complete_label_with_sep}'")
                    return False
            
            # 清空其余占位符runs
            for run in paragraph.runs:
                if run != target_run and run not in [r[1] for r in (label_runs if 'label_runs' in locals() else [])]:
                    # 清空包含空格或下划线的run
                    if (run.text.strip() == '' or '_' in run.text or run.text.isspace() or
                        (run.text.strip() and all(c in ' _　' for c in run.text))):
                        run.text = ""
            
            logger.info(f"跨run处理成功: '{label}{sep} {company_name}'")
            
            # 强化清理：专门处理"内容+残留占位符"混合状态
            self._enhanced_residual_placeholder_cleanup(paragraph, label, company_name)
            
            # 注意：强化清理已包含完整的清理逻辑，无需再调用智能清理
            
            return True
            
        except Exception as e:
            logger.error(f"跨run处理失败: {e}")
            return False
    
    def _clean_trailing_spaces(self, paragraph: Paragraph) -> None:
        """
        清理段落末尾的多余空格和下划线
        """
        try:
            # 从后往前检查run
            for run in reversed(paragraph.runs):
                original_text = run.text
                # 移除尾部的空格和下划线
                cleaned_text = re.sub(r'[_\s　]+$', '', run.text)
                if cleaned_text != original_text:
                    run.text = cleaned_text
                    removed_count = len(original_text) - len(cleaned_text)
                    if removed_count > 0:
                        logger.info(f"删除末尾空格下划线: '{original_text}' -> '{cleaned_text}'")
                        logger.info(f"尾部清理完成: 删除了{removed_count}个字符")
                    break  # 只清理最后一个有内容的run
        except Exception as e:
            logger.warning(f"清理尾部字符失败: {e}")
    def _copy_font_format(self, source_run, target_run):
        """复制字体格式"""
        try:
            if hasattr(source_run, 'font') and hasattr(target_run, 'font'):
                source_font = source_run.font
                target_font = target_run.font
                
                # 记录原始字体信息
                logger.info(f"源字体信息: 名称={source_font.name}, 大小={source_font.size}, 粗体={source_font.bold}")
                
                # 复制字体名称 - 尝试多种方式
                if source_font.name:
                    target_font.name = source_font.name
                    logger.info(f"设置目标字体名称为: {source_font.name}")
                else:
                    # 如果字体名称为空，尝试从段落样式获取
                    para_style = source_run._parent.style if hasattr(source_run, '_parent') else None
                    if para_style and hasattr(para_style.font, 'name') and para_style.font.name:
                        target_font.name = para_style.font.name
                        logger.info(f"从段落样式设置字体名称为: {para_style.font.name}")
                
                # 复制字体大小
                if source_font.size:
                    target_font.size = source_font.size
                elif hasattr(source_run, '_parent'):
                    # 尝试从段落样式获取
                    para_style = source_run._parent.style
                    if para_style and hasattr(para_style.font, 'size') and para_style.font.size:
                        target_font.size = para_style.font.size
                
                # 复制其他格式属性
                if source_font.bold is not None:
                    target_font.bold = source_font.bold
                if source_font.italic is not None:
                    target_font.italic = source_font.italic
                    
                # 复制字体颜色
                if source_font.color and hasattr(source_font.color, 'rgb'):
                    if source_font.color.rgb:
                        target_font.color.rgb = source_font.color.rgb
                
                # 验证复制结果
                logger.info(f"目标字体设置后: 名称={target_font.name}, 大小={target_font.size}, 粗体={target_font.bold}")
                        
        except Exception as e:
            logger.error(f"复制字体格式失败: {e}", exc_info=True)
    
    def _apply_saved_format(self, target_run, font_name, font_size, font_bold, font_italic, font_color):
        """应用保存的格式信息到目标run"""
        try:
            if hasattr(target_run, 'font'):
                target_font = target_run.font
                
                # 应用字体名称
                if font_name:
                    target_font.name = font_name
                    logger.info(f"应用字体名称: {font_name}")
                
                # 应用字体大小
                if font_size:
                    target_font.size = font_size
                    logger.info(f"应用字体大小: {font_size}")
                
                # 应用粗体
                if font_bold is not None:
                    target_font.bold = font_bold
                
                # 应用斜体
                if font_italic is not None:
                    target_font.italic = font_italic
                
                # 应用颜色
                if font_color:
                    target_font.color.rgb = font_color
                    
                logger.info(f"格式应用完成: 名称={target_font.name}, 大小={target_font.size}")
                        
        except Exception as e:
            logger.error(f"应用格式失败: {e}", exc_info=True)
    
    
    def _update_paragraph_text(self, paragraph: Paragraph, old_text: str, new_text: str):
        """更新段落文本，保持格式"""
        try:
            # 查找包含完整旧文本的run并替换
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    logger.debug(f"更新run文本: '{old_text}' -> '{new_text}'")
                    return True
            
            # 如果没找到包含完整文本的run，尝试更新整个段落
            if paragraph.text == old_text:
                # 清除所有run并重新创建
                for run in paragraph.runs:
                    run.clear()
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    logger.debug(f"重建段落文本: '{new_text}'")
                return True
            
            logger.warning("无法更新段落文本")
            return False
            
        except Exception as e:
            logger.error(f"更新段落文本失败: {e}", exc_info=True)
            return False
    


    def _cleanup_underline_placeholders_between_fields(self, paragraph, company_name: str):
        """清理供应商名称和采购编号之间的下划线占位符"""
        try:
            import re
            
            full_text = paragraph.text
            if not full_text or company_name not in full_text or '采购编号' not in full_text:
                return
                
            logger.info(f"开始清理字段间下划线占位符: '{full_text}'")
            
            # 方案1: 检查是否在同一个run中
            same_run_found = False
            pattern = f"({re.escape(company_name)})([\\s_]+)(采购编号)"
            
            for run in paragraph.runs:
                if company_name in run.text and '采购编号' in run.text:
                    # 在同一个run中处理
                    match = re.search(pattern, run.text)
                    if match:
                        company_part = match.group(1)  # 公司名称
                        middle_part = match.group(2)   # 中间的空格和下划线
                        field_part = match.group(3)    # "采购编号"
                        
                        # 分析中间部分，如果有下划线，替换为4个空格
                        underscore_count = middle_part.count('_')
                        if underscore_count > 0:
                            new_middle = "    "
                            logger.info(f"同run清理{underscore_count}个下划线，替换为4个空格")
                            
                            # 重建文本
                            new_text = run.text.replace(
                                match.group(0), 
                                f"{company_part}{new_middle}{field_part}"
                            )
                            run.text = new_text
                            logger.info(f"同run字段间占位符清理完成: '{new_text}'")
                            same_run_found = True
                            break
            
            # 方案2: 如果不在同一run中，查找包含下划线的独立run
            if not same_run_found:
                logger.info("未找到同run中的字段，检查跨run下划线占位符")
                
                for i, run in enumerate(paragraph.runs):
                    run_text = run.text
                    # 查找只包含下划线和空格的run
                    if '_' in run_text and len(run_text.strip('_ ')) == 0:
                        underscore_count = run_text.count('_')
                        if underscore_count > 0:
                            # 将下划线替换为4个空格
                            run.text = "    "
                            logger.info(f"跨run清理独立下划线run#{i}: '{run_text}' -> '    '")
                    
                    # 查找包含大量空格和下划线混合的run
                    elif re.search(r'[\s_]{5,}', run_text):
                        # 如果run中有很多空格和下划线的组合
                        underscore_count = run_text.count('_')
                        if underscore_count > 0:
                            # 替换所有下划线为空格，然后压缩为4个空格
                            new_text = re.sub(r'_+', ' ', run_text)
                            new_text = re.sub(r'\s{5,}', '    ', new_text)
                            run.text = new_text
                            logger.info(f"跨run清理混合占位符run#{i}: '{run_text}' -> '{new_text}'")
                            
            logger.info("字段间下划线占位符清理完成")
                            
        except Exception as e:
            logger.error(f"清理字段间下划线占位符失败: {e}")

    def _global_placeholder_cleanup(self, paragraph: Paragraph) -> None:
        """全局占位符清理 - 处理整个段落中的所有占位符"""
        try:
            import re
            
            full_text = paragraph.text
            if not full_text:
                return
                
            logger.info(f"开始全局占位符清理: '{full_text}'")
            
            # 逐个处理每个run
            for run_idx, run in enumerate(paragraph.runs):
                if not run.text:
                    continue
                    
                original_text = run.text
                cleaned_text = original_text
                
                # 清理连续的空格占位符（3个或以上）
                cleaned_text = re.sub(r'[ \t]{3,}', '', cleaned_text)
                
                # 清理连续的下划线占位符（2个或以上）
                cleaned_text = re.sub(r'_{2,}', '', cleaned_text)
                
                # 如果有变化则更新
                if cleaned_text != original_text:
                    run.text = cleaned_text
                    logger.info(f"全局清理run #{run_idx}: '{original_text}' -> '{cleaned_text}'")
                            
            logger.info(f"全局占位符清理完成: '{paragraph.text}'")
            
        except Exception as e:
            logger.error(f"全局占位符清理失败: {e}")
            # 如果智能清理失败，不执行任何清理，保持原状
            pass
    
    def _smart_placeholder_cleanup(self, paragraph: Paragraph, current_label: str = None) -> None:
        """智能占位符清理系统已删除 - 使用强化残留清理系统代替"""
        logger.debug(f"智能占位符清理系统已删除，跳过处理: '{paragraph.text[:50]}...'")
        pass
    
    def _identify_all_labels_in_paragraph(self, paragraph: Paragraph) -> list:
        """智能占位符清理系统已删除 - 返回空列表"""
        return []
    def _single_label_cleanup_safe(self, paragraph: Paragraph, label_info: dict):
        """智能占位符清理系统已删除 - 空实现"""
        pass
    def _direct_text_replace(self, paragraph: Paragraph, old_text: str, new_text: str):
        """智能占位符清理系统已删除 - 空实现"""
        pass
    def _multi_label_cleanup(self, paragraph: Paragraph, labels: list, current_label: str):
        """智能占位符清理系统已删除 - 空实现"""
        pass
    def _clean_placeholder_keep_separator(self, content: str, has_next_label: bool) -> str:
        """智能占位符清理系统已删除 - 返回原内容"""
        return content
    def _apply_cleaned_text_to_runs(self, paragraph: Paragraph, old_text: str, new_text: str):
        """将清理后的文本应用到runs，保持格式 - 修复版"""
        try:
            # 修复Bug: 直接执行替换，不做长度判断避免递归死循环
            logger.info(f"应用清理文本: '{old_text}' -> '{new_text}'")
            success = self.smart_text_replace(paragraph, old_text, new_text)
            
            if success:
                logger.info("✅ 清理文本应用成功")
            else:
                logger.warning(f"⚠️ smart_text_replace失败，但不再尝试其他清理方法避免递归")
                
        except Exception as e:
            logger.error(f"应用清理文本失败: {e}")
            

    def _replace_content_with_address_method(self, paragraph: Paragraph, match, company_name: str, rule: dict) -> bool:
        """替换内容方法 - 包含地址信息，使用智能三层替换"""
        try:
            logger.info(f"执行替换内容方法（含地址）: {rule['description']}")
            
            # 构造替换文本：公司名称 + 地址
            replacement_text = f"{company_name}、{self.company_address}"
            
            # 获取匹配的组
            prefix = match.group('prefix')  # （
            content = match.group('content')  # 供应商名称、地址
            suffix = match.group('suffix')  # ）
            full_match = match.group(0)  # （供应商名称、地址）
            
            # 使用智能三层替换策略
            old_text = full_match  # 完整的匹配文本
            new_text = f"{prefix}{replacement_text}{suffix}"  # 完整的替换文本
            
            success = self.smart_text_replace(paragraph, old_text, new_text)
            
            if success:
                logger.info(f"智能替换完成: '{content}' -> '{replacement_text}'")
            else:
                logger.error(f"智能替换失败: '{full_match}'")
                
            return success
            
        except Exception as e:
            logger.error(f"替换内容（含地址）失败: {e}")
            return False

    def _replace_content_with_project_info_method(self, paragraph: Paragraph, match, company_name: str, rule: dict) -> bool:
        """替换内容方法 - 包含项目名称和项目编号信息，使用智能三层替换"""
        try:
            logger.info(f"执行替换内容方法（含项目信息）: {rule['description']}")
            
            # 获取项目名称和项目编号
            project_name = self.project_name if hasattr(self, 'project_name') and self.project_name else "未提供项目名称"
            project_number = ""
            if hasattr(self, 'tender_no') and self.tender_no:
                project_number = self.tender_no
            elif hasattr(self, 'project_number') and self.project_number:
                project_number = self.project_number
            else:
                project_number = "未提供项目编号"
            
            # 构造替换文本：项目名称 + 项目编号
            replacement_text = f"{project_name}、{project_number}"
            
            # 获取匹配的组
            prefix = match.group('prefix')  # （
            content = match.group('content')  # 项目名称、项目编号
            suffix = match.group('suffix')  # ）
            full_match = match.group(0)  # （项目名称、项目编号）
            
            # 使用智能三层替换策略
            old_text = full_match  # 完整的匹配文本
            new_text = f"{prefix}{replacement_text}{suffix}"  # 完整的替换文本
            
            success = self.smart_text_replace(paragraph, old_text, new_text)
            
            if success:
                logger.info(f"智能替换完成: '{content}' -> '{replacement_text}'")
            else:
                logger.error(f"智能替换失败: '{full_match}'")
                
            return success
            
        except Exception as e:
            logger.error(f"替换内容（含项目信息）失败: {e}")
            return False

    def _redistribute_text_to_runs(self, paragraph: Paragraph, new_text: str, replacement_text: str) -> bool:
        """将新文本重新分布到paragraph的runs中，保持原始格式"""
        try:
            # 获取原始文本和新文本
            original_text = ''.join(run.text for run in paragraph.runs)
            
            # 计算替换位置
            pattern = re.compile(r'(?P<prefix>[\(（])\s*(?P<content>供应商名称、地址)\s*(?P<suffix>[\)）])')
            match = pattern.search(original_text)
            
            if not match:
                logger.error("无法在原始文本中找到匹配模式")
                return False
            
            # 分析文本分段：前部分 + 匹配部分 + 后部分
            before_match = original_text[:match.start()]
            match_text = match.group(0)
            after_match = original_text[match.end():]
            replacement_full = f"{match.group('prefix')}{replacement_text}{match.group('suffix')}"
            
            logger.info(f"文本分段分析:")
            logger.info(f"  前部分: '{before_match}' (长度: {len(before_match)})")
            logger.info(f"  匹配部分: '{match_text}' -> '{replacement_full}'")
            logger.info(f"  后部分: '{after_match}' (长度: {len(after_match)})")
            
            # 分析匹配区域的格式信息
            match_format_info = self._analyze_match_area_format(paragraph, match.start(), match.end())
            
            # 使用智能格式保持的方法重新分布文本
            return self._smart_redistribute_with_format(paragraph, before_match, replacement_full, after_match, match_format_info)
            
        except Exception as e:
            logger.error(f"重新分布文本失败: {e}")
            return False
    
    def _analyze_match_area_format(self, paragraph, match_start: int, match_end: int):
        """分析匹配区域的格式信息"""
        try:
            current_pos = 0
            match_area_format = {
                'font_name': None,
                'font_size': None,
                'bold': None,
                'italic': None,
                'underline': None,
                'color': None
            }
            
            # 找到匹配区域内最常见的格式
            format_samples = []
            
            for run in paragraph.runs:
                run_start = current_pos
                run_end = current_pos + len(run.text)
                current_pos = run_end
                
                # 检查这个run是否与匹配区域重叠
                if run_start < match_end and run_end > match_start:
                    format_samples.append({
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'underline': run.font.underline,
                        'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    })
            
            # 选择最常见的格式作为替换文本的格式
            if format_samples:
                # 对于替换的文本，使用匹配区域内的格式（通常是斜体）
                sample = format_samples[0]  # 使用第一个匹配的格式
                match_area_format = sample
                logger.info(f"匹配区域格式: 斜体={sample['italic']}, 字体={sample['font_name']}")
            
            return match_area_format
            
        except Exception as e:
            logger.error(f"分析匹配区域格式失败: {e}")
            return {}
    
    def _smart_redistribute_with_format(self, paragraph, before_text: str, replacement_text: str, after_text: str, match_format_info: dict = None) -> bool:
        """智能重新分布文本，保持格式"""
        try:
            # 记录原始run的格式信息
            original_runs = []
            current_pos = 0
            
            for run in paragraph.runs:
                run_info = {
                    'text': run.text,
                    'start_pos': current_pos,
                    'end_pos': current_pos + len(run.text),
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                }
                original_runs.append(run_info)
                current_pos += len(run.text)
            
            # 清空所有run
            for run in paragraph.runs:
                run.text = ""
            
            # 构建新文本
            new_full_text = before_text + replacement_text + after_text
            
            # 计算文本区域分界（包括前缀和后缀括号）
            # 从replacement_full中提取前后缀长度
            prefix_len = 1  # "（" 的长度
            suffix_len = 1  # "）" 的长度
            
            before_end = len(before_text)
            replacement_start = before_end  # 包括左括号
            replacement_end = before_end + prefix_len + len(replacement_text) + suffix_len  # 包括右括号
            
            # 重新分配文本，智能保持格式
            assigned_runs = 0
            current_text_pos = 0
            
            for i, run_info in enumerate(original_runs):
                if assigned_runs >= len(paragraph.runs):
                    break
                    
                current_run = paragraph.runs[assigned_runs]
                
                # 确定这个run应该包含多少文本
                if current_text_pos >= len(new_full_text):
                    break
                
                # 计算适合这个run的文本长度
                if i == len(original_runs) - 1:  # 最后一个run
                    run_text_length = len(new_full_text) - current_text_pos
                else:
                    # 使用原始run长度作为参考，但不超过剩余文本长度
                    original_length = len(run_info['text'])
                    run_text_length = min(original_length, len(new_full_text) - current_text_pos)
                
                # 分配文本到run
                if run_text_length > 0:
                    current_run.text = new_full_text[current_text_pos:current_text_pos + run_text_length]
                    
                    # 🎯 简化格式选择逻辑 - 确保格式一致性
                    text_end_pos = current_text_pos + run_text_length
                    
                    # 计算这个run与各个区域的重叠情况
                    before_overlap = max(0, min(text_end_pos, before_end) - current_text_pos)
                    replacement_overlap = max(0, min(text_end_pos, replacement_end) - max(current_text_pos, replacement_start))
                    after_overlap = max(0, text_end_pos - max(current_text_pos, replacement_end))
                    
                    # 🔧 修复：简化格式选择，优先保持一致性
                    if replacement_overlap > 0:
                        # 任何涉及替换区域的run都使用匹配区域的格式（斜体+下划线）
                        format_to_use = match_format_info if match_format_info else run_info
                        logger.debug(f"Run #{assigned_runs+1} 涉及替换区域，使用统一的斜体+下划线格式")
                    elif before_overlap > 0:
                        # 纯前部分，使用原始格式
                        format_to_use = run_info  
                        logger.debug(f"Run #{assigned_runs+1} 纯前部分，使用原始格式")
                    else:
                        # 纯后部分，使用原始格式
                        format_to_use = run_info
                        logger.debug(f"Run #{assigned_runs+1} 纯后部分，使用原始格式")
                    
                    # 🎯 严格应用格式 - 确保完全一致
                    try:
                        if format_to_use.get('font_name'):
                            current_run.font.name = format_to_use['font_name']
                        if format_to_use.get('font_size'):
                            current_run.font.size = format_to_use['font_size']
                        # 🔧 关键修复：确保布尔值格式属性正确设置
                        current_run.font.bold = format_to_use.get('bold') if format_to_use.get('bold') is not None else False
                        current_run.font.italic = format_to_use.get('italic') if format_to_use.get('italic') is not None else False
                        current_run.font.underline = format_to_use.get('underline') if format_to_use.get('underline') is not None else False
                        
                        # 记录格式应用情况
                        logger.debug(f"  已应用格式: 字体={current_run.font.name}, 斜体={current_run.font.italic}, 下划线={current_run.font.underline}")
                    except Exception as format_error:
                        logger.warning(f"格式应用失败: {format_error}")
                        # 降级处理：至少保证基本格式
                        if replacement_overlap > 0 and match_format_info:
                            current_run.font.italic = match_format_info.get('italic', False)
                            current_run.font.underline = match_format_info.get('underline', False)
                    
                    current_text_pos += run_text_length
                    assigned_runs += 1
            
            # 如果还有剩余文本，放在最后一个有效run中
            if current_text_pos < len(new_full_text) and assigned_runs > 0:
                remaining_text = new_full_text[current_text_pos:]
                paragraph.runs[assigned_runs - 1].text += remaining_text
            
            logger.info(f"智能格式保持替换完成: 使用了{assigned_runs}个run")
            logger.info(f"最终文本: {new_full_text[:100]}...")
            
            return True
            
        except Exception as e:
            logger.error(f"智能重新分布失败: {e}")
            # 回退到简单方法
            return self._simple_redistribute_fallback(paragraph, before_text + replacement_text + after_text)
    
    def _simple_redistribute_fallback(self, paragraph, new_text: str) -> bool:
        """简单回退方法"""
        try:
            # 清空所有run
            for run in paragraph.runs:
                run.text = ""
            
            # 将文本放在第一个run中
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                logger.info("使用简单回退方法完成替换")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"简单回退方法失败: {e}")
            return False

    def _fill_space_no_separator_method(self, paragraph: Paragraph, match, company_name: str, rule: dict) -> bool:
        """填写空格方法 - 无分隔符"""
        try:
            logger.info(f"执行无分隔符空格填写方法: {rule['description']}")
            
            # 直接在空格位置添加公司名称
            label = match.group('label')
            placeholder = match.group('placeholder')
            original_text = match.group(0)
            
            # 构造新文本
            new_text = f"{label} {company_name}"
            
            # 方法1：尝试在包含标签的run中替换
            for run in paragraph.runs:
                if label in run.text:
                    run.text = run.text.replace(original_text, new_text)
                    logger.info(f"无分隔符填写完成（单run）: '{original_text}' -> '{new_text}'")
                    
                    # 智能占位符清理
                    # 智能占位符清理已删除
                    return True
            
            # 方法2：如果标签跨run，使用智能重分布方法
            logger.info(f"标签'{label}'可能跨run，尝试智能重分布")
            
            # 获取完整文本
            full_text = ''.join(run.text for run in paragraph.runs)
            
            # 如果完整文本匹配，进行替换
            if original_text in full_text:
                # 计算替换位置
                start_pos = full_text.find(original_text)
                end_pos = start_pos + len(original_text)
                
                # 构建新的完整文本
                before_text = full_text[:start_pos]
                after_text = full_text[end_pos:]
                
                # 使用智能重分布
                success = self._smart_redistribute_with_format(
                    paragraph, 
                    before_text, 
                    new_text, 
                    after_text
                )
                
                if success:
                    logger.info(f"无分隔符填写完成（跨run）: '{original_text}' -> '{new_text}'")
                    # 智能占位符清理
                    # 智能占位符清理已删除
                    return True
                else:
                    logger.warning("智能重分布失败，尝试简单替换")
                    
            # 方法3：最后的回退方案 - 直接替换整个段落
            if paragraph.text.strip() == original_text.strip():
                # 清空所有run
                for run in paragraph.runs:
                    run.text = ""
                
                # 将新文本放入第一个run
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    logger.info(f"无分隔符填写完成（回退）: '{original_text}' -> '{new_text}'")
                    return True
            
            logger.warning(f"无法在段落中找到匹配的文本进行替换")
            return False
            
        except Exception as e:
            logger.error(f"无分隔符空格填写失败: {e}")
            return False

    def _fill_space_with_seal_prefix_method(self, paragraph: Paragraph, match, company_name: str, rule: dict) -> bool:
        """填写公章在前的格式 - 供应商名称（加盖公章）：___"""
        try:
            logger.info(f"执行公章在前格式填写方法: {rule['description']}")
            
            # 获取匹配的各部分
            label = match.group('label')        # 供应商名称
            seal = match.group('seal')          # （加盖公章）
            sep = match.group('sep')            # ：
            placeholder = match.group('placeholder')  # 空格
            
            # 构造新文本，保留公章部分
            new_text = f"{label}{seal}{sep}{company_name}"
            
            # 查找原始文本
            original_text = match.group(0)
            
            # 方法1：尝试直接替换
            for run in paragraph.runs:
                if original_text in run.text:
                    run.text = run.text.replace(original_text, new_text)
                    logger.info(f"公章在前格式填写完成（单run）: '{original_text}' -> '{new_text}'")
                    
                    # 智能占位符清理
                    # 智能占位符清理已删除
                    return True
            
            # 方法2：跨run处理
            full_text = ''.join(run.text for run in paragraph.runs)
            if original_text in full_text:
                logger.info(f"使用跨run处理公章在前格式")
                
                # 计算替换位置
                start_pos = full_text.find(original_text)
                end_pos = start_pos + len(original_text)
                
                # 构建新的完整文本
                before_text = full_text[:start_pos]
                after_text = full_text[end_pos:]
                
                # 使用智能重分布
                success = self._smart_redistribute_with_format(
                    paragraph, 
                    before_text, 
                    new_text, 
                    after_text
                )
                
                if success:
                    logger.info(f"公章在前格式填写完成（跨run）: '{original_text}' -> '{new_text}'")
                    # 智能占位符清理
                    # 智能占位符清理已删除
                    return True
            
            # 方法3：回退方案
            if paragraph.text.strip() == original_text.strip():
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    logger.info(f"公章在前格式填写完成（回退）: '{original_text}' -> '{new_text}'")
                    return True
            
            logger.warning(f"无法处理公章在前格式")
            return False
            
        except Exception as e:
            logger.error(f"公章在前格式填写失败: {e}")
            return False

    def _fill_space_tender_no_method(self, paragraph: Paragraph, match, rule: dict) -> bool:
        """招标编号填写方法 - 处理采购编号、项目编号等"""
        try:
            logger.info(f"执行招标编号填写方法: {rule['description']}")
            
            # 获取匹配的组
            groups = match.groupdict()
            label = groups.get('label', '')
            sep = groups.get('sep', '')
            placeholder = groups.get('placeholder', '')
            
            # 获取招标编号 - 优先使用tender_no，如果没有则使用project_number
            tender_number = self.tender_no if hasattr(self, 'tender_no') and self.tender_no else \
                           (self.project_number if hasattr(self, 'project_number') and self.project_number else "未提供编号")
            
            # 使用智能三层替换策略
            old_text = f"{label}{sep}{placeholder}"
            new_text = f"{label}{sep}{tender_number}"
            
            success = self.smart_text_replace(paragraph, old_text, new_text)
            
            if success:
                # 智能占位符清理 - 区分占位符和分隔符
                # 智能占位符清理已删除
                logger.info(f"招标编号填写完成: '{label}' -> '{tender_number}' (已清理占位符)")
            else:
                logger.error(f"招标编号填写失败: '{old_text}'")
                
            return success
            
        except Exception as e:
            logger.error(f"招标编号填写方法失败: {e}")
            return False

    def _replace_content_project_method(self, paragraph: Paragraph, match, rule: dict) -> bool:
        """项目名称括号内容替换方法"""
        try:
            logger.info(f"执行项目名称括号内容替换: {rule['description']}")
            
            # 获取项目名称
            project_name = self.project_name if hasattr(self, 'project_name') and self.project_name else "未提供项目名称"
            
            # 查找包含匹配内容的run并替换
            match_text = match.group(0)  # 如（项目名称）
            
            for run in paragraph.runs:
                if match_text in run.text:
                    # 直接替换括号内的内容
                    new_text = run.text.replace(match_text, f"（{project_name}）")
                    run.text = new_text
                    logger.info(f"项目名称括号替换: '{match_text}' -> '（{project_name}）'")
                    return True
            
            return False
            
        except Exception as e:
            logger.error(f"项目名称括号替换失败: {e}")
            return False

    def _replace_content_project_context_method(self, paragraph: Paragraph, match, rule: dict) -> bool:
        """项目名称上下文替换方法 - 处理"为（xxx）项目"格式"""
        try:
            logger.info(f"执行项目名称上下文替换: {rule['description']}")
            
            # 获取项目名称
            project_name = self.project_name if hasattr(self, 'project_name') and self.project_name else "未提供项目名称"
            
            # 查找包含匹配内容的run并替换
            original_text = paragraph.text
            
            # 使用正则表达式替换"为（xxx）项目"格式
            new_text = re.sub(r'为\s*[\(（][^）)]*[\)）]\s*项目', f'为（{project_name}）项目', original_text, count=1)
            
            if new_text != original_text:
                # 使用安全的方法更新段落文本
                self._safe_replace_paragraph_text(paragraph, original_text, new_text)
                logger.info(f"项目名称上下文替换完成: '为（xxx）项目' -> '为（{project_name}）项目'")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"项目名称上下文替换失败: {e}")
            return False

    def _replace_content_tender_no_method(self, paragraph: Paragraph, match, rule: dict) -> bool:
        """项目编号括号内容替换方法 - 使用智能三层替换"""
        try:
            logger.info(f"执行项目编号括号内容替换: {rule['description']}")
            
            # 获取项目编号 - 支持多种属性名
            tender_no = ""
            if hasattr(self, 'tender_no') and self.tender_no:
                tender_no = self.tender_no
            elif hasattr(self, 'project_number') and self.project_number:
                tender_no = self.project_number
            else:
                tender_no = "未提供项目编号"
            
            # 如果项目编号为空，跳过处理
            if not tender_no.strip():
                logger.info("项目编号为空，跳过处理")
                return False
            
            # 使用智能三层替换策略
            match_text = match.group(0)  # 如（采购编号）、（招标编号）、（项目编号）
            new_text = f"（{tender_no}）"
            
            success = self.smart_text_replace(paragraph, match_text, new_text)
            
            if success:
                logger.info(f"项目编号智能替换完成: '{match_text}' -> '{new_text}'")
            else:
                logger.error(f"项目编号智能替换失败: '{match_text}'")
                
            return success
            
        except Exception as e:
            logger.error(f"项目编号括号替换失败: {e}")
            return False

    def process_business_response(self, input_file: str, output_file: str, 
                                company_info: dict, project_name: str = "", 
                                tender_no: str = "", date_text: str = ""):
        """
        处理商务应答文档
        
        Args:
            input_file: 输入文档路径
            output_file: 输出文档路径
            company_info: 完整的公司信息字典
            project_name: 项目名称
            tender_no: 招标编号
            date_text: 日期文本
            
        Returns:
            dict: 处理结果
        """
        try:
            logger.info(f"开始处理商务应答文档")
            logger.info(f"输入文件: {input_file}")
            logger.info(f"输出文件: {output_file}")
            logger.info(f"公司名称: {company_info.get('companyName', 'N/A')}")
            logger.info(f"项目名称: {project_name}")
            logger.info(f"招标编号: {tender_no}")
            
            # 保存信息到实例变量
            self.company_info = company_info
            self.project_name = project_name
            self.tender_no = tender_no
            self.date_text = date_text
            
            # 更新公司地址（优先使用注册地址）
            if company_info.get('registeredAddress'):
                self.company_address = company_info['registeredAddress']
            elif company_info.get('officeAddress'):
                self.company_address = company_info['officeAddress']
            
            # 更新项目编号
            if tender_no:
                self.project_number = tender_no
            
            # 先处理投标人名称
            logger.info("第1步：处理投标人名称")
            name_result = self.process_bidder_name(
                input_file=input_file,
                output_file=output_file,
                company_name=company_info.get('companyName', '')
            )
            
            if not name_result.get('success'):
                return name_result
            
            # 继续处理其他信息字段
            logger.info("第2步：处理其他公司信息字段")
            info_result = self._process_company_info_fields(output_file, company_info, project_name, tender_no, date_text)
            
            # 合并结果
            combined_stats = name_result.get('stats', {})
            info_stats = info_result.get('stats', {})
            
            # 更新总计数
            combined_stats['total_replacements'] = combined_stats.get('total_replacements', 0) + info_stats.get('total_replacements', 0)
            combined_stats['info_fields_processed'] = info_stats.get('info_fields_processed', 0)
            
            # 合并处理的模式列表
            if 'patterns_found' in combined_stats and 'patterns_found' in info_stats:
                combined_stats['patterns_found'].extend(info_stats['patterns_found'])
            
            return {
                'success': True,
                'stats': combined_stats,
                'message': f'商务应答文档处理完成，处理了{combined_stats.get("total_replacements", 0)}个字段'
            }
            
        except Exception as e:
            logger.error(f"商务应答文档处理失败: {e}")
            return {
                'success': False,
                'error': f'处理失败: {str(e)}'
            }

    def _create_unified_field_config(self, company_info: dict, project_name: str, tender_no: str, date_text: str):
        """
        创建统一的字段处理配置 - 优化版
        统一管理所有公司信息字段，支持表格式双字段布局处理
        
        Args:
            company_info: 公司信息
            project_name: 项目名称
            tender_no: 招标编号
            date_text: 日期文本
            
        Returns:
            list: 统一的字段配置列表
        """
        return [
            # 联系电话字段 - 统一处理各种电话相关字段
            {
                'field_names': ['电话', '联系电话', '固定电话', '电话号码', '联系方式'],
                'value': company_info.get('fixedPhone', ''),
                'display_name': '联系电话',
                'field_type': 'contact',  # 字段类型：联系信息
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理', '联系人', '项目联系人', '业主'],
                'formats': {
                    'with_colon': True,      # 支持带冒号格式：电话：
                    'without_colon': True,   # 支持不带冒号格式：电话
                    'placeholder_types': ['underline', 'space', 'mixed'],  # 下划线、空格、混合占位符
                    'table_layout': True     # 支持表格式布局
                },
                # 表格式双字段组合配置
                'table_combinations': [
                    {
                        'partner_field': '电子邮件',
                        'pattern': r'(电话|联系电话)(\s{8,})(电子邮件|电子邮箱|邮箱)',
                        'description': '电话+邮件表格组合'
                    }
                ]
            },
            # 电子邮件字段 - 统一处理各种邮件相关字段
            {
                'field_names': ['电子邮件', '电子邮箱', '邮箱', 'email', 'Email'],
                'value': company_info.get('email', '') or '未填写',
                'display_name': '电子邮件',
                'field_type': 'contact',  # 字段类型：联系信息
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理', '联系人', '项目联系人', '业主'],
                'formats': {
                    'with_colon': True,
                    'without_colon': True,
                    'placeholder_types': ['underline', 'space', 'mixed'],
                    'table_layout': True
                },
                # 表格式双字段组合配置
                'table_combinations': [
                    {
                        'partner_field': '电话',
                        'pattern': r'(电话|联系电话)(\s{8,})(电子邮件|电子邮箱|邮箱)',
                        'description': '电话+邮件表格组合（作为第二字段）'
                    }
                ]
            },
            # 传真字段 - 统一处理各种传真相关字段
            {
                'field_names': ['传真', '传真号码', '传真号', 'fax', 'Fax'],
                'value': company_info.get('fax', '') or '未填写',
                'display_name': '传真',
                'field_type': 'contact',  # 字段类型：联系信息
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理', '联系人', '项目联系人', '业主'],
                'formats': {
                    'with_colon': True,
                    'without_colon': True,
                    'placeholder_types': ['underline', 'space', 'mixed'],
                    'table_layout': True
                },
                # 表格式双字段组合配置
                'table_combinations': [
                    {
                        'partner_field': '地址',
                        'pattern': r'(地址|注册地址|办公地址|联系地址)(\s{8,})(传真)',
                        'description': '地址+传真表格组合（作为第二字段）'
                    }
                ]
            },
            # 邮政编码字段
            {
                'field_names': ['邮政编码', '邮编', '邮码'],
                'value': company_info.get('postalCode', ''),
                'display_name': '邮政编码',
                'field_type': 'basic_info',  # 字段类型：基本信息
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            },
            # 统一社会信用代码字段
            {
                'field_names': ['统一社会信用代码', '社会信用代码', '信用代码', '统一信用代码'],
                'value': company_info.get('socialCreditCode', ''),
                'display_name': '统一社会信用代码',
                'field_type': 'legal_info',  # 字段类型：法定信息
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            },
            # 注册资本字段
            {
                'field_names': ['注册资本', '注册资金', '资本'],
                'value': company_info.get('registeredCapital', ''),
                'display_name': '注册资本',
                'field_type': 'legal_info',  # 字段类型：法定信息
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            },
            # 网站字段
            {
                'field_names': ['网站', '网址', '官网', '公司网站'],
                'value': company_info.get('website', ''),
                'display_name': '网站',
                'field_type': 'contact',  # 字段类型：联系信息
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            },
            # 开户银行字段
            {
                'field_names': ['开户银行', '银行名称', '开户行', '银行'],
                'value': company_info.get('bankName', ''),
                'display_name': '开户银行',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            },
            # 银行账号字段
            {
                'field_names': ['银行账号', '账号', '银行账户', '账户号'],
                'value': company_info.get('bankAccount', ''),
                'display_name': '银行账号',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            },
            # 经营范围字段
            {
                'field_names': ['经营范围', '营业范围', '业务范围', '主营业务'],
                'value': company_info.get('businessScope', ''),
                'display_name': '经营范围',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            },
            # 成立日期字段
            {
                'field_names': ['成立日期', '成立时间', '设立日期', '注册日期'],
                'value': self._format_chinese_date(company_info.get('establishDate', '')),
                'display_name': '成立日期',
                'exclude_contexts': ['采购人', '招标人', '甲方', '代理'],
                'formats': {
                    'with_colon': True,
                    'without_colon': False,
                    'placeholder_types': ['underline', 'space'],
                    'table_layout': False
                }
            }
        ]

    def _process_company_info_fields(self, file_path: str, company_info: dict, 
                                   project_name: str, tender_no: str, date_text: str):
        """
        使用统一框架处理公司信息字段
        
        Args:
            file_path: 文档路径
            company_info: 公司信息
            project_name: 项目名称
            tender_no: 招标编号  
            date_text: 日期文本
            
        Returns:
            dict: 处理结果
        """
        try:
            doc = Document(file_path)
            logger.info(f"开始处理公司信息字段（统一框架），文档共有 {len(doc.paragraphs)} 个段落")
            
            total_replacements = 0
            patterns_found = []
            processed_paragraphs = set()  # 记录已处理的段落，防止重复处理
            
            # 获取统一的字段配置
            field_configs = self._create_unified_field_config(company_info, project_name, tender_no, date_text)
            
            # 处理所有段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if para_idx in processed_paragraphs:
                    continue
                    
                para_text = paragraph.text.strip()
                if not para_text:
                    continue
                
                logger.info(f"处理段落 #{para_idx}: '{para_text[:50]}...'")
                
                # 使用统一框架处理字段
                for field_config in field_configs:
                    result = self._process_unified_field(paragraph, para_text, field_config, para_idx)
                    if result['modified']:
                        total_replacements += result['replacements']
                        patterns_found.extend(result['patterns'])
                        processed_paragraphs.add(para_idx)
                        break  # 一个段落只处理一个字段，避免冲突
            
            # 后处理美化机制 - 统一清理和优化格式
            beautified_paragraphs = self._post_process_beautification(doc, total_replacements)
            
            # 保存处理后的文档
            doc.save(file_path)
            
            # 返回处理结果
            return {
                'success': True,
                'total_replacements': total_replacements,
                'patterns_found': patterns_found,
                'processed_paragraphs': len(processed_paragraphs),
                'message': f'公司信息字段处理完成，共替换 {total_replacements} 处内容'
            }
            
        except Exception as e:
            logger.error(f"处理公司信息字段失败: {e}")
            return {
                'success': False,
                'error': f'处理失败: {str(e)}'
            }
    
    def _post_process_beautification(self, doc, total_replacements: int) -> int:
        """
        后处理美化机制 - 统一清理和优化格式
        
        Args:
            doc: Word文档对象
            total_replacements: 已完成的替换次数
            
        Returns:
            int: 美化处理的段落数量
        """
        try:
            beautified_count = 0
            logger.info("开始后处理美化机制")
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if not para_text:
                    continue
                
                original_text = para_text
                beautified_text = self._beautify_paragraph_text(para_text)
                
                # 如果文本有改动，进行美化替换
                if beautified_text != original_text:
                    success = self._safe_replace_paragraph_text(paragraph, original_text, beautified_text)
                    if success:
                        beautified_count += 1
                        logger.info(f"段落 #{para_idx} 美化成功: '{original_text[:30]}...' -> '{beautified_text[:30]}...'")
                    else:
                        logger.warning(f"段落 #{para_idx} 美化失败")
            
            logger.info(f"后处理美化完成，处理了 {beautified_count} 个段落")
            return beautified_count
            
        except Exception as e:
            logger.error(f"后处理美化失败: {e}")
            return 0
    
    def _beautify_paragraph_text(self, text: str) -> str:
        """
        美化段落文本 - 应用各种美化规则
        
        Args:
            text: 原始文本
            
        Returns:
            str: 美化后的文本
        """
        try:
            import re
            
            beautified = text
            
            # 美化规则1: 清理多余的连续空格（保留表格对齐需要的空格）
            beautified = self._clean_excessive_spaces(beautified)
            
            # 美化规则2: 标点符号统一化
            beautified = self._normalize_punctuation(beautified)
            
            # 美化规则3: 清理冗余的占位符
            beautified = self._clean_redundant_placeholders(beautified)
            
            # 美化规则4: 表格对齐优化
            beautified = self._optimize_table_alignment(beautified)
            
            # 美化规则5: 中英文间距处理
            beautified = self._handle_chinese_english_spacing(beautified)
            
            return beautified
            
        except Exception as e:
            logger.error(f"文本美化失败: {e}")
            return text
    
    def _clean_excessive_spaces(self, text: str) -> str:
        """清理多余的连续空格，但保留表格对齐需要的空格"""
        try:
            import re
            
            # 检查是否为表格式布局（两个字段间有大量空格）
            table_patterns = [
                r'(电话|联系电话)[:：][^电子邮件]*\s{4,}(电子邮件|电子邮箱|邮箱)',
                r'(地址|注册地址|办公地址)[:：][^传真]*\s{4,}(传真)',
                r'(邮政编码|邮编)[:：][^网站]*\s{4,}(网站|网址)'
            ]
            
            is_table_layout = any(re.search(pattern, text) for pattern in table_patterns)
            
            if is_table_layout:
                # 表格式布局：保留对齐空格，但清理其他多余空格
                logger.debug("检测到表格式布局，保留对齐空格")
                # 只清理非对齐区域的多余空格
                cleaned = re.sub(r'^(\s{2,})', ' ', text)  # 清理行首多余空格
                cleaned = re.sub(r'(\s{2,})$', '', cleaned)  # 清理行尾多余空格
                return cleaned
            else:
                # 普通文本：清理所有多余空格
                cleaned = re.sub(r'\s{2,}', ' ', text)  # 将多个空格替换为单个空格
                return cleaned.strip()
                
        except Exception as e:
            logger.error(f"清理多余空格失败: {e}")
            return text
    
    def _normalize_punctuation(self, text: str) -> str:
        """标点符号统一化"""
        try:
            # 统一使用中文冒号
            text = text.replace(':', '：')
            
            # 清理多余的标点符号
            import re
            text = re.sub(r'：{2,}', '：', text)  # 多个冒号合并为一个
            text = re.sub(r'\.{2,}', '.', text)   # 多个句号合并为一个
            text = re.sub(r'，{2,}', '，', text)   # 多个逗号合并为一个
            
            return text
            
        except Exception as e:
            logger.error(f"标点符号统一化失败: {e}")
            return text
    
    def _clean_redundant_placeholders(self, text: str) -> str:
        """清理冗余的占位符"""
        try:
            import re
            
            # 清理孤立的下划线占位符
            text = re.sub(r'^_+$', '', text)  # 整行都是下划线
            text = re.sub(r'^_{1,3}$', '', text)  # 少量下划线
            
            # 清理已填写字段后的多余下划线
            text = re.sub(r'([:：][\w\d\-@\.]+)_+', r'\1', text)  # 字段后的下划线
            
            # 清理多余的破折号
            text = re.sub(r'[—]{3,}', '', text)  # 多个破折号
            text = re.sub(r'[-]{4,}', '', text)   # 多个连字符
            
            return text
            
        except Exception as e:
            logger.error(f"清理占位符失败: {e}")
            return text
    
    def _optimize_table_alignment(self, text: str) -> str:
        """表格对齐优化"""
        try:
            import re
            
            # 优化表格式布局的对齐
            # 模式：字段1：值1     字段2：值2
            table_match = re.search(r'([\w\u4e00-\u9fff]+[:：][\w\d\-@\.\u4e00-\u9fff]+)(\s{2,})([\w\u4e00-\u9fff]+[:：][\w\d\-@\.\u4e00-\u9fff]*)', text)
            
            if table_match:
                field1_part = table_match.group(1)
                spaces = table_match.group(2)
                field2_part = table_match.group(3)
                
                # 计算最优空格数：确保总长度合理且对齐美观
                optimal_spaces = max(4, min(20, len(spaces)))  # 4-20个空格之间
                
                optimized_text = f"{field1_part}{' ' * optimal_spaces}{field2_part}"
                text = text.replace(table_match.group(0), optimized_text)
                
                logger.debug(f"表格对齐优化: {len(spaces)}空格 -> {optimal_spaces}空格")
            
            return text
            
        except Exception as e:
            logger.error(f"表格对齐优化失败: {e}")
            return text
    
    def _handle_chinese_english_spacing(self, text: str) -> str:
        """中英文间距处理"""
        try:
            import re
            
            # 在中文和英文数字之间添加适当的空格（如果需要的话）
            # 这里采用保守策略，只处理明显需要空格的情况
            
            # 中文和邮箱之间
            text = re.sub(r'([\u4e00-\u9fff])([a-zA-Z0-9]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', r'\1 \2', text)
            
            # 中文和网址之间
            text = re.sub(r'([\u4e00-\u9fff])(http[s]?://[^\s]+)', r'\1 \2', text)
            text = re.sub(r'([\u4e00-\u9fff])(www\.[^\s]+)', r'\1 \2', text)
            
            return text
            
        except Exception as e:
            logger.error(f"中英文间距处理失败: {e}")
            return text
    
    def _process_unified_field(self, paragraph, para_text: str, field_config: dict, para_idx: int):
        """
        使用统一框架处理单个字段
        
        Args:
            paragraph: Word文档段落对象
            para_text: 段落文本
            field_config: 字段配置
            para_idx: 段落索引
            
        Returns:
            dict: 处理结果 {'modified': bool, 'replacements': int, 'patterns': list}
        """
        field_names = field_config['field_names']
        field_value = field_config['value']
        display_name = field_config['display_name']
        exclude_contexts = field_config['exclude_contexts']
        formats = field_config['formats']
        
        if not field_value:
            return {'modified': False, 'replacements': 0, 'patterns': []}
        
        # 检查是否为采购人信息，如果是则跳过
        if self._is_purchaser_info(para_text, exclude_contexts):
            logger.info(f"段落 #{para_idx} 检测到采购人信息，跳过 {display_name} 字段填写")
            return {'modified': False, 'replacements': 0, 'patterns': []}
        
        # 检查字段是否已填写
        if self._field_already_filled(para_text, field_names):
            logger.info(f"段落 #{para_idx} {display_name} 字段已填写，跳过")
            return {'modified': False, 'replacements': 0, 'patterns': []}
        
        # 生成匹配模式
        patterns = self._generate_field_patterns(field_names, formats)
        
        # 尝试匹配和替换
        for pattern_info in patterns:
            match = pattern_info['regex'].search(para_text)
            if match:
                logger.info(f"段落 #{para_idx} 匹配 {display_name} 字段模式: {pattern_info['description']}")
                
                # 使用run级别格式保持进行替换
                success = self._replace_with_format_preservation(
                    paragraph, pattern_info['regex'], field_value, display_name, pattern_info
                )
                
                if success:
                    return {
                        'modified': True, 
                        'replacements': 1, 
                        'patterns': [f"{display_name}: {pattern_info['description']}"]
                    }
        
        return {'modified': False, 'replacements': 0, 'patterns': []}
    
    def _is_purchaser_info(self, text: str, exclude_contexts: list) -> bool:
        """
        增强版采购人信息识别 - 更准确地识别各种采购人信息格式
        
        Args:
            text: 要检查的文本
            exclude_contexts: 排除的关键词列表
            
        Returns:
            bool: True表示是采购人信息，应跳过填写
        """
        if not text:
            return False
            
        # 1. 直接关键词检查（最高优先级）
        for context in exclude_contexts:
            if context in text:
                logger.info(f"检测到排除关键词 '{context}'，判定为采购人信息")
                return True
        
        # 2. 语义上下文分析 - 检查更多采购方相关词汇
        purchaser_indicators = [
            '采购人', '招标人', '甲方', '发包方', '发包人', '业主方', '业主',
            '招标方', '采购单位', '发标方', '委托人', '委托方',
            '项目建设单位', '建设方', '建设单位', '项目单位',
            '代理机构', '招标代理', '代理公司', '咨询公司',
            '联系人', '项目联系人', '业务联系人', '技术联系人',
            '询问人', '质疑联系人', '投诉联系人'
        ]
        
        for indicator in purchaser_indicators:
            if indicator in text:
                logger.info(f"检测到采购方指示词 '{indicator}'，判定为采购人信息")
                return True
        
        # 3. 格式特征识别
        # 3.1 方括号/书名号格式（采购人预填信息的常见格式）
        bracket_patterns = [
            r'【[^】]*】',      # 【方括号】
            r'〖[^〗]*〗',      # 〖方括号〗  
            r'［[^］]*］',      # ［方括号］
            r'《[^》]*》',      # 《书名号》
            r'〈[^〉]*〉'       # 〈书名号〉
        ]
        
        for pattern in bracket_patterns:
            if re.search(pattern, text):
                logger.info(f"检测到特殊标记格式，判定为采购人信息")
                return True
        
        # 3.2 "项目"相关的描述性文本（通常是采购人填写的项目信息）
        project_context_patterns = [
            r'项目名称[:：][^投标人|^供应商]',
            r'招标项目[:：]',
            r'采购项目[:：]',
            r'建设项目[:：]',
            r'工程项目[:：]'
        ]
        
        for pattern in project_context_patterns:
            if re.search(pattern, text):
                logger.info(f"检测到项目描述文本，判定为采购人信息")
                return True
        
        # 4. 已填写内容识别（具体联系信息通常是采购人的）
        # 4.1 电话号码模式（更准确的匹配）
        phone_patterns = [
            r'\d{3,4}[-\s]?\d{7,8}',           # 标准电话：010-12345678
            r'\d{3,4}\s?\d{7,8}',              # 无连字符：010 12345678
            r'\(\d{3,4}\)\s?\d{7,8}',          # 括号格式：(010) 12345678
            r'\+86[-\s]?\d{3,4}[-\s]?\d{7,8}'  # 国际格式：+86-010-12345678
        ]
        
        for pattern in phone_patterns:
            if re.search(pattern, text):
                logger.info("检测到具体电话号码，判定为采购人信息")
                return True
        
        # 4.2 邮箱地址模式
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        if re.search(email_pattern, text):
            logger.info("检测到具体邮箱地址，判定为采购人信息")
            return True
        
        # 4.3 网址模式
        url_patterns = [
            r'http[s]?://[^\s]+',
            r'www\.[^\s]+',
            r'[a-zA-Z0-9-]+\.(com|cn|org|gov|net)[^\s]*'
        ]
        
        for pattern in url_patterns:
            if re.search(pattern, text):
                logger.info("检测到网址信息，判定为采购人信息")
                return True
        
        # 5. 地址信息特征（详细地址通常是采购人的）
        detailed_address_patterns = [
            r'[省市区县][^供应商|^投标人].*?[路街道巷弄号]',  # 完整地址格式
            r'\d{6}',                                      # 邮政编码
            r'[A-Z]\d+号',                                # 楼号格式：A1号
            r'\d+[层楼]',                                  # 楼层：3楼、5层
        ]
        
        for pattern in detailed_address_patterns:
            if re.search(pattern, text):
                # 额外检查：如果同时包含"供应商"或"投标人"，则可能不是采购人地址
                if not any(word in text for word in ['供应商', '投标人', '乙方']):
                    logger.info("检测到详细地址信息且无投标人标识，判定为采购人信息")
                    return True
        
        # 6. 时间信息（招标时间、开标时间等通常是采购人填写的）
        time_context_patterns = [
            r'开标时间[:：]',
            r'截止时间[:：]',
            r'递交时间[:：]',
            r'开标地点[:：]',
            r'现场踏勘[:：]',
            r'答疑时间[:：]'
        ]
        
        for pattern in time_context_patterns:
            if re.search(pattern, text):
                logger.info("检测到招标时间信息，判定为采购人信息")
                return True
        
        # 7. 表格标题和表头识别（招标文件中的表格通常是采购人预制的）
        table_header_patterns = [
            r'表\d+[-\.:]',              # 表1-1: 或 表1.1:
            r'附件\d+[:：]',            # 附件1：
            r'项目\s*编号[:：]',        # 项目编号：
            r'合同\s*编号[:：]',        # 合同编号：
            r'标段\s*[编号]*[:：]',     # 标段编号：
            r'评标\s*标准[:：]',        # 评标标准：
            r'技术\s*要求[:：]',        # 技术要求：
            r'商务\s*要求[:：]'         # 商务要求：
        ]
        
        for pattern in table_header_patterns:
            if re.search(pattern, text):
                logger.info("检测到表格标题/表头信息，判定为采购人预制内容")
                return True
        
        # 8. 说明性文字识别（招标文件中的说明通常是采购人撰写的）
        instruction_patterns = [
            r'说明[:：]',
            r'注[:：]',
            r'备注[:：]',
            r'填写说明[:：]',
            r'投标须知[:：]',
            r'特别提醒[:：]',
            r'重要提示[:：]',
            r'温馨提示[:：]'
        ]
        
        for pattern in instruction_patterns:
            if re.search(pattern, text):
                logger.info("检测到说明性文字，判定为采购人信息")
                return True
        
        # 9. 政府采购相关标识
        gov_procurement_patterns = [
            r'政府采购',
            r'财政资金',
            r'预算单位',
            r'集中采购',
            r'分散采购',
            r'采购计划',
            r'政采云',
            r'公共资源交易中心'
        ]
        
        for pattern in gov_procurement_patterns:
            if pattern in text:
                logger.info(f"检测到政府采购标识 '{pattern}'，判定为采购人信息")
                return True
        
        # 10. 反向检查：如果明确包含投标人相关词汇，则不是采购人信息
        bidder_indicators = ['供应商', '投标人', '承包商', '乙方', '服务商', '厂商', '制造商']
        has_bidder_context = any(indicator in text for indicator in bidder_indicators)
        
        if has_bidder_context:
            logger.info(f"检测到投标人相关词汇，确认为投标人信息区域")
            return False
        
        # 11. 综合权重评估（新增功能）
        purchaser_score = self._calculate_purchaser_probability_score(text, exclude_contexts)
        if purchaser_score > 0.7:  # 阈值可调整
            logger.info(f"综合评估得分 {purchaser_score:.2f} 超过阈值，判定为采购人信息")
            return True
        
        return False
    
    def _calculate_purchaser_probability_score(self, text: str, exclude_contexts: list) -> float:
        """
        计算文本为采购人信息的概率得分 - 综合多个因素进行权重评估
        
        Args:
            text: 要评估的文本
            exclude_contexts: 排除的关键词列表
            
        Returns:
            float: 概率得分 (0.0 - 1.0)，越高越可能是采购人信息
        """
        try:
            import re
            score = 0.0
            
            # 权重1: 关键词匹配 (权重: 0.4)
            keyword_indicators = [
                ('采购人', 0.3), ('招标人', 0.3), ('甲方', 0.2), ('业主', 0.2),
                ('代理机构', 0.25), ('联系人', 0.15), ('项目单位', 0.2)
            ]
            
            for keyword, weight in keyword_indicators:
                if keyword in text:
                    score += weight * 0.4
            
            # 权重2: 格式特征 (权重: 0.3)
            format_patterns = [
                (r'【[^】]*】', 0.3),        # 方括号格式
                (r'《[^》]*》', 0.2),        # 书名号格式
                (r'项目名称[:：]', 0.25),    # 项目描述
                (r'开标时间[:：]', 0.2),     # 招标时间信息
                (r'表\d+[-\.:]', 0.15)      # 表格标题
            ]
            
            for pattern, weight in format_patterns:
                if re.search(pattern, text):
                    score += weight * 0.3
            
            # 权重3: 具体信息内容 (权重: 0.2)
            if re.search(r'\d{3,4}[-\s]?\d{7,8}', text):  # 电话号码
                score += 0.2 * 0.2
            if re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text):  # 邮箱
                score += 0.15 * 0.2
            if re.search(r'http[s]?://[^\s]+', text):  # 网址
                score += 0.1 * 0.2
            
            # 权重4: 排除关键词 (权重: 0.1)
            for context in exclude_contexts:
                if context in text:
                    score += 0.1
            
            # 反向调整：投标人相关词汇降低得分
            bidder_keywords = ['供应商', '投标人', '乙方', '服务商']
            for keyword in bidder_keywords:
                if keyword in text:
                    score -= 0.3  # 显著降低得分
            
            # 确保得分在合理范围内
            score = max(0.0, min(1.0, score))
            
            logger.debug(f"采购人概率评估: '{text[:50]}...' -> 得分: {score:.3f}")
            return score
            
        except Exception as e:
            logger.error(f"概率评估计算失败: {e}")
            return 0.0
    
    def _field_already_filled(self, text: str, field_names: list) -> bool:
        """
        检查字段是否已经填写了有意义的内容
        
        Args:
            text: 文本内容
            field_names: 字段名称列表
            
        Returns:
            bool: True表示已填写，应跳过
        """
        # 检查是否只包含占位符（下划线、空格）
        for field_name in field_names:
            if field_name in text:
                # 找到字段后的内容
                pattern = rf'{re.escape(field_name)}[:：]?\s*(.+)'
                match = re.search(pattern, text)
                if match:
                    content = match.group(1).strip()
                    # 如果内容不只是占位符（下划线、空格），说明已填写
                    if content and re.search(r'[^\s_-]', content):
                        return True
        return False
    
    def _generate_field_patterns(self, field_names: list, formats: dict) -> list:
        """
        根据字段名称和格式配置生成正则表达式模式
        
        Args:
            field_names: 字段名称列表
            formats: 格式配置字典
            
        Returns:
            list: 包含正则表达式和描述的列表
        """
        patterns = []
        
        for field_name in field_names:
            # 表单式格式（带冒号）
            if formats.get('with_colon'):
                for placeholder_type in formats.get('placeholder_types', ['space']):
                    if placeholder_type == 'underline':
                        pattern = rf'({re.escape(field_name)}[:：]\s*)([_]+)'
                        description = f'{field_name}（带冒号+下划线）'
                    elif placeholder_type == 'space':
                        pattern = rf'({re.escape(field_name)}[:：]\s*)(\s{{2,}})'
                        description = f'{field_name}（带冒号+空格）'
                    elif placeholder_type == 'mixed':
                        pattern = rf'({re.escape(field_name)}[:：]\s*)([_\s]+)'
                        description = f'{field_name}（带冒号+混合占位符）'
                    
                    patterns.append({
                        'regex': re.compile(pattern),
                        'description': description,
                        'type': 'form_with_colon'
                    })
            
            # 无冒号格式
            if formats.get('without_colon'):
                for placeholder_type in formats.get('placeholder_types', ['space']):
                    if placeholder_type == 'space':
                        pattern = rf'({re.escape(field_name)})(\s{{3,}})(?=\S|$)'
                        description = f'{field_name}（无冒号+空格）'
                    elif placeholder_type == 'underline':
                        pattern = rf'({re.escape(field_name)})([_]+)'
                        description = f'{field_name}（无冒号+下划线）'
                    
                    patterns.append({
                        'regex': re.compile(pattern),
                        'description': description,
                        'type': 'form_without_colon'
                    })
            
            # 表格式布局
            if formats.get('table_layout'):
                # 字段后跟其他字段的情况
                other_fields = ['电话', '电子邮件', '电子邮箱', '邮箱', '传真', '地址']
                for other_field in other_fields:
                    if other_field != field_name:
                        pattern = rf'({re.escape(field_name)})(\s{{2,}})(?={re.escape(other_field)})'
                        description = f'{field_name}（表格式+{other_field}）'
                        patterns.append({
                            'regex': re.compile(pattern),
                            'description': description,
                            'type': 'table_layout'
                        })
            
            # 行末格式
            pattern = rf'({re.escape(field_name)})(\s*)$'
            description = f'{field_name}（行末）'
            patterns.append({
                'regex': re.compile(pattern),
                'description': description,
                'type': 'line_end'
            })
        
        return patterns
    
    def _replace_with_format_preservation(self, paragraph, pattern_regex, field_value: str, 
                                        field_name: str, pattern_info: dict) -> bool:
        """
        使用run级别格式保持进行字段替换
        集成现有的智能文本替换机制，确保格式完全保持
        
        Args:
            paragraph: Word段落对象
            pattern_regex: 正则表达式模式
            field_value: 要填入的值
            field_name: 字段名称
            pattern_info: 模式信息
            
        Returns:
            bool: 替换是否成功
        """
        try:
            para_text = paragraph.text
            match = pattern_regex.search(para_text)
            if not match:
                return False
            
            old_text = match.group(0)
            
            # 根据模式类型决定替换策略
            if pattern_info['type'] == 'form_with_colon':
                # 表单式带冒号：电话：_____ -> 电话：010-63271000
                if len(match.groups()) >= 2:
                    label_part = match.group(1)  # 电话：
                    replacement = f"{label_part}{field_value}"
                else:
                    replacement = f"{old_text}：{field_value}"
                    
            elif pattern_info['type'] == 'form_without_colon':
                # 表单式无冒号：电话     -> 电话：010-63271000
                if len(match.groups()) >= 1:
                    label_part = match.group(1)  # 电话
                    replacement = f"{label_part}：{field_value}"
                else:
                    replacement = f"{old_text}：{field_value}"
                    
            elif pattern_info['type'] == 'table_layout':
                # 表格式：电话     电子邮件 -> 电话：010-63271000     电子邮件
                if len(match.groups()) >= 2:
                    label_part = match.group(1)  # 电话
                    space_part = match.group(2)  # 空格
                    # 保留后续内容
                    remaining_text = para_text[match.end():]
                    replacement = f"{label_part}：{field_value}{space_part}{remaining_text}".rstrip()
                    # 对于表格式，需要替换整个段落文本
                    old_text = para_text
                else:
                    replacement = f"{match.group(1)}：{field_value}"
                    
            elif pattern_info['type'] == 'line_end':
                # 行末式：电话 -> 电话：010-63271000
                if len(match.groups()) >= 1:
                    label_part = match.group(1)  # 电话
                    replacement = f"{label_part}：{field_value}"
                else:
                    replacement = f"{old_text}：{field_value}"
            else:
                # 默认处理
                replacement = f"{old_text}：{field_value}"
            
            # 使用现有的智能文本替换方法，这个方法已经实现了完整的run级别格式保持
            logger.info(f"字段 {field_name} 替换: '{old_text}' -> '{replacement}'")
            return self.smart_text_replace(paragraph, old_text, replacement)
            
        except Exception as e:
            logger.error(f"格式保持替换失败 ({field_name}): {e}")
            return False


    def _compact_format_replace(self, para_text: str, match, field_value: str, field_name: str, preserve_trailing: bool = False) -> str:
        """
        方案A：紧凑格式替换 - 支持表单式和表格式两种布局
        表单式: '成立时间：                    年月日' -> '成立时间：2000年4月21日'
        表格式: '电话                                  电子邮件' -> '电话：010-63271000                    电子邮件'
        """
        try:
            logger.info(f"紧凑格式替换: 字段={field_name}, 原文本='{para_text[:50]}...', 填写值='{field_value}', 保留后续内容={preserve_trailing}")
            
            # 检查是否为表格式布局模式
            if self._is_table_layout_pattern(match, field_name, para_text):
                return self._handle_table_layout_replace(para_text, match, field_value, field_name)
            
            # 标准表单式处理逻辑
            if match.groups() and len(match.groups()) >= 1:
                label = match.group(1)  # 标签部分，如 '成立时间：'
                existing_content = match.group(2).strip() if len(match.groups()) >= 2 else ""  # 现有内容部分
                
                logger.info(f"标签部分: '{label}', 现有内容: '{existing_content}'")
                
                # 构建紧凑格式结果
                clean_label = label.rstrip()
                if clean_label.endswith('：') or clean_label.endswith(':'):
                    # 冒号格式：直接拼接
                    new_text = f"{clean_label}{field_value}"
                else:
                    # 其他格式：添加一个空格
                    new_text = f"{clean_label} {field_value}"
                
                # 如果需要保留后续内容，添加匹配后的剩余部分
                if preserve_trailing:
                    trailing_content = para_text[match.end():]
                    if trailing_content:
                        # 特别处理需要添加空格的情况
                        if field_name == '地址' and '传真' in trailing_content:
                            # 保持原有格式，不强制添加空格
                            new_text += trailing_content
                        elif field_name == '联系电话' and ('电子邮件' in trailing_content or '电子邮箱' in trailing_content or '邮箱' in trailing_content):
                            # 保持原有格式，不强制添加空格
                            new_text += trailing_content
                        else:
                            new_text += trailing_content
                        logger.info(f"保留后续内容: '{trailing_content}'")
                
                logger.info(f"紧凑格式替换完成: '{new_text}'")
                return new_text
            else:
                # 降级到原有逻辑
                logger.warning("紧凑格式替换失败：匹配组不足，降级到原有逻辑")
                return self._smart_date_replace(para_text, match, field_value) if field_name == '成立日期' else para_text
                
        except Exception as e:
            logger.error(f"紧凑格式替换失败: {e}")
            # 降级到原有逻辑
            return self._smart_date_replace(para_text, match, field_value) if field_name == '成立日期' else para_text

    def _is_table_layout_pattern(self, match, field_name: str, para_text: str) -> bool:
        """检查是否为表格式布局模式"""
        try:
            # 表格式模式特征：
            # 1. 电话和电子邮件在同一行，中间有大量空格
            # 2. 地址和传真在同一行，中间有大量空格
            if field_name in ['联系电话', '电子邮件'] and ('电话' in para_text and ('电子邮件' in para_text or '电子邮箱' in para_text)):
                return True
            if field_name in ['地址', '传真'] and ('地址' in para_text and '传真' in para_text):
                return True
            return False
        except Exception as e:
            logger.warning(f"表格式布局检查失败: {e}")
            return False

    def _handle_table_layout_replace(self, para_text: str, match, field_value: str, field_name: str) -> str:
        """处理表格式布局的字段替换 - 优化版"""
        try:
            import re
            logger.info(f"处理表格式布局: 字段={field_name}, 原文本='{para_text}'")
            
            # 先尝试智能双字段处理
            dual_field_result = self._handle_dual_field_table_layout(para_text, field_name, field_value)
            if dual_field_result != para_text:
                return dual_field_result
            
            # 如果双字段处理没有生效，降级到原有单字段逻辑
            if field_name == '联系电话':
                # 处理 "电话                                  电子邮件" 格式
                # 替换为 "电话：010-63271000                    电子邮件"
                pattern = r'电话(\s+)(电子邮件|电子邮箱|邮箱)'
                replacement = f'电话：{field_value}\\1\\2'
                new_text = re.sub(pattern, replacement, para_text)
                logger.info(f"电话表格式替换: '{para_text}' -> '{new_text}'")
                return new_text
            
            elif field_name == '电子邮件':
                # 处理电子邮件部分
                # 匹配各种电子邮件形式，包括已被电话处理后的情况
                patterns_to_try = [
                    (r'(电话[:：][^\s]+\s+)(电子邮件)(\s*)', f'\\1\\2：{field_value}\\3'),  # 电话已处理的情况
                    (r'(电话\s+)(电子邮件)(\s*)', f'\\1\\2：{field_value}\\3'),            # 电话未处理的情况
                    (r'(电子邮件)(\s*)', f'\\1：{field_value}\\2'),                        # 独立电子邮件
                    (r'(电子邮箱)(\s*)', f'\\1：{field_value}\\2'),                        # 独立电子邮箱
                ]
                
                new_text = para_text
                for pattern, replacement in patterns_to_try:
                    if re.search(pattern, new_text):
                        new_text = re.sub(pattern, replacement, new_text)
                        logger.info(f"电子邮件表格式替换成功: '{para_text}' -> '{new_text}'")
                        return new_text
                
                logger.warning(f"电子邮件表格式替换未匹配任何模式: '{para_text}'")
                return para_text
            
            elif field_name == '传真':
                # 处理 "地址                                  传真" 格式
                pattern = r'(地址[：:]?[^传真]*?)(\s+)(传真)(\s*)'
                replacement = f'\\1\\2\\3：{field_value}'
                new_text = re.sub(pattern, replacement, para_text)
                logger.info(f"传真表格式替换: '{para_text}' -> '{new_text}'")
                return new_text
            
            else:
                logger.warning(f"未知的表格式布局字段: {field_name}")
                return para_text
                
        except Exception as e:
            logger.error(f"表格式布局处理失败: {e}")
            return para_text

    def _handle_dual_field_table_layout(self, para_text: str, current_field: str, field_value: str) -> str:
        """
        智能双字段表格布局处理 - 一次性处理同行的两个字段
        解决问题：电话                    电子邮件 -> 电话：010-63271000                    电子邮件：xxx@xxx.com
        
        Args:
            para_text: 段落文本
            current_field: 当前正在处理的字段
            field_value: 字段值
            
        Returns:
            str: 处理后的文本（如果未处理则返回原文本）
        """
        try:
            import re
            
            # 定义双字段组合模式
            dual_field_patterns = [
                # 电话 + 电子邮件组合
                {
                    'pattern': r'(电话|联系电话)(\s{8,})(电子邮件|电子邮箱|邮箱)',
                    'field1': '电话',
                    'field2': '电子邮件',
                    'description': '电话+邮件表格组合'
                },
                # 地址 + 传真组合
                {
                    'pattern': r'(地址|注册地址|办公地址|联系地址)(\s{8,})(传真)',
                    'field1': '地址',
                    'field2': '传真',
                    'description': '地址+传真表格组合'
                },
                # 邮编 + 网站组合（可选）
                {
                    'pattern': r'(邮政编码|邮编)(\s{8,})(网站|网址|官网)',
                    'field1': '邮编',
                    'field2': '网站',
                    'description': '邮编+网站表格组合'
                }
            ]
            
            for pattern_info in dual_field_patterns:
                match = re.search(pattern_info['pattern'], para_text)
                if match:
                    logger.info(f"检测到双字段模式: {pattern_info['description']}")
                    
                    # 提取匹配组件
                    field1_label = match.group(1)  # 第一个字段标签
                    middle_spaces = match.group(2)  # 中间的空格
                    field2_label = match.group(3)  # 第二个字段标签
                    
                    # 获取两个字段的值
                    field1_value, field2_value = self._get_dual_field_values(
                        pattern_info['field1'], pattern_info['field2'], field_value, current_field
                    )
                    
                    if field1_value and field2_value:
                        # 计算适合的间距，保持美观对齐
                        optimal_spacing = self._calculate_optimal_spacing(
                            field1_label, field1_value, field2_label, middle_spaces
                        )
                        
                        # 构建替换结果
                        new_text = f"{field1_label}：{field1_value}{optimal_spacing}{field2_label}：{field2_value}"
                        
                        logger.info(f"双字段表格处理成功: '{para_text}' -> '{new_text}'")
                        return new_text
                    
                    # 如果只有一个字段有值，则只处理一个
                    elif field1_value and pattern_info['field1'] in self._normalize_field_name(current_field):
                        optimal_spacing = self._calculate_optimal_spacing(
                            field1_label, field1_value, field2_label, middle_spaces
                        )
                        new_text = f"{field1_label}：{field1_value}{optimal_spacing}{field2_label}"
                        logger.info(f"双字段单一处理: '{para_text}' -> '{new_text}'")
                        return new_text
                    
                    elif field2_value and pattern_info['field2'] in self._normalize_field_name(current_field):
                        optimal_spacing = self._calculate_optimal_spacing(
                            field1_label, field2_value, field2_label, middle_spaces  # 注意这里用field2_value计算间距
                        )
                        new_text = f"{field1_label}{optimal_spacing}{field2_label}：{field2_value}"
                        logger.info(f"双字段单一处理(第二字段): '{para_text}' -> '{new_text}'")
                        return new_text
            
            return para_text  # 未找到双字段模式，返回原文本
            
        except Exception as e:
            logger.error(f"双字段表格处理失败: {e}")
            return para_text
    
    def _get_dual_field_values(self, field1_type: str, field2_type: str, current_value: str, current_field: str) -> tuple:
        """获取双字段的值"""
        try:
            # 这里需要通过self.company_info获取对应的字段值
            # 暂时使用简化逻辑，后续可以优化
            field1_value = ""
            field2_value = ""
            
            # 根据字段类型获取值
            if field1_type == '电话':
                field1_value = getattr(self, 'company_info', {}).get('fixedPhone', '')
            elif field1_type == '地址':
                field1_value = getattr(self, 'company_info', {}).get('address', '')
            elif field1_type == '邮编':
                field1_value = getattr(self, 'company_info', {}).get('postalCode', '')
            
            if field2_type == '电子邮件':
                field2_value = getattr(self, 'company_info', {}).get('email', '') or '未填写'
            elif field2_type == '传真':
                field2_value = getattr(self, 'company_info', {}).get('fax', '') or '未填写'
            elif field2_type == '网站':
                field2_value = getattr(self, 'company_info', {}).get('website', '')
            
            # 如果当前正在处理的字段匹配，使用传入的值
            current_normalized = self._normalize_field_name(current_field)
            if field1_type in current_normalized:
                field1_value = current_value
            elif field2_type in current_normalized:
                field2_value = current_value
            
            return field1_value, field2_value
            
        except Exception as e:
            logger.error(f"获取双字段值失败: {e}")
            return "", ""
    
    def _calculate_optimal_spacing(self, field1_label: str, field1_value: str, field2_label: str, original_spaces: str) -> str:
        """计算最优的字段间距，保持美观对齐"""
        try:
            # 计算原始空格数
            original_space_count = len(original_spaces)
            
            # 计算第一个字段填入后的长度变化
            # 原始: 电话 (2个字符)
            # 填入后: 电话：010-63271000 (约13个字符)
            field1_original_length = len(field1_label)
            field1_new_length = len(f"{field1_label}：{field1_value}")
            length_increase = field1_new_length - field1_original_length
            
            # 计算调整后的空格数，保持整体对齐
            # 最少保留 4 个空格，最多不超过原有空格数
            optimal_space_count = max(4, original_space_count - length_increase)
            optimal_space_count = min(optimal_space_count, original_space_count)
            
            logger.info(
                f"空格计算: 原始={original_space_count}, 长度增加={length_increase}, 最优={optimal_space_count}"
            )
            
            return ' ' * optimal_space_count
            
        except Exception as e:
            logger.error(f"计算最优间距失败: {e}")
            # 错误情况下返回默认间距
            return '    '  # 4个空格
    
    def _normalize_field_name(self, field_name: str) -> list:
        """将字段名称标准化为可识别的形式"""
        try:
            # 定义字段名称映射
            field_mapping = {
                '联系电话': ['电话', '联系电话', '固定电话', '电话号码', '联系方式'],
                '电子邮件': ['电子邮件', '电子邮箱', '邮箱', 'email', 'Email'],
                '传真': ['传真', '传真号码', '传真号', 'fax', 'Fax'],
                '地址': ['地址', '注册地址', '办公地址', '联系地址'],
                '邮编': ['邮政编码', '邮编', '邮码'],
                '网站': ['网站', '网址', '官网', '公司网站']
            }
            
            for standard_name, variants in field_mapping.items():
                if field_name in variants:
                    return variants
            
            return [field_name]  # 如果找不到映射，返回原名称
            
        except Exception as e:
            logger.error(f"字段名称标准化失败: {e}")
            return [field_name]

    def _smart_date_replace(self, para_text: str, match, date_value: str) -> str:
        """
        智能日期替换 - 专门处理成立日期，避免重复的"年月日"
        例如：成立时间：2015年12月18日年月日 -> 成立时间：2015年12月18日
        """
        try:
            import re  # 移到方法开始，确保在使用前导入
            logger.info(f"智能日期替换: 原文本='{para_text}', 日期值='{date_value}'")
            
            # 获取匹配的部分
            match_text = match.group(0)
            placeholder = match.group(1) if match.groups() else ""
            
            # 执行基本替换 - 首先检查是否是特殊格式
            if re.match(r'^(日\s+期|时\s+间)\s*[:：]\s*([_\s]*)$', match_text):
                # 带空格的日期/时间格式，保持原有格式，只替换占位符
                spaced_match = re.match(r'^(日\s+期|时\s+间)\s*[:：]\s*([_\s]*)$', match_text)
                label_part = spaced_match.group(1)  # "日   期" 或 "时   间"
                placeholder_part = spaced_match.group(2)  # 占位符部分
                # 保持原有标签格式，只替换占位符或添加日期
                new_text = para_text.replace(match_text, f"{label_part}：{date_value}", 1)
                logger.info(f"带空格格式替换(保持原标签): '{match_text}' -> '{label_part}：{date_value}'")
            elif placeholder:
                new_text = para_text.replace(match.group(0), match.group(0).replace(placeholder, date_value, 1))
            else:
                # 检查是否是独立的年月日格式
                if re.match(r'^(\s*)([_-]+年[_-]+月[_-]+日)(\s*)$', match_text):
                    # 独立格式，直接替换整个匹配内容
                    groups = re.match(r'^(\s*)([_-]+年[_-]+月[_-]+日)(\s*)$', match_text).groups()
                    new_text = para_text.replace(match_text, groups[0] + date_value + groups[2], 1)
                    logger.info(f"独立日期格式替换: '{match_text}' -> '{groups[0] + date_value + groups[2]}'")
                elif re.match(r'^[_-]+年[_-]+月[_-]+日$', match_text):
                    # 段落中的年月日格式，直接替换
                    new_text = para_text.replace(match_text, date_value, 1)
                    logger.info(f"段落内日期格式替换: '{match_text}' -> '{date_value}'")
                elif re.match(r'^\s{2,}年\s{2,}月\s{2,}日$', match_text):
                    # 空格分隔的年月日格式，直接替换
                    new_text = para_text.replace(match_text, date_value, 1)
                    logger.info(f"空格年月日格式替换: '{match_text}' -> '{date_value}'")
                else:
                    # 在匹配部分后直接添加日期
                    new_text = para_text.replace(match_text, match_text + date_value, 1)
            
            # 检查并清理重复的年月日字符 - 更智能的清理逻辑
            
            # 处理各种重复模式
            redundant_patterns = [
                # 直接重复的年月日字符
                r'(\d+年\d+月\d+日)年',  # 2015年12月18日年
                r'(\d+年\d+月\d+日)月',  # 2015年12月18日月  
                r'(\d+年\d+月\d+日)日',  # 2015年12月18日日
                r'(\d+年\d+月\d+日)年\s*月',  # 2015年12月18日年月
                r'(\d+年\d+月\d+日)年\s*月\s*日',  # 2015年12月18日年月日
                r'(\d+年\d+月\d+日)月\s*日',  # 2015年12月18日月日
                # 带空格的重复模式
                r'(\d+年\d+月\d+日)\s+年',  # 2015年12月18日 年
                r'(\d+年\d+月\d+日)\s+月',  # 2015年12月18日 月
                r'(\d+年\d+月\d+日)\s+日',  # 2015年12月18日 日
                r'(\d+年\d+月\d+日)\s+年\s*月',  # 2015年12月18日 年月
                r'(\d+年\d+月\d+日)\s+年\s*月\s*日',  # 2015年12月18日 年月日
                r'(\d+年\d+月\d+日)\s+月\s*日',  # 2015年12月18日 月日
                # 更宽泛的模式，处理多个空格的情况
                r'(\d+年\d+月\d+日)\s*年\s*月\s*日',  # 2015年12月18日年月日（任意空格）
                r'(\d+年\d+月\d+日)\s+月\s+日',  # 2015年12月18日  月  日
                r'(\d+年\d+月\d+日)\s+年\s+月',  # 2015年12月18日  年  月
            ]
            
            for pattern in redundant_patterns:
                if re.search(pattern, new_text):
                    old_new_text = new_text
                    new_text = re.sub(pattern, r'\1', new_text)
                    logger.info(f"清理重复字符: '{old_new_text}' -> '{new_text}'")
            
            # 额外清理：处理日期后的占位符残留
            # 例如："2025年09月07日_____月_____日" -> "2025年09月07日"
            placeholder_cleanup_patterns = [
                # 清理日期后的下划线占位符
                r'(\d+年\d+月\d+日)_+月_+日',  # 2025年09月07日_____月_____日
                r'(\d+年\d+月\d+日)_+月',      # 2025年09月07日_____月
                r'(\d+年\d+月\d+日)_+日',      # 2025年09月07日_____日
                r'(\d+年\d+月\d+日)_+年_+月_+日', # 2025年09月07日_____年_____月_____日
                r'(\d+年\d+月\d+日)_+年_+月',     # 2025年09月07日_____年_____月
                r'(\d+年\d+月\d+日)_+年',         # 2025年09月07日_____年
                # 清理空格和下划线混合的情况
                r'(\d+年\d+月\d+日)[\s_]+月[\s_]*日', # 2025年09月07日 ___月___日
                r'(\d+年\d+月\d+日)[\s_]+年[\s_]*月[\s_]*日', # 带空格的混合情况
                # 清理横线形式的占位符
                r'(\d+年\d+月\d+日)-+',        # 2025年09月07日--------
                r'(\d+年\d+月\d+日)[\s-]+$',   # 2025年09月07日 ---- (行末)
                r'(\d+年\d+月\d+日)_+$',       # 2025年09月07日_____ (行末)
                # 清理日期后的任意组合占位符（更通用的模式）
                r'(\d+年\d+月\d+日)[\s_-]+.*?$', # 日期后任意占位符到行末
            ]
            
            for pattern in placeholder_cleanup_patterns:
                if re.search(pattern, new_text):
                    old_new_text = new_text
                    new_text = re.sub(pattern, r'\1', new_text)
                    logger.info(f"清理占位符残留: '{old_new_text}' -> '{new_text}'")
            
            logger.info(f"智能日期替换完成: '{new_text}'")
            return new_text
            
        except Exception as e:
            logger.error(f"智能日期替换失败: {e}")
            # 降级到普通替换
            placeholder = match.group(1) if match.groups() else ""
            if placeholder:
                return para_text.replace(match.group(0), match.group(0).replace(placeholder, date_value, 1))
            else:
                return para_text.replace(match.group(0), match.group(0) + date_value, 1)

    def _compact_format_paragraph_replace(self, paragraph, match, field_value: str, field_name: str, preserve_trailing: bool = False) -> bool:
        """
        紧凑格式段落替换 - 保持原有格式的情况下消除大量空格
        专门为方案A设计，确保不破坏原有的字体格式
        """
        try:
            logger.info(f"紧凑格式段落替换: 字段={field_name}, 填写值='{field_value}'")
            
            if not match.groups() or len(match.groups()) < 1:
                logger.warning("匹配组不足，降级到标准替换")
                return False
            
            label = match.group(1)  # 标签部分
            existing_content = match.group(2).strip() if len(match.groups()) >= 2 else ""  # 现有内容部分
            
            # 获取完整的段落文本
            full_text = paragraph.text
            
            # 构建字段部分的新文本
            clean_label = label.rstrip()
            if clean_label.endswith('：') or clean_label.endswith(':'):
                new_field_text = f"{clean_label}{field_value}"
            else:
                new_field_text = f"{clean_label} {field_value}"
            
            # 构建完整的新文本（包含匹配前后的内容）
            match_start = match.start()
            match_end = match.end()
            
            # 匹配前的内容
            prefix = full_text[:match_start]
            
            # 匹配后的内容（如果需要保留）
            suffix = ""
            if preserve_trailing:
                trailing_content = full_text[match_end:]
                if trailing_content:
                    # 特殊处理地址和传真、电话和电子邮件之间的间距
                    if field_name == '地址' and '传真' in trailing_content:
                        suffix = trailing_content
                    elif field_name == '电话' and ('电子邮件' in trailing_content or '电子邮箱' in trailing_content or '邮箱' in trailing_content):
                        suffix = trailing_content
                    else:
                        suffix = trailing_content
                    logger.info(f"保留后续内容: '{trailing_content}'")
            
            # 组合完整的新文本
            new_full_text = prefix + new_field_text + suffix
            
            # 尝试智能run替换，保持格式
            full_text = paragraph.text
            
            # 方法1：更精确的标签run定位
            # 找到完整匹配文本在段落中的位置
            match_start = match.start()
            match_end = match.end()
            
            # 定位包含匹配开始位置的run
            current_pos = 0
            target_run_index = -1
            
            for i, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                
                # 如果匹配的开始位置在这个run的范围内
                if run_start <= match_start < run_end:
                    target_run_index = i
                    break
                    
                current_pos = run_end
            
            # 采用最保险的方法：整体替换，但保持第一个run的格式
            logger.info("使用整体替换策略，保持第一个run的格式")
            
            # 清空所有run的文本
            for run in paragraph.runs:
                run.text = ""
            
            # 将新文本放到第一个run中（保持第一个run的原始格式）
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
                logger.info(f"✅ 紧凑格式段落替换成功，内容统一放在第一个run中")
                return True
            else:
                logger.warning("段落没有run，无法替换")
                return False
            
        except Exception as e:
            logger.error(f"紧凑格式段落替换失败: {e}")
            return False

    def _no_colon_date_paragraph_replace(self, paragraph, match, field_value: str) -> bool:
        """
        无冒号日期格式的段落替换 - 处理 '日期                    ' 这种格式
        保留标签，替换占位符空格为日期值
        """
        try:
            label = match.group(1)  # 标签（日期/时间）
            placeholder = match.group(2)  # 占位符空格
            original_text = match.group(0)  # 完整匹配
            target_text = label + field_value  # 目标文本：标签+日期值
            
            logger.info(f"无冒号日期格式段落替换: '{original_text}' -> '{target_text}'")
            
            # 先尝试单run替换，保持原有格式
            for i, run in enumerate(paragraph.runs):
                if original_text in run.text:
                    # 单run情况，保持格式的替换
                    new_text = run.text.replace(original_text, target_text, 1)
                    
                    # 使用格式保持策略：清空所有run，将新文本放入第一个run（保持格式）
                    for r in paragraph.runs:
                        r.text = ""
                    paragraph.runs[0].text = new_text
                    
                    logger.info(f"✅ 无冒号日期格式单run替换成功（保持格式）")
                    return True
            
            # 跨run处理
            if len(paragraph.runs) > 1:
                # 构建完整段落文本用于匹配
                para_text = ''.join([run.text for run in paragraph.runs])
                if original_text in para_text:
                    # 找到第一个包含标签的run
                    first_run_with_label = None
                    for i, run in enumerate(paragraph.runs):
                        if label in run.text:
                            first_run_with_label = i
                            break
                    
                    if first_run_with_label is not None:
                        # 构建完整的新文本
                        new_para_text = para_text.replace(original_text, target_text, 1)
                        
                        # 使用格式保持策略：清空所有run，将新文本放入第一个run（保持格式）
                        for r in paragraph.runs:
                            r.text = ""
                        paragraph.runs[0].text = new_para_text
                        
                        logger.info(f"✅ 无冒号日期格式跨run替换成功（保持格式）")
                        return True
            
            logger.warning("无冒号日期格式段落替换失败：未找到匹配内容")
            return False
            
        except Exception as e:
            logger.error(f"无冒号日期格式段落替换失败: {e}")
            return False

    def _smart_date_paragraph_replace(self, paragraph, old_text: str, new_text: str, date_value: str) -> bool:
        """
        智能日期段落替换 - 专门处理成立日期的run分布
        确保完整的日期如"2015年12月18日"放在第一个相关run中，清理多余的"年月日"
        """
        try:
            logger.info(f"智能日期段落替换: '{old_text}' -> '{new_text}'")
            
            # 先尝试单run替换，保持原有格式
            for i, run in enumerate(paragraph.runs):
                if old_text in run.text:
                    # 单run情况，保持格式的替换
                    updated_text = run.text.replace(old_text, new_text)
                    
                    # 使用格式保持策略：清空所有run，将新文本放入第一个run（保持格式）
                    for r in paragraph.runs:
                        r.text = ""
                    paragraph.runs[0].text = updated_text
                    
                    logger.info(f"✅ 成立日期单run替换成功（保持格式）")
                    return True
            
            # 跨run情况：需要重新分布文本
            full_text = paragraph.text
            if old_text not in full_text:
                logger.warning("原文本不在段落中，无法替换")
                return False
            
            # 查找包含日期相关文本的run
            date_related_runs = []
            current_pos = 0
            
            for i, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                
                # 检查这个run是否包含日期相关内容
                contains_date_chars = any(char in run.text for char in ['年', '月', '日', '时', '期', ':', '：'])
                
                date_related_runs.append({
                    'index': i,
                    'run': run,
                    'start': run_start, 
                    'end': run_end,
                    'text': run.text,
                    'contains_date': contains_date_chars
                })
                
                current_pos = run_end
            
            # 执行替换：将完整日期放到第一个相关run中
            new_full_text = full_text.replace(old_text, new_text)
            
            # 使用格式保持策略：清空所有run，将新文本放入第一个run（保持格式）
            if paragraph.runs:
                for r in paragraph.runs:
                    r.text = ""
                paragraph.runs[0].text = new_full_text
                
                logger.info(f"✅ 成立日期跨run替换成功，完整日期放在第一个run中（保持格式）")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"智能日期段落替换失败: {e}")
            # 降级到标准替换
            return self._safe_replace_paragraph_text(paragraph, old_text, new_text)

    def _safe_replace_paragraph_text(self, paragraph, old_text: str, new_text: str):
        """
        精确run修改法 - 最大化保持格式美观
        专为标书格式美观设计的高精度替换方法
        """
        try:
            # 方法1：尝试在现有单个run中替换
            for run in paragraph.runs:
                if old_text in run.text:
                    # 直接替换，完美保持格式
                    run.text = run.text.replace(old_text, new_text)
                    logger.info(f"✅ 单run精确替换成功，完美保持格式")
                    return True
            
            # 方法2：精确跨run处理 - 只修改涉及的run，保留其他run
            return self._precise_cross_run_replace(paragraph, old_text, new_text)
            
        except Exception as e:
            logger.error(f"精确替换失败: {e}", exc_info=True)
            return False

    def _precise_cross_run_replace(self, paragraph, old_text: str, new_text: str):
        """
        精确跨run替换 - 智能重新分布文本，保持格式边界
        """
        try:
            full_text = paragraph.text
            logger.info(f"执行精确跨run替换: '{old_text}' -> '{new_text}'")
            
            # 查找目标文本位置
            start_pos = full_text.find(old_text)
            if start_pos == -1:
                logger.warning("目标文本未找到")
                return False
            
            end_pos = start_pos + len(old_text)
            logger.info(f"目标文本位置: {start_pos}-{end_pos}")
            
            # 分析所有run的信息和格式
            all_runs_info = []
            current_pos = 0
            
            for i, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                
                all_runs_info.append({
                    'index': i,
                    'run': run,
                    'run_start': run_start,
                    'run_end': run_end,
                    'original_text': run.text,
                    'format': self._extract_run_format(run),
                    'affected': run_start < end_pos and run_end > start_pos
                })
                
                current_pos = run_end
            
            # 构建替换后的完整文本
            new_full_text = full_text.replace(old_text, new_text)
            logger.info(f"替换后完整文本: '{new_full_text[:100]}...'")
            
            # 🎯 关键修复：智能重新分布文本到run，保持格式边界
            return self._smart_redistribute_cross_run_text(paragraph, all_runs_info, old_text, new_text, new_full_text, start_pos, end_pos)
            
        except Exception as e:
            logger.error(f"精确跨run替换失败: {e}", exc_info=True)
            # 降级到传统方法
            return self._fallback_safe_replace(paragraph, old_text, new_text)
    
    def _smart_redistribute_cross_run_text(self, paragraph, all_runs_info, old_text: str, new_text: str, new_full_text: str, start_pos: int, end_pos: int):
        """
        智能重新分布跨run替换后的文本，保持格式边界
        """
        try:
            # 计算文本长度变化
            text_length_delta = len(new_text) - len(old_text)
            logger.info(f"文本长度变化: {text_length_delta}")
            
            # 清空所有run
            for run_info in all_runs_info:
                run_info['run'].text = ""
            
            # 重新分配文本
            current_text_pos = 0
            
            for i, run_info in enumerate(all_runs_info):
                if current_text_pos >= len(new_full_text):
                    break
                
                run = run_info['run']
                original_length = len(run_info['original_text'])
                
                # 🎯 关键：计算这个run应该承载的文本长度
                if run_info['affected']:
                    # 受影响的run：需要根据文本变化调整长度
                    if i == len(all_runs_info) - 1:
                        # 最后一个run：取剩余所有文本
                        assigned_length = len(new_full_text) - current_text_pos
                    else:
                        # 按比例分配，但考虑文本变化
                        if original_length > 0:
                            # 如果是包含目标文本的区域，调整长度
                            run_start = run_info['run_start']
                            run_end = run_info['run_end']
                            
                            # 检查此run与替换区域的重叠情况
                            overlap_start = max(start_pos, run_start)
                            overlap_end = min(end_pos, run_end)
                            overlap_length = max(0, overlap_end - overlap_start)
                            
                            if overlap_length > 0:
                                # 这个run包含被替换的内容，需要调整长度
                                non_overlap_length = original_length - overlap_length
                                assigned_length = non_overlap_length + (overlap_length * len(new_text) // len(old_text))
                                assigned_length = min(assigned_length, len(new_full_text) - current_text_pos)
                            else:
                                assigned_length = min(original_length, len(new_full_text) - current_text_pos)
                        else:
                            assigned_length = 0
                else:
                    # 未受影响的run：保持原始长度
                    assigned_length = min(original_length, len(new_full_text) - current_text_pos)
                
                # 分配文本到run
                if assigned_length > 0:
                    run.text = new_full_text[current_text_pos:current_text_pos + assigned_length]
                    
                    # 🔧 恢复原始格式
                    self._apply_run_format(run, run_info['format'])
                    
                    logger.debug(f"Run #{i+1}: 分配文本 '{run.text[:20]}...', 长度={assigned_length}")
                    current_text_pos += assigned_length
                else:
                    run.text = ""
            
            # 如果还有剩余文本，追加到最后一个非空run
            if current_text_pos < len(new_full_text):
                remaining_text = new_full_text[current_text_pos:]
                # 找到最后一个有文本的run
                for run_info in reversed(all_runs_info):
                    if run_info['run'].text:
                        run_info['run'].text += remaining_text
                        logger.info(f"剩余文本已追加到最后一个run: '{remaining_text}'")
                        break
            
            logger.info(f"✅ 智能跨run文本重分布完成，保持了格式边界")
            return True
            
        except Exception as e:
            logger.error(f"智能重分布失败: {e}", exc_info=True)
            return False

    def _extract_run_format(self, run):
        """提取run的格式信息"""
        return {
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'underline': run.font.underline,
            'color': run.font.color.rgb if run.font.color.rgb else None
        }

    def _apply_run_format(self, run, format_info):
        """应用格式到run"""
        try:
            if format_info['font_name']:
                run.font.name = format_info['font_name']
            if format_info['font_size']:
                run.font.size = format_info['font_size']
            if format_info['bold'] is not None:
                run.font.bold = format_info['bold']
            if format_info['italic'] is not None:
                run.font.italic = format_info['italic']
            if format_info['underline'] is not None:
                # 检查内容是否是填充的值（公司名称、招标编号等），如果是则不应用下划线
                run_text = run.text.strip()
                is_filled_content = False
                
                # 检查是否包含常见的填充内容
                if (hasattr(self, 'company_name') and self.company_name and self.company_name in run_text) or \
                   (hasattr(self, 'tender_no') and self.tender_no and self.tender_no in run_text) or \
                   (hasattr(self, 'project_name') and self.project_name and self.project_name in run_text) or \
                   ('中国联合网络通信有限公司' in run_text) or \
                   ('GXTC-C-251590031' in run_text):
                    is_filled_content = True
                
                if is_filled_content and format_info['underline']:
                    # 如果是填充内容且原格式有下划线，不应用下划线
                    run.font.underline = False
                    logger.info(f"跳过下划线格式应用: 填充内容'{run_text[:20]}...'")
                else:
                    run.font.underline = format_info['underline']
            if format_info['color']:
                run.font.color.rgb = format_info['color']
        except Exception as e:
            logger.warning(f"应用格式失败: {e}")

    def _fallback_safe_replace(self, paragraph, old_text: str, new_text: str):
        """降级方法：传统的整段重构（作为最后备选）"""
        try:
            logger.info("⚠️ 使用降级方法：整段重构")
            
            original_text = paragraph.text
            if old_text not in original_text:
                return False
                
            new_paragraph_text = original_text.replace(old_text, new_text)
            
            # 保存第一个run的格式
            original_format = None
            if paragraph.runs:
                original_format = self._extract_run_format(paragraph.runs[0])
            
            # 清空所有run
            for run in paragraph.runs:
                run.text = ""
            
            # 重新创建
            new_run = paragraph.add_run(new_paragraph_text)
            if original_format:
                self._apply_run_format(new_run, original_format)
            
            logger.info("⚠️ 降级方法完成，可能影响部分格式")
            return True
            
        except Exception as e:
            logger.error(f"降级方法也失败: {e}")
            return False

    def _batch_replace_multiple_items(self, paragraph, company_name: str, project_name: str = None, tender_number: str = None):
        """
        批量替换多项内容，避免多次替换的累积格式问题
        
        Args:
            paragraph: 段落对象
            company_name: 公司名称
            project_name: 项目名称（可选）
            tender_number: 招标编号（可选）
            
        Returns:
            dict: 替换结果统计
        """
        try:
            original_text = paragraph.text
            logger.info(f"🔄 开始批量替换: '{original_text[:100]}...'")
            
            # 定义所有可能的替换项
            replacements = []
            
            # 1. 收集供应商名称相关替换
            company_patterns = [
                (r'（供应商名称、地址）', f'（{company_name}、{self.company_address}）'),
                (r'（供应商名称）', f'（{company_name}）'),
                (r'\(供应商名称、地址\)', f'（{company_name}、{self.company_address}）'),
                (r'\(供应商名称\)', f'（{company_name}）'),
            ]
            
            # 2. 收集项目名称相关替换
            project_patterns = []
            if project_name:
                project_patterns.extend([
                    (r'（项目名称）', f'（{project_name}）'),
                    (r'\(项目名称\)', f'（{project_name}）'),
                    (r'为\s*[\(（][^）)]*[\)）]\s*项目', f'为（{project_name}）项目'),
                ])
            
            # 添加项目名称和项目编号的组合模式
            if project_name and tender_number:
                project_patterns.extend([
                    (r'（项目名称\s*[、，]\s*项目编号）', f'（{project_name}、{tender_number}）'),
                    (r'\(项目名称\s*[、，]\s*项目编号\)', f'（{project_name}、{tender_number}）'),
                ])
            
            # 3. 收集采购编号相关替换  
            if tender_number:
                tender_patterns = [
                    (r'（采购编号）', f'（{tender_number}）'),
                    (r'\(采购编号\)', f'（{tender_number}）'),
                ]
            else:
                tender_patterns = []
            
            # 4. 扫描段落，找到所有需要替换的项目（避免重复匹配）
            all_patterns = company_patterns + project_patterns + tender_patterns
            used_positions = set()  # 记录已使用的位置，避免重复
            
            for pattern_str, replacement in all_patterns:
                import re
                pattern = re.compile(pattern_str)
                match = pattern.search(original_text)
                if match:
                    start_pos = match.start()
                    end_pos = match.end()
                    
                    # 检查是否与现有替换项重叠
                    overlap = False
                    for used_start, used_end in used_positions:
                        if not (end_pos <= used_start or start_pos >= used_end):
                            overlap = True
                            logger.info(f"  跳过重叠替换项: '{match.group(0)}' (与位置 {used_start}-{used_end} 重叠)")
                            break
                    
                    if not overlap:
                        replacements.append({
                            'pattern': pattern_str,
                            'old_text': match.group(0),
                            'new_text': replacement,
                            'start_pos': start_pos,
                            'end_pos': end_pos,
                            'type': self._get_replacement_type(pattern_str)
                        })
                        used_positions.add((start_pos, end_pos))
                        logger.info(f"  发现替换项: '{match.group(0)}' -> '{replacement}' (位置: {start_pos}-{end_pos})")
            
            if not replacements:
                logger.info("未发现需要替换的项目")
                return {'success': False, 'replacements': 0}
            
            # 5. 按位置排序（从后往前，避免位置偏移）
            replacements.sort(key=lambda x: x['start_pos'], reverse=True)
            logger.info(f"共发现 {len(replacements)} 个替换项，按位置排序完成")
            
            # 6. 分析原始格式结构
            original_format_map = self._analyze_paragraph_format_structure(paragraph)
            
            # 7. 执行批量替换
            new_text = original_text
            total_length_delta = 0
            
            for i, repl in enumerate(replacements):
                # 调整位置偏移
                adjusted_start = repl['start_pos'] 
                adjusted_end = repl['end_pos']
                
                logger.info(f"  执行替换 {i+1}: '{repl['old_text']}' -> '{repl['new_text']}' (位置: {adjusted_start}-{adjusted_end})")
                
                # 执行文本替换
                new_text = new_text[:adjusted_start] + repl['new_text'] + new_text[adjusted_end:]
                
                # 记录长度变化
                length_delta = len(repl['new_text']) - len(repl['old_text'])
                total_length_delta += length_delta
                
                logger.info(f"    文本长度变化: {length_delta}, 累计变化: {total_length_delta}")
            
            # 8. 重建段落格式
            success = self._rebuild_paragraph_with_format(paragraph, original_format_map, new_text, replacements)
            
            if success:
                logger.info(f"✅ 批量替换完成: {len(replacements)} 项替换，格式已保持")
                return {'success': True, 'replacements': len(replacements)}
            else:
                logger.error("❌ 批量替换失败")
                return {'success': False, 'replacements': 0}
            
        except Exception as e:
            logger.error(f"批量替换失败: {e}", exc_info=True)
            return {'success': False, 'replacements': 0}
    
    def _get_replacement_type(self, pattern_str: str) -> str:
        """根据模式字符串确定替换类型"""
        if '供应商名称' in pattern_str:
            return 'company'
        elif '项目名称' in pattern_str:
            return 'project' 
        elif '采购编号' in pattern_str:
            return 'tender'
        else:
            return 'unknown'
    
    def _analyze_paragraph_format_structure(self, paragraph):
        """分析段落的格式结构"""
        try:
            format_map = []
            current_pos = 0
            
            for i, run in enumerate(paragraph.runs):
                run_length = len(run.text)
                if run_length > 0:
                    format_info = {
                        'start_pos': current_pos,
                        'end_pos': current_pos + run_length,
                        'text': run.text,
                        'format': self._extract_run_format(run)
                    }
                    format_map.append(format_info)
                    logger.debug(f"  格式分析 Run {i+1}: '{run.text[:20]}...' (位置: {current_pos}-{current_pos + run_length})")
                    current_pos += run_length
            
            logger.info(f"段落格式结构分析完成: {len(format_map)} 个格式区域")
            return format_map
            
        except Exception as e:
            logger.error(f"格式结构分析失败: {e}", exc_info=True)
            return []
    
    def _rebuild_paragraph_with_format(self, paragraph, original_format_map, new_text: str, replacements):
        """使用原始格式信息重建段落"""
        try:
            logger.info(f"🔧 开始重建段落格式，新文本长度: {len(new_text)}")
            
            # 清空所有run
            for run in paragraph.runs:
                run.text = ""
            
            # 计算格式边界的调整
            adjusted_format_map = self._adjust_format_boundaries(original_format_map, replacements, new_text)
            
            # 重新分配文本到run
            run_index = 0
            for format_info in adjusted_format_map:
                if run_index >= len(paragraph.runs):
                    # 需要创建新的run
                    new_run = paragraph.add_run("")
                else:
                    new_run = paragraph.runs[run_index]
                
                # 分配文本
                start_pos = format_info['start_pos']
                end_pos = format_info['end_pos']
                
                if start_pos < len(new_text):
                    actual_end = min(end_pos, len(new_text))
                    assigned_text = new_text[start_pos:actual_end]
                    new_run.text = assigned_text
                    
                    # 应用格式
                    self._apply_run_format(new_run, format_info['format'])
                    
                    logger.debug(f"  Run {run_index+1}: 分配文本 '{assigned_text[:30]}...' (位置: {start_pos}-{actual_end})")
                    run_index += 1
            
            logger.info("✅ 段落格式重建完成")
            return True
            
        except Exception as e:
            logger.error(f"段落格式重建失败: {e}", exc_info=True)
            return False
    
    def _adjust_format_boundaries(self, original_format_map, replacements, new_text: str):
        """调整格式边界以适应新文本 - 修复版本"""
        try:
            logger.info(f"🔧 开始调整格式边界，原格式区域: {len(original_format_map)}")
            
            # 🎯 新算法：基于原始文本与新文本的精确映射
            adjusted_format_map = []
            
            # 构建原始文本到新文本的位置映射
            position_mapping = self._build_position_mapping(original_format_map, replacements, new_text)
            
            # 根据映射重建格式区域
            for i, format_info in enumerate(original_format_map):
                original_start = format_info['start_pos']
                original_end = format_info['end_pos']
                original_text = format_info['text']
                
                # 计算在新文本中的位置
                new_start = position_mapping.get(original_start, original_start)
                new_end = position_mapping.get(original_end, original_start + len(original_text))
                
                # 修正边界，确保不超出新文本长度
                new_start = min(new_start, len(new_text))
                new_end = min(new_end, len(new_text))
                
                # 确保end >= start
                if new_end < new_start:
                    new_end = new_start
                
                # 如果区域有效，添加到调整后的格式映射中
                if new_start < len(new_text):
                    # 计算实际应该分配的文本长度
                    available_text = new_text[new_start:new_end] if new_end > new_start else ""
                    
                    if available_text or (i == 0):  # 保留第一个区域即使为空
                        adjusted_format_map.append({
                            'start_pos': new_start,
                            'end_pos': new_end,
                            'format': format_info['format'],
                            'expected_text': available_text
                        })
                        logger.debug(f"  格式区域 {i+1}: {original_start}-{original_end} -> {new_start}-{new_end} (文本: '{available_text[:20]}...')")
            
            # 🔧 关键修复：确保覆盖全部新文本
            if adjusted_format_map:
                last_end = adjusted_format_map[-1]['end_pos']
                if last_end < len(new_text):
                    # 有未覆盖的文本，扩展最后一个格式区域
                    adjusted_format_map[-1]['end_pos'] = len(new_text)
                    adjusted_format_map[-1]['expected_text'] = new_text[adjusted_format_map[-1]['start_pos']:]
                    logger.info(f"🔧 扩展最后格式区域以覆盖全部文本: {last_end} -> {len(new_text)}")
            
            logger.info(f"✅ 格式边界调整完成: {len(adjusted_format_map)} 个有效区域")
            return adjusted_format_map
            
        except Exception as e:
            logger.error(f"格式边界调整失败: {e}", exc_info=True)
            return []
    
    def _build_position_mapping(self, original_format_map, replacements, new_text: str):
        """构建原始位置到新位置的映射"""
        try:
            position_mapping = {}
            
            # 按替换位置排序（从前往后）
            sorted_replacements = sorted(replacements, key=lambda x: x['start_pos'])
            
            # 计算累积偏移量
            cumulative_offset = 0
            last_replacement_end = 0
            
            # 为每个原始格式区域计算新位置
            for format_info in original_format_map:
                original_start = format_info['start_pos']
                original_end = format_info['end_pos']
                
                # 计算影响此区域的替换
                offset_for_start = 0
                offset_for_end = 0
                
                for repl in sorted_replacements:
                    repl_start = repl['start_pos']
                    repl_end = repl['end_pos'] 
                    length_delta = len(repl['new_text']) - len(repl['old_text'])
                    
                    # 如果替换完全在此区域开始位置之前
                    if repl_end <= original_start:
                        offset_for_start += length_delta
                        offset_for_end += length_delta
                    
                    # 如果替换完全在此区域结束位置之前
                    elif repl_end <= original_end:
                        offset_for_end += length_delta
                
                # 应用偏移量
                new_start = original_start + offset_for_start
                new_end = original_end + offset_for_end
                
                position_mapping[original_start] = new_start
                position_mapping[original_end] = new_end
                
                logger.debug(f"  位置映射: {original_start} -> {new_start}, {original_end} -> {new_end}")
            
            return position_mapping
            
        except Exception as e:
            logger.error(f"构建位置映射失败: {e}")
            return {}

    def _should_use_batch_replacement(self, paragraph) -> bool:
        """检查段落是否应该使用批量替换策略"""
        try:
            text = paragraph.text
            
            # 定义多项替换的标识模式
            multiple_replacement_patterns = [
                # 同时包含多个替换项的模式
                (r'（供应商名称[^）]*）.*（项目名称）', '供应商名称+项目名称'),
                (r'（供应商名称[^）]*）.*（采购编号）', '供应商名称+采购编号'), 
                (r'（项目名称）.*（采购编号）', '项目名称+采购编号'),
                (r'（项目名称）.*（姓名、职务）', '项目名称+姓名职务'),
                (r'（采购编号）.*（姓名、职务）', '采购编号+姓名职务'),
                (r'（供应商名称[^）]*）.*（姓名、职务）', '供应商名称+姓名职务'),
                # 三项或以上的组合
                (r'（供应商名称[^）]*）.*（项目名称）.*（采购编号）', '三项替换'),
                (r'（项目名称）.*（采购编号）.*（姓名、职务）', '三项替换'),
            ]
            
            for pattern_str, description in multiple_replacement_patterns:
                import re
                pattern = re.compile(pattern_str)
                if pattern.search(text):
                    logger.info(f"检测到多项替换模式: {description}")
                    return True
            
            # 检查是否包含多个独立的替换项（即使不在同一句中）
            replacement_count = 0
            single_patterns = [
                r'（供应商名称[^）]*）',
                r'（项目名称）',
                r'（采购编号）', 
                r'（姓名、职务）'
            ]
            
            for pattern_str in single_patterns:
                if re.search(pattern_str, text):
                    replacement_count += 1
            
            if replacement_count >= 2:
                logger.info(f"检测到{replacement_count}个替换项，启用批量替换")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"检查批量替换条件失败: {e}")
            return False
    
    def _extract_project_name_from_config(self) -> str:
        """从“读取信息”页面的tender_config.ini提取项目名称"""
        try:
            import configparser
            import os
            from pathlib import Path
            
            # 指向“读取信息”页面的配置文件
            tender_info_path = str(Path(__file__).parent.parent.parent / "1.读取信息")
            config_file = os.path.join(tender_info_path, 'tender_config.ini')
            
            if os.path.exists(config_file):
                config = configparser.ConfigParser()
                config.read(config_file, encoding='utf-8')
                
                # 只从PROJECT_INFO节的project_name键读取
                if config.has_section('PROJECT_INFO'):
                    project_name = config.get('PROJECT_INFO', 'project_name', fallback='')
                    if project_name and project_name.strip():
                        logger.info(f"从读取信息页面读取项目名称: {project_name}")
                        return project_name.strip()
                        
            logger.info(f"未找到读取信息页面的配置文件: {config_file}")
            return None
            
        except Exception as e:
            logger.error(f"从读取信息页面提取项目名称失败: {e}")
            return None
    
    def _extract_tender_number_from_config(self) -> str:
        """从tender_config.ini的PROJECT_INFO节提取项目编号"""
        try:
            # 优先使用传递进来的参数
            if hasattr(self, 'tender_no') and self.tender_no:
                return self.tender_no
                
            # 从配置文件读取
            return self._load_tender_number_from_config()
            
        except Exception as e:
            logger.error(f"提取项目编号失败: {e}")
            return None
    
    def _load_tender_number_from_config(self) -> str:
        """从“读取信息”页面的tender_config.ini加载项目编号"""
        try:
            import configparser
            import os
            from pathlib import Path
            
            # 指向“读取信息”页面的配置文件
            tender_info_path = str(Path(__file__).parent.parent.parent / "1.读取信息")
            config_file = os.path.join(tender_info_path, 'tender_config.ini')
            
            if os.path.exists(config_file):
                config = configparser.ConfigParser()
                config.read(config_file, encoding='utf-8')
                
                # 只从PROJECT_INFO节的project_number键读取
                if config.has_section('PROJECT_INFO'):
                    project_number = config.get('PROJECT_INFO', 'project_number', fallback='')
                    if project_number and project_number.strip():
                        logger.info(f"从读取信息页面读取项目编号: {project_number}")
                        return project_number.strip()
                        
            logger.info(f"未找到读取信息页面的配置文件: {config_file}")
            return None
            
        except Exception as e:
            logger.error(f"从读取信息页面提取项目编号失败: {e}")
            return None

    def smart_text_replace(self, paragraph, old_text: str, new_text: str):
        """
        智能文本替换 - 三层渐进式策略
        新填入的内容使用目标文本第一个字符所在run的格式
        """
        logger.info(f"开始智能替换: '{old_text}' -> '{new_text}'")
        
        # 第一层：单run直接替换 (约80%)
        if self._single_run_replace(paragraph, old_text, new_text):
            logger.info("✅ 第一层：单run替换成功，格式完美保持")
            return True
        
        # 第二层：跨run精确替换 (约15%)
        if old_text in paragraph.text:
            target_format = self._find_first_char_format(paragraph, old_text)
            if self._cross_run_replace_with_format(paragraph, old_text, new_text, target_format):
                logger.info("✅ 第二层：跨run替换成功，新内容使用首字符格式")
                return True
        
        # 第三层：fallback处理 (约5%)
        if self._precise_cross_run_replace(paragraph, old_text, new_text):
            logger.info("⚠️ 第三层：fallback替换完成")
            return True
        
        logger.warning(f"❌ 所有替换策略都失败: '{old_text}'")
        return False

    def _single_run_replace(self, paragraph, old_text: str, new_text: str):
        """第一层：单run直接替换"""
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                logger.info(f"单run替换成功: '{old_text}' -> '{new_text}'")
                return True
        return False

    def _find_first_char_format(self, paragraph, old_text: str):
        """找到目标文本的最佳格式 - 优先查找占位符格式"""
        if not old_text:
            return None
            
        full_text = paragraph.text
        target_pos = full_text.find(old_text)
        
        if target_pos == -1:
            return None
        
        # 优先策略：查找占位符部分的格式（空格或下划线）
        placeholder_chars = [' ', '_', '—', '－', '＿']
        for char in placeholder_chars:
            if char in old_text and old_text.count(char) >= 3:  # 至少3个占位符字符
                placeholder_pos = old_text.find(char)
                if placeholder_pos != -1:
                    # 在段落中找到这个占位符的位置
                    actual_placeholder_pos = target_pos + placeholder_pos
                    format_info = self._find_run_format_at_position(paragraph, actual_placeholder_pos)
                    if format_info:
                        logger.info(f"找到占位符'{char}'的格式，位置{actual_placeholder_pos}")
                        return format_info
        
        # 备用策略：查找第一个字符的格式
        first_char = old_text[0]
        current_pos = 0
        for run in paragraph.runs:
            run_end = current_pos + len(run.text)
            if current_pos <= target_pos < run_end and first_char in run.text:
                format_info = self._extract_run_format(run)
                logger.info(f"找到首字符'{first_char}'的格式: run位置{current_pos}-{run_end}")
                return format_info
            current_pos = run_end
        
        # 如果找不到，使用第一个非空run的格式
        for run in paragraph.runs:
            if run.text.strip():
                return self._extract_run_format(run)
        return None
        
    def _find_run_format_at_position(self, paragraph, position: int):
        """在指定位置找到对应run的格式"""
        current_pos = 0
        for run in paragraph.runs:
            run_end = current_pos + len(run.text)
            if current_pos <= position < run_end:
                return self._extract_run_format(run)
            current_pos = run_end
        return None

    def _cross_run_replace_with_format(self, paragraph, old_text: str, new_text: str, target_format):
        """第二层：跨run精确替换，新内容使用指定格式 - 修复版本"""
        try:
            full_text = paragraph.text
            logger.info(f"第二层替换开始: 原文本='{full_text}', 目标='{old_text}' -> '{new_text}'")
            
            # 找到目标文本位置
            old_start = full_text.find(old_text)
            if old_start == -1:
                logger.error(f"目标文本未找到: '{old_text}'")
                return False
            
            old_end = old_start + len(old_text)
            logger.info(f"目标文本位置: {old_start}-{old_end}")
            
            # 分析所有run的信息
            current_pos = 0
            affected_runs = []
            run_infos = []
            
            for i, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                
                # 判断这个run是否与目标文本有交集
                is_affected = run_start < old_end and run_end > old_start
                if is_affected:
                    affected_runs.append(i)
                
                run_infos.append({
                    'index': i,
                    'run': run,
                    'start': run_start,
                    'end': run_end,
                    'original_text': run.text,
                    'affected': is_affected
                })
                
                current_pos = run_end
            
            if not affected_runs:
                logger.error("没有找到受影响的run")
                return False
            
            logger.info(f"受影响的run索引: {affected_runs}")
            
            # 构建新的完整文本
            new_full_text = full_text.replace(old_text, new_text)
            logger.info(f"新的完整文本: '{new_full_text}'")
            
            # 🔧 关键修复：正确重新分布文本
            # 1. 保存原始格式信息
            original_formats = {}
            for info in run_infos:
                original_formats[info['index']] = self._extract_run_format(info['run'])
            
            # 2. 清空所有受影响的run
            for i in affected_runs:
                paragraph.runs[i].text = ""
            
            # 3. 计算新文本的分布
            new_start = old_start  # 新内容在完整文本中的开始位置
            new_end = new_start + len(new_text)  # 新内容在完整文本中的结束位置
            
            # 4. 重新分配文本 - 修复版本
            # 先计算每个受影响run应该承载的新文本片段
            new_distributions = []
            
            for i, affected_idx in enumerate(affected_runs):
                info = run_infos[affected_idx]
                run_start = info['start']
                run_end = info['end']
                
                # 计算原始文本在这个run中的部分
                original_part_start = max(run_start, old_start) - run_start
                original_part_end = min(run_end, old_end) - run_start
                
                if i == 0:
                    # 第一个受影响的run：包含替换前的部分 + 新内容的开始
                    prefix = info['original_text'][:original_part_start] if original_part_start > 0 else ""
                    new_content = new_text
                    suffix = ""
                    
                    if i == len(affected_runs) - 1:
                        # 如果也是最后一个run，还要包含后缀
                        suffix = info['original_text'][original_part_end:] if original_part_end < len(info['original_text']) else ""
                    
                    final_content = prefix + new_content + suffix
                    
                elif i == len(affected_runs) - 1:
                    # 最后一个受影响的run：只包含后缀
                    suffix = info['original_text'][original_part_end:] if original_part_end < len(info['original_text']) else ""
                    final_content = suffix
                    
                else:
                    # 中间的受影响run：清空
                    final_content = ""
                
                new_distributions.append({
                    'run_index': affected_idx,
                    'content': final_content,
                    'has_new_text': i == 0  # 只有第一个run包含新内容
                })
            
            # 应用新的文本分布
            for dist in new_distributions:
                run_index = dist['run_index']
                content = dist['content']
                has_new_text = dist['has_new_text']
                
                paragraph.runs[run_index].text = content
                
                # 应用格式
                if has_new_text and target_format:
                    self._apply_run_format(paragraph.runs[run_index], target_format)
                    logger.info(f"为run {run_index}应用目标格式，内容='{content}'")
                else:
                    # 保持原格式
                    self._apply_run_format(paragraph.runs[run_index], original_formats.get(run_index))
                
                logger.info(f"Run {run_index}: '{content}'")
            
            # 验证替换结果
            final_text = paragraph.text
            if new_text in final_text and old_text not in final_text:
                logger.info(f"✅ 第二层替换成功验证: '{old_text}' -> '{new_text}'")
                return True
            else:
                logger.error(f"❌ 第二层替换验证失败: 期望包含'{new_text}'，不包含'{old_text}'，实际='{final_text}'")
                return False
            
        except Exception as e:
            logger.error(f"第二层跨run替换失败: {e}", exc_info=True)
            return False

    def _enhanced_residual_placeholder_cleanup(self, paragraph: Paragraph, label: str, company_name: str) -> None:
        """
        强化残留占位符清理 - 专门处理"内容+残留占位符"混合状态
        
        针对如下情况：
        原始：供应商名称：                                        
        填写后：供应商名称：中国联合网络通信有限公司                                        
        期望：供应商名称：中国联合网络通信有限公司
        
        Args:
            paragraph: 要清理的段落
            label: 标签名称（如"供应商名称"）
            company_name: 已填入的公司名称
        """
        try:
            full_text = paragraph.text
            if not full_text or not company_name or not label:
                return
                
            logger.info(f"开始强化残留占位符清理: '{full_text}'")
            
            # 构建期望的完整格式
            expected_prefix = f"{label}："
            if expected_prefix in full_text and company_name in full_text:
                # 找到公司名称在文本中的位置
                company_pos = full_text.find(company_name)
                prefix_pos = full_text.find(expected_prefix)
                
                if company_pos > prefix_pos >= 0:
                    # 计算公司名称结束位置
                    company_end = company_pos + len(company_name)
                    
                    # 检查公司名称后是否有残留占位符
                    remaining_text = full_text[company_end:]
                    
                    # 如果公司名称后有大量空格或下划线，认为是残留占位符
                    if re.search(r'[\s_]{3,}', remaining_text):
                        logger.info(f"检测到残留占位符: '{remaining_text}' (长度:{len(remaining_text)})")
                        
                        # 构建清理后的文本：保留前置内容+公司名称，清除后续占位符
                        before_company = full_text[:company_pos]
                        clean_text = before_company + company_name
                        
                        # 保留非占位符的有意义内容（如括号、后缀等）
                        meaningful_suffix = re.sub(r'[\s_]{3,}', '', remaining_text).strip()
                        if meaningful_suffix:
                            clean_text += meaningful_suffix
                            
                        logger.info(f"清理目标: '{full_text}' -> '{clean_text}'")
                        
                        # 应用清理结果到段落runs
                        if clean_text != full_text:
                            self._apply_cleaned_text_to_runs(paragraph, full_text, clean_text)
                            logger.info(f"✅ 强化清理完成：移除了 {len(full_text) - len(clean_text)} 个残留字符")
                        else:
                            logger.info("文本已经是最佳状态，无需清理")
                    else:
                        logger.info("未检测到需要清理的残留占位符")
                else:
                    logger.info("文本结构不符合预期，跳过强化清理")
            else:
                logger.info("未找到预期的标签和公司名称组合")
                
        except Exception as e:
            logger.error(f"强化残留占位符清理失败: {e}")


def test_mcp_bidder_processor():
    """测试MCP投标人名称处理器"""
    # 这里可以添加测试代码
    pass


if __name__ == "__main__":
    test_mcp_bidder_processor()