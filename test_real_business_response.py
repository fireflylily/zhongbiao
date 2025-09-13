#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试修复后的项目信息处理 - 使用BusinessResponseProcessor实际数据结构
"""

import sys
import logging
from pathlib import Path
from docx import Document

# 添加项目根目录到系统路径
sys.path.append(str(Path(__file__).parent))

from ai_tender_system.modules.business_response.info_filler import InfoFiller

def setup_test_logging():
    """设置测试日志"""
    # 设置根日志器
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    
    # 创建控制台处理器
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.DEBUG)
    
    # 创建格式器
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    console_handler.setFormatter(formatter)
    
    # 添加处理器到根日志器
    root_logger.addHandler(console_handler)
    
    return root_logger

def create_test_document_with_long_paragraph():
    """创建包含长段落的测试文档（模拟实际文档）"""
    doc = Document()
    
    # 添加包含括号格式的测试段落，特别是长段落
    test_paragraphs = [
        "根据贵方为（项目名称）项目采购采购货物及服务的竞争性磋商公告（采购编号），签字代表（姓名、职务）经正式授权并代表供应商（供应商名称、地址）提交下述文件正本一份及副本份：",
        "（项目名称）",  
        "（采购编号）",
        "供应商名称 ",
    ]
    
    for text in test_paragraphs:
        doc.add_paragraph(text)
    
    return doc

def test_business_response_data_structure():
    """测试BusinessResponseProcessor实际数据结构"""
    print("=" * 80)
    print("测试修复后的项目信息处理 - BusinessResponseProcessor数据结构")
    print("=" * 80)
    
    # 设置日志
    logger = setup_test_logging()
    
    # 创建InfoFiller实例
    filler = InfoFiller()
    
    # 创建测试文档
    doc = create_test_document_with_long_paragraph()
    
    # 使用BusinessResponseProcessor实际的数据结构（英文键名）
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'address': '北京市东城区王府井大街200号七层711室',
        'phone': '010-63271000',
        'email': 'lvhe@smartsteps.com',
        'fax': '010-63271000',
    }
    
    # 模拟BusinessResponseProcessor构造的project_info（参考processor.py第213-219行）
    project_info = {
        'projectName': '哈银消金2025年-2027年运营商数据采购项目',  # 英文键名!
        'projectNumber': 'GXTC-C-251590031',  # 英文键名!
        'date': '2025年9月13日',
        'purchaserName': '哈尔滨哈银消费金融有限责任公司',
        'projectOwner': '哈尔滨哈银消费金融有限责任公司'  # fallback
    }
    
    print(f"\n📋 测试数据（BusinessResponseProcessor格式）:")
    print(f"公司信息: {company_info}")
    print(f"项目信息: {project_info}")
    
    print(f"\n📄 测试文档段落:")
    for i, paragraph in enumerate(doc.paragraphs, 1):
        print(f"  {i}. '{paragraph.text}'")
    
    print(f"\n" + "=" * 80)
    print("开始执行信息填写处理")
    print("=" * 80)
    
    # 执行信息填写
    try:
        result = filler.fill_info(doc, company_info, project_info)
        
        print(f"\n" + "=" * 80)
        print("处理完成后的文档内容")
        print("=" * 80)
        
        for i, paragraph in enumerate(doc.paragraphs, 1):
            print(f"  {i}. '{paragraph.text}'")
        
        print(f"\n📊 最终统计结果: {result}")
        
        # 验证关键结果
        print(f"\n🔍 关键验证:")
        long_paragraph = doc.paragraphs[0].text
        if "哈银消金" in long_paragraph and "GXTC-C-251590031" in long_paragraph:
            print("✅ 长段落中的项目信息处理成功")
        else:
            print("❌ 长段落中的项目信息处理失败")
            
        if result.get('replacement_rules', 0) >= 2:
            print("✅ 替换规则正常工作")
        else:
            print("❌ 替换规则未正常工作")
        
    except Exception as e:
        print(f"\n❌ 测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_business_response_data_structure()