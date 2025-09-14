#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
格式隔离测试脚本 - 验证天然Run替换引擎的格式继承修复

测试目标:
1. 验证业务内容(公司名称等)不会继承装饰格式(下划线等)
2. 验证基本格式(字体等)正常保持
3. 验证非业务内容的格式继承正常工作
"""

import os
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent))

from docx import Document
from docx.shared import Inches
from ai_tender_system.modules.business_response.info_filler import InfoFiller


def create_test_document():
    """创建测试文档，包含各种格式的字段"""
    doc = Document()

    # 测试用例1: 带下划线格式的供应商名称
    p1 = doc.add_paragraph()
    run1 = p1.add_run("   ")
    run2 = p1.add_run("(供应商全称)")
    run2.font.underline = True  # 添加下划线
    run2.font.name = "宋体"
    run3 = p1.add_run("       授权 ")
    run4 = p1.add_run("(供应商代表姓名)")
    run4.font.name = "宋体"

    # 测试用例2: 带下划线的电话字段
    p2 = doc.add_paragraph()
    run1 = p2.add_run("联系电话：")
    run1.font.underline = True
    run1.font.name = "宋体"
    run2 = p2.add_run("___________")
    run2.font.underline = True

    # 测试用例3: 带删除线的邮箱字段
    p3 = doc.add_paragraph()
    run1 = p3.add_run("电子邮箱：")
    run1.font.strike = True
    run1.font.name = "宋体"
    run2 = p3.add_run("___________")
    run2.font.strike = True

    # 测试用例4: 普通格式的项目名称
    p4 = doc.add_paragraph()
    run1 = p4.add_run("项目名称：")
    run1.font.name = "宋体"
    run2 = p4.add_run("(项目名称)")
    run2.font.name = "宋体"

    return doc


def test_format_isolation():
    """测试格式隔离功能"""
    print("🧪 开始格式隔离测试")
    print("=" * 60)

    # 创建测试文档
    doc = create_test_document()
    test_file = "test_format_isolation_input.docx"
    doc.save(test_file)

    # 准备测试数据
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'authorizedPersonName': '吕贺',
        'fixedPhone': '010-63271000',
        'email': 'lvhe@smartsteps.com'
    }

    project_info = {
        'projectName': '哈银消金2025年-2027年运营商数据采购项目',
        'projectNumber': 'GXTC-C-251590031',
        'date': '2025年9月12日',
        'purchaserName': '哈银消费金融有限责任公司'
    }

    # 执行信息填写
    info_filler = InfoFiller()

    print("🔧 执行信息填写处理...")
    doc = Document(test_file)
    stats = info_filler.fill_info(doc, company_info, project_info)

    # 保存处理后的文档
    output_file = "test_format_isolation_output.docx"
    doc.save(output_file)

    print(f"📊 处理统计: {stats}")
    print()

    # 验证格式隔离效果
    print("🔍 验证格式隔离效果:")
    print("-" * 40)

    # 重新加载文档进行验证
    doc_verify = Document(output_file)

    test_results = []

    for i, paragraph in enumerate(doc_verify.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        print(f"段落 #{i+1}: '{text}'")

        # 检查每个run的格式
        for j, run in enumerate(paragraph.runs):
            if not run.text.strip():
                continue

            run_text = run.text
            has_underline = getattr(run.font.underline, 'value', False) if run.font.underline else False
            has_strike = getattr(run.font.strike, 'value', False) if run.font.strike else False
            font_name = run.font.name or "默认"

            print(f"  Run {j}: '{run_text}' | 字体: {font_name} | 下划线: {has_underline} | 删除线: {has_strike}")

            # 测试案例1: 检查公司名称是否清除了下划线
            if '智慧足迹数据科技有限公司' in run_text:
                test_case = {
                    'name': '公司名称格式隔离',
                    'content': run_text,
                    'expected_underline': False,
                    'actual_underline': has_underline,
                    'passed': not has_underline
                }
                test_results.append(test_case)

            # 测试案例2: 检查被授权人姓名是否清除了格式
            if '吕贺' in run_text:
                test_case = {
                    'name': '被授权人姓名格式隔离',
                    'content': run_text,
                    'expected_underline': False,
                    'actual_underline': has_underline,
                    'passed': not has_underline
                }
                test_results.append(test_case)

            # 测试案例3: 检查电话号码是否清除了下划线
            if '010-63271000' in run_text:
                test_case = {
                    'name': '电话号码格式隔离',
                    'content': run_text,
                    'expected_underline': False,
                    'actual_underline': has_underline,
                    'passed': not has_underline
                }
                test_results.append(test_case)

            # 测试案例4: 检查邮箱是否清除了删除线
            if 'lvhe@smartsteps.com' in run_text:
                test_case = {
                    'name': '邮箱格式隔离',
                    'content': run_text,
                    'expected_strike': False,
                    'actual_strike': has_strike,
                    'passed': not has_strike
                }
                test_results.append(test_case)

        print()

    # 输出测试结果
    print("📋 格式隔离测试结果:")
    print("=" * 60)

    passed_tests = 0
    total_tests = len(test_results)

    for i, test in enumerate(test_results, 1):
        status = "✅ PASS" if test['passed'] else "❌ FAIL"
        print(f"{i}. {test['name']}: {status}")
        print(f"   内容: '{test['content']}'")

        if 'expected_underline' in test:
            print(f"   下划线 - 期望: {test['expected_underline']}, 实际: {test['actual_underline']}")
        if 'expected_strike' in test:
            print(f"   删除线 - 期望: {test['expected_strike']}, 实际: {test['actual_strike']}")

        if test['passed']:
            passed_tests += 1
        print()

    # 总结
    print(f"📊 测试总结: {passed_tests}/{total_tests} 通过")
    if passed_tests == total_tests:
        print("🎉 所有格式隔离测试通过！")
        return True
    else:
        print(f"⚠️  {total_tests - passed_tests} 个测试失败，需要进一步调试")
        return False


def test_business_content_detection():
    """测试业务内容检测逻辑"""
    print("🔍 测试业务内容检测逻辑")
    print("=" * 40)

    info_filler = InfoFiller()

    # 测试用例
    test_cases = [
        # 应该识别为业务内容的
        ('智慧足迹数据科技有限公司', True, '公司名称'),
        ('lvhe@smartsteps.com', True, '邮箱地址'),
        ('www.smartsteps.com', True, '网站地址'),
        ('010-63271000', True, '电话号码'),
        ('北京市海淀区', True, '地址信息'),

        # 不应该识别为业务内容的
        ('项目名称', False, '字段标签'),
        ('联系电话：', False, '字段标签'),
        ('___________', False, '占位符'),
        ('（项目编号）', False, '字段标识'),
        ('授权代表', False, '普通文本'),
    ]

    passed = 0
    total = len(test_cases)

    for i, (text, expected, description) in enumerate(test_cases, 1):
        result = info_filler._is_business_content(text)
        status = "✅" if result == expected else "❌"

        print(f"{i:2d}. {status} '{text}' ({description})")
        print(f"     期望: {expected}, 实际: {result}")

        if result == expected:
            passed += 1
        print()

    print(f"📊 业务内容检测测试: {passed}/{total} 通过")
    return passed == total


if __name__ == '__main__':
    print("🚀 启动格式隔离修复验证测试")
    print("=" * 60)

    # 测试1: 业务内容检测
    detection_passed = test_business_content_detection()
    print()

    # 测试2: 格式隔离效果
    isolation_passed = test_format_isolation()
    print()

    # 总结
    if detection_passed and isolation_passed:
        print("🎉 格式隔离修复验证完全成功！")
        print("✅ 业务内容检测正确")
        print("✅ 格式隔离效果良好")
        print("✅ 下划线等装饰格式不会传播到业务内容")
        sys.exit(0)
    else:
        print("⚠️  格式隔离修复验证存在问题")
        if not detection_passed:
            print("❌ 业务内容检测逻辑需要调整")
        if not isolation_passed:
            print("❌ 格式隔离效果不理想")
        sys.exit(1)