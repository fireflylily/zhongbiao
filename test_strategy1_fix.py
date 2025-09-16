#!/usr/bin/env python3
"""
测试策略1修改效果 - 验证空格智能清理功能
"""

import os
import sys
import tempfile
from pathlib import Path
from docx import Document

# 添加项目路径到系统路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from ai_tender_system.modules.business_response.info_filler import InfoFiller

def create_test_document():
    """创建测试文档"""
    doc = Document()

    # 测试场景1：字段在段落中间
    doc.add_paragraph("电话                                  电子邮件                            ")

    # 测试场景2：字段在段落末尾
    doc.add_paragraph("供应商名称                                    ")

    # 测试场景3：多个字段在同一段落
    doc.add_paragraph("传真                     邮编                     地址                    ")

    # 测试场景4：字段后跟冒号（不应被策略1处理）
    doc.add_paragraph("电话：                                  ")

    return doc

def test_strategy1():
    """测试策略1的空格清理功能"""
    print("=" * 60)
    print("测试策略1插入式填空 - 空格智能清理")
    print("=" * 60)

    # 创建测试文档
    doc = create_test_document()

    # 保存临时文件
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()

    print(f"✅ 创建测试文档: {temp_file.name}")

    # 测试数据 - 匹配InfoFiller期望的字段名
    company_info = {
        'phone': '010-63271000',
        'email': 'lvhe@smartsteps.com',
        'companyName': '智慧足迹数据科技有限公司',
        'fax': '010-63271001',
        'postalCode': '100089',  # 修正：使用postalCode而不是zipCode
        'address': '北京市海淀区中关村'
    }
    project_info = {}

    # 初始化InfoFiller
    filler = InfoFiller()

    # 重新加载文档进行处理
    doc = Document(temp_file.name)

    print("\n📝 原始段落内容：")
    for i, para in enumerate(doc.paragraphs, 1):
        print(f"  段落{i}: '{para.text}'")
        print(f"  长度: {len(para.text)} 字符")

    # 处理文档
    print("\n🔄 开始处理...")
    filler.fill_info(doc, company_info, project_info)

    print("\n✨ 处理后段落内容：")
    for i, para in enumerate(doc.paragraphs, 1):
        print(f"  段落{i}: '{para.text}'")
        print(f"  长度: {len(para.text)} 字符")

        # 检查是否还有多余空格
        if '    ' in para.text:  # 4个或更多连续空格
            print(f"  ⚠️ 警告：仍有多余空格")
        else:
            print(f"  ✅ 空格处理正常")

    # 保存结果
    output_file = temp_file.name.replace('.docx', '_processed.docx')
    doc.save(output_file)
    print(f"\n💾 处理结果已保存: {output_file}")

    # 清理临时文件
    os.unlink(temp_file.name)

    return output_file

if __name__ == "__main__":
    try:
        output_file = test_strategy1()
        print(f"\n✅ 测试完成！请检查输出文件: {output_file}")
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()