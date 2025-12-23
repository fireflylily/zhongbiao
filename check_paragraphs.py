#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查哈银消金文档的段落内容
"""

from docx import Document


def check_paragraphs():
    """检查段落100-130的内容"""

    doc_path = "/Users/lvhe/Downloads/zhongbiao/zhongbiao/ai_tender_system/data/uploads/tender_processing/2025/12/20251221_180155_招标文件-哈银消金_b965f323.docx"

    print("=" * 100)
    print("📄 检查哈银消金文档段落内容")
    print("=" * 100)

    doc = Document(doc_path)

    print(f"\n文档总段落数: {len(doc.paragraphs)}")

    # 检查段落100-130
    print("\n" + "=" * 100)
    print("段落 100-130 内容")
    print("=" * 100)

    for i in range(100, min(131, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        word_count = len(text.replace(' ', '').replace('\n', ''))

        # 高亮关键段落
        marker = ""
        if i == 104:
            marker = "👉 [第二部分起始]"
        elif i == 120:
            marker = "👉 [第二部分结束?]"
        elif i == 121:
            marker = "👉 [第三部分?]"
        elif '第二部分' in text:
            marker = "⚠️  [包含'第二部分']"
        elif '第三部分' in text:
            marker = "⚠️  [包含'第三部分']"
        elif '第四部分' in text:
            marker = "⚠️  [包含'第四部分']"

        # 显示内容（截断到100字符）
        display_text = text[:100] + ('...' if len(text) > 100 else '')

        print(f"段落 {i:3d} ({word_count:4d}字) {marker}")
        print(f"  内容: {display_text if text else '(空段落)'}")
        print()

    # 统计段落104-120的字数
    print("=" * 100)
    print("统计段落104-120的内容")
    print("=" * 100)

    total_chars = 0
    non_empty_count = 0

    for i in range(104, 121):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text
            chars = len(text.replace(' ', '').replace('\n', ''))
            total_chars += chars
            if text.strip():
                non_empty_count += 1

    print(f"\n段落范围: 104-120 (共17个段落)")
    print(f"非空段落数: {non_empty_count}")
    print(f"总字数: {total_chars:,}字")

    # 提取内容（跳过标题段落104）
    content_chars = 0
    for i in range(105, 121):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text
            chars = len(text.replace(' ', '').replace('\n', ''))
            content_chars += chars

    print(f"\n内容段落: 105-120 (共16个段落)")
    print(f"内容字数: {content_chars:,}字")
    print(f"  (与日志中的409字{'一致 ✓' if content_chars == 409 else '不一致 ✗'})")

    # 检查是否有更多内容
    print("\n" + "=" * 100)
    print("检查段落121之后是否还有'第二部分'的内容")
    print("=" * 100)

    found_third = False
    for i in range(121, min(200, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()

        # 检查是否真的是"第三部分"的开始
        if '第三部分' in text and '评标' in text:
            found_third = True
            print(f"\n✓ 找到'第三部分': 段落 {i}")
            print(f"  内容: {text}")

            # 显示前后几个段落
            print(f"\n前后段落内容:")
            for j in range(max(i-2, 121), min(i+3, len(doc.paragraphs))):
                print(f"  段落 {j:3d}: {doc.paragraphs[j].text.strip()[:80]}")
            break

    if not found_third:
        print("\n⚠️  在段落121-200之间未找到'第三部分 评标办法'")


if __name__ == '__main__':
    check_paragraphs()
