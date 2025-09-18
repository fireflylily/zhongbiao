#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试标准表单式布局（电话：_____）处理
"""
import sys
import json
import tempfile
from pathlib import Path
import shutil

# 添加项目路径
script_dir = Path(__file__).parent
project_root = script_dir / 'ai_tender_system'
sys.path.insert(0, str(project_root))

try:
    from docx import Document
    import importlib.util
    
    print("🧪 测试标准表单式布局处理")
    print("=" * 40)
    
    # 创建一个简单的测试文档
    test_doc = Document()
    
    # 添加测试段落
    test_paragraphs = [
        "电话：                  ",
        "电子邮件：              ",
        "传真：                  ",
        "供应商名称：            ",
    ]
    
    for para_text in test_paragraphs:
        test_doc.add_paragraph(para_text)
    
    # 保存测试文档
    test_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
    output_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
    
    test_doc.save(test_file)
    shutil.copy2(test_file, output_file)
    
    print(f"📄 测试文档: {Path(test_file).name}")
    print(f"📤 输出文件: {Path(output_file).name}")
    
    # 读取公司数据
    company_file = project_root / "data/configs/companies/945150ca-68e1-4141-921b-fd4c48e07ebd.json"
    with open(company_file, 'r', encoding='utf-8') as f:
        company_data = json.load(f)
    
    print(f"\n📋 公司信息:")
    print(f"  电话: {company_data.get('fixedPhone', 'N/A')}")
    print(f"  邮件: {company_data.get('email', 'N/A')}")
    print(f"  传真: {company_data.get('fax', 'N/A')}")
    print(f"  公司名称: {company_data.get('companyName', 'N/A')}")
    
    # 显示原始内容
    print(f"\n📋 原始文档内容:")
    original_doc = Document(test_file)
    for i, para in enumerate(original_doc.paragraphs):
        if para.text.strip():
            print(f"  段落 #{i}: '{para.text}'")
    
    # 动态加载MCP处理器
    processor_file = script_dir / "2.填写标书/点对点应答/mcp_bidder_name_processor_enhanced 2.py"
    
    spec = importlib.util.spec_from_file_location("mcp_processor", processor_file)
    mcp_module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mcp_module)
    
    processor_class = getattr(mcp_module, 'MCPBidderNameProcessor')
    processor = processor_class()
    
    print(f"\n🚀 开始处理标准表单式格式...")
    
    # 执行处理
    result = processor.process_business_response(
        input_file=test_file,
        output_file=output_file,
        company_info=company_data,
        project_name='标准格式测试项目',
        tender_no='STD2025001',
        date_text='2025年9月12日'
    )
    
    print(f"✅ 处理完成: {result.get('message', '成功')}")
    
    if result.get('success'):
        # 验证处理结果
        print(f"\n🔍 验证处理结果:")
        
        processed_doc = Document(output_file)
        for i, para in enumerate(processed_doc.paragraphs):
            if para.text.strip():
                text = para.text
                print(f"  段落 #{i}: '{text}'")
                
                # 检查是否正确填充
                if '电话' in text:
                    filled = '010-63271000' in text
                    print(f"    📞 电话填充: {'✅' if filled else '❌'}")
                
                if '电子邮件' in text:
                    filled = 'lvhe@smartsteps.com' in text
                    print(f"    📧 邮件填充: {'✅' if filled else '❌'}")
                
                if '传真' in text:
                    filled = '010-63271000' in text
                    print(f"    📠 传真填充: {'✅' if filled else '❌'}")
                
                if '供应商名称' in text:
                    filled = '智慧足迹数据科技有限公司' in text
                    print(f"    🏢 公司填充: {'✅' if filled else '❌'}")
        
        # 统计
        stats = result.get('stats', {})
        print(f"\n📊 处理统计:")
        print(f"  总替换次数: {stats.get('total_replacements', 0)}")
        print(f"  字段处理数: {stats.get('info_fields_processed', 0)}")
        
    else:
        print(f"❌ 处理失败: {result.get('error', '未知错误')}")
    
    # 清理文件
    import os
    try:
        os.unlink(test_file)
        os.unlink(output_file)
        print(f"\n🗑️  测试文件已清理")
    except:
        pass
    
    print(f"\n✅ 标准表单式格式测试完成")
    
except Exception as e:
    print(f"❌ 测试错误: {e}")
    import traceback
    traceback.print_exc()