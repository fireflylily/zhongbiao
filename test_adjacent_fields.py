#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试紧邻字段格式处理
"""
import os
import sys
from docx import Document

# 添加项目根目录到Python路径
sys.path.append('/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/ai_tender_system')

def test_adjacent_field_processing():
    """测试紧邻字段处理"""
    from modules.point_to_point.processor import PointToPointProcessor
    
    processor = PointToPointProcessor()
    
    # 设置公司信息
    processor.company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'fixedPhone': '010-63271000',
        'email': 'lvhe@smartsteps.com',
        'address': '北京市东城区王府井大街200号七层711室',
        'postalCode': '100006'
    }
    
    print("=== 紧邻字段格式处理测试 ===")
    
    # 测试用例1: 地址+邮编紧邻格式
    test_case_1 = "地址: 北京市东城区王府井大街200号七层711室邮编：____________"
    print(f"\n测试1 - 地址+邮编紧邻格式:")
    print(f"输入: '{test_case_1}'")
    
    try:
        result1 = processor._handle_dual_field_table_layout(test_case_1, '地址', '北京市东城区王府井大街200号七层711室')
        print(f"结果: '{result1}'")
        if result1 != test_case_1:
            print("✅ 地址+邮编处理成功")
        else:
            print("❌ 地址+邮编处理无变化")
    except Exception as e:
        print(f"❌ 地址+邮编处理异常: {e}")
    
    # 测试用例2: 电话+邮箱紧邻格式
    test_case_2 = "电话：010-63271000电子邮箱："
    print(f"\n测试2 - 电话+邮箱紧邻格式:")
    print(f"输入: '{test_case_2}'")
    
    try:
        result2 = processor._handle_dual_field_table_layout(test_case_2, '电话', '010-63271000')
        print(f"结果: '{result2}'")
        if result2 != test_case_2:
            print("✅ 电话+邮箱处理成功")
        else:
            print("❌ 电话+邮箱处理无变化")
    except Exception as e:
        print(f"❌ 电话+邮箱处理异常: {e}")

def test_full_document_with_adjacent_fields():
    """创建并测试包含紧邻字段的完整文档"""
    
    # 创建测试文档，模拟用户的实际格式
    doc = Document()
    doc.add_paragraph("供应商名称：_____________")
    doc.add_paragraph("地址: 北京市东城区王府井大街200号七层711室邮编：____________")  # 紧邻格式
    doc.add_paragraph("电话：010-63271000电子邮箱：")  # 紧邻格式
    doc.add_paragraph("日期：____年____月____日")
    
    # 保存测试文档
    test_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/test_adjacent_fields.docx"
    doc.save(test_file)
    print(f"\n创建测试文档: {test_file}")
    
    # 使用处理器处理文档
    from modules.point_to_point.processor import PointToPointProcessor
    processor = PointToPointProcessor()
    
    output_file = test_file.replace('.docx', '_processed.docx')
    
    # 测试公司信息（包含邮政编码）
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'fixedPhone': '010-63271000', 
        'email': 'lvhe@smartsteps.com',
        'address': '北京市东城区王府井大街200号七层711室',
        'postalCode': '100006'
    }
    
    print(f"\n=== 完整文档处理测试 ===")
    print(f"公司信息: {company_info}")
    
    try:
        result = processor.process_business_response(
            input_file=test_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-2025-001",
            date_text="2025年9月12日"
        )
        
        print(f"\n处理结果: {result}")
        
        if result['success']:
            print(f"✅ 处理成功!")
            print(f"📄 输出文件: {output_file}")
            
            # 检查输出文档内容
            if os.path.exists(output_file):
                doc = Document(output_file)
                print(f"\n📋 处理后的文档内容:")
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        print(f"  {i+1}. {para.text}")
        else:
            print(f"❌ 处理失败: {result.get('error', 'Unknown error')}")
            
    except Exception as e:
        print(f"❌ 测试异常: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("🚀 开始测试紧邻字段格式处理")
    
    # 测试单独的紧邻字段处理
    test_adjacent_field_processing()
    
    # 测试完整文档处理
    test_full_document_with_adjacent_fields()
    
    print("\n🎉 测试完成")