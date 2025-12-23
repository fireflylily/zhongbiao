#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查显示0字的章节在Word文档中的实际内容
直接在文档中搜索章节标题，查看后面的内容
"""

import sys
from pathlib import Path

# 添加项目路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / 'ai_tender_system'))

from docx import Document

def check_zero_word_chapters():
    """检查0字章节的实际内容"""

    doc_path = "ai_tender_system/data/uploads/tender_processing/2025/11/20251117_160736_【招标方案】成都数据集团第一批数据产品供应商库（发售版）(1)_b6e72bac.docx"

    if not Path(doc_path).exists():
        doc_path = "ai_tender_system/data/uploads/tender_processing/2025/11/20251110_141642_【招标方案】成都数据集团第一批数据产品供应商库（发售版）(1)_75647a44.docx"

    print("=" * 100)
    print("检查显示0字的章节在Word文档中的实际内容")
    print("=" * 100)

    doc = Document(doc_path)

    # 需要检查的章节（从截图中看到的0字章节）
    chapters_to_check = [
        "第三章",
        "第四章",
        "第五章",
        "第六章",
        "第七章",
    ]

    print(f"\n文档总段落数: {len(doc.paragraphs)}\n")

    for chapter_title_prefix in chapters_to_check:
        print("=" * 100)
        print(f"检查章节: {chapter_title_prefix}")
        print("=" * 100)

        # 搜索章节标题
        chapter_idx = None
        chapter_full_title = None

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith(chapter_title_prefix):
                chapter_idx = idx
                chapter_full_title = text
                break

        if chapter_idx is None:
            print(f"❌ 未找到章节标题\n")
            continue

        print(f"章节标题: {chapter_full_title}")
        print(f"段落位置: {chapter_idx}")

        # 找下一个章节（以"第"开头且包含"章"的段落）
        next_chapter_idx = None
        next_chapter_title = None

        for idx in range(chapter_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[idx].text.strip()
            # 简单判断：以"第"开头，包含"章"
            if text.startswith("第") and "章" in text and len(text) < 50:
                next_chapter_idx = idx
                next_chapter_title = text
                break

        if next_chapter_idx is None:
            next_chapter_idx = len(doc.paragraphs) - 1
            next_chapter_title = "(文档末尾)"

        print(f"下一章节: {next_chapter_title}")
        print(f"下一章位置: {next_chapter_idx}")

        # 统计中间的内容
        content_paras = []
        for idx in range(chapter_idx + 1, next_chapter_idx):
            text = doc.paragraphs[idx].text.strip()
            if text:
                content_paras.append(text)

        total_text = '\n'.join(content_paras)
        word_count = len(total_text.replace(' ', '').replace('\n', ''))

        print(f"\n中间段落数: {next_chapter_idx - chapter_idx - 1}")
        print(f"非空段落数: {len(content_paras)}")
        print(f"总字数: {word_count:,}")

        if len(content_paras) > 0:
            print(f"\n前3个非空段落内容:")
            for i, para_text in enumerate(content_paras[:3]):
                preview = para_text[:80] + ('...' if len(para_text) > 80 else '')
                print(f"  {i+1}. {preview}")
        else:
            print("\n⚠️  章节标题后面没有任何内容，下一段就是下一章节")

        print()

    print("=" * 100)
    print("总结")
    print("=" * 100)
    print("\n如果章节后面确实没有内容（下一段就是下一章），那么显示0字是正确的。")
    print("如果章节后面有内容但显示0字，那就是解析BUG。\n")

if __name__ == '__main__':
    check_zero_word_chapters()
