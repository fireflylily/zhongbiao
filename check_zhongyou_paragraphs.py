#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查中邮保险文档段落26-100的内容
"""

import sys
from pathlib import Path

# 添加项目路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / 'ai_tender_system'))

from docx import Document

def main():
    doc_path = "ai_tender_system/data/uploads/tender_processing/2025/12/20251221_180214_中邮保险手机号实名认证服务采购项目竞争性磋商采购文件_fcfdcec9.docx"
    doc = Document(doc_path)

    print("=" * 100)
    print("中邮保险文档：段落26-100内容")
    print("=" * 100)
    print(f"文档: {Path(doc_path).name}\n")

    print("段落内容:")
    print("-" * 100)

    for i in range(26, min(100, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # 检查样式
        style_name = para.style.name if para.style else "None"

        # 检查大纲级别
        outline_level = None
        if hasattr(para, '_element') and hasattr(para._element, 'pPr') and para._element.pPr is not None:
            outline_lvl = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
            if outline_lvl is not None:
                outline_level = int(outline_lvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))

        # 只显示有文本的段落
        if text:
            outline_str = f"Outline={outline_level}" if outline_level is not None else ""
            print(f"{i:4d}. [{style_name:20s}] {outline_str:12s} {text[:80]}")

    print()

if __name__ == '__main__':
    main()
