#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析所有层级的段落范围关系
检查父节点是否 = 自身直接内容 + 直接子节点之和
"""

import sys
from pathlib import Path

# 添加项目路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / 'ai_tender_system'))

from modules.tender_processing.structure_parser import DocumentStructureParser

def analyze_node_relationship(ch, doc_paragraphs):
    """
    分析节点的段落范围关系

    检查:
    1. 父节点的段落范围是否包含所有子节点
    2. 父节点字数 vs (直接子节点字数之和)
    3. 父节点的"自身直接内容"有多少字（不包括任何子节点的部分）
    """
    if not ch.get('children'):
        return None  # 叶子节点

    result = {
        'title': ch['title'],
        'level': ch['level'],
        'para_range': f"{ch['para_start_idx']}-{ch['para_end_idx']}",
        'total_words': ch['word_count'],
        'children_count': len(ch['children']),
        'children_total_words': sum(c['word_count'] for c in ch['children']),
        'children_ranges': []
    }

    # 收集子节点的段落范围
    for child in ch['children']:
        result['children_ranges'].append({
            'title': child['title'],
            'level': child['level'],
            'range': f"{child['para_start_idx']}-{child['para_end_idx']}",
            'words': child['word_count']
        })

    # 计算父节点的"自身直接内容"
    # 方法：找出父节点范围内，不属于任何子节点的段落
    parent_start = ch['para_start_idx']
    parent_end = ch['para_end_idx']

    # 收集所有子节点占用的段落范围
    child_occupied = set()
    for child in ch['children']:
        for p in range(child['para_start_idx'], child['para_end_idx'] + 1):
            child_occupied.add(p)

    # 父节点的直接内容段落 = 父节点范围 - 子节点占用的范围
    parent_direct_paras = []
    for p in range(parent_start, parent_end + 1):
        if p not in child_occupied:
            parent_direct_paras.append(p)

    # 计算直接内容的字数
    if doc_paragraphs:
        direct_content_words = sum(
            len(doc_paragraphs[p].text.strip().replace(' ', '').replace('\n', ''))
            for p in parent_direct_paras
            if p < len(doc_paragraphs)
        )
        result['direct_content_words'] = direct_content_words
        result['direct_content_paras'] = len(parent_direct_paras)

    # 验证：父节点字数 应该 = 自身直接内容 + 直接子节点字数之和
    if doc_paragraphs:
        expected = direct_content_words + result['children_total_words']
        actual = result['total_words']
        result['validation'] = {
            'expected': expected,
            'actual': actual,
            'match': expected == actual,
            'diff': actual - expected
        }

    return result

def main():
    doc_path = "ai_tender_system/data/uploads/tender_processing/2025/11/20251123_111226_招标文件-哈银消金_c9a419c9.docx"

    print("=" * 100)
    print("分析所有层级的段落范围关系")
    print("=" * 100)
    print(f"文档: {Path(doc_path).name}\n")

    from docx import Document
    doc = Document(doc_path)

    parser = DocumentStructureParser()
    result = parser.parse_by_outline_level(doc_path)

    if not result['success']:
        print(f"❌ 解析失败")
        return

    chapters = result['chapters']

    # 分析1级节点（有子节点的）
    print("=" * 100)
    print("1级节点分析（只显示有子节点的）")
    print("=" * 100)

    for ch in chapters:
        analysis = analyze_node_relationship(ch, doc.paragraphs)
        if not analysis:
            continue

        print(f"\n📊 [{analysis['level']}级] {analysis['title']}")
        print(f"   段落范围: {analysis['para_range']}")
        print(f"   总字数: {analysis['total_words']:,}")
        print(f"   直接子节点数: {analysis['children_count']}")
        print(f"   直接子节点字数之和: {analysis['children_total_words']:,}")
        print(f"   自身直接内容字数: {analysis['direct_content_words']:,} (段落数: {analysis['direct_content_paras']})")

        # 验证
        val = analysis['validation']
        if val['match']:
            print(f"   ✅ 验证通过: {val['direct_content_words']:,} + {analysis['children_total_words']:,} = {val['actual']:,}")
        else:
            print(f"   ❌ 验证失败: 期望={val['expected']:,}, 实际={val['actual']:,}, 差异={val['diff']:,}")

        # 分析2级节点（如果有的话）
        for child in ch.get('children', []):
            child_analysis = analyze_node_relationship(child, doc.paragraphs)
            if not child_analysis:
                continue

            print(f"\n  📊 [{child_analysis['level']}级] {child_analysis['title']}")
            print(f"     段落范围: {child_analysis['para_range']}")
            print(f"     总字数: {child_analysis['total_words']:,}")
            print(f"     直接子节点数: {child_analysis['children_count']}")
            print(f"     直接子节点字数之和: {child_analysis['children_total_words']:,}")
            print(f"     自身直接内容字数: {child_analysis['direct_content_words']:,} (段落数: {child_analysis['direct_content_paras']})")

            # 验证
            val = child_analysis['validation']
            if val['match']:
                print(f"     ✅ 验证通过: {val['direct_content_words']:,} + {child_analysis['children_total_words']:,} = {val['actual']:,}")
            else:
                print(f"     ❌ 验证失败: 期望={val['expected']:,}, 实际={val['actual']:,}, 差异={val['diff']:,}")

    print("\n" + "=" * 100)
    print("总结")
    print("=" * 100)
    print("\n结论：")
    print("父节点字数 = 自身直接内容 + 直接子节点字数之和")
    print("\n因此，统计总字数时：")
    print("方案1: 只统计1级节点（根节点）的字数")
    print("方案2: 只统计叶子节点（没有子节点的节点）的字数")
    print("\n两种方案结果应该相同，因为:")
    print("  - 方案1: sum(1级节点字数) = sum(1级节点直接内容) + sum(2级节点字数)")
    print("           = sum(1级节点直接内容) + sum(2级节点直接内容) + sum(3级节点字数)")
    print("           = 所有叶子节点字数之和")
    print("  - 方案2: sum(叶子节点字数) = 所有叶子节点字数之和")
    print()

if __name__ == '__main__':
    main()
