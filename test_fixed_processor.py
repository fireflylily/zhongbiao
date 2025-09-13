#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试修复后的处理器
"""
import os
import sys
from docx import Document
from docx.shared import Inches

# 添加项目根目录到Python路径
sys.path.append('/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/ai_tender_system')

def create_test_document():
    """创建测试文档"""
    doc = Document()
    
    # 添加各种字段格式进行测试
    doc.add_paragraph("供应商名称：_____________")
    doc.add_paragraph("电话：_____________")
    doc.add_paragraph("电话                    电子邮件")  # 双字段表格式
    doc.add_paragraph("地址：_____________")
    doc.add_paragraph("传真：_____________")
    doc.add_paragraph("邮政编码：_____________")
    doc.add_paragraph("日期：____年____月____日")
    
    # 保存测试文档
    test_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/程序/test_template_fixed.docx"
    doc.save(test_file)
    print(f"创建测试文档: {test_file}")
    return test_file

def test_processor():
    """测试修复后的处理器"""
    from modules.point_to_point.processor import PointToPointProcessor
    
    # 创建测试文档
    input_file = create_test_document()
    output_file = input_file.replace('.docx', '_processed.docx')
    
    # 测试公司信息
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'fixedPhone': '010-63271000',
        'email': 'lvhe@smartsteps.com',
        'address': '北京市海淀区中关村大街1号',
        'fax': '010-63271001',
        'postalCode': '100080'
    }
    
    # 初始化处理器
    processor = PointToPointProcessor()
    
    print("=== 测试修复后的处理器 ===")
    print(f"公司信息: {company_info}")
    
    try:
        # 处理文档
        result = processor.process_business_response(
            input_file=input_file,
            output_file=output_file,
            company_info=company_info,
            project_name="测试项目",
            tender_no="TEST-2025-001",
            date_text="2025年9月12日"
        )
        
        print(f"\n处理结果: {result}")
        
        if result['success']:
            print(f"✅ 处理成功!")
            print(f"📄 输出文件: {output_file}")
            
            # 检查输出文档
            if os.path.exists(output_file):
                doc = Document(output_file)
                print(f"\n📋 处理后的文档内容:")
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        print(f"  {i+1}. {para.text}")
        else:
            print(f"❌ 处理失败: {result.get('error', 'Unknown error')}")
            
    except Exception as e:
        print(f"❌ 测试异常: {e}")
        import traceback
        traceback.print_exc()

def analyze_dual_field_processing():
    """专门分析双字段处理"""
    from modules.point_to_point.processor import PointToPointProcessor
    
    processor = PointToPointProcessor()
    
    # 设置公司信息
    processor.company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'fixedPhone': '010-63271000',
        'email': 'lvhe@smartsteps.com',
        'address': '北京市海淀区中关村大街1号',
        'fax': '010-63271001'
    }
    
    test_text = "电话                    电子邮件"
    current_field = "电话"
    field_value = "010-63271000"
    
    print(f"\n=== 双字段表格处理测试 ===")
    print(f"输入文本: '{test_text}'")
    print(f"当前字段: {current_field}")
    print(f"字段值: {field_value}")
    
    try:
        result = processor._handle_dual_field_table_layout(test_text, current_field, field_value)
        print(f"处理结果: '{result}'")
        
        if result != test_text:
            print("✅ 双字段处理成功")
        else:
            print("❌ 双字段处理无变化")
            
    except Exception as e:
        print(f"❌ 双字段处理异常: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("🚀 开始测试修复后的字段处理器")
    
    # 测试双字段处理
    analyze_dual_field_processing()
    
    # 测试完整处理器
    test_processor()
    
    print("\n🎉 测试完成")