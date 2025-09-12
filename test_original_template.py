#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试原始模板文件的字段处理执行结果
"""
import sys
import os
import json
from pathlib import Path
from datetime import datetime

# 添加项目路径
script_dir = Path(__file__).parent
project_root = script_dir / 'ai_tender_system'
sys.path.insert(0, str(project_root))

# 导入模块
try:
    from docx import Document
    
    # 测试文件路径
    template_file = "/Users/lvhe/Library/Mobile Documents/com~apple~CloudDocs/Work/智慧足迹2025/05投标项目/AI标书/项目/4-分段测试文件/采购人，项目名称，采购编号，（姓名，职务）（供应商名称，地址）传真，电子邮件，日期.docx"
    
    # 公司数据文件路径
    company_file = project_root / "data/configs/companies/945150ca-68e1-4141-921b-fd4c48e07ebd.json"
    
    print(f"📄 测试文件: {Path(template_file).name}")
    print(f"🏢 公司数据: {company_file.name}")
    
    # 检查文件存在性
    if not Path(template_file).exists():
        print(f"❌ 模板文件不存在: {template_file}")
        sys.exit(1)
    
    if not company_file.exists():
        print(f"❌ 公司数据文件不存在: {company_file}")
        sys.exit(1)
    
    # 读取公司数据
    with open(company_file, 'r', encoding='utf-8') as f:
        company_data = json.load(f)
    
    print(f"\n📋 公司关键信息:")
    print(f"  公司名称: {company_data.get('companyName', 'N/A')}")
    print(f"  固定电话: {company_data.get('fixedPhone', 'N/A')}")
    print(f"  电子邮件: {company_data.get('email', 'N/A')}")
    print(f"  传真: {company_data.get('fax', 'N/A')}")
    print(f"  法定代表人: {company_data.get('legalRepresentative', 'N/A')}")
    print(f"  注册地址: {company_data.get('registeredAddress', 'N/A')}")
    
    # 读取原始模板
    print(f"\n🔍 分析原始模板结构:")
    doc = Document(template_file)
    
    print(f"  总段落数: {len(doc.paragraphs)}")
    
    # 搜索目标字段
    target_keywords = {
        '采购人': '采购人信息',
        '项目名称': '项目名称',
        '采购编号': '采购编号', 
        '供应商名称': '供应商名称',
        '地址': '地址信息',
        '传真': '传真号码',
        '电子邮件': '电子邮件',
        '电话': '联系电话',
        '日期': '日期信息',
        '姓名': '联系人姓名',
        '职务': '职务信息'
    }
    
    found_paragraphs = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # 只处理非空段落
            # 检查是否包含目标关键词
            matched_keywords = []
            for keyword in target_keywords.keys():
                if keyword in text:
                    matched_keywords.append(keyword)
            
            if matched_keywords:
                found_paragraphs.append({
                    'index': i,
                    'text': text,
                    'keywords': matched_keywords,
                    'paragraph': paragraph
                })
                
                print(f"\n📍 段落 #{i} (包含: {', '.join(matched_keywords)}):")
                print(f"  内容: '{text}'")
                print(f"  Run结构 ({len(paragraph.runs)}个):")
                for j, run in enumerate(paragraph.runs):
                    if run.text:  # 只显示有内容的run
                        print(f"    Run {j}: '{run.text}'")
    
    print(f"\n🔧 模拟字段处理逻辑:")
    
    # 模拟MCP处理器的字段替换逻辑
    field_definitions = [
        {
            'patterns': [r'供应商名称[:：]\s*([_\s]*)', r'公司名称[:：]\s*([_\s]*)'],
            'value': company_data.get('companyName', ''),
            'field_name': '供应商名称'
        },
        {
            'patterns': [r'电话[:：]\s*([_\s]*)', r'联系电话[:：]\s*([_\s]*)', r'固定电话[:：]\s*([_\s]*)'],
            'value': company_data.get('fixedPhone', ''),
            'field_name': '联系电话'
        },
        {
            'patterns': [r'电子邮件[:：]\s*([_\s]*)', r'邮箱[:：]\s*([_\s]*)', r'email[:：]\s*([_\s]*)'],
            'value': company_data.get('email', ''),
            'field_name': '电子邮件'
        },
        {
            'patterns': [r'传真[:：]\s*([_\s]*)', r'fax[:：]\s*([_\s]*)'],
            'value': company_data.get('fax', ''),
            'field_name': '传真'
        },
        {
            'patterns': [r'地址[:：]\s*([_\s]*)', r'注册地址[:：]\s*([_\s]*)', r'办公地址[:：]\s*([_\s]*)'],
            'value': company_data.get('registeredAddress', ''),
            'field_name': '地址'
        }
    ]
    
    import re
    
    for para_info in found_paragraphs:
        text = para_info['text']
        print(f"\n🔄 处理段落 #{para_info['index']}: '{text}'")
        
        modified = False
        new_text = text
        
        for field_def in field_definitions:
            for pattern in field_def['patterns']:
                if re.search(pattern, text):
                    match = re.search(pattern, text)
                    if match:
                        # 模拟替换逻辑
                        replacement = f"{match.group().split(':')[0]}：{field_def['value']}"
                        new_text = re.sub(pattern, replacement, new_text)
                        modified = True
                        print(f"  ✅ 匹配字段 '{field_def['field_name']}': {pattern}")
                        print(f"     替换为: {replacement}")
                        break
        
        if modified:
            print(f"  📝 处理后文本: '{new_text}'")
        else:
            print(f"  ⚠️ 未找到匹配的字段模式")
    
    # 生成处理结果摘要
    print(f"\n📊 处理结果摘要:")
    print(f"  原始段落总数: {len(doc.paragraphs)}")
    print(f"  包含目标字段的段落: {len(found_paragraphs)}")
    
    # 统计各个字段的出现情况
    field_stats = {}
    for keyword in target_keywords.keys():
        count = sum(1 for para in found_paragraphs if keyword in para['keywords'])
        field_stats[keyword] = count
    
    print(f"\n🏷️ 字段出现统计:")
    for keyword, count in field_stats.items():
        status = "✅" if count > 0 else "❌"
        print(f"  {status} {keyword}: {count}次")
    
    print(f"\n✅ 原始模板分析完成")
    
except ImportError as e:
    print(f"❌ 导入错误: {e}")
    print("请确保安装了 python-docx: pip install python-docx")
except Exception as e:
    print(f"❌ 执行错误: {e}")
    import traceback
    traceback.print_exc()