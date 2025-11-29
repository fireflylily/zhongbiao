#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试脚本：分析商务应答Word文档结构，找出营业执照未插入的原因
"""

import sys
from pathlib import Path
from docx import Document

def analyze_document(doc_path):
    """分析Word文档结构"""
    print(f"📄 分析文档: {doc_path}")
    print("=" * 60)

    doc = Document(doc_path)

    # 1. 统计基本信息
    print(f"\n📊 基本统计:")
    print(f"  段落总数: {len(doc.paragraphs)}")
    print(f"  表格总数: {len(doc.tables)}")

    # 2. 查找"营业执照"相关段落
    print(f"\n🔍 查找'营业执照'关键词:")
    license_keywords = ['营业执照', '营业执照副本', '执照']

    found_in_paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for keyword in license_keywords:
            if keyword in text:
                found_in_paragraphs.append({
                    'index': idx,
                    'text': text,
                    'keyword': keyword,
                    'parent_type': type(para._parent).__name__
                })
                print(f"  ✅ 段落 #{idx}: '{text}'")
                print(f"     关键词: {keyword}")
                print(f"     父容器类型: {type(para._parent).__name__}")
                print(f"     段落对象: {para}")
                break

    if not found_in_paragraphs:
        print(f"  ❌ 未在段落中找到关键词")

    # 3. 查找表格中的"营业执照"
    print(f"\n📋 查找表格中的'营业执照':")
    found_in_tables = []
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                for keyword in license_keywords:
                    if keyword in cell_text:
                        found_in_tables.append({
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx,
                            'text': cell_text,
                            'keyword': keyword
                        })
                        print(f"  ✅ 表格 #{table_idx}, 行 {row_idx}, 列 {cell_idx}: '{cell_text}'")
                        print(f"     关键词: {keyword}")
                        break

    if not found_in_tables:
        print(f"  ❌ 未在表格中找到关键词")

    # 4. 显示前20个段落的内容（用于调试）
    print(f"\n📝 前20个段落内容:")
    for idx, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if text:
            print(f"  段落 #{idx}: {text[:80]}{'...' if len(text) > 80 else ''}")

    # 5. 总结
    print(f"\n📌 总结:")
    print(f"  段落中找到: {len(found_in_paragraphs)} 处")
    print(f"  表格中找到: {len(found_in_tables)} 处")

    if found_in_paragraphs:
        print(f"\n  第一个匹配的段落信息:")
        first = found_in_paragraphs[0]
        print(f"    位置: 段落 #{first['index']}")
        print(f"    文本: {first['text']}")
        print(f"    父容器: {first['parent_type']}")

    return {
        'paragraphs_found': found_in_paragraphs,
        'tables_found': found_in_tables
    }

if __name__ == '__main__':
    # 使用最近处理的商务应答文档
    # 您需要替换为实际的Word文档路径

    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        # 默认路径 - 请根据实际情况修改
        doc_path = input("请输入Word文档路径: ").strip()

    if not Path(doc_path).exists():
        print(f"❌ 文件不存在: {doc_path}")
        sys.exit(1)

    result = analyze_document(doc_path)

    print(f"\n" + "=" * 60)
    print(f"✅ 分析完成")
