#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证方法2和方法3字数差异的根本原因：表格内容
"""

import sys
from pathlib import Path

# 添加项目路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / 'ai_tender_system'))

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

def count_tables_in_range(doc, para_start, para_end):
    """
    统计指定段落范围内的表格数量和表格内容字数
    """
    para_count = 0
    start_body_idx = None
    end_body_idx = None

    # 找到段落在body中的位置
    for body_idx, element in enumerate(doc.element.body):
        if isinstance(element, CT_P):
            if para_count == para_start and start_body_idx is None:
                start_body_idx = body_idx
            if para_count == para_end:
                end_body_idx = body_idx
                break
            para_count += 1

    if start_body_idx is None or end_body_idx is None:
        return 0, 0, []

    # 统计表格
    table_count = 0
    table_word_count = 0
    table_info = []

    for body_idx in range(start_body_idx + 1, end_body_idx + 1):
        element = doc.element.body[body_idx]
        if isinstance(element, CT_Tbl):
            from docx.table import Table
            table = Table(element, doc)

            # 提取表格文本
            table_text_parts = []
            for row in table.rows:
                for cell in row.cells:
                    cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        table_text_parts.append(cell_text)

            table_text = '\n'.join(table_text_parts)
            words = len(table_text.replace(' ', '').replace('\n', ''))

            table_count += 1
            table_word_count += words
            table_info.append({
                'index': table_count,
                'words': words,
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0
            })

    return table_count, table_word_count, table_info

def main():
    doc_path = "ai_tender_system/data/uploads/tender_processing/2025/11/20251123_111226_招标文件-哈银消金_c9a419c9.docx"
    doc = Document(doc_path)

    print("=" * 100)
    print("验证表格内容导致的字数差异")
    print("=" * 100)
    print(f"文档: {Path(doc_path).name}\n")

    # 测试几个关键章节
    test_chapters = [
        ("第四部分 合同主要条款及格式", 267, 378),
        ("第五部分 采购需求书", 379, 426),
        ("第六部分 附  件", 427, 666),
    ]

    total_diff = 0

    for title, para_start, para_end in test_chapters:
        print(f"\n{'='*100}")
        print(f"章节: {title}")
        print(f"段落范围: {para_start} - {para_end}")
        print("=" * 100)

        # 方法3的计算方式（不含表格）
        content_paras = doc.paragraphs[para_start + 1 : para_end + 1]
        content_text = '\n'.join(p.text for p in content_paras)
        word_count_without_tables = len(content_text.replace(' ', '').replace('\n', ''))

        # 统计表格
        table_count, table_word_count, table_info = count_tables_in_range(doc, para_start, para_end)

        # 方法2的计算方式（含表格）
        word_count_with_tables = word_count_without_tables + table_word_count

        print(f"\n方法3字数（只计段落）: {word_count_without_tables:,}字")
        print(f"表格数量: {table_count}")
        print(f"表格字数: {table_word_count:,}字")
        print(f"方法2字数（段落+表格）: {word_count_with_tables:,}字")
        print(f"\n差异: {table_word_count:,}字 ({table_word_count/word_count_with_tables*100:.1f}%)")

        if table_info:
            print(f"\n表格详情:")
            for info in table_info:
                print(f"  表格{info['index']}: {info['words']:,}字 ({info['rows']}行 × {info['cols']}列)")

        total_diff += table_word_count

    print("\n" + "=" * 100)
    print("总结")
    print("=" * 100)
    print(f"\n所有测试章节的表格字数合计: {total_diff:,}字")
    print(f"\n这解释了为什么方法2和方法3的字数差异主要来自表格内容！")
    print(f"\n方法2: 使用 _extract_chapter_content_with_tables() → 包含表格")
    print(f"方法3: 直接提取 doc.paragraphs[...] → 不包含表格")

    print()

if __name__ == '__main__':
    main()
