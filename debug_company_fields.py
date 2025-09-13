#!/usr/bin/env python3
"""
调试脚本：检查商务应答文档中的公司信息字段
"""
import sys
import os
from pathlib import Path

# 添加项目路径到sys.path
script_dir = Path(__file__).parent
project_root = script_dir / 'ai_tender_system'
sys.path.insert(0, str(project_root))

try:
    from docx import Document
    
    # 找到最新的输出文档
    output_dir = project_root / 'data' / 'outputs'
    latest_file = None
    latest_time = 0
    
    for file in output_dir.glob('*.docx'):
        if file.stat().st_mtime > latest_time:
            latest_time = file.stat().st_mtime
            latest_file = file
    
    if not latest_file:
        print("❌ 没有找到输出文档")
        sys.exit(1)
    
    print(f"🔍 检查文档: {latest_file.name}")
    
    # 读取文档
    doc = Document(str(latest_file))
    
    # 搜索公司信息相关字段
    company_keywords = [
        '供应商名称', '公司名称', '法定代表人', '法人', 
        '电话', '联系电话', '固定电话',
        '邮件', '邮箱', 'email',
        '地址', '注册地址', '办公地址', '联系地址',
        '邮编', '邮政编码',
        '传真'
    ]
    
    print(f"\n📋 搜索公司信息字段:")
    found_fields = {}
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # 只处理非空段落
            for keyword in company_keywords:
                if keyword in text:
                    if keyword not in found_fields:
                        found_fields[keyword] = []
                    found_fields[keyword].append((i, text))
    
    # 显示找到的字段
    for keyword, matches in found_fields.items():
        print(f"\n🔑 关键词 '{keyword}':")
        for para_idx, text in matches:
            # 分析段落结构
            paragraph = doc.paragraphs[para_idx]
            print(f"  段落 #{para_idx}: '{text}'")
            
            # 检查是否有冒号分隔的字段格式
            if ':' in text or '：' in text:
                print(f"    ✓ 包含冒号分隔符")
                # 分析run结构
                print(f"    Run结构 ({len(paragraph.runs)}个):")
                for j, run in enumerate(paragraph.runs):
                    if run.text.strip():  # 只显示有内容的run
                        print(f"      Run {j}: '{run.text}'")
            else:
                print(f"    ⚠️ 无标准字段分隔符")
    
    # 特别检查"电话 电子邮件"这种格式
    print(f"\n🔍 检查组合字段格式:")
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if '电话' in text and ('邮件' in text or '邮箱' in text):
            print(f"  段落 #{i}: '{text}'")
            print(f"    Run分析:")
            for j, run in enumerate(paragraph.runs):
                print(f"      Run {j}: '{run.text}' (长度: {len(run.text)})")
    
    print(f"\n✅ 文档分析完成，共找到 {len(found_fields)} 种公司信息字段")
    
except Exception as e:
    print(f"❌ 执行错误: {e}")
    import traceback
    traceback.print_exc()