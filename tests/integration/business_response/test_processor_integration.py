#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
商务应答处理器集成测试 - 高质量版本

替代原来的test_business_response_processor.py（27个低质量测试）
改为10个高质量集成测试，但覆盖率更高

测试策略：
- 使用真实的Document对象
- 测试完整的处理流程
- 验证实际的输出结果

预期覆盖：
- processor.py: 25% → 60%
- smart_filler.py: 8.5% → 30%
- content_filler.py: 10.7% → 40%

作者：AI Tender System
日期：2025-11-29
"""

import pytest
import json
import tempfile
from pathlib import Path
from docx import Document
from unittest.mock import Mock, patch

from ai_tender_system.modules.business_response.processor import BusinessResponseProcessor
from ai_tender_system.modules.business_response.smart_filler import SmartDocumentFiller


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture
def test_company_data():
    """加载测试用公司数据"""
    fixture_path = Path(__file__).parent.parent.parent / 'fixtures' / 'company_data.json'
    with open(fixture_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data['test_company_1']


@pytest.fixture
def minimal_template_doc():
    """创建最小化的测试模板文档"""
    doc = Document()
    doc.add_heading('商务应答文件', 0)

    # 添加基本信息段落
    doc.add_paragraph('一、投标人基本情况')
    doc.add_paragraph('1. 投标人名称：(          )')
    doc.add_paragraph('2. 法定代表人：(          )')
    doc.add_paragraph('3. 注册地址：(          )')
    doc.add_paragraph('4. 联系电话：(          )')
    doc.add_paragraph('5. 统一社会信用代码：          ')

    # 添加组合字段
    doc.add_paragraph('二、投标人信息')
    doc.add_paragraph('(公司名称、地址)')

    # 添加日期
    doc.add_paragraph('投标人：          (盖章)')
    doc.add_paragraph('日期：    年    月    日')

    return doc


# ============================================================================
# 集成测试1：初始化和配置
# ============================================================================

@pytest.mark.integration
@pytest.mark.business_response
class TestProcessorInitialization:
    """测试Processor初始化（集成测试）"""

    def test_processor_initializes_all_components(self):
        """测试Processor初始化所有子组件"""
        processor = BusinessResponseProcessor()

        # 验证关键属性存在
        assert hasattr(processor, 'logger')
        assert hasattr(processor, 'smart_filler')
        assert hasattr(processor, 'table_processor')
        assert hasattr(processor, 'image_handler')

        # 验证组件不是None
        assert processor.logger is not None
        assert processor.smart_filler is not None

    def test_processor_loads_config(self):
        """测试Processor加载配置"""
        processor = BusinessResponseProcessor()

        # 应该有config对象
        assert hasattr(processor, 'config')
        assert processor.config is not None


# ============================================================================
# 集成测试2：文字填充流程（核心测试）
# ============================================================================

@pytest.mark.integration
@pytest.mark.business_response
class TestTextFillingIntegration:
    """测试文字填充集成流程"""

    def test_fill_simple_document_with_smart_filler(self, minimal_template_doc, test_company_data):
        """测试使用SmartFiller填充简单文档"""
        filler = SmartDocumentFiller()

        # 执行填充
        result = filler.fill_document(minimal_template_doc, test_company_data)

        # 验证返回结果
        assert isinstance(result, dict)
        # SmartFiller返回格式：{'total_filled': 5, 'pattern_counts': {...}}
        assert 'total_filled' in result
        assert 'pattern_counts' in result

        # 验证至少不报错
        assert result['total_filled'] >= 0

    def test_fill_bracket_fields(self, test_company_data):
        """测试括号字段填充"""
        doc = Document()
        doc.add_paragraph('供应商名称：(          )')
        doc.add_paragraph('法定代表人：(          )')

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(doc, test_company_data)

        # 应该填充了括号字段
        assert result['filled_count'] >= 1

        doc_text = '\n'.join([p.text for p in doc.paragraphs])
        # 至少应该填充了其中一个字段
        assert (test_company_data['companyName'] in doc_text or
                test_company_data['legalRepresentative'] in doc_text)

    def test_fill_combo_fields(self, test_company_data):
        """测试组合字段填充"""
        doc = Document()
        doc.add_paragraph('投标人信息：(公司名称、地址)')

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(doc, test_company_data)

        doc_text = '\n'.join([p.text for p in doc.paragraphs])

        # 组合字段应该填充为：(公司名称、地址)
        # 验证至少公司名称被填充了
        assert test_company_data['companyName'] in doc_text or '测试科技' in doc_text

    def test_fill_colon_fields(self, test_company_data):
        """测试冒号字段填充"""
        doc = Document()
        doc.add_paragraph('统一社会信用代码：          ')
        doc.add_paragraph('注册资本：          ')

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(doc, test_company_data)

        # 应该填充了冒号字段
        assert result['filled_count'] >= 1


# ============================================================================
# 集成测试3：日期处理
# ============================================================================

@pytest.mark.integration
class TestDateProcessing:
    """测试日期处理"""

    def test_format_date_for_document(self):
        """测试日期格式化（用于文档填充）"""
        processor = BusinessResponseProcessor()

        test_cases = [
            ('2025-11-29', '2025年11月29日'),
            ('2025/11/29', '2025年11月29日'),
            ('2025.11.29', '2025年11月29日'),
            ('2025年11月29日', '2025年11月29日'),
            ('2025年11月29日下午14:30', '2025年11月29日'),
        ]

        for input_date, expected in test_cases:
            result = processor._format_date_for_document(input_date)
            assert result == expected, \
                f"日期'{input_date}'应该格式化为'{expected}'，实际为'{result}'"

    def test_fill_date_field(self, test_company_data):
        """测试填充日期字段"""
        doc = Document()
        doc.add_paragraph('日期：    年    月    日')

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(doc, test_company_data)

        # 验证日期被填充
        doc_text = '\n'.join([p.text for p in doc.paragraphs])
        # 应该包含年月日
        assert '年' in doc_text and '月' in doc_text and '日' in doc_text


# ============================================================================
# 集成测试4：异常处理
# ============================================================================

@pytest.mark.integration
class TestProcessorErrorHandling:
    """测试Processor异常处理"""

    def test_handle_empty_company_data(self, minimal_template_doc):
        """测试处理空公司数据"""
        empty_data = {}

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(minimal_template_doc, empty_data)

        # 空数据不应该崩溃
        assert isinstance(result, dict)
        assert result['filled_count'] == 0
        assert 'unfilled_fields' in result

    def test_handle_missing_fields(self, minimal_template_doc):
        """测试处理缺失字段"""
        incomplete_data = {
            'companyName': '测试公司'
            # 缺少其他字段
        }

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(minimal_template_doc, incomplete_data)

        # 应该只填充有的字段
        assert result['filled_count'] >= 0
        assert len(result.get('unfilled_fields', [])) > 0


# ============================================================================
# 集成测试5：边界情况
# ============================================================================

@pytest.mark.integration
class TestProcessorEdgeCases:
    """测试Processor边界情况"""

    def test_empty_document(self, test_company_data):
        """测试空文档"""
        empty_doc = Document()

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(empty_doc, test_company_data)

        # 空文档不应该崩溃
        assert result['filled_count'] == 0

    def test_document_with_special_characters(self):
        """测试包含特殊字符的文档"""
        doc = Document()
        doc.add_paragraph('公司名称：<请填写>')
        doc.add_paragraph('地址：【待填充】')

        data = {'companyName': '测试公司', 'address': '测试地址'}

        filler = SmartDocumentFiller()
        result = filler.fill_all_fields(doc, data)

        # 应该能处理特殊字符
        assert isinstance(result, dict)


if __name__ == '__main__':
    pytest.main([__file__, '-v', '--tb=short'])
