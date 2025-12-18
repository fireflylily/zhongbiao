#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从商务应答模板自动提取测试用例

用途:
- 扫描 ai_tender_system/data/uploads/response_files/ 目录
- 提取所有括号字段
- 自动分类并生成测试用例JSON

使用方法:
    python tests/scripts/extract_test_cases_from_templates.py

输出:
    tests/data/business_response_test_cases_extracted.json (不覆盖原文件)

作者:AI Tender System
日期:2025-12-02
"""

import re
import json
import sys
from pathlib import Path
from datetime import datetime
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    print("❌ 缺少依赖包，请先安装：")
    print("   pip install python-docx")
    sys.exit(1)


class TemplateFieldExtractor:
    """模板字段提取器"""

    # 字段分类规则 (关键词匹配)
    FIELD_CATEGORIES = {
        'company_name': ['供应商', '公司', '单位', '企业', '应答人', '响应方', '投标人'],
        'address': ['地址'],
        'legal_person': ['法定代表人', '法人', '负责人'],
        'representative': ['授权', '代表', '委托', '被授权人'],
        'date': ['日期', '年', '月', '日'],
        'phone': ['电话', '联系方式', '手机'],
        'email': ['邮箱', 'email', 'Email', 'E-mail'],
    }

    # 需要过滤的噪音关键词
    NOISE_KEYWORDS = [
        '可选', '如有', '选填', '必填', '加盖', '盖章', '签字',
        '印刷体', '复印件', '原件', '份', '页', '项', '栏',
        '格式', '要求', '说明', '注', 'http', 'www', '查询',
        '截图', '报告', '证明', '材料', '文件', '附件'
    ]

    def __init__(self, template_dir):
        self.template_dir = Path(template_dir)
        self.all_fields = defaultdict(list)  # category -> [(field, source_info), ...]
        self.field_to_sources = defaultdict(list)  # field -> [source_info, ...]

    def is_noise_field(self, field):
        """判断是否是噪音字段"""
        # 过滤纯数字
        if field.isdigit():
            return True

        # 过滤太长的字段 (超过50个字符)
        if len(field) > 50:
            return True

        # 过滤包含噪音关键词的字段
        for noise in self.NOISE_KEYWORDS:
            if noise in field:
                return True

        return False

    def extract_project_name(self, doc):
        """从文档中提取项目名称"""
        # 尝试从前几段找到"项目名称"
        for para in doc.paragraphs[:20]:
            text = para.text.strip()
            if '项目名称' in text:
                # 提取项目名称
                match = re.search(r'项目名称[：:]\s*(.+?)[\n，,]?', text)
                if match:
                    return match.group(1).strip()

        return "未知项目"

    def extract_from_docx(self, docx_path):
        """从单个docx提取字段"""
        try:
            doc = Document(docx_path)
        except Exception as e:
            print(f"  ⚠️  无法读取 {docx_path.name}: {e}")
            return []

        # 提取项目信息
        path_parts = docx_path.parts
        project_id = path_parts[-2] if len(path_parts) >= 2 else "unknown"
        project_name = self.extract_project_name(doc)

        # 括号模式
        patterns = [
            r'\(([^)]+)\)',   # 英文括号
            r'（([^）]+)）',   # 中文括号
            r'\[([^\]]+)\]',  # 方括号
        ]

        extracted_fields = []
        for para_idx, para in enumerate(doc.paragraphs):
            for pattern in patterns:
                matches = re.findall(pattern, para.text)
                for match in matches:
                    field = match.strip()

                    # 过滤噪音
                    if self.is_noise_field(field):
                        continue

                    source_info = {
                        "type": "template",
                        "project_name": project_name,
                        "project_id": project_id,
                        "template_file": docx_path.name,
                        "template_path": str(docx_path.relative_to(Path.cwd())),
                        "paragraph_index": para_idx,
                        "extracted_date": datetime.now().strftime("%Y-%m-%d"),
                        "extract_method": "auto"
                    }

                    extracted_fields.append((field, source_info))

        return extracted_fields

    def classify_field(self, field):
        """分类字段"""
        field_lower = field.lower()

        for category, keywords in self.FIELD_CATEGORIES.items():
            for keyword in keywords:
                if keyword in field:
                    return category

        return None  # 返回None表示无法分类

    def map_to_standard_field(self, category):
        """映射到标准字段名"""
        mapping = {
            'company_name': 'companyName',
            'address': 'address',
            'legal_person': 'legalRepresentative',
            'representative': 'representativeName',
            'date': 'date',
            'phone': 'phone',
            'email': 'email',
        }
        return mapping.get(category)

    def scan_all_templates(self, max_files=50):
        """扫描所有模板"""
        docx_files = list(self.template_dir.rglob('*.docx'))

        # 过滤临时文件
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]

        print(f"📂 找到 {len(docx_files)} 个模板文件")

        if len(docx_files) > max_files:
            print(f"⚠️  文件数量过多，只处理最新的 {max_files} 个文件")
            # 按修改时间排序，取最新的
            docx_files = sorted(docx_files, key=lambda f: f.stat().st_mtime, reverse=True)[:max_files]

        total_extracted = 0
        for docx_file in docx_files:
            print(f"  📄 处理: {docx_file.name}")
            fields = self.extract_from_docx(docx_file)

            for field, source_info in fields:
                category = self.classify_field(field)
                if category:
                    self.all_fields[category].append((field, source_info))
                    self.field_to_sources[field].append(source_info)
                    total_extracted += 1

        print(f"✅ 提取完成，共 {total_extracted} 个字段(已分类)")
        print(f"📊 分类统计:")
        for category, fields in self.all_fields.items():
            print(f"   - {category}: {len(fields)} 个")

    def generate_test_cases_json(self):
        """生成测试用例JSON"""
        test_suites = {}

        # 为每个分类生成测试套件
        for category, field_list in self.all_fields.items():
            std_field = self.map_to_standard_field(category)
            if not std_field:
                continue

            # 去重：保留第一次出现的字段
            seen_fields = {}
            for field, source_info in field_list:
                if field not in seen_fields:
                    seen_fields[field] = source_info

            test_cases = []
            for idx, (field, source_info) in enumerate(seen_fields.items(), 1):
                test_cases.append({
                    "id": f"{category}_{idx:03d}_auto",
                    "field_alias": field,
                    "expected_standard_field": std_field,
                    "note": f"从{source_info['project_name']}模板自动提取",
                    "source": source_info
                })

            if test_cases:
                test_suites[f"field_recognition_{category}_extracted"] = {
                    "description": f"{category}字段识别测试(从模板自动提取)",
                    "test_cases": test_cases
                }

        # 统计信息
        total_cases = sum(len(suite['test_cases']) for suite in test_suites.values())
        by_source_type = {"template": total_cases}

        # 按项目统计
        by_project = defaultdict(int)
        for category_fields in self.all_fields.values():
            for field, source_info in category_fields:
                by_project[source_info['project_name']] += 1

        return {
            "version": "2.0-extracted",
            "last_updated": datetime.now().strftime("%Y-%m-%d"),
            "description": "从商务应答模板自动提取的测试用例",
            "extraction_info": {
                "extraction_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "template_count": len(list(self.template_dir.rglob('*.docx'))),
                "extracted_fields": total_cases
            },
            "test_suites": test_suites,
            "source_statistics": {
                "total_cases": total_cases,
                "by_source_type": by_source_type,
                "by_project": dict(by_project)
            }
        }


def main():
    """主函数"""
    print("=" * 70)
    print("📤 从商务应答模板提取测试用例")
    print("=" * 70)
    print()

    # 路径配置
    base_dir = Path(__file__).parent.parent.parent
    template_dir = base_dir / "ai_tender_system" / "data" / "uploads" / "response_files"
    output_file = base_dir / "tests" / "data" / "business_response_test_cases_extracted.json"

    # 检查模板目录是否存在
    if not template_dir.exists():
        print(f"❌ 模板目录不存在: {template_dir}")
        sys.exit(1)

    # 创建提取器
    extractor = TemplateFieldExtractor(template_dir)

    # 扫描模板
    extractor.scan_all_templates(max_files=20)  # 限制最多处理20个文件

    # 生成JSON
    print()
    print("📝 生成测试用例JSON...")
    test_cases_json = extractor.generate_test_cases_json()

    # 写入文件
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(test_cases_json, f, ensure_ascii=False, indent=2)

    print(f"✅ 已生成: {output_file}")
    print(f"📊 总计提取 {test_cases_json['source_statistics']['total_cases']} 个测试用例")
    print()
    print("📋 测试套件:")
    for suite_name, suite_data in test_cases_json['test_suites'].items():
        print(f"   - {suite_name}: {len(suite_data['test_cases'])} 个用例")

    print()
    print("=" * 70)
    print("✅ 提取完成！")
    print()
    print("📝 后续步骤:")
    print("  1. 查看生成的文件: cat tests/data/business_response_test_cases_extracted.json")
    print("  2. 人工审核测试用例,确认分类正确")
    print("  3. 合并到主文件: python tests/scripts/merge_test_cases.py")
    print("=" * 70)


if __name__ == "__main__":
    main()
