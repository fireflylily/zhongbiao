#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试长段落中项目信息处理问题
"""

import sys
import logging
from pathlib import Path
from docx import Document

# 添加项目根目录到系统路径
sys.path.append(str(Path(__file__).parent))

from ai_tender_system.modules.business_response.info_filler import InfoFiller

def setup_debug_logging():
    """设置调试日志"""
    # 设置根日志器为最详细的DEBUG级别
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    
    # 清除所有现有的处理器
    for handler in root_logger.handlers[:]:
        root_logger.removeHandler(handler)
    
    # 创建控制台处理器
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.DEBUG)
    
    # 创建格式器
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    console_handler.setFormatter(formatter)
    
    # 添加处理器到根日志器
    root_logger.addHandler(console_handler)
    
    return root_logger

def debug_long_paragraph():
    """专门调试长段落处理"""
    print("=" * 100)
    print("调试长段落中的项目信息处理")
    print("=" * 100)
    
    # 设置调试日志
    logger = setup_debug_logging()
    
    # 创建InfoFiller实例
    filler = InfoFiller()
    
    # 创建只包含长段落的测试文档
    doc = Document()
    long_text = "根据贵方为（项目名称）项目采购采购货物及服务的竞争性磋商公告（采购编号），签字代表（姓名、职务）经正式授权并代表供应商（供应商名称、地址）提交下述文件正本一份及副本份："
    doc.add_paragraph(long_text)
    
    # 准备项目信息（BusinessResponseProcessor格式）
    company_info = {
        'companyName': '智慧足迹数据科技有限公司',
        'address': '北京市东城区王府井大街200号七层711室',
    }
    
    project_info = {
        'projectName': '哈银消金2025年-2027年运营商数据采购项目',
        'projectNumber': 'GXTC-C-251590031',
        'purchaserName': '哈尔滨哈银消费金融有限责任公司',
    }
    
    print(f"\n📄 测试段落:")
    print(f"'{long_text}'")
    
    print(f"\n📋 项目信息:")
    for key, value in project_info.items():
        print(f"  {key}: {value}")
    
    print(f"\n" + "=" * 100)
    print("开始调试处理过程")
    print("=" * 100)
    
    # 手动调用段落处理来追踪详细过程
    try:
        paragraph = doc.paragraphs[0]
        
        print(f"\n🔍 段落处理前内容: '{paragraph.text}'")
        
        # 手动检查是否被跳过
        if filler._should_skip(paragraph.text):
            print("❌ 段落被跳过规则识别，不会被处理")
            return
        else:
            print("✅ 段落通过跳过规则检查")
        
        # 手动调用各个处理规则
        print(f"\n📝 尝试组合替换规则...")
        combo_result = filler._try_combination_rule(paragraph, project_info)
        print(f"组合替换结果: {combo_result}")
        
        print(f"\n📝 尝试单字段替换规则...")
        replacement_result = filler._try_replacement_rule(paragraph, project_info)
        print(f"单字段替换结果: {replacement_result}")
        
        print(f"\n📝 尝试填空规则...")
        fill_result = filler._try_fill_rule(paragraph, company_info)
        print(f"填空规则结果: {fill_result}")
        
        print(f"\n🔍 段落处理后内容: '{paragraph.text}'")
        
        # 完整处理测试
        print(f"\n" + "=" * 100)
        print("完整处理测试")
        print("=" * 100)
        
        result = filler.fill_info(doc, company_info, project_info)
        print(f"\n📊 最终结果: {result}")
        print(f"\n📄 最终段落内容: '{doc.paragraphs[0].text}'")
        
    except Exception as e:
        print(f"\n❌ 调试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_long_paragraph()